    # ── EXPORTACIÓN ──────────────────────────────────────────────
    with tab_exp:
        st.subheader("Exportar resultados del proyecto")
        col_w, col_d, col_x = st.columns(3)

        # ════════════════════════════════════════════════════════
        # MEMORIA TÉCNICA DOCX — 6 CAPÍTULOS COMPLETOS
        # ════════════════════════════════════════════════════════
        if col_w.button("📄 Generar Memoria DOCX"):
            try:
                from docx import Document
                from docx.shared import Pt, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import datetime

                doc = Document()
                for section in doc.sections:
                    section.top_margin    = Cm(2.5)
                    section.bottom_margin = Cm(2.5)
                    section.left_margin   = Cm(3.0)
                    section.right_margin  = Cm(2.0)

                def set_col_width(cell, width_cm):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(int(width_cm * 567)))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)

                def h(texto, nivel=1):
                    hh = doc.add_heading(texto, level=nivel)
                    hh.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    return hh

                def p(texto):
                    return doc.add_paragraph(texto)

                def tabla(encabezados, filas, anchos=None):
                    t = doc.add_table(rows=1, cols=len(encabezados))
                    t.style = 'Table Grid'
                    hdr = t.rows[0].cells
                    for i, hd_txt in enumerate(encabezados):
                        hdr[i].text = hd_txt
                        run = hdr[i].paragraphs[0].runs[0]
                        run.bold = True
                        run.font.size = Pt(10)
                        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if anchos:
                            set_col_width(hdr[i], anchos[i])
                    for fila in filas:
                        row = t.add_row().cells
                        for i, v in enumerate(fila):
                            row[i].text = str(v)
                            row[i].paragraphs[0].runs[0].font.size = Pt(10)
                            if anchos:
                                set_col_width(row[i], anchos[i])
                    return t

                # ── PORTADA ──────────────────────────────────────
                titulo = doc.add_heading('', 0)
                titulo.add_run(
                    'MEMORIA TÉCNICA DE CÁLCULO\n'
                    'MURO DE CONTENCIÓN EN VOLADIZO'
                ).bold = True
                titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph('')
                info = doc.add_paragraph()
                info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                info.add_run(
                    f'Proyecto: KonteWall — Diseño Estructural\n'
                    f'Norma de diseño: {norma_sel}\n'
                    f'Diseñó: ________________________________\n'
                    f'Empresa: _______________________________\n'
                    f'Fecha: {datetime.date.today().strftime("%d de %B de %Y")}\n'
                    f'Versión: KonteWall v1.0'
                )
                doc.add_page_break()

                # ── CAP 1: DATOS GENERALES ────────────────────────
                h('1. DATOS GENERALES Y PREDIMENSIONAMIENTO', 1)

                h('1.1 Materiales', 2)
                tabla(
                    ['Parámetro', 'Símbolo', 'Valor', 'Unidad'],
                    [
                        ["Resist. compresión concreto", "f'c", f"{fc:.0f}", "kg/cm²"],
                        ["Fluencia del acero",           "fy",  f"{fy:.0f}", "kg/cm²"],
                        ["Peso unitario concreto",       "γc",  f"{gc:.0f}", "kg/m³"],
                        ["Factor reducción flexión",     "φf",  f"{phif:.2f}", "—"],
                        ["Factor reducción cortante",    "φv",  f"{phic:.2f}", "—"],
                    ], [5.5, 2.5, 2.5, 2.0]
                )

                h('1.2 Suelo de relleno', 2)
                tabla(
                    ['Parámetro', 'Símbolo', 'Valor', 'Unidad'],
                    [
                        ["Peso unitario relleno",      "γr",  f"{gr:.0f}",   "kg/m³"],
                        ["Ángulo fricción interna",    "φ₁",  f"{ph1:.1f}",  "°"],
                        ["Inclinación terreno",        "α",   f"{alp:.1f}",  "°"],
                        ["Fricción suelo-muro",        "δ",   f"{delt:.2f}", "°"],
                        ["Sobrecarga uniforme",        "S/C", f"{SC:.0f}",   "kg/m²"],
                        ["Altura equiv. sobrecarga",   "hs",  f"{hs:.3f}",   "m"],
                    ], [5.5, 2.0, 2.5, 2.0]
                )

                h('1.3 Suelo de fundación', 2)
                tabla(
                    ['Parámetro', 'Símbolo', 'Valor', 'Unidad'],
                    [
                        ["Peso unitario fundación",    "γf",    f"{gf:.0f}",  "kg/m³"],
                        ["Ángulo de fricción",         "φ₂",    f"{ph2:.1f}", "°"],
                        ["Cohesión",                  "C",     f"{Ccoh:.0f}", "kg/m²"],
                        ["Cap. portante admisible",   "q_adm", f"{qadm:.2f}", "kg/cm²"],
                        ["Coef. fricción base-suelo", "μct",   f"{muct:.2f}", "—"],
                    ], [5.5, 2.0, 2.5, 2.0]
                )

                h('1.4 Geometría adoptada', 2)
                tabla(
                    ['Elemento', 'Símbolo', 'Valor', 'Rango recomendado'],
                    [
                        ["Altura total",              "Ht",  f"{Ht:.2f} m",  "—"],
                        ["Altura pantalla libre",     "Hp",  f"{Hp:.2f} m",  "Ht − hz"],
                        [f"Altura {T['zapata']}",     "hz",  f"{hz:.2f} m",  f"{Ht/12:.2f}–{Ht/10:.2f} m"],
                        ["Base total",                "B",   f"{B:.2f} m",   f"{Ht/2:.2f}–{2*Ht/3:.2f} m"],
                        ["Talón posterior",           "Bp",  f"{Bp:.2f} m",  "B − b − ct"],
                        [f"Punta / {T['punta']}",     "b",   f"{b:.2f} m",   "—"],
                        ["Base pantalla",             "ct",  f"{ct:.2f} m",  f"{Ht/12:.2f}–{Ht/10:.2f} m"],
                        ["Corona",                   "c",   f"{ccor:.2f} m","≥ max(Ht/24, 0.30m)"],
                        [f"{T['dentellón']}",         "hd",  f"{hd:.2f} m",  "0.0–0.80 m"],
                        ["Recubrimiento",             "r",   f"{rec:.1f} cm","7.5 cm exposición"],
                    ], [5.5, 2.0, 2.5, 3.5]
                )

                # ── CAP 2: COEFICIENTES DE EMPUJE ─────────────────
                doc.add_page_break()
                h('2. COEFICIENTES DE EMPUJE', 1)

                h('2.1 Coeficiente activo de Coulomb — Ka', 2)
                p(
                    f'Ka = {Ka:.5f}\n'
                    f'(φ₁={ph1}°, β=90°, α={alp}°, δ={delt:.2f}°)'
                )

                if sismo:
                    h('2.2 Coeficiente sísmico — Mononobe-Okabe (Kea)', 2)
                    p(
                        f'Kh = {Kh:.3f}  |  Kv = {Kv:.3f}\n'
                        f'θ = arctan(Kh / (1−Kv)) = {math.degrees(th_r):.4f}°\n'
                        f'Kea = {Kea:.5f}\n'
                        f'ΔKae = Kea − Ka = {Kea - Ka:.5f}'
                    )

                h('2.3 Coeficiente pasivo de Coulomb — Kp', 2)
                p(
                    f'Kp = {Kp:.4f}\n'
                    f'Empuje pasivo Ep = 0.5 × Kp × γf × hd² = {Ep:.2f} kg/m'
                )

                # ── CAP 3: ESTABILIDAD CON SOBRECARGA ─────────────
                doc.add_page_break()
                h('3. VERIFICACIÓN DE ESTABILIDAD', 1)
                h('3.1 Con sobrecarga' +
                  (' y sismo (Mononobe-Okabe)' if sismo else ''), 2)

                h('3.1.1 Fuerzas horizontales actuantes', 3)
                tabla(
                    ['Concepto', 'F [kg/m]', 'Brazo [m]', 'Ma [kg·m/m]'],
                    [
                        ['Eah — empuje tierra',              f"{Eah:.2f}",  f"{Hp/3:.3f}", f"{Ma_Eah:.2f}"],
                        ['Esch — empuje sobrecarga',         f"{Esch:.2f}", f"{Hp/2:.3f}", f"{Ma_Esch:.2f}"],
                        ['Es — incremento sísmico (MO)',     f"{Es:.2f}",   f"{Cy:.3f}",   f"{Ma_Es:.2f}"],
                        ['Eim — fuerza inercial muro+rel.',  f"{Eim:.2f}",  f"{Hp/2:.3f}", f"{Ma_Eim:.2f}"],
                        ['Ep — pasivo (resta)',              f"{-Ep:.2f}",  '—',           '—'],
                        ['TOTAL ΣFh',                        f"{sum_Fh:.2f}",'—',          f"{sum_Ma:.2f}"],
                    ], [6.0, 2.8, 2.5, 3.0]
                )

                h('3.1.2 Fuerzas verticales y momentos estabilizadores', 3)
                tabla(
                    ['Concepto', 'F [kg/m]', 'Brazo [m]', 'Mr [kg·m/m]'],
                    [
                        [f'W1 — {T["zapata"]}',          f"{W1:.2f}",  f"{x1:.3f}", f"{W1*x1:.2f}"],
                        [f'W2 — {T["pantalla"]} triang.',f"{W2:.2f}",  f"{x2:.3f}", f"{W2*x2:.2f}"],
                        [f'W3 — {T["pantalla"]} rect.',  f"{W3:.2f}",  f"{x3:.3f}", f"{W3*x3:.2f}"],
                        [f'W4 — {T["dentellón"]}',       f"{W4:.2f}",  f"{x4:.3f}", f"{W4*x4:.2f}"],
                        ['W5 — relleno posterior',       f"{W5:.2f}",  f"{x5:.3f}", f"{W5*x5:.2f}"],
                        ['WSC — sobrecarga',             f"{WSC:.2f}", f"{xSC:.3f}",f"{WSC*xSC:.2f}"],
                        ['Eav — comp. vert. activo',     f"{Eav:.2f}", f"{B:.3f}",  f"{Eav*B:.2f}"],
                        ['Escv — comp. vert. S/C',       f"{Escv:.2f}",f"{B:.3f}",  f"{Escv*B:.2f}"],
                        ['TOTAL ΣFv',                    f"{sum_Fv:.2f}",'—',       f"{sum_Mr:.2f}"],
                    ], [6.0, 2.8, 2.5, 3.0]
                )

                h('3.1.3 Verificaciones de estabilidad — Tabla D/C', 3)
                tabla(
                    ['Verificación', 'Valor calculado', 'Límite', 'D/C', 'Estado'],
                    [
                        [T['fs_volteo'], f"FS = {FS_volt:.3f}", "≥ 2.00",
                         f"{2.0/FS_volt:.3f}",  "CUMPLE" if FS_volt>=2.0 else "NO CUMPLE"],
                        [T['fs_desl'],   f"FS = {FS_desl:.3f}", "≥ 1.50",
                         f"{1.5/FS_desl:.3f}", "CUMPLE" if FS_desl>=1.5 else "NO CUMPLE"],
                        [T['excentr'],   f"e = {e_exc:.4f} m",  f"≤ {e_lim:.4f} m",
                         f"{e_exc/e_lim:.3f}", "CUMPLE" if e_exc<=e_lim else "NO CUMPLE"],
                        ["q₁ máx.",      f"{q1:.4f} kg/cm²",    f"≤ {qadm} kg/cm²",
                         f"{q1/qadm:.3f}","CUMPLE" if q1<=qadm else "NO CUMPLE"],
                        ["q₂ mín.",      f"{q2:.4f} kg/cm²",    "≥ 0",
                         "—",            "CUMPLE" if q2>=0 else "NO CUMPLE"],
                    ], [4.5, 3.0, 3.0, 2.0, 2.5]
                )

                h('3.2 Sin sobrecarga (verificación complementaria)', 2)
                tabla(
                    ['Verificación', 'Valor', 'Estado'],
                    [
                        [T['fs_volteo'], f"FS = {FSv_ns:.3f}", "OK" if FSv_ns>=2.0 else "REVISAR"],
                        [T['fs_desl'],   f"FS = {FSd_ns:.3f}", "OK" if FSd_ns>=1.5 else "REVISAR"],
                        [T['excentr'],   f"e = {e_ns:.4f} m",  "OK" if e_ns<=B/6   else "REVISAR"],
                        ["q₁",          f"{q1_ns:.4f} kg/cm²","OK" if q1_ns<=qadm else "REVISAR"],
                    ], [5.0, 4.0, 3.0]
                )

                # ── CAP 4: DISEÑO DE REFUERZO ──────────────────────
                doc.add_page_break()
                h('4. DISEÑO DE REFUERZO', 1)

                h(f'4.1 {T["pantalla"].title()} — Flexión en la base', 2)
                p(
                    f'Momento último de diseño:\n'
                    f'  Mu = 1.7×(0.5×Ka×γr×Hp²×Hp/3 + Ka×γr×hs×Hp×Hp/2)'
                    + (f' + 1.7×(Es×Cy + Eim×Hp/2)' if sismo else '') +
                    f'\n  Mu = {Mu_p:.2f} kg·m/m = {Mu_p/1000:.3f} t·m/m\n\n'
                    f'Parámetros:\n'
                    f'  d = ct×100 − r = {ct*100:.1f} − {rec:.1f} = {d_p:.2f} cm\n'
                    f'  Ru = Mu×100 / (φf×bw×d²) = {Ru_p:.5f} kg/cm²\n'
                    f'  ρ requerida = {rho_req:.6f}  |  ρ mín = {rho_min:.6f}\n'
                    f'  As req (cara interior) = {As_p:.3f} cm²/m\n'
                    f'  As mín (cara exterior) = {As_pe:.3f} cm²/m'
                )

                h('4.1.1 Tabla de opciones de armadura — cara interior', 3)
                df_t = tabla_barras(As_p)
                tabla(
                    list(df_t.columns),
                    [list(row) for _, row in df_t.iterrows()],
                    [2.5, 2.0, 2.5, 3.5]
                )
                p(
                    f'  → Adoptado cara interior: Ø {di_i} @ {s_i:.2f}m — '
                    f'As prov = {Ap_i:.3f} cm²/m\n'
                    f'  → Adoptado cara exterior:  Ø {di_e} @ {s_e:.2f}m — '
                    f'As prov = {Ap_e:.3f} cm²/m'
                )

                h('4.1.2 Acero horizontal (temperatura y retracción)', 3)
                p(
                    f'  Cara exterior (2/3): Ø 1/2" @ {s_h_ext:.2f}m\n'
                    f'  Cara interior (1/3): Ø 1/2" @ {s_h_int:.2f}m'
                )

                h('4.1.3 Verificación de cortante', 3)
                p(
                    f'  Vu = {Vu_pant:.2f} kg/m\n'
                    f'  φVc = φv×0.53×√f\'c×bw×d = {phiVc_p:.2f} kg/m\n'
                    f'  Estado: {"CUMPLE — no requiere estribos" if Vu_pant<=phiVc_p else "NO CUMPLE — aumentar ct"}'
                )

                h(f'4.1.4 {T["long_ancl"].title()}', 3)
                p(
                    f'  Ldh = {Ldh:.3f} m\n'
                    f'  Disponible hz = {hz:.2f} m\n'
                    f'  Estado: {"OK" if Ldh<=hz else "REVISAR — aumentar hz"}'
                )

                if dc > 0.05:
                    h(f'4.1.5 {T["cort_acero"].title()}', 3)
                    p(f'  Acero exterior puede cortarse a dc = {dc:.3f} m desde la corona.')

                h(f'4.2 {T["talon_post"].title()}', 2)
                p(
                    f'  Mu = {Mu_t:.2f} kg·m/m  |  d = {d_z:.2f} cm\n'
                    f'  As req = {As_t:.3f} cm²/m\n'
                    f'  → Longitudinal: Ø {diam_tl} @ {s_tl:.2f}m — As prov = {Ap_tl:.3f} cm²/m\n'
                    f'  → Transversal:  Ø 5/8" @ {s_tp:.2f}m\n'
                    f'  Vu = {Vu_t:.2f} kg/m  |  φVc = {phiVc_t:.2f} kg/m — '
                    f'{"CUMPLE" if Vu_t<=phiVc_t else "NO CUMPLE"}'
                )

                h(f'4.3 {T["talon_del"].title()} ({T["punta"]})', 2)
                p(
                    f'  Mu = {Mu_pu:.2f} kg·m/m  |  As req = {As_pu:.3f} cm²/m\n'
                    f'  → Longitudinal: Ø {diam_pu} @ {s_pu:.2f}m — As prov = {Ap_pu:.3f} cm²/m\n'
                    f'  → Transversal:  Ø 5/8" @ {s_pp:.2f}m'
                )

                # ── CAP 5: RESUMEN DE ARMADURA ─────────────────────
                doc.add_page_break()
                h('5. RESUMEN GENERAL DE ARMADURA', 1)
                tabla(
                    ['Elemento', 'Diámetro', 'Espaciamiento', 'As prov [cm²/m]'],
                    [
                        [f"{T['pantalla'].title()} — {T['cara_int']}",
                         f"Ø {di_i}", f"{s_i:.2f} m", f"{Ap_i:.3f}"],
                        [f"{T['pantalla'].title()} — {T['cara_ext']}",
                         f"Ø {di_e}", f"{s_e:.2f} m", f"{Ap_e:.3f}"],
                        [f"{T['pantalla'].title()} — horiz. ext.",
                         "Ø 1/2\"", f"{s_h_ext:.2f} m",
                         f"{As_prov('1/2\"', s_h_ext):.3f}"],
                        [f"{T['pantalla'].title()} — horiz. int.",
                         "Ø 1/2\"", f"{s_h_int:.2f} m",
                         f"{As_prov('1/2\"', s_h_int):.3f}"],
                        [f"{T['talon_post'].title()} — Longitudinal",
                         f"Ø {diam_tl}", f"{s_tl:.2f} m", f"{Ap_tl:.3f}"],
                        [f"{T['talon_post'].title()} — Transversal",
                         "Ø 5/8\"", f"{s_tp:.2f} m",
                         f"{As_prov('5/8\"', s_tp):.3f}"],
                        [f"{T['talon_del'].title()} — Longitudinal",
                         f"Ø {diam_pu}", f"{s_pu:.2f} m", f"{Ap_pu:.3f}"],
                        [f"{T['talon_del'].title()} — Transversal",
                         "Ø 5/8\"", f"{s_pp:.2f} m",
                         f"{As_prov('5/8\"', s_pp):.3f}"],
                    ], [6.5, 2.5, 3.5, 3.5]
                )

                # ── CAP 6: NOTAS Y FIRMA ───────────────────────────
                doc.add_page_break()
                h('6. NOTAS DE DISEÑO Y RESPONSABILIDAD', 1)
                p(
                    f'1. Herramienta de apoyo profesional KonteWall v1.0 — '
                    f'Konte, Construcción Consultoría y Tecnología.\n'
                    f'2. Norma de diseño activa: {norma_sel}.\n'
                    f'3. Factores de carga: 1.4D + 1.7L (gravitacional), '
                    f'análisis sísmico por Mononobe-Okabe (pseudoestático).\n'
                    f'4. Recubrimiento libre: {rec:.1f} cm '
                    f'(ambiente moderadamente agresivo).\n'
                    f'5. ρ_mín = {rho_min:.5f} según {norma_sel}.\n'
                    f'6. El ingeniero diseñador es responsable de verificar '
                    f'todos los resultados antes de su implementación.\n'
                    f'7. Fecha de generación: '
                    f'{datetime.date.today().strftime("%d/%m/%Y")}.'
                )
                doc.add_paragraph('\n\n\n')
                firma = doc.add_paragraph()
                firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
                firma.add_run(
                    '___________________________________\n'
                    'Ingeniero Diseñador\n'
                    'Matrícula Profesional: _______________\n'
                    f'Fecha: {datetime.date.today().strftime("%d/%m/%Y")}'
                )

                buf_w = io.BytesIO()
                doc.save(buf_w)
                buf_w.seek(0)
                st.download_button(
                    "⬇️ Descargar Memoria DOCX", buf_w,
                    file_name=f"Memoria_KonteWall_Ht{Ht:.1f}m.docx",
                    mime="application/vnd.openxmlformats-officedocument"
                          ".wordprocessingml.document"
                )
                st.success("✅ Memoria técnica generada — 6 capítulos.")

            except ImportError:
                st.error("Instala python-docx: pip install python-docx")
            except Exception as e_doc:
                st.error(f"Error DOCX: {e_doc}")

        # ════════════════════════════════════════════════════════
        # PLANO DXF — COMPLETO CON COTAS, ACERO Y TEXTOS
        # ════════════════════════════════════════════════════════
        if col_d.button("📐 Generar Plano DXF"):
            try:
                import ezdxf

                dxf_doc = ezdxf.new('R2010')
                dxf_doc.header['$INSUNITS'] = 6
                msp = dxf_doc.modelspace()

                for capa, color in [
                    ('CONCRETO', 7), ('ACERO', 5),
                    ('COTAS', 2),    ('RELLENO', 3), ('TEXTO', 1)
                ]:
                    dxf_doc.layers.add(capa, color=color)

                # Zapata
                msp.add_lwpolyline(
                    [(0,0),(B,0),(B,hz),(0,hz),(0,0)],
                    dxfattribs={'layer':'CONCRETO','lineweight':50,
                                'closed':True})
                # Pantalla
                msp.add_lwpolyline(
                    [(b,hz),(b+ct,hz),(b+ct,Ht),(b+(ct-ccor),Ht),(b,hz)],
                    dxfattribs={'layer':'CONCRETO','lineweight':50,
                                'closed':True})
                # Dentellón
                if hd > 0:
                    msp.add_lwpolyline(
                        [(b,0),(b+ct,0),(b+ct,-hd),(b,-hd),(b,0)],
                        dxfattribs={'layer':'CONCRETO','lineweight':35,
                                    'closed':True})
                # Hatch relleno
                hatch = msp.add_hatch(
                    color=30, dxfattribs={'layer':'RELLENO'})
                hatch.set_pattern_fill('ANSI31', scale=0.05)
                hatch.paths.add_polyline_path(
                    [(b+ct,hz),(B,hz),(B,Ht),(b+ct,Ht)],
                    is_closed=True)

                # Acero interior pantalla
                msp.add_line(
                    (b+ct-rec/100, hz+0.05),
                    (b+ct-rec/100, Ht-0.05),
                    dxfattribs={'layer':'ACERO','color':5,'lineweight':30})
                # Acero talón posterior
                msp.add_line(
                    (b+ct+rec/100, rec/100),
                    (B-rec/100, rec/100),
                    dxfattribs={'layer':'ACERO','color':5,'lineweight':30})
                # Acero punta
                msp.add_line(
                    (rec/100, rec/100),
                    (b-rec/100, rec/100),
                    dxfattribs={'layer':'ACERO','color':5,'lineweight':30})
                # Acero superior zapata
                msp.add_line(
                    (rec/100, hz-rec/100),
                    (B-rec/100, hz-rec/100),
                    dxfattribs={'layer':'ACERO','color':4,'lineweight':18})

                # Cotas verticales
                def cota_dxf(x, y1, y2, lbl):
                    msp.add_line((x,y1),(x,y2),
                        dxfattribs={'layer':'COTAS','color':2})
                    for yv in [y1, y2]:
                        msp.add_line((x-0.05,yv),(x+0.05,yv),
                            dxfattribs={'layer':'COTAS','color':2})
                    msp.add_text(lbl, dxfattribs={
                        'layer':'TEXTO','height':0.10,'color':1,
                        'insert':(x+0.08,(y1+y2)/2)})

                cota_dxf(B+0.30, 0,  Ht, f"Ht={Ht:.2f}m")
                cota_dxf(B+0.55, hz, Ht, f"Hp={Hp:.2f}m")
                cota_dxf(-0.25,  0,  hz, f"hz={hz:.2f}m")
                if hd > 0:
                    cota_dxf(-0.25, -hd, 0, f"hd={hd:.2f}m")

                # Cota horizontal B
                msp.add_line(
                    (0, -hd-0.30),(B, -hd-0.30),
                    dxfattribs={'layer':'COTAS','color':2})
                msp.add_text(f"B={B:.2f}m", dxfattribs={
                    'layer':'TEXTO','height':0.12,'color':1,
                    'insert':(B/2, -hd-0.45)})

                # Textos técnicos
                msp.add_text(
                    f"KonteWall v1.0 — {norma_sel}",
                    dxfattribs={'layer':'TEXTO','height':0.15,
                                'color':1,'insert':(0, Ht+0.30)})
                msp.add_text(
                    f"f'c={fc:.0f}kg/cm2  fy={fy:.0f}kg/cm2  "
                    f"Ka={Ka:.4f}  Kea={Kea:.4f}",
                    dxfattribs={'layer':'TEXTO','height':0.10,
                                'color':1,'insert':(0, Ht+0.15)})
                msp.add_text(
                    f"Pantalla int: O{di_i}@{s_i:.2f}m  "
                    f"ext: O{di_e}@{s_e:.2f}m  "
                    f"Talon: O{diam_tl}@{s_tl:.2f}m  "
                    f"Punta: O{diam_pu}@{s_pu:.2f}m",
                    dxfattribs={'layer':'TEXTO','height':0.09,
                                'color':1,'insert':(0, Ht+0.02)})

                buf_dxf = io.StringIO()
                dxf_doc.write(buf_dxf)
                st.download_button(
                    "⬇️ Descargar Plano DXF",
                    buf_dxf.getvalue().encode('utf-8'),
                    file_name=f"Plano_KonteWall_Ht{Ht:.1f}m.dxf",
                    mime="application/octet-stream"
                )
                st.success("✅ DXF generado — 5 capas: CONCRETO / ACERO / "
                           "COTAS / RELLENO / TEXTO.")

            except ImportError:
                st.error("Instala ezdxf: pip install ezdxf")
            except Exception as e_dxf:
                st.error(f"Error DXF: {e_dxf}")

        # ════════════════════════════════════════════════════════
        # PRESUPUESTO XLSX — 3 HOJAS COMPLETAS
        # ════════════════════════════════════════════════════════
        if col_x.button("📊 Exportar Presupuesto XLSX"):
            try:
                # Volúmenes
                vol_zapata   = round(B * hz, 3)
                vol_pantalla = round(0.5*(ct+ccor)*Hp, 3)
                vol_dent     = round(ct*hd, 3) if hd > 0 else 0.0
                vol_total    = round(vol_zapata+vol_pantalla+vol_dent, 3)

                def kg_ac(As_cm2, long_m):
                    return round(As_cm2/10000 * long_m * 7850, 2)

                kg_pi  = kg_ac(Ap_i,  Hp)
                kg_pe  = kg_ac(Ap_e,  Hp)
                kg_ph  = kg_ac(
                    As_prov("1/2\"",s_h_ext)+As_prov("1/2\"",s_h_int), 1.0)
                kg_tl  = kg_ac(Ap_tl, Bp)
                kg_pu  = kg_ac(Ap_pu, b)
                kg_tot = round(kg_pi+kg_pe+kg_ph+kg_tl+kg_pu, 2)
                vol_exc = round((B+0.50)*(Ht+hd+0.20), 3)

                df_cub = pd.DataFrame([
                    ["CONCRETO", f"Zapata",             f"{vol_zapata:.3f}",  "m³/m", f"f'c={fc:.0f}kg/cm²"],
                    ["CONCRETO", "Pantalla",             f"{vol_pantalla:.3f}","m³/m", ""],
                    ["CONCRETO", T['dentellón'],          f"{vol_dent:.3f}",   "m³/m", ""],
                    ["CONCRETO", "TOTAL CONCRETO",        f"{vol_total:.3f}",  "m³/m", ""],
                    ["ACERO",    f"Pant. Int Ø {di_i}",  f"{kg_pi:.2f}",      "kg/m", f"@{s_i:.2f}m"],
                    ["ACERO",    f"Pant. Ext Ø {di_e}",  f"{kg_pe:.2f}",      "kg/m", f"@{s_e:.2f}m"],
                    ["ACERO",    "Pant. Horiz Ø 1/2\"",  f"{kg_ph:.2f}",      "kg/m", ""],
                    ["ACERO",    f"Talón Ø {diam_tl}",   f"{kg_tl:.2f}",      "kg/m", f"@{s_tl:.2f}m"],
                    ["ACERO",    f"Punta Ø {diam_pu}",   f"{kg_pu:.2f}",      "kg/m", f"@{s_pu:.2f}m"],
                    ["ACERO",    "TOTAL ACERO",           f"{kg_tot:.2f}",     "kg/m", f"fy={fy:.0f}kg/cm²"],
                    ["EXCAV.",   "Corte y limpieza",      f"{vol_exc:.3f}",    "m³/m", "talud 1:0.5"],
                ], columns=["Capítulo","Ítem","Cantidad","Unidad","Nota"])

                df_est = pd.DataFrame({
                    "Verificación":
                        [T["fs_volteo"],T["fs_desl"],T["excentr"],"q₁","q₂"],
                    "Valor":
                        [f"FS={FS_volt:.3f}",f"FS={FS_desl:.3f}",
                         f"e={e_exc:.4f}m",f"{q1:.4f}kg/cm²",
                         f"{q2:.4f}kg/cm²"],
                    "Límite":
                        ["≥2.00","≥1.50",f"≤{e_lim:.4f}m",
                         f"≤{qadm}","≥0"],
                    "D/C":
                        [f"{2.0/FS_volt:.3f}",f"{1.5/FS_desl:.3f}",
                         f"{e_exc/e_lim:.3f}",f"{q1/qadm:.3f}","—"],
                    "Estado":
                        ["CUMPLE" if x else "NO CUMPLE"
                         for x in [FS_volt>=2.0,FS_desl>=1.5,
                                    e_exc<=e_lim,q1<=qadm,q2>=0]]
                })

                df_arm = pd.DataFrame({
                    "Elemento": [
                        f"{T['pantalla'].title()} — {T['cara_int']}",
                        f"{T['pantalla'].title()} — {T['cara_ext']}",
                        f"{T['pantalla'].title()} — horiz. ext.",
                        f"{T['pantalla'].title()} — horiz. int.",
                        f"{T['talon_post'].title()} — Long.",
                        f"{T['talon_post'].title()} — Trans.",
                        f"{T['talon_del'].title()} — Long.",
                        f"{T['talon_del'].title()} — Trans.",
                    ],
                    "Diámetro":
                        [f"Ø {di_i}",f"Ø {di_e}","Ø 1/2\"","Ø 1/2\"",
                         f"Ø {diam_tl}","Ø 5/8\"",f"Ø {diam_pu}","Ø 5/8\""],
                    "Espaciamiento [m]":
                        [s_i,s_e,s_h_ext,s_h_int,s_tl,s_tp,s_pu,s_pp],
                    "As prov [cm²/m]":
                        [round(Ap_i,3),round(Ap_e,3),
                         round(As_prov("1/2\"",s_h_ext),3),
                         round(As_prov("1/2\"",s_h_int),3),
                         round(Ap_tl,3),round(As_prov("5/8\"",s_tp),3),
                         round(Ap_pu,3),round(As_prov("5/8\"",s_pp),3)]
                })

                buf_xl = io.BytesIO()
                with pd.ExcelWriter(buf_xl, engine='xlsxwriter') as writer:
                    wb = writer.book
                    fmt_tit = wb.add_format({
                        'bold':True,'font_size':13,'align':'center',
                        'valign':'vcenter','bg_color':'#1e3a1e',
                        'font_color':'white'})
                    fmt_hdr = wb.add_format({
                        'bold':True,'bg_color':'#2e5c2e',
                        'font_color':'white','border':1,'align':'center'})
                    fmt_ok  = wb.add_format({
                        'bg_color':'#c8e6c9','bold':True,'border':1})
                    fmt_nok = wb.add_format({
                        'bg_color':'#ffcdd2','bold':True,'border':1})

                    # Hoja 1 — Cubicación
                    df_cub.to_excel(
                        writer, sheet_name='Cubicacion',
                        index=False, startrow=2)
                    ws1 = writer.sheets['Cubicacion']
                    ws1.merge_range(
                        'A1:E1',
                        f'KonteWall — Cubicación | Ht={Ht:.2f}m | {norma_sel}',
                        fmt_tit)
                    for ci, hdr_txt in enumerate(df_cub.columns):
                        ws1.write(2, ci, hdr_txt, fmt_hdr)
                    ws1.set_column('A:A', 12)
                    ws1.set_column('B:B', 30)
                    ws1.set_column('C:E', 14)

                    # Hoja 2 — Estabilidad
                    df_est.to_excel(
                        writer, sheet_name='Estabilidad',
                        index=False, startrow=2)
                    ws2 = writer.sheets['Estabilidad']
                    ws2.merge_range(
                        'A1:E1',
                        f'KonteWall — Estabilidad | {norma_sel}',
                        fmt_tit)
                    for ci, hdr_txt in enumerate(df_est.columns):
                        ws2.write(2, ci, hdr_txt, fmt_hdr)
                    ws2.set_column('A:E', 20)
                    for ri, row in df_est.iterrows():
                        fmt_e = fmt_ok if row['Estado']=='CUMPLE' else fmt_nok
                        ws2.write(ri+3, 4, row['Estado'], fmt_e)

                    # Hoja 3 — Armadura
                    df_arm.to_excel(
                        writer, sheet_name='Armadura',
                        index=False, startrow=2)
                    ws3 = writer.sheets['Armadura']
                    ws3.merge_range(
                        'A1:D1',
                        f'KonteWall — Armadura | {norma_sel}',
                        fmt_tit)
                    for ci, hdr_txt in enumerate(df_arm.columns):
                        ws3.write(2, ci, hdr_txt, fmt_hdr)
                    ws3.set_column('A:A', 38)
                    ws3.set_column('B:D', 18)

                buf_xl.seek(0)
                st.download_button(
                    "⬇️ Descargar Presupuesto XLSX", buf_xl,
                    file_name=f"Presupuesto_KonteWall_Ht{Ht:.1f}m.xlsx",
                    mime="application/vnd.openxmlformats-officedocument"
                          ".spreadsheetml.sheet"
                )
                st.success("✅ XLSX generado — 3 hojas: Cubicación / "
                           "Estabilidad / Armadura.")

            except ImportError:
                st.error("Instala xlsxwriter: pip install xlsxwriter")
            except Exception as e_xl:
                st.error(f"Error XLSX: {e_xl}")
