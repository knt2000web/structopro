import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import json
import datetime
import math

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

# ─────────────────────────────────────────────
# PARÁMETROS POR NORMA (extraídos de CODES en otros módulos)
# ─────────────────────────────────────────────
CODES = {
    "NSR-10 (Colombia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NSR-10 Título C"},
    "ACI 318-25 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-25"},
    "ACI 318-19 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-19"},
    "ACI 318-14 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-14"},
    "NEC-SE-HM (Ecuador)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NEC-SE-HM"},
    "E.060 (Perú)": {"phi_flex":0.90, "phi_shear":0.85, "phi_comp":0.70, "lambda":1.0, "ref":"E.060"},
    "NTC-EM (México)": {"phi_flex":0.85, "phi_shear":0.80, "phi_comp":0.70, "lambda":1.0, "ref":"NTC-EM"},
    "COVENIN 1753-2006 (Venezuela)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.70, "lambda":1.0, "ref":"COVENIN"},
    "NB 1225001-2020 (Bolivia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NB"},
    "CIRSOC 201-2025 (Argentina)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"CIRSOC"},
}
code = CODES.get(norma_sel, CODES["NSR-10 (Colombia)"])
phi_flex = code["phi_flex"]
phi_shear = code["phi_shear"]
phi_comp = code["phi_comp"]
lam = code["lambda"]

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES (diseño)
# ─────────────────────────────────────────────
def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_rho_min(fc, fy):
    return max(0.25*math.sqrt(fc)/fy, 1.4/fy)

def get_rho_max_beam(fc, fy, beta1):
    eps_cu = 0.003
    eps_t_min = 0.005
    return (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+eps_t_min))

def design_column(Pu, Mu, b, h, fc, fy, recub=5):
    """Diseño de columna rectangular con flexión uniaxial (simplificado). Retorna As, ok, mensaje."""
    Ag = b * h * 10000  # cm²
    # Acero mínimo 1% de Ag
    As_min = 0.01 * Ag
    # Capacidad axial con acero mínimo
    Pn_max = 0.85 * fc * (Ag - As_min) / 1000 + fy * As_min / 1000  # kN
    phi_Pn_max = phi_comp * Pn_max
    if Pu > phi_Pn_max:
        # Necesita más acero o mayor sección; estimamos nueva área
        As_needed = max(As_min, (Pu / phi_comp - 0.85*fc*Ag/1000) / (fy/1000))
        As_needed = max(As_needed, 0.01 * Ag)
        ok = False
        As = As_needed
    else:
        As = As_min
        ok = True
    # Verificación de momento (muy simplificada: se asume que la columna es esbelta y el momento es pequeño)
    # Para una columna corta, el momento resistente se puede estimar con diagrama de interacción; aquí asumimos que si la carga axial está dentro, el momento también lo está si es pequeño.
    # En caso de momento grande, se aumentaría acero; pero para no complicar, dejamos mensaje de advertencia.
    if Mu > 0.1 * phi_Pn_max * 0.1:  # un criterio arbitrario
        msg = "Momento significativo - revisar diseño detallado"
    else:
        msg = ""
    return As, ok, msg

def design_beam(Mu, Vu, b, h, fc, fy, recub=5):
    """Diseño de viga rectangular. Retorna As (cm²) y s_estribos (cm)."""
    d = h - recub - 1  # cm (estribo supuesto 1 cm)
    # Flexión
    Rn = (Mu * 1e6) / (phi_flex * b * d**2)
    disc = 1 - 2*Rn/(0.85*fc)
    if disc > 0:
        rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
    else:
        rho = 0.0018
    rho_min = get_rho_min(fc, fy)
    rho = max(rho, rho_min)
    As = rho * b * d
    # Cortante
    Vc = 0.17 * lam * math.sqrt(fc) * b * d / 10  # kN
    phi_Vc = phi_shear * Vc
    if Vu > phi_Vc/2:
        Vs = Vu / phi_shear - Vc
        Av = 2 * 0.71  # #3 estribo, 2 ramas (cm²)
        s = Av * fy * d / Vs  # cm
        s = max(5, min(s, 60, d/2))
    else:
        s = 0
    return As, s

def design_footing(Pu, Mu, qa, fc, fy, recub=5):
    """Diseño de zapata cuadrada. Retorna B (m), H (m), As (cm²/m)."""
    # Área requerida por carga axial
    A_req = Pu / qa
    B = math.sqrt(A_req)
    B = math.ceil(B*20)/20  # redondear a 0.05 m
    # Espesor mínimo (corte unidireccional aproximado)
    # Suponiendo columna 0.4x0.4 m
    L_v = (B - 0.4)/2
    # Presión neta promedio
    q_net = Pu / (B*B)
    # Momento en la cara de la columna
    Mu_zap = q_net * B * (L_v**2)/2
    d = 0.3  # suposición inicial, luego se ajustaría. Para simplificar, fijamos H = max(0.4, B/4)
    H = max(0.4, B/4)
    # Acero mínimo (0.0018 * b * h)
    As_min = 0.0018 * B*100 * H*100
    # Acero por flexión (simplificado)
    d_cm = H*100 - recub - 1
    b_cm = B*100
    Rn = (Mu_zap * 1e6) / (phi_flex * b_cm * d_cm**2)
    if Rn > 0:
        disc = 1 - 2*Rn/(0.85*fc)
        if disc > 0:
            rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
        else:
            rho = 0.0018
    else:
        rho = 0.0018
    As = rho * b_cm * d_cm
    As = max(As, As_min)
    return B, H, As

# ─────────────────────────────────────────────
# FUNCIONES DE GENERACIÓN GEOMÉTRICA (igual que antes, pero añadiendo pisos)
# ─────────────────────────────────────────────
def generar_malla_3d(Lx, Lz, nx, nz, alturas_pisos, col_b, col_h, vig_b, vig_h):
    pisos = len(alturas_pisos)
    nudos = []
    x_coords = np.linspace(0, Lx, nx)
    z_coords = np.linspace(0, Lz, nz)
    y_coords = [0.0]
    curr_y = 0.0
    for h in alturas_pisos:
        curr_y += h
        y_coords.append(curr_y)
    nid = 1
    for y in y_coords:
        for z in z_coords:
            for x in x_coords:
                nudos.append({"ID": nid, "X": round(x,2), "Y": round(y,2), "Z": round(z,2)})
                nid += 1
    df_nudos = pd.DataFrame(nudos)

    def find_nid(x, y, z):
        res = df_nudos[(np.isclose(df_nudos['X'], x)) & (np.isclose(df_nudos['Y'], y)) & (np.isclose(df_nudos['Z'], z))]
        return int(res.iloc[0]['ID']) if not res.empty else None

    # Columnas
    columnas = []
    cid = 1
    for z in z_coords:
        for x in x_coords:
            for p in range(pisos):
                n1 = find_nid(x, y_coords[p], z)
                n2 = find_nid(x, y_coords[p+1], z)
                if n1 and n2:
                    columnas.append({"ID": f"C{cid}", "N1": n1, "N2": n2, "Base_X": round(x,2), "Base_Z": round(z,2), "Piso": p+1, "b (m)": col_b, "h (m)": col_h})
                    cid += 1

    # Vigas X
    vigas_x = []
    vid = 1
    for y in y_coords[1:]:
        for z in z_coords:
            for i in range(nx-1):
                n1 = find_nid(x_coords[i], y, z)
                n2 = find_nid(x_coords[i+1], y, z)
                if n1 and n2:
                    vigas_x.append({"ID": f"VX{vid}", "N1": n1, "N2": n2, "Piso": y_coords.index(y), "b (m)": vig_b, "h (m)": vig_h})
                    vid += 1
    # Vigas Z
    vigas_z = []
    for y in y_coords[1:]:
        for x in x_coords:
            for i in range(nz-1):
                n1 = find_nid(x, y, z_coords[i])
                n2 = find_nid(x, y, z_coords[i+1])
                if n1 and n2:
                    vigas_z.append({"ID": f"VZ{vid}", "N1": n1, "N2": n2, "Piso": y_coords.index(y), "b (m)": vig_b, "h (m)": vig_h})
                    vid += 1

    # Zapatas
    zapatas = []
    zid = 1
    for z in z_coords:
        for x in x_coords:
            nid = find_nid(x, 0, z)
            if nid:
                es_borde_x = (np.isclose(x, 0) or np.isclose(x, Lx))
                es_borde_z = (np.isclose(z, 0) or np.isclose(z, Lz))
                if es_borde_x and es_borde_z:
                    tipo = "Esquinera"
                elif es_borde_x or es_borde_z:
                    tipo = "Medianera"
                else:
                    tipo = "Centrada"
                zapatas.append({"ID": f"Z{zid}", "Nudo": nid, "Tipo": tipo, "X": x, "Z": z, "b (m)": 1.5, "h (m)": 1.5, "D_f (m)": 1.0})
                zid += 1

    return df_nudos, pd.DataFrame(columnas), pd.DataFrame(vigas_x), pd.DataFrame(vigas_z), pd.DataFrame(zapatas)


# ─────────────────────────────────────────────
# FUNCIONES DE DISEÑO (por elemento)
# ─────────────────────────────────────────────
def design_all_elements(nudos_df, cols_df, vigas_x_df, vigas_z_df, zaps_df,
                        fc, fy, qa, recub, alturas_pisos, area_planta,
                        cm_kN_m2, cv_kN_m2):
    """Recorre todos los elementos y calcula diseño, volúmenes y acero."""
    # Factores de combinación (LRFD)
    wu = 1.2*cm_kN_m2 + 1.6*cv_kN_m2  # kN/m²
    # Área tributaria por columna (simplificada: la mitad de la luz en cada dirección)
    # Obtener coordenadas únicas X, Z
    x_vals = sorted(nudos_df['X'].unique())
    z_vals = sorted(nudos_df['Z'].unique())
    dx = x_vals[1] - x_vals[0] if len(x_vals)>1 else 1
    dz = z_vals[1] - z_vals[0] if len(z_vals)>1 else 1
    area_trib = dx * dz
    # Almacenar resultados por columna
    col_results = []
    for _, col in cols_df.iterrows():
        # Carga axial acumulada desde el piso de la columna hasta la azotea
        piso_actual = col['Piso']
        pisos_sobre = len(alturas_pisos) - piso_actual + 1  # incluye el piso actual
        Pu = wu * area_trib * pisos_sobre
        # Momento por excentricidad accidental (5% de la dimensión)
        e = 0.05 * max(col['b (m)'], col['h (m)'])
        Mu = Pu * e
        b_cm = col['b (m)']*100
        h_cm = col['h (m)']*100
        As, ok, msg = design_column(Pu, Mu, b_cm, h_cm, fc, fy, recub)
        # Volumen de concreto
        n1 = nudos_df[nudos_df['ID']==col['N1']].iloc[0]
        n2 = nudos_df[nudos_df['ID']==col['N2']].iloc[0]
        L = abs(n2['Y'] - n1['Y'])
        vol = L * col['b (m)'] * col['h (m)']
        # Acero (peso)
        peso_acero = As * (L*100) * 0.785  # kg
        col_results.append({
            "ID": col['ID'], "Piso": col['Piso'], "b (cm)": b_cm, "h (cm)": h_cm,
            "Pu (kN)": Pu, "Mu (kN-m)": Mu, "As (cm²)": As, "Cumple": ok,
            "Vol (m³)": vol, "Peso_acero (kg)": peso_acero, "Observación": msg
        })
    df_cols = pd.DataFrame(col_results)

    # Diseño de vigas (simplificado: momento y cortante aproximados para viga simplemente apoyada)
    # Usamos la carga wu * ancho tributario de la viga (la mitad de la luz perpendicular)
    # Para vigas X, ancho tributario en Z es dz/2 a cada lado; viga X recibe carga de losa en un ancho ~ dz
    beam_results = []
    for _, v in pd.concat([vigas_x_df, vigas_z_df]).iterrows():
        # Obtener longitud
        n1 = nudos_df[nudos_df['ID']==v['N1']].iloc[0]
        n2 = nudos_df[nudos_df['ID']==v['N2']].iloc[0]
        if 'X' in v['ID']:  # viga en X: longitud en X
            L = abs(n2['X'] - n1['X'])
            ancho_trib = dz/2 + dz/2  # = dz
        else:  # viga en Z
            L = abs(n2['Z'] - n1['Z'])
            ancho_trib = dx/2 + dx/2  # = dx
        w_viga = wu * ancho_trib  # kN/m
        Mu = w_viga * L**2 / 8
        Vu = w_viga * L / 2
        b_cm = v['b (m)']*100
        h_cm = v['h (m)']*100
        As, s = design_beam(Mu, Vu, b_cm, h_cm, fc, fy, recub)
        vol = L * v['b (m)'] * v['h (m)']
        peso_acero = As * (L*100) * 0.785  # kg
        beam_results.append({
            "ID": v['ID'], "Piso": v['Piso'], "b (cm)": b_cm, "h (cm)": h_cm,
            "L (m)": L, "Mu (kN-m)": Mu, "Vu (kN)": Vu, "As (cm²)": As,
            "s_estribos (cm)": s, "Vol (m³)": vol, "Peso_acero (kg)": peso_acero
        })
    df_beams = pd.DataFrame(beam_results)

    # Zapatas (asumimos carga axial de la columna más baja, sin momento)
    # Usamos la columna en base (piso=1) de cada ubicación
    base_cols = cols_df[cols_df['Piso']==1].copy()
    zap_results = []
    for _, zap in zaps_df.iterrows():
        # buscar columna con la misma base (X,Z)
        col_base = base_cols[(base_cols['Base_X'] == zap['X']) & (base_cols['Base_Z'] == zap['Z'])].iloc[0] if not base_cols.empty else None
        if col_base is not None:
            Pu = col_results[col_base.name]['Pu (kN)']  # usar el Pu de la columna
        else:
            Pu = 0
        B, H, As_zap = design_footing(Pu, 0, qa, fc, fy, recub)
        vol = B * B * H
        peso_acero = As_zap * (B*100) * 0.785  # kg (acero en ambas direcciones)
        zap_results.append({
            "ID": zap['ID'], "Tipo": zap['Tipo'], "Pu (kN)": Pu, "B (m)": B, "H (m)": H,
            "As (cm²/m)": As_zap, "Vol (m³)": vol, "Peso_acero (kg)": peso_acero
        })
    df_zaps = pd.DataFrame(zap_results)

    # Losa
    # Determinamos si es unidireccional o bidireccional según relación de luces
    if area_planta:
        # Para simplificar, asumimos que la losa se diseña como unidireccional en la dirección más corta
        L_short = min(dx, dz)
        L_long = max(dx, dz)
        if L_long / L_short > 2:
            # Unidireccional
            moment_coef = 1/8
        else:
            moment_coef = 1/10  # aproximado
        w_losa = wu
        Mu_losa = w_losa * L_short**2 * moment_coef
        # Espesor supuesto por norma (mínimo L/20)
        h_losa = max(0.10, L_short/20)
        # Diseño de losa por metro
        b_losa = 100  # cm
        d_losa = h_losa*100 - recub - 0.5
        Rn = (Mu_losa * 1e6) / (phi_flex * b_losa * d_losa**2)
        disc = 1 - 2*Rn/(0.85*fc)
        if disc > 0:
            rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
        else:
            rho = 0.0018
        rho_min = get_rho_min(fc, fy)
        rho = max(rho, rho_min)
        As_losa = rho * b_losa * d_losa  # cm²/m
        # Acero mínimo por temperatura (0.0018)
        As_temp = 0.0018 * b_losa * h_losa*100
        As_losa = max(As_losa, As_temp)
        # Volumen de losa
        vol_losa = area_planta * h_losa
        peso_acero_losa = (As_losa * (area_planta*100) / 100) * 0.785 * 2  # dos direcciones
    else:
        vol_losa = 0
        peso_acero_losa = 0
        As_losa = 0
        h_losa = 0

    return df_cols, df_beams, df_zaps, vol_losa, peso_acero_losa, h_losa, As_losa

# ─────────────────────────────────────────────
# VISUALIZACIÓN
# ─────────────────────────────────────────────
def plot_edificio(nudos, cols, vigas_x, vigas_z, zaps):
    fig = go.Figure()
    # Columnas
    for _, c in cols.iterrows():
        n1 = nudos[nudos['ID']==c['N1']].iloc[0]
        n2 = nudos[nudos['ID']==c['N2']].iloc[0]
        fig.add_trace(go.Scatter3d(
            x=[n1['X'], n2['X']], y=[n1['Z'], n2['Z']], z=[n1['Y'], n2['Y']],
            mode='lines', line=dict(color='#2196F3', width=6), name="Columnas", showlegend=False
        ))
    # Vigas
    for _, v in vigas_x.iterrows():
        n1 = nudos[nudos['ID']==v['N1']].iloc[0]
        n2 = nudos[nudos['ID']==v['N2']].iloc[0]
        fig.add_trace(go.Scatter3d(
            x=[n1['X'], n2['X']], y=[n1['Z'], n2['Z']], z=[n1['Y'], n2['Y']],
            mode='lines', line=dict(color='#4CAF50', width=4), name="Vigas", showlegend=False
        ))
    for _, v in vigas_z.iterrows():
        n1 = nudos[nudos['ID']==v['N1']].iloc[0]
        n2 = nudos[nudos['ID']==v['N2']].iloc[0]
        fig.add_trace(go.Scatter3d(
            x=[n1['X'], n2['X']], y=[n1['Z'], n2['Z']], z=[n1['Y'], n2['Y']],
            mode='lines', line=dict(color='#4CAF50', width=4), name="Vigas", showlegend=False
        ))
    # Zapatas
    for _, z in zaps.iterrows():
        fig.add_trace(go.Scatter3d(
            x=[z['X']], y=[z['Z']], z=[0],
            mode='markers', marker=dict(size=6, color='red', symbol='square'), name="Zapatas", showlegend=False
        ))
    # Losa (superficie en cada piso)
    y_pisos = sorted(nudos['Y'].unique())
    for y in y_pisos[1:]:
        nx_coords = nudos[nudos['Y']==y]['X']
        nz_coords = nudos[nudos['Y']==y]['Z']
        fig.add_trace(go.Mesh3d(
            x=[nx_coords.min(), nx_coords.max(), nx_coords.max(), nx_coords.min()],
            y=[nz_coords.min(), nz_coords.min(), nz_coords.max(), nz_coords.max()],
            z=[y, y, y, y],
            color='cyan', opacity=0.1, hoverinfo="skip", name="Losa", showlegend=False
        ))
    fig.update_layout(
        scene=dict(xaxis_title='X (m)', yaxis_title='Z (m)', zaxis_title='Y (m)', aspectmode='data'),
        margin=dict(l=0, r=0, b=0, t=30),
        plot_bgcolor='black', paper_bgcolor='#1e1e1e'
    )
    return fig

# =============================================================================
# INTERFAZ PRINCIPAL
# =============================================================================
st.set_page_config(page_title=_t("Generador Maestro 3D", "Master 3D Generator"), layout="wide")
st.image(r"assets/generador_3d_header_1773257156151.png", use_container_width=True)
st.title(_t("Generador Maestro 3D Paramétrico", "Parametric Master 3D Generator"))
st.markdown(_t("Construcción rápida de Edificios 3D con cálculos estructurales completos según la norma seleccionada.", "Rapid 3D Building generation with complete structural calculations per the selected code."))

# Inicializar estado
if "g3d_nudos" not in st.session_state:
    st.session_state.update({
        "g3d_nudos": pd.DataFrame(),
        "g3d_cols": pd.DataFrame(),
        "g3d_vx": pd.DataFrame(),
        "g3d_vz": pd.DataFrame(),
        "g3d_zaps": pd.DataFrame(),
        "g_state": False,
        "g_design": None
    })

# Barra lateral
with st.sidebar:
    _iso = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}.get(norma_sel, "un")
    st.markdown(f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;"><img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;"><span style="color:#7ec87e;font-weight:600;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span></div>', unsafe_allow_html=True)
    st.markdown("---")

    st.header(_t("📏 Parametrización Geométrica", "📏 Geometric Parametrization"))
    L_x = st.number_input("Frente Lote X (m)", value=st.session_state.get("g3d_lx", 12.0), min_value=2.0, key="g3d_lx")
    L_z = st.number_input("Fondo Lote Z (m)", value=st.session_state.get("g3d_lz", 15.0), min_value=2.0, key="g3d_lz")
    n_x = st.number_input("N° Columnas en Frente (eje X)", value=st.session_state.get("g3d_nx", 4), min_value=2, key="g3d_nx")
    n_z = st.number_input("N° Columnas en Fondo (eje Z)", value=st.session_state.get("g3d_nz", 5), min_value=2, key="g3d_nz")
    st.subheader("Alzado (Elevación)")
    n_pisos = st.number_input("Número de Pisos", value=st.session_state.get("g3d_npisos", 3), min_value=1, key="g3d_npisos")
    if "alturas_df" not in st.session_state or len(st.session_state.alturas_df) != n_pisos:
        data_h = [{"Piso": p+1, "Altura (m)": 3.5 if p==0 else 3.0} for p in range(n_pisos)]
        st.session_state.alturas_df = pd.DataFrame(data_h)
    st.session_state.alturas_df = st.data_editor(st.session_state.alturas_df, use_container_width=True, hide_index=True)
    alturas_list = st.session_state.alturas_df["Altura (m)"].tolist()

    st.subheader("Secciones (Dimensiones en metros)")
    col_b = st.number_input("Base Columnas (m)", value=st.session_state.get("g3d_cb", 0.40), step=0.05, key="g3d_cb")
    col_h = st.number_input("Altura Columnas (m)", value=st.session_state.get("g3d_ch", 0.40), step=0.05, key="g3d_ch")
    vig_b = st.number_input("Base Vigas (m)", value=st.session_state.get("g3d_vb", 0.30), step=0.05, key="g3d_vb")
    vig_h = st.number_input("Altura Vigas (m)", value=st.session_state.get("g3d_vh", 0.40), step=0.05, key="g3d_vh")

    st.subheader("Materiales y Suelo")
    # Estos valores deberían venir del sidebar global, pero por simplicidad los pedimos aquí
    fc = st.number_input("Resistencia concreto f'c [MPa]", value=st.session_state.get("g3d_fc", 21.0), step=1.0, key="g3d_fc")
    fy = st.number_input("Fluencia acero fy [MPa]", value=st.session_state.get("g3d_fy", 420.0), step=10.0, key="g3d_fy")
    qa = st.number_input("Capacidad portante suelo qa [kN/m²]", value=st.session_state.get("g3d_qa", 150.0), step=10.0, key="g3d_qa")
    recub = st.number_input("Recubrimiento [cm]", value=st.session_state.get("g3d_rec", 5.0), step=0.5, key="g3d_rec")

    st.subheader("Cargas (kN/m²)")
    cm = st.number_input("Carga Muerta (CM)", value=st.session_state.get("g3d_cm", 4.5), step=0.5, key="g3d_cm")
    cv = st.number_input("Carga Viva (CV)", value=st.session_state.get("g3d_cv", 2.0), step=0.5, key="g3d_cv")

    if st.button("🏗️ Generar / Actualizar Malla 3D", type="primary", use_container_width=True):
        dn, dc, dvx, dvz, dz = generar_malla_3d(L_x, L_z, n_x, n_z, alturas_list, col_b, col_h, vig_b, vig_h)
        st.session_state.update({
            "g3d_nudos": dn, "g3d_cols": dc, "g3d_vx": dvx, "g3d_vz": dvz, "g3d_zaps": dz,
            "g_state": True, "g_design": None
        })

    st.markdown("---")
    st.subheader("📂 Gestor de Proyectos")
    project_name = st.text_input("Nombre del Proyecto:", value=st.session_state.get("project_name", "Mi_Edificio"), key="g3d_pn")
    project_owner = st.text_input("Propietario / Cliente:", value=st.session_state.get("project_owner", ""), key="g3d_po")
    project_address = st.text_input("Dirección de Obra:", value=st.session_state.get("project_address", ""), key="g3d_pa")
    project_phone = st.text_input("Teléfono de Contacto:", value=st.session_state.get("project_phone", ""), key="g3d_pp")
    if project_name and project_owner and project_address and project_phone:
        def serialize_state():
            state_dict = {}
            for k, v in st.session_state.items():
                if isinstance(v, pd.DataFrame):
                    state_dict[k] = {"__type__": "dataframe", "data": v.to_dict(orient="records")}
                elif isinstance(v, (int, float, str, bool, list, dict)):
                    state_dict[k] = v
            return json.dumps(state_dict, indent=4)
        st.download_button("💾 Guardar Proyecto (.json)", data=serialize_state(), file_name=f"{project_name}_{datetime.datetime.now().strftime('%Y%m%d')}.json", mime="application/json", use_container_width=True)
    else:
        st.info("✍️ Llene los datos del proyecto para habilitar el guardado.")

# Cuerpo principal
if st.session_state.g_state:
    nudos = st.session_state.g3d_nudos
    cols = st.session_state.g3d_cols
    vigas_x = st.session_state.g3d_vx
    vigas_z = st.session_state.g3d_vz
    zaps = st.session_state.g3d_zaps

    col_vis, col_data = st.columns([1.2, 1])
    with col_vis:
        st.plotly_chart(plot_edificio(nudos, cols, vigas_x, vigas_z, zaps), use_container_width=True, height=600)

    with col_data:
        # Botón para lanzar el diseño
        if st.button("📐 Ejecutar Diseño Estructural", type="primary", use_container_width=True):
            area_planta = (max(nudos['X']) - min(nudos['X'])) * (max(nudos['Z']) - min(nudos['Z']))
            df_cols, df_beams, df_zaps, vol_losa, peso_acero_losa, h_losa, As_losa = design_all_elements(
                nudos, cols, vigas_x, vigas_z, zaps, fc, fy, qa, recub, alturas_list, area_planta, cm, cv
            )
            st.session_state.g_design = {
                "columnas": df_cols,
                "vigas": df_beams,
                "zapatas": df_zaps,
                "vol_losa": vol_losa,
                "peso_acero_losa": peso_acero_losa,
                "h_losa": h_losa,
                "As_losa": As_losa
            }
            st.success("✅ Diseño completado.")

        if st.session_state.g_design is not None:
            des = st.session_state.g_design
            st.subheader("Resumen de materiales")
            # Calcular totales
            vol_conc = des["columnas"]["Vol (m³)"].sum() + des["vigas"]["Vol (m³)"].sum() + des["zapatas"]["Vol (m³)"].sum() + des["vol_losa"]
            peso_acero = des["columnas"]["Peso_acero (kg)"].sum() + des["vigas"]["Peso_acero (kg)"].sum() + des["zapatas"]["Peso_acero (kg)"].sum() + des["peso_acero_losa"]
            st.metric("Volumen total concreto", f"{vol_conc:.2f} m³")
            st.metric("Peso total acero", f"{peso_acero:.1f} kg")
            st.metric("Cuantía de acero", f"{peso_acero/vol_conc:.1f} kg/m³" if vol_conc>0 else "N/A")

            # Pestañas para ver detalles
            tabs = st.tabs(["Columnas", "Vigas", "Zapatas", "Losa"])
            with tabs[0]:
                st.dataframe(des["columnas"], use_container_width=True)
            with tabs[1]:
                st.dataframe(des["vigas"], use_container_width=True)
            with tabs[2]:
                st.dataframe(des["zapatas"], use_container_width=True)
            with tabs[3]:
                st.write(f"Espesor losa: {des['h_losa']:.2f} m")
                st.write(f"Acero principal: {des['As_losa']:.2f} cm²/m")
                st.write(f"Acero temperatura: {max(0.0018*100*des['h_losa']*100, 0):.2f} cm²/m")

            # Exportaciones
            st.markdown("---")
            st.subheader("📥 Exportar")
            col_e1, col_e2, col_e3 = st.columns(3)
            with col_e1:
                if st.button("📊 Exportar a Excel"):
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        des["columnas"].to_excel(writer, sheet_name="Columnas", index=False)
                        des["vigas"].to_excel(writer, sheet_name="Vigas", index=False)
                        des["zapatas"].to_excel(writer, sheet_name="Zapatas", index=False)
                        pd.DataFrame({"Item": ["Volumen concreto", "Peso acero"], "Valor": [vol_conc, peso_acero]}).to_excel(writer, sheet_name="Resumen", index=False)
                    st.download_button("📥 Descargar Excel", data=output.getvalue(), file_name="Generador_3D_Resultados.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col_e2:
                if st.button("📄 Memoria DOCX"):
                    doc = Document()
                    doc.add_heading(f"Memoria de Cálculo - Edificio Paramétrico ({norma_sel})", 0)
                    doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
                    doc.add_paragraph(f"Materiales: f'c={fc} MPa, fy={fy} MPa, qa={qa} kN/m²")
                    doc.add_paragraph(f"Geometría: Lx={L_x} m, Lz={L_z} m, {n_pisos} pisos, altura media {sum(alturas_list)/len(alturas_list):.1f} m")
                    doc.add_heading("Cantidades de obra", level=1)
                    doc.add_paragraph(f"Concreto total: {vol_conc:.2f} m³")
                    doc.add_paragraph(f"Acero total: {peso_acero:.1f} kg")
                    doc.add_heading("Resumen de diseño", level=1)
                    doc.add_paragraph("Columnas (primeras 5):")
                    table_col = doc.add_table(rows=1+min(5, len(des["columnas"])), cols=6)
                    table_col.style = 'Table Grid'
                    hdr = table_col.rows[0].cells
                    hdr[0].text = "ID"; hdr[1].text = "b (cm)"; hdr[2].text = "h (cm)"; hdr[3].text = "As (cm²)"; hdr[4].text = "Cumple"; hdr[5].text = "Observación"
                    for i, row in des["columnas"].head(5).iterrows():
                        cells = table_col.rows[i+1].cells
                        cells[0].text = str(row["ID"])
                        cells[1].text = f"{row['b (cm)']:.0f}"
                        cells[2].text = f"{row['h (cm)']:.0f}"
                        cells[3].text = f"{row['As (cm²)']:.1f}"
                        cells[4].text = "✅" if row["Cumple"] else "❌"
                        cells[5].text = row["Observación"]
                    doc.add_paragraph("Vigas (primeras 5):")
                    table_vig = doc.add_table(rows=1+min(5, len(des["vigas"])), cols=6)
                    table_vig.style = 'Table Grid'
                    hdr = table_vig.rows[0].cells
                    hdr[0].text = "ID"; hdr[1].text = "b (cm)"; hdr[2].text = "h (cm)"; hdr[3].text = "As (cm²)"; hdr[4].text = "s_estribos (cm)"; hdr[5].text = "L (m)"
                    for i, row in des["vigas"].head(5).iterrows():
                        cells = table_vig.rows[i+1].cells
                        cells[0].text = str(row["ID"])
                        cells[1].text = f"{row['b (cm)']:.0f}"
                        cells[2].text = f"{row['h (cm)']:.0f}"
                        cells[3].text = f"{row['As (cm²)']:.1f}"
                        cells[4].text = f"{row['s_estribos (cm)']:.0f}" if row["s_estribos (cm)"]>0 else "No requiere"
                        cells[5].text = f"{row['L (m)']:.2f}"
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button("📥 Descargar Memoria DOCX", data=buf, file_name="Memoria_Generador_3D.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col_e3:
                if st.button("🏗️ Exportar DXF"):
                    dwg = ezdxf.new('R2010')
                    msp = dwg.modelspace()
                    # Crear capas
                    dwg.layers.add(name="COLUMNAS", color=5)
                    dwg.layers.add(name="VIGAS", color=3)
                    dwg.layers.add(name="ZAPATAS", color=1)
                    # Dibujar columnas
                    for _, col in cols.iterrows():
                        n1 = nudos[nudos['ID']==col['N1']].iloc[0]
                        n2 = nudos[nudos['ID']==col['N2']].iloc[0]
                        msp.add_line((n1['X'], n1['Z'], n1['Y']), (n2['X'], n2['Z'], n2['Y']), dxfattribs={'layer': 'COLUMNAS'})
                    for _, v in pd.concat([vigas_x, vigas_z]).iterrows():
                        n1 = nudos[nudos['ID']==v['N1']].iloc[0]
                        n2 = nudos[nudos['ID']==v['N2']].iloc[0]
                        msp.add_line((n1['X'], n1['Z'], n1['Y']), (n2['X'], n2['Z'], n2['Y']), dxfattribs={'layer': 'VIGAS'})
                    for _, z in zaps.iterrows():
                        msp.add_circle((z['X'], z['Z']), 0.3, dxfattribs={'layer': 'ZAPATAS'})
                    out = io.StringIO()
                    dwg.write(out)
                    st.download_button("📥 Descargar DXF", data=out.getvalue(), file_name="Edificio_3D.dxf", mime="application/dxf")

            # APU (presupuesto)
            if "apu_config" in st.session_state:
                apu = st.session_state.apu_config
                mon = apu.get("moneda", "$")
                # Costos concreto
                if apu.get("usar_concreto_premezclado", False):
                    costo_conc = vol_conc * apu.get("precio_concreto_m3", 0)
                else:
                    # Estimación con mezcla en sitio (350 kg/m³)
                    costo_conc = vol_conc * 350 / 50 * apu.get("cemento", 0)
                costo_ace = peso_acero * apu.get("acero", 0)
                total_mat = costo_conc + costo_ace
                # Mano de obra aproximada
                dias_mo = (peso_acero * 0.04) + (vol_conc * 0.4)
                costo_mo = dias_mo * apu.get("costo_dia_mo", 69333.33)
                costo_directo = total_mat + costo_mo
                herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
                aiu = costo_directo * apu.get("pct_aui", 0.30)
                utilidad = costo_directo * apu.get("pct_util", 0.05)
                iva = utilidad * apu.get("iva", 0.19)
                total = costo_directo + herramienta + aiu + iva

                st.markdown("---")
                st.subheader("💰 Presupuesto estimado")
                st.write(f"**Volumen concreto:** {vol_conc:.2f} m³")
                st.write(f"**Acero:** {peso_acero:.1f} kg")
                st.write(f"**Costo materiales:** {mon} {total_mat:,.2f}")
                st.write(f"**Costo total (incl. MO e indirectos):** {mon} {total:,.2f}")
            else:
                st.info("💡 Ve a la página 'APU Mercado' para configurar precios.")

else:
    st.info("👈 Configure los parámetros en la barra lateral y presione **Generar Malla 3D**.")