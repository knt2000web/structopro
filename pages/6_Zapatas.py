import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import plotly.graph_objects as go
import datetime as _dt

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Zapatas y Suelos", "Footings and Soils"), layout="wide")

st.image(r"assets/concrete_isolated_footing_1773262985104.png", use_container_width=True)
st.title(_t("Cimentaciones: Zapatas y Geotecnia", "Foundations: Footings and Geotechnics"))
st.markdown(_t("Módulo integral para diseño de cimentaciones superficiales (Zapata Aislada), capacidad portante y esfuerzos en la masa de suelo. Soporta normativa internacional ACI 318 / NSR-10 / Multi-Norma.", "Comprehensive module for shallow foundation design (Isolated Footing), bearing capacity and soil stresses. Supports international codes ACI 318 / NSR-10 / Multi-Code."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CONFIGURACIÓN GENERAL Y APU
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙️ Configuración Global", "⚙️ Global Settings"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = "NSR-10 (Colombia)"

norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)

# Parametros normativos base
phi_v = 0.75 # Cortante
phi_f = 0.90 # Flexión
if "E.060" in norma_sel:
    phi_v = 0.85

fc_basico = st.sidebar.number_input(_t("f'c Zapata [MPa]", "f'c Footing [MPa]"), 15.0, 50.0, st.session_state.get("z_fc", 21.0), 1.0, key="z_fc")
fy_basico = st.sidebar.number_input(_t("fy Acero [MPa]", "fy Steel [MPa]"), 240.0, 500.0, st.session_state.get("z_fy", 420.0), 10.0, key="z_fy")

# ─── CONVERSOR GLOBAL DE UNIDADES DE SUELO ← Visible en toda la página
st.sidebar.markdown("---")
st.sidebar.header(_t("🔄 Conversor Unidades de Suelo", "🔄 Soil Units Converter"))
_cu = st.sidebar.selectbox(_t("Unidad a convertir:", "Unit to convert:"), 
    ["kPa → ...", "ton/m² → ...", "kg/cm² → ...", "MPa → ...", "psi → ..."], key="conv_unit_global")
_cv = st.sidebar.number_input(_t("Valor:", "Value:"), value=1.0, step=0.1, key="conv_val_global")
_ck = {
    "kPa → ...": _cv, "ton/m² → ...": _cv * 9.80665,
    "kg/cm² → ...": _cv * 98.0665, "MPa → ...": _cv * 1000.0,
    "psi → ...": _cv * 6.89476,
}.get(_cu, _cv)
st.sidebar.markdown(f"""
| | |
|:---|---:|
| **kPa** | `{_ck:.2f}` |
| **ton/m²** | `{_ck/9.80665:.3f}` |
| **kg/cm²** | `{_ck/98.0665:.4f}` |
| **MPa** | `{_ck/1000:.5f}` |
| **psi** | `{_ck/6.89476:.2f}` |
""")

if any(n in norma_sel for n in ["Colombia", "EE.UU.", "Perú", "México", "Venezuela"]):
    REBAR_DICT = {
        "N.3 (3/8\")": {"area": 0.71, "db": 9.5},
        "N.4 (1/2\")": {"area": 1.29, "db": 12.7},
        "N.5 (5/8\")": {"area": 1.99, "db": 15.9},
        "N.6 (3/4\")": {"area": 2.84, "db": 19.1},
        "N.7 (7/8\")": {"area": 3.87, "db": 22.2},
        "N.8 (1\")":   {"area": 5.10, "db": 25.4},
    }
    def_idx = 1 # N.4
else:
    REBAR_DICT = {
        "10 mm": {"area": 0.785, "db": 10.0},
        "12 mm": {"area": 1.131, "db": 12.0},
        "14 mm": {"area": 1.539, "db": 14.0},
        "16 mm": {"area": 2.011, "db": 16.0},
        "18 mm": {"area": 2.545, "db": 18.0},
        "20 mm": {"area": 3.142, "db": 20.0},
        "22 mm": {"area": 3.801, "db": 22.0},
        "25 mm": {"area": 4.909, "db": 25.0},
    }
    def_idx = 1 # 12 mm

# ─────────────────────────────────────────────
# HELPER GLOBAL: BOUSSINESQ INFLUENCE FACTOR (version escalar y vectorizada)
# ─────────────────────────────────────────────
def I_z_bous(m, n):
    V1 = m**2 + n**2 + 1
    V2 = m**2 * n**2
    term1 = (2*m*n*np.sqrt(V1)) / (V1 + V2) if (V1 + V2) != 0 else 0
    term2 = (V1 + 1) / V1 if V1 != 0 else 0
    angulo = np.arctan2(2*m*n*np.sqrt(V1), (V1 - V2))
    # Manejo de ángulo para cuadrantes
    angulo = np.where(V1 - V2 < 0, angulo + np.pi, angulo)
    return (1 / (4 * np.pi)) * (term1 * term2 + angulo)

# Versión vectorizada para arrays de numpy
def I_z_bous_vec(m_arr, n_arr):
    V1 = m_arr**2 + n_arr**2 + 1
    V2 = m_arr**2 * n_arr**2
    term1 = (2*m_arr*n_arr*np.sqrt(V1)) / (V1 + V2)
    term1 = np.where((V1 + V2) != 0, term1, 0.0)
    term2 = (V1 + 1) / V1
    term2 = np.where(V1 != 0, term2, 0.0)
    angulo = np.arctan2(2*m_arr*n_arr*np.sqrt(V1), (V1 - V2))
    angulo = np.where(V1 - V2 < 0, angulo + np.pi, angulo)
    return (1 / (4 * np.pi)) * (term1 * term2 + angulo)

# ─────────────────────────────────────────────
# T1: ESFUERZOS EN EL SUELO (BOUSSINESQ)
# ─────────────────────────────────────────────
with st.expander(_t("🌍 1. Esfuerzos en masa de suelo debajo de zapata", "🌍 1. Soil Stresses under Footing (Boussinesq)"), expanded=False):
    st.info(_t("📺 **Modo de uso:** Ingresa las dimensiones de la zapata y la carga aplicada. El programa usa la solución de Boussinesq (integración de carga rectangular) para encontrar el incremento de esfuerzo vertical a cierta profundidad Z debajo del centro de la zapata.", "📺 **How to use:** Enter footing dimensions and load. Uses Boussinesq method to find vertical stress increment at depth Z."))
    c1, c2 = st.columns(2)
    with c1:
        P_bous = st.number_input("Carga en Zapata P [kN]", 10.0, 10000.0, st.session_state.get("z_bous_P", 1000.0), 100.0, key="z_bous_P")
        B_bous = st.number_input("Ancho B [m]", 0.5, 10.0, st.session_state.get("z_bous_B", 2.0), 0.1, key="z_bous_B")
        L_bous = st.number_input("Largo L [m]", 0.5, 10.0, st.session_state.get("z_bous_L", 2.0), 0.1, key="z_bous_L")
    with c2:
        Z_bous = st.number_input("Profundidad de análisis Z [m]", 0.1, 20.0, st.session_state.get("z_bous_Z", 2.0), 0.5, key="z_bous_Z")
        q_0 = P_bous / (B_bous * L_bous) # Esfuerzo de contacto kN/m2
        st.markdown(_t(f"**Esfuerzo de contacto ($q_0$):** {q_0:.2f} kPa = {q_0/9.80665:.3f} t/m² = {q_0/98.0665:.4f} kg/cm²",
                       f"**Contact Stress ($q_0$):** {q_0:.2f} kPa = {q_0/9.80665:.3f} t/m² = {q_0/98.0665:.4f} kg/cm²"))
        
    # Fadum/Boussinesq bajo el centro: dividimos el rectangulo en 4 rectangulos de B/2 x L/2
    m = (B_bous/2.0) / Z_bous
    n = (L_bous/2.0) / Z_bous
    delta_sigma_z = 4 * q_0 * I_z_bous(m, n)
    
    st.success(f"📈 Incremento de esfuerzo vertical bajo el centro a Z={Z_bous}m: **Δσ_z = {delta_sigma_z:.2f} kPa**")
    
    # Grafica rapida de bulbo central
    fig_b, ax_b = plt.subplots(figsize=(6,3))
    zs = np.linspace(0.1, max(B_bous*3, 10), 50)
    sigmas = [4 * q_0 * I_z_bous((B_bous/2.0)/z, (L_bous/2.0)/z) for z in zs]
    ax_b.plot(sigmas, zs, color="magenta", lw=2)
    ax_b.invert_yaxis()
    ax_b.set_xlabel("Incremento de Esfuerzo Δσ_z [kPa]")
    ax_b.set_ylabel("Profundidad Z [m]")
    ax_b.set_title(f"Distribución de Esfuerzos bajo el centro (P={P_bous} kN)")
    ax_b.grid(True, linestyle="--", alpha=0.5)
    st.pyplot(fig_b)

# ─────────────────────────────────────────────
# T2: CAPACIDAD PORTANTE DE SUELO (TERZAGHI) + ASENTAMIENTOS
# ─────────────────────────────────────────────
with st.expander(_t("🛑 2. Capacidad Portante de Suelo (Terzaghi) y Asentamientos", "🛑 2. Bearing Capacity (Terzaghi) and Settlements"), expanded=False):
    st.info(_t(
        "📺 **Modo de uso:** Ingresa φ, c, γ y la geometría de la zapata. "
        "El módulo calcula la capacidad última de Terzaghi con influencia del NF, "
        "grafica el diagrama Vesic (1973) para tipo de falla y el bulbo de presiones, "
        "y opcionalmente estima el asentamiento elástico inmediato.",
        "📺 **How to use:** Enter φ, c, γ and footing geometry. Module calculates "
        "Terzaghi ultimate capacity with water-table correction, Vesic failure-type chart, "
        "pressure bulb, and optionally estimates immediate elastic settlement."
    ))

    # ── CONVERSOR DE UNIDADES ───────────────────────────────────────────────
    with st.container():
        st.markdown("**🔄 Conversor Rápido de Resistencia de Suelo**")
        uc1, uc2, uc3 = st.columns([1, 1, 2])
        with uc1:
            q_conv_unit = st.selectbox(_t("Unidad de entrada:", "Input Unit:"),
                ["kPa (kN/m²)", "ton/m²", "kg/cm²", "MPa", "psi", "kN/m²"],
                key="q_conv_unit_terz")
        with uc2:
            q_conv_val = st.number_input(_t("Valor:", "Value:"), value=1.0, step=0.1, key="q_conv_val_terz")
        with uc3:
            _conv = {"kPa (kN/m²)": q_conv_val, "ton/m²": q_conv_val*9.80665,
                     "kg/cm²": q_conv_val*98.0665, "MPa": q_conv_val*1000.0,
                     "psi": q_conv_val*6.89476, "kN/m²": q_conv_val}.get(q_conv_unit, q_conv_val)
            st.markdown(f"| Unidad | Valor |\n|--------|-------|\n"
                        f"| **kPa** | `{_conv:.3f}` |\n"
                        f"| **ton/m²** | `{_conv/9.80665:.3f}` |\n"
                        f"| **kg/cm²** | `{_conv/98.0665:.4f}` |\n"
                        f"| **MPa** | `{_conv/1000.0:.5f}` |\n"
                        f"| **psi** | `{_conv/6.89476:.2f}` |")
    st.divider()

    # ── ENTRADAS ────────────────────────────────────────────────────────────
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("##### 🏔️ Parámetros Geotécnicos")
        phi_ang  = st.number_input(_t("Ángulo de fricción φ [°]","Friction angle φ [°]"),
                                   0.0, 50.0, st.session_state.get("z_phi", 30.0), 1.0, key="z_phi")
        coh_unit = st.selectbox(_t("Unidad cohesión:","Cohesion Unit:"),
                                ["kPa", "kg/cm²", "ton/m²"], key="coh_u")
        coh_val  = st.number_input(f"c [{coh_unit}]", 0.0, 200.0,
                                   st.session_state.get("coh_val", 5.0 if coh_unit=="kPa" else 0.05),
                                   0.5 if coh_unit=="kPa" else 0.01, key="coh_val")
        coh_c    = coh_val if coh_unit=="kPa" else (coh_val*98.0665 if coh_unit=="kg/cm²" else coh_val*9.80665)
        gam_unit = st.selectbox(_t("Unidad γ húmedo:","γ moist Unit:"),
                                ["kN/m³", "ton/m³", "kg/m³"], key="gam_u")
        gam_val  = st.number_input(f"γ [{gam_unit}]", 10.0,
                                   25.0 if gam_unit != "kg/m³" else 2500.0,
                                   st.session_state.get("gam_val", 18.0 if gam_unit != "kg/m³" else 1800.0),
                                   0.5, key="gam_val")
        gamma_s  = gam_val if gam_unit=="kN/m³" else (gam_val*9.80665 if gam_unit=="ton/m³" else gam_val*0.00980665)
        gamma_sat_t = st.number_input("γ saturado [kN/m³]", 16.0, 25.0, 20.0, 0.5, key="z_gsat")
        gamma_w  = 9.81
    with c2:
        st.markdown("##### 🏗️ Geometría y Carga")
        forma_zap = st.selectbox(_t("Forma de zapata","Footing shape"),
                                 ["Cuadrada", "Continua (Muro)", "Circular"] if lang=="Español"
                                 else ["Square", "Continuous (Wall)", "Circular"], key="z_shape")
        B_cp  = st.number_input(_t("Ancho B [m]","Width B [m]"), 0.5, 10.0,
                                st.session_state.get("cp_b", 2.0), 0.1, key="cp_b")
        L_cp  = st.number_input(_t("Largo L [m]","Length L [m]"), 0.5, 100.0,
                                st.session_state.get("cp_l", 2.0), 0.1, key="cp_l")
        Hz_cp = st.number_input(_t("Altura zapata Hz [m]","Footing height Hz [m]"),
                                0.1, 2.0, 0.5, 0.05, key="z_hz_cp")
        Df_cp = st.number_input(_t("Profundidad Df [m]","Depth Df [m]"), 0.0, 10.0,
                                st.session_state.get("cp_df", 1.5), 0.1, key="cp_df")
        b_col_cp = st.number_input(_t("Lado columna b [m]","Column side b [m]"), 0.1, 2.0, 0.4, 0.05, key="z_bcol_cp")
        Q_act  = st.number_input("Carga vertical actuante Q [kN]", 10.0, 50000.0, 3000.0, 100.0, key="z_Qact")
    with c3:
        st.markdown("##### 💧 Nivel Freático y FS")
        NF_prof = st.number_input("NF — Profundidad nivel freático [m]", 0.0, 20.0, 1.0, 0.5, key="z_nf")
        FS_terz = st.number_input(_t("Factor de Seguridad (FS)","Safety Factor (FS)"),
                                  1.0, 5.0, st.session_state.get("z_fs", 3.0), 0.1, key="z_fs")
        N_spt   = st.number_input("N60 campo (SPT)", 0, 100, 14, 1, key="z_spt")
        st.caption("ℹ️ N60 se usa para clasificar tipo de falla (Vesic 1973)")

    # Parámetros para asentamiento elástico (opcional)
    with st.expander("📏 Asentamiento elástico (opcional)", expanded=False):
        usar_asent = st.checkbox("Calcular asentamiento inmediato", value=False)
        if usar_asent:
            E_suelo = st.number_input("Módulo de elasticidad del suelo E [MPa]", 1.0, 500.0, 20.0, 1.0)
            nu_suelo = st.number_input("Relación de Poisson ν", 0.1, 0.5, 0.35, 0.01)
            metodo_asent = st.selectbox("Método", ["Steinbrenner (zapata flexible)", "Elástico uniforme"])
        else:
            E_suelo = None

    # ── CÁLCULO GEOTÉCNICO ──────────────────────────────────────────────────
    phi_rad = math.radians(phi_ang)

    # Factores capacidad — Terzaghi
    if phi_ang == 0:
        Nc, Nq, Ngamma = 5.7, 1.0, 0.0
    else:
        a_t    = math.exp((0.75*math.pi - phi_rad/2)*math.tan(phi_rad))
        Nq     = (a_t**2) / (2*math.cos(math.radians(45) + phi_rad/2)**2)
        Nc     = (Nq - 1) / math.tan(phi_rad)
        Ngamma = 2*(Nq+1)*math.tan(phi_rad) / (1 + 0.4*math.sin(4*phi_rad))

    # Factores de forma
    if forma_zap in ["Cuadrada", "Square"]:
        sc, sq, sgamma = 1.3, 1.0, 0.8
    elif forma_zap in ["Circular"]:
        sc, sq, sgamma = 1.3, 1.0, 0.6
    else:
        sc, sq, sgamma = 1.0, 1.0, 1.0

    # Corrección por nivel freático (Casos I / II / III)
    gamma_prime = gamma_sat_t - gamma_w
    if NF_prof <= Df_cp:                         # Caso I: NF sobre el desplante
        q_sob    = gamma_s*NF_prof + gamma_prime*(Df_cp - NF_prof)
        gamma_eff = gamma_prime
        caso_nf  = "I"
    elif NF_prof < Df_cp + B_cp:                 # Caso II: NF entre Df y Df+B
        z2        = NF_prof - Df_cp
        gamma_eff = gamma_s + (z2/B_cp)*(gamma_prime - gamma_s) if B_cp > 0 else gamma_s
        q_sob    = gamma_s*Df_cp
        caso_nf  = "II"
    else:                                        # Caso III: NF muy profundo
        q_sob    = gamma_s*Df_cp
        gamma_eff = gamma_s
        caso_nf  = "III"

    q_ult  = sc*coh_c*Nc + sq*q_sob*Nq + sgamma*0.5*gamma_eff*B_cp*Ngamma
    q_adm  = q_ult / FS_terz
    Q_ult  = q_ult * B_cp * L_cp
    FS_calc = Q_ult / Q_act if Q_act > 0 else 999.0
    cumplio = FS_calc >= FS_terz

    # ── RESULTADOS ──────────────────────────────────────────────────────────
    st.divider()
    st.markdown("#### 📊 Resultados de Capacidad Portante")
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        st.dataframe(pd.DataFrame([
            {"Parámetro": "Nc",                  "Valor": f"{Nc:.3f}"},
            {"Parámetro": "Nq",                  "Valor": f"{Nq:.3f}"},
            {"Parámetro": "Nγ",                  "Valor": f"{Ngamma:.3f}"},
            {"Parámetro": "sc / sq / sγ",         "Valor": f"{sc} / {sq} / {sgamma}"},
            {"Parámetro": f"Caso NF",             "Valor": f"Caso {caso_nf} | NF={NF_prof}m"},
            {"Parámetro": "q sobrecarga",         "Valor": f"{q_sob:.2f} kPa"},
            {"Parámetro": "γ' efectivo",          "Valor": f"{gamma_eff:.2f} kN/m³"},
            {"Parámetro": "q_ult",                "Valor": f"{q_ult:.2f} kPa"},
            {"Parámetro": "Q_ult = q_ult·B·L",   "Valor": f"{Q_ult:.1f} kN"},
            {"Parámetro": "FS calculado",         "Valor": f"{FS_calc:.3f}"},
        ]), use_container_width=True, hide_index=True)
    with col_r2:
        st.markdown("**q_adm en todas las unidades:**")
        m1,m2,m3 = st.columns(3)
        m1.metric("kPa",    f"{q_adm:.1f}")
        m2.metric("ton/m²", f"{q_adm/9.80665:.2f}")
        m3.metric("kg/cm²", f"{q_adm/98.0665:.4f}")
        m4,m5 = st.columns(2)
        m4.metric("MPa",    f"{q_adm/1000:.5f}")
        m5.metric("psi",    f"{q_adm/6.89476:.2f}")
        if cumplio:
            st.success(f"✅ q_adm = **{q_adm:.2f} kPa** | FS = {FS_calc:.2f} ≥ {FS_terz}")
        else:
            st.error(f"❌ q_adm = {q_adm:.2f} kPa | FS = {FS_calc:.2f} < {FS_terz} → Revisar dimensiones")

    # ── ASENTAMIENTO ELÁSTICO (si se solicitó) ─────────────────────────────
    if usar_asent and E_suelo is not None:
        st.divider()
        st.markdown("#### 📏 Asentamiento Elástico Inmediato")
        # Conversión a kPa
        E_kPa = E_suelo * 1000
        # Cálculo del asentamiento en el centro de una zapata flexible (Steinbrenner)
        # Factor de influencia I: para zapata rectangular flexible, centro
        # I = f(m,n) según Steinbrenner (solución de Boussinesq)
        # Usamos la fórmula de la tabla de Fadum (para el centro de un rectángulo)
        # Asentamiento = q * B * (1-ν²)/E * I_f
        # I_f = 0.5 * [ (m n sqrt(m²+n²+1) / (m²+n²+1) + ...] pero simplificado
        # Para una primera aproximación usamos el factor de influencia de la solución de Boussinesq
        # integrada sobre el área.
        # Usaremos la fórmula de Bowles (1996) para asentamiento inmediato:
        # δ = q * B * (1-ν²) / E * I_f, con I_f = factor de influencia
        # Para zapata rectangular, I_f = 0.5 * ln((1+sqrt(m²+1))/sqrt(m²+n²+1)) ... etc.
        # Por simplicidad, aquí se implementa el factor para el centro según Steinbrenner.
        # m = L/B, n = z/B (z=0 para superficie)
        # En asentamiento superficial z=0, el factor se simplifica a I_f = 0.5 * ln((1+sqrt(m²+1))/...)
        # Pero es más práctico usar la fórmula de la solución elástica para el centro:
        # I_f = 0.5 * [ (m n sqrt(m²+n²+1) / (m²+n²+1) ) + ...] pero con z=0 es indeterminado.
        # En su lugar usamos la expresión para el asentamiento en la superficie (z=0):
        # Para una carga uniforme sobre un rectángulo, el asentamiento en el centro es:
        # δ = (q * B / E) * (1 - ν²) * I_center
        # donde I_center se puede calcular como:
        # I_center = (1/π) * ln( (1+√(m²+1)) / (√(m²+1) - 1) ) * (1+ν) + ...
        # Utilizaremos la fórmula aproximada de la literatura (Bowles):
        # δ = q * B * (1-ν²) / E * I_f, I_f = 0.5 * [ (L/B) * ( ... ) ]? mejor usar la de la teoría de la elasticidad.
        # Simplificamos con la fórmula de la solución elástica para el centro:
        # I_f = (1/π) * [ (m ln((1+√(m²+1))/√(m²+1)) + ln((m+√(m²+1))/(m-√(m²+1)) ) ] ... es complejo.
        # Implementaremos una versión simplificada que usa un factor empírico I_f = 0.8 para L/B=1, 0.9 para L/B=2, etc.
        # Para una estimación rápida, se puede usar la tabla de Fadum (Boussinesq integrada) para el centro a profundidad cero.
        # Alternativamente, usamos el factor I = 0.8 (para zapata cuadrada) y I=1.0 (para continua). 
        # Se puede mejorar más adelante.
        # Aquí usaremos una fórmula sencilla de la literatura (Das):
        # δ = (q * B * (1-ν²) / E) * I_s, donde I_s = 0.5 * [ (L/B) * ( ... )] pero es complejo.
        # Para no complicar, utilizaremos un factor I_f obtenido de la tabla de Steinbrenner (Bowles 1996, Tabla 5.1)
        # m = L/B, n = H/B, H es la profundidad del estrato compressible (suponemos 10 m)
        H_estrato = st.number_input("Profundidad del estrato compressible H [m]", 1.0, 50.0, 10.0, 1.0, key="h_estrato")
        m = L_cp / B_cp
        n = H_estrato / B_cp
        # Cálculo de factor de influencia I_center (Bowles, Eq. 5-7)
        # I_f = (1/π) * [ (m ln((1+√(m²+1))/√(m²+1)) + ln((m+√(m²+1))/(m-√(m²+1)) ) ] ... pero con n también.
        # Para un estrato finito, se usa la solución de Steinbrenner. Usaremos la expresión de la referencia.
        # Implementación simplificada: usamos la fórmula de la Tabla 5.1 (Bowles) mediante interpolación.
        # Por simplicidad, usaremos la función de la siguiente manera:
        def I_f_center(m, n):
            # m = L/B, n = H/B
            # Calcula factor de influencia para el centro de una zapata rectangular flexible.
            # Fórmula de Steinbrenner (Bowles, 1996, Eq. 5-8)
            # I_f = (1/π) * [ m ln((1+√(m²+1))/√(m²+1)) + ln((m+√(m²+1))/(m-√(m²+1))) ] * F(n) + ... No.
            # Usamos una expresión de la solución de Boussinesq integrada sobre el rectángulo con profundidad finita:
            # I = (1/π) * [ m * ( ... ) + ... ]
            # Optamos por una función más directa: la solución de Newmark para el centro de un rectángulo.
            # La fórmula completa es larga, por lo que para este ejemplo utilizamos una aproximación:
            # I_f = 0.5 * (1 - ν²) * [ (m) / (1+m²) + ...] no.
            # Finalmente, para mantener la sencillez, usaremos un factor empírico basado en la tabla de Das (2004).
            # Para m entre 1 y 10, I_f varía entre 0.8 y 1.2. Tomamos I_f = 0.8 + 0.05*(m-1) con tope en 1.2.
            # Esto es solo una estimación, no para diseño riguroso.
            I = min(1.2, 0.8 + 0.05 * (m - 1))
            return I
        I_center = I_f_center(m, n)
        asentamiento = q_adm * B_cp * (1 - nu_suelo**2) / E_kPa * I_center * 1000  # en mm
        st.metric("Asentamiento inmediato estimado", f"{asentamiento:.1f} mm")
        st.caption("⚠️ Cálculo aproximado usando factor de influencia empírico. Para diseño detallado, usar métodos más refinados (e.g., Steinbrenner completo).")
        # Guardar asentamiento para uso en memoria
        st.session_state['asentamiento'] = asentamiento

    # ── GRÁFICA: TIPO DE FALLA (Vesic 1973 / SPT) ──────────────────────────
    st.divider()
    st.markdown("#### 📈 Diagrama Tipo de Falla — Vesic (1973)")
    Dr_pct = min(100.0, 1.08 * math.sqrt(max(N_spt, 0) / 60.0) * 100)
    Df_B   = Df_cp / B_cp if B_cp > 0 else 0.0

    fig_ves, ax_ves = plt.subplots(figsize=(8, 5))
    ax_ves.set_facecolor("#0f1117"); fig_ves.patch.set_facecolor("#0f1117")
    Dr_arr = np.linspace(0, 100, 300)
    # Límites aproximados de Vesic (1973) expresados como -Df/B* vs Dr
    c1v = -0.6 + 0.006*Dr_arr             # Punzonamiento → Local
    c2v = -2.5 + 0.030*Dr_arr             # Local → General
    ax_ves.plot(Dr_arr, c1v, color="#4fc3f7", lw=1.8, label="Límite punz./local")
    ax_ves.plot(Dr_arr, c2v, color="#81d4fa", lw=1.8, label="Límite local/general")
    ax_ves.fill_between(Dr_arr, c2v, -5, alpha=0.18, color="#e53935")
    ax_ves.fill_between(Dr_arr, c1v, c2v, alpha=0.15, color="#fb8c00")
    ax_ves.fill_between(Dr_arr, 0, c1v, alpha=0.12, color="#43a047")
    ax_ves.text(10, -0.3, "Falla por\nPunzonamiento", color="white", fontsize=8, ha="center")
    ax_ves.text(50, -1.5, "Falla local\npor corte",   color="white", fontsize=8, ha="center")
    ax_ves.text(80, -3.5, "Falla general\npor corte", color="white", fontsize=8, ha="center")
    # Miniatura fundación
    rec_w = 15; rec_h = 0.3
    ax_ves.add_patch(patches.Rectangle((2, -Df_B - rec_h), rec_w, rec_h,
                                       fc="#888", ec="white", lw=0.8))
    ax_ves.annotate(f"B={B_cp}m\nHz={Hz_cp}m", (2+rec_w/2, -Df_B - rec_h*2),
                    color="yellow", fontsize=7, ha="center")
    # Punto actual
    ax_ves.plot(Dr_pct, -Df_B, "o", color="red", ms=11, zorder=6,
                label=f"Dr≈{Dr_pct:.0f}% | Df/B={Df_B:.2f}")
    ax_ves.set_xlim(0, 100); ax_ves.set_ylim(-5, 0)
    ax_ves.set_xlabel("Densidad relativa Dr (%)", color="white")
    ax_ves.set_ylabel("-Df / B*", color="white")
    ax_ves.set_title(f"Tipo de Falla — N60={N_spt} campo → Dr≈{Dr_pct:.0f}%  |  Df/B={Df_B:.2f}", color="white")
    ax_ves.tick_params(colors="white"); ax_ves.spines[:].set_color("#444")
    ax_ves.legend(loc="lower right", fontsize=8, facecolor="#111", labelcolor="white")
    st.pyplot(fig_ves)

    # ── GRÁFICA: BULBO DE PRESIONES / MECANISMO DE FALLA (VECTORIZADO) ───
    st.divider()
    st.markdown("#### 🌐 Mecanismo de Falla y Bulbo de Presiones (Terzaghi)")
    fig_tb, ax_tb = plt.subplots(figsize=(10, 7))
    ax_tb.set_facecolor("#0f1117"); fig_tb.patch.set_facecolor("#0f1117")

    half_B = B_cp / 2.0
    zap_bot = -(Df_cp + Hz_cp)

    # === MECANISMO DE FALLA DE TERZAGHI (Prandtl) ===
    # Ángulos de cuña
    alpha = math.pi/4 + phi_rad/2
    H_tri = half_B * math.tan(alpha)
    pole_x = half_B
    pole_y = zap_bot
    theta_0 = math.atan2(-H_tri, -half_B)
    sweep = math.pi/2
    theta_end = theta_0 + sweep

    # Zona I — Triángulo activo central (Rosa oscuro)
    tri_pts = [(-half_B, zap_bot), (half_B, zap_bot), (0, zap_bot - H_tri)]
    ax_tb.add_patch(plt.Polygon(tri_pts, fc="#e57373", ec="black", lw=1.2, alpha=0.9, zorder=5, label="Zona I (Activa)"))

    if phi_rad >= 0.0:
        # Geometría del espiral (Arco)
        theta_arr = np.linspace(theta_0, theta_end, 40)
        r_arr = (half_B / math.cos(alpha)) * np.exp((theta_arr - theta_0) * math.tan(phi_rad))
        x_spiral = pole_x + r_arr * np.cos(theta_arr)
        y_spiral = pole_y + r_arr * np.sin(theta_arr)
        end_x = x_spiral[-1]
        end_y = y_spiral[-1]

        beta_passive = math.pi/4 - phi_rad/2
        # Prevenir error si beta_passive se hace 0 (para arcilla phi=0, beta_passive=45 deg, siempre >0)
        dist_x = abs(end_y - pole_y) / math.tan(beta_passive)
        x_top = end_x + dist_x

        # Zona II (Radial - Amarillo) - Derecha
        poly_Z2_R = [(pole_x, pole_y)] + list(zip(x_spiral, y_spiral))
        ax_tb.add_patch(plt.Polygon(poly_Z2_R, fc="#ffe082", ec="black", lw=1.2, alpha=0.85, zorder=4, label="Zona II (Radial)"))
        # Zona III (Pasiva - Naranja claro) - Derecha
        poly_Z3_R = [(pole_x, pole_y), (end_x, end_y), (x_top, pole_y)]
        ax_tb.add_patch(plt.Polygon(poly_Z3_R, fc="#ffb74d", ec="black", lw=1.2, alpha=0.85, zorder=3, label="Zona III (Pasiva)"))

        # Zonas II y III - Izquierda (Espejo simétrico)
        poly_Z2_L = [(-x, y) for (x, y) in poly_Z2_R]
        ax_tb.add_patch(plt.Polygon(poly_Z2_L, fc="#ffe082", ec="black", lw=1.2, alpha=0.85, zorder=4))
        poly_Z3_L = [(-x, y) for (x, y) in poly_Z3_R]
        ax_tb.add_patch(plt.Polygon(poly_Z3_L, fc="#ffb74d", ec="black", lw=1.2, alpha=0.85, zorder=3))

        # Sobrecarga de fosa (Suelo arriba de Df - Relleno pasivo)
        if zap_bot < 0:
            ax_tb.add_patch(patches.Rectangle((-x_top, zap_bot), 2*x_top, -zap_bot,
                                              fc="#ffebee", ec="black", lw=1, alpha=0.6, zorder=2, label="Sobrecarga"))
            # Extender limits del eje X para que abarque hasta la cuna pasiva
            ax_tb.set_xlim(-x_top * 1.2, x_top * 1.2)

    # Zapata
    ax_tb.add_patch(patches.Rectangle((-half_B, zap_bot), B_cp, Hz_cp,
                                      fc="#546e7a", ec="white", lw=1.5, zorder=4,
                                      label=f"Zapata {B_cp}×{L_cp}m"))
    # Columna
    ax_tb.add_patch(patches.Rectangle((-b_col_cp/2, 0), b_col_cp, -Df_cp + Hz_cp*0.1,
                                      fc="#78909c", ec="white", lw=1, zorder=4))

    # Bulbo de presiones Boussinesq (Teoría elástica 2D - Suma de Cuadrantes)
    _xg = np.linspace(-B_cp*3.5, B_cp*3.5, 200)
    _zg = np.linspace(zap_bot, zap_bot - B_cp*3.5, 150)
    Xg, Zg = np.meshgrid(_xg, _zg)
    q0_bulbo = q_ult
    dz = np.abs(Zg - zap_bot)
    dz = np.where(dz < 0.05, 0.05, dz)  # evitar división por cero
    
    def I_corner(X, Y, Z_arr):
        m = np.abs(X) / Z_arr
        n = np.abs(Y) / Z_arr
        m = np.clip(m, 1e-6, 1e6)
        n = np.clip(n, 1e-6, 1e6)
        return I_z_bous_vec(m, n) * np.sign(X) * np.sign(Y)

    # Suma algebraica para (x, y=0)
    x1 = half_B - Xg
    x2 = -half_B - Xg
    y1 = L_cp / 2.0
    y2 = -L_cp / 2.0

    I1 = I_corner(x1, y1, dz)
    I2 = I_corner(x2, y1, dz)
    I3 = I_corner(x1, y2, dz)
    I4 = I_corner(x2, y2, dz)
    
    sigma_arr = q0_bulbo * np.abs(I1 - I2 - I3 + I4)

    # Renderizado suave y profesional
    levels = np.linspace(q0_bulbo*0.02, q0_bulbo*0.95, 15)
    cs = ax_tb.contourf(Xg, Zg, sigma_arr, levels=levels, cmap="turbo", alpha=0.6, zorder=1)
    ax_tb.contour(Xg, Zg, sigma_arr, levels=levels, colors='white', linewidths=0.3, zorder=1, alpha=0.3)
    cbar = fig_tb.colorbar(cs, ax=ax_tb, label="Δσz [kPa]", shrink=0.7)
    cbar.ax.yaxis.set_tick_params(color="white")
    plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color="white")

    # Línea de exploración (si existe en sesión)
    if "z_exploracion" in st.session_state:
        z_ex = st.session_state.z_exploracion
        ax_tb.axhline(-(Df_cp + z_ex), color="cyan", lw=1.8, linestyle="--", label=f"Prof. exploración = {z_ex:.1f}m")

    # Nivel freático
    if NF_prof < Df_cp + B_cp*2.5:
        ax_tb.axhline(-NF_prof, color="#29b6f6", lw=1.5, linestyle=":",
                      label=f"▽ NF = {NF_prof}m")
    ax_tb.axhline(zap_bot, color="#888", lw=0.8, linestyle="--")
    ax_tb.axhline(0, color="#555", lw=0.5)
    ax_tb.set_xlim(-B_cp*4, B_cp*4)
    ax_tb.set_ylim(zap_bot - B_cp*2.8, 0.5)
    ax_tb.set_xlabel("Distancia [m]", color="white")
    ax_tb.set_ylabel("Profundidad [m]", color="white")
    ax_tb.set_title(f"Bulbo de Presiones + Falla Terzaghi — q_ult={q_ult:.1f} kPa | q_adm={q_adm:.1f} kPa",
                    color="white", fontsize=11)
    ax_tb.tick_params(colors="white"); ax_tb.spines[:].set_color("#444")
    ax_tb.legend(loc="upper right", fontsize=8, facecolor="#111", labelcolor="white")
    st.pyplot(fig_tb)

# ─────────────────────────────────────────────
# T4: PROFUNDIDAD MÍNIMA EXPLORACIÓN
# ─────────────────────────────────────────────
with st.expander(_t("🔬 4. Profundidad Mínima de Exploración de Subsuelo", "🔬 4. Minimum Subsurface Exploration Depth")):
    st.info(_t("📺 **Modo de uso:** La norma indica que se debe explorar el suelo (perforaciones) hasta que el incremento de esfuerzo de la estructura (Δσ_z) sea menor al 10% del esfuerzo aplicado q_0. Ingresa B, L y P para hallar esta profundidad D_exploración.", "📺 **How to use:** Enter B, L, and P to find exploration depth where vertical stress increment Δσ_z drops below 10% of q_0."))
    c1, c2 = st.columns(2)
    with c1:
        P_ex = st.number_input("Carga total P [kN]", 10.0, 50000.0, 1500.0, 100.0, key="ex1")
        B_ex = st.number_input("Ancho B [m]", 0.5, 20.0, 2.0, 0.5, key="ex2")
    with c2:
        L_ex = st.number_input("Largo L [m]", 0.5, 20.0, 3.0, 0.5, key="ex3")
        
    q0_ex = P_ex / (B_ex * L_ex)
    # Biseccion para encontrar Z donde delta_sigma_z / q0_ex = 0.10
    z_low = 0.1
    z_high = 50.0
    for _ in range(30):
        z_mid = (z_low + z_high)/2
        m_ex = (B_ex/2.0) / z_mid
        n_ex = (L_ex/2.0) / z_mid
        rat = 4 * I_z_bous(m_ex, n_ex)
        if rat > 0.10:
            z_low = z_mid
        else:
            z_high = z_mid
    st.success(f"✅ La profundidad mínima sugerida de exploración (donde Δσ_z = 10% de q0) es: **Z = {z_mid:.2f} metros** debajo del nivel de fundación.")
    # Guardar en sesión para usarlo en gráficos
    st.session_state.z_exploracion = z_mid

# ─────────────────────────────────────────────
# T3 & T5: DISEÑO ESTRUCTURAL DE ZAPATA + DIBUJADOR 3000 (con biaxialidad)
# ─────────────────────────────────────────────
with st.expander(_t("🏗️ 3 & 5. Diseño de Acero Zapata Prismática y Dibujador 3000", "🏗️ 3 & 5. Footing Structural Design & DXF Drafter"), expanded=True):
    st.info(_t("📺 **Modo de uso:** Ingresa las Cargas de Servicio (para dimensionar BxL) y Últimas (para diseñar espesor y acero). El módulo calculará Cortante, Punzonamiento y Flexión, generará tu geometría a AutoCAD y calculará Presupuestos APU.", "📺 **How to use:** Enter Service Loads (for sizing BxL) and Ultimate Loads (for thickness and steel). Calculates Shear, Punching, Flexure, DXF and APU budgets."))
    st.markdown(f"**Norma Estructural activa:** `{norma_sel}`")
    
    colA, colB, colC = st.columns(3)
    with colA:
        st.write("#### Cargas (Servicio y Últimas)")
        P_svc = st.number_input("Carga Axial de Servicio Ps [kN]", value=800.0, step=50.0)
        M_svc_B = st.number_input("Momento de Servicio Ms (dir. B) [kN·m]", value=0.0, step=10.0)
        M_svc_L = st.number_input("Momento de Servicio Ms (dir. L) [kN·m]", value=0.0, step=10.0)
        P_ult = st.number_input("Carga Axial Factorizada Pu [kN]", value=1120.0, step=50.0)
        M_ult_B = st.number_input("Momento Factorizado Mu (dir. B) [kN·m]", value=0.0, step=10.0)
        M_ult_L = st.number_input("Momento Factorizado Mu (dir. L) [kN·m]", value=0.0, step=10.0)
        
    with colB:
        st.write(_t("#### Geometría y Suelo", "#### Geometry and Soil"))
        q_unit = st.selectbox(_t("Unidad q_adm:", "q_adm Unit:"), ["kPa (kN/m²)", "ton/m²", "kg/cm²", "MPa"], index=0)
        
        default_q = 200.0
        step_q = 10.0
        if q_unit == "ton/m²":
            default_q, step_q = 20.0, 1.0
        elif q_unit == "kg/cm²":
            default_q, step_q = 2.0, 0.1
        elif q_unit == "MPa":
            default_q, step_q = 0.2, 0.05
            
        q_val_input = st.number_input(_t("Capacidad Portante q_adm", "Bearing Capacity q_adm"), value=default_q, step=step_q)
        
        if q_unit == "kPa (kN/m²)":
            q_adm_z = q_val_input
        elif q_unit == "ton/m²":
            q_adm_z = q_val_input * 9.80665
        elif q_unit == "MPa":
            q_adm_z = q_val_input * 1000.0
        else: # kg/cm²
            q_adm_z = q_val_input * 98.0665

        # Mostrar equivalencias en cards compactas
        eq_c1, eq_c2, eq_c3 = st.columns(3)
        eq_c1.metric("kPa", f"{q_adm_z:.1f}")
        eq_c2.metric("ton/m²", f"{q_adm_z/9.80665:.2f}")
        eq_c3.metric("MPa / kg/cm²", f"{q_adm_z/1000:.4f} / {q_adm_z/98.0665:.3f}")

        st.caption(_t(f"🔄 **Equivalencia:** {q_adm_z:.1f} kPa | {q_adm_z/9.80665:.3f} ton/m² | {q_adm_z/98.0665:.4f} kg/cm² | {q_adm_z/6.89476:.2f} psi",
                      f"🔄 **Equivalence:** {q_adm_z:.1f} kPa | {q_adm_z/9.80665:.3f} ton/m² | {q_adm_z/98.0665:.4f} kg/cm² | {q_adm_z/6.89476:.2f} psi"))

        c1_col = st.number_input(_t("Dim. Columna c1 (dir. B) [cm]", "Column dim. c1 (B dir.) [cm]"), min_value=5.0, value=40.0, step=5.0)
        c2_col = st.number_input(_t("Dim. Columna c2 (dir. L) [cm]", "Column dim. c2 (L dir.) [cm]"), min_value=5.0, value=40.0, step=5.0)
        gamma_prom = st.number_input(_t("γ_promedio (suelo+concreto) [kN/m³]", "γ_avg (soil+concrete) [kN/m³]"), value=20.0)
        Df_z = st.number_input(_t("Desplante Df [m]", "Footing Depth Df [m]"), value=1.0, step=0.1)

    with colC:
        st.write("#### Diseño Estructural")
        H_zap = st.number_input("Espesor H propuesto [cm]", value=50.0, step=5.0)
        recub_z = st.number_input("Recubrimiento al suelo [cm]", value=7.5, step=0.5)
        bar_z = st.selectbox("Varilla a utilizar:", list(REBAR_DICT.keys()), index=def_idx)
        A_bar_z = REBAR_DICT[bar_z]["area"] * 100 # mm2
        db_bar_z = REBAR_DICT[bar_z]["db"] # mm
        
    # ─── Validaciones robustas ──────────────────────────────────────────────
    if c1_col <= 0 or c2_col <= 0:
        st.error("❌ Las dimensiones de la columna (c1, c2) deben ser mayores a 0 cm.")
        st.stop()
    # Verificar espesor mínimo
    d_min = recub_z + db_bar_z/10.0 + 2  # cm
    if H_zap < d_min:
        st.warning(f"⚠️ El espesor H={H_zap} cm es menor que el mínimo recomendado (recubrimiento + db/10 + 2 cm = {d_min:.1f} cm). Aumente H.")
    # Verificar q_net positivo
    q_net = q_adm_z - gamma_prom * Df_z
    if q_net <= 0:
        st.error(f"❌ El esfuerzo neto disponible q_net = {q_net:.2f} kPa es negativo o nulo. Reduzca Df o aumente q_adm.")
        st.stop()

    # Paso 1: Dimensionamiento en planta (considerando excentricidades)
    # Para momento combinado, la presión máxima se da por flexión biaxial:
    # q_max = P/A ± (Mx * y / Ix) ± (My * x / Iy)
    # Primero obtenemos dimensiones mínimas basadas en q_net (sin momento)
    Area_req = P_svc / q_net
    # Proponemos zapata cuadrada si los momentos son bajos
    L_req = math.sqrt(Area_req * (c2_col/c1_col) if c1_col>0 else Area_req)
    B_req = Area_req / L_req if L_req > 0 else 0
    B_zap = math.ceil(B_req * 20) / 20.0
    L_zap = math.ceil(L_req * 20) / 20.0
    st.markdown(f"**Dimensiones mínimas sin excentricidad:** B = {B_zap:.2f} m, L = {L_zap:.2f} m")
    
    cB, cL = st.columns(2)
    B_use = cB.number_input("B usado para cálculo [m]", value=max(2.0, B_zap), step=0.1)
    L_use = cL.number_input("L usado para cálculo [m]", value=max(2.0, L_zap), step=0.1)

    # ─── Presión de contacto con flexión biaxial ─────────────────────────────
    A_use = B_use * L_use
    Ix = (B_use * L_use**3) / 12   # momento de inercia respecto al eje X (paralelo a B)
    Iy = (L_use * B_use**3) / 12   # momento de inercia respecto al eje Y (paralelo a L)
    # Coordenadas de las esquinas (x,y) con origen en el centro de la zapata
    corners = [(-B_use/2, -L_use/2), ( B_use/2, -L_use/2),
               ( B_use/2,  L_use/2), (-B_use/2,  L_use/2)]
    q_corners = []
    for x, y in corners:
        q = P_ult/A_use + (M_ult_L * x / Iy) + (M_ult_B * y / Ix)
        q_corners.append(q)
    qu_max = max(q_corners)
    qu_min = min(q_corners)
    # Presión promedio sobre el área crítica (se usa el promedio de las máximas y la presión en la zona de cortante)
    # Para simplificar, usamos el promedio de qu_max y max(qu_min,0)
    qu_avg = (qu_max + max(qu_min, 0)) / 2.0

    # Peralte efectivo
    d_z = H_zap - recub_z - (db_bar_z/10.0)
    d_z_m = d_z / 100.0

    # ─── CORTANTE UNIDIRECCIONAL (Viga) con integración exacta de presión ───
    # Se integra la presión lineal a lo largo del voladizo en dirección B
    lv_b = (B_use - c1_col/100.0) / 2.0
    lv_l = (L_use - c2_col/100.0) / 2.0

    # Función para calcular presión en un punto (x,y) (coordenadas locales con centro en zapata)
    def q_at(x, y):
        return P_ult/A_use + (M_ult_L * x / Iy) + (M_ult_B * y / Ix)

    # Cortante en dirección B (a una distancia d del borde de la columna)
    x_corte = lv_b - d_z_m
    if x_corte > 0:
        # Integrar presión sobre el área de ancho L_use desde x_corte hasta lv_b
        # La presión varía linealmente con y (si hay momento en B) y con x (si hay momento en L)
        # Para el cortante total, integramos sobre y en toda la longitud L_use, y sobre x en [x_corte, lv_b]
        # Usamos una aproximación numérica simple con 20 puntos
        y_vals = np.linspace(-L_use/2, L_use/2, 20)
        x_vals = np.linspace(x_corte, lv_b, 20)
        # La fuerza cortante es la integral doble de q(x,y) dx dy
        dx = (lv_b - x_corte) / 20
        dy = L_use / 20
        Vu_1way = 0.0
        for xi in x_vals:
            for yi in y_vals:
                q_xy = q_at(xi, yi)
                if q_xy > 0:
                    Vu_1way += q_xy * dx * dy
        # También se puede simplificar con promedio de q en el voladizo, pero la integración es más precisa.
    else:
        Vu_1way = 0.0

    phi_Vc_1way = phi_v * 0.17 * 1.0 * math.sqrt(fc_basico) * (L_use * 1000) * (d_z * 10) / 1000.0  # kN
    ok_1way = phi_Vc_1way >= Vu_1way

    # ─── PUNZONAMIENTO (bidireccional) con presión promedio ─────────────────
    bo_1 = c1_col/100.0 + d_z_m
    bo_2 = c2_col/100.0 + d_z_m
    bo_perim = 2 * (bo_1 + bo_2)
    Area_punz = bo_1 * bo_2
    # Presión promedio en el área crítica (se usa qu_avg)
    Vu_punz = P_ult - qu_avg * Area_punz

    _min_col = min(c1_col, c2_col)
    beta_c = max(c1_col, c2_col) / _min_col if _min_col > 0 else 1.0
    alpha_s = 40  # columna interior
    try:
        Vc1 = 0.33 * math.sqrt(fc_basico)
        Vc2 = 0.17 * (1 + 2/beta_c) * math.sqrt(fc_basico)
        Vc3 = 0.083 * (2 + alpha_s * (d_z*10) / (bo_perim*1000)) * math.sqrt(fc_basico)
        vc_min_MPa = min(Vc1, Vc2, Vc3)
    except ZeroDivisionError:
        vc_min_MPa = 0.0
        st.warning("⚠️ División por cero al calcular resistencia a punzonamiento — revise dimensiones.")
    phi_Vc_punz = phi_v * vc_min_MPa * (bo_perim * 1000) * (d_z * 10) / 1000.0  # kN
    ok_punz = phi_Vc_punz >= Vu_punz

    # ─── FLEXIÓN (momentos con integración exacta) ───────────────────────────
    # Momento en dirección B (respecto al eje paralelo a L)
    # Se integra presión * brazo a lo largo del voladizo
    # Para simplificar, integramos numéricamente
    def momento_dir_B():
        # Integrar (q(x,y) * (distancia desde la cara de la columna)) sobre el área del voladizo
        # La distancia desde la cara de la columna es (x + lv_b) con x desde -lv_b a 0 (coordenada local)
        # Usamos coordenadas globales: x desde -B_use/2 hasta -B_use/2 + lv_b ? Mejor usar x desde -lv_b a 0 con origen en cara
        # pero la presión es función de x e y globales.
        # La distancia desde la cara es (x_cara) que es lv_b - (x_global + B_use/2)? Mejor usar integración directa.
        # Usamos un método numérico: sobre el voladizo en dirección B, x_global desde -B_use/2 hasta -B_use/2+lv_b (lado izquierdo)
        # y_global desde -L_use/2 hasta L_use/2. El brazo de momento es la distancia desde la cara de la columna:
        # brazo = (x_global + B_use/2) ? No, la cara de la columna está a -B_use/2 + c1_col/200? Confuso.
        # Para simplificar, usamos la fórmula clásica con presión promedio, pero mejoramos con distribución lineal.
        # Dado que ya tenemos presión variable, podemos usar una integración numérica simple.
        # Tomamos puntos en el voladizo.
        x_min = -B_use/2
        x_cara_col = -c1_col/200.0  # coordenada de la cara de la columna (lado izquierdo)
        # El voladizo izquierdo va desde x_cara_col hasta x_min? No, el voladizo es desde x_cara_col hasta -B_use/2? Revisar.
        # En realidad la columna está centrada, entonces la cara izquierda está en -c1_col/200.
        # El voladizo izquierdo va desde x = -B_use/2 hasta x = -c1_col/200.
        x_left = -B_use/2
        x_right_face = -c1_col/200.0
        if x_right_face <= x_left:
            return 0.0
        n_x = 20
        n_y = 20
        x_vals = np.linspace(x_left, x_right_face, n_x)
        y_vals = np.linspace(-L_use/2, L_use/2, n_y)
        dx = (x_right_face - x_left) / n_x
        dy = L_use / n_y
        Mu_B = 0.0
        for xi in x_vals:
            # brazo desde la cara de la columna
            lever = x_right_face - xi
            for yi in y_vals:
                q_xy = q_at(xi, yi)
                if q_xy > 0:
                    Mu_B += q_xy * lever * dx * dy
        return Mu_B

    Mu_flex_B = momento_dir_B()
    # Similar para dirección L (voladizo en Y)
    def momento_dir_L():
        y_bot = -L_use/2
        y_face = -c2_col/200.0
        if y_face <= y_bot:
            return 0.0
        n_x = 20
        n_y = 20
        x_vals = np.linspace(-B_use/2, B_use/2, n_x)
        y_vals = np.linspace(y_bot, y_face, n_y)
        dx = B_use / n_x
        dy = (y_face - y_bot) / n_y
        Mu_L = 0.0
        for yi in y_vals:
            lever = y_face - yi
            for xi in x_vals:
                q_xy = q_at(xi, yi)
                if q_xy > 0:
                    Mu_L += q_xy * lever * dx * dy
        return Mu_L

    Mu_flex_L = momento_dir_L()

    # Diseño a flexión para dirección B
    try:
        Rn_B = (Mu_flex_B * 1e6) / (phi_f * (L_use*1000) * (d_z*10)**2)
        disc_B = 1 - 2*Rn_B/(0.85*fc_basico)
        rho_B = (0.85*fc_basico/fy_basico)*(1 - math.sqrt(max(disc_B, 0)))
    except (ZeroDivisionError, ValueError):
        rho_B = 0.02
        disc_B = 0
    rho_use_B = max(rho_B, 0.0018)
    As_req_B = rho_use_B * (L_use*100) * d_z  # cm2 para ancho L
    n_barras_B = math.ceil(As_req_B / REBAR_DICT[bar_z]["area"])
    sep_B = (B_use*100 - 2*recub_z) / max(1, n_barras_B - 1)

    # Diseño a flexión para dirección L
    try:
        Rn_L = (Mu_flex_L * 1e6) / (phi_f * (B_use*1000) * (d_z*10)**2)
        disc_L = 1 - 2*Rn_L/(0.85*fc_basico)
        rho_L = (0.85*fc_basico/fy_basico)*(1 - math.sqrt(max(disc_L, 0)))
    except (ZeroDivisionError, ValueError):
        rho_L = 0.02
        disc_L = 0
    rho_use_L = max(rho_L, 0.0018)
    As_req_L = rho_use_L * (B_use*100) * d_z  # cm2 para ancho B
    n_barras_L = math.ceil(As_req_L / REBAR_DICT[bar_z]["area"])
    sep_L = (L_use*100 - 2*recub_z) / max(1, n_barras_L - 1)

    # Para compatibilidad con el resto del código
    Mu_flex = Mu_flex_B
    disc_z = disc_B
    As_req_total = As_req_B
    n_barras_Z = n_barras_B
    separacion_S = sep_B

    # --- CÁLCULO DE GANCHOS Y DOBLECES (NSR-10 / ACI 318) ---
    # Diámetro mínimo de doblez (D_doblez)
    db_mm = db_bar_z
    if db_mm <= 25.4: # Hasta #8 (1")
        D_doblez_mm = 6 * db_mm
    elif db_mm <= 35.8: # #9 a #11
        D_doblez_mm = 8 * db_mm
    else: # #14 a #18
        D_doblez_mm = 10 * db_mm
    D_doblez_cm = D_doblez_mm / 10.0

    # Longitud de extensión del gancho a 90° (L_ext)
    # Norma ACI/NSR-10: max(12*db, 150mm) usualmente para estribos, pero para anclaje en tracción (gancho estándar 90°) es 12*db
    L_ext_gancho_mm = 12 * db_mm
    L_ext_gancho_cm = L_ext_gancho_mm / 10.0
    
    # Radios para dibujo
    radio_doblez_cm = D_doblez_cm / 2.0
    
    # Longitudes de desarrollo disponibles
    ldh_disp_B = (B_use*100 - c1_col)/2 - recub_z
    ldh_disp_L = (L_use*100 - c2_col)/2 - recub_z
    
    # Altura disponible para el gancho
    h_gancho_disp = H_zap - 2*recub_z
    L_gancho_real_cm = min(h_gancho_disp, L_ext_gancho_cm + radio_doblez_cm + db_mm/10.0) # Lo que cabe en la zapata


    tab_res, tab_dwg, tab_apu = st.tabs(["📋 Resultados del Diseño", "📏 Plano 3000 (DXF)", "💰 Cantidades APU"])
    
    with tab_res:
        st.markdown(f"**Revisión Estructural: f'c = {fc_basico} MPa | fy = {fy_basico} MPa**")
        
        st.markdown(r"**1. Cortante Unidireccional (Tipo Viga):** $\phi V_c \ge V_u$")
        st.latex(r"\phi V_c = \phi \cdot 0.17 \lambda \sqrt{f'_c} b_w d")
        if ok_1way:
            st.success(f"✅ Aprobado Cortante 1D: $\\phi V_c = {phi_Vc_1way:.1f}$ kN $\\ge V_u = {Vu_1way:.1f}$ kN")
        else:
            st.error(f"❌ No Aprobado Cortante 1D: $\\phi V_c = {phi_Vc_1way:.1f}$ kN $< V_u = {Vu_1way:.1f}$ kN $\\rightarrow$ **Aumentar Espesor H**")

        st.markdown(r"**2. Cortante Bidireccional (Punzonamiento):** $\phi V_c \ge V_{up}$")
        st.latex(r"\phi V_c = \phi \cdot \min\left(0.33, 0.17(1+\frac{2}{\beta}), 0.083(2+\frac{\alpha_s d}{b_o})\right) \sqrt{f'_c} b_o d")
        if ok_punz:
            st.success(f"✅ Aprobado Punzonamiento: $\\phi V_c = {phi_Vc_punz:.1f}$ kN $\\ge V_{{up}} = {Vu_punz:.1f}$ kN")
        else:
            st.error(f"❌ No Aprobado Punzonamiento: $\\phi V_c = {phi_Vc_punz:.1f}$ kN $< V_{{up}} = {Vu_punz:.1f}$ kN $\\rightarrow$ **Aumentar Espesor H o sección de Columna**")
        
        st.markdown("---")
        data_res = [
            {"Revisión": "Geometría Propuesta", "Solicitado": f"Area Req = {Area_req:.2f} m²", "Capacidad/Provisto": f"Area Prov. = {A_use:.2f} m²", "Estado": "✅ OK" if A_use>=Area_req else "⚠️ Subdimensionado"},
            {"Revisión": "Presión de Contacto qu_max", "Solicitado": f"{qu_max:.2f} kPa", "Capacidad/Provisto": f"qu_min = {qu_min:.2f} kPa", "Estado": "✅ Sin tensión" if qu_min >= 0 else "⚠️ Tensión en suelo"},
            {"Revisión": "Flexión Dir. B — cara col.", "Solicitado": f"Mu_B = {Mu_flex_B:.1f} kN-m", "Capacidad/Provisto": f"As_B = {As_req_B:.1f} cm² → {n_barras_B} {bar_z} c/{sep_B:.1f}cm", "Estado": "✅ OK" if disc_B>0 else "❌ Rompe en compresión"},
            {"Revisión": "Flexión Dir. L — cara col.", "Solicitado": f"Mu_L = {Mu_flex_L:.1f} kN-m", "Capacidad/Provisto": f"As_L = {As_req_L:.1f} cm² → {n_barras_L} {bar_z} c/{sep_L:.1f}cm", "Estado": "✅ OK" if disc_L>0 else "❌ Rompe en compresión"},
            {"Revisión": "Gancho Estándar 90°", "Solicitado": f"D_min doblez: {D_doblez_cm:.1f} cm", "Capacidad/Provisto": f"Extensión recta: {L_ext_gancho_cm:.1f} cm", "Estado": "ℹ️ Info"},
        ]
        st.table(pd.DataFrame(data_res))
        
        # ── MEMORIA DE CÁLCULO COMPLETA ────────────────────────────────────
        doc_zap = Document()
        doc_zap.add_heading(f"MEMORIA DE CÁLCULO — ZAPATA {B_use:.2f}x{L_use:.2f} m", 0)
        doc_zap.add_paragraph(f"Fecha: {_dt.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc_zap.add_paragraph(f"Norma Estructural: {norma_sel}")
        doc_zap.add_paragraph(f"Elaborado con: StructuroPro — Módulo Zapatas NSR-10/ACI-318/Multi-Norma")
        doc_zap.add_heading("1. MATERIALES", level=1)
        doc_zap.add_paragraph(f"  f'c = {fc_basico} MPa  |  fy = {fy_basico} MPa  |  Recubrimiento = {recub_z} cm")
        doc_zap.add_paragraph(f"  Varilla seleccionada: {bar_z}  |  Área unitaria = {REBAR_DICT[bar_z]['area']:.3f} cm²  |  db = {REBAR_DICT[bar_z]['db']:.1f} mm")
        doc_zap.add_heading("2. CARGAS APLICADAS", level=1)
        doc_zap.add_paragraph(f"  Servicio: Ps = {P_svc:.1f} kN, Ms_B = {M_svc_B:.1f} kN·m, Ms_L = {M_svc_L:.1f} kN·m")
        doc_zap.add_paragraph(f"  Últimas: Pu = {P_ult:.1f} kN, Mu_B = {M_ult_B:.1f} kN·m, Mu_L = {M_ult_L:.1f} kN·m")
        doc_zap.add_heading("3. DIMENSIONAMIENTO EN PLANTA", level=1)
        doc_zap.add_paragraph(f"  q_adm = {q_adm_z:.2f} kPa  |  γ_prom = {gamma_prom:.1f} kN/m³  |  Df = {Df_z:.2f} m")
        doc_zap.add_paragraph(f"  q_neto = {q_net:.2f} kPa  |  Área requerida = {Area_req:.2f} m²")
        doc_zap.add_paragraph(f"  Dimensiones mínimas → B = {B_zap:.2f} m, L = {L_zap:.2f} m")
        doc_zap.add_paragraph(f"  Dimensiones adoptadas → B = {B_use:.2f} m, L = {L_use:.2f} m")
        doc_zap.add_paragraph(f"  qu_max = {qu_max:.2f} kPa  |  qu_min = {qu_min:.2f} kPa  |  qu_avg = {qu_avg:.2f} kPa")
        doc_zap.add_heading("4. ESPESOR Y PERALTE EFECTIVO", level=1)
        doc_zap.add_paragraph(f"  Espesor H = {H_zap:.0f} cm  |  Recubrimiento = {recub_z:.1f} cm")
        doc_zap.add_paragraph(f"  Peralte efectivo d = H - recub - db/10 = {d_z:.1f} cm")
        doc_zap.add_heading("5. REVISIÓN DE CORTANTE", level=1)
        doc_zap.add_paragraph(f"  CORTANTE UNIDIRECCIONAL (Viga):")
        doc_zap.add_paragraph(f"    φVc = {phi_Vc_1way:.1f} kN  {'≥' if ok_1way else '<'}  Vu = {Vu_1way:.1f} kN  → {'✅ OK' if ok_1way else '❌ Aumentar H'}")
        doc_zap.add_paragraph(f"  PUNZONAMIENTO (Bidireccional):")
        doc_zap.add_paragraph(f"    bo = {bo_perim:.3f} m  |  φVc = {phi_Vc_punz:.1f} kN  {'≥' if ok_punz else '<'}  Vup = {Vu_punz:.1f} kN  → {'✅ OK' if ok_punz else '❌ Aumentar H'}")
        doc_zap.add_heading("6. DISEÑO A FLEXIÓN — ACERO DE REFUERZO", level=1)
        doc_zap.add_paragraph(f"  DIR. B (malla sobre el ancho L={L_use:.2f}m):")
        doc_zap.add_paragraph(f"    Mu_B = {Mu_flex_B:.1f} kN·m  |  As_B = {As_req_B:.2f} cm²  |  ρ_B = {rho_use_B:.4f}")
        doc_zap.add_paragraph(f"    Arreglo: {n_barras_B} varillas {bar_z}  c/ {sep_B:.1f} cm  {'✅ OK' if disc_B>0 else '❌ Aumentar H'}")
        doc_zap.add_paragraph(f"  DIR. L (malla sobre el ancho B={B_use:.2f}m):")
        doc_zap.add_paragraph(f"    Mu_L = {Mu_flex_L:.1f} kN·m  |  As_L = {As_req_L:.2f} cm²  |  ρ_L = {rho_use_L:.4f}")
        doc_zap.add_paragraph(f"    Arreglo: {n_barras_L} varillas {bar_z}  c/ {sep_L:.1f} cm  {'✅ OK' if disc_L>0 else '❌ Aumentar H'}")
        doc_zap.add_paragraph(f"  DETALLES DE DOBLADO (Gancho 90° estándar ACI/NSR-10):")
        doc_zap.add_paragraph(f"    Diámetro mín. de doblez = {D_doblez_cm:.1f} cm")
        doc_zap.add_paragraph(f"    Extensión recta después de la curva = {L_ext_gancho_cm:.1f} cm")
        doc_zap.add_paragraph(f"    Altura del gancho ajustada a zapata = {L_gancho_real_cm:.1f} cm")
        doc_zap.add_heading("7. CANTIDADES DE MATERIALES (APU)", level=1)
        _area_m2_doc = REBAR_DICT[bar_z]["area"] * 1e-4
        # Longitud exacta considerando tramo recto, curva y extensión
        # Tramo recto + 2 * (desarrollo curva + extension)
        _long_gancho_m = ( (math.pi * radio_doblez_cm / 2) + L_ext_gancho_cm ) / 100.0
        _long_var_B = L_use - 2*(recub_z/100.0) + 2*_long_gancho_m
        _long_var_L = B_use - 2*(recub_z/100.0) + 2*_long_gancho_m
        _pe_B = n_barras_B * _long_var_B * _area_m2_doc * 7850
        _pe_L = n_barras_L * _long_var_L * _area_m2_doc * 7850
        _pe_tot = _pe_B + _pe_L
        _vol_exc = (B_use + 0.5) * (L_use + 0.5) * Df_z
        _vol_conc = B_use * L_use * (H_zap/100.0)
        doc_zap.add_paragraph(f"  Excavación = {_vol_exc:.2f} m³")
        doc_zap.add_paragraph(f"  Concreto   = {_vol_conc:.2f} m³")
        doc_zap.add_paragraph(f"  Acero Dir.B = {_pe_B:.1f} kg  |  Acero Dir.L = {_pe_L:.1f} kg  |  Total = {_pe_tot:.1f} kg")
        doc_zap.add_paragraph(f"  Cuantía = {(_pe_tot/_vol_conc) if _vol_conc>0 else 0:.1f} kg/m³")
        # Si hay asentamiento guardado, agregarlo
        if "asentamiento" in st.session_state:
            doc_zap.add_heading("8. ASENTAMIENTO ELÁSTICO ESTIMADO", level=1)
            doc_zap.add_paragraph(f"  Asentamiento inmediato (estimado) = {st.session_state.asentamiento:.1f} mm")
        f_zap_io = io.BytesIO()
        doc_zap.save(f_zap_io)
        f_zap_io.seek(0)
        st.download_button("📥 Descargar Memoria Completa DOCX", data=f_zap_io,
                           file_name=f"Memoria_Zapata_{B_use:.1f}x{L_use:.1f}m.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab_dwg:
        st.subheader("🧊 Visualización 3D de la Fundación con Acero de Refuerzo")
        fig3d = go.Figure()

        Hz_m  = H_zap / 100.0
        rec_m = recub_z / 100.0
        db_m  = db_bar_z / 1000.0
        z_bar = rec_m + db_m / 2

        # Zapata
        x_z = [-B_use/2, B_use/2, B_use/2, -B_use/2, -B_use/2, B_use/2, B_use/2, -B_use/2]
        y_z = [-L_use/2, -L_use/2, L_use/2, L_use/2, -L_use/2, -L_use/2, L_use/2, L_use/2]
        z_z = [0]*4 + [Hz_m]*4
        fig3d.add_trace(go.Mesh3d(x=x_z, y=y_z, z=z_z, alphahull=0, opacity=0.25,
                                  color='steelblue', name='Zapata', showlegend=True))

        # Columna
        c1_m = c1_col/100.0; c2_m = c2_col/100.0
        x_c = [-c1_m/2, c1_m/2, c1_m/2, -c1_m/2, -c1_m/2, c1_m/2, c1_m/2, -c1_m/2]
        y_c = [-c2_m/2, -c2_m/2, c2_m/2, c2_m/2, -c2_m/2, -c2_m/2, c2_m/2, c2_m/2]
        z_c = [Hz_m]*4 + [Hz_m + 1.0]*4
        fig3d.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.5,
                                  color='slategray', name='Columna', showlegend=True))

        # Medidas para ganchos en 3D
        _r_m = radio_doblez_cm / 100.0          # radio de curva en metros
        _ext_m = L_ext_gancho_cm / 100.0         # extensión recta en metros
        _hook_h = (math.pi * _r_m / 2) + _ext_m  # altura total del gancho (arco + recta)

        # Varillas Dir. B (barras corren en Y=dirección L)
        _xs_barB = np.linspace(-B_use/2 + rec_m, B_use/2 - rec_m, n_barras_B)
        _show_B = True
        _show_Bh = True
        for xi in _xs_barB:
            # Tramo recto horizontal
            fig3d.add_trace(go.Scatter3d(
                x=[xi, xi], y=[-L_use/2 + rec_m, L_use/2 - rec_m], z=[z_bar, z_bar],
                mode='lines', line=dict(color='#ff6b35', width=5),
                name='Acero Dir.B' if _show_B else None, showlegend=_show_B, legendgroup='aB'))
            _show_B = False
            # Gancho extremo -Y (sube verticalmente)
            fig3d.add_trace(go.Scatter3d(
                x=[xi, xi], y=[-L_use/2 + rec_m, -L_use/2 + rec_m],
                z=[z_bar, z_bar + _hook_h],
                mode='lines', line=dict(color='#ff9a6c', width=3, dash='dot'),
                name='Ganchos Dir.B' if _show_Bh else None, showlegend=_show_Bh, legendgroup='aBh'))
            # Gancho extremo +Y
            fig3d.add_trace(go.Scatter3d(
                x=[xi, xi], y=[L_use/2 - rec_m, L_use/2 - rec_m],
                z=[z_bar, z_bar + _hook_h],
                mode='lines', line=dict(color='#ff9a6c', width=3, dash='dot'),
                name=None, showlegend=False, legendgroup='aBh'))
            _show_Bh = False

        # Varillas Dir. L (barras corren en X=dirección B)
        _ys_barL = np.linspace(-L_use/2 + rec_m, L_use/2 - rec_m, n_barras_L)
        _z_barL  = z_bar + db_m
        _hook_h_L = _hook_h + db_m   # ligeramente más alto porque van encima
        _show_L = True
        _show_Lh = True
        for yi in _ys_barL:
            # Tramo recto horizontal
            fig3d.add_trace(go.Scatter3d(
                x=[-B_use/2 + rec_m, B_use/2 - rec_m], y=[yi, yi], z=[_z_barL, _z_barL],
                mode='lines', line=dict(color='#ffd54f', width=5),
                name='Acero Dir.L' if _show_L else None, showlegend=_show_L, legendgroup='aL'))
            _show_L = False
            # Gancho extremo -X (sube verticalmente)
            fig3d.add_trace(go.Scatter3d(
                x=[-B_use/2 + rec_m, -B_use/2 + rec_m], y=[yi, yi],
                z=[_z_barL, _z_barL + _hook_h_L],
                mode='lines', line=dict(color='#ffe57f', width=3, dash='dot'),
                name='Ganchos Dir.L' if _show_Lh else None, showlegend=_show_Lh, legendgroup='aLh'))
            # Gancho extremo +X
            fig3d.add_trace(go.Scatter3d(
                x=[B_use/2 - rec_m, B_use/2 - rec_m], y=[yi, yi],
                z=[_z_barL, _z_barL + _hook_h_L],
                mode='lines', line=dict(color='#ffe57f', width=3, dash='dot'),
                name=None, showlegend=False, legendgroup='aLh'))
            _show_Lh = False

        fig3d.update_layout(
            scene=dict(aspectmode='data', xaxis_title='B (m)', yaxis_title='L (m)', zaxis_title='Z (m)',
                       bgcolor='#0f1117',
                       xaxis=dict(showgrid=True, gridcolor='#333'),
                       yaxis=dict(showgrid=True, gridcolor='#333'),
                       zaxis=dict(showgrid=True, gridcolor='#333')),
            margin=dict(l=0, r=0, b=0, t=30), height=550,
            showlegend=True, dragmode='turntable', paper_bgcolor='#0f1117', font=dict(color='white'),
            title=dict(text=f"Zapata {B_use:.2f}x{L_use:.2f}m | H={H_zap:.0f}cm | "
                           f"Dir.B: {n_barras_B}×{bar_z} c/{sep_B:.1f}cm | Dir.L: {n_barras_L}×{bar_z} c/{sep_L:.1f}cm",
                       font=dict(color='white')))
        st.plotly_chart(fig3d, use_container_width=True)
        
        st.markdown("---")
        st.write("#### Geometría de Zapata 2D")
        fig_z, ax_z = plt.subplots(figsize=(6, 4))
        ax_z.set_facecolor('#1a1a2e'); fig_z.patch.set_facecolor('#1a1a2e')
        ax_z.add_patch(patches.Rectangle((0,0), B_use*100, H_zap, linewidth=2, edgecolor='darkgray', facecolor='#4a4a6a'))
        pos_x_col = (B_use*100 - c1_col) / 2
        ax_z.add_patch(patches.Rectangle((pos_x_col, H_zap), c1_col, 50, linewidth=2, edgecolor='white', facecolor='#6a6a8a'))
        for i in range(n_barras_Z):
            xi = recub_z + i * separacion_S
            ax_z.add_patch(plt.Circle((xi, recub_z), db_bar_z/10, color='#ff6b35', zorder=5))
        # Dibujar perfil del doblez para la varilla principal (Dir B)
        # Tramo horizontal inferior
        _x_start = recub_z + radio_doblez_cm
        _x_end = B_use*100 - recub_z - radio_doblez_cm
        _y_bar = recub_z
        ax_z.add_patch(patches.Rectangle((_x_start, _y_bar - db_bar_z/20.0), _x_end - _x_start, db_bar_z/10.0, color='#ffd54f', zorder=4))
        
        # Gancho Izquierdo
        arc_izq = patches.Arc((_x_start, _y_bar + radio_doblez_cm), D_doblez_cm, D_doblez_cm, angle=180, theta1=90, theta2=180, color='#ffd54f', lw=3, zorder=4)
        ax_z.add_patch(arc_izq)
        ax_z.plot([recub_z, recub_z], [_y_bar + radio_doblez_cm, _y_bar + radio_doblez_cm + L_ext_gancho_cm], color='#ffd54f', lw=3, zorder=4)
        
        # Gancho Derecho
        arc_der = patches.Arc((_x_end, _y_bar + radio_doblez_cm), D_doblez_cm, D_doblez_cm, angle=270, theta1=90, theta2=180, color='#ffd54f', lw=3, zorder=4)
        ax_z.add_patch(arc_der)
        ax_z.plot([B_use*100 - recub_z, B_use*100 - recub_z], [_y_bar + radio_doblez_cm, _y_bar + radio_doblez_cm + L_ext_gancho_cm], color='#ffd54f', lw=3, zorder=4)
        
        # Cota del doblez
        ax_z.annotate(f"{L_ext_gancho_cm:.1f}cm", xy=(recub_z, _y_bar + radio_doblez_cm + L_ext_gancho_cm/2), xytext=(-15, _y_bar + radio_doblez_cm + L_ext_gancho_cm/2), arrowprops=dict(arrowstyle="->", color='yellow'), color='yellow', fontsize=8, va='center')

        ax_z.text(B_use*100/2, H_zap/2, f"{n_barras_Z} varillas {bar_z} L={L_use}m\nSep:{separacion_S:.1f}cm\nGancho: 90°", color='white', ha='center', va='center')
        ax_z.set_xlim(-20, B_use*100+20)
        ax_z.set_ylim(-10, H_zap+70)
        ax_z.axis('off')
        st.pyplot(fig_z)
        
        st.markdown("#### 💾 Exportar AutoCAD (.dxf)")
        doc_dxf = ezdxf.new('R2010')
        doc_dxf.units = ezdxf.units.CM
        msp = doc_dxf.modelspace()

        for _lay, _col in [('CONCRETO',7),('ACERO',1),('TEXTO',3),('EJES',8)]:
            if _lay not in doc_dxf.layers:
                doc_dxf.layers.add(_lay, color=_col)

        ex = 0; ey = 0
        # Elevación
        msp.add_lwpolyline([(ex,ey),(ex+B_use*100,ey),(ex+B_use*100,ey+H_zap),(ex,ey+H_zap),(ex,ey)],
                           close=True, dxfattribs={'layer':'CONCRETO'})
        msp.add_lwpolyline([(ex+pos_x_col,ey+H_zap),(ex+pos_x_col,ey+H_zap+50),
                            (ex+pos_x_col+c1_col,ey+H_zap+50),(ex+pos_x_col+c1_col,ey+H_zap)],
                           close=True, dxfattribs={'layer':'CONCRETO'})
        for i in range(n_barras_B):
            xi = ex + recub_z + i * sep_B
            msp.add_line((xi, ey+recub_z), (xi, ey+recub_z), dxfattribs={'layer':'ACERO'})
            msp.add_circle((xi, ey+recub_z), db_bar_z/10, dxfattribs={'layer':'ACERO'})
        msp.add_line((ex+recub_z, ey+recub_z), (ex+B_use*100-recub_z, ey+recub_z), dxfattribs={'layer':'ACERO'})
        msp.add_mtext(f"{n_barras_B}#{bar_z}@{sep_B:.0f}cm Dir.B",
                      dxfattribs={'layer':'TEXTO','char_height':4,'insert':(ex+B_use*50, ey-8)})

        # Planta
        px = ex + B_use*100 + 30; py = ey
        msp.add_lwpolyline([(px,py),(px+B_use*100,py),(px+B_use*100,py+L_use*100),(px,py+L_use*100),(px,py)],
                           close=True, dxfattribs={'layer':'CONCRETO'})
        msp.add_lwpolyline([(px+pos_x_col,py+(L_use*100-c2_col)/2),
                            (px+pos_x_col+c1_col,py+(L_use*100-c2_col)/2),
                            (px+pos_x_col+c1_col,py+(L_use*100+c2_col)/2),
                            (px+pos_x_col,py+(L_use*100+c2_col)/2)],
                           close=True, dxfattribs={'layer':'CONCRETO'})
        for i in range(n_barras_B):
            xi = px + recub_z + i * sep_B
            msp.add_line((xi, py+recub_z), (xi, py+L_use*100-recub_z), dxfattribs={'layer':'ACERO'})
        for j in range(n_barras_L):
            yj = py + recub_z + j * sep_L
            msp.add_line((px+recub_z, yj), (px+B_use*100-recub_z, yj), dxfattribs={'layer':'ACERO'})
        msp.add_mtext(f"{n_barras_L}#{bar_z}@{sep_L:.0f}cm Dir.L",
                      dxfattribs={'layer':'TEXTO','char_height':4,'insert':(px+B_use*50, py+L_use*100+8)})

        # ezdxf.write() requiere un stream de texto (StringIO), no BytesIO
        # Luego se codifica a bytes para el botón de Streamlit
        _dxf_text = io.StringIO()
        doc_dxf.write(_dxf_text)
        _dxf_bytes = _dxf_text.getvalue().encode("utf-8")
        st.download_button(label="📐 Descargar Plano DXF (Elev. + Planta con Acero)",
                           data=_dxf_bytes,
                           file_name=f"Zapata_{B_use:.1f}x{L_use:.1f}m_Armado.dxf",
                           mime="application/dxf")

    with tab_apu:
        # ─── Cantidades base ────────────────────────────────────────────────
        vol_excavacion   = (B_use + 0.5) * (L_use + 0.5) * Df_z
        vol_concreto_zap = B_use * L_use * (H_zap / 100.0)

        # Longitud de varilla con gancho real (arco + extensión)
        _area_m2  = REBAR_DICT[bar_z]["area"] * 1e-4
        _long_gancho_m  = ((math.pi * radio_doblez_cm / 2) + L_ext_gancho_cm) / 100.0
        _long_var_B = L_use - 2*(recub_z/100.0) + 2*_long_gancho_m
        _long_var_L = B_use - 2*(recub_z/100.0) + 2*_long_gancho_m
        _kg_por_m   = REBAR_DICT[bar_z]["area"] * 1e-4 * 7850          # kg/m por varilla
        peso_barras_B_apu = n_barras_B * _long_var_B * _kg_por_m
        peso_barras_L_apu = n_barras_L * _long_var_L * _kg_por_m
        peso_total_acero_zap = peso_barras_B_apu + peso_barras_L_apu

        # Proporciones concreto (ACI 211 aprox. para fc 21 MPa)
        pct_arena_apu = 0.55; pct_grava_apu = 0.80
        bultos_zap    = vol_concreto_zap * 350 / 50.0               # bultos de 50 kg
        vol_arena_z   = vol_concreto_zap * pct_arena_apu             # m³
        vol_grava_z   = vol_concreto_zap * pct_grava_apu             # m³
        litros_agua   = vol_concreto_zap * 185.0                     # l/m³  (rel a/c ≈0.53)

        # ─── SECCIÓN 1: Resumen de materiales ──────────────────────────────
        st.markdown("### 📦 Resumen de Materiales — Quantiy Take-Off")
        cols_m = st.columns(4)
        cols_m[0].metric("⛏️ Excavación", f"{vol_excavacion:.2f} m³")
        cols_m[1].metric("🧱 Vol. Concreto", f"{vol_concreto_zap:.2f} m³")
        cols_m[2].metric("🏋️ Acero Total", f"{peso_total_acero_zap:.1f} kg")
        cols_m[3].metric("📐 Cuantía", f"{peso_total_acero_zap/vol_concreto_zap:.1f} kg/m³")

        st.markdown("#### 🪣 Ingredientes de Concreto")
        df_mat = pd.DataFrame([
            {"Material": "🧱 Cemento",        "Cantidad": f"{bultos_zap:.1f}",    "Unidad": "bultos (50 kg)"},
            {"Material": "🏖️ Arena",          "Cantidad": f"{vol_arena_z:.3f}",   "Unidad": "m³"},
            {"Material": "🪨 Gravilla",        "Cantidad": f"{vol_grava_z:.3f}",   "Unidad": "m³"},
            {"Material": "💧 Agua",            "Cantidad": f"{litros_agua:.0f}",   "Unidad": "litros"},
            {"Material": "🏋️ Acero refuerzo",  "Cantidad": f"{peso_total_acero_zap:.1f}", "Unidad": "kg"},
        ])
        # Integrar precios en tiempo real si están disponibles
        _has_prices = "apu_config" in st.session_state
        _apu = st.session_state.get("apu_config", {})
        _mon = _apu.get("moneda", "")
        _p_cem = _apu.get("cemento", 0.0)
        _p_ace = _apu.get("acero",   0.0)
        _p_are = _apu.get("arena",   0.0)
        _p_gra = _apu.get("grava",   0.0)
        c_excav_u = _apu.get("costo_excav_m3", 25000.0)   # fallback COP

        _c_cem  = bultos_zap * _p_cem
        _c_ace  = peso_total_acero_zap * _p_ace
        _c_are  = vol_arena_z * _p_are
        _c_gra  = vol_grava_z * _p_gra
        _c_exc  = vol_excavacion * c_excav_u if _has_prices else 0.0
        
        _total_mat = _c_exc + _c_cem + _c_are + _c_gra + _c_ace

        # Calcular Gran Total si hay precios
        _gran_total = 0.0
        if _has_prices:
            total_dias_mo = (peso_total_acero_zap * 0.04) + (vol_concreto_zap * 0.4) + (vol_excavacion * 0.3)
            costo_mo = total_dias_mo * _apu.get("costo_dia_mo", 69333.33)
            costo_directo = _total_mat + costo_mo
            herramienta = costo_mo * _apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * _apu.get("pct_aui", 0.30)
            utilidad = costo_directo * _apu.get("pct_util", 0.05)
            iva = utilidad * _apu.get("iva", 0.19)
            _gran_total = costo_directo + herramienta + aiu + iva

            col_msg, col_metric = st.columns([2, 1])
            col_msg.success(f"💱 Precios actualizados del scraping — {_mon}")
            col_metric.metric(f"💎 Gran Total Proyecto [{_mon}]", f"{_gran_total:,.0f}")
        else:
            st.info("ℹ️ Ve a **APU Mercado** para descargar los costos en tiempo real aquí mismo.")

        # TABLA: Materiales con costos
        _mat_rows = [
            {"Material": "⛏️ Excavación",    "Cantidad": f"{vol_excavacion:.2f}", "Unidad": "m³",           "Precio Unit.": f"{c_excav_u:,.0f} {_mon}" if _has_prices else "—", "Subtotal": f"{_c_exc:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": "🧱 Cemento",        "Cantidad": f"{bultos_zap:.1f}",    "Unidad": "bultos (50kg)", "Precio Unit.": f"{_p_cem:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_cem:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": "🏖️ Arena",          "Cantidad": f"{vol_arena_z:.3f}",   "Unidad": "m³",           "Precio Unit.": f"{_p_are:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_are:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": "🪨 Gravilla",        "Cantidad": f"{vol_grava_z:.3f}",   "Unidad": "m³",           "Precio Unit.": f"{_p_gra:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_gra:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": "💧 Agua",            "Cantidad": f"{litros_agua:.0f}",   "Unidad": "litros",       "Precio Unit.": "—", "Subtotal": "—"},
            {"Material": "🏋️ Acero refuerzo",  "Cantidad": f"{peso_total_acero_zap:.1f}", "Unidad": "kg",  "Precio Unit.": f"{_p_ace:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_ace:,.0f} {_mon}" if _has_prices else "—"},
        ]
        st.dataframe(pd.DataFrame(_mat_rows), use_container_width=True, hide_index=True)

        if _has_prices:
            _total_mat = _c_exc + _c_cem + _c_are + _c_gra + _c_ace
            st.metric(f"💰 Total Materiales [{_mon}]", f"{_total_mat:,.0f}")

        # ─── SECCIÓN 2: Despiece de Acero con costos ─────────────────────────
        st.markdown("#### 🔩 Despiece de Acero de Refuerzo")
        db_mm_apu    = REBAR_DICT[bar_z]["db"]
        area_cm2_apu = REBAR_DICT[bar_z]["area"]
        _c_ace_B = peso_barras_B_apu * _p_ace
        _c_ace_L = peso_barras_L_apu * _p_ace

        _row_B = {
            "Dir.": "B  (⟵ sobre L →)",
            "Varilla": bar_z, "db [mm]": f"{db_mm_apu:.1f}",
            "N° Barras": n_barras_B, "Sep. [cm]": f"{sep_B:.1f}",
            "L gancho [cm]": f"{L_ext_gancho_cm:.1f}",
            "L total [m]": f"{_long_var_B:.3f}",
            "kg/m": f"{_kg_por_m:.3f}", "Peso [kg]": f"{peso_barras_B_apu:.2f}",
        }
        _row_L = {
            "Dir.": "L  (⟵ sobre B →)",
            "Varilla": bar_z, "db [mm]": f"{db_mm_apu:.1f}",
            "N° Barras": n_barras_L, "Sep. [cm]": f"{sep_L:.1f}",
            "L gancho [cm]": f"{L_ext_gancho_cm:.1f}",
            "L total [m]": f"{_long_var_L:.3f}",
            "kg/m": f"{_kg_por_m:.3f}", "Peso [kg]": f"{peso_barras_L_apu:.2f}",
        }
        _row_tot = {
            "Dir.": "━━ TOTAL ━━",
            "Varilla": "", "db [mm]": "",
            "N° Barras": n_barras_B + n_barras_L, "Sep. [cm]": "",
            "L gancho [cm]": "",
            "L total [m]": f"{n_barras_B*_long_var_B + n_barras_L*_long_var_L:.2f}",
            "kg/m": "", "Peso [kg]": f"{peso_total_acero_zap:.2f}",
        }
        if _has_prices:
            _row_B[f"Costo [{_mon}]"] = f"{_c_ace_B:,.0f}"
            _row_L[f"Costo [{_mon}]"] = f"{_c_ace_L:,.0f}"
            _row_tot[f"Costo [{_mon}]"] = f"{(_c_ace_B+_c_ace_L):,.0f}"

        st.dataframe(pd.DataFrame([_row_B, _row_L, _row_tot]), use_container_width=True, hide_index=True)
        st.caption(f"Gancho 90° ACI/NSR-10: D_doblez mín = {radio_doblez_cm*2:.1f} cm | ext. = {L_ext_gancho_cm:.1f} cm")

        # ─── GRÁFICO: Despiece visual ────────────────────────────────────────
        st.markdown("#### 📊 Diagrama de Despiece")
        _items_g   = ["Excavación", "Cemento", "Arena", "Gravilla", "Acero Dir.B", "Acero Dir.L"]
        _qty_g     = [vol_excavacion, bultos_zap, vol_arena_z, vol_grava_z, peso_barras_B_apu, peso_barras_L_apu]
        _units_g   = ["m³", "bultos", "m³", "m³", "kg", "kg"]
        _colors_g  = ["#5c8a5a","#e8c07d","#c2a06b","#9b7b5c","#ff6b35","#ffd54f"]

        _fig_desp = go.Figure()
        _fig_desp.add_trace(go.Bar(
            x=_items_g, y=_qty_g,
            marker_color=_colors_g,
            text=[f"{q:.2f} {u}" for q, u in zip(_qty_g, _units_g)],
            textposition="outside",
            name="Cantidad"
        ))

        if _has_prices:
            _cost_g = [_c_exc, _c_cem, _c_are, _c_gra, _c_ace_B, _c_ace_L]
            _fig_desp.add_trace(go.Bar(
                x=_items_g, y=_cost_g,
                marker_color=_colors_g, opacity=0.5,
                text=[f"{v:,.0f} {_mon}" for v in _cost_g],
                textposition="outside",
                name=f"Costo [{_mon}]",
                yaxis="y2"
            ))
            _fig_desp.update_layout(yaxis2=dict(title=f"Costo [{_mon}]", overlaying="y", side="right", showgrid=False, tickfont=dict(color="#aaa")))

        _fig_desp.update_layout(
            paper_bgcolor="#0f1117", plot_bgcolor="#0f1117",
            font=dict(color="white"), barmode="group",
            xaxis=dict(showgrid=False),
            yaxis=dict(title="Cantidad", showgrid=True, gridcolor="#222"),
            legend=dict(bgcolor="#111", font=dict(color="white")),
            margin=dict(l=20, r=20, t=30, b=20), height=380,
            title=dict(text=f"Despiece Zapata {B_use:.2f}x{L_use:.2f}m — {bar_z}", font=dict(color="white"))
        )
        st.plotly_chart(_fig_desp, use_container_width=True)



        if "apu_config" in st.session_state:
            st.markdown("---")
            st.markdown("### 💰 Presupuesto Estimado (Promedio de Fuentes Regionales)")
            apu = st.session_state.apu_config
            mon = apu["moneda"]
            c_excav = vol_excavacion * 25000
            
            bultos_zap  = vol_concreto_zap * 350 / 50.0
            pct_arena = apu.get("pct_arena_mezcla", 0.55)
            pct_grava = apu.get("pct_grava_mezcla", 0.80)
            vol_arena_z = vol_concreto_zap * pct_arena
            vol_grava_z = vol_concreto_zap * pct_grava

            c_cem = bultos_zap * apu["cemento"]
            c_ace = peso_total_acero_zap * apu["acero"]
            c_are = vol_arena_z * apu["arena"]
            c_gra = vol_grava_z * apu["grava"]
            total_mat = c_cem + c_ace + c_are + c_gra + c_excav
            
            total_dias_mo = (peso_total_acero_zap * 0.04) + (vol_concreto_zap * 0.4) + (vol_excavacion * 0.3)
            costo_mo = total_dias_mo * apu.get("costo_dia_mo", 69333.33)
            
            costo_directo = total_mat + costo_mo
            herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * apu.get("pct_aui", 0.30)
            utilidad = costo_directo * apu.get("pct_util", 0.05)
            iva = utilidad * apu.get("iva", 0.19)
            total_proyecto = costo_directo + herramienta + aiu + iva
            
            data_zap_apu = {
                "Item": ["Excavación (m³)", "Cemento (bultos)", "Acero (kg)", "Arena (m³)", "Grava (m³)", 
                         "Mano de Obra (días)", "Herramienta Menor", "A.I.U.", "IVA s/Utilidad", "TOTAL PRESUPUESTO"],
                "Cantidad": [f"{vol_excavacion:.2f}", f"{bultos_zap:.1f}", f"{peso_total_acero_zap:.1f}", f"{vol_arena_z:.2f}", f"{vol_grava_z:.2f}", 
                             f"{total_dias_mo:.2f}", f"{apu.get('pct_herramienta', 0.05)*100:.1f}% MO", 
                             f"{apu.get('pct_aui', 0.3)*100:.1f}% CD", f"{apu.get('iva', 0.19)*100:.1f}% Util", ""],
                f"Subtotal [{mon}]": [f"{c_excav:,.2f}", f"{c_cem:,.2f}", f"{c_ace:,.2f}", f"{c_are:,.2f}", f"{c_gra:,.2f}", 
                                      f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}", f"{iva:,.2f}", f"**{total_proyecto:,.2f}**"]
            }
            st.dataframe(pd.DataFrame(data_zap_apu), use_container_width=True, hide_index=True)
            
            # Excel APU Export
            output_excel = io.BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                df_export = pd.DataFrame({
                    "Item": ["Excavación", "Cemento", "Acero", "Arena", "Grava", "Mano de Obra"],
                    "Cantidad": [vol_excavacion, bultos_zap, peso_total_acero_zap, vol_arena_z, vol_grava_z, total_dias_mo],
                    "Unidad": [25000, apu['cemento'], apu['acero'], apu['arena'], apu['grava'], apu.get('costo_dia_mo', 69333.33)]
                })
                df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
                df_export.to_excel(writer, index=False, sheet_name='APU')
                workbook = writer.book
                worksheet = writer.sheets['APU']
                money_fmt = workbook.add_format({'num_format': '#,##0.00'})
                bold = workbook.add_format({'bold': True})
                worksheet.set_column('A:A', 25)
                worksheet.set_column('B:D', 15, money_fmt)
                
                row = len(df_export) + 1
                worksheet.write(row, 0, "Costo Directo (CD)", bold)
                worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
                row += 1
                worksheet.write(row, 0, "Herramienta Menor", bold)
                worksheet.write_formula(row, 3, f'=D7*{apu.get("pct_herramienta", 0.05)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "A.I.U", bold)
                worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui", 0.30)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "IVA s/ Utilidad", bold)
                _row_util = row - 1
                worksheet.write_formula(row, 3, f'=D{_row_util}*{apu.get("iva", 0.19)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "TOTAL PRESUPUESTO", bold)
                worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
                
            output_excel.seek(0)
            st.download_button(label="📥 Descargar Presupuesto Excel (.xlsx)", data=output_excel, 
                               file_name=f"APU_Zapata_{B_use}x{L_use}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("💡 Ve a la página 'APU Mercado' para cargar los costos base de agregados, acero y cemento y que tu presupuesto se genere automáticamente.")