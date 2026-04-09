import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go
import json
import datetime
import ifc_export
import os
from pathlib import Path
import qrcode
from PIL import Image
import tempfile

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state["idioma"] if "idioma" in st.session_state else "Español"
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Diagramas de Interacción Biaxial", "Biaxial Interaction Diagrams"), layout="wide")

# ─────────────────────────────────────────────
# PERSISTENCIA SUPABASE
# ─────────────────────────────────────────────
import requests

def get_supabase_rest_info():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return url, key
    except Exception as e:
        return None, None

def guardar_proyecto_supabase(nombre, estado_dict):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_columnas.json"
            db = {}
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            db[f"[COLUMNAS] {nombre}"] = {"nombre_proyecto": f"[COLUMNAS] {nombre}", "estado_json": json.dumps(estado_dict)}
            with open(db_path, "w", encoding="utf-8") as f: json.dump(db, f)
            return True, " Proyecto guardado (Local)"
        except Exception as e:
            return False, f" Error guardado local: {e}"
    
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "Prefer": "resolution=merge-duplicates"
    }
    payload = {
        "nombre_proyecto": nombre,
        "user_id": st.session_state.get("user_id", "anonimo"),
        "estado_json": json.dumps(estado_dict),
    }
    
    try:
        endpoint = f"{url}/rest/v1/proyectos?on_conflict=nombre_proyecto"
        res = requests.post(endpoint, headers=headers, json=payload)
        if res.status_code in [200, 201, 204]:
            return True, " Proyecto guardado en la nube"
        else:
            return False, f" Error API: {res.text}"
    except Exception as e:
        return False, f" Error al guardar: {e}"

def cargar_proyecto_supabase(nombre):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_columnas.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                match = db.get(f"[COLUMNAS] {nombre}") or db.get(nombre)
                if match:
                    estado = json.loads(match["estado_json"])
                    for k, v in estado.items(): st.session_state[k] = v
                    return True, f" Proyecto '{nombre}' cargado (Local)"
            return False, f" No se encontró el proyecto '{nombre}' localmente"
        except Exception as e:
            return False, f" Excepción al cargar local: {e}"
        
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Accept": "application/json"
    }
    
    try:
        endpoint = f"{url}/rest/v1/proyectos?nombre_proyecto=eq.{nombre}&select=*"
        res = requests.get(endpoint, headers=headers)
        
        if res.status_code == 200:
            data = res.json()
            if data and len(data) > 0:
                estado = json.loads(data[0]["estado_json"])
                for k, v in estado.items():
                    st.session_state[k] = v
                return True, f" Proyecto '{nombre}' cargado"
            else:
                return False, f" No se encontró el proyecto '{nombre}'"
        else:
            return False, f" Error al cargar (API): {res.text}"
    except Exception as e:
        return False, f" Excepción al cargar: {e}"

def listar_proyectos_supabase():
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_columnas.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                return sorted([k.replace("[COLUMNAS] ", "") for k in db.keys()])
            return []
        except Exception:
            return []
    
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Accept": "application/json"
    }
    
    try:
        endpoint = f"{url}/rest/v1/proyectos?select=nombre_proyecto"
        res = requests.get(endpoint, headers=headers)
        if res.status_code == 200:
            data = res.json()
            nombres = [item["nombre_proyecto"] for item in data if "nombre_proyecto" in item]
            return sorted(nombres)
        return []
    except Exception:
        return []

def capturar_estado_actual():
    claves = [
        "circ_pm_norma", "circ_pm_nivel_sismico", "circ_pm_fc_unit", "circ_pm_fc_mpa",
        "circ_pm_fy", "circ_pm_seccion_type", "circ_pm_b", "circ_pm_h", "circ_pm_D",
        "circ_pm_d_prime", "circ_pm_L", "circ_pm_unit_system", "circ_pm_rebar_type",
        "circ_pm_num_h", "circ_pm_num_v", "circ_pm_n_barras_circ",
        "circ_pm_col_type", "circ_pm_stirrup_type", "circ_pm_paso",
        "circ_pm_mux", "circ_pm_muy", "circ_pm_pu",
        "circ_pm_output_units", "circ_pm_k",
        # Factor k de esbeltez
        "circ_pm_k1",
        # APU concreto premezclado
        "apu_premix", "apu_premix_precio",
        "apu_moneda", "apu_cemento", "apu_acero", "apu_arena", "apu_grava", "apu_mo", "apu_aui",
        # Rótulo ICONTEC para DXF
        "dxf_empresa", "dxf_proyecto", "dxf_plano", "dxf_elaboro", "dxf_reviso", "dxf_aprobo",
        # QR del plano
        "qr_url",
    ]
    return {k: st.session_state[k] for k in claves if k in st.session_state}

# ─────────────────────────────────────────────
# FIX BUG #4: Manejo seguro de BASE_DIR para Streamlit Cloud
# ─────────────────────────────────────────────
try:
    BASE_DIR = Path(__file__).parent.parent
except NameError:
    BASE_DIR = Path(os.getcwd())

header_img_path = BASE_DIR / "assets" / "columnas_circulares_header.png"
if header_img_path.exists():
    st.image(str(header_img_path), use_container_width=False, width=700)
else:
    st.image("https://via.placeholder.com/700x100?text=Columnas+PM+Biaxial", width=700)

st.title(_t("Diagrama de Interacción P–M (Biaxial) y Diseño de Estribos", " P-M (Biaxial) Interaction Diagram & Tie Design"))
st.markdown(_t(
    "Generador interactivo de capacidad a flexocompresión **biaxial** para **Columnas Cuadradas, Rectangulares y Circulares**.",
    "Interactive **biaxial** flexure-compression capacity generator for **Square, Rectangular and Circular Columns**."
))

with st.expander(" ¿Cómo usar esta herramienta?"):
    st.markdown("""
    **Modo de Uso:**
    1. ** Sidebar:** Selecciona Norma, nivel sísmico, f'c, fy, geometría (cuadrada, rectangular o circular)
    2. ** Armadura:** Define varillas longitudinales y estribos (o espiral)
    3. ** Solicitaciones:** Ingresa Momentos (Mux, Muy) y Carga Axial (Pu)
    4. ** Resultados:** Diagrama P-M biaxial 3D + verificación Bresler
    5. ** Exportar:** Memoria DOCX completa, DXF con rótulo ICONTEC, Excel, etc.
    """)

# ─────────────────────────────────────────────
# DICCIONARIOS DE BARRAS
# ─────────────────────────────────────────────
REBAR_US = {
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
    "#5 (5/8\")": {"area": 1.99, "diam_mm": 15.88},
    "#6 (3/4\")": {"area": 2.84, "diam_mm": 19.05},
    "#7 (7/8\")": {"area": 3.87, "diam_mm": 22.23},
    "#8 (1\")":   {"area": 5.10, "diam_mm": 25.40},
    "#9 (1 1/8\")": {"area": 6.45, "diam_mm": 28.65},
    "#10 (1 1/4\")": {"area": 7.92, "diam_mm": 32.26},
}

REBAR_MM = {
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
    "14 mm": {"area": 1.539, "diam_mm": 14.0},
    "16 mm": {"area": 2.011, "diam_mm": 16.0},
    "18 mm": {"area": 2.545, "diam_mm": 18.0},
    "20 mm": {"area": 3.142, "diam_mm": 20.0},
    "22 mm": {"area": 3.801, "diam_mm": 22.0},
    "25 mm": {"area": 4.909, "diam_mm": 25.0},
    "28 mm": {"area": 6.158, "diam_mm": 28.0},
    "32 mm": {"area": 8.042, "diam_mm": 32.0},
}

STIRRUP_US = {
    "#2 (1/4\")": {"area": 0.32, "diam_mm": 6.35},
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
}

STIRRUP_MM = {
    "6 mm":  {"area": 0.283, "diam_mm": 6.0},
    "8 mm":  {"area": 0.503, "diam_mm": 8.0},
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
}

# ─────────────────────────────────────────────
# FUNCIONES DE DIBUJO PARA FIGURADO
# ─────────────────────────────────────────────
_MM_TO_BAR = {
    6.0:  "#2 (1/4\")",  6.35: "#2 (1/4\")",
    8.0:  "#2.5 (5/16\")",
    9.53: "#3 (3/8\")",  10.0: "#3 (3/8\")",
    12.0: "#4 (1/2\")",  12.70: "#4 (1/2\")",
    14.0: "#4.5 (9/16\")",
    15.88: "#5 (5/8\")", 16.0: "#5 (5/8\")",
    18.0: "#5.7 (11/16\")",
    19.05: "#6 (3/4\")", 20.0: "#6 (3/4\")",
    22.0: "#7 (7/8\")",  22.23: "#7 (7/8\")",
    25.0: "#8 (1\")",    25.40: "#8 (1\")",
    28.0: "#9 (1 1/8\")",28.65: "#9 (1 1/8\")",
    32.0: "#10 (1 1/4\")",32.26: "#10 (1 1/4\")",
}

def _bar_label(diam_mm):
    best = min(_MM_TO_BAR.keys(), key=lambda k: abs(k - diam_mm))
    if abs(best - diam_mm) <= 2.0:
        return f"{_MM_TO_BAR[best]}  Ø{diam_mm:.1f} mm"
    return f"Ø{diam_mm:.1f} mm"

def draw_longitudinal_bar(total_len_cm, straight_len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(6, total_len_cm/20), 2))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    ax.plot([0, straight_len_cm], [0, 0], 'k-', linewidth=2)
    ax.plot([0, 0], [0, hook_len_cm], 'k-', linewidth=2)
    ax.plot([straight_len_cm, straight_len_cm], [0, -hook_len_cm], 'k-', linewidth=2)
    ax.annotate(f"{straight_len_cm:.0f} cm", xy=(straight_len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 12db = {hook_len_cm:.0f} cm", xy=(0, hook_len_cm/2), ha='right', fontsize=8)
    ax.annotate(f"Gancho 12db", xy=(straight_len_cm, -hook_len_cm/2), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*0.2, straight_len_cm + hook_len_cm*0.2)
    ax.set_ylim(-hook_len_cm*1.2, hook_len_cm*1.2)
    ax.axis('off')
    ax.set_title(f"Varilla L1 — {label} — Longitud total {total_len_cm:.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_stirrup(b_cm, h_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    import math as _math
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(5, b_cm/12), max(4, h_cm/12)))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    x0, y0 = 0, 0
    ax.plot([x0, x0+b_cm], [y0, y0], 'k-', linewidth=2.5)
    ax.plot([x0+b_cm, x0+b_cm], [y0, y0+h_cm], 'k-', linewidth=2.5)
    ax.plot([x0+b_cm, x0], [y0+h_cm, y0+h_cm], 'k-', linewidth=2.5)
    ax.plot([x0, x0], [y0+h_cm, y0], 'k-', linewidth=2.5)
    angle_rad = _math.radians(45)
    vis_hook = min(hook_len_cm, b_cm/4.0, h_cm/4.0)
    hx = vis_hook * _math.cos(angle_rad)
    hy = -vis_hook * _math.sin(angle_rad)
    ax.plot([x0, x0 + hx], [y0+h_cm, y0+h_cm + hy], 'k-', linewidth=2.5)
    ax.plot([x0+b_cm, x0+b_cm - hx], [y0+h_cm, y0+h_cm + hy], 'k--', linewidth=1.5, alpha=0.5)
    ax.annotate(f"{b_cm:.0f} cm", xy=(b_cm/2, y0-0.8), ha='center', fontsize=9, fontweight='bold')
    ax.annotate(f"{h_cm:.0f} cm", xy=(x0-0.8, h_cm/2), ha='right', va='center', fontsize=9, fontweight='bold')
    ax.annotate(f"Gancho 135°\n6d_e = {6*bar_diam_mm/10:.1f} cm",
                xy=(x0 + hx + 0.2, y0+h_cm + hy - 0.2), fontsize=7.5, color='darkred',
                va='top', ha='left')
    ax.set_xlim(x0 - hook_len_cm*0.3, b_cm + hook_len_cm*0.6)
    ax.set_ylim(y0 - hook_len_cm*0.5, h_cm + hook_len_cm*0.9)
    ax.axis('off')
    ax.set_title(f"Estribo E1 — {label} — Perímetro {2*(b_cm+h_cm):.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_crosstie(len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(6, len_cm/15), 2))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    ax.plot([0, len_cm], [0, 0], 'k-', linewidth=2)
    ax.plot([0, -hook_len_cm*0.7], [0, -hook_len_cm*0.7], 'k-', linewidth=2)
    ax.plot([len_cm, len_cm + hook_len_cm*0.7], [0, -hook_len_cm*0.7], 'k-', linewidth=2)
    ax.annotate(f"{len_cm:.0f} cm", xy=(len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 135°", xy=(0, -hook_len_cm*0.5), ha='right', fontsize=8)
    ax.annotate(f"Gancho 135°", xy=(len_cm, -hook_len_cm*0.5), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*1.2, len_cm + hook_len_cm*1.2)
    ax.set_ylim(-hook_len_cm*1.5, hook_len_cm*0.5)
    ax.axis('off')
    ax.set_title(f"Crosstie C1 — {label} — Longitud {len_cm:.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_spiral(D_cm, paso_cm, bar_diam_mm, bar_name=None):
    """Dibujo esquemático de espiral para columnas circulares"""
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(6, 4))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    circle = plt.Circle((0, 0), D_cm/2, fill=False, edgecolor='black', linewidth=2)
    ax.add_patch(circle)
    theta = np.linspace(0, 4*np.pi, 200)
    r = D_cm/2 - bar_diam_mm/10
    x = r * np.cos(theta)
    y = r * np.sin(theta)
    ax.plot(x, y, 'k-', linewidth=1.5)
    ax.annotate(f"Espiral {label}", xy=(0, D_cm/2 + 2), ha='center', fontsize=9)
    ax.annotate(f"Paso = {paso_cm:.1f} cm", xy=(0, -D_cm/2 - 3), ha='center', fontsize=8)
    ax.set_xlim(-D_cm/2 - 5, D_cm/2 + 5)
    ax.set_ylim(-D_cm/2 - 8, D_cm/2 + 5)
    ax.axis('off')
    ax.set_title(f"Espiral — {label}", fontsize=9, fontweight='bold')
    return fig

# ─────────────────────────────────────────────
# PARÁMETROS POR NORMA (con límites por nivel sísmico)
# ─────────────────────────────────────────────
CODES = {
    "NSR-10 (Colombia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 4.0, "rho_max_des": 4.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["DMI — Disipación Mínima", "DMO — Disipación Moderada", "DES — Disipación Especial"],
        "ref": "NSR-10 Título C",
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary", "IMF — Intermediate", "SMF — Special"],
        "ref": "ACI 318-25",
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary", "IMF — Intermediate", "SMF — Special"],
        "ref": "ACI 318-19",
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary", "IMF — Intermediate", "SMF — Special"],
        "ref": "ACI 318-14",
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["GS — Grado Reducido", "GM — Grado Moderado", "GA — Grado Alto"],
        "ref": "NEC-SE-HM",
    },
    "E.060 (Perú)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 5.0, "rho_max_des": 5.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["PO — Pórtico Ordinario", "PM — Pórtico Moderado", "PE — Pórtico Especial"],
        "ref": "E.060 Perú",
    },
    "NTC-EM (México)": {
        "phi_tied": 0.70, "phi_spiral": 0.80, "phi_tension": 0.85,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 5.0, "rho_max_des": 5.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["MDL — Ductilidad Limitada", "MROD — Ductilidad Ordinaria", "MRLE — Ductilidad Alta"],
        "ref": "NTC-EM México",
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 5.0, "rho_max_des": 5.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["PO — Pórtico Ordinario", "PM — Pórtico Moderado", "PE — Pórtico Especial"],
        "ref": "COVENIN 1753",
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["DO — Diseño Ordinario", "DM — Ductilidad Moderada", "DE — Diseño Especial"],
        "ref": "NB 1225001",
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["GE — Grado Estándar", "GM — Ductilidad Moderada", "GA — Ductilidad Alta"],
        "ref": "CIRSOC 201",
    },
}

# ─────────────────────────────────────────────
# PRESENTACIONES DE CEMENTO POR PAÍS
# ─────────────────────────────────────────────
CEMENT_BAGS = {
    "NSR-10 (Colombia)": [{"label": "Cemento gris (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "ACI 318-25 (EE.UU.)": [{"label": "Type I/II (94 lb / 42.6 kg)", "kg": 42.6}, {"label": "Type III (47 lb / 21.3 kg)", "kg": 21.3}],
    "ACI 318-19 (EE.UU.)": [{"label": "Type I/II (94 lb / 42.6 kg)", "kg": 42.6}, {"label": "Type III (47 lb / 21.3 kg)", "kg": 21.3}],
    "ACI 318-14 (EE.UU.)": [{"label": "Type I/II (94 lb / 42.6 kg)", "kg": 42.6}, {"label": "Type III (47 lb / 21.3 kg)", "kg": 21.3}],
    "NEC-SE-HM (Ecuador)": [{"label": "Cemento Holcim (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "E.060 (Perú)": [{"label": "Cemento Andino (42.5 kg)", "kg": 42.5}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "NTC-EM (México)": [{"label": "Cemento Cemex (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "COVENIN 1753-2006 (Venezuela)": [{"label": "Cemento (42.5 kg)", "kg": 42.5}],
    "NB 1225001-2020 (Bolivia)": [{"label": "Cemento (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "CIRSOC 201-2025 (Argentina)": [{"label": "Cemento (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
}

# ─────────────────────────────────────────────
# TABLA DE DOSIFICACIÓN ACI 211
# ─────────────────────────────────────────────
MIX_DESIGNS = [
    {"fc_mpa": 14.0, "cem": 250, "agua": 205, "arena": 810, "grava": 1060, "wc": 0.82},
    {"fc_mpa": 17.0, "cem": 290, "agua": 200, "arena": 780, "grava": 1060, "wc": 0.69},
    {"fc_mpa": 21.0, "cem": 350, "agua": 193, "arena": 720, "grava": 1060, "wc": 0.55},
    {"fc_mpa": 25.0, "cem": 395, "agua": 193, "arena": 680, "grava": 1020, "wc": 0.49},
    {"fc_mpa": 28.0, "cem": 430, "agua": 190, "arena": 640, "grava": 1000, "wc": 0.44},
    {"fc_mpa": 35.0, "cem": 530, "agua": 185, "arena": 580, "grava":  960, "wc": 0.35},
    {"fc_mpa": 42.0, "cem": 620, "agua": 180, "arena": 520, "grava":  910, "wc": 0.29},
    {"fc_mpa": 56.0, "cem": 740, "agua": 175, "arena": 450, "grava":  850, "wc": 0.24},
]

def get_mix_for_fc(fc_mpa):
    if fc_mpa <= MIX_DESIGNS[0]["fc_mpa"]: return MIX_DESIGNS[0]
    if fc_mpa >= MIX_DESIGNS[-1]["fc_mpa"]: return MIX_DESIGNS[-1]
    for i in range(len(MIX_DESIGNS)-1):
        lo, hi = MIX_DESIGNS[i], MIX_DESIGNS[i+1]
        if lo["fc_mpa"] <= fc_mpa <= hi["fc_mpa"]:
            t = (fc_mpa - lo["fc_mpa"]) / (hi["fc_mpa"] - lo["fc_mpa"])
            return {k: lo[k] + t*(hi[k]-lo[k]) for k in ("cem", "agua", "arena", "grava", "wc")}
    return MIX_DESIGNS[-1]

def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_development_length(db_mm, fy, fc, lambda_=1.0, psi_t=1.0, psi_e=1.0, psi_s=1.0, psi_g=1.0, cb_ktr=2.5):
    if db_mm <= 0: return 0
    ld = (3/40) * (fy / (lambda_ * math.sqrt(fc))) * (psi_t * psi_e * psi_s * psi_g / cb_ktr) * db_mm
    return max(ld, 300)

# ─────────────────────────────────────────────
# FUNCIONES PARA CÁLCULO DE CAPACIDAD UNIAXIAL
# ─────────────────────────────────────────────
def compute_uniaxial_capacity(b, h, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza):
    eps_cu = 0.003
    eps_y = fy / Es
    beta_1 = get_beta1(fc)
    
    Ag = b * h
    Ast = sum([layer['As'] for layer in layers])
    Po_kN = (0.85 * fc * (Ag * 100 - Ast * 100) + fy * Ast * 100) / 1000.0
    
    c_vals = np.concatenate([np.linspace(1e-5, h, 120), np.linspace(h, h * 12, 60)])
    P_n_list = []; M_n_list = []; phi_P_n_list = []; phi_M_n_list = []
    eps_t_vals = []
    
    b_mm = b * 10
    h_mm = h * 10
    
    for c_cm in c_vals:
        c_mm = c_cm * 10
        a_mm = min(beta_1 * c_mm, h_mm)
        Cc = 0.85 * fc * a_mm * b_mm
        Mc = Cc * (h_mm / 2.0 - a_mm / 2.0)
        
        Ps = 0.0; Ms = 0.0; eps_t = 0.0
        for layer in layers:
            d_i_mm = layer['d'] * 10
            As_i = layer['As'] * 100
            eps_s = eps_cu * (c_mm - d_i_mm) / c_mm
            if d_i_mm >= max(l['d'] * 10 for l in layers):
                eps_t = eps_s
            fs = max(-fy, min(fy, Es * eps_s))
            if a_mm > d_i_mm and fs > 0:
                fs -= 0.85 * fc
            Ps += As_i * fs
            Ms += As_i * fs * (h_mm / 2.0 - d_i_mm)
        
        Pn = (Cc + Ps) / 1000.0
        Mn = abs((Mc + Ms) / 1_000_000.0)
        eps_t_tens = -eps_t
        
        if eps_t_tens <= eps_y:
            phi = phi_c_max
        elif eps_t_tens >= eps_full:
            phi = phi_tension
        else:
            phi = phi_c_max + (phi_tension - phi_c_max) * (eps_t_tens - eps_y) / (eps_full - eps_y)
        
        Pn_max_val = p_max_factor * Po_kN
        phi_Pn_max_val = phi_c_max * Pn_max_val
        
        Pn = min(Pn, Pn_max_val)
        phi_Pn = min(phi * Pn, phi_Pn_max_val)
        phi_Mn = phi * Mn
        
        P_n_list.append(Pn)
        M_n_list.append(Mn)
        phi_P_n_list.append(phi_Pn)
        phi_M_n_list.append(phi_Mn)
        eps_t_vals.append(eps_t_tens)
    
    c_balance = None
    P_balance = None
    M_balance = None
    for i, eps in enumerate(eps_t_vals):
        if abs(eps - eps_y) < 0.0001 or (i > 0 and eps_t_vals[i-1] <= eps_y <= eps):
            idx = i if abs(eps - eps_y) < 0.0001 else i-1
            c_balance = c_vals[idx]
            P_balance = P_n_list[idx] * factor_fuerza
            M_balance = M_n_list[idx] * factor_fuerza
            break
    
    return {
        'M_n': np.array(M_n_list) * factor_fuerza,
        'P_n': np.array(P_n_list) * factor_fuerza,
        'phi_M_n': np.array(phi_M_n_list) * factor_fuerza,
        'phi_P_n': np.array(phi_P_n_list) * factor_fuerza,
        'Po': Po_kN * factor_fuerza,
        'Pn_max': p_max_factor * Po_kN * factor_fuerza,
        'phi_Pn_max': phi_c_max * p_max_factor * Po_kN * factor_fuerza,
        'c_balance': c_balance,
        'P_balance': P_balance,
        'M_balance': M_balance,
    }

def compute_uniaxial_capacity_circular(D, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza):
    """Para sección circular"""
    eps_cu = 0.003
    eps_y = fy / Es
    beta_1 = get_beta1(fc)
    
    r = D / 2
    Ag = math.pi * r**2
    Ast = sum([layer['As'] for layer in layers])
    Po_kN = (0.85 * fc * (Ag * 100 - Ast * 100) + fy * Ast * 100) / 1000.0
    
    c_vals = np.concatenate([
        np.linspace(1e-5, D * 0.1, 40),
        np.linspace(D * 0.1, D * 0.5, 80),
        np.linspace(D * 0.5, D, 60),
        np.linspace(D, D * 12, 40)
    ])
    P_n_list = []; M_n_list = []; phi_P_n_list = []; phi_M_n_list = []
    eps_t_vals = []
    
    for c_cm in c_vals:
        c_mm = c_cm * 10
        if c_mm >= D * 10:
            a_mm = D * 10
        else:
            a_mm = beta_1 * c_mm
        a_cm = a_mm / 10
        if a_cm >= D:
            Ac_comp = Ag
        else:
            h_seg = a_cm
            _arg_ac = np.clip((r - h_seg) / r, -1.0, 1.0)
            Ac_comp = r**2 * math.acos(_arg_ac) - (r - h_seg) * math.sqrt(max(0.0, 2*r*h_seg - h_seg**2))
        
        Cc = 0.85 * fc * Ac_comp * 100
        if a_cm >= D:
            y_cent = r
        else:
            _arg_yc = np.clip((r - a_cm) / r, -1.0, 1.0)
            _angle = math.acos(_arg_yc)
            y_cent = (4*r * math.sin(_angle / 2)**3) / (3 * (_angle - math.sin(_angle) * math.cos(_angle)))
        Mc = Cc * (r - y_cent) / 10
        
        Ps = 0.0; Ms = 0.0; eps_t = 0.0
        for layer in layers:
            d_i_mm = layer['d'] * 10
            As_i = layer['As'] * 100
            eps_s = eps_cu * (c_mm - d_i_mm) / c_mm
            if d_i_mm >= max(l['d'] * 10 for l in layers):
                eps_t = eps_s
            fs = max(-fy, min(fy, Es * eps_s))
            Ps += As_i * fs
            Ms += As_i * fs * (r - d_i_mm/10) / 10
        
        Pn = (Cc + Ps) / 1000.0
        Mn = abs((Mc + Ms) / 1000.0)
        eps_t_tens = -eps_t
        
        if eps_t_tens <= eps_y:
            phi = phi_c_max
        elif eps_t_tens >= eps_full:
            phi = phi_tension
        else:
            phi = phi_c_max + (phi_tension - phi_c_max) * (eps_t_tens - eps_y) / (eps_full - eps_y)
        
        Pn_max_val = p_max_factor * Po_kN
        phi_Pn_max_val = phi_c_max * Pn_max_val
        
        Pn = min(Pn, Pn_max_val)
        phi_Pn = min(phi * Pn, phi_Pn_max_val)
        phi_Mn = phi * Mn
        
        P_n_list.append(Pn)
        M_n_list.append(Mn)
        phi_P_n_list.append(phi_Pn)
        phi_M_n_list.append(phi_Mn)
        eps_t_vals.append(eps_t_tens)
    
    c_balance = None
    P_balance = None
    M_balance = None
    for i, eps in enumerate(eps_t_vals):
        if abs(eps - eps_y) < 0.0001 or (i > 0 and eps_t_vals[i-1] <= eps_y <= eps):
            idx = i if abs(eps - eps_y) < 0.0001 else i-1
            c_balance = c_vals[idx]
            P_balance = P_n_list[idx] * factor_fuerza
            M_balance = M_n_list[idx] * factor_fuerza
            break
    
    return {
        'M_n': np.array(M_n_list) * factor_fuerza,
        'P_n': np.array(P_n_list) * factor_fuerza,
        'phi_M_n': np.array(phi_M_n_list) * factor_fuerza,
        'phi_P_n': np.array(phi_P_n_list) * factor_fuerza,
        'Po': Po_kN * factor_fuerza,
        'Pn_max': p_max_factor * Po_kN * factor_fuerza,
        'phi_Pn_max': phi_c_max * p_max_factor * Po_kN * factor_fuerza,
        'c_balance': c_balance,
        'P_balance': P_balance,
        'M_balance': M_balance,
    }

# ─────────────────────────────────────────────
# FUNCIÓN PARA VERIFICACIÓN BIAXIAL (BRESLER)
# ─────────────────────────────────────────────
def interp_pm_curve(M_query, phi_Mn_arr, phi_Pn_arr):
    """
    Interpola φPn dado φMn en la curva P-M NO monótona.
    La curva tiene dos ramas: compresión (M creciente) y tensión (M decreciente).
    Separamos ambas ramas e interpolamos en cada una, retornando el MAYOR φPn
    (el de la rama de compresión, que es el correcto para Bresler).
    """
    phi_Mn_arr = np.array(phi_Mn_arr)
    phi_Pn_arr = np.array(phi_Pn_arr)
    
    if len(phi_Mn_arr) == 0:
        return 0.0
    
    M_max = np.max(phi_Mn_arr)
    
    # Si M=0 → capacidad axial pura (Po)
    if M_query <= 0:
        return float(phi_Pn_arr[np.argmax(phi_Pn_arr)])
    
    # Si el momento pedido supera el máximo → falla total (marcador -1 para distinguirlo de Pn=0 real)
    if M_query > M_max:
        return -1.0  # sentinel: M excede el diagrama en este eje
    
    # Índice del punto de balance (donde φMn es máximo)
    idx_bal = int(np.argmax(phi_Mn_arr))
    
    # ── Rama COMPRESIÓN: Po → Balance (M creciente, P decreciente) ──
    Mc = phi_Mn_arr[:idx_bal + 1]
    Pc = phi_Pn_arr[:idx_bal + 1]
    
    # ── Rama TENSIÓN: Balance → Tracción pura (M decreciente, P→0) ──
    Mt = phi_Mn_arr[idx_bal:]
    Pt = phi_Pn_arr[idx_bal:]
    
    # Ordenar cada rama para np.interp (requiere x creciente)
    sc = np.argsort(Mc)
    st = np.argsort(Mt)
    
    P_comp = float(np.interp(M_query, Mc[sc], Pc[sc],
                              left=float(Pc[sc[0]]), right=0.0))
    P_tens = float(np.interp(M_query, Mt[st], Pt[st],
                              left=0.0, right=0.0))
    
    # Retornar el MAYOR (rama de compresión es siempre mayor)
    return max(P_comp, P_tens)


def biaxial_bresler(Pu, Mux, Muy, cap_x, cap_y, Po, phi_factor):
    phi_Pnx, phi_Pny, phi_P0 = None, None, None
    phi_Pnx = interp_pm_curve(abs(Mux), np.array(cap_x['phi_M_n']), np.array(cap_x['phi_P_n']))
    phi_Pny = interp_pm_curve(abs(Muy), np.array(cap_y['phi_M_n']), np.array(cap_y['phi_P_n']))
    phi_P0  = phi_factor * Po

    if phi_Pnx is None or phi_Pny is None:
        return {
            'phi_Pnx': 0.0,
            'phi_Pny': 0.0,
            'phi_P0':  phi_P0,
            'phi_Pni': 0.0,
            'ratio':   float('inf'),
            'ok':      False
        }

    # Detectar sentinel -1.0: el momento supera el diagrama en ese eje
    eje_excedido = []
    if phi_Pnx < 0:
        eje_excedido.append("X")
    if phi_Pny < 0:
        eje_excedido.append("Y")

    if eje_excedido:
        msg = f"Mux/Muy excede el diagrama P-M en eje {'&'.join(eje_excedido)}: aumentar sección o acero"
        return {
            'phi_Pnx': max(phi_Pnx, 0.0),
            'phi_Pny': max(phi_Pny, 0.0),
            'phi_P0':  phi_P0,
            'phi_Pni': 0.0,
            'ratio':   float('inf'),
            'ok':      False,
            'msg_exceso': msg,
        }

    if phi_Pnx > 0 and phi_Pny > 0 and phi_P0 > 0:
        inv = 1/phi_Pnx + 1/phi_Pny - 1/phi_P0
        phi_Pni = 1/inv if inv > 0 else 0.0
    else:
        phi_Pni = 0.0

    ratio = Pu / phi_Pni if phi_Pni > 0 else float('inf')
    ok    = Pu <= phi_Pni

    return {
        'phi_Pnx': phi_Pnx,
        'phi_Pny': phi_Pny,
        'phi_P0':  phi_P0,
        'phi_Pni': phi_Pni,
        'ratio':   ratio,
        'ok':      ok,
        'msg_exceso': None,
    }

# ─────────────────────────────────────────────
# FUNCIÓN PARA ESBELTEZ (NSR-10 C.10.10 / ACI 6.6)
# ─────────────────────────────────────────────
def check_slenderness(L, b, h, k, Pu, M1, M2, fc, fy, Es, factor_fuerza, es_circular=False):
    r = 0.25 * b if es_circular else min(b, h) / math.sqrt(12)
    kl = k * L
    kl_r = kl / r if r > 0 else 999
    
    if kl_r <= 22:
        slender = False
        classification = "Columna corta (sin efectos de segundo orden)"
        delta_ns = 1.0
    elif kl_r <= 100:
        slender = True
        classification = "Columna esbelta (requiere magnificación de momentos)"
        Ec = 4700 * math.sqrt(fc)
        # Bug Fix: para sección circular usar Ig = π*D⁴/64, no b*h³/12
        if es_circular:
            Ig = math.pi * b**4 / 64   # b == D cuando se llama para circular
        else:
            Ig = b * h**3 / 12
        Ig_mm4 = Ig * 1e4
        EI = 0.4 * Ec * Ig_mm4 / (1 + 0.0)
        Pc = math.pi**2 * EI / (kl * 1000)**2
        Pc = Pc / 1000
        Cm = 0.6 + 0.4 * (M1/M2) if abs(M2) > 0 else 1.0
        Cm = max(0.4, Cm)
        delta_ns = Cm / (1 - Pu / (0.75 * Pc))
        delta_ns = max(delta_ns, 1.0)
    else:
        slender = True
        classification = "Columna muy esbelta (kl/r > 100) — requiere análisis no lineal según NSR-10 C.10.10.7"
        delta_ns = 1.0
    
    return {
        'kl_r': kl_r,
        'slender': slender,
        'classification': classification,
        'delta_ns': delta_ns,
        'r': r,
        'kl': kl
    }

# =============================================================================
# SIDEBAR - ENTRADAS DEL USUARIO
# =============================================================================
st.sidebar.header(_t("0. Norma de Diseño", "0. Design Code"))
norma_options = list(CODES.keys())
norma_sel = st.sidebar.selectbox(_t("Norma", "Code"), norma_options, key="circ_pm_norma")
code = CODES[norma_sel]

nivel_sismico = st.sidebar.selectbox(
    _t("Nivel Sísmico / Ductilidad:", "Seismic / Ductility Level:"),
    code["seismic_levels"],
    key="circ_pm_nivel_sismico"
)

nivel_lower = nivel_sismico.lower()
es_des = any(k in nivel_lower for k in ["des", "disipación especial", "smf", "special", "ga", "ductilidad alta", "pe", "pórtico especial", "de", "diseño especial", "mrle", "alta"])
es_dmo = any(k in nivel_lower for k in ["dmo", "imf", "intermediate", "gm", "moderada", "pm", "mrod", "media", "moderado"]) and not es_des
es_dmi = not (es_des or es_dmo)

if es_des:
    rho_max = code["rho_max_des"]
elif es_dmo:
    rho_max = code["rho_max_dmo"]
else:
    rho_max = code["rho_max_dmi"]
rho_min = code["rho_min"]

st.sidebar.caption(f" {_t('Referencia', 'Reference')}: {code['ref']}")
st.sidebar.caption(f" ρ máx según nivel: {rho_max}% | ρ mín: {rho_min}%")

st.sidebar.header(_t("1. Materiales", "1. Materials"))
fc_unit = st.sidebar.radio(_t("Unidad de f'c:", "f'c Unit:"), ["MPa", "PSI", "kg/cm²"], horizontal=True, key="circ_pm_fc_unit")

if fc_unit == "PSI":
    psi_options = {"2500 PSI (≈ 17.2 MPa)": 2500, "3000 PSI (≈ 20.7 MPa)": 3000,
                   "3500 PSI (≈ 24.1 MPa)": 3500, "4000 PSI (≈ 27.6 MPa)": 4000,
                   "4500 PSI (≈ 31.0 MPa)": 4500, "5000 PSI (≈ 34.5 MPa)": 5000,
                   "Personalizado": None}
    psi_choice = st.sidebar.selectbox("Resistencia f'c [PSI]", list(psi_options.keys()), key="circ_pm_psi_choice")
    fc_psi = float(psi_options[psi_choice]) if psi_options[psi_choice] is not None else st.sidebar.number_input("f'c personalizado [PSI]", 2000.0, 12000.0, 3000.0, 100.0, key="circ_pm_fc_psi_custom")
    fc = fc_psi * 0.00689476
elif fc_unit == "kg/cm²":
    kgcm2_options = {"175 kg/cm² (≈ 17.2 MPa)": 175, "210 kg/cm² (≈ 20.6 MPa)": 210,
                     "250 kg/cm² (≈ 24.5 MPa)": 250, "280 kg/cm² (≈ 27.5 MPa)": 280,
                     "350 kg/cm² (≈ 34.3 MPa)": 350, "420 kg/cm² (≈ 41.2 MPa)": 420,
                     "Personalizado": None}
    kgcm2_choice = st.sidebar.selectbox("Resistencia f'c [kg/cm²]", list(kgcm2_options.keys()), key="circ_pm_kgcm2_choice")
    fc_kgcm2 = float(kgcm2_options[kgcm2_choice]) if kgcm2_options[kgcm2_choice] is not None else st.sidebar.number_input("f'c personalizado [kg/cm²]", 100.0, 1200.0, 210.0, 10.0, key="circ_pm_fc_kgcm2_custom")
    fc = fc_kgcm2 / 10.1972
else:
    fc = st.sidebar.number_input("Resistencia del Concreto (f'c) [MPa]", 15.0, 80.0, 21.0, 1.0, key="circ_pm_fc_mpa")

fy = st.sidebar.number_input("Fluencia del Acero (fy) [MPa]", 240.0, 500.0, 420.0, 10.0, key="circ_pm_fy")
Es = 200000.0

st.sidebar.header(_t("2. Geometría de la Sección", "2. Section Geometry"))
seccion_type = st.sidebar.selectbox(_t("Tipo de sección", "Section type"), 
                                    [_t("Rectangular / Cuadrada", "Rectangular / Square"), 
                                     _t("Circular (con espiral)", "Circular (with spiral)")],
                                    key="circ_pm_seccion_type")
es_circular = "Circular" in seccion_type

if es_circular:
    D = st.sidebar.number_input(_t("Diámetro (D) [cm]", "Diameter (D) [cm]"), 15.0, 150.0, 40.0, 5.0, key="circ_pm_D")
    b = D
    h = D
    st.sidebar.caption("ℹ Para columnas circulares se usa espiral en lugar de estribos")
else:
    b = st.sidebar.number_input(_t("Base (b) [cm]", "Width (b) [cm]"), 15.0, 150.0, 30.0, 5.0, key="circ_pm_b")
    h = st.sidebar.number_input(_t("Altura (h) [cm]", "Height (h) [cm]"), 15.0, 150.0, 40.0, 5.0, key="circ_pm_h")

d_prime = st.sidebar.number_input(_t("Recubrimiento al centroide (d') [cm]", "Cover to centroid (d') [cm]"), 2.0, 15.0, 5.0, 0.5, key="circ_pm_dprime")
L_col = st.sidebar.number_input(_t("Altura libre de la columna (L) [cm]", "Column clear height (L) [cm]"), 50.0, 1000.0, 300.0, 25.0, key="circ_pm_L")

st.sidebar.header(_t("3. Refuerzo Longitudinal", "3. Longitudinal Reinforcement"))
unit_system = st.sidebar.radio(_t("Sistema de Unidades de las Varillas:", "Bar Unit System:"), 
                               [_t("Pulgadas (EE. UU.)", "Inches (US)"), 
                                _t("Milímetros (SI)", "Millimeters (SI)")], 
                               key="circ_pm_unit_system")

rebar_dict = REBAR_US if "Pulgadas" in unit_system or "Inches" in unit_system else REBAR_MM
default_rebar = "#5 (5/8\")" if "Pulgadas" in unit_system else "16 mm"
rebar_type = st.sidebar.selectbox(_t("Diámetro de las Varillas", "Bar Diameter"), list(rebar_dict.keys()), key="circ_pm_rebar_type")
rebar_area = rebar_dict[rebar_type]["area"]
rebar_diam = rebar_dict[rebar_type]["diam_mm"]

if es_circular:
    n_barras = st.sidebar.number_input(_t("Número de varillas longitudinales", "Number of longitudinal bars"), 4, 20, 8, 2, key="circ_pm_n_barras_circ")
    Ast = n_barras * rebar_area
    layers = []
    angulos = np.linspace(0, 2*np.pi, n_barras, endpoint=False)
    radio_centro = D/2 - d_prime
    for i, ang in enumerate(angulos):
        x_pos = radio_centro * math.cos(ang)
        y_pos = radio_centro * math.sin(ang)
        layers.append({'d': D/2 + y_pos, 'As': rebar_area, 'x': x_pos, 'y': y_pos})
else:
    num_filas_h = st.sidebar.number_input(_t("# de filas Acero Horiz (Superior e Inferior)", "# of horizontal rows (Top & Bottom)"), 2, 15, 2, 1, key="circ_pm_num_h")
    num_filas_v = st.sidebar.number_input(_t("# de filas Acero Vert (Laterales)", "# of vertical rows (Sides)"), 2, 15, 2, 1, key="circ_pm_num_v")
    
    layers = []
    layers.append({'d': d_prime, 'As': num_filas_h * rebar_area})
    num_capas_intermedias = num_filas_v - 2
    if num_capas_intermedias > 0:
        espacio_h = (h - 2 * d_prime) / (num_capas_intermedias + 1)
        for i in range(1, num_capas_intermedias + 1):
            layers.append({'d': d_prime + i * espacio_h, 'As': 2 * rebar_area})
    layers.append({'d': h - d_prime, 'As': num_filas_h * rebar_area})
    n_barras_total = num_filas_h * 2 + num_capas_intermedias * 2
    Ast = sum([layer['As'] for layer in layers])
    n_barras = n_barras_total

Ag = (math.pi * (D/2)**2) if es_circular else (b * h)
cuantia = Ast / Ag * 100

st.sidebar.markdown(f"**Total varillas:** {n_barras} | **Ast:** {Ast:.2f} cm² | **ρ = {cuantia:.2f}%**")

if es_circular and n_barras < 6:
    st.sidebar.error(" NSR-10 C.10.9.2: El número mínimo de barras para columnas circulares es 6.")
elif not es_circular and n_barras < 4:
    st.sidebar.error(" NSR-10 C.10.9.2: El número mínimo de barras para columnas rectangulares es 4.")

if cuantia < rho_min or cuantia > rho_max:
    st.sidebar.error(f" CUANTÍA FUERA DE LÍMITES: ρ = {cuantia:.2f}% (límites: {rho_min}% - {rho_max}%)")
elif cuantia > 4.0:
    st.sidebar.warning(f"⚠ Alerta constructiva: ρ = {cuantia:.2f}% > 4%. NSR-10 recomienda no superar 4% por congestión.")

st.sidebar.header(_t("4. Refuerzo Transversal", "4. Transverse Reinforcement"))

if es_circular:
    col_type = _t("Espiral (Spiral)", "Spiral")
    stirrup_dict = STIRRUP_MM if "Milímetros" in unit_system else STIRRUP_US
    default_stirrup = "8 mm" if "Milímetros" in unit_system else "#3 (3/8\")"
    spiral_type = st.sidebar.selectbox(_t("Diámetro de la Espiral", "Spiral Diameter"), list(stirrup_dict.keys()), key="circ_pm_spiral_type")
    stirrup_area = stirrup_dict[spiral_type]["area"]
    stirrup_diam = stirrup_dict[spiral_type]["diam_mm"]
    paso_espiral = st.sidebar.number_input(_t("Paso de la espiral (s) [cm]", "Spiral pitch (s) [cm]"), 2.0, 20.0, 7.5, 0.5, key="circ_pm_paso")
else:
    col_type_options = [_t("Estribos (Tied)", "Tied"), _t("Espiral (Spiral)", "Spiral")]
    col_type = st.sidebar.selectbox(_t("Tipo de Columna", "Column Type"), col_type_options, key="circ_pm_col_type")
    stirrup_dict = STIRRUP_US if "Pulgadas" in unit_system else STIRRUP_MM
    default_stirrup = "#3 (3/8\")" if "Pulgadas" in unit_system else "8 mm"
    stirrup_type = st.sidebar.selectbox(_t("Diámetro del Estribo", "Stirrup Diameter"), list(stirrup_dict.keys()), key="circ_pm_stirrup_type")
    stirrup_area = stirrup_dict[stirrup_type]["area"]
    stirrup_diam = stirrup_dict[stirrup_type]["diam_mm"]

if "Espiral" in col_type or es_circular:
    phi_c_max = code["phi_spiral"]
    p_max_factor = code["pmax_spiral"]
else:
    phi_c_max = code["phi_tied"]
    p_max_factor = code["pmax_tied"]
phi_tension = code["phi_tension"]
eps_full = code["eps_tension_full"]

st.sidebar.header(_t("5. Solicitaciones (Biaxiales)", "5. Loads (Biaxial)"))
unidades_salida = st.sidebar.radio(_t("Unidades del Diagrama:", "Output Units:"), 
                                   [_t("KiloNewtons (kN, kN-m)", "KiloNewtons (kN, kN-m)"), 
                                    _t("Toneladas Fuerza (tonf, tonf-m)", "Tons Force (tonf, tonf-m)")], 
                                   key="circ_pm_output_units")

if unidades_salida == _t("Toneladas Fuerza (tonf, tonf-m)", "Tons Force (tonf, tonf-m)"):
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom = "tonf-m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom = "kN-m"

st.sidebar.markdown(f"Cargas últimas en **{unidad_fuerza}** y **{unidad_mom}**:")

if "circ_pm_mux" not in st.session_state:
    st.session_state["circ_pm_mux"] = round(45.0 * factor_fuerza, 2)
if "circ_pm_muy" not in st.session_state:
    st.session_state["circ_pm_muy"] = round(25.0 * factor_fuerza, 2)
if "circ_pm_pu" not in st.session_state:
    st.session_state["circ_pm_pu"] = round(2700.0 * factor_fuerza, 2)

Mux_input = st.sidebar.number_input(f"Momento Último Mux [{unidad_mom}]", value=st.session_state["circ_pm_mux"], step=round(10.0 * factor_fuerza, 2), key="circ_pm_mux")
Muy_input = st.sidebar.number_input(f"Momento Último Muy [{unidad_mom}]", value=st.session_state["circ_pm_muy"], step=round(10.0 * factor_fuerza, 2), key="circ_pm_muy")
Pu_input = st.sidebar.number_input(f"Carga Axial Última (Pu) [{unidad_fuerza}]", value=st.session_state["circ_pm_pu"], step=round(50.0 * factor_fuerza, 2), key="circ_pm_pu")

st.sidebar.header(_t("6. Esbeltez", "6. Slenderness"))
k_factor = st.sidebar.selectbox(_t("Factor de longitud efectiva (k)", "Effective length factor (k)"),
                                [("Ambos extremos articulados", 1.0),
                                 ("Un extremo articulado, otro empotrado", 0.7),
                                 ("Ambos extremos empotrados", 0.5),
                                 ("Voladizo (base empotrada, libre arriba)", 2.0)],
                                format_func=lambda x: x[0],
                                key="circ_pm_k")[1]

st.sidebar.markdown("---")
st.sidebar.subheader(" Guardar / Cargar Proyecto")

nombre_producido = st.session_state.get("nombre_proyecto_actual", "")

st.sidebar.markdown("**Nuevo Proyecto / Guardar**")
nombre_proy_guardar = st.sidebar.text_input("Nombre para guardar", value=nombre_producido, key="input_guardar_proy")

if st.sidebar.button(" Guardar Proyecto", use_container_width=True):
    if nombre_proy_guardar:
        ok, msg = guardar_proyecto_supabase(nombre_proy_guardar, capturar_estado_actual())
        if ok:
            st.session_state["nombre_proyecto_actual"] = nombre_proy_guardar
            st.sidebar.success(msg)
            st.rerun()
        else:
            st.sidebar.error(msg)
    else:
        st.sidebar.warning("Escribe un nombre de proyecto")

st.sidebar.markdown("**Cargar Proyecto Existente**")
lista_proyectos = listar_proyectos_supabase()

if lista_proyectos:
    idx_default = 0
    if nombre_producido in lista_proyectos:
        idx_default = lista_proyectos.index(nombre_producido)
        
    nombre_proy_cargar = st.sidebar.selectbox("Selecciona un proyecto", lista_proyectos, index=idx_default, key="select_cargar_proy")
    
    def on_cargar_columna_click():
        proy = st.session_state["select_cargar_proy"]
        if proy:
            ok, msg = cargar_proyecto_supabase(proy)
            if ok:
                st.session_state["nombre_proyecto_actual"] = proy
                st.session_state["__msg_cargar_col"] = (True, msg)
            else:
                st.session_state["__msg_cargar_col"] = (False, msg)

    st.sidebar.button(" Cargar", on_click=on_cargar_columna_click, use_container_width=True)

    if "__msg_cargar_col" in st.session_state:
        ok, msg = st.session_state.pop("__msg_cargar_col")
        if ok: st.sidebar.success(msg)
        else: st.sidebar.error(msg)
else:
    st.sidebar.info("No hay proyectos en la nube.")

st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    <br><br>
    <i>⚠ Nota Legal: Herramienta de apoyo profesional.</i>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# CÁLCULOS DE CAPACIDAD UNIAXIAL (X y Y)
# =============================================================================
layers_y = []
if es_circular:
    cap_x = compute_uniaxial_capacity_circular(D, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza)
    cap_y = cap_x
else:
    cap_x = compute_uniaxial_capacity(b, h, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza)
    layers_y.append({'d': d_prime, 'As': num_filas_v * rebar_area})
    # Bug Fix: eje Y depende de num_filas_v (dirección b), no de num_filas_h
    num_capas_intermedias_y = num_filas_v - 2
    if num_capas_intermedias_y > 0:
        espacio_y = (b - 2 * d_prime) / (num_capas_intermedias_y + 1)
        for i in range(1, num_capas_intermedias_y + 1):
            layers_y.append({'d': d_prime + i * espacio_y, 'As': 2 * rebar_area})
    layers_y.append({'d': b - d_prime, 'As': num_filas_v * rebar_area})
    
    cap_y = compute_uniaxial_capacity(h, b, d_prime, layers_y, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza)

# 2) Siempre conservador en Bresler Biaxial, usando phi_c_max
phi_factor = phi_c_max
bresler = biaxial_bresler(Pu_input, Mux_input, Muy_input, cap_x, cap_y, cap_x['Po'], phi_factor)

# ═══════════════════════════════════════════════════════════════
# BLOQUE: COMPRESIÓN AXIAL PURA — Verificación paso a paso
# ═══════════════════════════════════════════════════════════════
with st.expander("Compresión Axial Pura — Verificación Paso a Paso (NSR-10 C.9.3.2.2)", expanded=False):

    # Cálculo desglosado
    Ag_cm2  = Ag              # Ag ya está en cm²
    Ast_cm2 = Ast             # Ast ya está en cm²
    Anc_cm2 = Ag_cm2 - Ast_cm2

    Po_kN       = (0.85 * fc * Anc_cm2 + fy * Ast_cm2) / 10.0
    Pn_max_kN   = p_max_factor * Po_kN
    phi_Pn_max_kN = phi_c_max * Pn_max_kN

    Po_out     = Po_kN        * factor_fuerza
    Pn_max_out = Pn_max_kN    * factor_fuerza
    phi_Pn_out = phi_Pn_max_kN * factor_fuerza

    # Panel visual estilo consola
    st.markdown(f"""
    <div style="background:#1c2e1c;border-radius:10px;padding:14px 18px;
                font-family:monospace;font-size:13px;color:#e8f5e9;line-height:2.1">
    <b style="font-size:15px;color:#81c784"> Resistencia máxima a compresión axial pura:</b><br><br>
    <span style="color:#aaa">Po = [0.85·f'c·(Ag − Ast) + fy·Ast] / 1000</span><br>
    <span style="color:#aaa">Pn,máx = {p_max_factor:.2f} × Po &nbsp;&nbsp;|&nbsp;&nbsp; φPn,máx = φ × Pn,máx</span><br><br>
    <b>Ag</b> (área bruta)         = <b>{Ag_cm2:.2f} cm²</b><br>
    <b>Ast</b> (área acero)        = <b>{Ast_cm2:.2f} cm²</b><br>
    <b>Ag − Ast</b> (concreto neto) = <b>{Anc_cm2:.2f} cm²</b><br>
    <span style="color:#aaa">─────────────────────────────────────────</span><br>
    <b>Po</b>     = [0.85 × {fc:.1f} × {Anc_cm2:.2f} + {fy:.0f} × {Ast_cm2:.2f}] / 1000
             = <b>{Po_out:.1f} {unidad_fuerza}</b><br>
    <b>Pn,máx</b> = {p_max_factor:.2f} × {Po_out:.1f}
             = <b>{Pn_max_out:.1f} {unidad_fuerza}</b>
    <span style="color:#aaa;font-size:11px">
        {'→ Estribos: 0.80' if p_max_factor == 0.80 else '→ Espiral: 0.85'}  —  {code['ref']} C.9.3.2.2
    </span><br>
    <b style="color:#81c784">φPn,máx</b> = {phi_c_max:.2f} × {Pn_max_out:.1f}
             = <b style="color:#a5d6a7;font-size:16px">{phi_Pn_out:.1f} {unidad_fuerza}</b>
    <span style="color:#aaa;font-size:11px">
        {'→ φ = 0.65 estribos' if phi_c_max == 0.65 else '→ φ = 0.75 espiral'}
    </span><br>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    # Métricas visuales
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Po — Resistencia bruta",
                f"{Po_out:.1f} {unidad_fuerza}",
                help="Capacidad axial sin factores de reducción")
    col2.metric(f"Pn,máx  (×{p_max_factor:.2f})",
                f"{Pn_max_out:.1f} {unidad_fuerza}",
                help=f"Límite por excentricidad accidental — {code['ref']} C.9.3.2.2")
    col3.metric(f"φPn,máx  (φ={phi_c_max:.2f})",
                f"{phi_Pn_out:.1f} {unidad_fuerza}",
                delta="↑ Punto superior del diagrama P-M",
                help="Resistencia de diseño — punto (M=0, P=φPn,máx)")
    col4.metric("Utilización Pu / φPn,máx",
                f"{Pu_input / phi_Pn_max_kN * factor_fuerza / factor_fuerza * 100:.1f} %"
                if phi_Pn_max_kN > 0 else "—",
                help="Porcentaje de uso de la capacidad axial máxima")

    # Nota normativa
    tipo_col = "espiral" if (es_circular or ("Espiral" in col_type if not es_circular else True)) else "estribos"
    st.caption(
        f" **{code['ref']} C.9.3.2.2** — Para columnas con **{tipo_col}**, "
        f"la resistencia axial máxima se limita a **{p_max_factor:.0%}·Po** "
        f"para considerar excentricidades mínimas accidentales. "
        f"φ = **{phi_c_max:.2f}** para compresión con {tipo_col}."
    )

    # Semáforo Pu vs φPn,máx
    Pu_kN_check = Pu_input / factor_fuerza
    if Pu_kN_check > phi_Pn_max_kN:
        st.error(
            f" **Pu = {Pu_input:.1f} {unidad_fuerza}** supera **φPn,máx = {phi_Pn_out:.1f} {unidad_fuerza}**. "
            f"La columna NO puede soportar esta carga. Aumente la sección o el acero."
        )
    elif Pu_kN_check > 0.90 * phi_Pn_max_kN:
        st.warning(
            f"⚠ Pu representa el **{Pu_kN_check / phi_Pn_max_kN * 100:.1f}%** de φPn,máx. "
            f"Zona muy próxima al límite de capacidad axial."
        )
    else:
        st.success(
            f" Pu = {Pu_input:.1f} {unidad_fuerza} → **{Pu_kN_check / phi_Pn_max_kN * 100:.1f}%** "
            f"de φPn,máx = {phi_Pn_out:.1f} {unidad_fuerza}. Capacidad axial suficiente."
        )
# ═══════════════════════════════════════════════════════════════

if es_circular:
    slenderness = check_slenderness(L_col, D, D, k_factor, Pu_input, Mux_input, Mux_input, fc, fy, Es, factor_fuerza, es_circular=True)
else:
    slenderness = check_slenderness(L_col, b, h, k_factor, Pu_input, Mux_input, Mux_input, fc, fy, Es, factor_fuerza)

Mux_magnified = Mux_input * slenderness['delta_ns']
Muy_magnified = Muy_input * slenderness['delta_ns']

# =============================================================================
# CÁLCULO DE ESTRIBOS Y VERIFICACIÓN Ash
# =============================================================================
if not es_circular:
    recub_cm = max(d_prime - rebar_diam / 20.0, 2.5)
    # NSR-10 C.7.7.1 — recubrimiento mínimo para columnas
    _recub_min_nsr = 3.8  # cm — columnas expuestas o en ambiente normal
    if recub_cm < _recub_min_nsr:
        st.sidebar.warning(f"⚠ NSR-10 C.7.7.1: Recubrimiento calculado ({recub_cm:.1f} cm) < mínimo recomendado de {_recub_min_nsr} cm para columnas. Verifique d'.")
    bc = b - 2 * recub_cm
    hc = h - 2 * recub_cm
    Ach = bc * hc
    
    claro_libre_x = (b - 2 * d_prime) / (num_filas_h - 1) - rebar_diam / 10 if num_filas_h > 1 else 0
    claro_libre_y = (h - 2 * d_prime) / (num_filas_v - 1) - rebar_diam / 10 if num_filas_v > 1 else 0
    
    num_flejes_y = max(0, math.ceil((b - 2 * d_prime - 15) / 15)) if claro_libre_x > 15.0 else 0
    num_flejes_x = max(0, math.ceil((h - 2 * d_prime - 15) / 15)) if claro_libre_y > 15.0 else 0
    
    ramas_x = 2 + num_flejes_y
    ramas_y = 2 + num_flejes_x

    Ash_prov_x = ramas_x * stirrup_area
    Ash_prov_y = ramas_y * stirrup_area
    Ash_prov = min(Ash_prov_x, Ash_prov_y)

    s1 = 16 * rebar_diam / 10
    s2 = 48 * stirrup_diam / 10
    s3 = min(b, h)
    s_basico = min(s1, s2, s3)
    
    Lo_conf = max(max(b, h), L_col / 6.0, 45.0)
    
    if es_des:
        s_conf = min(8 * rebar_diam / 10, 24 * stirrup_diam / 10, min(b, h) / 3, 15.0)
        s_centro = min(6 * rebar_diam / 10, s_basico)
        s_centro = max(s_centro, s_conf)
    else:
        # DMO y DMI
        s_conf = min(8 * rebar_diam / 10, 24 * stirrup_diam / 10, min(b, h) / 3, 15.0)
        s_centro = s_basico
    
    fyt = fy
    
    Ash_req_1 = 0.3 * s_conf * bc * fc / fyt * (Ag / Ach - 1)
    Ash_req_2 = 0.09 * s_conf * bc * fc / fyt
    Ash_req = max(Ash_req_1, Ash_req_2)

    ash_ok = Ash_prov >= Ash_req
    
    n_est_por_Lo = math.ceil(Lo_conf / s_conf)
    n_estribos_zona = n_est_por_Lo * 2
    longitud_zona_libre = max(0, L_col - 2 * Lo_conf)
    n_estribos_centro = max(0, math.ceil(longitud_zona_libre / s_centro) - 1) if longitud_zona_libre > 0 else 0
    n_estribos_total = n_estribos_zona + n_estribos_centro + 1
    
    ld_mm = get_development_length(rebar_diam, fy, fc)
    splice_length_mm = 1.3 * ld_mm
    splice_zone_height = splice_length_mm / 10
    splice_start = max(L_col / 3, 0)
    splice_end = splice_start + splice_zone_height
    if splice_end > L_col:
        splice_end = L_col
        splice_start = max(0, L_col - splice_zone_height)
    
    # Para usar en despiece
    perim_estribo = 2 * (b - 2 * recub_cm) + 2 * (h - 2 * recub_cm) + 12 * stirrup_diam / 10
    long_fleje_x = h - 2 * recub_cm + 2 * 6 * stirrup_diam / 10
    long_fleje_y = b - 2 * recub_cm + 2 * 6 * stirrup_diam / 10
    perim_estribo += (num_flejes_x * long_fleje_x + num_flejes_y * long_fleje_y)
else:
    recub_cm = max(d_prime - rebar_diam / 20.0, 2.5)
    _recub_min_nsr = 3.8
    if recub_cm < _recub_min_nsr:
        st.sidebar.warning(f"⚠ NSR-10 C.7.7.1: Recubrimiento calculado ({recub_cm:.1f} cm) < mínimo de {_recub_min_nsr} cm.")
    dc = D - 2 * recub_cm
    Ach = math.pi * (dc/2)**2
    Ag_circ = math.pi * (D/2)**2
    
    rho_s_req = 0.45 * (Ag_circ / Ach - 1) * fc / fy
    rho_s_req = max(rho_s_req, 0.12 * fc / fy)
    
    area_spiral = stirrup_area
    rho_s_prov = (4 * area_spiral) / (dc * paso_espiral)
    ash_ok = rho_s_prov >= rho_s_req
    n_estribos_total = math.ceil(L_col / paso_espiral) + 1
    
    ld_mm = get_development_length(rebar_diam, fy, fc)
    splice_length_mm = 1.3 * ld_mm
    splice_zone_height = splice_length_mm / 10
    splice_start = max(L_col / 3, 0)
    splice_end = splice_start + splice_zone_height
    if splice_end > L_col:
        splice_end = L_col
        splice_start = max(0, L_col - splice_zone_height)
    
    # Para usar en despiece circular
    long_espiral_vuelta = math.sqrt((math.pi * dc)**2 + paso_espiral**2)
    long_espiral_total = long_espiral_vuelta * (L_col / paso_espiral)

# =============================================================================
# CANTIDADES DE MATERIALES
# =============================================================================
vol_concreto_m3 = (Ag / 10000) * (L_col / 100) if not es_circular else (math.pi * (D/2)**2 / 10000) * (L_col / 100)
peso_acero_long_kg = Ast * (L_col * 10) * 7.85e-3

if not es_circular:
    peso_unit_estribo = (perim_estribo / 100.0) * (stirrup_area * 100) * 7.85e-3
    peso_total_estribos_kg = n_estribos_total * peso_unit_estribo
else:
    peso_total_estribos_kg = long_espiral_total / 100 * (stirrup_area * 100) * 7.85e-3

peso_total_acero_kg = peso_acero_long_kg + peso_total_estribos_kg

# =============================================================================
# GRÁFICO 3D BIAXIAL
# =============================================================================
def create_biaxial_3d_plot(cap_x, cap_y, Pu, Mux, Muy, phi_factor_eval):
    Mx_vals = np.linspace(0, max(cap_x['phi_M_n']) * 1.1, 30)
    My_vals = np.linspace(0, max(cap_y['phi_M_n']) * 1.1, 30)
    Mx_grid, My_grid = np.meshgrid(Mx_vals, My_vals)
    P_grid = np.zeros_like(Mx_grid)
    
    for i in range(len(Mx_vals)):
        for j in range(len(My_vals)):
            res = biaxial_bresler(Pu, Mx_grid[i,j], My_grid[i,j], cap_x, cap_y, cap_x['Po'], phi_factor_eval)
            P_grid[i,j] = res['phi_Pni'] if res['phi_Pni'] > 0 else 0
    
    fig = go.Figure()
    fig.add_trace(go.Surface(x=Mx_grid, y=My_grid, z=P_grid, colorscale='Viridis', opacity=0.85,
                             name='Superficie de Interacción', showscale=True,
                             colorbar=dict(title=f"φPn [{unidad_fuerza}]")))
    fig.add_trace(go.Scatter3d(x=[Mux], y=[Muy], z=[Pu], mode='markers',
                               marker=dict(size=8, color='red', symbol='circle'),
                               name=f'Punto de Diseño (Mux={Mux:.1f}, Muy={Muy:.1f}, Pu={Pu:.1f})'))
    fig.update_layout(title=_t("Superficie de Interacción Biaxial (Método de Bresler)", "Biaxial Interaction Surface (Bresler Method)"),
                      scene=dict(xaxis_title=f"Mux [{unidad_mom}]", yaxis_title=f"Muy [{unidad_mom}]",
                                 zaxis_title=f"φPn [{unidad_fuerza}]", aspectmode='manual', aspectratio=dict(x=1, y=1, z=0.8)),
                      height=600, margin=dict(l=0, r=0, b=0, t=40))
    return fig

def create_pm_2d_plot(cap_x, Pu, Mu, unidad_mom, unidad_fuerza):
    fig, ax = plt.subplots(figsize=(8, 6))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.plot(cap_x['M_n'], cap_x['P_n'], 'b--', linewidth=1.5, label=r"Resistencia Nominal ($P_n, M_n$)")
    ax.plot(cap_x['phi_M_n'], cap_x['phi_P_n'], 'r-', linewidth=2.5, label=r"Resistencia de Diseño ($\phi P_n, \phi M_n$)")
    
    P_bal = cap_x.get('P_balance', None)
    P_ten = cap_x.get('P_tension', None)
    
    # Rellenar zonas de compresión, transición y tracción
    if P_bal is not None:
        ax.plot(cap_x['M_balance'], cap_x['P_balance'], 'ro', markersize=6, markeredgecolor='black', label=f"Falla Balanceada ($ε_t = ε_y$)")
        ax.axhline(P_bal, color='orange', linestyle='--', alpha=0.5)
    if P_ten is not None:
        ax.plot(cap_x['M_tension'], cap_x['P_tension'], 'go', markersize=6, markeredgecolor='black', label=f"Control a Tracción ($ε_t = 0.005$)")
        ax.axhline(P_ten, color='green', linestyle='--', alpha=0.5)
    
    ymax = max(cap_x['P_n']) * 1.05
    ymin = min(cap_x['P_n']) * 1.05
    xmax = max(cap_x['M_n']) * 1.15
    if P_bal is not None and P_ten is not None:
        ax.fill_between([0, xmax], P_bal, ymax, color='orange', alpha=0.07, label='Interacción Compresión')
        ax.fill_between([0, xmax], P_ten, P_bal, color='yellow', alpha=0.1, label='Transición')
        ax.fill_between([0, xmax], ymin, P_ten, color='green', alpha=0.07, label='Interacción Tracción')

    ax.plot(Mu, Pu, 'o', markersize=9, color='lime', markeredgecolor='black', zorder=6,
            label=f"Punto de Diseño (Mu={Mu:.1f}, Pu={Pu:.1f})")
    ax.axhline(y=0, color='gray', linestyle=':', alpha=0.5)
    ax.axvline(x=0, color='gray', linestyle=':', alpha=0.5)
    ax.set_xlabel(f"Momento Flector M [{unidad_mom}]")
    ax.set_ylabel(f"Carga Axial P [{unidad_fuerza}]")
    ax.set_title(f"Diagrama de Interacción P-M — {norma_sel}\n" +
                 f"Columna {b:.0f}×{h:.0f} cm | f'c={fc:.1f} MPa | fy={fy:.0f} MPa")
    ax.grid(True, linestyle=":", alpha=0.5)
    ax.legend(loc="upper right", fontsize=8)
    return fig

fig_pm_2d = create_pm_2d_plot(cap_x, Pu_input, Mux_input, unidad_mom, unidad_fuerza)
pm_2d_img = io.BytesIO()
fig_pm_2d.savefig(pm_2d_img, format='png', dpi=150, bbox_inches='tight')
pm_2d_img.seek(0)
fig_3d = create_biaxial_3d_plot(cap_x, cap_y, Pu_input, Mux_input, Muy_input, phi_factor)

try:
    pm_3d_img = io.BytesIO()
    fig_3d.write_image(pm_3d_img, format="png")
    pm_3d_img.seek(0)
    has_3d_img = True
except Exception:
    has_3d_img = False

# =============================================================================
# TABS PRINCIPALES
# =============================================================================
tab1, tab2, tab3, tab4 = st.tabs([
    _t(" Diagrama P-M Biaxial", " Biaxial P-M Diagram"),
    _t(" Sección & Estribos", " Section & Ties"),
    _t(" Cantidades y APU", " Quantities & APU"),
    _t(" Memoria de Cálculo", " Calculation Report")
])

# =============================================================================
# TAB 1: DIAGRAMA P-M BIAXIAL
# =============================================================================
with tab1:
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader(_t("Diagrama P-M 2D (Eje X)", " P-M 2D Diagram (X-Axis)"))
        st.pyplot(fig_pm_2d)
        st.subheader(_t("Superficie de Interacción Biaxial 3D", " Biaxial Interaction Surface 3D"))
        st.plotly_chart(fig_3d, use_container_width=True)
    with col2:
        st.subheader(_t("Verificación Biaxial (Bresler)", " Biaxial Verification (Bresler)"))
        st.markdown(f"""
        | Parámetro | Valor |
        |-----------|-------|
        | **φPnx** (para Mux={Mux_input:.1f}) | {bresler['phi_Pnx']:.2f} {unidad_fuerza} |
        | **φPny** (para Muy={Muy_input:.1f}) | {bresler['phi_Pny']:.2f} {unidad_fuerza} |
        | **φP0** (axial pura) | {bresler['phi_P0']:.2f} {unidad_fuerza} |
        | **φPni** (Bresler) | {bresler['phi_Pni']:.2f} {unidad_fuerza} |
        | **Pu** solicitante | {Pu_input:.2f} {unidad_fuerza} |
        | **Relación Pu/φPni** | {bresler['ratio']:.3f} |
        """)
        if bresler['ok']:
            st.success(f"**VERIFICACIÓN BIAXIAL CUMPLE**\n\nPu ({Pu_input:.1f}) ≤ φPni ({bresler['phi_Pni']:.1f})")
        else:
            st.error(f"**VERIFICACIÓN BIAXIAL NO CUMPLE**\n\nPu ({Pu_input:.1f}) > φPni ({bresler['phi_Pni']:.1f})")
            ratio = bresler['ratio']
            deficit = Pu_input - bresler['phi_Pni']
            st.markdown("** Recomendaciones para cumplir:**")
            recomendaciones = []
            if ratio > 1.5 and math.isfinite(ratio):
                pct = math.ceil((ratio**0.5 - 1) * 100)
                recomendaciones.append(f" **Aumentar sección:** La columna necesita ~{pct}% más de área — aumentar b y/o h en el sidebar.")
            if ratio <= 3:
                rho_obj = min(rho_max, cuantia * ratio**0.5)
                recomendaciones.append(f" **Aumentar acero longitudinal:** Añadir varillas o usar diámetro mayor — apuntar a ρ ≥ {rho_obj:.1f}%.")
            if ratio > 2:
                recomendaciones.append(f" **Reducir Pu:** La carga axial ({Pu_input:.0f} {unidad_fuerza}) supera {ratio:.1f}x la capacidad — revisar predimensionamiento.")
            if Mux_input > 0 or Muy_input > 0:
                recomendaciones.append(" **Reducir momentos Mux/Muy:** Considerar arriostrar la estructura o reducir excentricidades.")
            for rec in recomendaciones:
                st.markdown(f"- {rec}")
            st.markdown(f"""
| | |
|---|---|
| **Capacidad requerida** | φPni ≥ {Pu_input:.1f} {unidad_fuerza} |
| **Capacidad actual** | φPni = {bresler['phi_Pni']:.1f} {unidad_fuerza} |
| **Déficit** | {deficit:.1f} {unidad_fuerza} ({(ratio-1)*100:.0f}% sobre la capacidad) |
""")
            if ratio > 5:
                st.error("⚠ **Relación > 5x:** La sección es muy insuficiente. Se recomienda rediseñar completamente la geometría.")
        st.markdown("---")
        st.subheader(_t("Verificación de Esbeltez", " Slenderness Verification"))
        st.markdown(f"""
        | Parámetro | Valor | Estado |
        |-----------|-------|--------|
        | **kl/r** | {slenderness['kl_r']:.1f} | |
        | **Clasificación** | {slenderness['classification']} | |
        | **δns (magnificación)** | {slenderness['delta_ns']:.3f} | |
        | **Mux magnificado** | {Mux_magnified:.2f} {unidad_mom} | |
        | **Muy magnificado** | {Muy_magnified:.2f} {unidad_mom} | |
        """)
        if slenderness['kl_r'] > 100:
            st.warning("⚠ **kl/r > 100** — Se requiere análisis no lineal según NSR-10 C.10.10.7")
        st.markdown("---")
        st.subheader(_t("Verificación de Estribos / Espiral", " Tie / Spiral Verification"))
        if not es_circular:
            req_1_str = f"0.3 \\times {s_conf:.1f} \\times {bc:.1f} \\times ({fc:.1f}/{fyt:.0f}) \\times ({Ag:.1f}/{Ach:.1f} - 1) = {Ash_req_1:.2f} \\text{{ cm}}^2"
            req_2_str = f"0.09 \\times {s_conf:.1f} \\times {bc:.1f} \\times ({fc:.1f}/{fyt:.0f}) = {Ash_req_2:.2f} \\text{{ cm}}^2"
            
            st.markdown(f"**Cálculo de Ash requerido (NSR-10 C.21.3.5.4):**")
            st.latex(r"(a) \quad A_{sh} = " + req_1_str)
            st.latex(r"(b) \quad A_{sh} = " + req_2_str)
            st.markdown(f"**→ Rige: {Ash_req:.2f} cm²**")

            st.markdown(f"""
            | Parámetro | Valor | Requerido | Estado |
            |-----------|-------|-----------|--------|
            | **Claro Libre (Cx, Cy)** | {claro_libre_x:.1f} cm, {claro_libre_y:.1f} cm | ≤ 15 cm | {'' if claro_libre_x<=15 and claro_libre_y<=15 else '⚠ Crossties Requeridos'} |
            | **Apoyo lateral (Crossties)** | {num_flejes_x} en X, {num_flejes_y} en Y | NSR-10 C.7.10.5 | |
            | **Ramas Efectivas** | {ramas_x} ramas en X, {ramas_y} ramas en Y | | |
            | **Ash provisto** | {Ash_prov:.3f} cm² | ≥ {Ash_req:.2f} | {'' if ash_ok else ''} |
            | **s_conf** | {s_conf:.1f} cm | {'≤ 15' if es_des else '≤ 20' if es_dmo else '≤ min(b,h)'} | |
            | **Lo_conf** | {Lo_conf:.1f} cm | ≥ max(b,h,L/6,45) | {''} |
            | **N° estribos + Crossties** | {n_estribos_total} juegos de ramas | C.21.3.5 | |
            """)

            if not ash_ok:
                ratio_ash = Ash_prov / Ash_req if Ash_req > 0 else 1.0
                s_req1 = Ash_prov / (0.3 * bc * fc / fyt * (Ag/Ach - 1)) if (Ag/Ach - 1) > 0 else 999
                s_req2 = Ash_prov * fyt / (0.09 * bc * fc)
                s_correcto = min(s_req1, s_req2)
                
                if ratio_ash < 0.5:
                    st.error(f"⚠ **Déficit crítico de estribos.** Para cumplir con las estribos actuales, usar separación $s \\le {s_correcto:.1f}$ cm o proponer más ramas.")
                else:
                    st.warning(f"Para cumplir Ash con los estribos actuales → reducir separación a $s \\le {s_correcto:.1f}$ cm.")
        else:
            st.markdown(f"""
            | Parámetro | Valor | Requerido | Estado |
            |-----------|-------|-----------|--------|
            | **ρs requerido** | {rho_s_req:.4f} | | |
            | **ρs provisto** | {rho_s_prov:.4f} | ≥ {rho_s_req:.4f} | {'' if ash_ok else ''} |
            | **Paso espiral** | {paso_espiral:.1f} cm | ≤ min(D/5, 8 cm) | {'' if paso_espiral <= min(D/5, 8) else ''} |
            | **N° vueltas** | {n_estribos_total} | | |
            """)
        st.caption(f"Ref: {code['ref']} | Nivel Sísmico: {nivel_sismico}")

# =============================================================================
# TAB 2: SECCIÓN Y ESTRIBOS (con DXF y RÓTULO ICONTEC)
# =============================================================================
with tab2:
    st.subheader(_t("Visualización 3D de la Columna", " 3D Column Visualization"))
    fig3d_col = go.Figure()
    if es_circular:
        theta = np.linspace(0, 2*np.pi, 50)
        z = np.linspace(0, L_col, 20)
        theta_grid, z_grid = np.meshgrid(theta, z)
        x_grid = (D/2) * np.cos(theta_grid)
        y_grid = (D/2) * np.sin(theta_grid)
        fig3d_col.add_trace(go.Surface(x=x_grid, y=y_grid, z=z_grid, opacity=0.3, colorscale='Greys', showscale=False))
    else:
        x_c = [-b/2, b/2, b/2, -b/2, -b/2, b/2, b/2, -b/2]
        y_c = [-h/2, -h/2, h/2, h/2, -h/2, -h/2, h/2, h/2]
        z_c = [0, 0, 0, 0, L_col, L_col, L_col, L_col]
        fig3d_col.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
        
    # ── BARRAS LONGITUDINALES 3D ──────────────────────────
    z_barras = [0, L_col]
    if es_circular:
        radio_c = D/2 - d_prime
        for layer in layers:
            xb = layer.get('x', 0)
            yb = layer.get('y', 0)
            fig3d_col.add_trace(go.Scatter3d(
                x=[xb, xb], y=[yb, yb], z=z_barras,
                mode='lines',
                line=dict(color='#ff6b35', width=4),
                showlegend=False, name='Barra long.'
            ))
    else:
        r_bar3d = rebar_diam / 20.0
        xs3d = np.linspace(d_prime - b/2, b/2 - d_prime, num_filas_h) if num_filas_h > 1 else [0.0]
        ys3d_bot = -(h/2 - d_prime)
        ys3d_top =   h/2 - d_prime
        for x3 in xs3d:
            for y3 in [ys3d_bot, ys3d_top]:
                fig3d_col.add_trace(go.Scatter3d(
                    x=[x3, x3], y=[y3, y3], z=z_barras,
                    mode='lines',
                    line=dict(color='#ff6b35', width=4),
                    showlegend=False, name='Barra long.'
                ))
        if num_capas_intermedias > 0:
            esp3d = (h - 2*d_prime) / (num_capas_intermedias + 1)
            for ci in range(1, num_capas_intermedias + 1):
                yi = -(h/2) + d_prime + ci * esp3d
                for xi in [-(b/2 - d_prime), b/2 - d_prime]:
                    fig3d_col.add_trace(go.Scatter3d(
                        x=[xi, xi], y=[yi, yi], z=z_barras,
                        mode='lines',
                        line=dict(color='#ff6b35', width=4),
                        showlegend=False, name='Barra lat.'
                    ))

    # ── ESTRIBOS 3D ───────────────────────────────────────
    if not es_circular:
        paso_3d = s_conf if (es_des or es_dmo) else s_basico
        z_est = np.arange(0, L_col + paso_3d, paso_3d)
        bw = b/2 - recub_cm
        hw = h/2 - recub_cm
        estr_x = [-bw, bw, bw, -bw, -bw]
        estr_y = [-hw, -hw, hw, hw, -hw]
        for ze in z_est:
            fig3d_col.add_trace(go.Scatter3d(
                x=estr_x, y=estr_y,
                z=[ze]*5,
                mode='lines',
                line=dict(color='#00d4ff', width=2),
                showlegend=False, name='Estribo'
            ))
    else:
        rc_sp = D/2 - recub_cm
        theta_sp = np.linspace(0, 2*np.pi, 60)
        z_sp = np.arange(0, L_col + paso_espiral, paso_espiral)
        for ze in z_sp:
            fig3d_col.add_trace(go.Scatter3d(
                x=rc_sp * np.cos(theta_sp),
                y=rc_sp * np.sin(theta_sp),
                z=[ze]*len(theta_sp),
                mode='lines',
                line=dict(color='#00d4ff', width=2),
                showlegend=False, name='Espiral'
            ))

    fig3d_col.update_layout(scene=dict(aspectmode='manual', aspectratio=dict(x=1,y=1,z=1.2), xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                            height=450, margin=dict(l=0, r=0, b=0, t=0))
    st.plotly_chart(fig3d_col, use_container_width=True)
    st.markdown("---")
    st.subheader(_t("Sección Transversal", " Cross Section"))
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        fig_sec, ax_s = plt.subplots(figsize=(5, 5))
        fig_sec.patch.set_facecolor('#1e1e2e')
        for _ax in fig_sec.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
        ax_s.set_aspect('equal')
        ax_s.set_facecolor('#1a1a2e')
        fig_sec.patch.set_facecolor('#1a1a2e')
        if es_circular:
            circle = plt.Circle((0, 0), D/2, linewidth=2, edgecolor='white', facecolor='#4a4a6a', fill=True)
            ax_s.add_patch(circle)
            circle_rec = plt.Circle((0, 0), D/2 - recub_cm, linewidth=1.5, edgecolor='#00d4ff', facecolor='none', linestyle='--')
            ax_s.add_patch(circle_rec)
            radio_centro = D/2 - d_prime
            for layer in layers:
                x_pos = layer.get('x', 0)
                y_pos = layer.get('y', 0)
                ax_s.add_patch(plt.Circle((x_pos, y_pos), rebar_diam/20, color='#ff6b35', zorder=5))
            ax_s.set_xlim(-D/2 - 5, D/2 + 5)
            ax_s.set_ylim(-D/2 - 5, D/2 + 5)
        else:
            ax_s.add_patch(patches.Rectangle((0, 0), b, h, linewidth=2, edgecolor='white', facecolor='#4a4a6a'))
            ax_s.add_patch(patches.Rectangle((recub_cm, recub_cm), b-2*recub_cm, h-2*recub_cm,
                linewidth=1.5, edgecolor='#00d4ff', facecolor='none', linestyle='--'))
            r_bar = rebar_diam / 20.0
            xs = np.linspace(d_prime, b - d_prime, num_filas_h) if num_filas_h > 1 else [b/2]
            for x in xs:
                ax_s.add_patch(plt.Circle((x, h - d_prime), r_bar, color='#ff6b35', zorder=5))
                ax_s.add_patch(plt.Circle((x, d_prime), r_bar, color='#ff6b35', zorder=5))
            if num_capas_intermedias > 0:
                esp = (h - 2*d_prime) / (num_capas_intermedias + 1)
                for i in range(1, num_capas_intermedias + 1):
                    y_int = d_prime + i * esp
                    ax_s.add_patch(plt.Circle((d_prime, y_int), r_bar, color='#ff6b35', zorder=5))
                    ax_s.add_patch(plt.Circle((b - d_prime, y_int), r_bar, color='#ff6b35', zorder=5))
            ax_s.set_xlim(-5, b + 5)
            ax_s.set_ylim(-5, h + 5)
        ax_s.axis('off')
        ax_s.set_title(f"Sección {'Circular' if es_circular else 'Rectangular'} — {n_barras} varillas Ø{rebar_diam:.0f}mm", color='white', fontsize=9)
        st.pyplot(fig_sec)
    with col_s2:
        st.subheader(_t("Resumen de Verificaciones", " Verification Summary"))
        checks_data = {
            "Verificación": ["Cuantía longitudinal", "Verificación biaxial", "Esbeltez (kl/r ≤ 22)",
                             f"Ash {'espiral' if es_circular else 'estribos'}", "Longitud confinamiento Lo", "Separación máxima"],
            "Estado": ["" if rho_min <= cuantia <= rho_max else "", "" if bresler['ok'] else "",
                       "" if slenderness['kl_r'] <= 22 else "⚠", "" if ash_ok else "",
                       "",
                       "" if (es_des and s_conf <= 15) or (es_dmo and s_conf <= 20) or es_dmi else ""]
        }
        st.dataframe(pd.DataFrame(checks_data), use_container_width=True, hide_index=True)
        st.markdown("---")
        st.subheader(_t("Exportar Plano DXF (ICONTEC)", " Export DXF (ICONTEC)"))
        
        with st.expander(_t("Configurar Rótulo del Plano", "Configure Title Block"), expanded=True):
            col_r1, col_r2 = st.columns(2)
            dxf_empresa = col_r1.text_input("Empresa", "INGENIERÍA ESTRUCTURAL SAS")
            dxf_proyecto = col_r1.text_input("Proyecto", "Proyecto Estructural")
            dxf_plano = col_r2.text_input("N° Plano", "COL-001")
            dxf_elaboro = col_r1.text_input("Elaboró", "Ing. Diseñador")
            dxf_reviso = col_r2.text_input("Revisó", "Ing. Revisor")
            dxf_aprobo = col_r2.text_input("Aprobó", "Ing. Aprobador")
            qr_url = st.text_input(_t("URL para QR (proyecto en la nube)", "QR URL (cloud project)"), 
                                   value="https://github.com/tuproyecto", key="qr_url")
            
        try:
            doc_dxf = ezdxf.new('R2010', setup=True)
            doc_dxf.units = ezdxf.units.CM
            for lay, col in [('CONCRETO',7),('ESTRIBOS',4),('VARILLAS',1),
                             ('TEXTO',3),('EMPALME',6),('COTAS',5),('ROTULO',2),('MARGEN',8)]:
                if lay not in doc_dxf.layers:
                    doc_dxf.layers.add(lay, color=col)
            msp = doc_dxf.modelspace()

            # ── Zonas del plano A4 horizontal (29.7 x 21.0) ──
            ancho_plano = 29.7
            alto_plano  = 21.0
            rotulo_h    = 6.0
            rotulo_w    = 18.0

            # Margen exterior
            msp.add_lwpolyline(
                [(0.5,0.5),(ancho_plano-0.5,0.5),(ancho_plano-0.5,alto_plano-0.5),
                 (0.5,alto_plano-0.5),(0.5,0.5)],
                dxfattribs={'layer':'MARGEN','color':8})

            # ── Estilo de texto ──
            if 'ROMANS' not in doc_dxf.styles:
                try:    doc_dxf.styles.new('ROMANS', dxfattribs={'font':'romans.shx'})
                except: doc_dxf.styles.new('ROMANS', dxfattribs={'font':'txt.shx'})

            # ── ZONA 1: ALZADO LONGITUDINAL (Lado Izquierdo) ──
            alz_x0 = 1.0
            alz_y0 = 1.0
            alz_w  = 11.5
            alz_h  = 19.0
            
            # Dibujo del Alzado
            escala_alz = (alz_h - 2.5) / L_col
            dim_perfil = h if not es_circular else D
            dib_alz_w = dim_perfil * escala_alz
            escala_alz_x = escala_alz
            
            # Control de proporciones en alzado
            if dib_alz_w < 3.0:
                dib_alz_w = 3.0
                escala_alz_x = dib_alz_w / dim_perfil
            elif dib_alz_w > (alz_w - 4.0):
                dib_alz_w = alz_w - 4.0
                escala_alz_x = dib_alz_w / dim_perfil
                
            ax0 = alz_x0 + (alz_w - dib_alz_w) / 2
            ay0 = alz_y0 + 1.25

            # Contorno de columna en alzado
            msp.add_lwpolyline([(ax0, ay0), (ax0+dib_alz_w, ay0), (ax0+dib_alz_w, ay0+L_col*escala_alz), (ax0, ay0+L_col*escala_alz), (ax0, ay0)], dxfattribs={'layer':'CONCRETO'})
            
            # Símbolos de corte de losas superior e inferior
            msp.add_line((ax0-0.5, ay0), (ax0+dib_alz_w+0.5, ay0), dxfattribs={'layer':'CONCRETO'})
            msp.add_line((ax0-0.5, ay0+L_col*escala_alz), (ax0+dib_alz_w+0.5, ay0+L_col*escala_alz), dxfattribs={'layer':'CONCRETO'})
            msp.add_line((ax0-1.0, ay0-0.3), (ax0+dib_alz_w+1.0, ay0+0.3), dxfattribs={'layer':'CONCRETO'})
            
            # Varillas longitudinales y traslapos (Alzado)
            rec_alz = recub_cm * escala_alz_x
            # Barras principales
            msp.add_line((ax0+rec_alz, ay0), (ax0+rec_alz, ay0+L_col*escala_alz + 1.0), dxfattribs={'layer':'VARILLAS'})
            msp.add_line((ax0+dib_alz_w-rec_alz, ay0), (ax0+dib_alz_w-rec_alz, ay0+L_col*escala_alz + 1.0), dxfattribs={'layer':'VARILLAS'})
            
            # Dibujar traslapos simulados en la base (inclinación 1:6 normativa)
            msp.add_line((ax0+rec_alz, ay0-1.0), (ax0+rec_alz, ay0), dxfattribs={'layer':'VARILLAS'})
            msp.add_line((ax0+rec_alz, ay0), (ax0+rec_alz+0.15, ay0+0.4), dxfattribs={'layer':'VARILLAS'})
            msp.add_line((ax0+rec_alz+0.15, ay0+0.4), (ax0+rec_alz+0.15, ay0+2.5), dxfattribs={'layer':'VARILLAS'})
            
            msp.add_line((ax0+dib_alz_w-rec_alz, ay0-1.0), (ax0+dib_alz_w-rec_alz, ay0), dxfattribs={'layer':'VARILLAS'})
            msp.add_line((ax0+dib_alz_w-rec_alz, ay0), (ax0+dib_alz_w-rec_alz-0.15, ay0+0.4), dxfattribs={'layer':'VARILLAS'})
            msp.add_line((ax0+dib_alz_w-rec_alz-0.15, ay0+0.4), (ax0+dib_alz_w-rec_alz-0.15, ay0+2.5), dxfattribs={'layer':'VARILLAS'})

            # Distribución de estribos (Alzado)
            y_curr = 5.0 # arrancar 5cm arriba del nodo
            while y_curr <= L_col - 5.0:
                in_conf = (y_curr <= Lo_conf) or (y_curr >= L_col - Lo_conf)
                sep = s_conf if in_conf else s_basico
                ye = ay0 + y_curr * escala_alz
                msp.add_line((ax0+rec_alz, ye), (ax0+dib_alz_w-rec_alz, ye), dxfattribs={'layer':'ESTRIBOS'})
                y_curr += sep

            # Cota de altura L
            cx = ax0 - 1.2
            msp.add_line((cx, ay0), (cx, ay0+L_col*escala_alz), dxfattribs={'layer':'COTAS'})
            msp.add_line((ax0, ay0), (cx-0.3, ay0), dxfattribs={'layer':'COTAS'})
            msp.add_line((ax0, ay0+L_col*escala_alz), (cx-0.3, ay0+L_col*escala_alz), dxfattribs={'layer':'COTAS'})
            msp.add_text(f"L = {L_col:.0f} cm", dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.25,'insert':(cx-0.2, ay0+L_col*escala_alz/2),'align_point':(cx-0.2, ay0+L_col*escala_alz/2),'halign':1,'valign':2,'rotation':90})

            # Cota de zonas de confinamiento (Lo)
            cx_r = ax0 + dib_alz_w + 1.2
            if Lo_conf > 0:
                # Inferior
                msp.add_line((cx_r, ay0), (cx_r, ay0+Lo_conf*escala_alz), dxfattribs={'layer':'COTAS'})
                msp.add_line((ax0+dib_alz_w, ay0+Lo_conf*escala_alz), (cx_r+0.3, ay0+Lo_conf*escala_alz), dxfattribs={'layer':'COTAS'})
                msp.add_line((ax0+dib_alz_w, ay0), (cx_r+0.3, ay0), dxfattribs={'layer':'COTAS'})
                msp.add_text(f"Lo={Lo_conf:.0f}cm", dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.22,'insert':(cx_r+0.2, ay0+Lo_conf*escala_alz/2),'align_point':(cx_r+0.2, ay0+Lo_conf*escala_alz/2),'halign':1,'valign':2,'rotation':-90})
                
                # Superior
                msp.add_line((cx_r, ay0+L_col*escala_alz), (cx_r, ay0+(L_col-Lo_conf)*escala_alz), dxfattribs={'layer':'COTAS'})
                msp.add_line((ax0+dib_alz_w, ay0+(L_col-Lo_conf)*escala_alz), (cx_r+0.3, ay0+(L_col-Lo_conf)*escala_alz), dxfattribs={'layer':'COTAS'})
                msp.add_line((ax0+dib_alz_w, ay0+L_col*escala_alz), (cx_r+0.3, ay0+L_col*escala_alz), dxfattribs={'layer':'COTAS'})
                msp.add_text(f"Lo={Lo_conf:.0f}cm", dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.22,'insert':(cx_r+0.2, ay0+(L_col-Lo_conf/2)*escala_alz),'align_point':(cx_r+0.2, ay0+(L_col-Lo_conf/2)*escala_alz),'halign':1,'valign':2,'rotation':-90})

            # ── ZONA 2: SECCIÓN TRANSVERSAL (Top-Center) ──
            sec_zona_x0 = 13.5
            sec_zona_y0 = rotulo_h + 1.0  # 7.0
            sec_zona_ancho = 6.5
            sec_zona_alto  = alto_plano - sec_zona_y0 - 1.0 # 13.0
            
            dim_b = D if es_circular else b
            dim_h = D if es_circular else h
            escala = min((sec_zona_ancho - 1.5) / dim_b, (sec_zona_alto - 2.0) / dim_h, 1.0)
            escala = max(escala, 0.05)

            sec_w = dim_b * escala
            sec_h = dim_h * escala
            ox = sec_zona_x0 + (sec_zona_ancho - sec_w) / 2
            oy = sec_zona_y0 + (sec_zona_alto - sec_h) / 2

            if es_circular:
                cx, cy = ox + sec_w/2, oy + sec_h/2
                msp.add_circle((cx,cy), D/2*escala, dxfattribs={'layer':'CONCRETO'})
                r_esp = (D/2-recub_cm)*escala
                msp.add_circle((cx,cy), r_esp, dxfattribs={'layer':'ESTRIBOS'})
                hl = 4.0 * escala
                msp.add_lwpolyline([(cx, cy + r_esp), (cx - hl*0.5, cy + r_esp - hl*0.5)], dxfattribs={'layer':'ESTRIBOS'})
                for ang in np.linspace(0, 2*np.pi, n_barras, endpoint=False):
                    xb_c = cx + (D/2 - d_prime)*escala * math.cos(ang)
                    yb_c = cy + (D/2 - d_prime)*escala * math.sin(ang)
                    msp.add_circle((xb_c,yb_c), rebar_diam/20*escala, dxfattribs={'layer':'VARILLAS'})
            else:
                msp.add_lwpolyline([(ox,oy),(ox+sec_w,oy),(ox+sec_w,oy+sec_h),(ox,oy+sec_h),(ox,oy)], dxfattribs={'layer':'CONCRETO'})
                re_s = recub_cm * escala
                r_bar = rebar_diam/20 * escala
                dp_s  = d_prime * escala
                x_st_min, x_st_max = ox + re_s, ox + sec_w - re_s
                y_st_min, y_st_max = oy + re_s, oy + sec_h - re_s
                cx_min, cx_max = ox + dp_s, ox + sec_w - dp_s
                cy_min, cy_max = oy + dp_s, oy + sec_h - dp_s
                
                hl = 4.0 * escala  # gancho a 135 realista curvado
                pts_estribo = [
                    (x_st_min + hl*0.6, y_st_max - hl*0.9), (x_st_min, y_st_max - hl*0.3),
                    (x_st_min, y_st_min), (x_st_max, y_st_min), (x_st_max, y_st_max),
                    (x_st_min + hl*0.3, y_st_max), (x_st_min + hl*0.9, y_st_max - hl*0.6)
                ]
                msp.add_lwpolyline(pts_estribo, dxfattribs={'layer':'ESTRIBOS'})
                
                xs_b = np.linspace(cx_min, cx_max, num_filas_h) if num_filas_h > 1 else [ox + sec_w/2]
                for xb_c in xs_b:
                    msp.add_circle((xb_c, cy_max), r_bar, dxfattribs={'layer':'VARILLAS'})
                    msp.add_circle((xb_c, cy_min), r_bar, dxfattribs={'layer':'VARILLAS'})
                
                if num_filas_h > 2:
                    for xb_c in xs_b[1:-1]:
                        pts_grapa_v = [(xb_c - hl*0.5, y_st_max - hl*0.5), (xb_c, y_st_max), (xb_c, y_st_min), (xb_c + hl*0.5, y_st_min + hl*0.5)]
                        msp.add_lwpolyline(pts_grapa_v, dxfattribs={'layer':'ESTRIBOS'})

                if num_capas_intermedias > 0:
                    ys_h = np.linspace(cy_min, cy_max, num_filas_v)
                    for yi in ys_h[1:-1]:
                        msp.add_circle((cx_min, yi), r_bar, dxfattribs={'layer':'VARILLAS'})
                        msp.add_circle((cx_max, yi), r_bar, dxfattribs={'layer':'VARILLAS'})
                        pts_grapa_h = [(x_st_min + hl*0.5, yi + hl*0.5), (x_st_min, yi), (x_st_max, yi), (x_st_max - hl*0.5, yi - hl*0.5)]
                        msp.add_lwpolyline(pts_grapa_h, dxfattribs={'layer':'ESTRIBOS'})

                # Cotas dentro del plano (Sección)
                yc = oy + sec_h + 0.6
                msp.add_line((ox, yc),(ox+sec_w, yc), dxfattribs={'layer':'COTAS'})
                msp.add_line((ox, oy+sec_h),(ox, yc+0.3), dxfattribs={'layer':'COTAS'})
                msp.add_line((ox+sec_w, oy+sec_h),(ox+sec_w, yc+0.3), dxfattribs={'layer':'COTAS'})
                msp.add_text(f"b={b:.0f}cm", dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.27,'insert':(ox+sec_w/2, yc+0.12),'align_point':(ox+sec_w/2, yc+0.12),'halign':1,'valign':2})
                
                xc = ox + sec_w + 0.6
                msp.add_line((xc, oy),(xc, oy+sec_h), dxfattribs={'layer':'COTAS'})
                msp.add_line((ox+sec_w, oy),(xc+0.3, oy), dxfattribs={'layer':'COTAS'})
                msp.add_line((ox+sec_w, oy+sec_h),(xc+0.3, oy+sec_h), dxfattribs={'layer':'COTAS'})
                msp.add_text(f"h={h:.0f}cm", dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.27,'insert':(xc+0.12, oy+sec_h/2),'align_point':(xc+0.12, oy+sec_h/2),'halign':1,'valign':2})
                
                msp.add_text(f"rec={recub_cm:.0f}cm", dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.22,'insert':(ox+re_s*0.3, oy+re_s*0.3)})

            # ── ZONA 3: TEXTOS DESCRIPTIVOS (Top-Right Extremo) ──
            tx  = sec_zona_x0 + sec_zona_ancho + 0.5
            ty0 = alto_plano - 1.5
            th  = 0.30
            gap = 0.55

            if not es_circular:
                _s_corr_ash = min(Ash_prov * fy / (0.09 * recub_cm * fc) if recub_cm * fc > 0 else 15, 15.0)
                txt_ash = "CUMPLE" if ash_ok else f"NO CUMPLE - s<={_s_corr_ash:.0f}cm"
                txt_lo  = f"{Lo_conf:.0f} cm" if Lo_conf > 0 else "0 cm - CORREGIR (NSR-10)"
                color_ash = 3 if ash_ok else 1
                color_lo  = 7 if Lo_conf > 0 else 1
                lineas = [
                    (f"SECCION {b:.0f}x{h:.0f} cm", 7),
                    (f"fc={fc:.0f}MPa  fy={fy:.0f}MPa", 7),
                    (f"{n_barras_total} {_bar_label(rebar_diam)}", 7),
                    (f"Ast={Ast:.2f}cm2  rho={cuantia:.2f}%", 7),
                    (f"Estribos {_bar_label(stirrup_diam)}", 7),
                    (f"  @ {s_conf:.0f}cm (zona conf.)", 7),
                    (f"  @ {s_basico:.0f}cm (centro)", 7),
                    (f"Lo_conf = {txt_lo}", color_lo),
                    (f"Ash req={Ash_req:.2f} cm2", color_ash),
                    (f"Ash prov={Ash_prov:.2f} cm2", color_ash),
                    (f"Ash [{txt_ash}]", color_ash),
                    (f"Biaxial Pu/Pni = {bresler['ratio']:.3f}", 3 if bresler['ok'] else 1),
                    (f"{'CUMPLE' if bresler['ok'] else 'NO CUMPLE'}", 3 if bresler['ok'] else 1),
                ]
            else:
                lineas = [
                    (f"SECCION CIRCULAR D={D:.0f}cm", 7),
                    (f"fc={fc:.0f}MPa  fy={fy:.0f}MPa", 7),
                    (f"{n_barras} {_bar_label(rebar_diam)}", 7),
                    (f"Ast={Ast:.2f}cm2  rho={cuantia:.2f}%", 7),
                    (f"Espiral {_bar_label(stirrup_diam)}", 7),
                    (f"  paso={paso_espiral:.0f}cm", 7),
                    (f"rho_s req={rho_s_req:.4f}", 7),
                    (f"rho_s prov={rho_s_prov:.4f}", 7),
                    (f"Espiral [{'CUMPLE' if ash_ok else 'NO CUMPLE'}]", 3 if ash_ok else 1),
                    (f"Biaxial Pu/Pni={bresler['ratio']:.3f}", 3 if bresler['ok'] else 1),
                    (f"{'CUMPLE' if bresler['ok'] else 'NO CUMPLE'}", 3 if bresler['ok'] else 1),
                ]

            for i, (linea, color) in enumerate(lineas):
                yi = ty0 - i*gap
                if yi > rotulo_h + 0.3:
                    msp.add_text(linea,
                        dxfattribs={'layer':'TEXTO','style':'ROMANS','height':th,
                                    'insert':(tx, yi), 'color': color})

            # \u2500\u2500 Texto de armadura (zona derecha, l\u00ednea por l\u00ednea) \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500

            # \u2500\u2500 R\u00f3tulo ICONTEC \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
            rotulo_x = ancho_plano - rotulo_w - 0.5
            rotulo_y = 0.5
            rotulo_h_real = 3.6
            msp.add_lwpolyline(
                [(rotulo_x,rotulo_y),(rotulo_x+rotulo_w,rotulo_y),
                 (rotulo_x+rotulo_w,rotulo_y+rotulo_h_real),
                 (rotulo_x,rotulo_y+rotulo_h_real),(rotulo_x,rotulo_y)],
                dxfattribs={'layer':'ROTULO','color':2})

            # ── Tabla de Despiece (Resumen de Cantidades) ──
            tab_w = rotulo_w
            tab_x = rotulo_x
            tab_y_start = rotulo_y + rotulo_h_real + 0.5
            row_h = 0.8
            cols_w = [4.5, 3.5, 3.0, 3.5, 3.5]  # suma = 18.0
            
            if es_circular:
                _long_bar_m = (L_col + 2 * (ld_mm/10) + 2 * (12*rebar_diam/10)) / 100
                _peso_long = n_barras * _long_bar_m * (rebar_area * 100) * 7.85e-3
                _marca_long, _cant_long = "L1 (Long.)", n_barras
                _marca_trans, _cant_trans = "E1 (Espiral)", 1
                _long_trans_m = long_espiral_total / 100
                _peso_trans = peso_total_estribos_kg
                _vol_conc = vol_concreto_m3
            else:
                _long_bar_m = (L_col + 2 * (ld_mm/10) + 2 * (12*rebar_diam/10)) / 100
                _peso_long = n_barras_total * _long_bar_m * (rebar_area * 100) * 7.85e-3
                _marca_long, _cant_long = "L1 (Long.)", n_barras_total
                _marca_trans, _cant_trans = "E1 (Estribo)", n_estribos_total
                _long_trans_m = perim_estribo / 100
                _peso_trans = peso_total_estribos_kg
                _vol_conc = vol_concreto_m3

            filas_tabla = [
                [("RESUMEN DE MATERIALES - PEDIDO DE OBRA", 18.0)],
                [("MARCA", cols_w[0]), ("DIÁMETRO", cols_w[1]), ("CANT.", cols_w[2]), ("LONG. (m)", cols_w[3]), ("PESO (kg)", cols_w[4])],
                [(_marca_long, cols_w[0]), (f"{_bar_label(rebar_diam)}", cols_w[1]), (f"{_cant_long}", cols_w[2]), (f"{_long_bar_m:.2f}", cols_w[3]), (f"{_peso_long:.1f}", cols_w[4])],
                [(_marca_trans, cols_w[0]), (f"{_bar_label(stirrup_diam)}", cols_w[1]), (f"{_cant_trans}", cols_w[2]), (f"{_long_trans_m:.2f}", cols_w[3]), (f"{_peso_trans:.1f}", cols_w[4])],
                [("TOTAL ACERO", cols_w[0]+cols_w[1]+cols_w[2]+cols_w[3]), (f"{_peso_long + _peso_trans:.1f}", cols_w[4])],
                [("CONCRETO (m3)", cols_w[0]+cols_w[1]), (f"fc={fc:.0f}MPa", cols_w[2]+cols_w[3]), (f"{_vol_conc:.3f}", cols_w[4])],
            ]

            cy = tab_y_start + len(filas_tabla)*row_h
            for fila in filas_tabla:
                cx = tab_x
                cy -= row_h
                for texto, cw in fila:
                    msp.add_lwpolyline([(cx, cy), (cx+cw, cy), (cx+cw, cy+row_h), (cx, cy+row_h), (cx, cy)], dxfattribs={'layer':'ROTULO','color':8})
                    msp.add_text(texto, dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20, 'insert':(cx+cw/2, cy+row_h/2), 'align_point':(cx+cw/2, cy+row_h/2), 'halign':1,'valign':2,'color':7})
                    cx += cw

            campos = {
                "EMPRESA":  dxf_empresa,
                "PROYECTO": dxf_proyecto,
                "N. PLANO": dxf_plano,
                "ESCALA":   "1:20",
                "FECHA":    datetime.datetime.now().strftime("%d/%m/%Y"),
                "REVISION": "0",
                "ELABORO":  dxf_elaboro,
                "REVISO":   dxf_reviso,
                "APROBO":   dxf_aprobo,
                "HOJA":     "1/1"
            }
            celdas = [
                ("EMPRESA",  rotulo_x + 0.0,  rotulo_y + 2.4, 8.5, 1.2),
                ("PROYECTO", rotulo_x + 0.0,  rotulo_y + 1.2, 8.5, 1.2),
                ("N. PLANO", rotulo_x + 8.5,  rotulo_y + 2.4, 3.5, 1.2),
                ("ESCALA",   rotulo_x + 12.0, rotulo_y + 2.4, 2.5, 1.2),
                ("FECHA",    rotulo_x + 14.5, rotulo_y + 2.4, 3.5, 1.2),
                ("REVISION", rotulo_x + 8.5,  rotulo_y + 1.2, 3.5, 1.2),
                ("HOJA",     rotulo_x + 12.0, rotulo_y + 1.2, 6.0, 1.2),
                ("ELABORO",  rotulo_x + 0.0,  rotulo_y + 0.0, 6.0, 1.2),
                ("REVISO",   rotulo_x + 6.0,  rotulo_y + 0.0, 6.0, 1.2),
                ("APROBO",   rotulo_x + 12.0, rotulo_y + 0.0, 6.0, 1.2),
            ]
            for campo, cx2, cy2, cw, ch2 in celdas:
                msp.add_lwpolyline(
                    [(cx2,cy2),(cx2+cw,cy2),(cx2+cw,cy2+ch2),(cx2,cy2+ch2),(cx2,cy2)],
                    dxfattribs={'layer':'ROTULO'})
                # Etiqueta del campo (pequeña, arriba)
                msp.add_text(campo,
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.15,
                                'insert': (cx2 + 0.1, cy2 + ch2 - 0.22),
                                'color': 8})
                # Valor del campo (grande, centrado)
                msp.add_text(campos[campo],
                    dxfattribs={'layer':'TEXTO','style':'ROMANS',
                                'height': 0.35 if campo=="EMPRESA" else 0.25,
                                'insert': (cx2+cw/2, cy2+ch2/2),
                                'align_point': (cx2+cw/2, cy2+ch2/2),
                                'halign':1,'valign':2})

            # \u2500\u2500 Exportar \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
            import tempfile, os as _os
            with tempfile.NamedTemporaryFile(suffix='.dxf', delete=False) as tmp:
                tmp_path = tmp.name
            doc_dxf.saveas(tmp_path)
            with open(tmp_path, 'rb') as f:
                dxf_bytes = f.read()
            _os.unlink(tmp_path)

            st.download_button(
                label=_t(" Descargar DXF (ICONTEC)", " Download DXF (ICONTEC)"),
                data=dxf_bytes,
                file_name=f"Columna_{b:.0f}x{h:.0f}_ICONTEC.dxf",
                mime="application/dxf")

        except Exception as e:
            import traceback
            st.error(f"Error al generar DXF: {e}")
            st.code(traceback.format_exc(), language='python')
            st.info("Aseg\u00farate de tener instalado ezdxf: pip install ezdxf")








# =============================================================================
# TAB 3: CANTIDADES, DESPIECE Y APU
# =============================================================================
with tab3:
    st.subheader(f"Cantidades de Materiales — {'Circular' if es_circular else 'Rectangular'}, L={L_col:.0f} cm")
    col_c1, col_c2, col_c3, col_c4 = st.columns(4)
    col_c1.metric(_t("Concreto", "Concrete"), f"{vol_concreto_m3:.4f} m³")
    col_c2.metric(_t("Acero Total", "Total Steel"), f"{peso_total_acero_kg:.2f} kg")
    col_c3.metric(_t("Acero Longitudinal", "Long. Steel"), f"{peso_acero_long_kg:.2f} kg")
    col_c4.metric(_t("Acero Estribos", "Tie Steel"), f"{peso_total_estribos_kg:.2f} kg")
    st.markdown("---")
    st.subheader(_t("Despiece de Acero", " Bar Bending Schedule"))
    
    if es_circular:
        long_bar = (L_col + 2 * (ld_mm/10) + 2 * (12*rebar_diam/10)) / 100
        peso_long = n_barras * long_bar * (rebar_area * 100) * 7.85e-3
        
        despiece_data = {
            "Marca": ["L1", "E1 (Espiral)"],
            "Cantidad": [n_barras, 1],
            "Diámetro": [_bar_label(rebar_diam), _bar_label(stirrup_diam)],
            "Longitud (m)": [long_bar, long_espiral_total/100],
            "Longitud Total (m)": [n_barras * long_bar, long_espiral_total/100],
            "Peso (kg)": [peso_long, peso_total_estribos_kg]
        }
        long_bar_m = long_bar
    else:
        long_bar_m = (L_col + 2 * (ld_mm/10) + 2 * (12*rebar_diam/10)) / 100
        peso_long_total = n_barras_total * long_bar_m * (rebar_area * 100) * 7.85e-3
        
        despiece_data = {
            "Marca": ["L1", "E1"],
            "Cantidad": [n_barras_total, n_estribos_total],
            "Diámetro": [_bar_label(rebar_diam), _bar_label(stirrup_diam)],
            "Longitud (m)": [long_bar_m, perim_estribo/100],
            "Longitud Total (m)": [n_barras_total * long_bar_m, n_estribos_total * perim_estribo/100],
            "Peso (kg)": [peso_long_total, peso_total_estribos_kg]
        }
    
    df_despiece = pd.DataFrame(despiece_data)
    st.dataframe(df_despiece.style.format({"Longitud (m)": "{:.2f}", "Longitud Total (m)": "{:.2f}", "Peso (kg)": "{:.1f}"}), 
                 use_container_width=True)
    
    fig_bars, ax_bars = plt.subplots(figsize=(6, 4))
    fig_bars.patch.set_facecolor('#1e1e2e')
    for _ax in fig_bars.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax_bars.bar(df_despiece["Marca"], df_despiece["Peso (kg)"], color=['#ff6b35', '#4caf50'])
    ax_bars.set_xlabel(_t("Elemento", "Element"))
    ax_bars.set_ylabel(_t("Peso (kg)", "Weight (kg)"))
    ax_bars.set_title(_t("Distribución de pesos", "Weight distribution"))
    ax_bars.grid(True, alpha=0.3)
    st.pyplot(fig_bars)
    
    with st.expander(_t("Dibujo de Figurado para Taller", " Shop Drawing Details"), expanded=False):
        st.markdown(_t("Formas reales de las barras con ganchos y dimensiones.", "Actual bar shapes with hooks and dimensions."))
        hook_len_cm = 12 * rebar_diam / 10
        if es_circular:
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1 = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            st.pyplot(fig_l1)
            fig_spiral = draw_spiral(D, paso_espiral, stirrup_diam, _bar_label(stirrup_diam))
            st.pyplot(fig_spiral)
        else:
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1 = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            st.pyplot(fig_l1)
            inside_b = b - 2 * recub_cm
            inside_h = h - 2 * recub_cm
            hook_len_est = 12 * stirrup_diam / 10
            fig_e1 = draw_stirrup(inside_b, inside_h, hook_len_est, stirrup_diam, _bar_label(stirrup_diam))
            st.pyplot(fig_e1)
    
    with st.expander(_t("Presupuesto APU", " APU Budget"), expanded=False):
        st.markdown(_t("Ingrese precios unitarios para calcular el costo total.", "Enter unit prices to calculate total cost."))
        with st.form(key="apu_form"):
            if "apu_moneda" not in st.session_state: st.session_state["apu_moneda"] = "COP"
            moneda = st.text_input(_t("Moneda", "Currency"), value=st.session_state["apu_moneda"])
            col_apu1, col_apu2 = st.columns(2)
            with col_apu1:
                if "apu_cemento" not in st.session_state: st.session_state["apu_cemento"] = 28000.0
                if "apu_acero" not in st.session_state: st.session_state["apu_acero"] = 7500.0
                if "apu_arena" not in st.session_state: st.session_state["apu_arena"] = 120000.0
                if "apu_grava" not in st.session_state: st.session_state["apu_grava"] = 130000.0
                
                precio_cemento = st.number_input(_t("Precio por bulto cemento", "Price per cement bag"), value=st.session_state["apu_cemento"], step=1000.0)
                precio_acero = st.number_input(_t("Precio por kg acero", "Price per kg steel"), value=st.session_state["apu_acero"], step=100.0)
                precio_arena = st.number_input(_t("Precio por m³ arena", "Price per m³ sand"), value=st.session_state["apu_arena"], step=5000.0)
                precio_grava = st.number_input(_t("Precio por m³ grava", "Price per m³ gravel"), value=st.session_state["apu_grava"], step=5000.0)
            with col_apu2:
                if "apu_mo" not in st.session_state: st.session_state["apu_mo"] = 70000.0
                if "apu_aui" not in st.session_state: st.session_state["apu_aui"] = 30.0
                
                precio_mo = st.number_input(_t("Costo mano de obra (día)", "Labor cost per day"), value=st.session_state["apu_mo"], step=5000.0)
                pct_aui = st.number_input(_t("% A.I.U.", "% A.I.U."), value=st.session_state["apu_aui"], step=5.0) / 100.0
            st.markdown("---")
            usar_premezclado = st.checkbox(
                _t("Usar concreto premezclado (omite cemento/arena/grava)", "Use ready-mix concrete (skips cement/sand/gravel)"),
                value=st.session_state.get("apu_premix", False), key="apu_premix"
            )
            precio_premix_m3 = 0.0
            if usar_premezclado:
                precio_premix_m3 = st.number_input(
                    _t("Precio concreto premezclado / m³", "Ready-mix concrete price / m³"),
                    value=st.session_state.get("apu_premix_precio", 420000.0),
                    step=10000.0, key="apu_premix_precio"
                )
            submitted = st.form_submit_button(_t("Calcular Presupuesto", "Calculate Budget"))
            if submitted:
                st.session_state.apu_config = {"moneda": moneda, "cemento": precio_cemento, "acero": precio_acero,
                    "arena": precio_arena, "grava": precio_grava, "costo_dia_mo": precio_mo, "pct_aui": pct_aui,
                    "premix": usar_premezclado, "precio_premix_m3": precio_premix_m3}
                st.success(_t("Precios guardados.", "Prices saved."))
                st.rerun()
        if "apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mix = get_mix_for_fc(fc)
            bag_kg = CEMENT_BAGS.get(norma_sel, CEMENT_BAGS["NSR-10 (Colombia)"])[0]["kg"]
            bultos_col = vol_concreto_m3 * mix["cem"] / bag_kg
            _usar_premix = apu.get("premix", False)
            _precio_premix_m3 = apu.get("precio_premix_m3", 0.0)
            if _usar_premix:
                costo_cemento = 0.0
                costo_arena   = 0.0
                costo_grava   = 0.0
                costo_conc_premix = vol_concreto_m3 * _precio_premix_m3
            else:
                costo_cemento = bultos_col * apu["cemento"]
                costo_arena   = (mix["arena"] * vol_concreto_m3 / 1500) * apu["arena"]
                costo_grava   = (mix["grava"] * vol_concreto_m3 / 1600) * apu["grava"]
                costo_conc_premix = 0.0
            costo_acero = peso_total_acero_kg * apu["acero"]
            costo_mo = (peso_total_acero_kg * 0.04 + vol_concreto_m3 * 0.4) * apu["costo_dia_mo"]
            costo_directo = costo_cemento + costo_acero + costo_arena + costo_grava + costo_conc_premix + costo_mo
            aiu = costo_directo * apu["pct_aui"]
            total = costo_directo + aiu
            # ── Métricas cards ──────────────────────────────────────────────
            _c1, _c2, _c3 = st.columns(3)
            _c1.metric(_t(" Total Proyecto", " Total Project"), f"{total:,.0f} {apu['moneda']}")
            _c2.metric(_t(" Costo Directo", "Direct Cost"), f"{costo_directo:,.0f} {apu['moneda']}")
            _c3.metric(_t(" Mano de Obra", "Labor"), f"{costo_mo:,.0f} {apu['moneda']}")

            # ── Gráfica Plotly — Desglose de costos ────────────────────────
            import plotly.graph_objects as _go
            _items_label = (
                [_t("Concreto PM", "Ready-mix"), _t("Acero", "Steel"), _t("M.O.", "Labor"), "A.I.U."]
                if _usar_premix else
                [_t("Cemento", "Cement"), _t("Acero", "Steel"), _t("Arena", "Sand"),
                 _t("Grava", "Gravel"), _t("M.O.", "Labor"), "A.I.U."]
            )
            _items_val = (
                [costo_conc_premix, costo_acero, costo_mo, aiu]
                if _usar_premix else
                [costo_cemento, costo_acero, costo_arena, costo_grava, costo_mo, aiu]
            )
            _colors = ["#3fb950", "#79c0ff", "#ffa657", "#d2a8ff", "#58a6ff", "#f0883e"][:len(_items_label)]
            _fig_apu = _go.Figure(_go.Bar(
                x=_items_label, y=_items_val,
                marker_color=_colors,
                text=[f"{v:,.0f}" for v in _items_val],
                textposition='outside',
                textfont=dict(size=11, color='white')
            ))
            _fig_apu.update_layout(
                title=_t("Desglose de Costos — Columna", "Cost Breakdown — Column"),
                paper_bgcolor='#161b22', plot_bgcolor='#0d1117',
                font=dict(color='#cdd6f4'),
                xaxis=dict(gridcolor='#30363d'),
                yaxis=dict(gridcolor='#30363d', tickformat=',.0f'),
                margin=dict(t=40, b=20, l=20, r=20),
                height=320
            )
            st.plotly_chart(_fig_apu, use_container_width=True)
            if _usar_premix:
                st.info(_t(
                    f" Concreto premezclado: {vol_concreto_m3:.3f} m³ × {_precio_premix_m3:,.0f} {apu['moneda']}/m³ = {costo_conc_premix:,.0f} {apu['moneda']}",
                    f" Ready-mix concrete: {vol_concreto_m3:.3f} m³ × {_precio_premix_m3:,.0f} {apu['moneda']}/m³ = {costo_conc_premix:,.0f} {apu['moneda']}"
                ))

            output_excel = io.BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                if _usar_premix:
                    df_apu = pd.DataFrame({
                        "Item":     ["Concreto Premezclado", "Acero", "Mano de Obra", "A.I.U.", "TOTAL"],
                        "Cantidad": [vol_concreto_m3, peso_total_acero_kg,
                                     peso_total_acero_kg * 0.04 + vol_concreto_m3 * 0.4, "", ""],
                        "Unidad":   ["m³", "kg", "días", f"{apu['pct_aui']*100:.0f}%", ""],
                        "Subtotal": [costo_conc_premix, costo_acero, costo_mo, aiu, total]
                    })
                else:
                    df_apu = pd.DataFrame({
                        "Item":     ["Cemento", "Acero", "Arena", "Grava", "Mano de Obra", "A.I.U.", "TOTAL"],
                        "Cantidad": [bultos_col, peso_total_acero_kg, mix["arena"]*vol_concreto_m3/1500,
                                     mix["grava"]*vol_concreto_m3/1600,
                                     peso_total_acero_kg*0.04 + vol_concreto_m3*0.4, "", ""],
                        "Unidad":   [f"bultos ({bag_kg}kg)", "kg", "m³", "m³", "días",
                                     f"{apu['pct_aui']*100:.0f}%", ""],
                        "Subtotal": [costo_cemento, costo_acero, costo_arena, costo_grava, costo_mo, aiu, total]
                    })
                df_apu.to_excel(writer, index=False, sheet_name='APU')
                workbook = writer.book
                worksheet = writer.sheets['APU']
                money_fmt = workbook.add_format({'num_format': '#,##0.00'})
                worksheet.set_column('D:D', 15, money_fmt)
            output_excel.seek(0)
            st.download_button(_t("Descargar Presupuesto Excel", " Download Budget Excel"), 
                               data=output_excel, file_name=f"APU_Columna_{b:.0f}x{h:.0f}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# =============================================================================
# TAB 4: MEMORIA DE CÁLCULO COMPLETA
# =============================================================================
with tab4:
    st.subheader(_t("Generar Memoria de Cálculo Completa", " Generate Complete Calculation Report"))
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        btn_docx_col = st.button(_t("Generar Memoria DOCX", " Generate DOCX Report"), type="primary")
    with col_d2:
        try:
            if es_circular:
                _ifc_fname = f"Columna_circ_D{D:.0f}.ifc"
                buf_ifc_col = ifc_export.ifc_columna_circular(
                    D, L_col / 100, fc, fy, n_barras, rebar_type, rebar_diam, stirrup_diam,
                    Ast, recub_cm, Pu_input, max(abs(Mux_input), abs(Muy_input)),
                    cap_x.get('phi_Pn_max', 0) if 'cap_x' in locals() and cap_x else 0,
                    norma_sel, "Proyecto NSR-10"
                )
            else:
                _ifc_fname = f"Columna_{b:.0f}x{h:.0f}.ifc"
                buf_ifc_col = ifc_export.ifc_columna(
                    b, h, L_col / 100, fc, fy, n_barras, rebar_type, rebar_diam, stirrup_diam,
                    Ast, recub_cm, Pu_input, max(abs(Mux_input), abs(Muy_input)),
                    cap_x.get('phi_Pn_max', 0) if 'cap_x' in locals() and cap_x else 0,
                    norma_sel, "Proyecto NSR-10"
                )
            
            st.download_button("Exportar IFC (BIM)", data=buf_ifc_col,
                               file_name=_ifc_fname, mime="application/x-step", key="ifc_col")

        except ImportError:
            st.warning("⚠ La librería `ifcopenshell` no está instalada. Ejecuta `pip install ifcopenshell` para habilitar la exportación IFC/BIM.")
        except Exception as e:
            st.error(f"Error generando IFC: {e}")
            st.info("Asegúrate de que `ifc_export.py` y `ifcopenshell` estén disponibles en el entorno de ejecución.")

    if btn_docx_col:
        doc = Document()
        doc.add_heading(f"Memoria de Cálculo — Columna {'Circular' if es_circular else 'Rectangular'}", 0)
        doc.add_paragraph(f"Norma: {norma_sel}")
        doc.add_paragraph(f"Nivel Sísmico: {nivel_sismico}")
        doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Ingeniero: ")
        doc.add_heading("1. CRITERIOS NORMATIVOS", level=1)
        doc.add_paragraph(f"El diseño se realiza según la norma {norma_sel}. Nivel sísmico: {nivel_sismico}. "
                         f"Factores de reducción: φ compresión = {phi_c_max}, φ tensión = {phi_tension}.")
        doc.add_heading("2. PROPIEDADES MECÁNICAS DE MATERIALES", level=1)
        Ec = 4700 * math.sqrt(fc)
        doc.add_paragraph(f"Concreto: f'c = {fc:.1f} MPa, Ec = {Ec:.0f} MPa\nAcero: fy = {fy:.0f} MPa, Es = {Es:.0f} MPa")
        doc.add_heading("3. GEOMETRÍA Y ARMADO", level=1)
        if es_circular:
            doc.add_paragraph(f"Diámetro: D = {D:.0f} cm | Área bruta: Ag = {Ag:.1f} cm²")
        else:
            doc.add_paragraph(f"Base: b = {b:.0f} cm | Altura: h = {h:.0f} cm | Área bruta: Ag = {Ag:.1f} cm²")
        doc.add_paragraph(f"Refuerzo longitudinal: {n_barras} varillas {rebar_type} (Ast = {Ast:.2f} cm², ρ = {cuantia:.2f}%)")
        sec_img = io.BytesIO()
        fig_sec.savefig(sec_img, format='png', dpi=150, bbox_inches='tight')
        sec_img.seek(0)
        doc.add_picture(sec_img, width=Inches(4))
        doc.add_heading("4. DIAGRAMA DE INTERACCIÓN P-M", level=1)
        if cap_x['P_balance'] is not None and cap_x['M_balance'] is not None:
            doc.add_paragraph(f"Punto de balance (Eje X): Pb = {cap_x['P_balance']:.1f} {unidad_fuerza}, Mb = {cap_x['M_balance']:.1f} {unidad_mom}")
        if cap_y['P_balance'] is not None and cap_y['M_balance'] is not None:
            doc.add_paragraph(f"Punto de balance (Eje Y): Pb = {cap_y['P_balance']:.1f} {unidad_fuerza}, Mb = {cap_y['M_balance']:.1f} {unidad_mom}")
        doc.add_picture(pm_2d_img, width=Inches(5))
        
        fig_br, ax_br = plt.subplots(figsize=(5, 5))
        fig_br.patch.set_facecolor('#1e1e2e')
        for _ax in fig_br.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
        theta_vals = np.linspace(0, np.pi/2, 60)
        Mx_curve, My_curve = [], []
        for th in theta_vals:
            Mx_try = np.max(cap_x['phi_M_n']) * np.cos(th)
            My_try = np.max(cap_y['phi_M_n']) * np.sin(th)
            res_br = biaxial_bresler(Pu_input, Mx_try, My_try, cap_x, cap_y, cap_x['Po'], phi_factor)
            if res_br['ratio'] > 0:
                scale = 1.0 / res_br['ratio']
                Mx_curve.append(Mx_try * scale)
                My_curve.append(My_try * scale)
        ax_br.plot(Mx_curve, My_curve, 'b-', linewidth=2, label='Contorno Bresler (Pu fijo)')
        ax_br.plot(Mux_input, Muy_input, 'ro', markersize=10, label=f'Demanda ({Mux_input:.1f}, {Muy_input:.1f})')
        ax_br.set_xlabel(f'Mux [{unidad_mom}]')
        ax_br.set_ylabel(f'Muy [{unidad_mom}]')
        ax_br.set_title(f'Contorno Biaxial Bresler — Pu = {Pu_input:.1f} {unidad_fuerza}')
        ax_br.legend(); ax_br.grid(True, alpha=0.4)
        br2d_img = io.BytesIO()
        fig_br.savefig(br2d_img, format='png', dpi=150, bbox_inches='tight')
        br2d_img.seek(0)
        
        doc.add_heading("5. CONTORNO BIAXIAL 2D Y SUPERFICIE 3D", level=1)
        doc.add_picture(br2d_img, width=Inches(4.5))

        if has_3d_img:
            doc.add_picture(pm_3d_img, width=Inches(5))

        doc.add_heading("6. VERIFICACIÓN DE ESBELTEZ", level=1)
        doc.add_paragraph(f"Factor de longitud efectiva: k = {k_factor}\nRadio de giro: r = {slenderness['r']:.1f} cm\nkl/r = {slenderness['kl_r']:.1f}\nClasificación: {slenderness['classification']}\nFactor de magnificación δns = {slenderness['delta_ns']:.3f}")
        if slenderness['slender']:
            doc.add_paragraph(f"Momentos magnificados por esbeltez:\nMux mag = {Mux_magnified:.1f} {unidad_mom}\nMuy mag = {Muy_magnified:.1f} {unidad_mom}")
        doc.add_heading("7. VERIFICACIÓN BIAXIAL (MÉTODO DE BRESLER)", level=1)
        doc.add_paragraph(f"φPnx (para Mux={Mux_input:.1f}) = {bresler['phi_Pnx']:.2f} {unidad_fuerza}\nφPny (para Muy={Muy_input:.1f}) = {bresler['phi_Pny']:.2f} {unidad_fuerza}\nφP0 = {bresler['phi_P0']:.2f} {unidad_fuerza}\nφPni (Bresler) = {bresler['phi_Pni']:.2f} {unidad_fuerza}\nPu solicitante = {Pu_input:.2f} {unidad_fuerza}\n\n**Resultado: {'CUMPLE' if bresler['ok'] else 'NO CUMPLE'}**")
        doc.add_heading('8. DISE\u00d1O DE ESTRIBOS / ESPIRAL \u2014 CONFINAMIENTO S\u00cdSMICO', level=1)
        if not es_circular:
            # --- Glosario de parámetros ---
            p8a = doc.add_paragraph()
            p8a.add_run('\u00bfQu\u00e9 es cada par\u00e1metro?\n').bold = True
            p8a.add_run(
                f'\u2022 s_conf ({s_conf:.1f} cm): Separaci\u00f3n m\u00e1xima de estribos dentro de la zona de confinamiento Lo. '
                f'Es el espaciamiento centro a centro entre juegos de estribos en la regi\u00f3n m\u00e1s vulnerable.\n'
                f'\u2022 Lo_conf ({Lo_conf:.1f} cm): Longitud de la zona de confinamiento desde cada extremo de la columna. '
                f'En esta zona los estribos deben estar m\u00e1s juntos para garantizar ductilidad.\n'
                f'\u2022 Ash ({Ash_prov:.3f} cm\u00b2 prov. vs {Ash_req:.3f} cm\u00b2 req.): \u00c1rea total de acero transversal '
                f'(estribos + crossties) en un plano perpendicular a la direcci\u00f3n analizada.\n'
                f'\u2022 bc ({bc:.1f} cm) / hc ({hc:.1f} cm): Dimensiones del n\u00facleo confinado, '
                f'medidas de cara exterior a cara exterior de estribo.\n'
                f'\u2022 Ach ({Ach:.1f} cm\u00b2): \u00c1rea del n\u00facleo confinado = bc \u00d7 hc.\n'
            )
            # --- Texto literal NSR-10 ---
            p8b = doc.add_paragraph()
            p8b.add_run('Art\u00edculos aplicables NSR-10:\n').bold = True
            p8b.add_run(
                'NSR-10 C.21.3.5.1 \u2014 Longitud de confinamiento Lo:\n'
                '"Lo no debe ser menor que el mayor de: (a) la dimensi\u00f3n mayor de la secci\u00f3n transversal '
                'del elemento en el punto del empalme, (b) un sexto de la distancia libre entre pisos '
                'del elemento, ni (c) 450 mm."\n\n'
                'NSR-10 C.21.3.5.3 \u2014 Separaci\u00f3n m\u00e1xima en zona confinada:\n'
                '"La separaci\u00f3n s de los estribos rectangulares de confinamiento no debe exceder el menor de: '
                '(a) un cuarto de la dimensi\u00f3n m\u00ednima del elemento, (b) seis veces el di\u00e1metro de las barras '
                'longitudinales, ni (c) so = 100 + (350 \u2212 hx)/3 mm. so no debe exceder 150 mm ni ser menor que 100 mm."\n\n'
                'NSR-10 C.21.3.5.4 \u2014 \u00c1rea m\u00ednima de acero transversal Ash:\n'
                '"El \u00e1rea total de la secci\u00f3n transversal del refuerzo de confinamiento Ash, incluyendo crossties, '
                'no debe ser menor que el mayor de:\n'
                '   Ash \u2265 0.3 \u00b7 s \u00b7 bc \u00b7 (f\'c/fyt) \u00b7 (Ag/Ach \u2212 1)\n'
                '   Ash \u2265 0.09 \u00b7 s \u00b7 bc \u00b7 (f\'c/fyt)"\n\n'
                'NSR-10 C.7.10.5 \u2014 Apoyo lateral de barras (Crossties):\n'
                '"Toda barra longitudinal de esquina debe tener apoyo lateral continuo. '
                'Las barras intermedias deben tener apoyo lateral mediante crossties o ramas de estribo '
                'cuando la separaci\u00f3n libre entre barras apoyadas excede 150 mm."\n'
            )
            # --- Verificación numérica detallada ---
            p8c = doc.add_paragraph()
            p8c.add_run('Verificaci\u00f3n num\u00e9rica:\n').bold = True
            p8c.add_run(
                f'N\u00facleo confinado: bc = b \u2212 2\u00b7rec = {b:.0f} \u2212 2\u00d7{recub_cm:.1f} = {bc:.1f} cm\n'
                f'                  hc = h \u2212 2\u00b7rec = {h:.0f} \u2212 2\u00d7{recub_cm:.1f} = {hc:.1f} cm\n'
                f'Ach = bc \u00d7 hc = {bc:.1f} \u00d7 {hc:.1f} = {Ach:.1f} cm\u00b2\n'
                f'Ag = {Ag:.1f} cm\u00b2,  Ag/Ach = {Ag/Ach:.3f}\n\n'
                f'Lo_conf = max(h={h:.0f} cm, L/6={L_col/6:.1f} cm, 45 cm) = {Lo_conf:.1f} cm\n\n'
                f's_conf = min(8db={8*rebar_diam/10:.1f} cm, 24\u00d8e={24*stirrup_diam/10:.1f} cm, '
                f'min(b,h)/3={min(b,h)/3:.1f} cm, 15 cm) = {s_conf:.1f} cm\n\n'
                f'Ash req. 1 = 0.3 \u00d7 {s_conf:.1f} \u00d7 {bc:.1f} \u00d7 ({fc:.1f}/{fy:.0f}) \u00d7 ({Ag:.1f}/{Ach:.1f} \u2212 1) = {Ash_req_1:.3f} cm\u00b2\n'
                f'Ash req. 2 = 0.09 \u00d7 {s_conf:.1f} \u00d7 {bc:.1f} \u00d7 ({fc:.1f}/{fy:.0f}) = {Ash_req_2:.3f} cm\u00b2\n'
                f'Ash requerido = max({Ash_req_1:.3f}, {Ash_req_2:.3f}) = {Ash_req:.3f} cm\u00b2\n'
                f'Ash provisto  = {Ash_prov:.3f} cm\u00b2  \u2192  {"\u2714 CUMPLE" if ash_ok else "\u2718 NO CUMPLE"}\n\n'
                f'Crossties: {num_flejes_x} flejes en X, {num_flejes_y} flejes en Y\n'
                f'Ramas efectivas: {ramas_x} en X, {ramas_y} en Y\n'
                f'N\u00b0 estribos total: {n_estribos_total} juegos\n'
            )
            # --- Recomendaciones si NO cumple ---
            if not ash_ok:
                p8d = doc.add_paragraph()
                p8d.add_run('\u26a0 ACCI\u00d3N CORRECTIVA REQUERIDA \u2014 Ash insuficiente:\n').bold = True
                ratio_ash_doc = Ash_prov / Ash_req if Ash_req > 0 else 1.0
                s_corr1_doc = Ash_prov * fy / (0.3 * bc * fc * (Ag/Ach - 1)) if (Ag/Ach - 1) > 0 else 999
                s_corr2_doc = Ash_prov * fy / (0.09 * bc * fc) if bc * fc > 0 else 999
                s_correcto_doc = min(s_corr1_doc, s_corr2_doc)
                p8d.add_run(
                    f'D\u00e9ficit: Ash prov = {Ash_prov:.3f} cm\u00b2 < Ash req = {Ash_req:.3f} cm\u00b2 '
                    f'(ratio = {ratio_ash_doc:.2f})\n\n'
                    f'Opciones para corregir (aplicar al menos una):\n'
                    f'  1. REDUCIR SEPARACI\u00d3N: usar s \u2264 {s_correcto_doc:.1f} cm (mantener estribo \u00d8{stirrup_diam:.0f}mm)\n'
                    f'  2. AUMENTAR DI\u00c1METRO del estribo al siguiente comercial\n'
                    f'  3. AGREGAR CROSSTIES: aumentar ramas en el plano cr\u00edtico\n'
                    f'  4. AUMENTAR SECCI\u00d3N de la columna (b o h) para reducir Ag/Ach\n'
                )
        else:
            p8_circ = doc.add_paragraph()
            p8_circ.add_run('Par\u00e1metros de la Espiral:\n').bold = True
            p8_circ.add_run(
                f'\u2022 Espiral \u03c1s = {rho_s_prov:.4f}: Refuerzo helicoidal continuo. Confina el n\u00facleo circular.\n'
                f'\u2022 Paso s ({paso_espiral:.1f} cm): Distancia vertical entre espiras.\n'
                f'\u2022 \u03c1s requerido ({rho_s_req:.4f}): Cuant\u00eda volum\u00e9trica m\u00ednima seg\u00fan NSR-10.\n'
                f'\u2022 dc ({D - 2*recub_cm:.1f} cm): Di\u00e1metro del n\u00facleo confinado.\n'
            )
            p8_circ2 = doc.add_paragraph()
            p8_circ2.add_run('NSR-10 C.21.3.3.1 \u2014 Refuerzo en espiral:\n').bold = True
            p8_circ2.add_run(
                '"La cuant\u00eda volum\u00e9trica del refuerzo en espiral \u03c1s no debe ser menor que:\n'
                '   \u03c1s \u2265 0.45 \u00b7 (Ag/Ach \u2212 1) \u00b7 (f\'c/fyt)\n'
                '   \u03c1s \u2265 0.12 \u00b7 (f\'c/fyt)\n'
                'El paso de la espiral no debe exceder el menor de D/5 o 80 mm."\n\n'
                f'\u03c1s req = max({0.45*(Ag/(math.pi*((D-2*recub_cm)/2)**2)-1)*(fc/fy):.4f}, {0.12*fc/fy:.4f}) = {rho_s_req:.4f}\n'
                f'\u03c1s prov = 4\u00b7As/(dc\u00b7s) = {rho_s_prov:.4f}\n'
                f'\u2192 {"\u2714 CUMPLE" if ash_ok else "\u2718 NO CUMPLE"}\n'
            )
            if not ash_ok:
                dc_val = D - 2*recub_cm
                paso_corr = 4 * stirrup_area / (rho_s_req * dc_val) if rho_s_req * dc_val > 0 else 99
                p8_circ3 = doc.add_paragraph()
                p8_circ3.add_run('\u26a0 ACCI\u00d3N CORRECTIVA \u2014 Espiral insuficiente:\n').bold = True
                p8_circ3.add_run(
                    f'  1. REDUCIR PASO: usar s \u2264 {paso_corr:.1f} cm\n'
                    f'  2. AUMENTAR DI\u00c1METRO de la espiral\n'
                    f'  3. AUMENTAR DI\u00c1METRO de la columna D\n'
                )
        doc.add_heading('Ingeniero Responsable', level=1)
        doc.add_paragraph("Firma: ___________________________")
        doc.add_paragraph("Matrícula Profesional: _______________")
        doc.add_heading("8. LONGITUDES DE DESARROLLO Y EMPALMES", level=1)
        doc.add_paragraph(f"Longitud de desarrollo (ld): {ld_mm/10:.1f} cm\nLongitud de empalme (Clase B): {splice_length_mm/10:.1f} cm\nZona de empalme: {splice_start:.0f} cm - {splice_end:.0f} cm desde la base")
        doc.add_heading("9. SOLICITACIONES Y COMBINACIONES DE DISEÑO", level=1)
        doc.add_paragraph(f"Cargas últimas ingresadas para el diseño y verificación:\nPu = {Pu_input:.1f} {unidad_fuerza}\nMux = {Mux_input:.1f} {unidad_mom}\nMuy = {Muy_input:.1f} {unidad_mom}")
        doc.add_paragraph("Las cargas evaluadas deben corresponder a la envolvente máxima de combinaciones mayoradas según la norma (ej. 1.2D + 1.0E + 0.5L).")
        doc.add_heading("10. FIGURADO DE ACERO", level=1)
        hook_len_cm = 12 * rebar_diam / 10
        if es_circular:
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1_mem = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            l1_img = io.BytesIO()
            fig_l1_mem.savefig(l1_img, format='png', dpi=150, bbox_inches='tight')
            l1_img.seek(0)
            doc.add_picture(l1_img, width=Inches(5))
        else:
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1_mem = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            l1_img = io.BytesIO()
            fig_l1_mem.savefig(l1_img, format='png', dpi=150, bbox_inches='tight')
            l1_img.seek(0)
            doc.add_picture(l1_img, width=Inches(5))
            inside_b = b - 2 * recub_cm
            inside_h = h - 2 * recub_cm
            hook_len_est = 12 * stirrup_diam / 10
            fig_e1_mem = draw_stirrup(inside_b, inside_h, hook_len_est, stirrup_diam, _bar_label(stirrup_diam))
            e1_img = io.BytesIO()
            fig_e1_mem.savefig(e1_img, format='png', dpi=150, bbox_inches='tight')
            e1_img.seek(0)
            doc.add_picture(e1_img, width=Inches(4))
        doc.add_heading("11. CANTIDADES DE OBRA", level=1)
        doc.add_paragraph(f"Concreto: {vol_concreto_m3:.4f} m³\nAcero total: {peso_total_acero_kg:.1f} kg\nRelación acero/concreto: {peso_total_acero_kg/vol_concreto_m3:.1f} kg/m³")
        doc.add_heading('12. CONCLUSIONES Y RECOMENDACIONES', level=1)
        todo_ok = bresler['ok'] and ash_ok and slenderness['kl_r'] <= 22 and (rho_min <= cuantia <= rho_max)
        if todo_ok:
            p12 = doc.add_paragraph()
            p12.add_run(
                f'\u2714 La secci\u00f3n CUMPLE con todas las disposiciones sismorresistentes analizadas '
                f'seg\u00fan {norma_sel}, Nivel {nivel_sismico}.'
            ).bold = True
        else:
            doc.add_paragraph('Se encontraron los siguientes incumplimientos que deben corregirse antes de construir:\n')
            if not (rho_min <= cuantia <= rho_max):
                doc.add_paragraph(
                    f'1. CUANT\u00cdA FUERA DE L\u00cdMITES: \u03c1 = {cuantia:.2f}% (l\u00edmites {rho_min}% \u2013 {rho_max}%)\n'
                    f'   \u2192 Ajustar el n\u00famero o di\u00e1metro de barras longitudinales.'
                )
            if not bresler['ok']:
                deficit_12 = Pu_input - bresler['phi_Pni']
                doc.add_paragraph(
                    f'2. VERIFICACI\u00d3N BIAXIAL NO CUMPLE: Pu/Pni = {bresler["ratio"]:.3f} > 1.0\n'
                    f'   D\u00e9ficit de capacidad: {deficit_12:.1f} {unidad_fuerza}\n'
                    f'   \u2192 Opci\u00f3n A: Aumentar secci\u00f3n (b o h)\n'
                    f'   \u2192 Opci\u00f3n B: Aumentar cuant\u00eda longitudinal hacia {min(rho_max, cuantia*bresler["ratio"]**0.5):.1f}%\n'
                    f'   \u2192 Opci\u00f3n C: Revisar y reducir las cargas de dise\u00f1o'
                )
            if not ash_ok and not es_circular:
                s_c1 = Ash_prov * fy / (0.3 * bc * fc * (Ag/Ach - 1)) if (Ag/Ach - 1) > 0 else 999
                s_c2 = Ash_prov * fy / (0.09 * bc * fc) if bc * fc > 0 else 999
                s_correcto_12 = min(s_c1, s_c2)
                doc.add_paragraph(
                    f'3. ACERO TRANSVERSAL INSUFICIENTE: Ash prov = {Ash_prov:.3f} cm\u00b2 < Ash req = {Ash_req:.3f} cm\u00b2\n'
                    f'   \u2192 Reducir separaci\u00f3n a s \u2264 {s_correcto_12:.1f} cm, O aumentar di\u00e1metro del estribo, '
                    f'O agregar crossties'
                )
            if slenderness['kl_r'] > 22:
                doc.add_paragraph(
                    f'4. COLUMNA ESBELTA: kL/r = {slenderness["kl_r"]:.1f} > 22\n'
                    f'   \u2192 Se aplic\u00f3 magnificaci\u00f3n de momentos \u03b4ns = {slenderness["delta_ns"]:.3f}\n'
                    f'   \u2192 Si kL/r > 100, se requiere an\u00e1lisis no lineal (NSR-10 C.10.10.7)'
                )
        doc.add_heading("FIRMA Y SELLO DEL INGENIERO RESPONSABLE", level=1)
        doc.add_paragraph(" ")
        doc.add_paragraph(" ")
        doc.add_paragraph("_________________________________________")
        doc.add_paragraph("")
        doc.add_paragraph("Matrícula Profesional: _______________")
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.success(_t("Memoria generada exitosamente.", " Report generated successfully."))
        st.download_button(label=_t("Descargar Memoria DOCX", " Download DOCX Report"),
                           data=doc_mem, file_name=f"Memoria_Columna_{b:.0f}x{h:.0f}_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    st.markdown("---")
    st.subheader(_t("Exportar Verificaciones a Excel", " Export Verifications to Excel"))
    if st.button(_t("Exportar a Excel", " Export to Excel")):
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            df_verif = pd.DataFrame({
                "Verificación": ["Cuantía", "Biaxial", "Esbeltez", "Ash", "Lo_conf", "Separación"],
                "Valor": [f"{cuantia:.2f}%", f"{bresler['ratio']:.3f}", f"{slenderness['kl_r']:.1f}", 
                          f"{Ash_prov:.3f}/{Ash_req:.3f}" if not es_circular else f"{rho_s_prov:.4f}/{rho_s_req:.4f}",
                          f"{Lo_conf:.1f} cm requerido" if not es_circular else "N/A", f"{s_conf:.1f}" if not es_circular else f"{paso_espiral:.1f}"],
                "Límite": [f"{rho_min}% - {rho_max}%", "≤ 1.0", "≤ 22", "≥ 1.0", "≥ max(b,h,L/6,45)", "≤ 15/20"],
                "Cumple": ["SÍ" if rho_min <= cuantia <= rho_max else "NO",
                           "SÍ" if bresler['ok'] else "NO",
                           "SÍ" if slenderness['kl_r'] <= 22 else "NO",
                           "SÍ" if ash_ok else "NO",
                           "SÍ (por diseño)" if not es_circular else "SÍ",
                           "SÍ" if s_conf <= 15.0 or (es_dmo and s_conf <= 15.0) else "NO"]
            })
            df_verif.to_excel(writer, sheet_name='Verificaciones', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Verificaciones']
            green_format = workbook.add_format({'bg_color': '#C6EFCE', 'font_color': '#006100'})
            red_format = workbook.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006'})
            for row, val in enumerate(df_verif["Cumple"], start=2):
                cell = f"D{row}"
                if val == "SÍ":
                    worksheet.write(cell, val, green_format)
                else:
                    worksheet.write(cell, val, red_format)
            df_despiece.to_excel(writer, sheet_name='Despiece', index=False)
            df_cant = pd.DataFrame({
                "Material": ["Concreto", "Acero Longitudinal", "Acero Estribos", "Acero Total"],
                "Cantidad": [vol_concreto_m3, peso_acero_long_kg, peso_total_estribos_kg, peso_total_acero_kg],
                "Unidad": ["m³", "kg", "kg", "kg"]
            })
            df_cant.to_excel(writer, sheet_name='Cantidades', index=False)
        excel_buffer.seek(0)
        st.download_button(label=_t("Descargar Excel", " Download Excel"),
                           data=excel_buffer, file_name=f"Verificaciones_Columna_{b:.0f}x{h:.0f}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")