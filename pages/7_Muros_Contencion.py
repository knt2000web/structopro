import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import plotly.graph_objects as go

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Muros de Contención", "Retaining Walls"), layout="wide")

st.image(r"assets/retaining_wall_header_1773256923525.png", use_container_width=True)
st.title(_t("Muros de Contención y Estabilidad", "Retaining Walls and Stability"))
st.markdown(_t("Herramientas para revisar la estabilidad al volcamiento y deslizamiento de muros de contención de gravedad y en voladizo, considerando empujes de tierras y sobrecargas.", "Tools to verify overturning and sliding stability for gravity and cantilever retaining walls, considering earth pressures and surcharges."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CONFIGURACIÓN GENERAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙️ Geometría del Muro", "⚙️ Wall Geometry"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = "NSR-10 (Colombia)"
norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)
H_muro = st.sidebar.number_input("Altura Total H [m]", 1.0, 15.0, st.session_state.get("m_H", 4.0), 0.5, key="m_H")
B_base = st.sidebar.number_input("Ancho de Base B [m]", 0.5, 10.0, st.session_state.get("m_B", 2.5), 0.5, key="m_B")
espesor_base = st.sidebar.number_input("Espesor de la Base (Zapata) [m]", 0.2, 2.0, st.session_state.get("m_ebase", 0.5), 0.1, key="m_ebase")
pie_muro = st.sidebar.number_input("Longitud del Pie (Toe) [m]", 0.0, 5.0, st.session_state.get("m_pie", 0.6), 0.1, key="m_pie")
corona = st.sidebar.number_input("Ancho Corona (Top) [m]", 0.2, 2.0, st.session_state.get("m_corona", 0.3), 0.1, key="m_corona")
base_pantalla = st.sidebar.number_input("Ancho Base Pantalla [m]", 0.2, 3.0, st.session_state.get("m_bpantalla", 0.4), 0.1, key="m_bpantalla")

st.sidebar.header(_t("🌱 Propiedades del Suelo", "🌱 Soil Properties"))
gamma_s = st.sidebar.number_input(_t("Peso Unitario Suelo γ [kN/m³]", "Soil Unit Weight γ [kN/m³]"), 10.0, 22.0, st.session_state.get("m_gamma_s", 18.0), 0.5, key="m_gamma_s")
phi_ang = st.sidebar.number_input(_t("Ángulo de Fricción φ [°]", "Friction Angle φ [°]"), 20.0, 45.0, st.session_state.get("m_phi", 30.0), 1.0, key="m_phi")
c_base = st.sidebar.number_input(_t("Cohesión en la base c [kPa]", "Base cohesion c [kPa]"), 0.0, 100.0, st.session_state.get("m_c", 0.0), 5.0, key="m_c")
delta_ang = st.sidebar.number_input(_t("Fricción suelo-muro δ [°] (Fricción base)", "Soil-wall friction δ [°]"), 10.0, 40.0, st.session_state.get("m_delta", 20.0), 1.0, key="m_delta")
gamma_conc = 24.0 # kN/m3

# Cálculos Geométricos Base
talon = B_base - pie_muro - base_pantalla
if talon < 0:
    st.sidebar.error("Geometría inválida: El talón es negativo.")
    talon = 0

# ─────────────────────────────────────────────
# T1 & T2: ESTABILIDAD CON TERRAPLÉN Y SOBRECARGA
# ─────────────────────────────────────────────
with st.expander(_t("⚖️ Estabilidad al Volcamiento y Deslizamiento (Terraplén y Sobrecarga)", "⚖️ Overturning and Sliding Stability"), expanded=True):
    st.info(_t("📺 **Modo de uso:** Configura la geometría del muro y propiedades del suelo en el menú izquierdo. Aquí ingresa la inclinación del terraplén posterior y la sobrecarga (vehicular/estructural). Retornará los Factores de Seguridad según Rankine/Coulomb.", "📺 **How to use:** Set geometry and soil in the left menu. Enter backfill slope and surcharge here. Will output Safety Factors (Rankine/Coulomb)."))
    
    c1, c2 = st.columns(2)
    with c1:
        beta_ang = st.number_input("Inclinación del Terraplén β [°]", 0.0, phi_ang-0.1, st.session_state.get("m_beta", 0.0), 1.0, key="m_beta")
        q_sobrecarga = st.number_input("Sobrecarga uniforme q [kPa]", 0.0, 100.0, st.session_state.get("m_q", 10.0), 2.0, key="m_q")
    with c2:
        FS_v_min = st.number_input("FS Volcamiento Mínimo", value=st.session_state.get("m_fsv_min", 1.5), step=0.1, key="m_fsv_min")
        FS_d_min = st.number_input("FS Deslizamiento Mínimo", value=st.session_state.get("m_fsd_min", 1.5), step=0.1, key="m_fsd_min")
        
    # Coeficiente Activo Ka (Ecuación general para beta inclinado)
    phi_rad = math.radians(phi_ang)
    beta_rad = math.radians(beta_ang)
    if beta_ang > 0:
        num = math.cos(beta_rad) - math.sqrt(math.cos(beta_rad)**2 - math.cos(phi_rad)**2)
        den = math.cos(beta_rad) + math.sqrt(math.cos(beta_rad)**2 - math.cos(phi_rad)**2)
        Ka = math.cos(beta_rad) * (num / den)
    else:
        Ka = math.tan(math.radians(45) - phi_rad/2)**2

    # Altura virtual H' en la parte trasera del talón
    H_prima = H_muro + talon * math.tan(beta_rad)
    
    # Fuerzas Activas
    Pa_suelo = 0.5 * gamma_s * (H_prima**2) * Ka # kN/m
    Pah_suelo = Pa_suelo * math.cos(beta_rad)
    Pav_suelo = Pa_suelo * math.sin(beta_rad)
    y_pa = H_prima / 3.0
    
    # Fuerzas Sobrecarga (q x Ka x H')
    Pa_q = q_sobrecarga * H_prima * Ka
    Pah_q = Pa_q * math.cos(beta_rad)
    Pav_q = Pa_q * math.sin(beta_rad)
    y_q = H_prima / 2.0
    
    # Fuerzas Resistentes (Pesos)
    items = []
    # 1. Base
    W_base = B_base * espesor_base * gamma_conc
    x_base = B_base / 2.0
    items.append(("Base Concreto", W_base, x_base))
    
    # Pantalla rectangular
    W_pant_rect = corona * (H_muro - espesor_base) * gamma_conc
    x_pant_rect = pie_muro + (base_pantalla - corona) + corona/2.0
    items.append(("Pantalla Rectangular", W_pant_rect, x_pant_rect))
    
    # Pantalla triangular (si es variable)
    ancho_triang = base_pantalla - corona
    if ancho_triang > 0:
        W_pant_tri = 0.5 * ancho_triang * (H_muro - espesor_base) * gamma_conc
        x_pant_tri = pie_muro + (2.0/3.0) * ancho_triang
        items.append(("Pantalla Triangular", W_pant_tri, x_pant_tri))
        
    # Suelo sobre el talón
    W_suelo_talon = talon * (H_muro - espesor_base) * gamma_s
    x_suelo_talon = pie_muro + base_pantalla + talon/2.0
    items.append(("Suelo Talón (Rectángulo)", W_suelo_talon, x_suelo_talon))
    
    # Suelo cuña terraplen inclinado
    if beta_ang > 0:
        W_suelo_tri = 0.5 * talon * (talon*math.tan(beta_rad)) * gamma_s
        x_suelo_tri = pie_muro + base_pantalla + (2.0/3.0)*talon
        items.append(("Suelo Talón (Cuña Inclinada)", W_suelo_tri, x_suelo_tri))
        
    # Momentos y Sumatorias
    W_total = sum([w for name, w, x in items])
    Mr_pesos = sum([w * x for name, w, x in items])
    
    # Aportes verticales de la tierra activa (ayudan a estabilizar)
    Mr_Pav_suelo = Pav_suelo * B_base
    Mr_Pav_q = Pav_q * B_base
    
    Mr_total = Mr_pesos + Mr_Pav_suelo + Mr_Pav_q
    
    # Momentos de Volcamiento (respecto a esquina del pie)
    Mo_suelo = Pah_suelo * y_pa
    Mo_q = Pah_q * y_q
    Mo_total = Mo_suelo + Mo_q
    
    # Fuerzas Deslizantes
    Fd_total = Pah_suelo + Pah_q
    
    # Fuerzas Resistentes al deslizamiento
    N_total = W_total + Pav_suelo + Pav_q
    Fr_desl = N_total * math.tan(math.radians(delta_ang)) + c_base * B_base
    
    FS_volcamiento = Mr_total / Mo_total if Mo_total > 0 else 999
    FS_deslizamiento = Fr_desl / Fd_total if Fd_total > 0 else 999
    
    tab_r, tab_d, tab_a = st.tabs(["📊 Resultados de Estabilidad", "📐 Geometría DXF", "💰 Cantidades APU (1m)"])
    
    with tab_r:
        st.write("#### Factores de Empuje")
        st.write(f"Coeficiente Activo (Ka): **{Ka:.4f}**")
        st.write(f"Altura virtual H' trasera: **{H_prima:.2f} m**")
        
        st.markdown(r"**Verificación de Estabilidad**")
        st.markdown(r"1. **Volcamiento:** $\text{FS}_V = \frac{\sum M_r}{\sum M_o} \ge \text{FS}_{V,min}$")
        col_v, col_d = st.columns(2)
        col_v.metric("FS Volcamiento", f"{FS_volcamiento:.2f}")
        if FS_volcamiento >= FS_v_min:
            col_v.success(f"✅ Aprobado Volcamiento: {FS_volcamiento:.2f} $\\ge$ {FS_v_min:.1f}")
        else:
            col_v.error(f"❌ No Aprobado por Volcamiento: {FS_volcamiento:.2f} $<$ {FS_v_min:.1f} $\\rightarrow$ **Aumentar Talón o Base**")
            
        st.markdown(r"2. **Deslizamiento:** $\text{FS}_D = \frac{\sum F_r}{\sum F_d} \ge \text{FS}_{D,min}$")        
        col_d.metric("FS Deslizamiento", f"{FS_deslizamiento:.2f}")
        if FS_deslizamiento >= FS_d_min:
            col_d.success(f"✅ Aprobado Deslizamiento: {FS_deslizamiento:.2f} $\\ge$ {FS_d_min:.1f}")
        else:
            col_d.error(f"❌ No Aprobado por Deslizamiento: {FS_deslizamiento:.2f} $<$ {FS_d_min:.1f} $\\rightarrow$ **Aumentar Base o Usar Dentellón**")
            
        df_fuerzas = pd.DataFrame([
            {"Concepto": "Pah Suelo (H. Empuje)", "Fuerza [kN/m]": f"{Pah_suelo:.1f}", "Brazo [m]": f"{y_pa:.2f}", "Momento [kN·m/m]": f"{Mo_suelo:.1f}"},
            {"Concepto": "Pah Sobrecarga", "Fuerza [kN/m]": f"{Pah_q:.1f}", "Brazo [m]": f"{y_q:.2f}", "Momento [kN·m/m]": f"{Mo_q:.1f}"},
            {"Concepto": "Peso Total Muro+Suelo", "Fuerza [kN/m]": f"{W_total:.1f}", "Brazo [m]": f"{(Mr_pesos/W_total):.2f}", "Momento [kN·m/m]": f"{Mr_pesos:.1f}"},
        ])
        st.table(df_fuerzas)
        
        # Generar Memoria
        doc_muro = Document()
        doc_muro.add_heading(f"Memoria de Estabilidad — Muro H={H_muro}m", 0)
        doc_muro.add_paragraph(f"Geometría: Base = {B_base}m, Talón = {talon:.2f}m, Pie = {pie_muro:.2f}m")
        doc_muro.add_paragraph(f"Suelo: φ = {phi_ang}°, γ = {gamma_s} kN/m³, Inclinación β = {beta_ang}°")
        doc_muro.add_heading("Resultados", level=1)
        doc_muro.add_paragraph(f"FS Volcamiento: {FS_volcamiento:.2f} (Min: {FS_v_min})")
        doc_muro.add_paragraph(f"FS Deslizamiento: {FS_deslizamiento:.2f} (Min: {FS_d_min})")
        
        f_muro_io = io.BytesIO()
        doc_muro.save(f_muro_io)
        f_muro_io.seek(0)
        st.download_button("Descargar Memoria DOCX", data=f_muro_io, file_name="Estabilidad_Muro.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with tab_d:
        st.subheader("🧊 Visualización 3D del Muro (Tramos de 1m)")
        fig3d = go.Figure()
        
        # Base
        x_b = [-pie_muro, B_base-pie_muro, B_base-pie_muro, -pie_muro, -pie_muro, B_base-pie_muro, B_base-pie_muro, -pie_muro]
        z_b = [-1, -1, 0, 0, -1, -1, 0, 0] # Y normalizado a Z
        y_b = [0, 0, 0, 0, espesor_base, espesor_base, espesor_base, espesor_base] # Z normalizado a Y para el extrude
        fig3d.add_trace(go.Mesh3d(x=x_b, y=y_b, z=z_b, alphahull=0, opacity=0.4, color='gray', name='Zapata Base'))

        # Pantalla Rectangular/Triangular (Simplificada como Mesh)
        H_pantalla = H_muro - espesor_base
        x_p = [0, base_pantalla, base_pantalla-ancho_triang, corona, 0, base_pantalla, base_pantalla-ancho_triang, corona]
        z_p = [-1]*4 + [0]*4
        y_p = [espesor_base]*2 + [H_muro]*2 + [espesor_base]*2 + [H_muro]*2
        fig3d.add_trace(go.Mesh3d(x=x_p, y=y_p, z=z_p, alphahull=0, opacity=0.6, color='darkgray', name='Pantalla'))
        
        fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='X (m)', yaxis_title='H (m)', zaxis_title='Prof (m)'),
                            margin=dict(l=0, r=0, b=0, t=0), height=450, showlegend=False, dragmode='turntable')
        st.plotly_chart(fig3d, use_container_width=True)

        st.markdown("---")
        st.write("#### Geometría de Muro 2D")
        fig_m, ax_m = plt.subplots(figsize=(6, 5))
        ax_m.set_facecolor('#1a1a2e'); fig_m.patch.set_facecolor('#1a1a2e')
        
        # Dibujar base
        ax_m.add_patch(patches.Rectangle((0,0), B_base, espesor_base, linewidth=2, edgecolor='darkgray', facecolor='#4a4a6a'))
        # Dibujar pantalla
        pts_pantalla = [
            (pie_muro, espesor_base),
            (pie_muro + base_pantalla, espesor_base),
            (pie_muro + base_pantalla - corona + corona, H_muro), # top back (asume paramento vertical atras o adelante, dibujare paramento frontal inclinado)
            (pie_muro + base_pantalla - ancho_triang, H_muro)
        ]
        # Si asumimos paramento interior vertical (comun):
        pts_pant_vert = [
            (pie_muro, espesor_base),
            (pie_muro + base_pantalla, espesor_base),
            (pie_muro + base_pantalla, H_muro),
            (pie_muro + base_pantalla - corona, H_muro)
        ]
        pant = plt.Polygon(pts_pant_vert, edgecolor='white', facecolor='#6a6a8a', lw=2)
        ax_m.add_patch(pant)
        
        # Suelo
        ax_m.plot([-1, pie_muro], [espesor_base, espesor_base], color='green', lw=2, linestyle='--')
        ax_m.plot([pie_muro+base_pantalla, B_base], [H_muro, H_prima], color='saddlebrown', lw=3)
        ax_m.plot([B_base, B_base+2], [H_prima, H_prima + 2*math.tan(beta_rad)], color='saddlebrown', lw=3)
        
        ax_m.set_xlim(-1, B_base + 1)
        ax_m.set_ylim(-1, H_prima + 1)
        ax_m.axis('off')
        st.pyplot(fig_m)
        
        st.markdown("#### 💾 Exportar Autocad")
        doc_dxf = ezdxf.new('R2010')
        doc_dxf.units = ezdxf.units.M
        msp = doc_dxf.modelspace()
        # Elevacion Muro
        msp.add_lwpolyline([(0,0), (B_base,0), (B_base,espesor_base), (pie_muro+base_pantalla,espesor_base), 
                            (pie_muro+base_pantalla,H_muro), (pie_muro+base_pantalla-corona,H_muro), 
                            (pie_muro,espesor_base), (0,espesor_base), (0,0)], dxfattribs={'layer': 'CONCRETO', 'color': 7})
        msp.add_lwpolyline([(pie_muro+base_pantalla,H_muro), (B_base,H_prima)], dxfattribs={'layer': 'TERRENO', 'color': 3})
            
        out_stream = io.StringIO()
        doc_dxf.write(out_stream)
        st.download_button(label="Descargar Muro DXF", data=out_stream.getvalue(), file_name=f"Perfil_Muro_H{H_muro}.dxf", mime="application/dxf")

    with tab_a:
        vol_conc_muro = B_base * espesor_base + 0.5 * (corona + base_pantalla) * (H_muro - espesor_base)
        vol_excav_muro = B_base * espesor_base * 1.5 # Aprox asumiendo zanja ancha
        
        cuantia_asumida = 60.0 # kg/m3 para muros gravedad/voladizo ligero
        peso_acero_muro = vol_conc_muro * cuantia_asumida
        
        st.write(f"**Concreto (por metro lineal):** {vol_conc_muro:.2f} m³/m")
        st.write(f"**Acero Refuerzo (estimado):** {peso_acero_muro:.1f} kg/m")
        st.write(f"**Excavación:** {vol_excav_muro:.2f} m³/m")
        
        if "apu_config" in st.session_state:
            st.markdown("---")
            st.markdown("### 💰 Presupuesto Estimado (Promedio de Fuentes Regionales)")
            apu = st.session_state.apu_config
            mon = apu["moneda"]
            c_excav = vol_excav_muro * 20000 
            
            bultos_m = vol_conc_muro * 350 / 50.0 
            vol_arena_m = vol_conc_muro * 0.55
            vol_grava_m = vol_conc_muro * 0.8
            
            c_cem = bultos_m * apu["cemento"]
            c_ace = peso_acero_muro * apu["acero"]
            c_are = vol_arena_m * apu["arena"]
            c_gra = vol_grava_m * apu["grava"]
            total_mat = c_cem + c_ace + c_are + c_gra + c_excav
            
            # MO
            total_dias_mo = (peso_acero_muro * 0.04) + (vol_conc_muro * 0.4) + (vol_excav_muro * 0.3)
            costo_mo = total_dias_mo * apu.get("costo_dia_mo", 69333.33)
            
            # Indirectos
            costo_directo = total_mat + costo_mo
            herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * apu.get("pct_aui", 0.30)
            utilidad = costo_directo * apu.get("pct_util", 0.05)
            iva = utilidad * apu.get("iva", 0.19)
            
            total_proyecto = costo_directo + herramienta + aiu + iva
            
            data_muro_apu = {
                "Item": ["Excavación (m³)", "Cemento (bultos)", "Acero (kg)", "Arena (m³)", "Grava (m³)", 
                         "Mano de Obra (días)", "Herramienta Menor", "A.I.U.", "IVA s/Utilidad", "TOTAL PRESUPUESTO / ML"],
                "Cantidad": [f"{vol_excav_muro:.2f}", f"{bultos_m:.1f}", f"{peso_acero_muro:.1f}", f"{vol_arena_m:.2f}", f"{vol_grava_m:.2f}", 
                             f"{total_dias_mo:.2f}", f"{apu.get('pct_herramienta', 0.05)*100:.1f}% MO", 
                             f"{apu.get('pct_aui', 0.3)*100:.1f}% CD", f"{apu.get('iva', 0.19)*100:.1f}% Util", ""],
                f"Subtotal [{mon}]": [f"{c_excav:,.2f}", f"{c_cem:,.2f}", f"{c_ace:,.2f}", f"{c_are:,.2f}", f"{c_gra:,.2f}", 
                                      f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}", f"{iva:,.2f}", f"**{total_proyecto:,.2f}**"]
            }
            st.dataframe(pd.DataFrame(data_muro_apu), use_container_width=True, hide_index=True)
            
            # Excel APU
            output_excel = io.BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                df_export = pd.DataFrame({
                    "Item": ["Excavación", "Cemento", "Acero", "Arena", "Grava", "Mano de Obra"],
                    "Cantidad": [vol_excav_muro, bultos_m, peso_acero_muro, vol_arena_m, vol_grava_m, total_dias_mo],
                    "Unidad": [20000, apu['cemento'], apu['acero'], apu['arena'], apu['grava'], apu.get('costo_dia_mo', 69333.33)]
                })
                df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
                df_export.to_excel(writer, index=False, sheet_name='APU')
                workbook = writer.book
                worksheet = writer.sheets['APU']
                money_fmt = workbook.add_format({'num_format': '#,##0.00'})
                bold = workbook.add_format({'bold': True})
                worksheet.set_column('A:A', 25)
                worksheet.set_column('B:D', 15, money_fmt)
                
                row = len(df_export) + 1
                worksheet.write(row, 0, "Costo Directo (CD)", bold)
                worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
                row += 1
                worksheet.write(row, 0, "Herramienta Menor", bold)
                worksheet.write_formula(row, 3, f'=D7*{apu.get("pct_herramienta", 0.05)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "A.I.U", bold)
                worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui", 0.30)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "IVA s/ Utilidad", bold)
                worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_util", 0.05)/apu.get("pct_aui", 0.30)}*{apu.get("iva", 0.19)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "TOTAL PRESUPUESTO ML", bold)
                worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
                
            output_excel.seek(0)
            st.download_button(label="📥 Descargar Presupuesto Excel (.xlsx)", data=output_excel, 
                               file_name=f"APU_MuroContencion_H{H_muro:.1f}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("💡 Ve a la página 'APU Mercado' para cargar los costos en vivo de ferreterías.")

