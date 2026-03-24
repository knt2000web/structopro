import streamlit as st
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL y NORMA (al inicio)
# ─────────────────────────────────────────────
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es

# Obtener norma activa de session_state
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

# ─────────────────────────────────────────────
# BASES DE DATOS SÍSMICAS MULTI-NORMA (comprimida)
# ─────────────────────────────────────────────
SEISMIC_DATA = {
    "NSR-10 (Colombia)": {
        "ciudades": {
            "Bogotá D.C.": {"Aa": 0.20, "Av": 0.20},
            "Medellín": {"Aa": 0.25, "Av": 0.25},
            "Cali": {"Aa": 0.30, "Av": 0.30},
            "Barranquilla": {"Aa": 0.20, "Av": 0.20},
            "Bucaramanga": {"Aa": 0.30, "Av": 0.30},
            "Pereira": {"Aa": 0.35, "Av": 0.35},
            "Manizales": {"Aa": 0.35, "Av": 0.35},
        },
        "suelos": {
            "S1 (Roca)": {"Fa": 1.0, "Fv": 1.0, "Tp": 0.4, "TL": 3.0},
            "S2 (Suelo rígido)": {"Fa": 1.0, "Fv": 1.0, "Tp": 0.6, "TL": 3.0},
            "S3 (Suelo intermedio)": {"Fa": 1.2, "Fv": 1.2, "Tp": 0.8, "TL": 3.0},
            "S4 (Suelo blando)": {"Fa": 1.5, "Fv": 1.5, "Tp": 1.0, "TL": 3.0},
            "S5 (Suelo muy blando)": {"Fa": 2.0, "Fv": 2.0, "Tp": 1.2, "TL": 3.0},
        },
        "coef_reduccion": {"R": 6.0},
        "uso": {"I (Grupo I)": 1.0, "II (Grupo II)": 1.0, "III (Grupo III)": 1.1, "IV (Grupo IV)": 1.5},
    },
    "E.060 (Perú) / E.030": {
        "ciudades": {
            "Lima": {"Z": 0.45}, "Arequipa": {"Z": 0.40}, "Trujillo": {"Z": 0.35},
            "Cusco": {"Z": 0.30}, "Iquitos": {"Z": 0.25},
        },
        "suelos": {
            "S0 (Roca)": {"S": 1.0, "Tp": 0.3, "TL": 3.0},
            "S1 (Suelo rígido)": {"S": 1.0, "Tp": 0.4, "TL": 2.5},
            "S2 (Suelo intermedio)": {"S": 1.2, "Tp": 0.6, "TL": 2.0},
            "S3 (Suelo blando)": {"S": 1.4, "Tp": 0.9, "TL": 1.6},
        },
        "coef_reduccion": {"R": 6.0},
        "uso": {"A1": 1.5, "A2": 1.3, "B": 1.0, "C": 0.75},
    },
    "NEC-SE-HM (Ecuador)": {
        "ciudades": {"Quito": {"Z": 0.40}, "Guayaquil": {"Z": 0.45}, "Cuenca": {"Z": 0.35}},
        "suelos": {
            "A (Roca)": {"S": 1.0, "Tp": 0.2, "TL": 3.0},
            "B (Suelo rígido)": {"S": 1.0, "Tp": 0.4, "TL": 2.6},
            "C (Suelo intermedio)": {"S": 1.2, "Tp": 0.6, "TL": 2.0},
            "D (Suelo blando)": {"S": 1.4, "Tp": 0.9, "TL": 1.8},
        },
        "coef_reduccion": {"R": 6.0},
        "uso": {"I": 1.5, "II": 1.0, "III": 0.8},
    },
}
def get_seismic_params(norma, ciudad, perfil, uso):
    # Fallback si no existe la norma
    if norma not in SEISMIC_DATA:
        data = {"ciudades": {"Defecto": {"Z": 0.30}}, "suelos": {"Defecto": {"S": 1.0, "Tp": 0.6, "TL": 2.5}},
                "coef_reduccion": {"R": 6.0}, "uso": {"Categoría B": 1.0}}
    else:
        data = SEISMIC_DATA[norma]
    ciudad_data = data["ciudades"].get(ciudad, {"Aa": 0.20, "Av": 0.20} if "NSR-10" in norma else {"Z": 0.30})
    suelo_data = data["suelos"].get(perfil, list(data["suelos"].values())[0])
    uso_factor = data["uso"].get(uso, 1.0)
    R = data["coef_reduccion"]["R"]
    if "NSR-10" in norma:
        Aa = ciudad_data["Aa"]; Av = ciudad_data["Av"]
        Fa = suelo_data["Fa"]; Fv = suelo_data["Fv"]
        Tp = suelo_data["Tp"]; TL = suelo_data["TL"]
        def espectro(T):
            if T < Tp: return 2.5 * Aa * Fa
            elif T <= TL: return 2.5 * Aa * Fa * (Tp / T)
            else: return 2.5 * Aa * Fa * (Tp * TL / T**2)
        return {"Z": Aa, "U": uso_factor, "S": Fa * Fv, "R": R, "Tp": Tp, "TL": TL, "espectro": espectro}
    elif "E.060" in norma or "E.030" in norma:
        Z = ciudad_data["Z"]; S = suelo_data["S"]; Tp = suelo_data["Tp"]; TL = suelo_data["TL"]
        def espectro(T):
            if T < Tp: return 2.5 * Z * uso_factor * S
            elif T <= TL: return 2.5 * Z * uso_factor * S * (Tp / T)
            else: return 2.5 * Z * uso_factor * S * (Tp * TL / T**2)
        return {"Z": Z, "U": uso_factor, "S": S, "R": R, "Tp": Tp, "TL": TL, "espectro": espectro}
    else:
        Z = ciudad_data.get("Z", 0.30); S = suelo_data.get("S", 1.0)
        Tp = suelo_data.get("Tp", 0.6); TL = suelo_data.get("TL", 2.5)
        def espectro(T):
            if T < Tp: return 2.5 * Z * uso_factor * S
            elif T <= TL: return 2.5 * Z * uso_factor * S * (Tp / T)
            else: return 2.5 * Z * uso_factor * S * (Tp * TL / T**2)
        return {"Z": Z, "U": uso_factor, "S": S, "R": R, "Tp": Tp, "TL": TL, "espectro": espectro}

# ─────────────────────────────────────────────
# FACTORES MULTI-NORMA
# ─────────────────────────────────────────────
_NORM_FACTORS = {
    "NSR-10 (Colombia)":         (28, 21, 10, 12, 0.5525),
    "ACI 318-25 (EE.UU.)":       (28, 21, 10, 12, 0.5525),
    "E.060 (Perú)":              (25, 20, 10, 12, 0.5000),
    "NEC-SE-HM (Ecuador)":       (28, 21, 10, 12, 0.5525),
    "NTC-EM (México)":           (26, 20, 10, 12, 0.5200),
    "COVENIN 1753-2006 (Venezuela)": (28, 21, 10, 12, 0.5525),
    "NB 1225001-2020 (Bolivia)":  (28, 21, 10, 12, 0.5525),
    "CIRSOC 201-2025 (Argentina)": (28, 21, 10, 12, 0.5525),
}
_nf = _NORM_FACTORS.get(norma_sel, _NORM_FACTORS["NSR-10 (Colombia)"])
DIV_LOSA_MAC, DIV_LOSA_NERV, DIV_VIGA_SIMP, DIV_VIGA_CONT, COL_IDX = _nf

# ─────────────────────────────────────────────
# FUNCIONES DE CÁLCULO
# ─────────────────────────────────────────────
def predimensionar_losa(lmax, tipo="maciza"):
    divisor = DIV_LOSA_MAC if tipo == "maciza" else DIV_LOSA_NERV
    h = lmax / divisor
    h = math.ceil(h * 100 / 5) * 5 / 100
    return max(h, 0.10)

def predimensionar_viga(l_luz, posicion="continua"):
    div = DIV_VIGA_CONT if posicion == "continua" else DIV_VIGA_SIMP
    h = math.ceil((l_luz / div) * 100 / 5) * 5 / 100
    b = math.ceil((h / 2.0) * 100 / 5) * 5 / 100
    return max(h, 0.30), max(b, 0.25)

def predimensionar_columna(area_trib, q_tot, n_pisos, fc, uso_factor=1.0):
    Pu_ton = q_tot * area_trib * n_pisos * uso_factor
    Ac_cm2 = (Pu_ton * 1000.0) / (COL_IDX * fc)
    lado = math.ceil(math.sqrt(Ac_cm2) / 5.0) * 5.0
    return Pu_ton, Ac_cm2, max(lado, 30.0)

def predimensionar_zapata(Pu_ton, q_adm, lado_col, M=0):
    q_adm_tonm2 = q_adm * 10
    A_req = Pu_ton / q_adm_tonm2
    if M > 0:
        A_req *= 1.2
    B = math.ceil(math.sqrt(A_req) * 10) / 10.0
    B = max(B, lado_col/100 + 0.3)
    # Peralte mínimo por volado (regla práctica: volado/4)
    volado = (B - lado_col/100) / 2
    h_zap = max(0.25, volado / 4)   # en metros
    h_zap = math.ceil(h_zap * 100 / 5) * 5 / 100
    # Evitar profundidades exageradas para visualización (no afecta cálculo)
    return B, min(h_zap, 1.0)   # limitar a 1 m para visualización

def generate_mesh(Lx, Lz, n_pisos, alturas, lx, ly, col_b, col_h, vig_b, vig_h, nx=3, ny=3):
    """Genera nudos, columnas, vigas y zapatas para un modelo nx+1 × ny+1 columnas."""
    x_coords = np.linspace(0, lx, nx+1)
    z_coords = np.linspace(0, ly, ny+1)
    y_coords = [0.0]
    curr_y = 0.0
    for h in alturas:
        curr_y += h
        y_coords.append(curr_y)
    nudos = []
    nid = 1
    for y in y_coords:
        for z in z_coords:
            for x in x_coords:
                nudos.append({"ID": nid, "X": round(x,2), "Y": round(y,2), "Z": round(z,2)})
                nid += 1
    df_nudos = pd.DataFrame(nudos)
    def find_nid(x, y, z):
        res = df_nudos[(np.isclose(df_nudos['X'], x)) & (np.isclose(df_nudos['Y'], y)) & (np.isclose(df_nudos['Z'], z))]
        return int(res.iloc[0]['ID']) if not res.empty else None
    columnas = []
    cid = 1
    for z in z_coords:
        for x in x_coords:
            for p in range(n_pisos):
                n1 = find_nid(x, y_coords[p], z)
                n2 = find_nid(x, y_coords[p+1], z)
                if n1 and n2:
                    columnas.append({"ID": f"C{cid}", "N1": n1, "N2": n2, "Base_X": x, "Base_Z": z, "Piso": p+1, "b (m)": col_b, "h (m)": col_h})
                    cid += 1
    vigas_x = []; vigas_z = []; vid = 1
    for y in y_coords[1:]:
        for z in z_coords:
            for i in range(nx):
                n1 = find_nid(x_coords[i], y, z); n2 = find_nid(x_coords[i+1], y, z)
                if n1 and n2:
                    vigas_x.append({"ID": f"VX{vid}", "N1": n1, "N2": n2, "Piso": y_coords.index(y), "b (m)": vig_b, "h (m)": vig_h})
                    vid += 1
        for x in x_coords:
            for i in range(ny):
                n1 = find_nid(x, y, z_coords[i]); n2 = find_nid(x, y, z_coords[i+1])
                if n1 and n2:
                    vigas_z.append({"ID": f"VZ{vid}", "N1": n1, "N2": n2, "Piso": y_coords.index(y), "b (m)": vig_b, "h (m)": vig_h})
                    vid += 1
    zapatas = []
    zid = 1
    for z in z_coords:
        for x in x_coords:
            nid = find_nid(x, 0, z)
            if nid:
                es_borde_x = (np.isclose(x, 0) or np.isclose(x, lx))
                es_borde_z = (np.isclose(z, 0) or np.isclose(z, ly))
                if es_borde_x and es_borde_z:
                    tipo = "Esquina"
                elif es_borde_x or es_borde_z:
                    tipo = "Borde"
                else:
                    tipo = "Central"
                zapatas.append({"ID": f"Z{zid}", "Nudo": nid, "Tipo": tipo, "X": x, "Z": z})
                zid += 1
    return df_nudos, pd.DataFrame(columnas), pd.DataFrame(vigas_x), pd.DataFrame(vigas_z), pd.DataFrame(zapatas)

def _box_mesh(xc, yc, zc, bx, by, bz, color, opac, name, show_leg):
    """Genera un cubo Mesh3d centrado en (xc,yc,zc) con semiejes bx, by, bz."""
    return go.Mesh3d(
        x=[xc-bx, xc+bx, xc+bx, xc-bx, xc-bx, xc+bx, xc+bx, xc-bx],
        y=[yc-by, yc-by, yc+by, yc+by, yc-by, yc-by, yc+by, yc+by],
        z=[zc,    zc,    zc,    zc,    zc+bz, zc+bz, zc+bz, zc+bz],
        i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
        color=color, opacity=opac, name=name, showlegend=show_leg, flatshading=True)

def plot_edificio(nudos, cols, vigas_x, vigas_z, zaps, zap_dim_dict=None):
    """Modelo 3D sólido: columnas rectangular Mesh3d, vigas Mesh3d, losas semitransparentes, zapatas."""
    fig = go.Figure()
    col_added = False; beam_added = False; slab_added = False; zap_added = False

    # COLUMNAS — sólidos rectangulares
    for _, c in cols.iterrows():
        n1 = nudos[nudos["ID"]==c["N1"]].iloc[0]
        n2 = nudos[nudos["ID"]==c["N2"]].iloc[0]
        cb = float(c["b (m)"]); ch = float(c["h (m)"])
        xc = float(n1["X"]); yc = float(n1["Z"]); zc = float(n1["Y"])
        h_col = float(n2["Y"]) - float(n1["Y"])
        fig.add_trace(_box_mesh(xc, yc, zc, cb/2, ch/2, h_col, "#1565C0", 0.85, "Columna", not col_added))
        col_added = True

    # VIGAS — trazas sólidas rectangulares
    for _, v in pd.concat([vigas_x, vigas_z]).iterrows():
        n1 = nudos[nudos["ID"]==v["N1"]].iloc[0]; n2 = nudos[nudos["ID"]==v["N2"]].iloc[0]
        vb = float(v["b (m)"]); vh = float(v["h (m)"])
        y_piso = float(n1["Y"])
        if not np.isclose(n1["X"], n2["X"]):  # viga en X
            x0 = min(float(n1["X"]), float(n2["X"])); x1 = max(float(n1["X"]), float(n2["X"]))
            xc = (x0+x1)/2; yc = float(n1["Z"]); zc = y_piso - vh
            fig.add_trace(_box_mesh(xc, yc, zc, (x1-x0)/2, vb/2, vh, "#E65100", 0.85, "Viga X", not beam_added))
        else:  # viga en Y
            z0 = min(float(n1["Z"]), float(n2["Z"])); z1 = max(float(n1["Z"]), float(n2["Z"]))
            xc = float(n1["X"]); yc = (z0+z1)/2; zc = y_piso - vh
            fig.add_trace(_box_mesh(xc, yc, zc, vb/2, (z1-z0)/2, vh, "#E65100", 0.85, "Viga Y", False))
        beam_added = True

    # LOSAS — placas semitransparentes
    y_pisos = sorted(nudos["Y"].unique())
    for y in y_pisos[1:]:
        sub = nudos[nudos["Y"]==y]
        xmin,xmax = float(sub["X"].min()), float(sub["X"].max())
        zmin,zmax = float(sub["Z"].min()), float(sub["Z"].max())
        h_losa = 0.20
        fig.add_trace(go.Mesh3d(
            x=[xmin,xmax,xmax,xmin,xmin,xmax,xmax,xmin],
            y=[zmin,zmin,zmax,zmax,zmin,zmin,zmax,zmax],
            z=[float(y),float(y),float(y),float(y),float(y)+h_losa,float(y)+h_losa,float(y)+h_losa,float(y)+h_losa],
            i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
            color="#B0BEC5", opacity=0.20, name="Losa", showlegend=not slab_added, flatshading=True))
        slab_added = True

    # ZAPATAS — bloques naranja bajo la cota ±0
    if zap_dim_dict is not None:
        for _, zap in zaps.iterrows():
            tipo = zap["Tipo"]; B, H = zap_dim_dict.get(tipo, (1.2, 0.4))
            xc=float(zap["X"]); yc=float(zap["Z"])
            fig.add_trace(_box_mesh(xc, yc, -H, B/2, B/2, H, "#BF360C", 0.90, "Zapata", not zap_added))
            zap_added = True

    # Leyenda compacta
    for color, name in [("#1565C0","Columna"),("#E65100","Viga"),("#B0BEC5","Losa"),("#BF360C","Zapata")]:
        fig.add_trace(go.Scatter3d(x=[None],y=[None],z=[None],mode="markers",
            marker=dict(size=8,color=color), name=name))

    fig.update_layout(
        scene=dict(
            xaxis=dict(title="X (m)", showgrid=True, gridcolor="rgba(255,255,255,0.1)"),
            yaxis=dict(title="Y (m)", showgrid=True, gridcolor="rgba(255,255,255,0.1)"),
            zaxis=dict(title="Z (m)", showgrid=True, gridcolor="rgba(255,255,255,0.1)"),
            bgcolor="#0d1117", aspectmode="data",
            camera=dict(eye=dict(x=1.5, y=-1.5, z=1.3))
        ),
        paper_bgcolor="#161b22", font=dict(color="white"), height=700,
        margin=dict(l=0, r=0, b=0, t=40),
        legend=dict(x=1.02, y=0.5, bgcolor="rgba(0,0,0,0.5)", bordercolor="white", borderwidth=1),
        title=dict(text="Modelo Estructural 3D", font=dict(color="white", size=16))
    )
    return fig

# ─────────────────────────────────────────────
# SIDEBAR – CONFIGURACIÓN (se mantiene igual que en la versión anterior)
# ─────────────────────────────────────────────
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(f'<div style="background:#1e3a1e;border-radius:6px;padding:8px;margin-bottom:10px;"><img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;"><span style="color:#7ec87e;font-weight:600;">{_t("Normativa Activa:","Code:")} {norma_sel}</span></div>', unsafe_allow_html=True)

st.sidebar.header(_t("📊 Unidades de salida","📊 Output units"))
unidades_salida = st.sidebar.radio("Unidades de fuerza/momento:", ["kiloNewtons (kN, kN·m)", "Toneladas fuerza (tonf, tonf·m)"], key="pred_units")
if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom    = "tonf·m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom    = "kN·m"

with st.sidebar.expander(_t("1️⃣ GEOMETRÍA DEL PROYECTO", "1️⃣ PROJECT GEOMETRY"), expanded=True):
    num_stories = st.number_input(_t("Número de Pisos", "Number of Stories"), 1, 40, 5, 1)
    h_story = st.number_input(_t("Altura típica de entrepiso (m)", "Typical story height (m)"), 2.0, 6.0, 3.0, 0.1)
    lx = st.number_input(_t("Luz entre ejes X (m)", "Clear span X (m)"), 2.0, 15.0, 6.0, 0.1)
    ly = st.number_input(_t("Luz entre ejes Y (m)", "Clear span Y (m)"), 2.0, 15.0, 5.0, 0.1)
    nx_spans = int(st.number_input(_t("Vanos en dirección X (ejes-1)", "Bays in X direction"), 1, 8, 3, 1))
    ny_spans = int(st.number_input(_t("Vanos en dirección Y (ejes-1)", "Bays in Y direction"), 1, 8, 3, 1))

with st.sidebar.expander(_t("2️⃣ MATERIALES Y ZONA SÍSMICA", "2️⃣ MATERIALS & SEISMIC"), expanded=True):
    fc_val = st.selectbox(_t("Resistencia f'c (kg/cm²)", "Concrete f'c (kg/cm²)"), [210, 240, 280, 350, 420], index=2)
    fy_val = st.selectbox(_t("Fluencia acero fy (kg/cm²)", "Steel fy (kg/cm²)"), [2800, 4200, 5000], index=1)
    if norma_sel in SEISMIC_DATA:
        ciudades = list(SEISMIC_DATA[norma_sel]["ciudades"].keys())
        ciudad_sel = st.selectbox(_t("Ciudad / Localidad", "City / Location"), ciudades, index=0)
        perfiles = list(SEISMIC_DATA[norma_sel]["suelos"].keys())
        perfil_sel = st.selectbox(_t("Perfil de Suelo", "Soil Profile"), perfiles, index=2)
        usos = list(SEISMIC_DATA[norma_sel]["uso"].keys())
        uso_sel = st.selectbox(_t("Categoría de Uso", "Occupancy Category"), usos, index=2)
    else:
        ciudad_sel = "Defecto"; perfil_sel = "Defecto"; uso_sel = "Categoría B"
    seismic = get_seismic_params(norma_sel, ciudad_sel, perfil_sel, uso_sel)
    Z = seismic["Z"]; U = seismic["U"]; S = seismic["S"]; R = seismic["R"]
    Tp = seismic["Tp"]; TL = seismic["TL"]
    espectro_func = seismic["espectro"]
    st.info(f"**Parámetros sísmicos:** Z={Z:.3f}, U={U:.2f}, S={S:.2f}, R={R:.1f}, Tp={Tp:.2f}s, TL={TL:.2f}s")

with st.sidebar.expander(_t("3️⃣ CARGAS Y USO", "3️⃣ LOADS AND OCCUPANCY"), expanded=True):
    uso_edif_label = st.selectbox(_t("Uso Principal", "Main Occupancy"), 
                                  [_t("Residencial", "Residential"), _t("Oficinas", "Offices"), _t("Comercial", "Commercial"), _t("Almacenamiento", "Storage")], index=0)
    if "Residencial" in uso_edif_label:
        ll_estim = 200; dl_estim = 850
    elif "Oficinas" in uso_edif_label:
        ll_estim = 250; dl_estim = 900
    elif "Comercial" in uso_edif_label:
        ll_estim = 500; dl_estim = 950
    else:
        ll_estim = 600; dl_estim = 1000
    q_estimado = st.number_input(_t("Carga Total q_u (Tonf/m²)", "Total Approx Load q_u (Tonf/m²)"), 0.5, 3.0, (ll_estim + dl_estim)/1000.0, 0.1)

with st.sidebar.expander(_t("4️⃣ APU – PRECIOS", "4️⃣ APU – PRICES"), expanded=False):
    moneda = st.text_input("Moneda (ej. COP, USD)", value=st.session_state.get("apu_moneda_pred", "COP"), key="apu_moneda_pred")
    col1a, col2a = st.columns(2)
    with col1a:
        precio_cemento = st.number_input("Precio por bulto de cemento", value=st.session_state.get("apu_cemento_pred", 28000.0), step=1000.0, key="apu_cemento_pred")
        precio_acero = st.number_input("Precio por kg de acero", value=st.session_state.get("apu_acero_pred", 7500.0), step=100.0, key="apu_acero_pred")
        precio_concreto = st.number_input("Precio m³ concreto", value=st.session_state.get("apu_concreto_pred", 400000.0), step=10000.0, key="apu_concreto_pred")
        precio_excavacion = st.number_input("Precio m³ excavación", value=st.session_state.get("apu_exc_pred", 50000.0), step=5000.0, key="apu_exc_pred")
    with col2a:
        precio_mo = st.number_input("Costo mano de obra (día)", value=st.session_state.get("apu_mo_pred", 70000.0), step=5000.0, key="apu_mo_pred")
        pct_herramienta = st.number_input("% Herramienta menor", value=st.session_state.get("apu_herramienta_pred", 5.0), step=1.0, key="apu_herramienta_pred") / 100.0
        pct_aui = st.number_input("% A.I.U.", value=st.session_state.get("apu_aui_pred", 30.0), step=5.0, key="apu_aui_pred") / 100.0
        pct_util = st.number_input("% Utilidad", value=st.session_state.get("apu_util_pred", 5.0), step=1.0, key="apu_util_pred") / 100.0
        iva = st.number_input("IVA (%)", value=st.session_state.get("apu_iva_pred", 19.0), step=1.0, key="apu_iva_pred") / 100.0
    if st.button("Guardar precios APU"):
        st.session_state.apu_config_pred = {
            "moneda": moneda, "cemento": precio_cemento, "acero": precio_acero,
            "concreto": precio_concreto, "excavacion": precio_excavacion,
            "costo_dia_mo": precio_mo, "pct_herramienta": pct_herramienta,
            "pct_aui": pct_aui, "pct_util": pct_util, "iva": iva
        }
        st.success("Precios guardados")
        st.rerun()

with st.sidebar.expander(_t("5️⃣ AJUSTE MANUAL DE DIMENSIONES", "5️⃣ MANUAL DIMENSION OVERRIDE"), expanded=False):
    st.caption("Modifique las dimensiones calculadas para realizar verificación manual.")
    man_h_mac = st.number_input("Espesor Losa Maciza (cm)", 10, 50, 25, 1, key="man_h_mac") / 100.0
    man_h_ali = st.number_input("Espesor Losa Nervada (cm)", 10, 50, 30, 1, key="man_h_ali") / 100.0
    man_bvx   = st.number_input("Viga X – Ancho b (cm)", 20, 80, 25, 5, key="man_bvx") / 100.0
    man_hvx   = st.number_input("Viga X – Peralte h (cm)", 30, 120, 50, 5, key="man_hvx") / 100.0
    man_bvy   = st.number_input("Viga Y – Ancho b (cm)", 20, 80, 25, 5, key="man_bvy") / 100.0
    man_hvy   = st.number_input("Viga Y – Peralte h (cm)", 30, 120, 45, 5, key="man_hvy") / 100.0
    man_col_c = st.number_input("Columna Central – Lado (cm)", 20, 120, 40, 5, key="man_col_c")
    man_col_b = st.number_input("Columna Borde – Lado (cm)", 20, 100, 35, 5, key="man_col_b")
    man_col_e = st.number_input("Columna Esquina – Lado (cm)", 20, 80, 30, 5, key="man_col_e")
    usar_manual = st.checkbox("Usar estas dimensiones en Modelo 3D y Memoria", value=False, key="usar_manual")

with st.sidebar.expander(_t("6️⃣ FORMA DEL PREDIO (OPCIONAL)", "6️⃣ PLOT SHAPE (OPTIONAL)"), expanded=False):
    st.caption(_t("Dibuja el perímetro real del lote sobre la cuadrícula.", "Draw the real plot perimeter over the grid."))
    dibujar_predio = st.checkbox(_t("Activar límite de lote", "Enable plot boundary"), value=False, key="dib_predio")
    if dibujar_predio:
        coords_str = st.text_area(
            _t("Coordenadas X,Y (ej: 0,0; 10,0; 12,8; 0,10)", "X,Y coords (ex: 0,0; 10,0; 12,8; 0,10)"),
            value=f"0,0; {lx*nx_spans},0; {lx*nx_spans},{ly*ny_spans}; 0,{ly*ny_spans}"
        )
        try:
            ptos_predio = []
            for par in coords_str.split(";"):
                if par.strip():
                    px, py = map(float, par.split(","))
                    ptos_predio.append((px, py))
            if ptos_predio and ptos_predio[0] != ptos_predio[-1]:
                ptos_predio.append(ptos_predio[0])  # Cerrar polígono
        except:
            st.error("Formato inválido. Use X,Y separados por punto y coma (;)")
            ptos_predio = []

st.sidebar.markdown("---")
st.sidebar.markdown("""<div style="text-align: center; color: gray; font-size: 11px;">© 2026 Ing. Msc. César Augusto Giraldo Chaparro</div>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CÁLCULOS ESTRUCTURALES
# ─────────────────────────────────────────────
l_max = max(lx, ly)
h_mac = predimensionar_losa(l_max, "maciza")
h_ali = predimensionar_losa(l_max, "aligerada")
h_vx, b_vx = predimensionar_viga(lx, "continua")
h_vy, b_vy = predimensionar_viga(ly, "continua")
A_central = lx * ly
A_borde = A_central / 2.0
A_esquina = A_central / 4.0

# Columnas
Pu_c, Ac_c, lado_c = predimensionar_columna(A_central, q_estimado, num_stories, fc_val, 1.0)
Pu_b, Ac_b, lado_b = predimensionar_columna(A_borde, q_estimado, num_stories, fc_val, 1.2)
Pu_e, Ac_e, lado_e = predimensionar_columna(A_esquina, q_estimado, num_stories, fc_val, 1.4)

# Análisis sísmico
Peso_estimado_ton = q_estimado * (lx * ly) * num_stories
Ta = 0.1 * num_stories
C = espectro_func(Ta) if Ta > 0 else 2.5
V_basal_ton = (Z * U * S * C / R) * Peso_estimado_ton
V_basal_ton = max(V_basal_ton, 0.1 * Peso_estimado_ton)
alturas = [h_story * (i+1) for i in range(num_stories)]
F_i = [V_basal_ton * (h_i / sum(alturas)) for h_i in alturas]
n_cols = 9
M_base_c = V_basal_ton * (h_story * num_stories) / n_cols
V_basal_kN = V_basal_ton * 9.80665
K_col_estim = 30000
deriva_rel = V_basal_kN / (n_cols * K_col_estim) / h_story
warning_deriva = None
if deriva_rel > 0.01:
    warning_deriva = f"⚠️ Deriva estimada {deriva_rel*100:.2f}% > 1% → considerar vigas más peraltadas"

# Zapatas (con momento sísmico)
q_adm = st.sidebar.number_input("Capacidad admisible del suelo (kg/cm²)", 1.0, 5.0, 2.0, 0.1, key="q_adm")
B_zap_c, h_zap_c = predimensionar_zapata(Pu_c, q_adm, lado_c, M_base_c)
B_zap_b, h_zap_b = predimensionar_zapata(Pu_b, q_adm, lado_b, M_base_c*0.8)
B_zap_e, h_zap_e = predimensionar_zapata(Pu_e, q_adm, lado_e, M_base_c*0.6)

# Diccionario de dimensiones por tipo de zapata
zap_dim_dict = {"Central": (B_zap_c, h_zap_c), "Borde": (B_zap_b, h_zap_b), "Esquina": (B_zap_e, h_zap_e)}

# Guardar valores para ajuste manual
st.session_state["_calc_h_mac"] = int(h_mac * 100)
st.session_state["_calc_h_ali"] = int(h_ali * 100)
st.session_state["_calc_bvx"]   = int(b_vx * 100)
st.session_state["_calc_hvx"]   = int(h_vx * 100)
st.session_state["_calc_bvy"]   = int(b_vy * 100)
st.session_state["_calc_hvy"]   = int(h_vy * 100)
st.session_state["_calc_col_c"] = int(lado_c)
st.session_state["_calc_col_b"] = int(lado_b)
st.session_state["_calc_col_e"] = int(lado_e)

if usar_manual:
    h_mac  = man_h_mac
    h_ali  = man_h_ali
    b_vx   = man_bvx
    h_vx   = man_hvx
    b_vy   = man_bvy
    h_vy   = man_hvy
    lado_c = man_col_c
    lado_b = man_col_b
    lado_e = man_col_e

# Generar malla para visualización
alturas_list = [h_story] * num_stories
nudos, cols_df, vigas_x_df, vigas_z_df, zaps_df = generate_mesh(
    lx, ly, num_stories, alturas_list, lx, ly, lado_c/100, lado_c/100, b_vx, h_vx, nx=nx_spans, ny=ny_spans)

# Pestañas (se mantiene la misma estructura, pero con correcciones)
tab_res, tab_3d, tab_2d, tab_sism, tab_cim, tab_exp = st.tabs([
    "📊 Resultados", "🧊 Modelo 3D", "📐 Planta 2D", "🌍 Sísmico", "🏗️ Cimentación", "📤 Exportar"
])
# Las subtabs de exportar se definen dentro del with tab_exp
tab_dxf = tab_mem = tab_qty = tab_apu = None  # será redefinido dentro

with tab_res:
    st.subheader("Dimensiones Recomendadas")
    if warning_deriva:
        st.warning(warning_deriva)
    st.markdown("### 🟠 Losas")
    c1, c2 = st.columns(2)
    c1.metric("Espesor Losa Maciza", f"{h_mac*100:.0f} cm", "L/28")
    c2.metric("Espesor Losa Nervada", f"{h_ali*100:.0f} cm", "L/21")
    st.markdown("### 🟡 Vigas")
    c3, c4 = st.columns(2)
    c3.metric(f"Viga X (Luz {lx}m)", f"{b_vx*100:.0f} × {h_vx*100:.0f} cm", "b × h")
    c4.metric(f"Viga Y (Luz {ly}m)", f"{b_vy*100:.0f} × {h_vy*100:.0f} cm", "b × h")
    st.markdown("### 🏛️ Columnas")
    df_cols = pd.DataFrame({
        "Tipo": ["Central", "Borde", "Esquina"],
        "Área Trib. (m²)": [f"{A_central:.2f}", f"{A_borde:.2f}", f"{A_esquina:.2f}"],
        f"Pu Estimado ({unidad_fuerza})": [f"{Pu_c * factor_fuerza:.1f}", f"{Pu_b * factor_fuerza:.1f}", f"{Pu_e * factor_fuerza:.1f}"],
        "Sección Sugerida (cm)": [f"{lado_c:.0f} × {lado_c:.0f}", f"{lado_b:.0f} × {lado_b:.0f}", f"{lado_e:.0f} × {lado_e:.0f}"]
    })
    st.dataframe(df_cols, use_container_width=True, hide_index=True)
    
    st.markdown("### 🔗 Acero de Refuerzo Aproximado (ρ ≈ 1.2%)")
    ca1, ca2 = st.columns(2)
    # Cálculo para Columna Central (𝜌 = 1.2%)
    area_acero_c = (lado_c * lado_c) * 0.012  # cm2
    # Asumiendo varillas #5 (Ø15.9mm = 1.99 cm2) o #6 (Ø19.1mm = 2.84 cm2)
    n_var_c5 = math.ceil(area_acero_c / 1.99)
    n_var_c5 = max(n_var_c5 + (n_var_c5 % 2), 4) # par y mínimo 4
    
    # Estribos: Separación típica en zona de confinamiento (s = min(dimension/4, 10cm))
    s_conf = min(lado_c/4, 10.0)
    s_resto = min(lado_c/2, 20.0)
    # Para 1 piso de altura h_story
    l_conf = max(lado_c/100, h_story/6, 0.5)  # longitud confinada (m)
    n_estribos = math.ceil((l_conf*2)/(s_conf/100)) + math.ceil((h_story - 2*l_conf)/(s_resto/100))
    
    ca1.metric("Col Central: Acero Longitudinal", f"{n_var_c5} varillas #5 (Ø 5/8')", f"Área: {area_acero_c:.1f} cm²")
    ca2.metric("Col Central: Estribos por piso", f"{n_estribos} estribos #3 (Ø 3/8')", f"Separación: @{s_conf:.0f}/{s_resto:.0f} cm")
    
    st.caption("Cálculo estimado basado en cuantía volumétrica del 1.2%. El diseño final debe cumplir requisitos de flexocompresión NSR-10 C.10.")

with tab_3d:
    st.subheader("Modelo 3D - Cuadrícula Típica (3×3 Ejes)")
    fig3d = plot_edificio(nudos, cols_df, vigas_x_df, vigas_z_df, zaps_df, zap_dim_dict)
    st.plotly_chart(fig3d, use_container_width=True)

with tab_2d:
    st.subheader("Plano de Distribución y Áreas Tributarias")
    fig2d, ax = plt.subplots(figsize=(10, 8))
    fig2d.patch.set_facecolor('#1a1a2e')
    ax.set_facecolor('#1a1a2e')
    nx, ny = 2, 2
    for i in range(nx+1):
        ax.axvline(i*lx, color='gray', linestyle='-.', lw=1, alpha=0.5)
        ax.text(i*lx, ny*ly+1.0, f"{i+1}", color='white', ha='center', va='center', bbox=dict(boxstyle='circle', fc='#d9a05b', ec='white'))
    for j in range(ny+1):
        ax.axhline(j*ly, color='gray', linestyle='-.', lw=1, alpha=0.5)
        ax.text(-1.0, j*ly, f"{chr(65+j)}", color='white', ha='center', va='center', bbox=dict(boxstyle='circle', fc='#d9a05b', ec='white'))
    ax.add_patch(patches.Rectangle((lx/2, ly/2), lx, ly, fc='#00d4ff', alpha=0.2, ec='none'))
    ax.text(lx, ly, "AT Central", color='#00d4ff', ha='center', fontsize=9)
    ax.add_patch(patches.Rectangle((0, ly/2), lx/2, ly, fc='#ff9500', alpha=0.2, ec='none'))
    ax.add_patch(patches.Rectangle((0, 0), lx/2, ly/2, fc='#ff2a2a', alpha=0.2, ec='none'))
    for i in range(nx+1):
        for j in range(ny+1):
            if i in [0, nx] and j in [0, ny]:
                s = lado_e/100.
            elif i in [0, nx] or j in [0, ny]:
                s = lado_b/100.
            else:
                s = lado_c/100.
            ax.add_patch(patches.Rectangle((i*lx-s/2, j*ly-s/2), s, s, fc='#5a7bbf', ec='white', lw=1.5, zorder=3))
    for i in range(nx+1):
        ax.plot([i*lx, i*lx], [0, ny*ly], color='#d9a05b', lw=b_vy*20, alpha=0.7, zorder=2)
    for j in range(ny+1):
        ax.plot([0, nx*lx], [j*ly, j*ly], color='#d9a05b', lw=b_vx*20, alpha=0.7, zorder=2)
    # Dibujar predio si está activo
    if st.session_state.get("dib_predio", False) and 'ptos_predio' in locals() and ptos_predio:
        predio_x = [p[0] for p in ptos_predio]
        predio_y = [p[1] for p in ptos_predio]
        ax.plot(predio_x, predio_y, color='#00ffcc', linewidth=3, linestyle='-', zorder=5)
        ax.fill(predio_x, predio_y, color='#00ffcc', alpha=0.05, zorder=1)
        # Etiqueta de longitud en cada segmento
        for idx in range(len(ptos_predio)-1):
            x0, y0 = ptos_predio[idx]
            x1, y1 = ptos_predio[idx+1]
            dist = math.hypot(x1-x0, y1-y0)
            ax.text((x0+x1)/2, (y0+y1)/2, f"{dist:.2f}m", color='#00ffcc', fontsize=8, 
                    ha='center', va='center', bbox=dict(boxstyle='round', fc='#1a1a2e', ec='none', alpha=0.7))

    ax.set_aspect('equal')
    ax.axis('off')
    st.pyplot(fig2d)
    buf_2d = io.BytesIO()
    fig2d.savefig(buf_2d, format='png', dpi=150, bbox_inches='tight', facecolor='#1a1a2e')
    buf_2d.seek(0); plt.close(fig2d)

with tab_sism:
    st.subheader("Análisis Sísmico")
    st.markdown(f"**Norma:** {norma_sel}")
    st.markdown(f"**Ciudad:** {ciudad_sel} | **Perfil de suelo:** {perfil_sel} | **Categoría de uso:** {uso_sel}")
    st.markdown(f"**Parámetros:** Z = {Z:.3f} | U = {U:.2f} | S = {S:.2f} | R = {R:.1f} | Tp = {Tp:.2f}s | TL = {TL:.2f}s")
    st.markdown(f"**Periodo fundamental estimado:** Ta = {Ta:.2f}s")
    st.markdown(f"**Coeficiente espectral C = Sa(Ta) = {C:.3f}g**")
    st.markdown(f"**Peso total estimado:** {Peso_estimado_ton:.0f} ton")
    st.markdown(f"**Cortante basal (V):** {V_basal_ton:.2f} ton")
    st.markdown(f"**Deriva aproximada:** {deriva_rel*100:.2f}% (límite 1%)")
    T_vals = np.linspace(0.01, 3.0, 300)
    Sa_vals = [espectro_func(T) for T in T_vals]
    fig_esp, ax_esp = plt.subplots(figsize=(8,5))
    ax_esp.plot(T_vals, Sa_vals, 'r-', linewidth=2)
    ax_esp.fill_between(T_vals, 0, Sa_vals, alpha=0.3, color='red')
    ax_esp.set_xlabel("Periodo T (s)"); ax_esp.set_ylabel("Aceleración espectral Sa (g)")
    ax_esp.set_title(f"Espectro de diseño - {norma_sel}")
    ax_esp.grid(True, alpha=0.3)
    st.pyplot(fig_esp)
    fig_f, ax_f = plt.subplots(figsize=(6,4))
    ax_f.barh(alturas, F_i, color='#5a7bbf')
    ax_f.set_xlabel("Fuerza lateral (ton)"); ax_f.set_ylabel("Altura (m)")
    ax_f.set_title("Distribución de fuerzas sísmicas")
    ax_f.grid(True, axis='x', alpha=0.3)
    st.pyplot(fig_f)

with tab_cim:
    st.subheader("Predimensionamiento de Zapatas Aisladas")
    st.markdown(f"**Capacidad admisible del suelo:** q_adm = {q_adm:.1f} kg/cm²")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Central**")
        st.write(f"Pu = {Pu_c * factor_fuerza:.1f} {unidad_fuerza}")
        st.write(f"Lado columna = {lado_c:.0f} cm")
        st.write(f"Momento base ≈ {M_base_c:.1f} ton-m")
        st.write(f"B = {B_zap_c:.2f} m")
        st.write(f"Peralte h = {h_zap_c*100:.0f} cm")
    with col2:
        st.markdown("**Borde**")
        st.write(f"Pu = {Pu_b * factor_fuerza:.1f} {unidad_fuerza}")
        st.write(f"Lado columna = {lado_b:.0f} cm")
        st.write(f"B = {B_zap_b:.2f} m")
        st.write(f"Peralte h = {h_zap_b*100:.0f} cm")
    with col3:
        st.markdown("**Esquina**")
        st.write(f"Pu = {Pu_e * factor_fuerza:.1f} {unidad_fuerza}")
        st.write(f"Lado columna = {lado_e:.0f} cm")
        st.write(f"B = {B_zap_e:.2f} m")
        st.write(f"Peralte h = {h_zap_e*100:.0f} cm")
    fig_zap, ax_zap = plt.subplots(figsize=(4,4))
    ax_zap.add_patch(patches.Rectangle((-B_zap_c/2, -B_zap_c/2), B_zap_c, B_zap_c, fc='#c8c8c8', ec='black'))
    ax_zap.add_patch(patches.Rectangle((-lado_c/200, -lado_c/200), lado_c/100, lado_c/100, fc='#5a7bbf', ec='white'))
    ax_zap.set_aspect('equal'); ax_zap.axis('off')
    ax_zap.set_title(f"Zapata central {B_zap_c*100:.0f}×{B_zap_c*100:.0f} cm")
    st.pyplot(fig_zap)

with tab_exp:
    _sub_dxf, _sub_mem, _sub_qty, _sub_apu = st.tabs(["📏 DXF", "📄 Memoria", "📦 Cantidades", "💰 APU"])

with _sub_dxf:
    st.subheader("Descargar Plano Estructural DXF")
    doc_dxf = ezdxf.new('R2010'); msp = doc_dxf.modelspace()
    for lay, c in [('EJES',1), ('VIGAS',2), ('COLUMNAS',4), ('COTAS',3), ('ZAPATAS',5), ('TEXTO',7)]:
        if lay not in doc_dxf.layers: doc_dxf.layers.add(lay, color=c)
    nx, ny = 2, 2
    for i in range(nx+1):
        msp.add_line((i*lx, -1), (i*lx, ny*ly+1), dxfattribs={'layer':'EJES'})
    for j in range(ny+1):
        msp.add_line((-1, j*ly), (nx*lx+1, j*ly), dxfattribs={'layer':'EJES'})
    for i in range(nx+1):
        for j in range(ny+1):
            if i in [0, nx] and j in [0, ny]:
                s_col = lado_e/100.; B_zap = B_zap_e
            elif i in [0, nx] or j in [0, ny]:
                s_col = lado_b/100.; B_zap = B_zap_b
            else:
                s_col = lado_c/100.; B_zap = B_zap_c
            cx = i*lx; cy = j*ly
            msp.add_lwpolyline([(cx-B_zap/2, cy-B_zap/2), (cx+B_zap/2, cy-B_zap/2), (cx+B_zap/2, cy+B_zap/2), (cx-B_zap/2, cy+B_zap/2), (cx-B_zap/2, cy-B_zap/2)], dxfattribs={'layer':'ZAPATAS'})
            msp.add_lwpolyline([(cx-s_col/2, cy-s_col/2), (cx+s_col/2, cy-s_col/2), (cx+s_col/2, cy+s_col/2), (cx-s_col/2, cy+s_col/2), (cx-s_col/2, cy-s_col/2)], dxfattribs={'layer':'COLUMNAS'})
    for i in range(nx+1):
        msp.add_line((i*lx-b_vy/2, 0), (i*lx-b_vy/2, ny*ly), dxfattribs={'layer':'VIGAS'})
        msp.add_line((i*lx+b_vy/2, 0), (i*lx+b_vy/2, ny*ly), dxfattribs={'layer':'VIGAS'})
    for j in range(ny+1):
        msp.add_line((0, j*ly-b_vx/2), (nx*lx, j*ly-b_vx/2), dxfattribs={'layer':'VIGAS'})
        msp.add_line((0, j*ly+b_vx/2), (nx*lx, j*ly+b_vx/2), dxfattribs={'layer':'VIGAS'})
    _out_dxf = io.StringIO(); doc_dxf.write(_out_dxf)
    st.download_button("📥 Descargar DXF Planta", data=_out_dxf.getvalue().encode('utf-8'), file_name=f"Predimensionamiento_{lx}x{ly}.dxf", mime="application/dxf")

with _sub_mem:
    st.subheader("Generar Memoria de Cálculo")
    if st.button("🖨️ Descargar Reporte DOCX"):
        doc = Document()
        doc.add_heading(f"PREDIMENSIONAMIENTO ESTRUCTURAL — {norma_sel}", 0)
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_heading("01. Configuración del Proyecto", 1)
        for k, v in [("Pisos", num_stories), ("Altura Entrepiso", f"{h_story} m"), ("Luz X", f"{lx} m"), ("Luz Y", f"{ly} m"),
                     ("Ciudad", ciudad_sel), ("Perfil Suelo", perfil_sel), ("Categoría Uso", uso_sel),
                     ("f'c", f"{fc_val} kg/cm²"), ("fy", f"{fy_val} kg/cm²"), ("Carga Total", f"{q_estimado} Ton/m²")]:
            doc.add_paragraph(f"{k}: {v}", style='List Bullet')
        doc.add_heading("02. Análisis Sísmico", 1)
        doc.add_paragraph(f"Parámetros: Z={Z:.3f}, U={U:.2f}, S={S:.2f}, R={R:.1f}, Tp={Tp:.2f}s, TL={TL:.2f}s")
        doc.add_paragraph(f"Periodo Ta = {Ta:.2f}s, C = {C:.3f}g")
        doc.add_paragraph(f"Peso total = {Peso_estimado_ton:.0f} ton")
        doc.add_paragraph(f"Cortante basal V = {V_basal_ton:.2f} ton")
        doc.add_paragraph(f"Deriva aprox = {deriva_rel*100:.2f}%")
        doc.add_heading("03. Predimensionamiento Vigas y Losas", 1)
        doc.add_paragraph(f"Losa Maciza: {h_mac*100:.0f} cm")
        doc.add_paragraph(f"Losa Nervada: {h_ali*100:.0f} cm")
        doc.add_paragraph(f"Viga X: {b_vx*100:.0f}×{h_vx*100:.0f} cm")
        doc.add_paragraph(f"Viga Y: {b_vy*100:.0f}×{h_vy*100:.0f} cm")
        doc.add_heading("04. Columnas", 1)
        t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
        hdr = t.rows[0].cells; hdr[0].text="Tipo"; hdr[1].text="A.Trib(m²)"; hdr[2].text=f"Pu({unidad_fuerza})"; hdr[3].text="Lado(cm)"
        for t_n, a, pu, lado in [("Central", A_central, Pu_c*factor_fuerza, lado_c),
                                 ("Borde", A_borde, Pu_b*factor_fuerza, lado_b),
                                 ("Esquina", A_esquina, Pu_e*factor_fuerza, lado_e)]:
            r = t.add_row().cells; r[0].text=t_n; r[1].text=f"{a:.2f}"; r[2].text=f"{pu:.1f}"; r[3].text=f"{lado:.0f}"
        doc.add_heading("05. Cimentación", 1)
        doc.add_paragraph(f"Capacidad admisible: {q_adm:.1f} kg/cm²")
        doc.add_paragraph(f"Zapata central: {B_zap_c:.2f}×{B_zap_c:.2f} m, h={h_zap_c*100:.0f} cm")
        doc.add_paragraph(f"Zapata borde: {B_zap_b:.2f}×{B_zap_b:.2f} m, h={h_zap_b*100:.0f} cm")
        doc.add_paragraph(f"Zapata esquina: {B_zap_e:.2f}×{B_zap_e:.2f} m, h={h_zap_e*100:.0f} cm")
        try:
            buf_2d.seek(0); doc.add_heading("Esquema de Planta", 1); doc.add_picture(buf_2d, width=Inches(6.0))
        except: pass
        # Figurado Detallado de Acero (Columna Típica)
        doc.add_heading("06. Esquemas de Figurado y Detalle de Acero", 1)
        
        # Cuantía 1.2%
        area_acero = (lado_c**2) * 0.012
        n_var = max(math.ceil(area_acero / 1.99), 4)
        n_var += (n_var % 2) # par
        doc.add_paragraph(f"Cuantía transversal asumida: ρ = 1.2%")
        doc.add_paragraph(f"Refuerzo longitudinal sugerido: {n_var} varillas #5 (Ø 5/8')")
        
        # Dibujo Varilla Mejorado
        hook_len = max(15.0, 12 * 1.59) # gancho 12db o 15cm min
        straight_len = h_story * 100 - 2*4 # menos recubrimientos
        fig_col, ax_col = plt.subplots(figsize=(7,2.5))
        ax_col.plot([0, straight_len], [0,0], '#1565C0', linewidth=4, solid_capstyle='round')
        ax_col.plot([0,0], [0, hook_len], '#1565C0', linewidth=4, solid_capstyle='round')
        ax_col.plot([straight_len, straight_len], [0, -hook_len], '#1565C0', linewidth=4, solid_capstyle='round')
        # Cotas varilla
        ax_col.annotate(f"L={straight_len:.0f} cm", xy=(straight_len/2, 2.0), ha='center', fontsize=10, fontweight='bold')
        ax_col.annotate(f"Gancho={hook_len:.0f} cm", xy=(-2, hook_len/2), ha='right', va='center', fontsize=9)
        ax_col.annotate(f"Gancho={hook_len:.0f} cm", xy=(straight_len+2, -hook_len/2), ha='left', va='center', fontsize=9)
        ax_col.axis('equal'); ax_col.axis('off')
        ax_col.set_title(f"Varilla Longitudinal Principal (#5)", pad=15)
        buf_col = io.BytesIO(); fig_col.savefig(buf_col, format='png', dpi=200, bbox_inches='tight'); buf_col.seek(0)
        doc.add_picture(buf_col, width=Inches(5.0)); plt.close(fig_col)

        # Dibujo Estribo Mejorado
        inside_b = lado_c - 2*4 # recubrimiento 4cm
        inside_h = lado_c - 2*4
        fig_est, ax_est = plt.subplots(figsize=(4,4))
        # path del estribo con bordes redondeados
        rect = patches.FancyBboxPatch((0,0), inside_b, inside_h, boxstyle="round,pad=1.5,rounding_size=2", fill=False, edgecolor='#E65100', linewidth=3)
        ax_est.add_patch(rect)
        # ganchos 135° reales
        ax_est.plot([0, -5], [inside_h, inside_h+5], '#E65100', linewidth=3)
        ax_est.plot([2, -3], [inside_h+2, inside_h+7], '#E65100', linewidth=3)
        ax_est.annotate("Ganchos 135°\nL=7.5 cm", xy=(-4, inside_h+8), fontsize=8, color='#E65100')
        ax_est.annotate(f"{inside_b:.0f} cm", xy=(inside_b/2, -5), ha='center', fontsize=10, fontweight='bold')
        ax_est.annotate(f"{inside_h:.0f} cm", xy=(inside_b+3, inside_h/2), va='center', fontsize=10, fontweight='bold')
        ax_est.axis('equal'); ax_est.axis('off')
        long_estribo = 2*inside_b + 2*inside_h + 15
        ax_est.set_title(f"Estribo de Confinamiento (#3)\nLong. Total = {long_estribo:.0f} cm", pad=15)
        buf_est = io.BytesIO(); fig_est.savefig(buf_est, format='png', dpi=200, bbox_inches='tight'); buf_est.seek(0)
        doc.add_picture(buf_est, width=Inches(3.5)); plt.close(fig_est)
        doc_mem = io.BytesIO()
        doc.save(doc_mem); doc_mem.seek(0)
        st.download_button("📥 Descargar DOCX", data=doc_mem, file_name="Memoria_Predim.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with _sub_qty:
    st.subheader("Estimación de Cantidades de Materiales")
    n_cols_c = 1; n_cols_b = 4; n_cols_e = 4  # modelo 3x3
    vol_cols = (n_cols_c*(lado_c/100.)**2 + n_cols_b*(lado_b/100.)**2 + n_cols_e*(lado_e/100.)**2) * h_story * num_stories
    length_vigas_x = 2 * lx * 3
    length_vigas_y = 2 * ly * 3
    vol_vigas = (length_vigas_x * b_vx * h_vx + length_vigas_y * b_vy * h_vy) * num_stories
    area_losa = 2*lx * 2*ly * num_stories
    vol_losa = area_losa * h_mac
    # Zapatas (dimensiones individuales)
    vol_zap_central = n_cols_c * (B_zap_c**2 * h_zap_c)
    vol_zap_borde   = n_cols_b * (B_zap_b**2 * h_zap_b)
    vol_zap_esquina = n_cols_e * (B_zap_e**2 * h_zap_e)
    vol_zapatas = vol_zap_central + vol_zap_borde + vol_zap_esquina
    vol_conc_total = vol_cols + vol_vigas + vol_losa + vol_zapatas
    acero_estimado = vol_conc_total * 120  # kg
    # Dosificación (ACI 211)
    fc_mpa = fc_val * 0.0980665
    if fc_mpa <= 21: cem_kg, agua_L, arena_kg, grava_kg = 350,200,780,1020
    elif fc_mpa <= 28: cem_kg, agua_L, arena_kg, grava_kg = 430,190,640,1000
    else: cem_kg, agua_L, arena_kg, grava_kg = 530,185,580,960
    total_cem_kg = cem_kg * vol_conc_total
    bultos_cemento = math.ceil(total_cem_kg / 50)
    # Convertir arena y grava a m³ (densidades)
    arena_m3 = arena_kg * vol_conc_total / 1500   # kg a m³ (densidad arena ~1500 kg/m³)
    grava_m3 = grava_kg * vol_conc_total / 1600   # kg a m³ (densidad grava ~1600 kg/m³)
    agua_m3 = agua_L * vol_conc_total / 1000      # litros a m³ (opcional)
    vol_excavacion = vol_zapatas * 1.1  # m³
    cq1, cq2 = st.columns(2)
    with cq1:
        st.markdown("#### Concreto")
        st.write(f"Columnas: {vol_cols:.2f} m³")
        st.write(f"Vigas: {vol_vigas:.2f} m³")
        st.write(f"Losas: {vol_losa:.2f} m³")
        st.write(f"Zapatas: {vol_zapatas:.2f} m³")
        st.write(f"**Total concreto: {vol_conc_total:.2f} m³**")
        st.write(f"Cemento: {bultos_cemento} bultos (50 kg)")
        st.write(f"Arena: {arena_m3:.2f} m³")
        st.write(f"Grava: {grava_m3:.2f} m³")
        st.write(f"Agua: {agua_m3:.2f} m³ ({agua_L*vol_conc_total:.0f} L)")
    with cq2:
        st.markdown("#### Acero")
        st.write(f"Acero estimado: {acero_estimado:,.0f} kg")
        st.markdown("#### Excavación")
        st.write(f"Volumen excavación: {vol_excavacion:.2f} m³")
    fig_bar, ax_bar = plt.subplots(figsize=(8,4))
    ax_bar.bar(["Columnas", "Vigas", "Losas", "Zapatas"], [vol_cols, vol_vigas, vol_losa, vol_zapatas], color=['#5a7bbf', '#d9a05b', '#c8c8c8', '#a9a9a9'])
    ax_bar.set_ylabel("Volumen (m³)")
    ax_bar.set_title("Distribución de Concreto por Elemento")
    ax_bar.grid(True, alpha=0.3)
    st.pyplot(fig_bar)

with _sub_apu:
    st.subheader("Presupuesto APU")
    if "apu_config_pred" in st.session_state:
        apu = st.session_state.apu_config_pred
        mon = apu["moneda"]
        p_cem = apu["cemento"]; p_ace = apu["acero"]; p_conc = apu["concreto"]
        p_exc = apu["excavacion"]; p_mo = apu["costo_dia_mo"]
        pct_h = apu["pct_herramienta"]; pct_aui = apu["pct_aui"]
        pct_util = apu["pct_util"]; iva = apu["iva"]
    else:
        mon = "$"
        p_cem = 28000; p_ace = 7500; p_conc = 400000; p_exc = 50000; p_mo = 70000
        pct_h = 0.05; pct_aui = 0.30; pct_util = 0.05; iva = 0.19
    costo_concreto = vol_conc_total * p_conc
    costo_acero = acero_estimado * p_ace
    costo_excavacion = vol_excavacion * p_exc
    dias_mo = acero_estimado * 0.04 + vol_conc_total * 0.4 + vol_excavacion * 0.1
    costo_mo = dias_mo * p_mo
    costo_directo = costo_concreto + costo_acero + costo_excavacion + costo_mo
    herramienta = costo_mo * pct_h
    aiu = costo_directo * pct_aui
    utilidad = costo_directo * pct_util
    iva_util = utilidad * iva
    total_proyecto = costo_directo + herramienta + aiu + iva_util
    st.markdown("### 💰 Presupuesto Estimado")
    df_apu = pd.DataFrame({
        "Item": ["Concreto (m³)", "Acero (kg)", "Excavación (m³)", "Mano de Obra (días)", "Herramienta Menor", "A.I.U.", "IVA s/Utilidad", "TOTAL"],
        "Cantidad": [f"{vol_conc_total:.2f}", f"{acero_estimado:,.0f}", f"{vol_excavacion:.2f}", f"{dias_mo:.1f}", f"{pct_h*100:.1f}% MO", f"{pct_aui*100:.1f}% CD", f"{iva*100:.1f}% Util", ""],
        f"Subtotal ({mon})": [f"{costo_concreto:,.0f}", f"{costo_acero:,.0f}", f"{costo_excavacion:,.0f}", f"{costo_mo:,.0f}",
                              f"{herramienta:,.0f}", f"{aiu:,.0f}", f"{iva_util:,.0f}", f"{total_proyecto:,.0f}"]
    })
    st.dataframe(df_apu, use_container_width=True, hide_index=True)
    st.metric(f"💎 Gran Total Proyecto ({mon})", f"{total_proyecto:,.0f}")
    out_xl = io.BytesIO()
    with pd.ExcelWriter(out_xl, engine='xlsxwriter') as wr:
        df_apu.to_excel(wr, sheet_name='Presupuesto', index=False)
        pd.DataFrame({"Elemento": ["Columnas", "Vigas", "Losas", "Zapatas", "Total"],
                      "Volumen (m³)": [vol_cols, vol_vigas, vol_losa, vol_zapatas, vol_conc_total]}).to_excel(wr, sheet_name='Volumenes', index=False)
        pd.DataFrame({"Material": ["Cemento (bultos)", "Arena (m³)", "Grava (m³)", "Agua (m³)", "Acero (kg)", "Excavación (m³)"],
                      "Cantidad": [bultos_cemento, arena_m3, grava_m3, agua_m3, acero_estimado, vol_excavacion]}).to_excel(wr, sheet_name='Cantidades', index=False)
    out_xl.seek(0)
    st.download_button("📥 Descargar Excel", data=out_xl, file_name="Presupuesto_Predim.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")