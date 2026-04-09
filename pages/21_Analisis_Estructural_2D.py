import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import json

# ─────────────────────────────────────────────
# IDIOMA GLOBAL Y NORMA
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

st.set_page_config(page_title=_t("StructMaster 2D", "StructMaster 2D"), layout="wide")
st.title(_t("StructMaster 2D - Diseño Estructural Automatizado", "StructMaster 2D - Automated Structural Design"))
st.markdown(_t(f"Análisis Matricial Interactivo con Recomendaciones Automáticas de Acero y Cimentación según **{norma_sel}**.", 
               f"Interactive Matrix Analysis with Automated Steel and Foundation Recommendations per **{norma_sel}**."))

# ─────────────────────────────────────────────
# CODES (φ factores)
CODES = {
    "NSR-10 (Colombia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NSR-10"},
    "ACI 318-25 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-25"},
    "ACI 318-19 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-19"},
    "ACI 318-14 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-14"},
    "NEC-SE-HM (Ecuador)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NEC-SE-HM"},
    "E.060 (Perú)": {"phi_flex":0.90, "phi_shear":0.85, "phi_comp":0.70, "lambda":1.0, "ref":"E.060"},
    "NTC-EM (México)": {"phi_flex":0.85, "phi_shear":0.80, "phi_comp":0.70, "lambda":1.0, "ref":"NTC-EM"},
    "COVENIN 1753-2006 (Venezuela)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.70, "lambda":1.0, "ref":"COVENIN"},
    "NB 1225001-2020 (Bolivia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NB"},
    "CIRSOC 201-2025 (Argentina)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"CIRSOC"},
}
code = CODES.get(norma_sel, CODES["NSR-10 (Colombia)"])
phi_flex = code["phi_flex"]
phi_shear = code["phi_shear"]
phi_comp = code["phi_comp"]
lam = code["lambda"]

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES
def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_rho_min(fc, fy):
    return max(0.25*math.sqrt(fc)/fy, 1.4/fy)

def get_rho_max_beam(fc, fy, beta1):
    eps_cu = 0.003
    eps_t_min = 0.005
    return (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+eps_t_min))

def design_footing(Pu, Mu, qa, fc, fy, recub=5, B_init=1.0, tolerancia=0.05, max_iter=20):
    """Diseña zapata cuadrada con momento (flexión uniaxial). Retorna (B, H, As_req)"""
    # Para simplificar, asumimos zapata cuadrada con momento en una dirección.
    # Cálculo iterativo de B para que q_max ≤ qa
    B = B_init
    for _ in range(max_iter):
        A = B * B
        q_max = Pu/A + 6*Mu/(B*B*B)  # fórmula para zapata cuadrada
        if q_max <= qa * 1.01:
            break
        B *= 1.05
    # Espesor mínimo por corte (cortante unidireccional aproximado)
    # Suponemos columna 40x40 cm para calcular d, pero podría tomarse de la estructura.
    # Para simplificar, usamos regla empírica: H = max(0.4, B/4)
    H = max(0.4, B/4)
    d = H - recub/100 - 0.015  # m, asumiendo varilla 15 mm
    # Acero requerido (mínimo)
    As_min = 0.0018 * B * 100 * H * 100  # cm²
    # Momento de diseño (considerando presión neta)
    # Para zapata cuadrada, momento en la cara de la columna
    # Suponemos columna 40x40 cm, volado = (B - 0.4)/2
    L_v = (B - 0.4)/2
    if L_v > 0:
        # Presión de diseño para acero (usamos qu_promedio en volado)
        qu = Pu/A
        Mu_zap = qu * B * (L_v**2)/2
    else:
        Mu_zap = 0
    # Acero por flexión
    d_cm = d*100
    b_cm = B*100
    Rn = (Mu_zap * 1e6) / (phi_flex * b_cm * d_cm**2)
    if Rn > 0:
        disc = 1 - 2*Rn/(0.85*fc)
        if disc > 0:
            rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
        else:
            rho = 0.0018
    else:
        rho = 0.0018
    As_req = rho * b_cm * d_cm
    As_req = max(As_req, As_min)
    return B, H, As_req

def design_beam(Mu, Vu, b_cm, h_cm, fc, fy, recub=5):
    """Diseño de viga rectangular: acero longitudinal y estribos."""
    d_cm = h_cm - recub - 1  # asumiendo estribo 10 mm
    # Flexión
    Rn = (Mu * 1e6) / (phi_flex * b_cm * d_cm**2)
    disc = 1 - 2*Rn/(0.85*fc)
    if disc > 0:
        rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
    else:
        rho = 0.0018
    rho_min = get_rho_min(fc, fy)
    rho_max = get_rho_max_beam(fc, fy, get_beta1(fc))
    rho_use = max(rho, rho_min)
    As = rho_use * b_cm * d_cm
    # Cortante
    Vc = 0.17 * lam * math.sqrt(fc) * b_cm * d_cm / 10  # kN
    phi_Vc = phi_shear * Vc
    if Vu > phi_Vc/2:
        Vs = Vu / phi_shear - Vc
        # Asumimos estribos #3 (0.71 cm²) de 2 ramas
        Av = 2 * 0.71  # cm²
        s = Av * fy * d_cm / Vs  # cm
        s = max(5, min(s, 60, d_cm/2))
    else:
        s = 0  # no requiere estribos
    return As, s

def design_column(Pu, Mu, b_cm, h_cm, fc, fy, recub=5):
    """Diseño preliminar de columna rectangular (verificación P-M simplificada)."""
    # Acero mínimo por cuantía 1%
    Ag = b_cm * h_cm
    As_min = 0.01 * Ag
    # Verificación aproximada: capacidad de la columna con acero mínimo
    # Asumimos acero distribuido uniformemente. Simplificamos con fórmula de carga axial máxima.
    phi = phi_comp
    # Capacidad de la columna (sin considerar momento)
    Pn_max = 0.85 * fc * (Ag - As_min) + fy * As_min
    phi_Pn_max = phi * Pn_max
    # Para momento, se requiere diagrama P-M, pero lo dejamos como advertencia
    if Pu > phi_Pn_max:
        ok = False
    else:
        ok = True
    return As_min, ok

# ─────────────────────────────────────────────
# INICIALIZACIÓN DE DATOS (estado)
if "nudos_df" not in st.session_state:
    st.session_state.nudos_df = pd.DataFrame([{"ID":1,"X (m)":0.0,"Y (m)":0.0}, {"ID":2,"X (m)":0.0,"Y (m)":3.0}, {"ID":3,"X (m)":4.0,"Y (m)":3.0}, {"ID":4,"X (m)":4.0,"Y (m)":0.0}])
if "barras_df" not in st.session_state:
    st.session_state.barras_df = pd.DataFrame([
        {"Nudo I":1, "Nudo J":2, "E (MPa)":21000.0, "A (cm²)":1200.0, "I (cm⁴)":160000.0, "b (cm)":30.0, "h (cm)":40.0, "w_uniforme (kN/m)":0.0, "w_triangular (kN/m)":0.0, "direccion_w": "Y"},
        {"Nudo I":2, "Nudo J":3, "E (MPa)":21000.0, "A (cm²)":1200.0, "I (cm⁴)":250000.0, "b (cm)":30.0, "h (cm)":50.0, "w_uniforme (kN/m)":0.0, "w_triangular (kN/m)":0.0, "direccion_w": "Y"},
        {"Nudo I":3, "Nudo J":4, "E (MPa)":21000.0, "A (cm²)":1200.0, "I (cm⁴)":160000.0, "b (cm)":30.0, "h (cm)":40.0, "w_uniforme (kN/m)":0.0, "w_triangular (kN/m)":0.0, "direccion_w": "Y"}
    ])
if "apoyos_df" not in st.session_state:
    st.session_state.apoyos_df = pd.DataFrame([{"Nudo":1,"Fijo X":True,"Fijo Y":True,"Fijo Rz":True}, {"Nudo":4,"Fijo X":True,"Fijo Y":True,"Fijo Rz":True}])
if "cargas_nudo_df" not in st.session_state:
    st.session_state.cargas_nudo_df = pd.DataFrame([{"Nudo":2,"FX (kN)":20.0,"FY (kN)":-50.0,"Mz (kN-m)":0.0}, {"Nudo":3,"FX (kN)":0.0,"FY (kN)":-50.0,"Mz (kN-m)":0.0}])
if "resultados" not in st.session_state:
    st.session_state.resultados = None

# ─────────────────────────────────────────────
# TABS
tab_geo, tab_cargas, tab_res, tab_dis = st.tabs([
    _t("1. Geometría", "1. Geometry"),
    _t("2. Cargas y Suelo", "2. Loads & Soil"),
    _t("3. Resultados", "3. Results"),
    _t("4. Exportaciones", "4. Exports")
])

# ─────────────────────────────────────────────
# TAB 1: GEOMETRÍA
with tab_geo:
    st.subheader(_t("Nudos", "Nodes"))
    st.session_state.nudos_df = st.data_editor(st.session_state.nudos_df, num_rows="dynamic", use_container_width=True)
    st.subheader(_t("Elementos (Barras)", "Elements (Beams/Columns)"))
    st.session_state.barras_df = st.data_editor(st.session_state.barras_df, num_rows="dynamic", use_container_width=True)
    st.subheader(_t("Apoyos", "Supports"))
    st.session_state.apoyos_df = st.data_editor(st.session_state.apoyos_df, num_rows="dynamic", use_container_width=True)

    # Botón para guardar/cargar modelo
    col_save, col_load = st.columns(2)
    with col_save:
        model_data = {
            "nudos": st.session_state.nudos_df.to_dict(orient="records"),
            "barras": st.session_state.barras_df.to_dict(orient="records"),
            "apoyos": st.session_state.apoyos_df.to_dict(orient="records"),
            "cargas_nudo": st.session_state.cargas_nudo_df.to_dict(orient="records"),
            "norma": norma_sel
        }
        st.download_button(_t("Guardar Modelo", " Save Model"), data=json.dumps(model_data, indent=2), file_name="modelo.json", mime="application/json")
    with col_load:
        uploaded = st.file_uploader(_t("Cargar Modelo", " Load Model"), type=["json"])
        if uploaded:
            model = json.load(uploaded)
            st.session_state.nudos_df = pd.DataFrame(model["nudos"])
            st.session_state.barras_df = pd.DataFrame(model["barras"])
            st.session_state.apoyos_df = pd.DataFrame(model["apoyos"])
            st.session_state.cargas_nudo_df = pd.DataFrame(model["cargas_nudo"])
            st.rerun()

# ─────────────────────────────────────────────
# TAB 2: CARGAS Y SUELO
with tab_cargas:
    st.subheader(_t("Cargas nodales", "Nodal loads"))
    st.session_state.cargas_nudo_df = st.data_editor(st.session_state.cargas_nudo_df, num_rows="dynamic", use_container_width=True)

    st.subheader(_t("Cargas distribuidas en barras", "Distributed loads on bars"))
    # Las cargas distribuidas ya están en barras_df, se editan en la misma tabla
    st.dataframe(st.session_state.barras_df[["Nudo I","Nudo J","w_uniforme (kN/m)","w_triangular (kN/m)","direccion_w"]], use_container_width=True)
    st.info(_t("Edite las cargas distribuidas directamente en la tabla de barras (pestaña Geometría).", "Edit distributed loads directly in the bar table (Geometry tab)."))

    st.subheader(_t("Propiedades de materiales y suelo", "Material and soil properties"))
    col_mat1, col_mat2 = st.columns(2)
    with col_mat1:
        fc = st.number_input(_t("Resistencia concreto f'c [MPa]", "Concrete strength f'c [MPa]"), 15.0, 80.0, st.session_state.get("ana_fc", 21.0), 1.0, key="ana_fc")
        fy = st.number_input(_t("Fluencia acero fy [MPa]", "Steel yield fy [MPa]"), 200.0, 600.0, st.session_state.get("ana_fy", 420.0), 10.0, key="ana_fy")
        recub = st.number_input(_t("Recubrimiento [cm]", "Cover [cm]"), 2.0, 10.0, st.session_state.get("ana_rec", 5.0), 0.5, key="ana_rec")
    with col_mat2:
        qa_suelo = st.number_input(_t("Capacidad portante suelo qa [kN/m²]", "Soil bearing capacity qa [kN/m²]"), 10.0, 1000.0, st.session_state.get("ana_qa", 150.0), 10.0, key="ana_qa")

    if st.button(_t("Analizar estructura", " Analyze structure")):
        with st.spinner(_t("Analizando...", "Analyzing...")):
            # Copiamos datos actuales
            nodes = st.session_state.nudos_df.copy().dropna()
            elements = st.session_state.barras_df.copy().dropna()
            supports = st.session_state.apoyos_df.copy().dropna()
            loads = st.session_state.cargas_nudo_df.copy().dropna()

            # Validaciones
            if nodes.empty:
                st.error(_t("No hay nudos definidos.", "No nodes defined."))
                st.stop()
            if elements.empty:
                st.error(_t("No hay elementos definidos.", "No elements defined."))
                st.stop()
            # Mapa ID a índice
            node_ids = nodes['ID'].tolist()
            id_to_idx = {nid: i for i, nid in enumerate(node_ids)}
            # Verificar IDs de barras
            for _, e in elements.iterrows():
                if e['Nudo I'] not in id_to_idx or e['Nudo J'] not in id_to_idx:
                    st.error(_t(f"Barra {e.name}: nudo I o J no existe.", f"Bar {e.name}: node I or J does not exist."))
                    st.stop()
            # Verificar apoyos
            for _, s in supports.iterrows():
                if s['Nudo'] not in id_to_idx:
                    st.error(_t(f"Apoyo en nudo {s['Nudo']} no existe.", f"Support at node {s['Nudo']} does not exist."))
                    st.stop()

            num_nudos = len(nodes)
            num_gdl = num_nudos * 3
            # Inicializar matriz de rigidez y vector de fuerzas
            K = np.zeros((num_gdl, num_gdl))
            F = np.zeros(num_gdl)

            # Cargas nodales
            for _, c in loads.iterrows():
                if c['Nudo'] in id_to_idx:
                    n_idx = id_to_idx[c['Nudo']]
                    F[n_idx*3 + 0] += c.get('FX (kN)', 0.0)
                    F[n_idx*3 + 1] += c.get('FY (kN)', 0.0)
                    F[n_idx*3 + 2] += c.get('Mz (kN-m)', 0.0)

            # Ensamblaje de barras y generación de fuerzas de empotramiento por cargas distribuidas
            element_results = []
            for idx, e in elements.iterrows():
                # Coordenadas
                xi = nodes.loc[nodes['ID']==e['Nudo I'], 'X (m)'].values[0]
                yi = nodes.loc[nodes['ID']==e['Nudo I'], 'Y (m)'].values[0]
                xj = nodes.loc[nodes['ID']==e['Nudo J'], 'X (m)'].values[0]
                yj = nodes.loc[nodes['ID']==e['Nudo J'], 'Y (m)'].values[0]
                L = math.hypot(xj-xi, yj-yi)
                cx = (xj-xi)/L
                cy = (yj-yi)/L
                theta = math.atan2(yj-yi, xj-xi)

                # Propiedades
                E = e['E (MPa)'] * 1e3  # kPa
                A = e['A (cm²)'] / 1e4   # m²
                I = e['I (cm⁴)'] / 1e8   # m⁴
                # Matriz de rigidez local
                v1 = E*A/L
                v2 = 12*E*I/(L**3)
                v3 = 6*E*I/(L**2)
                v4 = 4*E*I/L
                v5 = 2*E*I/L
                k_local = np.zeros((6,6))
                k_local[0,0]=v1; k_local[0,3]=-v1; k_local[3,3]=v1
                k_local[1,1]=v2; k_local[1,2]=v3; k_local[1,4]=-v2; k_local[1,5]=v3
                k_local[2,1]=v3; k_local[2,2]=v4; k_local[2,4]=-v3; k_local[2,5]=v5
                k_local[4,1]=-v2; k_local[4,2]=-v3; k_local[4,4]=v2; k_local[4,5]=-v3
                k_local[5,1]=v3; k_local[5,2]=v5; k_local[5,4]=-v3; k_local[5,5]=v4

                # Matriz de rotación
                T = np.zeros((6,6))
                T[0,0]=cx; T[0,1]=cy
                T[1,0]=-cy; T[1,1]=cx
                T[2,2]=1
                T[3,3]=cx; T[3,4]=cy
                T[4,3]=-cy; T[4,4]=cx
                T[5,5]=1

                k_glob = T.T @ k_local @ T

                # Fuerzas de empotramiento por cargas distribuidas (solo en dirección local Y)
                w_uniform = e.get('w_uniforme (kN/m)', 0.0)
                w_tri = e.get('w_triangular (kN/m)', 0.0)
                direc = e.get('direccion_w', 'Y')
                f_fixed_local = np.zeros(6)
                if w_uniform != 0 and direc == 'Y':
                    # Carga uniforme en dirección Y local (transversal)
                    f_fixed_local[1] = w_uniform * L / 2.0
                    f_fixed_local[2] = w_uniform * L**2 / 12.0
                    f_fixed_local[4] = w_uniform * L / 2.0
                    f_fixed_local[5] = -w_uniform * L**2 / 12.0
                if w_tri != 0 and direc == 'Y':
                    # Triangular (máxima en nodo I)
                    f_fixed_local[1] += w_tri * L / 3.0
                    f_fixed_local[2] += w_tri * L**2 / 20.0
                    f_fixed_local[4] += w_tri * L / 6.0
                    f_fixed_local[5] += -w_tri * L**2 / 30.0
                # Rotar fuerzas a global
                f_fixed_glob = T.T @ f_fixed_local
                # Obtener grados de libertad
                i_idx = id_to_idx[e['Nudo I']]
                j_idx = id_to_idx[e['Nudo J']]
                dofs = [i_idx*3, i_idx*3+1, i_idx*3+2, j_idx*3, j_idx*3+1, j_idx*3+2]
                # Sumar a la matriz de rigidez
                for r in range(6):
                    for c in range(6):
                        K[dofs[r], dofs[c]] += k_glob[r,c]
                    F[dofs[r]] -= f_fixed_glob[r]

                element_results.append({
                    "ID": idx+1, "Nudo I": e['Nudo I'], "Nudo J": e['Nudo J'],
                    "L": L, "theta": theta, "E": E, "A": A, "I": I,
                    "b_cm": e.get('b (cm)', 30.0), "h_cm": e.get('h (cm)', 40.0),
                    "dofs": dofs, "T": T, "k_local": k_local, "f_fixed_local": f_fixed_local
                })

            # Eliminación de grados de libertad fijos (método directo)
            fixed_dofs = []
            for _, s in supports.iterrows():
                n_idx = id_to_idx[s['Nudo']]
                if s.get('Fijo X', False): fixed_dofs.append(n_idx*3 + 0)
                if s.get('Fijo Y', False): fixed_dofs.append(n_idx*3 + 1)
                if s.get('Fijo Rz', False): fixed_dofs.append(n_idx*3 + 2)
            free_dofs = [d for d in range(num_gdl) if d not in fixed_dofs]

            if len(free_dofs) == 0:
                st.error(_t("No hay grados de libertad libres. Verifique los apoyos.", "No free DOFs. Check supports."))
                st.stop()

            # Resolver sistema reducido
            Kff = K[np.ix_(free_dofs, free_dofs)]
            Ff = F[free_dofs]
            try:
                Uf = np.linalg.solve(Kff, Ff)
            except np.linalg.LinAlgError:
                st.error(_t("Matriz de rigidez singular. Verifique la estructura (apoyos insuficientes o rigidez cero).", "Singular stiffness matrix. Check structure (insufficient supports or zero stiffness)."))
                st.stop()

            # Desplazamientos totales
            U = np.zeros(num_gdl)
            U[free_dofs] = Uf
            # Reacciones
            R = K @ U
            # Fuerzas internas en elementos (en coordenadas locales)
            for e in element_results:
                u_glob = U[e['dofs']]
                u_loc = e['T'] @ u_glob
                # Sumar fuerzas de empotramiento a las deformaciones
                # Las fuerzas totales son: f = k_local * u_loc + f_fixed_local
                f_loc = e['k_local'] @ u_loc + e['f_fixed_local']
                e['N'] = f_loc[0]
                e['V'] = f_loc[1]
                e['M_i'] = f_loc[2]
                e['M_j'] = f_loc[5]

            # Guardar resultados
            st.session_state.resultados = {
                "U": U, "R": R, "elements": element_results,
                "nodes": nodes, "node_ids": node_ids, "id_to_idx": id_to_idx,
                "num_gdl": num_gdl, "free_dofs": free_dofs, "fixed_dofs": fixed_dofs
            }
            st.success(_t("Análisis completado.", "Analysis completed."))

# ─────────────────────────────────────────────
# TAB 3: RESULTADOS (solo si existe)
with tab_res:
    if st.session_state.resultados is None:
        st.info(_t("Realice el análisis en la pestaña 'Cargas y Suelo'.", "Run the analysis in 'Loads & Soil'tab."))
    else:
        res = st.session_state.resultados
        U = res["U"]
        R = res["R"]
        elements = res["elements"]
        nodes = res["nodes"]
        node_ids = res["node_ids"]
        id_to_idx = res["id_to_idx"]

        # Tabla de desplazamientos
        st.subheader(_t("Desplazamientos nodales", "Nodal displacements"))
        desp_data = []
        for nid in node_ids:
            idx = id_to_idx[nid]
            desp_data.append({
                "Nudo": nid,
                "dx (mm)": U[idx*3+0]*1000,
                "dy (mm)": U[idx*3+1]*1000,
                "rz (rad)": U[idx*3+2]
            })
        st.dataframe(pd.DataFrame(desp_data), use_container_width=True)

        # Tabla de reacciones
        st.subheader(_t("Reacciones en apoyos", "Support reactions"))
        react_data = []
        for _, s in st.session_state.apoyos_df.iterrows():
            nid = s['Nudo']
            if nid in id_to_idx:
                idx = id_to_idx[nid]
                react_data.append({
                    "Nudo": nid,
                    "Rx (kN)": R[idx*3+0],
                    "Ry (kN)": R[idx*3+1],
                    "Mz (kN-m)": R[idx*3+2]
                })
        st.dataframe(pd.DataFrame(react_data), use_container_width=True)

        # Tabla de fuerzas internas
        st.subheader(_t("Fuerzas internas en elementos", "Element internal forces"))
        forces_data = []
        for e in elements:
            forces_data.append({
                "Barra": e["ID"],
                "Axial I (kN)": e["N"],
                "Corte I (kN)": e["V"],
                "Momento I (kN-m)": e["M_i"],
                "Axial J (kN)": -e["N"],
                "Corte J (kN)": -e["V"],
                "Momento J (kN-m)": e["M_j"]
            })
        st.dataframe(pd.DataFrame(forces_data), use_container_width=True)

        # Gráfico de momentos flectores (BMD)
        st.subheader(_t("Diagrama de momentos flectores", "Bending moment diagram"))
        fig_bmd = go.Figure()
        for e in elements:
            xi = nodes.loc[nodes['ID']==e["Nudo I"], 'X (m)'].values[0]
            yi = nodes.loc[nodes['ID']==e["Nudo I"], 'Y (m)'].values[0]
            xj = nodes.loc[nodes['ID']==e["Nudo J"], 'X (m)'].values[0]
            yj = nodes.loc[nodes['ID']==e["Nudo J"], 'Y (m)'].values[0]
            L = e["L"]
            theta = e["theta"]
            Mi = e["M_i"]
            Mj = e["M_j"]
            # Escala para visualización
            max_mom = max(abs(Mi), abs(Mj))
            scale = max_mom / 50 if max_mom > 0 else 1
            # Dibujar barra
            fig_bmd.add_trace(go.Scatter(x=[xi, xj], y=[yi, yj], mode='lines', line=dict(color='white', width=2), showlegend=False))
            # Dibujar línea del momento (perpendicular a la barra)
            nx = -math.sin(theta)
            ny = math.cos(theta)
            # Puntos desplazados
            x_i = xi + nx * Mi/scale
            y_i = yi + ny * Mi/scale
            x_j = xj + nx * Mj/scale
            y_j = yj + ny * Mj/scale
            # Polígono relleno
            fig_bmd.add_trace(go.Scatter(
                x=[xi, x_i, x_j, xj, xi],
                y=[yi, y_i, y_j, yj, yi],
                fill="toself",
                fillcolor="rgba(0, 150, 255, 0.4)",
                line=dict(color="blue", width=1),
                name=f"BMD {e['ID']}",
                hoverinfo="text",
                text=f"M_i: {Mi:.1f} kN-m<br>M_j: {Mj:.1f} kN-m"
            ))
            # Etiquetas
            fig_bmd.add_annotation(x=x_i, y=y_i, text=f"{abs(Mi):.1f}", showarrow=False, font=dict(color="yellow", size=10))
            fig_bmd.add_annotation(x=x_j, y=y_j, text=f"{abs(Mj):.1f}", showarrow=False, font=dict(color="yellow", size=10))
        fig_bmd.update_layout(title=_t("Diagrama de momentos (escala indicativa)", "Moment diagram (indicative scale)"), 
                              plot_bgcolor='#1e1e1e', paper_bgcolor='#1e1e1e', font=dict(color='white'), height=600)
        st.plotly_chart(fig_bmd, use_container_width=True)

# ─────────────────────────────────────────────
# TAB 4: EXPORTACIONES (diseño automático y exportaciones)
with tab_dis:
    if st.session_state.resultados is None:
        st.info(_t("Realice el análisis en la pestaña 'Cargas y Suelo'.", "Run the analysis in 'Loads & Soil'tab."))
    else:
        res = st.session_state.resultados
        elements = res["elements"]
        R = res["R"]
        id_to_idx = res["id_to_idx"]

        st.header(_t("⚡ Diseño automático de elementos", "⚡ Automated element design"))

        # Preparar datos para exportación
        design_data = {
            "zapatas": [],
            "vigas": [],
            "columnas": []
        }

        # Diseño de zapatas (apoyos)
        st.subheader(_t("Zapatas", "Footings"))
        for _, s in st.session_state.apoyos_df.iterrows():
            nid = s['Nudo']
            if nid in id_to_idx:
                idx = id_to_idx[nid]
                Ry = abs(R[idx*3+1])
                Rx = abs(R[idx*3+0])
                Mz = abs(R[idx*3+2])
                # Momento transmitido a la zapata (puede ser debido a excentricidad)
                # Para simplificar, usamos Mz como momento de diseño
                Pu = Ry
                Mu = Mz
                B, H, As = design_footing(Pu, Mu, qa_suelo, fc, fy, recub)
                design_data["zapatas"].append({
                    "Nudo": nid,
                    "Pu (kN)": Pu,
                    "Mu (kN-m)": Mu,
                    "B (m)": B,
                    "H (m)": H,
                    "As (cm²)": As
                })
                st.success(f"**Zapata en nudo {nid}** → B = {B:.2f} m, H = {H:.2f} m, As ≈ {As:.1f} cm²/m")

        # Diseño de vigas y columnas
        st.subheader(_t("Vigas", "Beams"))
        for e in elements:
            # Determinar si es viga (ángulo cercano a 0° o 180°)
            deg = abs(np.degrees(e["theta"]))
            is_beam = (deg < 5) or (abs(deg - 180) < 5)
            if is_beam:
                # Momento máximo
                Mmax = max(abs(e["M_i"]), abs(e["M_j"]))
                # Cortante máximo
                Vmax = max(abs(e["V"]), abs(e["-V"]))  # corte en ambos extremos
                b = e["b_cm"]
                h = e["h_cm"]
                As, s = design_beam(Mmax, Vmax, b, h, fc, fy, recub)
                design_data["vigas"].append({
                    "Barra": e["ID"],
                    "Mmax (kN-m)": Mmax,
                    "Vmax (kN)": Vmax,
                    "b (cm)": b,
                    "h (cm)": h,
                    "As (cm²)": As,
                    "s_estribos (cm)": s if s > 0 else "No requiere"
                })
                st.info(f"**Viga {e['ID']}** (b={b:.0f} cm, h={h:.0f} cm): Mu={Mmax:.1f} kN-m → As ≈ {As:.1f} cm², cortante → s = {s:.1f} cm")
            else:
                # Columna
                # Carga axial y momento en extremos
                # Para simplificar, tomamos el máximo de ambos extremos
                P_axial = max(abs(e["N"]), abs(-e["N"]))
                M_max = max(abs(e["M_i"]), abs(e["M_j"]))
                b = e["b_cm"]
                h = e["h_cm"]
                As_min, ok = design_column(P_axial, M_max, b, h, fc, fy, recub)
                design_data["columnas"].append({
                    "Barra": e["ID"],
                    "Pu (kN)": P_axial,
                    "Mu (kN-m)": M_max,
                    "b (cm)": b,
                    "h (cm)": h,
                    "As_min (cm²)": As_min,
                    "Verificación": "OK" if ok else "Falla"
                })
                if ok:
                    st.success(f"**Columna {e['ID']}** (b={b:.0f} cm, h={h:.0f} cm): Pu={P_axial:.1f} kN, Mu={M_max:.1f} kN-m → As_min ≈ {As_min:.1f} cm²")
                else:
                    st.error(f"**Columna {e['ID']}** NO cumple: Pu={P_axial:.1f} kN > capacidad ≈ {0.85*fc*(b*h-As_min)+fy*As_min:.0f} kN")

        # Exportaciones
        st.markdown("---")
        st.subheader(_t("Exportar resultados", " Export results"))

        col_ex1, col_ex2, col_ex3, col_ex4 = st.columns(4)

        # 1. Excel
        with col_ex1:
            if st.button(_t("Exportar a Excel", " Export to Excel")):
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    # Reacciones
                    react = []
                    for _, s in st.session_state.apoyos_df.iterrows():
                        nid = s['Nudo']
                        if nid in id_to_idx:
                            idx = id_to_idx[nid]
                            react.append({"Nudo": nid, "Rx": R[idx*3+0], "Ry": R[idx*3+1], "Mz": R[idx*3+2]})
                    pd.DataFrame(react).to_excel(writer, sheet_name="Reacciones", index=False)
                    # Desplazamientos
                    desp = []
                    for nid in res["node_ids"]:
                        idx = id_to_idx[nid]
                        desp.append({"Nudo": nid, "dx (mm)": U[idx*3+0]*1000, "dy (mm)": U[idx*3+1]*1000, "rz (rad)": U[idx*3+2]})
                    pd.DataFrame(desp).to_excel(writer, sheet_name="Desplazamientos", index=False)
                    # Fuerzas internas
                    forces = []
                    for e in elements:
                        forces.append({"Barra": e["ID"], "Axial I": e["N"], "Corte I": e["V"], "Momento I": e["M_i"], "Momento J": e["M_j"]})
                    pd.DataFrame(forces).to_excel(writer, sheet_name="Fuerzas_internas", index=False)
                    # Diseño
                    pd.DataFrame(design_data["zapatas"]).to_excel(writer, sheet_name="Zapatas", index=False)
                    pd.DataFrame(design_data["vigas"]).to_excel(writer, sheet_name="Vigas", index=False)
                    pd.DataFrame(design_data["columnas"]).to_excel(writer, sheet_name="Columnas", index=False)
                output.seek(0)
                st.download_button(_t("Descargar Excel", " Download Excel"), data=output, file_name="resultados_estructura.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # 2. DXF
        with col_ex2:
            if st.button(_t("Exportar a DXF", " Export to DXF")):
                try:
                    from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                             dxf_rotulo, dxf_rotulo_campos)
                    _USE_H14 = True
                except ImportError:
                    _USE_H14 = False
                doc_dxf = ezdxf.new('R2010')
                doc_dxf.units = ezdxf.units.M
                if _USE_H14:
                    dxf_setup(doc_dxf, 50); dxf_add_layers(doc_dxf)
                msp = doc_dxf.modelspace()
                for lay, col14 in [('ESTRUCTURA',3),('ZAPATAS',4),('COTAS',2),('TEXTO',1)]:
                    if lay not in doc_dxf.layers: doc_dxf.layers.add(lay, color=col14)
                x_vals, y_vals = [], []
                for e in elements:
                    xi = nodes.loc[nodes['ID']==e["Nudo I"], 'X (m)'].values[0]
                    yi = nodes.loc[nodes['ID']==e["Nudo I"], 'Y (m)'].values[0]
                    xj = nodes.loc[nodes['ID']==e["Nudo J"], 'X (m)'].values[0]
                    yj = nodes.loc[nodes['ID']==e["Nudo J"], 'Y (m)'].values[0]
                    msp.add_line((xi, yi), (xj, yj), dxfattribs={'layer':'ESTRUCTURA'})
                    x_vals += [xi, xj]; y_vals += [yi, yj]
                for zap in design_data["zapatas"]:
                    nid = zap["Nudo"]
                    x = nodes.loc[nodes['ID']==nid, 'X (m)'].values[0]
                    y = nodes.loc[nodes['ID']==nid, 'Y (m)'].values[0]
                    B = zap["B (m)"]; half = B/2
                    msp.add_lwpolyline([(x-half,y-half),(x+half,y-half),(x+half,y+half),(x-half,y+half),(x-half,y-half)], close=True, dxfattribs={'layer':'ZAPATAS'})
                if _USE_H14:
                    _x0 = min(x_vals) if x_vals else 0
                    _y0 = min(y_vals) if y_vals else 0
                    _xw = max(x_vals)-_x0 if x_vals else 10
                    dxf_text(msp, _x0, (max(y_vals) if y_vals else 0)+0.5, f"StructMaster 2D – {norma_sel}", "EJES", h=0.025*50, ha="left")
                    _cam14 = dxf_rotulo_campos("Analisis Estructural 2D", norma_sel, "001")
                    dxf_rotulo(msp, _cam14, _x0, _y0-4, rot_w=max(_xw,10), rot_h=3, escala=50)
                else:
                    msp.add_text(_t("Estructura 2D","2D Structure"), dxfattribs={'layer':'TEXTO','height':0.2,'insert':(0,0)})
                out_dxf = io.StringIO()
                doc_dxf.write(out_dxf)
                st.download_button(_t("Descargar DXF"," Download DXF"), data=out_dxf.getvalue().encode('utf-8'), file_name="estructura_2d.dxf", mime="application/dxf")

        # 3. DOCX Memoria
        with col_ex3:
            if st.button(_t("Generar Memoria DOCX", " Generate DOCX Report")):
                doc = Document()
                doc.add_heading(_t("Memoria de análisis estructural 2D", "2D Structural Analysis Report"), 0)
                doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
                doc.add_paragraph(_t(f"Norma aplicada: {norma_sel}", f"Applied code: {norma_sel}"))
                doc.add_paragraph(_t(f"Materiales: f'c = {fc} MPa, fy = {fy} MPa", f"Materials: f'c = {fc} MPa, fy = {fy} MPa"))
                # Reacciones
                doc.add_heading(_t("Reacciones en apoyos", "Support reactions"), level=1)
                react_table = doc.add_table(rows=1+len(react), cols=4)
                react_table.style = 'Table Grid'
                hdr = react_table.rows[0].cells
                hdr[0].text = "Nudo"
                hdr[1].text = "Rx (kN)"
                hdr[2].text = "Ry (kN)"
                hdr[3].text = "Mz (kN-m)"
                for i, r in enumerate(react):
                    cells = react_table.rows[i+1].cells
                    cells[0].text = str(r["Nudo"])
                    cells[1].text = f"{r['Rx']:.1f}"
                    cells[2].text = f"{r['Ry']:.1f}"
                    cells[3].text = f"{r['Mz']:.1f}"
                # Desplazamientos
                doc.add_heading(_t("Desplazamientos", "Displacements"), level=1)
                desp_table = doc.add_table(rows=1+len(desp), cols=4)
                desp_table.style = 'Table Grid'
                hdr = desp_table.rows[0].cells
                hdr[0].text = "Nudo"
                hdr[1].text = "dx (mm)"
                hdr[2].text = "dy (mm)"
                hdr[3].text = "rz (rad)"
                for i, d in enumerate(desp):
                    cells = desp_table.rows[i+1].cells
                    cells[0].text = str(d["Nudo"])
                    cells[1].text = f"{d['dx (mm)']:.2f}"
                    cells[2].text = f"{d['dy (mm)']:.2f}"
                    cells[3].text = f"{d['rz (rad)']:.4f}"
                # Diseño
                doc.add_heading(_t("Diseño de elementos", "Element design"), level=1)
                if design_data["zapatas"]:
                    doc.add_heading(_t("Zapatas", "Footings"), level=2)
                    for z in design_data["zapatas"]:
                        doc.add_paragraph(f"Nudo {z['Nudo']}: B={z['B (m)']:.2f} m, H={z['H (m)']:.2f} m, As≈{z['As (cm²)']:.1f} cm²/m")
                if design_data["vigas"]:
                    doc.add_heading(_t("Vigas", "Beams"), level=2)
                    for v in design_data["vigas"]:
                        doc.add_paragraph(f"Barra {v['Barra']}: b={v['b (cm)']:.0f} cm, h={v['h (cm)']:.0f} cm, As={v['As (cm²)']:.2f} cm², estribos @ {v['s_estribos (cm)']} cm")
                if design_data["columnas"]:
                    doc.add_heading(_t("Columnas", "Columns"), level=2)
                    for c in design_data["columnas"]:
                        doc.add_paragraph(f"Barra {c['Barra']}: b={c['b (cm)']:.0f} cm, h={c['h (cm)']:.0f} cm, As_min={c['As_min (cm²)']:.2f} cm², verificación: {c['Verificación']}")

                buf_doc = io.BytesIO()
                doc.save(buf_doc)
                buf_doc.seek(0)
                st.download_button(_t("Descargar Memoria", " Download Report"), data=buf_doc, file_name="memoria_estructura_2d.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # 4. APU (presupuesto)
        with col_ex4:
            if st.button(_t("Presupuesto APU", " APU Budget")):
                if "apu_config" in st.session_state:
                    apu = st.session_state.apu_config
                    mon = apu.get("moneda", "$")
                    # Estimación simple de volumen de concreto y acero
                    vol_conc = 0.0
                    peso_acero = 0.0
                    # Zapatas
                    for z in design_data["zapatas"]:
                        vol_conc += z["B (m)"]**2 * z["H (m)"]  # aproximado
                        peso_acero += z["As (cm²)"] * z["B (m)"] * 100 * 0.785  # kg
                    # Vigas
                    for v in design_data["vigas"]:
                        vol_conc += v["b (cm)"] * v["h (cm)"] * v.get("L", 4) / 10000  # m³, L aprox
                        peso_acero += v["As (cm²)"] * v.get("L", 4) * 0.785
                    # Columnas
                    for c in design_data["columnas"]:
                        vol_conc += c["b (cm)"] * c["h (cm)"] * 3 / 10000  # altura típica 3 m
                        peso_acero += c["As_min (cm²)"] * 3 * 0.785
                    # Costos
                    costo_conc = vol_conc * apu.get("precio_concreto_m3", 0) if apu.get("usar_concreto_premezclado", False) else 0
                    costo_ace = peso_acero * apu.get("acero", 0)
                    st.write(f"Volumen de concreto estimado: {vol_conc:.2f} m³")
                    st.write(f"Peso de acero estimado: {peso_acero:.1f} kg")
                    st.write(f"**Costo total materiales:** {mon} {costo_conc + costo_ace:,.2f}")
                else:
                    st.info(_t("Configure precios en la página 'APU Mercado'.", "Configure prices in 'APU Mercado'page."))