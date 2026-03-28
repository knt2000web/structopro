import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import datetime
import json

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

st.set_page_config(page_title=_t("Generador Maestro 3D", "Master 3D Generator"), layout="wide")
st.image(r"assets/generador_3d_header_1773257156151.png", use_container_width=True)
st.title(_t("Generador Maestro 3D Paramétrico", "Parametric Master 3D Generator"))
st.markdown(_t("Construcción rápida de Edificios 3D con cálculos estructurales completos según la norma seleccionada, incluyendo análisis sísmico.", 
               "Rapid 3D Building generation with complete structural calculations per the selected code, including seismic analysis."))

# ─────────────────────────────────────────────
# PARÁMETROS POR NORMA (para diseño)
# ─────────────────────────────────────────────
CODES = {
    "NSR-10 (Colombia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NSR-10 Título C"},
    "ACI 318-25 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-25"},
    "ACI 318-19 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-19"},
    "ACI 318-14 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-14"},
    "NEC-SE-HM (Ecuador)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NEC-SE-HM"},
    "E.060 (Perú)": {"phi_flex":0.90, "phi_shear":0.85, "phi_comp":0.70, "lambda":1.0, "ref":"E.060"},
    "NTC-EM (México)": {"phi_flex":0.85, "phi_shear":0.80, "phi_comp":0.70, "lambda":1.0, "ref":"NTC-EM"},
    "COVENIN 1753-2006 (Venezuela)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.70, "lambda":1.0, "ref":"COVENIN"},
    "NB 1225001-2020 (Bolivia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NB"},
    "CIRSOC 201-2025 (Argentina)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"CIRSOC"},
}
code = CODES.get(norma_sel, CODES["NSR-10 (Colombia)"])
phi_flex = code["phi_flex"]
phi_shear = code["phi_shear"]
phi_comp = code["phi_comp"]
lam = code["lambda"]

# ─────────────────────────────────────────────
# BASES DE DATOS SÍSMICAS (extraídas del módulo de diseño sísmico)
# ─────────────────────────────────────────────
# (Se usa la misma estructura que en el módulo 10_Diseño_Sismico.py pero adaptada)

# Definir ciudades para cada norma (solo ejemplos representativos)
SEISMIC_DATA = {
    "NSR-10 (Colombia)": {
        "ciudades": {
            "Bogotá": {"Aa": 0.15, "Av": 0.20},
            "Medellín": {"Aa": 0.15, "Av": 0.20},
            "Cali": {"Aa": 0.25, "Av": 0.25},
            "Barranquilla": {"Aa": 0.10, "Av": 0.10},
            "Bucaramanga": {"Aa": 0.25, "Av": 0.25},
            "Cartagena": {"Aa": 0.10, "Av": 0.10},
            "Cúcuta": {"Aa": 0.35, "Av": 0.30},
            "Pereira": {"Aa": 0.25, "Av": 0.25},
            "Manizales": {"Aa": 0.25, "Av": 0.25},
            "Armenia": {"Aa": 0.25, "Av": 0.25},
            "Ibagué": {"Aa": 0.20, "Av": 0.20},
            "Neiva": {"Aa": 0.25, "Av": 0.25},
            "Pasto": {"Aa": 0.25, "Av": 0.25},
            "Villavicencio": {"Aa": 0.35, "Av": 0.30},
            "Santa Marta": {"Aa": 0.15, "Av": 0.10},
            "Otra / Custom": {"Aa": 0.20, "Av": 0.20}
        },
        "suelos": {
            "A": {"Fa": 0.8, "Fv": 0.8, "Tp": 0.4, "TL": 3.0},
            "B": {"Fa": 1.0, "Fv": 1.0, "Tp": 0.6, "TL": 3.0},
            "C": {"Fa": 1.2, "Fv": 1.6, "Tp": 0.8, "TL": 3.0},
            "D": {"Fa": 1.4, "Fv": 2.0, "Tp": 1.0, "TL": 3.0},
            "E": {"Fa": 2.5, "Fv": 3.2, "Tp": 1.2, "TL": 3.0},
        },
        "uso": {"I (Grupo I)": 1.0, "II (Grupo II)": 1.0, "III (Grupo III)": 1.1, "IV (Grupo IV)": 1.5},
        "R": 6.0  # coeficiente de reducción por ductilidad (pórticos resistentes a momentos)
    },
    "E.060 (Perú)": {
        "ciudades": {
            "Zona 4 (Lima)": {"Z": 0.45},
            "Zona 3 (Arequipa)": {"Z": 0.35},
            "Zona 2 (Cusco)": {"Z": 0.25},
            "Zona 1 (Selva)": {"Z": 0.10}
        },
        "suelos": {
            "S0": {"S": 0.8, "Tp": 0.3, "Tl": 3.0},
            "S1": {"S": 1.0, "Tp": 0.4, "Tl": 2.5},
            "S2": {"S": 1.15, "Tp": 0.6, "Tl": 2.0},
            "S3": {"S": 1.4, "Tp": 1.0, "Tl": 1.6},
        },
        "uso": {"A (Importante)": 1.5, "B (Común)": 1.3, "C (Menor)": 1.0},
        "R": 6.0
    },
    # Para otras normas se puede extender; por simplicidad, se usan valores por defecto si no están definidas
}
# Fallback para normas no listadas
if norma_sel not in SEISMIC_DATA:
    SEISMIC_DATA[norma_sel] = {
        "ciudades": {"Defecto": {"Z": 0.30}},
        "suelos": {"Defecto": {"S": 1.0, "Tp": 0.6, "Tl": 2.5}},
        "uso": {"Categoría B": 1.0},
        "R": 6.0
    }
seismic_info = SEISMIC_DATA[norma_sel]

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES
# ─────────────────────────────────────────────
def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_rho_min(fc, fy):
    return max(0.25*math.sqrt(fc)/fy, 1.4/fy)

def get_rho_max_beam(fc, fy, beta1):
    eps_cu = 0.003
    eps_t_min = 0.005
    return (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+eps_t_min))

def compute_spectrum(norma, params, T_domain):
    """
    Calcula el espectro de diseño para los parámetros dados.
    Devuelve Sa_vals (array) y key_periods.
    """
    Sa_vals = np.zeros_like(T_domain)
    key_periods = {'T0':0, 'Tc':0, 'TL':0}
    if "NSR" in norma:
        Aa = params['Aa']; Av = params['Av']; Fa = params['Fa']; Fv = params['Fv']; I = params['I']
        Tc = 0.48 * (Av * Fv) / (Aa * Fa)
        T0 = 0.1 * Tc
        TL = 2.4 * Fv
        key_periods = {'T0': T0, 'Tc': Tc, 'TL': TL}
        for i, t in enumerate(T_domain):
            if t < T0:
                sa = Aa * Fa * (1.0 + (t/T0)*(2.5 - 1.0))
            elif t <= Tc:
                sa = 2.5 * Aa * Fa
            elif t <= TL:
                sa = 1.2 * Av * Fv / t
            else:
                sa = 1.2 * Av * Fv * TL / (t**2)
            Sa_vals[i] = sa * I
    elif "E.060" in norma or "Perú" in norma:
        Z = params['Z']; S = params['S']; Tp = params['Tp']; Tl = params['Tl']; U = params['U']; R = params['R']
        key_periods = {'T0': 0.2*Tp, 'Tc': Tp, 'TL': Tl}
        for i, t in enumerate(T_domain):
            if t < 0.2*Tp:
                C = 1 + 1.5 * (t / (0.2*Tp))
            elif t <= Tp:
                C = 2.5
            elif t <= Tl:
                C = 2.5 * (Tp / t)
            else:
                C = 2.5 * (Tp * Tl / (t**2))
            Sa_vals[i] = (Z * U * C * S) / R
    else:
        # General ASCE / otros (simplificado)
        S_DS = params.get('S_DS', 0.6); S_D1 = params.get('S_D1', 0.4); TL = params.get('TL', 4.0)
        T0 = 0.2 * (S_D1 / S_DS)
        Ts = S_D1 / S_DS
        key_periods = {'T0': T0, 'Tc': Ts, 'TL': TL}
        for i, t in enumerate(T_domain):
            if t < T0:
                sa = S_DS * (0.4 + 0.6 * (t / T0))
            elif t <= Ts:
                sa = S_DS
            elif t <= TL:
                sa = S_D1 / t
            else:
                sa = (S_D1 * TL) / (t**2)
            Sa_vals[i] = sa
    return Sa_vals, key_periods

def design_column(Pu, Mu, b, h, fc, fy, recub=5):
    """Diseño de columna rectangular con flexión uniaxial (simplificado). Retorna As (cm²), ok, msg."""
    Ag = b * h * 10000  # cm²
    As_min = 0.01 * Ag
    # Capacidad axial con acero mínimo (sin momento)
    Pn_max = 0.85 * fc * (Ag - As_min) / 1000 + fy * As_min / 1000  # kN
    phi_Pn_max = phi_comp * Pn_max
    if Pu > phi_Pn_max:
        As_needed = max(As_min, (Pu / phi_comp - 0.85*fc*Ag/1000) / (fy/1000))
        As_needed = max(As_needed, 0.01 * Ag)
        ok = False
        As = As_needed
    else:
        As = As_min
        ok = True
    # Verificación de momento (muy simplificada: se aumenta As si el momento es significativo)
    # Criterio: Mu > 0.1 * Pu * (min(b,h)/100) -> momento relevante
    if Mu > 0.1 * Pu * (min(b, h)/100):
        # Aumentar As en 30% (arbitrario)
        As = As * 1.3
        msg = "Momento significativo - acero incrementado 30%"
    else:
        msg = ""
    return As, ok, msg

def design_beam(Mu, Vu, b, h, fc, fy, recub=5):
    """Diseño de viga rectangular. Retorna As (cm²) y s_estribos (cm)."""
    d = h - recub - 1  # cm (estribo supuesto 1 cm)
    # Flexión
    Rn = (Mu * 1e6) / (phi_flex * b * d**2)
    if Rn <= 0:
        rho = 0.0018
    else:
        disc = 1 - 2*Rn/(0.85*fc)
        if disc > 0:
            rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
        else:
            rho = 0.0018
    rho_min = get_rho_min(fc, fy)
    rho = max(rho, rho_min)
    As = rho * b * d
    # Cortante
    Vc = 0.17 * lam * math.sqrt(fc) * b * d / 10  # kN
    phi_Vc = phi_shear * Vc
    if Vu > phi_Vc/2:
        Vs = Vu / phi_shear - Vc
        Av = 2 * 0.71  # #3 estribo, 2 ramas (cm²)
        s = Av * fy * d / Vs  # cm
        s = max(5, min(s, 60, d/2))
    else:
        s = 0
    return As, s

def design_footing(Pu, Mu, qa, fc, fy, b_col, recub=5):
    """Diseño de zapata cuadrada. Retorna B (m), H (m), As (cm²/m)."""
    # Área requerida por carga axial
    A_req = Pu / qa
    B = math.sqrt(A_req)
    B = math.ceil(B*20)/20  # redondear a 0.05 m
    # Espesor mínimo (corte unidireccional aproximado)
    # Suponiendo columna cuadrada de lado b_col (cm)
    L_v = (B - b_col/100)/2
    # Presión neta promedio
    q_net = Pu / (B*B)
    # Momento en la cara de la columna
    Mu_zap = q_net * B * (L_v**2)/2
    # Añadir momento sísmico transferido a la zapata (se considera un 30% del momento de columna)
    Mu_zap += 0.3 * abs(Mu)
    # Peralte mínimo por flexión (estimado)
    d = max(0.25, L_v/4)  # cm (en metros)
    H = d + recub/100 + 0.01
    H = math.ceil(H*20)/20
    # Acero por flexión (simplificado)
    d_cm = H*100 - recub - 1
    b_cm = B*100
    Rn = (Mu_zap * 1e6) / (phi_flex * b_cm * d_cm**2)
    if Rn > 0:
        disc = 1 - 2*Rn/(0.85*fc)
        if disc > 0:
            rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
        else:
            rho = 0.0018
    else:
        rho = 0.0018
    As = rho * b_cm * d_cm
    As_min = 0.0018 * b_cm * H*100
    As = max(As, As_min)
    return B, H, As

def generate_mesh(Lx, Lz, nx, nz, alturas_pisos, col_b, col_h, vig_b, vig_h):
    """Genera malla 3D con nudos, columnas, vigas y zapatas (sin dimensiones de zapata aún)."""
    pisos = len(alturas_pisos)
    # Coordenadas de nudos
    x_coords = np.linspace(0, Lx, nx)
    z_coords = np.linspace(0, Lz, nz)
    y_coords = [0.0]
    curr_y = 0.0
    for h in alturas_pisos:
        curr_y += h
        y_coords.append(curr_y)
    nudos = []
    nid = 1
    for y in y_coords:
        for z in z_coords:
            for x in x_coords:
                nudos.append({"ID": nid, "X": round(x,2), "Y": round(y,2), "Z": round(z,2)})
                nid += 1
    df_nudos = pd.DataFrame(nudos)

    def find_nid(x, y, z):
        res = df_nudos[(np.isclose(df_nudos['X'], x)) & (np.isclose(df_nudos['Y'], y)) & (np.isclose(df_nudos['Z'], z))]
        return int(res.iloc[0]['ID']) if not res.empty else None

    # Columnas
    columnas = []
    cid = 1
    for z in z_coords:
        for x in x_coords:
            for p in range(pisos):
                n1 = find_nid(x, y_coords[p], z)
                n2 = find_nid(x, y_coords[p+1], z)
                if n1 and n2:
                    columnas.append({"ID": f"C{cid}", "N1": n1, "N2": n2, "Base_X": x, "Base_Z": z, "Piso": p+1, "b (m)": col_b, "h (m)": col_h})
                    cid += 1

    # Vigas X
    vigas_x = []
    vid = 1
    for y in y_coords[1:]:
        for z in z_coords:
            for i in range(nx-1):
                n1 = find_nid(x_coords[i], y, z)
                n2 = find_nid(x_coords[i+1], y, z)
                if n1 and n2:
                    vigas_x.append({"ID": f"VX{vid}", "N1": n1, "N2": n2, "Piso": y_coords.index(y), "b (m)": vig_b, "h (m)": vig_h})
                    vid += 1
    # Vigas Z
    vigas_z = []
    for y in y_coords[1:]:
        for x in x_coords:
            for i in range(nz-1):
                n1 = find_nid(x, y, z_coords[i])
                n2 = find_nid(x, y, z_coords[i+1])
                if n1 and n2:
                    vigas_z.append({"ID": f"VZ{vid}", "N1": n1, "N2": n2, "Piso": y_coords.index(y), "b (m)": vig_b, "h (m)": vig_h})
                    vid += 1

    # Zapatas (solo la base, luego se dimensionan)
    zapatas = []
    zid = 1
    for z in z_coords:
        for x in x_coords:
            nid = find_nid(x, 0, z)
            if nid:
                es_borde_x = (np.isclose(x, 0) or np.isclose(x, Lx))
                es_borde_z = (np.isclose(z, 0) or np.isclose(z, Lz))
                if es_borde_x and es_borde_z:
                    tipo = "Esquina"
                elif es_borde_x or es_borde_z:
                    tipo = "Borde"
                else:
                    tipo = "Central"
                zapatas.append({"ID": f"Z{zid}", "Nudo": nid, "Tipo": tipo, "X": x, "Z": z})
                zid += 1
    return df_nudos, pd.DataFrame(columnas), pd.DataFrame(vigas_x), pd.DataFrame(vigas_z), pd.DataFrame(zapatas)

def compute_seismic_forces(Peso, n_pisos, alturas, params, norma):
    """Calcula cortante basal y fuerzas laterales por piso."""
    # Período aproximado
    T_aprox = 0.1 * n_pisos  # fórmula empírica
    T_domain = np.linspace(0.01, 3.0, 200)
    Sa_vals, _ = compute_spectrum(norma, params, T_domain)
    Sa = np.interp(T_aprox, T_domain, Sa_vals) if T_aprox <= T_domain[-1] else Sa_vals[-1]
    V_basal = Sa * Peso  # Sa en g, Peso en kN
    # Distribución vertical (triangular)
    h_pisos = [sum(alturas[:i+1]) for i in range(n_pisos)]
    V_piso = [V_basal * (hi / sum(h_pisos)) for hi in h_pisos]
    return V_basal, V_piso, Sa, T_aprox

def design_all_elements(nudos_df, cols_df, vigas_x_df, vigas_z_df, zaps_df,
                        fc, fy, qa, recub, alturas_pisos, area_planta,
                        cm, cv, seismic_params, norma):
    """Diseño completo con cargas sísmicas."""
    # Factores de combinación (LRFD)
    wu_grav = 1.2*cm + 1.6*cv  # kN/m²
    # Área tributaria por columna (simplificada: igual para todas)
    x_vals = sorted(nudos_df['X'].unique())
    z_vals = sorted(nudos_df['Z'].unique())
    dx = x_vals[1] - x_vals[0] if len(x_vals)>1 else 1
    dz = z_vals[1] - z_vals[0] if len(z_vals)>1 else 1
    area_trib = dx * dz

    # Peso total de la estructura (para cortante basal)
    Peso_total = wu_grav * area_planta * len(alturas_pisos)  # kN
    # Cálculo sísmico
    V_basal, V_piso, Sa, T_aprox = compute_seismic_forces(Peso_total, len(alturas_pisos), alturas_pisos, seismic_params, norma)
    # Para simplificar, distribuimos las fuerzas laterales por piso y asumimos que cada columna recibe una fracción
    # de la fuerza de piso según su área tributaria (uniforme)
    n_cols_por_piso = len(cols_df[cols_df['Piso']==1])  # número de columnas en planta
    F_piso_por_col = [V / n_cols_por_piso for V in V_piso]  # kN por columna por piso

    # Resultados de diseño
    col_results = []
    for idx, col in cols_df.iterrows():
        piso = col['Piso']
        # Carga axial gravitacional acumulada desde el piso hasta la azotea
        pisos_sobre = len(alturas_pisos) - piso + 1
        Pu_grav = wu_grav * area_trib * pisos_sobre
        # Momento sísmico en la columna: considerar el momento de piso como V*altura/2 (simplificado)
        # Tomamos el cortante de piso (por columna) y la altura de entrepiso
        h_piso = alturas_pisos[piso-1]
        V_col = F_piso_por_col[piso-1]
        Mu_sismo = V_col * h_piso / 2  # aproximación
        # Combinación sísmica: 1.2D + 1.0E + 0.5L (usamos wu_grav ya incluye 1.2D+1.6L, pero para sismo combinamos)
        # Para simplificar, tomamos Pu = Pu_grav (ya mayorado) y Mu = Mu_sismo (mayorado por 1.0)
        Pu = Pu_grav
        Mu = Mu_sismo
        b_cm = col['b (m)']*100
        h_cm = col['h (m)']*100
        As, ok, msg = design_column(Pu, Mu, b_cm, h_cm, fc, fy, recub)
        # Volumen de concreto
        n1 = nudos_df[nudos_df['ID']==col['N1']].iloc[0]
        n2 = nudos_df[nudos_df['ID']==col['N2']].iloc[0]
        L = abs(n2['Y'] - n1['Y'])
        vol = L * col['b (m)'] * col['h (m)']
        # Peso de acero: As (cm²) * L (m) * 0.785 kg/m (factor de conversión)
        peso_acero = As * L * 0.785  # kg
        col_results.append({
            "ID": col['ID'], "Piso": piso, "b (cm)": b_cm, "h (cm)": h_cm,
            "Pu (kN)": Pu, "Mu (kN-m)": Mu, "As (cm²)": As, "Cumple": ok,
            "Vol (m³)": vol, "Peso_acero (kg)": peso_acero, "Observación": msg
        })
    df_cols = pd.DataFrame(col_results)

    # Diseño de vigas (similar, usando el cortante sísmico en vigas)
    beam_results = []
    for _, v in pd.concat([vigas_x_df, vigas_z_df]).iterrows():
        n1 = nudos_df[nudos_df['ID']==v['N1']].iloc[0]
        n2 = nudos_df[nudos_df['ID']==v['N2']].iloc[0]
        if 'X' in v['ID']:
            L = abs(n2['X'] - n1['X'])
            ancho_trib = dz/2 + dz/2  # = dz
        else:
            L = abs(n2['Z'] - n1['Z'])
            ancho_trib = dx/2 + dx/2  # = dx
        # Carga gravitacional en viga
        w_viga_grav = wu_grav * ancho_trib  # kN/m
        Mu_grav = w_viga_grav * L**2 / 8
        Vu_grav = w_viga_grav * L / 2
        # Carga sísmica en vigas: se asigna el cortante de piso distribuido a las vigas de ese piso
        piso = v['Piso']
        V_piso_total = V_piso[piso-1]
        # Número de vigas en ese piso en la dirección X o Z (contar)
        n_vigas_dir = len(vigas_x_df[vigas_x_df['Piso']==piso]) if 'X' in v['ID'] else len(vigas_z_df[vigas_z_df['Piso']==piso])
        V_sismo_viga = V_piso_total / n_vigas_dir  # cortante por viga
        # Momento sísmico en viga (aproximado: V*L/2)
        Mu_sismo = V_sismo_viga * L / 2
        # Combinación: Mu = Mu_grav + Mu_sismo
        Mu = Mu_grav + Mu_sismo
        Vu = Vu_grav + V_sismo_viga
        b_cm = v['b (m)']*100
        h_cm = v['h (m)']*100
        As, s = design_beam(Mu, Vu, b_cm, h_cm, fc, fy, recub)
        vol = L * v['b (m)'] * v['h (m)']
        peso_acero = As * L * 0.785  # kg
        beam_results.append({
            "ID": v['ID'], "Piso": piso, "b (cm)": b_cm, "h (cm)": h_cm,
            "L (m)": L, "Mu (kN-m)": Mu, "Vu (kN)": Vu, "As (cm²)": As,
            "s_estribos (cm)": s, "Vol (m³)": vol, "Peso_acero (kg)": peso_acero
        })
    df_beams = pd.DataFrame(beam_results)

    # Zapatas
    zap_results = []
    for _, zap in zaps_df.iterrows():
        # Buscar columna base (Piso 1) en la misma ubicación (X, Z)
        col_base = cols_df[(cols_df['Base_X'] == zap['X']) & (cols_df['Base_Z'] == zap['Z']) & (cols_df['Piso']==1)]
        if not col_base.empty:
            # Obtener el Pu y Mu de la columna base
            idx = col_base.index[0]
            Pu_zap = col_results[idx]['Pu (kN)']
            Mu_zap = col_results[idx]['Mu (kN-m)']
        else:
            Pu_zap = 0
            Mu_zap = 0
        b_columna = col_base.iloc[0]['b (m)']*100 if not col_base.empty else 30
        B, H, As_zap = design_footing(Pu_zap, Mu_zap, qa, fc, fy, b_columna, recub)
        vol = B * B * H
        peso_acero = As_zap * (B*100) * 0.785  # kg (acero en ambas direcciones)
        zap_results.append({
            "ID": zap['ID'], "Tipo": zap['Tipo'], "Pu (kN)": Pu_zap, "Mu (kN-m)": Mu_zap,
            "B (m)": B, "H (m)": H, "As (cm²/m)": As_zap, "Vol (m³)": vol, "Peso_acero (kg)": peso_acero
        })
    df_zaps = pd.DataFrame(zap_results)

    # Losa (diseño simplificado unidireccional)
    L_short = min(dx, dz)
    L_long = max(dx, dz)
    if L_long / L_short > 2:
        moment_coef = 1/8  # unidireccional
    else:
        moment_coef = 1/10  # bidireccional aproximado
    w_losa = wu_grav
    Mu_losa = w_losa * L_short**2 * moment_coef
    # Espesor por norma (mínimo L/20)
    h_losa = max(0.10, L_short/20)
    # Diseño de losa por metro
    b_losa = 100  # cm
    d_losa = h_losa*100 - recub - 0.5
    Rn = (Mu_losa * 1e6) / (phi_flex * b_losa * d_losa**2)
    disc = 1 - 2*Rn/(0.85*fc) if Rn>0 else -1
    if disc > 0:
        rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
    else:
        rho = 0.0018
    rho_min = get_rho_min(fc, fy)
    rho = max(rho, rho_min)
    As_losa = rho * b_losa * d_losa  # cm²/m
    As_temp = 0.0018 * b_losa * h_losa*100
    As_losa = max(As_losa, As_temp)
    vol_losa = area_planta * h_losa
    peso_acero_losa = (As_losa * (area_planta*100) / 100) * 0.785 * 2  # dos direcciones

    return df_cols, df_beams, df_zaps, vol_losa, peso_acero_losa, h_losa, As_losa, Sa, T_aprox, V_basal, V_piso

def plot_edificio(nudos, cols, vigas_x, vigas_z, zaps, zap_dim=None):
    """Dibuja modelo 3D con zapatas como sólidos si se proporcionan dimensiones."""
    fig = go.Figure()
    # Columnas
    for _, c in cols.iterrows():
        n1 = nudos[nudos['ID']==c['N1']].iloc[0]
        n2 = nudos[nudos['ID']==c['N2']].iloc[0]
        fig.add_trace(go.Scatter3d(
            x=[n1['X'], n2['X']], y=[n1['Z'], n2['Z']], z=[n1['Y'], n2['Y']],
            mode='lines', line=dict(color='#2196F3', width=6), name="Columnas", showlegend=False
        ))
    # Vigas
    for _, v in vigas_x.iterrows():
        n1 = nudos[nudos['ID']==v['N1']].iloc[0]
        n2 = nudos[nudos['ID']==v['N2']].iloc[0]
        fig.add_trace(go.Scatter3d(
            x=[n1['X'], n2['X']], y=[n1['Z'], n2['Z']], z=[n1['Y'], n2['Y']],
            mode='lines', line=dict(color='#4CAF50', width=4), name="Vigas", showlegend=False
        ))
    for _, v in vigas_z.iterrows():
        n1 = nudos[nudos['ID']==v['N1']].iloc[0]
        n2 = nudos[nudos['ID']==v['N2']].iloc[0]
        fig.add_trace(go.Scatter3d(
            x=[n1['X'], n2['X']], y=[n1['Z'], n2['Z']], z=[n1['Y'], n2['Y']],
            mode='lines', line=dict(color='#4CAF50', width=4), name="Vigas", showlegend=False
        ))
    # Zapatas (como sólidos si hay dimensiones)
    if zap_dim is not None:
        for i, zap in zaps.iterrows():
            if i < len(zap_dim):
                B = zap_dim[i]['B']
                H = zap_dim[i]['H']
            else:
                B, H = 1.2, 0.4
            xc = zap['X']
            zc = zap['Z']
            # Dibujar un prisma rectangular
            fig.add_trace(go.Mesh3d(
                x=[xc-B/2, xc+B/2, xc+B/2, xc-B/2, xc-B/2, xc+B/2, xc+B/2, xc-B/2],
                y=[zc-B/2, zc-B/2, zc+B/2, zc+B/2, zc-B/2, zc-B/2, zc+B/2, zc+B/2],
                z=[0, 0, 0, 0, -H, -H, -H, -H],
                i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
                color='#d95b0e', opacity=0.8, name="Zapatas", showlegend=False
            ))
    else:
        # Puntos como marcadores
        fig.add_trace(go.Scatter3d(
            x=zaps['X'], y=zaps['Z'], z=[0]*len(zaps), mode='markers',
            marker=dict(size=4, color='red', symbol='square'), name="Zapatas", showlegend=False
        ))
    # Losa (superficie en cada piso)
    y_pisos = sorted(nudos['Y'].unique())
    for y in y_pisos[1:]:
        nx_coords = nudos[nudos['Y']==y]['X']
        nz_coords = nudos[nudos['Y']==y]['Z']
        fig.add_trace(go.Mesh3d(
            x=[nx_coords.min(), nx_coords.max(), nx_coords.max(), nx_coords.min()],
            y=[nz_coords.min(), nz_coords.min(), nz_coords.max(), nz_coords.max()],
            z=[y, y, y, y],
            color='cyan', opacity=0.1, name="Losa", showlegend=False
        ))
    fig.update_layout(
        scene=dict(xaxis_title='X (m)', yaxis_title='Z (m)', zaxis_title='Y (m)', aspectmode='data'),
        margin=dict(l=0, r=0, b=0, t=30),
        plot_bgcolor='black', paper_bgcolor='#1e1e1e'
    )
    return fig

# ─────────────────────────────────────────────
# INTERFAZ DE USUARIO
# ─────────────────────────────────────────────
# Barra lateral
with st.sidebar:
    _iso = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}.get(norma_sel, "un")
    st.markdown(f'<div style="background:#1e3a1e;border-radius:6px;padding:8px;margin-bottom:10px;"><img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;"><span style="color:#7ec87e;font-weight:600;">{_t("Normativa Activa:","Code:")} {norma_sel}</span></div>', unsafe_allow_html=True)
    st.markdown("---")

    st.header(_t("📏 Geometría", "📏 Geometry"))
    L_x = st.number_input("Frente Lote X (m)", value=12.0, min_value=2.0, step=0.5, key="g3d_lx")
    L_z = st.number_input("Fondo Lote Z (m)", value=15.0, min_value=2.0, step=0.5, key="g3d_lz")
    n_x = st.number_input("N° Columnas en X", value=4, min_value=2, step=1, key="g3d_nx")
    n_z = st.number_input("N° Columnas en Z", value=5, min_value=2, step=1, key="g3d_nz")

    st.subheader(_t("Alzado", "Elevation"))
    n_pisos = st.number_input("Número de Pisos", value=3, min_value=1, step=1, key="g3d_npisos")
    if "alturas_df" not in st.session_state or len(st.session_state.alturas_df) != n_pisos:
        data_h = [{"Piso": p+1, "Altura (m)": 3.5 if p==0 else 3.0} for p in range(n_pisos)]
        st.session_state.alturas_df = pd.DataFrame(data_h)
    alturas_df = st.data_editor(st.session_state.alturas_df, use_container_width=True, hide_index=True)
    alturas_list = alturas_df["Altura (m)"].tolist()

    st.subheader(_t("Secciones", "Sections"))
    col_b = st.number_input("Base Columnas (m)", value=0.40, step=0.05, key="g3d_cb")
    col_h = st.number_input("Altura Columnas (m)", value=0.40, step=0.05, key="g3d_ch")
    vig_b = st.number_input("Base Vigas (m)", value=0.30, step=0.05, key="g3d_vb")
    vig_h = st.number_input("Altura Vigas (m)", value=0.40, step=0.05, key="g3d_vh")

    st.subheader(_t("Materiales y Suelo", "Materials & Soil"))
    fc = st.number_input("Resistencia concreto f'c [MPa]", value=21.0, step=1.0, key="g3d_fc")
    fy = st.number_input("Fluencia acero fy [MPa]", value=420.0, step=10.0, key="g3d_fy")
    qa = st.number_input("Capacidad admisible suelo qa [kN/m²]", value=150.0, step=10.0, key="g3d_qa")
    recub = st.number_input("Recubrimiento [cm]", value=5.0, step=0.5, key="g3d_rec")

    st.subheader(_t("Cargas (kN/m²)", "Loads (kN/m²)"))
    cm = st.number_input("Carga Muerta (CM)", value=4.5, step=0.5, key="g3d_cm")
    cv = st.number_input("Carga Viva (CV)", value=2.0, step=0.5, key="g3d_cv")

    st.subheader(_t("🌍 Parámetros Sísmicos", "🌍 Seismic Parameters"))
    # Selección de ciudad según norma
    ciudades = list(seismic_info["ciudades"].keys())
    ciudad_sel = st.selectbox(_t("Ciudad / Localidad", "City / Location"), ciudades, key="g3d_ciudad")
    perfil_sel = st.selectbox(_t("Perfil de Suelo", "Soil Profile"), list(seismic_info["suelos"].keys()), key="g3d_perfil")
    uso_sel = st.selectbox(_t("Categoría de Uso", "Occupancy Category"), list(seismic_info["uso"].keys()), key="g3d_uso")
    # Construir parámetros sísmicos según norma
    if "NSR-10" in norma_sel:
        ciudad_params = seismic_info["ciudades"][ciudad_sel]
        suelo_params = seismic_info["suelos"][perfil_sel]
        seismic_params = {
            "Aa": ciudad_params["Aa"], "Av": ciudad_params["Av"],
            "Fa": suelo_params["Fa"], "Fv": suelo_params["Fv"],
            "I": seismic_info["uso"][uso_sel],
            "R": seismic_info["R"]
        }
    elif "E.060" in norma_sel or "Perú" in norma_sel:
        ciudad_params = seismic_info["ciudades"][ciudad_sel]
        suelo_params = seismic_info["suelos"][perfil_sel]
        seismic_params = {
            "Z": ciudad_params["Z"],
            "S": suelo_params["S"], "Tp": suelo_params["Tp"], "Tl": suelo_params["Tl"],
            "U": seismic_info["uso"][uso_sel],
            "R": seismic_info["R"]
        }
    else:
        # Para otras normas, usar valores por defecto
        seismic_params = {"S_DS": 0.6, "S_D1": 0.4, "TL": 4.0, "R": 6.0, "U": 1.0}
    st.caption(f"R = {seismic_params.get('R', 6.0)} (coeficiente de reducción sísmica)")

    if st.button("🏗️ Generar / Actualizar Malla 3D", type="primary", use_container_width=True):
        nudos, cols, vigx, vigz, zaps = generate_mesh(L_x, L_z, n_x, n_z, alturas_list, col_b, col_h, vig_b, vig_h)
        st.session_state.update({
            "g3d_nudos": nudos, "g3d_cols": cols, "g3d_vx": vigx, "g3d_vz": vigz, "g3d_zaps": zaps,
            "g3d_ready": True, "g3d_design": None
        })

# Cuerpo principal
if st.session_state.get("g3d_ready", False):
    nudos = st.session_state.g3d_nudos
    cols = st.session_state.g3d_cols
    vigas_x = st.session_state.g3d_vx
    vigas_z = st.session_state.g3d_vz
    zaps = st.session_state.g3d_zaps

    col_vis, col_data = st.columns([1.2, 1])
    with col_vis:
        # Visualización inicial (sin dimensiones de zapatas aún)
        fig = plot_edificio(nudos, cols, vigas_x, vigas_z, zaps, None)
        st.plotly_chart(fig, use_container_width=True, height=600)

    with col_data:
        if st.button("📐 Ejecutar Diseño Estructural", type="primary", use_container_width=True):
            area_planta = (max(nudos['X']) - min(nudos['X'])) * (max(nudos['Z']) - min(nudos['Z']))
            df_cols, df_beams, df_zaps, vol_losa, peso_acero_losa, h_losa, As_losa, Sa, T_aprox, V_basal, V_piso = design_all_elements(
                nudos, cols, vigas_x, vigas_z, zaps,
                fc, fy, qa, recub, alturas_list, area_planta,
                cm, cv, seismic_params, norma_sel
            )
            st.session_state.g3d_design = {
                "columnas": df_cols,
                "vigas": df_beams,
                "zapatas": df_zaps,
                "vol_losa": vol_losa,
                "peso_acero_losa": peso_acero_losa,
                "h_losa": h_losa,
                "As_losa": As_losa,
                "Sa": Sa,
                "T_aprox": T_aprox,
                "V_basal": V_basal,
                "V_piso": V_piso,
                "seismic_params": seismic_params
            }
            st.success("✅ Diseño completado.")
            st.rerun()

        if st.session_state.get("g3d_design") is not None:
            des = st.session_state.g3d_design
            # Recalcular totales (los pesos ya están correctos)
            vol_conc = des["columnas"]["Vol (m³)"].sum() + des["vigas"]["Vol (m³)"].sum() + des["zapatas"]["Vol (m³)"].sum() + des["vol_losa"]
            peso_acero = des["columnas"]["Peso_acero (kg)"].sum() + des["vigas"]["Peso_acero (kg)"].sum() + des["zapatas"]["Peso_acero (kg)"].sum() + des["peso_acero_losa"]
            st.subheader("Resumen de materiales")
            col1, col2, col3 = st.columns(3)
            col1.metric("Volumen total concreto", f"{vol_conc:.2f} m³")
            col2.metric("Peso total acero", f"{peso_acero:.1f} kg")
            col3.metric("Cuantía de acero", f"{peso_acero/vol_conc:.1f} kg/m³" if vol_conc>0 else "N/A")

            # Mostrar información sísmica
            with st.expander("📊 Resultados del análisis sísmico"):
                st.write(f"Período fundamental aproximado: T = {des['T_aprox']:.3f} s")
                st.write(f"Aceleración espectral: Sa = {des['Sa']:.3f} g")
                st.write(f"Cortante basal: V = {des['V_basal']:.1f} kN")
                st.write("Fuerzas laterales por piso (kN):")
                for i, v in enumerate(des['V_piso']):
                    st.write(f"Piso {i+1}: {v:.1f} kN")

            tabs = st.tabs(["Columnas", "Vigas", "Zapatas", "Losa"])
            with tabs[0]:
                st.dataframe(des["columnas"], use_container_width=True)
            with tabs[1]:
                st.dataframe(des["vigas"], use_container_width=True)
            with tabs[2]:
                st.dataframe(des["zapatas"], use_container_width=True)
            with tabs[3]:
                st.write(f"Espesor losa: {des['h_losa']:.2f} m")
                st.write(f"Acero principal: {des['As_losa']:.2f} cm²/m")
                st.write(f"Acero temperatura: {max(0.0018*100*des['h_losa']*100, 0):.2f} cm²/m")

            # Exportaciones
            st.markdown("---")
            st.subheader("📥 Exportar")
            col_e1, col_e2, col_e3 = st.columns(3)
            with col_e1:
                if st.button("📊 Exportar a Excel"):
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        des["columnas"].to_excel(writer, sheet_name="Columnas", index=False)
                        des["vigas"].to_excel(writer, sheet_name="Vigas", index=False)
                        des["zapatas"].to_excel(writer, sheet_name="Zapatas", index=False)
                        pd.DataFrame({"Item": ["Volumen concreto", "Peso acero"], "Valor": [vol_conc, peso_acero]}).to_excel(writer, sheet_name="Resumen", index=False)
                    st.download_button("📥 Descargar Excel", data=output.getvalue(), file_name="Generador_3D_Resultados.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col_e2:
                if st.button("📄 Memoria DOCX"):
                    doc = Document()
                    doc.add_heading(f"Memoria de Cálculo - Edificio Paramétrico ({norma_sel})", 0)
                    doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
                    doc.add_paragraph(f"Materiales: f'c={fc} MPa, fy={fy} MPa, qa={qa} kN/m²")
                    doc.add_paragraph(f"Geometría: Lx={L_x} m, Lz={L_z} m, {n_pisos} pisos, altura media {sum(alturas_list)/len(alturas_list):.1f} m")
                    doc.add_heading("Análisis Sísmico", level=1)
                    doc.add_paragraph(f"Período fundamental T = {des['T_aprox']:.3f} s")
                    doc.add_paragraph(f"Aceleración espectral Sa = {des['Sa']:.3f} g")
                    doc.add_paragraph(f"Cortante basal V = {des['V_basal']:.1f} kN")
                    doc.add_paragraph("Fuerzas laterales por piso (kN): " + ", ".join([f"{v:.1f}" for v in des['V_piso']]))
                    doc.add_heading("Cantidades de obra", level=1)
                    doc.add_paragraph(f"Concreto total: {vol_conc:.2f} m³")
                    doc.add_paragraph(f"Acero total: {peso_acero:.1f} kg")
                    doc.add_heading("Resumen de diseño", level=1)
                    doc.add_paragraph("Columnas (primeras 5):")
                    table_col = doc.add_table(rows=1+min(5, len(des["columnas"])), cols=6)
                    table_col.style = 'Table Grid'
                    hdr = table_col.rows[0].cells
                    hdr[0].text = "ID"; hdr[1].text = "b (cm)"; hdr[2].text = "h (cm)"; hdr[3].text = "As (cm²)"; hdr[4].text = "Cumple"; hdr[5].text = "Observación"
                    for i, row in des["columnas"].head(5).iterrows():
                        cells = table_col.rows[i+1].cells
                        cells[0].text = str(row["ID"])
                        cells[1].text = f"{row['b (cm)']:.0f}"
                        cells[2].text = f"{row['h (cm)']:.0f}"
                        cells[3].text = f"{row['As (cm²)']:.1f}"
                        cells[4].text = "✅" if row["Cumple"] else "❌"
                        cells[5].text = row["Observación"]
                    doc.add_paragraph("Vigas (primeras 5):")
                    table_vig = doc.add_table(rows=1+min(5, len(des["vigas"])), cols=6)
                    table_vig.style = 'Table Grid'
                    hdr = table_vig.rows[0].cells
                    hdr[0].text = "ID"; hdr[1].text = "b (cm)"; hdr[2].text = "h (cm)"; hdr[3].text = "As (cm²)"; hdr[4].text = "s_estribos (cm)"; hdr[5].text = "L (m)"
                    for i, row in des["vigas"].head(5).iterrows():
                        cells = table_vig.rows[i+1].cells
                        cells[0].text = str(row["ID"])
                        cells[1].text = f"{row['b (cm)']:.0f}"
                        cells[2].text = f"{row['h (cm)']:.0f}"
                        cells[3].text = f"{row['As (cm²)']:.1f}"
                        cells[4].text = f"{row['s_estribos (cm)']:.0f}" if row["s_estribos (cm)"]>0 else "No requiere"
                        cells[5].text = f"{row['L (m)']:.2f}"
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button("📥 Descargar Memoria DOCX", data=buf, file_name="Memoria_Generador_3D.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col_e3:
                if st.button("🏗️ Exportar DXF"):
                    try:
                        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                                 dxf_rotulo, dxf_rotulo_campos)
                        _USE_H16 = True
                    except ImportError:
                        _USE_H16 = False
                    dwg = ezdxf.new('R2010')
                    dwg.units = ezdxf.units.M
                    if _USE_H16:
                        dxf_setup(dwg, 100); dxf_add_layers(dwg)
                    msp = dwg.modelspace()
                    dwg.layers.add(name="COLUMNAS", color=5)
                    dwg.layers.add(name="VIGAS", color=3)
                    dwg.layers.add(name="ZAPATAS", color=1)
                    for _, col16 in cols.iterrows():
                        n1 = nudos[nudos['ID']==col16['N1']].iloc[0]
                        n2 = nudos[nudos['ID']==col16['N2']].iloc[0]
                        msp.add_line((n1['X'], n1['Z'], n1['Y']), (n2['X'], n2['Z'], n2['Y']), dxfattribs={'layer':'COLUMNAS'})
                    for _, v16 in pd.concat([vigas_x, vigas_z]).iterrows():
                        n1 = nudos[nudos['ID']==v16['N1']].iloc[0]
                        n2 = nudos[nudos['ID']==v16['N2']].iloc[0]
                        msp.add_line((n1['X'], n1['Z'], n1['Y']), (n2['X'], n2['Z'], n2['Y']), dxfattribs={'layer':'VIGAS'})
                    for i16, zap16 in zaps.iterrows():
                        B16 = des["zapatas"].iloc[i16]['B (m)'] if i16 < len(des["zapatas"]) else 1.2
                        xc16 = zap16['X']; zc16 = zap16['Z']
                        msp.add_lwpolyline([(xc16-B16/2, zc16-B16/2),(xc16+B16/2, zc16-B16/2),(xc16+B16/2, zc16+B16/2),(xc16-B16/2, zc16+B16/2),(xc16-B16/2, zc16-B16/2)], dxfattribs={'layer':'ZAPATAS'})
                    if _USE_H16:
                        _x0_16 = float(nudos['X'].min()); _z0_16 = float(nudos['Z'].min())
                        _xw_16 = float(nudos['X'].max()) - _x0_16
                        dxf_text(msp, _x0_16, _z0_16-2, f"Edificio {n_pisos}P – {n_x}x{n_z} cols – {norma_sel}", "EJES", h=0.025*100, ha="left")
                        _cam16 = dxf_rotulo_campos(f"Generador Maestro 3D – {n_pisos} pisos", norma_sel, "001")
                        dxf_rotulo(msp, _cam16, _x0_16, _z0_16-8, rot_w=max(_xw_16,10), rot_h=4, escala=100)
                    out16 = io.StringIO()
                    dwg.write(out16)
                    st.download_button("📥 Descargar DXF", data=out16.getvalue(), file_name="Edificio_3D.dxf", mime="application/dxf")

            # APU (presupuesto)
            if "apu_config" in st.session_state:
                apu = st.session_state.apu_config
                mon = apu.get("moneda", "$")
                # Costos concreto
                if apu.get("usar_concreto_premezclado", False):
                    costo_conc = vol_conc * apu.get("precio_concreto_m3", 0)
                else:
                    # Estimación con mezcla en sitio (350 kg/m³)
                    costo_conc = vol_conc * 350 / 50 * apu.get("cemento", 0)
                costo_ace = peso_acero * apu.get("acero", 0)
                total_mat = costo_conc + costo_ace
                # Mano de obra aproximada
                dias_mo = (peso_acero * 0.04) + (vol_conc * 0.4)
                costo_mo = dias_mo * apu.get("costo_dia_mo", 69333.33)
                costo_directo = total_mat + costo_mo
                herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
                aiu = costo_directo * apu.get("pct_aui", 0.30)
                utilidad = costo_directo * apu.get("pct_util", 0.05)
                iva = utilidad * apu.get("iva", 0.19)
                total = costo_directo + herramienta + aiu + iva

                st.markdown("---")
                st.subheader("💰 Presupuesto estimado")
                st.write(f"**Volumen concreto:** {vol_conc:.2f} m³")
                st.write(f"**Acero:** {peso_acero:.1f} kg")
                st.write(f"**Costo materiales:** {mon} {total_mat:,.2f}")
                st.write(f"**Costo total (incl. MO e indirectos):** {mon} {total:,.2f}")
            else:
                st.info("💡 Ve a la página 'APU Mercado' para configurar precios.")
else:
    st.info("👈 Configure los parámetros en la barra lateral y presione **Generar Malla 3D**.")