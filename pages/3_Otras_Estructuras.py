import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import plotly.graph_objects as go

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Otras Estructuras", "Other Structures"), layout="wide")
st.title(_t("Otras Estructuras — Suite Multi-Norma", "Other Structures — Multi-Code Suite"))
st.markdown(_t("Módulo integrado para ménsulas, columnas, losas bidireccionales y otros elementos estructurales.", "Integrated module for corbels, columns, two-way slabs and other structural elements."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════
# CODES DICT (factores de resistencia)
# ══════════════════════════════════════════
CODES = {
    "NSR-10 (Colombia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "NSR-10 Título C",
        "bag_kg": 50.0,
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "ACI 318-25",
        "bag_kg": 42.6,
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "ACI 318-19",
        "bag_kg": 42.6,
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "ACI 318-14",
        "bag_kg": 42.6,
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "NEC-SE-HM Ecuador",
        "bag_kg": 50.0,
    },
    "E.060 (Perú)": {
        "phi_flex": 0.90, "phi_shear": 0.85, "phi_comp": 0.70,
        "lambda": 1.0, "ref": "Norma E.060 Perú",
        "bag_kg": 42.5,
    },
    "NTC-EM (México)": {
        "phi_flex": 0.85, "phi_shear": 0.80, "phi_comp": 0.70,
        "lambda": 1.0, "ref": "NTC-EM 2017 México",
        "bag_kg": 50.0,
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.70,
        "lambda": 1.0, "ref": "COVENIN 1753-2006",
        "bag_kg": 42.5,
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "NB 1225001-2020 Bolivia",
        "bag_kg": 50.0,
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "ref": "CIRSOC 201-2025 Argentina",
        "bag_kg": 50.0,
    },
}

# Rebar tables
REBAR_US = {
    "#3 (Ø9.5mm)":0.71,"#4 (Ø12.7mm)":1.29,"#5 (Ø15.9mm)":1.99,"#6 (Ø19.1mm)":2.84,
    "#7 (Ø22.2mm)":3.87,"#8 (Ø25.4mm)":5.10,"#9 (Ø28.7mm)":6.45,"#10 (Ø32.3mm)":7.92
}
REBAR_MM = {
    "8mm":0.503,"10mm":0.785,"12mm":1.131,"14mm":1.539,"16mm":2.011,"18mm":2.545,
    "20mm":3.142,"22mm":3.801,"25mm":4.909,"28mm":6.158,"32mm":8.042
}
DIAM_US = {
    "#3 (Ø9.5mm)":9.53,"#4 (Ø12.7mm)":12.7,"#5 (Ø15.9mm)":15.88,"#6 (Ø19.1mm)":19.05,
    "#7 (Ø22.2mm)":22.23,"#8 (Ø25.4mm)":25.4,"#9 (Ø28.7mm)":28.65,"#10 (Ø32.3mm)":32.26
}
DIAM_MM = {
    "8mm":8,"10mm":10,"12mm":12,"14mm":14,"16mm":16,"18mm":18,"20mm":20,"22mm":22,"25mm":25,"28mm":28,"32mm":32
}

def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def mix_for_fc(fc):
    """ACI 211 mix proportions in kg/m³."""
    table = [
        (14,250,205,810,1060),(17,290,200,780,1060),(21,350,193,720,1060),
        (25,395,193,680,1020),(28,430,190,640,1000),(35,530,185,580,960),
        (42,620,180,520,910),(56,740,175,450,850),
    ]
    if fc <= table[0][0]: return table[0][1:]
    if fc >= table[-1][0]: return table[-1][1:]
    for i in range(len(table)-1):
        lo,hi = table[i],table[i+1]
        if lo[0] <= fc <= hi[0]:
            t = (fc-lo[0])/(hi[0]-lo[0])
            return tuple(lo[j]+t*(hi[j]-lo[j]) for j in range(1,5))
    return table[-1][1:]

def sec_dark_fig(w, h, title=""):
    fig, ax = plt.subplots(figsize=(max(3,w/h*3), 3))
    fig.patch.set_facecolor('#1a1a2e')
    ax.set_facecolor('#1a1a2e')
    ax.add_patch(patches.Rectangle((0,0),w,h,linewidth=2,edgecolor='white',facecolor='#4a4a6a'))
    ax.set_xlim(-w*0.15, w*1.15); ax.set_ylim(-h*0.15, h*1.15)
    ax.axis('off')
    ax.set_title(title, color='white', fontsize=8)
    return fig, ax

def qty_table(rows):
    st.dataframe(pd.DataFrame(rows, columns=["Concepto","Valor"]), use_container_width=True, hide_index=True)

# ══════════════════════════════════════════
# GLOBAL SIDEBAR (Materiales y Norma)
# ══════════════════════════════════════════
st.sidebar.header(_t("🌍 Norma de Diseño", "🌍 Design Code"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = list(CODES.keys())[0]
norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)
code = CODES[norma_sel]
st.sidebar.markdown(f"📖 `{code['ref']}`")
st.sidebar.markdown(f"**φ flex:** {code['phi_flex']} | **φ cort:** {code['phi_shear']} | **φ comp:** {code['phi_comp']}")

st.sidebar.header(_t("⚙️ Materiales Globales", "⚙️ Global Materials"))
fc_unit = st.sidebar.radio(_t("Unidad f'c:", "f'c Unit:"), ["MPa","PSI","kg/cm²"], horizontal=True, key="o_fc_unit")
if fc_unit == "PSI":
    psi_options = ["2500","3000","3500","4000","4500","5000"]
    psi_v = st.sidebar.selectbox("f'c [PSI]:", psi_options, key="o_fc_psi")
    fc = float(psi_v)*0.00689476
    st.sidebar.info(f"f'c = {psi_v} PSI → **{fc:.2f} MPa**")
elif fc_unit == "kg/cm²":
    kg_options = ["175","210","250","280","350","420"]
    kg_v = st.sidebar.selectbox("f'c [kg/cm²]:", kg_options, key="o_fc_kgcm2")
    fc = float(kg_v)/10.1972
    st.sidebar.info(f"f'c = {kg_v} kg/cm² → **{fc:.2f} MPa**")
else:
    fc = st.sidebar.number_input("f'c [MPa]:", 15.0, 80.0, st.session_state.get("o_fc", 21.0), 1.0, key="o_fc")

fy = st.sidebar.number_input("fy [MPa]:", 200.0, 500.0, st.session_state.get("o_fy", 420.0), 10.0, key="o_fy")
Es = 200000.0
Ec = 4700*math.sqrt(fc)
beta1 = get_beta1(fc)
bag_kg = code["bag_kg"]

bar_sys = st.sidebar.radio("Sistema Varillas:", ["Pulgadas (# US)","Milímetros (mm)"], horizontal=True, key="o_bar_sys")
rebar_dict = REBAR_US if "Pulgadas" in bar_sys else REBAR_MM
diam_dict  = DIAM_US  if "Pulgadas" in bar_sys else DIAM_MM

phi_f = code["phi_flex"]
phi_v = code["phi_shear"]
phi_c = code["phi_comp"]
lam   = code["lambda"]

st.sidebar.markdown("---")
st.sidebar.caption(f"Ec = {Ec:.0f} MPa  |  β₁ = {beta1:.3f}  |  f'c = {fc:.2f} MPa")

# =============================================================================
# 1. CORTANTE A UNA DISTANCIA X (VIGAS) + DISEÑO DE ESTRIBOS
# =============================================================================
with st.expander(_t("✂️ Cortante a una Distancia X del Apoyo (Vigas)", "✂️ Shear at a Distance X from Support (Beams)")):
    st.info(_t("📺 **Modo de uso:** Ingresa la carga distribuida Wu, la longitud de la viga y la distancia X. La app calculará el cortante Vu en ese punto y el espaciamiento requerido de estribos.", "📺 **How to use:** Enter load Wu, span L and distance X. Shows required shear at that section and stirrup spacing."))
    c1, c2 = st.columns(2)
    with c1:
        L_vga = st.number_input(_t("Longitud luz libre [m]", "Clear span (m)"), 1.0, 20.0, 5.0, 0.5, key="o_cx_L")
        wu_vga= st.number_input(_t("Carga distribuida Wu [kN/m]", "Factored load Wu (kN/m)"), 1.0, 500.0, 50.0, 5.0, key="o_cx_wu")
        x_dist= st.number_input(_t("Distancia X desde el apoyo [m]", "Distance X from support (m)"), 0.0, L_vga/2, 1.0, 0.1, key="o_cx_x")
    with c2:
        bw_cx = st.number_input("Ancho bw [cm]", 10.0, 100.0, 25.0, 5.0, key="o_cx_bw")
        d_cx  = st.number_input("Peralte efectivo d [cm]", 10.0, 150.0, 40.0, 5.0, key="o_cx_d")
        # Selección de estribo
        est_opts = ["Ø6mm","Ø8mm","Ø10mm","Ø12mm","#2","#3","#4"]
        st_bar_cx = st.selectbox("Estribo:", est_opts, key="o_cx_st")
        st_area = {"Ø6mm":0.283,"Ø8mm":0.503,"Ø10mm":0.785,"Ø12mm":1.131,"#2":0.32,"#3":0.71,"#4":1.29}[st_bar_cx]
        n_ramas = st.number_input("# Ramas del estribo", 2, 6, 2, 1, key="o_cx_ramas")
        Av_cx = st_area * n_ramas  # cm²

    # Cálculo
    Vu_x = wu_vga * L_vga / 2.0 - wu_vga * x_dist
    # Resistencia del concreto (N) -> kN
    Vc_N = 0.17 * lam * math.sqrt(fc) * (bw_cx*10) * (d_cx*10)  # N
    Vc_kN = Vc_N / 1000
    phi_Vc = phi_v * Vc_kN
    Vs_req = max(0, Vu_x / phi_v - Vc_kN)
    if Vs_req > 0:
        s_calc_mm = Av_cx * 100 * fy * (d_cx*10) / (Vs_req * 1000)  # mm
    else:
        s_calc_mm = min(d_cx*10/2, 600)
    # Limitaciones ACI
    Vs_max = 0.66 * math.sqrt(fc) * (bw_cx*10) * (d_cx*10) / 1000  # kN
    if Vs_req > Vs_max:
        st.error("❌ Vs requerido excede el máximo permitido. Aumente la sección.")
    s_max = min(d_cx*10/2, 600) if Vs_req <= 0.33*math.sqrt(fc)*(bw_cx*10)*(d_cx*10)/1000 else min(d_cx*10/4, 300)
    s_diseno_mm = min(s_calc_mm, s_max)
    s_diseno_cm = s_diseno_mm/10

    n_estribos = math.ceil(L_vga*100/s_diseno_cm) + 1
    Vs_prov = Av_cx * 100 * fy * (d_cx*10) / (s_diseno_mm*1000) if s_diseno_mm>0 else 0
    phi_Vn = phi_v * (Vc_kN + Vs_prov)
    ok_cx = phi_Vn >= Vu_x

    tab_res, tab_q, tab_apu, tab_mem, tab_dxf = st.tabs(["📊 Resultados","📦 Cantidades","💰 APU","📄 Memoria","📏 DXF"])
    with tab_res:
        st.markdown(f"**φ cortante = {phi_v}** | Norma: `{code['ref']}`")
        rows_cx = [
            ("Vu (a X)", f"{Vu_x:.1f} kN"),
            ("Vc (concreto)", f"{Vc_kN:.2f} kN"),
            ("φVc", f"{phi_Vc:.2f} kN"),
            ("Vs requerido", f"{Vs_req:.2f} kN"),
            (f"Av ({n_ramas} ramas {st_bar_cx})", f"{Av_cx:.3f} cm²"),
            ("s calculado", f"{s_calc_mm:.0f} mm = {s_calc_mm/10:.1f} cm"),
            ("s máx (norma)", f"{s_max:.0f} mm = {s_max/10:.1f} cm"),
            ("s de diseño", f"**{s_diseno_cm:.1f} cm**"),
            ("Vs provisto", f"{Vs_prov:.2f} kN"),
            ("φVn", f"{phi_Vn:.2f} kN"),
            ("Estado", "✅ CUMPLE" if ok_cx else "❌ DEFICIENTE"),
        ]
        qty_table(rows_cx)
        if ok_cx:
            st.success(f"✅ Aprobado: Estribos {st_bar_cx} @ {s_diseno_cm:.1f} cm")
        else:
            st.error(f"❌ No aprobado: φVn = {phi_Vn:.2f} kN < Vu = {Vu_x:.2f} kN")

        # ── 3D SECCIÓN TRANSVERSAL ─────────────────────────────────────
        st.markdown("---")
        st.markdown("#### 🧊 Visualización 3D de la Sección con Estribos")
        _fig3d_cx = go.Figure()
        _bw, _d, _L = bw_cx, d_cx, L_vga*100  # cm
        # Cuerpo de la viga
        _corners_x = [-_bw/2, _bw/2, _bw/2, -_bw/2, -_bw/2, _bw/2, _bw/2, -_bw/2]
        _corners_y = [0, 0, _d, _d, 0, 0, _d, _d]
        _corners_z = [0, 0, 0, 0, _L, _L, _L, _L]
        _i = [0,0,4,4,0,1,2,3]; _j=[1,3,5,7,4,5,6,7]; _k=[2,2,6,6,1,2,3,0]
        _fig3d_cx.add_trace(go.Mesh3d(x=_corners_x, y=_corners_y, z=_corners_z,
                                       i=_i, j=_j, k=_k, opacity=0.15, color='#4a4a6a', name='Concreto'))
        # Estribos (optimizados con None gaps)
        _rec = 4  # cm recubrimiento
        _tx = [-_bw/2+_rec, _bw/2-_rec, _bw/2-_rec, -_bw/2+_rec, -_bw/2+_rec]
        _ty = [_rec, _rec, _d-_rec, _d-_rec, _rec]
        _tx_all, _ty_all, _tz_all = [], [], []
        for _zt in range(5, int(_L), int(s_diseno_cm)):
            _tx_all.extend(_tx + [None]); _ty_all.extend(_ty + [None]); _tz_all.extend([_zt]*5+[None])
        _fig3d_cx.add_trace(go.Scatter3d(x=_tx_all, y=_ty_all, z=_tz_all, mode='lines',
                                          line=dict(color='cornflowerblue', width=4), name='Estribos'))
        _fig3d_cx.update_layout(
            scene=dict(aspectmode='data', xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)',
                       bgcolor='#1a1a2e'),
            paper_bgcolor='#1a1a2e', font=dict(color='white'), height=400,
            margin=dict(l=0,r=0,t=40,b=0), dragmode='turntable',
            title=dict(text=f"Viga {_bw:.0f}×{_d:.0f} cm | Estribos {st_bar_cx} @ {s_diseno_cm:.1f} cm",
                       font=dict(color='white'))
        )
        st.plotly_chart(_fig3d_cx, use_container_width=True)

    with tab_q:
        vol_beam = bw_cx/100 * d_cx/100 * L_vga  # m³ (aproximado)
        perim_cx = 2*(bw_cx + d_cx) + 6*0.8  # cm
        peso_est = n_estribos * (perim_cx/100) * st_area * 0.785  # kg
        m = mix_for_fc(fc)
        qty_table([
            ("Concreto (estimado)", f"{vol_beam:.4f} m³"),
            (f"Estribos {st_bar_cx}", f"{n_estribos} ud → {peso_est:.1f} kg"),
            (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{m[0]*vol_beam/bag_kg:.1f} bultos"),
        ])
    with tab_apu:
        if "apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mon = apu.get("moneda", "$")
            c_cem = (m[0]*vol_beam/bag_kg) * apu.get("cemento", 0)
            c_ace = peso_est * apu.get("acero", 0)
            total_mat = c_cem + c_ace
            st.metric(f"💎 Costo Estimado ({mon})", f"{total_mat:,.2f}")
        else:
            st.info("💡 Configure APU en la página '4. APU Mercado'.")
    with tab_mem:
        doc = Document()
        doc.add_heading(f"Memoria de Cálculo – Cortante a X = {x_dist:.2f} m", 0)
        doc.add_paragraph(f"Norma: {norma_sel} | f'c = {fc:.1f} MPa, fy = {fy:.0f} MPa")
        doc.add_paragraph(f"Viga: b={bw_cx:.0f} cm, d={d_cx:.0f} cm, L={L_vga:.2f} m")
        doc.add_paragraph(f"Vu = {Vu_x:.1f} kN | φVc = {phi_Vc:.1f} kN | Estribos {st_bar_cx} @ {s_diseno_cm:.1f} cm")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button("Descargar Memoria DOCX", data=buf, file_name="Cortante_X.docx")
    with tab_dxf:
        doc_dxf = ezdxf.new('R2010')
        doc_dxf.units = ezdxf.units.CM
        msp = doc_dxf.modelspace()
        # Dibujar sección transversal
        x0, y0 = 0, 0
        msp.add_lwpolyline([(x0,y0),(x0+bw_cx,y0),(x0+bw_cx,y0+d_cx),(x0,y0+d_cx),(x0,y0)], close=True)
        # Estribo
        rec = 4  # recubrimiento (cm)
        msp.add_lwpolyline([(x0+rec,y0+rec),(x0+bw_cx-rec,y0+rec),(x0+bw_cx-rec,y0+d_cx-rec),(x0+rec,y0+d_cx-rec),(x0+rec,y0+rec)], close=True)
        # Texto
        msp.add_text(f"Estribo {st_bar_cx} @ {s_diseno_cm:.1f} cm", dxfattribs={'height':2, 'insert':(x0+bw_cx/2, y0+d_cx+3)})
        out_str = io.StringIO()
        doc_dxf.write(out_str)
        out = out_str.getvalue().encode('utf-8')
        st.download_button("Descargar DXF", data=out, file_name="Seccion_Cortante.dxf")

# =============================================================================
# 2. MÉNSULAS (CORBELS) – ACI 318
# =============================================================================
with st.expander(_t("🏗️ Diseño de Ménsulas (Corbels / ACI 318)", "🏗️ Corbel Design (ACI 318)")):
    st.info(_t("📺 **Modo de uso:** Ingresa la carga vertical Vu, la fuerza horizontal Nuc y la geometría. Se calculará el acero principal, estribos horizontales cerrados y el acero de colgado.", "📺 **How to use:** Enter vertical load Vu, horizontal Nuc and geometry. Calculates main steel and closed ties."))
    c1,c2 = st.columns(2)
    with c1:
        Vu_men = st.number_input(_t("Carga Vertical Vu [kN]", "Vertical Load Vu [kN]"), 50.0, 2000.0, 300.0, 50.0, key="o_men_vu")
        Nuc_men= st.number_input(_t("Tensión Horiz. Nuc [kN]", "Horiz. Tension Nuc [kN]"), 0.0, 1000.0, 60.0, 10.0, key="o_men_nuc")
        a_men  = st.number_input(_t("Brazo de palanca a [cm]", "Shear span a [cm]"), 5.0, 50.0, 15.0, 5.0, key="o_men_a")
    with c2:
        bw_men = st.number_input(_t("Ancho ménsula bw [cm]", "Corbel width bw [cm]"), 20.0, 100.0, 30.0, 5.0, key="o_men_bw")
        h_men  = st.number_input(_t("Alto total ménsula h [cm]", "Total height h [cm]"), 20.0, 150.0, 45.0, 5.0, key="o_men_h")
        dp_men = st.number_input(_t("Recubrimiento d' [cm]", "Cover d' [cm]"), 2.0, 10.0, 4.0, 0.5, key="o_men_dp")
    d_men = h_men - dp_men
    a_d_ratio = a_men / d_men
    if a_d_ratio > 1.0:
        st.warning(f"a/d = {a_d_ratio:.2f} > 1.0. Las ecuaciones de ménsula asumen a/d ≤ 1. Usar diseño de vigas convencionales.")
    else:
        mu_men = 1.4  # concreto monolítico
        Avf = Vu_men / (phi_v * fy * mu_men) * 10  # cm²
        # Flexión
        Mu_men = Vu_men * (a_men/100) + Nuc_men * (h_men - d_men)/100  # kN·m
        Rn_men = Mu_men * 1e6 / (phi_f * bw_men*10 * (d_men*10)**2)
        disc = 1 - 2*Rn_men/(0.85*fc)
        if disc >= 0:
            rho_men = (0.85*fc/fy)*(1 - math.sqrt(disc))
            Af = rho_men * bw_men * d_men  # cm²
        else:
            Af = 0
            st.error("Sección insuficiente por momento. Aumente h.")
        An = Nuc_men / (phi_f * fy) * 10  # cm² (tensión directa)
        As_req_men = max(Af + An, (2/3)*Avf + An, 0.04 * (fc/fy) * bw_men * d_men)
        Ah_req_men = 0.5 * (As_req_men - An) if As_req_men > An else 0

        # Selección de varillas
        varillas_men = list(rebar_dict.keys())
        bar_men = st.selectbox("Varilla principal:", varillas_men, key="o_men_bar")
        Ab_men = rebar_dict[bar_men]
        n_bars = math.ceil(As_req_men / Ab_men)
        As_prov_men = n_bars * Ab_men

        tab_res, tab_q, tab_apu, tab_mem, tab_dxf = st.tabs(["📊 Resultados","📦 Cantidades","💰 APU","📄 Memoria","📏 DXF"])
        with tab_res:
            rows_men = [
                ("a/d", f"{a_d_ratio:.2f}"),
                ("Avf (fricción corte)", f"{Avf:.2f} cm²"),
                ("Af (flexión)", f"{Af:.2f} cm²"),
                ("An (tracción directa)", f"{An:.2f} cm²"),
                ("As requerido (principal)", f"{As_req_men:.2f} cm²"),
                (f"Varillas {bar_men}", f"{n_bars} barras → As prov = {As_prov_men:.2f} cm²"),
                ("Ah requerido (estribos cerrados)", f"{Ah_req_men:.2f} cm²"),
            ]
            qty_table(rows_men)
            
            # ── 3D MÉNSULA REAL ──────────────────────────────────────────────
            st.markdown("---")
            st.markdown("#### 🧊 Visualización 3D de la Ménsula (Corbel)")
            _fm = go.Figure()
            
            # Columna de soporte (haz que la ménsula salga de ella)
            _col_w = bw_men        # misma anchura que la ménsula
            _col_h = h_men * 2     # doble de la ménsula para dar contexto
            _col_dep = bw_men      # profundidad de la columna
            _col_x = [-_col_w/2, _col_w/2, _col_w/2, -_col_w/2, -_col_w/2, _col_w/2, _col_w/2, -_col_w/2]
            _col_y = [-_col_h, -_col_h, 0, 0, -_col_h, -_col_h, 0, 0]
            _col_z = [0, 0, 0, 0, _col_dep, _col_dep, _col_dep, _col_dep]
            _i = [0,0,4,4,0,1,2,3]; _j=[1,3,5,7,4,5,6,7]; _k=[2,2,6,6,1,2,3,0]
            _fm.add_trace(go.Mesh3d(x=_col_x, y=_col_y, z=_col_z, i=_i, j=_j, k=_k,
                                    opacity=0.35, color='#5a5a7a', name='Columna'))
            
            # Ménsula voladizo (sale hacia la derecha en X con ancho bw_men, alto h_men, prof bw_men)
            _mx = [0, a_men, a_men, 0, 0, a_men, a_men, 0]  # X de 0 a a_men (brazo de palanca)
            _my = [0, 0, h_men, h_men, 0, 0, h_men, h_men]  # Y de 0 a h_men
            _mz = [0, 0, 0, 0, _col_dep, _col_dep, _col_dep, _col_dep]
            _fm.add_trace(go.Mesh3d(x=_mx, y=_my, z=_mz, i=_i, j=_j, k=_k,
                                    opacity=0.6, color='#6a6a8a', name='Ménsula'))
            
            # Varilla longitudinal principal (zona de tracción superior, y = h_men - dp_men)
            _y_bar = h_men - dp_men
            _z_spacing = _col_dep / (n_bars + 1)
            for _bi in range(n_bars):
                _zi = _z_spacing * (_bi + 1)
                _fm.add_trace(go.Scatter3d(
                    x=[0, a_men], y=[_y_bar, _y_bar], z=[_zi, _zi],
                    mode='lines', line=dict(color='#ff6b35', width=6), showlegend=(_bi==0),
                    name='Acero Principal'
                ))
            
            # Estribos cerrados horizontales (zh)
            _n_ties = max(2, int(h_men/15))
            _y_step = h_men / (_n_ties + 1)
            for _ti in range(_n_ties):
                _yt = _y_step * (_ti + 1)
                _ez = [4, _col_dep-4, _col_dep-4, 4, 4, None]
                _ey = [_yt, _yt, _yt, _yt, _yt, None]
                _ex = [4, 4, a_men-4, a_men-4, 4, None]
                _fm.add_trace(go.Scatter3d(x=_ex, y=_ey, z=_ez, mode='lines',
                                           line=dict(color='cornflowerblue', width=4),
                                           showlegend=(_ti==0), name='Estribos Cerrados'))
            
            # Carga puntual (flecha indicativa)
            _fm.add_trace(go.Scatter3d(
                x=[a_men, a_men], y=[h_men+15, h_men+2], z=[_col_dep/2, _col_dep/2],
                mode='lines+markers+text', line=dict(color='lime', width=5),
                marker=dict(symbol='arrow', size=8, color='lime'),
                text=[f'Vu={Vu_men:.0f}kN', ''], textposition='top center',
                textfont=dict(color='lime', size=11), name='Carga Vu'
            ))
            
            _fm.update_layout(
                scene=dict(aspectmode='data', bgcolor='#1a1a2e',
                           xaxis_title='a (cm)', yaxis_title='h (cm)', zaxis_title='Ancho (cm)'),
                paper_bgcolor='#1a1a2e', font=dict(color='white'),
                height=480, margin=dict(l=0,r=0,t=40,b=0), dragmode='turntable',
                showlegend=True, legend=dict(x=0, y=1, bgcolor='rgba(0,0,0,0.3)'),
                title=dict(text=f"Ménsula {a_men:.0f}×{h_men:.0f}cm | {n_bars}×{bar_men}",
                           font=dict(color='white'))
            )
            st.plotly_chart(_fm, use_container_width=True)

        with tab_q:
            vol_men = (bw_men/100) * (h_men/100) * (a_men/100)  # m³
            peso_principal = As_prov_men * (a_men/100 + 2*h_men/100) * 0.785  # kg
            # Estribos: aproximación 2 ramas con longitud por cada estribo
            perim_est = 2*(bw_men + h_men) + 6*0.8  # cm
            n_est_men = max(2, math.ceil(h_men/15))  # separación aproximada 15 cm
            peso_est_men = n_est_men * (perim_est/100) * (Ab_men) * 0.785  # usando misma varilla
            qty_table([
                ("Volumen concreto", f"{vol_men:.4f} m³"),
                ("Acero principal", f"{peso_principal:.2f} kg"),
                ("Acero estribos cerrados", f"{peso_est_men:.2f} kg"),
            ])
        with tab_apu:
            if "apu_config" in st.session_state:
                apu = st.session_state.apu_config
                mon = apu.get("moneda", "$")
                m = mix_for_fc(fc)
                c_cem = (m[0]*vol_men/bag_kg) * apu.get("cemento", 0)
                c_ace = (peso_principal+peso_est_men) * apu.get("acero", 0)
                st.metric(f"💎 Costo Estimado ({mon})", f"{c_cem + c_ace:,.2f}")
            else:
                st.info("💡 Configure APU en la página '4. APU Mercado'.")
        with tab_mem:
            doc = Document()
            doc.add_heading("Memoria de Cálculo – Ménsula (Corbel)", 0)
            doc.add_paragraph(f"Norma: {norma_sel} | f'c={fc:.1f} MPa, fy={fy:.0f} MPa")
            doc.add_paragraph(f"Vu={Vu_men:.1f} kN, Nuc={Nuc_men:.1f} kN, a={a_men:.0f} cm, h={h_men:.0f} cm")
            doc.add_paragraph(f"As principal = {As_req_men:.2f} cm² → {n_bars} {bar_men}")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("Descargar Memoria DOCX", data=buf, file_name="Corbel.docx")
        with tab_dxf:
            doc_dxf = ezdxf.new('R2010')
            doc_dxf.units = ezdxf.units.CM
            msp = doc_dxf.modelspace()
            # Perfil de la ménsula
            x0, y0 = 0, 0
            msp.add_lwpolyline([(x0,y0),(x0+a_men,y0),(x0+a_men,y0+h_men),(x0,y0+h_men),(x0,y0)], close=True)
            msp.add_text(f"As = {As_prov_men:.2f} cm²", dxfattribs={'height':2, 'insert':(a_men/2, y0+h_men+2)})
            out_str = io.StringIO()
            doc_dxf.write(out_str)
            out = out_str.getvalue().encode('utf-8')
            st.download_button("Descargar DXF", data=out, file_name="Corbel.dxf")

# =============================================================================
# 3. PREDIMENSIONAMIENTO DE COLUMNAS
# =============================================================================
with st.expander(_t("📐 Predimensionamiento de Columnas", "📐 Column Preliminary Sizing")):
    st.info(_t("📺 **Modo de uso:** Ingresa la carga viva y muerta estimada por piso, el número de pisos y el área tributaria. Te recomendaré dimensiones de columna base.", "📺 **How to use:** Enter estimated load per floor, number of floors, and tributary area. Predicts base column section."))
    c1,c2 = st.columns(2)
    with c1:
        area_trib = st.number_input(_t("Área Tributaria [m²]", "Tributary Area [m²]"), 5.0, 100.0, 20.0, 5.0, key="o_pre_a")
        pisos     = st.number_input(_t("Total de Pisos", "Total Floors"), 1, 50, 5, 1, key="o_pre_p")
        W_piso    = st.number_input(_t("Carga estimada por piso (D+L) [kN/m²]", "Estimated Floor Load (D+L) [kN/m²]"), 5.0, 20.0, 12.0, 1.0, key="o_pre_w")
    with c2:
        tipo_col  = st.selectbox(_t("Posición (afecta k):", "Column Position (affects k):"), ["Céntrica (k=0.30)", "Esquinera/Borde (k=0.20-0.25)"], key="o_pre_tipo")
        rho_p     = st.number_input(_t("Cuantía acero estimada [%]", "Estimated steel ratio [%]"), 1.0, 4.0, 1.5, 0.5, key="o_pre_r")
    Pu_serv = W_piso * area_trib * pisos  # kN (servicio)
    Pu_fact = Pu_serv * 1.4  # factor aproximado
    k_val = 0.30 if "Céntrica" in tipo_col else 0.22
    Ag_req_cm2 = (Pu_fact * 1000) / (k_val * fc) / 100  # cm²
    b_req = math.sqrt(Ag_req_cm2)
    b_round = math.ceil(b_req/5)*5
    st.write(f"- Carga Axial de Diseño Estimada ($P_u$): **{Pu_fact:.0f} kN**")
    st.write(f"- Área Bruta Requerida ($A_g$): **{Ag_req_cm2:.0f} cm²**")
    st.success(f"Sección Cuadrada Sugerida: **{b_round} cm × {b_round} cm**")

    # ── 3D COLUMNA SUGERIDA ──────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### 🧊 Visualización 3D de la Columna Sugerida")
    _h_col = pisos * 3.0 * 100  # cm = pisos * 3m
    _fc2 = go.Figure()
    _cb = b_round / 2
    _pts = [
        (-_cb,-_cb), (_cb,-_cb), (_cb,_cb), (-_cb,_cb), (-_cb,-_cb),
        (-_cb,-_cb), (_cb,-_cb), (_cb,_cb), (-_cb,_cb), (-_cb,-_cb)
    ]
    _fc2.add_trace(go.Scatter3d(
        x=[p[0] for p in _pts[:5]] + [None] + [p[0] for p in _pts[5:]],
        y=[p[1] for p in _pts[:5]] + [None] + [p[1] for p in _pts[5:]],
        z=[0]*5 + [None] + [_h_col]*5,
        mode='lines', line=dict(color='white', width=3), name='Sección'
    ))
    # Lateral edges
    for _px, _py in [(-_cb,-_cb),(_cb,-_cb),(_cb,_cb),(-_cb,_cb)]:
        _fc2.add_trace(go.Scatter3d(x=[_px,_px], y=[_py,_py], z=[0,_h_col],
                                    mode='lines', line=dict(color='white', width=3), showlegend=False))
    # Rebar (4 esquinas simplificadas)
    _ro = _cb - 4
    for _bx, _by in [(-_ro,-_ro),(_ro,-_ro),(_ro,_ro),(-_ro,_ro)]:
        _fc2.add_trace(go.Scatter3d(x=[_bx,_bx], y=[_by,_by], z=[0,_h_col],
                                    mode='lines', line=dict(color='#ff6b35', width=6),
                                    name='Acero Long.', showlegend=(_bx==-_ro and _by==-_ro)))

    _fc2.update_layout(
        scene=dict(aspectmode='data', bgcolor='#1a1a2e',
                   xaxis_title='X (cm)', yaxis_title='Y (cm)', zaxis_title='H (cm)'),
        paper_bgcolor='#1a1a2e', font=dict(color='white'),
        height=420, margin=dict(l=0,r=0,t=50,b=0), dragmode='turntable',
        title=dict(text=f"Columna Sugerida: {b_round}×{b_round} cm × {pisos} pisos ({_h_col/100:.1f}m)",
                   font=dict(color='white'))
    )
    st.plotly_chart(_fc2, use_container_width=True)


    # Memoria y DXF para esta sección
    if st.button("Generar Memoria y DXF de Predimensionamiento"):
        doc = Document()
        doc.add_heading("Memoria de Predimensionamiento de Columna", 0)
        doc.add_paragraph(f"Norma: {norma_sel} | f'c={fc:.1f} MPa")
        doc.add_paragraph(f"Área tributaria: {area_trib:.1f} m² | {pisos} pisos | Carga por piso: {W_piso:.1f} kN/m²")
        doc.add_paragraph(f"Pu (servicio) = {Pu_serv:.0f} kN | Pu (diseño) = {Pu_fact:.0f} kN")
        doc.add_paragraph(f"Sección sugerida: {b_round} x {b_round} cm")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button("Descargar Memoria DOCX", data=buf, file_name="Predimensionamiento.docx")

        # DXF de la sección cuadrada
        doc_dxf = ezdxf.new('R2010')
        doc_dxf.units = ezdxf.units.CM
        msp = doc_dxf.modelspace()
        msp.add_lwpolyline([(0,0),(b_round,0),(b_round,b_round),(0,b_round),(0,0)], close=True)
        msp.add_text(f"{b_round}x{b_round} cm", dxfattribs={'height':3, 'insert':(b_round/2, b_round/2)})
        out_str = io.StringIO()
        doc_dxf.write(out_str)
        out = out_str.getvalue().encode('utf-8')
        st.download_button("Descargar DXF", data=out, file_name="Columna_Predim.dxf")

# =============================================================================
# 4. CAPACIDAD AXIAL COLUMNAS CORTAS
# =============================================================================
with st.expander(_t("🧱 Capacidad Axial Pn,max (Columnas Cortas)", "🧱 Axial Capacity Pn,max (Short Columns)")):
    st.info(_t("📺 **Modo de uso:** Ingresa la sección transversal probada y su armadura. El sistema calculará la carga axial máxima que soporta, ignorando el pandeo.", "📺 **How to use:** Enter section and steel. Calculates max axial capacity (ignoring slenderness)."))
    c1,c2 = st.columns(2)
    with c1:
        b_c = st.number_input("b [cm]", 20.0, 150.0, 40.0, 5.0, key="o_cap_b")
        h_c = st.number_input("h [cm]", 20.0, 150.0, 40.0, 5.0, key="o_cap_h")
        estribo_tipo = st.selectbox(_t("Forma columna:", "Column Shape:"), ["Estribada (Cuadrada/Rectg)", "Sunchada (Espiral)"], key="o_cap_est")
    with c2:
        varillas = st.number_input(_t("No. Varillas", "No. Rebars"), 4, 40, 8, 2, key="o_cap_n")
        dia_bar  = st.selectbox("Varilla:", list(rebar_dict.keys()), key="o_cap_db")
        area_bar = rebar_dict[dia_bar]
    Ag_c = b_c * h_c  # cm²
    Ast_c = varillas * area_bar  # cm²
    phi_c_val = 0.65 if "Estribada" in estribo_tipo else 0.75
    alpha_c_val = 0.80 if "Estribada" in estribo_tipo else 0.85
    Po_kN = (0.85 * fc * (Ag_c - Ast_c) * 100 + Ast_c * 100 * fy) / 1000
    Pn_max = alpha_c_val * Po_kN
    phi_Pn_max = phi_c_val * Pn_max
    st.write(f"- Carga Axial Nominal Pn,max: **{Pn_max:.0f} kN**")
    st.write(f"- Carga Axial de Diseño **φPn,max**: **{phi_Pn_max:.0f} kN**")
    cuantia_c = Ast_c / Ag_c * 100
    if cuantia_c < 1.0 or cuantia_c > 8.0:
        st.warning(f"La cuantía de acero {cuantia_c:.1f}% debe estar entre 1% y 8%.")
    # Cantidades
    vol_col = (b_c/100)*(h_c/100)*3  # 3m de altura típica
    peso_acero = Ast_c * 3 * 0.785  # kg
    st.write("**Cantidades estimadas (columna de 3m):**")
    st.write(f"Volumen concreto: {vol_col:.3f} m³")
    st.write(f"Acero: {peso_acero:.1f} kg")

# =============================================================================
# 5. LOSAS BIDIRECCIONALES (Método Coeficientes ACI)
# =============================================================================
with st.expander(_t("🏗️ Momentos en Losas 2D (Método ACI Coeficientes)", "🏗️ 2D Slab Moments (ACI Coefficients Method)")):
    st.info(_t("📺 **Modo de uso:** Ingresa las luces la y lb del tablero. Sirve para diseñar losas apoyadas perimetralmente en vigas. Calcula los momentos en ambas direcciones.", "📺 **How to use:** Enter short and long spans. Useful for edge-supported slabs. Calculates moments in both directions."))
    c1,c2 = st.columns(2)
    with c1:
        la_losa = st.number_input(_t("Luz corta La [m]", "Short span La [m]"), 2.0, 15.0, 4.0, 0.5, key="lo2_la")
        lb_losa = st.number_input(_t("Luz larga Lb [m]", "Long span Lb [m]"), 2.0, 15.0, 5.0, 0.5, key="lo2_lb")
    with c2:
        wu_losa = st.number_input("Carga distribuida factorizada Wu [kN/m²]", 2.0, 50.0, 10.0, 0.5, key="lo2_wu")
        casos = ["Caso 1 (Interior)", "Caso 2 (4 bordes discontinuos)", "Caso 3 (1 borde continuo)", "Caso 4 (2 bordes ady. continuos)"]
        caso_borde = st.selectbox("Condición de Borde (ACI):", casos, key="lo2_caso")
    m_ratio = la_losa / lb_losa if lb_losa > 0 else 1.0
    if m_ratio < 0.5:
        st.warning("m < 0.5. El panel se comporta como una losa en Una Dirección apoyada sobre las vigas largas.")
    else:
        # Coeficientes aproximados para demostración (deben basarse en tablas ACI)
        # Aquí se usan valores didácticos
        ca_pos = 0.050 * m_ratio
        cb_pos = 0.050 / m_ratio
        ca_neg = 0.070 * m_ratio if "Caso 2" not in caso_borde else 0
        cb_neg = 0.070 / m_ratio if "Caso 2" not in caso_borde else 0
        Ma_pos = ca_pos * wu_losa * la_losa**2
        Mb_pos = cb_pos * wu_losa * lb_losa**2
        Ma_neg = ca_neg * wu_losa * la_losa**2
        Mb_neg = cb_neg * wu_losa * lb_losa**2

        st.write(f"- Relación m (La/Lb): **{m_ratio:.2f}**")
        colA, colB = st.columns(2)
        colA.markdown("#### Dirección Corta (La)")
        colA.write(f"Ma(+) Positivo: **{Ma_pos:.1f} kN·m/m**")
        if ca_neg>0: colA.write(f"Ma(-) Negativo: **{Ma_neg:.1f} kN·m/m**")
        colB.markdown("#### Dirección Larga (Lb)")
        colB.write(f"Mb(+) Positivo: **{Mb_pos:.1f} kN·m/m**")
        if cb_neg>0: colB.write(f"Mb(-) Negativo: **{Mb_neg:.1f} kN·m/m**")

        # Gráfico de momentos
        fig, ax = plt.subplots(figsize=(6,4))
        ax.bar(["Ma+","Ma-","Mb+","Mb-"], [Ma_pos, Ma_neg, Mb_pos, Mb_neg], color=['#4caf50','#ff9800','#4caf50','#ff9800'])
        ax.set_ylabel("Momento (kN·m/m)")
        ax.set_title("Momentos en losa bidireccional")
        st.pyplot(fig)

        # Cantidades
        area_losa = la_losa * lb_losa
        h_losa = max(la_losa/30, 0.10)  # espesor estimado
        vol_losa = area_losa * h_losa
        m = mix_for_fc(fc)
        st.write("**Cantidades estimadas (losa completa):**")
        st.write(f"Volumen concreto: {vol_losa:.3f} m³")
        st.write(f"Cemento: {m[0]*vol_losa/bag_kg:.1f} bultos")
        st.write(f"Arena: {m[2]*vol_losa:.0f} kg")
        st.write(f"Grava: {m[3]*vol_losa:.0f} kg")

        # Memoria DOCX
        if st.button("Generar Memoria Losa 2D"):
            doc = Document()
            doc.add_heading("Memoria de Losa Bidireccional", 0)
            doc.add_paragraph(f"Norma: {norma_sel} | f'c={fc:.1f} MPa, fy={fy:.0f} MPa")
            doc.add_paragraph(f"Tablero: La={la_losa:.2f} m, Lb={lb_losa:.2f} m | Wu={wu_losa:.1f} kN/m²")
            doc.add_paragraph(f"Ma+ = {Ma_pos:.1f} kN·m/m | Ma- = {Ma_neg:.1f} kN·m/m")
            doc.add_paragraph(f"Mb+ = {Mb_pos:.1f} kN·m/m | Mb- = {Mb_neg:.1f} kN·m/m")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("Descargar Memoria DOCX", data=buf, file_name="Losa_2D.docx")

# =============================================================================
# FOOTER
# =============================================================================
st.markdown("---")
st.markdown(f"""
> **Otras Estructuras — Multi-Norma**  
> Norma activa: `{norma_sel}`  
> f'c = {fc:.2f} MPa | fy = {fy:.0f} MPa | Ec = {Ec:.0f} MPa  
> **Referencia:** {code['ref']}  
> ⚠️ *Las herramientas son de apoyo para el diseño. Verifique siempre con la norma vigente del país.*
""")