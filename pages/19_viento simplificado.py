import streamlit as st
import pandas as pd
import numpy as np
import math
import io
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches
import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL (compartido con la app)
# ─────────────────────────────────────────────
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es

st.set_page_config(page_title=_t("Carga de Viento", "Wind Load"), layout="wide")
st.title(_t("🌬️ Carga de Viento – Método Simplificado NSR-10", "🌬️ Wind Load – Simplified Method NSR-10"))
st.markdown(_t(
    "Diseño de presiones de viento para edificaciones bajas (h ≤ 18 m) con cubiertas inclinadas ≤ 45°.\n"
    "Sigue los lineamientos de la **NSR-10, Capítulo B.6** (procedimiento simplificado).",
    "Wind pressure design for low-rise buildings (h ≤ 18 m) with roofs slope ≤ 45°.\n"
    "Based on **NSR-10, Chapter B.6** (simplified procedure)."
))

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES (basadas en NSR-10)
# ─────────────────────────────────────────────
# Tabla de factores de ajuste λ (exposición B, C, D) – interpolación lineal
LAMBDA_TABLE = {
    "B": {  # Exposición B
        4.5: 0.84, 6.0: 0.91, 7.5: 0.96, 9.0: 1.00, 10.5: 1.04,
        12.0: 1.07, 13.5: 1.10, 15.0: 1.12, 16.5: 1.15, 18.0: 1.17
    },
    "C": {  # Exposición C
        4.5: 1.21, 6.0: 1.29, 7.5: 1.35, 9.0: 1.40, 10.5: 1.45,
        12.0: 1.49, 13.5: 1.53, 15.0: 1.56, 16.5: 1.59, 18.0: 1.62
    },
    "D": {  # Exposición D
        4.5: 1.47, 6.0: 1.55, 7.5: 1.61, 9.0: 1.66, 10.5: 1.70,
        12.0: 1.74, 13.5: 1.78, 15.0: 1.81, 16.5: 1.84, 18.0: 1.87
    }
}

def get_lambda_factor(expo, h):
    """
    Devuelve el factor de ajuste λ según exposición y altura media h (m).
    Interpola linealmente entre los valores tabulados.
    """
    if expo not in LAMBDA_TABLE:
        return 1.0
    table = LAMBDA_TABLE[expo]
    heights = sorted(table.keys())
    if h <= heights[0]:
        return table[heights[0]]
    if h >= heights[-1]:
        return table[heights[-1]]
    for i in range(len(heights)-1):
        if heights[i] <= h <= heights[i+1]:
            t = (h - heights[i]) / (heights[i+1] - heights[i])
            return table[heights[i]] + t*(table[heights[i+1]] - table[heights[i]])
    return 1.0

# Coeficientes ps10 (SPRFV) de la Figura B.6.4-2 (para muros y cubiertas)
PS10_SPRFV = {
    "Zona A":  {0:0.50, 5:0.48, 10:0.46, 15:0.44, 20:0.42, 25:0.40, 30:0.38, 35:0.36, 40:0.34, 45:0.32},
    "Zona B":  {0:0.40, 5:0.38, 10:0.36, 15:0.34, 20:0.32, 25:0.30, 30:0.28, 35:0.26, 40:0.24, 45:0.22},
    "Zona C":  {0:0.33, 5:0.32, 10:0.31, 15:0.30, 20:0.29, 25:0.28, 30:0.27, 35:0.26, 40:0.25, 45:0.24},
    "Zona D":  {0:0.25, 5:0.24, 10:0.23, 15:0.22, 20:0.21, 25:0.20, 30:0.19, 35:0.18, 40:0.17, 45:0.16},
    "Zona E":  {0:-0.60, 5:-0.58, 10:-0.56, 15:-0.54, 20:-0.52, 25:-0.50, 30:-0.48, 35:-0.46, 40:-0.44, 45:-0.42},
    "Zona F":  {0:-0.34, 5:-0.33, 10:-0.32, 15:-0.31, 20:-0.30, 25:-0.29, 30:-0.28, 35:-0.27, 40:-0.26, 45:-0.25},
    "Zona G":  {0:-0.41, 5:-0.40, 10:-0.39, 15:-0.38, 20:-0.37, 25:-0.36, 30:-0.35, 35:-0.34, 40:-0.33, 45:-0.32},
    "Zona H":  {0:-0.26, 5:-0.25, 10:-0.24, 15:-0.23, 20:-0.22, 25:-0.21, 30:-0.20, 35:-0.19, 40:-0.18, 45:-0.17},
}

def get_ps10_SPRFV(zona, theta):
    """Obtiene el coeficiente ps10 (kN/m²) para una zona y ángulo θ (grados)."""
    if zona not in PS10_SPRFV:
        return 0.0
    table = PS10_SPRFV[zona]
    angles = sorted(table.keys())
    if theta <= angles[0]:
        return table[angles[0]]
    if theta >= angles[-1]:
        return table[angles[-1]]
    for i in range(len(angles)-1):
        if angles[i] <= theta <= angles[i+1]:
            t = (theta - angles[i]) / (angles[i+1] - angles[i])
            return table[angles[i]] + t*(table[angles[i+1]] - table[angles[i]])
    return 0.0

# Coeficientes pnet10 (C&R) de la Figura B.6.4-3 (para zonas 1 a 5)
PNET10_CR = {
    "Zona1": {
        "pos": {0:0.05, 5:0.05, 10:0.04, 15:0.04, 20:0.04, 25:0.04, 30:0.04, 35:0.04, 40:0.04, 45:0.04},
        "neg": {0:-0.12, 5:-0.12, 10:-0.11, 15:-0.11, 20:-0.10, 25:-0.10, 30:-0.10, 35:-0.09, 40:-0.09, 45:-0.09}
    },
    "Zona2": {
        "pos": {0:0.04, 5:0.04, 10:0.04, 15:0.04, 20:0.04, 25:0.04, 30:0.04, 35:0.04, 40:0.04, 45:0.04},
        "neg": {0:-0.13, 5:-0.13, 10:-0.12, 15:-0.12, 20:-0.11, 25:-0.11, 30:-0.11, 35:-0.10, 40:-0.10, 45:-0.10}
    },
    "Zona3": {
        "pos": {0:0.05, 5:0.05, 10:0.05, 15:0.05, 20:0.05, 25:0.05, 30:0.05, 35:0.05, 40:0.05, 45:0.05},
        "neg": {0:-0.30, 5:-0.29, 10:-0.28, 15:-0.27, 20:-0.26, 25:-0.25, 30:-0.24, 35:-0.23, 40:-0.22, 45:-0.21}
    },
    "Zona4": {
        "pos": {0:0.10, 5:0.10, 10:0.09, 15:0.09, 20:0.09, 25:0.09, 30:0.08, 35:0.08, 40:0.08, 45:0.08},
        "neg": {0:-0.11, 5:-0.11, 10:-0.10, 15:-0.10, 20:-0.09, 25:-0.09, 30:-0.09, 35:-0.08, 40:-0.08, 45:-0.08}
    },
    "Zona5": {
        "pos": {0:0.11, 5:0.11, 10:0.10, 15:0.10, 20:0.10, 25:0.10, 30:0.09, 35:0.09, 40:0.09, 45:0.09},
        "neg": {0:-0.14, 5:-0.14, 10:-0.13, 15:-0.13, 20:-0.12, 25:-0.12, 30:-0.12, 35:-0.11, 40:-0.11, 45:-0.11}
    }
}

def get_pnet10_CR(zona, theta, area_efectiva):
    """
    Obtiene el coeficiente pnet10 (kN/m²) para una zona (1-5), ángulo θ y área efectiva (m²).
    Ajuste simplificado por área efectiva.
    """
    if zona not in PNET10_CR:
        return 0.0, 0.0
    tab = PNET10_CR[zona]
    # Interpolación para θ
    angles = sorted(tab["pos"].keys())
    if theta <= angles[0]:
        pos = tab["pos"][angles[0]]
        neg = tab["neg"][angles[0]]
    elif theta >= angles[-1]:
        pos = tab["pos"][angles[-1]]
        neg = tab["neg"][angles[-1]]
    else:
        for i in range(len(angles)-1):
            if angles[i] <= theta <= angles[i+1]:
                t = (theta - angles[i]) / (angles[i+1] - angles[i])
                pos = tab["pos"][angles[i]] + t*(tab["pos"][angles[i+1]] - tab["pos"][angles[i]])
                neg = tab["neg"][angles[i]] + t*(tab["neg"][angles[i+1]] - tab["neg"][angles[i]])
                break
    # Ajuste por área efectiva (simplificado: para áreas > 10 m², se reduce un 20%)
    if area_efectiva > 10:
        factor = 0.8
        pos *= factor
        neg *= factor
    return pos, neg

# ─────────────────────────────────────────────
# FUNCIONES DE DIBUJO ESQUEMÁTICO (reemplazan las imágenes externas)
# ─────────────────────────────────────────────
def draw_figure_b642(theta_deg, presiones):
    """
    Dibuja un esquema de la Figura B.6.4-2 (SPRFV) mostrando las zonas A-H.
    Muestra además los valores de presión calculados.
    """
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_facecolor('#f0f0f0')
    ax.set_xlim(-1, 10)
    ax.set_ylim(-1, 7)
    ax.axis('off')

    # Dibujar el perfil del edificio
    w = 8
    h_muro = 3
    h_cumbrera = 5 if theta_deg > 0 else h_muro
    # Muros
    ax.add_patch(patches.Rectangle((0, 0), w, h_muro, edgecolor='black', facecolor='lightgray', alpha=0.5))
    # Cubierta
    if theta_deg > 0:
        ax.plot([0, w/2, w], [h_muro, h_cumbrera, h_muro], 'b-', linewidth=2)
    else:
        ax.plot([0, w], [h_muro, h_muro], 'b-', linewidth=2)

    # Anotar zonas según NSR-10 (posiciones aproximadas)
    # Zona A: barlovento, parte inferior del muro
    ax.annotate(f"A\n{presiones['A']:.2f}", xy=(w*0.1, h_muro*0.5), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona B: sotavento, parte inferior del muro
    ax.annotate(f"B\n{presiones['B']:.2f}", xy=(w*0.9, h_muro*0.5), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona C: interior del muro (se puede indicar con texto)
    ax.annotate(f"C\n{presiones['C']:.2f}", xy=(w*0.5, h_muro*0.5), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona D: interior de cubierta
    ax.annotate(f"D\n{presiones['D']:.2f}", xy=(w*0.5, h_muro+0.5), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona E: final de cubierta barlovento
    ax.annotate(f"E\n{presiones['E']:.2f}", xy=(w*0.2, h_muro+1.2), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona F: final de cubierta sotavento
    ax.annotate(f"F\n{presiones['F']:.2f}", xy=(w*0.8, h_muro+1.2), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona G: interior cubierta barlovento
    ax.annotate(f"G\n{presiones['G']:.2f}", xy=(w*0.4, h_muro+0.8), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))
    # Zona H: interior cubierta sotavento
    ax.annotate(f"H\n{presiones['H']:.2f}", xy=(w*0.6, h_muro+0.8), ha='center', fontsize=9, bbox=dict(boxstyle="round,pad=0.3", fc='white', ec='black'))

    ax.set_title(f"Figura B.6.4-2 – Zonas de presión (θ = {theta_deg:.1f}°)", fontsize=12)
    return fig

def draw_figure_b643(theta_deg, pnet_fachada, pnet_cubierta):
    """
    Dibuja un esquema de la Figura B.6.4-3 (C&R) mostrando zonas 1-5.
    """
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_facecolor('#f0f0f0')
    ax.set_xlim(-1, 10)
    ax.set_ylim(-1, 7)
    ax.axis('off')

    # Dibujar perfil
    w = 8
    h_muro = 3
    h_cumbrera = 5 if theta_deg > 0 else h_muro
    ax.add_patch(patches.Rectangle((0, 0), w, h_muro, edgecolor='black', facecolor='lightgray', alpha=0.5))
    if theta_deg > 0:
        ax.plot([0, w/2, w], [h_muro, h_cumbrera, h_muro], 'b-', linewidth=2)
    else:
        ax.plot([0, w], [h_muro, h_muro], 'b-', linewidth=2)

    # Zonas de fachada (4 y 5)
    # Zona 4: interior fachada
    ax.annotate(f"Zona 4 (interior)\n+{pnet_fachada['Interior']['+']:.2f}\n-{abs(pnet_fachada['Interior']['-']):.2f}",
                xy=(w*0.5, h_muro*0.5), ha='center', fontsize=8,
                bbox=dict(boxstyle="round,pad=0.3", fc='lightyellow', ec='black'))
    # Zona 5: final fachada (cerca del borde)
    ax.annotate(f"Zona 5 (final)\n+{pnet_fachada['Final']['+']:.2f}\n-{abs(pnet_fachada['Final']['-']):.2f}",
                xy=(w*0.1, h_muro*0.5), ha='center', fontsize=8,
                bbox=dict(boxstyle="round,pad=0.3", fc='lightyellow', ec='black'))
    # Zonas de cubierta (1,2,3)
    # Zona 1: interior
    ax.annotate(f"Zona 1 (interior)\n+{pnet_cubierta['Interior']['+']:.2f}\n-{abs(pnet_cubierta['Interior']['-']):.2f}",
                xy=(w*0.5, h_muro+0.8), ha='center', fontsize=8,
                bbox=dict(boxstyle="round,pad=0.3", fc='lightyellow', ec='black'))
    # Zona 2: final
    ax.annotate(f"Zona 2 (final)\n+{pnet_cubierta['Final']['+']:.2f}\n-{abs(pnet_cubierta['Final']['-']):.2f}",
                xy=(w*0.8, h_muro+1.2), ha='center', fontsize=8,
                bbox=dict(boxstyle="round,pad=0.3", fc='lightyellow', ec='black'))
    # Zona 3: esquina
    ax.annotate(f"Zona 3 (esquina)\n+{pnet_cubierta['Esquinas']['+']:.2f}\n-{abs(pnet_cubierta['Esquinas']['-']):.2f}",
                xy=(w*0.95, h_muro+0.5), ha='center', fontsize=8,
                bbox=dict(boxstyle="round,pad=0.3", fc='lightyellow', ec='black'))

    ax.set_title(f"Figura B.6.4-3 – Zonas para componentes y revestimientos (θ = {theta_deg:.1f}°)", fontsize=12)
    return fig

# ─────────────────────────────────────────────
# SIDEBAR – PARÁMETROS GENERALES
# ─────────────────────────────────────────────
st.sidebar.header(_t("Norma de Diseño", "Design Code"))
norma_sel = st.sidebar.selectbox(
    _t("Seleccione la norma:", "Select code:"),
    ["NSR-10 (Colombia)", "ACI 318-25 (EE.UU.)", "ACI 318-19 (EE.UU.)", "ACI 318-14 (EE.UU.)",
     "NEC-SE-HM (Ecuador)", "E.060 (Perú)", "NTC-EM (México)", "COVENIN 1753-2006 (Venezuela)",
     "NB 1225001-2020 (Bolivia)", "CIRSOC 201-2025 (Argentina)"],
    index=0
)

st.sidebar.markdown("---")
st.sidebar.header(_t("Datos de entrada", "Input data"))
V = st.sidebar.number_input(_t("Velocidad básica del viento V [m/s]", "Basic wind speed V [m/s]"), 10.0, 60.0, 36.0, 1.0)
I = st.sidebar.selectbox(_t("Factor de importancia I", "Importance factor I"), ["I (Grupo I)", "II (Grupo II)", "III (Grupo III)", "IV (Grupo IV)"], index=1)
I_val = {"I (Grupo I)":0.87, "II (Grupo II)":1.00, "III (Grupo III)":1.15, "IV (Grupo IV)":1.15}[I]
expo = st.sidebar.selectbox(_t("Categoría de exposición", "Exposure category"), ["B", "C", "D"], index=1)
hr = st.sidebar.number_input(_t("Altura de cumbrera hr [m]", "Ridge height hr [m]"), 2.0, 18.0, 4.68, 0.1)
he = st.sidebar.number_input(_t("Altura de cornisa he [m]", "Eave height he [m]"), 2.0, 18.0, 4.50, 0.1)
W = st.sidebar.number_input(_t("Ancho del edificio W [m] (perpendicular a cumbrera)", "Building width W [m] (perpendicular to ridge)"), 5.0, 100.0, 10.0, 1.0)
L = st.sidebar.number_input(_t("Longitud del edificio L [m] (paralelo a cumbrera)", "Building length L [m] (parallel to ridge)"), 5.0, 100.0, 20.0, 1.0)
tipo_cubierta = st.sidebar.selectbox(_t("Tipo de cubierta", "Roof type"), ["Plana", "1 agua", "2 aguas"], index=2)
theta = st.sidebar.number_input(_t("Ángulo de inclinación θ [°]", "Roof slope θ [°]"), 0.0, 45.0, 2.0, 0.5)
Kzt = st.sidebar.number_input(_t("Factor topográfico Kzt", "Topographic factor Kzt"), 0.8, 1.5, 1.0, 0.05)
region_huracan = st.sidebar.checkbox(_t("¿Región propensa a huracanes?", "Hurricane-prone region?"), value=False)

# C&R áreas efectivas
area_fachada = st.sidebar.number_input(_t("Área efectiva en fachada [m²]", "Effective area for walls [m²]"), 0.5, 50.0, 6.75, 1.0)
area_cubierta = st.sidebar.number_input(_t("Área efectiva en cubierta [m²]", "Effective area for roof [m²]"), 0.5, 50.0, 12.0, 1.0)
area_alero = st.sidebar.number_input(_t("Área efectiva en alero [m²]", "Effective area for eaves [m²]"), 0.0, 50.0, 0.0, 1.0)

# ─────────────────────────────────────────────
# VERIFICACIONES PREVIAS
# ─────────────────────────────────────────────
h_media = (hr + he) / 2
if h_media > 18.0:
    st.error(_t("La altura media del edificio supera los 18 m. Este método simplificado solo aplica para h ≤ 18 m.", "Mean building height exceeds 18 m. This simplified method applies only for h ≤ 18 m."))
if theta > 45.0:
    st.error(_t("El ángulo de cubierta supera los 45°. Este método solo aplica para θ ≤ 45°.", "Roof slope exceeds 45°. This method applies only for θ ≤ 45°."))
if tipo_cubierta == "2 aguas" and theta > 27.0 and region_huracan:
    st.warning(_t("En regiones propensas a huracanes, el ángulo máximo para cubiertas a dos aguas es 27° (NSR-10 B.6.4.1.2(d)).", "In hurricane-prone regions, the maximum roof slope for gable roofs is 27° (NSR-10 B.6.4.1.2(d))."))

if norma_sel != "NSR-10 (Colombia)":
    st.warning(_t("El método simplificado de viento está implementado solo para NSR-10 (Colombia). Para otras normas se recomienda utilizar métodos analíticos completos.", "The simplified wind method is implemented only for NSR-10 (Colombia). For other codes, it is recommended to use full analytical methods."))

# ─────────────────────────────────────────────
# CÁLCULOS
# ─────────────────────────────────────────────
lambda_ = get_lambda_factor(expo, h_media)

# SPRFV
zonas = ["A", "B", "C", "D", "E", "F", "G", "H"]
presiones = {}
for zona in zonas:
    ps10 = get_ps10_SPRFV(f"Zona {zona}", theta)
    ps = lambda_ * Kzt * I_val * ps10
    presiones[zona] = ps

# Aplicar presiones mínimas NSR-10 B.6.4.2.1.1
for zona in ["A", "B", "C", "D"]:
    if presiones[zona] < 0.4:
        presiones[zona] = 0.4
# Para zonas E,F,G,H no se aplica mínimo positivo, pero se asegura que no sean positivas
# (ya son negativas, pero si por interpolación dieran positivas, se anulan)
for zona in ["E", "F", "G", "H"]:
    if presiones[zona] > 0:
        presiones[zona] = 0.0

# C&R
pos4, neg4 = get_pnet10_CR("Zona4", theta, area_fachada)
pos5, neg5 = get_pnet10_CR("Zona5", theta, area_fachada)
pnet_fachada = {
    "Interior": {"+": pos4, "-": neg4},
    "Final":   {"+": pos5, "-": neg5}
}

pos1, neg1 = get_pnet10_CR("Zona1", theta, area_cubierta)
pos2, neg2 = get_pnet10_CR("Zona2", theta, area_cubierta)
pos3, neg3 = get_pnet10_CR("Zona3", theta, area_cubierta)
pnet_cubierta = {
    "Interior": {"+": pos1, "-": neg1},
    "Final":   {"+": pos2, "-": neg2},
    "Esquinas": {"+": pos3, "-": neg3}
}

# Presiones mínimas C&R (NSR-10 B.6.4.2.2.1)
for tipo in pnet_fachada.values():
    if tipo["+"] < 0.4:
        tipo["+"] = 0.4
    if abs(tipo["-"]) < 0.4:
        tipo["-"] = -0.4
for tipo in pnet_cubierta.values():
    if tipo["+"] < 0.4:
        tipo["+"] = 0.4
    if abs(tipo["-"]) < 0.4:
        tipo["-"] = -0.4

# ─────────────────────────────────────────────
# MOSTRAR RESULTADOS
# ─────────────────────────────────────────────
st.markdown("---")
st.header(_t("Resultados", "Results"))
st.subheader(_t("Sistema Principal de Resistencia al Viento (SPRFV)", "Main Wind Force Resisting System (MWFRS)"))

col1, col2 = st.columns(2)
with col1:
    st.markdown("**Presiones netas de diseño (ps) [kN/m²]**")
    df_sp = pd.DataFrame({
        "Zona": zonas,
        "ps (kN/m²)": [f"{presiones[z]:.2f}" for z in zonas]
    })
    st.dataframe(df_sp, use_container_width=True, hide_index=True)

with col2:
    # Gráfico de barras
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=zonas,
        y=[presiones[z] for z in zonas],
        name="ps",
        marker_color=['#4a6ea8' if z in ['A','B','C','D'] else '#a84a6e' for z in zonas]
    ))
    fig.update_layout(
        title="Presiones de viento (SPRFV)",
        yaxis_title="Presión (kN/m²)",
        template="plotly_dark"
    )
    st.plotly_chart(fig, use_container_width=True)

# Dibujar esquema de la Figura B.6.4-2 con valores reales
st.markdown("---")
st.subheader(_t("Esquema de zonas SPRFV (Figura B.6.4-2)", "SPRFV zones diagram (Figure B.6.4-2)"))
fig_b642 = draw_figure_b642(theta, presiones)
st.pyplot(fig_b642)

st.markdown("---")
st.subheader(_t("Componentes y Revestimientos (C&R)", "Components & Cladding (C&C)"))

st.markdown("**Fachadas**")
df_fachada = pd.DataFrame({
    "Ubicación": ["Interior", "Final"],
    "Presión positiva (+) [kN/m²]": [pnet_fachada["Interior"]["+"], pnet_fachada["Final"]["+"]],
    "Presión negativa (-) [kN/m²]": [pnet_fachada["Interior"]["-"], pnet_fachada["Final"]["-"]]
})
# Formatear solo las columnas numéricas
df_fachada["Presión positiva (+) [kN/m²]"] = df_fachada["Presión positiva (+) [kN/m²]"].map(lambda x: f"{x:.2f}")
df_fachada["Presión negativa (-) [kN/m²]"] = df_fachada["Presión negativa (-) [kN/m²]"].map(lambda x: f"{x:.2f}")
st.dataframe(df_fachada, use_container_width=True, hide_index=True)

st.markdown("**Cubierta**")
df_cubierta = pd.DataFrame({
    "Ubicación": ["Interior", "Final", "Esquinas"],
    "Presión positiva (+) [kN/m²]": [pnet_cubierta["Interior"]["+"], pnet_cubierta["Final"]["+"], pnet_cubierta["Esquinas"]["+"]],
    "Presión negativa (-) [kN/m²]": [pnet_cubierta["Interior"]["-"], pnet_cubierta["Final"]["-"], pnet_cubierta["Esquinas"]["-"]]
})
df_cubierta["Presión positiva (+) [kN/m²]"] = df_cubierta["Presión positiva (+) [kN/m²]"].map(lambda x: f"{x:.2f}")
df_cubierta["Presión negativa (-) [kN/m²]"] = df_cubierta["Presión negativa (-) [kN/m²]"].map(lambda x: f"{x:.2f}")
st.dataframe(df_cubierta, use_container_width=True, hide_index=True)

# Dibujar esquema de la Figura B.6.4-3 con valores reales
st.markdown("---")
st.subheader(_t("Esquema de zonas C&R (Figura B.6.4-3)", "C&C zones diagram (Figure B.6.4-3)"))
fig_b643 = draw_figure_b643(theta, pnet_fachada, pnet_cubierta)
st.pyplot(fig_b643)

st.caption(_t("Nota: Las presiones mostradas ya incluyen el factor de ajuste λ, Kzt, I y los mínimos exigidos por NSR-10.", "Note: Pressures already include adjustment factor λ, Kzt, I and the minimum requirements of NSR-10."))

# ─────────────────────────────────────────────
# EXPORTAR RESULTADOS (DOCX)
# ─────────────────────────────────────────────
st.markdown("---")
if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
    doc = Document()
    doc.add_heading(_t("Memoria de Cálculo – Carga de Viento", "Wind Load Calculation Report"), 0)
    doc.add_paragraph(f"Norma: {norma_sel}")
    doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_heading("1. Datos de entrada", level=1)
    doc.add_paragraph(f"Velocidad básica V = {V} m/s\n"
                      f"Factor de importancia I = {I_val:.2f} (Grupo {I.split('(')[1].split(')')[0]})\n"
                      f"Exposición = {expo}\n"
                      f"Altura media h = {h_media:.2f} m\n"
                      f"Geometría: W = {W} m, L = {L} m\n"
                      f"Cubierta: {tipo_cubierta}, θ = {theta}°\n"
                      f"Kzt = {Kzt:.2f}\n"
                      f"Región propensa a huracanes: {'Sí' if region_huracan else 'No'}\n"
                      f"Áreas efectivas: fachada = {area_fachada} m², cubierta = {area_cubierta} m², alero = {area_alero} m²")
    doc.add_heading("2. Factor de ajuste λ", level=1)
    doc.add_paragraph(f"Para exposición {expo} y altura {h_media:.2f} m → λ = {lambda_:.3f}")
    doc.add_heading("3. Presiones de diseño SPRFV", level=1)
    table = doc.add_table(rows=1+len(zonas), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Zona"; hdr[1].text = "ps (kN/m²)"
    for i, zona in enumerate(zonas):
        row = table.rows[i+1].cells
        row[0].text = zona
        row[1].text = f"{presiones[zona]:.2f}"
    doc.add_heading("4. Presiones de diseño C&R", level=1)
    doc.add_paragraph("Fachadas:")
    doc.add_paragraph(f"Interior: +{pnet_fachada['Interior']['+']:.2f} / -{abs(pnet_fachada['Interior']['-']):.2f} kN/m²")
    doc.add_paragraph(f"Final:   +{pnet_fachada['Final']['+']:.2f} / -{abs(pnet_fachada['Final']['-']):.2f} kN/m²")
    doc.add_paragraph("Cubierta:")
    doc.add_paragraph(f"Interior: +{pnet_cubierta['Interior']['+']:.2f} / -{abs(pnet_cubierta['Interior']['-']):.2f} kN/m²")
    doc.add_paragraph(f"Final:   +{pnet_cubierta['Final']['+']:.2f} / -{abs(pnet_cubierta['Final']['-']):.2f} kN/m²")
    doc.add_paragraph(f"Esquinas:+{pnet_cubierta['Esquinas']['+']:.2f} / -{abs(pnet_cubierta['Esquinas']['-']):.2f} kN/m²")
    doc.add_paragraph("Referencia: NSR-10 Capítulo B.6 (Procedimiento simplificado)")
    doc_mem = io.BytesIO()
    doc.save(doc_mem)
    doc_mem.seek(0)
    st.download_button(_t("Descargar Memoria DOCX", "Download DOCX"), data=doc_mem,
                       file_name="Memoria_Viento.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")