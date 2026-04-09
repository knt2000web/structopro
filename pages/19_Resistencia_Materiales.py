import streamlit as st
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es

st.set_page_config(page_title=_t("Resistencia de Materiales", "Mechanics of Materials"), layout="wide")
st.title(_t("Resistencia de Materiales", "Mechanics of Materials"))
st.markdown(_t("Análisis de Esfuerzos Secundarios y Propiedades Geométricas Universales.", 
               "Secondary Stress Analysis and Universal Geometric Properties."))

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES PARA EXPORTACIÓN
# ─────────────────────────────────────────────
def export_plot_as_png(fig, filename="plot.png"):
    """Convierte una figura de matplotlib a bytes PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    return buf

def export_plotly_as_png(fig):
    """Convierte una figura de plotly a bytes PNG (requiere kaleido)."""
    # Nota: kaleido puede no estar instalado; usamos una alternativa con orca?
    # Por simplicidad, usamos el enfoque de guardar la imagen a bytes.
    # Si kaleido no está, se puede pedir al usuario que descargue desde el menú de Plotly.
    # No lo implementamos aquí para evitar dependencias extra.
    return None

# =============================================================================
# TAB 1: CÍRCULO DE MOHR
# =============================================================================
tab_mohr, tab_centroide = st.tabs([
    _t("⭕ 1. Círculo de Mohr (Transf. Esfuerzos)", "⭕ 1. Mohr's Circle (Stress Transf.)"),
    _t(" 2. Centroides e Inercias (Secc. Compuestas)", " 2. Centroids & Inertias (Composite Sections)")
])

with tab_mohr:
    st.header(_t("Círculo de Mohr para Estado Plano de Esfuerzos", "Mohr's Circle for Plane Stress State"))
    
    col_m1, col_m2, col_m3 = st.columns([1, 1, 1.5])
    with col_m1:
        st.subheader(_t("Esfuerzos Iniciales", "Initial Stresses"))
        sig_x = st.number_input(r"$\sigma_x$ (tracción +, comp -)", -1000.0, 1000.0, st.session_state.get("mo_sx", 50.0), key="mo_sx")
        sig_y = st.number_input(r"$\sigma_y$ (tracción +, comp -)", -1000.0, 1000.0, st.session_state.get("mo_sy", -10.0), key="mo_sy")
        tau_xy = st.number_input(r"$\tau_{xy}$ (corte)", -1000.0, 1000.0, st.session_state.get("mo_txy", 40.0), key="mo_txy")
        u_opts = ["MPa", "psi", "ksi", "kgf/cm²"]
        unidades_m = st.selectbox(_t("Unidades", "Units"), u_opts, 
                                  index=u_opts.index(st.session_state.get("mo_unit", "MPa")),
                                  key="mo_unit")
        
    # Cálculos
    sig_avg = (sig_x + sig_y) / 2.0
    R = np.sqrt(((sig_x - sig_y)/2.0)**2 + tau_xy**2)
    sig_1 = sig_avg + R
    sig_2 = sig_avg - R
    tau_max = R
    
    if (sig_x - sig_y) != 0:
        theta_p1_rad = np.arctan2(2*tau_xy, (sig_x - sig_y)) / 2.0
    else:
        theta_p1_rad = np.pi/4 if tau_xy > 0 else -np.pi/4
    theta_p1_deg = np.degrees(theta_p1_rad)
    theta_p2_deg = theta_p1_deg + 90.0
    theta_s_deg = theta_p1_deg - 45.0
    
    with col_m2:
        st.subheader(_t("Esfuerzos Principales", "Principal Stresses"))
        st.markdown(f"**$\sigma_1$ ({_t('Máximo', 'Max')}):** {sig_1:.2f} {unidades_m}")
        st.markdown(f"**$\sigma_2$ ({_t('Mínimo', 'Min')}):** {sig_2:.2f} {unidades_m}")
        st.markdown(f"**$\tau_{{max}}$ ({_t('Corte Máx', 'Max Shear')}):** {tau_max:.2f} {unidades_m}")
        st.markdown(f"**{_t('Centro', 'Center')} ($\sigma_{{avg}}$):** {sig_avg:.2f} {unidades_m}")
        st.markdown(f"**$\theta_{{p1}}$ ({_t('Plano Principal 1', 'Principal Plane 1')}):** {theta_p1_deg:.2f}°")
        st.markdown(f"**$\theta_{{p2}}$ ({_t('Plano Principal 2', 'Principal Plane 2')}):** {theta_p2_deg:.2f}°")
        st.markdown(f"**$\theta_{{s}}$ ({_t('Corte Máx', 'Max Shear')}):** {theta_s_deg:.2f}°")
        
    with col_m3:
        # Gráfico interactivo con Plotly
        theta = np.linspace(0, 2*np.pi, 200)
        x_circle = sig_avg + R * np.cos(theta)
        y_circle = R * np.sin(theta)
        
        fig_mohr = go.Figure()
        fig_mohr.add_trace(go.Scatter(x=x_circle, y=y_circle, mode='lines', name=_t('Círculo Mohr', 'Mohr Circle'), line=dict(color='blue', width=2)))
        fig_mohr.add_hline(y=0, line_color="black", line_width=1)
        fig_mohr.add_vline(x=0, line_color="black", line_width=1)
        # Puntos del estado inicial (σ_x, -τ_xy) y (σ_y, τ_xy)
        fig_mohr.add_trace(go.Scatter(x=[sig_x, sig_y], y=[-tau_xy, tau_xy], mode='markers+lines', 
                                      name=_t('Estado Inicial', 'Initial State'), marker=dict(size=10, color='red'),
                                      line=dict(dash='dash', color='red')))
        fig_mohr.add_trace(go.Scatter(x=[sig_avg], y=[0], mode='markers', name=_t('Centro', 'Center'), marker=dict(size=8, color='black')))
        fig_mohr.add_trace(go.Scatter(x=[sig_1, sig_2], y=[0, 0], mode='markers', name=_t('Esf. Principales', 'Principal Stresses'), marker=dict(size=10, color='green')))
        fig_mohr.add_trace(go.Scatter(x=[sig_avg, sig_avg], y=[tau_max, -tau_max], mode='markers', name=_t('Tau Max', 'Max Tau'), marker=dict(size=10, color='purple')))
        
        fig_mohr.update_layout(title=_t("Círculo de Mohr Interactivo", "Interactive Mohr's Circle"), 
                               xaxis_title=rf"$\sigma$ Normal [{unidades_m}]", 
                               yaxis_title=rf"$\tau$ Cortante [{unidades_m}] ({_t('abajo + según convención', 'down + per convention')})",
                               yaxis=dict(autorange="reversed"),  # convención clásica: tau positivo hacia abajo
                               plot_bgcolor='white', hovermode="x")
        st.plotly_chart(fig_mohr, use_container_width=True)
        
        # Exportación
        st.markdown("---")
        st.subheader(_t("Exportar", "Export"))
        col_exp1, col_exp2 = st.columns(2)
        with col_exp1:
            # Exportar a DOCX
            if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
                # Crear una imagen estática del círculo (usando matplotlib)
                fig_static, ax = plt.subplots(figsize=(6, 5))
                ax.plot(x_circle, y_circle, 'b-', lw=2)
                ax.axhline(0, color='k', lw=1)
                ax.axvline(0, color='k', lw=1)
                ax.plot([sig_x, sig_y], [-tau_xy, tau_xy], 'r--', marker='o', markersize=5)
                ax.plot(sig_avg, 0, 'ko', markersize=8)
                ax.plot([sig_1, sig_2], [0, 0], 'go', markersize=8)
                ax.plot([sig_avg, sig_avg], [tau_max, -tau_max], 'mo', markersize=8)
                ax.set_xlabel(f"σ Normal [{unidades_m}]")
                ax.set_ylabel(f"τ Cortante [{unidades_m}]")
                ax.set_title("Círculo de Mohr")
                ax.grid(True, alpha=0.3)
                ax.set_aspect('equal')
                buf_img = export_plot_as_png(fig_static)
                
                doc = Document()
                doc.add_heading(_t("Círculo de Mohr - Análisis de Esfuerzos", "Mohr's Circle - Stress Analysis"), 0)
                doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
                doc.add_paragraph(_t("Datos de entrada:", "Input data:"))
                doc.add_paragraph(f"σ_x = {sig_x:.2f} {unidades_m}")
                doc.add_paragraph(f"σ_y = {sig_y:.2f} {unidades_m}")
                doc.add_paragraph(f"τ_xy = {tau_xy:.2f} {unidades_m}")
                doc.add_heading(_t("Resultados", "Results"), level=1)
                doc.add_paragraph(f"σ₁ = {sig_1:.2f} {unidades_m}")
                doc.add_paragraph(f"σ₂ = {sig_2:.2f} {unidades_m}")
                doc.add_paragraph(f"τ_max = {tau_max:.2f} {unidades_m}")
                doc.add_paragraph(f"θ_p1 = {theta_p1_deg:.2f}°")
                doc.add_paragraph(f"θ_p2 = {theta_p2_deg:.2f}°")
                doc.add_paragraph(f"θ_s = {theta_s_deg:.2f}°")
                doc.add_picture(buf_img, width=Inches(5))
                buf_doc = io.BytesIO()
                doc.save(buf_doc)
                buf_doc.seek(0)
                st.download_button(_t("Descargar Memoria DOCX", "Download DOCX Report"), data=buf_doc,
                                   file_name="Mohr_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col_exp2:
            # Exportar la figura como PNG (desde Plotly requiere kaleido, pero podemos usar matplotlib)
            if st.button(_t("Exportar Figura PNG", "Export Figure as PNG")):
                fig_static, ax = plt.subplots(figsize=(6, 5))
                ax.plot(x_circle, y_circle, 'b-', lw=2)
                ax.axhline(0, color='k', lw=1); ax.axvline(0, color='k', lw=1)
                ax.plot([sig_x, sig_y], [-tau_xy, tau_xy], 'r--', marker='o', markersize=5)
                ax.plot(sig_avg, 0, 'ko', markersize=8)
                ax.plot([sig_1, sig_2], [0, 0], 'go', markersize=8)
                ax.plot([sig_avg, sig_avg], [tau_max, -tau_max], 'mo', markersize=8)
                ax.set_xlabel(f"σ Normal [{unidades_m}]")
                ax.set_ylabel(f"τ Cortante [{unidades_m}]")
                ax.set_title("Círculo de Mohr")
                ax.grid(True, alpha=0.3)
                ax.set_aspect('equal')
                buf_img = export_plot_as_png(fig_static)
                st.download_button(_t("Descargar PNG", "Download PNG"), data=buf_img,
                                   file_name="Mohr_Circle.png", mime="image/png")

# =============================================================================
# TAB 2: CENTROIDES E INERCIAS
# =============================================================================
with tab_centroide:
    st.header(_t("Cálculo de Propiedades de Secciones Compuestas", "Composite Section Properties Calculation"))
    st.write(_t("Utilice el sistema global de coordenadas. Añada figuras simples para ensamblar su sección (puede usar anchos o altos negativos para restar áreas vacías).", 
                "Use the global coordinate system. Add simple primitive shapes to build your section (use negative width/height to subtract voids)."))
    
    if "shapes_list" not in st.session_state:
        st.session_state.shapes_list = []
        
    col_s1, col_s2 = st.columns([1.5, 2])
    
    with col_s1:
        st.subheader(_t("Agregar Figura", "Add Shape"))
        tipo_opts = ["Rectángulo", "Círculo", "Triángulo Rectángulo"]
        s_tipo = st.selectbox(_t("Forma", "Shape"), tipo_opts, 
                              index=tipo_opts.index(st.session_state.get("ge_tipo", "Rectángulo")),
                              key="ge_tipo")
        s_x = st.number_input(_t("Posición origen local X_0", "Local origin X_0"), -1000.0, 1000.0, st.session_state.get("ge_x", 0.0), 1.0, key="ge_x")
        s_y = st.number_input(_t("Posición origen local Y_0", "Local origin Y_0"), -1000.0, 1000.0, st.session_state.get("ge_y", 0.0), 1.0, key="ge_y")
        
        if s_tipo == "Rectángulo":
            s_b = st.number_input(_t("Base b (+ sólido, - hueco)", "Base b (+ solid, - void)"), -1000.0, 1000.0, st.session_state.get("ge_b", 10.0), 1.0, key="ge_b")
            s_h = st.number_input(_t("Altura h (+ sólido, - hueco)", "Height h (+ solid, - void)"), -1000.0, 1000.0, st.session_state.get("ge_h", 20.0), 1.0, key="ge_h")
            s_prop = {"tipo": "Rect", "x": s_x, "y": s_y, "b": s_b, "h": s_h}
        elif s_tipo == "Círculo":
            s_r = st.number_input(_t("Radio r (+ sólido, - radio hueco)", "Radius r (+ solid, - void)"), -1000.0, 1000.0, st.session_state.get("ge_r", 10.0), 1.0, key="ge_r")
            s_prop = {"tipo": "Circ", "x": s_x, "y": s_y, "r": s_r}
        else:
            s_b_t = st.number_input(_t("Base Triángulo b", "Triangle base b"), -1000.0, 1000.0, st.session_state.get("ge_bt", 10.0), 1.0, key="ge_bt")
            s_h_t = st.number_input(_t("Altura Triángulo h", "Triangle height h"), -1000.0, 1000.0, st.session_state.get("ge_ht", 10.0), 1.0, key="ge_ht")
            s_prop = {"tipo": "Triang", "x": s_x, "y": s_y, "b": s_b_t, "h": s_h_t}
            
        if st.button(_t("Añadir Figura", "Add Shape")):
            st.session_state.shapes_list.append(s_prop)
            st.rerun()
            
        if st.button(_t("Limpiar Todo", "Clear All")):
            st.session_state.shapes_list = []
            st.rerun()
            
    with col_s2:
        st.subheader(_t("Conjunto y Resultados", "Assembly and Results"))
        if not st.session_state.shapes_list:
            st.info(_t("Añade figuras a la lista para comenzar a calcular.", "Add shapes to the list to start calculations."))
        else:
            # Mostrar lista con opción de eliminar individualmente
            df_shapes = []
            sum_A = 0.0
            sum_Ax = 0.0
            sum_Ay = 0.0
            shape_list_for_delete = []
            for idx, s in enumerate(st.session_state.shapes_list):
                if s["tipo"] == "Rect":
                    A = s["b"] * s["h"]
                    xc = s["x"] + s["b"]/2.0
                    yc = s["y"] + s["h"]/2.0
                    ixc = (s["b"] * s["h"]**3)/12.0
                    iyc = (s["h"] * s["b"]**3)/12.0
                elif s["tipo"] == "Circ":
                    r = s["r"]
                    A = np.pi * r**2 * (1 if r>0 else -1)
                    r_abs = abs(r)
                    xc = s["x"]
                    yc = s["y"]
                    ixc = (np.pi * r_abs**4)/4.0 * (1 if r>0 else -1)
                    iyc = ixc
                elif s["tipo"] == "Triang": # Triángulo rectángulo con vértice recto en (x,y)
                    A = 0.5 * s["b"] * s["h"]
                    xc = s["x"] + s["b"]/3.0
                    yc = s["y"] + s["h"]/3.0
                    ixc = (s["b"] * s["h"]**3)/36.0
                    iyc = (s["h"] * s["b"]**3)/36.0
                
                sum_A += A
                sum_Ax += A * xc
                sum_Ay += A * yc
                df_shapes.append({"Fig": f"F{idx+1}", "Tipo": s["tipo"], "A": A, "xc": xc, "yc": yc, "Ixc": ixc, "Iyc": iyc})
                shape_list_for_delete.append(idx)
            
            if sum_A == 0:
                st.error(_t("Área total cero. No se puede calcular centroide.", "Total area zero. Cannot compute centroid."))
                st.stop()
            
            y_bar = sum_Ay / sum_A
            x_bar = sum_Ax / sum_A
            
            # Teorema de Steiner
            Ix_total = 0.0
            Iy_total = 0.0
            for i, s in enumerate(st.session_state.shapes_list):
                # Recalcular propiedades para cada figura (puede reutilizar df_shapes, pero mejor recalcular)
                if s["tipo"] == "Rect":
                    A = s["b"] * s["h"]
                    xc = s["x"] + s["b"]/2.0
                    yc = s["y"] + s["h"]/2.0
                    ixc = (s["b"] * s["h"]**3)/12.0
                    iyc = (s["h"] * s["b"]**3)/12.0
                elif s["tipo"] == "Circ":
                    r = s["r"]
                    A = np.pi * r**2 * (1 if r>0 else -1)
                    r_abs = abs(r)
                    xc = s["x"]
                    yc = s["y"]
                    ixc = (np.pi * r_abs**4)/4.0 * (1 if r>0 else -1)
                    iyc = ixc
                else:  # Triang
                    A = 0.5 * s["b"] * s["h"]
                    xc = s["x"] + s["b"]/3.0
                    yc = s["y"] + s["h"]/3.0
                    ixc = (s["b"] * s["h"]**3)/36.0
                    iyc = (s["h"] * s["b"]**3)/36.0
                dy = yc - y_bar
                dx = xc - x_bar
                Ix_total += ixc + A * dy**2
                Iy_total += iyc + A * dx**2
            
            # Mostrar tabla
            st.dataframe(pd.DataFrame(df_shapes))
            
            # Botones para eliminar figuras individuales
            st.subheader(_t("Eliminar figuras", "Delete shapes"))
            cols_del = st.columns(5)
            for idx in shape_list_for_delete:
                with cols_del[idx % 5]:
                    if st.button(f"F{idx+1}", key=f"del_{idx}"):
                        del st.session_state.shapes_list[idx]
                        st.rerun()
            
            # Resultados
            st.markdown(f"**{_t('Área Total ($A$):', 'Total Area ($A$):')}** {sum_A:.2f}")
            st.markdown(f"**{_t('Centroide Global ($\\bar{{x}}$):', 'Global Centroid ($\\bar{{x}}$):')}** {x_bar:.2f}")
            st.markdown(f"**{_t('Centroide Global ($\\bar{{y}}$):', 'Global Centroid ($\\bar{{y}}$):')}** {y_bar:.2f}")
            st.markdown(f"**{_t('Inercia Eje Centroide X ($I_{{xx}}$):', 'Centroidal Moment of Inertia X ($I_{{xx}}$):')}** {Ix_total:.2f}")
            st.markdown(f"**{_t('Inercia Eje Centroide Y ($I_{{yy}}$):', 'Centroidal Moment of Inertia Y ($I_{{yy}}$):')}** {Iy_total:.2f}")
            
            # Gráfico de la sección
            fig_c, ax_c = plt.subplots(figsize=(6,6))
            ax_c.axhline(0, color='black', lw=1)
            ax_c.axvline(0, color='black', lw=1)
            for s in st.session_state.shapes_list:
                if s["tipo"] == "Rect":
                    ax_c.add_patch(patches.Rectangle((s["x"], s["y"]), s["b"], s["h"],
                                    alpha=0.5, edgecolor='blue', facecolor='cyan' if s["b"]*s["h"]>0 else 'white',
                                    hatch='//' if s["b"]*s["h"]>0 else ''))
                elif s["tipo"] == "Circ":
                    ax_c.add_patch(patches.Circle((s["x"], s["y"]), abs(s["r"]),
                                    alpha=0.5, edgecolor='green', facecolor='lightgreen' if s["r"]>0 else 'white'))
                elif s["tipo"] == "Triang":
                    pts = np.array([[s["x"], s["y"]], [s["x"]+s["b"], s["y"]], [s["x"], s["y"]+s["h"]]])
                    ax_c.add_patch(patches.Polygon(pts, alpha=0.5, edgecolor='red', facecolor='salmon' if s["b"]*s["h"]>0 else 'white'))
            ax_c.plot(x_bar, y_bar, 'r+', markersize=12, markeredgewidth=2, label=_t('Centroide', 'Centroid'))
            ax_c.axhline(y_bar, color='red', linestyle='--', alpha=0.5)
            ax_c.axvline(x_bar, color='red', linestyle='--', alpha=0.5)
            ax_c.legend()
            ax_c.autoscale_view()
            ax_c.set_aspect('equal', 'box')
            st.pyplot(fig_c)
            
            # Exportaciones
            st.markdown("---")
            st.subheader(_t("Exportar", "Export"))
            col_exp1, col_exp2, col_exp3 = st.columns(3)
            with col_exp1:
                if st.button(_t("Memoria DOCX", "DOCX Report")):
                    # Capturar figura
                    buf_img = export_plot_as_png(fig_c)
                    doc = Document()
                    doc.add_heading(_t("Sección Compuesta - Propiedades Geométricas", "Composite Section - Geometric Properties"), 0)
                    doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
                    doc.add_heading(_t("Figuras que componen la sección", "Shapes composing the section"), level=1)
                    # Tabla en doc
                    table = doc.add_table(rows=1+len(df_shapes), cols=7)
                    table.style = 'Table Grid'
                    hdr = table.rows[0].cells
                    hdr[0].text = _t("Figura", "Shape")
                    hdr[1].text = _t("Tipo", "Type")
                    hdr[2].text = _t("Área", "Area")
                    hdr[3].text = "xc"
                    hdr[4].text = "yc"
                    hdr[5].text = "Ixc"
                    hdr[6].text = "Iyc"
                    for i, row in enumerate(df_shapes):
                        cells = table.rows[i+1].cells
                        cells[0].text = row["Fig"]
                        cells[1].text = row["Tipo"]
                        cells[2].text = f"{row['A']:.2f}"
                        cells[3].text = f"{row['xc']:.2f}"
                        cells[4].text = f"{row['yc']:.2f}"
                        cells[5].text = f"{row['Ixc']:.2f}"
                        cells[6].text = f"{row['Iyc']:.2f}"
                    doc.add_heading(_t("Resultados globales", "Global results"), level=1)
                    doc.add_paragraph(f"{_t('Área total A:', 'Total area A:')} {sum_A:.2f}")
                    doc.add_paragraph(f"{_t('Centroide X:', 'Centroid X:')} {x_bar:.2f}")
                    doc.add_paragraph(f"{_t('Centroide Y:', 'Centroid Y:')} {y_bar:.2f}")
                    doc.add_paragraph(f"{_t('Ixx (centroide):', 'Ixx (centroidal):')} {Ix_total:.2f}")
                    doc.add_paragraph(f"{_t('Iyy (centroide):', 'Iyy (centroidal):')} {Iy_total:.2f}")
                    doc.add_picture(buf_img, width=Inches(5))
                    buf_doc = io.BytesIO()
                    doc.save(buf_doc)
                    buf_doc.seek(0)
                    st.download_button(_t("Descargar Memoria DOCX", "Download DOCX Report"), data=buf_doc,
                                       file_name="Seccion_Compuesta.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col_exp2:
                if st.button(_t("Exportar Excel", "Export Excel")):
                    # Crear Excel con dos hojas: figuras y propiedades
                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                        df_shapes.to_excel(writer, sheet_name=_t("Figuras", "Shapes"), index=False)
                        df_res = pd.DataFrame({
                            _t("Propiedad", "Property"): [_t("Área total", "Total area"), _t("Centroide X", "Centroid X"), _t("Centroide Y", "Centroid Y"), _t("Ixx", "Ixx"), _t("Iyy", "Iyy")],
                            _t("Valor", "Value"): [sum_A, x_bar, y_bar, Ix_total, Iy_total]
                        })
                        df_res.to_excel(writer, sheet_name=_t("Propiedades", "Properties"), index=False)
                    output.seek(0)
                    st.download_button(_t("Descargar Excel", "Download Excel"), data=output,
                                       file_name="Seccion_Propiedades.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col_exp3:
                if st.button(_t("Exportar DXF", "Export DXF")):
                    try:
                        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                                 dxf_rotulo, dxf_rotulo_campos)
                        _USE_H12 = True
                    except ImportError:
                        _USE_H12 = False
                    doc_dxf = ezdxf.new('R2010')
                    doc_dxf.units = ezdxf.units.CM
                    if _USE_H12:
                        dxf_setup(doc_dxf, 20); dxf_add_layers(doc_dxf)
                    msp = doc_dxf.modelspace()
                    for s in st.session_state.shapes_list:
                        if s["tipo"] == "Rect":
                            x0r, y0r = s["x"], s["y"]
                            w, h = s["b"], s["h"]
                            w_abs = abs(w); h_abs = abs(h)
                            x0_ = x0r if w >= 0 else x0r + w
                            y0_ = y0r if h >= 0 else y0r + h
                            msp.add_lwpolyline([(x0_, y0_), (x0_+w_abs, y0_), (x0_+w_abs, y0_+h_abs), (x0_, y0_+h_abs), (x0_, y0_)], close=True, dxfattribs={'layer':'CONCRETO'})
                        elif s["tipo"] == "Circ":
                            cx, cy = s["x"], s["y"]
                            r_abs = abs(s["r"])
                            msp.add_circle((cx, cy), r_abs, dxfattribs={'layer':'CONCRETO'})
                        elif s["tipo"] == "Triang":
                            pts = [(s["x"], s["y"]), (s["x"]+s["b"], s["y"]), (s["x"], s["y"]+s["h"])]
                            msp.add_lwpolyline(pts, close=True, dxfattribs={'layer':'CONCRETO'})
                    msp.add_point((x_bar, y_bar), dxfattribs={'layer':'EJES'})
                    if _USE_H12:
                        TH12 = 0.025*20
                        dxf_text(msp, x_bar+2, y_bar+2, f"Centroide: ({x_bar:.2f}, {y_bar:.2f})", "TEXTO", h=TH12)
                        dxf_text(msp, x_bar, y_bar-5, f"Ixx={Ix_total:.2f}  Iyy={Iy_total:.2f}", "TEXTO", h=TH12*0.8)
                        norma_res = st.session_state.get("norma_sel", "NSR-10")
                        _cam12 = dxf_rotulo_campos(f"Seccion Compuesta – Inercias", norma_res, "001")
                        dxf_rotulo(msp, _cam12, 0, y_bar-30, rot_w=max(sum_A**0.5*3,20), rot_h=12, escala=20)
                    else:
                        msp.add_text(f"Centroid: ({x_bar:.2f}, {y_bar:.2f})", dxfattribs={'layer':'TEXTO', 'height':1, 'insert':(x_bar+2, y_bar+2)})
                    out_dxf = io.BytesIO()
                    doc_dxf.write(out_dxf)
                    st.download_button(_t("Descargar DXF", "Download DXF"), data=out_dxf.getvalue(),
                                       file_name="Seccion_Compuesta.dxf", mime="application/dxf")