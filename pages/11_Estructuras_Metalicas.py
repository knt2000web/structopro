import streamlit as st
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es

norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")

st.set_page_config(page_title=_t("Estructuras Metálicas", "Steel Structures"), layout="wide")
st.image(r"assets/steel_header_1773257206595.png", use_container_width=True)
st.title(_t("Diseño de Estructuras Metálicas", "Steel Structure Design"))
st.markdown(_t(f"Cálculo de Propiedades, Compresión y Flexión de Perfiles Laminados en Caliente y Conformados en Frío. Adaptado a la terminología de **{norma_sel}**.", 
               "Properties, Compression, and Flexure computation for Hot-Rolled and Cold-Formed Steel Sections."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# Diccionario de Términos Regionales
if "NTC-EM" in norma_sel or "México" in norma_sel:
    term_W = "Viga IPR"
    term_Col = "Castillo / Columna"
    term_C = "Polín / Perfil C"
    term_Tubo = "HSS / PTR"
elif "CIRSOC 201" in norma_sel or "CIRSOC 301" in norma_sel or "Argentina" in norma_sel or "Bolivia" in norma_sel:
    term_W = "Perfil Doble T"
    term_Col = "Columna"
    term_C = "Costanera"
    term_Tubo = "Tubo Estructural"
elif "Perú" in norma_sel or "Ecuador" in norma_sel:
    term_W = "Viga Tipo I / H"
    term_Col = "Columna"
    term_C = "Correa / Perfil C"
    term_Tubo = "Tubo LAF"
else:
    term_W = "Perfil W / Tipo I"
    term_Col = "Columna / Poste"
    term_C = "Perfil C"
    term_Tubo = "Perfil Tubular Rect."

# ─────────────────────────────────────────────
# CONFIGURACIÓN MATERIAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙️ Materiales", "⚙️ Materials"))
Fy_adm = st.sidebar.number_input(_t("Esfuerzo de Fluencia Fy [MPa]", "Yield Stress Fy [MPa]"), 100.0, 500.0, st.session_state.get("st_fy", 250.0), 10.0, key="st_fy")
Fu_adm = st.sidebar.number_input(_t("Esfuerzo Último Fu [MPa]", "Ultimate Stress Fu [MPa]"), 300.0, 600.0, st.session_state.get("st_fu", 400.0), 10.0, key="st_fu")
E_steel = st.sidebar.number_input(_t("Módulo Elasticidad E [MPa]", "Modulus of Elasticity E [MPa]"), 100000.0, 250000.0, st.session_state.get("st_E", 200000.0), 1000.0, key="st_E")
G_steel = E_steel / (2 * (1 + 0.3))  # Shear modulus approx
peso_esp_acero = 7850.0  # kg/m3

# Variables globales para despiece y costos (se guardan en session_state)
if "steel_despiece" not in st.session_state:
    st.session_state.steel_despiece = []
if "steel_costo_total" not in st.session_state:
    st.session_state.steel_costo_total = 0.0

def agregar_al_despiece(tipo, seccion, longitud_m, peso_kg, observacion):
    st.session_state.steel_despiece.append({
        "Tipo": tipo,
        "Sección (mm)": seccion,
        "Longitud (m)": longitud_m,
        "Peso (kg)": peso_kg,
        "Observación": observacion
    })
    st.session_state.steel_costo_total += peso_kg

# Funciones de dibujo 2D (se mantienen)
def plot_W(d, bf, tw, tf):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((-bf/2, d/2 - tf), bf, tf, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((-bf/2, -d/2), bf, tf, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((-tw/2, -d/2 + tf), tw, d - 2*tf, facecolor='steelblue', edgecolor='black'))
    ax.set_xlim(-bf, bf); ax.set_ylim(-d, d)
    ax.axis('off'); return fig

def plot_3d_W(d, bf, tw, tf, L_m):
    """Crea un modelo 3D simple del perfil W con Plotly."""
    # Convertir a cm para visualización
    d_cm = d / 10
    bf_cm = bf / 10
    tw_cm = tw / 10
    tf_cm = tf / 10
    L_cm = L_m * 100

    # Coordenadas de las partes: web, ala superior, ala inferior
    # Web: rectángulo de ancho tw y alto d-2*tf, centrado
    web_z = [-tw_cm/2, tw_cm/2, tw_cm/2, -tw_cm/2]
    web_y = [0, 0, L_cm, L_cm]
    web_x = [-d_cm/2 + tf_cm, -d_cm/2 + tf_cm, d_cm/2 - tf_cm, d_cm/2 - tf_cm]
    # Para que sea un sólido, usamos mesh. Simplificamos con líneas.
    fig = go.Figure()
    # Web como superficie lateral (líneas)
    for i in range(len(web_x)-1):
        fig.add_trace(go.Scatter3d(x=[web_x[i], web_x[i+1]], y=[web_y[i], web_y[i+1]], z=[web_z[i], web_z[i+1]],
                                   mode='lines', line=dict(color='gray', width=4), showlegend=False))
    # Alas superior e inferior (rectángulos)
    # Ala superior: en z = d_cm/2 - tf_cm/2? Mejor representar como dos planos.
    # Simplemente añadimos líneas de contorno
    ala_sup_x = [-bf_cm/2, bf_cm/2, bf_cm/2, -bf_cm/2, -bf_cm/2]
    ala_sup_y = [0, 0, L_cm, L_cm, 0]
    ala_sup_z = [d_cm/2, d_cm/2, d_cm/2, d_cm/2, d_cm/2]
    fig.add_trace(go.Scatter3d(x=ala_sup_x, y=ala_sup_y, z=ala_sup_z, mode='lines', line=dict(color='darkred', width=4), name='Ala superior'))
    ala_inf_x = [-bf_cm/2, bf_cm/2, bf_cm/2, -bf_cm/2, -bf_cm/2]
    ala_inf_y = [0, 0, L_cm, L_cm, 0]
    ala_inf_z = [-d_cm/2, -d_cm/2, -d_cm/2, -d_cm/2, -d_cm/2]
    fig.add_trace(go.Scatter3d(x=ala_inf_x, y=ala_inf_y, z=ala_inf_z, mode='lines', line=dict(color='darkred', width=4), name='Ala inferior'))
    fig.update_layout(scene=dict(aspectmode='data', xaxis_title='Ancho (cm)', yaxis_title='Largo (cm)', zaxis_title='Alto (cm)'),
                      margin=dict(l=0, r=0, b=0, t=0), height=400, showlegend=False)
    return fig

def plot_T(d, bf, tw, tf):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((-bf/2, d - tf), bf, tf, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((-tw/2, 0), tw, d - tf, facecolor='steelblue', edgecolor='black'))
    ax.set_xlim(-bf, bf); ax.set_ylim(-d/2, d*1.2)
    ax.axis('off'); return fig

def plot_L(h, b, t):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((0, 0), t, h, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((t, 0), b-t, t, facecolor='steelblue', edgecolor='black'))
    ax.set_xlim(-10, max(b,h)+10); ax.set_ylim(-10, max(b,h)+10)
    ax.axis('off'); return fig

def plot_C(h, b, d_lip, t):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((0, 0), t, h, facecolor='darkgray', edgecolor='black'))
    ax.add_patch(patches.Rectangle((0, h-t), b, t, facecolor='darkgray', edgecolor='black'))
    ax.add_patch(patches.Rectangle((0, 0), b, t, facecolor='darkgray', edgecolor='black'))
    if d_lip > 0:
        ax.add_patch(patches.Rectangle((b-t, h-d_lip), t, d_lip, facecolor='darkgray', edgecolor='black'))
        ax.add_patch(patches.Rectangle((b-t, 0), t, d_lip, facecolor='darkgray', edgecolor='black'))
    ax.set_xlim(-b*0.5, b*1.5); ax.set_ylim(-h*0.2, h*1.2)
    ax.axis('off'); return fig

def plot_Tubo(h, b, t):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((0, 0), b, h, facecolor='none', edgecolor='black', lw=3))
    ax.add_patch(patches.Rectangle((t, t), b-2*t, h-2*t, facecolor='white', edgecolor='black', lw=2))
    ax.add_patch(patches.Rectangle((0, 0), b, h, facecolor='darkgray', alpha=0.5))
    ax.set_xlim(-b*0.2, b*1.2); ax.set_ylim(-h*0.2, h*1.2)
    ax.axis('off'); return fig

# Pestañas principales
tab_P, tab_C, tab_F, tab_CF, tab_E = st.tabs([
    _t(f"1. Propiedades {term_W}", "1. W-Shape Properties"),
    _t("2. LAMINADOS: Compresión", "2. HOT-ROLLED: Compression"),
    _t(f"3. LAMINADOS: Flexión ({term_W})", "3. HOT-ROLLED: Flexure"),
    _t("4. CONFORMADOS EN FRÍO", "4. COLD-FORMED"),
    _t("💾 Exportaciones Globales", "💾 Global Exports")
])

# ─────────────────────────────────────────────
# TAB 1: PROPIEDADES PERFIL W
# ─────────────────────────────────────────────
with tab_P:
    st.header(_t(f"Propiedades de Sección - {term_W}", f"Section Properties - {term_W}"))
    p1, p2, p3 = st.columns([1,2,1])
    with p1:
        st.subheader("Dimensiones [mm]")
        dw = st.number_input("Peralte total (d) [mm]", 50.0, 1500.0, st.session_state.get("st_p_d", 300.0), 1.0, key="st_p_d")
        bfw = st.number_input("Ancho patín (bf) [mm]", 50.0, 800.0, st.session_state.get("st_p_bf", 150.0), 1.0, key="st_p_bf")
        tfw = st.number_input("Espesor patín (tf) [mm]", 2.0, 100.0, st.session_state.get("st_p_tf", 10.0), 0.5, key="st_p_tf")
        tww = st.number_input("Espesor alma (tw) [mm]", 2.0, 100.0, st.session_state.get("st_p_tw", 6.0), 0.5, key="st_p_tw")
        L_ref = st.number_input("Longitud de referencia (L) [m]", 0.5, 20.0, 3.0, 0.5, key="st_p_L")
    with p2:
        # Cálculos básicos
        A_w = 2 * (bfw * tfw) + (dw - 2*tfw)*tww
        Ix_w = (tww*(dw - 2*tfw)**3)/12.0 + 2 * ((bfw*tfw**3)/12.0 + (bfw*tfw)*(dw/2.0 - tfw/2.0)**2)
        Iy_w = 2*((tfw*bfw**3)/12.0) + ((dw-2*tfw)*tww**3)/12.0
        rx_w = math.sqrt(Ix_w / A_w)
        ry_w = math.sqrt(Iy_w / A_w)
        Sx_w = Ix_w / (dw/2.0)
        Zx_w = bfw*tfw*(dw-tfw) + (tww*(dw-2*tfw)**2)/4.0
        peso_lin = (A_w / 1e6) * peso_esp_acero
        peso_total = peso_lin * L_ref

        st.write(f"**Área (A):** {A_w:.2f} mm² ($= {A_w/100:.2f} cm^2$)")
        st.write(f"**Momento Inercia Eje Fuerte (Ix):** {Ix_w/1e4:.2f} cm⁴")
        st.write(f"**Momento Inercia Eje Débil (Iy):** {Iy_w/1e4:.2f} cm⁴")
        st.write(f"**Radio de Giro (rx):** {rx_w/10:.2f} cm  |  **(ry):** {ry_w/10:.2f} cm")
        st.write(f"**Módulo Plástico Fuerte (Zx):** {Zx_w/1e3:.2f} cm³")
        st.write(f"⚖️ **Peso Lineal Estimado:** {peso_lin:.2f} kg/m")
        st.write(f"**Peso total para L={L_ref:.2f} m:** {peso_total:.2f} kg")

        # Opción para agregar al despiece
        if st.button(_t("Agregar este perfil al despiece", "Add this profile to the list")):
            agregar_al_despiece(
                tipo=f"{term_W} (propiedades)",
                seccion=f"{dw}x{bfw}x{tww}x{tfw} mm",
                longitud_m=L_ref,
                peso_kg=peso_total,
                observacion=_t("Perfil W calculado", "W-shape calculated")
            )
            st.success(_t("Perfil agregado al despiece", "Profile added to list"))

    with p3:
        st.pyplot(plot_W(dw, bfw, tww, tfw))
        st.markdown("---")
        st.subheader(_t("Visualización 3D", "3D Visualization"))
        fig3d = plot_3d_W(dw, bfw, tww, tfw, L_ref)
        st.plotly_chart(fig3d, use_container_width=True)

# ─────────────────────────────────────────────
# TAB 2: LAMINADOS - COMPRESIÓN
# ─────────────────────────────────────────────
with tab_C:
    st.header(_t(f"Resistencia a Compresión Axial", f"Axial Compression Resistance"))
    comp_opts = [term_W, "Perfil T", "Perfil L (Angular)"]
    tipo_comp = st.selectbox("Seleccione Perfil Laminado en Caliente a evaluar:", comp_opts,
                             index=comp_opts.index(st.session_state.get("st_c_tipo", comp_opts[0])),
                             key="st_c_tipo")
    col_c1, col_c2, col_c3 = st.columns([1,1.5,1.5])
    with col_c1:
        st.subheader("Datos Compresión")
        L_c_m = st.number_input("Longitud No Arriostrada (Lc) [m]", 0.5, 20.0, st.session_state.get("st_c_L", 3.0), 0.5, key="st_c_L")
        P_u = st.number_input("Carga Axial Requerida Pu [kN]", 10.0, 10000.0, st.session_state.get("st_c_Pu", 500.0), 10.0, key="st_c_Pu")
    with col_c2:
        if tipo_comp == term_W:
            d_c = st.number_input(f"{tipo_comp} - Peralte d [mm]", 50.0, 1000.0, st.session_state.get("st_cw_d", 300.0), key="st_cw_d")
            bf_c = st.number_input(f"{tipo_comp} - Ancho bf [mm]", 50.0, 500.0, st.session_state.get("st_cw_bf", 150.0), key="st_cw_bf")
            tf_c = st.number_input(f"{tipo_comp} - tf [mm]", 2.0, 50.0, st.session_state.get("st_cw_tf", 10.0), key="st_cw_tf")
            tw_c = st.number_input(f"{tipo_comp} - tw [mm]", 2.0, 50.0, st.session_state.get("st_cw_tw", 6.0), key="st_cw_tw")
            A_c = 2 * (bf_c * tf_c) + (d_c - 2*tf_c)*tw_c
            Iy_c = 2*((tf_c*bf_c**3)/12.0) + ((d_c-2*tf_c)*tw_c**3)/12.0
            r_c = math.sqrt(Iy_c / A_c)
            fig_c = plot_W(d_c, bf_c, tw_c, tf_c)
            peso_lin_c = (A_c / 1e6) * peso_esp_acero
        elif tipo_comp == "Perfil T":
            d_c = st.number_input("Perfil T - Peralte d [mm]", 50.0, 500.0, st.session_state.get("st_ct_d", 150.0), key="st_ct_d")
            bf_c = st.number_input("Perfil T - Ancho bf [mm]", 50.0, 500.0, st.session_state.get("st_ct_bf", 150.0), key="st_ct_bf")
            tf_c = st.number_input("Perfil T - tf [mm]", 2.0, 50.0, st.session_state.get("st_ct_tf", 10.0), key="st_ct_tf")
            tw_c = st.number_input("Perfil T - tw [mm]", 2.0, 50.0, st.session_state.get("st_ct_tw", 8.0), key="st_ct_tw")
            A_c = (bf_c * tf_c) + (d_c - tf_c)*tw_c
            Iy_c = ((tf_c*bf_c**3)/12.0) + ((d_c-tf_c)*tw_c**3)/12.0
            r_c = math.sqrt(Iy_c / A_c)
            fig_c = plot_T(d_c, bf_c, tw_c, tf_c)
            peso_lin_c = (A_c / 1e6) * peso_esp_acero
        else:  # Perfil L
            b_c = st.number_input("Perfil L - Altura/Base b [mm]", 20.0, 300.0, st.session_state.get("st_cl_b", 100.0), key="st_cl_b")
            t_c = st.number_input("Perfil L - Espesor t [mm]", 2.0, 30.0, st.session_state.get("st_cl_t", 10.0), key="st_cl_t")
            A_c = (2 * b_c - t_c) * t_c
            r_c = 0.2 * b_c
            fig_c = plot_L(b_c, b_c, t_c)
            peso_lin_c = (A_c / 1e6) * peso_esp_acero
    with col_c3:
        Lc_mm = L_c_m * 1000.0
        esbeltez = Lc_mm / r_c
        Fe = (math.pi**2 * E_steel) / (esbeltez**2)
        if esbeltez <= 4.71 * math.sqrt(E_steel/Fy_adm):
            Fcr = (0.658**(Fy_adm/Fe)) * Fy_adm
        else:
            Fcr = 0.877 * Fe
        Pn = Fcr * A_c / 1000.0
        phi_Pn = 0.90 * Pn
        st.markdown(f"**Relación de Esbeltez L/r_min:** {esbeltez:.2f} (Límite ≈ 200)")
        st.markdown(f"**Esfuerzo Crítico (Fcr):** {Fcr:.2f} MPa")
        st.markdown(f"**Resistencia Axial de Diseño ($\phi P_n$):** <span style='color:blue;font-size:22px'>{phi_Pn:.1f} kN</span>", unsafe_allow_html=True)
        if P_u <= phi_Pn:
            st.success(f"✅ ¡Aprobado! (FS = {phi_Pn/P_u:.2f})")
        else:
            st.error("❌ No Aprobado por Compresión Axial / Pandeo.")
        st.pyplot(fig_c)

        # Agregar al despiece si se desea
        peso_total_c = peso_lin_c * L_c_m
        if st.button(_t("Agregar este elemento al despiece", "Add this element to the list")):
            agregar_al_despiece(
                tipo=f"{tipo_comp} (compresión)",
                seccion=f"{d_c if 'd' in locals() else b_c}x{bf_c if 'bf_c' in locals() else t_c} mm",
                longitud_m=L_c_m,
                peso_kg=peso_total_c,
                observacion=_t(f"Resistencia φPn={phi_Pn:.1f} kN vs Pu={P_u:.1f} kN", f"Strength φPn={phi_Pn:.1f} kN vs Pu={P_u:.1f} kN")
            )
            st.success(_t("Elemento agregado", "Element added"))

# ─────────────────────────────────────────────
# TAB 3: LAMINADOS - FLEXIÓN (PERFIL W)
# ─────────────────────────────────────────────
with tab_F:
    st.header(_t(f"Resistencia a Flexión - {term_W}", f"Flexural Resistance - {term_W}"))
    st.write(_t("Evaluación del momento flector considerando Pandeo Lateral Torsional (LTB).", "Bending moment evaluation considering Lateral Torsional Buckling (LTB)."))
    f_c1, f_c2, f_c3 = st.columns(3)
    with f_c1:
        Mu = st.number_input("Momento Último Solicitante Mu [kN-m]", 10.0, 5000.0, st.session_state.get("st_f_Mu", 150.0), 10.0, key="st_f_Mu")
        Lb_m = st.number_input("Longitud No Arriostrada Lateramente Lb [m]", 0.5, 20.0, st.session_state.get("st_f_Lb", 2.0), 0.5, key="st_f_Lb")
        Cb = st.number_input("Coeficiente de Momento (Cb)", 1.0, 3.0, st.session_state.get("st_f_Cb", 1.0), 0.1, key="st_f_Cb")
    with f_c2:
        d_f = st.number_input(f"Peralte d [mm]", 100.0, 1500.0, st.session_state.get("st_fw_d", 300.0), 1.0, key="st_fw_d")
        bf_f = st.number_input(f"Ancho bf [mm]", 50.0, 800.0, st.session_state.get("st_fw_bf", 150.0), 1.0, key="st_fw_bf")
        tf_f = st.number_input(f"Espesor tf [mm]", 2.0, 50.0, st.session_state.get("st_fw_tf", 10.0), 0.5, key="st_fw_tf")
        tw_f = st.number_input(f"Espesor tw [mm]", 2.0, 50.0, st.session_state.get("st_fw_tw", 6.0), 0.5, key="st_fw_tw")
        L_viga = st.number_input("Longitud de la viga (L) [m]", 1.0, 20.0, 6.0, 0.5, key="st_f_L")
    with f_c3:
        # Propiedades para LTB
        h0 = d_f - tf_f
        cw = ((h0**2)*tf_f*bf_f**3)/24.0
        Iy_f = 2*((tf_f*bf_f**3)/12.0) + ((d_f-2*tf_f)*tw_f**3)/12.0
        J_tor = (2*bf_f*tf_f**3 + (d_f-tf_f)*tw_f**3)/3.0
        r_ts = bf_f / min(math.sqrt(12*(1+(1/6)*(d_f*tw_f)/(bf_f*tf_f))), 1000)
        Zx_f = bf_f*tf_f*(d_f-tf_f) + (tw_f*(d_f-2*tf_f)**2)/4.0
        Sx_f = (2 * ((bf_f*tf_f**3)/12.0 + (bf_f*tf_f)*(d_f/2.0 - tf_f/2.0)**2) + (tw_f*(d_f - 2*tf_f)**3)/12.0) / (d_f/2.0)
        Lp = 1.76 * r_ts * math.sqrt(E_steel/Fy_adm) / 1000.0
        Lr = np.pi * r_ts * math.sqrt(E_steel/(0.7*Fy_adm)) / 1000.0
        Mp = Fy_adm * Zx_f / 1e6
        st.write(f"**Límite Fluencia (Lp):** {Lp:.2f} m")
        st.write(f"**Límite Pandeo (Lr):** {Lr:.2f} m")
        st.write(f"**Momento Plástico (Mp):** {Mp:.1f} kN-m")
        Lbmm = Lb_m * 1000.0
        if Lb_m <= Lp:
            Mn = Mp
            estado = "Fluencia (Zona Plástica)"
        elif Lb_m <= Lr:
            Mn = Cb * (Mp - (Mp - 0.7*Fy_adm*Sx_f/1e6) * ((Lb_m - Lp)/(Lr - Lp)))
            Mn = min(Mn, Mp)
            estado = "LTB Inelástico"
        else:
            Fcr_ltb = (Cb * np.pi**2 * E_steel / (Lbmm/r_ts)**2) * math.sqrt(1 + 0.078*(J_tor/Sx_f)*(Lbmm/r_ts)**2)
            Mn = Fcr_ltb * Sx_f / 1e6
            Mn = min(Mn, Mp)
            estado = "LTB Elástico"
        phi_Mn = 0.90 * Mn
        st.markdown(f"**Zona Comportamiento:** {estado}")
        st.markdown(f"**Capacidad Momento ($\phi M_n$):** <span style='color:blue;font-size:22px'>{phi_Mn:.1f} kN-m</span>", unsafe_allow_html=True)
        if Mu <= phi_Mn:
            st.success(f"✅ ¡Viga Cumple a Flexión! (FS={phi_Mn/Mu:.2f})")
        else:
            st.error("❌ Viga falla por Flexión / LTB.")

        # Peso y despiece
        A_f = 2 * (bf_f * tf_f) + (d_f - 2*tf_f)*tw_f
        peso_lin_f = (A_f / 1e6) * peso_esp_acero
        peso_total_f = peso_lin_f * L_viga
        st.write(f"⚖️ **Peso estimado de la viga:** {peso_total_f:.2f} kg")

        if st.button(_t("Agregar esta viga al despiece", "Add this beam to the list")):
            agregar_al_despiece(
                tipo=f"{term_W} (flexión)",
                seccion=f"{d_f}x{bf_f}x{tw_f}x{tf_f} mm",
                longitud_m=L_viga,
                peso_kg=peso_total_f,
                observacion=_t(f"φMn={phi_Mn:.1f} kN-m vs Mu={Mu:.1f} kN-m", f"φMn={phi_Mn:.1f} kN-m vs Mu={Mu:.1f} kN-m")
            )
            st.success(_t("Viga agregada al despiece", "Beam added to list"))

# ─────────────────────────────────────────────
# TAB 4: CONFORMADOS EN FRÍO (Cold-Formed)
# ─────────────────────────────────────────────
with tab_CF:
    st.header(_t("Perfiles Conformados en Frío (Cold-Formed)", "Cold-Formed Sections"))
    cf_opts = [f"{term_C} con labios (rígido)", f"{term_C} sin labios (U)", term_Tubo]
    tipo_cf = st.selectbox("Seleccione Perfil Conformado en Frío:", cf_opts,
                           index=cf_opts.index(st.session_state.get("st_cf_tipo", cf_opts[0])),
                           key="st_cf_tipo")
    col_cf1, col_cf2 = st.columns([1,2])
    with col_cf1:
        h_cf = st.number_input("Altura total h [mm]", 20.0, 400.0, st.session_state.get("st_cf_h", 150.0), 1.0, key="st_cf_h")
        b_cf = st.number_input("Ancho base b [mm]", 20.0, 200.0, st.session_state.get("st_cf_b", 50.0), 1.0, key="st_cf_b")
        t_cf = st.number_input("Espesor de lámina t [mm]", 0.5, 10.0, st.session_state.get("st_cf_t", 2.0), 0.1, key="st_cf_t")
        L_cf = st.number_input("Longitud [m]", 0.5, 20.0, 2.0, 0.5, key="st_cf_L")
        if "labios" in tipo_cf:
            d_lip = st.number_input("Pestaña/Labio d [mm]", 0.0, 50.0, st.session_state.get("st_cf_l", 15.0), key="st_cf_l")
            A_cf = (h_cf + 2*b_cf + 2*d_lip - 4*t_cf) * t_cf
            fig_cf = plot_C(h_cf, b_cf, d_lip, t_cf)
        elif tipo_cf == term_Tubo:
            d_lip = 0
            A_cf = 2*(h_cf + b_cf - 2*t_cf) * t_cf
            fig_cf = plot_Tubo(h_cf, b_cf, t_cf)
        else:
            d_lip = 0
            A_cf = (h_cf + 2*b_cf - 2*t_cf) * t_cf
            fig_cf = plot_C(h_cf, b_cf, 0, t_cf)
        peso_lin_cf = (A_cf / 1e6) * peso_esp_acero
        peso_total_cf = peso_lin_cf * L_cf
        st.write(f"**Área Bruta (A):** {A_cf:.2f} mm²")
        st.write(f"⚖️ **Peso Lineal Estimado:** {peso_lin_cf:.2f} kg/m")
        st.write(f"**Peso total:** {peso_total_cf:.2f} kg")
    with col_cf2:
        st.subheader("Verificación de Compresión")
        Pu_cf = st.number_input("Carga Axial Pu [kN]", 5.0, 500.0, st.session_state.get("st_cf_Pu", 50.0), key="st_cf_Pu")
        w_t_ratio = max(h_cf/t_cf, b_cf/t_cf)
        if w_t_ratio > 200:
            st.warning("⚠️ Relación ancho/espesor > 200. Riesgo de pandeo local muy severo.")
            Ae = 0.5 * A_cf
        elif w_t_ratio > 50:
            Ae = 0.8 * A_cf
        else:
            Ae = A_cf
        st.write(f"**Área Efectiva por Pandeo Local (Ae):** aprox {Ae:.2f} mm²")
        if tipo_cf == term_Tubo:
            r_cf = b_cf * 0.38
        else:
            r_cf = b_cf * 0.30
        esbeltez_cf = (L_cf*1000.0) / r_cf
        st.write(f"**Relación Esbeltez Global L/r:** {esbeltez_cf:.2f}")
        Fe_cf = (math.pi**2 * E_steel) / (max(esbeltez_cf, 1)**2)
        if esbeltez_cf <= 4.71 * math.sqrt(E_steel/Fy_adm):
            Fcr_cf = (0.658**(Fy_adm/Fe_cf)) * Fy_adm
        else:
            Fcr_cf = 0.877 * Fe_cf
        Pn_cf = (Ae * Fcr_cf) / 1000.0
        phi_Pn_cf = 0.85 * Pn_cf
        st.markdown(f"**Resistencia Axial Compresión ($\phi P_n$):** <span style='color:blue;font-size:22px'>{phi_Pn_cf:.1f} kN</span>", unsafe_allow_html=True)
        if Pu_cf <= phi_Pn_cf:
            st.success("✅ ¡El perfil conformado en frío CUMPLE!")
        else:
            st.error("❌ El perfil falla por compresión / pandeo global o local.")
        st.pyplot(fig_cf)

        if st.button(_t("Agregar este perfil al despiece", "Add this profile to the list")):
            agregar_al_despiece(
                tipo=f"{tipo_cf} (conformado)",
                seccion=f"{h_cf}x{b_cf}x{t_cf} mm",
                longitud_m=L_cf,
                peso_kg=peso_total_cf,
                observacion=_t(f"φPn={phi_Pn_cf:.1f} kN vs Pu={Pu_cf:.1f} kN", f"φPn={phi_Pn_cf:.1f} kN vs Pu={Pu_cf:.1f} kN")
            )
            st.success(_t("Perfil agregado al despiece", "Profile added to list"))

# ─────────────────────────────────────────────
# TAB 5: EXPORTACIONES GLOBALES
# ─────────────────────────────────────────────
with tab_E:
    st.header(_t("💾 Exportaciones Globales", "💾 Global Exports"))

    # Mostrar despiece acumulado
    st.subheader(_t("📋 Despiece de Acero", "📋 Steel Cutting List"))
    if st.session_state.steel_despiece:
        df_desp = pd.DataFrame(st.session_state.steel_despiece)
        st.dataframe(df_desp.style.format({"Longitud (m)": "{:.2f}", "Peso (kg)": "{:.2f}"}), use_container_width=True, hide_index=False)
        st.write(f"**Peso total acumulado:** {st.session_state.steel_costo_total:.2f} kg")
        # Gráfico de barras
        fig_bars, ax_bars = plt.subplots(figsize=(8, 4))
        # Agrupar por tipo
        df_group = df_desp.groupby("Tipo")["Peso (kg)"].sum().reset_index()
        ax_bars.bar(df_group["Tipo"], df_group["Peso (kg)"], color='steelblue')
        ax_bars.set_xlabel(_t("Tipo de perfil", "Profile type"))
        ax_bars.set_ylabel(_t("Peso (kg)", "Weight (kg)"))
        ax_bars.set_title(_t("Distribución de pesos por tipo", "Weight distribution by type"))
        ax_bars.grid(True, alpha=0.3)
        st.pyplot(fig_bars)
        bars_img = io.BytesIO()
        fig_bars.savefig(bars_img, format='png', dpi=150, bbox_inches='tight')
        bars_img.seek(0)
    else:
        st.info(_t("Aún no hay elementos en el despiece. Agrega perfiles desde las pestañas anteriores.", "No elements in the list yet. Add profiles from the previous tabs."))

    st.markdown("---")
    st.subheader(_t("📄 Memoria de Cálculo (DOCX)", "📄 Calculation Report (DOCX)"))
    if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
        doc = Document()
        doc.add_heading(_t(f"Memoria de Estructuras Metálicas - {norma_sel}", f"Steel Structures Report - {norma_sel}"), 0)
        doc.add_paragraph(_t(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"))
        doc.add_heading(_t("1. Materiales", "1. Materials"), level=1)
        doc.add_paragraph(f"Fy = {Fy_adm:.1f} MPa, Fu = {Fu_adm:.1f} MPa, E = {E_steel:.0f} MPa")
        doc.add_heading(_t("2. Despiece de Acero", "2. Steel Cutting List"), level=1)
        if st.session_state.steel_despiece:
            table = doc.add_table(rows=1+len(st.session_state.steel_despiece), cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = _t("Tipo", "Type")
            hdr[1].text = _t("Sección (mm)", "Section (mm)")
            hdr[2].text = _t("Longitud (m)", "Length (m)")
            hdr[3].text = _t("Peso (kg)", "Weight (kg)")
            hdr[4].text = _t("Observación", "Observation")
            for i, row in enumerate(st.session_state.steel_despiece):
                cells = table.rows[i+1].cells
                cells[0].text = row["Tipo"]
                cells[1].text = row["Sección (mm)"]
                cells[2].text = f"{row['Longitud (m)']:.2f}"
                cells[3].text = f"{row['Peso (kg)']:.1f}"
                cells[4].text = row["Observación"]
            doc.add_paragraph(f"{_t('Peso total:', 'Total weight:')} {st.session_state.steel_costo_total:.2f} kg")
            # Insertar gráfico
            doc.add_picture(bars_img, width=Inches(5))
        else:
            doc.add_paragraph(_t("No se han agregado elementos.", "No elements added."))
        doc.add_heading(_t("3. Notas de diseño", "3. Design notes"), level=1)
        doc.add_paragraph(_t("Los cálculos se realizaron según los principios de resistencia LRFD (AISC 360).", "Calculations performed using LRFD principles (AISC 360)."))
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button(_t("📥 Descargar Memoria", "📥 Download Report"), data=doc_mem, file_name="Memoria_Acero.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.subheader(_t("💰 Presupuesto APU", "💰 APU Budget"))
    if "apu_config" in st.session_state:
        apu = st.session_state.apu_config
        mon = apu["moneda"]
        precio_kg = st.number_input(_t(f"Costo por Kg de Acero Estructural Suministrado y Armado [{mon}/kg]", f"Steel supply and erection cost per kg [{mon}/kg]"),
                                    value=st.session_state.get("st_price_kg", 8000.0 if mon=="COP$" else 4.0), step=100.0, key="st_price_kg")
        costo_acero = st.session_state.steel_costo_total * precio_kg
        # Mano de obra (rendimiento aproximado 0.05 días/kg)
        dias_mo = st.session_state.steel_costo_total * 0.05
        costo_mo = dias_mo * apu.get("costo_dia_mo", 69333.33)
        # Herramienta, AIU, etc.
        herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
        aiu = costo_acero * apu.get("pct_aui", 0.30)
        utilidad = costo_acero * apu.get("pct_util", 0.05)
        iva = utilidad * apu.get("iva", 0.19)
        total_proyecto = costo_acero + costo_mo + herramienta + aiu + iva

        data_apu = {
            _t("Item", "Item"): [_t("Acero estructural", "Structural steel"), _t("Mano de obra (montaje)", "Labor (erection)"),
                                 _t("Herramienta menor", "Minor tools"), _t("A.I.U.", "A.I.U."),
                                 _t("IVA s/Utilidad", "VAT on profit"), _t("TOTAL", "TOTAL")],
            _t("Cantidad", "Quantity"): [f"{st.session_state.steel_costo_total:.1f} kg", f"{dias_mo:.2f} días", f"{apu.get('pct_herramienta',0.05)*100:.1f}% MO",
                                         f"{apu.get('pct_aui',0.30)*100:.1f}%", f"{apu.get('iva',0.19)*100:.1f}%", ""],
            f"Subtotal [{mon}]": [f"{costo_acero:,.2f}", f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}",
                                  f"{iva:,.2f}", f"**{total_proyecto:,.2f}**"]
        }
        st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)

        # Exportar Excel APU
        output_excel = io.BytesIO()
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            df_export = pd.DataFrame({
                "Item": ["Acero estructural", "Mano de obra"],
                "Cantidad": [st.session_state.steel_costo_total, dias_mo],
                "Unidad": [precio_kg, apu.get("costo_dia_mo", 69333.33)]
            })
            df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
            df_export.to_excel(writer, index=False, sheet_name='APU Acero')
            workbook = writer.book
            worksheet = writer.sheets['APU Acero']
            money_fmt = workbook.add_format({'num_format': '#,##0.00'})
            worksheet.set_column('A:A', 25)
            worksheet.set_column('B:D', 15, money_fmt)
            row = len(df_export) + 1
            worksheet.write(row, 0, "Costo Directo (CD)", workbook.add_format({'bold': True}))
            worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
            row += 1
            worksheet.write(row, 0, "Herramienta Menor", workbook.add_format({'bold': True}))
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_herramienta",0.05)}', money_fmt)
            row += 1
            worksheet.write(row, 0, "A.I.U", workbook.add_format({'bold': True}))
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui",0.30)}', money_fmt)
            row += 1
            worksheet.write(row, 0, "IVA s/ Utilidad", workbook.add_format({'bold': True}))
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_util",0.05)}*{apu.get("iva",0.19)}', money_fmt)
            row += 1
            worksheet.write(row, 0, "TOTAL PRESUPUESTO", workbook.add_format({'bold': True}))
            worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
        output_excel.seek(0)
        st.download_button(_t("📥 Descargar Presupuesto Excel", "📥 Download Budget Excel"), data=output_excel,
                           file_name="APU_Acero.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info(_t("💡 Ve a la página 'APU Mercado' para cargar los costos en vivo.", "💡 Go to the 'Market APU' page to load live costs."))

    st.markdown("---")
    st.subheader(_t("📐 DXF de Sección", "📐 Section DXF"))
    st.write(_t("Genera un archivo DXF con la sección transversal del último perfil W definido en la pestaña de propiedades.", "Generates a DXF file with the cross-section of the last W-shape defined in the properties tab."))
    try:
        doc_dxf = ezdxf.new('R2010')
        msp = doc_dxf.modelspace()
        # Usar las dimensiones de la pestaña de propiedades (si existen)
        if 'dw' in locals() and 'bfw' in locals():
            p1 = (-bfw/2, dw/2)
            points = [
                (-bfw/2, dw/2), (bfw/2, dw/2), (bfw/2, dw/2 - tfw),
                (tww/2, dw/2 - tfw), (tww/2, -dw/2 + tfw), (bfw/2, -dw/2 + tfw),
                (bfw/2, -dw/2), (-bfw/2, -dw/2), (-bfw/2, -dw/2 + tfw),
                (-tww/2, -dw/2 + tfw), (-tww/2, dw/2 - tfw), (-bfw/2, dw/2 - tfw), (-bfw/2, dw/2)
            ]
            msp.add_lwpolyline(points, dxfattribs={'layer': 'PERFIL_W', 'color': 5, 'closed': True})
            out_dxf = io.BytesIO()
            doc_dxf.write(out_dxf)
            st.download_button(_t("📥 Descargar Perfil_W.dxf", "📥 Download Perfil_W.dxf"), data=out_dxf.getvalue(),
                               file_name=f"Perfil_W_{dw}x{bfw}.dxf", mime="application/dxf")
        else:
            st.warning(_t("Primero defina un perfil W en la pestaña de propiedades.", "First define a W-shape in the properties tab."))
    except Exception as e:
        st.error(f"Error generando DXF: {e}")