"""
08_Placa_facil.py  —  StructoPro · Módulo Placa Fácil
Losa Aligerada con Perfil Colmena y Bloquelón Santafé
NSR-10 / E.060 / ACI 318-25
Prompt Maestro v5.1 — 100% implementado:
  N1 Motor calculo [OK]  N2 Persistencia [OK]  N3 Robustez [OK]
  N4 Entregables [OK]   N5 Comercial [OK]
"""

import streamlit as st
import numpy as np
import plotly.graph_objects as go
import pandas as pd
import math
import io
import json
import os
import base64
from datetime import datetime

# ── Imports opcionales protegidos ─────────────────────────────────────────────
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _MPL_OK = True
except ImportError:
    _MPL_OK = False

try:
    import ezdxf
    from ezdxf.math import Vec2
    _DXF_EXT = True          # ezdxf disponible
except ImportError:
    _DXF_EXT = False          # usar generador interno sin dependencias

try:
    from docx import Document as _DocxDoc
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm as rl_cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, Image as RLImage)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle as RLPStyle
    from reportlab.lib.enums import TA_CENTER as RL_CENTER, TA_LEFT as RL_LEFT
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

try:
    from normas_referencias import mostrar_referencias_norma
except ImportError:
    def mostrar_referencias_norma(*a, **kw): pass

# ══════════════════════════════════════════════════════════════════════════════
# PERSISTENCIA  —  save_state / load_state
# ══════════════════════════════════════════════════════════════════════════════
STATE_FILE = "placa_facil_state.json"
_PERSIST_KEYS = [
    "empresa", "proyecto", "ingeniero",
    "pf_elaboro", "pf_reviso", "pf_aprobo",
    "pf_proyecto_nombre", "pf_proyecto_cliente", "pf_proyecto_dir",
    "pf_plano_num", "pf_escala", "pf_norma",
    "logo_bytes", "user_role", "idioma",
]

def save_state():
    try:
        snap = {}
        for k in _PERSIST_KEYS:
            val = st.session_state.get(k)
            if k == "logo_bytes" and isinstance(val, (bytes, bytearray)):
                snap[k] = {"__type__": "bytes_b64",
                            "data": base64.b64encode(val).decode()}
            elif val is not None:
                try:
                    json.dumps(val)
                    snap[k] = val
                except (TypeError, ValueError):
                    snap[k] = str(val)
        with open(STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(snap, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def load_state():
    if not os.path.exists(STATE_FILE):
        return
    try:
        with open(STATE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        for k, v in data.items():
            if k not in st.session_state:
                if isinstance(v, dict) and v.get("__type__") == "bytes_b64":
                    st.session_state[k] = base64.b64decode(v["data"])
                else:
                    st.session_state[k] = v
    except Exception:
        pass


# Inicializar estado
if "pf_state_loaded" not in st.session_state:
    load_state()
    st.session_state["pf_state_loaded"] = True

for _k in ["empresa", "proyecto", "ingeniero",
           "pf_elaboro", "pf_reviso", "pf_aprobo",
           "pf_proyecto_nombre", "pf_proyecto_cliente", "pf_proyecto_dir",
           "pf_plano_num", "pf_escala"]:
    if _k not in st.session_state:
        st.session_state[_k] = ""

if "user_role" not in st.session_state:
    st.session_state["user_role"] = "free"
if "idioma" not in st.session_state:
    st.session_state["idioma"] = "Español"

# ── Admin automático por email (corre en cada recarga de página) ───────────
_PF_ADMIN_EMAILS = {"civcesar@gmail.com"}
try:
    _pf_auth_top = st.session_state.get("auth_user")
    _pf_email_top = (getattr(_pf_auth_top, "email", None) or "").lower().strip() if _pf_auth_top else ""
except Exception:
    _pf_email_top = ""
if _pf_email_top in _PF_ADMIN_EMAILS:
    st.session_state["user_role"] = "admin"

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURACION
# ══════════════════════════════════════════════════════════════════════════════
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es

st.set_page_config(page_title=_t("Placa Facil", "Easy Slab"),
                   layout="wide", page_icon=":material/layers:")

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTES DE MATERIALES
# ══════════════════════════════════════════════════════════════════════════════
BLOCK_DATA = {
    "nombre": "Bloquelón Santafé",
    "largo": 0.80, "ancho": 0.23, "alto": 0.08,
    "peso_unitario": 13, "rendimiento_por_m2": 5.18,
    "color": "#C19A6B", "pestana": 0.015,
}
PROFILE_DATA = {
    "nombre": "Perfil Colmena",
    "alto_total": 0.090, "ancho_alma": 0.040, "espesor_alma": 0.002,
    "ancho_ala": 0.020, "espesor_ala": 0.002,
    "b_total": 0.130,
    "peso_por_m": 9.8, "color": "#C0C0C0",
}
MESH_DATA = {
    "nombre": "Malla Electrosoldada Q2",
    "diametro": 0.00635, "espaciado_largo": 0.25,
    "espaciado_corto": 0.40, "traslapo": 0.15, "peso_por_m2": 3.0,
}
CONCRETE_DATA = {"resistencia": 21, "densidad": 2400, "cemento_por_m3": 350}
CARGA_VIVA = {"Residencial": 1.8, "Comercial": 2.5, "Industrial": 3.5, "Oficinas": 2.4}
NORMAS_PLACA = {
    "NSR-10 (Colombia)": {
        "luz_max": 4.2, "topping_min": 0.04, "concreto_min": 21,
        "requiere_viga_borde": True,
        "ref": "NSR-10 Cap. C.21 y Titulo E",
    },
    "E.060 (Perú)": {
        "luz_max": 4.2, "topping_min": 0.04, "concreto_min": 21,
        "requiere_viga_borde": True, "ref": "E.060 / NTE E.030",
    },
    "ACI 318-25 (EE.UU.)": {
        "luz_max": 4.2, "topping_min": 0.04, "concreto_min": 21,
        "requiere_viga_borde": True, "ref": "ACI 318-25 Cap. 7",
    },
}

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS GRAFICOS
# ══════════════════════════════════════════════════════════════════════════════
def _plotly_light(fig, height=380):
    fig.update_layout(
        height=height,
        plot_bgcolor="#ffffff",
        paper_bgcolor="#ffffff",
        font=dict(color="#1a1a1a", size=12),
        xaxis=dict(gridcolor="#e0e0e0", linecolor="#cccccc", zerolinecolor="#e0e0e0"),
        yaxis=dict(gridcolor="#e0e0e0", linecolor="#cccccc", zerolinecolor="#e0e0e0"),
        legend=dict(bgcolor="rgba(255,255,255,0.9)", bordercolor="#cccccc", borderwidth=1),
    )
    return fig


def fig_to_docx_white(fig):
    """Convierte figura matplotlib a PNG modo claro sin alterar la pantalla."""
    if not _MPL_OK:
        return None
    orig_fc = fig.get_facecolor()
    axes_backup = []
    for ax in fig.get_axes():
        axes_backup.append({
            "fc": ax.get_facecolor(),
            "title": ax.title.get_color(),
            "xlabel": ax.xaxis.label.get_color(),
            "ylabel": ax.yaxis.label.get_color(),
            "spines": {s: ax.spines[s].get_edgecolor() for s in ax.spines},
        })
    fig.patch.set_facecolor("white")
    for ax in fig.get_axes():
        ax.set_facecolor("#f8f9fa")
        ax.title.set_color("#1a1a1a")
        ax.xaxis.label.set_color("#1a1a1a")
        ax.yaxis.label.set_color("#1a1a1a")
        ax.tick_params(colors="#1a1a1a")
        for spine in ax.spines.values():
            spine.set_edgecolor("#cccccc")
        for lbl in ax.get_xticklabels() + ax.get_yticklabels():
            lbl.set_color("#1a1a1a")
        lg = ax.get_legend()
        if lg:
            lg.get_frame().set_facecolor("white")
            for t in lg.get_texts():
                t.set_color("#1a1a1a")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, facecolor="white",
                transparent=False, bbox_inches="tight")
    buf.seek(0)
    png_bytes = buf.read()
    fig.patch.set_facecolor(orig_fc)
    for ax, bk in zip(fig.get_axes(), axes_backup):
        ax.set_facecolor(bk["fc"])
        ax.title.set_color(bk["title"])
        ax.xaxis.label.set_color(bk["xlabel"])
        ax.yaxis.label.set_color(bk["ylabel"])
        for sname, sc in bk["spines"].items():
            ax.spines[sname].set_edgecolor(sc)
    return png_bytes

# ══════════════════════════════════════════════════════════════════════════════
# BANNER HEADER
# ══════════════════════════════════════════════════════════════════════════════
_emp_hdr = st.session_state.get("empresa", "") or "StructoPro"
_proy_hdr = st.session_state.get("proyecto", "") or "Placa Facil"
st.markdown(f"""
<div style="background:linear-gradient(135deg,#0d1b2a 0%,#1e3a5f 60%,#2e6da4 100%);
  padding:24px 36px;border-radius:14px;margin-bottom:16px;">
 <div style="display:flex;align-items:center;gap:18px;">
  <div style="background:rgba(255,255,255,0.1);border-radius:10px;padding:10px;">
    <span style="font-size:38px;line-height:1;">&#x1F3D7;</span>
  </div>
  <div>
   <h1 style="color:white;margin:0;font-size:2rem;font-weight:800;">{_emp_hdr} — Placa Facil</h1>
   <p style="color:#90caf9;margin:4px 0 0;font-size:0.95rem;">
     Proyecto: {_proy_hdr} &nbsp;·&nbsp; NSR-10 / E.060 / ACI 318
     &nbsp;·&nbsp; Perfil Colmena + Bloquelón Santafé
   </p>
  </div>
 </div>
</div>""", unsafe_allow_html=True)

# ── Onboarding (solo si faltan datos) ─────────────────────────────────────────
_has_id  = bool(st.session_state.get("empresa") and st.session_state.get("proyecto"))
_s1 = "[Listo]" if _has_id else "Paso 1:"
_s2 = "Paso 2:"
_s3 = "Paso 3:"
if not _has_id:
    st.info(
        "Guia de inicio rapido:\n\n"
        f"{_s1} Configura empresa y proyecto — abre 'Identidad del Proyecto' en el panel lateral\n\n"
        f"{_s2} Ajusta geometria, materiales y precios en el panel lateral\n\n"
        f"{_s3} Revisa resultados en los tabs y exporta DOCX, DXF, IFC o PDF"
    )

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR — IDENTIDAD, ROL, LOGO, NORMA, PARAMETROS
# ══════════════════════════════════════════════════════════════════════════════
st.sidebar.markdown("## Configuracion")

# ── Rol de usuario ─────────────────────────────────────────────────────────────
with st.sidebar.expander("Rol de Usuario", expanded=False):
    _roles = ["free", "pro", "admin"]
    _rol_idx = _roles.index(st.session_state.get("user_role", "free"))
    _rol_nuevo = st.radio(
        "Selecciona tu rol:", _roles, index=_rol_idx,
        key="pf_rol_radio",
        help="free: calculos + Excel/CSV. pro: todos los entregables. admin: acceso total.",
    )
    if _rol_nuevo != st.session_state.get("user_role"):
        st.session_state["user_role"] = _rol_nuevo
        save_state()
        st.rerun()
    _color_rol = {"admin": "#1b5e20", "pro": "#0d47a1", "free": "#b71c1c"}
    st.markdown(
        f'<div style="background:{_color_rol[_rol_nuevo]};color:white;border-radius:6px;'
        f'padding:5px 10px;font-size:12px;text-align:center;font-weight:600;margin-top:6px;">'
        f'Rol activo: {_rol_nuevo.upper()}</div>',
        unsafe_allow_html=True,
    )

# ── Identidad del proyecto ─────────────────────────────────────────────────────
with st.sidebar.expander("Identidad del Proyecto", expanded=False):
    _id_emp  = st.text_input("Empresa / Firma",
                             value=st.session_state.get("empresa", ""),
                             key="sid_empresa_pf", placeholder="________________")
    _id_proy = st.text_input("Nombre del Proyecto",
                             value=st.session_state.get("proyecto", ""),
                             key="sid_proyecto_pf", placeholder="________________")
    _id_ing  = st.text_input("Ingeniero Responsable",
                             value=st.session_state.get("ingeniero", ""),
                             key="sid_ingeniero_pf", placeholder="________________")
    _c1, _c2 = st.columns(2)
    _id_elab = _c1.text_input("Elaboro",  value=st.session_state.get("pf_elaboro", ""),  key="sid_elab_pf")
    _id_rev  = _c2.text_input("Reviso",   value=st.session_state.get("pf_reviso", ""),   key="sid_rev_pf")
    _id_apb  = st.text_input("Aprobo",    value=st.session_state.get("pf_aprobo", ""),   key="sid_apb_pf")
    if st.button("Guardar identidad", key="btn_id_pf", use_container_width=True):
        st.session_state["empresa"]    = _id_emp
        st.session_state["proyecto"]   = _id_proy
        st.session_state["ingeniero"]  = _id_ing
        st.session_state["pf_elaboro"] = _id_elab
        st.session_state["pf_reviso"]  = _id_rev
        st.session_state["pf_aprobo"]  = _id_apb
        save_state()
        st.success("Identidad guardada")
        st.rerun()

    # Logo
    st.markdown("---")
    st.markdown("**Logo de empresa**")
    st.caption("Aparece en portada de PDF y DOCX. PNG/JPG, fondo blanco recomendado.")
    _logo_file = st.file_uploader("Subir logo (PNG/JPG)",
                                  type=["png", "jpg", "jpeg"],
                                  key="pf_logo_upload")
    if _logo_file is not None:
        _raw = _logo_file.read()
        if _raw:
            st.session_state["logo_bytes"] = _raw
            save_state()
            st.success("Logo cargado")
    if st.session_state.get("logo_bytes"):
        try:
            st.image(st.session_state["logo_bytes"], width=150,
                     caption="Logo activo — aparece en PDF y DOCX")
        except Exception:
            st.caption("Logo cargado (no previsualizable)")
        _la, _lb = st.columns(2)
        if _la.button("Quitar logo", key="pf_rm_logo", use_container_width=True):
            st.session_state.pop("logo_bytes", None)
            save_state()
            st.rerun()
        _lb.download_button("Descargar", data=st.session_state["logo_bytes"],
                            file_name="logo.png", mime="image/png",
                            use_container_width=True, key="pf_dl_logo")
    else:
        st.caption("Sin logo — documentos solo texto.")

st.sidebar.markdown("---")

# ── Norma ──────────────────────────────────────────────────────────────────────
norma_sel = st.sidebar.selectbox(
    _t("Norma de diseno", "Design code"),
    list(NORMAS_PLACA.keys()), index=0, key="pf_norma_sel")
mostrar_referencias_norma(norma_sel, "placa_facil")
norma = NORMAS_PLACA[norma_sel]

# ── Datos del proyecto (plano) ─────────────────────────────────────────────────
st.sidebar.header(_t("Datos del proyecto", "Project data"))
proyecto_nombre    = st.sidebar.text_input(
    _t("Nombre del proyecto", "Project name"),
    value=st.session_state.get("pf_proyecto_nombre", ""),
    placeholder="________________", key="pf_proy_nom")
proyecto_direccion = st.sidebar.text_input(
    _t("Direccion de obra", "Site address"),
    value=st.session_state.get("pf_proyecto_dir", ""),
    placeholder="________________", key="pf_proy_dir")
proyecto_cliente   = st.sidebar.text_input(
    _t("Cliente / Propietario", "Client / Owner"),
    value=st.session_state.get("pf_proyecto_cliente", ""),
    placeholder="________________", key="pf_proy_cli")

# Sincronizar al session_state en cada rerun
st.session_state["pf_proyecto_nombre"]  = proyecto_nombre
st.session_state["pf_proyecto_dir"]     = proyecto_direccion
st.session_state["pf_proyecto_cliente"] = proyecto_cliente

num_pisos = st.sidebar.number_input(_t("Numero de pisos", "Number of floors"), 1, 10, 1, 1)
uso_placa = st.sidebar.selectbox(_t("Uso de la placa", "Slab use"), list(CARGA_VIVA.keys()), index=0)
carga_viva_kn = CARGA_VIVA[uso_placa]

# ── Geometria ──────────────────────────────────────────────────────────────────
st.sidebar.header(_t("Geometria de la placa", "Slab geometry"))
Lx = st.sidebar.number_input(_t("Luz X (m)", "Span X (m)"), 2.0, 12.0, 6.0, 0.1)
Ly = st.sidebar.number_input(_t("Luz Y (m)", "Span Y (m)"), 2.0, 12.0, 5.0, 0.1)
orientacion = st.sidebar.selectbox(
    _t("Direccion de los perfiles", "Profile direction"),
    ["Paralelo a X", "Paralelo a Y"], index=0)

# ── Diseno ─────────────────────────────────────────────────────────────────────
st.sidebar.header(_t("Parametros de diseno", "Design parameters"))
espesor_torta    = st.sidebar.number_input(
    _t("Espesor torta concreto (cm)", "Concrete topping (cm)"), 4.0, 10.0, 5.0, 0.5) / 100.0
fc_concreto      = st.sidebar.number_input(
    _t("Resistencia f'c (MPa)", "Concrete f'c (MPa)"), 18.0, 35.0, 21.0, 0.5)
perfil_espaciado = st.sidebar.number_input(
    _t("Separacion entre perfiles (cm)", "Profile spacing (cm)"), 70.0, 100.0, 89.0, 1.0) / 100.0
incluir_vigas    = st.sidebar.checkbox(
    _t("Incluir vigas de borde", "Include edge beams"), value=True)
viga_b = st.sidebar.number_input(
    _t("Ancho viga borde (cm)", "Edge beam width (cm)"), 10.0, 30.0, 15.0, 1.0) / 100.0
viga_h = st.sidebar.number_input(
    _t("Altura viga borde (cm)", "Edge beam height (cm)"), 15.0, 40.0, 20.0, 1.0) / 100.0

# ── Sismico ────────────────────────────────────────────────────────────────────
st.sidebar.header(_t("Parametros sismicos", "Seismic parameters"))
zona_sismica = st.sidebar.selectbox(
    _t("Zona sismica (Aa)", "Seismic zone"),
    ["I (Aa<0.10)", "II (0.10<=Aa<0.20)", "III (0.20<=Aa<0.30)", "IV (Aa>=0.30)"], index=1)
sistema_estructural = st.sidebar.selectbox(
    _t("Sistema estructural", "Structural system"),
    ["DMO (Desempeno Minimo)", "DES (Desempeno Especial)", "DMI (Desempeno Intermedio)"], index=0)
is_high_seismic = ("IV" in zona_sismica) or ("III" in zona_sismica)
if is_high_seismic and "DMO" in sistema_estructural:
    st.sidebar.warning("Zona alta: DMO no permitido segun NSR-10 C.21. Use DES o DMI.")

# ── Desperdicios ───────────────────────────────────────────────────────────────
st.sidebar.header(_t("Factores de desperdicio", "Waste factors"))
desp_bloques  = st.sidebar.number_input("Bloques (%)",   0.0, 20.0, 5.0,  1.0) / 100.0
desp_concreto = st.sidebar.number_input("Concreto (%)",  0.0, 20.0, 10.0, 1.0) / 100.0
desp_malla    = st.sidebar.number_input("Malla (%)",     0.0, 20.0, 10.0, 1.0) / 100.0
desp_perfiles = st.sidebar.number_input("Perfiles (%)",  0.0, 20.0, 5.0,  1.0) / 100.0

# ── APU precios (referencia Q1-2026 Colombia) ──────────────────────────────────
st.sidebar.header("APU — Precios unitarios")
moneda         = st.sidebar.text_input("Moneda", "COP", key="pf_moneda")
precio_bloque  = st.sidebar.number_input(
    "Precio bloque (und)", 3000.0, 20000.0, 7200.0, 100.0,
    help="Ref. mercado Colombia Q1-2026: COP 7,200")
precio_perfil  = st.sidebar.number_input(
    "Precio perfil (m lineal)", 15000.0, 60000.0, 28000.0, 1000.0,
    help="Ref. mercado Colombia Q1-2026: COP 28,000")
precio_malla   = st.sidebar.number_input(
    "Precio malla (m2)", 6000.0, 25000.0, 11000.0, 500.0,
    help="Ref. mercado Colombia Q1-2026: COP 11,000")
precio_concreto= st.sidebar.number_input(
    "Precio concreto (m3)", 200000.0, 700000.0, 450000.0, 10000.0,
    help="Ref. mercado Colombia Q1-2026: COP 450,000")
precio_mo      = st.sidebar.number_input(
    "Costo MO (dia)", 40000.0, 200000.0, 70000.0, 5000.0,
    help="Ref. SMLMV 2026 Colombia")
pct_herramienta= st.sidebar.number_input("% Herramienta menor", 0.0, 20.0, 5.0, 1.0) / 100.0
pct_aui        = st.sidebar.number_input("% A.I.U.", 0.0, 50.0, 30.0, 5.0) / 100.0
pct_util       = st.sidebar.number_input("% Utilidad", 0.0, 20.0, 5.0, 1.0) / 100.0
iva_pct        = st.sidebar.number_input("% IVA s/Utilidad", 0.0, 30.0, 19.0, 1.0) / 100.0

# Indicador de precios
_price_nota = "Precios de referencia Q1-2026 — verificar con cotizacion actual antes de contratar."
st.sidebar.markdown(
    f'<div style="background:#1a1a1a;border-radius:6px;padding:6px 10px;margin:4px 0 8px;'
    f'font-size:10px;color:#e65100;">[!] {_price_nota}</div>',
    unsafe_allow_html=True)
st.session_state["_price_status_label"] = _price_nota
st.session_state["_price_date"] = datetime.now().strftime("%d/%m/%Y")

# ── Datos plano ────────────────────────────────────────────────────────────────
st.sidebar.header("Datos del plano")
plano_numero = st.sidebar.text_input("Numero de plano",
    value=st.session_state.get("pf_plano_num", "PL-01"), key="pf_plano")
escala_plano = st.sidebar.text_input("Escala",
    value=st.session_state.get("pf_escala", "1:50"), key="pf_escala_inp")
elaboro  = st.session_state.get("pf_elaboro", "") or "________________"
revisado = st.session_state.get("pf_reviso",  "") or "________________"
aprobado = st.session_state.get("pf_aprobo",  "") or "________________"
st.session_state["pf_plano_num"] = plano_numero
st.session_state["pf_escala"]    = escala_plano

st.sidebar.markdown(f'<div style="text-align:center;color:gray;font-size:10px;">'
                    f'© {datetime.now().year} {_emp_hdr} — Uso profesional exclusivo</div>',
                    unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# FUNCIONES DE CALCULO
# ══════════════════════════════════════════════════════════════════════════════
def profile_inertia():
    h = PROFILE_DATA["alto_total"]; bw = PROFILE_DATA["ancho_alma"]
    tw = PROFILE_DATA["espesor_alma"]; bf = PROFILE_DATA["ancho_ala"]
    tf = PROFILE_DATA["espesor_ala"]
    hc = h*100; bwc = bw*100; twc = tw*100; bfc = bf*100; tfc = tf*100
    Aw = bwc*twc; Af = bfc*tfc
    yw = hc/2; yf = tfc/2
    Iw = (bwc*twc**3)/12; If = (bfc*tfc**3)/12
    dw = yw-(hc/2); df = yf-(hc/2)
    return ((Iw+Aw*dw**2)+2*(If+Af*df**2))/10000


def select_rebars(As_req_cm2):
    table = [
        {"nom": '#3 (3/8")', "area": 0.71},
        {"nom": '#4 (1/2")', "area": 1.29},
        {"nom": '#5 (5/8")', "area": 1.99},
        {"nom": '#6 (3/4")', "area": 2.87},
        {"nom": '#7 (7/8")', "area": 3.87},
        {"nom": '#8 (1")',   "area": 5.10},
    ]
    best = None
    for b1 in table:
        for b2 in table:
            t = b1["area"] + b2["area"]
            if t >= As_req_cm2 and (best is None or t < best["total"]):
                best = {"b1": b1, "b2": b2, "total": t}
    if best is None:
        best = {"b1": table[-1], "b2": table[-1], "total": table[-1]["area"]*2}
    return best

# ══════════════════════════════════════════════════════════════════════════════
# CALCULOS PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════
area_total = Lx * Ly * num_pisos
if orientacion == "Paralelo a X":
    perfil_largo = Lx; perfil_ancho = Ly
else:
    perfil_largo = Ly; perfil_ancho = Lx

n_profiles = math.ceil(perfil_ancho / perfil_espaciado) + 1
longitud_total_perfiles = n_profiles * perfil_largo
longitud_total_perfiles_desp = longitud_total_perfiles * (1 + desp_perfiles) * num_pisos

n_hileras = n_profiles - 1
bloques_por_hilera = math.ceil((Lx if orientacion=="Paralelo a X" else Ly) / BLOCK_DATA["largo"])
n_bloques = n_hileras * bloques_por_hilera
n_bloques_desp = math.ceil(n_bloques * (1 + desp_bloques)) * num_pisos

vol_torta = area_total * espesor_torta
vol_vigas = (2*(Lx+Ly)) * viga_b * viga_h * num_pisos if incluir_vigas else 0
vol_concreto_total = vol_torta + vol_vigas
vol_concreto_total_desp = vol_concreto_total * (1 + desp_concreto)
area_malla = area_total * (1 + desp_malla + 0.15)

peso_bloques_kg  = n_bloques * BLOCK_DATA["peso_unitario"]
peso_concreto_kg = vol_concreto_total * CONCRETE_DATA["densidad"]
peso_perfiles_kg = longitud_total_perfiles * PROFILE_DATA["peso_por_m"] * num_pisos
peso_total_kg    = peso_bloques_kg + peso_concreto_kg + peso_perfiles_kg
carga_muerta_kgm2 = peso_total_kg / (Lx*Ly*num_pisos) if (Lx*Ly*num_pisos) > 0 else 0

total_cemento_kg = CONCRETE_DATA["cemento_por_m3"] * vol_concreto_total_desp
bultos_cemento   = math.ceil(total_cemento_kg / 50)

# Estructural
g = 9.81
Wd_kn  = carga_muerta_kgm2 * g / 1000
Wu_kn  = 1.2*Wd_kn + 1.6*carga_viva_kn
Wu_lin = Wu_kn * perfil_espaciado
Mu     = Wu_lin * perfil_largo**2 / 8

I_real     = profile_inertia()
E_concreto = 4700 * math.sqrt(fc_concreto) * 1e6
Wserv_lin  = (Wd_kn + carga_viva_kn) * perfil_espaciado
delta_max  = perfil_largo / 360
delta_calc = (5*Wserv_lin*perfil_largo**4)/(384*E_concreto*I_real) if I_real > 0 else 0
cumple_deflexion = delta_calc <= delta_max

As_min   = 0.0018 * BLOCK_DATA["ancho"] * espesor_torta
As_malla = MESH_DATA["diametro"]**2 * math.pi / 4 * 1000 / MESH_DATA["espaciado_largo"]
cumple_malla = As_malla >= As_min

Vu = Wu_lin * perfil_largo / 2
Vc = 0.17 * math.sqrt(fc_concreto*1000) * BLOCK_DATA["ancho"] * (BLOCK_DATA["alto"]+espesor_torta) / 1000
cumple_cortante = Vu <= Vc

# Viga de borde
W_beam_kn = Wu_kn * (perfil_espaciado/2 + 0.5)
beam_span  = max(Lx, Ly)
Mu_beam    = W_beam_kn * beam_span**2 / 8
Vu_beam    = W_beam_kn * beam_span / 2

d_beam = viga_h - 0.05
Rn     = Mu_beam / (0.9*viga_b*d_beam**2)*1e6 if (viga_b*d_beam**2) > 0 else 0
rho    = (0.85*fc_concreto/420)*(1-math.sqrt(max(0,1-2*Rn/(0.85*fc_concreto)))) if Rn > 0 else 0
As_beam     = max(rho*viga_b*d_beam*1e4, 0.0018*viga_b*viga_h*1e4)
rebar_combo = select_rebars(As_beam)
As_prov_beam = rebar_combo["total"]
ref_beam    = f"1{rebar_combo['b1']['nom']} + 1{rebar_combo['b2']['nom']} (As={As_prov_beam:.2f} cm2)"

stirrup_factor = 0.25 if (is_high_seismic and "DMO" not in sistema_estructural) else 0.50
confinement_zone_length = 2.0*max(viga_h,0.50) if stirrup_factor == 0.25 else 0.0
s_beam = round(min(stirrup_factor*d_beam, 0.30), 2)

Vc_beam = 0.17*math.sqrt(fc_concreto*1000)*viga_b*d_beam/1000
cumple_cortante_beam = Vu_beam <= Vc_beam
cumple_flexion_beam  = As_prov_beam >= As_beam

x_vals = np.linspace(0, beam_span, 100)
M_vals = W_beam_kn * x_vals * (beam_span - x_vals) / 2
V_vals = W_beam_kn * (beam_span/2 - x_vals)

# Seccion compuesta
beff    = perfil_espaciado
E_acero = 200e9
n_mod   = E_acero / E_concreto
A_prof  = PROFILE_DATA["ancho_alma"]*PROFILE_DATA["alto_total"] + 2*PROFILE_DATA["ancho_ala"]*PROFILE_DATA["espesor_ala"]
y_prof  = PROFILE_DATA["alto_total"] / 2.0
A_slab  = beff * espesor_torta
y_slab  = PROFILE_DATA["alto_total"] + espesor_torta/2.0
A_pt    = A_prof * n_mod
y_c     = (A_pt*y_prof + A_slab*y_slab) / (A_pt + A_slab)
I_comp  = (profile_inertia()*n_mod + A_pt*(y_prof-y_c)**2 +
           (beff*espesor_torta**3)/12 + A_slab*(y_slab-y_c)**2)
S_comp  = I_comp/y_c if y_c > 0 else 0.0

# Longitud desarrollo
fy=420.0; db_ref=0.0127
ld = max((fy*db_ref)/(1.1*math.sqrt(fc_concreto))/1000, 12*db_ref, 0.30)

# Verificaciones normativas
verificaciones = [
    {"item":"Luz maxima perfiles","referencia":norma["ref"],
     "requerido":f"<= {norma['luz_max']:.2f} m","calculado":f"{perfil_largo:.2f} m",
     "cumple":perfil_largo<=norma["luz_max"],
     "obs":"Ok" if perfil_largo<=norma["luz_max"] else "Excede — requiere viga intermedia"},
    {"item":"Espesor torta concreto","referencia":"NSR-10 C.21.6.4.1",
     "requerido":f">= {norma['topping_min']*100:.0f} cm","calculado":f"{espesor_torta*100:.1f} cm",
     "cumple":espesor_torta>=norma["topping_min"],
     "obs":"Ok" if espesor_torta>=norma["topping_min"] else "Incrementar espesor"},
    {"item":"Resistencia concreto","referencia":"NSR-10 C.21.3.1",
     "requerido":f">= {norma['concreto_min']} MPa","calculado":f"{fc_concreto:.1f} MPa",
     "cumple":fc_concreto>=norma["concreto_min"],
     "obs":"Ok" if fc_concreto>=norma["concreto_min"] else "Usar concreto mayor resistencia"},
    {"item":"Deflexion (L/360)","referencia":"NSR-10 C.9.5.2",
     "requerido":f"<= {delta_max*1000:.1f} mm","calculado":f"{delta_calc*1000:.1f} mm",
     "cumple":cumple_deflexion,"obs":"Ok" if cumple_deflexion else "Aumentar peralte o reducir luz"},
    {"item":"Cortante en apoyo (placa)","referencia":"NSR-10 C.11",
     "requerido":f"Vu <= {Vc:.2f} kN","calculado":f"Vu = {Vu:.2f} kN",
     "cumple":cumple_cortante,"obs":"Ok" if cumple_cortante else "Aumentar seccion de viga"},
    {"item":"Acero minimo de malla","referencia":"NSR-10 C.7.12",
     "requerido":f"As >= {As_min*100:.2f} cm2/m","calculado":f"As = {As_malla:.2f} cm2/m",
     "cumple":cumple_malla,"obs":"Ok" if cumple_malla else "Aumentar diametro o reducir espaciamiento"},
    {"item":"Viga borde - Momento","referencia":"NSR-10 C.21.6",
     "requerido":f"As >= {As_beam:.2f} cm2","calculado":f"As prov = {As_prov_beam:.2f} cm2",
     "cumple":cumple_flexion_beam,"obs":"Ok" if cumple_flexion_beam else "Aumentar acero"},
    {"item":"Viga borde - Cortante","referencia":"NSR-10 C.11",
     "requerido":f"Vu <= {Vc_beam:.2f} kN","calculado":f"Vu = {Vu_beam:.2f} kN",
     "cumple":cumple_cortante_beam,"obs":"Ok" if cumple_cortante_beam else "Aumentar seccion o estribos"},
    {"item":"Altura total de la placa","referencia":"Practica constructiva",
     "requerido":">= 13 cm","calculado":f"{(BLOCK_DATA['alto']+espesor_torta)*100:.1f} cm",
     "cumple":(BLOCK_DATA["alto"]+espesor_torta)>=0.13,
     "obs":"Ok" if (BLOCK_DATA["alto"]+espesor_torta)>=0.13 else "Aumentar espesor"},
    {"item":"Vigas de borde","referencia":"NSR-10 C.21.6.4",
     "requerido":"Incluidas" if norma.get("requiere_viga_borde") else "Opcional",
     "calculado":"Si" if incluir_vigas else "No",
     "cumple":incluir_vigas or not norma.get("requiere_viga_borde"),
     "obs":"Ok" if (incluir_vigas or not norma.get("requiere_viga_borde")) else "Incluir vigas de borde"},
    {"item":"Zona sismica y confinamiento","referencia":"NSR-10 C.21",
     "requerido":f"Estribos {'d/4' if stirrup_factor==0.25 else 'd/2'} @ {s_beam*100:.0f} cm",
     "calculado":f"Zona {zona_sismica} | {sistema_estructural}",
     "cumple":not (is_high_seismic and "DMO" in sistema_estructural),
     "obs":"Confinamiento requerido" if stirrup_factor==0.25 else "Ok"},
]

# APU
costo_bloques  = n_bloques_desp * precio_bloque
costo_perfiles = longitud_total_perfiles_desp * precio_perfil
costo_malla    = area_malla * precio_malla
costo_concreto = vol_concreto_total_desp * precio_concreto
dias_mo        = area_total * 0.8
costo_mo       = dias_mo * precio_mo
costo_directo  = costo_bloques + costo_perfiles + costo_malla + costo_concreto + costo_mo
herramienta    = costo_mo * pct_herramienta
aiu            = costo_directo * pct_aui
utilidad       = costo_directo * pct_util
iva_util       = utilidad * iva_pct
total_proyecto = costo_directo + herramienta + aiu + iva_util
costo_maciza   = vol_concreto_total * 1.2 * precio_concreto
ahorro         = max(0, costo_maciza - total_proyecto)
sobrecosto     = max(0, total_proyecto - costo_maciza)

actividades   = ["Instalacion de perfiles","Colocacion de bloques",
                 "Colocacion de malla","Fundida de concreto","Curado"]
duracion_dias = [0.5*area_total/100, 1.0*area_total/100,
                 0.3*area_total/100, 1.5*area_total/100, 0.5]
cronograma    = pd.DataFrame({"Actividad": actividades,
                               "Duracion (dias)": [max(1, round(d)) for d in duracion_dias]})

# Alertas de luz
if perfil_largo > 3.5 and perfil_largo <= norma["luz_max"]:
    st.warning(f"Luz de perfiles {perfil_largo:.2f} m — cerca del limite {norma['luz_max']} m.")
elif perfil_largo > norma["luz_max"]:
    st.error(f"Luz de perfiles {perfil_largo:.2f} m excede limite {norma['luz_max']} m. Requiere viga intermedia.")

# ══════════════════════════════════════════════════════════════════════════════
# MODELO 3D
# ══════════════════════════════════════════════════════════════════════════════
def create_3d_model(Lx, Ly, orientacion, n_profiles, sep, perfil_largo,
                    esp_t, inc_vigas, vb, vh):
    fig = go.Figure()

    # ── Helper: caja sólida con hover ─────────────────────────────────────────
    def _box(x0, y0, z0, dx, dy, dz, color, op, name, htext):
        xs = [x0, x0+dx, x0+dx, x0,    x0,    x0+dx, x0+dx, x0   ]
        ys = [y0, y0,    y0+dy, y0+dy,  y0,    y0,    y0+dy, y0+dy]
        zs = [z0, z0,    z0,    z0,     z0+dz, z0+dz, z0+dz, z0+dz]
        ii = [0, 0, 4, 4, 1, 5, 2, 6, 3, 7, 0, 4]
        jj = [1, 2, 5, 6, 5, 6, 6, 7, 7, 4, 4, 7]
        kk = [2, 3, 6, 7, 6, 7, 7, 4, 4, 5, 1, 3]
        fig.add_trace(go.Mesh3d(
            x=xs, y=ys, z=zs, i=ii, j=jj, k=kk,
            color=color, opacity=op,
            name=name, showlegend=False,
            hovertemplate=htext + "<extra></extra>",
            flatshading=True,
            lighting=dict(ambient=0.7, diffuse=0.6, specular=0.3),
        ))

    # ── Dimensiones perfil colmena ────────────────────────────────────────────
    h_p = PROFILE_DATA["alto_total"]        # 0.090 m
    ba  = PROFILE_DATA["ancho_ala"]         # 0.020 m (ala c/lado del alma)
    ta  = PROFILE_DATA["espesor_ala"]       # 0.002 m
    bw  = PROFILE_DATA["ancho_alma"]        # 0.040 m (anchura de cada alma)
    tw  = PROFILE_DATA["espesor_alma"]      # 0.002 m
    bt  = PROFILE_DATA["b_total"]           # 0.130 m total
    # Sección: |ala_izq|alma_izq|hueco|alma_der|ala_der|
    # bt = ba + tw + (bw - tw) + tw + ba  ≈ 2*ba + bw + 2*tw  pero usamos bt directo

    pc = "#B0B0B0"   # color perfil claro
    pd = "#909090"   # color perfil oscuro (alas)

    p_hover = (
        "<b>Perfil Colmena</b><br>"
        f"h = {h_p*1000:.0f} mm<br>"
        f"b_total = {bt*1000:.0f} mm<br>"
        f"alma = {bw*1000:.0f} mm × {tw*1000:.1f} mm<br>"
        f"ala  = {ba*1000:.0f} mm × {ta*1000:.1f} mm<br>"
        f"Peso = {PROFILE_DATA['peso_por_m']:.1f} kg/m | Lámina Cal.18"
    )

    def _add_profile_x(yc, length):
        """Perfil corriendo en X, centrado en yc."""
        # Alma izquierda (exterior izq)
        _box(0, yc - bt/2,          0,      length, tw, h_p, pc, 0.97, "Perfil Colmena", p_hover)
        # Alma derecha  (exterior der)
        _box(0, yc + bt/2 - tw,     0,      length, tw, h_p, pc, 0.97, "Perfil Colmena", p_hover)
        # Ala inferior
        _box(0, yc - bt/2,          0,      length, bt, ta,  pd, 0.97, "Perfil Colmena", p_hover)
        # Ala superior
        _box(0, yc - bt/2,          h_p-ta, length, bt, ta,  pd, 0.97, "Perfil Colmena", p_hover)

    def _add_profile_y(xc, length):
        """Perfil corriendo en Y, centrado en xc."""
        _box(xc - bt/2,          0, 0,      tw, length, h_p, pc, 0.97, "Perfil Colmena", p_hover)
        _box(xc + bt/2 - tw,     0, 0,      tw, length, h_p, pc, 0.97, "Perfil Colmena", p_hover)
        _box(xc - bt/2,          0, 0,      bt, length, ta,  pd, 0.97, "Perfil Colmena", p_hover)
        _box(xc - bt/2,          0, h_p-ta, bt, length, ta,  pd, 0.97, "Perfil Colmena", p_hover)

    # ── Bloquelón Santafé ─────────────────────────────────────────────────────
    BL = BLOCK_DATA["largo"]    # 0.80 m
    BW_b = BLOCK_DATA["ancho"]  # 0.23 m
    BH = BLOCK_DATA["alto"]     # 0.08 m
    bc = "#C4884A"

    b_hover = (
        "<b>Bloquelón Santafé</b><br>"
        f"L×A×H = {BL*100:.0f}×{BW_b*100:.0f}×{BH*100:.0f} cm<br>"
        f"Peso = {BLOCK_DATA['peso_unitario']} kg/und<br>"
        f"Rend. = {BLOCK_DATA['rendimiento_por_m2']:.2f} und/m²<br>"
        "Material: Arcilla cocida"
    )

    def _blocks_strip_x(yc1, yc2, max_b=50):
        """Hilera de bloquelones entre dos perfiles paralelos a X.
        largo=0.80m cruza en Y (de perfil a perfil)
        ancho=0.23m va en X (a lo largo del perfil).
        """
        y0b = yc1 + bt/2       # borde exterior ala perfil izq
        y1b = yc2 - bt/2       # borde exterior ala perfil der
        span_y = y1b - y0b     # ≈ BL = 0.80m (bloquelón cruza de perfil a perfil)
        if span_y < 0.05:
            return
        # Cantidad de bloques en X a paso de BW_b (0.23m), limitados
        nb = min(math.ceil(Lx / BW_b), max_b)
        for j in range(nb):
            x0b = j * BW_b
            x1b = min(x0b + BW_b, Lx)
            _box(x0b, y0b, ta, x1b - x0b, span_y, BH, bc, 0.92, "Bloquelón", b_hover)

    def _blocks_strip_y(xc1, xc2, max_b=50):
        """Hilera de bloquelones entre dos perfiles paralelos a Y.
        largo=0.80m cruza en X (de perfil a perfil)
        ancho=0.23m va en Y (a lo largo del perfil).
        """
        x0b = xc1 + bt/2
        x1b = xc2 - bt/2
        span_x = x1b - x0b     # ≈ BL = 0.80m
        if span_x < 0.05:
            return
        nb = min(math.ceil(Ly / BW_b), max_b)
        for j in range(nb):
            y0b = j * BW_b
            y1b = min(y0b + BW_b, Ly)
            _box(x0b, y0b, ta, span_x, y1b - y0b, BH, bc, 0.92, "Bloquelón", b_hover)

    # ── Renderizar según orientación ──────────────────────────────────────────
    max_blk_per_bay = max(3, 200 // max(n_profiles - 1, 1))

    if orientacion == "Paralelo a X":
        yp = np.linspace(0, Ly, n_profiles)
        for yc in yp:
            _add_profile_x(yc, Lx)
        for i in range(len(yp) - 1):
            _blocks_strip_x(yp[i], yp[i+1], max_b=max_blk_per_bay)
    else:
        xp = np.linspace(0, Lx, n_profiles)
        for xc in xp:
            _add_profile_y(xc, Ly)
        for i in range(len(xp) - 1):
            _blocks_strip_y(xp[i], xp[i+1], max_b=max_blk_per_bay)

    # ── Torta de concreto (semitransparente) ──────────────────────────────────
    z_torta = h_p + BH + ta
    t_hover = (
        "<b>Torta de Concreto</b><br>"
        f"f'c = 21 MPa<br>"
        f"e = {esp_t*100:.0f} cm<br>"
        f"Malla Q2 Ø{MESH_DATA['diametro']*1000:.1f}mm<br>"
        f"@{MESH_DATA['espaciado_largo']*100:.0f}cm"
    )
    _box(0, 0, z_torta, Lx, Ly, esp_t, "#C8C8C8", 0.32, "Torta concreto", t_hover)

    # ── Malla electrosoldada ──────────────────────────────────────────────────
    ms_l = MESH_DATA["espaciado_largo"]
    ms_c = MESH_DATA["espaciado_corto"]
    z_malla = z_torta + 0.004
    mx, my, mz = [], [], []
    for ym in np.arange(0, Ly + ms_l, ms_l):
        mx += [0, Lx, None]; my += [ym, ym, None]; mz += [z_malla, z_malla, None]
    for xm in np.arange(0, Lx + ms_c, ms_c):
        mx += [xm, xm, None]; my += [0, Ly, None]; mz += [z_malla, z_malla, None]
    fig.add_trace(go.Scatter3d(
        x=mx, y=my, z=mz, mode="lines",
        line=dict(color="#333333", width=1),
        name="Malla Q2",
        showlegend=True,
        hovertemplate=(
            f"<b>Malla Electrosoldada Q2</b><br>"
            f"Ø {MESH_DATA['diametro']*1000:.1f} mm<br>"
            f"@{ms_l*100:.0f}cm long × @{ms_c*100:.0f}cm trans<br>"
            f"Peso = {MESH_DATA['peso_por_m2']:.1f} kg/m²<extra></extra>"
        ),
    ))

    # ── Vigas de borde ────────────────────────────────────────────────────────
    if inc_vigas:
        v_hover = (
            "<b>Viga de Borde</b><br>"
            f"b = {vb*100:.0f} cm  h = {vh*100:.0f} cm<br>"
            "f'c = 21 MPa | Concreto reforzado"
        )
        # Sur y Norte
        _box(0, -vb,     -vh, Lx, vb, vh, "#6A6A6A", 0.88, "Viga borde", v_hover)
        _box(0, Ly,      -vh, Lx, vb, vh, "#6A6A6A", 0.88, "Viga borde", v_hover)
        # Oeste y Este
        _box(-vb, -vb,   -vh, vb, Ly+2*vb, vh, "#6A6A6A", 0.88, "Viga borde", v_hover)
        _box(Lx,  -vb,   -vh, vb, Ly+2*vb, vh, "#6A6A6A", 0.88, "Viga borde", v_hover)

    # ── Leyenda manual ────────────────────────────────────────────────────────
    legends = [
        (f"Perfil Colmena h={h_p*1000:.0f}mm b={bt*1000:.0f}mm", "#B0B0B0", "square"),
        (f"Bloquelón Santafé {BL*100:.0f}×{BW_b*100:.0f}×{BH*100:.0f}cm",  "#C4884A", "square"),
        (f"Torta f'c=21MPa e={esp_t*100:.0f}cm", "#C8C8C8", "square"),
    ]
    if inc_vigas:
        legends.append((f"Viga borde {vb*100:.0f}×{vh*100:.0f}cm", "#6A6A6A", "square"))
    for lname, lcolor, lsym in legends:
        fig.add_trace(go.Scatter3d(
            x=[None], y=[None], z=[None], mode="markers",
            marker=dict(size=10, color=lcolor, symbol=lsym),
            name=lname, showlegend=True,
        ))

    # ── Layout ────────────────────────────────────────────────────────────────
    h_total = z_torta + esp_t
    fig.update_layout(
        scene=dict(
            xaxis=dict(title="X (m)", backgroundcolor="#f4f4f4",
                       gridcolor="#cccccc", showbackground=True, nticks=6),
            yaxis=dict(title="Y (m)", backgroundcolor="#efefef",
                       gridcolor="#cccccc", showbackground=True, nticks=6),
            zaxis=dict(title="Z (m)", backgroundcolor="#eaeaea",
                       gridcolor="#cccccc", showbackground=True, nticks=4),
            aspectmode="data",
            bgcolor="#ffffff",
            camera=dict(
                up=dict(x=0, y=0, z=1),
                eye=dict(x=1.5, y=-2.0, z=1.0),
                center=dict(x=0, y=0, z=0),
            ),
        ),
        legend=dict(
            x=0.01, y=0.99,
            bgcolor="rgba(255,255,255,0.90)",
            bordercolor="#aaaaaa", borderwidth=1,
            font=dict(size=11, color="#111111"),
            traceorder="normal",
        ),
        margin=dict(l=0, r=0, b=0, t=40),
        height=560,
        paper_bgcolor="#ffffff",
        font=dict(color="#111111", size=12),
        title=dict(
            text=(
                f"<b>Placa Aligerada</b> — Perfil Colmena + Bloquelón Santafé<br>"
                f"<sup>Lx={Lx:.2f}m × Ly={Ly:.2f}m | "
                f"h_total={h_total*100:.0f}cm | {n_profiles} perfiles@{sep:.2f}m | "
                f"Torta={esp_t*100:.0f}cm</sup>"
            ),
            x=0.5, xanchor="center",
            font=dict(size=13, color="#111111"),
        ),
    )
    return fig


# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DXF PROFESIONAL (sin ezdxf — stdlib pura, Color 7 ICONTEC)
# ══════════════════════════════════════════════════════════════════════════════
def _dxf_writer():
    """Mini-writer DXF R2010 ASCII."""
    layers = {}
    entities = []
    handle = [200]

    def _h():
        handle[0] += 1
        return f"{handle[0]:X}"

    def layer(name, color=7, lw=25):
        layers[name] = (color, lw)

    def line(x1,y1,x2,y2,lay="0"):
        entities.append(
            f"  0\nLINE\n  5\n{_h()}\n100\nAcDbEntity\n  8\n{lay}\n"
            f"100\nAcDbLine\n 10\n{x1:.5f}\n 20\n{y1:.5f}\n 30\n0.0\n"
            f" 11\n{x2:.5f}\n 21\n{y2:.5f}\n 31\n0.0\n")

    def poly(pts,lay="0",closed=True):
        n=len(pts); flag=1 if closed else 0
        b=(f"  0\nLWPOLYLINE\n  5\n{_h()}\n100\nAcDbEntity\n  8\n{lay}\n"
           f"100\nAcDbPolyline\n 90\n{n}\n 70\n{flag}\n")
        for x,y in pts: b+=f" 10\n{x:.5f}\n 20\n{y:.5f}\n"
        entities.append(b)

    def circle(cx,cy,r,lay="0"):
        entities.append(
            f"  0\nCIRCLE\n  5\n{_h()}\n100\nAcDbEntity\n  8\n{lay}\n"
            f"100\nAcDbCircle\n 10\n{cx:.5f}\n 20\n{cy:.5f}\n 30\n0.0\n 40\n{r:.5f}\n")

    def txt(x,y,s,h=0.10,lay="0",ha=0,va=0):
        s=str(s).replace("\n"," ")
        ap=f" 11\n{x:.5f}\n 21\n{y:.5f}\n 31\n0.0\n" if (ha or va) else ""
        entities.append(
            f"  0\nTEXT\n  5\n{_h()}\n100\nAcDbEntity\n  8\n{lay}\n"
            f"100\nAcDbText\n 10\n{x:.5f}\n 20\n{y:.5f}\n 30\n0.0\n"
            f" 40\n{h:.5f}\n  1\n{s}\n 72\n{ha}\n 73\n{va}\n{ap}100\nAcDbText\n")

    def render(xmin=-2.0, ymin=-9.0, xmax=22.0, ymax=8.0):
        # Centro de la vista y altura de vista para VPORT
        cx = (xmin + xmax) / 2.0
        cy = (ymin + ymax) / 2.0
        vw = xmax - xmin   # ancho total
        vh = ymax - ymin   # alto total
        aspect = vw / vh if vh > 0 else 1.5

        b=io.StringIO()
        # ── HEADER ──────────────────────────────────────────────────────────
        b.write("  0\nSECTION\n  2\nHEADER\n")
        b.write("  9\n$ACADVER\n  1\nAC1024\n")
        b.write("  9\n$INSUNITS\n 70\n6\n")        # 6 = metros
        b.write("  9\n$MEASUREMENT\n 70\n1\n")      # 1 = métrico
        b.write("  9\n$LUNITS\n 70\n2\n")           # unidades decimales
        b.write("  9\n$LUPREC\n 70\n4\n")
        b.write(f"  9\n$EXTMIN\n 10\n{xmin:.4f}\n 20\n{ymin:.4f}\n 30\n0.0\n")
        b.write(f"  9\n$EXTMAX\n 10\n{xmax:.4f}\n 20\n{ymax:.4f}\n 30\n0.0\n")
        b.write(f"  9\n$LIMMIN\n 10\n{xmin:.4f}\n 20\n{ymin:.4f}\n")
        b.write(f"  9\n$LIMMAX\n 10\n{xmax:.4f}\n 20\n{ymax:.4f}\n")
        b.write(f"  9\n$VIEWCTR\n 10\n{cx:.4f}\n 20\n{cy:.4f}\n")
        b.write(f"  9\n$VIEWSIZE\n 40\n{vh:.4f}\n")
        b.write("  0\nENDSEC\n")
        # ── TABLES ──────────────────────────────────────────────────────────
        b.write("  0\nSECTION\n  2\nTABLES\n")
        # LTYPE
        b.write("  0\nTABLE\n  2\nLTYPE\n 70\n2\n")
        for lt in [("Continuous",""),("DASHDOT","")]:
            b.write(f"  0\nLTYPE\n  2\n{lt[0]}\n 70\n0\n  3\n\n 72\n65\n 73\n0\n 40\n0.0\n")
        b.write("  0\nENDTAB\n")
        # VPORT — viewport activo con zoom al extents
        b.write("  0\nTABLE\n  2\nVPORT\n 70\n1\n")
        b.write("  0\nVPORT\n")
        b.write("  2\n*ACTIVE\n")
        b.write(" 70\n0\n")
        b.write(" 10\n0.0\n 20\n0.0\n")         # lower-left corner (0,0)
        b.write(" 11\n1.0\n 21\n1.0\n")         # upper-right corner (1,1)
        b.write(f" 12\n{cx:.5f}\n 22\n{cy:.5f}\n")  # view center
        b.write(" 13\n0.0\n 23\n0.0\n")         # snap base
        b.write(" 14\n1.0\n 24\n1.0\n")         # snap spacing
        b.write(" 15\n1.0\n 25\n1.0\n")         # grid spacing
        b.write(" 16\n0.0\n 26\n0.0\n 36\n1.0\n")  # view direction (plan view)
        b.write(" 17\n0.0\n 27\n0.0\n 37\n0.0\n")  # view target
        b.write(f" 40\n{vh:.5f}\n")               # VIEW HEIGHT (clave para zoom)
        b.write(f" 41\n{aspect:.5f}\n")            # aspect ratio
        b.write(" 42\n50.0\n 43\n0.0\n 44\n0.0\n")
        b.write(" 50\n0.0\n 51\n0.0\n")
        b.write(" 71\n0\n 72\n100\n 73\n1\n 74\n3\n 75\n0\n 76\n1\n 77\n0\n 78\n0\n")
        b.write("  0\nENDTAB\n")
        # LAYER
        b.write(f"  0\nTABLE\n  2\nLAYER\n 70\n{len(layers)}\n")
        for nm,(col,lw) in layers.items():
            b.write(f"  0\nLAYER\n  2\n{nm}\n 70\n0\n 62\n{col}\n  6\nContinuous\n370\n{lw}\n")
        b.write("  0\nENDTAB\n")
        # STYLE
        b.write("  0\nTABLE\n  2\nSTYLE\n 70\n1\n")
        b.write("  0\nSTYLE\n  2\nSTANDARD\n 70\n0\n 40\n0.0\n 41\n1.0\n 42\n0.10\n  3\ntxt\n  4\n\n")
        b.write("  0\nENDTAB\n  0\nENDSEC\n")
        # ── BLOCKS ──────────────────────────────────────────────────────────
        b.write("  0\nSECTION\n  2\nBLOCKS\n")
        b.write("  0\nBLOCK\n  8\n0\n  2\n*Model_Space\n 70\n0\n 10\n0.0\n 20\n0.0\n 30\n0.0\n  3\n*Model_Space\n  1\n\n")
        b.write("  0\nENDBLK\n  8\n0\n  0\nENDSEC\n")
        # ── ENTITIES ────────────────────────────────────────────────────────
        b.write("  0\nSECTION\n  2\nENTITIES\n")
        for e in entities: b.write(e)
        b.write("  0\nENDSEC\n  0\nEOF\n")
        return b.getvalue().encode("utf-8")

    return layer, line, poly, circle, txt, render


def generate_dxf_profesional():
    """DXF profesional con 6 detalles escalados 1:10 — ICONTEC 2289 Color 7."""
    layer, line, poly, circle, txt, render = _dxf_writer()

    for nm, lw in [("CONTORNO",50),("CONCRETO",50),("PERFILES",35),
                   ("ACERO_LONG",35),("ACERO_TRANS",25),("MALLA",18),
                   ("COTAS",18),("TEXTO",18),("EJES",13),
                   ("ROTULO",35),("NOTAS",18),("ESTRIBOS",25)]:
        layer(nm, 7, lw)

    TH=0.10; THB=0.13; THT=0.18; THG=0.24
    S = 10.0   # escala detalle: 1 mm real -> 0.010 m en plano (= 10:1)

    def cota_h(x1,y1,x2,y2,offset,lab):
        line(x1,y1+offset*0.85,x2,y1+offset*0.85,lay="COTAS")
        line(x1,y1,x1,y1+offset*0.85,lay="COTAS")
        line(x2,y2,x2,y2+offset*0.85,lay="COTAS")
        txt((x1+x2)/2,y1+offset*1.0,lab,h=TH,lay="COTAS",ha=1)

    def cota_v(x1,y1,x2,y2,offset,lab):
        line(x1+offset*0.85,y1,x1+offset*0.85,y2,lay="COTAS")
        line(x1,y1,x1+offset*0.85,y1,lay="COTAS")
        line(x2,y2,x2+offset*0.85,y2,lay="COTAS")
        txt(x1+offset*1.0,(y1+y2)/2,lab,h=TH,lay="COTAS",ha=1)

    # ZONA 1: Planta general
    poly([(0,0),(Lx,0),(Lx,Ly),(0,Ly)], lay="CONTORNO")
    if incluir_vigas:
        hb = viga_b/2
        for y0 in [0,Ly]: poly([(0,y0-hb),(Lx,y0-hb),(Lx,y0+hb),(0,y0+hb)],lay="ACERO_LONG")
        for x0 in [0,Lx]: poly([(x0-hb,0),(x0-hb,Ly),(x0+hb,Ly),(x0+hb,0)],lay="ACERO_LONG")
    if orientacion == "Paralelo a X":
        for y in np.linspace(0, Ly, n_profiles): line(0,y,Lx,y,lay="PERFILES")
    else:
        for x in np.linspace(0, Lx, n_profiles): line(x,0,x,Ly,lay="PERFILES")
    xm = MESH_DATA["espaciado_largo"]
    while xm < Lx-0.01: line(xm,0,xm,Ly,lay="MALLA"); xm += MESH_DATA["espaciado_largo"]
    ym = MESH_DATA["espaciado_corto"]
    while ym < Ly-0.01: line(0,ym,Lx,ym,lay="MALLA"); ym += MESH_DATA["espaciado_corto"]
    line(-0.2,Ly/2,Lx+0.2,Ly/2,lay="EJES")
    line(Lx/2,-0.2,Lx/2,Ly+0.2,lay="EJES")
    ya = Ly/2
    line(-0.35,ya,0,ya,lay="TEXTO"); line(Lx,ya,Lx+0.35,ya,lay="TEXTO")
    txt(-0.50,ya,"A",h=0.16,lay="TEXTO"); txt(Lx+0.38,ya,"A'",h=0.16,lay="TEXTO")
    cota_h(0,0,Lx,0,-0.40,f"Lx={Lx:.2f} m")
    cota_v(0,0,0,Ly,-0.40,f"Ly={Ly:.2f} m")
    txt(Lx/2,Ly+0.60,"PLANTA ESTRUCTURAL - PLACA ALIGERADA",h=THG,lay="TEXTO",ha=1)
    txt(Lx/2,Ly+0.42,f"Luz X={Lx:.2f}m  Luz Y={Ly:.2f}m  {norma_sel}  f'c={fc_concreto:.0f} MPa",h=TH,lay="TEXTO",ha=1)
    txt(Lx/2,Ly+0.30,f"Viga borde: {ref_beam}  Est.@{s_beam*100:.0f}cm",h=TH*0.85,lay="TEXTO",ha=1)
    txt(Lx/2,Ly+0.18,f"Malla Q2 O{MESH_DATA['diametro']*1000:.0f}mm @{MESH_DATA['espaciado_largo']*100:.0f}cm perp.",h=TH*0.85,lay="TEXTO",ha=1)

    # ZONA 2: Corte A-A' escalado x10
    ox2 = Lx + 2.0; oy2 = 0.0
    h_pr2 = PROFILE_DATA["alto_total"] * S
    h_bl2 = BLOCK_DATA["alto"] * S
    hw2   = PROFILE_DATA["ancho_alma"] / 2 * S
    hf2   = PROFILE_DATA["ancho_ala"] * S
    t2s   = PROFILE_DATA["espesor_alma"] * S
    tc2s  = 0.0015 * S
    esp2  = perfil_espaciado * S
    vb2   = viga_b * S
    vh2   = viga_h * S
    tor2  = espesor_torta * S
    z0 = oy2; z1 = z0+h_pr2; z2 = z1+h_bl2; z3 = z2+tor2
    zvb2 = z0-vh2 if incluir_vigas else z0
    xst2 = ox2+(vb2 if incluir_vigas else 0)
    n_bay2 = 2
    xend2 = xst2+n_bay2*esp2+(vb2 if incluir_vigas else 0)

    for i in range(n_bay2+1):
        xp2 = xst2+i*esp2
        poly([(xp2-hw2-hf2,z0),(xp2+hw2+hf2,z0),(xp2+hw2+hf2,z0+t2s),(xp2-hw2-hf2,z0+t2s)],lay="PERFILES")
        poly([(xp2-hw2-hf2,z1-t2s),(xp2+hw2+hf2,z1-t2s),(xp2+hw2+hf2,z1),(xp2-hw2-hf2,z1)],lay="PERFILES")
        poly([(xp2-hw2-tc2s,z0+t2s),(xp2-hw2,z0+t2s),(xp2-hw2,z1-t2s),(xp2-hw2-tc2s,z1-t2s)],lay="PERFILES")
        poly([(xp2+hw2,z0+t2s),(xp2+hw2+tc2s,z0+t2s),(xp2+hw2+tc2s,z1-t2s),(xp2+hw2,z1-t2s)],lay="PERFILES")
    for i in range(n_bay2):
        bx_l = xst2+i*esp2+hw2+hf2; bx_r = xst2+(i+1)*esp2-hw2-hf2
        if bx_r > bx_l+0.01:
            poly([(bx_l,z1),(bx_r,z1),(bx_r,z2),(bx_l,z2)],lay="CONCRETO")
            for zh in [z1+h_bl2*0.33,z1+h_bl2*0.67]: line(bx_l+0.06,zh,bx_r-0.06,zh,lay="CONCRETO")
    poly([(xst2,z2),(xst2+n_bay2*esp2,z2),(xst2+n_bay2*esp2,z3),(xst2,z3)],lay="CONCRETO")
    xm2 = xst2; r_m = MESH_DATA["diametro"]*S*0.5
    while xm2 <= xst2+n_bay2*esp2:
        circle(xm2,z2+tor2*0.25,r_m,lay="MALLA"); xm2 += MESH_DATA["espaciado_largo"]*S
    line(xst2,z2+tor2*0.25,xst2+n_bay2*esp2,z2+tor2*0.25,lay="MALLA")
    if incluir_vigas:
        cov2=0.04*S; rb2=0.008*S
        for xv2 in [ox2,xend2-vb2]:
            poly([(xv2,zvb2),(xv2+vb2,zvb2),(xv2+vb2,z3),(xv2,z3)],lay="CONCRETO")
            for bz2 in [zvb2+cov2+rb2,z3-cov2-rb2]:
                for bx2 in [xv2+cov2+rb2,xv2+vb2-cov2-rb2]: circle(bx2,bz2,rb2,lay="ACERO_LONG")
            poly([(xv2+cov2,zvb2+cov2),(xv2+vb2-cov2,zvb2+cov2),
                  (xv2+vb2-cov2,z3-cov2),(xv2+cov2,z3-cov2)],lay="ESTRIBOS")
    cx2c = xend2+0.8
    cota_v(cx2c,z0,cx2c,z1,0.7,f"{PROFILE_DATA['alto_total']*1000:.0f}mm")
    cota_v(cx2c,z1,cx2c,z2,0.7,f"{BLOCK_DATA['alto']*1000:.0f}mm")
    cota_v(cx2c,z2,cx2c,z3,0.7,f"{espesor_torta*100:.0f}cm torta")
    if incluir_vigas: cota_v(cx2c,zvb2,cx2c,z0,0.7,f"{viga_h*100:.0f}cm VB")
    cota_h(xst2,z0,xst2+esp2,z0,-1.0,f"@{perfil_espaciado*100:.0f}cm")
    txt((ox2+xend2)/2,z3+0.40,"CORTE A-A'",h=THG,lay="TEXTO",ha=1)
    txt((ox2+xend2)/2,z3+0.25,f"Concreto f'c={fc_concreto:.0f} MPa  Malla Q2 O{MESH_DATA['diametro']*1000:.0f}mm",h=TH*0.85,lay="TEXTO",ha=1)
    txt((ox2+xend2)/2,z3+0.12,"Escala 1:10  (cotas en mm o cm segun rotulo)",h=TH*0.75,lay="NOTAS",ha=1)

    # Posicion base de zonas de detalle
    _oy_det = zvb2 - 1.5

    # ZONA 3: Detalle Perfil Colmena x10
    ox3=0.0; oy3=_oy_det; sc=0.010
    h3=90*sc; hwc=40*sc; hfc=20*sc; tc3=2*sc; tc2_3=1.5*sc
    pts_p=[
        (ox3-hwc-hfc,oy3),(ox3+hwc+hfc,oy3),(ox3+hwc+hfc,oy3+tc3),(ox3+hwc+tc2_3,oy3+tc3),
        (ox3+hwc+tc2_3,oy3+h3-tc3),(ox3+hwc+hfc,oy3+h3-tc3),(ox3+hwc+hfc,oy3+h3),
        (ox3-hwc-hfc,oy3+h3),(ox3-hwc-hfc,oy3+h3-tc3),(ox3-hwc-tc2_3,oy3+h3-tc3),
        (ox3-hwc-tc2_3,oy3+tc3),(ox3-hwc-hfc,oy3+tc3),
    ]
    poly(pts_p,lay="PERFILES")
    txt(ox3,oy3+h3+0.15,"DETALLE PERFIL COLMENA",h=THT,lay="TEXTO",ha=1)
    cota_h(ox3-hwc-hfc,oy3,ox3-hwc,oy3,-0.15,"20.0")
    cota_h(ox3-hwc,oy3,ox3+hwc,oy3,-0.15,"80.0")
    cota_h(ox3+hwc,oy3,ox3+hwc+hfc,oy3,-0.15,"20.0")
    cota_v(ox3+hwc+hfc,oy3,ox3+hwc+hfc,oy3+h3,0.12,"90.0")
    txt(ox3,oy3-0.22,"(medidas en milimetros - escala 1:10)",h=TH*0.8,lay="NOTAS",ha=1)
    txt(ox3,oy3-0.32,f"Espesor=1.5mm  Peso={PROFILE_DATA['peso_por_m']:.1f} kg/m",h=TH*0.8,lay="NOTAS",ha=1)

    # ZONA 4: Perfil Colmena Reforzado x10
    ox4=3.3; oy4=_oy_det
    vb3=viga_b*S; vh3=viga_h*S
    poly([(ox4-vb3,oy4-vh3),(ox4,oy4-vh3),(ox4,oy4+h3),(ox4-vb3,oy4+h3)],lay="CONCRETO")
    pts_p4=[(px+ox4,py) for px,py in pts_p]
    poly(pts_p4,lay="PERFILES")
    poly([(ox4+hwc+hfc,oy4-vh3),(ox4+hwc+hfc+vb3,oy4-vh3),
          (ox4+hwc+hfc+vb3,oy4+h3),(ox4+hwc+hfc,oy4+h3)],lay="CONCRETO")
    cov3=0.04*S; rb3=0.008*S
    for xv4 in [ox4-vb3,ox4+hwc+hfc]:
        xe4=xv4+vb3
        for bz4 in [oy4-vh3+cov3+rb3,oy4+h3-cov3-rb3]:
            for bx4 in [xv4+cov3+rb3,xe4-cov3-rb3]: circle(bx4,bz4,rb3,lay="ACERO_LONG")
        poly([(xv4+cov3,oy4-vh3+cov3),(xe4-cov3,oy4-vh3+cov3),
              (xe4-cov3,oy4+h3-cov3),(xv4+cov3,oy4+h3-cov3)],lay="ESTRIBOS")
    cota_h(ox4-vb3,oy4,ox4,oy4,-0.18,f"{viga_b*100:.0f}")
    cota_h(ox4-hwc,oy4,ox4+hwc,oy4,-0.18,"80")
    cota_h(ox4+hwc,oy4,ox4+hwc+hfc,oy4,-0.18,"20")
    cota_h(ox4+hwc+hfc,oy4,ox4+hwc+hfc+vb3,oy4,-0.18,f"{viga_b*100:.0f}")
    cota_v(ox4+hwc+hfc+vb3,oy4-vh3,ox4+hwc+hfc+vb3,oy4+h3,0.15,
           f"{int((viga_h+PROFILE_DATA['alto_total'])*1000)}")
    txt(ox4+(hwc+hfc)*0.2,oy4+h3+0.15,"PERFIL COLMENA REFORZADO - VIGA BORDE",h=THT,lay="TEXTO",ha=1)
    txt(ox4-vb3/2,oy4-vh3*0.5,"2#4",h=TH*0.85,lay="TEXTO",ha=1)
    txt(ox4,oy4-vh3-0.18,f"As={As_beam:.2f}/{As_prov_beam:.1f} cm2  Est.@{s_beam*100:.0f}cm",h=TH*0.85,lay="NOTAS",ha=1)
    txt(ox4,oy4-vh3-0.30,"(medidas en milimetros - escala 1:10)",h=TH*0.75,lay="NOTAS",ha=1)

    # ZONA 5: Bloquelon Santafe x5
    S5=5.0; ox5=ox4+hwc+hfc+vb3+1.8; oy5=_oy_det
    L5=BLOCK_DATA["largo"]*S5; W5=BLOCK_DATA["ancho"]*S5; H5=BLOCK_DATA["alto"]*S5
    ang_bl=math.radians(30); cos30=math.cos(ang_bl); sin30=math.sin(ang_bl)
    def iso5(x,y,z): return ox5+x+y*0.5*cos30, oy5+z+y*0.5*sin30
    poly([iso5(0,0,0),iso5(L5,0,0),iso5(L5,0,H5),iso5(0,0,H5)],lay="CONCRETO")
    poly([iso5(0,0,H5),iso5(L5,0,H5),iso5(L5,W5,H5),iso5(0,W5,H5)],lay="CONCRETO")
    poly([iso5(L5,0,0),iso5(L5,W5,0),iso5(L5,W5,H5),iso5(L5,0,H5)],lay="CONCRETO")
    hw5=W5*0.35; hl5=L5*0.42; cx1b5=L5*0.05; cy1b5=W5*0.18
    poly([iso5(cx1b5,cy1b5,H5),iso5(cx1b5+hl5,cy1b5,H5),iso5(cx1b5+hl5,cy1b5+hw5,H5),iso5(cx1b5,cy1b5+hw5,H5)],lay="CONCRETO")
    cx2b5=L5-cx1b5-hl5
    poly([iso5(cx2b5,cy1b5,H5),iso5(cx2b5+hl5,cy1b5,H5),iso5(cx2b5+hl5,cy1b5+hw5,H5),iso5(cx2b5,cy1b5+hw5,H5)],lay="CONCRETO")
    txt(ox5+L5/2,oy5+H5+0.20,"DETALLE BLOQUELON SANTAFE",h=THT,lay="TEXTO",ha=1)
    cota_h(ox5,oy5-0.25,ox5+L5,oy5-0.25,0.0,f"Largo={BLOCK_DATA['largo']*100:.0f}cm")
    txt(ox5+L5+0.15,oy5+H5/2,f"H={BLOCK_DATA['alto']*100:.0f}cm",h=TH,lay="COTAS")
    txt(ox5+L5/2,oy5-0.40,f"Ancho={BLOCK_DATA['ancho']*100:.0f}cm  Peso={BLOCK_DATA['peso_unitario']}kg  5.18 ud/m2",h=TH*0.8,lay="NOTAS",ha=1)
    txt(ox5+L5/2,oy5-0.52,"(medidas en centimetros - escala 1:5)",h=TH*0.75,lay="NOTAS",ha=1)

    # ZONA 6: Vigueta Coronamiento x10
    ox6=0.0; oy6=_oy_det-vh3-2.0
    B_vg=0.065*S; H_vg=0.25*S; t_vg=0.005*S
    for pts6 in [
        [(ox6,oy6),(ox6+B_vg*2+t_vg,oy6),(ox6+B_vg*2+t_vg,oy6+t_vg),(ox6,oy6+t_vg)],
        [(ox6,oy6),(ox6,oy6+H_vg),(ox6+t_vg,oy6+H_vg),(ox6+t_vg,oy6)],
        [(ox6+B_vg*2,oy6),(ox6+B_vg*2,oy6+H_vg),(ox6+B_vg*2+t_vg,oy6+H_vg),(ox6+B_vg*2+t_vg,oy6)],
    ]:
        poly(pts6,lay="PERFILES")
    rb6=0.008*S
    circle(ox6+t_vg/2,oy6+H_vg*0.5,rb6,lay="ACERO_LONG")
    circle(ox6+B_vg*2+t_vg/2,oy6+H_vg*0.5,rb6,lay="ACERO_LONG")
    txt(ox6-0.15,oy6+H_vg*0.5,"1#5 Corrida",h=TH*0.8,lay="TEXTO",ha=2)
    line(ox6-0.01,oy6+H_vg*0.5,ox6+t_vg+rb6,oy6+H_vg*0.5,lay="TEXTO")
    ox6s=ox6+B_vg*2+t_vg+0.60
    poly([(ox6s,oy6+H_vg),(ox6s+B_vg*2+t_vg,oy6+H_vg),(ox6s+B_vg*2+t_vg,oy6+H_vg+espesor_torta*S),(ox6s,oy6+H_vg+espesor_torta*S)],lay="CONCRETO")
    for pts6s in [
        [(ox6s,oy6),(ox6s+B_vg*2+t_vg,oy6),(ox6s+B_vg*2+t_vg,oy6+t_vg),(ox6s,oy6+t_vg)],
        [(ox6s,oy6),(ox6s,oy6+H_vg),(ox6s+t_vg,oy6+H_vg),(ox6s+t_vg,oy6)],
        [(ox6s+B_vg*2,oy6),(ox6s+B_vg*2,oy6+H_vg),(ox6s+B_vg*2+t_vg,oy6+H_vg),(ox6s+B_vg*2+t_vg,oy6)],
    ]:
        poly(pts6s,lay="PERFILES")
    cota_v(ox6-0.16,oy6,ox6-0.16,oy6+H_vg,-0.15,f"{int(0.25*1000)}mm")
    txt(ox6+B_vg,oy6-0.18,"Lamina Cal.18  1#5 c/ala  f'c=21 MPa",h=TH*0.8,lay="NOTAS")
    txt(ox6+B_vg,oy6+H_vg+0.18,"DETALLE VIGUETA CORONAMIENTO",h=THT,lay="TEXTO")
    txt(ox6+B_vg,oy6-0.30,"(medidas en milimetros - escala 1:10)",h=TH*0.75,lay="NOTAS")

    # TABLA RESUMEN CANTIDADES
    tx0=cx2c+1.2; ty0=Ly+0.5
    cw_tab=[2.0,1.2,0.9]; rh_tab=0.30; tw_tab=sum(cw_tab)
    cantidades_tab=[
        ("Bloques Santafe",str(n_bloques_desp),"und"),
        ("Perfiles Colmena",f"{longitud_total_perfiles_desp:.1f}","m"),
        ("Malla Q2",f"{area_malla:.1f}","m2"),
        ("Concreto f'c",f"{vol_concreto_total_desp:.2f}","m3"),
        ("Cemento 50 kg",str(bultos_cemento),"bultos"),
    ]
    nr_qty=len(cantidades_tab)+1
    poly([(tx0,ty0-rh_tab*nr_qty),(tx0+tw_tab,ty0-rh_tab*nr_qty),(tx0+tw_tab,ty0),(tx0,ty0)],lay="ROTULO")
    for ci,cw in enumerate(cw_tab[:-1]):
        xc=tx0+sum(cw_tab[:ci+1]); line(xc,ty0-rh_tab*nr_qty,xc,ty0,lay="ROTULO")
    line(tx0,ty0-rh_tab,tx0+tw_tab,ty0-rh_tab,lay="ROTULO")
    for ci,hdr in enumerate(["MATERIAL","CANT.","UD."]): txt(tx0+sum(cw_tab[:ci])+0.06,ty0-rh_tab*0.4,hdr,h=TH,lay="TEXTO")
    for ri,(mat,cant,uni) in enumerate(cantidades_tab):
        ry=ty0-rh_tab*(ri+2); line(tx0,ry,tx0+tw_tab,ry,lay="ROTULO")
        txt(tx0+0.06,ry+rh_tab*0.25,mat,h=TH*0.85,lay="TEXTO")
        txt(tx0+cw_tab[0]+0.06,ry+rh_tab*0.25,cant,h=TH*0.85,lay="TEXTO")
        txt(tx0+cw_tab[0]+cw_tab[1]+0.04,ry+rh_tab*0.25,uni,h=TH*0.85,lay="TEXTO")
    txt(tx0+tw_tab/2,ty0+0.18,"RESUMEN CANTIDADES",h=THT,lay="TEXTO",ha=1)
    txt(tx0+tw_tab/2,ty0+0.06,f"Area total: {area_total:.2f} m2  Pisos: {num_pisos}",h=TH*0.85,lay="NOTAS",ha=1)

    # DESPIECE ACERO — PEDIDO FERRETERIA
    ty_des=ty0-rh_tab*nr_qty-0.6
    cw_des=[0.55,0.85,0.75,0.55,0.80,1.25]; tw_des=sum(cw_des)
    hdrs_des=["MCA","DIAM","L (m)","N","KG","OBSERV."]
    long_est=2*(viga_b+viga_h)+0.25
    n_est_v=math.ceil(perfil_largo/s_beam)+1
    n_v=4 if incluir_vigas else 2
    despiece_rows=[
        ("L1",rebar_combo["b1"]["nom"][:4],f"{perfil_largo:.2f}",str(4*n_v),f"{perfil_largo*4*n_v*0.994:.1f}","Long. viga borde"),
        ("E1","1/4\"",f"{long_est:.2f}",str(n_est_v*n_v),f"{long_est*n_est_v*n_v*0.62:.1f}","Estribos VB"),
    ]
    nr_des=len(despiece_rows)+1
    poly([(tx0,ty_des-rh_tab*nr_des),(tx0+tw_des,ty_des-rh_tab*nr_des),(tx0+tw_des,ty_des),(tx0,ty_des)],lay="ROTULO")
    for ci,cw in enumerate(cw_des[:-1]):
        xc=tx0+sum(cw_des[:ci+1]); line(xc,ty_des-rh_tab*nr_des,xc,ty_des,lay="ROTULO")
    line(tx0,ty_des-rh_tab,tx0+tw_des,ty_des-rh_tab,lay="ROTULO")
    for ci,hdr in enumerate(hdrs_des): txt(tx0+sum(cw_des[:ci])+0.04,ty_des-rh_tab*0.4,hdr,h=TH*0.85,lay="TEXTO")
    for ri,row_d in enumerate(despiece_rows):
        ry=ty_des-rh_tab*(ri+2); line(tx0,ry,tx0+tw_des,ry,lay="ROTULO")
        for ci,val in enumerate(row_d): txt(tx0+sum(cw_des[:ci])+0.04,ry+rh_tab*0.25,str(val),h=TH*0.85,lay="TEXTO")
    txt(tx0+tw_des/2,ty_des+0.18,"DESPIECE ACERO - PEDIDO FERRETERIA",h=THT,lay="TEXTO",ha=1)

    # ROTULO ICONTEC 2289
    rw_r=2.80; rh_r=0.90
    _xmax_cont=max(tx0+tw_des,cx2c+0.8,ox5+L5+1.0)
    _ymin_cont=min(oy6-0.5,_oy_det-vh3-0.5)
    rx=_xmax_cont-rw_r; ry_r=_ymin_cont-rh_r-0.20
    poly([(rx,ry_r),(rx+rw_r,ry_r),(rx+rw_r,ry_r+rh_r),(rx,ry_r+rh_r)],lay="ROTULO")
    cws_r=[0.55,0.85,0.70,0.70]; cx_rr=rx
    for cw in cws_r[:-1]: cx_rr+=cw; line(cx_rr,ry_r,cx_rr,ry_r+rh_r,lay="ROTULO")
    rrow=rh_r/5
    for i in range(1,5): line(rx,ry_r+i*rrow,rx+rw_r,ry_r+i*rrow,lay="ROTULO")
    _emp_r=st.session_state.get("empresa","________________") or "________________"
    _proy_r=proyecto_nombre or st.session_state.get("proyecto","________________") or "________________"
    _cli_r=proyecto_cliente or "________________"
    _ing_r=st.session_state.get("ingeniero","________________") or "________________"
    def _cr(i): return rx+sum(cws_r[:i])+0.05
    campos_r=[
        (_cr(0),ry_r+0.05,"ELABORO:",TH),(_cr(0),ry_r+0.01,elaboro,TH*0.85),
        (_cr(1),ry_r+0.05,"REVISO:",TH),(_cr(1),ry_r+0.01,revisado,TH*0.85),
        (_cr(2),ry_r+0.05,"APROBO:",TH),(_cr(2),ry_r+0.01,aprobado,TH*0.85),
        (_cr(3),ry_r+0.05,"FECHA:",TH),(_cr(3),ry_r+0.01,datetime.now().strftime("%d/%m/%Y"),TH*0.85),
        (_cr(0),ry_r+rrow+0.05,"NORMA:",TH),(_cr(0),ry_r+rrow+0.01,norma_sel[:26],TH*0.85),
        (_cr(1),ry_r+rrow+0.05,"ESCALA:",TH),(_cr(1),ry_r+rrow+0.01,"VAR.",TH*0.85),
        (_cr(2),ry_r+rrow+0.05,"HOJA:",TH),(_cr(2),ry_r+rrow+0.01,"1/1",TH*0.85),
        (_cr(3),ry_r+rrow+0.05,"PLANO:",TH),(_cr(3),ry_r+rrow+0.01,plano_numero,TH*0.85),
        (_cr(0),ry_r+rrow*2+0.05,"CLIENTE:",TH),(_cr(0),ry_r+rrow*2+0.01,_cli_r[:26],TH*0.85),
        (_cr(1),ry_r+rrow*2+0.05,f"f'c={fc_concreto:.0f}MPa",TH),
        (_cr(2),ry_r+rrow*2+0.05,"Mu placa:",TH),(_cr(2),ry_r+rrow*2+0.01,f"{Mu:.2f} kNm",TH*0.85),
        (_cr(3),ry_r+rrow*2+0.05,"Deflexion:",TH),
        (_cr(3),ry_r+rrow*2+0.01,f"{delta_calc*1000:.1f}mm {'OK' if cumple_deflexion else 'EXC'}",TH*0.85),
        (_cr(0),ry_r+rrow*3+0.05,"EMPRESA:",TH),(_cr(0),ry_r+rrow*3+0.01,_emp_r[:26],TH*0.85),
        (_cr(1),ry_r+rrow*3+0.05,"Ingeniero:",TH),(_cr(1),ry_r+rrow*3+0.01,_ing_r[:22],TH*0.85),
        (_cr(2),ry_r+rrow*3+0.05,"As viga:",TH),(_cr(2),ry_r+rrow*3+0.01,f"{As_beam:.1f}/{As_prov_beam:.1f} cm2",TH*0.85),
        (_cr(3),ry_r+rrow*3+0.05,"Vu placa:",TH),(_cr(3),ry_r+rrow*3+0.01,f"{Vu:.2f}kN",TH*0.85),
        (_cr(0),ry_r+rrow*4+0.08,"PROYECTO:",THG),(_cr(0),ry_r+rrow*4+0.01,_proy_r[:35],THB),
    ]
    for (fx,fy,fs,fh) in campos_r: txt(fx,fy,fs,h=fh,lay="ROTULO")
    _xmin_f=-1.0; _ymin_f=ry_r-0.3; _xmax_f=_xmax_cont+0.5; _ymax_f=Ly+1.8
    return render(_xmin_f,_ymin_f,_xmax_f,_ymax_f)

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR IFC (STEP ASCII sin dependencias externas)
# ══════════════════════════════════════════════════════════════════════════════
def generate_ifc():
    """IFC 2x3 STEP ASCII con geometria correcta: perfiles extruidos en X/Y."""
    now  = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    emp  = st.session_state.get("empresa","StructoPro") or "StructoPro"
    ing  = st.session_state.get("ingeniero","Ingeniero") or "Ingeniero"
    proy = proyecto_nombre or st.session_state.get("proyecto","PlacaFacil") or "PlacaFacil"

    h_pr   = PROFILE_DATA["alto_total"]
    h_bl   = BLOCK_DATA["alto"]
    b_prof = PROFILE_DATA["b_total"]

    lines = [
        "ISO-10303-21;",
        "HEADER;",
        f"FILE_DESCRIPTION(('IFC2X3 PlacaFacil {proy}'),'2;1');",
        f"FILE_NAME('{proy}.ifc','{now}','',(''),(''),'{emp}','');",
        "FILE_SCHEMA(('IFC2X3'));",
        "ENDSEC;",
        "DATA;",
        "#1=IFCPROJECT('Proj01',$,$,$,$,$,$,(#100),#200);",
        "#100=IFCGEOMETRICREPRESENTATIONCONTEXT($,'Model',3,1.0E-5,#101,$);",
        "#101=IFCAXIS2PLACEMENT3D(#102,$,$);",
        "#102=IFCCARTESIANPOINT((0.,0.,0.));",
        "#200=IFCUNITASSIGNMENT((#201,#202,#203));",
        "#201=IFCSIUNIT(*,.LENGTHUNIT.,$,.METRE.);",
        "#202=IFCSIUNIT(*,.AREAUNIT.,$,.SQUARE_METRE.);",
        "#203=IFCSIUNIT(*,.VOLUMEUNIT.,$,.CUBIC_METRE.);",
        f"#10=IFCORGANIZATION($,'{emp}',$,$,$);",
        f"#11=IFCPERSON($,'{ing}',$,$,$,$,$,$);",
        "#12=IFCPERSONANDORGANIZATION(#11,#10,$);",
        "#13=IFCOWNERHISTORY(#12,#10,$,.ADDED.,$,$,$,0);",
        "#20=IFCSITE('Site01',#13,'Sitio',$,$,#21,$,$,.ELEMENT.,$,$,$,$,$);",
        "#21=IFCLOCALPLACEMENT($,#101);",
        "#30=IFCBUILDING('Build01',#13,'Edificio',$,$,#31,$,$,.ELEMENT.,$,$,$);",
        "#31=IFCLOCALPLACEMENT(#21,#101);",
        "#40=IFCBUILDINGSTOREY('Piso1',#13,'Piso 1',$,$,#41,$,$,.ELEMENT.,0.);",
        "#41=IFCLOCALPLACEMENT(#31,#101);",
        "#50=IFCRELAGGREGATES('RA1',#13,$,$,#1,(#20));",
        "#51=IFCRELAGGREGATES('RA2',#13,$,$,#20,(#30));",
        "#52=IFCRELAGGREGATES('RA3',#13,$,$,#30,(#40));",
        f"#60=IFCMATERIAL('Concreto fc={fc_concreto:.0f}MPa');",
        "#61=IFCMATERIAL('Perfil Colmena Acero Galvanizado');",
        "#62=IFCMATERIAL('Bloquelon Santafe Arcilla');",
        "#70=IFCDIRECTION((0.,0.,1.));",
        "#71=IFCDIRECTION((1.,0.,0.));",
        "#72=IFCDIRECTION((0.,1.,0.));",
    ]

    eid = 500
    def _eid():
        nonlocal eid; eid += 1; return eid

    def _pt(x, y, z):
        e = _eid()
        lines.append(f"#{e}=IFCCARTESIANPOINT(({x:.5f},{y:.5f},{z:.5f}));")
        return e

    def _ax3(pt_e, ax_e=None, rd_e=None):
        e = _eid()
        a = f"#{ax_e}" if ax_e else "$"
        r = f"#{rd_e}" if rd_e else "$"
        lines.append(f"#{e}=IFCAXIS2PLACEMENT3D(#{pt_e},{a},{r});")
        return e

    def _lp(ax3_e):
        e = _eid()
        lines.append(f"#{e}=IFCLOCALPLACEMENT(#41,#{ax3_e});")
        return e

    def _rp(xd, yd):
        e = _eid()
        lines.append(f"#{e}=IFCRECTANGLEPROFILEDEF(.AREA.,$,$,{xd:.5f},{yd:.5f});")
        return e

    def _ext(prof_e, ax3_e, depth):
        e = _eid()
        lines.append(f"#{e}=IFCEXTRUDEDAREASOLID(#{prof_e},#{ax3_e},#70,{depth:.5f});")
        return e

    def _shp(sol_e):
        er = _eid(); ep = _eid()
        lines.append(f"#{er}=IFCSHAPEREPRESENTATION(#100,'Body','SweptSolid',(#{sol_e}));")
        lines.append(f"#{ep}=IFCPRODUCTDEFINITIONSHAPE($,$,(#{er}));")
        return ep

    all_ents = []

    # LOSA (topping de concreto) — z=h_pr+h_bl hasta +espesor_torta
    z_slab = h_pr + h_bl
    e_sp = _pt(Lx/2, Ly/2, z_slab)
    e_sax = _ax3(e_sp)
    e_slp = _lp(e_sax)
    e_srp = _rp(Lx, Ly)
    e_sext = _ext(e_srp, e_sax, espesor_torta)
    e_sshp = _shp(e_sext)
    e_slab = _eid()
    lines.append(f"#{e_slab}=IFCSLAB('Losa01',#13,'Losa Aligerada PlacaFacil',"
                 f"'Perfil Colmena + Bloquelon + Torta',"
                 f"'Placa Aligerada',#{e_slp},#{e_sshp},$,.FLOOR.);")
    all_ents.append(f"#{e_slab}")

    # PERFILES METALICOS
    # Paralelo a X: AXIS=(1,0,0)=#71  REFD=(0,1,0)=#72
    #   LocalZ=globalX, LocalX=globalY, LocalY=globalZ
    #   Seccion: XDim=b_prof(en Y), YDim=h_pr(en Z)
    #   Origen en (0, yp, h_pr/2) -> z de 0 a h_pr
    #   Extrusion depth=Lx en localZ=globalX
    # Paralelo a Y: AXIS=(0,1,0)=#72  REFD=(1,0,0)=#71
    #   LocalY=(0,1,0)x(1,0,0)=(0,0,-1) -> altura en -Z
    #   Origen en (xp, 0, h_pr) -> extrusion depth=Ly

    if orientacion == "Paralelo a X":
        for idx, yp in enumerate(np.linspace(0, Ly, n_profiles)):
            ep_pt = _pt(0., yp, h_pr/2)
            ep_ax = _ax3(ep_pt, 71, 72)
            ep_lp = _lp(ep_ax)
            ep_rp = _rp(b_prof, h_pr)
            ep_ext = _ext(ep_rp, ep_ax, Lx)
            ep_shp = _shp(ep_ext)
            em = _eid()
            lines.append(f"#{em}=IFCMEMBER('Prof{idx:03d}',#13,'Perfil Colmena {idx+1}',"
                         f"'Perfil metalico colmena','Perfil Colmena',#{ep_lp},#{ep_shp},$);")
            all_ents.append(f"#{em}")
    else:
        for idx, xp in enumerate(np.linspace(0, Lx, n_profiles)):
            ep_pt = _pt(xp, 0., h_pr)
            ep_ax = _ax3(ep_pt, 72, 71)
            ep_lp = _lp(ep_ax)
            ep_rp = _rp(b_prof, h_pr)
            ep_ext = _ext(ep_rp, ep_ax, Ly)
            ep_shp = _shp(ep_ext)
            em = _eid()
            lines.append(f"#{em}=IFCMEMBER('Prof{idx:03d}',#13,'Perfil Colmena {idx+1}',"
                         f"'Perfil metalico colmena','Perfil Colmena',#{ep_lp},#{ep_shp},$);")
            all_ents.append(f"#{em}")

    # VIGAS DE BORDE
    if incluir_vigas:
        for by, bnom in [(0., "Viga_Sur"), (Ly-viga_b, "Viga_Norte")]:
            eb_pt = _pt(0., by+viga_b/2, -viga_h/2)
            eb_ax = _ax3(eb_pt, 71, 72)
            eb_lp = _lp(eb_ax)
            eb_rp = _rp(viga_b, viga_h)
            eb_ext = _ext(eb_rp, eb_ax, Lx)
            eb_shp = _shp(eb_ext)
            eb = _eid()
            lines.append(f"#{eb}=IFCBEAM('{bnom}',#13,'{bnom}',"
                         f"'Viga borde concreto','{ref_beam}',#{eb_lp},#{eb_shp},$);")
            all_ents.append(f"#{eb}")
        for bx, bnom in [(0., "Viga_Oeste"), (Lx-viga_b, "Viga_Este")]:
            eb_pt = _pt(bx+viga_b/2, 0., -viga_h/2)
            eb_ax = _ax3(eb_pt, 72, 71)
            eb_lp = _lp(eb_ax)
            eb_rp = _rp(viga_b, viga_h)
            eb_ext = _ext(eb_rp, eb_ax, Ly)
            eb_shp = _shp(eb_ext)
            eb = _eid()
            lines.append(f"#{eb}=IFCBEAM('{bnom}',#13,'{bnom}',"
                         f"'Viga borde concreto','{ref_beam}',#{eb_lp},#{eb_shp},$);")
            all_ents.append(f"#{eb}")

    # RELACION ESPACIAL
    if all_ents:
        e_rel = _eid()
        lines.append(f"#{e_rel}=IFCRELCONTAINEDINSPATIALSTRUCTURE('Rel04',#13,$,$,"
                     f"({','.join(all_ents)}),#40);")

    # CANTIDADES
    e_qA  = _eid(); lines.append(f"#{e_qA}=IFCQUANTITYAREA('GrossFloorArea',$,$,{area_total:.4f},$);")
    e_qV  = _eid(); lines.append(f"#{e_qV}=IFCQUANTITYVOLUME('GrossVolume',$,$,{vol_concreto_total_desp:.4f},$);")
    e_qty = _eid(); lines.append(f"#{e_qty}=IFCELEMENTQUANTITY('Qty01',#13,'Cantidades',$,'BaseQuantities',(#{e_qA},#{e_qV}));")
    e_qrl = _eid(); lines.append(f"#{e_qrl}=IFCRELDEFINESBYPROPERTIES('QRel01',#13,$,$,(#{e_slab}),#{e_qty});")
    e_p1  = _eid(); lines.append(f"#{e_p1}=IFCPROPERTYSINGLEVALUE('NumPerfiles',$,IFCINTEGER({n_profiles}),$);")
    e_p2  = _eid(); lines.append(f"#{e_p2}=IFCPROPERTYSINGLEVALUE('NumBloques',$,IFCINTEGER({n_bloques_desp}),$);")
    e_p3  = _eid(); lines.append(f"#{e_p3}=IFCPROPERTYSINGLEVALUE('BultosCemento',$,IFCINTEGER({bultos_cemento}),$);")
    e_pst = _eid(); lines.append(f"#{e_pst}=IFCPROPERTYSET('Pset_Placa',#13,'Datos PlacaFacil',$,(#{e_p1},#{e_p2},#{e_p3}));")
    e_prl = _eid(); lines.append(f"#{e_prl}=IFCRELDEFINESBYPROPERTIES('PRel01',#13,$,$,(#{e_slab}),#{e_pst});")

    lines += ["ENDSEC;", "END-ISO-10303-21;"]
    return "\n".join(lines).encode("utf-8")


def _make_pt(lines, _eid, x, y, z):
    eid_val = _eid()
    lines.append(f"#{eid_val}=IFCCARTESIANPOINT(({x:.4f},{y:.4f},{z:.4f}));")
    return eid_val

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DOCX
# ══════════════════════════════════════════════════════════════════════════════
def generate_memory():
    if not _DOCX_OK:
        return None
    doc = _DocxDoc()
    _AZUL  = RGBColor(0x1e, 0x3a, 0x5f)
    _GRIS  = RGBColor(0x44, 0x44, 0x44)
    _BLANC = RGBColor(0xFF, 0xFF, 0xFF)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _shd(cell, hex_c):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"),hex_c); tcPr.append(shd)

    def _h(text, lv=1):
        p = doc.add_heading(text, level=lv)
        for r in p.runs: r.font.color.rgb = _AZUL
        return p

    def _tbl_hdr(tbl, cols):
        row = tbl.rows[0]
        for i,c in enumerate(cols):
            cell=row.cells[i]; cell.text=c; _shd(cell,"1E3A5F")
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.bold=True; r.font.color.rgb=_BLANC; r.font.size=Pt(9)
                par.alignment=WD_ALIGN_PARAGRAPH.CENTER

    def _row(tbl, vals, bold=False, bg=None):
        row=tbl.add_row()
        for i,v in enumerate(vals):
            cell=row.cells[i]; cell.text=str(v)
            if bg: _shd(cell,bg)
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.size=Pt(9)
                    if bold: r.font.bold=True

    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3.0); sec.right_margin=Cm(2.5)

    # PORTADA
    logo_b = st.session_state.get("logo_bytes")
    if logo_b:
        try:
            _lp = doc.add_paragraph(); _lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            _lp.add_run().add_picture(io.BytesIO(logo_b), width=Cm(5.0))
            doc.add_paragraph()
        except Exception:
            pass
    _emp_d = st.session_state.get("empresa","________________") or "________________"
    t_port = doc.add_paragraph(); t_port.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = t_port.add_run(_emp_d); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=_AZUL
    doc.add_paragraph()
    t2 = doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("MEMORIA DE CALCULO — PLACA FACIL"); r2.font.size=Pt(16); r2.font.bold=True
    doc.add_paragraph()
    for lbl, val in [
        ("Proyecto:", proyecto_nombre or st.session_state.get("proyecto","________________") or "________________"),
        ("Cliente:", proyecto_cliente or "________________"),
        ("Norma:", norma_sel),
        ("Ingeniero:", st.session_state.get("ingeniero","________________") or "________________"),
        ("Fecha:", datetime.now().strftime("%d/%m/%Y")),
        ("Precios ref.:", st.session_state.get("_price_status_label","Verificar cotizacion")),
    ]:
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(lbl+" "); rl.font.bold=True; rl.font.color.rgb=_AZUL
        p.add_run(val)
    doc.add_page_break()

    # CAP 1
    _h("1. Datos de Entrada")
    t1=doc.add_table(rows=1,cols=2); t1.style="Table Grid"
    _tbl_hdr(t1,["Parametro","Valor"])
    for lbl,val in [
        ("Luz X",f"{Lx:.2f} m"),("Luz Y",f"{Ly:.2f} m"),
        ("Orientacion perfiles",orientacion),("Separacion perfiles",f"{perfil_espaciado*100:.0f} cm"),
        ("Espesor torta",f"{espesor_torta*100:.1f} cm"),("Altura bloquelón",f"{BLOCK_DATA['alto']*100:.0f} cm"),
        ("Altura perfil colmena",f"{PROFILE_DATA['alto_total']*1000:.0f} mm"),
        ("f'c concreto",f"{fc_concreto:.1f} MPa"),("Num. pisos",str(num_pisos)),
        ("Uso placa",f"{uso_placa} — Wl={carga_viva_kn} kN/m2"),
        ("Zona sismica",zona_sismica),("Sistema estructural",sistema_estructural),
        ("Vigas de borde","Si" if incluir_vigas else "No"),
        ("Viga borde (bxh)",f"{viga_b*100:.0f}x{viga_h*100:.0f} cm"),
    ]:
        _row(t1,[lbl,val])
    doc.add_paragraph()

    # CAP 2
    _h("2. Materiales")
    t2b=doc.add_table(rows=1,cols=3); t2b.style="Table Grid"
    _tbl_hdr(t2b,["Elemento","Especificacion","Referencia"])
    for row in [
        ("Bloquelon Santafe",f"{BLOCK_DATA['largo']*100:.0f}x{BLOCK_DATA['ancho']*100:.0f}x{BLOCK_DATA['alto']*100:.0f}cm",""),
        ("Perfil Colmena",f"h={PROFILE_DATA['alto_total']*1000:.0f}mm | {PROFILE_DATA['peso_por_m']:.1f}kg/m","Acero estructural"),
        ("Malla electrosoldada Q2",f"O{MESH_DATA['diametro']*1000:.0f}mm @{MESH_DATA['espaciado_largo']*100:.0f}cm","NSR-10 C.7.12"),
        ("Concreto",f"f'c={fc_concreto:.0f} MPa | Ec={E_concreto/1e9:.2f} GPa","NSR-10 C.8.5"),
        ("Acero refuerzo","fy=420 MPa","NSR-10 C.3.5"),
    ]:
        _row(t2b,row)
    doc.add_paragraph()

    # CAP 3
    _h("3. Cargas de Diseno")
    doc.add_paragraph(f"Carga muerta: Wd = {Wd_kn:.3f} kN/m2")
    doc.add_paragraph(f"Carga viva:   Wl = {carga_viva_kn:.2f} kN/m2")
    doc.add_paragraph(f"Carga ultima: Wu = {Wu_kn:.3f} kN/m2 (NSR-10 B.2.4: 1.2D + 1.6L)")
    doc.add_paragraph()

    # CAP 4
    _h("4. Seccion Compuesta y Desarrollo")
    for txt_p in [
        f"Modulo elasticidad concreto: Ec = {E_concreto/1e9:.2f} GPa (NSR-10 C.8.5)",
        f"Relacion modular n = Es/Ec = {n_mod:.1f}",
        f"Inercia seccion compuesta: I_comp = {I_comp*1e8:.4f} x 10-8 m4",
        f"Centroide compuesto desde base: y_c = {y_c*100:.2f} cm",
        f"Modulo elastico inferior: S_comp = {S_comp*1e6:.4f} x 10-6 m3",
        f"Longitud de desarrollo ld (NSR-10 C.12.2.2): {ld*100:.1f} cm",
    ]:
        doc.add_paragraph(txt_p)
    doc.add_paragraph()

    # CAP 5
    _h("5. Resultados Estructurales")
    doc.add_paragraph(f"Momento ultimo placa: Mu = {Mu:.2f} kN·m")
    doc.add_paragraph(f"Deflexion calculada: d = {delta_calc*1000:.2f} mm (limite = {delta_max*1000:.2f} mm | {'CUMPLE' if cumple_deflexion else 'NO CUMPLE'})")
    doc.add_paragraph(f"Cortante en apoyo: Vu = {Vu:.2f} kN | Vc = {Vc:.2f} kN | {'CUMPLE' if cumple_cortante else 'NO CUMPLE'}")
    doc.add_paragraph(f"Acero malla: As = {As_malla:.2f} cm2/m (minimo = {As_min*100:.2f} cm2/m)")
    doc.add_paragraph()

    # CAP 5.1 — Grafica diagramas modo claro
    if _MPL_OK:
        try:
            fig_dc, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
            ax1.plot(x_vals, M_vals, linewidth=2, color="#1565c0")
            ax1.fill_between(x_vals, 0, M_vals, alpha=0.25, color="#1565c0")
            ax1.set_xlabel("Distancia (m)"); ax1.set_ylabel("Momento (kN·m)")
            ax1.set_title("Diagrama de Momento — Viga Borde"); ax1.grid(True, alpha=0.3)
            ax2.plot(x_vals, V_vals, linewidth=2, color="#c62828")
            ax2.fill_between(x_vals, 0, V_vals, alpha=0.25, color="#c62828")
            ax2.set_xlabel("Distancia (m)"); ax2.set_ylabel("Cortante (kN)")
            ax2.set_title("Diagrama de Cortante — Viga Borde"); ax2.grid(True, alpha=0.3)
            plt.tight_layout()
            png = fig_to_docx_white(fig_dc)
            plt.close(fig_dc)
            if png:
                _ip = doc.add_paragraph(); _ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
                _ip.add_run().add_picture(io.BytesIO(png), width=Cm(15.0))
                _cap = doc.add_paragraph("Figura 1 — Diagramas de Momento y Cortante — Viga de Borde")
                _cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in _cap.runs: r.font.size=Pt(8); r.font.italic=True; r.font.color.rgb=_GRIS
        except Exception:
            pass

    # CAP 6 — Viga de borde
    _h("6. Diseno Viga de Borde")
    doc.add_paragraph(f"Momento: Mu = {Mu_beam:.2f} kN·m | As req = {As_beam:.2f} cm2")
    doc.add_paragraph(f"Refuerzo seleccionado: {ref_beam}")
    doc.add_paragraph(f"Estribos: O1/4\" @ {s_beam*100:.0f} cm (factor sismico: {'d/4' if stirrup_factor==0.25 else 'd/2'})")
    if confinement_zone_length > 0:
        doc.add_paragraph(f"Zona confinamiento: {confinement_zone_length*100:.0f} cm desde apoyo (NSR-10 C.21)")
    doc.add_paragraph()

    # CAP 7 — Verificaciones
    _h("7. Verificaciones Normativas")
    tv=doc.add_table(rows=1,cols=5); tv.style="Table Grid"
    _tbl_hdr(tv,["Parametro","Articulo NSR-10","Calculado","Limite","Estado"])
    for v in verificaciones:
        _row(tv,[v["item"],v["referencia"],v["calculado"],v["requerido"],
                 "CUMPLE" if v["cumple"] else "NO CUMPLE"],
             bg="D6E4F0" if not v["cumple"] else None)
    doc.add_paragraph()

    # Grafica distribucion de costos
    if _MPL_OK:
        try:
            labels_c = ["Bloques","Perfiles","Malla","Concreto","Mano de obra"]
            values_c = [costo_bloques,costo_perfiles,costo_malla,costo_concreto,costo_mo]
            colors_c = ["#C19A6B","#C0C0C0","#4CAF50","#607D8B","#FF9800"]
            fig_cos, (axb, axp) = plt.subplots(1, 2, figsize=(10, 4))
            axb.barh(labels_c, values_c, color=colors_c)
            axb.set_xlabel(f"Costo ({moneda})"); axb.set_title("Costo por componente")
            axb.grid(True, alpha=0.3, axis="x")
            axp.pie(values_c, labels=labels_c, colors=colors_c,
                    autopct="%1.1f%%", startangle=140,
                    textprops={"fontsize": 8})
            axp.set_title("Distribucion de costos")
            plt.tight_layout()
            png2 = fig_to_docx_white(fig_cos)
            plt.close(fig_cos)
            if png2:
                _ip2 = doc.add_paragraph(); _ip2.alignment=WD_ALIGN_PARAGRAPH.CENTER
                _ip2.add_run().add_picture(io.BytesIO(png2), width=Cm(15.0))
                _cap2 = doc.add_paragraph("Figura 2 — Distribucion de Costos por Componente")
                _cap2.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in _cap2.runs: r.font.size=Pt(8); r.font.italic=True; r.font.color.rgb=_GRIS
        except Exception:
            pass

    # CAP 8 — APU
    _h("8. Cuantificacion y APU")
    ta=doc.add_table(rows=1,cols=3); ta.style="Table Grid"
    _tbl_hdr(ta,["Concepto",f"Costo ({moneda})","Nota"])
    for lbl,val,nota in [
        ("Bloques",f"{costo_bloques:,.0f}","c/desperdicio"),
        ("Perfiles Colmena",f"{costo_perfiles:,.0f}","c/desperdicio"),
        ("Malla electrosoldada",f"{costo_malla:,.0f}","c/traslapo"),
        ("Concreto",f"{costo_concreto:,.0f}","c/desperdicio"),
        ("Mano de obra",f"{costo_mo:,.0f}","rendimiento 0.8 dias/m2"),
        ("Herramienta menor",f"{herramienta:,.0f}",f"{pct_herramienta*100:.0f}% s/MO"),
        ("A.I.U.",f"{aiu:,.0f}",f"{pct_aui*100:.0f}% s/costo directo"),
        ("Utilidad",f"{utilidad:,.0f}",f"{pct_util*100:.0f}% s/costo directo"),
        ("IVA s/Utilidad",f"{iva_util:,.0f}",f"{iva_pct*100:.0f}%"),
        ("TOTAL PROYECTO",f"{total_proyecto:,.0f}",""),
    ]:
        bold = (lbl=="TOTAL PROYECTO")
        _row(ta,[lbl,val,nota],bold=bold,bg="1E3A5F" if bold else None)
    doc.add_paragraph(f"Costo por m2: {moneda} {total_proyecto/area_total:,.0f}")
    doc.add_paragraph()

    # CAP 9 — Responsables y referencias
    _h("9. Firmas y Responsables")
    tr=doc.add_table(rows=1,cols=2); tr.style="Table Grid"
    _tbl_hdr(tr,["Rol","Nombre"])
    for rol,key in [("Empresa/Firma","empresa"),("Ingeniero Responsable","ingeniero"),
                    ("Elaboro","pf_elaboro"),("Reviso","pf_reviso"),("Aprobo","pf_aprobo")]:
        _row(tr,[rol,st.session_state.get(key,"________________") or "________________"])
    doc.add_paragraph()
    _h("10. Referencias Normativas")
    for ref in [
        f"{norma_sel} — Norma de diseno aplicable",
        "NSR-10 C.9.5.2 — Control de deflexiones",
        "NSR-10 C.7.12 — Acero minimo de temperatura y retraccion",
        "NSR-10 C.11 — Diseno por cortante",
        "NSR-10 C.12.2.2 — Longitud de desarrollo",
        "NSR-10 C.21 — Confinamiento sismico",
        f"Precios referencia: {st.session_state.get('_price_date','Q1-2026')} — verificar con cotizacion",
    ]:
        doc.add_paragraph(f"- {ref}")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR PDF
# ══════════════════════════════════════════════════════════════════════════════
def generate_pdf():
    if not _PDF_OK:
        return None
    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=3*rl_cm, rightMargin=2.5*rl_cm,
                                topMargin=2.5*rl_cm, bottomMargin=2.5*rl_cm)
    styles = getSampleStyleSheet()
    AZUL   = rl_colors.HexColor("#1e3a5f")
    AZUL_CL= rl_colors.HexColor("#d6e4f0")
    GRIS   = rl_colors.HexColor("#444444")
    sT  = RLPStyle("tit",parent=styles["Title"],textColor=AZUL,fontSize=18,alignment=RL_CENTER)
    sS  = RLPStyle("sub",parent=styles["Normal"],textColor=GRIS,fontSize=11,alignment=RL_CENTER)
    sH  = RLPStyle("h2",parent=styles["Heading2"],textColor=AZUL,fontSize=13,spaceBefore=12,spaceAfter=6)
    sB  = RLPStyle("body",parent=styles["Normal"],fontSize=9,textColor=GRIS,spaceAfter=3)

    story = []
    # Logo
    logo_b = st.session_state.get("logo_bytes")
    if logo_b:
        try:
            story.append(RLImage(io.BytesIO(logo_b), width=4*rl_cm, height=2*rl_cm))
            story.append(RLSpacer(1,0.3*rl_cm))
        except Exception: pass

    _emp_p = st.session_state.get("empresa","________________") or "________________"
    _proy_p = proyecto_nombre or st.session_state.get("proyecto","________________") or "________________"
    story.append(Paragraph(_emp_p, sT))
    story.append(Paragraph("RESUMEN EJECUTIVO — PLACA FACIL", sS))
    story.append(HRFlowable(width="100%",thickness=1.5,color=AZUL,spaceAfter=8))

    meta=[["Proyecto:", _proy_p],["Norma:",norma_sel],
          ["Cliente:", proyecto_cliente or "________________"],
          ["Ingeniero:", st.session_state.get("ingeniero","________________") or "________________"],
          ["Fecha:", datetime.now().strftime("%d/%m/%Y")],
          ["Estado precios:", st.session_state.get("_price_status_label","Verificar cotizacion")],]
    mt = Table(meta, colWidths=[3.5*rl_cm,12*rl_cm])
    mt.setStyle(TableStyle([
        ("FONTSIZE",(0,0),(-1,-1),9),("TEXTCOLOR",(0,0),(0,-1),AZUL),
        ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]))
    story.append(mt); story.append(RLSpacer(1,0.6*rl_cm))

    story.append(Paragraph("1. Resultados Clave", sH))
    td=[["Parametro","Valor","Estado"]]
    for v in verificaciones:
        td.append([v["item"], v["calculado"], "CUMPLE" if v["cumple"] else "NO CUMPLE"])
    vt = Table(td, colWidths=[7*rl_cm,4*rl_cm,3*rl_cm])
    vt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),AZUL),("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white,AZUL_CL]),
        ("GRID",(0,0),(-1,-1),0.4,rl_colors.lightgrey),("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]))
    story.append(vt); story.append(RLSpacer(1,0.4*rl_cm))

    story.append(Paragraph("2. Resumen Economico", sH))
    ed=[["Concepto",f"Costo ({moneda})"]]
    for lbl,val in [("Costo directo",f"{costo_directo:,.0f}"),("A.I.U.",f"{aiu:,.0f}"),
                    ("Utilidad",f"{utilidad:,.0f}"),("IVA s/Utilidad",f"{iva_util:,.0f}"),
                    ("TOTAL PROYECTO",f"{total_proyecto:,.0f}"),
                    ("Costo por m2",f"{total_proyecto/area_total:,.0f}"),
                    ("Comparacion vs maciza",f"{'Ahorro' if ahorro>0 else 'Sobrecosto'}: {max(ahorro,sobrecosto):,.0f}"),]:
        ed.append([lbl,val])
    et = Table(ed, colWidths=[8*rl_cm,6*rl_cm])
    et.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),AZUL),("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white,AZUL_CL]),
        ("GRID",(0,0),(-1,-1),0.4,rl_colors.lightgrey),("ALIGN",(1,1),(-1,-1),"RIGHT"),
    ]))
    story.append(et); story.append(RLSpacer(1,0.4*rl_cm))

    story.append(Paragraph("3. Firmas y Responsables", sH))
    rd=[["Rol","Nombre"]]
    for rol,key in [("Empresa/Firma","empresa"),("Ingeniero","ingeniero"),
                    ("Elaboro","pf_elaboro"),("Reviso","pf_reviso"),("Aprobo","pf_aprobo")]:
        rd.append([rol, st.session_state.get(key,"________________") or "________________"])
    rt = Table(rd, colWidths=[5*rl_cm,9*rl_cm])
    rt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),AZUL),("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),8),
        ("GRID",(0,0),(-1,-1),0.4,rl_colors.lightgrey),
    ]))
    story.append(rt)

    doc_pdf.build(story)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# TABS PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════
tabs = st.tabs(["Resultados","Modelo 3D","DXF e IFC","Memoria DOCX",
                "Cantidades","APU","Resumen y PDF"])

# ── Tab Resultados ─────────────────────────────────────────────────────────────
with tabs[0]:
    st.subheader("Resultados del diseno")
    c1,c2,c3 = st.columns(3)
    with c1:
        st.metric("Area total",f"{area_total:.2f} m2")
        st.metric("N° perfiles",f"{n_profiles}")
        st.metric("Long. perfiles",f"{longitud_total_perfiles_desp:.1f} m")
        st.metric("N° bloques",f"{n_bloques_desp} und")
    with c2:
        st.metric("Vol. concreto",f"{vol_concreto_total_desp:.2f} m3")
        st.metric("Area malla",f"{area_malla:.2f} m2")
        st.metric("Carga muerta",f"{carga_muerta_kgm2:.0f} kg/m2")
        st.metric("Cemento",f"{bultos_cemento} bultos")
    with c3:
        st.metric("Mu por perfil",f"{Mu:.2f} kN·m")
        st.metric("Deflexion",f"{delta_calc*1000:.1f} mm")
        st.metric("Limite deflexion",f"{delta_max*1000:.1f} mm")
        st.metric("Vu placa",f"{Vu:.2f} kN")

    st.markdown("#### Viga de borde")
    st.write(f"Momento: Mu = {Mu_beam:.2f} kN·m | As req = {As_beam:.2f} cm2 | As prov = {As_prov_beam:.2f} cm2")
    st.write(f"Cortante: Vu = {Vu_beam:.2f} kN | Vc = {Vc_beam:.2f} kN")
    st.write(f"Refuerzo: {ref_beam}, estribos @ {s_beam*100:.0f} cm")
    if confinement_zone_length > 0:
        st.info(f"Zona de confinamiento: {confinement_zone_length*100:.0f} cm desde apoyo | Est. @ {s_beam*100:.0f} cm (NSR-10 C.21)")

    st.markdown("#### Seccion compuesta y longitud de desarrollo")
    cs1,cs2,cs3 = st.columns(3)
    with cs1:
        st.metric("I_comp",f"{I_comp*1e8:.4f} x10-8 m4")
        st.metric("y_c centroide",f"{y_c*100:.2f} cm")
    with cs2:
        st.metric("S_comp",f"{S_comp*1e6:.4f} x10-6 m3")
        st.metric("n modular",f"{n_mod:.1f}")
    with cs3:
        st.metric("Ec concreto",f"{E_concreto/1e9:.2f} GPa")
        st.metric("ld desarrollo",f"{ld*100:.1f} cm")

    st.markdown("#### Diagramas viga de borde")
    if _MPL_OK:
        fig_d, (ax1,ax2) = plt.subplots(1,2,figsize=(10,4))
        fig_d.patch.set_facecolor("white")
        for ax in [ax1,ax2]: ax.set_facecolor("#f8f9fa")
        ax1.plot(x_vals,M_vals,'b-',linewidth=2)
        ax1.fill_between(x_vals,0,M_vals,alpha=0.25,color='blue')
        ax1.set_xlabel("Distancia (m)"); ax1.set_ylabel("Momento (kN·m)")
        ax1.set_title("Diagrama de Momento — Viga Borde"); ax1.grid(True,alpha=0.3)
        ax2.plot(x_vals,V_vals,'r-',linewidth=2)
        ax2.fill_between(x_vals,0,V_vals,alpha=0.25,color='red')
        ax2.set_xlabel("Distancia (m)"); ax2.set_ylabel("Cortante (kN)")
        ax2.set_title("Diagrama de Cortante — Viga Borde"); ax2.grid(True,alpha=0.3)
        plt.tight_layout()
        st.pyplot(fig_d); plt.close()
    else:
        st.info("Instala matplotlib para ver los diagramas.")

    st.markdown("#### Verificaciones normativas")
    for v in verificaciones:
        if v["cumple"]:
            st.success(f"[CUMPLE] {v['item']}: {v['calculado']} — {v['obs']}")
        else:
            st.error(f"[NO CUMPLE] {v['item']}: {v['calculado']} — {v['obs']}")
        st.caption(f"Referencia: {v['referencia']}")

# ── Tab Modelo 3D ──────────────────────────────────────────────────────────────
with tabs[1]:
    st.subheader("Modelo 3D de la placa")
    fig_3d = create_3d_model(Lx,Ly,orientacion,n_profiles,perfil_espaciado,perfil_largo,
                             espesor_torta,incluir_vigas,viga_b,viga_h)
    st.plotly_chart(fig_3d, use_container_width=True)

# ── Tab DXF e IFC ──────────────────────────────────────────────────────────────
with tabs[2]:
    st.subheader("Exportar planos — DXF e IFC")

    # ── Admin automático por email ─────────────────────────────────────────────
    _ADMIN_EMAILS_PF = {"civcesar@gmail.com"}
    try:
        _pf_auth = st.session_state.get("auth_user")
        _pf_email = (getattr(_pf_auth, "email", None) or "").lower().strip() if _pf_auth else ""
    except Exception:
        _pf_email = ""
    if _pf_email in _ADMIN_EMAILS_PF:
        st.session_state["user_role"] = "admin"

    _rol_t = st.session_state.get("user_role","free")
    _puede = _rol_t in ("admin","pro")

    col_dxf, col_ifc = st.columns(2)

    with col_dxf:
        st.markdown("#### Plano DXF profesional")
        st.caption("6 detalles ICONTEC 2289: Planta, Corte A-A', Perfil Colmena, "
                   "Perfil Reforzado, Bloquelon Santafe, Vigueta Coronamiento. "
                   "Todo Color 7. Sin dependencias externas.")
        if not _puede:
            st.button("Generar DXF", disabled=True,
                      help="Requiere rol Pro o Admin", use_container_width=True)
        else:
            if st.button("Generar DXF", key="btn_dxf", use_container_width=True, type="primary"):
                with st.spinner("Generando DXF..."):
                    dxf_data = generate_dxf_profesional()
                _slug = (proyecto_nombre or "Placa").replace(" ","_")
                st.download_button(
                    "Descargar DXF", data=dxf_data,
                    file_name=f"PlacaFacil_{_slug}_{datetime.now().strftime('%Y%m%d')}.dxf",
                    mime="application/dxf", use_container_width=True)

    with col_ifc:
        st.markdown("#### Modelo IFC 2x3")
        st.caption("Losa, vigas de borde, perfiles metalicos y bloquelones "
                   "como entidades IFC con atributos de material y propietario.")
        if not _puede:
            st.button("Generar IFC", disabled=True,
                      help="Requiere rol Pro o Admin", use_container_width=True)
        else:
            if st.button("Generar IFC", key="btn_ifc", use_container_width=True, type="primary"):
                with st.spinner("Generando IFC..."):
                    ifc_data = generate_ifc()
                _slug2 = (proyecto_nombre or "Placa").replace(" ","_")
                st.download_button(
                    "Descargar IFC", data=ifc_data,
                    file_name=f"PlacaFacil_{_slug2}_{datetime.now().strftime('%Y%m%d')}.ifc",
                    mime="application/octet-stream", use_container_width=True)

    if not _puede:
        st.warning("Rol Pro o Admin requerido para exportar DXF e IFC. "
                   "Cambia tu rol en 'Rol de Usuario' en el panel lateral.")

# ── Tab Memoria DOCX ───────────────────────────────────────────────────────────
with tabs[3]:
    st.subheader("Memoria de calculo — DOCX")
    _puede3 = st.session_state.get("user_role","free") in ("admin","pro")
    if not _puede3:
        st.button("Generar DOCX", disabled=True,
                  help="Requiere rol Pro o Admin", use_container_width=True)
        st.warning("Cambia tu rol en 'Rol de Usuario' en el panel lateral para desbloquear.")
    elif not _DOCX_OK:
        st.button("Generar DOCX", disabled=True,
                  help="Instala: pip install python-docx", use_container_width=True)
    else:
        if st.button("Generar Memoria DOCX", key="btn_docx",
                     use_container_width=True, type="primary"):
            with st.spinner("Generando DOCX con graficas..."):
                buf = generate_memory()
            if buf:
                _slug3 = (proyecto_nombre or "Placa").replace(" ","_")
                st.download_button(
                    "Descargar Memoria DOCX", data=buf,
                    file_name=f"Memoria_PlacaFacil_{_slug3}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            else:
                st.error("Error generando DOCX.")

# ── Tab Cantidades ─────────────────────────────────────────────────────────────
with tabs[4]:
    st.subheader("Cantidades de materiales")
    df_qty = pd.DataFrame({
        "Material": ["Bloques","Perfiles Colmena","Malla electrosoldada","Concreto","Cemento"],
        "Cantidad": [f"{n_bloques_desp} und",f"{longitud_total_perfiles_desp:.1f} m",
                     f"{area_malla:.2f} m2",f"{vol_concreto_total_desp:.2f} m3",
                     f"{bultos_cemento} bultos"],
        "Desperdicio incluido": ["Si","Si","Si (traslapo)","Si","Si"],
    })
    st.dataframe(df_qty, use_container_width=True)
    st.markdown("#### Desglose concreto")
    st.write(f"- Torta: {vol_torta:.3f} m3")
    st.write(f"- Vigas de borde: {vol_vigas:.3f} m3")
    st.write(f"- Total sin desperdicio: {vol_concreto_total:.3f} m3")
    st.write(f"- Total con desperdicio: {vol_concreto_total_desp:.3f} m3")
    fig_qty = go.Figure(go.Bar(
        x=["Bloques","Perfiles","Malla","Concreto","Mano de obra"],
        y=[costo_bloques,costo_perfiles,costo_malla,costo_concreto,costo_mo],
        marker_color=["#C19A6B","#C0C0C0","#4CAF50","#607D8B","#FF9800"]))
    _plotly_light(fig_qty)
    fig_qty.update_layout(title="Costo por componente",
                          yaxis_title=f"Costo ({moneda})")
    st.plotly_chart(fig_qty, use_container_width=True)

# ── Tab APU ────────────────────────────────────────────────────────────────────
with tabs[5]:
    st.subheader("Analisis de Precios Unitarios (APU)")
    df_apu = pd.DataFrame({
        "Item": ["Bloques","Perfiles","Malla","Concreto","Mano de obra",
                 "Herramienta","A.I.U.","Utilidad","IVA s/Utilidad","TOTAL"],
        "Costo": [costo_bloques,costo_perfiles,costo_malla,costo_concreto,costo_mo,
                  herramienta,aiu,utilidad,iva_util,total_proyecto],
    })
    df_apu[f"Costo ({moneda})"] = df_apu["Costo"].apply(lambda x: f"{x:,.0f}")
    st.dataframe(df_apu[["Item",f"Costo ({moneda})"]], use_container_width=True)
    c_a1,c_a2,c_a3 = st.columns(3)
    c_a1.metric("Costo directo",f"{moneda} {costo_directo:,.0f}")
    c_a2.metric("TOTAL PROYECTO",f"{moneda} {total_proyecto:,.0f}")
    c_a3.metric("Costo por m2",f"{moneda} {total_proyecto/area_total:,.0f}")
    st.markdown("#### Cronograma estimado")
    st.dataframe(cronograma, use_container_width=True)
    fig_g = go.Figure(go.Bar(
        x=cronograma["Duracion (dias)"], y=cronograma["Actividad"],
        orientation="h", marker_color="#1565c0"))
    _plotly_light(fig_g, height=300)
    fig_g.update_layout(title="Cronograma de actividades",
                        xaxis_title="Dias")
    st.plotly_chart(fig_g, use_container_width=True)

    # Analisis de sensibilidad
    st.markdown("---")
    with st.expander("Analisis de Sensibilidad de Precios", expanded=False):
        _total_b = total_proyecto
        if _total_b > 0:
            _sc1, _sc2 = st.columns([2,3])
            with _sc1:
                _mat_s = st.selectbox("Material a variar:",
                    ["Bloques","Perfiles","Malla","Concreto","Mano de obra","Todos"],
                    key="pf_sens_mat")
                _var_p = st.slider("Variacion (%)", -50, 100, 15, 5, key="pf_sens_var")
            with _sc2:
                _wmap = {"Bloques":0.20,"Perfiles":0.25,"Malla":0.08,
                         "Concreto":0.30,"Mano de obra":0.10,"Todos":1.0}
                _delta = _total_b * _wmap.get(_mat_s,0.20) * (_var_p/100.0)
                _nuevo = _total_b + _delta
                _col_s = "#c62828" if _delta > 0 else "#2e7d32"
                _sig = "+" if _delta > 0 else ""
                st.markdown(
                    f'<div style="background:#f5f5f5;border-radius:8px;padding:14px;margin-top:6px;">'
                    f'<p style="margin:0;color:#555;font-size:0.9rem;">Si <b>{_mat_s}</b> varia <b>{_sig}{_var_p}%</b>:</p>'
                    f'<h3 style="color:{_col_s};margin:6px 0;">{_sig}{moneda} {_delta:,.0f}</h3>'
                    f'<p style="margin:0;color:#333;">Impacto: <b>{_sig}{(_delta/_total_b*100):.1f}%</b></p>'
                    f'<p style="margin:4px 0 0;font-size:0.85rem;color:#777;">'
                    f'Base: {moneda} {_total_b:,.0f} → Ajustado: {moneda} {_nuevo:,.0f}</p></div>',
                    unsafe_allow_html=True)

# ── Tab Resumen y PDF ──────────────────────────────────────────────────────────
with tabs[6]:
    st.subheader("Resumen ejecutivo")
    cr1,cr2 = st.columns(2)
    with cr1:
        st.metric("Costo total",f"{moneda} {total_proyecto:,.0f}")
        st.metric("Costo por m2",f"{moneda} {total_proyecto/area_total:,.0f}")
        cumplidas = sum(1 for v in verificaciones if v["cumple"])
        st.metric("Verificaciones cumplidas",f"{cumplidas}/{len(verificaciones)}")
    with cr2:
        if sobrecosto > 0:
            st.warning(f"Sobrecosto vs placa maciza: {moneda} {sobrecosto:,.0f}")
        else:
            st.success(f"Ahorro vs placa maciza: {moneda} {ahorro:,.0f}")
        st.metric("Deflexion",f"{'OK' if cumple_deflexion else 'EXCEDE'} — {delta_calc*1000:.1f} mm")
        st.metric("Zona sismica", zona_sismica)

    st.markdown("---")
    _puede_pdf = st.session_state.get("user_role","free") in ("admin","pro")
    if not _puede_pdf:
        st.button("Generar PDF", disabled=True,
                  help="Requiere rol Pro o Admin", use_container_width=True)
        st.warning("Cambia tu rol en 'Rol de Usuario' en el panel lateral.")
    elif not _PDF_OK:
        st.button("Generar PDF", disabled=True,
                  help="Instala: pip install reportlab", use_container_width=True)
    else:
        if st.button("Generar PDF Resumen", key="btn_pdf",
                     use_container_width=True, type="primary"):
            with st.spinner("Generando PDF..."):
                pdf_data = generate_pdf()
            if pdf_data:
                _slug4 = (proyecto_nombre or "Placa").replace(" ","_")
                st.download_button(
                    "Descargar PDF", data=pdf_data,
                    file_name=f"Resumen_PlacaFacil_{_slug4}_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf", use_container_width=True)

    # Excel / CSV libre
    st.markdown("---")
    st.markdown("#### Exportacion libre (todos los roles)")
    ex1, ex2 = st.columns(2)
    with ex1:
        df_export = pd.DataFrame([
            {"Concepto":"Area total","Valor":f"{area_total:.2f}","Unidad":"m2"},
            {"Concepto":"Perfiles","Valor":f"{longitud_total_perfiles_desp:.1f}","Unidad":"m"},
            {"Concepto":"Bloques","Valor":str(n_bloques_desp),"Unidad":"und"},
            {"Concepto":"Malla","Valor":f"{area_malla:.2f}","Unidad":"m2"},
            {"Concepto":"Concreto","Valor":f"{vol_concreto_total_desp:.2f}","Unidad":"m3"},
            {"Concepto":"Cemento","Valor":str(bultos_cemento),"Unidad":"bultos"},
            {"Concepto":"Mu placa","Valor":f"{Mu:.2f}","Unidad":"kN·m"},
            {"Concepto":"Deflexion","Valor":f"{delta_calc*1000:.2f}","Unidad":"mm"},
            {"Concepto":"Total proyecto","Valor":f"{total_proyecto:,.0f}","Unidad":moneda},
        ])
        csv_b = df_export.to_csv(index=False).encode("utf-8")
        st.download_button("Descargar CSV", data=csv_b,
                           file_name=f"PlacaFacil_{datetime.now().strftime('%Y%m%d')}.csv",
                           mime="text/csv", use_container_width=True)
    with ex2:
        if st.button("Limpiar cache de estado", key="btn_clear_pf",
                     use_container_width=True, type="secondary"):
            for k in _PERSIST_KEYS:
                st.session_state.pop(k, None)
            if os.path.exists(STATE_FILE):
                os.remove(STATE_FILE)
            st.rerun()
