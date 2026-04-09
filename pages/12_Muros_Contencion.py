import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import plotly.graph_objects as go
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Muros de Contención", "Retaining Walls"), layout="wide")

st.image(r"assets/retaining_wall_header_1773256923525.png", use_container_width=True)
st.title(_t("Muros de Contención y Estabilidad", "Retaining Walls and Stability"))
st.markdown(_t("Herramientas para revisar la estabilidad al volcamiento y deslizamiento de muros de contención de gravedad y en voladizo, considerando empujes de tierras y sobrecargas, con diseño estructural completo y despiece de acero.", "Tools to verify overturning and sliding stability for gravity and cantilever retaining walls, considering earth pressures and surcharges, with full structural design and steel bending schedule."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    <br><br>
    <i>⚠ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# DICCIONARIOS DE BARRAS
# ─────────────────────────────────────────────
REBAR_US = {
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
    "#5 (5/8\")": {"area": 1.99, "diam_mm": 15.88},
    "#6 (3/4\")": {"area": 2.84, "diam_mm": 19.05},
    "#7 (7/8\")": {"area": 3.87, "diam_mm": 22.23},
    "#8 (1\")":   {"area": 5.10, "diam_mm": 25.40},
}
REBAR_MM = {
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
    "14 mm": {"area": 1.539, "diam_mm": 14.0},
    "16 mm": {"area": 2.011, "diam_mm": 16.0},
    "18 mm": {"area": 2.545, "diam_mm": 18.0},
    "20 mm": {"area": 3.142, "diam_mm": 20.0},
    "22 mm": {"area": 3.801, "diam_mm": 22.0},
    "25 mm": {"area": 4.909, "diam_mm": 25.0},
}
STIRRUP_US = {
    "#2 (1/4\")": {"area": 0.32, "diam_mm": 6.35},
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
}
STIRRUP_MM = {
    "6 mm": {"area": 0.283, "diam_mm": 6.0},
    "8 mm": {"area": 0.503, "diam_mm": 8.0},
}

# ─────────────────────────────────────────────
# PARÁMETROS POR NORMA
# ─────────────────────────────────────────────
CODES = {
    "NSR-10 (Colombia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "NSR-10 Título C", "bag_kg": 50.0
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "ACI 318-25", "bag_kg": 42.6
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "ACI 318-19", "bag_kg": 42.6
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "ACI 318-14", "bag_kg": 42.6
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "NEC-SE-HM Ecuador", "bag_kg": 50.0
    },
    "E.060 (Perú)": {
        "phi_flex": 0.90, "phi_shear": 0.85, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "Norma E.060 Perú", "bag_kg": 42.5
    },
    "NTC-EM (México)": {
        "phi_flex": 0.85, "phi_shear": 0.80, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "NTC-EM México", "bag_kg": 50.0
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "COVENIN 1753-2006", "bag_kg": 42.5
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "NB 1225001-2020", "bag_kg": 50.0
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "lambda": 1.0,
        "rho_min": 0.0018, "ref": "CIRSOC 201-2025", "bag_kg": 50.0
    },
}

# Funciones auxiliares
def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def mix_for_fc(fc):
    """ACI 211 mix proportions: (cement_kg, water_L, sand_kg, gravel_kg) per m³"""
    table = [
        (14, 250, 205, 810, 1060), (17, 290, 200, 780, 1060), (21, 350, 193, 720, 1060),
        (25, 395, 193, 680, 1020), (28, 430, 190, 640, 1000), (35, 530, 185, 580, 960),
        (42, 620, 180, 520, 910), (56, 740, 175, 450, 850),
    ]
    if fc <= table[0][0]: return table[0][1:]
    if fc >= table[-1][0]: return table[-1][1:]
    for i in range(len(table)-1):
        lo, hi = table[i], table[i+1]
        if lo[0] <= fc <= hi[0]:
            t = (fc - lo[0]) / (hi[0] - lo[0])
            return tuple(lo[j] + t*(hi[j]-lo[j]) for j in range(1,5))
    return table[-1][1:]

def get_development_length(db_mm, fy, fc, lambda_=1.0, psi_t=1.0, psi_e=1.0, psi_s=1.0, psi_g=1.0, cb_ktr=2.5):
    """Longitud de desarrollo a tracción (ACI 318-19 25.4.2.3)"""
    if db_mm <= 0: return 0
    ld = (3/40) * (fy / (lambda_ * math.sqrt(fc))) * (psi_t * psi_e * psi_s * psi_g / cb_ktr) * db_mm
    return max(ld, 300)  # mínimo 300 mm

# ─────────────────────────────────────────────
# CONFIGURACIÓN GENERAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙ Norma y Materiales", "⚙ Code and Materials"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = list(CODES.keys())[0]
norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagpedia.net/data/flags/mini/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)
code = CODES[norma_sel]
phi_f = code["phi_flex"]
phi_v = code["phi_shear"]
lam = code["lambda"]
bag_kg = code["bag_kg"]

st.sidebar.header(_t(" Materiales Concreto/Acero", " Concrete/Steel Materials"))
fc_unit = st.sidebar.radio(_t("Unidad f'c:", "f'c Unit:"), ["MPa", "PSI", "kg/cm²"], horizontal=True, key="m_fc_unit")
if fc_unit == "PSI":
    fc_psi = st.sidebar.number_input("f'c [PSI]", 2000.0, 12000.0, 4000.0, 100.0, key="m_fc_psi")
    fc = fc_psi * 0.00689476
elif fc_unit == "kg/cm²":
    fc_kgcm2 = st.sidebar.number_input("f'c [kg/cm²]", 100.0, 1200.0, 280.0, 10.0, key="m_fc_kgcm2")
    fc = fc_kgcm2 / 10.1972
else:
    fc = st.sidebar.number_input(_t("Resistencia del Concreto (f'c) [MPa]", "Concrete Strength (f'c) [MPa]"), 15.0, 80.0, 28.0, 1.0, key="m_fc_mpa")
fy = st.sidebar.number_input(_t("Fluencia del Acero (fy) [MPa]", "Steel Yield (fy) [MPa]"), 240.0, 500.0, 420.0, 10.0, key="m_fy")
Es = 200000.0
beta1 = get_beta1(fc)
rho_min = code["rho_min"]

st.sidebar.header(_t(" Geometría del Muro", " Wall Geometry"))
H_muro = st.sidebar.number_input("Altura Total H [m]", 1.0, 15.0, 4.0, 0.5, key="m_H")
B_base = st.sidebar.number_input("Ancho de Base B [m]", 0.5, 10.0, 2.5, 0.5, key="m_B")
espesor_base = st.sidebar.number_input("Espesor de la Base (Zapata) [m]", 0.2, 2.0, 0.5, 0.1, key="m_ebase")
pie_muro = st.sidebar.number_input("Longitud del Pie (Toe) [m]", 0.0, 5.0, 0.6, 0.1, key="m_pie")
corona = st.sidebar.number_input("Ancho Corona (Top) [m]", 0.2, 2.0, 0.3, 0.1, key="m_corona")
base_pantalla = st.sidebar.number_input("Ancho Base Pantalla [m]", 0.2, 3.0, 0.4, 0.1, key="m_bpantalla")
talon = B_base - pie_muro - base_pantalla
if talon < 0:
    st.sidebar.error("Geometría inválida: Talón negativo.")
    talon = 0

st.sidebar.header(_t(" Propiedades del Suelo", " Soil Properties"))
gamma_s = st.sidebar.number_input(_t("Peso Unitario Suelo γ [kN/m³]", "Soil Unit Weight γ [kN/m³]"), 10.0, 22.0, 18.0, 0.5, key="m_gamma_s")
phi_ang = st.sidebar.number_input(_t("Ángulo de Fricción φ [°]", "Friction Angle φ [°]"), 20.0, 45.0, 30.0, 1.0, key="m_phi")
c_base = st.sidebar.number_input(_t("Cohesión en la base c [kPa]", "Base cohesion c [kPa]"), 0.0, 100.0, 0.0, 5.0, key="m_c")
delta_ang = st.sidebar.number_input(_t("Fricción suelo-muro δ [°] (Fricción base)", "Soil-wall friction δ [°]"), 10.0, 40.0, 20.0, 1.0, key="m_delta")
q_adm = st.sidebar.number_input(_t("Capacidad Portante Admisible q_adm [kPa]", "Allowable Bearing Capacity q_adm [kPa]"), 50.0, 500.0, 150.0, 10.0, key="m_qadm")
gamma_conc = 24.0

st.sidebar.header(_t(" Refuerzo (Acero)", " Reinforcement (Steel)"))
bar_sys = st.sidebar.radio(_t("Sistema Varillas:", "Rebar System:"), ["Pulgadas (# US)","Milímetros (mm)"], horizontal=True, key="m_bar_sys")
rebar_dict = REBAR_US if "Pulgadas" in bar_sys else REBAR_MM
stirrup_dict = STIRRUP_US if "Pulgadas" in bar_sys else STIRRUP_MM
default_rebar = "#5 (5/8\")" if "Pulgadas" in bar_sys else "16 mm"
rebar_vert = st.sidebar.selectbox(_t("Varilla principal (pantalla)", "Main rebar (wall)"), list(rebar_dict.keys()), index=list(rebar_dict.keys()).index(default_rebar), key="m_rebar")
rebar_area = rebar_dict[rebar_vert]["area"]  # cm²
rebar_diam = rebar_dict[rebar_vert]["diam_mm"]
stirrup_sel = st.sidebar.selectbox(_t("Estribo (pantalla)", "Stirrup (wall)"), list(stirrup_dict.keys()), index=0, key="m_stirrup")
stirrup_area = stirrup_dict[stirrup_sel]["area"]  # cm²
stirrup_diam = stirrup_dict[stirrup_sel]["diam_mm"]

# Sobrecarga y talud
st.sidebar.header(_t(" Sobrecarga y Talud", " Surcharge and Slope"))
beta_ang = st.sidebar.number_input("Inclinación del Terraplén β [°]", 0.0, phi_ang-0.1, 0.0, 1.0, key="m_beta")
q_sobrecarga = st.sidebar.number_input("Sobrecarga uniforme q [kPa]", 0.0, 100.0, 10.0, 2.0, key="m_q")
FS_v_min = st.sidebar.number_input("FS Volcamiento Mínimo", 1.0, 3.0, 1.5, 0.1, key="m_fsv_min")
FS_d_min = st.sidebar.number_input("FS Deslizamiento Mínimo", 1.0, 3.0, 1.5, 0.1, key="m_fsd_min")

# Longitud total del muro (para despiece)
st.sidebar.header(_t(" Longitud Total del Muro", " Total Wall Length"))
L_total = st.sidebar.number_input(_t("Longitud total L [m]", "Total length L [m]"), 1.0, 100.0, 10.0, 1.0, key="m_Ltotal")

# =============================================================================
# CÁLCULOS GEOTÉCNICOS Y ESTRUCTURALES (por metro lineal)
# =============================================================================
phi_rad = math.radians(phi_ang)
beta_rad = math.radians(beta_ang)
if beta_ang > 0:
    num = math.cos(beta_rad) - math.sqrt(math.cos(beta_rad)**2 - math.cos(phi_rad)**2)
    den = math.cos(beta_rad) + math.sqrt(math.cos(beta_rad)**2 - math.cos(phi_rad)**2)
    Ka = math.cos(beta_rad) * (num / den)
else:
    Ka = math.tan(math.radians(45) - phi_rad/2)**2

H_prima = H_muro + talon * math.tan(beta_rad) if talon > 0 else H_muro

# Empuje activo
Pa_suelo = 0.5 * gamma_s * (H_prima**2) * Ka
Pah_suelo = Pa_suelo * math.cos(beta_rad)
Pav_suelo = Pa_suelo * math.sin(beta_rad)
y_pa = H_prima / 3.0
Pa_q = q_sobrecarga * H_prima * Ka
Pah_q = Pa_q * math.cos(beta_rad)
Pav_q = Pa_q * math.sin(beta_rad)
y_q = H_prima / 2.0

# Pesos y momentos resistentes
items = []
# Base
W_base = B_base * espesor_base * gamma_conc
x_base = B_base / 2.0
items.append(("Base", W_base, x_base))
# Pantalla rectangular
W_pant_rect = corona * (H_muro - espesor_base) * gamma_conc
x_pant_rect = pie_muro + (base_pantalla - corona) + corona/2.0
items.append(("Pantalla Rect", W_pant_rect, x_pant_rect))
# Pantalla triangular
ancho_triang = base_pantalla - corona
if ancho_triang > 0:
    W_pant_tri = 0.5 * ancho_triang * (H_muro - espesor_base) * gamma_conc
    x_pant_tri = pie_muro + (2.0/3.0) * ancho_triang
    items.append(("Pantalla Triang", W_pant_tri, x_pant_tri))
# Suelo sobre talón
if talon > 0:
    W_suelo_talon = talon * (H_muro - espesor_base) * gamma_s
    x_suelo_talon = pie_muro + base_pantalla + talon/2.0
    items.append(("Suelo Talón", W_suelo_talon, x_suelo_talon))
# Cuña inclinada
if beta_ang > 0 and talon > 0:
    W_suelo_tri = 0.5 * talon * (talon * math.tan(beta_rad)) * gamma_s
    x_suelo_tri = pie_muro + base_pantalla + (2.0/3.0) * talon
    items.append(("Suelo Cuña", W_suelo_tri, x_suelo_tri))

W_total = sum(w for _, w, _ in items)
Mr_pesos = sum(w * x for _, w, x in items)
Mr_Pav_suelo = Pav_suelo * B_base
Mr_Pav_q = Pav_q * B_base
Mr_total = Mr_pesos + Mr_Pav_suelo + Mr_Pav_q
Mo_suelo = Pah_suelo * y_pa
Mo_q = Pah_q * y_q
Mo_total = Mo_suelo + Mo_q
FS_volc = Mr_total / Mo_total if Mo_total > 0 else 999

Fd_total = Pah_suelo + Pah_q
N_total = W_total + Pav_suelo + Pav_q
Fr_desl = N_total * math.tan(math.radians(delta_ang)) + c_base * B_base
FS_desl = Fr_desl / Fd_total if Fd_total > 0 else 999

# Presión de contacto
x_r = (Mr_total - Mo_total) / N_total if N_total > 0 else 0
e = B_base/2 - x_r
q_max = N_total / B_base * (1 + 6*e/B_base)
q_min = N_total / B_base * (1 - 6*e/B_base)
ok_bearing = q_max <= q_adm

# Diseño estructural de la pantalla
H_pantalla = H_muro - espesor_base
Pa_pantalla = 0.5 * gamma_s * H_pantalla**2 * Ka
Pah_pantalla = Pa_pantalla * math.cos(beta_rad)
y_pantalla = H_pantalla / 3.0
Pa_q_pantalla = q_sobrecarga * H_pantalla * Ka
Pah_q_pantalla = Pa_q_pantalla * math.cos(beta_rad)
y_q_pantalla = H_pantalla / 2.0
Mu_pantalla = Pah_pantalla * y_pantalla + Pah_q_pantalla * y_q_pantalla  # kN·m/m

# Peralte efectivo
d_pantalla_cm = H_pantalla * 100 - 5  # recubrimiento 5 cm
if d_pantalla_cm <= 0:
    d_pantalla_cm = 10
b = 100  # cm
Mu_Nmm = Mu_pantalla * 1e6
Rn = Mu_Nmm / (phi_f * b * d_pantalla_cm**2)
disc = 1 - 2*Rn/(0.85*fc)
if disc > 0:
    rho = (0.85*fc/fy) * (1 - math.sqrt(disc))
else:
    rho = rho_min
rho = max(rho, rho_min)
As_vert_req = rho * b * d_pantalla_cm  # cm²/m
As_bar = rebar_area
s_vert = As_bar / As_vert_req * 100  # cm
s_vert = min(s_vert, 45)
if s_vert < 10:
    s_vert = 10
n_barras_vert = math.ceil(100 / s_vert)
As_vert_prov = n_barras_vert * As_bar

# Acero horizontal (mínimo)
rho_h_min = 0.0020
As_hor_req = rho_h_min * b * d_pantalla_cm
s_hor = As_bar / As_hor_req * 100
s_hor = min(s_hor, 45)
if s_hor < 10:
    s_hor = 10
n_barras_hor = math.ceil(H_pantalla * 100 / s_hor)  # por metro

# Estribos
Vu_pantalla = Pah_pantalla + Pah_q_pantalla  # kN/m
Vc = 0.17 * lam * math.sqrt(fc) * b * d_pantalla_cm / 10  # kN
phi_Vc = phi_v * Vc
if Vu_pantalla > phi_Vc/2:
    Av_min = 0.062 * math.sqrt(fc) * b * s_vert / fy  # cm²
    s_est = 30 if Av_min <= stirrup_area * 2 else 15
    n_est = math.ceil(H_pantalla * 100 / s_est) if s_est > 0 else 0
else:
    s_est = 0
    n_est = 0

# Diseño de la base
q_max_base = N_total / B_base * (1 + 6*e/B_base)
q_min_base = N_total / B_base * (1 - 6*e/B_base)
# Pie
x_pie = pie_muro
q_pie = q_min_base + (q_max_base - q_min_base) * (x_pie / B_base)
W_pie = x_pie * espesor_base * gamma_conc
q_avg_pie = (q_min_base + q_pie) / 2
Mu_pie = q_avg_pie * x_pie**2 / 2 - W_pie * x_pie/2
# Talón
x_talon = talon
q_talon = q_pie + (q_max_base - q_min_base) * (x_talon / B_base)
q_avg_talon = (q_pie + q_talon) / 2
W_suelo_talon = talon * H_pantalla * gamma_s
W_conc_talon = talon * espesor_base * gamma_conc
Mu_talon = (W_suelo_talon + W_conc_talon) * talon/2 - q_avg_talon * talon**2 / 2
Mu_base = max(abs(Mu_pie), abs(Mu_talon))

d_base = espesor_base * 100 - 5
Rn_base = Mu_base * 1e6 / (phi_f * b * d_base**2)
disc_base = 1 - 2*Rn_base/(0.85*fc)
if disc_base > 0:
    rho_base = (0.85*fc/fy) * (1 - math.sqrt(disc_base))
else:
    rho_base = rho_min
rho_base = max(rho_base, rho_min)
As_base_req = rho_base * b * d_base
s_base = As_bar / As_base_req * 100
s_base = min(s_base, 45)
n_barras_base = math.ceil(B_base * 100 / s_base)

# =============================================================================
# CANTIDADES TOTALES (para longitud total L_total)
# =============================================================================
vol_conc_ml = B_base * espesor_base + 0.5 * (corona + base_pantalla) * H_pantalla
vol_conc_total = vol_conc_ml * L_total

# Acero vertical
long_vert = H_pantalla + 0.3  # +anclaje (m)
n_barras_vert_total = n_barras_vert * L_total  # cada metro tiene n_barras_vert
peso_vert_total = n_barras_vert_total * long_vert * As_bar * 0.785  # kg

# Acero horizontal
long_hor = L_total  # cada barra corre toda la longitud
n_barras_hor_total = n_barras_hor  # por metro de altura, total para toda la altura
peso_hor_total = n_barras_hor_total * long_hor * As_bar * 0.785

# Estribos
if s_est > 0:
    perim_est = 2 * (b/100 + d_pantalla_cm/100)  # m
    n_est_total = n_est * L_total  # por metro lineal
    peso_est_total = n_est_total * perim_est * stirrup_area * 0.785
else:
    peso_est_total = 0

# Acero base
n_barras_base_total = n_barras_base * L_total
long_base = B_base
peso_base_total = n_barras_base_total * long_base * As_bar * 0.785

peso_total_acero = peso_vert_total + peso_hor_total + peso_est_total + peso_base_total

# Dosificación concreto
cement_kg, water_L, sand_kg, gravel_kg = mix_for_fc(fc)
total_cement_kg = cement_kg * vol_conc_total
total_water_L = water_L * vol_conc_total
total_sand_kg = sand_kg * vol_conc_total
total_gravel_kg = gravel_kg * vol_conc_total
bags_cement = total_cement_kg / bag_kg

# Longitudes de desarrollo y empalmes
ld_vert_mm = get_development_length(rebar_diam, fy, fc, lam, 1.0, 1.0, 1.0, 1.0, 2.5)
splice_length_mm = 1.3 * ld_vert_mm

# =============================================================================
# DESPIECE DE ACERO (BENDING SCHEDULE)
# =============================================================================
despiece = []

# Barra vertical
despiece.append({
    "Marca": "V1",
    "Cantidad": n_barras_vert_total,
    "Diámetro (mm)": rebar_diam,
    "Longitud (m)": long_vert,
    "Longitud Total (m)": n_barras_vert_total * long_vert,
    "Peso (kg)": peso_vert_total,
    "Observación": f"Vertical principal, espaciamiento {s_vert:.0f} cm"
})

# Barra horizontal
despiece.append({
    "Marca": "H1",
    "Cantidad": n_barras_hor_total,
    "Diámetro (mm)": rebar_diam,
    "Longitud (m)": long_hor,
    "Longitud Total (m)": n_barras_hor_total * long_hor,
    "Peso (kg)": peso_hor_total,
    "Observación": f"Horizontal, espaciamiento {s_hor:.0f} cm"
})

if s_est > 0:
    despiece.append({
        "Marca": "E1",
        "Cantidad": n_est_total,
        "Diámetro (mm)": stirrup_diam,
        "Longitud (m)": perim_est,
        "Longitud Total (m)": n_est_total * perim_est,
        "Peso (kg)": peso_est_total,
        "Observación": f"Estribo rectangular @ {s_est:.0f} cm"
    })

# Barra base (longitudinal)
despiece.append({
    "Marca": "B1",
    "Cantidad": n_barras_base_total,
    "Diámetro (mm)": rebar_diam,
    "Longitud (m)": long_base,
    "Longitud Total (m)": n_barras_base_total * long_base,
    "Peso (kg)": peso_base_total,
    "Observación": f"Base longitudinal, espaciamiento {s_base:.0f} cm"
})

df_despiece = pd.DataFrame(despiece)

# =============================================================================
# INTERFAZ
# =============================================================================
tab1, tab2, tab3, tab4 = st.tabs([" Estabilidad y Estructura", " Sección y DXF", " Visualización 3D", " Cantidades, APU y Despiece"])

with tab1:
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Factores de Seguridad")
        st.metric("FS Volcamiento", f"{FS_volc:.2f}", delta=f"vs {FS_v_min}")
        if FS_volc >= FS_v_min:
            st.success(f" Aprobado: {FS_volc:.2f} ≥ {FS_v_min:.1f}")
        else:
            st.error(f" No aprobado: {FS_volc:.2f} < {FS_v_min:.1f} → Aumentar base")
        st.metric("FS Deslizamiento", f"{FS_desl:.2f}", delta=f"vs {FS_d_min}")
        if FS_desl >= FS_d_min:
            st.success(f" Aprobado: {FS_desl:.2f} ≥ {FS_d_min:.1f}")
        else:
            st.error(f" No aprobado: {FS_desl:.2f} < {FS_d_min:.1f} → Aumentar base o usar dentellón")
        st.metric("Presión Máxima q_max", f"{q_max:.1f} kPa", delta=f"q_adm={q_adm} kPa")
        if ok_bearing:
            st.success(f" Capacidad portante OK: {q_max:.1f} ≤ {q_adm:.1f} kPa")
        else:
            st.error(f" Capacidad insuficiente: {q_max:.1f} > {q_adm:.1f} kPa → Aumentar base")
    with col2:
        st.subheader("Resultados Estructurales")
        st.write(f"**Momento en pantalla (Mu):** {Mu_pantalla:.1f} kN·m/m")
        st.write(f"**Acero vertical requerido:** {As_vert_req:.2f} cm²/m → {rebar_vert} @ {s_vert:.0f} cm (As={As_vert_prov:.2f} cm²/m)")
        st.write(f"**Acero horizontal:** {rebar_vert} @ {s_hor:.0f} cm")
        if s_est > 0:
            st.write(f"**Estribos:** {stirrup_sel} @ {s_est:.0f} cm (mínimo)")
        else:
            st.write(f"**Estribos:** No requeridos por corte")
        st.write(f"**Momento en base (Mu):** {Mu_base:.1f} kN·m/m")
        st.write(f"**Acero base:** {rebar_vert} @ {s_base:.0f} cm (As={As_base_req:.2f} cm²/m)")
        
        # Tabla de fuerzas
        df_fuerzas = pd.DataFrame([
            ("Peso total", f"{W_total:.1f} kN/m", f"{Mr_pesos/W_total:.2f} m", f"{Mr_pesos:.1f} kN·m/m"),
            ("Empuje suelo (Pah)", f"{Pah_suelo:.1f} kN/m", f"{y_pa:.2f} m", f"{Mo_suelo:.1f} kN·m/m"),
            ("Empuje sobrecarga", f"{Pah_q:.1f} kN/m", f"{y_q:.2f} m", f"{Mo_q:.1f} kN·m/m"),
            ("Fricción base", f"{Fr_desl:.1f} kN/m", "", ""),
        ], columns=["Concepto", "Fuerza [kN/m]", "Brazo [m]", "Momento [kN·m/m]"])
        st.dataframe(df_fuerzas, use_container_width=True, hide_index=True)

    # Memoria DOCX
    st.markdown("---")
    col_mem1, col_mem2 = st.columns(2)
    with col_mem1:
        if st.button(_t(" Generar Memoria DOCX", " Generate DOCX Report")):
            # Crear figura de la sección para incrustar
            fig_mem, ax_mem = plt.subplots(figsize=(6, 4))
            ax_mem.set_facecolor('#1a1a2e'); fig_mem.patch.set_facecolor('#1a1a2e')
            ax_mem.add_patch(patches.Rectangle((0,0), B_base, espesor_base, facecolor='#4a4a6a', edgecolor='white'))
            pts_pant = [(pie_muro, espesor_base), (pie_muro+base_pantalla, espesor_base),
                        (pie_muro+base_pantalla, H_muro), (pie_muro+base_pantalla-corona, H_muro)]
            ax_mem.add_patch(patches.Polygon(pts_pant, facecolor='#6a6a8a', edgecolor='white'))
            ax_mem.plot([pie_muro+base_pantalla, B_base], [H_muro, H_prima], color='saddlebrown', lw=2)
            ax_mem.set_xlim(-1, B_base+1); ax_mem.set_ylim(-1, H_prima+1)
            ax_mem.axis('off')
            buf_img = io.BytesIO()
            fig_mem.savefig(buf_img, format='png', dpi=150)
            buf_img.seek(0)
            
            doc = Document()
            doc.add_heading(f"Memoria de Cálculo – Muro de Contención H={H_muro:.1f} m", 0)
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph(f"Norma: {norma_sel} | f'c={fc:.1f} MPa, fy={fy:.0f} MPa")
            doc.add_heading("1. Geometría y Suelos", level=1)
            doc.add_paragraph(f"Altura H={H_muro:.2f} m, Base B={B_base:.2f} m, Pie={pie_muro:.2f} m, Talón={talon:.2f} m")
            doc.add_paragraph(f"Ángulo fricción φ={phi_ang}°, γ_suelo={gamma_s} kN/m³, c_base={c_base} kPa")
            doc.add_heading("2. Resultados de Estabilidad", level=1)
            doc.add_paragraph(f"FS Volcamiento = {FS_volc:.2f} (min {FS_v_min})")
            doc.add_paragraph(f"FS Deslizamiento = {FS_desl:.2f} (min {FS_d_min})")
            doc.add_paragraph(f"q_max = {q_max:.1f} kPa ≤ q_adm={q_adm} kPa → {'OK' if ok_bearing else 'NO'}")
            doc.add_heading("3. Diseño Estructural", level=1)
            doc.add_paragraph(f"Momento en pantalla Mu={Mu_pantalla:.1f} kN·m/m → As={As_vert_prov:.2f} cm²/m ({rebar_vert} @ {s_vert:.0f} cm)")
            doc.add_paragraph(f"Momento en base Mu={Mu_base:.1f} kN·m/m → As={As_base_req:.2f} cm²/m ({rebar_vert} @ {s_base:.0f} cm)")
            doc.add_heading("4. Cantidades Totales (L={:.1f} m)".format(L_total), level=1)
            doc.add_paragraph(f"Concreto: {vol_conc_total:.3f} m³")
            doc.add_paragraph(f"Cemento: {total_cement_kg:.1f} kg ({bags_cement:.1f} bultos de {bag_kg:.0f} kg)")
            doc.add_paragraph(f"Arena: {total_sand_kg:.1f} kg")
            doc.add_paragraph(f"Grava: {total_gravel_kg:.1f} kg")
            doc.add_paragraph(f"Acero total: {peso_total_acero:.1f} kg")
            doc.add_heading("5. Despiece de Acero", level=1)
            # Tabla en DOCX
            table = doc.add_table(rows=1+len(despiece), cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Marca"; hdr[1].text = "Cantidad"; hdr[2].text = "Ø (mm)"; hdr[3].text = "L (m)"; hdr[4].text = "L total (m)"; hdr[5].text = "Peso (kg)"
            for i, row in enumerate(despiece):
                cells = table.rows[i+1].cells
                cells[0].text = row["Marca"]
                cells[1].text = f"{int(row['Cantidad'])}"
                cells[2].text = f"{row['Diámetro (mm)']}"
                cells[3].text = f"{row['Longitud (m)']:.2f}"
                cells[4].text = f"{row['Longitud Total (m)']:.2f}"
                cells[5].text = f"{row['Peso (kg)']:.1f}"
            # Insertar figura
            doc.add_picture(buf_img, width=Inches(5))
            buf_doc = io.BytesIO()
            doc.save(buf_doc)
            buf_doc.seek(0)
            st.download_button("Descargar Memoria DOCX", data=buf_doc, file_name=f"Muro_{H_muro:.1f}m.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab2:
    st.subheader(_t(" Sección Transversal y DXF", " Cross Section and DXF"))
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_facecolor('#1a1a2e'); fig.patch.set_facecolor('#1a1a2e')
    # Base
    ax.add_patch(patches.Rectangle((0,0), B_base, espesor_base, facecolor='#4a4a6a', edgecolor='white', lw=2))
    # Pantalla
    pts = [(pie_muro, espesor_base), (pie_muro+base_pantalla, espesor_base),
           (pie_muro+base_pantalla, H_muro), (pie_muro+base_pantalla-corona, H_muro)]
    ax.add_patch(patches.Polygon(pts, facecolor='#6a6a8a', edgecolor='white', lw=2))
    # Suelo
    ax.plot([pie_muro+base_pantalla, B_base], [H_muro, H_prima], color='saddlebrown', lw=2)
    ax.plot([B_base, B_base+2], [H_prima, H_prima + 2*math.tan(beta_rad)], color='saddlebrown', lw=2)
    ax.set_xlim(-1, B_base+2)
    ax.set_ylim(-0.5, H_prima+1)
    ax.axis('off')
    st.pyplot(fig)
    
    st.markdown("####  Exportar DXF (Planta + Elevación)")
    try:
        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                 dxf_dim_horiz, dxf_dim_vert, dxf_rotulo,
                                 dxf_leyenda, dxf_rotulo_campos)
        _USE_H = True
    except ImportError:
        _USE_H = False
    doc_dxf = ezdxf.new('R2010')
    doc_dxf.units = ezdxf.units.M
    if _USE_H:
        dxf_setup(doc_dxf, 50)      # escala 1:50 para muros
        dxf_add_layers(doc_dxf)
    msp = doc_dxf.modelspace()
    for lay, col in [('CONCRETO',7), ('ACERO',1), ('TEXTO',3), ('SUELO',2)]:
        if lay not in doc_dxf.layers:
            doc_dxf.layers.add(lay, color=col)
    # Planta (vista superior de la base)
    msp.add_lwpolyline([(0,0), (B_base,0), (B_base,1), (0,1), (0,0)], close=True, dxfattribs={'layer':'CONCRETO'})
    msp.add_line((pie_muro,0), (pie_muro,1), dxfattribs={'layer':'ACERO'})
    msp.add_line((pie_muro+base_pantalla,0), (pie_muro+base_pantalla,1), dxfattribs={'layer':'ACERO'})
    for i in range(int(B_base*10)+1):
        x = i/10
        msp.add_line((x,0.4), (x,0.6), dxfattribs={'layer':'ACERO'})
    for j in range(0,11):
        y = j/10
        msp.add_line((0.2,y), (B_base-0.2,y), dxfattribs={'layer':'ACERO'})
    # Elevación
    off_x = B_base + 2
    msp.add_lwpolyline([(off_x,0), (off_x+B_base,0), (off_x+B_base,espesor_base), (off_x,espesor_base), (off_x,0)], close=True, dxfattribs={'layer':'CONCRETO'})
    pts_elev = [(off_x+pie_muro, espesor_base), (off_x+pie_muro+base_pantalla, espesor_base),
                (off_x+pie_muro+base_pantalla, H_muro), (off_x+pie_muro+base_pantalla-corona, H_muro)]
    msp.add_lwpolyline(pts_elev, close=True, dxfattribs={'layer':'CONCRETO'})
    # Acero vertical
    for i in range(int(100/s_vert)+1):
        x = off_x + pie_muro + i * s_vert/100
        if x <= off_x+pie_muro+base_pantalla:
            msp.add_line((x, espesor_base), (x, H_muro), dxfattribs={'layer':'ACERO'})
    # Estribos horizontales
    if s_est > 0:
        n_est_elev = int(H_muro / (s_est/100)) + 1
        for i in range(n_est_elev):
            z = espesor_base + i * s_est/100
            if z <= H_muro:
                msp.add_line((off_x+pie_muro, z), (off_x+pie_muro+base_pantalla, z), dxfattribs={'layer':'ACERO'})
    # Suelo inclinado
    msp.add_line((off_x+pie_muro+base_pantalla, H_muro), (off_x+B_base, H_prima), dxfattribs={'layer':'SUELO'})
    msp.add_text(f"Empalme Clase B  Lap={splice_length_mm/10:.0f} cm",
                 dxfattribs={'layer':'TEXTO','height':0.18,'insert':(off_x+pie_muro+base_pantalla+0.15, H_muro/3)})
    if _USE_H:
        TH = 0.025 * 50
        # Cotas elevación
        dxf_dim_horiz(msp, off_x, off_x+B_base, -0.6, f"B = {B_base:.2f} m", 50)
        dxf_dim_vert(msp, off_x-0.6, 0, H_muro, f"H = {H_muro:.2f} m", 50)
        dxf_dim_vert(msp, off_x-1.0, 0, espesor_base, f"eb = {espesor_base:.2f} m", 50)
        # Títulos
        dxf_text(msp, B_base/2, 1.3, "PLANTA", "EJES", h=TH*1.2, ha="center")
        dxf_text(msp, off_x+B_base/2, H_muro+0.4, "ELEVACION", "EJES", h=TH*1.2, ha="center")
        dxf_text(msp, off_x+B_base/2, H_muro+0.7,
                 f"Muro H={H_muro:.1f}m  B={B_base:.1f}m  |  fc={fc:.0f}MPa  fy={fy:.0f}MPa",
                 "TEXTO", h=TH, ha="center")
        # Leyenda
        dxf_leyenda(msp, off_x+B_base+0.3, H_muro-0.5, [
            ("CONCRETO", "Concreto"),
            ("ACERO",    f"Acero {rebar_vert} @ {s_vert:.0f}cm"),
            ("SUELO",    "Superficie suelo"),
        ], 50)
        # Rótulo
        _cam = dxf_rotulo_campos(f"Muro Contencion H={H_muro:.1f}m  B={B_base:.1f}m", norma_sel, "001")
        dxf_rotulo(msp, _cam, off_x, -4.5, rot_w=9, rot_h=3, escala=50)
    else:
        msp.add_text(f"Empalme Clase B\nLap={splice_length_mm/10:.0f} cm",
                     dxfattribs={'layer':'TEXTO','height':0.18,'insert':(off_x+pie_muro+base_pantalla+0.2, H_muro/3)})
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix='.dxf', delete=False) as tmp__out:
        tmp_path__out = tmp__out.name
    doc_dxf.saveas(tmp_path__out)
    with open(tmp_path__out, 'rb') as f__out:
        bytes__out = f__out.read()
    os.unlink(tmp_path__out)    st.download_button("Descargar DXF", data=bytes__out, file_name=f"Muro_{H_muro:.1f}m.dxf")

with tab3:
    st.subheader(_t(" Visualización 3D", " 3D Visualization"))
    fig3d = go.Figure()
    # Base
    x_base = [0, B_base, B_base, 0, 0, B_base, B_base, 0]
    y_base = [0, 0, 1, 1, 0, 0, 1, 1]
    z_base = [0, 0, 0, 0, espesor_base, espesor_base, espesor_base, espesor_base]
    fig3d.add_trace(go.Mesh3d(x=x_base, y=y_base, z=z_base, alphahull=0, opacity=0.4, color='gray', name='Base'))
    # Pantalla
    x_pant = [pie_muro, pie_muro+base_pantalla, pie_muro+base_pantalla, pie_muro, pie_muro, pie_muro+base_pantalla, pie_muro+base_pantalla, pie_muro]
    y_pant = [0, 0, 1, 1, 0, 0, 1, 1]
    z_pant = [espesor_base, espesor_base, H_muro, H_muro, espesor_base, espesor_base, H_muro, H_muro]
    fig3d.add_trace(go.Mesh3d(x=x_pant, y=y_pant, z=z_pant, alphahull=0, opacity=0.6, color='darkgray', name='Pantalla'))
    # Suelo (simplificado)
    x_soil = [pie_muro+base_pantalla, B_base, B_base+1, pie_muro+base_pantalla+1]
    y_soil = [0, 0, 1, 1]
    z_soil = [H_muro, H_prima, H_prima+0.5, H_muro+0.5]
    fig3d.add_trace(go.Mesh3d(x=x_soil, y=y_soil, z=z_soil, alphahull=0, opacity=0.3, color='saddlebrown', name='Suelo'))
    fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='X (m)', yaxis_title='Y (m)', zaxis_title='Z (m)'),
                        margin=dict(l=0, r=0, b=0, t=0), height=450, paper_bgcolor='#1a1a2e')
    st.plotly_chart(fig3d, use_container_width=True)

with tab4:
    st.subheader(_t(" Cantidades Totales y Despiece", " Total Quantities and Bending Schedule"))
    col1, col2, col3 = st.columns(3)
    col1.metric(_t("Volumen Concreto", "Concrete Volume"), f"{vol_conc_total:.3f} m³")
    col2.metric(_t("Acero Total", "Total Steel"), f"{peso_total_acero:.1f} kg")
    col3.metric(_t("Cuantía Acero", "Steel Ratio"), f"{peso_total_acero/vol_conc_total:.1f} kg/m³")
    
    st.markdown("---")
    st.markdown(_t("####  Dosificación de Concreto (f'c = {:.1f} MPa)", "####  Concrete Mix Design (f'c = {:.1f} MPa)").format(fc))
    df_mix = pd.DataFrame([
        (_t("Cemento", "Cement"), f"{total_cement_kg:.1f} kg", f"{bags_cement:.1f} bultos"),
        (_t("Agua", "Water"), f"{total_water_L:.0f} L", ""),
        (_t("Arena", "Sand"), f"{total_sand_kg:.1f} kg", f"{total_sand_kg/1600:.2f} m³"),
        (_t("Grava", "Gravel"), f"{total_gravel_kg:.1f} kg", f"{total_gravel_kg/1600:.2f} m³"),
    ], columns=[_t("Material","Material"), _t("Total","Total"), _t("Observación","Note")])
    st.dataframe(df_mix, use_container_width=True, hide_index=True)
    
    st.markdown("---")
    st.markdown(_t("####  Despiece de Acero (Bending Schedule)", "####  Steel Bending Schedule"))
    st.dataframe(df_despiece.style.format({"Longitud (m)": "{:.2f}", "Longitud Total (m)": "{:.2f}", "Peso (kg)": "{:.1f}"}), use_container_width=True, hide_index=False)
    
    # Gráfico de barras de pesos
    fig_bars, ax_bars = plt.subplots(figsize=(8, 4))
    ax_bars.bar(df_despiece["Marca"], df_despiece["Peso (kg)"], color='#ff6b35')
    ax_bars.set_xlabel(_t("Marca", "Mark"))
    ax_bars.set_ylabel(_t("Peso (kg)", "Weight (kg)"))
    ax_bars.set_title(_t("Peso por tipo de barra", "Weight per bar type"))
    ax_bars.grid(True, alpha=0.3)
    st.pyplot(fig_bars)
    
    # Exportar despiece a Excel
    output_excel_desp = io.BytesIO()
    with pd.ExcelWriter(output_excel_desp, engine='xlsxwriter') as writer:
        df_despiece.to_excel(writer, sheet_name="Despiece", index=False)
        df_mix.to_excel(writer, sheet_name="Dosificacion", index=False)
        workbook = writer.book
        worksheet = writer.sheets["Despiece"]
        money_fmt = workbook.add_format({'num_format': '#,##0.00'})
        worksheet.set_column('A:E', 15, money_fmt)
    output_excel_desp.seek(0)
    st.download_button(_t(" Descargar Despiece y Dosificación (Excel)", " Download Bending Schedule and Mix Design (Excel)"), 
                       data=output_excel_desp, file_name=f"Despiece_Muro_{H_muro:.1f}m.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    # APU
    if "apu_config" in st.session_state:
        apu = st.session_state.apu_config
        mon = apu.get("moneda", "$")
        if apu.get("usar_concreto_premezclado", False):
            precio_concreto_m3 = apu.get("precio_concreto_m3", 0)
            c_conc = vol_conc_total * precio_concreto_m3
            c_cem = 0; c_are = 0; c_gra = 0
        else:
            pct_arena = apu.get("pct_arena_mezcla", 0.55)
            pct_grava = apu.get("pct_grava_mezcla", 0.80)
            vol_arena = vol_conc_total * pct_arena
            vol_grava = vol_conc_total * pct_grava
            c_cem = bags_cement * apu.get("cemento", 0)
            c_are = vol_arena * apu.get("arena", 0)
            c_gra = vol_grava * apu.get("grava", 0)
            c_conc = 0
        c_ace = peso_total_acero * apu.get("acero", 0)
        total_mat = c_cem + c_ace + c_are + c_gra + c_conc
        total_dias_mo = (peso_total_acero * 0.04) + (vol_conc_total * 0.4)
        costo_mo = total_dias_mo * apu.get("costo_dia_mo", 69333.33)
        costo_directo = total_mat + costo_mo
        herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
        aiu = costo_directo * apu.get("pct_aui", 0.30)
        utilidad = costo_directo * apu.get("pct_util", 0.05)
        iva = utilidad * apu.get("iva", 0.19)
        total_proyecto = costo_directo + herramienta + aiu + iva
        
        st.markdown("---")
        st.markdown(_t("###  Presupuesto Total", "###  Total Budget"))
        data_apu = {
            "Item": [_t("Cemento (bultos)", "Cement (bags)"), _t("Acero (kg)", "Steel (kg)"), 
                     _t("Arena (m³)", "Sand (m³)"), _t("Grava (m³)", "Gravel (m³)"),
                     _t("Mano de Obra (días)", "Labor (days)"), _t("Herramienta Menor", "Minor Tools"), 
                     _t("A.I.U.", "A.I.U."), _t("IVA s/Utilidad", "VAT on Profit"), _t("TOTAL", "TOTAL")],
            "Cantidad": [f"{bags_cement:.1f}", f"{peso_total_acero:.1f}", 
                         f"{vol_arena:.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                         f"{vol_grava:.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                         f"{total_dias_mo:.2f}", f"{apu.get('pct_herramienta', 0.05)*100:.1f}% MO", 
                         f"{apu.get('pct_aui', 0.3)*100:.1f}% CD", f"{apu.get('iva', 0.19)*100:.1f}% Util", ""],
            f"Subtotal [{mon}]": [f"{c_cem:,.2f}", f"{c_ace:,.2f}", f"{c_are:,.2f}", f"{c_gra:,.2f}", 
                                  f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}", f"{iva:,.2f}", f"**{total_proyecto:,.2f}**"]
        }
        st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)
        st.metric(f" Gran Total Proyecto [{mon}]", f"{total_proyecto:,.0f}")
        
        # Excel APU total
        output_excel_apu = io.BytesIO()
        with pd.ExcelWriter(output_excel_apu, engine='xlsxwriter') as writer:
            df_export_apu = pd.DataFrame({
                "Item": ["Cemento", "Acero", "Arena", "Grava", "Mano de Obra"],
                "Cantidad": [bags_cement, peso_total_acero, vol_arena, vol_grava, total_dias_mo],
                "Unidad": [apu.get('cemento',0), apu.get('acero',0), apu.get('arena',0), apu.get('grava',0), apu.get('costo_dia_mo', 69333.33)]
            })
            if apu.get("usar_concreto_premezclado", False):
                df_pre = pd.DataFrame({"Item": ["Concreto Premezclado"], "Cantidad": [vol_conc_total], "Unidad": [apu.get("precio_concreto_m3", 0)]})
                df_export_apu = pd.concat([df_export_apu, df_pre], ignore_index=True)
            df_export_apu["Subtotal"] = df_export_apu["Cantidad"] * df_export_apu["Unidad"]
            df_export_apu.to_excel(writer, index=False, sheet_name='APU')
            workbook = writer.book
            worksheet = writer.sheets['APU']
            money_fmt = workbook.add_format({'num_format': '#,##0.00'})
            bold = workbook.add_format({'bold': True})
            worksheet.set_column('A:A', 25)
            worksheet.set_column('B:D', 15, money_fmt)
            row = len(df_export_apu) + 1
            worksheet.write(row, 0, "Costo Directo (CD)", bold)
            worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
            row += 1
            worksheet.write(row, 0, "Herramienta Menor", bold)
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_herramienta", 0.05)}', money_fmt)
            row += 1
            worksheet.write(row, 0, "A.I.U", bold)
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui", 0.30)}', money_fmt)
            row += 1
            worksheet.write(row, 0, "IVA s/ Utilidad", bold)
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_util", 0.05)}*{apu.get("iva", 0.19)}', money_fmt)
            row += 1
            worksheet.write(row, 0, "TOTAL PRESUPUESTO", bold)
            worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
        output_excel_apu.seek(0)
        st.download_button(_t(" Descargar Presupuesto Excel", " Download Budget Excel"), data=output_excel_apu, 
                           file_name=f"APU_Muro_{H_muro:.1f}m.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info(_t(" Ve a la página 'APU Mercado' para cargar los costos en vivo.", " Go to the 'Market APU' page to load live costs."))