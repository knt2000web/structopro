import streamlit as st
import pandas as pd
import numpy as np
import math
import io
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
try:
    from ifc_export import ifc_grupo_pilotes
except ImportError:
    pass

# Plotly
try:
    import plotly.graph_objects as go
    _PLOTLY_AVAILABLE = True
except ImportError:
    _PLOTLY_AVAILABLE = False

# Librerías opcionales
try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import ezdxf
    _DXF_AVAILABLE = True
except ImportError:
    _DXF_AVAILABLE = False

try:
    import ifcopenshell
    _IFC_AVAILABLE = True
except ImportError:
    _IFC_AVAILABLE = False

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Pilotes y Cimentación Profunda", "Piles and Deep Foundations"), layout="wide")

# ─────────────────────────────────────────────
# PERSISTENCIA SUPABASE
# ─────────────────────────────────────────────
import requests
import json
import streamlit as st

def get_supabase_rest_info():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return url, key
    except Exception as e:
        return None, None

def guardar_proyecto_supabase(nombre, estado_dict, ref):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = f"db_proyectos_{ref.lower()}.json"
            db = {}
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            db[f"[{ref.upper()}] {nombre}"] = {"nombre_proyecto": f"[{ref.upper()}] {nombre}", "estado_json": json.dumps(estado_dict)}
            with open(db_path, "w", encoding="utf-8") as f: json.dump(db, f)
            return True, " Proyecto guardado (Local)"
        except Exception as e:
            return False, f" Error local: {e}"
    
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "Prefer": "resolution=merge-duplicates"
    }
    payload = {
        "nombre_proyecto": f"[{ref.upper()}] {nombre}",
        "user_id": st.session_state.get("user_id", "anonimo"),
        "estado_json": json.dumps(estado_dict),
    }
    try:
        endpoint = f"{url}/rest/v1/proyectos?on_conflict=nombre_proyecto"
        res = requests.post(endpoint, headers=headers, json=payload)
        if res.status_code in [200, 201, 204]: return True, " Proyecto en la nube"
        return False, f" Error API: {res.text}"
    except Exception as e: return False, f" Error API: {e}"

def cargar_proyecto_supabase(nombre, ref):
    url, key = get_supabase_rest_info()
    full_name = f"[{ref.upper()}] {nombre}" if not nombre.startswith("[") else nombre
    if not url or not key:
        try:
            import os
            db_path = f"db_proyectos_{ref.lower()}.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                match = db.get(full_name) or db.get(nombre)
                if match:
                    estado = json.loads(match["estado_json"])
                    for k, v in estado.items(): st.session_state[k] = v
                    return True, " Proyecto cargado (Local)"
            return False, " No encontrado"
        except Exception as e: return False, str(e)
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        endpoint = f"{url}/rest/v1/proyectos?nombre_proyecto=eq.{full_name}&select=*"
        res = requests.get(endpoint, headers=headers)
        if res.status_code == 200:
            data = res.json()
            if data:
                estado = json.loads(data[0]["estado_json"])
                for k, v in estado.items(): st.session_state[k] = v
                return True, " Proyecto cargado"
        return False, " No encontrado"
    except Exception as e: return False, str(e)

def listar_proyectos_supabase(ref):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = f"db_proyectos_{ref.lower()}.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                return sorted([k.replace(f"[{ref.upper()}] ", "") for k in db.keys()])
            return []
        except: return []
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        endpoint = f"{url}/rest/v1/proyectos?select=nombre_proyecto"
        res = requests.get(endpoint, headers=headers)
        if res.status_code == 200:
            return sorted([i["nombre_proyecto"].replace(f"[{ref.upper()}] ", "") for i in res.json() if f"[{ref.upper()}]" in i.get("nombre_proyecto","")])
        return []
    except: return []

def capturar_estado_module(ref):
    if ref == "zapatas":
        match = ["z_", "cp_", "apu_", "z_tipo_idx", "fc_basico", "fy_basico"]
    else:
        match = ["pi_", "p_", "apu_", "fc_pi"]
    return {k: st.session_state[k] for k in st.session_state if any(k.startswith(m) for m in match)}

# ─────────────────────────────────────────────


st.markdown('''
<div style="display:flex; align-items:center; gap:20px; margin-bottom:10px; background:#1e2530; padding:15px; border-radius:10px; border-left:5px solid #29b6f6;">
    <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 60" width="80" height="48">
        <line x1="10" y1="20" x2="90" y2="20" stroke="#29b6f6" stroke-width="4"/>
        <rect x="35" y="5" width="30" height="15" fill="#29b6f6" opacity="0.8"/>
        <line x1="40" y1="20" x2="40" y2="55" stroke="#29b6f6" stroke-width="6" stroke-linecap="round"/>
        <line x1="60" y1="20" x2="60" y2="55" stroke="#29b6f6" stroke-width="6" stroke-linecap="round"/>
        <line x1="10" y1="35" x2="90" y2="35" stroke="#fff" stroke-width="1" stroke-dasharray="2,2" opacity="0.3"/>
        <line x1="10" y1="50" x2="90" y2="50" stroke="#fff" stroke-width="1" stroke-dasharray="2,2" opacity="0.3"/>
    </svg>
    <div>
        <h1 style="margin:0; padding:0; font-size:28px;">Cimentaciones Profundas: Pilotes</h1>
        <p style="margin:0; padding:0; color:#aaa; font-size:14px;">Módulo integral NSR-10 / ACI 318 (Diseño por capas, eficiencia de grupo, APU e IFC)</p>
    </div>
</div>
''', unsafe_allow_html=True)

with st.expander(_t(" ¿Qué hace este módulo? — Guía rápida", "Quick Guide"), expanded=False):
    st.markdown('''
    ###  Pilotes — Verificaciones Geotécnicas y Estructurales
    | Módulo | Qué calcula | Métodos |
    |--------|------------|-------|
    | **Geotécnico** | Capacidad por punta y fuste por cada estrato | Meyerhof, Vesic, α, β |
    | **Grupo** | Distribución de cargas y eficiencia de conjunto | Converse-Labarre |
    | **Estructural** | Refuerzo de pilote (circular/cuadrado), cortante | ACI 318-14 §18.7 |
    | **Encepado** | Dimensionamiento del dado de pilotes | ACI 318 / NSR-10 |
    ''')

st.sidebar.header(_t("⚙ Configuración Global", "⚙ Global Settings"))
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
st.sidebar.success(f"{_t('Norma Activa:', 'Active Code:')} {norma_sel}")

st.sidebar.markdown("---")
st.sidebar.subheader(" Guardar / Cargar Proyecto")
nombre_producido = st.session_state.get(f"np_{"pilotes"}", "")

st.sidebar.markdown("**Nuevo Proyecto / Guardar**")
nombre_proy_guardar = st.sidebar.text_input("Nombre para guardar", value=nombre_producido, key=f"input_guardar_{"pilotes"}")

if st.sidebar.button(" Guardar Proyecto", use_container_width=True, key=f"btn_save_{"pilotes"}"):
    if nombre_proy_guardar:
        ok, msg = guardar_proyecto_supabase(nombre_proy_guardar, capturar_estado_module("pilotes"), "pilotes")
        if ok:
            st.session_state[f"np_{"pilotes"}"] = nombre_proy_guardar
            st.sidebar.success(msg)
        else:
            st.sidebar.error(msg)
    else:
        st.sidebar.warning("Escribe un nombre de proyecto")

st.sidebar.markdown("**Cargar Proyecto Existente**")
lista_proyectos = listar_proyectos_supabase("pilotes")
if lista_proyectos:
    idx_def = lista_proyectos.index(nombre_producido) if nombre_producido in lista_proyectos else 0
    nombre_proy_cargar = st.sidebar.selectbox("Selecciona un proyecto", lista_proyectos, index=idx_def, key=f"sel_load_{"pilotes"}")
    
    def on_cargar_click():
        proy = st.session_state[f"sel_load_{"pilotes"}"]
        if proy:
            ok, msg = cargar_proyecto_supabase(proy, "pilotes")
            st.session_state[f"__msg_cargar_{"pilotes"}"] = (ok, msg)
            if ok: st.session_state[f"np_{"pilotes"}"] = proy

    st.sidebar.button(" Cargar", on_click=on_cargar_click, use_container_width=True, key=f"btn_load_{"pilotes"}")

    if f"__msg_cargar_{"pilotes"}" in st.session_state:
        ok, msg = st.session_state.pop(f"__msg_cargar_{"pilotes"}")
        if ok: st.sidebar.success(msg)
        else: st.sidebar.error(msg)


# Tabla de estratigrafía base (perfil por capas)
if "perfil_suelo" not in st.session_state:
    st.session_state.perfil_suelo = pd.DataFrame([
        {"Estrato": 1, "Espesor (m)": 2.0, "Tipo": "Arcilla", "γ (kN/m³)": 17.5, "c (kPa)": 40.0, "φ (°)": 0.0, "N60": 4},
        {"Estrato": 2, "Espesor (m)": 4.0, "Tipo": "Arena",   "γ (kN/m³)": 18.5, "c (kPa)": 0.0,  "φ (°)": 32.0, "N60": 15},
        {"Estrato": 3, "Espesor (m)": 8.0, "Tipo": "Arcilla", "γ (kN/m³)": 18.0, "c (kPa)": 90.0, "φ (°)": 0.0, "N60": 12},
    ])

tab_geo, tab_grup, tab_est, tab_apu, tab_mem = st.tabs([
    _t("1. Estratigrafía y Pilote", "1. Stratigraphy"),
    _t("2. Grupo y Encepado", "2. Pile Group & Cap"),
    _t("3. Diseño Estructural", "3. Structural Design"),
    _t("4. APU y Cantidades", "4. APU"),
    _t("5. Entregables", "5. Deliverables")
])

with tab_geo:
    st.subheader(_t("1.1 Definición del Perfil de Suelo (Por Capas)", "1.1 Soil Profile (Layered)"))
    st.info("Ingresa la estratigrafía del terreno. La capacidad por fricción sumará cada estrato atravesado.")
    st.session_state.perfil_suelo = st.data_editor(
        st.session_state.perfil_suelo,
        num_rows="dynamic",
        column_config={
            "Tipo": st.column_config.SelectboxColumn("Tipo de Suelo", options=["Arcilla", "Arena", "Grava", "Roca"], required=True),
            "Espesor (m)": st.column_config.NumberColumn(min_value=0.1, max_value=50.0, format="%.1f", required=True),
        },
        use_container_width=True
    )
    
    st.subheader(_t("1.2 Geometría del Pilote Único", "1.2 Single Pile Geometry"))
    col1, col2, col3 = st.columns(3)
    with col1:
        tipo_seccion = st.selectbox("Sección", ["Circular", "Cuadrada"])
        D_pilote = st.number_input("Dimensión D [m]", 0.3, 2.5, 0.60, 0.05)
    with col2:
        L_pilote = st.number_input("Longitud L_p [m]", 3.0, 50.0, 12.0, 0.5)
        NF_prof = st.number_input("Freático [m]", 0.0, 100.0, 2.0, 0.5)
    with col3:
        FS_global = st.number_input("Factor de Seguridad", 2.0, 4.0, 3.0, 0.1)

    prof_total = st.session_state.perfil_suelo['Espesor (m)'].sum()
    if L_pilote > prof_total:
        st.error(f"El pilote ({L_pilote}m) atraviesa más suelo del definido ({prof_total:.1f}m). Agrega estratos.")
    else:
        st.success(f"Pilote apoyado dentro de la estratigrafía (Prof. Mínima Estudiada {prof_total:.1f}m)")

        # ──── CÁLCULOS GEOTÉCNICOS ────────────────────────────────────────
        import numpy as np
        import math

        df = st.session_state.perfil_suelo.copy()
        gamma_w = 9.81

        # 1. Acumular profundidades por capa
        df["z_top"] = [0.0] + list(df["Espesor (m)"].cumsum()[:-1])
        df["z_bot"] = df["z_top"] + df["Espesor (m)"]

        # 2. Esfuerzo vertical efectivo en la base de cada capa
        sigma_v = []
        sv = 0.0
        for _, row in df.iterrows():
            gamma_ef = row["γ (kN/m³)"] - gamma_w if row["z_bot"] > NF_prof else row["γ (kN/m³)"]
            sv += gamma_ef * row["Espesor (m)"]
            sigma_v.append(sv)
        df["σv' base (kPa)"] = sigma_v

        # 3. Función encapsulada de cálculo geotécnico y Optimizador
        def calcular_capacidad(Lp_eval, df_eval, D_pil, tipo_sec, NF, FS_val):
            Aw_fuste = math.pi * D_pil if tipo_sec == "Circular" else 4 * D_pil
            Aw_punta = math.pi * (D_pil/2)**2 if tipo_sec == "Circular" else D_pil**2
            Qs_tot = 0.0
            filas_fs = []
            
            for _, row in df_eval.iterrows():
                if row["z_top"] >= Lp_eval: break
                h_activo = min(row["z_bot"], Lp_eval) - row["z_top"]
                if h_activo <= 0: continue
                
                if row["Tipo"] == "Arcilla":
                    alpha = 0.55 if row["c (kPa)"] < 50 else (0.55 - 0.1*(row["c (kPa)"]-50)/50)
                    alpha = max(0.20, min(alpha, 0.55))
                    qs = alpha * row["c (kPa)"]
                elif row["Tipo"] in ["Arena", "Grava"]:
                    K0 = 1 - math.sin(math.radians(row["φ (°)"]))
                    beta = K0 * math.tan(math.radians(row["φ (°)"]))
                    beta = max(0.15, min(beta, 0.40))
                    # EC-2 Fix: Usar gamma efectivo por capa bajo nivel freático
                    g_ef = row["γ (kN/m³)"] - 9.81 if row["z_bot"] > NF else row["γ (kN/m³)"]
                    sv_m = max(0.0, row["σv' base (kPa)"] - g_ef * h_activo / 2.0)
                    qs = beta * sv_m
                else:  # Roca
                    qs = 0.1 * row["c (kPa)"]
                
                Qs_c = qs * Aw_fuste * h_activo
                Qs_tot += Qs_c
                filas_fs.append({"Estrato": int(row["Estrato"]), "Tipo": row["Tipo"], "H activo (m)": round(h_activo, 2), "qs (kPa)": round(qs, 2), "Qs capa (kN)": round(Qs_c, 1)})
                
            _df_p = df_eval[df_eval["z_bot"] >= Lp_eval - 1e-6]
            f_punta = _df_p.iloc[0] if not _df_p.empty else df_eval.iloc[-1]
            phi_p, c_p, sv_p = f_punta["φ (°)"], f_punta["c (kPa)"], f_punta["σv' base (kPa)"]
            
            if f_punta["Tipo"] == "Roca":
                Qp_tot = 3.0 * c_p * Aw_punta * 1000
            elif phi_p > 0:
                Nq_m = (math.exp(math.pi * math.tan(math.radians(phi_p))) * math.tan(math.radians(45 + phi_p/2))**2)
                Nq_m = min(Nq_m, 60)
                Qp_tot = min(sv_p * Nq_m, 11000) * Aw_punta
            else:
                Qp_tot = 9.0 * c_p * Aw_punta
                
            Qu_tot = Qp_tot + Qs_tot
            Qa_tot = Qu_tot / FS_val
            return Qp_tot, Qs_tot, Qu_tot, Qa_tot, filas_fs, f_punta

        # --- Análisis Pilote Único Actual ---
        Qp, Qs_total, Qu, Qadm, filas_fuste, fila_punta = calcular_capacidad(L_pilote, df, D_pilote, tipo_seccion, NF_prof, FS_global)

        # Advertencia de estrato blando
        if fila_punta["Tipo"] == "Arcilla" and fila_punta.get("N60", 100) < 5:
            st.warning(f"⚠️ **Seguridad:** El pilote (L={L_pilote}m) termina apoyado sobre el Estrato {int(fila_punta['Estrato'])} (Arcilla Blanda, N60={fila_punta.get('N60',0)}). Considere profundizar la cimentación hasta un estrato competente o roca.")

        # Mostrar resultados del pilote único
        st.divider()
        st.subheader("1.3 Resultados de Capacidad Portante del Pilote")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Q_punta (kN)", f"{Qp:.1f}")
        m2.metric("Q_fuste total (kN)", f"{Qs_total:.1f}")
        m3.metric("Q_última (kN)", f"{Qu:.1f}")
        m4.metric(f"Q_adm (kN) — FS={FS_global}", f"{Qadm:.1f}")

        with st.expander("Tabla de Fricción Lateral por Capa", expanded=False):
            st.dataframe(pd.DataFrame(filas_fuste), use_container_width=True, hide_index=True)

        # 6. Optimizador de Profundidad y Perfil Visual
        st.divider()
        st.subheader("1.4 Optimizador y Perfil Estratigráfico Geotécnico")
        
        col_opt1, col_opt2 = st.columns([1.5, 2.5])
        with col_opt1:
            Pu_req = st.number_input("Carga Axial Requerida (por pilote) [kN]", 0.0, 50000.0, 1000.0, 50.0)
            
            # Gráfico de Perfil
            if _PLOTLY_AVAILABLE:
                fig_pf = go.Figure()
                color_map = {"Arcilla": "#8B4513", "Arena": "#DEB887", "Grava": "#8e9eab", "Roca": "#546e7a"}
                for _, r_env in df.iterrows():
                    fig_pf.add_shape(type="rect", 
                        x0=0.1, y0=r_env["z_top"], x1=1.9, y1=r_env["z_bot"],
                        line=dict(color="black", width=1), fillcolor=color_map.get(r_env["Tipo"], "#999"), opacity=0.7)
                    if (r_env["z_bot"] - r_env["z_top"]) >= 0.5:
                        props = f"γ={r_env['γ (kN/m³)']}"
                        if r_env['c (kPa)'] > 0: props += f" c={r_env['c (kPa)']}"
                        if r_env['φ (°)'] > 0: props += f" φ={r_env['φ (°)']}°"
                        texto_capa = f"<b>{r_env['Tipo']}</b><br><span style='font-size:10px'>{props}</span>"
                        fig_pf.add_annotation(x=1.5, y=(r_env["z_top"]+r_env["z_bot"])/2, text=texto_capa, showarrow=False, font=dict(color="white"), align="center")
                
                # Dibujar pilote actual sobrepuesto
                fig_pf.add_shape(type="rect", x0=0.8, y0=0, x1=1.2, y1=L_pilote, line=dict(color="white", width=2), fillcolor="#b0bec5")
                fig_pf.add_annotation(x=1.0, y=L_pilote, text="▶ Base", showarrow=True, arrowhead=1, ax=40, ay=0, font=dict(color="white"))
                
                fig_pf.update_layout(
                    title="Perfil del Suelo vs Pilote", 
                    xaxis=dict(visible=False, range=[0, 2]), 
                    yaxis=dict(title="Profundidad [m]", range=[prof_total, 0]), 
                    height=450, margin=dict(l=20, r=20, t=40, b=20), 
                    paper_bgcolor="#1e2530", plot_bgcolor="#1e2530", font=dict(color="white")
                )
                st.plotly_chart(fig_pf, use_container_width=True)

        with col_opt2:
            if _PLOTLY_AVAILABLE and Pu_req > 0:
                Lp_vals = np.arange(3.0, prof_total + 0.1, 0.5)
                Qa_vals = []
                for lv in Lp_vals:
                    out_qp, out_qs, out_qu, out_qa, out_fil, out_fp = calcular_capacidad(lv, df, D_pilote, tipo_seccion, NF_prof, FS_global)
                    Qa_vals.append(out_qa)
                
                fig_opt = go.Figure()
                fig_opt.add_trace(go.Scatter(x=Lp_vals, y=Qa_vals, mode='lines+markers', name='Capacidad Q_adm', line=dict(color='#00e676', width=3)))
                fig_opt.add_hline(y=Pu_req, line_dash="dash", line_color="orange", annotation_text=f"Requerido: {Pu_req} kN")
                
                idx_cumple = next((i for i, v in enumerate(Qa_vals) if v >= Pu_req), -1)
                if idx_cumple != -1:
                    Lp_opt = Lp_vals[idx_cumple]
                    fig_opt.add_vline(x=Lp_opt, line_dash="dot", line_color="#29b6f6", annotation_text=f"L_mín = {Lp_opt} m")
                    st.success(f"✔️ Para la carga requerida de {Pu_req} kN, la longitud mínima recomendada es **{Lp_opt} m**.")
                else:
                    st.error(f"❌ La carga de {Pu_req} kN sobrepasa la capacidad máxima (L={prof_total}m). Profundice la exploración o aumente diámetro.")

                fig_opt.update_layout(title="Curva de Optimización: Q_adm vs L_p", xaxis_title="Longitud Pilote L_p [m]", yaxis_title="Capacidad Admisible Q_adm [kN]", height=400, margin=dict(l=20, r=20, t=40, b=20), paper_bgcolor="#1e2530", plot_bgcolor="#1e2530", font=dict(color="white"))
                st.plotly_chart(fig_opt, use_container_width=True)
            elif _PLOTLY_AVAILABLE:
                st.info("Ingresa una carga métrica mayor a 0 para trazar el cruce de optimización.")


with tab_grup:
    st.subheader(_t("2.1 Cargas y Configuración del Grupo", "2.1 Loads and Group Configuration"))
    
    col_g1, col_g2, col_g3 = st.columns(3)
    with col_g1:
        P_ult_grupo = st.number_input("Carga Axial Estructural P_u [kN]", min_value=10.0, max_value=50000.0, value=2000.0, step=100.0)
    with col_g2:
        m_filas = st.number_input("N° Filas de Pilotes (m)", min_value=1, max_value=10, value=2)
    with col_g3:
        n_cols = st.number_input("N° Columnas de Pilotes (n)", min_value=1, max_value=10, value=2)
    
    col_g4, col_g5, col_g6 = st.columns(3)
    with col_g4:
        s_ratio = st.number_input("Espaciamiento (S_ratio = S / D)", min_value=2.0, max_value=6.0, value=3.0, step=0.1)
        # S real = S_ratio * D_pilote
        S_metros = s_ratio * D_pilote
    with col_g5:
        # Dimensionado aproximado del encepado: margen perimetral ~ 0.5D a 1.0D
        margen = max(0.30, D_pilote * 0.5)
        B_dado = (n_cols - 1) * S_metros + D_pilote + 2 * margen
        L_dado = (m_filas - 1) * S_metros + D_pilote + 2 * margen
        st.info(f"**S = {S_metros:.2f} m**\n\nEncepado: **{B_dado:.2f}m × {L_dado:.2f}m**")
    with col_g6:
        H_dado = st.number_input("Peralte del Encepado (H) [m]", 0.40, 3.00, 0.80, 0.10)
        
    st.divider()

    st.subheader(_t("2.2 Eficiencia de Grupo (Converse-Labarre)", "2.2 Group Efficiency (Converse-Labarre)"))
    n_pilotes = m_filas * n_cols
    
    if n_pilotes == 1:
        eficiencia = 1.0
        Q_grupo = Qadm
        st.success(f"Pilote único. Eficiencia η = 1.0. Capacidad = {Q_grupo:.1f} kN")
    else:
        # Converse-Labarre
        theta = math.degrees(math.atan(D_pilote / S_metros))
        termino = ((n_cols - 1) * m_filas + (m_filas - 1) * n_cols) / (m_filas * n_cols)
        eficiencia = 1.0 - (theta / 90.0) * termino
        eficiencia = max(0.5, min(eficiencia, 1.0))
        
        Q_grupo = n_pilotes * Qadm * eficiencia
        peso_dado = B_dado * L_dado * H_dado * 24.0  # Concreto 24 kN/m3
        P_total_act = P_ult_grupo + peso_dado
        
        col_r1, col_r2, col_r3 = st.columns(3)
        col_r1.metric("N° Pilotes", f"{n_pilotes}", delta=f"{m_filas}×{n_cols}")
        col_r2.metric("Eficiencia Converse-Labarre (η)", f"{eficiencia:.2f}", delta=f"{(eficiencia*100):.1f}% admisible" if eficiencia > 0.70 else "⚠ Eficiencia Baja", delta_color="normal")
        col_r3.metric("Capacidad Total del Grupo", f"{Q_grupo:.1f} kN")
        
        st.markdown(f"**Carga Total Actuante (Incl. Dado):** {P_total_act:.1f} kN")
        if P_total_act <= Q_grupo:
            st.success(f"**CUMPLE.** FS > {FS_global:.1f} para todo el grupo bajo carga {P_ult_grupo:.1f} kN.")
        else:
            st.error(f"**NO CUMPLE.** La carga actuante de {P_total_act:.1f} kN sobrepasa la capacidad del grupo conectada ({Q_grupo:.1f} kN). Aumente Q_adm mejorando la geometría o adicione pilotes.")

    # Grafica esquematica
    if _PLOTLY_AVAILABLE:
        fig_g = go.Figure()

        # Borde del dado
        bx = [-B_dado/2, B_dado/2, B_dado/2, -B_dado/2, -B_dado/2]
        by = [-L_dado/2, -L_dado/2, L_dado/2, L_dado/2, -L_dado/2]
        fig_g.add_trace(go.Scatter(x=bx, y=by, mode='lines',
            line=dict(color='#ffaa00', width=2), name="Encepado / Dado"))

        # Dibujar pilotes centrados
        offset_x = (n_cols - 1) * S_metros / 2.0
        offset_y = (m_filas - 1) * S_metros / 2.0

        px, py = [], []
        for i in range(n_cols):
            for j in range(m_filas):
                px.append(-offset_x + i * S_metros)
                py.append(-offset_y + j * S_metros)

        # Según tipo dibujamos el pilote exacto
        if tipo_seccion == "Circular":
            fig_g.add_trace(go.Scatter(x=px, y=py, mode='markers',
                marker=dict(color='#29b6f6', size=(D_pilote*40),
                            line=dict(color='white', width=1)),
                name="Pilote Circular"))
        else:
            for x, y in zip(px, py):
                hx, hy = D_pilote/2, D_pilote/2
                fig_g.add_trace(go.Scatter(
                    x=[x-hx, x+hx, x+hx, x-hx, x-hx],
                    y=[y-hy, y-hy, y+hy, y+hy, y-hy],
                    mode='lines', fill='toself',
                    fillcolor='rgba(41,182,246,0.6)',
                    line=dict(color='white', width=1),
                    showlegend=False))
            fig_g.add_trace(go.Scatter(x=[None], y=[None], mode='markers',
                marker=dict(symbol='square', color='#29b6f6', size=15),
                name="Pilote Cuadrado"))

        fig_g.update_layout(
            title="Planta Esquemática — Distribución de Pilotes",
            xaxis=dict(title="X [m]", scaleanchor="y", scaleratio=1, gridcolor="#333"),
            yaxis=dict(title="Y [m]", gridcolor="#333"),
            paper_bgcolor="#0f1117", plot_bgcolor="#0f1117",
            font=dict(color="white"),
            height=400, margin=dict(l=20, r=20, t=50, b=20)
        )
        st.plotly_chart(fig_g, use_container_width=True)

        # --- Visor 3D del Grupo ---
        st.markdown("---")
        st.subheader("Visualización 3D del Grupo de Cimentación")
        if _PLOTLY_AVAILABLE:
            fig_g3d = go.Figure()
            # Encepado (Dado) Transparente
            _Hd_cm = H_dado * 100
            _Bd_cm = B_dado * 100
            _Ld_cm = L_dado * 100
            _xD = [-_Bd_cm/2, _Bd_cm/2, _Bd_cm/2, -_Bd_cm/2, -_Bd_cm/2, _Bd_cm/2, _Bd_cm/2, -_Bd_cm/2]
            _yD = [-_Ld_cm/2, -_Ld_cm/2, _Ld_cm/2, _Ld_cm/2, -_Ld_cm/2, -_Ld_cm/2, _Ld_cm/2, _Ld_cm/2]
            _zD = [0, 0, 0, 0, _Hd_cm, _Hd_cm, _Hd_cm, _Hd_cm]  # Base del dado en Z=0
            _I_box, _J_box, _K_box = [0,0,4,4,0,1,2,3], [1,3,5,7,4,5,6,7], [2,2,6,6,1,2,3,0]
            fig_g3d.add_trace(go.Mesh3d(x=_xD, y=_yD, z=_zD, i=_I_box, j=_J_box, k=_K_box, opacity=0.3, color='#ffaa00', name='Encepado'))

            # Pilotes extruidos hacia abajo
            _Lp_cm = L_pilote * 100
            # max prof limit for visual proportion
            _Lp_disp = min(_Lp_cm, 600)
            _D_cm = D_pilote * 100
            offset_x_cm = offset_x * 100
            offset_y_cm = offset_y * 100

            for i in range(n_cols):
                for j in range(m_filas):
                    cx = -offset_x_cm + i * (S_metros * 100)
                    cy = -offset_y_cm + j * (S_metros * 100)
                    
                    if tipo_seccion == "Circular":
                        _u = np.linspace(0, 2*np.pi, 20)
                        _v = np.linspace(-_Lp_disp, 0, 2)
                        _U, _V = np.meshgrid(_u, _v)
                        _X = cx + (_D_cm/2) * np.cos(_U)
                        _Y = cy + (_D_cm/2) * np.sin(_U)
                        _Z = _V
                        fig_g3d.add_trace(go.Surface(x=_X, y=_Y, z=_Z, colorscale=[[0, '#29b6f6'], [1, '#29b6f6']], showscale=False, opacity=0.7, name=f'Pilote {i}-{j}'))
                    else:
                        hx, hy = _D_cm/2, _D_cm/2
                        _xP = [cx-hx, cx+hx, cx+hx, cx-hx, cx-hx, cx+hx, cx+hx, cx-hx]
                        _yP = [cy-hy, cy-hy, cy+hy, cy+hy, cy-hy, cy-hy, cy+hy, cy+hy]
                        _zP = [-_Lp_disp, -_Lp_disp, -_Lp_disp, -_Lp_disp, 0, 0, 0, 0]
                        fig_g3d.add_trace(go.Mesh3d(x=_xP, y=_yP, z=_zP, i=_I_box, j=_J_box, k=_K_box, color='#29b6f6', opacity=0.8, name=f'Pilote {i}-{j}'))

            fig_g3d.update_layout(scene=dict(aspectmode='data', xaxis_title='X (cm)', yaxis_title='Y (cm)', zaxis_title='Z (cm)'), paper_bgcolor="#0f1117", height=550, margin=dict(l=0, r=0, b=0, t=30))
            st.plotly_chart(fig_g3d, use_container_width=True)
            if _Lp_cm > 600:
                st.caption(f"Nota visual: La profundidad de los pilotes en el diagrama ha sido truncada a 6.0m por proporción. Longitud real calculada: {L_pilote}m.")
        # --- P2: ASENTAMIENTO DE GRUPO ---
        st.markdown("---")
        st.subheader("Asentamiento Estimado del Grupo (Bloque Equivalente)")
        col_s1, col_s2 = st.columns(2)
        E_s_pil = col_s1.number_input("Módulo elástico terreno Es [MPa]", 1.0, 200.0, 30.0, 5.0, key="es_pil")
        nu_s_pil = col_s2.number_input("Relación de Poisson terreno ν", 0.1, 0.49, 0.35, 0.05, key="nu_pil")
        
        # Bloque rígido equivalente a 2/3 L_pilote
        B_eq = offset_x * 2 + D_pilote + (0.5)*L_pilote * 0.25 # dispersion approx
        L_eq = offset_y * 2 + D_pilote + (0.5)*L_pilote * 0.25
        Area_eq = max(B_eq * L_eq, 0.1)
        
        q_eq_kpa = P_total_act / Area_eq
        S_m = (q_eq_kpa * math.sqrt(Area_eq) * (1 - nu_s_pil**2)) / (E_s_pil * 1000)
        S_mm = S_m * 1000
        
        st.metric("Asentamiento Total Grupo", f"{S_mm:.1f} mm", delta=f"{25.0 - S_mm:.1f} mm Disp", delta_color="normal"if S_mm<=25.0 else "inverse")
        if S_mm <= 25.0:
            st.success("Asentamiento dentro de tolerancias convencionales (< 25 mm).")
        else:
            st.error("Asentamiento excesivo (> 25 mm). Considere aumentar separación (S), ensanchar grupo o profundizar pilotes.")

with tab_est:
    st.subheader(_t("3.1 Refuerzo Longitudinal y Confinamiento", "3.1 Longitudinal Reinforcement & Confinement"))
    
    col_e1, col_e2, col_e3 = st.columns(3)
    with col_e1:
        st.markdown("**1. Parámetros del Material**")
        fc_pilote = st.number_input("Resistencia del Concreto f'c [MPa]", 21.0, 60.0, 28.0, 1.0)
        fy_pilote = st.number_input("Fluencia del Acero fy [MPa]", 420.0, 550.0, 420.0, 10.0)
        recub_pil = st.number_input("Recubrimiento Libre [cm]", 5.0, 15.0, 7.5, 0.5)
    with col_e2:
        st.markdown("**2. Acero Longitudinal**")
        barras_long = st.selectbox("Barra Longitudinal", ["#5 (5/8\")", "#6 (3/4\")", "#7 (7/8\")", "#8 (1\")", "#10", "#11"], index=1)
        # Diccionario interno simplificado
        _dict_ab = {"#3 (3/8\")": 0.71, "#4 (1/2\")": 1.29, "#5 (5/8\")": 1.99, "#6 (3/4\")": 2.84, "#7 (7/8\")": 3.87, "#8 (1\")": 5.10, "#10": 8.19, "#11": 10.06}
        _dict_db = {"#3 (3/8\")": 0.95, "#4 (1/2\")": 1.27, "#5 (5/8\")": 1.59, "#6 (3/4\")": 1.91, "#7 (7/8\")": 2.22, "#8 (1\")": 2.54, "#10": 3.23, "#11": 3.58}
        ab_long = _dict_ab[barras_long]
        n_barras_p = st.number_input("N° de Barras", 4, 30, max(6, int(D_pilote*10)), 1)
    with col_e3:
        st.markdown("**3. Acero Transversal (Estribo/Espiral)**")
        barras_trans = st.selectbox("Barra Confinamiento", ["#3 (3/8\")", "#4 (1/2\")", "#5 (5/8\")"], index=0)
        tipo_trans = st.radio("Tipo", ["Espiral Continua", "Estribo Cerrado"], horizontal=True)
        s_trans_cm = st.number_input("Espaciamiento (Paso) [cm]", 5.0, 30.0, 10.0, 1.0)
        
    st.divider()
    
    st.markdown("#### Cumplimiento Normativo (ACI 318 / NSR-10)")
    
    # Cálculos
    Ag_cm2 = (math.pi * (D_pilote*100)**2 / 4) if tipo_seccion == "Circular" else (D_pilote*100)**2
    As_prov_cm2 = n_barras_p * ab_long
    rho_prov = As_prov_cm2 / Ag_cm2
    
    # Cuantía min/max según NSR-10 C.10 / ACI 318 
    # Generalmente Pilotes min 1% a max 8%
    rho_min = 0.01
    rho_max = 0.08
    
    col_c1, col_c2, col_c3 = st.columns(3)
    # Check Cuantía
    with col_c1:
        st.metric("Área Bruta Ag", f"{Ag_cm2:.1f} cm²")
        st.metric("Acero Provisto As", f"{As_prov_cm2:.1f} cm²")
        
        estado_cuantia = " OK" if rho_min <= rho_prov <= rho_max else " INCUMPLE"
        st.markdown(f"**Cuantía ρ:** {(rho_prov*100):.2f}% ({estado_cuantia})")
        st.caption(f"Límites normativos: 1% a 8%")
        
    # Check Separación Longitudinal Libre s_libre
    with col_c2:
        db_long_cm = _dict_db[barras_long]
        perim_int_cm = math.pi * ((D_pilote*100) - 2*recub_pil - 2*_dict_db[barras_trans]) if tipo_seccion == "Circular" else 4 * ((D_pilote*100) - 2*recub_pil - 2*_dict_db[barras_trans])
        s_libre_long = (perim_int_cm / n_barras_p) - db_long_cm
        
        s_lim_min = max(4.0, 1.5 * db_long_cm) # ACI límite práctico
        estado_slibre = " OK" if s_libre_long >= s_lim_min else " MUY JUNTAS"
        st.metric("Espaciamiento libre", f"{s_libre_long:.1f} cm")
        st.markdown(f"**Verificación:** {estado_slibre}")
        st.caption(f"Límite min: {s_lim_min:.1f} cm")

    # Transversal Espiral/Estribo volumetrico
    with col_c3:
        # Cuantía volumétrica requerida rho_s (Solo circular/espiral simplificado)
        if tipo_seccion == "Circular" and tipo_trans == "Espiral Continua":
            Ach_cm2 = math.pi * ((D_pilote*100) - 2*recub_pil)**2 / 4
            rho_s_req = 0.45 * ((Ag_cm2 / Ach_cm2) - 1) * (fc_pilote / fy_pilote)
            
            d_out = (D_pilote*100) - 2*recub_pil
            db_t_cm = _dict_db[barras_trans]
            ab_t_cm2 = _dict_ab[barras_trans]
            d_core = d_out - db_t_cm
            
            vol_esp_1_paso = math.pi * d_core * ab_t_cm2
            vol_core_1_paso = (math.pi * d_out**2 / 4) * s_trans_cm
            rho_s_prov = vol_esp_1_paso / vol_core_1_paso
            
            estado_espiral = " OK (Espiral Densa)" if rho_s_prov >= rho_s_req else "⚠ Aumentar Espiral u optar por estribos simples si no es zona sísmica D."
            st.metric("Cuantía Volumétrica ρ_s", f"{(rho_s_prov*100):.2f}%")
            st.markdown(f"Requerido ACI §18.7.5: **{(rho_s_req*100):.2f}%**")
            st.markdown(estado_espiral)
        else:
            st.info("Verificación volumétrica requerida principalmente en pilotes circulares con espiral en zona de amenaza sísmica alta.")
            
        # --- VISUALIZACION 3D DEL PILOTE INDIVIDUAL ---
        st.divider()
        st.markdown("####  Visualización 3D del Pilote (Estructural)")
        if _PLOTLY_AVAILABLE:
            import numpy as np
            fig_p3d = go.Figure()
            _D = D_pilote * 100 # cm
            _L = min(L_pilote * 100, 300) # limitarlo a 3m en la visualizacion 3D para que no quede muy delgado
            _R = _D / 2
            _R_core = _R - recub_pil

            if tipo_seccion == "Circular":
                # Concreto (Cilindro transparente)
                _u = np.linspace(0, 2 * np.pi, 25)
                _v = np.linspace(0, _L, 2)
                _U, _V = np.meshgrid(_u, _v)
                _X = _R * np.cos(_U)
                _Y = _R * np.sin(_U)
                _Z = _V
                fig_p3d.add_trace(go.Surface(x=_X, y=_Y, z=_Z, opacity=0.15, colorscale=[[0, '#4a4a6a'], [1, '#4a4a6a']], showscale=False, name='Concreto'))

                # Refuerzo transversal (Espiral o Estribos circulares)
                _X_t, _Y_t, _Z_t = [], [], []
                if tipo_trans == "Espiral Continua":
                    _theta_t = np.linspace(0, 2 * math.pi * (_L / s_trans_cm), int(20 * (_L / s_trans_cm)))
                    _X_t = _R_core * np.cos(_theta_t)
                    _Y_t = _R_core * np.sin(_theta_t)
                    _Z_t = _theta_t / (2 * math.pi) * s_trans_cm
                    fig_p3d.add_trace(go.Scatter3d(x=_X_t, y=_Y_t, z=_Z_t, mode='lines', line=dict(color='cornflowerblue', width=4), name='Espiral'))
                else: # Estribo Cerrado Circular
                    for _z in np.arange(s_trans_cm, _L, s_trans_cm):
                        _c_u = np.linspace(0, 2*np.pi, 20)
                        _cx = _R_core * np.cos(_c_u)
                        _cy = _R_core * np.sin(_c_u)
                        _cz = np.full_like(_c_u, _z)
                        _X_t.extend(list(_cx) + [None])
                        _Y_t.extend(list(_cy) + [None])
                        _Z_t.extend(list(_cz) + [None])
                    fig_p3d.add_trace(go.Scatter3d(x=_X_t, y=_Y_t, z=_Z_t, mode='lines', line=dict(color='cornflowerblue', width=4), name='Estribos'))

                # Refuerzo longitudinal
                _rad_l = _R_core - 0.5 - math.sqrt(ab_long/math.pi)
                for _i in range(n_barras_p):
                    _ang = 2 * math.pi * _i / n_barras_p
                    _bx = _rad_l * math.cos(_ang)
                    _by = _rad_l * math.sin(_ang)
                    fig_p3d.add_trace(go.Scatter3d(x=[_bx, _bx], y=[_by, _by], z=[0, _L], mode='lines', line=dict(color='#ff6b35', width=6), showlegend=(_i==0), name='Acero Long.'))

            else: # Cuadrado
                # Concreto (Prisma transparente)
                _W = _D
                _X = [-_W/2, _W/2, _W/2, -_W/2, -_W/2, _W/2, _W/2, -_W/2]
                _Y = [-_W/2, -_W/2, _W/2, _W/2, -_W/2, -_W/2, _W/2, _W/2]
                _Z = [0, 0, 0, 0, _L, _L, _L, _L]
                _I = [0,0,4,4,0,1,2,3]; _J = [1,3,5,7,4,5,6,7]; _K = [2,2,6,6,1,2,3,0]
                fig_p3d.add_trace(go.Mesh3d(x=_X, y=_Y, z=_Z, i=_I, j=_J, k=_K, opacity=0.15, color='#4a4a6a', name='Concreto'))

                # Refuerzo transversal (Estribos cuadrados)
                _W_c = _W - 2 * recub_pil
                _X_t, _Y_t, _Z_t = [], [], []
                for _z in np.arange(s_trans_cm, _L, s_trans_cm):
                    _X_t.extend([-_W_c/2, _W_c/2, _W_c/2, -_W_c/2, -_W_c/2, None])
                    _Y_t.extend([-_W_c/2, -_W_c/2, _W_c/2, _W_c/2, -_W_c/2, None])
                    _Z_t.extend([_z, _z, _z, _z, _z, None])
                fig_p3d.add_trace(go.Scatter3d(x=_X_t, y=_Y_t, z=_Z_t, mode='lines', line=dict(color='cornflowerblue', width=4), name='Estribos'))

                # Refuerzo longitudinal
                _bars_per_face = n_barras_p // 4 + 1
                _spacing = _W_c / max(1, (_bars_per_face - 1))
                _b_coords = set()
                # Bottom & Top
                for _i in range(_bars_per_face):
                    _b_coords.add((-_W_c/2 + _i*_spacing, -_W_c/2))
                    _b_coords.add((-_W_c/2 + _i*_spacing,  _W_c/2))
                # Left & Right
                for _i in range(_bars_per_face):
                    _b_coords.add((-_W_c/2, -_W_c/2 + _i*_spacing))
                    _b_coords.add(( _W_c/2, -_W_c/2 + _i*_spacing))
                
                _coord_list = list(_b_coords)
                while len(_coord_list) > n_barras_p:
                   _coord_list.pop()

                for _i, (_bx, _by) in enumerate(_coord_list):
                    fig_p3d.add_trace(go.Scatter3d(x=[_bx, _bx], y=[_by, _by], z=[0, _L], mode='lines', line=dict(color='#ff6b35', width=6), showlegend=(_i==0), name='Acero Long.'))

            fig_p3d.update_layout(
                scene=dict(aspectmode='data', xaxis_title='X (cm)', yaxis_title='Y (cm)', zaxis_title='Z (cm)', bgcolor='#1a1a2e'),
                paper_bgcolor='#1a1a2e', font=dict(color='white'),
                height=500, margin=dict(l=0,r=0,t=40,b=0), dragmode='turntable',
                title=dict(text=f"Pilote {tipo_seccion} ({_D:.0f}cm) | Vista Modelo", font=dict(color='white'))
            )
            st.plotly_chart(fig_p3d, use_container_width=True)
            st.caption("El visor 3D Estructural ilustra el acero longitudinal y transversal de **UNA (1) sola unidad representativa** de la sección.")

# --- P3: DIAGRAMA ESTRUCTURAL P-M ---
        st.divider()
        st.subheader("Interacción P-M (Flexocompresión Biaxial)")
        Mu_req = st.number_input("Momento Actuante Mayorado Mu [kN·m]", 0.0, 10000.0, 50.0, 10.0, key="mu_pil")
        
        # Conectar con 01_Columnas_PM matemáticamente
        import importlib.util
        import sys
        from pathlib import Path
        try:
            pm_path = Path(__file__).parent / "01_Columnas_PM.py"
            spec = importlib.util.spec_from_file_location("pm_calc", pm_path)
            pm_calc = importlib.util.module_from_spec(spec)
            sys.modules["pm_calc"] = pm_calc
            spec.loader.exec_module(pm_calc)
            
            # Set layers for rebar
            # We have n_barras_p, ab_long, recub_pil, D_pilote
            # we distribute them equally around circle or square
            layers_pil = []
            if tipo_seccion == "Circular":
                rad_c = (D_pilote*100/2) - recub_pil - 0.5 - math.sqrt(ab_long/math.pi)
                for i in range(n_barras_p):
                    ang = 2 * math.pi * i / n_barras_p
                    y_pos = rad_c * math.cos(ang)
                    d_i = (D_pilote*100/2) - y_pos
                    layers_pil.append({'As': ab_long, 'd': d_i})
                
                res_pm = pm_calc.compute_uniaxial_capacity_circular(
                    D=D_pilote*100, d_prime=recub_pil, layers=layers_pil,
                    fc=fc_pilote, fy=fy_pilote, Es=200000, 
                    phi_c_max=0.75, phi_tension=0.90, eps_full=0.005, 
                    p_max_factor=0.85, factor_fuerza=1.0)
            else:
                # Square pile
                bars_per_face = n_barras_p // 4 + 1
                spacing = ((D_pilote*100) - 2*recub_pil) / (bars_per_face - 1) if bars_per_face > 1 else 0
                for i in range(n_barras_p):
                    # Simplified discrete layer distribution
                    _denom = max(n_barras_p - 1, 1)  # EF-3: evita div/0 si solo 1 barra
                    layers_pil.append({'As': ab_long, 'd': recub_pil + (D_pilote*100 - 2*recub_pil)*(i/_denom)})
                
                res_pm = pm_calc.compute_uniaxial_capacity(
                    b=D_pilote*100, h=D_pilote*100, d_prime=recub_pil, layers=layers_pil,
                    fc=fc_pilote, fy=fy_pilote, Es=200000, 
                    phi_c_max=0.65, phi_tension=0.90, eps_full=0.005, 
                    p_max_factor=0.80, factor_fuerza=1.0)

            # Plottings
            fig_pm = go.Figure()
            fig_pm.add_trace(go.Scatter(x=res_pm['M_n'], y=res_pm['P_n'], mode='lines', name='Curva Nominal (P_n, M_n)', line=dict(color='#888', dash='dash')))
            fig_pm.add_trace(go.Scatter(x=res_pm['phi_M_n'], y=res_pm['phi_P_n'], mode='lines', name='Curva de Diseño (φP_n, φM_n)', line=dict(color='#29b6f6', width=3)))
            
            # The point
            P_pil_act = P_ult_grupo / max(1, n_cols * m_filas)
            ok_pm = False
            _interp_warn = ""
            try:
                p_cap = float(pm_calc.interp_pm_curve(Mu_req, res_pm['phi_M_n'], res_pm['phi_P_n']))
                ok_pm = P_pil_act <= p_cap and p_cap > 0
            except Exception as _e_interp:
                _interp_warn = f"Error en interpolación P-M: {_e_interp}"
            
            fig_pm.add_trace(go.Scatter(x=[Mu_req], y=[P_pil_act], mode='markers', name='Carga Actuante (Pu, Mu)', marker=dict(symbol='x', color='red', size=12)))
            
            fig_pm.update_layout(title="Diagrama P-M del Pilote Individual", xaxis_title="Momento (kN·m)", yaxis_title="Carga Axial (kN)", paper_bgcolor="#0f1117", plot_bgcolor="#0f1117", font=dict(color="white"))
            st.plotly_chart(fig_pm, use_container_width=True)
            
            if _interp_warn:
                st.warning(f"⚠ {_interp_warn} — verifique que Mu esté dentro del rango del diagrama.")
            elif ok_pm:
                st.success("**CUMPLE.** La combinación Pu-Mu reside dentro de la envolvente segura.")
            else:
                st.error("**NO CUMPLE.** La combinación Pu-Mu excede la capacidad estructural del pilote.")
            
        except Exception as e:
            st.warning(f"No se pudo cargar generador P-M: {e}")
            
    with tab_apu:
        st.subheader(_t("4.1 Cantidades de Obra y Análisis de Precios Unitarios", "4.1 Quantities & Unit Price Analysis"))

        # ── Geometría base heredada de tabs anteriores ──────────────────────────
        Ag_m2_apu = (math.pi * D_pilote**2 / 4) if tipo_seccion == "Circular" else D_pilote**2
        n_pil_apu = m_filas * n_cols

        # Metrado pilotes
        vol_exc_pil   = round(Ag_m2_apu * L_pilote * n_pil_apu, 3)
        vol_conc_pil  = round(vol_exc_pil * 1.05, 3)            # +5% desperdicios
        long_total_pil = L_pilote * n_pil_apu                   # m lineales totales

        # Metrado encepado (dado)
        vol_exc_dado   = round(B_dado * L_dado * H_dado, 3)
        vol_conc_dado  = round(vol_exc_dado * 1.05, 3)
        perim_dado_apu = 2 * (B_dado + L_dado)
        area_form_dado = round(perim_dado_apu * H_dado, 2)

        # Metrado acero pilotes
        kg_m_long_apu  = ab_long * 0.785          # kg/m barra longitudinal
        long_barra_m   = L_pilote + max(0.40, H_dado - 0.15)   # anclaje en dado
        peso_long_total = round(long_barra_m * n_barras_p * kg_m_long_apu * n_pil_apu * 1.05, 1)

        kg_m_trans_apu = _dict_ab[barras_trans] * 0.785
        if tipo_seccion == "Circular" and tipo_trans == "Espiral Continua":
            long_esp_1pil = (math.pi * (D_pilote - (2 * recub_pil / 100)) / (s_trans_cm / 100)) * L_pilote
            peso_trans_total = round(long_esp_1pil * kg_m_trans_apu * n_pil_apu * 1.05, 1)
        else:
            # Recalcular perímetro interior (independiente del scope del Tab 3)
            if tipo_seccion == "Circular":
                _perim_int_cm_apu = math.pi * ((D_pilote * 100) - 2 * recub_pil - 2 * _dict_db[barras_trans])
            else:
                _perim_int_cm_apu = 4 * ((D_pilote * 100) - 2 * recub_pil - 2 * _dict_db[barras_trans])
            perim_int_m_apu = _perim_int_cm_apu / 100.0
            n_estribos_apu = int(L_pilote / (s_trans_cm / 100.0)) + 1
            peso_trans_total = round(n_estribos_apu * perim_int_m_apu * kg_m_trans_apu * n_pil_apu * 1.05, 1)

        peso_acero_pilotes = round(peso_long_total + peso_trans_total, 1)

        # ── 4.1 Cuadro de Cantidades ─────────────────────────────────────────────
        st.markdown("#####  Cuadro de Cantidades de Obra")

        filas_cant = [
            {"Nº": 1, "Descripción": "Excavación pilotes (broca/hélice)",      "Unidad": "m³",  "Cantidad": vol_exc_pil},
            {"Nº": 2, "Descripción": "Estabilización / Lodos Bentoníticos",     "Unidad": "m³",  "Cantidad": vol_exc_pil},
            {"Nº": 3, "Descripción": f"Concreto pilotes f′c={fc_pilote:.0f} MPa","Unidad": "m³", "Cantidad": vol_conc_pil},
            {"Nº": 4, "Descripción": "Acero Refuerzo pilotes (Long.+Trans.)",   "Unidad": "kg",  "Cantidad": peso_acero_pilotes},
            {"Nº": 5, "Descripción": "Excavación mecánica encepado",            "Unidad": "m³",  "Cantidad": vol_exc_dado},
            {"Nº": 6, "Descripción": f"Concreto encepado f′c={fc_pilote:.0f} MPa","Unidad": "m³","Cantidad": vol_conc_dado},
            {"Nº": 7, "Descripción": "Formaleta encepado (tableros metálicos)", "Unidad": "m²",  "Cantidad": area_form_dado},
        ]
        df_cant = pd.DataFrame(filas_cant)

        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Pilotes en Grupo",       f"{n_pil_apu} ({m_filas}×{n_cols})")
        col_m2.metric("Vol. Excavación Total",  f"{vol_exc_pil + vol_exc_dado:.2f} m³")
        col_m3.metric("Ml Totales Hincados",    f"{long_total_pil:.1f} ml")

        st.dataframe(df_cant.style.format({"Cantidad": "{:,.2f}"}),
                     use_container_width=True, hide_index=True)

        st.divider()

        # ── 4.2 Precios Unitarios (Editables) ────────────────────────────────────
        st.markdown("#####  Base de Precios Unitarios (COP) — Editable")
        st.caption("Los precios son paramétricos. Ajusta según la región y el mercado local.")

        precios_default = pd.DataFrame([
            {"Nº": 1, "Descripción": "Excavación pilotes (broca/hélice)",       "Unidad": "m³",  "Precio Unitario (COP)": 78_000},
            {"Nº": 2, "Descripción": "Estabilización / Lodos Bentoníticos",      "Unidad": "m³",  "Precio Unitario (COP)": 48_000},
            {"Nº": 3, "Descripción": f"Concreto pilotes f′c={fc_pilote:.0f} MPa","Unidad": "m³",  "Precio Unitario (COP)": 490_000},
            {"Nº": 4, "Descripción": "Acero Refuerzo pilotes (Long.+Trans.)",    "Unidad": "kg",  "Precio Unitario (COP)": 5_600},
            {"Nº": 5, "Descripción": "Excavación mecánica encepado",             "Unidad": "m³",  "Precio Unitario (COP)": 36_000},
            {"Nº": 6, "Descripción": f"Concreto encepado f′c={fc_pilote:.0f} MPa","Unidad": "m³", "Precio Unitario (COP)": 460_000},
            {"Nº": 7, "Descripción": "Formaleta encepado (tableros metálicos)",  "Unidad": "m²",  "Precio Unitario (COP)": 40_000},
        ])

        df_precios = st.data_editor(
            precios_default,
            column_config={
                "Precio Unitario (COP)": st.column_config.NumberColumn(
                    "Precio Unitario (COP)", min_value=0, format="$ %d", required=True
                )
            },
            num_rows="fixed",
            use_container_width=True,
            hide_index=True,
            key="editor_precios_pilotes"
        )

        # Merge cantidades + precios
        df_pres = df_cant.copy()
        df_pres["Precio Unitario (COP)"] = df_precios["Precio Unitario (COP)"].values
        df_pres["Subtotal (COP)"]        = df_pres["Cantidad"] * df_pres["Precio Unitario (COP)"]

        costo_directo = df_pres["Subtotal (COP)"].sum()

        st.divider()

        # ── 4.3 Configuración de AIU ──────────────────────────────────────────────
        st.markdown("##### ⚙ Administración, Imprevistos y Utilidad (A.I.U.)")
        col_aiu1, col_aiu2, col_aiu3 = st.columns(3)
        with col_aiu1:
            pct_admin = st.slider("Administración (%)", 5, 25, 15, 1, key="aiu_admin_pil")
        with col_aiu2:
            pct_impr  = st.slider("Imprevistos (%)",    2, 10, 5,  1, key="aiu_impr_pil")
        with col_aiu3:
            pct_util  = st.slider("Utilidad (%)",       5, 20, 10, 1, key="aiu_util_pil")

        pct_aiu_total = pct_admin + pct_impr + pct_util
        v_admin = costo_directo * pct_admin / 100
        v_impr  = costo_directo * pct_impr  / 100
        v_util  = costo_directo * pct_util  / 100
        v_aiu   = v_admin + v_impr + v_util
        costo_total = costo_directo + v_aiu

        st.divider()

        # ── 4.4 Presupuesto Consolidado ───────────────────────────────────────────
        st.markdown("#####  Presupuesto Consolidado")

        # KPI cards — HTML para evitar truncamiento en valores COP grandes
        st.markdown(f"""
        <div style="display:grid;grid-template-columns:repeat(5,1fr);gap:10px;margin-bottom:12px">
          <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;padding:14px 16px">
            <div style="font-size:11px;color:#8b949e;margin-bottom:4px">Costo Directo</div>
            <div style="font-size:1.15rem;font-weight:700;color:#e6edf3">${costo_directo:,.0f}</div>
          </div>
          <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;padding:14px 16px">
            <div style="font-size:11px;color:#8b949e;margin-bottom:4px">Administración ({pct_admin}%)</div>
            <div style="font-size:1.15rem;font-weight:700;color:#e6edf3">${v_admin:,.0f}</div>
          </div>
          <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;padding:14px 16px">
            <div style="font-size:11px;color:#8b949e;margin-bottom:4px">Imprevistos ({pct_impr}%)</div>
            <div style="font-size:1.15rem;font-weight:700;color:#e6edf3">${v_impr:,.0f}</div>
          </div>
          <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;padding:14px 16px">
            <div style="font-size:11px;color:#8b949e;margin-bottom:4px">Utilidad ({pct_util}%)</div>
            <div style="font-size:1.15rem;font-weight:700;color:#e6edf3">${v_util:,.0f}</div>
          </div>
          <div style="background:#1e3a5f;border:1px solid #1f6feb;border-radius:8px;padding:14px 16px">
            <div style="font-size:11px;color:#79c0ff;margin-bottom:4px">TOTAL (A.I.U.={pct_aiu_total}%)</div>
            <div style="font-size:1.15rem;font-weight:700;color:#79c0ff">${costo_total:,.0f}</div>
            <div style="font-size:10px;color:#58a6ff;margin-top:4px">${costo_total/n_pil_apu:,.0f} / pilote</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        # Tabla detallada
        df_final = df_pres[["Nº","Descripción","Unidad","Cantidad","Precio Unitario (COP)","Subtotal (COP)"]].copy()

        # Agregar fila resumen AIU al final
        resumen_rows = pd.DataFrame([
            {"Nº": "", "Descripción": "── SUBTOTAL COSTO DIRECTO",       "Unidad": "",  "Cantidad": "—", "Precio Unitario (COP)": "—", "Subtotal (COP)": f"${costo_directo:,.0f}"},
            {"Nº": "", "Descripción": f"  Administración ({pct_admin}%)", "Unidad": "%", "Cantidad": "—", "Precio Unitario (COP)": "—", "Subtotal (COP)": f"${v_admin:,.0f}"},
            {"Nº": "", "Descripción": f"  Imprevistos ({pct_impr}%)",    "Unidad": "%", "Cantidad": "—", "Precio Unitario (COP)": "—", "Subtotal (COP)": f"${v_impr:,.0f}"},
            {"Nº": "", "Descripción": f"  Utilidad ({pct_util}%)",       "Unidad": "%", "Cantidad": "—", "Precio Unitario (COP)": "—", "Subtotal (COP)": f"${v_util:,.0f}"},
            {"Nº": "", "Descripción": "TOTAL PRESUPUESTO",               "Unidad": "",  "Cantidad": "—", "Precio Unitario (COP)": "—", "Subtotal (COP)": f"${costo_total:,.0f}"},
        ])
        df_tabla_final = pd.concat([df_final, resumen_rows], ignore_index=True)

        # Formatear columnas numéricas solo en las filas de detalle
        df_detalle_fmt = df_final.copy()
        df_detalle_fmt["Cantidad"]              = df_detalle_fmt["Cantidad"].map(lambda x: f"{x:,.2f}")
        df_detalle_fmt["Precio Unitario (COP)"] = df_detalle_fmt["Precio Unitario (COP)"].map(lambda x: f"${x:,.0f}")
        df_detalle_fmt["Subtotal (COP)"]        = df_detalle_fmt["Subtotal (COP)"].map(lambda x: f"${x:,.0f}")
        df_tabla_final = pd.concat([df_detalle_fmt, resumen_rows], ignore_index=True)

        st.dataframe(
            df_tabla_final.style.apply(lambda row: [
                "background-color: #1e3a5f; font-weight:bold; color:#29b6f6"
                if str(row["Descripción"]).startswith("TOTAL") else
                "background-color: #1a2a3a; font-weight:bold; color:#e6edf3"
                if str(row["Descripción"]).startswith("──") else ""
            ] * len(row), axis=1),
            use_container_width=True, hide_index=True
        )

        st.divider()

        # ── 4.5 Indicadores de Costo Técnico ─────────────────────────────────────
        st.markdown("#####  Indicadores de Eficiencia de Costo")
        try:
            _Q_grupo_safe = Q_grupo
        except NameError:
            _Q_grupo_safe = 1.0
        costo_por_kN   = costo_total / max(_Q_grupo_safe, 1.0)            # COP/kN capacidad
        costo_ml_pil   = costo_directo / max(long_total_pil, 0.01) # COP/ml pilote
        costo_m3_conc  = (df_pres[df_pres["Descripción"].str.contains("pilotes", case=False)]["Subtotal (COP)"].sum()) / max(vol_conc_pil, 0.01)

        ic1, ic2, ic3 = st.columns(3)
        ic1.metric("Costo por kN de Capacidad",  f"${costo_por_kN:,.0f} COP/kN",
                   help=f"Costo Total / Q_grupo ({_Q_grupo_safe:.1f} kN)")
        ic2.metric("Costo por ml de Pilote",     f"${costo_ml_pil:,.0f} COP/ml",
                   help=f"Costo Directo / {long_total_pil:.1f} ml totales")
        ic3.metric("Costo Directo / m³ Concreto Pilote", f"${costo_m3_conc:,.0f} COP/m³" if vol_conc_pil > 0 else "—")

        # ── 4.6 Gráficas ─────────────────────────────────────────────────────────
        if _PLOTLY_AVAILABLE:
            col_pie1, col_pie2 = st.columns(2)

            with col_pie1:
                st.markdown("**Distribución por Actividad**")
                fig_pie = go.Figure(go.Pie(
                    labels=df_pres["Descripción"].str[:32],
                    values=df_pres["Subtotal (COP)"],
                    hole=0.45,
                    textinfo="percent",
                    hovertemplate="%{label}<br>$%{value:,.0f}<extra></extra>",
                    marker=dict(colors=[
                        "#29b6f6","#0288d1","#01579b","#80d8ff","#4fc3f7","#b3e5fc","#e1f5fe"
                    ])
                ))
                fig_pie.update_layout(
                    paper_bgcolor="#0f1117", font=dict(color="white"),
                    height=350, margin=dict(t=20, b=10, l=10, r=10),
                    legend=dict(font=dict(size=10))
                )
                st.plotly_chart(fig_pie, use_container_width=True)

            with col_pie2:
                st.markdown("**Desglose Costo Total**")
                fig_donut = go.Figure(go.Pie(
                    labels=["Costo Directo", f"Admin ({pct_admin}%)", f"Imprevistos ({pct_impr}%)", f"Utilidad ({pct_util}%)"],
                    values=[costo_directo, v_admin, v_impr, v_util],
                    hole=0.55,
                    textinfo="percent",
                    hovertemplate="%{label}<br>$%{value:,.0f}<extra></extra>",
                    marker=dict(colors=["#29b6f6","#ff9800","#ef5350","#66bb6a"])
                ))
                fig_donut.add_annotation(
                    text=f"${costo_total/1_000_000:.1f}M",
                    x=0.5, y=0.5, showarrow=False,
                    font=dict(size=18, color="white"), xanchor="center"
                )
                fig_donut.update_layout(
                    paper_bgcolor="#0f1117", font=dict(color="white"),
                    height=350, margin=dict(t=20, b=10, l=10, r=10),
                    legend=dict(font=dict(size=10))
                )
                st.plotly_chart(fig_donut, use_container_width=True)

        # ── 4.7 Export CSV ───────────────────────────────────────────────────────
        st.divider()
        st.markdown("#####  Exportar Presupuesto")
        col_exp1, col_exp2 = st.columns([1, 3])
        with col_exp1:
            csv_out = df_tabla_final.to_csv(index=False).encode("utf-8-sig")
            st.download_button(
                label="⬇ Descargar CSV",
                data=csv_out,
                file_name=f"APU_Pilotes_{n_pil_apu}u_D{D_pilote}m_L{L_pilote}m.csv",
                mime="text/csv",
                key="btn_csv_apu_pil"
            )
        with col_exp2:
            st.info(
                f" **Resumen:** Grupo de **{n_pil_apu} pilotes** {tipo_seccion.lower()}s "
                f"Ø{D_pilote}m × {L_pilote}m. "
                f"Presupuesto total: **${costo_total:,.0f} COP** "
                f"(A.I.U. = {pct_aiu_total}%)."
            )


with tab_mem:
    st.subheader(_t("5.1 Memorias, Planos y BIM", "5.1 Reports, Drawings & BIM"))
    
    col_dxf, col_docx, col_ifc = st.columns(3)
    
    # ─── Rescates de seguridad por si L_pilote_m falló en Tab 1 ───
    try: _Qp = Qp
    except NameError: _Qp = 0.0
    try: _Qs = Qs_total
    except NameError: _Qs = 0.0
    try: _efi = eficiencia
    except NameError: _efi = 1.0
    
    with col_docx:
        st.markdown('#####  Memoria de Cálculo')
        if not _DOCX_AVAILABLE:
            st.error("Librería `python-docx` requerida.")
        else:
            if st.button("Generar DOCX (Ingeniería)", key="btn_docx"):
                try:
                    doc = Document()
                    doc.add_heading(f"Memoria de Cimentación Profunda", 0)
                    doc.add_paragraph("Generado por Konte Ingeniería - Módulo NSR-10/ACI 318")
                    
                    doc.add_heading("1. Configuración Geométrica", level=1)
                    doc.add_paragraph(f"• Tipo de Pilote: {tipo_seccion}")
                    doc.add_paragraph(f"• Diámetro/Lado: {D_pilote:.2f} m")
                    doc.add_paragraph(f"• Longitud (L): {L_pilote:.2f} m")
                    doc.add_paragraph(f"• Total de Pilotes: {n_pilotes} ({m_filas}x{n_cols})")
                    doc.add_paragraph(f"• Encepado (Dado): {B_dado:.2f} m × {L_dado:.2f} m × {H_dado:.2f} m")
                    
                    doc.add_heading("2. Desempeño Geotécnico", level=1)
                    doc.add_paragraph(f"• Capacidad por Punta (Meyerhof): {_Qp:.1f} kN")
                    doc.add_paragraph(f"• Capacidad por Fuste (Tomlinson/Beta): {_Qs:.1f} kN")
                    doc.add_paragraph(f"• Eficiencia Converse-Labarre: {(_efi*100):.1f}%")
                    
                    doc.add_heading("3. Perfil Estratigráfico Relevante", level=1)
                    for _, rem in df.iterrows():
                        doc.add_paragraph(f"• Estrato {int(rem['Estrato'])}: {rem['Tipo']} | Prof: {rem['z_top']}m a {rem['z_bot']}m")
                        prp = f"   γ={rem['γ (kN/m³)']}"
                        if rem['c (kPa)'] > 0: prp += f", c={rem['c (kPa)']} kPa"
                        if rem['φ (°)'] > 0: prp += f", φ={rem['φ (°)']}°"
                        doc.add_paragraph(prp)
                    
                    doc.add_heading("4. Refuerzo Estructural ACI-318", level=1)
                    doc.add_paragraph(f"• Concreto f'c: {fc_pilote:.0f} MPa")
                    doc.add_paragraph(f"• Acero Longitudinal: {n_barras_p} barras {barras_long}")
                    doc.add_paragraph(f"• Acero Confinamiento: {barras_trans} c/{s_trans_cm} cm ({tipo_trans})")
                    
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    st.download_button(
                        label=" Descargar Memoria.docx",
                        data=doc_io,
                        file_name=f"Memoria_Pilotes_Konte.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="btn_dl_docx"
                    )
                    st.success("Memoria DOCX Generada.")
                except Exception as e:
                    st.error(f"Error DOCX: {e}")

    with col_dxf:
        st.markdown('#####  Plano en Planta (DXF)')
        if not _DXF_AVAILABLE:
            st.error("Librería `ezdxf` no instalada.")
        else:
            if st.button("Generar Plano DXF", key="btn_dxf"):
                try:
                    doc_dxf = ezdxf.new(dxfversion='R2010')
                    msp = doc_dxf.modelspace()
                    
                    # Dibujar contorno del encepado a escala
                    msp.add_lwpolyline([
                        (-B_dado/2, -L_dado/2), (B_dado/2, -L_dado/2),
                        (B_dado/2, L_dado/2), (-B_dado/2, L_dado/2),
                        (-B_dado/2, -L_dado/2)
                    ], dxfattribs={'layer': 'ENCEPADO', 'color': 3}) # Color 3=Green
                    
                    # Pilotes
                    offset_x = (n_cols - 1) * (S_metros) / 2.0
                    offset_y = (m_filas - 1) * (S_metros) / 2.0
                    
                    for i in range(n_cols):
                        for j in range(m_filas):
                            cx = -offset_x + i * S_metros
                            cy = -offset_y + j * S_metros
                            if tipo_seccion == "Circular":
                                msp.add_circle((cx, cy), D_pilote/2, dxfattribs={'layer': 'PILOTES', 'color': 4}) # Cyan
                            else:
                                Dp2 = D_pilote/2
                                msp.add_lwpolyline([
                                    (cx-Dp2, cy-Dp2), (cx+Dp2, cy-Dp2),
                                    (cx+Dp2, cy+Dp2), (cx-Dp2, cy+Dp2),
                                    (cx-Dp2, cy-Dp2)
                                ], dxfattribs={'layer': 'PILOTES', 'color': 4})
                    
                    # Textos Informativos
                    # DXF fix: API ezdxf >=0.17 usa set_placement() en vez de set_pos()
                    _t1 = msp.add_text(f"GRUPO DE {n_pilotes} PILOTES", dxfattribs={'height': 0.15, 'layer': 'TEXTOS'})
                    _t1.set_placement((-B_dado/2, L_dado/2 + 0.3))
                    _t2 = msp.add_text(f"DADO: {B_dado:.2f}x{L_dado:.2f}x{H_dado:.2f}m", dxfattribs={'height': 0.15, 'layer': 'TEXTOS'})
                    _t2.set_placement((-B_dado/2, L_dado/2 + 0.1))

                    # ROTULO PERIMETRAL ICONTEC
                    rot_w = max(4.0, B_dado/1.5); rot_h = max(1.0, B_dado/6); row_h = rot_h / 6
                    rot_x = B_dado/2 + 1.0; rot_y = -L_dado/2
                    # Marco
                    if 'ROTULO' not in doc_dxf.layers: doc_dxf.layers.add('ROTULO', color=6)
                    msp.add_lwpolyline([(rot_x, rot_y), (rot_x + rot_w, rot_y), (rot_x + rot_w, rot_y + rot_h), (rot_x, rot_y + rot_h)], close=True, dxfattribs={'layer': 'ROTULO'})
                    msp.add_line((rot_x + rot_w*0.35, rot_y), (rot_x + rot_w*0.35, rot_y + rot_h), dxfattribs={'layer': 'ROTULO'})
                    import datetime as _dt_dxf
                    _rotulo_filas = [
                        ("PROYECTO", f"Grupo {n_pilotes} Pilotes"),
                        ("DADO", f"{B_dado:.2f}x{L_dado:.2f}x{H_dado:.2f}m"),
                        ("DISEÑO", "StructuroPro"),
                        ("FECHA", _dt_dxf.date.today().strftime('%d/%m/%Y')),
                        ("PLANO", "PIL-001"), ("REV", "R0")
                    ]
                    for _i_rot, (c, v) in enumerate(_rotulo_filas):
                        yr = rot_y + rot_h - (_i_rot + 1) * row_h
                        if _i_rot > 0: msp.add_line((rot_x, yr + row_h), (rot_x + rot_w, yr + row_h), dxfattribs={'layer': 'ROTULO'})
                        _t_c = msp.add_text(c, dxfattribs={'height': row_h*0.35, 'layer': 'ROTULO'})
                        _t_c.set_placement((rot_x + rot_w*0.02, yr + row_h*0.25))
                        _t_v = msp.add_text(v, dxfattribs={'height': row_h*0.45, 'layer': 'ROTULO'})
                        _t_v.set_placement((rot_x + rot_w*0.38, yr + row_h*0.20))
                    
                    # Añadir Estratigrafía al DXF (Debajo del Rótulo)
                    _tx_est = msp.add_text("PERFIL ESTRATIGRAFICO:", dxfattribs={'height': row_h*0.4, 'layer': 'TEXTOS'})
                    _tx_est.set_placement((rot_x, rot_y - row_h*1.5))
                    for i_est, (_, rem) in enumerate(df.iterrows()):
                        props_txt = f"[{rem['z_top']}-{rem['z_bot']}m] {rem['Tipo']} | gam={rem['γ (kN/m³)']}"
                        if rem['c (kPa)'] > 0: props_txt += f" c={rem['c (kPa)']}"
                        if rem['φ (°)'] > 0: props_txt += f" phi={rem['φ (°)']}"
                        _tx_p = msp.add_text(props_txt, dxfattribs={'height': row_h*0.35, 'layer': 'TEXTOS'})
                        _tx_p.set_placement((rot_x, rot_y - row_h * (i_est + 2.5)))

                    dxf_io = io.StringIO()
                    doc_dxf.write(dxf_io)
                    dxf_io.seek(0)
                    st.download_button(
                        label=" Descargar Plano.dxf",
                        data=dxf_io.getvalue().encode('utf-8'),
                        file_name=f"Plano_Pilotes_{n_pilotes}u.dxf",
                        mime="application/dxf",
                        key="btn_dl_dxf"
                    )
                    st.success("Plano DXF Geométrico Generado.")
                except Exception as e:
                    st.error(f"Error DXF: {e}")

    with col_ifc:
        st.markdown('#####  Integración BIM 3D (IFC)')
        if not _IFC_AVAILABLE:
            st.warning("`ifcopenshell` no vinculado en este entorno Python.")
            st.info("Pip install ifcopenshell requerido para activar BIM 3D.")
        else:
            if st.button("Exportar Modelo BIM (.ifc)", help="BIM 3D con Encepado + Pilotes", key="btn_ifc"):
                try:
                    if 'ifc_grupo_pilotes' not in globals():
                        st.error("No se detectó ifc_grupo_pilotes en ifc_export.py")
                    else:
                        fc_d = 21.0 # default dado
                        # Exportación usando la nueva función BIM
                        ifc_io = ifc_grupo_pilotes(
                            B_dado_m=B_dado, L_dado_m=L_dado, H_dado_m=H_dado,
                            tipo_seccion=tipo_seccion, D_pilote_m=D_pilote, L_pilote_m=L_pilote,
                            m_filas=m_filas, n_cols=n_cols, S_metros=S_metros,
                            fc_dado=fc_d, fc_pilote=fc_pilote,
                            nombre_proyecto=f"Grupo de {n_pilotes} Pilotes"
                        )
                        st.download_button(
                            label=" Descargar Modelo.ifc",
                            data=ifc_io,
                            file_name=f"Modelo_BIM_{n_pilotes}Pilotes.ifc",
                            mime="application/octet-stream",
                            key="btn_dl_ifc"
                        )
                        st.success("Modelo IFC4 Generado con Éxito.")
                except Exception as e:
                    st.error(f"Error Genérico IFC: {e}")

# ==============================================================================
# P2 & P3: CUADRO DE MANDO Y ASENTAMIENTOS DIFERENCIALES
# ==============================================================================
st.divider()
st.subheader("Cuadro de Mando — Registro de Grupos de Pilotes")

etiqueta_pi = st.text_input("Etiqueta del grupo (ej. P-1, Eje 3-B)", value="P-1", key="pi_etiqueta")
if st.button("Agregar al proyecto", key="pi_agregar"):
    try:
        if "registro_pilotes" not in st.session_state:
            st.session_state.registro_pilotes = []
        st.session_state.registro_pilotes.append({
            "Grupo": etiqueta_pi, "N Pilotes": n_pilotes,
            "D [m]": D_pilote, "L [m]": L_pilote,
            "Qadm [kN]": round(Qadm, 1), "Qgrupo [kN]": round(Q_grupo, 1),
            "Eficiencia [%]": round(eficiencia * 100, 1),
            "Asentamiento [mm]": round(S_mm, 1),
            "Estado": " CUMPLE" if P_total_act <= Q_grupo else " NO CUMPLE"
        })
    except NameError:
        st.warning("⚠ Completa los cálculos y geometría antes de guardar el grupo.")

if st.session_state.get("registro_pilotes"):
    st.dataframe(pd.DataFrame(st.session_state.registro_pilotes),
                 use_container_width=True, hide_index=True)
    if st.button("Limpiar registro", key="pi_limpiar"):
        st.session_state.registro_pilotes = []
        st.rerun()

if len(st.session_state.get("registro_pilotes", [])) >= 2:
    asentamientos = [r["Asentamiento [mm]"] for r in st.session_state.registro_pilotes]
    delta_s = max(asentamientos) - min(asentamientos)
    col_d1, col_d2 = st.columns(2)
    col_d1.metric("Δ Asentamiento diferencial máx.", f"{delta_s:.1f} mm",
                  delta=f"{19.0 - delta_s:.1f} mm Margen",
                  delta_color="normal" if delta_s <= 19 else "inverse")
    if delta_s > 19:
        st.error("⚠ Asentamiento diferencial > 19 mm (NSR-10 H.3.4). "
                 "Revisar rigidez de la superestructura o nivelar profundidades.")
    else:
        st.success(f"Δs = {delta_s:.1f} mm ≤ 19 mm — Dentro de tolerancia NSR-10.")