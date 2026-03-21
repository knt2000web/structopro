import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Columnas Circulares P-M", "Circular Columns P-M"), layout="wide")

st.image(r"assets/circular_concrete_column_technical_render.png", use_container_width=False, width=700)
st.title(_t("Diagrama P-M para Columnas Circulares — Multi-Norma", "P-M Diagram for Circular Columns — Multi-Code"))
st.markdown(_t("Generación del diagrama de interacción **P–M** para columnas circulares de concreto reforzado. Soporta normativa de **Colombia, EE.UU., Ecuador, Perú, México, Venezuela, Bolivia y Argentina**.", "Generation of the **P-M** interaction diagram for reinforced concrete circular columns. Supports codes from **Colombia, USA, Ecuador, Peru, Mexico, Venezuela, Bolivia and Argentina**."))

with st.expander(_t("📺 ¿Cómo usar esta herramienta?", "📺 How to use this tool?")):
    st.markdown(_t("""
    **Modo de Uso:**
    1. **📍 Sidebar (Menú Izquierdo):** Selecciona la Norma de tu país, el nivel sísmico, y las propiedades del concreto y el acero.
    2. **🏗️ Geometría y Armadura:** Define el Diámetro D de la columna, el diámetro de varilla y cantidad total de varillas longitudinales. El acero se distribuirá simétricamente en el perímetro.
    3. **🚦 Solicitaciones:** Ingresa el Momento Último (Mu) y la Carga Axial (Pu) solicitantes.
    4. **📊 P-M:** Revisa que el punto verde de solicitación quede DENTRO de la curva roja (Resistencia de Diseño).
    """, """
    **How to use:**
    1. **📍 Sidebar (Left Menu):** Select the Country Code, seismic level, and concrete and steel properties.
    2. **🏗️ Geometry and Reinforcement:** Define Column Diameter D, rebar diameter and total number of longitudinal rebars. Steel will be distributed symmetrically on the perimeter.
    3. **🚦 Demands:** Enter the required Ultimate Moment (Mu) and Axial Load (Pu).
    4. **📊 P-M:** Check that the green demand point falls INSIDE the red curve (Design Strength).
    """))

# ─────────────────────────────────────────────
# DICCIONARIOS DE BARRAS
# ─────────────────────────────────────────────
REBAR_US = {
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
    "#5 (5/8\")": {"area": 1.99, "diam_mm": 15.88},
    "#6 (3/4\")": {"area": 2.84, "diam_mm": 19.05},
    "#7 (7/8\")": {"area": 3.87, "diam_mm": 22.23},
    "#8 (1\")":   {"area": 5.10, "diam_mm": 25.40},
    "#9 (1 1/8\")": {"area": 6.45, "diam_mm": 28.65},
    "#10 (1 1/4\")": {"area": 7.92, "diam_mm": 32.26},
}

REBAR_MM = {
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
    "14 mm": {"area": 1.539, "diam_mm": 14.0},
    "16 mm": {"area": 2.011, "diam_mm": 16.0},
    "18 mm": {"area": 2.545, "diam_mm": 18.0},
    "20 mm": {"area": 3.142, "diam_mm": 20.0},
    "22 mm": {"area": 3.801, "diam_mm": 22.0},
    "25 mm": {"area": 4.909, "diam_mm": 25.0},
    "28 mm": {"area": 6.158, "diam_mm": 28.0},
    "32 mm": {"area": 8.042, "diam_mm": 32.0},
}

STIRRUP_US = {
    "#2 (1/4\")": {"area": 0.32, "diam_mm": 6.35},
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
}

STIRRUP_MM = {
    "6 mm":  {"area": 0.283, "diam_mm": 6.0},
    "8 mm":  {"area": 0.503, "diam_mm": 8.0},
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
}

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# PARÁMETROS POR NORMA
# ─────────────────────────────────────────────
CODES = {
    "NSR-10 (Colombia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 4.0, "eps_tension_full": 0.005, "ref": "NSR-10 Título C", "bag_kg": 50.0
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0, "eps_tension_full": 0.005, "ref": "ACI 318-25", "bag_kg": 42.6
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0, "eps_tension_full": 0.005, "ref": "ACI 318-19", "bag_kg": 42.6
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0, "eps_tension_full": 0.005, "ref": "ACI 318-14", "bag_kg": 42.6
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0, "eps_tension_full": 0.005, "ref": "NEC-SE-HM (Ecuador)", "bag_kg": 50.0
    },
    "E.060 (Perú)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 6.0, "eps_tension_full": 0.005, "ref": "Norma E.060 (Perú)", "bag_kg": 42.5
    },
    "NTC-EM (México)": {
        "phi_tied": 0.70, "phi_spiral": 0.80, "phi_tension": 0.85, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 6.0, "eps_tension_full": 0.005, "ref": "NTC-EM México", "bag_kg": 50.0
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 6.0, "eps_tension_full": 0.005, "ref": "COVENIN 1753-2006 (Venezuela)", "bag_kg": 42.5
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0, "eps_tension_full": 0.005, "ref": "NB 1225001-2020 (Bolivia)", "bag_kg": 50.0
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90, "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0, "eps_tension_full": 0.005, "ref": "CIRSOC 201-2025 (Argentina)", "bag_kg": 50.0
    },
}

def mix_for_fc(fc):
    """ACI 211 mix proportions: (cement_kg, water_L, sand_kg, gravel_kg) per m³"""
    table = [
        (14, 250, 205, 810, 1060),
        (17, 290, 200, 780, 1060),
        (21, 350, 193, 720, 1060),
        (25, 395, 193, 680, 1020),
        (28, 430, 190, 640, 1000),
        (35, 530, 185, 580, 960),
        (42, 620, 180, 520, 910),
        (56, 740, 175, 450, 850),
    ]
    if fc <= table[0][0]:
        return table[0][1:]
    if fc >= table[-1][0]:
        return table[-1][1:]
    for i in range(len(table)-1):
        lo, hi = table[i], table[i+1]
        if lo[0] <= fc <= hi[0]:
            t = (fc - lo[0]) / (hi[0] - lo[0])
            return tuple(lo[j] + t*(hi[j] - lo[j]) for j in range(1, 5))
    return table[-1][1:]

def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_development_length(db_mm, fy, fc, lambda_=1.0, psi_t=1.0, psi_e=1.0, psi_s=1.0, psi_g=1.0, cb_ktr=2.5):
    """Longitud de desarrollo a tracción para barras rectas (ACI 318-19 25.4.2.3)"""
    # Unidades: db_mm, fy MPa, fc MPa
    # Retorna mm
    if db_mm <= 0: return 0
    ld = (3/40) * (fy / (lambda_ * math.sqrt(fc))) * (psi_t * psi_e * psi_s * psi_g / cb_ktr) * db_mm
    return max(ld, 300)  # mínimo 300 mm

# ─────────────────────────────────────────────
# INPUTS
# ─────────────────────────────────────────────
st.sidebar.header(_t("0. Norma de Diseño", "0. Design Code"))

if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = list(CODES.keys())[0]

norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)
code = CODES[norma_sel]
bag_kg = code["bag_kg"]

st.sidebar.header(_t("1. Materiales", "1. Materials"))
fc_unit = st.sidebar.radio(_t("Unidad de f'c:", "f'c Unit:"), ["MPa", "PSI", "kg/cm²"], horizontal=True, key="circ_fc_unit")
if fc_unit == "PSI":
    fc_psi = st.sidebar.number_input("f'c [PSI]", 2000.0, 12000.0, st.session_state.get("circ_fc_psi", 4000.0), 100.0, key="circ_fc_psi")
    fc = fc_psi * 0.00689476
elif fc_unit == "kg/cm²":
    fc_kgcm2 = st.sidebar.number_input("f'c [kg/cm²]", 100.0, 1200.0, st.session_state.get("circ_fc_kgcm2", 280.0), 10.0, key="circ_fc_kgcm2")
    fc = fc_kgcm2 / 10.1972
else:
    fc = st.sidebar.number_input(_t("Resistencia del Concreto (f'c) [MPa]", "Concrete Strength (f'c) [MPa]"), 15.0, 80.0, st.session_state.get("circ_fc_mpa", 28.0), 1.0, key="circ_fc_mpa")
fy = st.sidebar.number_input(_t("Fluencia del Acero (fy) [MPa]", "Steel Yield (fy) [MPa]"), 240.0, 500.0, st.session_state.get("circ_fy", 420.0), 10.0, key="circ_fy")
Es = 200000.0

st.sidebar.header(_t("2. Geometría", "2. Geometry"))
D_col = st.sidebar.number_input(_t("Diámetro Exterior (D) [cm]", "Outer Diameter (D) [cm]"), min_value=20.0, value=st.session_state.get("circ_D", 60.0), step=5.0, key="circ_D")
recub = st.sidebar.number_input(_t("Recubrimiento Libre (al estribo) [cm]", "Clear cover (to tie) [cm]"), min_value=1.5, value=st.session_state.get("circ_recub", 4.0), step=0.5, key="circ_recub")
L_col = st.sidebar.number_input(_t("Altura libre columna (L) [cm]", "Clear column height (L) [cm]"), min_value=50.0, value=st.session_state.get("circ_L", 300.0), step=25.0, key="circ_L")

st.sidebar.header(_t("3. Refuerzo Longitudinal", "3. Longitudinal Reinforcement"))
unit_system = st.sidebar.radio(_t("Sistema de Varillas:", "Rebar System:"), 
                               ["Pulgadas (EE. UU.)", "Milímetros (SI)"] if lang == "Español" else ["Inches (US)", "Millimeters (SI)"],
                               key="circ_unit_system")
rebar_dict = REBAR_US if "Pulgadas" in unit_system or "Inches" in unit_system else REBAR_MM
default_rebar = "#6 (3/4\")" if "Pulgadas" in unit_system or "Inches" in unit_system else "20 mm"
rebar_type = st.sidebar.selectbox(_t("Diámetro de Varillas Long.", "Longitudinal Rebars Dia."), list(rebar_dict.keys()), 
                                 index=list(rebar_dict.keys()).index(st.session_state.circ_rebar_type) if "circ_rebar_type" in st.session_state and st.session_state.circ_rebar_type in rebar_dict else list(rebar_dict.keys()).index(default_rebar),
                                 key="circ_rebar_type")
rebar_area  = rebar_dict[rebar_type]["area"]    # cm²
rebar_diam  = rebar_dict[rebar_type]["diam_mm"] # mm

n_barras_total = st.sidebar.number_input("Total de Varillas", min_value=4, max_value=40, value=st.session_state.get("circ_n_bars", 8), step=1, key="circ_n_bars")
Ast = n_barras_total * rebar_area # cm²
Ag = math.pi * (D_col / 2.0)**2     # cm²

st.sidebar.header(_t("4. Estribo / Zuncho", "4. Tie / Spiral"))
stirrup_dict = STIRRUP_US if "Pulgadas" in unit_system or "Inches" in unit_system else STIRRUP_MM
default_stirrup = "#3 (3/8\")" if "Pulgadas" in unit_system or "Inches" in unit_system else "10 mm"
stirrup_type = st.sidebar.selectbox(_t("Diámetro del Estribo/Espiral", "Tie/Spiral Diameter"), list(stirrup_dict.keys()), 
                                   index=list(stirrup_dict.keys()).index(st.session_state.circ_stirrup_type) if "circ_stirrup_type" in st.session_state and st.session_state.circ_stirrup_type in stirrup_dict else list(stirrup_dict.keys()).index(default_stirrup),
                                   key="circ_stirrup_type")
stirrup_area = stirrup_dict[stirrup_type]["area"]     # cm²
stirrup_diam = stirrup_dict[stirrup_type]["diam_mm"]  # mm

col_type_options = ["Espiral (Zunchada)", "Estribos (Tied)"] if lang == "Español" else ["Spiral", "Tied"]
col_type = st.sidebar.selectbox(_t("Tipo de Confinamiento", "Confinement Type"), 
                                col_type_options,
                                index=col_type_options.index(st.session_state.circ_col_type) if "circ_col_type" in st.session_state and st.session_state.circ_col_type in col_type_options else 0,
                                key="circ_col_type")
if "Estrib" in col_type or "Tied" in col_type:
    phi_c_max  = code["phi_tied"]
    p_max_factor = code["pmax_tied"]
else:
    phi_c_max  = code["phi_spiral"]
    p_max_factor = code["pmax_spiral"]
phi_tension   = code["phi_tension"]

st.sidebar.header(_t("5. Verificación", "5. Demands"))
M_u_input = st.sidebar.number_input(_t("Momento Último (Mu) [kN-m]", "Ultimate Moment (Mu) [kN-m]"), value=st.session_state.get("circ_mu", 100.0), step=10.0, key="circ_mu")
P_u_input = st.sidebar.number_input(_t("Carga Axial Última (Pu) [kN]", "Ultimate Axial Load (Pu) [kN]"), value=st.session_state.get("circ_pu", 800.0), step=50.0, key="circ_pu")

# Longitud de desarrollo y empalme
lambda_factor = 1.0  # concreto normal
psi_t = 1.0  # barra inferior (por defecto)
psi_e = 1.0  # sin epoxy
psi_s = 1.0  # barra >= #6? Usamos 1.0 por simplicidad
psi_g = 0.75 if fy <= 420 else 1.0  # psi_g (fy≤420 -> 0.75)
cb_ktr = 2.5  # recubrimiento efectivo (suponemos favorable)
ld_mm = get_development_length(rebar_diam, fy, fc, lambda_factor, psi_t, psi_e, psi_s, psi_g, cb_ktr)
lap_length_mm = 1.3 * ld_mm  # empalme clase B
splice_zone_height = lap_length_mm / 10  # cm

# Ubicación del empalme (a 1/3 de la altura desde la base, por ejemplo)
splice_start = L_col / 3
splice_end = splice_start + splice_zone_height

# ─────────────────────────────────────────────
# CÁLCULOS P-M
# ─────────────────────────────────────────────
eps_cu = 0.003
eps_y  = fy / Es

beta_1 = get_beta1(fc)

D_mm = D_col * 10
R_mm = D_mm / 2.0
As_mm2 = rebar_area * 100

Po_kN = (0.85 * fc * (Ag * 100 - Ast * 100) + fy * Ast * 100) / 1000.0

# Posiciones de las barras (coordenadas en mm, centro en (0,0))
Rs_mm = R_mm - (recub * 10) - stirrup_diam - (rebar_diam / 2.0)
bar_coords = []  # lista de (x_mm, y_mm)
for i in range(n_barras_total):
    theta_i = (2 * math.pi * i) / n_barras_total
    x_mm = Rs_mm * math.cos(theta_i)
    y_mm = Rs_mm * math.sin(theta_i)
    bar_coords.append((x_mm, y_mm))

# Cálculo del diagrama
c_vals = np.concatenate([np.linspace(1e-5, D_col, 120), np.linspace(D_col, D_col * 12, 60)])
P_n_list = []; M_n_list = []; phi_P_n_list = []; phi_M_n_list = []
eps_t_list = []

for c_cm in c_vals:
    c_mm = c_cm * 10
    a_mm = min(beta_1 * c_mm, D_mm)
    
    # Área de concreto comprimida y su centroide
    d_c = R_mm - a_mm
    if d_c >= R_mm:
        Ac = 0.0; yc = 0.0
    elif d_c <= -R_mm:
        Ac = math.pi * R_mm**2; yc = 0.0
    else:
        alpha = math.acos(d_c / R_mm)
        Ac = R_mm**2 * (alpha - math.sin(alpha) * math.cos(alpha))
        yc = (2 * R_mm**3 * math.sin(alpha)**3) / (3 * Ac) if Ac > 0 else 0
    Cc = 0.85 * fc * Ac  # N
    Mc = Cc * yc         # N·mm
    
    Ps = 0.0; Ms = 0.0; eps_t = 0.0
    for (x_mm, y_mm) in bar_coords:
        di = R_mm - y_mm
        eps_s = eps_cu * (c_mm - di) / c_mm
        dt = R_mm + Rs_mm
        eps_t = eps_cu * (c_mm - dt) / c_mm
        
        fs = max(-fy, min(fy, Es * eps_s))
        if a_mm > di and fs > 0:
            fs -= 0.85 * fc
        Ps += As_mm2 * fs
        Ms += As_mm2 * fs * y_mm
    
    Pn = (Cc + Ps) / 1000.0
    Mn = abs((Mc + Ms) / 1_000_000.0)
    eps_t_list.append(eps_t)
    
    # factor phi
    eps_t_tens = -eps_t
    if eps_t_tens <= eps_y:
        phi = phi_c_max
    elif eps_t_tens >= code["eps_tension_full"]:
        phi = phi_tension
    else:
        phi = phi_c_max + (phi_tension - phi_c_max) * (eps_t_tens - eps_y) / (code["eps_tension_full"] - eps_y)
    
    Pn_max_val = p_max_factor * Po_kN
    phi_Pn_max_val = phi_c_max * Pn_max_val
    
    Pn_adj = min(Pn, Pn_max_val)
    phi_Pn_adj = min(phi * Pn_adj, phi_Pn_max_val)
    phi_Mn_adj = phi * Mn
    
    P_n_list.append(Pn_adj); M_n_list.append(Mn)
    phi_P_n_list.append(phi_Pn_adj); phi_M_n_list.append(phi_Mn_adj)

P_n_arr = np.array(P_n_list); M_n_arr = np.array(M_n_list)
phi_P_n_arr = np.array(phi_P_n_list); phi_M_n_arr = np.array(phi_M_n_list)

Pn_max = p_max_factor * Po_kN
phi_Pn_max_disp = phi_c_max * Pn_max

# Verificación del punto de diseño
idx = np.where(phi_P_n_arr >= P_u_input)[0]
if len(idx) > 0:
    idx_upper = idx[-1]
    if idx_upper < len(phi_M_n_arr)-1:
        P1 = phi_P_n_arr[idx_upper]; M1 = phi_M_n_arr[idx_upper]
        P2 = phi_P_n_arr[idx_upper+1]; M2 = phi_M_n_arr[idx_upper+1]
        if P2 != P1:
            M_lim = M1 + (M2 - M1) * (P_u_input - P1) / (P2 - P1)
        else:
            M_lim = M1
    else:
        M_lim = phi_M_n_arr[idx_upper]
else:
    M_lim = 0.0
ok_design = M_u_input <= M_lim

# ─────────────────────────────────────────────
# CANTIDADES DE MATERIALES
# ─────────────────────────────────────────────
vol_concreto_m3 = (math.pi * (D_col/100)**2 / 4) * (L_col/100)
long_varilla_m = L_col/100 + 0.3  # +0.3m para traslapo/anclaje
peso_acero_long_kg = Ast * long_varilla_m * 0.785  # kg

# Estribos / espiral
if "Estrib" in col_type or "Tied" in col_type:
    # Separación de estribos según ACI 318-25 25.7.2
    s_max_mm = min(16 * rebar_diam, 48 * stirrup_diam, D_col * 10)
    s_estribo_cm = min(s_max_mm / 10, 30)  # limitamos a 30 cm para simplificar
    perim_estribo_m = math.pi * (D_col/100 - 2*recub/100)
    n_estribos = math.ceil(L_col/100 / (s_estribo_cm/100)) + 1
    long_estribos_m = n_estribos * perim_estribo_m
    long_estribos_m += n_estribos * 0.1  # ganchos
else:
    # Espiral
    s_estribo_cm = min(8 * rebar_diam / 10, 10)  # paso máximo 8*db o 10 cm
    paso_espiral = s_estribo_cm / 100
    diam_espiral_m = D_col/100 - 2*recub/100
    long_estribos_m = math.sqrt((math.pi * diam_espiral_m)**2 + paso_espiral**2) * (L_col/100 / paso_espiral)
peso_acero_estribos_kg = long_estribos_m * stirrup_area * 0.785
peso_total_acero_kg = peso_acero_long_kg + peso_acero_estribos_kg
relacion_acero_kg_m3 = peso_total_acero_kg / vol_concreto_m3 if vol_concreto_m3 > 0 else 0
cuantia = Ast / Ag * 100
ok_rho = (code["rho_min"] <= cuantia <= code["rho_max"])

# Mezcla para dosificación
cement_kg, water_L, sand_kg, gravel_kg = mix_for_fc(fc)
total_cement_kg = cement_kg * vol_concreto_m3
total_water_L = water_L * vol_concreto_m3
total_sand_kg = sand_kg * vol_concreto_m3
total_gravel_kg = gravel_kg * vol_concreto_m3
bags_cement = total_cement_kg / bag_kg

# ─────────────────────────────────────────────
# LAYOUT 
# ─────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs(["📊 Diagrama P–M", "🔲 Sección Transversal", "📦 Costos (APU) y Cantidades"])

with tab1:
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader(f"Gráfica P–M — {norma_sel}")
        fig, ax = plt.subplots(figsize=(8, 6))
        fig.patch.set_facecolor('#1a1a2e')
        ax.set_facecolor('#1a1a2e')

        idx_M_nom = np.argmax(M_n_arr); M_nom_max = M_n_arr[idx_M_nom]; P_at_Mnom = P_n_arr[idx_M_nom]
        idx_M_dis = np.argmax(phi_M_n_arr); M_dis_max = phi_M_n_arr[idx_M_dis]

        ax.plot(M_n_arr, P_n_arr, label=r"Resistencia Nominal ($P_n, M_n$)", color="blue", linestyle="--")
        ax.plot(phi_M_n_arr, phi_P_n_arr, label=r"Resistencia de Diseño ($\phi P_n, \phi M_n$)", color="red", linewidth=2)
        ax.plot(M_u_input, P_u_input, 'o', markersize=9, color='lime', markeredgecolor='black', zorder=6, label="Punto de Diseño (Mu, Pu)")
        
        ax.annotate(f"{Pn_max:.0f} kN", xy=(0, Pn_max), xytext=(5, 5), textcoords="offset points", color='blue')
        ax.annotate(f"{phi_Pn_max_disp:.0f} kN", xy=(0, phi_Pn_max_disp), xytext=(5, -15), textcoords="offset points", color='red')
        ax.annotate(f"[{M_u_input:g} kN·m : {P_u_input:g} kN]", xy=(M_u_input, P_u_input), xytext=(6, 6), textcoords="offset points", color='green', weight='bold')
        
        ax.plot([M_nom_max, M_nom_max], [0, P_at_Mnom], color='gray', linestyle='--', alpha=0.5)
        ax.plot([M_dis_max, M_dis_max], [0, phi_P_n_arr[idx_M_dis]], color='gray', linestyle='--', alpha=0.5)

        ax.axhline(0, color='white', linewidth=1)
        ax.set_xlabel(_t("Momento Flector M [kN·m]", "Bending Moment M [kN·m]"), color='white')
        ax.set_ylabel(_t("Carga Axial P [kN]", "Axial Load P [kN]"), color='white')
        ax.set_xlim(left=0)
        ax.set_title(_t(f"Columna Circular Ø{D_col:.0f} cm | f'c={fc:.1f} MPa | {n_barras_total} varillas {rebar_type}", 
                        f"Circular Column Ø{D_col:.0f} cm | f'c={fc:.1f} MPa | {n_barras_total} rebars {rebar_type}"), color='white')
        ax.grid(True, linestyle=":", alpha=0.5)
        ax.legend(loc='best')
        ax.tick_params(colors='white')
        st.pyplot(fig)

        # Guardar figura para la memoria
        img_buffer = io.BytesIO()
        fig.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)

    with col2:
        st.subheader(_t("Resumen de Diseño", "Design Summary"))
        st.markdown(f"**Norma:** {norma_sel}")
        st.write(f"**Cuantía ρ:** {cuantia:.2f}%")
        st.markdown(r"**Verificación Normativa:** $\rho_{min} \le \rho \le \rho_{max}$")
        if ok_rho:
            st.success(f"✅ {_t('Aprobado Cuantía:', 'Steel ratio OK:')} {code['rho_min']:.1f}% $\\le$ {cuantia:.2f}% $\\le$ {code['rho_max']:.1f}%")
        elif cuantia > code["rho_max"]:
            st.error(f"❌ {_t('No Aprobado por Cuantía Máxima:', 'Exceeds maximum steel ratio:')} $\\rho = {cuantia:.2f}\% > \\rho_{{max}} = {code['rho_max']:.1f}\%$ $\\rightarrow$ {_t('Aumentar Diámetro D', 'Increase Diameter D')}")
        else:
            st.error(f"❌ {_t('No Aprobado por Cuantía Mínima:', 'Below minimum steel ratio:')} $\\rho = {cuantia:.2f}\% < \\rho_{{min}} = {code['rho_min']:.1f}\%$ $\\rightarrow$ {_t('Aumentar acero longitudinal', 'Increase longitudinal steel')}")
        
        st.markdown(f"**{_t('Capacidad Máxima Pn,máx:', 'Maximum Nominal Capacity Pn,max:')}** {Pn_max:.0f} kN")
        st.markdown(f"**{_t('Resistencia de Diseño φPn,máx:', 'Design Strength φPn,max:')}** {phi_Pn_max_disp:.0f} kN")
        st.markdown("---")
        st.metric(_t("As Total Provisto", "Total Provided As"), f"{Ast:.2f} cm²")
        
        st.markdown("---")
        st.markdown(f"**{_t('Estribos / Espiral:', 'Ties / Spiral:')}**")
        st.write(f"- {_t('Diámetro:', 'Diameter:')} {stirrup_type}")
        st.write(f"- {_t('Separación de diseño:', 'Design spacing:')} {s_estribo_cm:.1f} cm")
        st.write(f"- {_t('Peso estimado:', 'Estimated weight:')} {peso_acero_estribos_kg:.1f} kg")
        
        if not ok_design:
            st.error(f"❌ {_t('El punto de diseño (Mu, Pu) está fuera de la curva de resistencia.', 'The design point (Mu, Pu) is outside the strength curve.')}")

with tab2:
    st.subheader(_t("🧊 Visualización 3D en Tiempo Real", "🧊 Real-time 3D Visualization"))
    fig3d = go.Figure()

    # Cilindro de concreto
    R_cm = D_col / 2.0
    theta_3d = np.linspace(0, 2 * np.pi, 48)
    x_c = R_cm * np.cos(theta_3d)
    y_c = R_cm * np.sin(theta_3d)
    z_top = L_col
    z_bot = 0

    # Mallas laterales
    x_mesh = np.concatenate([x_c, x_c])
    y_mesh = np.concatenate([y_c, y_c])
    z_mesh = np.concatenate([np.full(len(theta_3d), z_bot), np.full(len(theta_3d), z_top)])
    i_idx, j_idx, k_idx = [], [], []
    for i in range(len(theta_3d)-1):
        i_idx.extend([i, i+1, i+len(theta_3d)])
        j_idx.extend([i+1, i+len(theta_3d)+1, i+len(theta_3d)+1])
        k_idx.extend([i+len(theta_3d), i+len(theta_3d), i+1])
    i_idx.extend([len(theta_3d)-1, 0, len(theta_3d)-1+len(theta_3d)])
    j_idx.extend([0, len(theta_3d), len(theta_3d)])
    k_idx.extend([len(theta_3d)-1+len(theta_3d), len(theta_3d)-1+len(theta_3d), 0])

    fig3d.add_trace(go.Mesh3d(x=x_mesh, y=y_mesh, z=z_mesh, i=i_idx, j=j_idx, k=k_idx,
                              alphahull=-1, opacity=0.15, color='gray', name=_t('Concreto', 'Concrete')))

    # Varillas longitudinales
    diam_reb_cm = rebar_diam / 10.0
    line_width = max(4, diam_reb_cm * 3)
    for (x_mm, y_mm) in bar_coords:
        x_cm = x_mm / 10
        y_cm = y_mm / 10
        fig3d.add_trace(go.Scatter3d(x=[x_cm, x_cm], y=[y_cm, y_cm], z=[z_bot, z_top],
                                     mode='lines', line=dict(color='darkred', width=line_width),
                                     name=_t('Varilla', 'Rebar'), showlegend=False))

    # Zona de empalme (anillo semitransparente)
    if splice_zone_height > 0:
        z_center = splice_start + splice_zone_height/2
        splice_radius = R_cm + 2  # ligeramente mayor que la columna
        # Crear un cilindro hueco (usamos un Mesh3d)
        theta_ring = np.linspace(0, 2*np.pi, 36)
        x_ring = splice_radius * np.cos(theta_ring)
        y_ring = splice_radius * np.sin(theta_ring)
        # Superficie lateral del anillo
        for i in range(len(theta_ring)-1):
            # Crear un pequeño rectángulo en 3D para cada segmento vertical
            x_pts = [x_ring[i], x_ring[i+1], x_ring[i+1], x_ring[i]]
            y_pts = [y_ring[i], y_ring[i+1], y_ring[i+1], y_ring[i]]
            z_pts = [splice_start, splice_start, splice_end, splice_end]
            fig3d.add_trace(go.Mesh3d(x=x_pts, y=y_pts, z=z_pts, alphahull=0, opacity=0.3, color='yellow', showlegend=False))

    # Espiral / Estribos
    tie_color = 'cornflowerblue'
    tie_width = max(2, (stirrup_diam/10.0) * 3)
    if "Espiral" in col_type or "Spiral" in col_type:
        # Espiral continua
        z_spiral = np.linspace(0, z_top, 300)
        angle = 2 * np.pi * z_spiral / (s_estribo_cm / 10)
        R_ext_cm = R_cm - recub
        x_spiral = R_ext_cm * np.cos(angle)
        y_spiral = R_ext_cm * np.sin(angle)
        fig3d.add_trace(go.Scatter3d(x=x_spiral, y=y_spiral, z=z_spiral, mode='lines',
                                     line=dict(color=tie_color, width=tie_width), name=_t('Espiral', 'Spiral')))
    else:
        # Estribos individuales
        R_ext_cm = R_cm - recub
        theta_ring = np.linspace(0, 2*np.pi, 50)
        x_ring = R_ext_cm * np.cos(theta_ring)
        y_ring = R_ext_cm * np.sin(theta_ring)
        n_rings = int(z_top / (s_estribo_cm/10)) + 1
        for k in range(n_rings):
            z_pos = k * (s_estribo_cm/10)
            if z_pos <= z_top:
                fig3d.add_trace(go.Scatter3d(x=x_ring, y=y_ring, z=np.full_like(theta_ring, z_pos),
                                             mode='lines', line=dict(color=tie_color, width=tie_width),
                                             name=_t('Estribo', 'Tie'), showlegend=(k==0)))

    fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='X (cm)', yaxis_title='Y (cm)', zaxis_title='L (cm)'),
                        margin=dict(l=0, r=0, b=0, t=0), height=450, showlegend=False, dragmode='turntable',
                        paper_bgcolor='#1a1a2e', scene_bgcolor='#1a1a2e')
    st.plotly_chart(fig3d, use_container_width=True)
    st.markdown("---")

    st.subheader(_t("Sección Transversal 2D", "2D Cross Section"))
    fig_s, ax_s = plt.subplots(figsize=(5,5))
    ax_s.set_facecolor('#1a1a2e'); fig_s.patch.set_facecolor('#1a1a2e')
    # Concreto
    ax_s.add_patch(plt.Circle((0,0), D_col/2, edgecolor='white', facecolor='#4a4a6a', lw=2))
    # Estribo / espiral
    ax_s.add_patch(plt.Circle((0,0), D_col/2 - recub, edgecolor='#00d4ff', facecolor='none', lw=2, linestyle='--'))
    # Varillas
    for (x_mm, y_mm) in bar_coords:
        x_cm = x_mm / 10
        y_cm = y_mm / 10
        ax_s.add_patch(plt.Circle((x_cm, y_cm), rebar_diam/20, color='#ff6b35', zorder=5))
    # Cotas
    ax_s.annotate('', xy=(-D_col/2, -D_col/2-5), xytext=(D_col/2, -D_col/2-5), arrowprops=dict(arrowstyle='<->', color='white'))
    ax_s.text(0, -D_col/2-8, f"Ø Exterior = {D_col:.0f} cm", color='white', ha='center', va='top')
    ax_s.set_xlim(-D_col/2-10, D_col/2+10)
    ax_s.set_ylim(-D_col/2-10, D_col/2+10)
    ax_s.set_aspect('equal')
    ax_s.axis('off')
    st.pyplot(fig_s)

    # DXF mejorado con empalmes y ganchos
    st.markdown("---")
    st.markdown(_t("#### 💾 Exportar Plano (AutoCAD DXF)", "#### 💾 Export Drawing (AutoCAD DXF)"))
    doc_dxf = ezdxf.new('R2010')
    doc_dxf.units = ezdxf.units.CM
    msp = doc_dxf.modelspace()
    # Capas
    for lay, col in [('CONCRETO',7), ('ESTRIBOS',4), ('VARILLAS',1), ('TEXTO',3), ('COTAS',2), ('EMPALME',6)]:
        if lay not in doc_dxf.layers:
            doc_dxf.layers.add(lay, color=col)

    # Planta (círculos)
    msp.add_circle((0,0), D_col/2, dxfattribs={'layer': 'CONCRETO'})
    msp.add_circle((0,0), D_col/2 - recub, dxfattribs={'layer': 'ESTRIBOS'})
    for (x_mm, y_mm) in bar_coords:
        x_cm = x_mm / 10
        y_cm = y_mm / 10
        msp.add_circle((x_cm, y_cm), rebar_diam/20, dxfattribs={'layer': 'VARILLAS'})
    msp.add_text(f"{n_barras_total} {rebar_type}", dxfattribs={'layer': 'TEXTO', 'height': 2, 'insert': (0, D_col/2+5)})
    
    # Vista en elevación (lado derecho)
    off_x = D_col + 20
    # Concreto (rectángulo)
    msp.add_lwpolyline([(off_x, 0), (off_x+D_col, 0), (off_x+D_col, L_col), (off_x, L_col), (off_x, 0)], 
                       close=True, dxfattribs={'layer': 'CONCRETO'})
    # Varillas longitudinales (líneas verticales) con ganchos en los extremos
    hook_len = 12 * rebar_diam / 10  # cm, gancho a 90° de longitud 12db
    for (x_mm, y_mm) in bar_coords:
        x_cm = x_mm / 10
        x_pos = off_x + x_cm + D_col/2
        # Varilla principal
        msp.add_line((x_pos, 0), (x_pos, L_col), dxfattribs={'layer': 'VARILLAS'})
        # Gancho inferior (90°)
        msp.add_line((x_pos, 0), (x_pos + hook_len, 0), dxfattribs={'layer': 'VARILLAS'})
        msp.add_line((x_pos, 0), (x_pos - hook_len, 0), dxfattribs={'layer': 'VARILLAS'})
        # Gancho superior (90°)
        msp.add_line((x_pos, L_col), (x_pos + hook_len, L_col), dxfattribs={'layer': 'VARILLAS'})
        msp.add_line((x_pos, L_col), (x_pos - hook_len, L_col), dxfattribs={'layer': 'VARILLAS'})
    
    # Estribos en elevación (líneas horizontales con ganchos)
    n_est_elev = int(L_col / s_estribo_cm) + 1
    for i in range(n_est_elev):
        z = i * s_estribo_cm
        if z <= L_col:
            # Línea horizontal del estribo
            msp.add_line((off_x, z), (off_x + D_col, z), dxfattribs={'layer': 'ESTRIBOS'})
            # Ganchos a 135° (simplificados como pequeñas líneas inclinadas)
            hook_ang = 45  # grados
            hook_len = 6 * stirrup_diam / 10  # cm
            # Extremo izquierdo
            x_left = off_x
            msp.add_line((x_left, z), (x_left - hook_len*math.cos(math.radians(hook_ang)), z + hook_len*math.sin(math.radians(hook_ang))), dxfattribs={'layer': 'ESTRIBOS'})
            msp.add_line((x_left, z), (x_left - hook_len*math.cos(math.radians(hook_ang)), z - hook_len*math.sin(math.radians(hook_ang))), dxfattribs={'layer': 'ESTRIBOS'})
            # Extremo derecho
            x_right = off_x + D_col
            msp.add_line((x_right, z), (x_right + hook_len*math.cos(math.radians(hook_ang)), z + hook_len*math.sin(math.radians(hook_ang))), dxfattribs={'layer': 'ESTRIBOS'})
            msp.add_line((x_right, z), (x_right + hook_len*math.cos(math.radians(hook_ang)), z - hook_len*math.sin(math.radians(hook_ang))), dxfattribs={'layer': 'ESTRIBOS'})
    
    # Zona de empalme (sombreado)
    if splice_zone_height > 0:
        # Relleno de la zona de empalme con líneas diagonales
        rect = msp.add_lwpolyline([(off_x, splice_start), (off_x+D_col, splice_start), (off_x+D_col, splice_end), (off_x, splice_end), (off_x, splice_start)],
                                   close=True, dxfattribs={'layer': 'EMPALME'})
        # Hatch (no es trivial en ezdxf, pero podemos agregar un texto)
        msp.add_text(_t(f"EMPALME Clase B\nLap = {lap_length_mm/10:.1f} cm", f"SPLICE Class B\nLap = {lap_length_mm/10:.1f} cm"),
                     dxfattribs={'layer': 'TEXTO', 'height': 2, 'insert': (off_x + D_col/2, (splice_start + splice_end)/2)})
    
    # Tabla de armado
    table_text = (f"ARMADO LONGITUDINAL:\n{n_barras_total} {rebar_type}\nAs = {Ast:.2f} cm²\n"
                  f"Lap = {lap_length_mm/10:.1f} cm (Clase B)\n\n"
                  f"ESTRIBOS:\n{stirrup_type} @ {s_estribo_cm:.1f} cm\n"
                  f"Long total = {long_estribos_m:.2f} m\nPeso = {peso_acero_estribos_kg:.1f} kg\n\n"
                  f"GANCHOS: 90° en longitudinales (12db)\n135° en estribos (6db)")
    msp.add_text(table_text, dxfattribs={'layer': 'TEXTO', 'height': 2, 'insert': (off_x + D_col + 10, L_col - 20)})
    
    out_stream = io.BytesIO()
    doc_dxf.write(out_stream)
    st.download_button(label=_t("Descargar Plano DXF (Planta + Elevación con Empalmes y Ganchos)", "Download DXF Drawing (Plan + Elevation with Splices and Hooks)"),
                       data=out_stream.getvalue(), file_name=f"Columna_Circular_D{D_col:.0f}cm.dxf", mime="application/dxf")

with tab3:
    st.subheader(_t("📦 Cantidades de Materiales y Costos", "📦 Material Quantities and Costs"))
    col_c1, col_c2, col_c3 = st.columns(3)
    col_c1.metric(_t("Volumen Concreto", "Concrete Volume"), f"{vol_concreto_m3:.3f} m³")
    col_c2.metric(_t("Peso Total Acero", "Total Steel Weight"), f"{peso_total_acero_kg:.1f} kg")
    col_c3.metric(_t("Ratio de Acero", "Steel Ratio"), f"{relacion_acero_kg_m3:.1f} kg/m³")
    
    st.markdown("---")
    st.markdown(_t("#### 📏 Detalle de Acero", "#### 📏 Steel Detail"))
    det_acero = [
        ("Longitudinal", f"{n_barras_total} {rebar_type}", f"{Ast:.2f} cm²", f"{peso_acero_long_kg:.1f} kg"),
        ("Estribos/Espiral", stirrup_type, f"@ {s_estribo_cm:.1f} cm", f"{peso_acero_estribos_kg:.1f} kg"),
        ("TOTAL", "", "", f"{peso_total_acero_kg:.1f} kg")
    ]
    df_acero = pd.DataFrame(det_acero, columns=[_t("Elemento","Element"), _t("Detalle","Detail"), _t("Cantidad","Quantity"), _t("Peso","Weight")])
    st.dataframe(df_acero, use_container_width=True, hide_index=True)
    
    st.markdown("---")
    st.markdown(_t("#### 🧱 Dosificación de Concreto (f'c = {:.1f} MPa)", "#### 🧱 Concrete Mix Design (f'c = {:.1f} MPa)").format(fc))
    mix_data = [
        (_t("Cemento", "Cement"), f"{total_cement_kg:.1f} kg", f"{bags_cement:.1f} bultos de {bag_kg:.0f} kg"),
        (_t("Agua", "Water"), f"{total_water_L:.0f} L", ""),
        (_t("Arena", "Sand"), f"{total_sand_kg:.1f} kg", f"{total_sand_kg/1600:.2f} m³ aprox."),
        (_t("Grava", "Gravel"), f"{total_gravel_kg:.1f} kg", f"{total_gravel_kg/1600:.2f} m³ aprox."),
    ]
    df_mix = pd.DataFrame(mix_data, columns=[_t("Material","Material"), _t("Total","Total"), _t("Observación","Note")])
    st.dataframe(df_mix, use_container_width=True, hide_index=True)

    if "apu_config" in st.session_state:
        st.markdown("---")
        st.markdown(_t("### 💰 Presupuesto Estimado (Promedio de Fuentes Regionales)", "### 💰 Estimated Budget (Regional Sources)"))
        apu = st.session_state.apu_config
        mon = apu.get("moneda", "$")
        
        if apu.get("usar_concreto_premezclado", False):
            precio_concreto_m3 = apu.get("precio_concreto_m3", 0)
            c_conc = vol_concreto_m3 * precio_concreto_m3
            c_cem = 0
            c_are = 0
            c_gra = 0
        else:
            pct_arena = apu.get("pct_arena_mezcla", 0.55)  # m³/m³
            pct_grava = apu.get("pct_grava_mezcla", 0.80)
            vol_arena = vol_concreto_m3 * pct_arena
            vol_grava = vol_concreto_m3 * pct_grava
            c_cem = bags_cement * apu.get("cemento", 0)
            c_are = vol_arena * apu.get("arena", 0)
            c_gra = vol_grava * apu.get("grava", 0)
            c_conc = 0
        
        c_ace = peso_total_acero_kg * apu.get("acero", 0)
        total_mat = c_cem + c_ace + c_are + c_gra + c_conc
        
        total_dias_mo = (peso_total_acero_kg * 0.04) + (vol_concreto_m3 * 0.4)
        costo_mo = total_dias_mo * apu.get("costo_dia_mo", 69333.33)
        
        costo_directo = total_mat + costo_mo
        herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
        aiu = costo_directo * apu.get("pct_aui", 0.30)
        utilidad = costo_directo * apu.get("pct_util", 0.05)
        iva = utilidad * apu.get("iva", 0.19)
        total_proyecto = costo_directo + herramienta + aiu + iva
        
        data_apu = {
            "Item": [_t("Cemento (bultos)", "Cement (bags)"), _t("Acero (kg)", "Steel (kg)"), 
                     _t("Arena (m³)", "Sand (m³)"), _t("Grava (m³)", "Gravel (m³)"),
                     _t("Concreto Premezclado (m³)", "Ready-mix Concrete (m³)") if apu.get("usar_concreto_premezclado", False) else _t("Concreto (preparado)", "Concrete (prepared)"),
                     _t("Mano de Obra (días)", "Labor (days)"), _t("Herramienta Menor", "Minor Tools"), 
                     _t("A.I.U.", "A.I.U."), _t("IVA s/Utilidad", "VAT on Profit"), _t("TOTAL PRESUPUESTO", "TOTAL BUDGET")],
            "Cantidad": [f"{bags_cement:.1f}" if not apu.get("usar_concreto_premezclado", False) else "-", 
                         f"{peso_total_acero_kg:.1f}",
                         f"{vol_arena:.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                         f"{vol_grava:.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                         f"{vol_concreto_m3:.2f}" if apu.get("usar_concreto_premezclado", False) else "-",
                         f"{total_dias_mo:.2f}",
                         f"{apu.get('pct_herramienta', 0.05)*100:.1f}% MO",
                         f"{apu.get('pct_aui', 0.3)*100:.1f}% CD",
                         f"{apu.get('iva', 0.19)*100:.1f}% Util", ""],
            f"Subtotal [{mon}]": [f"{c_cem:,.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                                  f"{c_ace:,.2f}",
                                  f"{c_are:,.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                                  f"{c_gra:,.2f}" if not apu.get("usar_concreto_premezclado", False) else "-",
                                  f"{c_conc:,.2f}" if apu.get("usar_concreto_premezclado", False) else "-",
                                  f"{costo_mo:,.2f}",
                                  f"{herramienta:,.2f}",
                                  f"{aiu:,.2f}",
                                  f"{iva:,.2f}",
                                  f"**{total_proyecto:,.2f}**"]
        }
        st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)
        
        # Excel APU
        output_excel = io.BytesIO()
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            df_export = pd.DataFrame({
                "Item": [_t("Cemento", "Cement"), _t("Acero", "Steel"), _t("Arena", "Sand"), _t("Grava", "Gravel"), _t("Mano de Obra", "Labor")],
                "Cantidad": [bags_cement if not apu.get("usar_concreto_premezclado", False) else 0,
                             peso_total_acero_kg,
                             vol_arena if not apu.get("usar_concreto_premezclado", False) else 0,
                             vol_grava if not apu.get("usar_concreto_premezclado", False) else 0,
                             total_dias_mo],
                "Unidad": [apu.get('cemento', 0), apu.get('acero', 0), apu.get('arena', 0), apu.get('grava', 0), apu.get('costo_dia_mo', 69333.33)]
            })
            if apu.get("usar_concreto_premezclado", False):
                df_pre = pd.DataFrame({"Item": ["Concreto Premezclado"], "Cantidad": [vol_concreto_m3], "Unidad": [apu.get("precio_concreto_m3", 0)]})
                df_export = pd.concat([df_export, df_pre], ignore_index=True)
            df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
            df_export.to_excel(writer, index=False, sheet_name='APU')
            workbook = writer.book
            worksheet = writer.sheets['APU']
            money_fmt = workbook.add_format({'num_format': '#,##0.00'})
            bold = workbook.add_format({'bold': True})
            worksheet.set_column('A:A', 25)
            worksheet.set_column('B:D', 15, money_fmt)
            row = len(df_export) + 1
            worksheet.write(row, 0, _t("Costo Directo (CD)", "Direct Cost (DC)"), bold)
            worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
            row += 1
            worksheet.write(row, 0, _t("Herramienta Menor", "Minor Tools"), bold)
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_herramienta", 0.05)}', money_fmt)
            row += 1
            worksheet.write(row, 0, _t("A.I.U.", "A.I.U."), bold)
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui", 0.30)}', money_fmt)
            row += 1
            worksheet.write(row, 0, _t("IVA s/ Utilidad", "VAT on Profit"), bold)
            worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_util", 0.05)}*{apu.get("iva", 0.19)}', money_fmt)
            row += 1
            worksheet.write(row, 0, _t("TOTAL PRESUPUESTO", "TOTAL BUDGET"), bold)
            worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
            
        output_excel.seek(0)
        st.download_button(label=_t("📥 Descargar Presupuesto Excel (.xlsx)", "📥 Download Budget Excel (.xlsx)"), 
                           data=output_excel, file_name=f"APU_ColumnaCirc_D{D_col:.0f}.xlsx", 
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info(_t("💡 Ve a la página 'APU Mercado' en la barra lateral para calcular presupuestos.", "💡 Go to the 'Market APU' page in the sidebar to calculate budgets."))

    st.markdown("---")
    st.markdown(_t("#### 📄 Generar Memoria de Cálculo Completa (DOCX)", "#### 📄 Generate Complete Calculation Report (DOCX)"))
    if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
        doc_word = Document()
        doc_word.add_heading(_t(f"Memoria de Cálculo — Columna Circular Ø{D_col:.0f} cm", f"Calculation Report — Circular Column Ø{D_col:.0f} cm"), 0)
        doc_word.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
        doc_word.add_paragraph(f"{_t('Norma Utilizada:', 'Code Used:')} {norma_sel}")
        doc_word.add_paragraph(f"{_t('Nivel Sísmico:', 'Seismic Level:')} {st.session_state.get('v_nivel_sis', 'N/A')}")
        
        doc_word.add_heading(_t("1. Materiales", "1. Materials"), level=1)
        doc_word.add_paragraph(f"f'c = {fc:.1f} MPa | fy = {fy:.0f} MPa | Es = {Es:.0f} MPa")
        doc_word.add_paragraph(f"Ec = {4700*math.sqrt(fc):.0f} MPa | β₁ = {beta_1:.3f}")
        
        doc_word.add_heading(_t("2. Geometría y Armadura", "2. Geometry and Reinforcement"), level=1)
        doc_word.add_paragraph(f"{_t('Diámetro D:', 'Diameter D:')} {D_col:.0f} cm | {_t('Altura L:', 'Height L:')} {L_col:.0f} cm | {_t('Recubrimiento:', 'Cover:')} {recub:.1f} cm")
        doc_word.add_paragraph(f"{_t('Varillas longitudinales:', 'Longitudinal rebars:')} {n_barras_total} {rebar_type} (As = {Ast:.2f} cm²)")
        doc_word.add_paragraph(f"{_t('Estribos / Espiral:', 'Ties / Spiral:')} {stirrup_type} @ {s_estribo_cm:.1f} cm")
        doc_word.add_paragraph(f"{_t('Longitud de desarrollo ld:', 'Development length ld:')} {ld_mm/10:.1f} cm")
        doc_word.add_paragraph(f"{_t('Longitud de empalme Clase B:', 'Splice length Class B:')} {lap_length_mm/10:.1f} cm")
        
        doc_word.add_heading(_t("3. Solicitaciones", "3. Demands"), level=1)
        doc_word.add_paragraph(f"Mu = {M_u_input:.2f} kN·m | Pu = {P_u_input:.2f} kN")
        
        doc_word.add_heading(_t("4. Resultados del Diagrama P-M", "4. P-M Diagram Results"), level=1)
        doc_word.add_paragraph(f"{_t('Carga axial nominal máxima Pn,máx:', 'Maximum nominal axial load Pn,max:')} {Pn_max:.1f} kN")
        doc_word.add_paragraph(f"{_t('Resistencia de diseño φPn,máx:', 'Design strength φPn,max:')} {phi_Pn_max_disp:.1f} kN")
        doc_word.add_paragraph(f"{_t('Cuantía de acero ρ:', 'Steel ratio ρ:')} {cuantia:.2f}% ({_t('rango', 'range')} {code['rho_min']:.1f}% – {code['rho_max']:.1f}%)")
        
        # Insertar gráfico P-M
        img_buffer.seek(0)
        doc_word.add_picture(img_buffer, width=Inches(5))
        
        doc_word.add_heading(_t("5. Verificación", "5. Verification"), level=1)
        if ok_design:
            doc_word.add_paragraph(_t("✅ El punto de diseño (Mu, Pu) se encuentra dentro de la curva de resistencia.", "✅ The design point (Mu, Pu) lies inside the strength curve."))
        else:
            doc_word.add_paragraph(_t("❌ El punto de diseño (Mu, Pu) está fuera de la curva de resistencia. Se debe aumentar la sección o el acero.", "❌ The design point (Mu, Pu) is outside the strength curve. Increase section or steel."))
        
        doc_word.add_heading(_t("6. Cantidades de Materiales", "6. Material Quantities"), level=1)
        doc_word.add_paragraph(f"{_t('Volumen de concreto:', 'Concrete volume:')} {vol_concreto_m3:.3f} m³")
        doc_word.add_paragraph(f"{_t('Cemento:', 'Cement:')} {total_cement_kg:.1f} kg ({bags_cement:.1f} bultos de {bag_kg:.0f} kg)")
        doc_word.add_paragraph(f"{_t('Agua:', 'Water:')} {total_water_L:.0f} L")
        doc_word.add_paragraph(f"{_t('Arena:', 'Sand:')} {total_sand_kg:.1f} kg")
        doc_word.add_paragraph(f"{_t('Grava:', 'Gravel:')} {total_gravel_kg:.1f} kg")
        doc_word.add_paragraph(f"{_t('Acero longitudinal:', 'Longitudinal steel:')} {peso_acero_long_kg:.1f} kg")
        doc_word.add_paragraph(f"{_t('Acero de estribos:', 'Tie steel:')} {peso_acero_estribos_kg:.1f} kg")
        doc_word.add_paragraph(f"{_t('Acero total:', 'Total steel:')} {peso_total_acero_kg:.1f} kg")
        
        f_io = io.BytesIO()
        doc_word.save(f_io)
        f_io.seek(0)
        st.download_button(label=_t("Descargar Memoria DOCX", "Download DOCX Report"), data=f_io, 
                           file_name=f"Memoria_ColumnaCircular_{D_col:.0f}cm.docx", 
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")