import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Mampostería y Morteros", "Masonry and Mortars"), layout="wide")
st.title(_t("Mampostería y Dosificación de Morteros", "Masonry and Mortar Dosing"))
st.markdown(_t("Herramienta integral para el cálculo de cantidades de materiales en tabiquería/mampostería (ladrillos y mortero) y la dosificación exacta de mezclas según el volumen requerido.", 
               "Comprehensive tool for calculating partition/masonry materials (bricks and mortar) and exact mortar mix dosing based on required volume."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CONFIGURACIÓN GENERAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙️ Configuración Global", "⚙️ Global Settings"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = "NSR-10 (Colombia)"

norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Normativa Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)

# ─────────────────────────────────────────────
# BASE DE DATOS DE LADRILLOS (MULTINORMA)
# ─────────────────────────────────────────────
BRICK_DB = {
    "NSR-10 (Colombia)": {
        "Ladrillo Común Macizo (25x12x6)": {"L":25.0, "W":12.0, "H":6.0},
        "Ladrillo Portante (29x12x9)":     {"L":29.0, "W":12.0, "H":9.0},
        "Bloque Arcilla N5 (30x12x20)":    {"L":30.0, "W":12.0, "H":20.0},
        "Bloque Cemento (40x20x20)":       {"L":40.0, "W":20.0, "H":20.0},
    },
    "ACI 318-25 (EE.UU.)": {
        "Standard Brick (19.4x9.2x5.7)":   {"L":19.4, "W":9.2, "H":5.7},
        "Modular Brick (19.4x9.2x5.7)":    {"L":19.4, "W":9.2, "H":5.7},
        "Utility Brick (29.5x9.2x9.2)":    {"L":29.5, "W":9.2, "H":9.2},
    },
    "ACI 318-19 (EE.UU.)": {
        "Standard Brick (19.4x9.2x5.7)":   {"L":19.4, "W":9.2, "H":5.7},
        "Modular Brick (19.4x9.2x5.7)":    {"L":19.4, "W":9.2, "H":5.7},
        "Utility Brick (29.5x9.2x9.2)":    {"L":29.5, "W":9.2, "H":9.2},
    },
    "ACI 318-14 (EE.UU.)": {
        "Standard Brick (19.4x9.2x5.7)":   {"L":19.4, "W":9.2, "H":5.7},
        "Modular Brick (19.4x9.2x5.7)":    {"L":19.4, "W":9.2, "H":5.7},
        "Utility Brick (29.5x9.2x9.2)":    {"L":29.5, "W":9.2, "H":9.2},
    },
    "NEC-SE-HM (Ecuador)": {
        "Ladrillo Mambrón (25x12x7)":      {"L":25.0, "W":12.0, "H":7.0},
        "Bloque Hormigón (40x20x20)":      {"L":40.0, "W":20.0, "H":20.0},
    },
    "E.060 (Perú)": {
        "King Kong 18 Huecos (24x13x9)":   {"L":24.0, "W":13.0, "H":9.0},
        "Pandereta (23x11x9)":             {"L":23.0, "W":11.0, "H":9.0},
        "Macizo (24x11.5x5.5)":            {"L":24.0, "W":11.5, "H":5.5},
    },
    "NTC-EM (México)": {
        "Rojo Recocido (24x12x5)":         {"L":24.0, "W":12.0, "H":5.0},
        "Tabicón (28x14x7)":               {"L":28.0, "W":14.0, "H":7.0},
        "Block Hormigón (40x20x15)":       {"L":40.0, "W":15.0, "H":20.0},
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "Bloque Arcilla (25x15x10)":       {"L":25.0, "W":15.0, "H":10.0},
        "Bloque Concreto (40x20x15)":      {"L":40.0, "W":15.0, "H":20.0},
    },
    "NB 1225001-2020 (Bolivia)": {
        "Ladrillo Gambote (25x12x6)":      {"L":25.0, "W":12.0, "H":6.0},
        "Ladrillo 6 Huecos (24x15x10)":    {"L":24.0, "W":15.0, "H":10.0},
    },
    "CIRSOC 201-2025 (Argentina)": {
        "Ladrillo Común (24x11.5x5.5)":    {"L":24.0, "W":11.5, "H":5.5},
        "Ladrillo Hueco 8 (33x8x18)":      {"L":33.0, "W":8.0, "H":18.0},
        "Ladrillo Hueco 12 (33x12x18)":    {"L":33.0, "W":12.0, "H":18.0},
        "Ladrillo Hueco 18 (33x18x18)":    {"L":33.0, "W":18.0, "H":18.0},
    }
}

# Obtener diccionario de ladrillos para la norma seleccionada
if "ACI" in norma_sel:
    current_dict = BRICK_DB["ACI 318-25 (EE.UU.)"]  # usar cualquier versión ACI, todas iguales
elif norma_sel in BRICK_DB:
    current_dict = BRICK_DB[norma_sel]
else:
    current_dict = BRICK_DB["NSR-10 (Colombia)"]  # fallback

lista_ladrillos = list(current_dict.keys()) + [_t("Típico Personalizado...", "Custom Size...")]

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES
# ─────────────────────────────────────────────
def mix_for_fc(fc):
    """Placeholder para futura integración con APU (no usada directamente aquí)."""
    return (350, 200, 800, 1000)  # cemento_kg, agua_L, arena_kg, grava_kg

# ─────────────────────────────────────────────
# T1: CÁLCULO DE TABIQUE Y MAMPOSTERÍA
# ─────────────────────────────────────────────
with st.expander(_t("🧱 1. Cantidades de Mampostería (Ladrillos y Juntas)", "🧱 1. Masonry Wall Quantities (Bricks and Joints)"), expanded=True):
    st.info(_t("📺 **Modo de uso:** Ingresa las dimensiones del muro a construir y el tipo de aparejo/ladrillo (filtrado automáticamente por la norma de tu país). El sistema calculará el número exacto de ladrillos por metro cuadrado, el total, y el volumen de mortero requerido para las juntas.", 
               "📺 **How to use:** Enter wall dimensions and brick type (filtered by active country code). The system will calculate bricks per square meter, total bricks, and joint mortar volume."))
    
    col1, col2, col3 = st.columns([1, 1.2, 1])
    
    with col1:
        st.write(_t("#### Geometría del Muro", "#### Wall Geometry"))
        L_muro = st.number_input(_t("Largo del Muro (L) [m]", "Wall Length (L) [m]"), 0.5, 100.0, st.session_state.get("mam_L", 5.0), 0.5, key="mam_L")
        H_muro = st.number_input(_t("Altura del Muro (H) [m]", "Wall Height (H) [m]"), 0.5, 10.0, st.session_state.get("mam_H", 2.5), 0.1, key="mam_H")
        area_muro = L_muro * H_muro
        st.markdown(f"**{_t('Área Bruta del Muro:', 'Gross Wall Area:')}** {area_muro:.2f} m²")
    
    with col2:
        st.write(_t("#### Dimensiones del Ladrillo", "#### Brick Dimensions"))
        ladrillo_sel = st.selectbox(_t("Tipo de Ladrillo (Norma Activa)", "Brick Type (Active Code)"), 
                                     lista_ladrillos,
                                     index=lista_ladrillos.index(st.session_state.get("mam_lad_sel", lista_ladrillos[0])) if st.session_state.get("mam_lad_sel", lista_ladrillos[0]) in lista_ladrillos else 0,
                                     key="mam_lad_sel")
        
        if ladrillo_sel == _t("Típico Personalizado...", "Custom Size..."):
            dimL_base = 24.0
            dimW_base = 12.0
            dimH_base = 6.0
        else:
            dimL_base = current_dict[ladrillo_sel]["L"]
            dimW_base = current_dict[ladrillo_sel]["W"]
            dimH_base = current_dict[ladrillo_sel]["H"]
            
        c2_1, c2_2, c2_3 = st.columns(3)
        with c2_1:
            dimL = st.number_input("Largo [cm]", 5.0, 60.0, st.session_state.get("mam_dimL", float(dimL_base)), 1.0, disabled=(ladrillo_sel != _t("Típico Personalizado...", "Custom Size...")), key="mam_dimL")
        with c2_2:
            dimW = st.number_input("Ancho [cm]", 5.0, 40.0, st.session_state.get("mam_dimW", float(dimW_base)), 1.0, disabled=(ladrillo_sel != _t("Típico Personalizado...", "Custom Size...")), key="mam_dimW")
        with c2_3:
            dimH = st.number_input("Alto [cm]", 3.0, 40.0, st.session_state.get("mam_dimH", float(dimH_base)), 1.0, disabled=(ladrillo_sel != _t("Típico Personalizado...", "Custom Size...")), key="mam_dimH")
        
        disp_opts = ["Soga (Grosor = Ancho)", "Tizón (Grosor = Largo)", "Canto (Grosor = Alto)"] if lang=="Español" else ["Stretcher (Thick = Width)", "Header (Thick = Length)", "Shiner (Thick = Height)"]
        disposicion = st.radio(_t("Disposición del muro:", "Wall Disposition:"), 
                               disp_opts, 
                               index=disp_opts.index(st.session_state.get("mam_disp", disp_opts[0])) if st.session_state.get("mam_disp", disp_opts[0]) in disp_opts else 0,
                               horizontal=True,
                               key="mam_disp")
        # Adapt dimensions to perspective
        if lang == "Español":
            if "Soga" in disposicion:
                ladrillo_frente_L = dimL
                ladrillo_frente_H = dimH
                espesor_muro = dimW
            elif "Tizón" in disposicion:
                ladrillo_frente_L = dimW
                ladrillo_frente_H = dimH
                espesor_muro = dimL
            else:
                ladrillo_frente_L = dimL
                ladrillo_frente_H = dimW
                espesor_muro = dimH
        else:
            if "Stretcher" in disposicion:
                ladrillo_frente_L = dimL
                ladrillo_frente_H = dimH
                espesor_muro = dimW
            elif "Header" in disposicion:
                ladrillo_frente_L = dimW
                ladrillo_frente_H = dimH
                espesor_muro = dimL
            else:
                ladrillo_frente_L = dimL
                ladrillo_frente_H = dimW
                espesor_muro = dimH
    
    with col3:
        st.write(_t("#### Juntas y Desperdicios", "#### Joints & Waste"))
        junta_h = st.number_input(_t("Espesor junta Hz. [cm]", "Horizontal joint [cm]"), 0.5, 3.0, st.session_state.get("mam_jh", 1.5), 0.1, key="mam_jh")
        junta_v = st.number_input(_t("Espesor junta Vt. [cm]", "Vertical joint [cm]"), 0.5, 3.0, st.session_state.get("mam_jv", 1.5), 0.1, key="mam_jv")
        desp_lad = st.number_input(_t("Desperdicio Ladrillos [%]", "Brick waste [%]"), 0.0, 20.0, st.session_state.get("mam_desp_l", 5.0), 1.0, key="mam_desp_l")
        desp_mor = st.number_input(_t("Desperdicio Mortero [%]", "Mortar waste [%]"), 0.0, 25.0, st.session_state.get("mam_desp_m", 10.0), 1.0, key="mam_desp_m")
        
    # --- CALCULO MATEMATICO TABIQUE ---
    # Convertir a metros
    L_m = ladrillo_frente_L / 100.0
    H_m = ladrillo_frente_H / 100.0
    W_m = espesor_muro / 100.0
    Jh_m = junta_h / 100.0
    Jv_m = junta_v / 100.0
    
    # Ladrillos por m² neto
    ladrillos_por_m2_neto = 1.0 / ((L_m + Jv_m) * (H_m + Jh_m))
    ladrillos_totales_netos = ladrillos_por_m2_neto * area_muro
    ladrillos_pedidos = math.ceil(ladrillos_totales_netos * (1.0 + desp_lad/100.0))
    
    # Volumen de mortero
    vol_muro_1m2 = 1.0 * 1.0 * W_m
    vol_ladrillos_1m2 = ladrillos_por_m2_neto * (L_m * H_m * W_m)
    vol_mortero_1m2_neto = vol_muro_1m2 - vol_ladrillos_1m2
    vol_mortero_total_neto = vol_mortero_1m2_neto * area_muro
    vol_mortero_pedido = vol_mortero_total_neto * (1.0 + desp_mor/100.0)
    
    st.markdown("---")
    res_1, res_2, res_3 = st.columns(3)
    res_1.metric(label=_t("Ladrillos por m² (S/Desp.)", "Bricks per m² (w/o waste)"), value=f"{ladrillos_por_m2_neto:.1f} uds")
    res_2.metric(label=_t("Ladrillos Totales (Con Desp.)", "Total Bricks (incl. waste)"), value=f"{ladrillos_pedidos} uds")
    res_3.metric(label=_t("Vol. Mortero p/ Muro (Con Desp.)", "Mortar Vol. (incl. waste)"), value=f"{vol_mortero_pedido:.3f} m³")

# ─────────────────────────────────────────────
# T2: DOSIFICACIÓN DE MORTEROS
# ─────────────────────────────────────────────
with st.expander(_t("💧 2. Dosificación y Diseño de Morteros (Cemento, Arena, Agua)", "💧 2. Mortar Dosing (Cement, Sand, Water)"), expanded=True):
    st.info(_t("📺 **Modo de uso:** Ingresa la proporción volumétrica deseada del mortero (ej. 1:3 para pegue muy resistente, o 1:4 para pegue estándar) y el volumen a producir. El sistema usa equivalencias teóricas empíricas para entregar bultos de cemento, m³ de arena y litros de agua.", 
               "📺 **How to use:** Enter mortar volumetric ratio (e.g. 1:3 for high strength, 1:4 for standard masonry) and volume. Returns bags of cement, sand volume, and water in liters."))
    
    # Datos empíricos para 1 m³ de mortero (cemento kg, arena m³, agua L)
    DOSIFICACIONES = {
        "1:1 (Revoques impermeables / muy ricos)": [908, 0.71, 250],
        "1:2 (Pañetes muros gruesos)": [610, 0.95, 250],
        "1:3 (Mampostería estructural / Pegue fuerte)": [454, 1.05, 250],
        "1:4 (Mampostería no estructural / Pañetes)": [364, 1.10, 250],
        "1:5 (Pegue baja resistencia / Plantillas)": [302, 1.13, 247],
        "1:6 (Morteros pobres)": [260, 1.16, 245],
        "1:8 (Rellenos)": [200, 1.18, 245]
    }
    DOSIFICACIONES_EN = {
        "1:1 (Waterproof renders)": [908, 0.71, 250],
        "1:2 (Thick plasters)": [610, 0.95, 250],
        "1:3 (Structural masonry / Strong joint)": [454, 1.05, 250],
        "1:4 (Non-structural masonry / Plasters)": [364, 1.10, 250],
        "1:5 (Low strength joints / Leveling)": [302, 1.13, 247],
        "1:6 (Poor mortars)": [260, 1.16, 245],
        "1:8 (Fillings)": [200, 1.18, 245]
    }
    
    dc1, dc2, dc3 = st.columns(3)
    with dc1:
        mezcla_opts = list(DOSIFICACIONES_EN.keys()) if lang=="English" else list(DOSIFICACIONES.keys())
        mezcla_sel = st.selectbox(_t("Proporción Cemento:Arena", "Cement:Sand Ratio"), 
                                  mezcla_opts, 
                                  index=mezcla_opts.index(st.session_state.get("mam_mezcla_sel", mezcla_opts[3])) if st.session_state.get("mam_mezcla_sel", mezcla_opts[3]) in mezcla_opts else 3,
                                  key="mam_mezcla_sel")
        if lang=="English":
            val_mezcla = DOSIFICACIONES_EN[mezcla_sel]
        else:
            val_mezcla = DOSIFICACIONES[mezcla_sel]
    with dc2:
        vol_producir = st.number_input(_t("Volumen de Mortero a Fabricar [m³]", "Mortar Volume to Produce [m³]"), 
                                       0.0, 1000.0, st.session_state.get("mam_vol_prod", float(vol_mortero_pedido)), 0.1, key="mam_vol_prod")
    with dc3:
        bulto_opts = ["50 kg", "42.5 kg"]
        bulto_kg = st.selectbox(_t("Peso Bulto Cemento", "Cement Bag Weight"), bulto_opts, 
                                 index=bulto_opts.index(st.session_state.get("mam_bulto_kg", "50 kg")),
                                 key="mam_bulto_kg")
        val_bulto_kg = 50.0 if bulto_kg == "50 kg" else 42.5
    
    if vol_producir > 0:
        cemento_kg_total = vol_producir * val_mezcla[0]
        bultos_cemento = math.ceil(cemento_kg_total / val_bulto_kg)
        arena_m3_total = vol_producir * val_mezcla[1]
        agua_litros_total = vol_producir * val_mezcla[2]
    else:
        cemento_kg_total = 0
        bultos_cemento = 0
        arena_m3_total = 0
        agua_litros_total = 0
    
    df_mort = pd.DataFrame([
        {_t("Material", "Material"): _t("Cemento", "Cement"), _t("Cantidad", "Quantity"): f"{bultos_cemento} {_t('Bultos', 'Bags')} ({cemento_kg_total:.1f} kg)"},
        {_t("Material", "Material"): _t("Arena de Peña/Sitio", "Sand"), _t("Cantidad", "Quantity"): f"{arena_m3_total:.2f} m³"},
        {_t("Material", "Material"): _t("Agua", "Water"), _t("Cantidad", "Quantity"): f"{agua_litros_total:.1f} {_t('Litros', 'Liters')}"},
    ])
    st.table(df_mort)

# ─────────────────────────────────────────────
# T3: PESO SUPERFICIAL DEL MURO
# ─────────────────────────────────────────────
with st.expander(_t("⚖️ 3. Peso Superficial de Muros de Mampostería", "⚖️ 3. Surface Weight of Masonry Walls"), expanded=False):
    st.info(_t("📺 **Modo de uso:** Calcula el peso real de 1 m² de muro (kg/m²) tomando en cuenta la densidad del material, el porcentaje de vacíos o huecos de la pieza, y el peso de las juntas de mortero.", 
               "📺 **How to use:** Calculates the actual weight of 1 m² of wall (kg/m²) accounting for material density, void percentage of the brick, and mortar joints weight."))
    
    col_w1, col_w2, col_w3 = st.columns(3)
    
    with col_w1:
        dens_opts = [
            _t("Arcilla Cocida (1800 kg/m³)", "Fired Clay (1800 kg/m³)"),
            _t("Concreto de Peso Normal (2200 kg/m³)", "Normal Weight Concrete (2200 kg/m³)"),
            _t("Concreto Liviano (1500 kg/m³)", "Lightweight Concrete (1500 kg/m³)"),
            _t("Sílico-Calcáreo (1900 kg/m³)", "Calcium-Silicate (1900 kg/m³)")
        ]
        densidad_pieza = st.selectbox(_t("Material de la Pieza (Densidad)", "Block Material (Density)"), 
                                      dens_opts,
                                      index=dens_opts.index(st.session_state.get("mam_densidad_pieza", dens_opts[0])) if st.session_state.get("mam_densidad_pieza", dens_opts[0]) in dens_opts else 0,
                                      key="mam_densidad_pieza")
        # Extraer densidad numérica
        if _t("Arcilla", "Fired Clay") in densidad_pieza:
            d_p_v = 1800.0
        elif _t("Normal", "Normal") in densidad_pieza:
            d_p_v = 2200.0
        elif _t("Liviano", "Lightweight") in densidad_pieza:
            d_p_v = 1500.0
        elif _t("Sílico", "Calcium") in densidad_pieza:
            d_p_v = 1900.0
        else:
            d_p_v = 1800.0
        
        huecos_pct = st.number_input(_t("Porcentaje de Huecos de la Pieza [%]", "Percentage of voids in piece [%]"), 0.0, 70.0, st.session_state.get("mam_huecos_pct", 30.0), 5.0, key="mam_huecos_pct")
        
    with col_w2:
        densidad_mortero = st.number_input(_t("Densidad del Mortero [kg/m³]", "Mortar Density [kg/m³]"), 1000.0, 3000.0, st.session_state.get("mam_dens_mor", 2000.0), 50.0, key="mam_dens_mor")
        st.write(f"**{_t('Volumen Neto Ladrillo/m²:', 'Net Brick Volume/m²:')}** {(vol_ladrillos_1m2 * (1 - huecos_pct/100.0)):.4f} m³/m²")
        st.write(f"**{_t('Volumen Mortero/m²:', 'Mortar Volume/m²:')}** {vol_mortero_1m2_neto:.4f} m³/m²")
        
    with col_w3:
        peso_ladrillo_m2 = (vol_ladrillos_1m2 * (1 - huecos_pct/100.0)) * d_p_v
        peso_mortero_m2 = vol_mortero_1m2_neto * densidad_mortero
        peso_total_m2 = peso_ladrillo_m2 + peso_mortero_m2
        peso_total_kn = peso_total_m2 * 9.81 / 1000.0
        
        st.markdown(f"### **{_t('Peso del Muro:', 'Wall Weight:')}** <span style='color:blue'>{peso_total_m2:.1f} kg/m²</span>", unsafe_allow_html=True)
        st.markdown(f"### **{_t('Equivalencia:', 'Equivalent:')}** <span style='color:green'>{peso_total_kn:.2f} kN/m²</span>", unsafe_allow_html=True)
        st.write(_t("*Usualmente empleado como Carga Muerta (D) en análisis estructural.*", "*Commonly used as Dead Load (D) in structural analysis.*"))

# ─────────────────────────────────────────────
# T4: DIAGRAMAS, DXF, APU Y EXPORTACIONES
# ─────────────────────────────────────────────
st.markdown("---")
tab_diag, tab_3d, tab_dxf, tab_mem, tab_apu = st.tabs([
    "📐 " + _t("Diagrama Ladrillo 2D", "2D Brick Diagram"),
    "🧊 " + _t("Muro 3D", "3D Wall"),
    "📏 " + _t("DXF (Planta/Elevación)", "DXF (Plan/Elevation)"),
    "📄 " + _t("Memoria DOCX", "DOCX Report"),
    "💰 " + _t("Presupuesto APU", "APU Budget")
])

with tab_diag:
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.set_facecolor('#1a1a2e'); fig.patch.set_facecolor('#1a1a2e')
    ax.add_patch(patches.Rectangle((0,0), dimL, dimH, linewidth=2, edgecolor='black', facecolor='#e2725b'))
    ax.annotate(f'L = {dimL} cm', xy=(dimL/2, -1.5), color='white', ha='center', fontsize=10)
    ax.annotate(f'H = {dimH} cm', xy=(-1.5, dimH/2), color='white', va='center', rotation=90, fontsize=10)
    ax.annotate(f'W = {dimW} cm ({_t("Profundidad", "Depth")})', xy=(dimL/2, dimH/2), color='white', ha='center', fontsize=9)
    ax.set_xlim(-5, dimL + 5)
    ax.set_ylim(-5, dimH + 5)
    ax.axis('off')
    st.pyplot(fig)
    # Guardar figura para memoria
    img_buffer = io.BytesIO()
    fig.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
    img_buffer.seek(0)

with tab_3d:
    st.write(_t("Visualización volumétrica del muro proyectado.", "Volumetric visualization of the projected wall."))
    # Crear una rejilla de ladrillos en la superficie (simple patrón)
    fig3d = go.Figure()
    # Cuerpo principal del muro (caja)
    X_m = L_muro
    Z_m = H_muro
    Y_m = espesor_muro / 100.0
    x_w = [0, X_m, X_m, 0, 0, X_m, X_m, 0]
    y_w = [0, 0, Y_m, Y_m, 0, 0, Y_m, Y_m]
    z_w = [0, 0, 0, 0, Z_m, Z_m, Z_m, Z_m]
    fig3d.add_trace(go.Mesh3d(x=x_w, y=y_w, z=z_w, alphahull=0, opacity=0.7, color='#e2725b', name=_t('Muro', 'Wall')))
    # Patrón de ladrillos (líneas en la cara frontal)
    # Calcular cantidad de ladrillos en ancho y alto
    n_lad_x = int(X_m / (ladrillo_frente_L/100 + junta_v/100)) + 1
    n_lad_z = int(Z_m / (ladrillo_frente_H/100 + junta_h/100)) + 1
    # Dibujar líneas verticales (juntas)
    for i in range(1, n_lad_x):
        x = i * (ladrillo_frente_L/100 + junta_v/100)
        if x <= X_m:
            fig3d.add_trace(go.Scatter3d(x=[x, x], y=[0, Y_m], z=[0, Z_m], mode='lines',
                                         line=dict(color='black', width=1), showlegend=False))
    # Líneas horizontales
    for j in range(1, n_lad_z):
        z = j * (ladrillo_frente_H/100 + junta_h/100)
        if z <= Z_m:
            fig3d.add_trace(go.Scatter3d(x=[0, X_m], y=[0, Y_m], z=[z, z], mode='lines',
                                         line=dict(color='black', width=1), showlegend=False))
    fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='L (m)', yaxis_title=_t('Espesor (m)', 'Thickness (m)'), zaxis_title='H (m)'),
                        margin=dict(l=0, r=0, b=0, t=0), height=450, showlegend=False, dragmode='turntable',
                        paper_bgcolor='#1a1a2e', scene_bgcolor='#1a1a2e')
    st.plotly_chart(fig3d, use_container_width=True)

with tab_dxf:
    st.markdown(_t("#### 💾 Exportar plano AutoCAD (DXF)", "#### 💾 Export AutoCAD drawing (DXF)"))
    try:
        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                 dxf_dim_horiz, dxf_dim_vert, dxf_rotulo,
                                 dxf_leyenda, dxf_rotulo_campos)
        _USE_H = True
    except ImportError:
        _USE_H = False
    doc_dxf = ezdxf.new('R2010')
    doc_dxf.units = ezdxf.units.M
    if _USE_H:
        dxf_setup(doc_dxf, 50)
        dxf_add_layers(doc_dxf)
    msp = doc_dxf.modelspace()
    for lay, col in [('CONCRETO',7), ('MURO',4), ('ACERO',1), ('TEXTO',3), ('COTAS',2)]:
        if lay not in doc_dxf.layers:
            doc_dxf.layers.add(lay, color=col)
    # Planta
    msp.add_lwpolyline([(0,0), (L_muro,0), (L_muro, espesor_muro/100), (0, espesor_muro/100), (0,0)],
                       close=True, dxfattribs={'layer':'MURO'})
    # Elevacion desplazada
    off_x = L_muro + 2
    msp.add_lwpolyline([(off_x,0), (off_x+L_muro,0), (off_x+L_muro, H_muro), (off_x, H_muro), (off_x,0)],
                       close=True, dxfattribs={'layer':'MURO'})
    # Juntas verticales
    n_vert = int(L_muro / (ladrillo_frente_L/100 + junta_v/100)) + 1
    for i in range(1, n_vert):
        x = off_x + i * (ladrillo_frente_L/100 + junta_v/100)
        if x <= off_x+L_muro:
            msp.add_line((x, 0), (x, H_muro), dxfattribs={'layer':'ACERO'})
    # Juntas horizontales
    n_hor = int(H_muro / (ladrillo_frente_H/100 + junta_h/100)) + 1
    for j in range(1, n_hor):
        z = j * (ladrillo_frente_H/100 + junta_h/100)
        if z <= H_muro:
            msp.add_line((off_x, z), (off_x+L_muro, z), dxfattribs={'layer':'ACERO'})
    if _USE_H:
        TH = 0.025 * 50
        dxf_dim_horiz(msp, off_x, off_x+L_muro, -0.4, f"L = {L_muro:.2f} m", 50)
        dxf_dim_vert(msp, off_x-0.5, 0, H_muro, f"H = {H_muro:.2f} m", 50)
        dxf_text(msp, L_muro/2, espesor_muro/100+0.2, "PLANTA", "EJES", h=TH, ha="center")
        dxf_text(msp, off_x+L_muro/2, H_muro+0.4, "ELEVACION", "EJES", h=TH, ha="center")
        dxf_leyenda(msp, off_x+L_muro+0.3, H_muro-0.3, [
            ("MURO", f"Muro {ladrillo_sel[:20]}"),
            ("ACERO", f"Juntas h={junta_h}cm v={junta_v}cm"),
        ], 50)
        _cam = dxf_rotulo_campos(f"Muro Mamposteria {L_muro:.1f}x{H_muro:.1f}m", norma_sel, "001")
        dxf_rotulo(msp, _cam, 0, -4.5, rot_w=9, rot_h=3, escala=50)
    else:
        msp.add_text(f"L = {L_muro:.2f} m", dxfattribs={'layer':'TEXTO','height':0.1,'insert':(off_x+L_muro/2, -0.2)})
        msp.add_text(f"H = {H_muro:.2f} m", dxfattribs={'layer':'TEXTO','height':0.1,'insert':(off_x-0.5, H_muro/2)})
    _out = io.StringIO()
    doc_dxf.write(_out)
    st.download_button(_t("📥 Descargar DXF", "📥 Download DXF"), data=_out.getvalue().encode('utf-8'),
                       file_name=f"Muro_{L_muro}x{H_muro}.dxf", mime="application/dxf")

with tab_mem:
    st.markdown(_t("#### 📄 Generar Memoria de Cálculo (DOCX)", "#### 📄 Generate Calculation Report (DOCX)"))
    if st.button(_t("Generar Memoria", "Generate Report")):
        doc = Document()
        doc.add_heading(_t(f"Memoria de Cálculo – Muro de Mampostería {L_muro:.2f} x {H_muro:.2f} m", 
                           f"Calculation Report – Masonry Wall {L_muro:.2f} x {H_muro:.2f} m"), 0)
        doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
        doc.add_paragraph(_t(f"Norma Activa: {norma_sel}", f"Active Code: {norma_sel}"))
        doc.add_heading(_t("1. Datos de entrada", "1. Input data"), level=1)
        doc.add_paragraph(f"{_t('Largo del muro:', 'Wall length:')} {L_muro:.2f} m")
        doc.add_paragraph(f"{_t('Altura del muro:', 'Wall height:')} {H_muro:.2f} m")
        doc.add_paragraph(f"{_t('Tipo de ladrillo:', 'Brick type:')} {ladrillo_sel}")
        doc.add_paragraph(f"{_t('Disposición:', 'Disposition:')} {disposicion}")
        doc.add_paragraph(f"{_t('Dimensiones:', 'Dimensions:')} L={dimL} cm, W={dimW} cm, H={dimH} cm")
        doc.add_paragraph(f"{_t('Juntas:', 'Joints:')} h={junta_h} cm, v={junta_v} cm")
        doc.add_paragraph(f"{_t('Desperdicios:', 'Waste:')} ladrillos {desp_lad:.0f}%, mortero {desp_mor:.0f}%")
        doc.add_heading(_t("2. Resultados", "2. Results"), level=1)
        doc.add_paragraph(f"{_t('Ladrillos por m²:', 'Bricks per m²:')} {ladrillos_por_m2_neto:.1f} uds")
        doc.add_paragraph(f"{_t('Ladrillos totales (con desperdicio):', 'Total bricks (with waste):')} {ladrillos_pedidos} uds")
        doc.add_paragraph(f"{_t('Volumen de mortero (con desperdicio):', 'Mortar volume (with waste):')} {vol_mortero_pedido:.3f} m³")
        doc.add_heading(_t("3. Dosificación del mortero", "3. Mortar mix"), level=1)
        doc.add_paragraph(f"{_t('Proporción:', 'Ratio:')} {mezcla_sel}")
        doc.add_paragraph(f"{_t('Cemento:', 'Cement:')} {bultos_cemento} {_t('bultos de', 'bags of')} {bulto_kg} ({cemento_kg_total:.1f} kg)")
        doc.add_paragraph(f"{_t('Arena:', 'Sand:')} {arena_m3_total:.2f} m³")
        doc.add_paragraph(f"{_t('Agua:', 'Water:')} {agua_litros_total:.0f} L")
        doc.add_heading(_t("4. Peso superficial del muro", "4. Surface weight of wall"), level=1)
        doc.add_paragraph(f"{_t('Peso total:', 'Total weight:')} {peso_total_m2:.1f} kg/m² ({peso_total_kn:.2f} kN/m²)")
        # Insertar figura del ladrillo
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(3))
        buf_doc = io.BytesIO()
        doc.save(buf_doc)
        buf_doc.seek(0)
        st.download_button(_t("📥 Descargar Memoria DOCX", "📥 Download DOCX Report"), data=buf_doc, 
                           file_name=f"Memoria_Muro_{L_muro}x{H_muro}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_apu:
    st.write(_t("Basado en el cálculo de Muro de Mampostería y la mezcla de Mortero seleccionada.", 
                "Based on the Masonry Wall calculation and selected Mortar mixture."))
    # Inputs APU
    st.markdown(_t("### 🏷️ Precios y Rendimientos", "### 🏷️ Prices and Yields"))
    col_a1, col_a2 = st.columns(2)
    with col_a1:
        precio_ladrillo = st.number_input(_t("Precio por unidad de ladrillo [moneda local]", "Price per brick [local currency]"), 
                                          value=st.session_state.get("precio_ladrillo", 1200.0), step=100.0, key="precio_ladrillo")
        rendimiento_dia = st.number_input(_t("Rendimiento de cuadrilla [m²/día]", "Crew output [m²/day]"), 
                                          value=st.session_state.get("rendimiento_dia", 12.0), step=1.0, key="rendimiento_dia")
        costo_dia_cuadrilla = st.number_input(_t("Costo diario de la cuadrilla (oficial+ayudante) [moneda]", "Daily crew cost (skilled+helper) [currency]"), 
                                               value=st.session_state.get("costo_dia_cuadrilla", 150000.0), step=10000.0, key="costo_dia_cuadrilla")
    with col_a2:
        if "apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mon = apu.get("moneda", "$")
            precio_cemento = apu.get("cemento", 0)
            precio_arena = apu.get("arena", 0)
            st.info(_t(f"Precios de cemento y arena obtenidos de la configuración APU global.", 
                       f"Cement and sand prices loaded from global APU configuration."))
        else:
            mon = "$"
            precio_cemento = st.number_input(_t("Precio del bulto de cemento [moneda]", "Cement bag price [currency]"), value=32000.0, step=1000.0, key="precio_cem_apu")
            precio_arena = st.number_input(_t("Precio del m³ de arena [moneda]", "Sand price per m³ [currency]"), value=70000.0, step=5000.0, key="precio_are_apu")
    
    # Cálculos de costos
    costo_ladrillos = ladrillos_pedidos * precio_ladrillo
    costo_cemento = bultos_cemento * precio_cemento
    costo_arena = arena_m3_total * precio_arena
    total_materiales = costo_ladrillos + costo_cemento + costo_arena
    
    dias_cuadrilla = area_muro / rendimiento_dia if rendimiento_dia > 0 else 0
    costo_mo = dias_cuadrilla * costo_dia_cuadrilla
    
    # Aplicar factores indirectos (tomados de APU si existen, sino valores por defecto)
    pct_herramienta = apu.get("pct_herramienta", 0.05) if "apu_config" in st.session_state else 0.05
    pct_aui = apu.get("pct_aui", 0.30) if "apu_config" in st.session_state else 0.30
    pct_util = apu.get("pct_util", 0.05) if "apu_config" in st.session_state else 0.05
    iva = apu.get("iva", 0.19) if "apu_config" in st.session_state else 0.19
    
    costo_directo = total_materiales + costo_mo
    herramienta = costo_mo * pct_herramienta
    aiu = costo_directo * pct_aui
    utilidad = costo_directo * pct_util
    iva_total = utilidad * iva
    total_proyecto = costo_directo + herramienta + aiu + iva_total
    
    # Gráfico de cantidades (barras)
    st.markdown(_t("#### 📊 Cantidades de materiales", "#### 📊 Material quantities"))
    fig_q, ax_q = plt.subplots(figsize=(6, 3))
    ax_q.bar([_t("Ladrillos (uds)", "Bricks (units)"), _t("Cemento (bultos)", "Cement (bags)"), _t("Arena (m³)", "Sand (m³)")],
             [ladrillos_pedidos, bultos_cemento, arena_m3_total], color=['#e2725b', '#4caf50', '#ff9800'])
    ax_q.set_ylabel(_t("Cantidad", "Quantity"))
    ax_q.grid(True, alpha=0.3)
    st.pyplot(fig_q)
    
    # Tabla de costos
    st.markdown(_t("#### 💰 Resumen de costos", "#### 💰 Cost summary"))
    data_apu = {
        _t("Item", "Item"): [
            _t("Ladrillos (uds)", "Bricks (units)"), 
            _t("Cemento (bultos)", "Cement (bags)"), 
            _t("Arena (m³)", "Sand (m³)"),
            _t("Mano de Obra (días)", "Labor (days)"),
            _t("Herramienta Menor", "Minor Tools"),
            _t("A.I.U.", "A.I.U."),
            _t("IVA s/Utilidad", "VAT on Profit"),
            _t("TOTAL", "TOTAL")
        ],
        _t("Cantidad", "Quantity"): [
            f"{ladrillos_pedidos}", 
            f"{bultos_cemento}", 
            f"{arena_m3_total:.2f}",
            f"{dias_cuadrilla:.2f}",
            f"{pct_herramienta*100:.1f}% MO",
            f"{pct_aui*100:.1f}% CD",
            f"{iva*100:.1f}% Util",
            ""
        ],
        f"Subtotal [{mon}]": [
            f"{costo_ladrillos:,.2f}",
            f"{costo_cemento:,.2f}",
            f"{costo_arena:,.2f}",
            f"{costo_mo:,.2f}",
            f"{herramienta:,.2f}",
            f"{aiu:,.2f}",
            f"{iva_total:,.2f}",
            f"**{total_proyecto:,.2f}**"
        ]
    }
    st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)
    st.metric(f"💎 Gran Total Proyecto [{mon}]", f"{total_proyecto:,.0f}")
    
    # Exportar Excel
    output_excel = io.BytesIO()
    with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
        df_export = pd.DataFrame({
            "Item": ["Ladrillos", "Cemento", "Arena", "Mano de Obra"],
            "Cantidad": [ladrillos_pedidos, bultos_cemento, arena_m3_total, dias_cuadrilla],
            "Unidad/Costo": [precio_ladrillo, precio_cemento, precio_arena, costo_dia_cuadrilla]
        })
        df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad/Costo"]
        df_export.to_excel(writer, index=False, sheet_name='APU Muro')
        # Agregar página de dosificación
        df_dosi = pd.DataFrame({
            "Material": ["Cemento", "Arena", "Agua"],
            "Cantidad": [f"{bultos_cemento} bultos ({cemento_kg_total:.1f} kg)", f"{arena_m3_total:.2f} m³", f"{agua_litros_total:.0f} L"]
        })
        df_dosi.to_excel(writer, sheet_name='Dosificación', index=False)
        workbook = writer.book
        worksheet = writer.sheets['APU Muro']
        money_fmt = workbook.add_format({'num_format': '#,##0.00'})
        worksheet.set_column('A:A', 25)
        worksheet.set_column('B:D', 15, money_fmt)
    output_excel.seek(0)
    st.download_button(_t("📥 Descargar Presupuesto Excel", "📥 Download Excel Budget"), data=output_excel, 
                       file_name=f"APU_Muro_{L_muro}x{H_muro}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")