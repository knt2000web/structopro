import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import json

# 
# IDIOMA GLOBAL Y NORMA
try:
    from normas_referencias import mostrar_referencias_norma
except ImportError:
    def mostrar_referencias_norma(*a, **kw): pass
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

mostrar_referencias_norma(norma_sel, "analisis_3d")

st.set_page_config(page_title=_t("StructMaster 3D", "StructMaster 3D"), layout="wide")
st.title(_t("StructMaster 3D - Análisis Matricial Espacial", "StructMaster 3D - Spatial Matrix Analysis"))
st.markdown(_t("Entorno Tridimensional Interactivo para Pórticos y Armaduras Espaciales (6 GDL por Nudo).", 
               "Interactive 3D Environment for Space Frames and Trusses (6 DOF per Node)."))

# 
# CODES (φ factores)
CODES = {
    "NSR-10 (Colombia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NSR-10"},
    "ACI 318-25 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-25"},
    "ACI 318-19 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-19"},
    "ACI 318-14 (EE.UU.)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"ACI 318-14"},
    "NEC-SE-HM (Ecuador)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NEC-SE-HM"},
    "E.060 (Perú)": {"phi_flex":0.90, "phi_shear":0.85, "phi_comp":0.70, "lambda":1.0, "ref":"E.060"},
    "NTC-EM (México)": {"phi_flex":0.85, "phi_shear":0.80, "phi_comp":0.70, "lambda":1.0, "ref":"NTC-EM"},
    "COVENIN 1753-2006 (Venezuela)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.70, "lambda":1.0, "ref":"COVENIN"},
    "NB 1225001-2020 (Bolivia)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"NB"},
    "CIRSOC 201-2025 (Argentina)": {"phi_flex":0.90, "phi_shear":0.75, "phi_comp":0.65, "lambda":1.0, "ref":"CIRSOC"},
}
code = CODES.get(norma_sel, CODES["NSR-10 (Colombia)"])
phi_flex = code["phi_flex"]
phi_shear = code["phi_shear"]
phi_comp = code["phi_comp"]
lam = code["lambda"]

# 
# FUNCIONES DE DISEÑO (similares a las del 2D, adaptadas a 3D)
def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_rho_min(fc, fy):
    return max(0.25*math.sqrt(fc)/fy, 1.4/fy)

def get_rho_max_beam(fc, fy, beta1):
    eps_cu = 0.003
    eps_t_min = 0.005
    return (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+eps_t_min))

def design_beam(Mu, Vu, b_cm, h_cm, fc, fy, recub=5):
    """Diseño de viga rectangular (flexión uniaxial + cortante)."""
    d_cm = h_cm - recub - 1  # asumiendo estribo 10 mm
    # Flexión
    Rn = (Mu * 1e6) / (phi_flex * b_cm * d_cm**2)
    disc = 1 - 2*Rn/(0.85*fc)
    if disc > 0:
        rho = (0.85*fc/fy)*(1 - math.sqrt(disc))
    else:
        rho = 0.0018
    rho_min = get_rho_min(fc, fy)
    rho_use = max(rho, rho_min)
    As = rho_use * b_cm * d_cm
    # Cortante
    Vc = 0.17 * lam * math.sqrt(fc) * b_cm * d_cm / 10  # kN
    phi_Vc = phi_shear * Vc
    if Vu > phi_Vc/2:
        Vs = Vu / phi_shear - Vc
        Av = 2 * 0.71  # #3 estribo, 2 ramas
        s = Av * fy * d_cm / Vs  # cm
        s = max(5, min(s, 60, d_cm/2))
    else:
        s = 0
    return As, s

def design_column(Pu, Mu_y, Mu_z, b_cm, h_cm, fc, fy, recub=5):
    """Diseño preliminar de columna rectangular con flexión biaxial."""
    # Acero mínimo por cuantía 1%
    Ag = b_cm * h_cm
    As_min = 0.01 * Ag
    # Verificación simplificada usando fórmula de interacción biaxial (Bresler)
    # Para simplificar, usamos una aproximación: solo chequeo de carga axial máxima
    # y advertencia de momento.
    phi = phi_comp
    Pn_max = 0.85 * fc * (Ag - As_min) + fy * As_min
    phi_Pn_max = phi * Pn_max
    # Verificar si la carga axial excede la capacidad
    ok = Pu <= phi_Pn_max
    # Verificar si el momento es grande (advertencia)
    if Mu_y > 0.1 * phi_Pn_max * 0.1 and Mu_z > 0.1 * phi_Pn_max * 0.1:
        warning = "Momento biaxial significativo - se requiere análisis detallado"
    else:
        warning = ""
    return As_min, ok, warning

# 
# INICIALIZACIÓN DE DATOS (estado)
if "nudos3d_df" not in st.session_state:
    st.session_state.nudos3d_df = pd.DataFrame([
        {"ID":1,"X":0.0,"Y":0.0,"Z":0.0}, {"ID":2,"X":4.0,"Y":0.0,"Z":0.0},
        {"ID":3,"X":4.0,"Y":0.0,"Z":3.0}, {"ID":4,"X":0.0,"Y":0.0,"Z":3.0},
        {"ID":5,"X":0.0,"Y":3.0,"Z":0.0}, {"ID":6,"X":4.0,"Y":3.0,"Z":0.0},
        {"ID":7,"X":4.0,"Y":3.0,"Z":3.0}, {"ID":8,"X":0.0,"Y":3.0,"Z":3.0}
    ])
if "barras3d_df" not in st.session_state:
    st.session_state.barras3d_df = pd.DataFrame([
        # Base
        {"N I":1, "N J":2, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":2, "N J":3, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":3, "N J":4, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":4, "N J":1, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        # Columnas
        {"N I":1, "N J":5, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":80, "J (cm⁴)":150, "Iy (cm⁴)":3000, "Iz (cm⁴)":3000,
         "b (cm)":40, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":2, "N J":6, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":80, "J (cm⁴)":150, "Iy (cm⁴)":3000, "Iz (cm⁴)":3000,
         "b (cm)":40, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":3, "N J":7, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":80, "J (cm⁴)":150, "Iy (cm⁴)":3000, "Iz (cm⁴)":3000,
         "b (cm)":40, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":4, "N J":8, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":80, "J (cm⁴)":150, "Iy (cm⁴)":3000, "Iz (cm⁴)":3000,
         "b (cm)":40, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        # Techo
        {"N I":5, "N J":6, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":6, "N J":7, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":7, "N J":8, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0},
        {"N I":8, "N J":5, "E (MPa)":2e5, "G (MPa)":7.7e4, "A (cm²)":50, "J (cm⁴)":100, "Iy (cm⁴)":800, "Iz (cm⁴)":2000,
         "b (cm)":30, "h (cm)":40, "wX (kN/m)":0, "wY (kN/m)":0, "wZ (kN/m)":0}
    ])
if "apoyos3d_df" not in st.session_state:
    st.session_state.apoyos3d_df = pd.DataFrame([
        {"Nudo":1,"Fx":True,"Fy":True,"Fz":True,"Mx":True,"My":True,"Mz":True},
        {"Nudo":2,"Fx":True,"Fy":True,"Fz":True,"Mx":False,"My":False,"Mz":False},
        {"Nudo":3,"Fx":True,"Fy":True,"Fz":True,"Mx":False,"My":False,"Mz":False},
        {"Nudo":4,"Fx":True,"Fy":True,"Fz":True,"Mx":False,"My":False,"Mz":False}
    ])
if "cargas3d_df" not in st.session_state:
    st.session_state.cargas3d_df = pd.DataFrame([
        {"Nudo":6, "FX (kN)":20.0, "FY (kN)":0.0, "FZ (kN)":0.0, "MX":0.0, "MY":0.0, "MZ":0.0},
        {"Nudo":7, "FX (kN)":20.0, "FY (kN)":0.0, "FZ (kN)":0.0, "MX":0.0, "MY":0.0, "MZ":0.0}
    ])
if "resultados3d" not in st.session_state:
    st.session_state.resultados3d = None

# 
# FUNCIÓN DE DIBUJO 3D
def plot_frame_3d(df_n, df_b, df_sup, df_load, U_amp=None, amp=20.0):
    """Dibuja la estructura con opción de deformada amplificada."""
    fig = go.Figure()
    # Nudos
    fig.add_trace(go.Scatter3d(
        x=df_n['X'], y=df_n['Z'], z=df_n['Y'],
        mode='markers+text', text=df_n['ID'], textposition="top center",
        marker=dict(size=6, color='blue', symbol='circle'), name="Nudos"
    ))
    # Barras
    for idx, b in df_b.iterrows():
        try:
            ni = df_n[df_n['ID'] == b['N I']].iloc[0]
            nj = df_n[df_n['ID'] == b['N J']].iloc[0]
            xi, yi, zi = ni['X'], ni['Y'], ni['Z']
            xj, yj, zj = nj['X'], nj['Y'], nj['Z']
            if U_amp is not None:
                # Obtener desplazamientos
                i_idx = id_to_idx.get(b['N I'], None) if 'id_to_idx' in globals() else None
                j_idx = id_to_idx.get(b['N J'], None)
                if i_idx is not None and j_idx is not None:
                    xi += U_amp[i_idx*6+0] * amp
                    yi += U_amp[i_idx*6+1] * amp
                    zi += U_amp[i_idx*6+2] * amp
                    xj += U_amp[j_idx*6+0] * amp
                    yj += U_amp[j_idx*6+1] * amp
                    zj += U_amp[j_idx*6+2] * amp
            fig.add_trace(go.Scatter3d(
                x=[xi, xj], y=[zi, zj], z=[yi, yj],
                mode='lines', line=dict(color='silver', width=4), name=f"B{idx+1}", hoverinfo="text",
                text=f"Barra {idx+1}: {b['N I']}-{b['N J']}"
            ))
        except: pass
    # Apoyos
    for _, s in df_sup.iterrows():
        try:
            n = df_n[df_n['ID'] == s['Nudo']].iloc[0]
            fig.add_trace(go.Scatter3d(
                x=[n['X']], y=[n['Z']], z=[n['Y']], mode='markers',
                marker=dict(size=8, color='red', symbol='diamond'), name="Apoyo"
            ))
        except: pass
    # Cargas (flechas)
    for _, c in df_load.iterrows():
        try:
            n = df_n[df_n['ID'] == c['Nudo']].iloc[0]
            # Vector resultante
            fx = c.get('FX (kN)', 0)
            fy = c.get('FY (kN)', 0)
            fz = c.get('FZ (kN)', 0)
            if fx != 0 or fy != 0 or fz != 0:
                fig.add_trace(go.Scatter3d(
                    x=[n['X']], y=[n['Z']], z=[n['Y']], mode='markers+text',
                    text=[f"{fx:.0f},{fy:.0f},{fz:.0f}"], textposition="bottom center",
                    textfont=dict(color="green", size=10), marker=dict(size=12, color='green', symbol='x'),
                    name="Carga"
                ))
        except: pass
    fig.update_layout(
        scene=dict(
            xaxis_title='X (m)', yaxis_title='Z (m)', zaxis_title='Y (m)',
            xaxis=dict(showgrid=True, gridcolor='gray'),
            yaxis=dict(showgrid=True, gridcolor='gray'),
            zaxis=dict(showgrid=True, gridcolor='gray'),
            aspectmode='data'
        ),
        margin=dict(r=10, l=10, b=10, t=30),
        plot_bgcolor='black', paper_bgcolor='#1e1e1e'
    )
    return fig

# 
# INTERFAZ PRINCIPAL
st.info("**Tip:** Todos los cambios en tablas se reflejan en tiempo real en el gráfico 3D.")
c_ui1, c_ui2 = st.columns([1, 1.2])

with c_ui1:
    st.subheader("1. Nodos (X, Y, Z)")
    st.session_state.nudos3d_df = st.data_editor(st.session_state.nudos3d_df, num_rows="dynamic", use_container_width=True)
    st.subheader("2. Barras (conectividad y propiedades)")
    st.session_state.barras3d_df = st.data_editor(st.session_state.barras3d_df, num_rows="dynamic", use_container_width=True)
    st.subheader("3. Apoyos (fijos o libres)")
    st.session_state.apoyos3d_df = st.data_editor(st.session_state.apoyos3d_df, num_rows="dynamic", use_container_width=True)
    st.subheader("4. Cargas nodales (fuerzas y momentos)")
    st.session_state.cargas3d_df = st.data_editor(st.session_state.cargas3d_df, num_rows="dynamic", use_container_width=True)

with c_ui2:
    # Gráfico en tiempo real
    fig_live = plot_frame_3d(st.session_state.nudos3d_df, st.session_state.barras3d_df,
                              st.session_state.apoyos3d_df, st.session_state.cargas3d_df)
    st.plotly_chart(fig_live, use_container_width=True, height=800)

st.markdown("---")
st.header(_t(" Análisis y diseño estructural 3D", " 3D Structural Analysis and Design"))

# 
# BOTÓN DE CÁLCULO
if st.button(_t(" Ejecutar análisis espacial", " Run spatial analysis"), type="primary"):
    with st.spinner(_t("Ensamblando matriz de rigidez espacial...", "Assembling spatial stiffness matrix...")):
        nodes = st.session_state.nudos3d_df.copy().dropna()
        elements = st.session_state.barras3d_df.copy().dropna()
        supports = st.session_state.apoyos3d_df.copy().dropna()
        loads = st.session_state.cargas3d_df.copy().dropna()

        # Validaciones
        if nodes.empty or elements.empty:
            st.error(_t("No hay nudos o elementos definidos.", "No nodes or elements defined."))
            st.stop()
        node_ids = nodes['ID'].tolist()
        id_to_idx = {nid: i for i, nid in enumerate(node_ids)}
        # Verificar IDs en barras, apoyos, cargas
        for _, e in elements.iterrows():
            if e['N I'] not in id_to_idx or e['N J'] not in id_to_idx:
                st.error(_t(f"Barra {e.name}: nudo I o J no existe.", f"Bar {e.name}: node I or J does not exist."))
                st.stop()
        for _, s in supports.iterrows():
            if s['Nudo'] not in id_to_idx:
                st.error(_t(f"Apoyo en nudo {s['Nudo']} no existe.", f"Support at node {s['Nudo']} does not exist."))
                st.stop()
        for _, c in loads.iterrows():
            if c['Nudo'] not in id_to_idx:
                st.error(_t(f"Carga en nudo {c['Nudo']} no existe.", f"Load at node {c['Nudo']} does not exist."))
                st.stop()

        num_nudos = len(nodes)
        num_gdl = num_nudos * 6  # 6 DOF por nudo
        K = np.zeros((num_gdl, num_gdl))
        F = np.zeros(num_gdl)

        # Cargas nodales
        for _, c in loads.iterrows():
            n_idx = id_to_idx[c['Nudo']]
            F[n_idx*6 + 0] += c.get('FX (kN)', 0.0)
            F[n_idx*6 + 1] += c.get('FY (kN)', 0.0)  # Y = elevación
            F[n_idx*6 + 2] += c.get('FZ (kN)', 0.0)  # Z = profundidad
            F[n_idx*6 + 3] += c.get('MX', 0.0)
            F[n_idx*6 + 4] += c.get('MY', 0.0)
            F[n_idx*6 + 5] += c.get('MZ', 0.0)

        # Ensamblaje de barras con cargas distribuidas
        element_internal = []  # guardará datos para diseño
        for idx, e in elements.iterrows():
            i_idx = id_to_idx[e['N I']]
            j_idx = id_to_idx[e['N J']]
            xi = nodes.loc[nodes['ID']==e['N I'], 'X'].values[0]
            yi = nodes.loc[nodes['ID']==e['N I'], 'Y'].values[0]
            zi = nodes.loc[nodes['ID']==e['N I'], 'Z'].values[0]
            xj = nodes.loc[nodes['ID']==e['N J'], 'X'].values[0]
            yj = nodes.loc[nodes['ID']==e['N J'], 'Y'].values[0]
            zj = nodes.loc[nodes['ID']==e['N J'], 'Z'].values[0]
            dx, dy, dz = xj-xi, yj-yi, zj-zi
            L = math.hypot(dx, dy, dz)
            if L == 0: continue

            # Propiedades
            E = e['E (MPa)'] * 1e3   # kN/m²
            G = e['G (MPa)'] * 1e3
            A = e['A (cm²)'] / 1e4   # m²
            J = e['J (cm⁴)'] / 1e8   # m⁴
            Iy = e['Iy (cm⁴)'] / 1e8
            Iz = e['Iz (cm⁴)'] / 1e8
            b_cm = e.get('b (cm)', 30)
            h_cm = e.get('h (cm)', 40)

            # Matriz de rigidez local (12x12) – estándar para elemento 3D
            a = E*A/L
            b = 12*E*Iz/(L**3)
            c = 6*E*Iz/(L**2)
            d = 4*E*Iz/L
            e_v = 2*E*Iz/L
            f = 12*E*Iy/(L**3)
            g = 6*E*Iy/(L**2)
            h = 4*E*Iy/L
            i = 2*E*Iy/L
            j = G*J/L

            k_local = np.zeros((12,12))
            # Axial
            k_local[0,0] = a; k_local[0,6] = -a
            k_local[6,0] = -a; k_local[6,6] = a
            # Shear y (local y)
            k_local[1,1] = b; k_local[1,5] = c; k_local[1,7] = -b; k_local[1,11] = c
            k_local[5,1] = c; k_local[5,5] = d; k_local[5,7] = -c; k_local[5,11] = e_v
            k_local[7,1] = -b; k_local[7,5] = -c; k_local[7,7] = b; k_local[7,11] = -c
            k_local[11,1] = c; k_local[11,5] = e_v; k_local[11,7] = -c; k_local[11,11] = d
            # Shear z (local z)
            k_local[2,2] = f; k_local[2,4] = -g; k_local[2,8] = -f; k_local[2,10] = -g
            k_local[4,2] = -g; k_local[4,4] = h; k_local[4,8] = g; k_local[4,10] = i
            k_local[8,2] = -f; k_local[8,4] = g; k_local[8,8] = f; k_local[8,10] = g
            k_local[10,2] = -g; k_local[10,4] = i; k_local[10,8] = g; k_local[10,10] = h
            # Torsión
            k_local[3,3] = j; k_local[3,9] = -j
            k_local[9,3] = -j; k_local[9,9] = j

            # Matriz de transformación (12x12) a global
            cx, cy, cz = dx/L, dy/L, dz/L
            # Vector auxiliar para definir el sistema local (evitar singularidad para miembros verticales)
            # Definir el vector z local = eje del elemento, luego construir el sistema de referencia
            # Según el método de vectores perpendiculares (opción simple: usar proyección en plano XY si no es vertical)
            if abs(cx) < 1e-6 and abs(cz) < 1e-6:  # elemento vertical en Y
                # Vector x' = (1,0,0) como referencia
                x_ref = np.array([1.0, 0.0, 0.0])
            else:
                x_ref = np.array([0.0, 1.0, 0.0])
            # Calcular vector perpendicular al elemento
            v = np.array([cx, cy, cz])
            # Vector y' = cross(x_ref, v) / norma
            y_prime = np.cross(x_ref, v)
            norm_y = np.linalg.norm(y_prime)
            if norm_y < 1e-6:
                # fallback: usar otro vector
                x_ref = np.array([0.0, 0.0, 1.0])
                y_prime = np.cross(x_ref, v)
                norm_y = np.linalg.norm(y_prime)
            y_prime /= norm_y
            # Vector z' = cross(v, y_prime)
            z_prime = np.cross(v, y_prime)
            z_prime /= np.linalg.norm(z_prime)
            # Matriz de rotación 3x3
            Lambda = np.vstack([v, y_prime, z_prime]).T
            # Matriz de transformación 12x12
            T = np.zeros((12,12))
            T[0:3,0:3] = Lambda
            T[3:6,3:6] = Lambda
            T[6:9,6:9] = Lambda
            T[9:12,9:12] = Lambda

            k_glob = T.T @ k_local @ T

            # Fuerzas de empotramiento por cargas distribuidas (uniforme en direcciones globales)
            # Convertimos wX, wY, wZ de la barra a fuerzas equivalentes nodales en coordenadas globales.
            # Usamos integración directa: fuerza total en cada dirección dividida en partes iguales.
            wX = e.get('wX (kN/m)', 0.0)
            wY = e.get('wY (kN/m)', 0.0)
            wZ = e.get('wZ (kN/m)', 0.0)
            F_fixed_local = np.zeros(12)
            if wX != 0:
                F_fixed_local[0] = wX * L / 2.0
                F_fixed_local[6] = wX * L / 2.0
            if wY != 0:
                F_fixed_local[1] = wY * L / 2.0
                F_fixed_local[7] = wY * L / 2.0
            if wZ != 0:
                F_fixed_local[2] = wZ * L / 2.0
                F_fixed_local[8] = wZ * L / 2.0
            # Rotar a global y restar del vector de fuerzas (signo: F_fixed = -K*U_emp)
            F_fixed_glob = T.T @ F_fixed_local
            dofs = [
                i_idx*6, i_idx*6+1, i_idx*6+2, i_idx*6+3, i_idx*6+4, i_idx*6+5,
                j_idx*6, j_idx*6+1, j_idx*6+2, j_idx*6+3, j_idx*6+4, j_idx*6+5
            ]
            for r in range(12):
                for c in range(12):
                    K[dofs[r], dofs[c]] += k_glob[r, c]
                F[dofs[r]] -= F_fixed_glob[r]

            # Guardar datos del elemento para postprocesamiento
            element_internal.append({
                "ID": idx+1,
                "N I": e['N I'], "N J": e['N J'],
                "L": L,
                "b_cm": b_cm, "h_cm": h_cm,
                "dofs": dofs,
                "T": T, "k_local": k_local,
                "F_fixed_local": F_fixed_local,
                "wX": wX, "wY": wY, "wZ": wZ
            })

        # Eliminación de grados de libertad fijos
        fixed_dofs = []
        for _, s in supports.iterrows():
            n_idx = id_to_idx[s['Nudo']]
            if s.get('Fx', False): fixed_dofs.append(n_idx*6 + 0)
            if s.get('Fy', False): fixed_dofs.append(n_idx*6 + 1)
            if s.get('Fz', False): fixed_dofs.append(n_idx*6 + 2)
            if s.get('Mx', False): fixed_dofs.append(n_idx*6 + 3)
            if s.get('My', False): fixed_dofs.append(n_idx*6 + 4)
            if s.get('Mz', False): fixed_dofs.append(n_idx*6 + 5)
        free_dofs = [d for d in range(num_gdl) if d not in fixed_dofs]

        if len(free_dofs) == 0:
            st.error(_t("No hay grados de libertad libres. Verifique los apoyos.", "No free DOFs. Check supports."))
            st.stop()

        Kff = K[np.ix_(free_dofs, free_dofs)]
        Ff = F[free_dofs]
        try:
            Uf = np.linalg.solve(Kff, Ff)
        except np.linalg.LinAlgError:
            st.error(_t("Matriz de rigidez singular. Verifique apoyos y rigideces.", "Singular stiffness matrix. Check supports and stiffnesses."))
            st.stop()

        U = np.zeros(num_gdl)
        U[free_dofs] = Uf
        R = K @ U

        # Fuerzas internas en cada elemento (coordenadas locales)
        for e in element_internal:
            u_glob = U[e['dofs']]
            u_loc = e['T'] @ u_glob
            f_loc = e['k_local'] @ u_loc + e['F_fixed_local']
            # Signos: N, Vy, Vz, Mx, My, Mz en extremo i y j según convención
            e['Ni'] = f_loc[0]; e['Vy_i'] = f_loc[1]; e['Vz_i'] = f_loc[2]; e['Mx_i'] = f_loc[3]; e['My_i'] = f_loc[4]; e['Mz_i'] = f_loc[5]
            e['Nj'] = -f_loc[6]; e['Vy_j'] = -f_loc[7]; e['Vz_j'] = -f_loc[8]; e['Mx_j'] = f_loc[9]; e['My_j'] = f_loc[10]; e['Mz_j'] = f_loc[11]

        st.session_state.resultados3d = {
            "U": U, "R": R, "elements": element_internal,
            "nodes": nodes, "id_to_idx": id_to_idx, "node_ids": node_ids,
            "num_gdl": num_gdl, "free_dofs": free_dofs, "fixed_dofs": fixed_dofs,
            "K": K, "F": F
        }
        st.success(_t("Análisis completado.", "Analysis completed."))

# 
# TABLA DE RESULTADOS Y EXPORTACIONES
if st.session_state.resultados3d is not None:
    res = st.session_state.resultados3d
    U = res["U"]
    R = res["R"]
    elements = res["elements"]
    nodes = res["nodes"]
    node_ids = res["node_ids"]
    id_to_idx = res["id_to_idx"]

    st.subheader(_t("Resultados del análisis", "Analysis results"))
    # Desplazamientos
    desp_data = []
    for nid in node_ids:
        idx = id_to_idx[nid]
        desp_data.append({
            "Nudo": nid,
            "dx (mm)": U[idx*6+0]*1000, "dy (mm)": U[idx*6+1]*1000, "dz (mm)": U[idx*6+2]*1000,
            "rx (rad)": U[idx*6+3], "ry (rad)": U[idx*6+4], "rz (rad)": U[idx*6+5]
        })
    st.dataframe(pd.DataFrame(desp_data), use_container_width=True)

    # Reacciones en apoyos
    react_data = []
    for _, s in st.session_state.apoyos3d_df.iterrows():
        nid = s['Nudo']
        if nid in id_to_idx:
            idx = id_to_idx[nid]
            react_data.append({
                "Nudo": nid,
                "Rx (kN)": R[idx*6+0] if s.get('Fx', False) else 0,
                "Ry (kN)": R[idx*6+1] if s.get('Fy', False) else 0,
                "Rz (kN)": R[idx*6+2] if s.get('Fz', False) else 0,
                "Mx (kN-m)": R[idx*6+3] if s.get('Mx', False) else 0,
                "My (kN-m)": R[idx*6+4] if s.get('My', False) else 0,
                "Mz (kN-m)": R[idx*6+5] if s.get('Mz', False) else 0
            })
    st.dataframe(pd.DataFrame(react_data), use_container_width=True)

    # Fuerzas internas
    forces_data = []
    for e in elements:
        forces_data.append({
            "Barra": e["ID"],
            "Axial I (kN)": e["Ni"], "Vy I (kN)": e["Vy_i"], "Vz I (kN)": e["Vz_i"],
            "Mx I (kN-m)": e["Mx_i"], "My I (kN-m)": e["My_i"], "Mz I (kN-m)": e["Mz_i"],
            "Axial J (kN)": e["Nj"], "Vy J (kN)": e["Vy_j"], "Vz J (kN)": e["Vz_j"],
            "Mx J (kN-m)": e["Mx_j"], "My J (kN-m)": e["My_j"], "Mz J (kN-m)": e["Mz_j"]
        })
    st.dataframe(pd.DataFrame(forces_data), use_container_width=True)

    # Visualización de deformada
    st.subheader(_t("Deformada amplificada (escala 20x)", "Deformed shape (scale 20x)"))
    fig_def = plot_frame_3d(nodes, st.session_state.barras3d_df,
                            st.session_state.apoyos3d_df, st.session_state.cargas3d_df,
                            U_amp=U, amp=20.0)
    fig_def.update_layout(title="Deformada (rojo) vs original (gris)")
    st.plotly_chart(fig_def, use_container_width=True, height=600)

    # Materiales para diseño
    col_mat1, col_mat2 = st.columns(2)
    with col_mat1:
        fc = st.number_input(_t("Resistencia concreto f'c [MPa]", "Concrete strength f'c [MPa]"), 15.0, 80.0, 21.0, 1.0, key="3d_fc")
        fy = st.number_input(_t("Fluencia acero fy [MPa]", "Steel yield fy [MPa]"), 200.0, 600.0, 420.0, 10.0, key="3d_fy")
        recub = st.number_input(_t("Recubrimiento [cm]", "Cover [cm]"), 2.0, 10.0, 5.0, 0.5, key="3d_rec")
    with col_mat2:
        qa_suelo = st.number_input(_t("Capacidad portante suelo qa [kN/m²]", "Soil bearing capacity qa [kN/m²]"), 10.0, 1000.0, 150.0, 10.0, key="3d_qa")
        st.markdown("---")
        st.markdown(_t("Se asumirá un peralte típico para zapatas de 0.5 m y armadura mínima.", "Typical footing depth 0.5 m and minimum reinforcement assumed."))

    # Diseño automático
    st.subheader(_t("Diseño automático de elementos", "Automated element design"))
    design_data = {"vigas": [], "columnas": []}
    for e in elements:
        # Determinar tipo: si la barra es aproximadamente horizontal o vertical
        # Usamos la dirección del vector local x (global)
        dx = nodes.loc[nodes['ID']==e["N I"], 'X'].values[0] - nodes.loc[nodes['ID']==e["N J"], 'X'].values[0]
        dy = nodes.loc[nodes['ID']==e["N I"], 'Y'].values[0] - nodes.loc[nodes['ID']==e["N J"], 'Y'].values[0]
        dz = nodes.loc[nodes['ID']==e["N I"], 'Z'].values[0] - nodes.loc[nodes['ID']==e["N J"], 'Z'].values[0]
        L = math.hypot(dx, dy, dz)
        # Ángulo con la horizontal (despreciando Z)
        ang = math.degrees(math.atan2(abs(dy), math.hypot(dx, dz)))
        is_beam = ang < 20  # si el ángulo con la horizontal es pequeño
        if is_beam:
            # Momento flector máximo (consideramos la mayor combinación de My y Mz)
            M_max = max(abs(e["My_i"]), abs(e["My_j"]), abs(e["Mz_i"]), abs(e["Mz_j"]))
            V_max = max(abs(e["Vy_i"]), abs(e["Vy_j"]), abs(e["Vz_i"]), abs(e["Vz_j"]))
            b = e["b_cm"]
            h = e["h_cm"]
            As, s = design_beam(M_max, V_max, b, h, fc, fy, recub)
            design_data["vigas"].append({
                "Barra": e["ID"],
                "Mmax (kN-m)": M_max,
                "Vmax (kN)": V_max,
                "b (cm)": b,
                "h (cm)": h,
                "As (cm²)": As,
                "s_estribos (cm)": s if s > 0 else "No requiere"
            })
        else:
            # Columna
            Pu = max(abs(e["Ni"]), abs(e["Nj"]))  # carga axial de compresión (tomamos la mayor)
            My_max = max(abs(e["My_i"]), abs(e["My_j"]))
            Mz_max = max(abs(e["Mz_i"]), abs(e["Mz_j"]))
            b = e["b_cm"]
            h = e["h_cm"]
            As_min, ok, warn = design_column(Pu, My_max, Mz_max, b, h, fc, fy, recub)
            design_data["columnas"].append({
                "Barra": e["ID"],
                "Pu (kN)": Pu,
                "My (kN-m)": My_max,
                "Mz (kN-m)": Mz_max,
                "b (cm)": b,
                "h (cm)": h,
                "As_min (cm²)": As_min,
                "Verificación": "OK" if ok else "Falla",
                "Observación": warn
            })

    # Mostrar tablas de diseño
    if design_data["vigas"]:
        st.markdown("#### Vigas")
        st.dataframe(pd.DataFrame(design_data["vigas"]), use_container_width=True)
    if design_data["columnas"]:
        st.markdown("#### Columnas")
        st.dataframe(pd.DataFrame(design_data["columnas"]), use_container_width=True)

    # Exportaciones
    st.markdown("---")
    st.subheader(_t("Exportaciones", "Exports"))
    col_e1, col_e2, col_e3, col_e4 = st.columns(4)

    # Excel
    with col_e1:
        if st.button(_t("Exportar a Excel", "Export to Excel")):
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                pd.DataFrame(desp_data).to_excel(writer, sheet_name="Desplazamientos", index=False)
                pd.DataFrame(react_data).to_excel(writer, sheet_name="Reacciones", index=False)
                pd.DataFrame(forces_data).to_excel(writer, sheet_name="Fuerzas_internas", index=False)
                if design_data["vigas"]:
                    pd.DataFrame(design_data["vigas"]).to_excel(writer, sheet_name="Diseño_vigas", index=False)
                if design_data["columnas"]:
                    pd.DataFrame(design_data["columnas"]).to_excel(writer, sheet_name="Diseño_columnas", index=False)
                # También exportar datos de entrada
                st.session_state.nudos3d_df.to_excel(writer, sheet_name="Nudos", index=False)
                st.session_state.barras3d_df.to_excel(writer, sheet_name="Barras", index=False)
                st.session_state.apoyos3d_df.to_excel(writer, sheet_name="Apoyos", index=False)
                st.session_state.cargas3d_df.to_excel(writer, sheet_name="Cargas", index=False)
            output.seek(0)
            st.download_button(_t("Descargar Excel", "Download Excel"), data=output, file_name="estructura_3d.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # DXF (planos 2D)
    with col_e2:
        if st.button(_t("Exportar a DXF", "Export to DXF")):
            try:
                from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                         dxf_rotulo, dxf_rotulo_campos)
                _USE_H15 = True
            except ImportError:
                _USE_H15 = False
            doc_dxf = ezdxf.new('R2010')
            doc_dxf.units = ezdxf.units.M
            if _USE_H15:
                dxf_setup(doc_dxf, 50); dxf_add_layers(doc_dxf)
            msp = doc_dxf.modelspace()
            for lay, col15 in [('ESTRUCTURA',3),('COTAS',2),('TEXTO',1)]:
                if lay not in doc_dxf.layers: doc_dxf.layers.add(lay, color=col15)
            # Planta (vista en XZ)
            x_vals15, z_vals15 = [], []
            for e in elements:
                ni = nodes[nodes['ID']==e["N I"]].iloc[0]
                nj = nodes[nodes['ID']==e["N J"]].iloc[0]
                msp.add_line((ni['X'], ni['Z']), (nj['X'], nj['Z']), dxfattribs={'layer':'ESTRUCTURA'})
                x_vals15 += [ni['X'], nj['X']]; z_vals15 += [ni['Z'], nj['Z']]
            # Alzado (vista XY)
            off_x = (max(x_vals15) - min(x_vals15) + 2) if x_vals15 else 10
            for e in elements:
                ni = nodes[nodes['ID']==e["N I"]].iloc[0]
                nj = nodes[nodes['ID']==e["N J"]].iloc[0]
                msp.add_line((ni['X']+off_x, ni['Y']), (nj['X']+off_x, nj['Y']), dxfattribs={'layer':'ESTRUCTURA'})
            if _USE_H15:
                _x0_15 = min(x_vals15) if x_vals15 else 0
                _z015 = min(z_vals15) if z_vals15 else 0
                _xw_15 = max(x_vals15)-_x0_15 if x_vals15 else 10
                dxf_text(msp, _x0_15, (max(z_vals15)+0.5 if z_vals15 else 1), f"StructMaster 3D – {norma_sel}", "EJES", h=0.025*50, ha="left")
                _cam15 = dxf_rotulo_campos("Analisis Estructural 3D", norma_sel, "001")
                dxf_rotulo(msp, _cam15, _x0_15, _z015-4, rot_w=max(_xw_15,10), rot_h=3, escala=50)
            out_dxf = io.StringIO()
            doc_dxf.write(out_dxf)
            st.download_button(_t("Descargar DXF","Download DXF"), data=out_dxf.getvalue().encode('utf-8'), file_name="estructura_3d.dxf", mime="application/dxf")

    # Memoria DOCX
    with col_e3:
        if st.button(_t("Memoria DOCX", "DOCX Report")):
            doc = Document()
            doc.add_heading(_t("Análisis estructural 3D", "3D Structural Analysis"), 0)
            doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
            doc.add_paragraph(_t(f"Norma aplicada: {norma_sel}", f"Applied code: {norma_sel}"))
            doc.add_heading(_t("Resultados de desplazamientos", "Displacement results"), level=1)
            # Tabla de desplazamientos
            table = doc.add_table(rows=1+len(desp_data), cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Nudo"
            hdr[1].text = "dx (mm)"
            hdr[2].text = "dy (mm)"
            hdr[3].text = "dz (mm)"
            for i, d in enumerate(desp_data):
                cells = table.rows[i+1].cells
                cells[0].text = str(d["Nudo"])
                cells[1].text = f"{d['dx (mm)']:.2f}"
                cells[2].text = f"{d['dy (mm)']:.2f}"
                cells[3].text = f"{d['dz (mm)']:.2f}"
            # Diseño
            if design_data["vigas"]:
                doc.add_heading(_t("Diseño de vigas", "Beam design"), level=1)
                v_table = doc.add_table(rows=1+len(design_data["vigas"]), cols=6)
                v_table.style = 'Table Grid'
                vhdr = v_table.rows[0].cells
                vhdr[0].text = "Barra"
                vhdr[1].text = "Mu (kN-m)"
                vhdr[2].text = "b (cm)"
                vhdr[3].text = "h (cm)"
                vhdr[4].text = "As (cm²)"
                vhdr[5].text = "s estribos (cm)"
                for i, v in enumerate(design_data["vigas"]):
                    cells = v_table.rows[i+1].cells
                    cells[0].text = str(v["Barra"])
                    cells[1].text = f"{v['Mmax (kN-m)']:.1f}"
                    cells[2].text = f"{v['b (cm)']:.0f}"
                    cells[3].text = f"{v['h (cm)']:.0f}"
                    cells[4].text = f"{v['As (cm²)']:.1f}"
                    cells[5].text = str(v["s_estribos (cm)"])
            if design_data["columnas"]:
                doc.add_heading(_t("Diseño de columnas", "Column design"), level=1)
                c_table = doc.add_table(rows=1+len(design_data["columnas"]), cols=7)
                c_table.style = 'Table Grid'
                chdr = c_table.rows[0].cells
                chdr[0].text = "Barra"
                chdr[1].text = "Pu (kN)"
                chdr[2].text = "My (kN-m)"
                chdr[3].text = "Mz (kN-m)"
                chdr[4].text = "b (cm)"
                chdr[5].text = "h (cm)"
                chdr[6].text = "As_min (cm²)"
                for i, c in enumerate(design_data["columnas"]):
                    cells = c_table.rows[i+1].cells
                    cells[0].text = str(c["Barra"])
                    cells[1].text = f"{c['Pu (kN)']:.1f}"
                    cells[2].text = f"{c['My (kN-m)']:.1f}"
                    cells[3].text = f"{c['Mz (kN-m)']:.1f}"
                    cells[4].text = f"{c['b (cm)']:.0f}"
                    cells[5].text = f"{c['h (cm)']:.0f}"
                    cells[6].text = f"{c['As_min (cm²)']:.1f}"
            buf_doc = io.BytesIO()
            doc.save(buf_doc)
            buf_doc.seek(0)
            st.download_button(_t("Descargar Memoria", "Download Report"), data=buf_doc, file_name="memoria_estructura_3d.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # APU
    with col_e4:
        if st.button(_t("APU (presupuesto)", " APU Budget")):
            if "apu_config" in st.session_state:
                apu = st.session_state.apu_config
                mon = apu.get("moneda", "$")
                # Estimación de volúmenes de concreto y acero
                vol_conc = 0.0
                peso_acero = 0.0
                # Zapatas (simplificado: un dado de 1x1x0.5 m por cada apoyo restringido en traslación)
                # Contar apoyos con restricción en Y (vertical)
                n_zapatas = sum(1 for _, s in st.session_state.apoyos3d_df.iterrows() if s.get('Fy', False))
                vol_conc += n_zapatas * 1.0 * 1.0 * 0.5
                # Acero de zapatas: aprox 10 kg/m³
                peso_acero += vol_conc * 10
                # Vigas y columnas
                for e in elements:
                    # Volumen de concreto
                    vol_conc += (e["b_cm"] * e["h_cm"] / 10000) * e["L"]
                    # Acero longitudinal: usar As de diseño (si está disponible) o mínimo
                    if is_beam:
                        # Buscar en design_data
                        v = next((v for v in design_data["vigas"] if v["Barra"] == e["ID"]), None)
                        As = v["As (cm²)"] if v else 0
                    else:
                        c = next((c for c in design_data["columnas"] if c["Barra"] == e["ID"]), None)
                        As = c["As_min (cm²)"] if c else 0
                    peso_acero += As * e["L"] * 0.785  # kg
                # Costos
                if apu.get("usar_concreto_premezclado", False):
                    costo_conc = vol_conc * apu.get("precio_concreto_m3", 0)
                else:
                    # Suponemos mezcla en sitio con proporciones por defecto (350 kg/m³)
                    costo_conc = vol_conc * 350 / 50 * apu.get("cemento", 0)
                costo_ace = peso_acero * apu.get("acero", 0)
                total_mat = costo_conc + costo_ace
                # Mano de obra estimada (días)
                dias = (peso_acero * 0.04) + (vol_conc * 0.4)
                costo_mo = dias * apu.get("costo_dia_mo", 69333.33)
                costo_directo = total_mat + costo_mo
                herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
                aiu = costo_directo * apu.get("pct_aui", 0.30)
                utilidad = costo_directo * apu.get("pct_util", 0.05)
                iva = utilidad * apu.get("iva", 0.19)
                total = costo_directo + herramienta + aiu + iva

                st.markdown(f"**{_t('Volumen de concreto estimado', 'Estimated concrete volume')}:** {vol_conc:.2f} m³")
                st.markdown(f"**{_t('Peso de acero estimado', 'Estimated steel weight')}:** {peso_acero:.1f} kg")
                st.markdown(f"**{_t('Costo total materiales', 'Total material cost')}:** {mon} {total_mat:,.2f}")
                st.markdown(f"**{_t('Costo total incluyendo MO e indirectos', 'Total cost incl. labor & overhead')}:** {mon} {total:,.2f}")
            else:
                st.info(_t("Configure precios en la página 'APU Mercado'.", "Set prices in 'Market APU'page."))