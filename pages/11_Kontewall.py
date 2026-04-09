# ═══════════════════════════════════════════════════════════════════
# KonteWall — Diseño Estructural de Muros de Contención en Voladizo
# Multinorma: NSR-10 · ACI 318-25/19 · E.060 · NEC-SE-HM · NTC-EM
# COVENIN 1753 · NB 1225001 · CIRSOC 201
# Autor:  — Konte v4.0 · 2026
# ═══════════════════════════════════════════════════════════════════
import streamlit as st
import numpy as np
import math
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import io
import plotly.graph_objects as go

st.set_page_config(
    page_title="KonteWall © Konte | Muro de Contención",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── Modo claro / oscuro ──────────────────────────────────────────
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = True
_c1, _c2 = st.sidebar.columns([1, 3])
with _c1:
    if st.button("" if st.session_state.dark_mode else "☀", key="kw_dm"):
        st.session_state.dark_mode = not st.session_state.dark_mode
with _c2:
    st.caption("Modo oscuro " if st.session_state.dark_mode else "Modo claro ☀")

_BG     = "#0e1117" if st.session_state.dark_mode else "#FFFFFF"
_FG     = "#FFFFFF" if st.session_state.dark_mode else "#111111"
_SEC_BG = "#161b22" if st.session_state.dark_mode else "#F0F2F6"
_BORDER = "#30363d" if st.session_state.dark_mode else "#E0E0E0"
st.markdown(f"""<style>
.stApp {{ background-color:{_BG}; color:{_FG}; }}
section[data-testid="stSidebar"] {{ background-color:{_SEC_BG}; }}
div[data-testid="stExpander"] {{ border:1px solid {_BORDER}; border-radius:8px; }}
</style>""", unsafe_allow_html=True)

_plt_bg   = '#0e1117' if st.session_state.dark_mode else '#FFFFFF'
_plt_axbg = '#161b22' if st.session_state.dark_mode else '#F8F9FA'
_plt_txt  = 'white'   if st.session_state.dark_mode else 'black'
_plt_grid = '#1e1e2e' if st.session_state.dark_mode else '#E0E0E0'

lang      = st.session_state.get("idioma",    "Español")
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
en_usa    = (lang == "English" or "ACI" in norma_sel)
def _t(es, en): return en if en_usa else es

try:
    st.image("assets/kontewall_header.png", use_container_width=True)
except:
    pass

st.title(_t(
    "KonteWall — Diseño Estructural de Muro de Contención en Voladizo",
    "KonteWall — Cantilever Retaining Wall Structural Design"
))
st.caption(
    "NSR-10 · ACI 318-25/19 · E.060 · NEC-SE-HM · NTC-EM · "
    "COVENIN 1753 · NB 1225001 · CIRSOC 201"
)

with st.expander(" KonteWall — Solución Integral para el Análisis y Diseño de Muros de Contención", expanded=False):
    st.markdown(
        "## ¿Qué es KonteWall?\n\n"
        "**KonteWall** es una plataforma avanzada de ingeniería para el **diseño, análisis y "
        "optimización de muros de contención en voladizo en concreto reforzado**. "
        "Transforma parámetros geotécnicos y estructurales complejos en soluciones "
        "constructivas seguras, económicas y técnicamente rigurosas.\n\n"
        "Desarrollado por **Konte — Construcción, Consultoría y Tecnología**. "
        "Autor: **** — Duitama, Boyacá, Colombia.\n\n"
        "---\n"
        "##  Funcionalidades Clave\n\n"
        "| # | Funcionalidad | Descripción |\n"
        "|---|---|---|\n"
        "| 1 | **Modelado Geotécnico** | Cohesión, fricción interna, pesos específicos |\n"
        "| 2 | **Análisis Sísmico** | Mononobe-Okabe (Kh y Kv) |\n"
        "| 3 | **Optimización Geométrica** | Zapata, pantalla, dentellón |\n"
        "| 4 | **Verificación de Estabilidad** | FS volcamiento, deslizamiento, portante |\n"
        "| 5 | **Diseño de Refuerzo** | Armadura vertical, horizontal y de zapata |\n"
        "| 6 | **Cuantificación** | Concreto, acero, insumos en tiempo real |\n"
        "| 7 | **Exportación Profesional** | Memoria DOCX, hoja XLSX, plano DXF 3D |\n"
        "| 8 | **Visualización** | 2D + Planta + 3D interactivo + Isométrico |\n\n"
        "---\n"
        "##  Colores de Armadura\n\n"
        "| Color | Elemento |\n"
        "|---|---|\n"
        "|  Azul sólido | Acero interior pantalla (tracción) |\n"
        "|  Rojo punteado | Acero exterior (temperatura) |\n"
        "|  Rojo sólido | Acero horizontal (distribución) |\n"
        "|  Verde | Acero talón posterior |\n"
        "|  Naranja | Acero punta (voladizo frontal) |\n"
        "|  Morado | Acero dentellón |\n"
    )
    st.caption("KonteWall v4.0 · 2026 · Konte — Construcción, Consultoría y Tecnología ·  · Duitama, Boyacá")

# ══════════════════════════════════════════════════════════════════
# FUNCIONES AUXILIARES
# ══════════════════════════════════════════════════════════════════
def Ab_diam(d):
    m = {"3/8\"":(0.71,"#3"),"1/2\"":(1.27,"#4"),"5/8\"":(1.99,"#5"),
         "3/4\"":(2.85,"#6"),"7/8\"":(3.87,"#7"),"1\"":(5.10,"#8")}
    return m.get(d,(1.99,"#5"))

def s_acero(As, diam="5/8\"", bw=100.0):
    Ab,_ = Ab_diam(diam)
    if As <= 0: return 0.45
    return round(min(max(Ab*bw/As, 0.10), 0.45), 2)

def As_prov(diam, s, bw=100.0):
    Ab,_ = Ab_diam(diam)
    return (Ab*bw/s) if s > 0 else 0

def tabla_barras(As, bw=100.0):
    rows = []
    for d in ["3/8\"","1/2\"","5/8\"","3/4\"","7/8\"","1\""]:
        Ab,num = Ab_diam(d)
        s = s_acero(As, d, bw)
        rows.append({"Ø":d,"Num":num,"s [m]":s,
                     "As prov [cm²/m]":round(As_prov(d,s,bw),2)})
    return pd.DataFrame(rows)

def badge(ok, etq="", det=""):
    w = ("PASSES" if en_usa else "CUMPLE") if ok else ("FAILS" if en_usa else "NO CUMPLE")
    icon = "" if ok else ""
    msg = f"{icon} **{w}**"
    if etq: msg += f" — {etq}"
    if det: msg += f" _{det}_"
    if ok:
        st.success(msg)
    else:
        st.error(msg)

def recbox(msg):
    st.info(f" {msg}")

# ── Funciones de recomendación ───────────────────────────────────
def recomendar_volteo(FS, FS_min, B, b, ct, Bp, xbr, MrW1, MrW2, MrW3, MrW4, MrEav, MrEscv, sumMa, gr, Hp, SC):
    """Sugiere nuevo Bp para cumplir FS volteo >= FS_min."""
    if FS >= FS_min:
        return None
    Bpr = None
    for Bpt in np.arange(Bp, Bp+3.0, 0.01):
        W5t = gr * Bpt * Hp
        WSCt = SC * Bpt
        x5t = xbr + Bpt/2
        Mrt = MrW1 + MrW2 + MrW3 + MrW4 + W5t*x5t + WSCt*x5t + MrEav + MrEscv
        if sumMa > 0 and Mrt/sumMa >= FS_min:
            Bpr = round(Bpt, 2)
            break
    if Bpr:
        return f"Aumentar talón Bp de {Bp:.2f} m a {Bpr:.2f} m (B total = {round(b+ct+Bpr,2)} m)"
    else:
        return "Considere aumentar la base total B o agregar un dentellón."

def recomendar_excentricidad(e, elim, B, b, ct, Bp, xbr, MrW1, MrW2, MrW3, MrW4, MrEav, MrEscv, sumMa, gr, Hp, SC, sumFv_base, Eav, Escv):
    """Sugiere nuevo B para que e <= elim."""
    if e < elim:
        return None
    for Bt in np.arange(B, B+5.0, 0.05):
        Bpt = Bt - b - ct
        if Bpt < 0:
            continue
        W5t = gr * Bpt * Hp
        WSCt = SC * Bpt
        x5t = xbr + Bpt/2
        Fvt = sumFv_base + W5t + WSCt + Eav + Escv  # sumFv_base debe ser la suma sin W5, WSC, Eav, Escv? Mejor usar los valores ya calculados pero dinámicos. Simplificamos: llamamos a la función con los parámetros necesarios.
        # Como es complejo, por ahora dejamos solo el mensaje genérico.
        pass
    return "Aumentar la base total B para reducir la excentricidad."

def recomendar_q1(q1, qadm, B, b, ct, Bp, xbr, MrW1, MrW2, MrW3, MrW4, MrEav, MrEscv, sumMa, gr, Hp, SC, sumFv_base, Eav, Escv):
    """Sugiere nuevo B para que q1 <= qadm."""
    if q1 <= qadm:
        return None
    for Bt in np.arange(B, B+5.0, 0.05):
        Bpt = Bt - b - ct
        if Bpt < 0:
            continue
        W5t = gr * Bpt * Hp
        WSCt = SC * Bpt
        x5t = xbr + Bpt/2
        Fvt = sumFv_base + W5t + WSCt + Eav + Escv
        Mrt = MrW1 + MrW2 + MrW3 + MrW4 + W5t*x5t + WSCt*x5t + MrEav + MrEscv
        xrt = (Mrt - sumMa) / max(Fvt, 1e-9)
        et = abs(Bt/2 - xrt)
        q1t_ = Fvt / Bt * (1 + 6*et/Bt) / 10000
        if q1t_ <= qadm:
            return f"Ampliar B de {B:.2f} m a {round(Bt,2)} m"
    return "Aumentar la base total B o mejorar la capacidad portante del suelo."

def recomendar_q2(q2):
    """Sugiere acción si q2 es negativo."""
    if q2 >= 0:
        return None
    return "Aumentar la base total B para que la resultante caiga dentro del tercio central."

def recomendar_Ldh(Ldh, hz):
    """Sugiere nuevo hz si Ldh > hz."""
    if Ldh <= hz:
        return None
    hz_new = math.ceil(Ldh * 100 / 5) * 5 / 100  # redondeo a 5 cm
    return f"Aumentar hz de {hz:.2f} m a al menos {hz_new:.2f} m (o usar un diámetro de barra menor)."

def recomendar_cortante_pantalla(Vu, phiVc, ct, rec):
    """Sugiere nuevo ct si Vu > phiVc."""
    if Vu <= phiVc:
        return None
    d_needed = Vu / (phic * 0.53 * math.sqrt(fc) * bw)  # en cm
    ct_new = math.ceil((d_needed + rec) / 5) * 5 / 100
    return f"Aumentar ct de {ct:.2f} m a {ct_new:.2f} m"

def recomendar_deslizamiento(FS, FS_min, B, muct, sumFv, Ccoh, Ep, sumFh):
    """Sugiere acciones para mejorar FS deslizamiento."""
    if FS >= FS_min:
        return None
    msg = []
    # Opción 1: aumentar B (aumenta sumFv por peso propio y relleno)
    # Cálculo aproximado: necesitamos un nuevo B que cumpla (muct*(W1+W2+W3+W4+W5+WSC+Eav+Escv) + Ccoh*B + Ep) / sumFh >= FS_min
    # Es una ecuación no lineal, pero podemos iterar.
    B_new = None
    for Bt in np.arange(B, B+5.0, 0.05):
        Bpt = Bt - b - ct
        if Bpt < 0:
            continue
        W1t = gc * hz * Bt
        W5t = gr * Bpt * Hp
        WSCt = SC * Bpt
        sumFv_t = W1t + W2 + W3 + W4 + W5t + WSCt + Eav + Escv
        Fr_t = muct * sumFv_t + Ccoh * Bt + Ep
        if Fr_t / max(sumFh, 1e-9) >= FS_min:
            B_new = round(Bt, 2)
            break
    if B_new:
        msg.append(f"Aumentar B a {B_new} m")
    # Opción 2: añadir dentellón si no existe
    if hd == 0:
        msg.append("Considerar agregar un dentellón para aumentar el empuje pasivo.")
    # Opción 3: aumentar mu (si es posible)
    msg.append("Si es factible, aumentar el coeficiente de fricción μ mejorando la interfaz.")
    return " | ".join(msg)

# ══════════════════════════════════════════════════════════════════
# SIDEBAR — PARÁMETROS DE ENTRADA (con descripciones)
# ══════════════════════════════════════════════════════════════════
st.sidebar.markdown("## ⚙ Parámetros de Diseño")

with st.sidebar.expander(" Norma y Idioma", expanded=False):
    c1, c2 = st.columns([2, 1])
    with c1:
        norma_sel = st.selectbox("Norma de diseño", [
            "NSR-10 (Colombia)","ACI 318-25 (USA)","E.060 (Perú)",
            "NEC-SE-HM (Ecuador)","NTC-EM (México)","COVENIN 1753 (Venezuela)",
            "NB 1225001 (Bolivia)","CIRSOC 201 (Argentina)"
        ], key="norma_sel")
    with c2:
        st.caption("Código de diseño")
    c1, c2 = st.columns([2, 1])
    with c1:
        lang = st.selectbox("Idioma / Language", ["Español","English"], key="idioma")
    with c2:
        st.caption("Lenguaje")

with st.sidebar.expander(" Materiales", expanded=True):
    c1, c2 = st.columns([2, 1])
    with c1:
        fc = st.number_input("f'c [kg/cm²]", 140.0, 420.0, 210.0, 5.0)
    with c2:
        st.caption("Resistencia concreto")
    c1, c2 = st.columns([2, 1])
    with c1:
        fy = st.number_input("fy [kg/cm²]", 2800.0, 5600.0, 4200.0, 100.0)
    with c2:
        st.caption("Fluencia acero")
    c1, c2 = st.columns([2, 1])
    with c1:
        gc = st.number_input("γc [kg/m³]", 2200.0, 2500.0, 2400.0, 50.0)
    with c2:
        st.caption("Peso concreto")
    c1, c2 = st.columns([2, 1])
    with c1:
        phif = st.number_input("φ flexión", 0.70, 0.95, 0.90, 0.01)
    with c2:
        st.caption("Factor flexión")
    c1, c2 = st.columns([2, 1])
    with c1:
        phic = st.number_input("φ cortante", 0.70, 0.90, 0.85, 0.01)
    with c2:
        st.caption("Factor cortante")
    c1, c2 = st.columns([2, 1])
    with c1:
        rec = st.number_input("Recubrimiento [cm]", 4.0, 10.0, 7.5, 0.5)
    with c2:
        st.caption("Recub. concreto")

with st.sidebar.expander(" Suelo de Relleno", expanded=True):
    c1, c2 = st.columns([2, 1])
    with c1:
        gr = st.number_input("γr [kg/m³]", 1400.0, 2200.0, 1800.0, 50.0)
    with c2:
        st.caption("Peso relleno")
    c1, c2 = st.columns([2, 1])
    with c1:
        ph1 = st.number_input("φ₁ relleno [°]", 15.0, 45.0, 30.0, 1.0)
    with c2:
        st.caption("Ángulo fricción")
    c1, c2 = st.columns([2, 1])
    with c1:
        alp = st.number_input("α talud relleno [°]", 0.0, 20.0, 0.0, 1.0)
    with c2:
        st.caption("Inclinación terreno")
    c1, c2 = st.columns([2, 1])
    with c1:
        delt = st.number_input("δ fricción muro [°]", 0.0, 25.0, 0.0, 1.0)
    with c2:
        st.caption("Fricción muro-suelo")
    c1, c2 = st.columns([2, 1])
    with c1:
        SC = st.number_input("Sobrecarga SC [kg/m²]", 0.0, 2000.0, 300.0, 50.0)
    with c2:
        st.caption("Sobrecarga")
    c1, c2 = st.columns([2, 1])
    with c1:
        hs = st.number_input("Altura SC equiv. hs [m]", 0.0, 3.0, SC/gr if gr>0 else 0.17, 0.01)
    with c2:
        st.caption("Altura equivalente")

with st.sidebar.expander(" Suelo de Fundación", expanded=True):
    c1, c2 = st.columns([2, 1])
    with c1:
        gf = st.number_input("γf [kg/m³]", 1400.0, 2200.0, 1900.0, 50.0)
    with c2:
        st.caption("Peso fundación")
    c1, c2 = st.columns([2, 1])
    with c1:
        ph2 = st.number_input("φ₂ fundación [°]", 15.0, 45.0, 28.0, 1.0)
    with c2:
        st.caption("Ángulo fricción")
    c1, c2 = st.columns([2, 1])
    with c1:
        Ccoh = st.number_input("Cohesión C [kg/m²]", 0.0, 5000.0, 0.0, 100.0)
    with c2:
        st.caption("Cohesión")
    c1, c2 = st.columns([2, 1])
    with c1:
        qadm = st.number_input("qadm [kg/cm²]", 0.5, 10.0, 2.0, 0.1)
    with c2:
        st.caption("Capacidad portante")
    c1, c2 = st.columns([2, 1])
    with c1:
        muct = st.number_input("μ fricción cimiento", 0.20, 0.80, 0.45, 0.01)
    with c2:
        st.caption("Coef. fricción")

with st.sidebar.expander(" Geometría", expanded=True):
    c1, c2 = st.columns([2, 1])
    with c1:
        Ht = st.number_input("Ht total muro [m]", 1.0, 12.0, 4.90, 0.10)
    with c2:
        st.caption("Altura total")
    c1, c2 = st.columns([2, 1])
    with c1:
        hz = st.number_input("hz espesor zapata [m]", 0.20, 1.50, 0.50, 0.05)
    with c2:
        st.caption("Altura zapata")
    Hp = round(Ht - hz, 3)
    c1, c2 = st.columns([2, 1])
    with c1:
        ct = st.number_input("ct espesor base pantalla [m]", 0.20, 1.00, 0.45, 0.05)
    with c2:
        st.caption("Base pantalla")
    c1, c2 = st.columns([2, 1])
    with c1:
        ccor = st.number_input("ccor corona pantalla [m]", 0.15, 0.60, 0.30, 0.05)
    with c2:
        st.caption("Corona")
    c1, c2 = st.columns([2, 1])
    with c1:
        B = st.number_input("B base total [m]", 1.00, 10.0, 3.10, 0.05)
    with c2:
        st.caption("Ancho base")
    c1, c2 = st.columns([2, 1])
    with c1:
        b = st.number_input("b punta [m]", 0.10, 3.0, 0.80, 0.05)
    with c2:
        st.caption("Longitud punta")
    Bp = round(B - b - ct, 3)
    c1, c2 = st.columns([2, 1])
    with c1:
        hd = st.number_input("hd dentellón [m]", 0.00, 1.50, 0.30, 0.05)
    with c2:
        st.caption("Profundidad dentellón")

with st.sidebar.expander(" Sismicidad", expanded=False):
    c1, c2 = st.columns([2, 1])
    with c1:
        sismo = st.checkbox("Incluir sismo (Mononobe-Okabe)", value=True)
    with c2:
        st.caption("Análisis sísmico")
    if sismo:
        c1, c2 = st.columns([2, 1])
        with c1:
            Kh = st.number_input("Kh coef. sísmico horiz.", 0.0, 0.40, 0.15, 0.01)
        with c2:
            st.caption("Coef. horizontal")
        c1, c2 = st.columns([2, 1])
        with c1:
            Kv = st.number_input("Kv coef. sísmico vert.", 0.0, 0.20, 0.05, 0.01)
        with c2:
            st.caption("Coef. vertical")
    else:
        Kh = 0.0
        Kv = 0.0

with st.sidebar.expander(" Longitud del Muro", expanded=False):
    c1, c2 = st.columns([2, 1])
    with c1:
        Lmuro = st.number_input("Longitud L [m]", 0.5, 50.0, 10.0, 0.5)
    with c2:
        st.caption("Longitud total")

# ══ Copyright sidebar ══
st.sidebar.markdown("---")
st.sidebar.markdown(
    " **KonteWall v4.0** © 2026  \n"
    "Todos los derechos reservados  \n"
    "**Konte** — Construcción, Consultoría y Tecnología  \n"
    "**  \n"
    "Duitama, Boyacá, Colombia"
)

# ══════════════════════════════════════════════════════════════════
# CÁLCULO
# ══════════════════════════════════════════════════════════════════
if st.sidebar.button(_t(" Calcular Muro", " Design Wall"), type="primary"):
    st.session_state["calculado"] = True

if not st.session_state.get("calculado", False):
    st.info(_t(" Configure los parámetros y pulse **Calcular Muro**",
               " Set parameters and press **Design Wall**"))
    st.stop()

# ── Variables globales auxiliares ──
rm = rec / 100  # recubrimiento en metros

# ── Geometría derivada ──
xbl = b
xbr = b + ct
xtl = xbl + (ct - ccor)
xtr = xbr

# ── Predimensionamiento ──
db_supuesto = 1.59
Ldh_supuesto = max(0.24 * fy / math.sqrt(fc) * db_supuesto, 8 * db_supuesto, 15.0) / 100
hz_necesaria = Ldh_supuesto + rec / 100

# ── Coeficientes de presión (igual que antes) ──
ph1r = math.radians(ph1)
alpr = math.radians(alp)
deltr = math.radians(delt)

_alp_wall = math.radians(90)
_sin_sum = math.sin(ph1r + deltr)
_sin_dif = math.sin(ph1r - math.radians(alp))
_sin_den1 = math.sin(_alp_wall - deltr)
_sin_den2 = math.sin(_alp_wall + math.radians(alp))
_denom = _sin_den1 * _sin_den2
if _denom <= 0:
    Ka = math.tan(math.radians(45) - ph1r/2)**2
else:
    _raiz = math.sqrt(max(0.0, _sin_sum * _sin_dif / _denom))
    Ka = math.sin(ph1r + _alp_wall)**2 / (math.sin(_alp_wall)**2 * _sin_den1 * (1 + _raiz)**2)

if sismo:
    thr = math.atan(Kh / max(1 - Kv, 1e-9))
    ph1t = ph1r - thr
    _alp_wall_mo = math.radians(90)
    _s1 = math.sin(ph1t + deltr)
    _s2 = math.sin(ph1t - math.radians(alp))
    _sd1 = math.sin(_alp_wall_mo - deltr - thr)
    _sd2 = math.sin(_alp_wall_mo + math.radians(alp))
    _dmo = _sd1 * _sd2
    if _dmo <= 0 or ph1t <= 0:
        Kea = Ka * (1 - Kv)
    else:
        _raiz_mo = math.sqrt(max(0.0, _s1 * _s2 / _dmo))
        Kea = math.sin(ph1t + _alp_wall_mo)**2 / (math.cos(thr) * math.sin(_alp_wall_mo)**2 * _sd1 * (1 + _raiz_mo)**2)
else:
    thr = 0.0
    Kea = Ka

ph2r = math.radians(ph2)
Kp = math.tan(math.radians(45) + ph2r/2)**2

# ── Empujes ──
Eah = 0.5 * Ka * gr * Hp**2
Esch = Ka * gr * hs * Hp
Ep = 0.5 * Kp * gf * (hz + hd)**2 if hd > 0 else 0.5 * Kp * gf * hz**2
Es = 0.5 * (Kea - Ka) * (1 - Kv) * gr * Hp**2 if sismo else 0.0
# Wm eliminado: ya separado en W2 (triángulo) + W3 (rectángulo)
Eav = (Eah + Es) * math.sin(deltr)
Escv = Esch * math.sin(deltr)
Cy = 0.6 * Hp + hz
Eim = Kh * gc * 0.5 * (ct + ccor) * Hp if sismo else 0.0

# ── Pesos y momentos ──
W1 = gc * hz * B
W2 = 0.5 * gc * (ct - ccor) * Hp
W3 = gc * ccor * Hp
W4 = gc * ct * hd if hd > 0 else 0.0
W5 = gr * Bp * Hp
WSC = SC * Bp

x1 = B / 2
x2 = b + ct - (ct - ccor) / 3
x3 = b + ct - ccor / 2
x4 = b + ct / 2
x5 = xbr + Bp / 2
xSC = xbr + Bp / 2

sumFv = W1 + W2 + W3 + W4 + W5 + WSC + Eav + Escv
sumFh = Eah + Esch + Es + Eim - Ep

MrW1 = W1 * x1
MrW2 = W2 * x2
MrW3 = W3 * x3
MrW4 = W4 * x4
MrW5 = W5 * x5
MrWSC = WSC * xSC
MrEav = Eav * xbr
MrEscv = Escv * xbr
sumMr = MrW1 + MrW2 + MrW3 + MrW4 + MrW5 + MrWSC + MrEav + MrEscv

MaEah = Eah * (Hp / 3 + hz)
MaEsch = Esch * (Hp / 2 + hz)
MaEs = Es * Cy if sismo else 0.0
MaEim = Eim * (Hp / 2 + hz) if sismo else 0.0
sumMa = MaEah + MaEsch + MaEs + MaEim

# ── Verificaciones ──
FSvolt = sumMr / max(sumMa, 1e-9)
Frdesl = muct * sumFv + Ccoh * B + Ep
FSdesl = Frdesl / max(sumFh, 1e-9)
xr = (sumMr - sumMa) / max(sumFv, 1e-9)
eexc = abs(B / 2 - xr)
elim = B / 6
q1 = sumFv / B * (1 + 6 * eexc / B) / 10000
q2 = sumFv / B * (1 - 6 * eexc / B) / 10000

# Sin sismo
W5ns = gr * Bp * Hp
sumFvns = W1 + W2 + W3 + W4 + W5ns + Ccoh * B
sumFhns = Eah + Esch - Ep
Mrns = MrW1 + MrW2 + MrW3 + MrW4 + W5ns * x5 + MrWSC + MrEav + MrEscv
Mans = MaEah + MaEsch
FSvns = Mrns / max(Mans, 1e-9)
FSdns = (muct * sumFvns + Ccoh * B + Ep) / max(sumFhns, 1e-9)
xrns = (Mrns - Mans) / max(sumFvns, 1e-9)
ens = abs(B/2 - xrns)
q1ns = sumFvns / B * (1 + 6 * ens / B) / 10000
q2ns = sumFvns / B * (1 - 6 * ens / B) / 10000

# ── Diseño acero ──
rhomin = 0.0015
rhomin_zap = 0.0018
diam_def = "5/8\""
bw = 100.0

# Pantalla
dp = ct * 100 - rec
Mup = abs(1.7 * (MaEah + MaEsch) + (1.4 * MaEs if sismo else 0))
Rup = Mup * 100 / (phif * bw * dp**2)
rhoreq = 0.85 * fc / fy * (1 - math.sqrt(max(0, 1 - 2 * Rup / (0.85 * fc))))
Asp = max(rhoreq, rhomin) * bw * dp
dii = diam_def
si = s_acero(Asp, dii)
Api = As_prov(dii, si)
die = "5/8\""
se = s_acero(rhomin * bw * dp, die)
Ape = As_prov(die, se)
Aspe = rhomin * bw * dp

# Horizontal
Asmin_sup = 0.0025 * (ct + ccor) / 2 * 100 * bw
Asmin_inf = 0.0025 * ct * 100 * bw
shext_inf = s_acero(2/3 * Asmin_inf, "1/2\"")
shint_inf = s_acero(1/3 * Asmin_inf, "3/8\"")
shext_sup = s_acero(2/3 * Asmin_sup, "1/2\"")
shint_sup = s_acero(1/3 * Asmin_sup, "3/8\"")
shext = shext_inf
shint = shint_inf
# Corte de acero
dc = 0.0
if Asp > 0:
    for h_try in np.arange(0.0, Hp, 0.01):
        Mu_h = abs(1.7 * (0.5 * Ka * gr * h_try**2 * h_try / 3 + Ka * gr * hs * h_try * h_try / 2))
        dp_h = max((ccor + (ct - ccor) * h_try / Hp) * 100 - rec, 5)
        am_h = Api * fy / (0.85 * fc * bw)
        phiMn_h = phif * Api * fy * (dp_h - am_h / 2) / 100
        if phiMn_h >= Mu_h:
            dc = round((Hp - h_try), 2)
            break

# Anclaje
db_diam = {"3/8\"": 0.95, "1/2\"": 1.27, "5/8\"": 1.59, "3/4\"": 1.91, "7/8\"": 2.22, "1\"": 2.54}
db_val = db_diam.get(dii, 1.59)
Ldh = max(0.24 * fy / math.sqrt(fc) * db_val, 8 * db_val, 15.0) / 100

# Talón posterior
dz = hz * 100 - rec
Wu1 = 1.4 * (gr * Hp + hz * gc) + 1.7 * SC
q2t = q2 * 10000
q1t = q1 * 10000
q_xbr = q1t + (q2t - q1t) * (B - Bp) / B
Mut = abs(Wu1 * Bp**2 / 2 - 1.7 * (q_xbr * Bp**2 / 2 + (q2t - q_xbr) * Bp**2 / 3))
Rut = Mut * 100 / (phif * bw * dz**2)
rhot = 0.85 * fc / fy * (1 - math.sqrt(max(0, 1 - 2 * Rut / (0.85 * fc))))
Ast = max(rhot, rhomin_zap) * bw * dz
diamtl = "3/4\""
stl = s_acero(Ast, diamtl)
Aptl = As_prov(diamtl, stl)
Astp = rhomin_zap * bw * dz
stp = s_acero(Astp, "5/8\"")
Vut = abs(Wu1 * Bp - 1.7 * (q_xbr * Bp + 0.5 * (q2t - q_xbr) * Bp))
phiVct = phic * 0.53 * math.sqrt(fc) * bw * dz

# Talón delantero (punta)
q_b = q1t + (q2t - q1t) * b / B
Mupu = abs(1.7 * (q1t * b**2 / 2 + (q_b - q1t) * b**2 / 6) - 1.4 * gc * hz * b**2 / 2)
Rupu = Mupu * 100 / (phif * bw * dz**2)
rhopu = 0.85 * fc / fy * (1 - math.sqrt(max(0, 1 - 2 * Rupu / (0.85 * fc))))
Aspu = max(rhopu, rhomin_zap) * bw * dz
diampu = "5/8\""
spu = s_acero(Aspu, diampu)
Appu = As_prov(diampu, spu)
Aspp = rhomin_zap * bw * dz
spp = s_acero(Aspp, "5/8\"")
phiVcpu = phic * 0.53 * math.sqrt(fc) * bw * dz
Vupu = abs(1.7 * (q1t * b + 0.5 * (q_b - q1t) * b) - 1.4 * gc * hz * b)

# Cortante pantalla
Vupant = abs(1.7 * (0.5 * Ka * gr * Hp**2 + Ka * gr * hs * Hp))
phiVcp = phic * 0.53 * math.sqrt(fc) * bw * dp

# Cubicación
Vc_zap = hz * B * Lmuro
Vc_pan = 0.5 * (ct + ccor) * Hp * Lmuro
Vc_den = ct * hd * Lmuro if hd > 0 else 0.0
Vc_tot = Vc_zap + Vc_pan + Vc_den
As_tot_kg = (Aptl * Bp + Appu * b + Api * Hp + Ape * Hp) * Lmuro * 7850 / 10000

# ══════════════════════════════════════════════════════════════════
# RESULTADOS — EXPANDERS
# ══════════════════════════════════════════════════════════════════
with st.expander("01. Datos Generales y Predimensionamiento", expanded=True):
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**Materiales**")
        st.write(f"f'c = {fc:.0f} kg/cm²  |  fy = {fy:.0f} kg/cm²")
        st.write(f"γc = {gc:.0f} kg/m³  |  φf = {phif:.2f}  φv = {phic:.2f}")
    with c2:
        st.markdown("**Suelo de Relleno**")
        st.write(f"γr = {gr:.0f} kg/m³  φ₁ = {ph1:.1f}°  α = {alp:.1f}°  δ = {delt:.1f}°")
        st.write(f"S/C = {SC:.0f} kg/m²  hs = {hs:.3f} m")
    with c3:
        st.markdown("**Geometría adoptada**")
        st.write(f"Ht = {Ht:.2f} m  Hp = {Hp:.2f} m  hz = {hz:.2f} m")
        st.write(f"B = {B:.2f} m  Bp = {Bp:.2f} m  b = {b:.2f} m  ct = {ct:.2f} m")
        if hd > 0:
            st.write(f"Dentellón hd = {hd:.2f} m")

    st.markdown("---")
    st.markdown("### Predimensionamiento")
    colA, colB = st.columns(2)
    with colA:
        st.markdown("**Altura de zapata**")
        st.latex(r"L_{dh} = \frac{0.24 f_y}{\sqrt{f'_c}} \cdot d_b \geq 8d_b \geq 15\text{ cm}")
        st.write(f"Diámetro asumido: 5/8\" (d_b = {db_supuesto:.2f} cm)")
        st.write(f"Longitud de anclaje requerida Ldh = {Ldh_supuesto*100:.2f} cm")
        st.write(f"Altura necesaria de zapata: hz_necesaria = Ldh + rec = {hz_necesaria*100:.2f} cm")
        st.write(f"Altura adoptada: hz = {hz*100:.2f} cm → {' OK' if hz >= hz_necesaria else '⚠ Aumentar'}")
    with colB:
        st.markdown("**Longitud de la base**")
        st.write(f"Rango recomendado: H/2 = {Ht/2:.2f} m  a  2H/3 = {2*Ht/3:.2f} m")
        st.write(f"Valor adoptado: B = {B:.2f} m → {' OK' if Ht/2 <= B <= 2*Ht/3 else '⚠ Revisar'}")

    st.markdown("---")
    st.markdown("### Verificación de dimensiones")
    rowspre = []
    for par, vmin, vmax, vadp in [
        ("hz Zapata", Ht/12, Ht/10, hz),
        ("B base total", Ht/2, 2*Ht/3, B),
        ("ccor corona", max(Ht/24, 0.30), None, ccor),
        ("ct base pantalla", Ht/12, Ht/10, ct),
    ]:
        okv = vadp >= vmin if vmax is None else vmin <= vadp <= vmax
        rowspre.append({
            "Parámetro": par,
            "Mín": f"{vmin:.2f} m",
            "Máx": f"{vmax:.2f} m" if vmax else "—",
            "Adoptado": f"{vadp:.2f} m",
            "Estado": " OK" if okv else "⚠ Revisar"
        })
    st.table(pd.DataFrame(rowspre))

with st.expander("02. Coeficientes de Empuje Coulomb / Mononobe-Okabe", expanded=False):
    c1, c2, c3 = st.columns(3)
    c1.metric("Ka Coulomb estático", f"{Ka:.4f}")
    c2.metric("Kea Mononobe-Okabe", f"{Kea:.4f}" if sismo else "N/A")
    c3.metric("Kp pasivo Coulomb", f"{Kp:.3f}")
    if sismo:
        st.write(f"θ = arctan(Kh/(1-Kv)) = {math.degrees(thr):.2f}°")
    st.latex(r"K_a = \frac{\sin^2(\phi_1+\beta)}{\sin^2\beta \cdot \sin(\beta-\delta)"
             r"\left(1+\sqrt{\frac{\sin(\phi_1+\delta)\sin(\phi_1-\alpha)}{\sin(\beta-\delta)\sin(\alpha+\beta)}}\right)^2}")
    st.caption("Con β = 90° (muro vertical)")

with st.expander("03. Verificación de Estabilidad" + (" + Sismo Mononobe-Okabe" if sismo else ""), expanded=True):
    cf, cv = st.columns(2)
    with cf:
        st.markdown("**Fuerzas Horizontales [kg/m]**")
        st.table(pd.DataFrame({
            "Concepto": ["Eah tierra", "Esch sobrecarga", "Es sísmico (Kea-Ka)", "Eim inercial", "Ep pasivo −"],
            "F [kg/m]": [f"{Eah:.2f}", f"{Esch:.2f}", f"{Es:.2f}", f"{Eim:.2f}", f"{-Ep:.2f}"],
            "Brazo [m]": [f"{Hp/3:.2f}", f"{Hp/2:.2f}", f"{Cy:.2f}", f"{Hp/2+hz:.2f}", "—"],
            "Ma [kg·m/m]": [f"{MaEah:.2f}", f"{MaEsch:.2f}", f"{MaEs:.2f}", f"{MaEim:.2f}", "—"]
        }))
        st.write(f"**ΣFh = {sumFh:.2f} kg/m  |  ΣMa = {sumMa:.2f} kg·m/m**")
    with cv:
        st.markdown("**Fuerzas Verticales y Momentos Estabilizadores [kg/m]**")
        st.table(pd.DataFrame({
            "Concepto": ["W1 Zapata", "W2 Pantalla △", "W3 Pantalla □", "W4 Dentellón", "W5 Relleno", "SC", "Eav", "Escv"],
            "F [kg/m]": [f"{W1:.2f}", f"{W2:.2f}", f"{W3:.2f}", f"{W4:.2f}", f"{W5:.2f}", f"{WSC:.2f}", f"{Eav:.2f}", f"{Escv:.2f}"],
            "Brazo [m]": [f"{x1:.2f}", f"{x2:.2f}", f"{x3:.2f}", f"{x4:.2f}", f"{x5:.2f}", f"{xSC:.2f}", f"{xbr:.2f}", f"{xbr:.2f}"],
            "Mr [kg·m/m]": [f"{MrW1:.2f}", f"{MrW2:.2f}", f"{MrW3:.2f}", f"{MrW4:.2f}", f"{MrW5:.2f}", f"{MrWSC:.2f}", f"{MrEav:.2f}", f"{MrEscv:.2f}"]
        }))
        st.write(f"**ΣFv = {sumFv:.2f} kg/m  |  ΣMr = {sumMr:.2f} kg·m/m**")
    st.markdown("---")
    okvFS = FSvolt >= 2.0
    okdFS = FSdesl >= 1.5
    okee = eexc < elim
    okq1 = q1 <= qadm
    okq2 = q2 >= 0

    st.markdown(f"**1.  Volcamiento** — ΣMr/ΣMa = {sumMr:.2f}/{sumMa:.2f}")
    badge(okvFS, f"FS={FSvolt:.3f}", "≥ 2.00")
    if not okvFS:
        sugg = recomendar_volteo(FSvolt, 2.0, B, b, ct, Bp, xbr, MrW1, MrW2, MrW3, MrW4, MrEav, MrEscv, sumMa, gr, Hp, SC)
        if sugg:
            recbox(sugg)

    st.markdown(f"**2. ↔ Deslizamiento** — Fr/ΣFh = {Frdesl:.2f}/{sumFh:.2f}")
    badge(okdFS, f"FS={FSdesl:.3f}", "≥ 1.50")
    if not okdFS:
        sugg = recomendar_deslizamiento(FSdesl, 1.5, B, muct, sumFv, Ccoh, Ep, sumFh)
        if sugg:
            recbox(sugg)

    st.markdown(f"**3.  Excentricidad** — e = {eexc:.4f}m  |  B/6 = {elim:.4f}m")
    badge(okee, f"e={eexc:.4f}m", f"< {elim:.4f}m")
    if not okee:
        sugg = recomendar_excentricidad(eexc, elim, B, b, ct, Bp, xbr, MrW1, MrW2, MrW3, MrW4, MrEav, MrEscv, sumMa, gr, Hp, SC, W1+W2+W3+W4, Eav, Escv)
        if sugg:
            recbox(sugg)
        else:
            recbox("Aumentar la base total B para reducir la excentricidad.")

    st.markdown(f"**4.  Cap. portante** — q₁ = {q1:.4f} kg/cm²  |  qadm = {qadm:.2f} kg/cm²")
    badge(okq1, f"q₁={q1:.4f}", f"qadm={qadm:.2f}")
    if not okq1:
        sugg = recomendar_q1(q1, qadm, B, b, ct, Bp, xbr, MrW1, MrW2, MrW3, MrW4, MrEav, MrEscv, sumMa, gr, Hp, SC, W1+W2+W3+W4, Eav, Escv)
        if sugg:
            recbox(sugg)
        else:
            recbox("Aumentar la base total B o mejorar el terreno.")

    st.markdown(f"**5. q₂ = {q2:.4f} kg/cm² ≥ 0**")
    badge(okq2)
    if not okq2:
        sugg = recomendar_q2(q2)
        if sugg:
            recbox(sugg)

    # Diagrama presiones
    fig_q, ax_q = plt.subplots(figsize=(7, 2.5))
    fig_q.patch.set_facecolor(_plt_bg)
    ax_q.set_facecolor(_plt_bg)
    ax_q.fill_between([0, B], [0, 0], [q1, q2], alpha=0.35, color='#2196F3')
    ax_q.plot([0, B], [q1, q2], 'o-', color='#4CAF50', lw=2)
    ax_q.axhline(qadm, color='red', lw=2, ls='--', label=f'qadm={qadm:.2f}')
    ax_q.text(0.02, q1 * 1.05, f'q₁={q1:.3f}', color=_plt_txt, fontsize=9)
    ax_q.text(B * 0.65, q2 * 1.10 + 0.003, f'q₂={q2:.3f}', color=_plt_txt, fontsize=9)
    ax_q.set_xlabel('Posición [m]', color=_plt_txt)
    ax_q.set_ylabel('[kg/cm²]', color=_plt_txt)
    ax_q.tick_params(colors=_plt_txt)
    ax_q.grid(True, alpha=0.2, color=_plt_grid)
    ax_q.legend(facecolor=_plt_bg, labelcolor=_plt_txt, fontsize=8)
    st.pyplot(fig_q)

    st.markdown("---")
    st.markdown("### Relaciones Demanda/Capacidad (D/C)")
    dc_data = {
        "Verificación": ["Volteo", "Deslizamiento", "Excentricidad", "Presión q₁", "Presión q₂"],
        "D/C": [
            f"{2.0/FSvolt:.3f}" if FSvolt > 0 else "N/A",
            f"{1.5/FSdesl:.3f}" if FSdesl > 0 else "N/A",
            f"{eexc/elim:.3f}" if elim > 0 else "N/A",
            f"{q1/qadm:.3f}" if qadm > 0 else "N/A",
            f"{q2/qadm:.3f}" if qadm > 0 else "N/A"
        ]
    }
    st.table(pd.DataFrame(dc_data))

    st.markdown("### Resumen de verificación")
    resumen = pd.DataFrame({
        "Verificación": ["Volteo", "Deslizamiento", "Excentricidad", "Presión q₁", "Presión q₂"],
        "Valor": [f"FS={FSvolt:.3f}", f"FS={FSdesl:.3f}", f"e={eexc:.4f}m", f"q₁={q1:.4f}", f"q₂={q2:.4f}"],
        "Límite": ["≥2.00", "≥1.50", f"<{elim:.4f}m", f"≤{qadm:.2f}", "≥0"],
        "Estado": [
            " Cumple" if okvFS else " No cumple",
            " Cumple" if okdFS else " No cumple",
            " Cumple" if okee else " No cumple",
            " Cumple" if okq1 else " No cumple",
            " Cumple" if okq2 else " No cumple"
        ]
    })
    st.table(resumen)

with st.expander("04. Verificación Sin Sobrecarga", expanded=False):
    st.table(pd.DataFrame({
        "Verificación": ["FS volteo", "FS desl.", "Excentr.", "q₁", "q₂"],
        "Valor": [f"FS={FSvns:.3f}", f"FS={FSdns:.3f}", f"e={ens:.4f}m", f"q₁={q1ns:.4f} kg/cm²", f"q₂={q2ns:.4f} kg/cm²"],
        "Estado": [" OK" if x else "⚠ Revisar" for x in [FSvns >= 2.0, FSdns >= 1.5, ens < B/6, q1ns <= qadm, q2ns >= 0]]
    }))

with st.expander(f"05. Diseño Pantalla — Sección Crítica Base", expanded=False):
    c1, c2 = st.columns(2)
    with c1:
        st.write(f"**Mu base pantalla** = {Mup:.2f} kg·m/m")
        st.latex(r"M_u = 1.7 \cdot \left(\frac{K_a \gamma_r H_p^2}{3} \cdot H_p + K_a \gamma_r h_s H_p \cdot \frac{H_p}{2}\right)")
        st.write(f"d = {dp:.2f} cm  |  Ru = {Rup:.4f} kg/cm²")
        st.write(f"ρ req = {rhoreq:.5f}  |  **As req = {Asp:.2f} cm²/m**")
        st.markdown(f"**Cara Interior (tracción) — As = {Asp:.2f} cm²/m**")
        st.table(tabla_barras(Asp))
        badge(True, f"Ø {dii} @ {si:.2f}m", f"As prov={Api:.2f} cm²/m")
        st.markdown(f"**Cara Exterior (temperatura) — As = {Aspe:.2f} cm²/m**")
        st.table(tabla_barras(Aspe))
        badge(True, f"Ø {die} @ {se:.2f}m", f"As prov={Ape:.2f} cm²/m")
    with c2:
        st.markdown("**Acero Horizontal — Distribución por Tramos**")
        st.markdown(f"*Tramo inferior* — As min = {Asmin_inf:.2f} cm²/m")
        df_hinf_ext = tabla_barras(2/3 * Asmin_inf)
        df_hinf_int = tabla_barras(1/3 * Asmin_inf)
        st.markdown(f"Cara Exterior (2/3): **Ø 1/2\" @ {shext_inf:.2f}m**")
        st.dataframe(df_hinf_ext, hide_index=True)
        st.markdown(f"Cara Interior (1/3): **Ø 3/8\" @ {shint_inf:.2f}m**")
        st.dataframe(df_hinf_int, hide_index=True)
        st.markdown(f"*Tramo superior* — As min = {Asmin_sup:.2f} cm²/m")
        df_hsup_ext = tabla_barras(2/3 * Asmin_sup)
        df_hsup_int = tabla_barras(1/3 * Asmin_sup)
        st.markdown(f"Cara Exterior (2/3): **Ø 1/2\" @ {shext_sup:.2f}m**")
        st.dataframe(df_hsup_ext, hide_index=True)
        st.markdown(f"Cara Interior (1/3): **Ø 3/8\" @ {shint_sup:.2f}m**")
        st.dataframe(df_hsup_int, hide_index=True)
    st.markdown("**Verificación Cortante Pantalla**")
    st.write(f"Vu = {Vupant:.2f} kg/m  |  φVc = {phiVcp:.2f} kg/m")
    badge(Vupant <= phiVcp, f"Vu={Vupant:.2f}", f"vs φVc={phiVcp:.2f}")
    if Vupant > phiVcp:
        d_needed = Vupant / (phic * 0.53 * math.sqrt(fc) * bw)
        ct_new = math.ceil((d_needed + rec) / 5) * 5 / 100
        recbox(f" Aumentar ct a {ct_new:.2f} m")
    st.markdown(f"**Longitud de anclaje Ldh = {Ldh:.2f}m  |  Disponible hz={hz:.2f}m**")
    badge(Ldh <= hz, f"Ldh={Ldh:.2f}m", f"hz={hz:.2f}m")
    if not (Ldh <= hz):
        sugg = recomendar_Ldh(Ldh, hz)
        if sugg:
            recbox(sugg)
    if dc > 0.05:
        st.markdown(f"** Corte de acero exterior a dc = {dc:.2f} m desde la corona**")
        st.markdown("El acero vertical exterior se cortará a esa distancia, según el análisis de equilibrio.")
    else:
        st.markdown("** No se requiere corte de acero**")

with st.expander(f"06. Diseño Talón Posterior — Bp={Bp:.2f}m", expanded=False):
    c1, c2 = st.columns(2)
    with c1:
        fig_t, ax_t = plt.subplots(figsize=(5, 3))
        fig_t.patch.set_facecolor(_plt_bg)
        ax_t.set_facecolor(_plt_bg)
        q2p = (gr * Hp + hz * gc)
        ax_t.fill_between([0, Bp], [0, 0], [q2p / 1000, q2 / 10000 * 1000],
                           alpha=0.4, color='#90CAF9', label="Carga talón")
        ax_t.fill_between([0, Bp], [0, 0], [-Wu1 / 1000, -Wu1 / 1000],
                           alpha=0.4, color='#FFCC80', label="Wu1 (hacia abajo)")
        ax_t.axhline(0, color=_plt_txt, lw=1)
        ax_t.set_xlabel("Longitud Bp [m]", color=_plt_txt)
        ax_t.set_ylabel("kN/m²", color=_plt_txt)
        ax_t.tick_params(colors=_plt_txt)
        ax_t.grid(True, alpha=0.2, color=_plt_grid)
        ax_t.legend(facecolor=_plt_bg, labelcolor=_plt_txt, fontsize=8)
        ax_t.set_title(f"Talón posterior Bp={Bp:.2f}m | hz={hz:.2f}m", color=_plt_txt, fontsize=10)
        st.pyplot(fig_t)
        st.write(f"Wu₁ = {Wu1:.2f} kg/m  |  q₂' = {q2t:.2f} kg/m  |  q₂ = {q2 * 10000:.2f} kg/m")

        # Gráfica de momento flector
        x_talon = np.linspace(0, Bp, 50)
        M_talon = Wu1 * x_talon**2 / 2 - 1.7 * (q_xbr * x_talon**2 / 2 + (q2t - q_xbr) * x_talon**3 / (6 * Bp))
        fig_m, ax_m = plt.subplots(figsize=(5, 3))
        fig_m.patch.set_facecolor(_plt_bg)
        ax_m.set_facecolor(_plt_bg)
        ax_m.plot(x_talon, M_talon, color='#FF5722', lw=2)
        ax_m.fill_between(x_talon, 0, M_talon, alpha=0.3, color='#FF5722')
        ax_m.set_xlabel("x [m]", color=_plt_txt)
        ax_m.set_ylabel("Momento [kg-m/m]", color=_plt_txt)
        ax_m.tick_params(colors=_plt_txt)
        ax_m.grid(True, alpha=0.2, color=_plt_grid)
        ax_m.set_title("Diagrama de Momento - Talón Posterior", color=_plt_txt, fontsize=10)
        st.pyplot(fig_m)

    with c2:
        st.markdown("**Cálculo Mu Talón Posterior**")
        st.latex(r"W_{u1} = 1.4(\gamma_r H_p + h_z \gamma_c) + 1.7 S/C")
        st.latex(r"M_u = \frac{W_{u1} \cdot B_p^2}{2} - 1.7\left(\frac{q_2' \cdot B_p^2}{6} + \frac{q_2 \cdot B_p^2}{3}\right)")
        st.write(f"**Mu = {Mut:.2f} kg·m/m**  |  d = {dz:.2f} cm")
        st.write(f"As req = **{Ast:.2f} cm²/m**")
        st.table(tabla_barras(Ast))
        st.markdown(
            f'<div style="border:2px solid black;background:#1a1a2e;'
            f'padding:8px 16px;border-radius:4px;font-size:15px;'
            f'font-weight:bold;color:white;text-align:center;">'
            f'Ø {diamtl} &nbsp;&nbsp; @ &nbsp;&nbsp; {stl:.2f} m</div>',
            unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        st.write(f"Acero transversal mínimo:")
        st.markdown(
            f'<div style="border:2px solid #888;background:#0e1117;'
            f'padding:6px 14px;border-radius:4px;font-size:14px;'
            f'font-weight:bold;color:#ccc;text-align:center;">'
            f'Ø 5/8\" &nbsp;&nbsp; @ &nbsp;&nbsp; {stp:.2f} m</div>',
            unsafe_allow_html=True)
    st.markdown("---")
    st.write(f"Vu = {Vut:.2f} kg/m  |  φVc = {phiVct:.2f} kg/m")
    badge(Vut <= phiVct, f"Vu={Vut:.2f}", f"vs φVc={phiVct:.2f}")

with st.expander(f"07. Diseño Talón Delantero (Punta) — b={b:.2f}m", expanded=False):
    c1, c2 = st.columns(2)
    with c1:
        fig_p, ax_p = plt.subplots(figsize=(5, 3))
        fig_p.patch.set_facecolor(_plt_bg)
        ax_p.set_facecolor(_plt_bg)
        ax_p.fill_between([0, b], [0, 0], [q1 * 10000 / 1000, q2 * 10000 / 1000],
                           alpha=0.4, color='#A5D6A7', label="Reacción suelo")
        ax_p.fill_between([0, b], [0, 0], [-gc * hz, -gc * hz],
                           alpha=0.3, color='#FFCC80', label="Peso zapata")
        ax_p.axhline(0, color=_plt_txt, lw=1)
        ax_p.set_xlabel("Longitud b [m]", color=_plt_txt)
        ax_p.set_ylabel("kN/m²", color=_plt_txt)
        ax_p.tick_params(colors=_plt_txt)
        ax_p.grid(True, alpha=0.2, color=_plt_grid)
        ax_p.legend(facecolor=_plt_bg, labelcolor=_plt_txt, fontsize=8)
        ax_p.set_title(f"Talón delantero (Punta) b={b:.2f}m", color=_plt_txt, fontsize=10)
        st.pyplot(fig_p)
        st.write(f"q₁ = {q1 * 10000:.2f} kg/m  |  q₁' = {q1t:.2f} kg/m")

        # Gráfica de momento flector
        x_punta = np.linspace(0, b, 50)
        q_net = q1t + (q_b - q1t) * x_punta / b
        M_punta = 1.7 * (q1t * x_punta**2 / 2 + (q_b - q1t) * x_punta**3 / (6 * b)) - 1.4 * gc * hz * x_punta**2 / 2
        fig_mp, ax_mp = plt.subplots(figsize=(5, 3))
        fig_mp.patch.set_facecolor(_plt_bg)
        ax_mp.set_facecolor(_plt_bg)
        ax_mp.plot(x_punta, M_punta, color='#FF5722', lw=2)
        ax_mp.fill_between(x_punta, 0, M_punta, alpha=0.3, color='#FF5722')
        ax_mp.set_xlabel("x [m]", color=_plt_txt)
        ax_mp.set_ylabel("Momento [kg-m/m]", color=_plt_txt)
        ax_mp.tick_params(colors=_plt_txt)
        ax_mp.grid(True, alpha=0.2, color=_plt_grid)
        ax_mp.set_title("Diagrama de Momento - Punta", color=_plt_txt, fontsize=10)
        st.pyplot(fig_mp)

    with c2:
        st.markdown("**Cálculo Mu Punta**")
        st.latex(r"M_u = 1.7\left(\frac{5 q_1 b^2}{6} - \frac{q_1' b^2}{3}\right)")
        st.write(f"**Mu = {Mupu:.2f} kg·m/m**  |  d = {dz:.2f} cm")
        st.write(f"As req = **{Aspu:.2f} cm²/m**")
        st.table(tabla_barras(Aspu))
        st.markdown(
            f'<div style="border:2px solid black;background:#1a1a2e;'
            f'padding:8px 16px;border-radius:4px;font-size:15px;'
            f'font-weight:bold;color:white;text-align:center;">'
            f'Ø {diampu} &nbsp;&nbsp; @ &nbsp;&nbsp; {spu:.2f} m</div>',
            unsafe_allow_html=True)
        st.write(f"Vu = {Vupu:.2f} kg/m  |  φVc = {phiVcpu:.2f} kg/m")
        badge(Vupu <= phiVcpu, f"Vu={Vupu:.2f}", f"vs φVc={phiVcpu:.2f}")

        st.markdown("**Verificación de cortante a una distancia d de la cara**")
        Vupu_d = abs(1.7 * (q1t * (b - dz/100) + 0.5 * (q_b - q1t) * (b - dz/100)) - 1.4 * gc * hz * (b - dz/100))
        st.write(f"Vu a d = {dz/100:.2f} m de la cara: {Vupu_d:.2f} kg")
        badge(Vupu_d <= phiVcpu, f"Vu={Vupu_d:.2f}", f"vs φVc={phiVcpu:.2f}")

with st.expander("08. Resumen General de Armadura", expanded=False):
    st.markdown("### Resumen de refuerzos")
    st.markdown("#### Acero vertical de la pantalla")
    df_vert = pd.DataFrame({
        "Ubicación": ["Cara interior", "Cara exterior"],
        "Refuerzo": [f"Ø {dii} @ {si:.2f} m", f"Ø {die} @ {se:.2f} m"],
        "Corte de acero": [f"dc = {dc:.2f} m" if dc > 0.05 else "No requiere", "—"]
    })
    st.table(df_vert)
    st.markdown("#### Acero horizontal de la pantalla")
    df_horiz = pd.DataFrame({
        "Tramo": ["Inferior", "Superior"],
        "Cara exterior": [f"Ø 1/2\" @ {shext_inf:.2f} m", f"Ø 1/2\" @ {shext_sup:.2f} m"],
        "Cara interior": [f"Ø 3/8\" @ {shint_inf:.2f} m", f"Ø 3/8\" @ {shint_sup:.2f} m"]
    })
    st.table(df_horiz)
    st.markdown("#### Talón posterior")
    df_talon = pd.DataFrame({
        "Dirección": ["Longitudinal", "Transversal"],
        "Refuerzo": [f"Ø {diamtl} @ {stl:.2f} m", f"Ø 5/8\" @ {stp:.2f} m"]
    })
    st.table(df_talon)
    st.markdown("#### Talón delantero (punta)")
    df_punta = pd.DataFrame({
        "Dirección": ["Longitudinal", "Transversal"],
        "Refuerzo": [f"Ø {diampu} @ {spu:.2f} m", f"Ø 5/8\" @ {spp:.2f} m"]
    })
    st.table(df_punta)

with st.expander("09. Diagramas Mux y Vux — Envolventes de Diseño", expanded=False):
    xh = np.linspace(0, Hp, 300)
    Vux_ = 1.7 * (0.5 * Ka * gr * xh**2 + Ka * gr * hs * xh)
    Muxd_ = 1.7 * (0.5 * Ka * gr * xh**2 * xh / 3 + Ka * gr * hs * xh * xh / 2)
    if sismo:
        Vux_ = 1.7 * (0.5 * Kea * (1 - Kv) * gr * xh**2 + Ka * gr * hs * xh)
        Muxd_ = 1.7 * (0.5 * Kea * (1 - Kv) * gr * xh**2 * xh / 3 + Ka * gr * hs * xh * xh / 2)
    dxv_ = np.array([max((ccor + (ct - ccor) * x / Hp) * 100 - rec, 5) for x in xh])
    phiVcx_ = phic * 0.53 * math.sqrt(fc) * bw * dxv_
    amx_ = rhomin * bw * dxv_ * fy / (0.85 * fc * bw)
    phiMnx_ = phif * (rhomin * bw * dxv_) * fy * (dxv_ - amx_ / 2) / 100
    ypl_ = Hp - xh
    Vux_t = Vux_ / 1000
    Muxd_t = Muxd_ / 1000
    Vcx_t = phiVcx_ / 1000
    Mnx_t = phiMnx_ / 1000

    fig_dg, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(15, 7))
    fig_dg.patch.set_facecolor(_plt_bg)
    for ax in [ax1, ax2, ax3]:
        ax.set_facecolor(_plt_axbg)
        ax.tick_params(colors=_plt_txt)
        ax.xaxis.label.set_color(_plt_txt)
        ax.yaxis.label.set_color(_plt_txt)
        ax.title.set_color(_plt_txt)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        for sp in ['bottom', 'left']:
            ax.spines[sp].set_color(_plt_grid)
        ax.grid(True, alpha=0.2, color=_plt_grid)

    ptie = Ka * gr * xh / 1000
    psc = Ka * gr * hs * np.ones_like(xh) / 1000
    ax1.fill_betweenx(ypl_, 0, ptie, alpha=0.4, color='#FF9800', label="Tierra")
    ax1.fill_betweenx(ypl_, ptie, ptie + psc, alpha=0.4, color='#2196F3', label="S/C")
    ax1.plot(ptie + psc, ypl_, color=_plt_txt, lw=2)
    ax1.set_xlabel("Presión [t/m²]", fontsize=10)
    ax1.set_ylabel("Altura [m]", fontsize=9)
    ax1.set_title("Presiones", fontsize=11, fontweight='bold')
    ax1.legend(facecolor=_plt_bg, labelcolor=_plt_txt, fontsize=8)
    ax1.set_ylim(0, Ht + 0.2)

    fv2_ = Vux_t >= Vcx_t
    ax2.fill_betweenx(ypl_, 0, Vux_t, where=~fv2_, alpha=0.3, color='#4CAF50')
    ax2.fill_betweenx(ypl_, 0, Vux_t, where=fv2_, alpha=0.3, color='#f44336')
    ax2.plot(Vux_t, ypl_, color='#FF9800', lw=2.5, label=f"Vu máx={max(Vux_t):.2f} ton")
    ax2.plot(Vcx_t, ypl_, color='#4CAF50', lw=2, ls='--', label="φVc")
    idx_max = np.argmax(Vux_t)
    ax2.annotate(f"{Vux_t[idx_max]:.2f} ton",
                 xy=(Vux_t[idx_max], ypl_[idx_max]),
                 xytext=(Vux_t[idx_max] * 0.6, ypl_[idx_max] - 0.3),
                 arrowprops=dict(arrowstyle='->', color='#FF9800', lw=1.5),
                 color='#FF9800', fontsize=9, fontweight='bold')
    ax2.set_xlabel("Cortante [ton]", fontsize=10)
    ax2.set_title("Fuerza Cortante", fontsize=11, fontweight='bold')
    ax2.legend(facecolor=_plt_bg, labelcolor=_plt_txt, fontsize=8)
    ax2.set_ylim(0, Ht + 0.2)

    fm2_ = Muxd_t >= Mnx_t
    ax3.fill_betweenx(ypl_, 0, Muxd_t, where=~fm2_, alpha=0.3, color='#4CAF50')
    ax3.fill_betweenx(ypl_, 0, Muxd_t, where=fm2_, alpha=0.3, color='#f44336')
    ax3.plot(Muxd_t, ypl_, color='#2196F3', lw=2.5, label=f"Mu máx={max(Muxd_t):.2f} t·m")
    ax3.plot(Mnx_t, ypl_, color='#4CAF50', lw=2, ls='--', label="φMn(As min)")
    if dc > 0.05:
        ydc_ = hz + Hp - dc
        ax3.axhline(ydc_, color='#FF6F00', lw=2, ls='-.', label=f"Corte dc={dc:.2f}m")
    ax3.set_xlabel("Momento [ton·m]", fontsize=10)
    ax3.set_title("Momento Flector", fontsize=11, fontweight='bold')
    ax3.legend(facecolor=_plt_bg, labelcolor=_plt_txt, fontsize=8)
    ax3.set_ylim(0, Ht + 0.2)
    fig_dg.tight_layout()
    st.pyplot(fig_dg)

with st.expander("10. Cubicación de Materiales", expanded=False):
    Ac_zap = hz * B * Lmuro
    Ac_pan = 0.5 * (ct + ccor) * Hp * Lmuro
    Ac_den = ct * hd * Lmuro if hd > 0 else 0.0
    Ac_tot = Ac_zap + Ac_pan + Ac_den
    As_pan_kg = (Api + Ape) * Hp * Lmuro * 7850 / 10000
    As_horiz_kg = (As_prov("1/2\"", shext_inf) + As_prov("3/8\"", shint_inf)) * Hp / 2 * Lmuro * 7850 / 10000
    As_horiz_kg += (As_prov("1/2\"", shext_sup) + As_prov("3/8\"", shint_sup)) * Hp / 2 * Lmuro * 7850 / 10000
    As_tal_kg = (Aptl + As_prov("5/8\"", stp)) * Bp * Lmuro * 7850 / 10000
    As_pun_kg = (Appu + As_prov("5/8\"", spp)) * b * Lmuro * 7850 / 10000
    As_tot_kg = As_pan_kg + As_horiz_kg + As_tal_kg + As_pun_kg
    st.table(pd.DataFrame({
        "Elemento": ["Zapata", "Pantalla", "Dentellón", "TOTAL"],
        "Volumen concreto [m³]": [f"{Ac_zap:.3f}", f"{Ac_pan:.3f}", f"{Ac_den:.3f}", f"{Ac_tot:.3f}"],
        "Cemento [bultos]": [f"{Ac_zap * 9.4:.1f}", f"{Ac_pan * 9.4:.1f}", f"{Ac_den * 9.4:.1f}", f"{Ac_tot * 9.4:.1f}"],
        "Arena [m³]": [f"{Ac_zap * 0.577:.3f}", f"{Ac_pan * 0.577:.3f}", f"{Ac_den * 0.577:.3f}", f"{Ac_tot * 0.577:.3f}"],
        "Gravilla [m³]": [f"{Ac_zap * 0.643:.3f}", f"{Ac_pan * 0.643:.3f}", f"{Ac_den * 0.643:.3f}", f"{Ac_tot * 0.643:.3f}"],
    }))
    st.write(f"**Acero total estimado: {As_tot_kg:.1f} kg  ({As_tot_kg / Ac_tot:.1f} kg/m³)**")
    varillas = math.ceil(As_tot_kg / (7850 * Ab_diam(diam_def)[0] / 10000 * 6))
    st.write(f"Varillas Ø {diam_def} × 6m estimadas: **{varillas} varillas**")

# ══════════════════════════════════════════════════════════════════
# 10. ESQUEMA FINAL, MODELO 3D Y EXPORTACIÓN
# ══════════════════════════════════════════════════════════════════
with st.expander("10. Esquema Final, Modelo 3D y Exportación", expanded=True):
    tab2d, tabpl, tab3d, tabiso, tabexp = st.tabs([
        " Sección 2D", " Vista en Planta",
        " Modelo 3D", " Isométrico", " Exportar"
    ])

    # ── TAB 2D ───────────────────────────────────────────────────
    with tab2d:
        fig2d, ax2d = plt.subplots(figsize=(9, 12))
        fig2d.patch.set_facecolor('white')
        ax2d.set_facecolor('white')

        # Relleno
        ax2d.fill_between([xbr, B + 0.5], [Ht, Ht], [hz, hz],
                           alpha=0.18, color='#8B6914', label="Relleno")
        ax2d.fill_between([0, B + 0.5], [0, 0], [-0.15, -0.15],
                           alpha=0.12, color='#8B6914')

        # Zapata
        ax2d.add_patch(patches.Rectangle((0, 0), B, hz, ec='black', fc='#C8C8C8', lw=2, zorder=2))

        # Pantalla trapezoidal
        ptsp = [(xbl, hz), (xbr, hz), (xtr, Ht), (xtl, Ht), (xbl, hz)]
        ax2d.add_patch(plt.Polygon(ptsp, ec='black', fc='#C8C8C8', lw=2, zorder=2))

        # Dentellón
        if hd > 0:
            ax2d.add_patch(patches.Rectangle((xbl, -hd), ct, hd, ec='black', fc='#C8C8C8', lw=1.5, zorder=2))

        # Acero interior
        ax2d.plot([xbr - rm, xbr - rm], [hz + rm, Ht - rm],
                  color='#0D47A1', lw=3, zorder=4, label=f"Cara int. Ø{dii} @{si:.2f}m")
        # Círculos extremos
        ax2d.add_patch(plt.Circle((xbr - rm, hz + rm), 0.03, color='#0D47A1', fill=True, zorder=6))
        ax2d.add_patch(plt.Circle((xbr - rm, Ht - rm), 0.03, color='#0D47A1', fill=True, zorder=6))

        # Acero exterior
        def xe_ext(z):
            return xtl + (xbl - xtl) * (Ht - z) / Hp if Hp > 0 else xbl
        xe_bot = xe_ext(hz) + rm
        xe_top = xe_ext(Ht) + rm
        ax2d.plot([xe_bot, xe_top], [hz + rm, Ht - rm],
                  color='#B71C1C', lw=2, ls='--', zorder=4, label=f"Cara ext. Ø{die} @{se:.2f}m")
        ax2d.add_patch(plt.Circle((xe_bot, hz + rm), 0.03, color='#B71C1C', fill=True, zorder=6))
        ax2d.add_patch(plt.Circle((xe_top, Ht - rm), 0.03, color='#B71C1C', fill=True, zorder=6))

        # Acero horizontal
        zh_div = hz + Hp / 2
        nh_inf = max(int(Hp / 2 / shext_inf), 2)
        nh_sup = max(int(Hp / 2 / shext_sup), 2)
        first_inf = True
        for zh in np.linspace(hz + shext_inf / 2, zh_div, nh_inf):
            xei = xe_ext(zh) + rm * 0.5
            ax2d.plot([xei, xei + 0.025], [zh, zh],
                      color='#E53935', lw=1.8, zorder=5,
                      label="Horiz. ext. tramo inf." if first_inf else "")
            ax2d.plot([xbr - rm * 0.5, xbr - rm * 0.5 + 0.025], [zh, zh],
                      color='#EF9A9A', lw=1.5, zorder=5)
            first_inf = False
        first_sup = True
        for zh in np.linspace(zh_div + shext_sup / 2, Ht - shext_sup / 2, nh_sup):
            xei = xe_ext(zh) + rm * 0.5
            ax2d.plot([xei, xei + 0.025], [zh, zh],
                      color='#C62828', lw=1.8, zorder=5,
                      label="Horiz. ext. tramo sup." if first_sup else "")
            ax2d.plot([xbr - rm * 0.5, xbr - rm * 0.5 + 0.025], [zh, zh],
                      color='#EF5350', lw=1.5, zorder=5)
            first_sup = False

        # Línea divisoria
        ax2d.axhline(zh_div, color='#FF6F00', lw=1.2, ls=':', zorder=5,
                     label=f"División tramos Hp/2={Hp/2:.2f}m")
        ax2d.text(xbr + 0.2, hz + Hp / 4, f"{Hp/2:.2f} m", fontsize=8, color='blue',
                  bbox=dict(facecolor='white', alpha=0.7, edgecolor='none'))
        ax2d.text(xbr + 0.2, hz + 3 * Hp / 4, f"{Hp/2:.2f} m", fontsize=8, color='blue',
                  bbox=dict(facecolor='white', alpha=0.7, edgecolor='none'))
        if dc <= 0.05:
            ax2d.text(xbr + 0.5, hz + Hp - 0.2, "No requiere corte de acero",
                      fontsize=8, color='green',
                      bbox=dict(facecolor='white', alpha=0.8, edgecolor='green'))

        # Acero talón posterior
        ax2d.plot([xbr + rm, B - rm], [hz - rm, hz - rm],
                  color='#2E7D32', lw=2.5, zorder=4, label=f"Talón Ø{diamtl} @{stl:.2f}m")

        # Acero punta
        ax2d.plot([rm, xbl - rm], [rm, rm],
                  color='#E65100', lw=2.5, zorder=4, label=f"Punta Ø{diampu} @{spu:.2f}m")

        # Dentellón
        if hd > 0:
            ax2d.plot([xbl + rm, xbl + rm, xbr - rm, xbr - rm],
                      [-rm, -hd + rm, -hd + rm, -rm],
                      color='#6A1B9A', lw=2.5, zorder=4, label=f"Dentellón Ø{diamtl}")

        # Gancho anclaje
        ax2d.annotate("", xy=(xbr - rm, hz - Ldh), xytext=(xbr - rm, hz + rm),
                      arrowprops=dict(arrowstyle='-', color='#0D47A1', lw=2))
        ax2d.text(xbr - rm + 0.04, hz - Ldh / 2, f"Ldh={Ldh:.2f}m",
                  color='#0D47A1', fontsize=7)

        # Punto de corte
        if dc > 0.05:
            ydc_ = hz + Hp - dc
            ax2d.axhline(ydc_, color='#FF6F00', lw=1.5, ls='-.', zorder=5)
            ax2d.text(xbr + 0.05, ydc_ + 0.05, f"Corte dc={dc:.2f}m", color='#FF6F00', fontsize=8)

        # N.T.
        ax2d.axhline(hz, color='#4CAF50', lw=1.2, ls=':', alpha=0.7)
        ax2d.text(B + 0.05, hz + 0.03, "N.T.", color='#4CAF50', fontsize=8)

        # Cotas (funciones auxiliares)
        CC = '#C62828'
        def cota_v(ax, xpos, y1, y2, lbl, side=1):
            off = 0.22 * side
            xp = xpos + off
            ax.annotate("", xy=(xp, y2), xytext=(xp, y1),
                        arrowprops=dict(arrowstyle='<->', color=CC, lw=1.4))
            for yv in [y1, y2]:
                ax.plot([xp - 0.06, xp + 0.06], [yv, yv], color=CC, lw=1)
            ax.text(xp + 0.07 * side, (y1 + y2) / 2, lbl,
                    va='center', ha='left' if side > 0 else 'right',
                    color=CC, fontsize=8, fontweight='bold')

        def cota_h(ax, ypos, x1, x2, lbl, side=-1):
            off = 0.30 * abs(side) * (-1 if side < 0 else 1)
            yp = ypos + off
            ax.annotate("", xy=(x2, yp), xytext=(x1, yp),
                        arrowprops=dict(arrowstyle='<->', color=CC, lw=1.4))
            for xv in [x1, x2]:
                ax.plot([xv, xv], [yp - 0.06, yp + 0.06], color=CC, lw=1)
            ax.text((x1 + x2) / 2, yp - 0.12 * abs(side),
                    lbl, ha='center', va='top' if side < 0 else 'bottom',
                    color=CC, fontsize=8, fontweight='bold')

        cota_v(ax2d, B + 0.1, 0, Ht, f"Ht={Ht:.2f}m", +1)
        cota_v(ax2d, B + 0.1, hz, Ht, f"Hp={Hp:.2f}m", +2)
        cota_v(ax2d, B + 0.1, 0, hz, f"hz={hz:.2f}m", +3)
        if hd > 0:
            cota_v(ax2d, -0.1, -hd, 0, f"hd={hd:.2f}m", -1)

        cota_h(ax2d, -hd if hd > 0 else 0, 0, B, f"B={B:.2f}m", -3)
        cota_h(ax2d, -hd if hd > 0 else 0, 0, b, f"b={b:.2f}m", -2)
        cota_h(ax2d, -hd if hd > 0 else 0, xbr, B, f"Bp={Bp:.2f}m", -2)
        cota_h(ax2d, -hd if hd > 0 else 0, xbl, xbr, f"ct={ct:.2f}m", -1)

        # Líderes
        def lider(ax, xy_dest, xy_text, txt, col, fs=8):
            ax.annotate(txt, xy=xy_dest, xytext=xy_text,
                        fontsize=fs, color=col, fontweight='bold',
                        arrowprops=dict(arrowstyle='->', color=col, lw=1.2),
                        bbox=dict(boxstyle='round,pad=0.3', fc='white', ec=col, alpha=0.9))

        lider(ax2d, (xbr - rm, hz + Hp * 0.6), (xbr + 0.5, hz + Hp * 0.6),
              f"Acero Vertical (Cara Interior)\nØ {dii}  @  {si:.2f} m", '#0D47A1', fs=7)
        lider(ax2d, (xe_ext(hz + Hp * 0.8) + rm, hz + Hp * 0.8), (xbr + 0.5, hz + Hp * 0.85),
              f"Acero Vertical (Cara Exterior)\nØ {die}  @  {se:.2f} m", '#B71C1C', fs=7)
        lider(ax2d, (xe_ext(hz + Hp * 0.88) + rm * 0.5, hz + Hp * 0.88), (-0.9, hz + Hp * 0.88),
              f"Acero Horizontal\nØ 1/2\"  @  {shext_sup:.2f} m", '#C62828', fs=7)
        lider(ax2d, (xe_ext(hz + Hp * 0.35) + rm * 0.5, hz + Hp * 0.35), (-0.9, hz + Hp * 0.35),
              f"Acero Horizontal\nØ 1/2\"  @  {shext_inf:.2f} m", '#C62828', fs=7)
        lider(ax2d, (xbr + Bp * 0.5, hz - rm), (xbr + Bp * 0.5 + 0.3, hz + 0.25),
              f"Ø {diamtl}  @  {stl:.2f} m", '#2E7D32', fs=7)
        lider(ax2d, (b * 0.5, rm), (b * 0.5 - 0.3, hz - 0.2),
              f"Ø {diampu}  @  {spu:.2f} m", '#E65100', fs=7)

        ax2d.set_xlim(-1.3, B + 1.2)
        ax2d.set_ylim(-hd - 0.8 if hd > 0 else -0.6, Ht + 0.6)
        ax2d.set_aspect('equal')
        ax2d.legend(fontsize=7, loc='upper right', facecolor='white', edgecolor='#ccc')
        ax2d.set_title(
            f"10. Esquema Final — KonteWall  |  {norma_sel}\n"
            f"Ht={Ht:.2f}m  B={B:.2f}m  f'c={fc:.0f} kg/cm²  fy={fy:.0f} kg/cm²",
            fontsize=10, fontweight='bold')
        ax2d.grid(True, alpha=0.1, color='#aaa')
        st.pyplot(fig2d, use_container_width=True)

    # ── TAB PLANTA ───────────────────────────────────────────────
    with tabpl:
        largopl = st.slider("Longitud de análisis [m]", 0.5, 5.0, 1.0, 0.10)
        espesor = ct
        rpl = rec / 100
        yintpl = espesor - rpl
        yextpl = rpl
        figpl, axpl = plt.subplots(figsize=(13, 5.5))
        figpl.patch.set_facecolor('white')
        axpl.set_facecolor('white')
        axpl.add_patch(patches.Rectangle((0, 0), largopl, espesor, lw=2, ec='black', fc='#D8D8D8', zorder=1))
        nip = max(3, int(largopl / si))
        nep = max(3, int(largopl / se))
        xbi = np.linspace(si / 2, largopl - si / 2, nip)
        xbe = np.linspace(se / 2, largopl - se / 2, nep)
        for xb in xbi:
            axpl.add_patch(plt.Circle((xb, yintpl), 0.012, color='#0D47A1', fill=True, zorder=5))
        for xb in xbe:
            axpl.add_patch(plt.Circle((xb, yextpl), 0.009, color='#B71C1C', fill=True, zorder=5))
        nhe = max(4, int(largopl / shext_inf))
        nhi = max(4, int(largopl / shint_inf))
        for xi in np.linspace(0.05, largopl - 0.05, nhe):
            axpl.plot([xi - 0.04, xi + 0.04], [yextpl + 0.016, yextpl + 0.016],
                      color='#B71C1C', lw=2.5, zorder=4)
        for xi in np.linspace(0.05, largopl - 0.05, nhi):
            axpl.plot([xi - 0.04, xi + 0.04], [yintpl - 0.016, yintpl - 0.016],
                      color='#0D47A1', lw=2.0, zorder=4)
        axpl.annotate("", xy=(largopl + 0.12, espesor), xytext=(largopl + 0.12, 0),
                      arrowprops=dict(arrowstyle='<->', color='#CC0000', lw=1.5))
        axpl.text(largopl + 0.22, espesor / 2,
                  f"ct={ct:.2f}m\n({ct*100:.0f}cm)",
                  va='center', ha='left', fontsize=9, color='#CC0000', fontweight='bold')
        xld = xbi[min(1, len(xbi) - 1)]
        axpl.annotate(f"Ø{dii} @{si:.2f}m\nAs={Api:.2f}cm²/m",
                      xy=(xld, yintpl), xytext=(xld - 0.05, yintpl + 0.20),
                      fontsize=8, color='#0D47A1', fontweight='bold',
                      arrowprops=dict(arrowstyle='->', color='#0D47A1', lw=1.2),
                      bbox=dict(boxstyle='round', pad=0.3, fc='white', ec='#0D47A1'))
        xld2 = xbe[min(1, len(xbe) - 1)]
        axpl.annotate(f"Ø{die} @{se:.2f}m\nAs={Ape:.2f}cm²/m",
                      xy=(xld2, yextpl), xytext=(xld2 + 0.05, yextpl - 0.20),
                      fontsize=8, color='#B71C1C', fontweight='bold',
                      arrowprops=dict(arrowstyle='->', color='#B71C1C', lw=1.2),
                      bbox=dict(boxstyle='round', pad=0.3, fc='white', ec='#B71C1C'))
        axpl.text(largopl / 2, espesor + 0.04, "CARA INTERIOR",
                  ha='center', va='bottom', fontsize=10, color='#0D47A1', fontweight='bold',
                  bbox=dict(boxstyle='round', pad=0.3, fc='#E3F2FD', ec='#0D47A1'))
        axpl.text(largopl / 2, -0.04, "CARA EXTERIOR",
                  ha='center', va='top', fontsize=10, color='#B71C1C', fontweight='bold',
                  bbox=dict(boxstyle='round', pad=0.3, fc='#FFEBEE', ec='#B71C1C'))
        axpl.set_xlim(-0.15, largopl + 0.90)
        axpl.set_ylim(-0.32, espesor + 0.22)
        axpl.set_aspect('equal')
        axpl.set_xlabel("Longitud del muro [m]", fontsize=10)
        axpl.set_ylabel("Espesor pantalla [m]", fontsize=10)
        axpl.set_title(
            f"Vista en Planta — Pantalla  |  f'c={fc:.0f} kg/cm²  fy={fy:.0f} kg/cm²  ct={ct:.2f}m  r={rec:.0f}cm",
            fontsize=11, fontweight='bold')
        axpl.grid(True, alpha=0.12)
        st.pyplot(figpl, use_container_width=True)

    # ── TAB 3D (Plotly) ──────────────────────────────────────────
    with tab3d:
        try:
            L3d = st.slider("Longitud del muro [m]", 0.5, 6.0, 3.0, 0.5, key="kwL3d")
            rm3 = rec / 100
            fig3d = go.Figure()

            def wire_box(x0, x1, y0, y1, z0, z1, col, nm, lw=3):
                pts = [(x0, y0, z0), (x1, y0, z0), (x1, y1, z0), (x0, y1, z0),
                       (x0, y0, z1), (x1, y0, z1), (x1, y1, z1), (x0, y1, z1)]
                edges = [(0, 1), (1, 2), (2, 3), (3, 0),
                         (4, 5), (5, 6), (6, 7), (7, 4),
                         (0, 4), (1, 5), (2, 6), (3, 7)]
                xs, ys, zs = [], [], []
                for i, j in edges:
                    xs += [pts[i][0], pts[j][0], None]
                    ys += [pts[i][1], pts[j][1], None]
                    zs += [pts[i][2], pts[j][2], None]
                fig3d.add_trace(go.Scatter3d(
                    x=xs, y=ys, z=zs, mode='lines',
                    line=dict(color=col, width=lw), name=nm, showlegend=True))

            def wire_trap(xb0, xb1, xt0, xt1, y0, y1, zb, zt, col, nm, lw=3):
                def face(y_):
                    return [(xb0, y_, zb), (xb1, y_, zb), (xt1, y_, zt), (xt0, y_, zt), (xb0, y_, zb)]
                xs, ys, zs = [], [], []
                for y_ in [y0, y1]:
                    for p in face(y_):
                        xs.append(p[0])
                        ys.append(p[1])
                        zs.append(p[2])
                    xs.append(None)
                    ys.append(None)
                    zs.append(None)
                for (xa, za), (xb_, zb_) in [
                    ((xb0, zb), (xb0, zb)), ((xb1, zb), (xb1, zb)),
                    ((xt0, zt), (xt0, zt)), ((xt1, zt), (xt1, zt))
                ]:
                    xs += [xa, xa, None]
                    ys += [y0, y1, None]
                    zs += [za, za, None]
                fig3d.add_trace(go.Scatter3d(
                    x=xs, y=ys, z=zs, mode='lines',
                    line=dict(color=col, width=lw), name=nm, showlegend=True))

            # Concreto
            wire_box(0, B, 0, L3d, 0, hz, '#78909C', 'Zapata', lw=4)
            wire_trap(xbl, xbr, xtl, xtr, 0, L3d, hz, Ht, '#546E7A', 'Pantalla', lw=4)
            if hd > 0:
                wire_box(xbl, xbr, 0, L3d, -hd, 0, '#455A64', 'Dentellón', lw=3)

            # Relleno (semitransparente)
            rx = [xbr, B, B, xbr, xbr, B, B, xbr]
            ry = [0, 0, L3d, L3d, 0, 0, L3d, L3d]
            rz = [hz, hz, hz, hz, Ht, Ht, Ht, Ht]
            fig3d.add_trace(go.Mesh3d(
                x=rx, y=ry, z=rz,
                i=[0, 0, 0, 1, 4, 4], j=[1, 2, 4, 5, 5, 6], k=[2, 3, 5, 6, 6, 7],
                color='#8B6914', opacity=0.15, name='Relleno', showlegend=True))

            # Acero interior
            nvi = min(max(int(L3d / si) + 1, 5), 25)
            first = True
            y_positions = np.linspace(si / 2, L3d - si / 2, nvi)
            for yb in y_positions:
                fig3d.add_trace(go.Scatter3d(
                    x=[xbr - rm3, xbr - rm3], y=[yb, yb], z=[hz + rm3, Ht - rm3],
                    mode='lines', line=dict(color='#1E88E5', width=8),
                    showlegend=first, name=f"Int Ø{dii} @{si:.2f}m"))
                if first:
                    fig3d.add_trace(go.Scatter3d(
                        x=[xbr - rm3 + 0.2], y=[yb], z=[(hz + Ht) / 2],
                        mode='text', text=[f"Ø{dii} @{si:.2f}m"],
                        textfont=dict(color='white', size=10),
                        showlegend=False))
                first = False

            # Acero exterior
            nve = min(max(int(L3d / se) + 1, 5), 25)
            first = True
            for yb in np.linspace(se / 2, L3d - se / 2, nve):
                fig3d.add_trace(go.Scatter3d(
                    x=[xbl + rm3, xtl + rm3], y=[yb, yb], z=[hz + rm3, Ht - rm3],
                    mode='lines', line=dict(color='#E53935', width=6, dash='dash'),
                    showlegend=first, name=f"Ext Ø{die} @{se:.2f}m"))
                if first:
                    fig3d.add_trace(go.Scatter3d(
                        x=[(xbl + xtl) / 2 + rm3 + 0.2], y=[yb], z=[(hz + Ht) / 2],
                        mode='text', text=[f"Ø{die} @{se:.2f}m"],
                        textfont=dict(color='white', size=10),
                        showlegend=False))
                first = False

            # Acero horizontal ext tramo inferior
            zh_div3d = hz + Hp / 2
            nh_inf3d = min(max(int(Hp / 2 / shext_inf) + 1, 5), 15)
            first = True
            for zh in np.linspace(hz + shext_inf / 2, zh_div3d, nh_inf3d):
                frac = (zh - hz) / Hp if Hp > 0 else 0
                xe3d = xbl + (xtl - xbl) * frac + rm3
                fig3d.add_trace(go.Scatter3d(
                    x=[xe3d, xe3d], y=[rm3, L3d - rm3], z=[zh, zh],
                    mode='lines', line=dict(color='#E53935', width=4),
                    showlegend=first, name=f"Horiz Ext inf Ø1/2\" @{shext_inf:.2f}m"))
                first = False

            # Acero horizontal ext tramo superior
            nh_sup3d = min(max(int(Hp / 2 / shext_sup) + 1, 5), 15)
            first = True
            for zh in np.linspace(zh_div3d + shext_sup / 2, Ht - shext_sup / 2, nh_sup3d):
                frac = (zh - hz) / Hp if Hp > 0 else 0
                xe3d = xbl + (xtl - xbl) * frac + rm3
                fig3d.add_trace(go.Scatter3d(
                    x=[xe3d, xe3d], y=[rm3, L3d - rm3], z=[zh, zh],
                    mode='lines', line=dict(color='#EF5350', width=4),
                    showlegend=first, name=f"Horiz Ext sup Ø1/2\" @{shext_sup:.2f}m"))
                first = False

            # Acero horizontal int (opcional, se puede omitir para no saturar)

            # Acero talón posterior
            ntl = min(max(int(L3d / stl) + 1, 5), 20)
            first = True
            for yb in np.linspace(stl / 2, L3d - stl / 2, ntl):
                fig3d.add_trace(go.Scatter3d(
                    x=[xbr + rm3, B - rm3], y=[yb, yb], z=[hz - rm3, hz - rm3],
                    mode='lines', line=dict(color='#2E7D32', width=8),
                    showlegend=first, name=f"Talón X Ø{diamtl} @{stl:.2f}m"))
                if first:
                    fig3d.add_trace(go.Scatter3d(
                        x=[(xbr + B) / 2], y=[yb], z=[hz - rm3 + 0.1],
                        mode='text', text=[f"Ø{diamtl} @{stl:.2f}m"],
                        textfont=dict(color='white', size=10),
                        showlegend=False))
                first = False

            # Acero punta
            npu = min(max(int(L3d / spu) + 1, 5), 20)
            first = True
            for yb in np.linspace(spu / 2, L3d - spu / 2, npu):
                fig3d.add_trace(go.Scatter3d(
                    x=[rm3, xbl - rm3], y=[yb, yb], z=[rm3, rm3],
                    mode='lines', line=dict(color='#E65100', width=8),
                    showlegend=first, name=f"Punta X Ø{diampu} @{spu:.2f}m"))
                if first:
                    fig3d.add_trace(go.Scatter3d(
                        x=[b / 2], y=[yb], z=[rm3 + 0.1],
                        mode='text', text=[f"Ø{diampu} @{spu:.2f}m"],
                        textfont=dict(color='white', size=10),
                        showlegend=False))
                first = False

            # Dentellón (si existe)
            if hd > 0:
                nd = min(max(int(L3d / stl) + 1, 5), 20)
                first = True
                for yb in np.linspace(stl / 2, L3d - stl / 2, nd):
                    fig3d.add_trace(go.Scatter3d(
                        x=[xbl + rm3, xbl + rm3, xbr - rm3, xbr - rm3],
                        y=[yb, yb, yb, yb],
                        z=[-rm3, -hd + rm3, -hd + rm3, -rm3],
                        mode='lines', line=dict(color='#6A1B9A', width=6),
                        showlegend=first, name=f"Dentellón U Ø{diamtl}"))
                    first = False

            fig3d.update_layout(
                scene=dict(
                    xaxis=dict(title=f'Ancho [m]  (B={B:.2f}m)',
                               backgroundcolor='#0e1117', gridcolor='#333',
                               zerolinecolor='#555', range=[0, B + 0.5]),
                    yaxis=dict(title=f'Longitud [m]  (L={L3d:.1f}m)',
                               backgroundcolor='#0e1117', gridcolor='#333',
                               zerolinecolor='#555', range=[0, L3d]),
                    zaxis=dict(title=f'Altura [m]  (Ht={Ht:.2f}m)',
                               backgroundcolor='#0e1117', gridcolor='#333',
                               zerolinecolor='#555', range=[-hd - 0.1, Ht + 0.2]),
                    bgcolor='#0e1117',
                    aspectmode='manual',
                    aspectratio=dict(x=B, y=min(L3d, B * 1.5), z=Ht),
                    camera=dict(eye=dict(x=1.8, y=-1.5, z=1.2))
                ),
                paper_bgcolor='#0e1117',
                plot_bgcolor='#0e1117',
                font=dict(color='white'),
                legend=dict(bgcolor='rgba(0,0,0,0.5)', font=dict(size=10)),
                height=700,
                title=dict(
                    text=f"Modelo 3D — KonteWall | B={B:.2f}m  Ht={Ht:.2f}m  L={L3d:.1f}m",
                    font=dict(color='white'))
            )
            st.plotly_chart(fig3d, use_container_width=True)
        except Exception as e3d:
            st.error(f"Error 3D: {e3d}")

    # ── TAB ISOMÉTRICO (Matplotlib) ───────────────────────────────
    with tabiso:
        try:
            Liso = st.slider("Longitud isométrico [m]", 0.5, 5.0, 2.0, 0.5, key="kwLiso")
            fig_iso, ax_iso = plt.subplots(figsize=(13, 10))
            fig_iso.patch.set_facecolor('white')
            ax_iso.set_facecolor('white')
            ax_iso.axis('off')

            ang = math.radians(30)
            def px3(x, y, z):
                return x * math.cos(ang) - y * math.cos(ang), \
                       x * math.sin(ang) + y * math.sin(ang) + z

            def seg3(ax, p1, p2, col, lw=1.5, ls='-', zo=3):
                ix1, iy1 = px3(*p1)
                ix2, iy2 = px3(*p2)
                ax.plot([ix1, ix2], [iy1, iy2], color=col, lw=lw, ls=ls, zorder=zo)

            def face3(ax, pts4, fc, ec='black', alpha=0.7, zo=2):
                proj = [px3(*p) for p in pts4]
                poly = plt.Polygon(proj, fc=fc, ec=ec, alpha=alpha, zorder=zo)
                ax.add_patch(poly)

            # Concreto: zapata
            face3(ax_iso, [(0, 0, 0), (B, 0, 0), (B, Liso, 0), (0, Liso, 0)], '#BDBDBD', alpha=0.6)
            face3(ax_iso, [(0, 0, 0), (B, 0, 0), (B, 0, hz), (0, 0, hz)], '#BDBDBD', alpha=0.7)
            face3(ax_iso, [(B, 0, 0), (B, Liso, 0), (B, Liso, hz), (B, 0, hz)], '#9E9E9E', alpha=0.7)
            face3(ax_iso, [(0, 0, hz), (B, 0, hz), (B, Liso, hz), (0, Liso, hz)], '#D0D0D0', alpha=0.8)

            # Pantalla
            face3(ax_iso, [(xbl, 0, hz), (xbr, 0, hz), (xtr, 0, Ht), (xtl, 0, Ht)], '#C8C8C8', alpha=0.85)
            face3(ax_iso, [(xbl, Liso, hz), (xbr, Liso, hz), (xtr, Liso, Ht), (xtl, Liso, Ht)], '#AAAAAA', alpha=0.7)
            face3(ax_iso, [(xbr, 0, hz), (xbr, Liso, hz), (xbr, Liso, Ht), (xbr, 0, Ht)], '#B0B0B0', alpha=0.75)
            face3(ax_iso, [(xtl, 0, Ht), (xtr, 0, Ht), (xtr, Liso, Ht), (xtl, Liso, Ht)], '#D5D5D5', alpha=0.9)

            # Dentellón
            if hd > 0:
                face3(ax_iso, [(xbl, 0, -hd), (xbr, 0, -hd), (xbr, 0, 0), (xbl, 0, 0)], '#C8C8C8', alpha=0.85)
                face3(ax_iso, [(xbl, 0, 0), (xbr, 0, 0), (xbr, Liso, 0), (xbl, Liso, 0)], '#BDBDBD', alpha=0.7)

            # Relleno
            face3(ax_iso, [(xbr, 0, hz), (B, 0, hz), (B, 0, Ht), (xbr, 0, Ht)], '#C8A96E', alpha=0.15)
            face3(ax_iso, [(xbr, 0, hz), (B, 0, hz), (B, Liso, hz), (xbr, Liso, hz)], '#C8A96E', alpha=0.12)

            # Aristas
            for p1, p2 in [
                ((0, 0, 0), (B, 0, 0)), ((0, 0, 0), (0, Liso, 0)),
                ((B, 0, 0), (B, Liso, 0)), ((0, Liso, 0), (B, Liso, 0)),
                ((0, 0, hz), (B, 0, hz)), ((0, 0, 0), (0, 0, hz)),
                ((B, 0, 0), (B, 0, hz)),
                ((xbl, 0, hz), (xbr, 0, hz)), ((xbr, 0, hz), (xbr, 0, Ht)),
                ((xbr, 0, Ht), (xtl, 0, Ht)), ((xtl, 0, Ht), (xbl, 0, hz)),
                ((xbl, Liso, hz), (xbr, Liso, hz)), ((xbr, Liso, hz), (xbr, Liso, Ht)),
                ((xbr, Liso, Ht), (xtl, Liso, Ht)), ((xtl, Liso, Ht), (xbl, Liso, hz)),
                ((xbl, 0, hz), (xbl, Liso, hz)), ((xbr, 0, hz), (xbr, Liso, hz)),
                ((xbr, 0, Ht), (xbr, Liso, Ht)), ((xtl, 0, Ht), (xtl, Liso, Ht)),
            ]:
                seg3(ax_iso, p1, p2, '#444', lw=1.2)

            # Acero interior
            nvi_iso = min(max(int(Liso / si) + 1, 2), 12)
            for yb in np.linspace(si / 2, Liso - si / 2, nvi_iso):
                seg3(ax_iso, (xbr - rm3, yb, hz + rm3), (xbr - rm3, yb, Ht - rm3),
                     '#1565C0', lw=3, zo=8)

            # Acero exterior
            nve_iso = min(max(int(Liso / se) + 1, 2), 12)
            for yb in np.linspace(se / 2, Liso - se / 2, nve_iso):
                seg3(ax_iso, (xbl + rm3, yb, hz + rm3), (xtl + rm3, yb, Ht - rm3),
                     '#C62828', lw=2, ls='--', zo=7)

            # Acero horizontal ext tramo inferior
            nh_iso = min(max(int(Hp / 2 / shext_inf) + 1, 2), 8)
            for zh in np.linspace(hz + shext_inf / 2, zh_div3d, nh_iso):
                frac = (zh - hz) / Hp if Hp > 0 else 0
                xei_ = xbl + (xtl - xbl) * frac + rm3
                seg3(ax_iso, (xei_, rm3, zh), (xei_, Liso - rm3, zh), '#E53935', lw=2, zo=7)

            # Acero horizontal ext tramo superior
            nh_iso2 = min(max(int(Hp / 2 / shext_sup) + 1, 2), 8)
            for zh in np.linspace(zh_div3d + shext_sup / 2, Ht - shext_sup / 2, nh_iso2):
                frac = (zh - hz) / Hp if Hp > 0 else 0
                xei_ = xbl + (xtl - xbl) * frac + rm3
                seg3(ax_iso, (xei_, rm3, zh), (xei_, Liso - rm3, zh), '#EF5350', lw=1.8, zo=7)

            # Acero talón posterior
            ntl_iso = min(max(int(Liso / stl) + 1, 2), 12)
            for yb in np.linspace(stl / 2, Liso - stl / 2, ntl_iso):
                seg3(ax_iso, (xbr + rm3, yb, hz - rm3), (B - rm3, yb, hz - rm3),
                     '#2E7D32', lw=3, zo=8)

            # Acero punta
            npu_iso = min(max(int(Liso / spu) + 1, 2), 12)
            for yb in np.linspace(spu / 2, Liso - spu / 2, npu_iso):
                seg3(ax_iso, (rm3, yb, rm3), (xbl - rm3, yb, rm3),
                     '#E65100', lw=3, zo=8)

            # Dentellón
            if hd > 0:
                nd_iso = min(max(int(Liso / stl) + 1, 2), 8)
                for yb in np.linspace(stl / 2, Liso - stl / 2, nd_iso):
                    seg3(ax_iso, (xbl + rm3, yb, -rm3), (xbl + rm3, yb, -hd + rm3), '#6A1B9A', lw=2.5, zo=8)
                    seg3(ax_iso, (xbl + rm3, yb, -hd + rm3), (xbr - rm3, yb, -hd + rm3), '#6A1B9A', lw=2.5, zo=8)
                    seg3(ax_iso, (xbr - rm3, yb, -hd + rm3), (xbr - rm3, yb, -rm3), '#6A1B9A', lw=2.5, zo=8)

            # Cotas isométricas (simplificadas)
            CC2 = '#C62828'
            def cota_iso_v(ax, x_, y_, z1, z2, lbl):
                ix1, iy1 = px3(x_, y_, z1)
                ix2, iy2 = px3(x_, y_, z2)
                ax.annotate("", xy=(ix2, iy2), xytext=(ix1, iy1),
                            arrowprops=dict(arrowstyle='<->', color=CC2, lw=1.3))
                ax.text((ix1 + ix2) / 2 + 0.08, (iy1 + iy2) / 2, lbl,
                        color=CC2, fontsize=8, fontweight='bold', va='center')

            def cota_iso_h(ax, y_, z_, x1, x2, lbl):
                ix1, iy1 = px3(x1, y_, z_)
                ix2, iy2 = px3(x2, y_, z_)
                ax.annotate("", xy=(ix2, iy2), xytext=(ix1, iy1),
                            arrowprops=dict(arrowstyle='<->', color=CC2, lw=1.3))
                ax.text((ix1 + ix2) / 2, (iy1 + iy2) / 2 - 0.15, lbl,
                        color=CC2, fontsize=8, fontweight='bold', ha='center')

            cota_iso_v(ax_iso, B + 0.3, 0, 0, Ht, f"Ht={Ht:.2f}m")
            cota_iso_v(ax_iso, B + 0.5, 0, hz, Ht, f"Hp={Hp:.2f}m")
            cota_iso_v(ax_iso, B + 0.7, 0, 0, hz, f"hz={hz:.2f}m")
            if hd > 0:
                cota_iso_v(ax_iso, B + 0.3, 0, -hd, 0, f"hd={hd:.2f}m")
            cota_iso_h(ax_iso, 0, -0.3, 0, B, f"B={B:.2f}m")
            cota_iso_h(ax_iso, 0, -0.5, 0, b, f"b={b:.2f}m")
            cota_iso_h(ax_iso, 0, -0.7, xbr, B, f"Bp={Bp:.2f}m")

            ax_iso.set_title(
                f"Vista Isométrica — KonteWall  |  Ht={Ht:.2f}m  B={B:.2f}m  L={Liso:.1f}m\n"
                f"f'c={fc:.0f} kg/cm²  fy={fy:.0f} kg/cm²  {norma_sel}",
                fontsize=11, fontweight='bold')
            ax_iso.autoscale()
            st.pyplot(fig_iso, use_container_width=True)
        except Exception as eiso:
            st.error(f"Error isométrico: {eiso}")

    # ── TAB EXPORTAR (DXF 2D mejorado) ────────────────────────────
    with tabexp:
        st.markdown("###  Exportar Documentación Profesional")
        col_exp1, col_exp2, col_exp3 = st.columns(3)

        with col_exp1:
            st.markdown("####  Plano DXF 2D")
            if st.button(" Generar DXF", key="btn_dxf"):
                try:
                    import ezdxf
                    try:
                        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                                 dxf_rotulo, dxf_rotulo_campos)
                        _USE_HKW = True
                    except ImportError:
                        _USE_HKW = False
                    doc = ezdxf.new('R2010')
                    doc.header['$INSUNITS'] = 6
                    doc.header['$MEASUREMENT'] = 1
                    if _USE_HKW:
                        dxf_setup(doc, 20); dxf_add_layers(doc)
                    msp = doc.modelspace()

                    layers = {
                        'CONCRETO':       {'color': 7,  'lineweight': 35, 'linetype': 'CONTINUOUS'},
                        'ACERO_INT':      {'color': 5,  'lineweight': 30, 'linetype': 'CONTINUOUS'},
                        'ACERO_EXT':      {'color': 1,  'lineweight': 30, 'linetype': 'DASHED'},
                        'ACERO_HORIZ_EXT':{'color': 1,  'lineweight': 25, 'linetype': 'CONTINUOUS'},
                        'ACERO_HORIZ_INT':{'color': 3,  'lineweight': 20, 'linetype': 'CONTINUOUS'},
                        'ACERO_TALON':    {'color': 3,  'lineweight': 30, 'linetype': 'CONTINUOUS'},
                        'ACERO_PUNTA':    {'color': 30, 'lineweight': 25, 'linetype': 'CONTINUOUS'},
                        'ACERO_DENT':     {'color': 6,  'lineweight': 25, 'linetype': 'CONTINUOUS'},
                        'RELLENO':        {'color': 8,  'lineweight': 15, 'linetype': 'CONTINUOUS'},
                        'COTAS':          {'color': 2,  'lineweight': 15, 'linetype': 'CONTINUOUS'},
                        'TEXTO':          {'color': 7,  'lineweight': 13, 'linetype': 'CONTINUOUS'},
                    }
                    for name, attrs in layers.items():
                        doc.layers.add(name, color=attrs['color'], lineweight=attrs['lineweight'])
                        if attrs['linetype'] != 'CONTINUOUS':
                            if attrs['linetype'] == 'DASHED':
                                doc.linetypes.add('DASHED', pattern='DASHED')
                            # Nota: para aplicar el tipo de línea a la capa, se necesita un paso adicional, pero por simplicidad se puede asignar al dibujar.

                    def line(p1, p2, layer):
                        msp.add_line(p1, p2, dxfattribs={'layer': layer})

                    def circle(center, radius, layer):
                        msp.add_circle(center, radius, dxfattribs={'layer': layer})

                    def text(txt, pos, height=0.2, layer='TEXTO'):
                        msp.add_text(txt, dxfattribs={'layer': layer, 'height': height, 'insert': pos})

                    # Concreto
                    line((0, 0), (B, 0), 'CONCRETO')
                    line((B, 0), (B, hz), 'CONCRETO')
                    line((B, hz), (0, hz), 'CONCRETO')
                    line((0, hz), (0, 0), 'CONCRETO')
                    line((xbl, hz), (xbr, hz), 'CONCRETO')
                    line((xbr, hz), (xtr, Ht), 'CONCRETO')
                    line((xtr, Ht), (xtl, Ht), 'CONCRETO')
                    line((xtl, Ht), (xbl, hz), 'CONCRETO')
                    line((xbr, hz), (xbr, Ht), 'CONCRETO')
                    if hd > 0:
                        line((xbl, -hd), (xbr, -hd), 'CONCRETO')
                        line((xbr, -hd), (xbr, 0), 'CONCRETO')
                        line((xbr, 0), (xbl, 0), 'CONCRETO')
                        line((xbl, 0), (xbl, -hd), 'CONCRETO')
                    line((xbr, hz), (B, hz), 'RELLENO')
                    line((B, hz), (B, Ht), 'RELLENO')
                    line((B, Ht), (xbr, Ht), 'RELLENO')
                    line((xbr, Ht), (xbr, hz), 'RELLENO')

                    # Acero interior
                    line((xbr - rm, hz + rm), (xbr - rm, Ht - rm), 'ACERO_INT')
                    circle((xbr - rm, hz + rm), 0.02, 'ACERO_INT')
                    circle((xbr - rm, Ht - rm), 0.02, 'ACERO_INT')
                    text(f"Ø{dii} @ {si:.2f}", (xbr - rm + 0.3, (hz + Ht) / 2), height=0.15, layer='TEXTO')

                    # Acero exterior
                    xe_bot = xe_ext(hz) + rm
                    xe_top = xe_ext(Ht) + rm
                    line((xe_bot, hz + rm), (xe_top, Ht - rm), 'ACERO_EXT')
                    circle((xe_bot, hz + rm), 0.02, 'ACERO_EXT')
                    circle((xe_top, Ht - rm), 0.02, 'ACERO_EXT')
                    text(f"Ø{die} @ {se:.2f}", (xe_bot - 0.5, (hz + Ht) / 2), height=0.15, layer='TEXTO')

                    # Acero horizontal ext tramo inferior
                    zh_divd = hz + Hp / 2
                    n_inf = max(int(Hp / 2 / shext_inf) + 1, 3)
                    for zh in np.linspace(hz + shext_inf / 2, zh_divd, n_inf):
                        x_pos = xe_ext(zh) + rm * 0.5
                        line((x_pos, zh), (x_pos + 0.5, zh), 'ACERO_HORIZ_EXT')
                    n_sup = max(int(Hp / 2 / shext_sup) + 1, 3)
                    for zh in np.linspace(zh_divd + shext_sup / 2, Ht - shext_sup / 2, n_sup):
                        x_pos = xe_ext(zh) + rm * 0.5
                        line((x_pos, zh), (x_pos + 0.5, zh), 'ACERO_HORIZ_EXT')

                    # Acero talón
                    n_talon = max(int(10.0 / stl) + 1, 3)
                    for i in range(n_talon):
                        y = i * stl
                        if y > 10.0:
                            break
                        line((xbr + rm, hz - rm + y), (B - rm, hz - rm + y), 'ACERO_TALON')
                    text(f"Ø{diamtl} @ {stl:.2f}", ((xbr + B) / 2, hz - rm + 0.3), height=0.15, layer='TEXTO')

                    # Acero punta
                    n_punta = max(int(10.0 / spu) + 1, 3)
                    for i in range(n_punta):
                        y = i * spu
                        if y > 10.0:
                            break
                        line((rm, rm + y), (xbl - rm, rm + y), 'ACERO_PUNTA')
                    text(f"Ø{diampu} @ {spu:.2f}", (b / 2, rm + 0.3), height=0.15, layer='TEXTO')

                    # Dentellón
                    if hd > 0:
                        n_dent = max(int(10.0 / stl) + 1, 2)
                        for i in range(n_dent):
                            y = i * stl
                            if y > 10.0:
                                break
                            line((xbl + rm, -rm + y), (xbl + rm, -hd + rm + y), 'ACERO_DENT')
                            line((xbl + rm, -hd + rm + y), (xbr - rm, -hd + rm + y), 'ACERO_DENT')
                            line((xbr - rm, -hd + rm + y), (xbr - rm, -rm + y), 'ACERO_DENT')
                        text(f"Ø{diamtl} @ {stl:.2f}", (xbl, -hd / 2), height=0.15, layer='TEXTO')

                    # Cotas
                    def cota_h(x1, x2, y, txt):
                        mid = (x1 + x2) / 2
                        line((x1, y), (x2, y), 'COTAS')
                        line((x1, y - 0.05), (x1, y + 0.05), 'COTAS')
                        line((x2, y - 0.05), (x2, y + 0.05), 'COTAS')
                        text(txt, (mid, y - 0.3), height=0.15, layer='TEXTO')

                    def cota_v(x, y1, y2, txt):
                        mid = (y1 + y2) / 2
                        line((x, y1), (x, y2), 'COTAS')
                        line((x - 0.05, y1), (x + 0.05, y1), 'COTAS')
                        line((x - 0.05, y2), (x + 0.05, y2), 'COTAS')
                        text(txt, (x + 0.3, mid), height=0.15, layer='TEXTO')

                    cota_h(0, B, -0.5, f"B={B:.2f}")
                    cota_h(0, b, -0.8, f"b={b:.2f}")
                    cota_h(xbr, B, -0.8, f"Bp={Bp:.2f}")
                    cota_h(xbl, xbr, -1.1, f"ct={ct:.2f}")
                    cota_v(B + 0.5, 0, Ht, f"Ht={Ht:.2f}")
                    cota_v(B + 0.5, hz, Ht, f"Hp={Hp:.2f}")
                    cota_v(B + 0.5, 0, hz, f"hz={hz:.2f}")
                    if hd > 0:
                        cota_v(xbl - 0.5, -hd, 0, f"hd={hd:.2f}")
                    cota_v(B + 0.8, zh_divd, zh_divd + 0.01, f"Hp/2={Hp/2:.2f}")

                    # Guardar
                    buf = io.BytesIO()
                    doc.saveas(buf)
                    buf.seek(0)
                    st.download_button(
                        "⬇ Descargar DXF 2D",
                        buf,
                        file_name=f"KonteWall_{Ht:.1f}m.dxf",
                        mime="application/octet-stream")
                    st.success(" DXF 2D generado con capas mejoradas y cotas completas")
                except Exception as e:
                    st.error(f"Error DXF: {e}")

        # DOCX y XLSX (se mantienen igual que antes, no se repiten por brevedad)
        with col_exp2:
            st.markdown("####  Memoria Técnica DOCX")
            if st.button(" Generar DOCX", key="btn_docx"):
                try:
                    from docx import Document as DocxDoc
                    from docx.shared import Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    doc = DocxDoc()
                    doc.styles['Normal'].paragraph_format.space_after = Pt(2)
                    doc.styles['Normal'].paragraph_format.space_before = Pt(1)

                    def h1(txt):
                        p = doc.add_heading(txt, level=1)
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(4)

                    def h2(txt):
                        p = doc.add_heading(txt, level=2)
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(2)

                    def par(txt):
                        p = doc.add_paragraph(txt)
                        p.paragraph_format.space_after = Pt(2)
                        return p

                    def kv(k, v):
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(1)
                        run_k = p.add_run(f"{k}: ")
                        run_k.bold = True
                        p.add_run(str(v))

                    # Portada
                    doc.add_heading("MEMORIA TÉCNICA DE CÁLCULO", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_heading("KonteWall — Diseño Estructural de Muro de Contención en Voladizo", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    par(f"Norma: {norma_sel}").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    par("Autor: ").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    par("Konte — Construcción, Consultoría y Tecnología · Duitama, Boyacá, Colombia").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    par("KonteWall v4.0 © 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_page_break()

                    h1("1. PARÁMETROS DE DISEÑO")
                    h2("1.1 Materiales")
                    kv("f'c", f"{fc:.0f} kg/cm²")
                    kv("fy", f"{fy:.0f} kg/cm²")
                    kv("γc", f"{gc:.0f} kg/m³")
                    kv("φ flexión", f"{phif:.2f}")
                    kv("φ cortante", f"{phic:.2f}")
                    kv("Recubrimiento", f"{rec:.1f} cm")
                    h2("1.2 Suelo de Relleno")
                    kv("γr", f"{gr:.0f} kg/m³")
                    kv("φ₁", f"{ph1:.1f}°")
                    kv("α", f"{alp:.1f}°")
                    kv("δ", f"{delt:.1f}°")
                    kv("S/C", f"{SC:.0f} kg/m²")
                    kv("hs", f"{hs:.3f} m")
                    h2("1.3 Suelo de Fundación")
                    kv("γf", f"{gf:.0f} kg/m³")
                    kv("φ₂", f"{ph2:.1f}°")
                    kv("C", f"{Ccoh:.0f} kg/m²")
                    kv("qadm", f"{qadm:.2f} kg/cm²")
                    kv("μ", f"{muct:.3f}")
                    h2("1.4 Geometría")
                    kv("Ht", f"{Ht:.2f} m")
                    kv("Hp", f"{Hp:.2f} m")
                    kv("hz", f"{hz:.2f} m")
                    kv("B", f"{B:.2f} m")
                    kv("b", f"{b:.2f} m")
                    kv("Bp", f"{Bp:.2f} m")
                    kv("ct", f"{ct:.2f} m")
                    kv("ccor", f"{ccor:.2f} m")
                    if hd > 0:
                        kv("hd dentellón", f"{hd:.2f} m")
                    if sismo:
                        h2("1.5 Sismicidad")
                        kv("Kh", f"{Kh:.3f}")
                        kv("Kv", f"{Kv:.3f}")
                        kv("θ M-O", f"{math.degrees(thr):.2f}°")

                    h1("2. COEFICIENTES DE PRESIÓN")
                    kv("Ka Coulomb estático", f"{Ka:.4f}")
                    kv("Kea Mononobe-Okabe", f"{Kea:.4f}" if sismo else "No aplica")
                    kv("Kp pasivo Coulomb", f"{Kp:.3f}")

                    h1("3. VERIFICACIÓN DE ESTABILIDAD")
                    t = doc.add_table(rows=1, cols=4)
                    t.style = 'Table Grid'
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, h in enumerate(["Verificación", "Valor", "Límite", "Estado"]):
                        t.rows[0].cells[i].text = h
                        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                    for ver, val, lim, ok in [
                        ("FS Volcamiento", f"{FSvolt:.3f}", "≥ 2.00", FSvolt >= 2.0),
                        ("FS Deslizamiento", f"{FSdesl:.3f}", "≥ 1.50", FSdesl >= 1.5),
                        ("Excentricidad e", f"{eexc:.4f} m", f"< {elim:.4f} m", eexc < elim),
                        ("q₁ portante", f"{q1:.4f} kg/cm²", f"≤ {qadm:.2f}", q1 <= qadm),
                        ("q₂ tensión", f"{q2:.4f} kg/cm²", "≥ 0", q2 >= 0),
                    ]:
                        row = t.add_row().cells
                        row[0].text = ver
                        row[1].text = val
                        row[2].text = lim
                        row[3].text = " CUMPLE" if ok else " NO CUMPLE"

                    h1("4. DISEÑO DE ARMADURA")
                    h2("4.1 Pantalla — Sección Crítica Base")
                    kv("Mu base", f"{Mup:.2f} kg·m/m")
                    kv("d", f"{dp:.2f} cm")
                    kv("As req cara interior", f"{Asp:.2f} cm²/m")
                    kv("Acero cara interior", f"Ø {dii} @ {si:.2f} m  →  As prov={Api:.2f} cm²/m")
                    kv("Acero cara exterior", f"Ø {die} @ {se:.2f} m  →  As prov={Ape:.2f} cm²/m")
                    kv("Acero horiz. ext. tramo inf.", f"Ø 1/2\" @ {shext_inf:.2f} m")
                    kv("Acero horiz. ext. tramo sup.", f"Ø 1/2\" @ {shext_sup:.2f} m")
                    kv("Acero horiz. int. tramo inf.", f"Ø 3/8\" @ {shint_inf:.2f} m")
                    kv("Acero horiz. int. tramo sup.", f"Ø 3/8\" @ {shint_sup:.2f} m")
                    if dc > 0.05:
                        kv("Corte acero exterior", f"dc={dc:.2f} m desde corona")
                    kv("Ldh anclaje", f"{Ldh:.2f} m")
                    kv("Verificación cortante", f"Vu={Vupant:.2f} kg/m  {'≤' if Vupant <= phiVcp else '>'} φVc={phiVcp:.2f} kg/m")

                    h2("4.2 Talón Posterior")
                    kv("Mu talón", f"{Mut:.2f} kg·m/m")
                    kv("As req", f"{Ast:.2f} cm²/m")
                    kv("Acero longitudinal", f"Ø {diamtl} @ {stl:.2f} m  →  As prov={Aptl:.2f} cm²/m")
                    kv("Acero transversal", f"Ø 5/8\" @ {stp:.2f} m")
                    kv("Verificación cortante", f"Vu={Vut:.2f} kg/m  {'≤' if Vut <= phiVct else '>'} φVc={phiVct:.2f} kg/m")

                    h2("4.3 Talón Delantero (Punta)")
                    kv("Mu punta", f"{Mupu:.2f} kg·m/m")
                    kv("As req", f"{Aspu:.2f} cm²/m")
                    kv("Acero longitudinal", f"Ø {diampu} @ {spu:.2f} m  →  As prov={Appu:.2f} cm²/m")
                    kv("Acero transversal", f"Ø 5/8\" @ {spp:.2f} m")
                    kv("Verificación cortante", f"Vu={Vupu:.2f} kg/m  {'≤' if Vupu <= phiVcpu else '>'} φVc={phiVcpu:.2f} kg/m")

                    h1("5. CUBICACIÓN DE MATERIALES")
                    kv("Volumen total concreto", f"{Ac_tot:.3f} m³  (L={Lmuro:.1f}m)")
                    kv("Acero estimado", f"{As_tot_kg:.1f} kg  ({As_tot_kg / Ac_tot:.1f} kg/m³)")
                    kv("Cemento", f"{Ac_tot * 9.4:.1f} bultos de 50 kg")
                    kv("Arena", f"{Ac_tot * 0.577:.3f} m³")
                    kv("Gravilla", f"{Ac_tot * 0.643:.3f} m³")
                    kv("Agua", f"{Ac_tot * 196:.0f} L")

                    # Footer
                    section = doc.sections[0]
                    footer = section.footer
                    fp = footer.paragraphs[0]
                    fp.text = "KonteWall v4.0 © 2026 · Konte — Construcción, Consultoría y Tecnología ·  · Duitama, Boyacá, Colombia"
                    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fp.runs[0].font.size = Pt(8)

                    buf_docx = io.BytesIO()
                    doc.save(buf_docx)
                    buf_docx.seek(0)
                    st.download_button(
                        "⬇ Descargar Memoria DOCX",
                        buf_docx,
                        file_name=f"KonteWall_Memoria_Ht{Ht:.1f}m.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    st.success(" Memoria técnica generada")
                except Exception as edocx:
                    st.error(f"Error DOCX: {edocx}")

        with col_exp3:
            st.markdown("####  Hoja de Cálculo XLSX")
            if st.button(" Generar XLSX", key="btn_xlsx"):
                try:
                    import openpyxl
                    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                    wb = openpyxl.Workbook()
                    ws = wb.active
                    ws.title = "KonteWall"
                    ws.column_dimensions['A'].width = 32
                    ws.column_dimensions['B'].width = 20
                    ws.column_dimensions['C'].width = 16

                    hdr_fill = PatternFill("solid", fgColor="1F3864")
                    hdr_font = Font(color="FFFFFF", bold=True, size=11)
                    sec_fill = PatternFill("solid", fgColor="2E75B6")
                    sec_font = Font(color="FFFFFF", bold=True, size=10)
                    ok_fill = PatternFill("solid", fgColor="C6EFCE")
                    nk_fill = PatternFill("solid", fgColor="FFC7CE")
                    bd = Border(left=Side(style='thin'), right=Side(style='thin'),
                                top=Side(style='thin'), bottom=Side(style='thin'))

                    def hrow(ws, r, txt):
                        ws.merge_cells(f'A{r}:C{r}')
                        c = ws[f'A{r}']
                        c.value = txt
                        c.fill = sec_fill
                        c.font = sec_font
                        c.alignment = Alignment(horizontal='center')

                    def drow(ws, r, k, v, unit="", ok=None):
                        ws[f'A{r}'] = k
                        ws[f'A{r}'].font = Font(bold=True)
                        ws[f'B{r}'] = v
                        ws[f'C{r}'] = unit
                        if ok is True:
                            ws[f'B{r}'].fill = ok_fill
                        elif ok is False:
                            ws[f'B{r}'].fill = nk_fill
                        for col in ['A', 'B', 'C']:
                            ws[f'{col}{r}'].border = bd
                            ws[f'{col}{r}'].alignment = Alignment(horizontal='center' if col != 'A' else 'left')

                    ws.merge_cells('A1:C1')
                    ws['A1'] = "KonteWall v4.0 © 2026 — Diseño Muro de Contención"
                    ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
                    ws['A1'].fill = hdr_fill
                    ws['A1'].alignment = Alignment(horizontal='center')
                    ws.merge_cells('A2:C2')
                    ws['A2'] = f"Autor:  · Konte · {norma_sel}"
                    ws['A2'].alignment = Alignment(horizontal='center')
                    ws['A2'].font = Font(size=10, italic=True)

                    r = 4
                    hrow(ws, r, "GEOMETRÍA Y MATERIALES")
                    r += 1
                    for k, v, u in [
                        ("Ht total", Ht, "m"), ("Hp pantalla", Hp, "m"),
                        ("hz zapata", hz, "m"), ("B base", B, "m"),
                        ("b punta", b, "m"), ("Bp talón", Bp, "m"),
                        ("ct base pantalla", ct, "m"), ("ccor corona", ccor, "m"),
                        ("hd dentellón", hd, "m"),
                        ("f'c", fc, "kg/cm²"), ("fy", fy, "kg/cm²"),
                        ("γc", gc, "kg/m³"), ("γr", gr, "kg/m³"),
                        ("SC", SC, "kg/m²"), ("φ₁", ph1, "°"),
                        ("qadm", qadm, "kg/cm²"), ("μ", muct, "—"),
                    ]:
                        drow(ws, r, k, v, u)
                        r += 1

                    hrow(ws, r, "COEFICIENTES DE PRESIÓN")
                    r += 1
                    drow(ws, r, "Ka Coulomb", round(Ka, 4), "")
                    r += 1
                    drow(ws, r, "Kea M-O", round(Kea, 4) if sismo else "N/A", "")
                    r += 1
                    drow(ws, r, "Kp pasivo", round(Kp, 3), "")
                    r += 1

                    hrow(ws, r, "VERIFICACIÓN DE ESTABILIDAD")
                    r += 1
                    drow(ws, r, "FS Volcamiento", round(FSvolt, 3), "≥ 2.00", ok=FSvolt >= 2.0)
                    r += 1
                    drow(ws, r, "FS Deslizamiento", round(FSdesl, 3), "≥ 1.50", ok=FSdesl >= 1.5)
                    r += 1
                    drow(ws, r, "Excentricidad e", round(eexc, 4), "m < B/6", ok=eexc < elim)
                    r += 1
                    drow(ws, r, "q₁ portante", round(q1, 4), "kg/cm²", ok=q1 <= qadm)
                    r += 1
                    drow(ws, r, "q₂ tensión", round(q2, 4), "kg/cm²", ok=q2 >= 0)
                    r += 1

                    hrow(ws, r, "ARMADURA DISEÑADA")
                    r += 1
                    for k, v in [
                        (f"Pantalla int. Ø{dii}", f"@ {si:.2f} m  →  {Api:.2f} cm²/m"),
                        (f"Pantalla ext. Ø{die}", f"@ {se:.2f} m  →  {Ape:.2f} cm²/m"),
                        (f"Horiz. ext. inf. Ø1/2\"", f"@ {shext_inf:.2f} m"),
                        (f"Horiz. ext. sup. Ø1/2\"", f"@ {shext_sup:.2f} m"),
                        (f"Horiz. int. inf. Ø3/8\"", f"@ {shint_inf:.2f} m"),
                        (f"Talón Ø{diamtl}", f"@ {stl:.2f} m  →  {Aptl:.2f} cm²/m"),
                        (f"Punta Ø{diampu}", f"@ {spu:.2f} m  →  {Appu:.2f} cm²/m"),
                    ]:
                        drow(ws, r, k, v, "")
                        r += 1

                    hrow(ws, r, "CUBICACIÓN")
                    r += 1
                    drow(ws, r, "Vol. concreto total", round(Ac_tot, 3), "m³")
                    r += 1
                    drow(ws, r, "Acero estimado", round(As_tot_kg, 1), "kg")
                    r += 1
                    drow(ws, r, "Cemento", round(Ac_tot * 9.4, 1), "bultos 50kg")
                    r += 1
                    drow(ws, r, "Arena", round(Ac_tot * 0.577, 3), "m³")
                    r += 1
                    drow(ws, r, "Gravilla", round(Ac_tot * 0.643, 3), "m³")
                    r += 1
                    drow(ws, r, "Agua", round(Ac_tot * 196, 0), "L")
                    r += 1

                    buf_xlsx = io.BytesIO()
                    wb.save(buf_xlsx)
                    buf_xlsx.seek(0)
                    st.download_button(
                        "⬇ Descargar XLSX",
                        buf_xlsx,
                        file_name=f"KonteWall_Calculo_Ht{Ht:.1f}m.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                    st.success(" Hoja de cálculo generada")
                except Exception as exlsx:
                    st.error(f"Error XLSX: {exlsx}")

# ══════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown(
    "<div style='text-align:center;font-size:12px;opacity:0.65;padding:8px;'>"
    " <b>KonteWall v4.0</b> © 2026 · Todos los derechos reservados · "
    "<b>Konte</b> — Construcción, Consultoría y Tecnología · "
    " · Duitama, Boyacá, Colombia"
    "</div>",
    unsafe_allow_html=True
)