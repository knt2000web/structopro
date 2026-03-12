import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import math
import io
import ezdxf
from docx import Document

# ─────────────────────────────────────────────
# IDIOMA GLOBAL Y NORMA
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

st.set_page_config(page_title=_t("StructMaster 2D", "StructMaster 2D"), layout="wide")
st.title(_t("StructMaster 2D - Diseño Estructural Automatizado", "StructMaster 2D - Automated Structural Design"))
st.markdown(_t(f"Análisis Matricial Interactivo con Recomendaciones Automáticas de Acero y Cimentación según **{norma_sel}**.", 
               f"Interactive Matrix Analysis with Automated Steel and Foundation Recommendations per **{norma_sel}**."))

# ─────────────────────────────────────────────
# INIT DATAFRAMES
if "nudos_df" not in st.session_state:
    st.session_state.nudos_df = pd.DataFrame([{"ID":1,"X (m)":0.0,"Y (m)":0.0}, {"ID":2,"X (m)":0.0,"Y (m)":3.0}, {"ID":3,"X (m)":4.0,"Y (m)":3.0}, {"ID":4,"X (m)":4.0,"Y (m)":0.0}])
if "barras_df" not in st.session_state:
    st.session_state.barras_df = pd.DataFrame([
        {"Nudo I":1, "Nudo J":2, "E (MPa)":21000.0, "A (cm²)":1200.0, "I (cm⁴)":160000.0}, # Columna 30x40
        {"Nudo I":2, "Nudo J":3, "E (MPa)":21000.0, "A (cm²)":1200.0, "I (cm⁴)":250000.0}, # Viga 30x50
        {"Nudo I":3, "Nudo J":4, "E (MPa)":21000.0, "A (cm²)":1200.0, "I (cm⁴)":160000.0}  # Columna 30x40
    ])
if "apoyos_df" not in st.session_state:
    st.session_state.apoyos_df = pd.DataFrame([{"Nudo":1,"Fijo X":True,"Fijo Y":True,"Fijo Rz":True}, {"Nudo":4,"Fijo X":True,"Fijo Y":True,"Fijo Rz":True}])
if "cargas_nudo_df" not in st.session_state:
    st.session_state.cargas_nudo_df = pd.DataFrame([{"Nudo":2,"FX (kN)":20.0,"FY (kN)":-50.0,"Mz (kN-m)":0.0}, {"Nudo":3,"FX (kN)":0.0,"FY (kN)":-50.0,"Mz (kN-m)":0.0}])

# ─────────────────────────────────────────────
# TABS PRINCIPALES
tab_gr, tab_mat, tab_res, tab_dis = st.tabs([
    _t("1. Geometría", "1. Geometry"), 
    _t("2. Cargas y Suelo", "2. Loads & Soil"),
    _t("3. Fuerzas Internas", "3. Internal Forces"),
    _t("4. Diseño Automático", "4. Auto-Design")
])

# ─────────────────────────────────────────────
# FUNCION DIBUJO GEOMETRIA
def plot_frame(df_n, df_b, df_sup, df_load):
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=df_n['X (m)'], y=df_n['Y (m)'], mode='markers+text', text=df_n['ID'], textposition="top right", marker=dict(size=10, color='blue'), name="Nudos"))
    for idx, b in df_b.iterrows():
        try:
            n_i = df_n[df_n['ID'] == b['Nudo I']].iloc[0]; n_j = df_n[df_n['ID'] == b['Nudo J']].iloc[0]
            fig.add_trace(go.Scatter(x=[n_i['X (m)'], n_j['X (m)']], y=[n_i['Y (m)'], n_j['Y (m)']], mode='lines+text', text=[f"Bar {idx+1}", ""], textposition="middle right", line=dict(color='black', width=3), name=f"B{idx+1}"))
        except: pass
    for idx, s in df_sup.iterrows():
        try:
            n = df_n[df_n['ID'] == s['Nudo']].iloc[0]
            fig.add_trace(go.Scatter(x=[n['X (m)']], y=[n['Y (m)']], mode='markers', marker=dict(size=14, color='red', symbol='triangle-up'), name="Apoyo"))
        except: pass
    for idx, c in df_load.iterrows():
        try:
            n = df_n[df_n['ID'] == c['Nudo']].iloc[0]
            if c['FX (kN)'] != 0: fig.add_annotation(x=n['X (m)'], y=n['Y (m)'], ax=n['X (m)']-np.sign(c['FX (kN)'])*1, ay=n['Y (m)'], xref="x", yref="y", axref="x", ayref="y", showarrow=True, arrowhead=2, arrowcolor="green", text=f"{c['FX (kN)']} kN")
            if c['FY (kN)'] != 0: fig.add_annotation(x=n['X (m)'], y=n['Y (m)'], ax=n['X (m)'], ay=n['Y (m)']-np.sign(c['FY (kN)'])*1, xref="x", yref="y", axref="x", ayref="y", showarrow=True, arrowhead=2, arrowcolor="green", text=f"{c['FY (kN)']} kN")
        except: pass
    fig.update_layout(title="Lienzo Estructural Vectorizado", xaxis=dict(title="X (m)", showgrid=True, scaleanchor="y", scaleratio=1), yaxis=dict(title="Y (m)"), plot_bgcolor='#1e1e1e', paper_bgcolor='#1e1e1e', font=dict(color='white'), showlegend=False, dragmode='pan', margin=dict(l=20, r=20, t=40, b=20))
    return fig

# ─────────────────────────────────────────────
# TAB 1 & 2: INPUTS
with tab_gr:
    c_geo1, c_geo2 = st.columns([1, 2])
    with c_geo1:
        st.subheader("Nudos (X, Y)")
        st.session_state.nudos_df = st.data_editor(st.session_state.nudos_df, num_rows="dynamic", use_container_width=True)
        st.subheader("Elementos (Barras)")
        st.session_state.barras_df = st.data_editor(st.session_state.barras_df, num_rows="dynamic", use_container_width=True)
    with c_geo2:
        st.plotly_chart(plot_frame(st.session_state.nudos_df, st.session_state.barras_df, st.session_state.apoyos_df, st.session_state.cargas_nudo_df), use_container_width=True)

with tab_mat:
    c_mat1, c_mat2 = st.columns([1.5, 1])
    with c_mat1:
        st.subheader("Restricciones en Nudos (Apoyos)")
        st.session_state.apoyos_df = st.data_editor(st.session_state.apoyos_df, num_rows="dynamic", use_container_width=True)
        st.subheader("Fuerzas Nodales")
        st.session_state.cargas_nudo_df = st.data_editor(st.session_state.cargas_nudo_df, num_rows="dynamic", use_container_width=True)
    with c_mat2:
        st.subheader("Mecánica de Suelos y Materiales")
        st.info("Requerido para el Diseño Automático de Zapatas y Refuerzo.")
        qa_suelo = st.number_input("Capacidad Portante Suelo (qa) [kN/m²]", 10.0, 1000.0, 150.0, 10.0) # approx 15 ton/m2
        fc_concreto = st.number_input("Resistencia Concreto f'c [MPa]", 15.0, 60.0, 21.0, 1.0)
        fy_acero = st.number_input("Fluencia Acero fy [MPa]", 200.0, 600.0, 420.0, 10.0)
        recubrimiento = st.number_input("Recubrimiento Vigas/Col [cm]", 2.0, 10.0, 4.0, 0.5)

# ─────────────────────────────────────────────
# CALCULO ENGINE (Ejecuta automaticamente o con boton)
# ─────────────────────────────────────────────
nodes = st.session_state.nudos_df.copy().dropna()
elements = st.session_state.barras_df.copy().dropna()
supports = st.session_state.apoyos_df.copy().dropna()
loads = st.session_state.cargas_nudo_df.copy().dropna()

num_nudos = len(nodes)
num_gdl = num_nudos * 3
node_ids = nodes['ID'].tolist()
id_to_idx = {nid: i for i, nid in enumerate(node_ids)}

F = np.zeros(num_gdl)
for idx, c in loads.iterrows():
    if c['Nudo'] in id_to_idx:
        n_idx = id_to_idx[c['Nudo']]
        F[n_idx*3 + 0] += c.get('FX (kN)', 0.0)
        F[n_idx*3 + 1] += c.get('FY (kN)', 0.0)
        F[n_idx*3 + 2] += c.get('Mz (kN-m)', 0.0)

K = np.zeros((num_gdl, num_gdl))
element_results = [] # Almacenara k_local, T, dofs, L, angulo, etc

for idx, e in elements.iterrows():
    if e['Nudo I'] not in id_to_idx or e['Nudo J'] not in id_to_idx: continue
    i_idx = id_to_idx[e['Nudo I']]; j_idx = id_to_idx[e['Nudo J']]
    xi = nodes.loc[nodes['ID']==e['Nudo I'], 'X (m)'].values[0]
    yi = nodes.loc[nodes['ID']==e['Nudo I'], 'Y (m)'].values[0]
    xj = nodes.loc[nodes['ID']==e['Nudo J'], 'X (m)'].values[0]
    yj = nodes.loc[nodes['ID']==e['Nudo J'], 'Y (m)'].values[0]
    
    L = np.sqrt((xj-xi)**2 + (yj-yi)**2)
    if L == 0: continue
    cx = (xj-xi)/L; cy = (yj-yi)/L
    theta_rad = math.atan2(yj-yi, xj-xi)
    
    E = e['E (MPa)'] * 1e3; A_m2 = e['A (cm²)'] / 1e4; I_m4 = e['I (cm⁴)'] / 1e8
    k_local = np.zeros((6,6))
    v1 = E*A_m2/L; v2 = 12*E*I_m4/(L**3); v3 = 6*E*I_m4/(L**2); v4 = 4*E*I_m4/L; v5 = 2*E*I_m4/L
    k_local[0,0]=v1; k_local[0,3]=-v1; k_local[3,3]=v1; k_local[3,0]=-v1
    k_local[1,1]=v2; k_local[1,2]=v3; k_local[1,4]=-v2; k_local[1,5]=v3
    k_local[2,1]=v3; k_local[2,2]=v4; k_local[2,4]=-v3; k_local[2,5]=v5
    k_local[4,1]=-v2; k_local[4,2]=-v3; k_local[4,4]=v2; k_local[4,5]=-v3
    k_local[5,1]=v3; k_local[5,2]=v5; k_local[5,4]=-v3; k_local[5,5]=v4
    
    T = np.zeros((6,6))
    T[0,0]=cx; T[0,1]=cy; T[1,0]=-cy; T[1,1]=cx; T[2,2]=1
    T[3,3]=cx; T[3,4]=cy; T[4,3]=-cy; T[4,4]=cx; T[5,5]=1
    
    k_glob = T.T @ k_local @ T
    dofs = [i_idx*3, i_idx*3+1, i_idx*3+2, j_idx*3, j_idx*3+1, j_idx*3+2]
    for r in range(6):
        for c in range(6): K[dofs[r], dofs[c]] += k_glob[r,c]
            
    element_results.append({
        "ID": idx+1, "Nudo I": e['Nudo I'], "Nudo J": e['Nudo J'], "L": L, "theta_rad": theta_rad, 
        "k_local": k_local, "T": T, "dofs": dofs, "A_cm2": e['A (cm²)'], "I_cm4": e['I (cm⁴)'], "xi":xi, "yi":yi, "xj":xj, "yj":yj
    })

K_solve = np.copy(K)
penalty = 1e12
apoyos_idx = []
for idx, s in supports.iterrows():
    if s['Nudo'] in id_to_idx:
        n_idx = id_to_idx[s['Nudo']]
        if s.get('Fijo X', False): K_solve[n_idx*3+0, n_idx*3+0] += penalty
        if s.get('Fijo Y', False): K_solve[n_idx*3+1, n_idx*3+1] += penalty
        if s.get('Fijo Rz', False): K_solve[n_idx*3+2, n_idx*3+2] += penalty

try:
    U = np.linalg.solve(K_solve, F)
    R_rx = K @ U
    calc_success = True
except:
    calc_success = False

# ─────────────────────────────────────────────
# TAB 3: FUERZAS INTERNAS Y GRAFICOS BMD
with tab_res:
    if not calc_success:
        st.error("Error Matemático Estructural (Matriz Singular). Revise los apoyos.")
    else:
        st.subheader("Resultados Analíticos: Fuerzas en Elementos")
        
        # Calculate local forces f = k_local * T * U_element
        df_forzas = []
        fig_bmd = go.Figure()
        
        for e in element_results:
            u_global = U[e['dofs']]
            u_local = e['T'] @ u_global
            f_local = e['k_local'] @ u_local # [Ni, Vi, Mi, Nj, Vj, Mj]
            
            Ni = f_local[0]; Vi = f_local[1]; Mi = f_local[2]
            Nj = f_local[3]; Vj = f_local[4]; Mj = f_local[5]
            
            e['Ni'] = Ni; e['Vi'] = Vi; e['Mi'] = Mi
            e['Nj'] = Nj; e['Vj'] = Vj; e['Mj'] = Mj
            
            df_forzas.append({
                "Barra": e['ID'], "Axial I (kN)": Ni, "V Corte I (kN)": Vi, "Momento I (kN-m)": Mi,
                "Axial J (kN)": Nj, "V Corte J (kN)": Vj, "Momento J (kN-m)": Mj
            })
            
            # --- BMD PLOTTING ALGORITHM ---
            # Momento linearly drawn perpendicular to the bar
            scale_BMD = 0.05 # Scaling factor for visualization
            
            # Coordinates of bar ends
            L = e['L']; theta = e['theta_rad']
            
            # Sign convention standard frame: M > 0 is drawn on tension side.
            mi_plot = Mi * scale_BMD
            mj_plot = -Mj * scale_BMD # Reverse Mj for continuous plotting assuming equilibrium
            
            # Vector perp
            nx = -math.sin(theta); ny = math.cos(theta)
            
            # Plot Polygon for BMD
            p_xi = e['xi']; p_yi = e['yi']
            p_xj = e['xj']; p_yj = e['yj']
            
            # Offsets
            o_xi = p_xi + nx * mi_plot; o_yi = p_yi + ny * mi_plot
            o_xj = p_xj + nx * mj_plot; o_yj = p_yj + ny * mj_plot
            
            # Dibujar Area del Momento
            fig_bmd.add_trace(go.Scatter(
                x=[p_xi, o_xi, o_xj, p_xj, p_xi],
                y=[p_yi, o_yi, o_yj, p_yj, p_yi],
                fill="toself", fillcolor="rgba(0, 150, 255, 0.4)", line=dict(color="blue", width=1),
                name=f"BMD {e['ID']}", hoverinfo="text", 
                text=f"M_i: {Mi:.1f} kN-m<br>M_j: {-Mj:.1f} kN-m"
            ))
            # Dibujar Linea Barra Fuerte
            fig_bmd.add_trace(go.Scatter(x=[p_xi, p_xj], y=[p_yi, p_yj], mode='lines', line=dict(color='white', width=4), showlegend=False))
            # Textos
            fig_bmd.add_annotation(x=o_xi, y=o_yi, text=f"{abs(Mi):.1f}", showarrow=False, font=dict(color="yellow", size=10))
            fig_bmd.add_annotation(x=o_xj, y=o_yj, text=f"{abs(Mj):.1f}", showarrow=False, font=dict(color="yellow", size=10))

        st.dataframe(pd.DataFrame(df_forzas).style.format("{:.2f}").background_gradient(cmap='coolwarm'), use_container_width=True)
        
        st.subheader("Diagrama de Momentos Flectores (BMD)")
        fig_bmd.update_layout(title="BMD (Tension Side Plot)", plot_bgcolor='#1e1e1e', paper_bgcolor='#1e1e1e', font=dict(color='white'), yaxis_scaleanchor="x", showlegend=False, height=600)
        st.plotly_chart(fig_bmd, use_container_width=True)

# ─────────────────────────────────────────────
# TAB 4: DISEÑO AUTOMATIZADO (ZAPATAS, VIGAS, COLUMNAS)
with tab_dis:
    st.header(_t("⚡ Reporte de Diseño Estructural Automatizado", "⚡ Automated Structural Design Report"))
    if not calc_success:
        st.warning("Resuelve la estructura primero.")
    else:
        st.markdown(f"**Normativa Aplicada:** {norma_sel} | **$f'_c$:** {fc_concreto} MPa | **$f_y$:** {fy_acero} MPa | **$q_a$ Suelo:** {qa_suelo} kN/m²", unsafe_allow_html=True)
        doc_mem = Document()
        doc_mem.add_heading(f"Memoria de Diseño Estructural Auto - {norma_sel}", 0)
        doc_mem.add_paragraph(f"Propiedades del Material: fc={fc_concreto} MPa, fy={fy_acero} MPa, qa={qa_suelo} kN/m²")
        
        autores_c1, autores_c2, autores_c3 = st.columns(3)
        doc_mem.add_heading("1. Diseño de Zapatas (Cimentación)", level=1)
        
        with autores_c1:
            st.subheader("🪨 Dimensionamiento de Cimentación")
            zapatas_dxf_data = [] # Array dicts [x, y, B, L]
            for idx, s in supports.iterrows():
                nid = s['Nudo']
                if nid in id_to_idx:
                    n_idx = id_to_idx[nid]
                    Ry = abs(R_rx[n_idx*3+1]) # Reaccion vertical abs
                    # Req area = Pu / qa
                    Arear = Ry / qa_suelo
                    # Dimension cuadrada B = sqrt(A)
                    B_calc = math.sqrt(Arear) if Arear > 0 else 0.5
                    # Redondear a multiplos de 0.05m superio
                    B_final = math.ceil(max(B_calc, 0.8) / 0.05) * 0.05
                    # Espesor zapatas (h) min 0.30m, regla dedo D/4
                    H_zap = max(0.30, math.ceil((B_final/4)/0.05)*0.05)
                    
                    st.success(f"**Zapata en Nudo {nid}**: R_y = {Ry:.1f} kN")
                    st.markdown(f"- Área req: **{Arear:.2f} m²**")
                    st.markdown(f"- Dims Auto: **{B_final:.2f}m x {B_final:.2f}m x {H_zap:.2f}m**")
                    st.markdown(f"- Acero Sugerido: Parrilla **#4 @ 15 cm** ambas dirs.")
                    
                    doc_mem.add_paragraph(f"Zapata Nudo {nid}: Ry={Ry:.1f} kN. Dimensions: {B_final} x {B_final} x {H_zap} m. Parrilla: #4 @ 15 cm.")
                    n_x = nodes.loc[nodes['ID']==nid, 'X (m)'].values[0]
                    n_y = nodes.loc[nodes['ID']==nid, 'Y (m)'].values[0]
                    zapatas_dxf_data.append({"x": n_x, "y": n_y, "B": B_final})
                    
        with autores_c2:
            st.subheader("📏 Diseño Vigas A Flexión")
            doc_mem.add_heading("2. Diseño de Vigas (Flexión)", level=1)
            # Find horizontal bars (theta near 0 or 180)
            for e in element_results:
                deg = abs(np.degrees(e['theta_rad']))
                is_beam = (deg < 5) or (abs(deg - 180) < 5)
                if is_beam:
                    M_max = max(abs(e['Mi']), abs(e['Mj']))
                    # Approx b and h from Area and Inertia assuming rectangular 
                    A_cm2 = e['A_cm2']
                    # b*h = A_cm2; b*h^3/12 = I_cm4. h^2 = 12*I/A 
                    h_cm = math.sqrt(12 * e['I_cm4'] / A_cm2)
                    b_cm = A_cm2 / h_cm
                    
                    d_eff = h_cm - recubrimiento - 1.0 # 1cm estribo
                    
                    # ACI Bending As calculation
                    phi = 0.90
                    Rn = (M_max * 1e6) / (phi * (b_cm*10) * (d_eff*10)**2) # MPa
                    rho = (0.85 * fc_concreto / fy_acero) * (1 - math.sqrt(1 - (2*Rn)/(0.85*fc_concreto)))
                    if np.isnan(rho) or rho < 0.0033: rho = 0.0033 # rho min
                    if rho > 0.025: st.error(f"Viga {e['ID']}: Falla frágil. Aumentar sección."); rho_ok = False
                    else: rho_ok = True
                    As_req = rho * b_cm * d_eff # cm2
                    
                    st.info(f"**Viga {e['ID']} (b={b_cm:.0f}, h={h_cm:.0f}cm):** $M_u$ = {M_max:.1f} kN-m")
                    st.markdown(f"- Acero Requerido: **$A_s$ = {As_req:.2f} cm²**")
                    cant_v5 = math.ceil(As_req / 1.99)
                    st.markdown(f"- Armado Tentativo: **{cant_v5} vars #5**")
                    doc_mem.add_paragraph(f"Viga {e['ID']}: Mu={M_max:.1f} kN-m. b={b_cm:.0f}cm, h={h_cm:.0f}cm. As_req={As_req:.2f} cm2. D={cant_v5} vars #5.")
                    
        with autores_c3:
            st.subheader("🏛️ Revisión de Columnas")
            doc_mem.add_heading("3. Carga en Columnas", level=1)
            for e in element_results:
                deg = abs(np.degrees(e['theta_rad']))
                is_col = (abs(deg - 90) < 5) or (abs(deg - 270) < 5)
                if is_col:
                    P_max = max(abs(e['Ni']), abs(e['Nj']))
                    M_max_c = max(abs(e['Mi']), abs(e['Mj']))
                    A_cm2 = e['A_cm2']
                    h_cm = math.sqrt(12 * e['I_cm4'] / A_cm2); b_cm = A_cm2 / h_cm
                    
                    # Axial max limit ACI phi_Pn max
                    phi_c = 0.65
                    Pn_max = phi_c * 0.80 * (0.85 * fc_concreto * (A_cm2*100) / 1000.0) # Approx
                    rho_c = 0.01 # 1% min normative steel inside column
                    As_min_c = rho_c * A_cm2
                    
                    st.warning(f"**Columna {e['ID']} ({b_cm:.0f}x{h_cm:.0f}cm):** $P_u$={P_max:.1f} kN, $M_u$={M_max_c:.1f} kN-m")
                    if P_max > Pn_max: st.error("❌ Excede carga máxima.")
                    st.markdown(f"- Acero Mín. Longitudinal ($\rho=1\%$): **{As_min_c:.2f} cm²**")
                    st.markdown(f"- Estribos conf.: **#3 @ 10 cm** nudos, **@ 20 cm** luz.")
                    doc_mem.add_paragraph(f"Columna {e['ID']}: Pu={P_max:.1f} kN, Mu={M_max_c:.1f} kN-m. As min {rho_c*100}% = {As_min_c:.2f} cm2.")

        st.markdown("---")
        # EXPORTS EXCEL, DOCX, DXF
        st.subheader("📥 Exportaciones del Modelo Resolvido y Diseñado")
        col_ex1, col_ex2, col_ex3 = st.columns(3)
        
        # WORD MEMORIA
        f_doc = io.BytesIO(); doc_mem.save(f_doc); f_doc.seek(0)
        col_ex1.download_button("📥 Descargar Memoria.docx", data=f_doc, file_name="StructMaster2D_Memoria.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        # EXCEL RESULTADOS
        output_excel = io.BytesIO()
        df_reacciones = []; df_desp = []
        for nid in node_ids: n_idx=id_to_idx[nid]; df_desp.append({"Nudo": nid, "Dx (mm)": U[n_idx*3+0]*1000, "Dy (mm)": U[n_idx*3+1]*1000, "Rz (rad)": U[n_idx*3+2]})
        for idx, s in supports.iterrows(): n_idx=id_to_idx[s['Nudo']]; df_reacciones.append({"Nudo":s['Nudo'], "Rx":R_rx[n_idx*3+0],"Ry":R_rx[n_idx*3+1],"Mz":R_rx[n_idx*3+2]})
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            pd.DataFrame(df_reacciones).to_excel(writer, index=False, sheet_name='Reacciones')
            pd.DataFrame(df_desp).to_excel(writer, index=False, sheet_name='Desplazamientos Nodos')
        output_excel.seek(0)
        col_ex2.download_button("📥 Descargar Matrices.xlsx", data=output_excel, file_name="StructMaster2D_Resultados.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        
        # DXF STRUCTURAL PLAN
        try:
            doc_dxf = ezdxf.new('R2010'); msp = doc_dxf.modelspace()
            # Draw frame
            for e in element_results: msp.add_lwpolyline([(e['xi'], e['yi']), (e['xj'], e['yj'])], dxfattribs={'layer': 'VIGAS_COLUMNAS', 'color': 3})
            # Draw Zapatas
            for zap in zapatas_dxf_data:
                zx = zap["x"]; zy = zap["y"]; hz = zap["B"]/2.0
                msp.add_lwpolyline([(zx-hz, zy-hz), (zx+hz, zy-hz), (zx+hz, zy+hz), (zx-hz, zy+hz), (zx-hz, zy-hz)], dxfattribs={'layer': 'ZAPATAS_AUTO', 'color': 4, 'closed': True})
            out_dxf = io.StringIO(); doc_dxf.write(out_dxf)
            col_ex3.download_button("📥 Descargar Planos Auto Zapatas.dxf", data=out_dxf.getvalue(), file_name="StructMaster_Planos.dxf", mime="application/dxf")
        except: pass
