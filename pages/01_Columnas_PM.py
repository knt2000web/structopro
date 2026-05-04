import streamlit as st
from normas_referencias import mostrar_referencias_norma
st.set_page_config(
    page_title="StructoPro — Columnas P-M",
    page_icon=None,
    layout="wide"
)

#  Utilidad: Color AutoCAD segun # de cuartos de pulgada 
def _color_acero_dxf(db_mm: float) -> int:
    """Diferenciación visual por diámetro — colores AutoCAD estándar.
    En papel/ploteo usar pluma Color 7. En pantalla diferencia diámetros.
    NSR-10 / práctica colombiana de talleres."""
    if db_mm <= 9.5:   return 3   # Verde  → No.3 (3/8")
    elif db_mm <= 12.7: return 2  # Amarillo → No.4 (1/2")
    elif db_mm <= 15.9: return 1  # Rojo   → No.5 (5/8")
    elif db_mm <= 19.1: return 5  # Azul   → No.6 (3/4")
    elif db_mm <= 22.2: return 6  # Magenta → No.7 (7/8")
    else:               return 4  # Cian   → No.8+ (1"+)
# 
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math

# === GRAPAS CROSS-TIES — funciones exactas n=12 (ÚNICA FUENTE DE VERDAD) =====
import math as _gm

def generar_puntos_grapa_x(X_v, Y_f, R, L_cola):
    """Grapa H. Barras en (±X_v, Y_f). Tramo recto por Y_f+R."""
    pts = []; n = 12
    pf = (-X_v + R*_gm.cos(_gm.radians(225)), Y_f + R*_gm.sin(_gm.radians(225)))
    pts.append((pf[0] + L_cola*(-_gm.sin(_gm.radians(225))),
                pf[1] + L_cola*( _gm.cos(_gm.radians(225)))))
    for i in range(n, -1, -1):
        a = _gm.radians(90 + 135*(i/n))
        pts.append((-X_v + R*_gm.cos(a), Y_f + R*_gm.sin(a)))
    pts.append((X_v, Y_f + R))
    for i in range(1, n+1):
        a = _gm.radians(90 - 135*(i/n))
        pts.append((X_v + R*_gm.cos(a), Y_f + R*_gm.sin(a)))
    pf = (X_v + R*_gm.cos(_gm.radians(-45)), Y_f + R*_gm.sin(_gm.radians(-45)))
    pts.append((pf[0] + L_cola*( _gm.sin(_gm.radians(-45))),
                pf[1] + L_cola*(-_gm.cos(_gm.radians(-45)))))
    return pts

def generar_puntos_grapa_y(X_f, Y_v, R, L_cola):
    """Grapa V. Barras en (X_f, ±Y_v). Tramo recto por X_f+R."""
    pts = []; n = 12
    pf = (X_f + R*_gm.cos(_gm.radians(-135)), -Y_v + R*_gm.sin(_gm.radians(-135)))
    pts.append((pf[0] + L_cola*( _gm.sin(_gm.radians(-135))),
                pf[1] + L_cola*(-_gm.cos(_gm.radians(-135)))))
    for i in range(n, -1, -1):
        a = _gm.radians(-135*(i/n))
        pts.append((X_f + R*_gm.cos(a), -Y_v + R*_gm.sin(a)))
    pts.append((X_f + R, Y_v))
    for i in range(1, n+1):
        a = _gm.radians(135*(i/n))
        pts.append((X_f + R*_gm.cos(a), Y_v + R*_gm.sin(a)))
    pf = (X_f + R*_gm.cos(_gm.radians(135)), Y_v + R*_gm.sin(_gm.radians(135)))
    pts.append((pf[0] + L_cola*(-_gm.sin(_gm.radians(135))),
                pf[1] + L_cola*( _gm.cos(_gm.radians(135)))))
    return pts
# =============================================================================
import io
# ── Imports protegidos (Prompt Maestro UX v6.0) ──────────────────────────────
try:
    import ezdxf
    DXFEXT = True
except ImportError:
    ezdxf = None
    DXFEXT = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCXOK = True
except ImportError:
    DOCXOK = False
# ─────────────────────────────────────────────────────────────────────────────
try:
    import requests
    REQOK = True
except ImportError:
    REQOK = False

import plotly.graph_objects as go
import json
import datetime
try:
    import ifc_export
except ImportError:
    ifc_export = None
import os
from pathlib import Path
try:
    import qrcode
    from PIL import Image
    HAS_QR = True
except ImportError:
    HAS_QR = False
import tempfile

# 
# IDIOMA GLOBAL
lang = st.session_state["idioma"] if "idioma" in st.session_state else "Español"
def _t(es, en):
    return en if lang == "English" else es
# 


# 
# PERSISTENCIA SUPABASE
# 
import requests


# =============================================================================
# FEATURE D — Tabla de combinaciones de carga (Diagrama P-M multi-punto)
# =============================================================================
def parse_combo_table(df_combos, factor_fuerza):
    puntos = []
    for i, row in df_combos.iterrows():
        try:
            pu  = float(row.get('Pu',  row.get('pu',  0))) / factor_fuerza
            mux = float(row.get('Mux', row.get('mux', 0))) / factor_fuerza
            muy = float(row.get('Muy', row.get('muy', 0))) / factor_fuerza
            lbl = str(row.get('Combinacion', row.get('combo', f'C{i+1}')))
            puntos.append((pu, mux, muy, lbl))
        except Exception:
            pass
    return puntos


def plot_pm_combos(cap_x, Pu_single, Mux_single, puntos_combo,
                   factor_fuerza, unidad_fuerza, unidad_mom):
    COLORS = ['#f87171','#fb923c','#facc15','#4ade80','#34d399',
              '#60a5fa','#a78bfa','#f472b6','#e879f9','#94a3b8']
    fig, ax = plt.subplots(figsize=(7, 8))
    fig.patch.set_facecolor('#1e1e2e'); ax.set_facecolor('#14142a')
    ax.tick_params(colors='#cdd6f4', labelsize=8)
    for sp in ['top','right']: ax.spines[sp].set_visible(False)
    for sp in ['bottom','left']: ax.spines[sp].set_color('#555')
    ax.xaxis.label.set_color('#cdd6f4'); ax.yaxis.label.set_color('#cdd6f4')
    Mn_c = [v*factor_fuerza for v in cap_x['phi_Mn']]
    Pn_c = [v*factor_fuerza for v in cap_x['phi_Pn']]
    ax.plot(Mn_c, Pn_c, color='#60a5fa', lw=2.5, label='Diagrama \u03c6P-\u03c6M')
    ax.fill_betweenx(Pn_c, 0, Mn_c, alpha=0.08, color='#60a5fa')
    ax.plot(abs(Mux_single)*factor_fuerza, Pu_single*factor_fuerza,
            'D', ms=9, color='#facc15', zorder=8,
            markeredgecolor='white', markeredgewidth=0.8,
            label=f'Punto actual Pu={Pu_single*factor_fuerza:.1f} {unidad_fuerza}')
    for i, (pu, mux, muy, lbl) in enumerate(puntos_combo):
        c = COLORS[i % len(COLORS)]
        ax.plot(abs(mux)*factor_fuerza, pu*factor_fuerza, 'o', ms=7, color=c, zorder=7,
                markeredgecolor='white', markeredgewidth=0.5,
                label=f'{lbl}: ({abs(mux)*factor_fuerza:.1f}, {pu*factor_fuerza:.1f})')
    ax.axhline(0, color='#555', lw=0.8); ax.axvline(0, color='#555', lw=0.8)
    ax.set_xlabel(f'Momento M [{unidad_mom}]', fontsize=9)
    ax.set_ylabel(f'Carga Axial P [{unidad_fuerza}]', fontsize=9)
    ax.set_title('Diagrama P-M \u2014 Todas las Combinaciones de Carga',
                 color='white', fontsize=10, fontweight='bold')
    ax.legend(loc='upper right', fontsize=6.5, framealpha=0.3,
              facecolor='#1e1e2e', edgecolor='#555', labelcolor='#cdd6f4', ncol=1)
    ax.grid(True, linestyle=':', alpha=0.2, color='#888'); fig.tight_layout()
    return fig


# =============================================================================
# FEATURE E — Alzado 2D de Confinamiento
# =============================================================================
def plot_alzado_confinamiento(L_col_cm, Lo_cm, s_conf_cm, s_centro_cm,
                               b_cm, h_cm, es_circular,
                               rebar_diam_mm, stirrup_diam_mm, norma_sel, coderef):
    from matplotlib.lines import Line2D
    W = h_cm; L = L_col_cm; Lo = Lo_cm; s1 = s_conf_cm; s2 = s_centro_cm; cov = 3.8
    fig, ax = plt.subplots(figsize=(5, 10))
    fig.patch.set_facecolor('#1e1e2e'); ax.set_facecolor('#14142a')
    ax.tick_params(colors='#cdd6f4', labelsize=7)
    for sp in ax.spines.values(): sp.set_color('#555')
    ax.xaxis.label.set_color('#cdd6f4'); ax.yaxis.label.set_color('#cdd6f4')
    ax.add_patch(plt.Rectangle((0, 0), W, L, lw=2, edgecolor='#94a3b8', facecolor='#2d3748', zorder=1))
    for y0 in [0, L-Lo]:
        ax.add_patch(plt.Rectangle((cov, y0), W-2*cov, Lo, lw=0, facecolor='#f9731620', zorder=2))
        ax.add_patch(plt.Rectangle((cov, y0), W-2*cov, Lo, lw=1.5, edgecolor='#f97316',
                                    facecolor='none', linestyle='--', zorder=3))
    # estribos zona Lo inf
    y = s1
    while y <= Lo - s1*0.5:
        ax.plot([cov, W-cov], [y, y], color='#f97316', lw=1.2, zorder=4); y += s1
    # estribos zona central
    y = Lo + s2
    while y <= L - Lo - s2*0.5:
        ax.plot([cov, W-cov], [y, y], color='#60a5fa', lw=1.0, linestyle='--', zorder=4); y += s2
    # estribos zona Lo sup
    y = L - Lo + s1
    while y <= L - s1*0.5:
        ax.plot([cov, W-cov], [y, y], color='#f97316', lw=1.2, zorder=4); y += s1
    # barras longitudinales
    for xb in [cov + rebar_diam_mm/20, W - cov - rebar_diam_mm/20]:
        ax.plot([xb, xb], [cov, L-cov], color='#ff6b35', lw=2.5, zorder=5)
    # cotas
    def cota(y1, y2, texto, lado='r'):
        xc = W+6 if lado=='r' else -6; dx=3
        for y in [y1,y2]: ax.plot([xc-dx, xc+dx], [y,y], color='#a3e635', lw=0.8)
        ax.annotate('', xy=(xc,y2), xytext=(xc,y1),
                    arrowprops=dict(arrowstyle='<->', color='#a3e635', lw=0.9))
        ax.text(xc+dx+1, (y1+y2)/2, texto, color='#a3e635', fontsize=7, va='center')
    cota(0, Lo,   f'Lo={Lo:.0f} cm\ns\u2081={s1:.0f} cm')
    cota(Lo, L-Lo, f'Libre\ns\u2082={s2:.0f} cm')
    cota(L-Lo, L, f'Lo={Lo:.0f} cm\ns\u2081={s1:.0f} cm')
    cota(0, L,    f'L={L:.0f} cm', lado='l')
    ax.text(W/2, Lo/2, 'ZONA\nCONFINADA', color='#f97316', fontsize=7, ha='center', va='center', fontweight='bold', zorder=6)
    ax.text(W/2, L-Lo/2, 'ZONA\nCONFINADA', color='#f97316', fontsize=7, ha='center', va='center', fontweight='bold', zorder=6)
    ax.text(W/2, L/2, 'ZONA\nCENTRAL', color='#60a5fa', fontsize=7, ha='center', va='center', zorder=6)
    leg = [Line2D([0],[0], color='#f97316', lw=1.5, label=f'Estribos s\u2081={s1:.0f} cm (confinamiento)'),
           Line2D([0],[0], color='#60a5fa', lw=1.0, linestyle='--', label=f'Estribos s\u2082={s2:.0f} cm (zona libre)'),
           Line2D([0],[0], color='#ff6b35', lw=2.5, label='Barras longitudinales')]
    ax.legend(handles=leg, loc='upper center', bbox_to_anchor=(0.5,-0.04), fontsize=7,
              framealpha=0.3, facecolor='#1e1e2e', edgecolor='#555', labelcolor='#cdd6f4', ncol=1)
    sec = f"Circular D={h_cm:.0f} cm" if es_circular else f"Rectangular {b_cm:.0f}\xd7{h_cm:.0f} cm"
    ax.set_title(f'Alzado Confinamiento \u2014 {sec}\n{norma_sel}', color='white', fontsize=9, fontweight='bold')
    ax.set_xlim(-10, W+22); ax.set_ylim(-10, L+10); ax.set_aspect('equal'); ax.axis('off')
    fig.tight_layout()
    return fig

def get_supabase_rest_info():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return url, key
    except Exception as e:
        return None, None

def guardar_proyecto_supabase(nombre, estado_dict):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_columnas.json"
            db = {}
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            db[f"[COLUMNAS] {nombre}"] = {"nombre_proyecto": f"[COLUMNAS] {nombre}", "estado_json": json.dumps(estado_dict)}
            with open(db_path, "w", encoding="utf-8") as f: json.dump(db, f)
            return True, " Proyecto guardado (Local)"
        except Exception as e:
            return False, f" Error guardado local: {e}"
    
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "Prefer": "resolution=merge-duplicates"
    }
    payload = {
        "nombre_proyecto": nombre,
        "user_id": st.session_state.get("user_id", "anonimo"),
        "estado_json": json.dumps(estado_dict),
    }
    
    try:
        endpoint = f"{url}/rest/v1/proyectos?on_conflict=nombre_proyecto"
        res = requests.post(endpoint, headers=headers, json=payload)
        if res.status_code in [200, 201, 204]:
            return True, " Proyecto guardado en la nube"
        else:
            return False, f" Error API: {res.text}"
    except Exception as e:
        return False, f" Error al guardar: {e}"

def cargar_proyecto_supabase(nombre):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_columnas.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                match = db.get(f"[COLUMNAS] {nombre}") or db.get(nombre)
                if match:
                    estado = json.loads(match["estado_json"])
                    for k, v in estado.items(): st.session_state[k] = v
                    return True, f" Proyecto '{nombre}' cargado (Local)"
            return False, f" No se encontró el proyecto '{nombre}' localmente"
        except Exception as e:
            return False, f" Excepción al cargar local: {e}"
        
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Accept": "application/json"
    }
    
    try:
        endpoint = f"{url}/rest/v1/proyectos?nombre_proyecto=eq.{nombre}&select=*"
        res = requests.get(endpoint, headers=headers)
        
        if res.status_code == 200:
            data = res.json()
            if data and len(data) > 0:
                estado = json.loads(data[0]["estado_json"])
                for k, v in estado.items():
                    st.session_state[k] = v
                return True, f" Proyecto '{nombre}' cargado"
            else:
                return False, f" No se encontró el proyecto '{nombre}'"
        else:
            return False, f" Error al cargar (API): {res.text}"
    except Exception as e:
        return False, f" Excepción al cargar: {e}"

def listar_proyectos_supabase():
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_columnas.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                return sorted([k.replace("[COLUMNAS] ", "") for k in db.keys()])
            return []
        except Exception:
            return []
    
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Accept": "application/json"
    }
    
    try:
        endpoint = f"{url}/rest/v1/proyectos?select=nombre_proyecto"
        res = requests.get(endpoint, headers=headers)
        if res.status_code == 200:
            data = res.json()
            nombres = [item["nombre_proyecto"] for item in data if "nombre_proyecto" in item]
            return sorted(nombres)
        return []
    except Exception:
        return []

# 
# PANEL DE BIENVENIDA — info normativa visible siempre en cada módulo
# 
def _panel_normativo(titulo, descripcion, inputs_requeridos, fc, fy,
                     rhomin, rhomax, phif, phiv, normasel, nivelsis, coderef):
    st.markdown(
        f"""<div style="background:#1a2a1a;border-left:4px solid #4caf50;
        border-radius:8px;padding:14px 18px;margin-bottom:16px;">
        <span style="color:#81c784;font-weight:700;font-size:15px;"> {titulo}</span><br>
        <span style="color:#aaa;font-size:13px;">{descripcion}</span>
        </div>""",
        unsafe_allow_html=True,
    )
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("φ flexión/axial", f"{phif:.2f}/{phiv:.2f}", help=coderef)
    c2.metric("φ cortante", "0.75", help=coderef)
    c3.metric("ρ mín", f"{rhomin:.4f}", help="Cuantía mínima longitudinal")
    c4.metric("ρ máx", f"{rhomax:.4f}", help="Cuantía máxima")
    c5.metric("fc / fy", f"{fc:.0f} / {fy:.0f} MPa", help="Materiales activos")
    st.caption(f" Norma: **{norma_sel}** | Nivel sísmico: **{nivelsis}** | Ref: {coderef}")
    if inputs_requeridos:
        with st.expander(" Datos requeridos para este módulo", expanded=False):
            for item in inputs_requeridos:
                st.markdown(f"- {item}")
    st.markdown("---")


def capturar_estado_actual():
    claves = [
        "c_pm_norma", "c_pm_nivel_sismico", "c_pm_fc_unit", "c_pm_fc_mpa", "c_pm_psichoice", "c_pm_kgcm2choice",
        "c_pm_fy", "c_pm_seccion_type", "c_pm_b", "c_pm_h", "c_pm_D",
        "c_pm_d_prime", "c_pm_L", "c_pm_unit_system", "c_pm_rebar_type",
        "c_pm_num_h", "c_pm_num_v", "c_pm_n_barras_circ",
        "c_pm_col_type", "c_pm_stirrup_type", "c_pm_paso_circ", "c_pm_paso_rect",
        "c_pm_mux", "c_pm_muy", "c_pm_pu",
        "c_pm_output_units", "c_pm_k",
        # Factor k de esbeltez
        "c_pm_k", "c_pm_beta_dns", "c_pm_portico_tipo",
        # APU concreto premezclado
        "col_apu_premix", "col_apu_premix_p",
        "col_apu_moneda", "col_apu_cemento", "col_apu_acero", "col_apu_arena", "col_apu_grava", "col_apu_mo", "col_apu_aui",
        # Rótulo ICONTEC para DXF
        "dxf_empresa", "dxf_proyecto", "dxf_plano", "dxf_elaboro", "dxf_reviso", "dxf_aprobo",
        "qr_url",
        "cpm_empresa","cpm_proyecto_nombre","cpm_ingeniero",
    ]
    return {k: st.session_state[k] for k in claves if k in st.session_state}

# 
# FIX BUG #4: Manejo seguro de BASE_DIR para Streamlit Cloud
# 
st.markdown("""<div style="width:100%;overflow:hidden;border-radius:14px;margin-bottom:20px;box-shadow:0 4px 32px #0008;">
<svg viewBox="0 0 1100 220" xmlns="http://www.w3.org/2000/svg" style="width:100%;display:block;background:linear-gradient(135deg,#07091a 0%,#0d1b3e 100%);">
  <!-- Grid lines -->
  <g opacity="0.07" stroke="#60a5fa" stroke-width="0.5">
    <line x1="0" y1="55" x2="1100" y2="55"/><line x1="0" y1="110" x2="1100" y2="110"/>
    <line x1="0" y1="165" x2="1100" y2="165"/>
    <line x1="220" y1="0" x2="220" y2="220"/><line x1="440" y1="0" x2="440" y2="220"/>
    <line x1="660" y1="0" x2="660" y2="220"/>
  </g>
  <!-- Top accent bar -->
  <rect x="0" y="0" width="1100" height="3" fill="#3b82f6" opacity="0.9"/>
  <!-- Bottom accent bar -->
  <rect x="0" y="217" width="1100" height="3" fill="#8b5cf6" opacity="0.7"/>

  <!--  SECTION RECTANGULAR  -->
  <g transform="translate(45,25)">
    <text x="35" y="-6" text-anchor="middle" font-family="monospace" font-size="9" fill="#6b7280" letter-spacing="1.5">SECCION RECT.</text>
    <rect x="0" y="0" width="70" height="165" rx="2" fill="#2d3748" stroke="#718096" stroke-width="1.5"/>
    <rect x="8" y="8" width="54" height="149" rx="1" fill="none" stroke="#00d4ff" stroke-width="2" stroke-dasharray="none" opacity="0.85"/>
    <!-- Bars -->
    <circle cx="12" cy="12" r="5.5" fill="#e8a838"/><circle cx="35" cy="12" r="5.5" fill="#e8a838"/>
    <circle cx="58" cy="12" r="5.5" fill="#e8a838"/><circle cx="12" cy="84" r="5.5" fill="#e8a838"/>
    <circle cx="58" cy="84" r="5.5" fill="#e8a838"/><circle cx="12" cy="153" r="5.5" fill="#e8a838"/>
    <circle cx="35" cy="153" r="5.5" fill="#e8a838"/><circle cx="58" cy="153" r="5.5" fill="#e8a838"/>
    <!-- Stirrups -->
    <rect x="8" y="38" width="54" height="0" fill="none" stroke="#00d4ff" stroke-width="0.8" opacity="0.5"/>
    <rect x="8" y="68" width="54" height="0" fill="none" stroke="#00d4ff" stroke-width="0.8" opacity="0.5"/>
    <rect x="8" y="98" width="54" height="0" fill="none" stroke="#00d4ff" stroke-width="0.8" opacity="0.5"/>
    <rect x="8" y="128" width="54" height="0" fill="none" stroke="#00d4ff" stroke-width="0.8" opacity="0.5"/>
    <!-- Dim b -->
    <line x1="0" y1="178" x2="70" y2="178" stroke="#60a5fa" stroke-width="0.8"/>
    <text x="35" y="190" text-anchor="middle" font-family="monospace" font-size="10" fill="#93c5fd">b = 40 cm</text>
    <!-- Dim h -->
    <line x1="78" y1="0" x2="78" y2="165" stroke="#60a5fa" stroke-width="0.8"/>
    <text x="88" y="86" font-family="monospace" font-size="10" fill="#93c5fd">h=60</text>
  </g>

  <!--  SECTION CIRCULAR  -->
  <g transform="translate(200,35)">
    <text x="52" y="-6" text-anchor="middle" font-family="monospace" font-size="9" fill="#6b7280" letter-spacing="1.5">SECCION CIRC.</text>
    <circle cx="52" cy="60" r="52" fill="#2d3748" stroke="#718096" stroke-width="1.5"/>
    <circle cx="52" cy="60" r="45" fill="none" stroke="#00d4ff" stroke-width="2" opacity="0.85"/>
    <circle cx="52" cy="15" r="5.5" fill="#e8a838"/><circle cx="89" cy="33" r="5.5" fill="#e8a838"/>
    <circle cx="97" cy="72" r="5.5" fill="#e8a838"/><circle cx="77" cy="104" r="5.5" fill="#e8a838"/>
    <circle cx="52" cy="107" r="5.5" fill="#e8a838"/><circle cx="27" cy="104" r="5.5" fill="#e8a838"/>
    <circle cx="7" cy="72" r="5.5" fill="#e8a838"/><circle cx="15" cy="33" r="5.5" fill="#e8a838"/>
    <text x="52" y="135" text-anchor="middle" font-family="monospace" font-size="10" fill="#93c5fd">D = 60 cm</text>
  </g>

  <!--  P-M DIAGRAM  -->
  <g transform="translate(360,18)">
    <text x="95" y="-4" text-anchor="middle" font-family="monospace" font-size="9" fill="#6b7280" letter-spacing="1.5">DIAGRAMA P-M (NSR-10 C.10)</text>
    <line x1="18" y1="185" x2="18" y2="8" stroke="#60a5fa" stroke-width="1.2"/>
    <line x1="18" y1="185" x2="195" y2="185" stroke="#60a5fa" stroke-width="1.2"/>
    <polygon points="18,4 14,13 22,13" fill="#60a5fa"/>
    <polygon points="199,185 190,181 190,189" fill="#60a5fa"/>
    <text x="9" y="10" font-family="monospace" font-size="11" fill="#60a5fa">P</text>
    <text x="200" y="189" font-family="monospace" font-size="11" fill="#60a5fa">M</text>
    <path d="M18,22 C18,22 38,26 75,48 C115,72 158,130 168,158 C176,178 162,185 152,185 L18,185 Z"
          fill="none" stroke="#3b82f6" stroke-width="2.2"/>
    <path d="M18,38 C18,38 34,43 67,63 C103,85 143,140 152,164 C159,182 148,185 139,185 L18,185 Z"
          fill="#3b82f615" stroke="#8b5cf6" stroke-width="1.4" stroke-dasharray="5,3"/>
    <circle cx="128" cy="144" r="5" fill="#f59e0b"/>
    <line x1="128" y1="144" x2="128" y2="185" stroke="#f59e0b" stroke-width="0.8" stroke-dasharray="3,3"/>
    <line x1="18" y1="144" x2="128" y2="144" stroke="#f59e0b" stroke-width="0.8" stroke-dasharray="3,3"/>
    <text x="130" y="140" font-family="monospace" font-size="8" fill="#fbbf24">Bal.</text>
    <line x1="22" y1="200" x2="50" y2="200" stroke="#3b82f6" stroke-width="2"/>
    <text x="53" y="204" font-family="monospace" font-size="9" fill="#93c5fd">Nominal</text>
    <line x1="100" y1="200" x2="128" y2="200" stroke="#8b5cf6" stroke-width="1.4" stroke-dasharray="5,3"/>
    <text x="131" y="204" font-family="monospace" font-size="9" fill="#a78bfa">Reducido</text>
  </g>

  <!--  TEXT BLOCK  -->
  <g transform="translate(622,0)">
    <rect x="0" y="28" width="4" height="165" rx="2" fill="#3b82f6"/>
    <text x="18" y="68" font-family="Arial,sans-serif" font-size="34" font-weight="bold" fill="#ffffff">COLUMNAS</text>
    <text x="18" y="96" font-family="Arial,sans-serif" font-size="18" font-weight="300" fill="#93c5fd" letter-spacing="2">DIAGRAMAS P-M BIAXIAL</text>
    <rect x="18" y="104" width="440" height="1" fill="#3b82f6" opacity="0.5"/>
    <!-- Tags -->
    <rect x="18" y="115" width="108" height="22" rx="11" fill="#1e3a5f" stroke="#3b82f6" stroke-width="1"/>
    <text x="72" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#60a5fa">NSR-10 / ACI 318</text>
    <rect x="132" y="115" width="112" height="22" rx="11" fill="#2d1f4e" stroke="#8b5cf6" stroke-width="1"/>
    <text x="188" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#a78bfa">CONFINAMIENTO DMO</text>
    <rect x="250" y="115" width="98" height="22" rx="11" fill="#1a3a2a" stroke="#10b981" stroke-width="1"/>
    <text x="299" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#34d399">IFC / DXF ICONTEC</text>
    <rect x="354" y="115" width="88" height="22" rx="11" fill="#3a1a1a" stroke="#ef4444" stroke-width="1"/>
    <text x="398" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#f87171">ESPIRAL / ESTRIBO</text>
    <!-- Description -->
    <text x="18" y="158" font-family="Arial,sans-serif" font-size="11" fill="#64748b">Verificacion de capacidad a flexocompresion biaxial segun</text>
    <text x="18" y="174" font-family="Arial,sans-serif" font-size="11" fill="#64748b">NSR-10 C.10 / ACI 318-19 Cap. 10 — Planos ICONTEC en</text>
    <text x="18" y="190" font-family="Arial,sans-serif" font-size="11" fill="#64748b">formatos Carta, Oficio, Medio Pliego y Pliego. BIM IFC4.</text>
  </g>
</svg></div>""", unsafe_allow_html=True)


st.title(_t("Diagrama de Interacción P–M (Biaxial) y Diseño de Estribos", " P-M (Biaxial) Interaction Diagram & Tie Design"))
st.markdown(_t(
    "Generador interactivo de capacidad a flexocompresión **biaxial** para **Columnas Cuadradas, Rectangulares y Circulares**.",
    "Interactive **biaxial** flexure-compression capacity generator for **Square, Rectangular and Circular Columns**."
))

with st.expander(" ¿Cómo usar este módulo? — Guía Profesional", expanded=False):
    st.markdown("""
    ### Metodologia y Flujo de Diseno — Flexocompresion Biaxial
    Este modulo evalua la capacidad resistente de columnas **Cuadradas, Rectangulares y Circulares**, modelando la topologia tridimensional del acero bajo las prescripciones de NSR-10, ACI 318-19, NEC, E.060, NTC-EM, COVENIN y CIRSOC.

    ---

    #### 1. Configuracion Inicial (Barra Lateral)
    - **Norma y nivel sismico:** Seleccione la normativa y el nivel de disipacion (DES / DMO / DMI). Los factores phi, rho_max, phi_espiral y los requisitos de confinamiento C.21 se ajustan automaticamente.
    - **Identidad del proyecto:** Empresa, ingeniero responsable, elaboro, reviso, aprobo. Aparece en el rotulo ICONTEC del DXF y en la portada del DOCX.
    - **Modo Revision:** Active "Solo revision — sin paneles de entrada" en la barra lateral para obtener una vista condensada de resultados sin expanders de configuracion. Util cuando un tercero revisa el diseno sin modificar parametros.
    - **Guardar / Cargar proyecto:** Persiste todos los parametros en la nube (Supabase) o localmente en JSON. Permite retomar el diseno sin reingreso de datos.

    ---

    #### 2. Geometria y Materiales
    - Defina el tipo de seccion: **Rectangular/Cuadrada** (b x h) o **Circular** (D).
    - Ingrese recubrimiento d', longitud libre Lc, sistema de unidades (SI / Imperial) y unidades de salida (kN-m o tonf-m).
    - Distribuya el acero longitudinal en caras X-Y (rectangular) o contorno circular.
    - Para el refuerzo transversal: elija entre **Estribos cerrados** o **Espiral continua**; defina barra y paso.
    - Los **flejes (cross-ties)** se generan automaticamente cuando la separacion libre entre barras supera 15 cm (NSR-10 C.21.6.4.2).
    - El modulo verifica el **espaciamiento libre minimo** entre barras longitudinales (NSR-10 C.10.8.1) y emite advertencia si las barras no caben fisicamente en la seccion.
    - Para **concreto de alta resistencia** (fc >= 55 MPa) se despliega un aviso indicando que se recomienda validar con un modelo de confinamiento no lineal (Mander) dado que la hipotesis de deformacion ultima ecu = 0.003 puede ser no conservadora.

    ---

    #### 3. Pre-dimensionamiento Automatico — Auto-Sizing
    Cuando las dimensiones de la seccion aun no estan definidas, utilice el panel **Auto-Sizing** en la pestana Configuracion:
    - Defina la cuantia objetivo rho (1 % a 4 %).
    - Presione **Calcular Secciones**: el modulo barre una grilla de secciones tipicas colombianas (25 a 80 cm x 25 a 100 cm en pasos de 5 cm) y reporta las **5 secciones mas compactas** que satisfacen Pu, Mux y Muy con verificacion simplificada de carga axial y momento uniaxial.
    - La seccion sugerida es un punto de partida; siempre confirme la eleccion con el **Diagrama P-M completo** en la pestana correspondiente.

    ---

    #### 4. Demanda de Cargas y Esbeltez
    - Ingrese los esfuerzos ultimos factorados: **Pu** (axial), **Mux** (momento eje X), **Muy** (momento eje Y).
    - El modulo aplica automaticamente la **excentricidad minima** NSR-10 C.10.3.6: emin = max(0.10 h, 1.5 cm).
    - **Esbeltez:** Calcule kL/r con el factor k y la luz libre Lc. El modulo clasifica la columna (corta / esbelta) y aplica el factor delta_ns de magnificacion de momentos para marcos no desplazables (NSR-10 C.10.10). Para marcos desplazables, ingrese los momentos ya amplificados por analisis P-Delta.
    - **Beta_dns** ajustable para considerar efectos de fluencia diferida.

    ---

    #### 5. Combinaciones de Carga LRFD
    La pestana **Alzado y Combos** incluye:
    - **Tabla editable de combinaciones:** Ingrese hasta n combinaciones (Pu, Mux, Muy, etiqueta). El modulo grafica todos los puntos simultameante sobre el diagrama P-M uniaxial y la superficie biaxial 3D.
    - **Reporte de combo critico:** Identifica automaticamente la combinacion con mayor DCR = Mux / phi_Mn y la resalta.
    - Para un flujo sistematico, se recomienda ingresar al menos las combinaciones LRFD fundamentales de NSR-10 A.2: 1.4D, 1.2D+1.6L, 1.2D+1.0L+1.0E, 0.9D+1.0E y 1.2D+1.6L+0.5Lr. Los valores por defecto de la tabla corresponden a esas combinaciones escaladas a partir de Pu y Mu del caso base.

    ---

    #### 6. Diagrama P-M y Verificacion Biaxial (Bresler)
    - **Curvas P-M uniaxiales** para el eje X y el eje Y: envolvente nominal (Mn, Pn) y reducida (phi_Mn, phi_Pn), punto de balance y demanda graficada.
    - **Criterio de Bresler:** Calcula phi_Pni mediante la superficie de carga reciproca: `1/phi_Pni = 1/phi_Pnx + 1/phi_Pny - 1/phi_P0`. El **DCR = Pu / phi_Pni <= 1.0** es la verificacion de aprobacion. El calculo usa `np.nan` como centinela interno para la interpolacion, garantizando que valores fuera del diagrama se traten de forma robusta.
    - **Superficie 3D interactiva (Plotly):** Envolvente biaxial completa con todos los puntos de carga en el espacio P-Mx-My.
    - **Exportar CSV del diagrama P-M:** Disponible en la pestana Diagrama P-M. Descarga un archivo con las columnas (Pn, phi_Mn_x, phi_Mn_y) de ambos ejes para verificacion independiente en Excel o software externo.

    ---

    #### 7. Diseno Transversal Sismico (C.21)
    - Calcula la **zona de confinamiento lo** segun nivel sismico (DES / DMO / DMI) y la separacion maxima s en la zona confinada.
    - Verifica el **area minima de acero transversal Ash** con ambas formulas NSR-10 C.21.6.4 y reporta la que rige.
    - **Alzado de confinamiento 2D:** Zonas diferenciadas (confinada / libre), separaciones rotuladas, compatible para incluir en planos.
    - Tipos de gancho sismico: 135-135 (DES), 135-90 alternado (DMO), 90-90 (DMI).

    ---

    #### 8. Verificaciones Sismicas Adicionales (DES / DMO)
    Las siguientes verificaciones se activan automaticamente para niveles DES y DMO:

    - **Columna fuerte / Viga debil (NSR-10 C.21.6.1):** Ingrese la suma de momentos nominales de las vigas que llegan al nodo (Sigma_Mnv) en la barra lateral. El modulo calcula la suma de momentos nominales de la columna (Sigma_Mnc) y verifica Sigma_Mnc >= 1.2 * Sigma_Mnv.
    - **Nodo Viga-Columna (NSR-10 C.21.7.4.1):** Ingrese el area de acero de traccion de las vigas en el nodo (As_vigas) y el cortante de diseno de la columna (Vu_col). Seleccione el tipo de confinamiento del nodo (interior / borde / esquina) para obtener el factor gamma (1.70 / 1.25 / 1.00). El modulo calcula la demanda de cortante en el nodo Vu_j y la verifica contra phi_Vn = 0.85 * gamma * sqrt(fc) * Aj.
    - Ambas verificaciones aparecen en la tabla de cumplimiento de la **Memoria DOCX**, con semaforo de estado (CUMPLE / NO CUMPLE) y nota de advertencia si el criterio no se satisface.

    ---

    #### 9. Despiece de Acero y Figurado para Taller
    - **Tabla de despiece completa:** Marca, cantidad, diametro, longitud unitaria, longitud total y peso (kg) para: L1 barra longitudinal recta, L1A barra de arranque (longitud + 1.3 ld), E1 estribo perimetral, GX/GY flejes eje X e Y, Espiral zuncho continuo.
    - **Grafico de distribucion de pesos:** Barras por elemento con valor kg rotulado.
    - **Figurado para taller:** Dibujos acotados de varilla longitudinal con ganchos, estribo con ganchos 135, flejes internos y espiral.

    ---

    #### 10. Exportaciones BIM y Documentacion Tecnica
    - **Plano DXF (ICONTEC):** Seccion transversal acotada + alzado de confinamiento + tabla de hierros + rotulo ICONTEC. Formatos Carta, Oficio, Medio Pliego y Pliego. Compatible con AutoCAD y LibreCAD.
    - **Modelo IFC4 (BIM):** IfcColumn con geometria 3D de barras longitudinales y estribos posicionados (LOD 350), listo para Revit, ArchiCAD o BIMvision. Incluye Pset con cuantia y verificacion Bresler.
    - **Memoria DOCX:** Portada marca blanca con logo y marca de agua diagonal (empresa / StructoPro) + parametros + diagrama P-M eje X + seccion con estribos + tabla de despiece + alzado de confinamiento + tabla de verificaciones normativas + referencias + desglose APU. La marca de agua solo aparece en los documentos generados; no afecta la interfaz.
    - **CSV Diagrama P-M:** Descarga directa de los puntos numericos del diagrama (nominal y reducido, ejes X e Y) para revision independiente.

    ---

    #### 11. Presupuesto APU
    - Ingrese precios unitarios de cemento, acero, arena, grava, agua, encofrado y mano de obra.
    - Soporte para **concreto premezclado** (precio por m3). El jornal de MO se actualiza automaticamente desde el **SMLMV vigente** de Colombia.
    - Desglose APU en grafico de barras interactivo (Plotly) con transferencia al modulo de Presupuesto WBS.

    ---

    **Flujo recomendado:** Norma y nivel sismico -> Geometria y materiales -> Auto-Sizing si es predimensionamiento -> Cargas (Pu, Mux, Muy) -> Verificar DCR <= 1.0 en Bresler -> Combinaciones LRFD -> Estribos sismicos -> Nodo V-C y columna fuerte (DES/DMO) -> Exportar DXF + IFC4 + DOCX + CSV.
    """)

# 
# DICCIONARIOS DE BARRAS
# 
REBAR_US = {
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
    "#5 (5/8\")": {"area": 1.99, "diam_mm": 15.88},
    "#6 (3/4\")": {"area": 2.84, "diam_mm": 19.05},
    "#7 (7/8\")": {"area": 3.87, "diam_mm": 22.23},
    "#8 (1\")":   {"area": 5.10, "diam_mm": 25.40},
    "#9 (1 1/8\")": {"area": 6.45, "diam_mm": 28.65},
    "#10 (1 1/4\")": {"area": 7.92, "diam_mm": 32.26},
}

REBAR_MM = {
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
    "14 mm": {"area": 1.539, "diam_mm": 14.0},
    "16 mm": {"area": 2.011, "diam_mm": 16.0},
    "18 mm": {"area": 2.545, "diam_mm": 18.0},
    "20 mm": {"area": 3.142, "diam_mm": 20.0},
    "22 mm": {"area": 3.801, "diam_mm": 22.0},
    "25 mm": {"area": 4.909, "diam_mm": 25.0},
    "28 mm": {"area": 6.158, "diam_mm": 28.0},
    "32 mm": {"area": 8.042, "diam_mm": 32.0},
}

STIRRUP_US = {
    "#2 (1/4\")": {"area": 0.32, "diam_mm": 6.35},
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
}

STIRRUP_MM = {
    "6 mm":  {"area": 0.283, "diam_mm": 6.0},
    "8 mm":  {"area": 0.503, "diam_mm": 8.0},
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
}

# 
# FUNCIONES DE DIBUJO PARA FIGURADO
# 
_MM_TO_BAR = {
    6.0:  "#2 (1/4\")",  6.35: "#2 (1/4\")",
    8.0:  "#2.5 (5/16\")",
    9.53: "#3 (3/8\")",  10.0: "#3 (3/8\")",
    12.0: "#4 (1/2\")",  12.70: "#4 (1/2\")",
    14.0: "#4.5 (9/16\")",
    15.88: "#5 (5/8\")", 16.0: "#5 (5/8\")",
    18.0: "#5.7 (11/16\")",
    19.05: "#6 (3/4\")", 20.0: "#6 (3/4\")",
    22.0: "#7 (7/8\")",  22.23: "#7 (7/8\")",
    25.0: "#8 (1\")",    25.40: "#8 (1\")",
    28.0: "#9 (1 1/8\")",28.65: "#9 (1 1/8\")",
    32.0: "#10 (1 1/4\")",32.26: "#10 (1 1/4\")",
}

def _bar_label(diam_mm):
    """Retorna la designación colombiana NSR-10: #N (fracción") — Ø mm"""
    best = min(_MM_TO_BAR.keys(), key=lambda k: abs(k - diam_mm))
    if abs(best - diam_mm) <= 2.0:
        return f"{_MM_TO_BAR[best]}  Ø{diam_mm:.1f} mm"
    return f"Ø{diam_mm:.1f} mm"

def _bar_label_short(diam_mm):
    """Retorna solo la designación colombiana: #N (fracción"), sin mm — para títulos y gráficos."""
    best = min(_MM_TO_BAR.keys(), key=lambda k: abs(k - diam_mm))
    if abs(best - diam_mm) <= 2.0:
        return _MM_TO_BAR[best]
    return f"Ø{diam_mm:.1f} mm"

def configurar_pdf_comercial(fig):
    fig.patch.set_facecolor('white')
    for _ax in fig.get_axes():
        _ax.set_facecolor('white')
        _ax.tick_params(colors='black')
        for spine in _ax.spines.values():
            spine.set_edgecolor('black')
            spine.set_linewidth(1.5)
        _ax.xaxis.label.set_color('black')
        _ax.yaxis.label.set_color('black')
        if _ax.get_title(): _ax.title.set_color('black')

def draw_longitudinal_bar(total_len_cm, straight_len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(6, total_len_cm/20), 2))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    _bar_color = '#ff6b35'
    ax.plot([0, straight_len_cm], [0, 0], color=_bar_color, linewidth=3)
    ax.plot([0, 0], [0, hook_len_cm], color=_bar_color, linewidth=3)
    ax.plot([straight_len_cm, straight_len_cm], [0, -hook_len_cm], color=_bar_color, linewidth=3)
    ax.plot([0, straight_len_cm], [0, 0], 'o', color=_bar_color, markersize=5)
    ax.annotate(f"{straight_len_cm:.0f} cm", xy=(straight_len_cm/2, 0.4),
                ha='center', fontsize=9, fontweight='bold', color='#cdd6f4')
    ax.annotate(f"Gancho 12db\n= {hook_len_cm:.0f} cm", xy=(-0.3, hook_len_cm/2),
                ha='right', va='center', fontsize=8, color='#ffd700')
    ax.annotate(f"Gancho 12db\n= {hook_len_cm:.0f} cm", xy=(straight_len_cm+0.3, -hook_len_cm/2),
                ha='left', va='center', fontsize=8, color='#ffd700')
    ax.set_xlim(-hook_len_cm*0.2, straight_len_cm + hook_len_cm*0.2)
    ax.set_ylim(-hook_len_cm*1.2, hook_len_cm*1.2)
    ax.axis('off')
    ax.set_title(f"Varilla L1 — {label} — Longitud total {total_len_cm:.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_stirrup(b_cm, h_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    import math as _math
    import numpy as _np
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(5, b_cm/12), max(4, h_cm/12)))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes():
        _ax.set_facecolor('#14142a')
        _ax.tick_params(colors='#cdd6f4')
        _ax.xaxis.label.set_color('#cdd6f4')
        _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')

    x0, y0 = 0.0, 0.0
    # Radio doblez 3db (NSR-10 C.7.2), mínimo 0.3 cm
    r = max(min(bar_diam_mm / 10.0 * 3.0, b_cm * 0.08, h_cm * 0.08), 0.3)

    def _arc(cx, cy, rad, a0, a1, n=10):
        angs = _np.linspace(_math.radians(a0), _math.radians(a1), n)
        return [(cx + rad*_math.cos(a), cy + rad*_math.sin(a)) for a in angs]

    # ── Perímetro con 4 esquinas redondeadas ─────────────────────────────
    # Arranca desde (x0, y0+h_cm-r) = pie del arco sup-izq, lado izquierdo
    # Sentido: baja por izq → inf-izq → inf → inf-der → sube por der → sup-der → sup → arco sup-izq
    pts = []
    pts.append((x0, y0 + h_cm - r))           # inicio lado izq (donde nace gancho 1)
    pts += _arc(x0+r, y0+r,       r, 180, 270) # esquina inf-izq
    pts.append((x0+b_cm-r, y0))               # lado inferior
    pts += _arc(x0+b_cm-r, y0+r,  r, 270, 360) # esquina inf-der
    pts.append((x0+b_cm, y0+h_cm-r))          # lado der
    pts += _arc(x0+b_cm-r, y0+h_cm-r, r, 0, 90)  # esquina sup-der
    pts.append((x0+r, y0+h_cm))               # lado sup (desde donde nace gancho 2)
    pts += _arc(x0+r, y0+h_cm-r,  r, 90, 180) # arco sup-izq
    pts.append((x0, y0+h_cm-r))               # cierra en el mismo punto de inicio

    ax.plot([p[0] for p in pts], [p[1] for p in pts], color='white', linewidth=2.5)

    # ── Dos ganchos 135° — ambos salen de esquina SUP-IZQ hacia NÚCLEO (+X, -Y)
    vis = min(hook_len_cm, b_cm*0.32, h_cm*0.32)
    dx =  vis * _math.cos(_math.radians(45))   # +X  (hacia interior)
    dy = -vis * _math.sin(_math.radians(45))   # -Y  (hacia interior)

    # Gancho 1: desde pie del arco sup-izq en el lado IZQUIERDO
    ax.plot([x0,     x0+dx],            [y0+h_cm-r, y0+h_cm-r+dy], color='white', linewidth=2.5)
    ax.plot(x0+dx, y0+h_cm-r+dy, 'o',  color='#ff4444', markersize=5, zorder=5)

    # Gancho 2: desde inicio del lado SUPERIOR (derecha → izquierda)
    ax.plot([x0+r,   x0+r+dx],          [y0+h_cm,   y0+h_cm+dy],   color='white', linewidth=2.5)
    ax.plot(x0+r+dx, y0+h_cm+dy, 'o',  color='#ff4444', markersize=5, zorder=5)

    ax.annotate(f"{b_cm:.0f} cm", xy=(b_cm/2, y0-0.8),
                ha='center', fontsize=9, fontweight='bold', color='#cdd6f4')
    ax.annotate(f"{h_cm:.0f} cm", xy=(x0-0.8, h_cm/2),
                ha='right', va='center', fontsize=9, fontweight='bold', color='#cdd6f4')
    ax.annotate(f"Gancho 135\u00b0\next.=6d\u1d49={6*bar_diam_mm/10:.1f}cm | r=2d\u1d49={2*bar_diam_mm/10:.1f}cm",
                xy=(x0+dx+0.2, y0+h_cm-r+dy-0.1), fontsize=7.5, color='#ff9999',
                va='top', ha='left')
    ax.annotate("alterna estribo\na estribo",
                xy=(x0+r+dx*0.5, y0+h_cm+dy*0.5), fontsize=6.5, color='#888',
                va='center', ha='left', style='italic')
    ax.set_xlim(x0 - hook_len_cm*0.3, b_cm + hook_len_cm*0.6)
    ax.set_ylim(y0 - hook_len_cm*0.5, h_cm + hook_len_cm*1.1)
    ax.axis('off')
    ax.set_title(f"Estribo E1 \u2014 {label} \u2014 Per\u00edmetro {2*(b_cm+h_cm):.0f} cm",
                 fontsize=9, fontweight='bold', color='white')
    return fig


def draw_stirrup_with_ties(b_cm, h_cm, recub_cm, hook_len_cm, bar_diam_mm,
                           n_ties_x=0, n_ties_y=0,
                           nivel_sismico_str="DES", bar_name=None,
                           bar_coords=None, long_bar_diam_mm=None):
    """Estribo perimetral + flejes NSR-10 C.21.6.4.2.
    bar_coords: lista [(x,y),...] en cm de todas las varillas.
    Si supera 15 cm entre varillas no arriostradas genera fleje automático.
    """
    import math as _m, numpy as _np
    Ancho_Estribo = b_cm - (2 * recub_cm)
    Alto_Estribo  = h_cm - (2 * recub_cm)
    
    label   = bar_name if bar_name else _bar_label(bar_diam_mm)
    _is_des = "DES" in nivel_sismico_str.upper() or "ESPECIAL" in nivel_sismico_str.upper()
    _is_dmo = "DMO" in nivel_sismico_str.upper() or "MODERADA" in nivel_sismico_str.upper()
    _ang1   = 135 if (_is_des or _is_dmo) else 90
    MAX_LIB = 15.0   # cm — NSR-10 C.21.6.4.2 / ACI 318-19 §25.7.2.3

    fig, ax = plt.subplots(figsize=(max(5, Ancho_Estribo/10), max(4, Alto_Estribo/10)))
    fig.patch.set_facecolor('#1e1e2e')
    for _a in fig.get_axes():
        _a.set_facecolor('#14142a'); _a.tick_params(colors='#cdd6f4')
        _a.xaxis.label.set_color('#cdd6f4'); _a.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')

    r  = max(min(bar_diam_mm/10.0*3, Ancho_Estribo*0.08, Alto_Estribo*0.08), 0.3)
    hk = hook_len_cm

    db_cm = (long_bar_diam_mm / 10.0) if long_bar_diam_mm else (bar_diam_mm / 10.0 * 1.5)
    dst_cm = bar_diam_mm / 10.0
    R_grapa = (db_cm / 2.0) + (dst_cm / 2.0)

    def _arc_pts(cx, cy, rad, a0, a1, n=8):
        angs = _np.linspace(_m.radians(a0), _m.radians(a1), n)
        return [(cx+rad*_m.cos(a), cy+rad*_m.sin(a)) for a in angs]

    import math as _mg
    # Centro del estribo en coords locales (origen = esquina inf-izq)
    _Cx = Ancho_Estribo / 2.0
    _Cy = Alto_Estribo  / 2.0
    # xv_max / yv_max en coords centradas: semiancho hasta CENTRO de la varilla
    _xv_max_m = _Cx - dst_cm - (db_cm / 2.0)   # = Ancho/2 - dst - radio_varilla
    _yv_max_m = _Cy - dst_cm - (db_cm / 2.0)
    _n_g = 10

    def _get_grapa_pts(x1, y1, x2, y2, orientation='X'):
        _Lc = hk
        if orientation == 'X':
            _yv_c = y1 - _Cy
            _pc = generar_puntos_grapa_x(_xv_max_m, _yv_c, R_grapa, _Lc)
            return [(x + _Cx, y + _Cy) for x, y in _pc]
        else:
            _xv_c = x1 - _Cx
            _pc = generar_puntos_grapa_y(_xv_c, _yv_max_m, R_grapa, _Lc)
            return [(x + _Cx, y + _Cy) for x, y in _pc]

    # ── Estribo perimetral ───────────────────────────────────────────────────
    pts = []
    pts.append((0, Alto_Estribo-r)); pts += _arc_pts(r, r, r, 180, 270)
    pts.append((Ancho_Estribo-r, 0)); pts += _arc_pts(Ancho_Estribo-r, r, r, 270, 360)
    pts.append((Ancho_Estribo, Alto_Estribo-r)); pts += _arc_pts(Ancho_Estribo-r, Alto_Estribo-r, r, 0, 90)
    pts.append((r, Alto_Estribo)); pts += _arc_pts(r, Alto_Estribo-r, r, 90, 180)
    pts.append((0, Alto_Estribo-r))
    ax.plot([p[0] for p in pts], [p[1] for p in pts],
            color='#00d4ff', linewidth=2.5, zorder=3)
    # Ganchos 135° estribo perimetral
    _dk = hk * 0.707
    ax.annotate("", xy=(0+_dk, Alto_Estribo-r-_dk), xytext=(0, Alto_Estribo-r),
                arrowprops=dict(arrowstyle="-", color='#00d4ff', lw=2.5))
    ax.annotate("", xy=(r+_dk, Alto_Estribo-_dk),   xytext=(r, Alto_Estribo),
                arrowprops=dict(arrowstyle="-", color='#00d4ff', lw=2.5))

    # ── Flejes pantalla: líneas directas hacia interior (DXF/IFC usan generar_puntos_grapa) ──
    color_transversal = '#00d4ff'
    _hk_vis = min(hk * 0.65, Ancho_Estribo * 0.15, Alto_Estribo * 0.15)
    _gk = _hk_vis * 0.707

    # Flejes horizontales (dir-X)
    if n_ties_x > 0:
        esp_fx = Alto_Estribo / (n_ties_x + 1)
        for _i in range(1, n_ties_x + 1):
            _yf = esp_fx * _i
            ax.plot([0, Ancho_Estribo], [_yf, _yf], color=color_transversal, linewidth=2.0, zorder=4)
            ax.plot([0, _gk], [_yf, _yf + _gk], color=color_transversal, linewidth=2.0, zorder=4)
            if _i % 2 == 0:
                ax.plot([Ancho_Estribo, Ancho_Estribo - _gk], [_yf, _yf + _gk],
                        color=color_transversal, linewidth=2.0, zorder=4)
            else:
                ax.plot([Ancho_Estribo, Ancho_Estribo - _gk], [_yf, _yf - _gk],
                        color=color_transversal, linewidth=2.0, zorder=4)

    # Flejes verticales (dir-Y)
    if n_ties_y > 0:
        esp_fy = Ancho_Estribo / (n_ties_y + 1)
        for _j in range(1, n_ties_y + 1):
            _xf = esp_fy * _j
            ax.plot([_xf, _xf], [0, Alto_Estribo], color=color_transversal, linewidth=2.0, zorder=4)
            ax.plot([_xf, _xf + _gk], [0, _gk], color=color_transversal, linewidth=2.0, zorder=4)
            if _j % 2 == 0:
                ax.plot([_xf, _xf + _gk], [Alto_Estribo, Alto_Estribo - _gk],
                        color=color_transversal, linewidth=2.0, zorder=4)
            else:
                ax.plot([_xf, _xf - _gk], [Alto_Estribo, Alto_Estribo - _gk],
                        color=color_transversal, linewidth=2.0, zorder=4)

    # ── Cotas dinámicas ──
    # Definir offsets de recubrimiento para cotas y límites de ejes
    _shift_x = recub_cm
    _shift_y = recub_cm
    # Cotas de sección total (b y h) y núcleo (Ancho_Estribo y Alto_Estribo)
    ax.annotate(f"b = {b_cm:.1f} cm", xy=(Ancho_Estribo/2, -_shift_y - hk*0.55),
                ha='center', fontsize=9, fontweight='bold', color='#93c5fd')
    ax.annotate(f"nucleo = {Ancho_Estribo:.1f} cm", xy=(Ancho_Estribo/2, -hk*0.25),
                ha='center', fontsize=8, color='#cdd6f4')
    ax.annotate(f"h = {h_cm:.1f} cm", xy=(-_shift_x - hk*0.55, Alto_Estribo/2),
                ha='right', va='center', fontsize=9, fontweight='bold', color='#93c5fd',
                rotation=90, rotation_mode='anchor')
    ax.annotate(f"nucleo = {Alto_Estribo:.1f} cm", xy=(-hk*0.2, Alto_Estribo/2),
                ha='right', va='center', fontsize=8, color='#cdd6f4',
                rotation=90, rotation_mode='anchor')
    _gancho_str = f"Gancho {_ang1}°"
    _n_tx = n_ties_x + n_ties_y
    ax.set_title(
        f"Sección Transversal — Estribos + Flejes | {label}\n"
        f"{_gancho_str} ({nivel_sismico_str}) | {_n_tx} flete(s) NSR-10 C.21.6.4.2",
        fontsize=9, fontweight='bold', color='white')
    # El estribo se dibuja en (0..Ancho_Estribo, 0..Alto_Estribo)
    # El recubrimiento se representa desplazando el origen del plot para mostrar la seccion completa
    _shift_x = recub_cm   # offset para mostrar concreto exterior al estribo
    _shift_y = recub_cm
    ax.set_xlim(-_shift_x - hk*0.8,  Ancho_Estribo + _shift_x + hk*1.5)
    ax.set_ylim(-_shift_y - hk*1.2,  Alto_Estribo  + _shift_y + hk*1.5)
    ax.axis('off')
    # ── Recubrimiento: rectángulo exterior (sección bruta) ─────────────────
    _rect_outer = plt.Rectangle((-_shift_x, -_shift_y), b_cm, h_cm,
                                 fill=False, edgecolor='#555566',
                                 linewidth=1.0, linestyle='--', zorder=1)
    ax.add_patch(_rect_outer)
    # ── Barras longitudinales (círculos rellenos) ──────────────────────────
    if bar_coords:
        _lbd_cm = (long_bar_diam_mm / 10.0) if long_bar_diam_mm else (bar_diam_mm / 10.0 * 1.5)
        _r_bar  = max(_lbd_cm / 2.0, 0.4)
        for (_xb, _yb) in bar_coords:
            _circ = plt.Circle((_xb, _yb), _r_bar,
                                facecolor='#f5a623', edgecolor='#1e1e2e',
                                linewidth=0.8, zorder=6)
            ax.add_patch(_circ)

    handles, labels_l = ax.get_legend_handles_labels()
    if handles:
        ax.legend(handles, labels_l, fontsize=8, loc='upper right',
                  facecolor='#1e1e2e', labelcolor='white', edgecolor='#444')
    return fig


def draw_crosstie(len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(6, len_cm/15), 2))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    _ct_color = '#00d4ff'
    ax.plot([0, len_cm], [0, 0], color=_ct_color, linewidth=2.5)
    ax.plot([0, -hook_len_cm*0.7], [0, -hook_len_cm*0.7], color=_ct_color, linewidth=2.5)
    ax.plot([len_cm, len_cm + hook_len_cm*0.7], [0, -hook_len_cm*0.7], color=_ct_color, linewidth=2.5)
    ax.annotate(f"{len_cm:.0f} cm", xy=(len_cm/2, 0.3), ha='center', fontsize=8, color='#cdd6f4')
    ax.annotate(f"Gancho 135°", xy=(0, -hook_len_cm*0.5), ha='right', fontsize=8, color='#ffd700')
    ax.annotate(f"Gancho 135°", xy=(len_cm, -hook_len_cm*0.5), ha='left', fontsize=8, color='#ffd700')
    ax.set_xlim(-hook_len_cm*1.2, len_cm + hook_len_cm*1.2)
    ax.set_ylim(-hook_len_cm*1.5, hook_len_cm*0.5)
    ax.axis('off')
    ax.set_title(f"Crosstie C1 — {label} — Longitud {len_cm:.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_spiral(D_cm, paso_cm, bar_diam_mm, bar_name=None):
    """Dibujo esquemático de espiral para columnas circulares"""
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(6, 4))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    circle = plt.Circle((0, 0), D_cm/2, fill=False, edgecolor='#718096', linewidth=2)
    ax.add_patch(circle)
    theta = np.linspace(0, 4*np.pi, 200)
    r = D_cm/2 - bar_diam_mm/10
    x = r * np.cos(theta)
    y = r * np.sin(theta)
    ax.plot(x, y, color='#a78bfa', linewidth=1.8)
    ax.annotate(f"Espiral {label}", xy=(0, D_cm/2 + 2), ha='center', fontsize=9, color='#cdd6f4')
    ax.annotate(f"Paso = {paso_cm:.1f} cm", xy=(0, -D_cm/2 - 3), ha='center', fontsize=8)
    ax.set_xlim(-D_cm/2 - 5, D_cm/2 + 5)
    ax.set_ylim(-D_cm/2 - 8, D_cm/2 + 5)
    ax.axis('off')
    ax.set_title(f"Espiral — {label}", fontsize=9, fontweight='bold')
    return fig

# 
# PARÁMETROS POR NORMA (con límites por nivel sísmico)
# 
CODES = {
    "NSR-10 (Colombia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 4.0, "rho_max_des": 4.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["DMI — Disipación Mínima", "DMO — Disipación Moderada", "DES — Disipación Especial"],
        "ref": "NSR-10 Título C",
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary", "IMF — Intermediate", "SMF — Special"],
        "ref": "ACI 318-25",
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary", "IMF — Intermediate", "SMF — Special"],
        "ref": "ACI 318-19",
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary", "IMF — Intermediate", "SMF — Special"],
        "ref": "ACI 318-14",
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["GS — Grado Reducido", "GM — Grado Moderado", "GA — Grado Alto"],
        "ref": "NEC-SE-HM",
    },
    "E.060 (Perú)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 5.0, "rho_max_des": 5.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["PO — Pórtico Ordinario", "PM — Pórtico Moderado", "PE — Pórtico Especial"],
        "ref": "E.060 Perú",
    },
    "NTC-EM (México)": {
        "phi_tied": 0.70, "phi_spiral": 0.80, "phi_tension": 0.85,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 5.0, "rho_max_des": 5.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["MDL — Ductilidad Limitada", "MROD — Ductilidad Ordinaria", "MRLE — Ductilidad Alta"],
        "ref": "NTC-EM México",
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 6.0, "rho_max_dmo": 5.0, "rho_max_des": 5.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["PO — Pórtico Ordinario", "PM — Pórtico Moderado", "PE — Pórtico Especial"],
        "ref": "COVENIN 1753",
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["DO — Diseño Ordinario", "DM — Ductilidad Moderada", "DE — Diseño Especial"],
        "ref": "NB 1225001",
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max_dmi": 8.0, "rho_max_dmo": 6.0, "rho_max_des": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["GE — Grado Estándar", "GM — Ductilidad Moderada", "GA — Ductilidad Alta"],
        "ref": "CIRSOC 201",
    },
}

# =============================================================================
# RECUBRIMIENTO MÍNIMO EN COLUMNAS POR NORMA (cm)
# =============================================================================
COVER_MIN_COL = {
    "NSR-10 (Colombia)":              3.8,   # NSR-10 C.7.7.1
    "ACI 318-25 (EE.UU.)":            3.8,   # ACI 318-25 §26.4.1.1
    "ACI 318-19 (EE.UU.)":            3.8,   # ACI 318-19 §26.4.1.1
    "ACI 318-14 (EE.UU.)":            3.8,   # ACI 318-14 §7.7.1
    "NEC-SE-HM (Ecuador)":            3.8,   # NEC-SE-HM §7.7
    "E.060 (Perú)":                   4.0,   # E.060 Art. 7.7
    "NTC-EM (México)":                3.5,   # NTC-EM §2.1.4
    "COVENIN 1753-2006 (Venezuela)":  3.8,   # COVENIN 1753
    "NB 1225001-2020 (Bolivia)":      3.8,   # NB 1225001
    "CIRSOC 201-2025 (Argentina)":    3.8,   # CIRSOC 201
}
COVER_REF_COL = {
    "NSR-10 (Colombia)":              "NSR-10 C.7.7.1",
    "ACI 318-25 (EE.UU.)":            "ACI 318-25 §26.4.1.1",
    "ACI 318-19 (EE.UU.)":            "ACI 318-19 §26.4.1.1",
    "ACI 318-14 (EE.UU.)":            "ACI 318-14 §7.7.1",
    "NEC-SE-HM (Ecuador)":            "NEC-SE-HM §7.7",
    "E.060 (Perú)":                   "E.060 Art. 7.7",
    "NTC-EM (México)":                "NTC-EM §2.1.4",
    "COVENIN 1753-2006 (Venezuela)":  "COVENIN 1753 §7.7",
    "NB 1225001-2020 (Bolivia)":      "NB 1225001 §7.7",
    "CIRSOC 201-2025 (Argentina)":    "CIRSOC 201 §7.7",
}

# 
# PRESENTACIONES DE CEMENTO POR PAÍS
# 
CEMENT_BAGS = {
    "NSR-10 (Colombia)": [{"label": "Cemento gris (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "ACI 318-25 (EE.UU.)": [{"label": "Type I/II (94 lb / 42.6 kg)", "kg": 42.6}, {"label": "Type III (47 lb / 21.3 kg)", "kg": 21.3}],
    "ACI 318-19 (EE.UU.)": [{"label": "Type I/II (94 lb / 42.6 kg)", "kg": 42.6}, {"label": "Type III (47 lb / 21.3 kg)", "kg": 21.3}],
    "ACI 318-14 (EE.UU.)": [{"label": "Type I/II (94 lb / 42.6 kg)", "kg": 42.6}, {"label": "Type III (47 lb / 21.3 kg)", "kg": 21.3}],
    "NEC-SE-HM (Ecuador)": [{"label": "Cemento Holcim (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "E.060 (Perú)": [{"label": "Cemento Andino (42.5 kg)", "kg": 42.5}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "NTC-EM (México)": [{"label": "Cemento Cemex (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "COVENIN 1753-2006 (Venezuela)": [{"label": "Cemento (42.5 kg)", "kg": 42.5}],
    "NB 1225001-2020 (Bolivia)": [{"label": "Cemento (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
    "CIRSOC 201-2025 (Argentina)": [{"label": "Cemento (50 kg)", "kg": 50.0}, {"label": "Bolsa pequeña (25 kg)", "kg": 25.0}],
}

# 
# TABLA DE DOSIFICACIÓN ACI 211
# 
MIX_DESIGNS = [
    {"fc_mpa": 14.0, "cem": 250, "agua": 205, "arena": 810, "grava": 1060, "wc": 0.82},
    {"fc_mpa": 17.0, "cem": 290, "agua": 200, "arena": 780, "grava": 1060, "wc": 0.69},
    {"fc_mpa": 21.0, "cem": 350, "agua": 193, "arena": 720, "grava": 1060, "wc": 0.55},
    {"fc_mpa": 25.0, "cem": 395, "agua": 193, "arena": 680, "grava": 1020, "wc": 0.49},
    {"fc_mpa": 28.0, "cem": 430, "agua": 190, "arena": 640, "grava": 1000, "wc": 0.44},
    {"fc_mpa": 35.0, "cem": 530, "agua": 185, "arena": 580, "grava":  960, "wc": 0.35},
    {"fc_mpa": 42.0, "cem": 620, "agua": 180, "arena": 520, "grava":  910, "wc": 0.29},
    {"fc_mpa": 56.0, "cem": 740, "agua": 175, "arena": 450, "grava":  850, "wc": 0.24},
]

def get_mix_for_fc(fc_mpa):
    if fc_mpa <= MIX_DESIGNS[0]["fc_mpa"]: return MIX_DESIGNS[0]
    if fc_mpa >= MIX_DESIGNS[-1]["fc_mpa"]: return MIX_DESIGNS[-1]
    for i in range(len(MIX_DESIGNS)-1):
        lo, hi = MIX_DESIGNS[i], MIX_DESIGNS[i+1]
        if lo["fc_mpa"] <= fc_mpa <= hi["fc_mpa"]:
            t = (fc_mpa - lo["fc_mpa"]) / (hi["fc_mpa"] - lo["fc_mpa"])
            return {k: lo[k] + t*(hi[k]-lo[k]) for k in ("cem", "agua", "arena", "grava", "wc")}
    return MIX_DESIGNS[-1]

def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_development_length(db_mm, fy, fc, lambda_=1.0, psi_t=1.0, psi_e=1.0, psi_s=1.0, psi_g=1.0, cb_ktr=2.5):
    if db_mm <= 0: return 0
    ld = (3/40) * (fy / (lambda_ * math.sqrt(fc))) * (psi_t * psi_e * psi_s * psi_g / cb_ktr) * db_mm
    return max(ld, 300)

# 
# FUNCIONES PARA CÁLCULO DE CAPACIDAD UNIAXIAL
# 
def compute_uniaxial_capacity(b, h, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza):
    eps_cu = 0.003
    eps_y = fy / Es
    beta_1 = get_beta1(fc)
    
    Ag = b * h
    Ast = sum([layer['As'] for layer in layers])
    Po_kN = (0.85 * fc * (Ag * 100 - Ast * 100) + fy * Ast * 100) / 1000.0
    
    c_vals = np.concatenate([np.linspace(1e-5, h, 120), np.linspace(h, h * 12, 60)])
    P_n_list = []; M_n_list = []; phi_P_n_list = []; phi_M_n_list = []
    eps_t_vals = []
    
    b_mm = b * 10
    h_mm = h * 10
    
    for c_cm in c_vals:
        c_mm = c_cm * 10
        a_mm = min(beta_1 * c_mm, h_mm)
        Cc = 0.85 * fc * a_mm * b_mm
        Mc = Cc * (h_mm / 2.0 - a_mm / 2.0)
        
        Ps = 0.0; Ms = 0.0; eps_t = 0.0
        for layer in layers:
            d_i_mm = layer['d'] * 10
            As_i = layer['As'] * 100
            eps_s = eps_cu * (c_mm - d_i_mm) / c_mm
            if d_i_mm >= max(l['d'] * 10 for l in layers):
                eps_t = eps_s
            fs = max(-fy, min(fy, Es * eps_s))
            if a_mm > d_i_mm and fs > 0:
                fs -= 0.85 * fc
            Ps += As_i * fs
            Ms += As_i * fs * (h_mm / 2.0 - d_i_mm)
        
        Pn = (Cc + Ps) / 1000.0
        Mn = abs((Mc + Ms) / 1_000_000.0)
        eps_t_tens = -eps_t
        
        if eps_t_tens <= eps_y:
            phi = phi_c_max
        elif eps_t_tens >= eps_full:
            phi = phi_tension
        else:
            phi = phi_c_max + (phi_tension - phi_c_max) * (eps_t_tens - eps_y) / (eps_full - eps_y)
        
        Pn_max_val = p_max_factor * Po_kN
        phi_Pn_max_val = min(phi_c_max*Pn_max_val, 0.80*phi_c_max*Po_kN)  # ACI 318-19 §22.4.2.1
        
        Pn = min(Pn, Pn_max_val)
        phi_Pn = min(phi * Pn, phi_Pn_max_val)
        phi_Mn = phi * Mn
        
        P_n_list.append(Pn)
        M_n_list.append(Mn)
        phi_P_n_list.append(phi_Pn)
        phi_M_n_list.append(phi_Mn)
        eps_t_vals.append(eps_t_tens)
    
    c_balance = None
    P_balance = None
    M_balance = None
    for i, eps in enumerate(eps_t_vals):
        if abs(eps - eps_y) < 0.0001 or (i > 0 and eps_t_vals[i-1] <= eps_y <= eps):
            idx = i if abs(eps - eps_y) < 0.0001 else i-1
            c_balance = c_vals[idx]
            P_balance = P_n_list[idx] * factor_fuerza
            M_balance = M_n_list[idx] * factor_fuerza
            break
    
    return {
        'M_n': np.array(M_n_list) * factor_fuerza,
        'P_n': np.array(P_n_list) * factor_fuerza,
        'phi_M_n': np.array(phi_M_n_list) * factor_fuerza,
        'phi_P_n': np.array(phi_P_n_list) * factor_fuerza,
        'Po': Po_kN * factor_fuerza,
        'Pn_max': p_max_factor * Po_kN * factor_fuerza,
        'phi_Pn_max': phi_c_max * p_max_factor * Po_kN * factor_fuerza,
        'c_balance': c_balance,
        'P_balance': P_balance,
        'M_balance': M_balance,
    }

def compute_uniaxial_capacity_circular(D, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza):
    """Para sección circular"""
    eps_cu = 0.003
    eps_y = fy / Es
    beta_1 = get_beta1(fc)
    
    r = D / 2
    Ag = math.pi * r**2
    Ast = sum([layer['As'] for layer in layers])
    Po_kN = (0.85 * fc * (Ag * 100 - Ast * 100) + fy * Ast * 100) / 1000.0
    
    c_vals = np.concatenate([
        np.linspace(1e-5, D * 0.1, 40),
        np.linspace(D * 0.1, D * 0.5, 80),
        np.linspace(D * 0.5, D, 60),
        np.linspace(D, D * 12, 40)
    ])
    P_n_list = []; M_n_list = []; phi_P_n_list = []; phi_M_n_list = []
    eps_t_vals = []
    
    for c_cm in c_vals:
        c_mm = c_cm * 10
        if c_mm >= D * 10:
            a_mm = D * 10
        else:
            a_mm = beta_1 * c_mm
        a_cm = a_mm / 10
        if a_cm >= D:
            Ac_comp = Ag
        else:
            h_seg = a_cm
            _arg_ac = np.clip((r - h_seg) / r, -1.0, 1.0)
            Ac_comp = r**2 * math.acos(_arg_ac) - (r - h_seg) * math.sqrt(max(0.0, 2*r*h_seg - h_seg**2))
        
        Cc = 0.85 * fc * Ac_comp * 100
        if a_cm >= D:
            y_cent = r
        else:
            _arg_yc = np.clip((r - a_cm) / r, -1.0, 1.0)
            _angle = math.acos(_arg_yc)
            y_cent = (4*r * math.sin(_angle / 2)**3) / (3 * (_angle - math.sin(_angle) * math.cos(_angle)))
        Mc = Cc * (r * 10 - y_cent * 10)
        
        Ps = 0.0; Ms = 0.0; eps_t = 0.0
        for layer in layers:
            d_i_mm = layer['d'] * 10
            As_i = layer['As'] * 100
            eps_s = eps_cu * (c_mm - d_i_mm) / c_mm
            if d_i_mm >= max(l['d'] * 10 for l in layers):
                eps_t = eps_s
            fs = max(-fy, min(fy, Es * eps_s))
            Ps += As_i * fs
            Ms += As_i * fs * (r * 10 - d_i_mm)
        
        Pn = (Cc + Ps) / 1000.0
        Mn = abs((Mc + Ms) / 1000000.0)
        eps_t_tens = -eps_t
        
        if eps_t_tens <= eps_y:
            phi = phi_c_max
        elif eps_t_tens >= eps_full:
            phi = phi_tension
        else:
            phi = phi_c_max + (phi_tension - phi_c_max) * (eps_t_tens - eps_y) / (eps_full - eps_y)
        
        Pn_max_val = p_max_factor * Po_kN
        phi_Pn_max_val = min(phi_c_max*Pn_max_val, 0.80*phi_c_max*Po_kN)  # ACI 318-19 §22.4.2.1
        
        Pn = min(Pn, Pn_max_val)
        phi_Pn = min(phi * Pn, phi_Pn_max_val)
        phi_Mn = phi * Mn
        
        P_n_list.append(Pn)
        M_n_list.append(Mn)
        phi_P_n_list.append(phi_Pn)
        phi_M_n_list.append(phi_Mn)
        eps_t_vals.append(eps_t_tens)
    
    c_balance = None
    P_balance = None
    M_balance = None
    for i, eps in enumerate(eps_t_vals):
        if abs(eps - eps_y) < 0.0001 or (i > 0 and eps_t_vals[i-1] <= eps_y <= eps):
            idx = i if abs(eps - eps_y) < 0.0001 else i-1
            c_balance = c_vals[idx]
            P_balance = P_n_list[idx] * factor_fuerza
            M_balance = M_n_list[idx] * factor_fuerza
            break
    
    return {
        'M_n': np.array(M_n_list) * factor_fuerza,
        'P_n': np.array(P_n_list) * factor_fuerza,
        'phi_M_n': np.array(phi_M_n_list) * factor_fuerza,
        'phi_P_n': np.array(phi_P_n_list) * factor_fuerza,
        'Po': Po_kN * factor_fuerza,
        'Pn_max': p_max_factor * Po_kN * factor_fuerza,
        'phi_Pn_max': phi_c_max * p_max_factor * Po_kN * factor_fuerza,
        'c_balance': c_balance,
        'P_balance': P_balance,
        'M_balance': M_balance,
    }

# 
# FUNCIÓN PARA VERIFICACIÓN BIAXIAL (BRESLER)
# 
def interp_pm_curve(M_query, phi_Mn_arr, phi_Pn_arr):
    """
    Interpola φPn dado φMn en la curva P-M NO monótona.
    La curva tiene dos ramas: compresión (M creciente) y tensión (M decreciente).
    Separamos ambas ramas e interpolamos en cada una, retornando el MAYOR φPn
    (el de la rama de compresión, que es el correcto para Bresler).
    """
    phi_Mn_arr = np.array(phi_Mn_arr)
    phi_Pn_arr = np.array(phi_Pn_arr)
    
    if len(phi_Mn_arr) == 0:
        return 0.0
    
    M_max = np.max(phi_Mn_arr)
    
    # Si M=0 → capacidad axial pura (Po)
    if M_query <= 0:
        return float(phi_Pn_arr[np.argmax(phi_Pn_arr)])
    
    # Si el momento pedido supera el máximo → sentinel -1.0 (distinguible de Pn=0 real)
    if M_query > M_max:
        return np.nan  # sentinel: M excede el diagrama en este eje
    
    # Índice del punto de balance (donde φMn es máximo)
    idx_bal = int(np.argmax(phi_Mn_arr))
    
    #  Rama COMPRESIÓN: Po → Balance (M creciente, P decreciente) 
    Mc = phi_Mn_arr[:idx_bal + 1]
    Pc = phi_Pn_arr[:idx_bal + 1]
    
    #  Rama TENSIÓN: Balance → Tracción pura (M decreciente, P→0) 
    Mt = phi_Mn_arr[idx_bal:]
    Pt = phi_Pn_arr[idx_bal:]
    
    # Ordenar cada rama para np.interp (requiere x creciente)
    sc = np.argsort(Mc)
    st = np.argsort(Mt)
    
    P_comp = float(np.interp(M_query, Mc[sc], Pc[sc],
                              left=float(Pc[sc[0]]), right=0.0))
    P_tens = float(np.interp(M_query, Mt[st], Pt[st],
                              left=0.0, right=0.0))
    
    # Retornar el MAYOR (rama de compresión es siempre mayor)
    return max(P_comp, P_tens)


def biaxial_bresler(Pu, Mux, Muy, cap_x, cap_y, Po, phi_factor):
    phi_Pnx, phi_Pny, phi_P0 = None, None, None
    phi_Pnx = interp_pm_curve(abs(Mux), np.array(cap_x['phi_M_n']), np.array(cap_x['phi_P_n']))
    phi_Pny = interp_pm_curve(abs(Muy), np.array(cap_y['phi_M_n']), np.array(cap_y['phi_P_n']))
    phi_P0  = phi_factor * Po

    if phi_Pnx is None or phi_Pny is None:
        return {
            'phi_Pnx': 0.0,
            'phi_Pny': 0.0,
            'phi_P0':  phi_P0,
            'phi_Pni': 0.0,
            'ratio':   float('inf'),
            'ok':      False
        }

    # Detectar sentinel np.nan: el momento supera el diagrama en ese eje
    eje_excedido = []
    if np.isnan(phi_Pnx):
        eje_excedido.append("X")
    if np.isnan(phi_Pny):
        eje_excedido.append("Y")

    if eje_excedido:
        msg = f"Mux/Muy excede el diagrama P-M en eje {'&'.join(eje_excedido)}: aumentar sección o acero"
        return {
            'phi_Pnx': 0.0,
            'phi_Pny': 0.0,
            'phi_P0':  phi_P0,
            'phi_Pni': 0.0,
            'ratio':   float('inf'),
            'ok':      False,
            'msg_exceso': msg,
        }

    if phi_Pnx > 0 and phi_Pny > 0 and phi_P0 > 0:
        inv = 1/phi_Pnx + 1/phi_Pny - 1/phi_P0
        phi_Pni = 1/inv if inv > 0 else 0.0
    else:
        phi_Pni = 0.0

    ratio = Pu / phi_Pni if phi_Pni > 0 else float('inf')
    ok    = Pu <= phi_Pni

    return {
        'phi_Pnx': phi_Pnx,
        'phi_Pny': phi_Pny,
        'phi_P0':  phi_P0,
        'phi_Pni': phi_Pni,
        'ratio':   ratio,
        'ok':      ok,
        'msg_exceso': None,
    }

# 
# FUNCIÓN PARA ESBELTEZ (NSR-10 C.10.10 / ACI 6.6)
# 
def check_slenderness(L, b, h, k, beta_dns, Pu, M1, M2, fc, fy, Es, factor_fuerza, es_circular=False):
    r = 0.25 * b if es_circular else min(b, h) / math.sqrt(12)
    kl = k * L
    kl_r = kl / r if r > 0 else 999
    
    if kl_r <= 22:
        slender = False
        classification = "Columna corta (sin efectos de segundo orden)"
        delta_ns = 1.0
    elif kl_r <= 100:
        slender = True
        classification = "Columna esbelta (requiere magnificación de momentos)"
        Ec = 4700 * math.sqrt(fc)
        # Bug Fix: para sección circular usar Ig = π*D⁴/64, no b*h³/12
        if es_circular:
            Ig = math.pi * b**4 / 64   # b == D cuando se llama para circular
        else:
            Ig = b * h**3 / 12
        Ig_mm4 = Ig * 1e4
        EI = 0.4 * Ec * Ig_mm4 / (1 + beta_dns)
        Pc = math.pi**2 * EI / (kl * 1000)**2
        Pc = Pc / 1000
        Cm = 0.6 + 0.4 * (M1/M2) if abs(M2) > 0 else 1.0
        Cm = max(0.4, Cm)
        # Bug Fix N3: Prevenir división por cero si Pu se acerca a la carga crítica (Pandeo)
        denom = 1 - Pu / (0.75 * Pc)
        if denom <= 0.01:
            denom = 0.01
        delta_ns = Cm / denom
        delta_ns = max(delta_ns, 1.0)
    else:
        slender = True
        classification = "Columna muy esbelta (kl/r > 100) — requiere análisis no lineal según NSR-10 C.10.10.7"
        delta_ns = 1.0
    
    return {
        'kl_r': kl_r,
        'slender': slender,
        'classification': classification,
        'delta_ns': delta_ns,
        'r': r,
        'kl': kl
    }

# =============================================================================
# SIDEBAR - ENTRADAS DEL USUARIO
# =============================================================================


# =============================================================================
# PERSISTENCIA LOCAL — Columnas PM
# =============================================================================
import base64 as _b64, json as _json, os as _os

_COL_PERSIST_KEYS = [
    "cpm_empresa","cpm_proyecto_nombre","cpm_ingeniero","cpm_elaboro","cpm_reviso","cpm_aprobo",
    "cpm_plano_num","cpm_escala","col_logo_bytes","user_role","nombre_proyecto_actual",
    "c_pm_norma","c_pm_nivel_sismico","c_pm_fc_unit","c_pm_fc_mpa","c_pm_fy",
    "c_pm_seccion_type","c_pm_b","c_pm_h","c_pm_D","c_pm_dprime","c_pm_L",
    "c_pm_unit_system","c_pm_rebar_type","c_pm_num_h","c_pm_num_v","c_pm_n_barras_circ",
    "c_pm_col_type","c_pm_stirrup_type","c_pm_spiral_type","c_pm_paso_circ","c_pm_paso_rect",
    "c_pm_output_units","c_pm_mux","c_pm_muy","c_pm_pu","c_pm_cm_chk","c_pm_m1x","c_pm_m1y",
    "c_pm_beta_dns","c_pm_k_sel",
]

def save_state_col():
    snap = {}
    for k in _COL_PERSIST_KEYS:
        val = st.session_state.get(k)
        if k == "col_logo_bytes" and isinstance(val, (bytes, bytearray)):
            snap[k] = {"type": "bytes_b64", "data": _b64.b64encode(val).decode()}
        elif val is not None:
            try:    snap[k] = val
            except: snap[k] = str(val)
    try:
        with open("col_state.json", "w", encoding="utf-8") as f:
            _json.dump(snap, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def load_state_col():
    if not _os.path.exists("col_state.json"):
        return
    try:
        with open("col_state.json", "r", encoding="utf-8") as f:
            data = _json.load(f)
        for k, v in data.items():
            if k not in st.session_state:
                if isinstance(v, dict) and v.get("type") == "bytes_b64":
                    st.session_state[k] = _b64.b64decode(v["data"])
                else:
                    st.session_state[k] = v
    except Exception:
        pass

if "col_state_loaded" not in st.session_state:
    load_state_col()
    st.session_state["col_state_loaded"] = True


# =============================================================================
# PASO 1 — SIDEBAR: Identidad, logo, norma, rol, plano, guardar/cargar
#           NO contiene inputs de calculo. Solo identidad y configuracion global.
# =============================================================================

with st.sidebar.expander("Identidad del Proyecto", expanded=False):
    st.session_state["cpm_empresa"]         = st.text_input("Empresa / Firma",
        value=st.session_state.get("cpm_empresa",""), key="_wb_e")
    st.session_state["cpm_proyecto_nombre"] = st.text_input("Nombre del Proyecto",
        value=st.session_state.get("cpm_proyecto_nombre",""), key="_wb_p")
    st.session_state["cpm_ingeniero"]       = st.text_input("Ingeniero Responsable",
        value=st.session_state.get("cpm_ingeniero",""), key="_wb_i")
    _c1s, _c2s = st.columns(2)
    st.session_state["cpm_elaboro"] = _c1s.text_input("Elaboro",
        value=st.session_state.get("cpm_elaboro",""), key="_wb_elab")
    st.session_state["cpm_reviso"]  = _c2s.text_input("Reviso",
        value=st.session_state.get("cpm_reviso",""),  key="_wb_rev")
    st.session_state["cpm_aprobo"]  = st.text_input("Aprobo",
        value=st.session_state.get("cpm_aprobo",""), key="_wb_apr")
    if st.button("Guardar identidad", key="btn_id_col", use_container_width=True):
        save_state_col()
        st.success("Identidad guardada")

st.sidebar.markdown("**Logo de empresa**")
st.sidebar.caption("Aparece en portada de DOCX y PDF. PNG/JPG, fondo blanco recomendado.")
_logo_file = st.sidebar.file_uploader("Subir logo", type=["png","jpg","jpeg"], key="col_logo_upload")
if _logo_file is not None:
    _raw = _logo_file.read()
    if _raw:
        st.session_state["col_logo_bytes"] = _raw
        save_state_col()
        st.sidebar.success("Logo cargado")
if st.session_state.get("col_logo_bytes"):
    try:
        st.sidebar.image(st.session_state["col_logo_bytes"], width=140, caption="Logo activo")
    except Exception:
        st.sidebar.caption("Logo cargado (no previsualizable)")
    _la, _lb = st.sidebar.columns(2)
    if _la.button("Quitar logo", key="col_rm_logo", use_container_width=True):
        st.session_state.pop("col_logo_bytes", None)
        save_state_col(); st.rerun()
    _lb.download_button("Descargar", data=st.session_state["col_logo_bytes"],
                        file_name="logo.png", mime="image/png",
                        use_container_width=True, key="col_dl_logo")
else:
    st.sidebar.caption("Sin logo — documentos solo texto.")

st.sidebar.markdown("---")

# ── Modo Revisión (solo lectura) ─────────────────────────────────────────
if "c_pm_modo_revision" not in st.session_state:
    st.session_state["c_pm_modo_revision"] = False

_col_rev1, _col_rev2 = st.sidebar.columns([3, 1])
with _col_rev1:
    st.sidebar.markdown(
        "<span style='font-size:0.85em;color:#aaa'>Modo Revisión</span>",
        unsafe_allow_html=True)
modo_revision = st.sidebar.toggle(
    "🔍 Modo Revisión",
    value=st.session_state.get("c_pm_modo_revision", False),
    key="c_pm_modo_revision",
    help="Oculta todos los paneles de entrada. Solo muestra resultados y verificaciones. "
         "Ideal para revisores externos."
)
if modo_revision:
    st.sidebar.info(
        "**Modo Revisión activo** — Paneles de entrada ocultos.  \n"
        "Desactiva el toggle para editar parámetros.")
st.sidebar.markdown("---")
# ── Fin Modo Revisión toggle ──────────────────────────────────────────────

if not modo_revision:
    st.sidebar.header("0. Norma de Diseno")
norma_options = list(CODES.keys())
norma_sel = st.sidebar.selectbox("Norma", norma_options, key="c_pm_norma",
                                  disabled=modo_revision)
mostrar_referencias_norma(norma_sel, "columnas_pm")
code = CODES[norma_sel]

nivel_sismico = st.sidebar.selectbox(
    "Nivel Sismico / Ductilidad:",
    code["seismic_levels"],
    key="c_pm_nivel_sismico"
)
nivel_lower = nivel_sismico.lower()
es_des = any(k in nivel_lower for k in ["des","disipacion especial","smf","special","ga","ductilidad alta","pe","portico especial","de","diseno especial","mrle","alta"])
es_dmo = any(k in nivel_lower for k in ["dmo","imf","intermediate","gm","moderada","pm","mrod","media","moderado"]) and not es_des
es_dmi = not (es_des or es_dmo)

# ── ΣMnv nodo (columna fuerte/viga débil) — solo visible en DES/DMO ──────
if es_des or es_dmo:
    st.sidebar.markdown("---")
    st.sidebar.markdown("**Columna Fuerte / Viga Débil — C.21.6.1**")
    st.sidebar.number_input(
        "ΣMnv vigas en nodo [kN·m]",
        min_value=0.0, max_value=50000.0,
        value=float(st.session_state.get("c_pm_Mn_viga_kNm", 0.0)),
        step=10.0,
        key="c_pm_Mn_viga_kNm",
        help="Suma de momentos nominales de TODAS las vigas que llegan al nodo (NSR-10 C.21.6.1). "
             "Se requiere ΣMnc ≥ 1.20·ΣMnv. Ingrese 0 para omitir la verificación."
    )
    # ── Nodo Viga-Columna C.21.7 ──────────────────────────────────────────
    st.sidebar.markdown("---")
    st.sidebar.markdown("**Nodo Viga-Columna — C.21.7.4.1**")
    st.sidebar.number_input(
        "As vigas en nodo [cm²]",
        min_value=0.0, max_value=500.0,
        value=float(st.session_state.get("c_pm_As_vigas_nodo", 0.0)),
        step=0.5,
        key="c_pm_As_vigas_nodo",
        help="Acero de tracción total de las vigas que llegan al nodo (sum As, cm²). "
             "Ingrese 0 para omitir la verificación del nodo."
    )
    st.sidebar.number_input(
        "Vcol diseño [kN]",
        min_value=0.0, max_value=5000.0,
        value=float(st.session_state.get("c_pm_Vu_col_nodo", 0.0)),
        step=5.0,
        key="c_pm_Vu_col_nodo",
        help="Cortante de diseño en la columna (Vu [kN]) para el cuerpo libre del nodo."
    )
    st.sidebar.selectbox(
        "Confinamiento del nodo",
        ["4_caras", "3_caras", "otros"],
        index=["4_caras", "3_caras", "otros"].index(
            st.session_state.get("c_pm_nodo_conf", "3_caras")),
        key="c_pm_nodo_conf",
        format_func=lambda x: {
            "4_caras": "4 caras (γ=1.70) — Nodo interior",
            "3_caras": "3 caras (γ=1.25) — Nodo borde",
            "otros":   "Otras (γ=1.00) — Nodo esquina"
        }[x],
        help="NSR-10 C.21.7.4.1 — El confinamiento define el coeficiente γ de resistencia al cortante."
    )
# ── Fin ΣMnv y Nodo ───────────────────────────────────────────────────────

if es_des:   rho_max = code["rho_max_des"]
elif es_dmo: rho_max = code["rho_max_dmo"]
else:        rho_max = code["rho_max_dmi"]
rho_min = code["rho_min"]

st.sidebar.caption(f"Referencia: {code['ref']}")
st.sidebar.caption(f"rho max segun nivel: {rho_max}% | rho min: {rho_min}%")

with st.sidebar.expander("Rol de Usuario", expanded=False):
    _roles    = ["free", "pro", "admin"]
    _rol_idx  = _roles.index(st.session_state.get("user_role","free"))
    _rol_nuevo = st.radio("Selecciona tu rol", _roles, index=_rol_idx, key="col_rol_radio",
                          help="free: calculos y CSV.  pro: todos los entregables.  admin: acceso total.")
    if _rol_nuevo != st.session_state.get("user_role","free"):
        st.session_state["user_role"] = _rol_nuevo
        save_state_col(); st.rerun()
    _crol = {"admin":"#1b5e20","pro":"#0d47a1","free":"#b71c1c"}
    st.markdown(
        f'<div style="background:{_crol[_rol_nuevo]};color:white;border-radius:6px;'
        f'padding:5px 10px;font-size:12px;text-align:center;font-weight:600;margin-top:6px">'
        f'Rol activo: {_rol_nuevo.upper()}</div>', unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.header("Datos del Plano")
_plano_num    = st.sidebar.text_input("Numero de plano",
    value=st.session_state.get("cpm_plano_num","COL-001"), key="cpm_plano_inp")
_escala_plano = st.sidebar.text_input("Escala",
    value=st.session_state.get("cpm_escala","1:50"), key="cpm_escala_inp")
st.session_state["cpm_plano_num"] = _plano_num
st.session_state["cpm_escala"]    = _escala_plano

st.sidebar.markdown("---")
st.sidebar.subheader("Guardar / Cargar Proyecto")
_nombre_producido    = st.session_state.get("nombre_proyecto_actual","")
_nombre_proy_guardar = st.sidebar.text_input("Nombre para guardar",
    value=_nombre_producido, key="col_input_proy")
if st.sidebar.button("Guardar Proyecto", use_container_width=True):
    if _nombre_proy_guardar:
        _ok, _msg = guardar_proyecto_supabase(_nombre_proy_guardar, capturar_estado_actual())
        if _ok:
            st.session_state["nombre_proyecto_actual"] = _nombre_proy_guardar
            st.sidebar.success(_msg); st.rerun()
        else:
            st.sidebar.error(_msg)
    else:
        st.sidebar.warning("Escribe un nombre de proyecto")

st.sidebar.markdown("**Cargar Proyecto Existente**")
_lista_proy = listar_proyectos_supabase()
if _lista_proy:
    _idx_def = 0
    if _nombre_producido in _lista_proy:
        _idx_def = _lista_proy.index(_nombre_producido)
    _nombre_proy_cargar = st.sidebar.selectbox("Selecciona un proyecto",
        _lista_proy, index=_idx_def, key="col_sel_proy")
    def _on_cargar_col():
        _p = st.session_state["col_sel_proy"]
        if _p:
            _ok2, _msg2 = cargar_proyecto_supabase(_p)
            if _ok2:
                st.session_state["nombre_proyecto_actual"] = _p
                st.session_state["__msg_cargar_col"] = (True, _msg2)
            else:
                st.session_state["__msg_cargar_col"] = (False, _msg2)
    st.sidebar.button("Cargar", on_click=_on_cargar_col, use_container_width=True)
    if "__msg_cargar_col" in st.session_state:
        _ok3, _msg3 = st.session_state.pop("__msg_cargar_col")
        if _ok3: st.sidebar.success(_msg3)
        else:    st.sidebar.error(_msg3)
else:
    st.sidebar.info("No hay proyectos guardados.")

st.sidebar.markdown("---")
_emp_foot = st.session_state.get("cpm_empresa","") or "StructoPro"
st.sidebar.markdown(
    f'<div style="text-align:center;color:gray;font-size:10px">'
    f'&copy; 2026 {_emp_foot}<br>Uso profesional exclusivo</div>',
    unsafe_allow_html=True)


# =============================================================================
# PASO 2 — Leer variables desde session_state con defaults
#           SIN crear widgets. Los widgets se crean en render_config_tab()
# =============================================================================

fc_unit     = st.session_state.get("c_pm_fc_unit", "MPa")
if fc_unit == "PSI":
    _fc_psi = float(st.session_state.get("c_pm_fc_psi_val", 3000.0))
    fc      = _fc_psi * 0.00689476
elif fc_unit == "kg/cm2":
    _fc_kg  = float(st.session_state.get("c_pm_fc_kgcm2_val", 210.0))
    fc      = _fc_kg / 10.1972
else:
    fc      = float(st.session_state.get("c_pm_fc_mpa", 21.0))

fy          = float(st.session_state.get("c_pm_fy", 420.0))
Es          = 200000.0

# ── Advertencia HSC — ACI 318-19 §22.2.2.4.3 / NSR-10 C.22.2.2 ──────────
if fc >= 55.0:
    st.warning(
        f"⚠️ **Concreto de Alta Resistencia (f'c = {fc:.1f} MPa ≥ 55 MPa)**  \n"
        "ACI 318-19 §22.2.2.4.3 y NSR-10 C.22.2.2 requieren consideraciones adicionales:  \n"
        "• β₁ mínimo = 0.65 (ya aplicado automáticamente).  \n"
        "• Para f'c ≥ 70 MPa se recomienda modelo no-lineal confinado (Mander et al. 1988) "
        "que el módulo no implementa aún.  \n"
        "• Verifique que el proveedor garantice la resistencia con ensayos según NTC 673."
    )
# ── Fin advertencia HSC ───────────────────────────────────────────────────

seccion_type = st.session_state.get("c_pm_seccion_type", "Rectangular / Cuadrada")
es_circular  = "Circular" in seccion_type
D            = float(st.session_state.get("c_pm_D", 40.0))
b            = float(st.session_state.get("c_pm_b", 30.0))
h            = float(st.session_state.get("c_pm_h", 40.0))
if es_circular:
    b = D; h = D
d_prime      = float(st.session_state.get("c_pm_dprime", 5.0))
L_col        = float(st.session_state.get("c_pm_L", 300.0))

limite_d = D/2 if es_circular else min(b/2, h/2)
if d_prime >= limite_d:
    st.error(
        f"Recubrimiento d' ({d_prime:.1f} cm) supera el nucleo de la seccion "
        f"(limite: {limite_d:.1f} cm). Corrija en la pestana Configuracion."
    )
    st.stop()

unit_system  = st.session_state.get("c_pm_unit_system", "Milimetros (SI)")
_usar_us     = "Pulgadas" in unit_system or "Inches" in unit_system
rebar_dict   = REBAR_US if _usar_us else REBAR_MM
rebar_type   = st.session_state.get("c_pm_rebar_type", list(rebar_dict.keys())[2])
if rebar_type not in rebar_dict:
    rebar_type = list(rebar_dict.keys())[2]
rebar_area   = rebar_dict[rebar_type]["area"]
rebar_diam   = rebar_dict[rebar_type]["diam_mm"]

if es_circular:
    n_barras     = int(st.session_state.get("c_pm_n_barras_circ", 8))
    Ast          = n_barras * rebar_area
    layers       = []
    angulos      = np.linspace(0, 2*np.pi, n_barras, endpoint=False)
    radio_centro = D/2 - d_prime
    for i, ang in enumerate(angulos):
        x_pos = radio_centro * math.cos(ang)
        y_pos = radio_centro * math.sin(ang)
        layers.append({"d": D/2 + y_pos, "As": rebar_area, "x": x_pos, "y": y_pos})
    num_filas_h = n_barras
    num_filas_v = n_barras
else:
    num_filas_h = int(st.session_state.get("c_pm_num_h", 2))
    num_filas_v = int(st.session_state.get("c_pm_num_v", 2))
    layers      = []
    layers.append({"d": d_prime, "As": num_filas_h * rebar_area})
    _num_int    = num_filas_v - 2
    num_capas_intermedias = _num_int  # alias para compatibilidad con tabs
    if _num_int > 0:
        _esp_h  = (h - 2 * d_prime) / (_num_int + 1)
        for i in range(1, _num_int + 1):
            layers.append({"d": d_prime + i * _esp_h, "As": 2 * rebar_area})
    layers.append({"d": h - d_prime, "As": num_filas_h * rebar_area})
    n_barras_total = num_filas_h * 2 + (num_filas_v - 2) * 2
    n_barras       = n_barras_total

    # --- Verificación espaciamiento físico NSR-10 C.10.8.1 / ACI 318-19 25.2.3 ---
    # stirrup_diam puede no estar definido aún → leerlo de session_state de forma segura
    _st_cfg    = st.session_state
    _strdcfg_v = STIRRUP_US if "Pulgadas" in _st_cfg.get("cpm_unitsystem", "") else STIRRUP_MM
    _stir_key  = _st_cfg.get("cpm_stirrup_type", list(_strdcfg_v.keys())[0])
    _stir_key  = _stir_key if _stir_key in _strdcfg_v else list(_strdcfg_v.keys())[0]
    _stir_diam_v = _strdcfg_v[_stir_key]["diam_mm"]   # mm
    _db_cm       = rebar_diam / 10.0                   # diámetro barra en cm
    _stir_cm     = _stir_diam_v / 10.0                 # diámetro estribo en cm
    _recub_util  = d_prime - _stir_cm - _db_cm / 2.0   # recubrimiento libre al nucleo
    _ancho_util_x = b - 2.0 * _recub_util              # ancho neto cara X
    _ancho_util_y = h - 2.0 * _recub_util              # alto  neto cara Y
    _esp_min_cm  = max(_db_cm, 2.5)                    # NSR-10 C.10.8.1
    _esp_libre_x = (_ancho_util_x - num_filas_h * _db_cm) / max(num_filas_h - 1, 1)
    _esp_libre_y = (_ancho_util_y - num_filas_v * _db_cm) / max(num_filas_v - 1, 1)
    _barras_no_caben = []
    if num_filas_h > 1 and _esp_libre_x < _esp_min_cm:
        _barras_no_caben.append(
            f'Cara X ({num_filas_h} barras): espacio libre {_esp_libre_x:.1f} cm '
            f'< min {_esp_min_cm:.1f} cm')
    if num_filas_v > 1 and _esp_libre_y < _esp_min_cm:
        _barras_no_caben.append(
            f'Cara Y ({num_filas_v} barras): espacio libre {_esp_libre_y:.1f} cm '
            f'< min {_esp_min_cm:.1f} cm')
    if _barras_no_caben:
        st.sidebar.error(
            'AVISO C.10.8.1: Las barras NO caben fisicamente en la seccion.\n'
            + '\n'.join(_barras_no_caben)
            + '\nReduzca numero de barras o aumente la seccion.')
    # --- Fin verificación espaciamiento ---
    Ast            = sum([l["As"] for l in layers])

Ag      = (math.pi * (D/2)**2) if es_circular else (b * h)
cuantia = Ast / Ag * 100

if es_circular:
    col_type     = "Espiral (Spiral)"
    stirrup_dict = STIRRUP_US if _usar_us else STIRRUP_MM
    spiral_type  = st.session_state.get("c_pm_spiral_type", list(stirrup_dict.keys())[0])
    if spiral_type not in stirrup_dict: spiral_type = list(stirrup_dict.keys())[0]
    stirrup_area = stirrup_dict[spiral_type]["area"]
    stirrup_diam = stirrup_dict[spiral_type]["diam_mm"]
    paso_espiral = float(st.session_state.get("c_pm_paso_circ", 7.5))
else:
    col_type     = st.session_state.get("c_pm_col_type", "Estribos (Tied)")
    stirrup_dict = STIRRUP_US if _usar_us else STIRRUP_MM
    stirrup_type = st.session_state.get("c_pm_stirrup_type", list(stirrup_dict.keys())[0])
    if stirrup_type not in stirrup_dict: stirrup_type = list(stirrup_dict.keys())[0]
    stirrup_area = stirrup_dict[stirrup_type]["area"]
    stirrup_diam = stirrup_dict[stirrup_type]["diam_mm"]
    if stirrup_diam < 9.5: stirrup_diam = 9.53
    paso_espiral = float(st.session_state.get("c_pm_paso_rect", 7.5))

if "Espiral" in col_type or es_circular:
    phi_c_max    = code["phi_spiral"]
    p_max_factor = code["pmax_spiral"]
else:
    phi_c_max    = code["phi_tied"]
    p_max_factor = code["pmax_tied"]
phi_tension  = code["phi_tension"]
eps_full     = code["eps_tension_full"]

unidades_salida = st.session_state.get("c_pm_output_units", "KiloNewtons (kN, kN-m)")
if "Toneladas" in unidades_salida or "tonf" in unidades_salida:
    factor_fuerza = 0.1019716; unidad_fuerza = "tonf"; unidad_mom = "tonf-m"
else:
    factor_fuerza = 1.0; unidad_fuerza = "kN"; unidad_mom = "kN-m"

Mux_input  = float(st.session_state.get("c_pm_mux",  round(45.0   * factor_fuerza, 2)))
Muy_input  = float(st.session_state.get("c_pm_muy",  round(25.0   * factor_fuerza, 2)))
Pu_input   = float(st.session_state.get("c_pm_pu",   round(2700.0 * factor_fuerza, 2)))
aplicar_cm = bool(st.session_state.get("c_pm_cm_chk", False))
if aplicar_cm:
    M1x_input = float(st.session_state.get("c_pm_m1x", 0.0))
    M1y_input = float(st.session_state.get("c_pm_m1y", 0.0))
else:
    M1x_input = Mux_input
    M1y_input = Muy_input

_dim_x_emin = h if not es_circular else D
_dim_y_emin = b if not es_circular else D
e_min_x_cm  = max(0.10 * _dim_x_emin, 1.5)
e_min_y_cm  = max(0.10 * _dim_y_emin, 1.5)
_Mux_emin   = Pu_input * e_min_x_cm / 100.0
_Muy_emin   = Pu_input * e_min_y_cm / 100.0
_Mux_safe   = Mux_input if abs(Mux_input) >= _Mux_emin else (math.copysign(_Mux_emin, Mux_input) if Mux_input != 0 else _Mux_emin)
_Muy_safe   = Muy_input if abs(Muy_input) >= _Muy_emin else (math.copysign(_Muy_emin, Muy_input) if Muy_input != 0 else _Muy_emin)
_Mux_safe   = _Mux_safe if abs(_Mux_safe) > 1e-9 else _Mux_emin
_Muy_safe   = _Muy_safe if abs(_Muy_safe) > 1e-9 else _Muy_emin
Mux_plot    = _Mux_safe
Muy_plot    = _Muy_safe

beta_dns = float(st.session_state.get("c_pm_beta_dns", 0.6))
_k_options = {
    "Ambos extremos articulados":               1.0,
    "Un extremo articulado, otro empotrado":     0.7,
    "Ambos extremos empotrados":                0.5,
    "Voladizo (base empotrada, libre arriba)":  2.0,
}
_k_sel   = st.session_state.get("c_pm_k_sel", "Ambos extremos articulados")
if _k_sel not in _k_options: _k_sel = "Ambos extremos articulados"
k_factor = _k_options[_k_sel]

# Tipo de pórtico — No desplazable vs Desplazable (NSR-10 C.10.10 / ACI 318-19 §6.6.4)
_portico_tipo  = st.session_state.get("c_pm_portico_tipo", "No desplazable (Pórtico arriostrado)")
es_desplazable = "Desplazable" in _portico_tipo  # True → Mu ya amplificado por P-Δ global

# =============================================================================
# CÁLCULOS DE CAPACIDAD UNIAXIAL (X y Y)
# =============================================================================
layers_y = []
if es_circular:
    cap_x = compute_uniaxial_capacity_circular(D, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza)
    cap_y = cap_x
else:
    # Eje X: flexión sobre eje X, peralte = h, ancho = b
    cap_x = compute_uniaxial_capacity(b, h, d_prime, layers, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza)
    # Eje Y: flexión sobre eje Y, peralte = b, ancho = h
    # Extremos (caras paralelas a h, separadas en dirección b): num_filas_v barras
    layers_y.append({'d': d_prime, 'As': num_filas_v * rebar_area})
    # Laterales (caras a lo largo de h): num_filas_h barras totales menos las 2 esquinas
    num_capas_intermedias_y = num_filas_h - 2
    if num_capas_intermedias_y > 0:
        espacio_y = (b - 2 * d_prime) / (num_capas_intermedias_y + 1)
        for i in range(1, num_capas_intermedias_y + 1):
            layers_y.append({'d': d_prime + i * espacio_y, 'As': 2 * rebar_area})
    layers_y.append({'d': b - d_prime, 'As': num_filas_v * rebar_area})
    cap_y = compute_uniaxial_capacity(h, b, d_prime, layers_y, fc, fy, Es, phi_c_max, phi_tension, eps_full, p_max_factor, factor_fuerza)

# 2) Siempre conservador en Bresler Biaxial, usando phi_c_max
phi_factor = phi_c_max
bresler = biaxial_bresler(Pu_input, Mux_input, Muy_input, cap_x, cap_y, cap_x['Po'], phi_factor)

# 
# BLOQUE: COMPRESIÓN AXIAL PURA — Verificación paso a paso
# 
with st.expander("Compresión Axial Pura — Verificación Paso a Paso (NSR-10 C.9.3.2.2)", expanded=False):

    # Cálculo desglosado
    Ag_cm2  = Ag              # Ag ya está en cm²
    Ast_cm2 = Ast             # Ast ya está en cm²
    Anc_cm2 = Ag_cm2 - Ast_cm2

    Po_kN       = (0.85 * fc * Anc_cm2 + fy * Ast_cm2) / 10.0
    Pn_max_kN   = p_max_factor * Po_kN
    phi_Pn_max_kN = phi_c_max * Pn_max_kN

    Po_out     = Po_kN        * factor_fuerza
    Pn_max_out = Pn_max_kN    * factor_fuerza
    phi_Pn_out = phi_Pn_max_kN * factor_fuerza

    # Panel visual estilo consola
    st.markdown(f"""
    <div style="background:#1c2e1c;border-radius:10px;padding:14px 18px;
                font-family:monospace;font-size:13px;color:#e8f5e9;line-height:2.1">
    <b style="font-size:15px;color:#81c784"> Resistencia máxima a compresión axial pura:</b><br><br>
    <span style="color:#aaa">Po = [0.85·f'c·(Ag − Ast) + fy·Ast] / 1000</span><br>
    <span style="color:#aaa">Pn,máx = {p_max_factor:.2f} × Po &nbsp;&nbsp;|&nbsp;&nbsp; φPn,máx = φ × Pn,máx</span><br><br>
    <b>Ag</b> (área bruta)         = <b>{Ag_cm2:.2f} cm²</b><br>
    <b>Ast</b> (área acero)        = <b>{Ast_cm2:.2f} cm²</b><br>
    <b>Ag − Ast</b> (concreto neto) = <b>{Anc_cm2:.2f} cm²</b><br>
    <span style="color:#aaa"></span><br>
    <b>Po</b>     = [0.85 × {fc:.1f} × {Anc_cm2:.2f} + {fy:.0f} × {Ast_cm2:.2f}] / 1000
             = <b>{Po_out:.1f} {unidad_fuerza}</b><br>
    <b>Pn,máx</b> = {p_max_factor:.2f} × {Po_out:.1f}
             = <b>{Pn_max_out:.1f} {unidad_fuerza}</b>
    <span style="color:#aaa;font-size:11px">
        {'→ Estribos: 0.80' if p_max_factor == 0.80 else '→ Espiral: 0.85'}  —  {code['ref']} C.9.3.2.2
    </span><br>
    <b style="color:#81c784">φPn,máx</b> = {phi_c_max:.2f} × {Pn_max_out:.1f}
             = <b style="color:#a5d6a7;font-size:16px">{phi_Pn_out:.1f} {unidad_fuerza}</b>
    <span style="color:#aaa;font-size:11px">
        {'→ φ = 0.65 estribos' if phi_c_max == 0.65 else '→ φ = 0.75 espiral'}
    </span><br>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    # Métricas visuales
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Po — Resistencia bruta",
                f"{Po_out:.1f} {unidad_fuerza}",
                help="Capacidad axial sin factores de reducción")
    col2.metric(f"Pn,máx  (×{p_max_factor:.2f})",
                f"{Pn_max_out:.1f} {unidad_fuerza}",
                help=f"Límite por excentricidad accidental — {code['ref']} C.9.3.2.2")
    col3.metric(f"φPn,máx  (φ={phi_c_max:.2f})",
                f"{phi_Pn_out:.1f} {unidad_fuerza}",
                delta="↑ Punto superior del diagrama P-M",
                help="Resistencia de diseño — punto (M=0, P=φPn,máx)")
    col4.metric("Utilización Pu / φPn,máx",
                f"{Pu_input / phi_Pn_max_kN * factor_fuerza / factor_fuerza * 100:.1f} %"
                if phi_Pn_max_kN > 0 else "—",
                help="Porcentaje de uso de la capacidad axial máxima")

    # Nota normativa
    tipo_col = "espiral" if (es_circular or ("Espiral" in col_type if not es_circular else True)) else "estribos"
    st.caption(
        f" **{code['ref']} C.9.3.2.2** — Para columnas con **{tipo_col}**, "
        f"la resistencia axial máxima se limita a **{p_max_factor:.0%}·Po** "
        f"para considerar excentricidades mínimas accidentales. "
        f"φ = **{phi_c_max:.2f}** para compresión con {tipo_col}."
    )

    # Semáforo Pu vs φPn,máx
    Pu_kN_check = Pu_input / factor_fuerza
    if Pu_kN_check > phi_Pn_max_kN:
        st.error(
            f" **Pu = {Pu_input:.1f} {unidad_fuerza}** supera **φPn,máx = {phi_Pn_out:.1f} {unidad_fuerza}**. "
            f"La columna NO puede soportar esta carga. Aumente la sección o el acero."
        )
    elif Pu_kN_check > 0.90 * phi_Pn_max_kN:
        st.warning(
            f" Pu representa el **{Pu_kN_check / phi_Pn_max_kN * 100:.1f}%** de φPn,máx. "
            f"Zona muy próxima al límite de capacidad axial."
        )
    else:
        st.success(
            f" Pu = {Pu_input:.1f} {unidad_fuerza} → **{Pu_kN_check / phi_Pn_max_kN * 100:.1f}%** "
            f"de φPn,máx = {phi_Pn_out:.1f} {unidad_fuerza}. Capacidad axial suficiente."
        )
# 

if es_circular:
    slender_x = check_slenderness(L_col, D, D, k_factor, beta_dns, Pu_input, M1x_input, Mux_input, fc, fy, Es, factor_fuerza, es_circular=True)
    slender_y = check_slenderness(L_col, D, D, k_factor, beta_dns, Pu_input, M1y_input, Muy_input, fc, fy, Es, factor_fuerza, es_circular=True)
else:
    slender_x = check_slenderness(L_col, h, b, k_factor, beta_dns, Pu_input, M1x_input, Mux_input, fc, fy, Es, factor_fuerza)
    slender_y = check_slenderness(L_col, b, h, k_factor, beta_dns, Pu_input, M1y_input, Muy_input, fc, fy, Es, factor_fuerza)

slenderness = slender_x if slender_x['kl_r'] >= slender_y['kl_r'] else slender_y

# Marco Desplazable: Mu ya viene amplificado por P-Δ del modelo global (ETABS/SAP2000).
# El módulo solo clasifica kL/r, NO aplica δns adicional (NSR-10 C.10.10.7 / ACI §6.6.4.6).
# Marco No Desplazable: se aplica δns = Cm / (1 - Pu/0.75Pc) normalmente.
if es_desplazable:
    Mux_magnified = Mux_input   # δns = 1.0 — P-Δ ya incluido en Mu
    Muy_magnified = Muy_input
else:
    Mux_magnified = Mux_input * slender_x['delta_ns']
    Muy_magnified = Muy_input * slender_y['delta_ns']

# =============================================================================
# CÁLCULO DE ESTRIBOS Y VERIFICACIÓN Ash
# =============================================================================
if not es_circular:
    recub_cm = max(d_prime - rebar_diam / 20.0, 2.5)
    # Recubrimiento mínimo dinámico por norma
    _recub_min_nsr = COVER_MIN_COL.get(norma_sel, 3.8)
    _cover_ref_nsr = COVER_REF_COL.get(norma_sel, "C.7.7.1")
    if recub_cm < _recub_min_nsr:
        st.sidebar.warning(f"️ {_cover_ref_nsr}: Recubrimiento calculado ({recub_cm:.1f} cm) < mínimo de {_recub_min_nsr} cm para columnas según {norma_sel}. Verifique d'.")
    bc = b - 2 * recub_cm
    hc = h - 2 * recub_cm
    Ach = bc * hc
    
    claro_libre_x = (b - 2 * d_prime) / (num_filas_h - 1) - rebar_diam / 10 if num_filas_h > 1 else 0
    claro_libre_y = (h - 2 * d_prime) / (num_filas_v - 1) - rebar_diam / 10 if num_filas_v > 1 else 0
    
    # 1 varilla intermedia = 1 grapa cobrada en APU
    num_flejes_x = max(0, num_filas_v - 2)
    num_flejes_y = max(0, num_filas_h - 2)
    
    ramas_x = 2 + num_flejes_y
    ramas_y = 2 + num_flejes_x

    Ash_prov_x = ramas_x * stirrup_area
    Ash_prov_y = ramas_y * stirrup_area
    Ash_prov = min(Ash_prov_x, Ash_prov_y)

    s1 = 16 * rebar_diam / 10
    s2 = 48 * stirrup_diam / 10
    s3 = min(b, h)
    s3_db = 6*rebar_diam/10  # cm NSR-10 C.21.6.4.3
    s_basico = min(s1, s2, s3)
    
    Lo_conf = max(max(b, h), L_col / 6.0, 45.0)
    
    if es_des:
        s_conf = min(8*rebar_diam/10, 24*stirrup_diam/10, min(b,h)/3, 15.0, s3_db)
        s_centro = min(6*rebar_diam/10, s_basico)
        s_centro = max(s_centro, s_conf)
    else:
        # DMO y DMI
        s_conf = min(8*rebar_diam/10, 24*stirrup_diam/10, min(b,h)/3, 15.0, s3_db)
        s_centro = s_basico
    
    fyt = fy
    
    Ash_req_1 = 0.3 * s_conf * bc * fc / fyt * (Ag / Ach - 1)
    Ash_req_2 = 0.09 * s_conf * bc * fc / fyt
    Ash_req = max(Ash_req_1, Ash_req_2)

    ash_ok = Ash_prov >= Ash_req

    # ── AUTO CROSS-TIES NSR-10 C.21.6.4.4 / ACI 318-19 Table 18.7.5.4 ──────
    _n_flejes_auto  = 0
    _Ash_prov_auto  = Ash_prov
    _tipo_flete     = ""
    if not ash_ok and not es_circular and stirrup_area > 0:
        _n_ramas_min    = math.ceil(Ash_req / stirrup_area)
        _flejes_min     = max(0, _n_ramas_min - 2)
        _n_flejes_auto  = max(0, _flejes_min - max(num_flejes_x, num_flejes_y))
        num_flejes_x    = max(num_flejes_x, _flejes_min)
        num_flejes_y    = max(num_flejes_y, _flejes_min)
        ramas_x         = 2 + num_flejes_y
        ramas_y         = 2 + num_flejes_x
        Ash_prov_x      = ramas_x * stirrup_area
        Ash_prov_y      = ramas_y * stirrup_area
        Ash_prov        = min(Ash_prov_x, Ash_prov_y)
        _Ash_prov_auto  = Ash_prov
        ash_ok          = Ash_prov >= Ash_req
        if es_des:
            _tipo_flete = _t(
                "Gancho 135°-135° ambos extremos (NSR-10 C.21.6.4.2a — DES)",
                "135°-135° hooks both ends (NSR-10 C.21.6.4.2a — DES)")
        elif es_dmo:
            _tipo_flete = _t(
                "Gancho 135°-90° alternados en capas contiguas (NSR-10 C.21.5.3.3 — DMO)",
                "Alternating 135°-90° hooks in adjacent layers (NSR-10 C.21.5.3.3 — DMO)")
        else:
            _tipo_flete = _t(
                "Gancho 90°-90° permitido (NSR-10 DMI / ACI 318)",
                "90°-90° hook permitted (NSR-10 DMI / ACI 318)")
    
    n_est_por_Lo = math.ceil(Lo_conf / s_conf)
    n_estribos_zona = n_est_por_Lo * 2
    longitud_zona_libre = max(0, L_col - 2 * Lo_conf)
    n_estribos_centro = max(0, math.ceil(longitud_zona_libre / s_centro) - 1) if longitud_zona_libre > 0 else 0
    n_estribos_total = n_estribos_zona + n_estribos_centro + 1
    
    ld_mm = get_development_length(rebar_diam, fy, fc)
    splice_length_cm = 1.3 * (ld_mm / 10.0)
    splice_zone_height = splice_length_cm
    splice_start = max(L_col / 3, 0)
    splice_end = splice_start + splice_zone_height
    if splice_end > L_col:
        splice_end = L_col
        splice_start = max(0, L_col - splice_zone_height)
    
    # Para usar en despiece
    perim_estribo = 2 * (b - 2 * recub_cm) + 2 * (h - 2 * recub_cm) + 12 * stirrup_diam / 10
    long_fleje_x = h - 2 * recub_cm + 2 * 6 * stirrup_diam / 10
    long_fleje_y = b - 2 * recub_cm + 2 * 6 * stirrup_diam / 10
    perim_estribo += (num_flejes_x * long_fleje_x + num_flejes_y * long_fleje_y)
else:
    recub_cm = max(d_prime - rebar_diam / 20.0, 2.5)
    _recub_min_nsr = COVER_MIN_COL.get(norma_sel, 3.8)
    _cover_ref_nsr = COVER_REF_COL.get(norma_sel, "C.7.7.1")
    if recub_cm < _recub_min_nsr:
        st.sidebar.warning(f"️ {_cover_ref_nsr}: Recubrimiento calculado ({recub_cm:.1f} cm) < mínimo de {_recub_min_nsr} cm según {norma_sel}.")
    dc = D - 2 * recub_cm
    Ach = math.pi * (dc/2)**2
    Ag_circ = math.pi * (D/2)**2
    
    rho_s_req = 0.45 * (Ag_circ / Ach - 1) * fc / fy
    rho_s_req = max(rho_s_req, 0.12 * fc / fy)
    
    area_spiral = stirrup_area
    rho_s_prov = (4 * area_spiral) / (dc * paso_espiral)
    ash_ok = rho_s_prov >= rho_s_req
    n_estribos_total = math.ceil(L_col / paso_espiral) + 1
    
    ld_mm = get_development_length(rebar_diam, fy, fc)
    splice_length_cm = 1.3 * (ld_mm / 10.0)
    splice_zone_height = splice_length_cm
    splice_start = max(L_col / 3, 0)
    splice_end = splice_start + splice_zone_height
    if splice_end > L_col:
        splice_end = L_col
        splice_start = max(0, L_col - splice_zone_height)
    
    # Para usar en despiece circular
    long_espiral_vuelta = math.sqrt((math.pi * dc)**2 + paso_espiral**2)
    long_espiral_total = long_espiral_vuelta * (L_col / paso_espiral)

# =============================================================================
# CANTIDADES DE MATERIALES
# =============================================================================
vol_concreto_m3 = (Ag / 10000) * (L_col / 100) if not es_circular else (math.pi * (D/2)**2 / 10000) * (L_col / 100)
peso_acero_long_kg = Ast * L_col * 7.85e-3      # Ast(cm²)×L(cm)×ρ(g/cm³)/1000→kg

if not es_circular:
    peso_unit_estribo = (perim_estribo / 100.0) * (stirrup_area * 100) * 7.85e-3
    peso_total_estribos_kg = n_estribos_total * peso_unit_estribo
else:
    peso_total_estribos_kg = long_espiral_total / 100 * (stirrup_area * 100) * 7.85e-3

peso_total_acero_kg = peso_acero_long_kg + peso_total_estribos_kg

# =============================================================================
# GRÁFICO 3D BIAXIAL
# =============================================================================

# =============================================================================
#  gen_df_cap_3d — Nube de puntos 3D de la superficie de interacción
#  Método: interpolación elíptica entre cap_x y cap_y  (~0.01 s)
# =============================================================================
def gen_df_cap_3d(cap_x, cap_y, n_theta=24, n_P=80):
    """Superficie P-M 3D por barrido radial θ=0→360°.
    Genera columnas nominales (P,Mx,My) Y de diseño (phi_Pn,phi_Mx,phi_My)
    para que plot_pm_3d use ambas capas desde el mismo DataFrame.
    """
    # [FIX] scipy→np.interp

    def _arr(d, *keys):
        for k in keys:
            v = d.get(k)
            if v is not None and len(v) > 0:
                return np.array(v)
        return np.array([])

    Px   = _arr(cap_x, 'P_n',    'phi_P_n', 'phi_Pn')
    Mx_  = _arr(cap_x, 'M_n',    'phi_M_n', 'phi_Mn')
    Py   = _arr(cap_y, 'P_n',    'phi_P_n', 'phi_Pn')
    My_  = _arr(cap_y, 'M_n',    'phi_M_n', 'phi_Mn')
    phiPx  = _arr(cap_x, 'phi_P_n', 'phi_Pn', 'P_n')
    phiMx_ = _arr(cap_x, 'phi_M_n', 'phi_Mn', 'M_n')
    phiPy  = _arr(cap_y, 'phi_P_n', 'phi_Pn', 'P_n')
    phiMy_ = _arr(cap_y, 'phi_M_n', 'phi_Mn', 'M_n')

    if len(Px) == 0 or len(Py) == 0:
        return pd.DataFrame({'P':[],'Mx':[],'My':[],'phi_Pn':[],'phi_Mx':[],'phi_My':[]})

    P_min = min(Px.min(), Py.min()); P_max = max(Px.max(), Py.max())
    P_grid = np.linspace(P_min, P_max, n_P)

    _sx=np.argsort(Px)
    _sy=np.argsort(Py)
    _spx=np.argsort(phiPx)
    _spy=np.argsort(phiPy)

    Mxc  = np.maximum(np.interp(P_grid,np.array(Px)[_sx],np.array(Mx_)[_sx],left=0.,right=0.),0.)
    Myc  = np.maximum(np.interp(P_grid,np.array(Py)[_sy],np.array(My_)[_sy],left=0.,right=0.),0.)
    pMxc = np.maximum(np.interp(P_grid,np.array(phiPx)[_spx],np.array(phiMx_)[_spx],left=0.,right=0.),0.)
    pMyc = np.maximum(np.interp(P_grid,np.array(phiPy)[_spy],np.array(phiMy_)[_spy],left=0.,right=0.),0.)

    rows = []
    for theta in np.linspace(0, 2*np.pi, n_theta, endpoint=False):
        ct = np.cos(theta);  st = np.sin(theta)
        ca = abs(ct);         sa = abs(st)
        for i, P in enumerate(P_grid):
            d  = ca**2/(Mxc[i] +1e-9)**2 + sa**2/(Myc[i] +1e-9)**2
            Mn = 1.0/np.sqrt(d +1e-18) if d > 0 else 0.0
            dp = ca**2/(pMxc[i]+1e-9)**2 + sa**2/(pMyc[i]+1e-9)**2
            Mp = 1.0/np.sqrt(dp+1e-18) if dp > 0 else 0.0
            if Mn > 1e-3:
                rows.append({'P': float(P), 'Mx': Mn*ct, 'My': Mn*st,
                             'phi_Pn': float(P), 'phi_Mx': Mp*ct, 'phi_My': Mp*st})
    phi_P_max = max(phiPx.max(), phiPy.max()) if len(phiPx) > 0 else P_max
    phi_P_min = min(phiPx.min(), phiPy.min()) if len(phiPx) > 0 else P_min
    rows.append({'P':float(P_max),'Mx':0.,'My':0.,'phi_Pn':float(phi_P_max),'phi_Mx':0.,'phi_My':0.})
    rows.append({'P':float(P_min),'Mx':0.,'My':0.,'phi_Pn':float(phi_P_min),'phi_Mx':0.,'phi_My':0.})
    return pd.DataFrame(rows)


def plot_pm_3d(df_cap_3d, df_combos, factor_fuerza=1.0,
               unidad_fuerza='kN', unidad_mom='kN·m'):
    """Diagrama P-M Biaxial 3D — calidad comercial (estilo ETABS/CSiCol).
    Escala consistente: nominal vs diseñoφ vienen del mismo df_cap_3d.
    """
    import plotly.graph_objects as _go3

    fig = _go3.Figure()
    _L  = dict(ambient=0.5, diffuse=0.8, fresnel=0.2, specular=0.5, roughness=0.1)
    _FF = factor_fuerza

    # ── Superficie Nominal (gris translúcido) ───────────────────────────────────
    fig.add_trace(_go3.Mesh3d(
        x=df_cap_3d['Mx'] * _FF,
        y=df_cap_3d['My'] * _FF,
        z=df_cap_3d['P']  * _FF,
        alphahull=0, opacity=0.12, color='#a0aab5',
        flatshading=False, lighting=_L,
        name='Nominal (Pn, Mn)', hoverinfo='skip', showlegend=True,
    ))

    # ── Superficie Diseño φ (usa phi_Mx/phi_My/phi_Pn del df) ───────────────────
    _hp = 'phi_Mx' in df_cap_3d.columns and 'phi_My' in df_cap_3d.columns
    fig.add_trace(_go3.Mesh3d(
        x=df_cap_3d['phi_Mx' if _hp else 'Mx'] * _FF,
        y=df_cap_3d['phi_My' if _hp else 'My'] * _FF,
        z=df_cap_3d['phi_Pn' if _hp else 'P']  * _FF,
        alphahull=0, opacity=0.55, color='#2ecc71',
        flatshading=False, lighting=_L,
        name='φPn, φMn — Diseño',
        hovertemplate=(
            '<b>φPn:</b> %{z:.1f} ' + unidad_fuerza + '<br>'
            '<b>φMnx:</b> %{x:.1f} ' + unidad_mom + '<br>'
            '<b>φMny:</b> %{y:.1f} ' + unidad_mom + '<extra></extra>'
        ),
        contour=dict(show=True, color='rgba(255,255,255,0.55)', width=2),
        showlegend=True,
    ))

    # ── Combinaciones de carga — MISMA escala nominal ─────────────────────────
    if df_combos is not None and len(df_combos) > 0:
        _cp  = next((c for c in ['Pu','P','pu']    if c in df_combos.columns), None)
        _cmx = next((c for c in ['Mux','Mx','mux'] if c in df_combos.columns), None)
        _cmy = next((c for c in ['Muy','My','muy'] if c in df_combos.columns), None)
        _clb = next((c for c in ['Combo','combo','Combinacion'] if c in df_combos.columns), None)
        if _cp and _cmx and _cmy:
            _P  = df_combos[_cp]  * _FF
            _Mx = df_combos[_cmx] * _FF
            _My = df_combos[_cmy] * _FF
            _lb = df_combos[_clb].tolist() if _clb else ['']*len(_P)
            fig.add_trace(_go3.Scatter3d(
                x=_Mx, y=_My, z=_P,
                mode='markers+text',
                marker=dict(size=7, color='#e74c3c', symbol='circle',
                            line=dict(color='white', width=1.5)),
                text=_lb, textposition='top center',
                textfont=dict(size=10, color='white'),
                name='Cargas (Pu, Mu)',
                hovertemplate='<b>%{text}</b><br>Pu: %{z:.1f}<br>Mx: %{x:.1f}<br>My: %{y:.1f}<extra></extra>',
            ))
        _pmax = float(df_cap_3d['P'].max()) * _FF * 1.1
        _pmin = float(df_cap_3d['P'].min()) * _FF * 1.1
        fig.add_trace(_go3.Scatter3d(
            x=[0,0], y=[0,0], z=[_pmin, _pmax],
            mode='lines',
            line=dict(color='rgba(255,255,255,0.45)', width=2, dash='dash'),
            name='Eje Neutro (0,0)', hoverinfo='skip',
        ))

    # ── Escenario CAD limpio ─────────────────────────────────────────────────
    _ax = dict(showbackground=False,
               gridcolor='rgba(255,255,255,0.1)',
               zerolinecolor='rgba(255,255,255,0.55)', zerolinewidth=2,
               title_font=dict(size=12, color='#bdc3c7'),
               tickfont=dict(size=10, color='#7f8c8d'))
    fig.update_layout(
        scene=dict(
            xaxis_title=f'Mx [{unidad_mom}]',
            yaxis_title=f'My [{unidad_mom}]',
            zaxis_title=f'P [{unidad_fuerza}]',
            xaxis=_ax, yaxis=_ax, zaxis=_ax,
            aspectmode='data',
            camera=dict(eye=dict(x=1.5, y=1.5, z=0.8)),
            bgcolor='rgba(0,0,0,0)',
        ),
        paper_bgcolor='#1e1e2e', plot_bgcolor='#1e1e2e',
        margin=dict(l=0, r=0, b=0, t=30),
        legend=dict(orientation='h', yanchor='bottom', y=1.02,
                    xanchor='right', x=1, font=dict(color='white')),
        hovermode='closest',
    )
    return fig


def create_pm_2d_plot(cap_x, Pu, Mu, unidad_mom, unidad_fuerza):
    fig, ax = plt.subplots(figsize=(8, 6))
    fig.patch.set_facecolor('#1e1e2e')
    for _ax in fig.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.plot(cap_x['M_n'], cap_x['P_n'], 'b--', linewidth=1.5, label=r"Resistencia Nominal ($P_n, M_n$)")
    ax.plot(cap_x['phi_M_n'], cap_x['phi_P_n'], 'r-', linewidth=2.5, label=r"Resistencia de Diseño ($\phi P_n, \phi M_n$)")
    
    P_bal = cap_x.get('P_balance', None)
    P_ten = cap_x.get('P_tension', None)
    
    # Rellenar zonas de compresión, transición y tracción
    if P_bal is not None:
        ax.plot(cap_x['M_balance'], cap_x['P_balance'], 'ro', markersize=6, markeredgecolor='black', label=f"Falla Balanceada ($ε_t = ε_y$)")
        ax.axhline(P_bal, color='orange', linestyle='--', alpha=0.5)
    if P_ten is not None:
        ax.plot(cap_x['M_tension'], cap_x['P_tension'], 'go', markersize=6, markeredgecolor='black', label=f"Control a Tracción ($ε_t = 0.005$)")
        ax.axhline(P_ten, color='green', linestyle='--', alpha=0.5)
    
    ymax = max(cap_x['P_n']) * 1.05
    ymin = min(cap_x['P_n']) * 1.05
    xmax = max(cap_x['M_n']) * 1.15
    if P_bal is not None and P_ten is not None:
        ax.fill_between([0, xmax], P_bal, ymax, color='orange', alpha=0.07, label='Interacción Compresión')
        ax.fill_between([0, xmax], P_ten, P_bal, color='yellow', alpha=0.1, label='Transición')
        ax.fill_between([0, xmax], ymin, P_ten, color='green', alpha=0.07, label='Interacción Tracción')

    ax.plot(Mu, Pu, 'o', markersize=9, color='lime', markeredgecolor='black', zorder=6,
            label=f"Punto de Diseño (Mu={Mu:.1f}, Pu={Pu:.1f})")
    ax.axhline(y=0, color='gray', linestyle=':', alpha=0.5)
    ax.axvline(x=0, color='gray', linestyle=':', alpha=0.5)
    ax.set_xlabel(f"Momento Flector M [{unidad_mom}]")
    ax.set_ylabel(f"Carga Axial P [{unidad_fuerza}]")
    ax.set_title(f"Diagrama de Interacción P-M — {norma_sel}\n" +
                 f"Columna {b:.0f}×{h:.0f} cm | f'c={fc:.1f} MPa | fy={fy:.0f} MPa")
    ax.grid(True, linestyle=":", alpha=0.5)
    ax.legend(loc="upper right", fontsize=8)
    return fig

fig_pm_2d = create_pm_2d_plot(cap_x, Pu_input, Mux_input, unidad_mom, unidad_fuerza)
pm_2d_img = io.BytesIO()
fig_pm_2d.savefig(pm_2d_img, facecolor='white', edgecolor='none', format='png', dpi=150)
pm_2d_img.seek(0)
# Generar superficie 3D — elipse entre cap_x y cap_y (36 ángulos × 80 niveles ≈ 0.01 s)
df_cap_3d = gen_df_cap_3d(cap_x, cap_y, n_theta=36, n_P=80)
# Figura tab1: punto actual (Pu, Mux, Muy) como única carga
_df_single = pd.DataFrame({
    'Combinacion': [_t('Punto actual', 'Current point')],
    'Pu':  [Pu_input  / factor_fuerza],
    'Mux': [Mux_input / factor_fuerza],
    'Muy': [Muy_input / factor_fuerza],
})
fig_3d = plot_pm_3d(df_cap_3d, _df_single, factor_fuerza,
                    unidad_fuerza=unidad_fuerza, unidad_mom=unidad_mom)

try:
    pm_3d_img = io.BytesIO()
    fig_3d.write_image(pm_3d_img, format="png")
    pm_3d_img.seek(0)
    has_3d_img = True
except Exception:
    has_3d_img = False

# =============================================================================
# TABS PRINCIPALES
# =============================================================================

# =============================================================================
# PASO 4 — def render_config_tab()
#           Todos los widgets de calculo van aqui. Se llama SOLO desde tab0.
# =============================================================================
def render_config_tab():
    st.markdown("### Parametros de Diseno")
    col_izq, col_der = st.columns(2)

    with col_izq:
        st.subheader("1. Materiales")
        _fc_unit_w = st.radio("Unidad de f'c:", ["MPa", "PSI", "kg/cm2"],
                              horizontal=True, key="c_pm_fc_unit")
        if _fc_unit_w == "PSI":
            _psi_opts = {
                "2500 PSI (17.2 MPa)": 2500.0, "3000 PSI (20.7 MPa)": 3000.0,
                "3500 PSI (24.1 MPa)": 3500.0, "4000 PSI (27.6 MPa)": 4000.0,
                "4500 PSI (31.0 MPa)": 4500.0, "5000 PSI (34.5 MPa)": 5000.0,
                "Personalizado": None
            }
            _psi_ch = st.selectbox("Resistencia f'c [PSI]", list(_psi_opts.keys()), key="c_pm_psi_choice")
            if _psi_opts[_psi_ch] is not None:
                st.session_state["c_pm_fc_psi_val"] = _psi_opts[_psi_ch]
                st.caption(f"fc = {_psi_opts[_psi_ch] * 0.00689476:.2f} MPa")
            else:
                _pv = st.number_input("f'c personalizado [PSI]", 2000.0, 12000.0, 3000.0, 100.0, key="c_pm_fc_psi_custom")
                st.session_state["c_pm_fc_psi_val"] = _pv
        elif _fc_unit_w == "kg/cm2":
            _kg_opts = {
                "175 kg/cm2 (17.2 MPa)": 175.0, "210 kg/cm2 (20.6 MPa)": 210.0,
                "250 kg/cm2 (24.5 MPa)": 250.0, "280 kg/cm2 (27.5 MPa)": 280.0,
                "350 kg/cm2 (34.3 MPa)": 350.0, "420 kg/cm2 (41.2 MPa)": 420.0,
                "Personalizado": None
            }
            _kg_ch = st.selectbox("Resistencia f'c [kg/cm2]", list(_kg_opts.keys()), key="c_pm_kgcm2_choice")
            if _kg_opts[_kg_ch] is not None:
                st.session_state["c_pm_fc_kgcm2_val"] = _kg_opts[_kg_ch]
                st.caption(f"fc = {_kg_opts[_kg_ch] / 10.1972:.2f} MPa")
            else:
                _kv = st.number_input("f'c personalizado [kg/cm2]", 100.0, 1200.0, 210.0, 10.0, key="c_pm_fc_kgcm2_custom")
                st.session_state["c_pm_fc_kgcm2_val"] = _kv
        else:
            st.number_input("Resistencia del Concreto f'c [MPa]", 15.0, 80.0, 21.0, 1.0, key="c_pm_fc_mpa")
        st.number_input("Fluencia del Acero fy [MPa]", 240.0, 500.0, 420.0, 10.0, key="c_pm_fy")

        st.markdown("---")
        st.subheader("2. Geometria de la Seccion")
        st.selectbox("Tipo de seccion",
                     ["Rectangular / Cuadrada", "Circular (con espiral)"],
                     key="c_pm_seccion_type")
        _es_circ_cfg = "Circular" in st.session_state.get("c_pm_seccion_type", "Rectangular / Cuadrada")
        if _es_circ_cfg:
            st.number_input("Diametro D [cm]", 15.0, 150.0, 40.0, 5.0, key="c_pm_D")
            st.caption("Para pilares circulares de gran diametro (puentes, silos) use el modulo Pilares Circulares.")
        else:
            _ca, _cb = st.columns(2)
            _ca.number_input("Base b [cm]",   15.0, 150.0, 30.0, 5.0, key="c_pm_b")
            _cb.number_input("Altura h [cm]", 15.0, 150.0, 40.0, 5.0, key="c_pm_h")
        st.number_input("Recubrimiento al centroide d' [cm]", 2.0, 15.0, 5.0, 0.5, key="c_pm_dprime")
        st.number_input("Altura libre de la columna L [cm]", 50.0, 1000.0, 300.0, 25.0, key="c_pm_L")

        st.markdown("---")
        st.subheader("6. Esbeltez")
        st.radio(
            _t("Tipo de Pórtico", "Frame Type"),
            options=["No desplazable (Pórtico arriostrado)", "Desplazable (Pórtico sin arriostre)"],
            help=_t(
                "No desplazable: el módulo calcula δns y amplifica Mu. "
                "Desplazable: Mu ya viene amplificado por P-Δ desde ETABS/SAP2000; "
                "el módulo solo clasifica la esbeltez y documenta el tipo de pórtico.",
                "Non-sway: module computes δns and amplifies Mu. "
                "Sway: Mu already includes P-Δ amplification from global analysis; "
                "module only classifies slenderness and documents frame type."
            ),
            horizontal=False,
            key="c_pm_portico_tipo"
        )
        st.number_input(
            _t("Factor de carga sostenida (β_dns)", "Sustained load ratio (β_dns)"),
            0.0, 1.0, 0.6, 0.1,
            help=_t("Relación M_sostenido / M_total. Default 0.6 (solo aplica a pórticos No Desplazables).",
                    "Ratio M_sustained / M_total. Default 0.6 (applies to Non-Sway frames only)."),
            key="c_pm_beta_dns"
        )
        st.selectbox(
            _t("Factor de longitud efectiva k", "Effective length factor k"),
            ["Ambos extremos articulados",
             "Un extremo articulado, otro empotrado",
             "Ambos extremos empotrados",
             "Voladizo (base empotrada, libre arriba)"],
            key="c_pm_k_sel"
        )

    with col_der:
        st.subheader("3. Refuerzo Longitudinal")
        st.radio("Sistema de unidades de varillas:",
                 ["Milimetros (SI)", "Pulgadas (EE. UU.)"],
                 horizontal=True, key="c_pm_unit_system")
        _rbd_cfg  = REBAR_US if ("Pulgadas" in st.session_state.get("c_pm_unit_system","") or
                                  "Inches"  in st.session_state.get("c_pm_unit_system","")) else REBAR_MM
        _rb_keys  = list(_rbd_cfg.keys())
        _rb_def   = st.session_state.get("c_pm_rebar_type", _rb_keys[2] if len(_rb_keys) > 2 else _rb_keys[0])
        if _rb_def not in _rb_keys: _rb_def = _rb_keys[0]
        st.selectbox("Diametro de las varillas", _rb_keys,
                     index=_rb_keys.index(_rb_def), key="c_pm_rebar_type")

        _es_circ_rb = "Circular" in st.session_state.get("c_pm_seccion_type","")
        if _es_circ_rb:
            st.number_input("Numero de varillas longitudinales", 4, 20, 8, 2, key="c_pm_n_barras_circ")
        else:
            _cr1, _cr2 = st.columns(2)
            _cr1.number_input("Filas horizontales (sup e inf)", 2, 15, 2, 1, key="c_pm_num_h")
            _cr2.number_input("Filas verticales (laterales)",   2, 15, 2, 1, key="c_pm_num_v")

        st.markdown("---")
        st.subheader("4. Refuerzo Transversal")
        _strd_cfg = STIRRUP_US if ("Pulgadas" in st.session_state.get("c_pm_unit_system","") or
                                    "Inches"  in st.session_state.get("c_pm_unit_system","")) else STIRRUP_MM
        if _es_circ_rb:
            _sp_keys = list(_strd_cfg.keys())
            _sp_def  = st.session_state.get("c_pm_spiral_type", _sp_keys[0])
            if _sp_def not in _sp_keys: _sp_def = _sp_keys[0]
            st.selectbox("Diametro de la espiral", _sp_keys,
                         index=_sp_keys.index(_sp_def), key="c_pm_spiral_type")
            st.number_input("Paso de la espiral s [cm]", 2.0, 20.0, 7.5, 0.5, key="c_pm_paso_circ")
        else:
            st.selectbox("Tipo de columna",
                         ["Estribos (Tied)", "Espiral (Spiral)"], key="c_pm_col_type")
            _st_keys = list(_strd_cfg.keys())
            _st_def  = st.session_state.get("c_pm_stirrup_type", _st_keys[0])
            if _st_def not in _st_keys: _st_def = _st_keys[0]
            st.selectbox("Diametro del estribo", _st_keys,
                         index=_st_keys.index(_st_def), key="c_pm_stirrup_type")

        st.markdown("---")
        st.subheader("5. Solicitaciones Biaxiales")
        st.radio("Unidades del diagrama:",
                 ["KiloNewtons (kN, kN-m)", "Toneladas Fuerza (tonf, tonf-m)"],
                 horizontal=True, key="c_pm_output_units")
        _ff_w = 0.1019716 if "Toneladas" in st.session_state.get("c_pm_output_units","") else 1.0
        _uf_w = "tonf"   if "Toneladas" in st.session_state.get("c_pm_output_units","") else "kN"
        _um_w = "tonf-m" if "Toneladas" in st.session_state.get("c_pm_output_units","") else "kN-m"
        _cs1, _cs2, _cs3 = st.columns(3)
        _cs1.number_input(f"Pu [{_uf_w}]",  value=round(2700.0*_ff_w,2), step=round(50.0*_ff_w,2),  key="c_pm_pu")
        _cs2.number_input(f"Mux [{_um_w}]", value=round(45.0*_ff_w,2),   step=round(10.0*_ff_w,2),  key="c_pm_mux")
        _cs3.number_input(f"Muy [{_um_w}]", value=round(25.0*_ff_w,2),   step=round(10.0*_ff_w,2),  key="c_pm_muy")
        st.checkbox("Especificar Curvatura Simple M1/M2", value=False,
                    help="Permite calcular Cm diferente a 1.0", key="c_pm_cm_chk")
        if st.session_state.get("c_pm_cm_chk"):
            _cm1, _cm2 = st.columns(2)
            _cm1.number_input(f"Momento menor M1x [{_um_w}]", value=0.0, step=round(10.0*_ff_w,2), key="c_pm_m1x")
            _cm2.number_input(f"Momento menor M1y [{_um_w}]", value=0.0, step=round(10.0*_ff_w,2), key="c_pm_m1y")

        st.markdown("---")
        st.subheader("Resumen de la Seccion")
        st.info(
            f"f'c = {st.session_state.get('c_pm_fc_mpa', 21.0)} MPa  |  "
            f"fy = {st.session_state.get('c_pm_fy', 420.0)} MPa  |  "
            f"Norma: {norma_sel}  |  Nivel: {nivel_sismico}  |  "
            f"rho min: {rho_min}%  |  rho max: {rho_max}%"
        )


# =============================================================================
# PASO 5 — HEADER BANNER + METRICAS SUPERIORES
# =============================================================================

_emp_hdr  = st.session_state.get("cpm_empresa","") or "StructoPro"
_proy_hdr = st.session_state.get("cpm_proyecto_nombre","") or "Diseno de Columnas P-M"

st.markdown(
    f"""<div style="background:linear-gradient(135deg,#0d1b2a 0%,#1e3a5f 60%,#2e6da4 100%);
    padding:20px 32px;border-radius:12px;margin-bottom:14px">
    <div style="display:flex;align-items:center;gap:16px">
    <div style="background:rgba(255,255,255,0.1);border-radius:8px;padding:8px">
    <span style="font-size:30px;line-height:1">&#9632;</span></div>
    <div>
    <h1 style="color:white;margin:0;font-size:1.8rem;font-weight:800">{_emp_hdr} &mdash; Columnas P-M</h1>
    <p style="color:#90caf9;margin:4px 0 0;font-size:0.9rem">
    Proyecto: {_proy_hdr} &nbsp;|&nbsp; {norma_sel} &nbsp;|&nbsp; {nivel_sismico}
    </p></div></div></div>""",
    unsafe_allow_html=True
)

if abs(Mux_input) < _Mux_emin or abs(Muy_input) < _Muy_emin:
    st.info(
        f"Excentricidad minima NSR-10 C.10.3.6: "
        f"ex={e_min_x_cm:.2f} cm -> Mux >= {_Mux_emin:.2f} {unidad_mom}  |  "
        f"ey={e_min_y_cm:.2f} cm -> Muy >= {_Muy_emin:.2f} {unidad_mom}"
    )

_ratio_ok = bresler["ok"]
_cur_ok   = (rho_min <= cuantia <= rho_max)
_slend_ok = slenderness["kl_r"] <= 100

_m1, _m2, _m3, _m4, _m5 = st.columns(5)
_m1.metric("Pu", f"{Pu_input:.1f} {unidad_fuerza}")
_m2.metric("phi·Pni (Bresler)", f"{bresler['phi_Pni']:.1f} {unidad_fuerza}",
           delta="CUMPLE" if _ratio_ok else "NO CUMPLE",
           delta_color="normal" if _ratio_ok else "inverse")
_m3.metric("Ratio Pu / phi·Pni", f"{bresler['ratio']:.3f}",
           delta_color="normal" if _ratio_ok else "inverse")
_m4.metric("Cuantia rho", f"{cuantia:.2f}%",
           delta="OK" if _cur_ok else "FUERA DE LIMITES",
           delta_color="normal" if _cur_ok else "inverse")
_m5.metric("Esbeltez kL/r", f"{slenderness['kl_r']:.1f}",
           delta="Corta" if slenderness["kl_r"] <= 22 else (
               "Esbelta" if slenderness["kl_r"] <= 100 else "Muy esbelta"),
           delta_color="normal" if _slend_ok else "inverse")

st.markdown("---")

# =============================================================================
# FUNCIÓN AUXILIAR — VERIFICACIÓN NODO VIGA-COLUMNA NSR-10 C.21.7
# =============================================================================
def verificar_nodo_viga_columna(b_col_cm, h_col_cm, fc_mpa, fy_mpa,
                                 As_vigas_cm2, Vu_col_kN, confinamiento="3_caras"):
    """
    Verificación de la resistencia al cortante del nodo — NSR-10 C.21.7.4.1

    Parámetros:
      b_col_cm, h_col_cm : dimensiones columna [cm]
      fc_mpa             : resistencia concreto [MPa]
      fy_mpa             : fluencia acero [MPa]
      As_vigas_cm2       : Acero de tracción de vigas que llegan al nodo [cm²]
      Vu_col_kN          : Cortante de diseño en la columna [kN]
      confinamiento      : '4_caras' | '3_caras' | 'otros'

    Retorna dict con Vn, Vu_j, ratio, ok, detalles
    """
    # Coeficiente γ según confinamiento (NSR-10 C.21.7.4.1)
    gamma_map = {"4_caras": 1.7, "3_caras": 1.25, "otros": 1.0}
    gamma = gamma_map.get(confinamiento, 1.0)

    # Área efectiva del nodo Aj [mm²] → [cm²]
    Aj_cm2 = b_col_cm * h_col_cm
    Aj_mm2 = Aj_cm2 * 100.0

    # Resistencia nominal: Vn = γ·√f'c·Aj  [kN]
    Vn_kN  = gamma * math.sqrt(fc_mpa) * Aj_mm2 / 1000.0  # N → kN
    phi_Vn = 0.85 * Vn_kN  # φ = 0.85 para cortante en nodo (NSR-10 C.9.3.2.3)

    # Cortante último en el nodo (aprox. por cuerpo libre ACI 318R-19 Fig. RC.21.7.4)
    # Vu_j = As_vigas·fy - Vcol  (resultante de tracciones menos corte de columna)
    Vu_j_kN = (As_vigas_cm2 * fy_mpa / 100.0) - Vu_col_kN  # kN
    Vu_j_kN = abs(Vu_j_kN)  # valor absoluto (sentido de cortante)

    ratio  = Vu_j_kN / phi_Vn if phi_Vn > 0 else float('inf')
    ok     = Vu_j_kN <= phi_Vn

    return {
        "gamma":     gamma,
        "Aj_cm2":    Aj_cm2,
        "Vn_kN":     Vn_kN,
        "phi_Vn_kN": phi_Vn,
        "Vu_j_kN":   Vu_j_kN,
        "ratio":     ratio,
        "ok":        ok,
        "confinamiento": confinamiento,
    }
# =============================================================================
# FUNCIÓN AUXILIAR — AUTO-SIZING DE SECCIÓN NSR-10 / ACI 318-19
# =============================================================================
def auto_size_column(Pu_kN, Mux_kNm, Muy_kNm, fc_mpa, fy_mpa,
                     nivel_sismico="DES", rho_objetivo=0.02):
    """
    Sugiere la sección mínima de columna rectangular que satisface Pu, Mux, Muy.

    Algoritmo:
      1. Para cada b×h en la grilla de secciones típicas colombianas
      2. Estima Ag mínimo por carga axial: Ag ≥ Pu/(0.40·f'c) (regla Nilson)
      3. Calcula cuantía necesaria con ρ_objetivo
      4. Verifica diagrama P-M uniaxial simplificado
      5. Retorna lista de secciones factibles ordenadas por área
    """
    import math
    factor_phi = 0.65  # φ compresión controlada

    # Grilla de secciones típicas (cm) — múltiplos de 5cm
    secciones = []
    for b_cm in range(25, 85, 5):
        for h_cm in range(b_cm, 105, 5):  # h ≥ b siempre
            Ag_cm2 = b_cm * h_cm

            # Filtro rápido: capacidad axial mínima
            Po_est = 0.85 * fc_mpa * Ag_cm2 / 100.0  # kN (sin acero, aprox)
            phi_Pmax = factor_phi * 0.80 * Po_est
            if phi_Pmax < Pu_kN * 0.5:
                continue  # Sección demasiado pequeña — saltar

            # Cuantía con ρ_objetivo
            Ast_cm2 = rho_objetivo * Ag_cm2
            Po = (0.85 * fc_mpa * (Ag_cm2 - Ast_cm2) / 100.0 
                  + fy_mpa * Ast_cm2 / 100.0)  # kN
            phi_Pn_max = factor_phi * 0.80 * Po

            if phi_Pn_max < Pu_kN:
                continue

            # Verificación momento simplificada (línea recta approx)
            # φMn ≈ φ·As·fy·(h - 2·d')·0.5/1000 + φ·Cc·jd
            d_prime = max(5.0, b_cm * 0.10)
            jd = 0.85 * (h_cm - 2 * d_prime)  # brazo de palanca aprox
            phi_Mn_x = factor_phi * (Ast_cm2/2 * fy_mpa / 100.0) * jd / 100.0  # kN·m
            jd_y = 0.85 * (b_cm - 2 * d_prime)
            phi_Mn_y = factor_phi * (Ast_cm2/2 * fy_mpa / 100.0) * jd_y / 100.0  # kN·m

            # Bresler simplificado
            if phi_Mn_x <= 0 or phi_Mn_y <= 0:
                continue
            ratio_x = Mux_kNm / phi_Mn_x if phi_Mn_x > 0 else 999
            ratio_y = Muy_kNm / phi_Mn_y if phi_Mn_y > 0 else 999

            # Criterio de aceptación conservador
            if ratio_x <= 1.0 and ratio_y <= 1.0 and phi_Pn_max >= Pu_kN:
                esbeltez_ok = True  # filtro básico
                secciones.append({
                    "b": b_cm, "h": h_cm,
                    "Ag": Ag_cm2,
                    "rho": rho_objetivo,
                    "phi_Pn_max": phi_Pn_max,
                    "phi_Mn_x": phi_Mn_x,
                    "phi_Mn_y": phi_Mn_y,
                    "ratio_x": ratio_x,
                    "ratio_y": ratio_y,
                    "Ast_cm2": Ast_cm2,
                })

    # Ordenar por área (mínima primero)
    secciones.sort(key=lambda s: s["Ag"])
    return secciones[:5]  # top 5 más compactas
# =============================================================================
# PASO 6 — TABS PRINCIPALES
# =============================================================================
tab0, tab1, tab2, tab2b, tab3, tab4 = st.tabs([
    "Configuracion",
    "Diagrama P-M Biaxial",
    "Seccion y Estribos 3D",
    "Alzado y Combos de Carga",
    "Cantidades y APU",
    "Memoria de Calculo",
])

with tab0:
    if not modo_revision:
        render_config_tab()
    else:
        # ── Modo Revisión: mostrar resumen de parámetros en lugar de formularios ──
        st.info("**Modo Revisión activo** — Los paneles de entrada están ocultos. "
                "Desactiva el toggle en el sidebar para editar parámetros.")
        _rev_cols = st.columns(3)
        with _rev_cols[0]:
            st.metric("f'c", f"{fc:.1f} MPa")
            st.metric("fy", f"{fy:.0f} MPa")
            st.metric("Sección", f"{'Ø'+str(D) if es_circular else str(b)+'×'+str(h)} cm")
        with _rev_cols[1]:
            st.metric("d'", f"{d_prime:.1f} cm")
            st.metric("Barras", f"{n_barras if es_circular else n_barras_total} Ø{rebar_diam:.1f} mm")
            st.metric("Estribo", f"Ø{stirrup_diam:.1f} mm @ {sconf:.1f} cm")
        with _rev_cols[2]:
            st.metric("Pu", f"{Pu_input/factor_fuerza:.1f} {unidad_fuerza}")
            st.metric("Mux", f"{Mux_input/factor_mom:.1f} {unidad_mom}")
            st.metric("Muy", f"{Muy_input/factor_mom:.1f} {unidad_mom}")
        st.markdown("---")
    # ── AUTO-SIZING ───────────────────────────────────────────────────────────
    st.markdown("---")
    with st.expander(_t("🔍 Auto-Sizing — Sección Mínima Sugerida",
                         "🔍 Auto-Sizing — Minimum Section Suggestion"), expanded=False):
        st.caption(_t(
            "Dado Pu, Mux, Muy actuales, calcula las 5 secciones más compactas que cumplen "
            "la capacidad P-M con cuantía objetivo (verificación preliminar, confirme con diagrama completo).",
            "Given current Pu, Mux, Muy, finds the 5 most compact sections meeting P-M capacity "
            "at target reinforcement ratio (preliminary check — confirm with full P-M diagram)."
        ))
        _rho_obj = st.slider(_t("Cuantía objetivo ρ [%]","Target ratio ρ [%]"),
                              1.0, 4.0, 2.0, 0.5, key="c_pm_rho_autosize") / 100.0
        if st.button(_t("▶ Calcular Secciones","▶ Calculate Sections"), key="c_pm_run_autosize"):
            _as_results = auto_size_column(
                Pu_input / factor_fuerza,
                Mux_input / factor_mom,
                Muy_input / factor_mom,
                fc, fy,
                nivel_sismico=nivel_sismico,
                rho_objetivo=_rho_obj
            )
            if _as_results:
                _as_header = [_t("b [cm]","b [cm]"), _t("h [cm]","h [cm]"),
                               _t("Ag [cm²]","Ag [cm²]"), _t("Ast [cm²]","Ast [cm²]"),
                               _t("φPn,max [kN]","φPn,max [kN]"),
                               _t("φMnx [kN·m]","φMnx [kN·m]"), _t("Ratio x","Ratio x"),
                               _t("Ratio y","Ratio y")]
                _as_rows = []
                for _s in _as_results:
                    _as_rows.append({
                        _as_header[0]: f"{_s['b']:.0f}",
                        _as_header[1]: f"{_s['h']:.0f}",
                        _as_header[2]: f"{_s['Ag']:.0f}",
                        _as_header[3]: f"{_s['Ast_cm2']:.2f}",
                        _as_header[4]: f"{_s['phi_Pn_max']:.1f}",
                        _as_header[5]: f"{_s['phi_Mn_x']:.1f}",
                        _as_header[6]: f"{_s['ratio_x']:.2f}",
                        _as_header[7]: f"{_s['ratio_y']:.2f}",
                    })
                import pandas as _pd_as
                _df_as = _pd_as.DataFrame(_as_rows)
                st.dataframe(_df_as, use_container_width=True)
                _best = _as_results[0]
                st.success(
                    _t(f"✅ Sección mínima sugerida: **{_best['b']:.0f}×{_best['h']:.0f} cm** "
                       f"(Ag={_best['Ag']:.0f} cm²) con ρ={_rho_obj*100:.1f}%",
                       f"✅ Minimum suggested section: **{_best['b']:.0f}×{_best['h']:.0f} cm** "
                       f"(Ag={_best['Ag']:.0f} cm²) with ρ={_rho_obj*100:.1f}%"))
                st.caption(_t(
                    "⚠️ Verificación preliminar por carga axial y momento uniaxial simplificado. "
                    "Configure la sección seleccionada en el panel Configuración y valide con el "
                    "Diagrama P-M completo.",
                    "⚠️ Preliminary check based on axial load and simplified uniaxial moment. "
                    "Set the selected section in Configuration and validate with the full P-M diagram."))
            else:
                st.warning(_t(
                    "No se encontraron secciones en la grilla 25–80 cm × 25–100 cm con ρ objetivo. "
                    "Aumente la cuantía o revise las cargas de diseño.",
                    "No sections found in 25–80 cm × 25–100 cm grid at target ρ. "
                    "Increase ratio or check design loads."))
    # ── VISTA EN VIVO ─────────────────────────────────────────────────────────
    st.markdown("---")
    st.subheader(_t(" Vista en Vivo — Superficie Biaxial 3D + Verificación",
                    " Live View — 3D Biaxial Surface + Verification"))
    _ca_live, _cb_live = st.columns([1, 2])
    with _ca_live:
        st.metric(
            label=_t("Ratio Pu / φPni", "Ratio Pu / φPni"),
            value=f"{bresler['ratio']:.3f}",
            delta=_t("CUMPLE", "PASS") if bresler['ok'] else _t("NO CUMPLE", "FAIL"),
            delta_color="normal" if bresler['ok'] else "inverse",
        )
        st.metric(_t("φPni (Bresler)", "φPni (Bresler)"),
                  f"{bresler['phi_Pni']:.1f} {unidad_fuerza}")
        st.metric(_t("Pu aplicado", "Applied Pu"),
                  f"{Pu_input:.1f} {unidad_fuerza}")
        st.markdown("---")
        if not es_circular:
            if ash_ok:
                st.success(f"OK Ash\n{Ash_prov:.2f} ≥ {Ash_req:.2f} cm²")
            else:
                st.error(f" Ash DÉFICIT\n{Ash_prov:.2f} < {Ash_req:.2f} cm²")
                st.info(_t(
                    f"Auto-solución: +{_n_flejes_auto} flejes/cara añadidos\nAsh_prov = {_Ash_prov_auto:.2f} cm²",
                    f"Auto-fix: +{_n_flejes_auto} ties/face added\nAsh_prov = {_Ash_prov_auto:.2f} cm²",
                ))
        st.metric(_t("Esbeltez kL/r", "Slenderness kL/r"),
                  f"{slenderness['kl_r']:.1f}",
                  delta=slenderness['classification'],
                  delta_color="normal" if slenderness['kl_r'] <= 22 else "inverse")
    with _cb_live:
        st.plotly_chart(fig_3d, use_container_width=True, key="fig_3d_live_tab0")

with tab1:
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader(_t("Diagrama P-M 2D (Eje X)", " P-M 2D Diagram (X-Axis)"))
        configurar_pdf_comercial(fig_pm_2d)
        st.pyplot(fig_pm_2d)
        plt.close(fig_pm_2d)

        # ── Exportar CSV del diagrama P-M ──────────────────────────────────────
        try:
            import pandas as _pd_csv_pm
            _n_pts = min(len(cap_x.get("Pn", [])), len(cap_y.get("Pn", [])),
                         len(cap_x.get("phiMn", [])), len(cap_y.get("phiMn", [])))
            if _n_pts > 0:
                _df_pm_csv = _pd_csv_pm.DataFrame({
                    f"Pn [{unidad_fuerza}]":          [round(v, 3) for v in cap_x["Pn"][:_n_pts]],
                    f"phiMn_x [{unidad_mom}]":        [round(v, 3) for v in cap_x["phiMn"][:_n_pts]],
                    f"Mn_x_nominal [{unidad_mom}]":   [round(v, 3) for v in cap_x["Mn"][:_n_pts]],
                    f"phiMn_y [{unidad_mom}]":        [round(v, 3) for v in cap_y["phiMn"][:_n_pts]],
                    f"Mn_y_nominal [{unidad_mom}]":   [round(v, 3) for v in cap_y["Mn"][:_n_pts]],
                })
                _csv_pm_bytes = _df_pm_csv.to_csv(index=False).encode("utf-8")
                st.download_button(
                    label=_t("Descargar puntos del diagrama P-M (CSV)",
                              "Download P-M diagram points (CSV)"),
                    data=_csv_pm_bytes,
                    file_name="diagrama_PM_puntos.csv",
                    mime="text/csv",
                    key="c_pm_dl_pm_csv",
                    help=_t(
                        "Descarga los puntos numericos de ambas curvas P-M (eje X e eje Y) "
                        "en valores nominales y reducidos. Util para verificacion independiente.",
                        "Downloads numeric points of both P-M curves (X and Y axes) "
                        "nominal and reduced values. Useful for independent verification."
                    )
                )
        except Exception:
            pass
        # ── Fin CSV P-M ────────────────────────────────────────────────────────

        st.subheader(_t(" Superficie de Interacción Biaxial 3D — Punto actual",
                        " Biaxial 3D Interaction Surface — Current point"))
        st.plotly_chart(fig_3d, use_container_width=True)
        st.caption(_t(
            "▶ Arrastra para rotar · Scroll = zoom · Doble-clic = resetear. "
            "Verde = envolvente φ · Gris = nominal. "
            "Cambia b, h, ρ en el sidebar → diagrama en vivo.",
            "▶ Drag to rotate · Scroll = zoom · Double-click = reset. "
            "Green = design envelope φ · Grey = nominal. "
            "Change b, h, ρ in sidebar → live updates."))
    with col2:
        st.subheader(_t("Verificación Biaxial (Bresler)", " Biaxial Verification (Bresler)"))
        st.markdown(f"""
        | Parámetro | Valor |
        |-----------|-------|
        | **φPnx** (para Mux={Mux_input:.1f}) | {bresler['phi_Pnx']:.2f} {unidad_fuerza} |
        | **φPny** (para Muy={Muy_input:.1f}) | {bresler['phi_Pny']:.2f} {unidad_fuerza} |
        | **φP0** (axial pura) | {bresler['phi_P0']:.2f} {unidad_fuerza} |
        | **φPni** (Bresler) | {bresler['phi_Pni']:.2f} {unidad_fuerza} |
        | **Pu** solicitante | {Pu_input:.2f} {unidad_fuerza} |
        | **Relación Pu/φPni** | {bresler['ratio']:.3f} |
        """)
        if bresler['ok']:
            st.success(f"**VERIFICACIÓN BIAXIAL CUMPLE**\n\nPu ({Pu_input:.1f}) ≤ φPni ({bresler['phi_Pni']:.1f})")
        else:
            st.error(f"**VERIFICACIÓN BIAXIAL NO CUMPLE**\n\nPu ({Pu_input:.1f}) > φPni ({bresler['phi_Pni']:.1f})")
            ratio = bresler['ratio']
            deficit = Pu_input - bresler['phi_Pni']
            st.markdown("** Recomendaciones para cumplir:**")
            recomendaciones = []
            if ratio > 1.5 and math.isfinite(ratio):
                pct = math.ceil((ratio**0.5 - 1) * 100)
                recomendaciones.append(f" **Aumentar sección:** La columna necesita ~{pct}% más de área — aumentar b y/o h en el sidebar.")
            if ratio <= 3:
                rho_obj = min(rho_max, cuantia * ratio**0.5)
                recomendaciones.append(f" **Aumentar acero longitudinal:** Añadir varillas o usar diámetro mayor — apuntar a ρ ≥ {rho_obj:.1f}%.")
            if ratio > 2:
                recomendaciones.append(f" **Reducir Pu:** La carga axial ({Pu_input:.0f} {unidad_fuerza}) supera {ratio:.1f}x la capacidad — revisar predimensionamiento.")
            if Mux_input > 0 or Muy_input > 0:
                recomendaciones.append(" **Reducir momentos Mux/Muy:** Considerar arriostrar la estructura o reducir excentricidades.")
            for rec in recomendaciones:
                st.markdown(f"- {rec}")
            st.markdown(f"""
| | |
|---|---|
| **Capacidad requerida** | φPni ≥ {Pu_input:.1f} {unidad_fuerza} |
| **Capacidad actual** | φPni = {bresler['phi_Pni']:.1f} {unidad_fuerza} |
| **Déficit** | {deficit:.1f} {unidad_fuerza} ({(ratio-1)*100:.0f}% sobre la capacidad) |
""")
            if ratio > 5:
                st.error(" **Relación > 5x:** La sección es muy insuficiente. Se recomienda rediseñar completamente la geometría.")
        st.markdown("---")
        st.subheader(_t("Verificación de Esbeltez", "Slenderness Verification"))
        st.markdown(f"""
        | Parámetro | Valor | Estado |
        |-----------|-------|--------|
        | **Tipo de Pórtico** | {"Desplazable (P-Δ en Mu)" if es_desplazable else "No Desplazable (δns aplicado)"} | |
        | **kl/r** | {slenderness['kl_r']:.1f} | |
        | **Clasificación** | {slenderness['classification']} | |
        | **δns (magnificación)** | {"1.000 (no aplica)" if es_desplazable else f"{slenderness['delta_ns']:.3f}"} | |
        | **Mux amplificado** | {Mux_magnified:.2f} {unidad_mom} | |
        | **Muy amplificado** | {Muy_magnified:.2f} {unidad_mom} | |
        """)
        if slenderness['kl_r'] > 100:
            st.warning(f"️ **kl/r > 100** — Se requiere análisis no lineal de segundo orden según {norma_sel}.")
        elif slenderness['kl_r'] > 22:
            if es_desplazable:
                st.info(
                    f"ℹ️ **Marco Desplazable (Sway Frame) — kL/r = {slenderness['kl_r']:.1f}.** "
                    f"Los momentos Mux y Muy ingresados se asumen **ya amplificados por P-Δ** del modelo global. "
                    f"No se aplica magnificación δns adicional. Ref: {norma_sel} C.10.10.7 / ACI 318-19 §6.6.4.6."
                )
            else:
                st.info(
                    f"ℹ️ **Marco No Desplazable — δns = {slenderness['delta_ns']:.3f} aplicado.** "
                    f"Momentos amplificados: Mux = {Mux_magnified:.2f} {unidad_mom} | "
                    f"Muy = {Muy_magnified:.2f} {unidad_mom}. "
                    f"Ref: {norma_sel} C.10.10 / ACI 318-19 §6.6.4.5."
                )
        st.markdown("---")
        st.subheader(_t("Verificación de Estribos / Espiral", "Tie / Spiral Verification"))
        if not es_circular:
            req_1_str = f"0.3 \\times {s_conf:.1f} \\times {bc:.1f} \\times ({fc:.1f}/{fyt:.0f}) \\times ({Ag:.1f}/{Ach:.1f} - 1) = {Ash_req_1:.2f} \\text{{ cm}}^2"
            req_2_str = f"0.09 \\times {s_conf:.1f} \\times {bc:.1f} \\times ({fc:.1f}/{fyt:.0f}) = {Ash_req_2:.2f} \\text{{ cm}}^2"
            
            st.markdown(f"**Cálculo de Ash requerido (NSR-10 C.21.3.5.4):**")
            st.latex(r"(a) \quad A_{sh} = " + req_1_str)
            st.latex(r"(b) \quad A_{sh} = " + req_2_str)
            st.markdown(f"**→ Rige: {Ash_req:.2f} cm²**")

            st.markdown(f"""
            | Parámetro | Valor | Requerido | Estado |
            |-----------|-------|-----------|--------|
            | **Claro Libre (Cx, Cy)** | {claro_libre_x:.1f} cm, {claro_libre_y:.1f} cm | ≤ 15 cm | {'' if claro_libre_x<=15 and claro_libre_y<=15 else ' Crossties Requeridos'} |
            | **Apoyo lateral (Crossties)** | {num_flejes_x} en X, {num_flejes_y} en Y | NSR-10 C.7.10.5 | |
            | **Ramas Efectivas** | {ramas_x} ramas en X, {ramas_y} ramas en Y | | |
            | **Ash provisto** | {Ash_prov:.3f} cm² | ≥ {Ash_req:.2f} | {'' if ash_ok else ''} |
            | **s_conf** | {s_conf:.1f} cm | {'≤ 15' if es_des else '≤ 20' if es_dmo else '≤ min(b,h)'} | |
            | **Lo_conf** | {Lo_conf:.1f} cm | ≥ max(b,h,L/6,45) | {''} |
            | **N° estribos + Crossties** | {n_estribos_total} juegos de ramas | C.21.3.5 | |
            """)

            if _n_flejes_auto > 0:
                st.success(
                    f" **Auto-solución aplicada:** Se añadieron **{_n_flejes_auto} flete(s)/cara** "
                    f"para cumplir Ash_req = {Ash_req:.2f} cm².\n\n"
                    f"- Ash provisto corregido: **{_Ash_prov_auto:.2f} cm²** ≥ {Ash_req:.2f} cm²\n"
                    f"- Ramas efectivas: **{ramas_x} en X · {ramas_y} en Y**\n"
                    f"- Tipo gancho: {_tipo_flete}"
                )
            if not ash_ok:
                ratio_ash = Ash_prov / Ash_req if Ash_req > 0 else 1.0
                s_req1 = Ash_prov / (0.3 * bc * fc / fyt * (Ag/Ach - 1)) if (Ag/Ach - 1) > 0 else 999
                s_req2 = Ash_prov * fyt / (0.09 * bc * fc)
                s_correcto = min(s_req1, s_req2)
                st.error(
                    f" **Déficit persistente** aún con {num_flejes_x} flejes/cara.\n\n"
                    f"Opciones: reducir separación a **s ≤ {s_correcto:.1f} cm**, "
                    f"aumentar diámetro de estribo, o aumentar b/h de la sección."
                )
        else:
            st.markdown(f"""
            | Parámetro | Valor | Requerido | Estado |
            |-----------|-------|-----------|--------|
            | **ρs requerido** | {rho_s_req:.4f} | | |
            | **ρs provisto** | {rho_s_prov:.4f} | ≥ {rho_s_req:.4f} | {'' if ash_ok else ''} |
            | **Paso espiral** | {paso_espiral:.1f} cm | ≤ min(D/5, 8 cm) | {'' if paso_espiral <= min(D/5, 8) else ''} |
            | **N° vueltas** | {n_estribos_total} | | |
            """)
        st.caption(f"Ref: {code['ref']} | Nivel Sísmico: {nivel_sismico}")

# =============================================================================
# TAB 2: SECCIÓN Y ESTRIBOS (con DXF y RÓTULO ICONTEC)
# =============================================================================
with tab2:
    st.subheader(_t("Visualización 3D de la Columna", "3D Column Visualization"))
    fig3d_col = go.Figure()
    if es_circular:
        theta = np.linspace(0, 2*np.pi, 50)
        z = np.linspace(0, L_col, 20)
        theta_grid, z_grid = np.meshgrid(theta, z)
        x_grid = (D/2) * np.cos(theta_grid)
        y_grid = (D/2) * np.sin(theta_grid)
        fig3d_col.add_trace(go.Surface(x=x_grid, y=y_grid, z=z_grid, opacity=0.3, colorscale='Greys', showscale=False))
    else:
        x_c = [-b/2, b/2, b/2, -b/2, -b/2, b/2, b/2, -b/2]
        y_c = [-h/2, -h/2, h/2, h/2, -h/2, -h/2, h/2, h/2]
        z_c = [0, 0, 0, 0, L_col, L_col, L_col, L_col]
        fig3d_col.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
        
    #  BARRAS LONGITUDINALES 3D 
    z_barras = [0, L_col]
    if es_circular:
        offset_real_cm = recub_cm + (stirrup_diam / 10.0) + (rebar_diam / 20.0)
        radio_plot = D/2 - offset_real_cm
        ang_step = 2 * math.pi / n_barras
        for i_bar in range(n_barras):
            xb = radio_plot * math.cos(i_bar * ang_step)
            yb = radio_plot * math.sin(i_bar * ang_step)
            fig3d_col.add_trace(go.Scatter3d(
                x=[xb, xb], y=[yb, yb], z=z_barras,
                mode='lines',
                line=dict(color='#ff6b35', width=4),
                showlegend=False, name='Barra long.'
            ))
    else:
        offset_real_cm = recub_cm + (stirrup_diam / 10.0) + (rebar_diam / 20.0)
        r_bar3d = rebar_diam / 20.0
        xs3d = np.linspace(offset_real_cm - b/2, b/2 - offset_real_cm, num_filas_h) if num_filas_h > 1 else [0.0]
        ys3d_bot = -(h/2 - offset_real_cm)
        ys3d_top =   h/2 - offset_real_cm
        for x3 in xs3d:
            for y3 in [ys3d_bot, ys3d_top]:
                fig3d_col.add_trace(go.Scatter3d(
                    x=[x3, x3], y=[y3, y3], z=z_barras,
                    mode='lines',
                    line=dict(color='#ff6b35', width=4),
                    showlegend=False, name='Barra long.'
                ))
        if num_capas_intermedias > 0:
            esp3d = (h - 2*offset_real_cm) / (num_capas_intermedias + 1)
            for ci in range(1, num_capas_intermedias + 1):
                yi = -(h/2) + offset_real_cm + ci * esp3d
                for xi in [-(b/2 - offset_real_cm), b/2 - offset_real_cm]:
                    fig3d_col.add_trace(go.Scatter3d(
                        x=[xi, xi], y=[yi, yi], z=z_barras,
                        mode='lines',
                        line=dict(color='#ff6b35', width=4),
                        showlegend=False, name='Barra lat.'
                    ))

    #  ESTRIBOS + FLEJES EN CRUZ 3D 
    if not es_circular:
        # Sincronizar alturas 3D EXACTAMENTE con la cantidad del APU
        z_est = []
        _esp_real_conf = Lo_conf / n_est_por_Lo if n_est_por_Lo > 0 else s_conf
        # Zona inferior confinada (incluye z=0 y z=Lo_conf)
        for i in range(n_est_por_Lo + 1):
            z_est.append(i * _esp_real_conf)
        # Zona central (distribuida uniformemente cerrando el gap superior)
        if n_estribos_centro > 0:
            esp_real_centro = (longitud_zona_libre + _esp_real_conf) / (n_estribos_centro + 1)
            for i in range(1, n_estribos_centro + 1):
                z_est.append(Lo_conf + i * esp_real_centro)
        # Zona superior confinada (hasta L_col para encajar exacto con APU)
        if n_est_por_Lo > 0:
            for i in range(n_est_por_Lo):
                z_est.append(L_col - Lo_conf + (i+1) * _esp_real_conf)
        bw = b/2 - recub_cm    # semiancho eje estribo
        hw = h/2 - recub_cm    # semialto  eje estribo
        _lv = min(bw, hw) * 0.18  # longitud del gancho 135° en 3D
        _estr_shown = False
        _flet_shown = False
        for _iz, ze in enumerate(z_est):
            # ── Estribo perimetral (rectángulo redondeado simplificado) ────
            _n_arc = 6
            _r3 = min(bw, hw) * 0.12
            _ex, _ey = [], []
            for _seg, (_x0,_y0,_x1,_y1,_a0,_a1) in enumerate([
                (-bw+_r3,-hw, bw-_r3,-hw, 180,270),
                (bw-_r3,-hw, bw,-hw+_r3, 270,360),
                (bw,-hw+_r3, bw, hw-_r3, None,None),
                (bw-_r3,hw, -bw+_r3,hw, 0,90),
                (-bw,hw-_r3,-bw,-hw+_r3, 90,180),
            ]):
                if _a0 is not None:
                    _cxs = [_x0+(_x1-_x0)/2, _x0+(_x1-_x0)/2][0]
                    _cys = [_y0+(_y1-_y0)/2, _y0+(_y1-_y0)/2][0]
                    # esquina: centro del arco
                    if _seg==0:   _cx,_cy = -bw+_r3, -hw+_r3
                    elif _seg==1: _cx,_cy =  bw-_r3, -hw+_r3
                    elif _seg==3: _cx,_cy =  bw-_r3,  hw-_r3
                    elif _seg==4: _cx,_cy = -bw+_r3,  hw-_r3
                    else:         _cx,_cy = 0,0
                    for _k in range(_n_arc+1):
                        _a = np.radians(_a0+(_a1-_a0)*_k/_n_arc)
                        _ex.append(_cx + _r3*np.cos(_a))
                        _ey.append(_cy + _r3*np.sin(_a))
                else:
                    _ex += [_x0,_x1]; _ey += [_y0,_y1]
            _ex.append(_ex[0]); _ey.append(_ey[0])  # cerrar
            # Ganchos 135° alternados (j%2)
            _s45 = _lv * 0.707
            if _iz % 2 == 0:
                _hkx = [-bw, -bw+_s45];  _hky = [hw-_r3, hw-_r3+_s45]   # sup-izq
            else:
                _hkx = [ bw,  bw-_s45];  _hky = [-hw+_r3,-hw+_r3-_s45]  # inf-der (espejo)
            fig3d_col.add_trace(go.Scatter3d(
                x=_ex+_hkx, y=_ey+_hky, z=[ze]*(len(_ex)+len(_hkx)),
                mode='lines', line=dict(color='#00d4ff', width=2),
                showlegend=not _estr_shown, name='Estribo perimetral',
                legendgroup='estribo'
            ))
            _estr_shown = True
            # ── Flejes en X (Anclados a varillas reales) ──
            if num_flejes_x > 0 and (num_filas_v - 2) > 0:
                n_int_3d = num_filas_v - 2
                n_ties_3d = min(num_flejes_x, n_int_3d)
                idx_ties_3d = np.linspace(1, n_int_3d, n_ties_3d).round().astype(int)
                esp_3d_y = (h - 2*offset_real_cm) / (n_int_3d + 1)
                for _nfx in idx_ties_3d:
                    _yf = -(h/2) + offset_real_cm + esp_3d_y * _nfx
                    _hk_ang = _s45 if _iz % 2 == 0 else -_s45
                    fig3d_col.add_trace(go.Scatter3d(x=[-bw, bw], y=[_yf, _yf], z=[ze, ze], mode='lines', line=dict(color='#00d4ff', width=2.5), showlegend=not _flet_shown, name=f'Flete en X ({n_ties_3d})', legendgroup='flejeX'))
                    fig3d_col.add_trace(go.Scatter3d(x=[-bw,-bw+_s45], y=[_yf,_yf+_hk_ang], z=[ze,ze], mode='lines', line=dict(color='#00d4ff', width=2.5), showlegend=False, legendgroup='flejeX'))
                    fig3d_col.add_trace(go.Scatter3d(x=[bw,bw-_s45], y=[_yf,_yf-_hk_ang], z=[ze,ze], mode='lines', line=dict(color='#00d4ff', width=2.5), showlegend=False, legendgroup='flejeX'))
                    _flet_shown = True
            # ── Flejes en Y (Anclados a varillas reales) ──
            if num_flejes_y > 0 and (num_filas_h - 2) > 0:
                n_int_3d = num_filas_h - 2
                n_ties_3d = min(num_flejes_y, n_int_3d)
                idx_ties_3d = np.linspace(1, n_int_3d, n_ties_3d).round().astype(int)
                esp_3d_x = (b - 2*offset_real_cm) / (n_int_3d + 1)
                for _nfy in idx_ties_3d:
                    _xf = -(b/2) + offset_real_cm + esp_3d_x * _nfy
                    _hk_ang = _s45 if _iz % 2 == 0 else -_s45
                    fig3d_col.add_trace(go.Scatter3d(x=[_xf, _xf], y=[-hw, hw], z=[ze, ze], mode='lines', line=dict(color='#00d4ff', width=2.5), showlegend=not _flet_shown, name=f'Flete en Y ({n_ties_3d})', legendgroup='flejeY'))
                    fig3d_col.add_trace(go.Scatter3d(x=[_xf,_xf+_hk_ang], y=[-hw,-hw+_s45], z=[ze,ze], mode='lines', line=dict(color='#00d4ff', width=2.5), showlegend=False, legendgroup='flejeY'))
                    fig3d_col.add_trace(go.Scatter3d(x=[_xf,_xf-_hk_ang], y=[hw,hw-_s45], z=[ze,ze], mode='lines', line=dict(color='#00d4ff', width=2.5), showlegend=False, legendgroup='flejeY'))
                    _flet_shown = True
    else:
        rc_sp = D/2 - recub_cm
        theta_sp = np.linspace(0, 2*np.pi, 36)
        z_sp = np.arange(0, L_col + paso_espiral, paso_espiral)
        for ze in z_sp:
            fig3d_col.add_trace(go.Scatter3d(
                x=rc_sp * np.cos(theta_sp),
                y=rc_sp * np.sin(theta_sp),
                z=[ze]*len(theta_sp),
                mode='lines',
                line=dict(color='#00d4ff', width=2),
                showlegend=False, name='Espiral'
            ))

    # Ajustar aspecto para que no se deforme la seccion real (b vs h)
    # Correct proportional 3D aspect: scale b,h,L to each other so section looks real
    _dim_max = max(b if not es_circular else D, h if not es_circular else D, L_col)
    ar_x = (D / _dim_max) if es_circular else (b / _dim_max)
    ar_y = (D / _dim_max) if es_circular else (h / _dim_max)
    ar_z = L_col / _dim_max
    # ── Zonas confinadas (Lo_conf) — cajas semitransparentes ─────────────────
    if not es_circular:
        _Lo = Lo_conf   # cm
        _bw3 = b/2 - recub_cm
        _hw3 = h/2 - recub_cm
        for _z0_zona, _z1_zona, _color_zona in [
            (0,     _Lo,          'rgba(255,100,0,0.15)'),   # base
            (L_col-_Lo, L_col,    'rgba(255,100,0,0.15)'),   # tope
        ]:
            _xz = [-b/2, b/2, b/2, -b/2, -b/2, b/2,  b/2, -b/2]
            _yz = [-h/2,-h/2, h/2,  h/2, -h/2,-h/2,  h/2,  h/2]
            _zz = [_z0_zona]*4 + [_z1_zona]*4
            fig3d_col.add_trace(go.Mesh3d(
                x=_xz, y=_yz, z=_zz,
                alphahull=0, opacity=0.25,
                color='orange', showlegend=False,
                name='Zona confinada Lo',
                hovertemplate=f'Zona confinada Lo={_Lo:.0f}cm<extra></extra>',
            ))
        # Etiquetas Lo en el alzado
        for _zl, _lbl in [(0, f'Lo={_Lo:.0f}cm'), (L_col, f'Lo={_Lo:.0f}cm')]:
            fig3d_col.add_trace(go.Scatter3d(
                x=[b/2+2], y=[0], z=[_zl + _Lo/2],
                mode='text', text=[_lbl],
                textfont=dict(color='orange', size=10),
                showlegend=False,
            ))

    fig3d_col.update_layout(
        scene=dict(
            aspectmode='manual',
            aspectratio=dict(x=ar_x, y=ar_y, z=ar_z),
            xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)',
            xaxis=dict(range=[-b/2-5, b/2+5] if not es_circular else [-D/2-5, D/2+5]),
            yaxis=dict(range=[-h/2-5, h/2+5] if not es_circular else [-D/2-5, D/2+5]),
        ),
        height=500, margin=dict(l=0, r=0, b=0, t=30),
        title=dict(text="Visualizacion 3D de la Columna — Proporciones Reales", font=dict(size=13))
    )
    st.plotly_chart(fig3d_col, use_container_width=True)

    # ─────────────────────────────────────────────────────────────────
    # RESUMEN DE ESTRIBOS Y VERIFICACIONES BAJO EL 3D
    # ─────────────────────────────────────────────────────────────────
    _col3d_a, _col3d_b = st.columns(2)

    with _col3d_a:
        st.subheader(_t("Distribucion de Estribos (NSR-10)", "Stirrup Distribution (NSR-10)"))
        # Calcular cantidad real dibujada en el 3D (zona-aware)
        _n_conf_bot = sum(1 for z in z_est if z <= Lo_conf)          if not es_circular else 0
        _n_conf_top = sum(1 for z in z_est if z >= L_col - Lo_conf)  if not es_circular else 0
        _n_medio    = sum(1 for z in z_est if Lo_conf < z < L_col - Lo_conf) if not es_circular else 0
        _n_3d_total = len(z_est) if not es_circular else 0

        _estr_rows = {
            _t("Zona", "Zone"): [
                _t("Confinada inferior", "Bottom confined"),
                _t("Zona central",       "Mid zone"),
                _t("Confinada superior", "Top confined"),
                _t("TOTAL dibujado 3D",  "TOTAL drawn 3D"),
                _t("TOTAL APU (calculo)","TOTAL APU (calc)"),
            ],
            _t("Long. zona", "Zone length"): [
                f"{Lo_conf:.0f} cm",
                f"{max(0, L_col - 2*Lo_conf):.0f} cm",
                f"{Lo_conf:.0f} cm",
                "", ""
            ],
            _t("Sep.", "Spacing"): [
                f"s = {s_conf:.0f} cm",
                f"s = {s_basico:.0f} cm",
                f"s = {s_conf:.0f} cm",
                "", ""
            ],
            _t("Nº estribos", "# Stirrups"): [
                str(_n_conf_bot),
                str(_n_medio),
                str(_n_conf_top),
                str(_n_3d_total),
                str(n_estribos_total),
            ],
        }
        st.dataframe(pd.DataFrame(_estr_rows), use_container_width=True, hide_index=True)

        if not es_circular:
            _diff = n_estribos_total - _n_3d_total
            if _diff > 0:
                st.info(f"[INFO] El modelo 3D dibuja {_n_3d_total} estribos. El APU calcula {n_estribos_total}. "
                        f"Diferencia: {_diff} estribos en los extremos (margen de 5 cm en cada tope).")
            else:
                st.success(f"[OK] El modelo 3D coincide con el APU: {n_estribos_total} estribos totales.")

    with _col3d_b:
        st.subheader(_t("Resumen de Verificaciones", "Verification Summary"))
        _checks_3d = {
            _t("Verificacion", "Check"): [
                _t("Cuantia longitudinal", "Long. ratio"),
                _t("Verificacion biaxial", "Biaxial check"),
                _t("Esbeltez (kl/r <= 22)", "Slenderness"),
                _t(f"Ash {'espiral' if es_circular else 'estribos'}", f"Ash {'spiral' if es_circular else 'stirrups'}"),
                _t("Longitud confinamiento Lo", "Conf. length Lo"),
                _t("Separacion maxima", "Max spacing"),
            ],
            _t("Valor", "Value"): [
                f"rho={cuantia*100:.2f}% [{rho_min*100:.1f}-{rho_max*100:.1f}%]",
                f"phi*Mn/Mu = {bresler.get('ratio', 0):.2f}" if bresler.get('ok') else "No cumple",
                f"kl/r = {slenderness['kl_r']:.1f}",
                f"Ash_prov >= Ash_req" if ash_ok else f"Ash insuficiente",
                f"Lo = {Lo_conf:.0f} cm",
                f"s_conf={s_conf:.0f} cm | s_bas={s_basico:.0f} cm",
            ],
            _t("Estado", "Status"): [
                "[OK]" if rho_min <= cuantia <= rho_max else "[!]",
                "[OK]" if bresler['ok'] else "[!]",
                "[OK]" if slenderness['kl_r'] <= 22 else "[!]",
                "[OK]" if ash_ok else "[!]",
                "[OK]",
                "[OK]" if (es_des and s_conf <= 15) or (es_dmo and s_conf <= 20) or es_dmi else "[!]",
            ],
        }
        st.dataframe(pd.DataFrame(_checks_3d), use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader(_t("Seccion Transversal", "Cross Section"))
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        fig_sec, ax_s = plt.subplots(figsize=(5, 5))
        fig_sec.patch.set_facecolor('#1e1e2e')
        for _ax in fig_sec.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
        ax_s.set_aspect('equal')
        ax_s.set_facecolor('#1a1a2e')
        fig_sec.patch.set_facecolor('#1a1a2e')
        if es_circular:
            circle = plt.Circle((0, 0), D/2, linewidth=2, edgecolor='white', facecolor='#4a4a6a', fill=True)
            ax_s.add_patch(circle)
            circle_rec = plt.Circle((0, 0), D/2 - recub_cm, linewidth=1.5, edgecolor='#00d4ff', facecolor='none', linestyle='--')
            ax_s.add_patch(circle_rec)
            offset_real_cm = recub_cm + (stirrup_diam / 10.0) + (rebar_diam / 20.0)
            radio_plot = D/2 - offset_real_cm
            ang_step = 2 * math.pi / n_barras
            for i_bar in range(n_barras):
                x_pos = radio_plot * math.cos(i_bar * ang_step)
                y_pos = radio_plot * math.sin(i_bar * ang_step)
                ax_s.add_patch(plt.Circle((x_pos, y_pos), rebar_diam/20, color='#ff6b35', zorder=5))
            ax_s.set_xlim(-D/2 - 5, D/2 + 5)
            ax_s.set_ylim(-D/2 - 5, D/2 + 5)
        else:
            ax_s.add_patch(patches.Rectangle((0, 0), b, h, linewidth=2, edgecolor='white', facecolor='#4a4a6a'))
            ax_s.add_patch(patches.Rectangle((recub_cm, recub_cm), b-2*recub_cm, h-2*recub_cm,
                linewidth=1.5, edgecolor='#00d4ff', facecolor='none', linestyle='--'))
            offset_real_cm = recub_cm + (stirrup_diam / 10.0) + (rebar_diam / 20.0)
            r_bar = rebar_diam / 20.0
            xs = np.linspace(offset_real_cm, b - offset_real_cm, num_filas_h) if num_filas_h > 1 else [b/2]
            for i_x, x in enumerate(xs):
                ax_s.add_patch(plt.Circle((x, h - offset_real_cm), r_bar, color='#ff6b35', zorder=5))
                ax_s.add_patch(plt.Circle((x, offset_real_cm), r_bar, color='#ff6b35', zorder=5))
                
            if num_capas_intermedias > 0:
                esp = (h - 2*offset_real_cm) / (num_capas_intermedias + 1)
                for i in range(1, num_capas_intermedias + 1):
                    y_int = offset_real_cm + i * esp
                    ax_s.add_patch(plt.Circle((offset_real_cm, y_int), r_bar, color='#ff6b35', zorder=5))
                    ax_s.add_patch(plt.Circle((b - offset_real_cm, y_int), r_bar, color='#ff6b35', zorder=5))
            # ── LÍNEAS PUNTEADAS ANCLADAS A LAS VARILLAS EXACTAS ──
            ys_reales = np.linspace(offset_real_cm, h - offset_real_cm, num_filas_v)
            if num_flejes_x > 0 and len(ys_reales) > 2:
                ys_int = ys_reales[1:-1]
                idx = np.linspace(0, len(ys_int)-1, min(num_flejes_x, len(ys_int))).round().astype(int)
                for i in idx:
                    ax_s.plot([offset_real_cm, b - offset_real_cm], [ys_int[i], ys_int[i]], color='#00d4ff', linewidth=1.5, linestyle='--', zorder=4)

            xs_reales = np.linspace(offset_real_cm, b - offset_real_cm, num_filas_h)
            if num_flejes_y > 0 and len(xs_reales) > 2:
                xs_int = xs_reales[1:-1]
                idx = np.linspace(0, len(xs_int)-1, min(num_flejes_y, len(xs_int))).round().astype(int)
                for i in idx:
                    ax_s.plot([xs_int[i], xs_int[i]], [offset_real_cm, h - offset_real_cm], color='#00d4ff', linewidth=1.5, linestyle='--', zorder=4)
            ax_s.set_xlim(-5, b + 5)
            ax_s.set_ylim(-5, h + 5)
        ax_s.axis('off')
        ax_s.set_title(f"Sección {'Circular' if es_circular else 'Rectangular'} — {n_barras} varillas {_bar_label_short(rebar_diam)}", color='white', fontsize=9)
        configurar_pdf_comercial(fig_sec)
        st.pyplot(fig_sec)
        plt.close(fig_sec)
    with col_s2:
        st.subheader(_t("Resumen de Grapas (Cross-Ties)", "Cross-Tie Summary"))
        _grapa_rows = {
            _t("Elemento", "Element"): [],
            _t("Cant.", "Qty"): [],
            _t("Long. (m)", "Length (m)"): [],
            _t("Peso (kg)", "Weight (kg)"): [],
        }
        # Estribo perimetral
        _le_r = (2*(b-2*recub_cm) + 2*(h-2*recub_cm) + 12*stirrup_diam/10) / 100
        _pe_r = n_estribos_total * _le_r * (stirrup_area*100) * 7.85e-3
        _grapa_rows[_t("Elemento","Element")].append("E1 - Estribo perim.")
        _grapa_rows[_t("Cant.","Qty")].append(str(n_estribos_total))
        _grapa_rows[_t("Long. (m)","Length (m)")].append(f"{_le_r:.2f}")
        _grapa_rows[_t("Peso (kg)","Weight (kg)")].append(f"{_pe_r:.1f}")
        # Grapas X
        if num_flejes_x > 0:
            _cgx = n_estribos_total * num_flejes_x
            _lgx = long_fleje_x / 100
            _pgx = _cgx * _lgx * (stirrup_area*100) * 7.85e-3
            _grapa_rows[_t("Elemento","Element")].append(f"GX ({num_flejes_x}/estribo)")
            _grapa_rows[_t("Cant.","Qty")].append(str(_cgx))
            _grapa_rows[_t("Long. (m)","Length (m)")].append(f"{_lgx:.2f}")
            _grapa_rows[_t("Peso (kg)","Weight (kg)")].append(f"{_pgx:.1f}")
        # Grapas Y
        if num_flejes_y > 0:
            _cgy = n_estribos_total * num_flejes_y
            _lgy = long_fleje_y / 100
            _pgy = _cgy * _lgy * (stirrup_area*100) * 7.85e-3
            _grapa_rows[_t("Elemento","Element")].append(f"GY ({num_flejes_y}/estribo)")
            _grapa_rows[_t("Cant.","Qty")].append(str(_cgy))
            _grapa_rows[_t("Long. (m)","Length (m)")].append(f"{_lgy:.2f}")
            _grapa_rows[_t("Peso (kg)","Weight (kg)")].append(f"{_pgy:.1f}")
        st.dataframe(pd.DataFrame(_grapa_rows), use_container_width=True, hide_index=True)
        st.markdown("---")
        st.subheader(_t("Exportar Plano DXF (ICONTEC)", "Export DXF (ICONTEC)"))

        # ── CONFIGURACION DEL PLANO ─────────────────────────────────────────
        with st.expander(_t("Configurar Plano y Rotulo", "Configure Plot & Title Block"), expanded=True):
            col_cfg1, col_cfg2 = st.columns(2)
            with col_cfg1:
                dxf_empresa  = st.text_input("Empresa",  "INGENIERIA ESTRUCTURAL SAS", key="col_dxf_emp")
                dxf_proyecto = st.text_input("Proyecto", "Proyecto Estructural",        key="col_dxf_proy")
                dxf_plano    = st.text_input("N Plano",  "COL-001",                     key="col_dxf_num")
                dxf_elaboro  = st.text_input("Elaboro",  "Ing. Disenador",              key="col_dxf_elab")
            with col_cfg2:
                dxf_reviso   = st.text_input("Reviso",   "Ing. Revisor",                key="col_dxf_rev")
                dxf_aprobo   = st.text_input("Aprobo",   "Ing. Aprobador",              key="col_dxf_apr")
                # Selector de tamanio de papel colombiano (ICONTEC)
                papel_opciones = {
                    "Carta  (216 x 279 mm)":     (21.6,  27.9,  "CARTA"),
                    "Oficio (216 x 330 mm)":     (21.6,  33.0,  "OFICIO"),
                    "Medio Pliego (500 x 707 mm)":(50.0,  70.7,  "MEDIO PLIEGO"),
                    "Pliego       (707 x 1000 mm)":(70.7, 100.0, "PLIEGO"),
                }
                papel_sel = st.selectbox("Tamano de Papel (ICONTEC)", list(papel_opciones.keys()), index=0, key="col_dxf_papel")
                ANCHO_PLANO, ALTO_PLANO, PAPEL_LABEL = papel_opciones[papel_sel]

        # ── GENERACION DXF ──────────────────────────────────────────────────
        if st.button(_t("Generar Plano DXF ICONTEC", "Generate DXF ICONTEC"), key="col_btn_dxf"):
            try:
                from ezdxf.enums import TextEntityAlignment
                doc_dxf = ezdxf.new('R2010', setup=True)
                # 1. Configurar unidades a milímetros estrictos
                doc_dxf.header['$INSUNITS'] = 5 # units.CM (Critico para que no desparezca el corte)

                # Lineweights ICONTEC (en mm x 100 para ezdxf)
                # ICONTEC NTC 1033 / ISO 128: 0.18, 0.25, 0.35, 0.50, 0.70 mm
                LW = {
                    'CONCRETO':    50,   # 0.50 mm  contorno exterior
                    'ACERO_LONG':  35,   # 0.35 mm  barras longitudinales
                    'ACERO_TRANS': 25,   # 0.25 mm  estribos / flejes
                    'DOBLEZ':      25,   # 0.25 mm  diagramas de doblez
                    'COTAS':       18,   # 0.18 mm  lineas de cota
                    'TEXTO':       18,   # 0.18 mm  anotaciones
                    'EJES':        13,   # 0.13 mm  ejes de simetria
                    'ROTULO':      35,   # 0.35 mm  cuadro del rotulo
                    'MARGEN':      50,   # 0.50 mm  marco exterior
                }
                COLORES = {
                    'CONCRETO': 7,   # blanco
                    'ACERO_LONG': 7,   # rojo
                    'ACERO_TRANS': 7,   # cian
                    'DOBLEZ': 7,   # magenta
                    'COTAS': 7,   # amarillo
                    'TEXTO': 7,   # blanco
                    'EJES': 7,   # gris
                    'ROTULO': 7,   # gris
                    'MARGEN': 7,   # blanco
                }
                # Capas con grosores (lw) aumentados para legibilidad en PDF y DWG
                capas = {
                    'MARCO': {'color': 7, 'lw': 25},  # Borde de hoja a la mitad
                    'TEXTO': {'color': 7, 'lw': 35},  # Letras nítidas
                    'ACERO_LONG': {'color': 7, 'lw': 50},  
                    'ACERO_TRANS': {'color': 7, 'lw': 35},  
                    'COTAS': {'color': 7, 'lw': 25},  
                    'CONCRETO': {'color': 7, 'lw': 25}, # Borde de concreto a la mitad
                    'DOBLEZ': {'color': 7, 'lw': 35},
                    'EJES': {'color': 7, 'lw': 18, 'linetype': 'DASHDOT'},
                    'ROTULO': {'color': 7, 'lw': 25}, # Cajetín y tablas a la mitad
                    'MARGEN': {'color': 7, 'lw': 25}
                }

                def _color_acero_dxf_col(d_mm):
                    if d_mm <= 8.0: return 5   # azul 
                    if d_mm <= 10.0: return 4  # cyan 
                    if d_mm <= 12.0: return 3  # verde 
                    if d_mm <= 16.0: return 6  # magenta (reemplaza amarillo para evitar que no se vea al imprimir)
                    if d_mm <= 20.0: return 1  # rojo
                    if d_mm <= 25.0: return 6  # magenta
                    return 1 # por defecto
                    

                for nombre, props in capas.items():
                    if nombre not in doc_dxf.layers:
                        doc_dxf.layers.new(nombre, dxfattribs={'color': props['color'], 'lineweight': props['lw']})
                    else:
                        _layer = doc_dxf.layers.get(nombre)
                        _layer.dxf.color = props['color']
                        _layer.dxf.lineweight = props['lw']

                msp = doc_dxf.modelspace()

                # Estilo texto (Usamos Arial.ttf para que el PDF lo renderice nítido y suave, no entrecortado como los SHX)
                if 'ROMANS' not in doc_dxf.styles:
                    try:    doc_dxf.styles.new('ROMANS', dxfattribs={'font':'Arial.ttf'})
                    except: doc_dxf.styles.new('ROMANS', dxfattribs={'font':'Arial.ttf'})

                # ── ESCALAS DISPONIBLES (serie normalizada ICONTEC) ─────────
                ESCALAS = [200, 100, 50, 25, 20, 10]

                # Margen y rotulo
                MARGEN   = 0.5       # cm (Restaurado a 0.5 para que el PDF no recorte las lineas superior y derecha formando una L)
                ROT_H    = 3.5       # altura rotulo ajustada para no dejar espacio vacio abajo
                ROT_W    = ANCHO_PLANO - 2 * MARGEN
                AREA_W   = ANCHO_PLANO - 2 * MARGEN
                AREA_H   = ALTO_PLANO  - 2 * MARGEN - ROT_H - 0.5  # area util dibujo

                # FACTOR DE ESCALA PARA TEXTOS Y TABLAS SEGÚN EL PAPEL
                K_DXF = min(1.0, ANCHO_PLANO / 50.0)

                # ── ZONA ALZADO (left 45% del area) ─────────────────────────
                ALZ_W = AREA_W * 0.35
                ALZ_H = AREA_H

                # Calcular escala automatica para el alzado
                escala_den = 200  # empezar por la mas pequena
                for den in reversed(ESCALAS):
                    drawn_h = (L_col / den)        # en cm de plano
                    drawn_w = ((D if es_circular else max(b, h)) / den)
                    if drawn_h <= ALZ_H - 2.5 and drawn_w <= ALZ_W - 2.5:
                        escala_den = den
                        break
                ESCALA = 1.0 / escala_den
                ESCALA_LABEL = f"1:{escala_den}"

                # Ejes del alzado
                AX0 = MARGEN + (ALZ_W - (D if es_circular else b) * ESCALA) / 2
                AY0 = MARGEN + ROT_H + 0.8
                ALZ_WDRAW = (D if es_circular else b) * ESCALA
                ALZ_HDRAW = L_col * ESCALA

                # Marco exterior
                msp.add_lwpolyline(
                    [(MARGEN, MARGEN), (ANCHO_PLANO-MARGEN, MARGEN),
                     (ANCHO_PLANO-MARGEN, ALTO_PLANO-MARGEN),
                     (MARGEN, ALTO_PLANO-MARGEN), (MARGEN, MARGEN)],
                    dxfattribs={'layer': 'MARGEN'})

                # Eje vertical central de la columna (linetype CENTER)
                cx_alz = AX0 + ALZ_WDRAW / 2
                msp.add_line(
                    (cx_alz, AY0 - 0.5),
                    (cx_alz, AY0 + ALZ_HDRAW + 0.5),
                    dxfattribs={'layer': 'EJES', 'linetype': 'DASHDOT'})

                # Contorno columna alzado
                if es_circular:
                    msp.add_lwpolyline(
                        [(AX0, AY0), (AX0 + ALZ_WDRAW, AY0),
                         (AX0 + ALZ_WDRAW, AY0 + ALZ_HDRAW),
                         (AX0, AY0 + ALZ_HDRAW), (AX0, AY0)],
                        dxfattribs={'layer': 'CONCRETO'})
                else:
                    msp.add_lwpolyline(
                        [(AX0, AY0), (AX0 + ALZ_WDRAW, AY0),
                         (AX0 + ALZ_WDRAW, AY0 + ALZ_HDRAW),
                         (AX0, AY0 + ALZ_HDRAW), (AX0, AY0)],
                        dxfattribs={'layer': 'CONCRETO'})

                # Lineas de corte en nodos
                ext = 0.5
                msp.add_line((AX0 - ext, AY0), (AX0 + ALZ_WDRAW + ext, AY0), dxfattribs={'layer': 'CONCRETO'})
                msp.add_line((AX0 - ext, AY0 + ALZ_HDRAW), (AX0 + ALZ_WDRAW + ext, AY0 + ALZ_HDRAW), dxfattribs={'layer': 'CONCRETO'})

                # Barras longitudinales (alzado) - una por cara visible
                rec_s = recub_cm * ESCALA
                db_s  = rebar_diam * ESCALA / 20  # radio barra en plano
                for xb in [AX0 + rec_s, AX0 + ALZ_WDRAW - rec_s]:
                    msp.add_line(
                        (xb, AY0),
                        (xb, AY0 + ALZ_HDRAW),
                        dxfattribs={'layer': 'ACERO_LONG', 'color': _color_acero_dxf_col(rebar_diam)})

                # Estribos (alzado) - distribucion real EXACTA (Sincronizada con APU)
                y_pos_dxf = []
                _esp_real_conf = Lo_conf / n_est_por_Lo if n_est_por_Lo > 0 else s_conf
                for i in range(n_est_por_Lo + 1):
                    y_pos_dxf.append(i * _esp_real_conf)
                if n_estribos_centro > 0:
                    _esp_real_centro = longitud_zona_libre / (n_estribos_centro + 1)
                    for i in range(1, n_estribos_centro + 1):
                        y_pos_dxf.append(Lo_conf + i * _esp_real_centro)
                if n_est_por_Lo > 0:
                    for i in range(n_est_por_Lo):
                        y_pos_dxf.append(L_col - Lo_conf + i * _esp_real_conf)

                for y_curr in y_pos_dxf:
                    in_conf = (y_curr <= Lo_conf) or (y_curr >= L_col - Lo_conf)
                    ye = AY0 + y_curr * ESCALA
                    msp.add_line(
                        (AX0 + rec_s, ye),
                        (AX0 + ALZ_WDRAW - rec_s, ye),
                        dxfattribs={'layer': 'ACERO_TRANS', 'color': _color_acero_dxf_col(stirrup_diam)})

                # Cotas alzado - L total
                cx_cota = AX0 - 1.4
                msp.add_line((cx_cota, AY0), (cx_cota, AY0 + ALZ_HDRAW), dxfattribs={'layer': 'COTAS'})
                for yy in [AY0, AY0 + ALZ_HDRAW]:
                    msp.add_line((AX0, yy), (cx_cota - 0.2, yy), dxfattribs={'layer': 'COTAS'})
                msp.add_text(f"L = {L_col:.0f} cm",
                    dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.22,
                                'insert': (cx_cota - 0.15, AY0 + ALZ_HDRAW / 2),
                                'align_point': (cx_cota - 0.15, AY0 + ALZ_HDRAW / 2),
                                'halign': 1, 'valign': 2, 'rotation': 90})

                # Cota zona confinamiento
                if Lo_conf > 0:
                    cxr = AX0 + ALZ_WDRAW + 1.2
                    for (y_ini, y_fin, txt) in [
                        (AY0, AY0 + Lo_conf * ESCALA, f"Lo={Lo_conf:.0f}cm"),
                        (AY0 + (L_col - Lo_conf) * ESCALA, AY0 + ALZ_HDRAW, f"Lo={Lo_conf:.0f}cm")
                    ]:
                        msp.add_line((cxr, y_ini), (cxr, y_fin), dxfattribs={'layer': 'COTAS'})
                        for yy in [y_ini, y_fin]:
                            msp.add_line((AX0 + ALZ_WDRAW, yy), (cxr + 0.2, yy), dxfattribs={'layer': 'COTAS'})
                        msp.add_text(txt,
                            dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF),
                                        'insert': (cxr + 0.15, (y_ini + y_fin) / 2),
                                        'align_point': (cxr + 0.15, (y_ini + y_fin) / 2),
                                        'halign': 1, 'valign': 2, 'rotation': -90})

                # Etiqueta zona conf / zona basica
                yz_mid = AY0 + Lo_conf * ESCALA + (L_col * ESCALA - 2 * Lo_conf * ESCALA) / 2
                msp.add_text(f"s={s_basico:.0f}cm",
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.22,
                                'insert': (AX0 + ALZ_WDRAW / 2, yz_mid)})

                # ── ZONA SECCION TRANSVERSAL (centro) ───────────────────────
                SEC_X0 = MARGEN + ALZ_W + 0.5
                SEC_Y0 = AY0
                SEC_AVAILABLE_W = AREA_W * 0.30
                SEC_AVAILABLE_H = AREA_H * 0.55

                dim_b = D if es_circular else b
                dim_h = D if es_circular else h
                # Escala correcta para que la seccion entre en SEC_AVAILABLE
                escala_sec = min((SEC_AVAILABLE_W - 1.0) / dim_b, (SEC_AVAILABLE_H - 2.0) / dim_h)
                escala_sec = min(escala_sec, 1.0 / 5)  # max 1:5
                escala_sec_int = max(5, int(round(1/escala_sec)))
                sec_w = dim_b * escala_sec
                sec_h = dim_h * escala_sec
                # Centrado dentro de la zona asignada a la seccion
                ox = SEC_X0 + (SEC_AVAILABLE_W - sec_w) / 2
                oy = SEC_Y0 + (SEC_AVAILABLE_H - sec_h) / 2

                # Titulo seccion
                msp.add_text("SECCION TRANSVERSAL",
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.28 * max(0.6, K_DXF),
                                'insert': (ox + sec_w / 2, oy + sec_h + 1.0),
                                'align_point': (ox + sec_w / 2, oy + sec_h + 1.0),
                                'halign': 1, 'valign': 2})
                msp.add_text(f"(ESCALA 1: {escala_sec_int})",
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.22,
                                'insert': (ox + sec_w / 2, oy + sec_h + 0.65),
                                'align_point': (ox + sec_w / 2, oy + sec_h + 0.65),
                                'halign': 1, 'valign': 2})

                if es_circular:
                    cxc = ox + sec_w / 2
                    cyc = oy + sec_h / 2
                    msp.add_circle((cxc, cyc), D / 2 * escala_sec, dxfattribs={'layer': 'CONCRETO'})
                    r_esp = (D / 2 - recub_cm) * escala_sec
                    msp.add_circle((cxc, cyc), r_esp, dxfattribs={'layer': 'ACERO_TRANS', 'color': _color_acero_dxf_col(stirrup_diam)})
                    r_bar = rebar_diam / 20 * escala_sec
                    offset_real_cm = recub_cm + (stirrup_diam / 10.0) + (rebar_diam / 20.0)
                    for ang in np.linspace(0, 2 * math.pi, n_barras, endpoint=False):
                        xb_c = cxc + (D / 2 - offset_real_cm) * escala_sec * math.cos(ang)
                        yb_c = cyc + (D / 2 - offset_real_cm) * escala_sec * math.sin(ang)
                        # HATCH sólido NSR-10 / ICONTEC (no círculo vacío)
                        _hatch_c = msp.add_hatch(color=_color_acero_dxf_col(rebar_diam), dxfattribs={'layer': 'ACERO_LONG'})
                        _hatch_c.set_solid_fill()
                        _ep_c = _hatch_c.paths.add_edge_path()
                        _ep_c.add_arc((xb_c, yb_c), r_bar, 0, 360)
                        hatch = msp.add_hatch(color=_color_acero_dxf_col(rebar_diam))
                        hatch.paths.add_edge_path().add_ellipse((xb_c, yb_c), major_axis=(r_bar, 0), ratio=1.0)
                else:
                    # Contorno seccion
                    msp.add_lwpolyline(
                        [(ox, oy), (ox + sec_w, oy), (ox + sec_w, oy + sec_h), (ox, oy + sec_h), (ox, oy)],
                        dxfattribs={'layer': 'CONCRETO'})

                    # Estribo con esquinas redondeadas y dos ganchos 135° en esquina SUP-IZQ
                    # NSR-10 C.7.1.4 / ACI 318-19 25.3.4 — gancho sísmico 135°, ext. 6db al núcleo
                    re_s  = recub_cm * escala_sec
                    r_bar = rebar_diam / 20 * escala_sec
                    offset_real_cm = recub_cm + (stirrup_diam / 10.0) + (rebar_diam / 20.0)
                    offset_s  = offset_real_cm * escala_sec
                    xm, xM = ox + re_s, ox + sec_w - re_s
                    ym, yM = oy + re_s, oy + sec_h - re_s
                    bw_e  = xM - xm    # ancho estribo escalado
                    hw_e  = yM - ym    # alto estribo escalado
                    # Radio doblez = 3db (NSR-10 C.7.2)
                    _rd = max(min(3.0*stirrup_diam/10.0*escala_sec, bw_e*0.12, hw_e*0.12), 0.03)
                    # Longitud gancho: 6db (NSR-10 C.7.1.4 sismo)
                    L_gancho = max(6.0*stirrup_diam/10.0*escala_sec, _rd*1.2)
                    _n = 8  # puntos por arco de 90°

                    def _adxf(cx, cy, rad, a0, a1):
                        return [(cx+rad*math.cos(math.radians(a0+(a1-a0)*i/_n)),
                                 cy+rad*math.sin(math.radians(a0+(a1-a0)*i/_n)))
                                for i in range(_n+1)]

                    # Perímetro: arranca en (xm, yM-_rd) = pie del arco sup-izq
                    _pe = []
                    _pe.append((xm, yM-_rd))                          # inicio lado izq
                    _pe += _adxf(xm+_rd, ym+_rd, _rd, 180, 270)      # arco inf-izq
                    _pe.append((xM-_rd, ym))                          # lado inf
                    _pe += _adxf(xM-_rd, ym+_rd, _rd, 270, 360)      # arco inf-der
                    _pe.append((xM, yM-_rd))                          # lado der
                    _pe += _adxf(xM-_rd, yM-_rd, _rd, 0,   90)       # arco sup-der
                    _pe.append((xm+_rd, yM))                          # lado sup
                    _pe += _adxf(xm+_rd, yM-_rd, _rd, 90, 180)       # arco sup-izq
                    _pe.append((xm, yM-_rd))                          # cierra
                    msp.add_lwpolyline(_pe, dxfattribs={'layer': 'ACERO_TRANS',
                                       'color': _color_acero_dxf_col(stirrup_diam)})

                    # Ganchos 135°: dirección (+X, -Y) = hacia núcleo desde esquina sup-izq
                    _dk = L_gancho * math.cos(math.radians(45))       # componente = mismo en X e Y
                    # Gancho 1: desde (xm, yM-_rd) — pie del arco sup-izq en lado IZQUIERDO
                    msp.add_line((xm, yM-_rd), (xm+_dk, yM-_rd-_dk),
                                 dxfattribs={'layer':'ACERO_TRANS','color':_color_acero_dxf_col(stirrup_diam)})
                    # Gancho 2: desde (xm+_rd, yM) — inicio del lado SUPERIOR
                    msp.add_line((xm+_rd, yM), (xm+_rd+_dk, yM-_dk),
                                 dxfattribs={'layer':'ACERO_TRANS','color':_color_acero_dxf_col(stirrup_diam)})

                    # Barras en esquinas y caras
                    cxm, cxM = ox + offset_s, ox + sec_w - offset_s
                    cym, cyM = oy + offset_s, oy + sec_h - offset_s
                    xs_bar = np.linspace(cxm, cxM, num_filas_h) if num_filas_h > 1 else [ox + sec_w / 2]
                    ys_bar = np.linspace(cym, cyM, num_filas_v) if num_filas_v > 1 else [oy + sec_h / 2]
                    for xb in xs_bar:
                        for yb in [cym, cyM]:
                            # HATCH sólido Color 7 — ICONTEC 2289 (sin círculo vacío)
                            _h1 = msp.add_hatch(color=_color_acero_dxf_col(rebar_diam), dxfattribs={'layer': 'ACERO_LONG'})
                            _h1.set_solid_fill()
                            _h1.paths.add_edge_path().add_ellipse((xb, yb), major_axis=(r_bar, 0), ratio=1.0)
                    for yb in ys_bar[1:-1]:
                        for xb in [cxm, cxM]:
                            # HATCH sólido Color 7 — ICONTEC 2289
                            _h2 = msp.add_hatch(color=_color_acero_dxf_col(rebar_diam), dxfattribs={'layer': 'ACERO_LONG'})
                            _h2.set_solid_fill()
                            _h2.paths.add_edge_path().add_ellipse((xb, yb), major_axis=(r_bar, 0), ratio=1.0)

                    # Grapas interiores con ganchos a 135° en ambos extremos
                    _n_hx = 8
                    def _arc_dxf(cx, cy, rad, a0, a1):
                        return [(cx+rad*math.cos(math.radians(a)), cy+rad*math.sin(math.radians(a)))
                                for a in np.linspace(a0, a1, _n_hx+1)]
                    
                    # ── GRAPAS DXF ANCLADAS A LAS VARILLAS REALES ──
                    _R_grapa_dxf = ((rebar_diam / 20.0) + (stirrup_diam / 20.0)) * escala_sec
                    _Cx_dxf  = (cxm + cxM) / 2.0
                    _Cy_dxf  = (cym + cyM) / 2.0
                    _xv_dxf  = (cxM - cxm) / 2.0  # Centro exacto de varillas laterales
                    _yv_dxf  = (cyM - cym) / 2.0  # Centro exacto de varillas sup/inf

                    # GRAPAS DXF ANCLADAS A CADA VARILLA INTERMEDIA (SIN CANDADOS)
                    if num_flejes_x > 0 and len(ys_bar) > 2:
                        for y_coord in ys_bar[1:-1]:
                            _yv_c = y_coord - _Cy_dxf
                            _pc = generar_puntos_grapa_x(_xv_dxf, _yv_c, _R_grapa_dxf, L_gancho)
                            _pts_gx = [(_x + _Cx_dxf, _y + _Cy_dxf) for _x, _y in _pc]
                            msp.add_lwpolyline(_pts_gx, dxfattribs={'layer': 'ACERO_TRANS', 'color': _color_acero_dxf_col(stirrup_diam)})

                    if num_flejes_y > 0 and len(xs_bar) > 2:
                        for x_coord in xs_bar[1:-1]:
                            _xv_c = x_coord - _Cx_dxf
                            _pc = generar_puntos_grapa_y(_xv_c, _yv_dxf, _R_grapa_dxf, L_gancho)
                            _pts_gy = [(_x + _Cx_dxf, _y + _Cy_dxf) for _x, _y in _pc]
                            msp.add_lwpolyline(_pts_gy, dxfattribs={'layer': 'ACERO_TRANS', 'color': _color_acero_dxf_col(stirrup_diam)})

                    # Cotas seccion
                    yc_dim = oy - 0.7
                    msp.add_line((ox, yc_dim), (ox + sec_w, yc_dim), dxfattribs={'layer': 'COTAS'})
                    for xx in [ox, ox + sec_w]:
                        msp.add_line((xx, oy), (xx, yc_dim - 0.2), dxfattribs={'layer': 'COTAS'})
                    msp.add_text(f"b = {b:.0f} cm",
                        dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.22,
                                    'insert': (ox + sec_w / 2, yc_dim - 0.25),
                                    'align_point': (ox + sec_w / 2, yc_dim - 0.25),
                                    'halign': 1, 'valign': 2})

                    xc_dim = ox + sec_w + 0.7
                    msp.add_line((xc_dim, oy), (xc_dim, oy + sec_h), dxfattribs={'layer': 'COTAS'})
                    for yy in [oy, oy + sec_h]:
                        msp.add_line((ox + sec_w, yy), (xc_dim + 0.2, yy), dxfattribs={'layer': 'COTAS'})
                    msp.add_text(f"h = {h:.0f} cm",
                        dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.22,
                                    'insert': (xc_dim + 0.15, oy + sec_h / 2),
                                    'align_point': (xc_dim + 0.15, oy + sec_h / 2),
                                    'halign': 1, 'valign': 2, 'rotation': 90})

                    # Recubrimiento con flecha
                    # Texto recubrimiento FUERA de la sección (abajo-izquierda con offset)
                    _txt_h_rec = 0.18 * max(0.5, K_DXF)
                    msp.add_text(f"rec. = {recub_cm:.0f} cm",
                        dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': _txt_h_rec,
                                    'insert': (ox + re_s, oy - _txt_h_rec * 3.0)})

                # ── TABLA DE DESPIECE (zona derecha DENTRO de Carta) ─────────────
                # Carta: 21.6cm ancho. Margen derecho: 20.6cm. Alzado+Sec ocupa ~13cm.
                TAB_X0 = SEC_X0 + SEC_AVAILABLE_W + 0.3 - 2.0  # Desplazado 2 cm a la izq
                # Clamp: si la tabla no cabe a la derecha, colocar debajo de la sección
                TAB_MAX_W = (ANCHO_PLANO - MARGEN - TAB_X0)
                if TAB_MAX_W < 4.0:  # mínimo 4cm para tabla legible
                    TAB_X0 = MARGEN
                    TAB_Y0_OVERRIDE = oy - 2.0  # debajo del alzado
                    TAB_MAX_W = ANCHO_PLANO - 2 * MARGEN
                else:
                    TAB_Y0_OVERRIDE = None
                # Columnas: Marca(1.4) Diam(1.6) Cant(0.9) L(1.0) Forma(2.0) Peso(1.1) = 8.0cm
                # Ajustar proporcionalmente si no caben
                _cws_base = [1.4, 1.6, 0.9, 1.0, 2.0, 1.1]
                _scale_w = min(1.0, TAB_MAX_W / sum(_cws_base))
                COL_WS = [w * _scale_w for w in _cws_base]
                TAB_TW = sum(COL_WS)  # Ancho total de la tabla
                ROW_H  = 0.42
                K_DXF = min(1.0, ANCHO_PLANO / 50.0)  # Restaurar escala real
                # Tabla empieza en la parte superior de la zona de dibujo
                TAB_Y0 = AY0 + AREA_H - 0.5 - 1.0  # Desplazado 1 cm abajo
                HEADERS = ["MARCA", "DIAMETRO", "CANT.", "L (m)", "FORMA", "PESO kg"]

                # Calcular datos barras y separar Grapas para pedido de Ferreteria
                if es_circular:
                    _lb  = (L_col + 2 * (ld_mm / 10) + 2 * (12 * rebar_diam / 10)) / 100
                    _pw  = n_barras * _lb * (rebar_area * 100) * 7.85e-3
                    _lt  = long_espiral_total / 100
                    _pt  = peso_total_estribos_kg
                    barras_despiece = [
                        ("L1 - Long.", _bar_label(rebar_diam), str(n_barras), f"{_lb:.2f}", "Recta + gancho sup.", f"{_pw:.1f}"),
                        ("E1 - Espiral", _bar_label(stirrup_diam), "1 esp.", f"{_lt:.2f}", "Espiral continua", f"{_pt:.1f}"),
                    ]
                else:
                    _lb  = (L_col + 2 * (ld_mm / 10) + 2 * (12 * rebar_diam / 10)) / 100
                    _pw  = n_barras_total * _lb * (rebar_area * 100) * 7.85e-3
                    _le  = (2*(b-2*recub_cm) + 2*(h-2*recub_cm) + 12*stirrup_diam/10) / 100
                    _pt_e = n_estribos_total * _le * (stirrup_area * 100) * 7.85e-3
                    barras_despiece = [
                        ("L1 - Long.",    _bar_label(rebar_diam),    str(n_barras_total),   f"{_lb:.2f}",  "Recta + gan.180",   f"{_pw:.1f}"),
                        ("E1 - Perimetral",_bar_label(stirrup_diam), str(n_estribos_total),  f"{_le:.2f}",  "Cerrado 135",       f"{_pt_e:.1f}"),
                    ]
                    # Grapas X
                    if num_flejes_x > 0:
                        _lgx = long_fleje_x / 100
                        _cgx = n_estribos_total * num_flejes_x
                        _pgx = _cgx * _lgx * (stirrup_area * 100) * 7.85e-3
                        barras_despiece.append(("GX - Grapas X", _bar_label(stirrup_diam), str(_cgx), f"{_lgx:.2f}", "Grapa 135-135", f"{_pgx:.1f}"))
                    # Grapas Y
                    if num_flejes_y > 0:
                        _lgy = long_fleje_y / 100
                        _cgy = n_estribos_total * num_flejes_y
                        _pgy = _cgy * _lgy * (stirrup_area * 100) * 7.85e-3
                        barras_despiece.append(("GY - Grapas Y", _bar_label(stirrup_diam), str(_cgy), f"{_lgy:.2f}", "Grapa 135-135", f"{_pgy:.1f}"))

                def draw_despiece_table(x0, y_top):
                    TH = 0.16  # Altura texto tabla (pequeño para no desbordarse)
                    # Titulo
                    msp.add_text("DESPIECE DE ACERO - ICONTEC 2289",
                        dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.22,
                                    'insert': (x0 + TAB_TW / 2, y_top + 0.25),
                                    'align_point': (x0 + TAB_TW / 2, y_top + 0.25),
                                    'halign': 1, 'valign': 2})
                    # Header row
                    y = y_top
                    cx = x0
                    for hdr, cw in zip(HEADERS, COL_WS):
                        msp.add_lwpolyline(
                            [(cx, y - ROW_H), (cx + cw, y - ROW_H),
                             (cx + cw, y), (cx, y), (cx, y - ROW_H)],
                            dxfattribs={'layer': 'ROTULO'})
                        msp.add_text(hdr,
                            dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': TH,
                                        'insert': (cx + cw / 2, y - ROW_H / 2),
                                        'align_point': (cx + cw / 2, y - ROW_H / 2),
                                        'halign': 1, 'valign': 2})
                        cx += cw
                    y -= ROW_H
                    # Data rows
                    for row in barras_despiece:
                        cx = x0
                        for val, cw in zip(row, COL_WS):
                            msp.add_lwpolyline(
                                [(cx, y - ROW_H), (cx + cw, y - ROW_H),
                                 (cx + cw, y), (cx, y), (cx, y - ROW_H)],
                                dxfattribs={'layer': 'ROTULO'})
                            # Truncar texto al ancho de columna (aprox 8 chars por cm)
                            max_chars = max(4, int(cw * 7))
                            txt_disp = str(val)[:max_chars]
                            msp.add_text(txt_disp,
                                dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': TH,
                                            'insert': (cx + 0.05, y - ROW_H / 2),
                                            'align_point': (cx + 0.05, y - ROW_H / 2),
                                            'halign': 0, 'valign': 2})
                            cx += cw
                        y -= ROW_H
                    # Totales
                    _tot_acero = _pw + _pt_e + sum(
                        _cgx * (long_fleje_x/100) * (stirrup_area*100) * 7.85e-3 if num_flejes_x > 0 else 0
                        for _cgx in [n_estribos_total * num_flejes_x]
                    ) + sum(
                        _cgy * (long_fleje_y/100) * (stirrup_area*100) * 7.85e-3 if num_flejes_y > 0 else 0
                        for _cgy in [n_estribos_total * num_flejes_y]
                    )
                    cx = x0
                    totales = [("TOTAL ACERO", "", "", "", "", f"{_tot_acero:.1f} kg"),
                               ("CONCRETO", f"f'c={fc:.0f}MPa", f"{vol_concreto_m3:.3f}m3", "", "", "")]
                    for tot in totales:
                        for val, cw in zip(tot, COL_WS):
                            msp.add_lwpolyline(
                                [(cx, y - ROW_H), (cx + cw, y - ROW_H),
                                 (cx + cw, y), (cx, y), (cx, y - ROW_H)],
                                dxfattribs={'layer': 'ROTULO'})
                            msp.add_text(str(val)[:max(4,int(cw*7))],
                                dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': TH,
                                            'insert': (cx + 0.05, y - ROW_H / 2),
                                            'align_point': (cx + 0.05, y - ROW_H / 2),
                                            'halign': 0, 'valign': 2})
                            cx += cw
                        cx = x0
                        y -= ROW_H
                    return y  # y final de la tabla

                y_after_table = draw_despiece_table(TAB_X0, TAB_Y0)

                # ── DIAGRAMA DE GANCHOS (DOBLECES) ──────────────────────────
                # Dibuja el esquema de cada tipo de gancho a escala 1:5
                ESCALA_DOBL = 1 / 5
                DOBL_X = TAB_X0 - 10.0  # Movido 2 cm a la derecha
                DOBL_Y = y_after_table - 1.5 + 5.0  # Subido 5 cm
                db_dobl = stirrup_diam  # mm

                # Gancho 135 grados (sismico - estribo)
                msp.add_text("Gancho 135 (estrib.)",
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF),
                                'insert': (DOBL_X, DOBL_Y)})
                # Trazo simplificado del gancho 135 (L pequeña con gancho)
                pts_g135 = [
                    (DOBL_X + 0.2, DOBL_Y - 0.5),
                    (DOBL_X + 0.2, DOBL_Y - 1.2),
                    (DOBL_X + 0.6, DOBL_Y - 0.8)
                ]
                msp.add_lwpolyline(pts_g135, dxfattribs={'layer': 'DOBLEZ'})
                msp.add_text(f"ext.={6*stirrup_diam:.0f}mm",
                    dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.18,
                                'insert': (DOBL_X + 0.8, DOBL_Y - 0.7)})

                DOBL_X2 = DOBL_X + 3.5
                # Gancho 180 (barra longitudinal)
                msp.add_text("Gancho 180 (long.)",
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF),
                                'insert': (DOBL_X2, DOBL_Y)})
                # Trazo simplificado U
                pts_g180 = [
                    (DOBL_X2 + 0.2, DOBL_Y - 0.3),
                    (DOBL_X2 + 0.2, DOBL_Y - 1.0),
                    (DOBL_X2 + 0.7, DOBL_Y - 1.0),
                    (DOBL_X2 + 0.7, DOBL_Y - 0.6)
                ]
                msp.add_lwpolyline(pts_g180, dxfattribs={'layer': 'DOBLEZ'})
                msp.add_text(f"ext.={4*rebar_diam:.0f}mm",
                    dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.18,
                                'insert': (DOBL_X2 + 0.9, DOBL_Y - 0.7)})

                if num_flejes_x > 0 or num_flejes_y > 0:
                    DOBL_X3 = DOBL_X2 + 3.5
                    msp.add_text("Grapa (ext. 135)",
                        dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF),
                                    'insert': (DOBL_X3, DOBL_Y)})
                    # Trazo simplificado grapa
                    pts_grapa = [
                        (DOBL_X3 + 0.2 + 0.4, DOBL_Y - 1.0 + 0.4),
                        (DOBL_X3 + 0.2,       DOBL_Y - 1.0),
                        (DOBL_X3 + 2.0,       DOBL_Y - 1.0),
                        (DOBL_X3 + 2.0 - 0.4, DOBL_Y - 1.0 + 0.4)
                    ]
                    msp.add_lwpolyline(pts_grapa, dxfattribs={'layer': 'DOBLEZ'})
                    msp.add_text(f"ext.={6*stirrup_diam:.0f}mm",
                        dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.18,
                                    'insert': (DOBL_X3 + 0.5, DOBL_Y - 0.7)})

                # --- NUEVA TABLA EMPALMES Y MATERIALES (DXF) ---
                Y_MAT = DOBL_Y - 2.5
                msp.add_text("TABLA DE EMPALMES (A TENSION) Ld", dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.25 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT)})
                _ld_cm = (1.3 * 0.02 * 420) / (21**0.5) * rebar_diam / 10
                msp.add_text(f"Varilla principal ({_bar_label(rebar_diam)}): Ld = {_ld_cm:.0f} cm", dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT - 0.4)})
                
                Y_MAT -= 1.2
                msp.add_text(f"DOSIFICACION m3 CONCRETO (f'c={fc:.0f}MPa)", dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.25 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT)})
                if fc >= 21:
                    _cem, _are, _gra, _agu = 350, 0.56, 0.84, 180
                    _txt_mat = f"Cemento: {_cem} kg (7 sacos) | Arena: {_are} m3 | Grava: {_gra} m3 | Agua: {_agu} L"
                    msp.add_text(_txt_mat, dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT - 0.4)})
                    
                    Y_MAT -= 0.8
                    msp.add_text(f"TOTALES PARA ESTA COLUMNA ({vol_concreto_m3:.2f} m3):", dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT)})
                    _txt_tot = f"Cemento: {_cem*vol_concreto_m3:.0f}kg | Are: {_are*vol_concreto_m3:.2f}m3 | Gra: {_gra*vol_concreto_m3:.2f}m3 | Ag: {_agu*vol_concreto_m3:.0f}L"
                    msp.add_text(_txt_tot, dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT - 0.3)})
                else:
                    msp.add_text("Requiere diseno de mezcla.", dxfattribs={'layer': 'COTAS', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF), 'insert': (DOBL_X, Y_MAT - 0.4)})


                # ── CUADRO VERIFICACIONES (debajo del dibujo alzado) ───────
                # ── CUADRO VERIFICACIONES (debajo de la seccion, zona derecha) ──
                VER_X = TAB_X0 - 3.0  # Movido 3 cm a la izquierda
                # Posicionar BAJO el cuadro de doblez/dosificacion
                VER_Y = AY0 + (AREA_H * 0.35) + 8.0  # Subido 8 cm
                ver_rows = []
                if not es_circular:
                    ver_rows = [
                        ("Seccion", f"{b:.0f} x {h:.0f} cm", 7),
                        (f"f'c / fy", f"{fc:.0f} / {fy:.0f} MPa", 7),
                        ("Armado", f"{n_barras_total} {_bar_label(rebar_diam)}", 7),
                        ("Ast / rho", f"{Ast:.2f} cm2 / {cuantia:.2f}%", 7),
                        ("Estribo", f"{_bar_label(stirrup_diam)} c/{s_conf:.0f}cm conf.", 7),
                        ("",        f"c/{s_basico:.0f}cm zona media", 7),
                        ("Ash req/prov", f"{Ash_req:.2f}/{Ash_prov:.2f} cm2", 3 if ash_ok else 1),
                        ("Ash",     "CUMPLE" if ash_ok else "NO CUMPLE", 3 if ash_ok else 1),
                        ("Biaxial Pu/Pni", f"{bresler['ratio']:.3f}", 3 if bresler['ok'] else 1),
                        ("Verificacion", "CUMPLE" if bresler['ok'] else "NO CUMPLE", 3 if bresler['ok'] else 1),
                    ]
                else:
                    ver_rows = [
                        ("Seccion", f"Circ. D={D:.0f} cm", 7),
                        (f"f'c / fy", f"{fc:.0f} / {fy:.0f} MPa", 7),
                        ("Armado", f"{n_barras} {_bar_label(rebar_diam)}", 7),
                        ("Ast / rho", f"{Ast:.2f} cm2 / {cuantia:.2f}%", 7),
                        ("Espiral", f"{_bar_label(stirrup_diam)} paso={paso_espiral:.0f}cm", 7),
                        ("rho_s req/prov", f"{rho_s_req:.4f}/{rho_s_prov:.4f}", 3 if ash_ok else 1),
                        ("Espiral", "CUMPLE" if ash_ok else "NO CUMPLE", 3 if ash_ok else 1),
                        ("Biaxial Pu/Pni", f"{bresler['ratio']:.3f}", 3 if bresler['ok'] else 1),
                        ("Verificacion", "CUMPLE" if bresler['ok'] else "NO CUMPLE", 3 if bresler['ok'] else 1),
                    ]

                msp.add_text("VERIFICACIONES NSR-10 / ACI 318",
                    dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.26,
                                'insert': (VER_X, VER_Y)})
                for i, (lbl, val, col) in enumerate(ver_rows):
                    y_row = VER_Y - 0.5 - i * 0.48
                    if lbl:
                        msp.add_text(f"{lbl}:",
                            dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.20 * max(0.6, K_DXF),
                                        'insert': (VER_X, y_row)})
                    _attr_val = {'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.20, 'insert': (VER_X + TAB_TW * 0.42, y_row)}
                    if col != 7: _attr_val['color'] = col
                    msp.add_text(val, dxfattribs=_attr_val)

                # ── ROTULO ICONTEC (inferior) ────────────────────────────────
                ROT_X = MARGEN
                ROT_Y = MARGEN
                # Marco exterior rotulo
                msp.add_lwpolyline(
                    [(ROT_X, ROT_Y), (ROT_X + ROT_W, ROT_Y),
                     (ROT_X + ROT_W, ROT_Y + ROT_H),
                     (ROT_X, ROT_Y + ROT_H), (ROT_X, ROT_Y)],
                    dxfattribs={'layer': 'ROTULO'})

                # Definicion de celdas del rotulo
                _sello_sismico = "DISIPACIÓN: " + nivel_sismico.replace('–', '-').replace('—', '-')
                celdas_rot = [
                    # (etiqueta, valor, x_rel, y_rel, cw, ch)
                    ("SISMO",    _sello_sismico, 0.0, 3.0, ROT_W, 0.5), # Sello visual superpuesto superior
                    ("EMPRESA",  dxf_empresa,  0.0,   2.0, ROT_W * 0.48, 1.0),
                    ("PROYECTO", dxf_proyecto, 0.0,   1.0, ROT_W * 0.48, 1.0),
                    ("CONTENIDO", f"Columna {'Circ.' if es_circular else 'Rect.'} - Despiece", 0.0, 0.0, ROT_W * 0.48, 1.0),
                    ("N. PLANO",  dxf_plano,   ROT_W * 0.48, 2.0, ROT_W * 0.18, 1.0),
                    ("ESCALA",   ESCALA_LABEL, ROT_W * 0.48, 1.0, ROT_W * 0.18, 1.0),
                    ("FECHA",    datetime.datetime.now().strftime("%d/%m/%Y"), ROT_W * 0.48, 0.0, ROT_W * 0.18, 1.0),
                    ("REVISION", "0",          ROT_W * 0.66, 2.0, ROT_W * 0.12, 1.0),
                    ("HOJA",     "1/1",        ROT_W * 0.66, 1.0, ROT_W * 0.12, 1.0),
                    ("PAPEL",    PAPEL_LABEL,  ROT_W * 0.66, 0.0, ROT_W * 0.12, 1.0),
                    ("ELABORO",  "",  ROT_W * 0.78, 2.0, ROT_W * 0.0733, 1.0),
                    ("REVISO",   "",   ROT_W * 0.8533, 2.0, ROT_W * 0.0733, 1.0),
                    ("APROBO",   "",   ROT_W * 0.9266, 2.0, ROT_W * 0.0734, 1.0),
                    ("ACERO kg", f"{peso_total_acero_kg:.1f}", ROT_W * 0.78, 0.0, ROT_W * 0.11, 2.0),
                    ("CONC. m3", f"{vol_concreto_m3:.3f}",    ROT_W * 0.89, 0.0, ROT_W * 0.11, 2.0),
                ]
                for etiq, valor, xr, yr, cw, ch in celdas_rot:
                    cx2 = ROT_X + xr
                    cy2 = ROT_Y + yr
                    msp.add_lwpolyline(
                        [(cx2, cy2), (cx2 + cw, cy2), (cx2 + cw, cy2 + ch),
                         (cx2, cy2 + ch), (cx2, cy2)],
                        dxfattribs={'layer': 'ROTULO'})
                    msp.add_text(etiq,
                        dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS', 'height': 0.14,
                                    'insert': (cx2 + 0.08, cy2 + ch - 0.18)})
                    msp.add_text(valor,
                        dxfattribs={'layer': 'TEXTO', 'style': 'ROMANS',
                                    'height': 0.30 if etiq in ("EMPRESA", "PROYECTO") else 0.22,
                                    'insert': (cx2 + cw / 2, cy2 + ch / 2 - 0.04),
                                    'align_point': (cx2 + cw / 2, cy2 + ch / 2 - 0.04),
                                    'halign': 1, 'valign': 2})

                # Linea divisoria sobre el rotulo
                msp.add_line(
                    (MARGEN, ROT_H + MARGEN),
                    (ANCHO_PLANO - MARGEN, ROT_H + MARGEN),
                    dxfattribs={'layer': 'MARGEN'})

                # ── EXPORTAR DXF ────────────────────────────────────────────
                import tempfile, os as _os
                with tempfile.NamedTemporaryFile(suffix='.dxf', delete=False) as tmp:
                    tmp_path = tmp.name
                doc_dxf.saveas(tmp_path)
                with open(tmp_path, 'rb') as f_tmp:
                    dxf_bytes = f_tmp.read()
                _os.unlink(tmp_path)

                nombre_dxf = f"Columna_{'Circ' if es_circular else f'{b:.0f}x{h:.0f}'}_{PAPEL_LABEL.replace(' ','_')}.dxf"
                st.download_button(
                    label=_t("Descargar DXF ICONTEC", "Download DXF ICONTEC"),
                    data=dxf_bytes,
                    file_name=nombre_dxf,
                    mime="application/dxf",
                    key="col_dxf_dl")

                # ── EXPORTAR PDF (renderizado) ───────────────────────────────
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as _mpdf
                    from ezdxf.addons.drawing import RenderContext, Frontend
                    from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
                    # ¡AQUÍ SE IMPORTA ColorPolicy!
                    from ezdxf.addons.drawing.config import Configuration, BackgroundPolicy, ColorPolicy 
                    
                    # Tamaño en pulgadas (1 cm = 0.3937 inch)
                    fig_w = ANCHO_PLANO * 0.3937
                    fig_h = ALTO_PLANO  * 0.3937
                    fig_pdf, ax_pdf = _mpdf.subplots(figsize=(fig_w, fig_h))
                    fig_pdf.patch.set_facecolor('white')
                    ax_pdf.set_facecolor('white')
                    
                    # Forzar a que los ejes ocupen EXACTAMENTE todo el tamaño de la figura (sin bordes blancos extra de matplotlib)
                    fig_pdf.subplots_adjust(left=0, right=1, bottom=0, top=1)
                    
                    _ctx     = RenderContext(doc_dxf)
                    # Es CRÍTICO adjust_figure=False para evitar que ezdxf sobrescriba el tamaño de la hoja (Carta/Oficio)
                    _backend = MatplotlibBackend(ax_pdf, adjust_figure=False)
                    
                    # PLOTEO A COLOR CON GROSORES MULTIPLICADOS PARA IMPRESIÓN
                    _config = Configuration.defaults().with_changes(
                        background_policy=BackgroundPolicy.WHITE,
                        color_policy=ColorPolicy.COLOR, # Color real (el color 7 se vuelve negro automáticamente sobre fondo blanco)
                        lineweight_scaling=2.2  # Balanceado (2.2) para que se vea claro sin empastarse
                    )
                    
                    Frontend(_ctx, _backend, config=_config).draw_layout(msp, finalize=True)
                    ax_pdf.set_aspect('equal')
                    ax_pdf.axis('off')
                    
                    # Límites reales EXACTOS de la hoja (1:1)
                    ax_pdf.set_xlim(0, ANCHO_PLANO)
                    ax_pdf.set_ylim(0, ALTO_PLANO)
                    
                    # ── Watermark / Firma digital ─────────────────────────────────
                    _wm_user  = st.session_state.get("user_email", "StructoPro")
                    _wm_fecha = __import__('datetime').datetime.now().strftime('%d/%m/%Y %H:%M')
                    _wm_text  = f"StructoPro  |  {_wm_user}  |  {_wm_fecha}  |  Generado con StructoPro — Solo para uso técnico"
                    ax_pdf.text(
                        ANCHO_PLANO / 2, ALTO_PLANO / 2, _wm_text,
                        fontsize=7, color='#c0c0c0', alpha=0.18,
                        ha='center', va='center', rotation=35,
                        fontfamily='monospace', zorder=999,
                        transform=ax_pdf.transData,
                        clip_on=False
                    )
                    # ── Fin Watermark ──────────────────────────────────────────────

                    bio_pdf_col = io.BytesIO()
                    # Resolución alta (DPI=300) y SIN bbox_inches='tight' para forzar el tamaño real del papel (Carta/Oficio)
                    fig_pdf.savefig(bio_pdf_col, format='pdf', dpi=300, facecolor='white')
                    bio_pdf_col.seek(0)
                    _mpdf.close(fig_pdf)
                    
                    nombre_pdf = nombre_dxf.replace('.dxf', '.pdf')
                    st.download_button(
                        label=_t("Descargar PDF Imprimible", "Download Printable PDF"),
                        data=bio_pdf_col.getvalue(),
                        file_name=nombre_pdf,
                        mime="application/pdf",
                        key="col_pdf_dl")
                    st.success(_t(
                        f"Plano PDF generado | Papel: {PAPEL_LABEL} | Escala: {ESCALA_LABEL} | Ploteo Monocromático HD",
                        f"PDF Plot generated | Paper: {PAPEL_LABEL} | Scale: {ESCALA_LABEL} | HD Monochrome Plot"))
                except Exception as e_pdf:
                    st.warning(f"PDF no disponible (verifique versión de ezdxf): {e_pdf}")

            except Exception as e:
                import traceback
                st.error(f"Error al generar DXF: {e}")
                st.code(traceback.format_exc(), language='python')
                st.info("Asegurate de tener instalado ezdxf: pip install ezdxf")








# =============================================================================
# =============================================================================
# TAB 3: CANTIDADES, DESPIECE Y APU
# =============================================================================
# =============================================================================
# TAB 2B: ALZADO DE CONFINAMIENTO + TABLA DE COMBINACIONES DE CARGA
# =============================================================================
with tab2b:
    col_alz, col_combo = st.columns([1, 1], gap="large")

    # ── ALZADO 2D DE CONFINAMIENTO ────────────────────────────────────────────
    with col_alz:
        st.subheader(_t(" Alzado de Confinamiento", " Confinement Elevation"))
        s_centro_alz = st.number_input(
            _t("Separación zona central s₂ (cm)", "Central zone spacing s₂ (cm)"),
            5.0, 50.0, float(min(s_conf * 2.0, 30.0)), 1.0, key="s2_alz_col")
        fig_alz = plot_alzado_confinamiento(
            L_col_cm=L_col, Lo_cm=Lo_conf, s_conf_cm=s_conf, s_centro_cm=s_centro_alz,
            b_cm=b if not es_circular else D, h_cm=h if not es_circular else D,
            es_circular=es_circular,
            rebar_diam_mm=rebar_diam, stirrup_diam_mm=stirrup_diam,
            norma_sel=norma_sel, coderef="C.7.7.1")
        st.pyplot(fig_alz); plt.close(fig_alz)
        st.caption(
            f"Lo = {Lo_conf:.0f} cm | s₁ = {s_conf:.0f} cm (confinamiento) | "
            f"s₂ = {s_centro_alz:.0f} cm (zona libre) | Ref: {norma_sel}")

    # ── TABLA DE COMBINACIONES DE CARGA ──────────────────────────────────────
    with col_combo:
        st.subheader(_t(" Combinaciones P-M", " P-M Load Combos"))
        st.caption(_t("Ingrese combinaciones manualmente o genere automaticamente desde cargas de servicio (NSR-10 A.2).",
                      "Enter load combinations manually or auto-generate from service loads (NSR-10 A.2)."))

        # ── Generador LRFD automatico NSR-10 A.2 ─────────────────────────────
        with st.expander(_t("Generador LRFD automatico — NSR-10 A.2",
                             "Automatic LRFD Generator — NSR-10 A.2"), expanded=False):
            st.caption(_t(
                "Ingrese cargas de servicio por tipo de carga. "
                "El modulo genera las combinaciones NSR-10 A.2 y las muestra en tabla "
                "para que las copie a la tabla de combinaciones principal.",
                "Enter service loads by load type. The module generates NSR-10 A.2 "
                "combinations for reference — copy values to the main combination table."
            ))
            _lc1, _lc2, _lc3 = st.columns(3)
            _lc4, _lc5, _lc6 = st.columns(3)
            _PD  = _lc1.number_input(_t("PD Carga Muerta [kN]","PD Dead Load [kN]"),
                                      value=0.0, step=50.0, key="c_pm_lrfd_PD")
            _PL  = _lc2.number_input(_t("PL Carga Viva [kN]","PL Live Load [kN]"),
                                      value=0.0, step=50.0, key="c_pm_lrfd_PL")
            _PE  = _lc3.number_input(_t("PE Sismo [kN]","PE Seismic [kN]"),
                                      value=0.0, step=50.0, key="c_pm_lrfd_PE")
            _MxD = _lc4.number_input(_t(f"MxD Muerta [{unidad_mom}]",f"MxD Dead [{unidad_mom}]"),
                                      value=0.0, step=10.0, key="c_pm_lrfd_MxD")
            _MxL = _lc5.number_input(_t(f"MxL Viva [{unidad_mom}]",f"MxL Live [{unidad_mom}]"),
                                      value=0.0, step=10.0, key="c_pm_lrfd_MxL")
            _MxE = _lc6.number_input(_t(f"MxE Sismo [{unidad_mom}]",f"MxE Seismic [{unidad_mom}]"),
                                      value=0.0, step=10.0, key="c_pm_lrfd_MxE")
            _MyD = _lc1.number_input(_t(f"MyD Muerta [{unidad_mom}]",f"MyD Dead [{unidad_mom}]"),
                                      value=0.0, step=10.0, key="c_pm_lrfd_MyD")
            _MyL = _lc2.number_input(_t(f"MyL Viva [{unidad_mom}]",f"MyL Live [{unidad_mom}]"),
                                      value=0.0, step=10.0, key="c_pm_lrfd_MyL")
            _MyE = _lc3.number_input(_t(f"MyE Sismo [{unidad_mom}]",f"MyE Seismic [{unidad_mom}]"),
                                      value=0.0, step=10.0, key="c_pm_lrfd_MyE")

            # Combinaciones NSR-10 A.2 — ec. 2.3.1.1 a 2.3.1.7
            # (alphaP_D, alphaP_L, alphaP_E, alphaMx_D, alphaMx_L, alphaMx_E,
            #  alphaMy_D, alphaMy_L, alphaMy_E)
            _NSR_COMBOS = [
                ("1.4D",            1.4,  0.0,  0.0,  1.4,  0.0,  0.0,  1.4,  0.0,  0.0),
                ("1.2D+1.6L",       1.2,  1.6,  0.0,  1.2,  1.6,  0.0,  1.2,  1.6,  0.0),
                ("1.2D+1.0L+1.0E",  1.2,  1.0,  1.0,  1.2,  1.0,  1.0,  1.2,  1.0,  1.0),
                ("1.2D+1.0L-1.0E",  1.2,  1.0, -1.0,  1.2,  1.0, -1.0,  1.2,  1.0, -1.0),
                ("0.9D+1.0E",       0.9,  0.0,  1.0,  0.9,  0.0,  1.0,  0.9,  0.0,  1.0),
                ("0.9D-1.0E",       0.9,  0.0, -1.0,  0.9,  0.0, -1.0,  0.9,  0.0, -1.0),
                ("1.2D+1.6L+0.5Lr", 1.2,  1.6,  0.0,  1.2,  1.6,  0.0,  1.2,  1.6,  0.0),
            ]

            if st.button(_t("Calcular combinaciones NSR-10 A.2",
                             "Calculate NSR-10 A.2 combinations"), key="c_pm_gen_lrfd"):
                _rows_lrfd = []
                for (_nm, _apD, _apL, _apE,
                     _amxD, _amxL, _amxE,
                     _amyD, _amyL, _amyE) in _NSR_COMBOS:
                    _Pu_c  = (_apD * _PD  + _apL * _PL  + _apE * _PE)
                    _Mux_c = (_amxD * _MxD + _amxL * _MxL + _amxE * _MxE)
                    _Muy_c = (_amyD * _MyD + _amyL * _MyL + _amyE * _MyE)
                    _rows_lrfd.append({
                        _t("Combinacion","Combination"): _nm,
                        f"Pu [{unidad_fuerza}]":  round(_Pu_c * factor_fuerza, 1),
                        f"Mux [{unidad_mom}]": round(abs(_Mux_c), 2),
                        f"Muy [{unidad_mom}]": round(abs(_Muy_c), 2),
                    })
                import pandas as _pd_lrfd
                _df_lrfd = _pd_lrfd.DataFrame(_rows_lrfd)
                # Resaltar fila critica (max |Pu| + |Mux| combinado)
                _crit_idx = (_df_lrfd[f"Pu [{unidad_fuerza}]"].abs() +
                             _df_lrfd[f"Mux [{unidad_mom}]"].abs()).idxmax()
                st.dataframe(_df_lrfd, use_container_width=True, hide_index=True)
                st.info(_t(
                    f"Combinacion critica sugerida (mayor Pu + Mux): "
                    f"**{_df_lrfd.iloc[_crit_idx][_t('Combinacion','Combination')]}** "
                    f"con Pu = {_df_lrfd.iloc[_crit_idx][f'Pu [{unidad_fuerza}]']:.1f} {unidad_fuerza} "
                    f"y Mux = {_df_lrfd.iloc[_crit_idx][f'Mux [{unidad_mom}]']:.2f} {unidad_mom}.",
                    f"Suggested critical combination (max Pu + Mux): "
                    f"**{_df_lrfd.iloc[_crit_idx][_t('Combinacion','Combination')]}** "
                    f"Pu = {_df_lrfd.iloc[_crit_idx][f'Pu [{unidad_fuerza}]']:.1f} {unidad_fuerza} "
                    f"Mux = {_df_lrfd.iloc[_crit_idx][f'Mux [{unidad_mom}]']:.2f} {unidad_mom}."
                ))
                # Exportar CSV de combinaciones
                _csv_lrfd = _df_lrfd.to_csv(index=False).encode("utf-8")
                st.download_button(
                    _t("Descargar combinaciones CSV","Download combinations CSV"),
                    data=_csv_lrfd,
                    file_name="combinaciones_LRFD_NSR10.csv",
                    mime="text/csv",
                    key="c_pm_dl_lrfd_csv"
                )
        # ── Fin Generador LRFD ────────────────────────────────────────────────
        _combos_def = pd.DataFrame({
            "Combinacion": ["1.4D", "1.2D+1.6L", "1.2D+1.0L+1.0E", "0.9D+1.0E", "1.2D+1.6L+0.5Lr"],
            "Pu":  [round(Pu_input*0.70,1), round(Pu_input,1), round(Pu_input*0.85,1), round(Pu_input*0.50,1), round(Pu_input*0.95,1)],
            "Mux": [round(abs(Mux_input)*0.50,1), round(abs(Mux_input),1), round(abs(Mux_input)*1.20,1), round(abs(Mux_input)*0.80,1), round(abs(Mux_input)*0.90,1)],
            "Muy": [round(abs(Muy_input)*0.50,1), round(abs(Muy_input),1), round(abs(Muy_input)*1.20,1), round(abs(Muy_input)*0.80,1), round(abs(Muy_input)*0.90,1)],
        })
        df_combos = st.data_editor(
            _combos_def, num_rows="dynamic", use_container_width=True, key="col_combos_pm",
            column_config={
                "Combinacion": st.column_config.TextColumn("Combinación", width="medium"),
                "Pu":  st.column_config.NumberColumn(f"Pu [{unidad_fuerza}]",  min_value=0),
                "Mux": st.column_config.NumberColumn(f"Mux [{unidad_mom}]", min_value=0),
                "Muy": st.column_config.NumberColumn(f"Muy [{unidad_mom}]", min_value=0),
            })
        puntos_combo = parse_combo_table(df_combos, factor_fuerza)
        if puntos_combo and len(cap_x.get('phi_Mn', [])) > 0:
            fig_pm_c = plot_pm_combos(cap_x, Pu_col, Mux_col, puntos_combo,
                                      factor_fuerza, unidad_fuerza, unidad_mom)
            st.pyplot(fig_pm_c); plt.close(fig_pm_c)
            # ── Superficie 3D interactiva con TODOS los combos ───────────
            with st.expander(_t(' Superficie 3D — Todos los combos',
                                ' 3D Surface — All load combinations'), expanded=True):
                fig_3d_combos = plot_pm_3d(df_cap_3d, df_combos, factor_fuerza,
                                           unidad_fuerza=unidad_fuerza, unidad_mom=unidad_mom)
                st.plotly_chart(fig_3d_combos, use_container_width=True)
                st.caption(_t(' Dentro   Fuera de la superficie φ',
                              ' Inside   Outside design surface φ'))
            # Tabla de estado por combinación
            try:
                _sp3=np.argsort(cap_x['phi_Pn']); _pPn3=np.array(cap_x['phi_Pn'])[_sp3]; _pMn3=np.array(cap_x['phi_Mn'])[_sp3]
                rows_ok = []
                for (pu, mux, muy, lbl) in puntos_combo:
                    phi_at = float(np.interp(pu,_pPn3,_pMn3))
                    dentro = abs(mux) <= phi_at
                    rows_ok.append({"Combo": lbl,
                                    f"Pu [{unidad_fuerza}]": f"{pu*factor_fuerza:.1f}",
                                    f"Mux [{unidad_mom}]": f"{abs(mux)*factor_fuerza:.1f}",
                                    "Estado": "OK Dentro" if dentro else "FALLO Fuera"})
                st.dataframe(pd.DataFrame(rows_ok), use_container_width=True, hide_index=True)
                ratios = [(abs(mux)/float(_fi(pu)) if float(_fi(pu))>0 else 999, lbl) for pu,mux,muy,lbl in puntos_combo]
                critica = max(ratios, key=lambda r: r[0])
                msg = f"OK Todas dentro — combo crítica: `{critica[1]}` (M/φM={critica[0]:.3f})" if critica[0]<=1 else f"AVISO Combo fuera del diagrama: `{critica[1]}` (M/φM={critica[0]:.3f})"
                (st.success if critica[0]<=1 else st.warning)(msg)
            except Exception:
                pass
        else:
            st.info(_t("Complete la tabla para ver los puntos en el diagrama P-M.",
                       "Fill the table to plot points on the P-M diagram."))

with tab3:
    st.subheader(f"Cantidades de Materiales — {'Circular'if es_circular else 'Rectangular'}, L={L_col:.0f} cm")
    col_c1, col_c2, col_c3, col_c4 = st.columns(4)
    col_c1.metric(_t("Concreto", "Concrete"), f"{vol_concreto_m3:.4f} m³")
    col_c2.metric(_t("Acero Total", "Total Steel"), f"{peso_total_acero_kg:.2f} kg")
    col_c3.metric(_t("Acero Longitudinal", "Long. Steel"), f"{peso_acero_long_kg:.2f} kg")
    col_c4.metric(_t("Acero Estribos", "Tie Steel"), f"{peso_total_estribos_kg:.2f} kg")
    # P-3 Encofrado / P-1 Agua / P-4 Cuantía
    _dim_b = D if es_circular else b
    _dim_h = D if es_circular else h
    _area_enc = (3.14159 * _dim_b / 100) * (L_col / 100) if es_circular else (2 * (_dim_b + _dim_h) / 100) * (L_col / 100)
    _mix_c = get_mix_for_fc(fc)
    _litros_agua = _mix_c.get("agua", 180) * vol_concreto_m3
    _col_e1, _col_e2, _col_e3 = st.columns(3)
    _col_e1.metric(_t("Encofrado", "Formwork"), f"{_area_enc:.2f} m²")
    _col_e2.metric(_t("Agua concreto", "Mixing water"), f"{_litros_agua:.0f} L")
    _col_e3.metric(_t("Cuantía ρ", "Steel ratio ρ"), f"{cuantia:.2f}%")
    st.markdown("---")
    st.subheader(_t("Despiece de Acero", "Bar Bending Schedule"))
    
    if es_circular:
        long_bar = (L_col + 2 * (ld_mm/10) + 2 * (12*rebar_diam/10)) / 100
        peso_long = n_barras * long_bar * (rebar_area * 100) * 7.85e-3
        
        despiece_data = {
            "Marca": ["L1", "E1 (Espiral)"],
            "Cantidad": [n_barras, 1],
            "Diámetro": [_bar_label(rebar_diam), _bar_label(stirrup_diam)],
            "Longitud (m)": [long_bar, long_espiral_total/100],
            "Longitud Total (m)": [n_barras * long_bar, long_espiral_total/100],
            "Peso (kg)": [peso_long, peso_total_estribos_kg]
        }
        long_bar_m = long_bar
    else:
        _tiene_empalme = st.checkbox(_t("Incluir empalme inferior (arranque desde cimentación)", "Include lap splice at base"), value=False, key="col_empalme")
        _long_empalme = (1.3 * ld_mm/10) / 100 if _tiene_empalme else 0.0
        long_bar_m = (L_col + 2 * (ld_mm/10) + 2 * (12*rebar_diam/10)) / 100
        long_bar_arranque = long_bar_m + _long_empalme
        peso_long_total = n_barras_total * long_bar_m * (rebar_area * 100) * 7.85e-3
        _marcas = ["L1"]
        _cants = [n_barras_total]
        _diams = [_bar_label(rebar_diam)]
        _longs = [long_bar_m]
        _longs_t = [n_barras_total * long_bar_m]
        _pesos = [peso_long_total]
        
        if _tiene_empalme:
            _marcas = ["L1 (cuerpo)", "L1A (arranque)"]
            _cants = [n_barras_total, n_barras_total]
            _diams = [_bar_label(rebar_diam), _bar_label(rebar_diam)]
            _longs = [long_bar_m, long_bar_arranque]
            _longs_t = [n_barras_total * long_bar_m, n_barras_total * long_bar_arranque]
            peso_arranque = n_barras_total * long_bar_arranque * (rebar_area * 100) * 7.85e-3
            _pesos = [peso_long_total, peso_arranque]

        _marcas.append("E1 Perimetral")
        _cants.append(n_estribos_total)
        _diams.append(_bar_label(stirrup_diam))
        _longs.append(perim_estribo/100)
        _longs_t.append(n_estribos_total * perim_estribo/100)
        _pesos.append(peso_total_estribos_kg)

        if num_flejes_x > 0:
            _marcas.append("GX - Grapas X")
            _cants.append(n_estribos_total * num_flejes_x)
            _diams.append(_bar_label(stirrup_diam))
            _longs.append(long_fleje_x / 100)
            _longs_t.append((n_estribos_total * num_flejes_x) * (long_fleje_x / 100))
            _pesos.append((n_estribos_total * num_flejes_x) * (long_fleje_x / 100) * (stirrup_area * 100) * 7.85e-3)

        if num_flejes_y > 0:
            _marcas.append("GY - Grapas Y")
            _cants.append(n_estribos_total * num_flejes_y)
            _diams.append(_bar_label(stirrup_diam))
            _longs.append(long_fleje_y / 100)
            _longs_t.append((n_estribos_total * num_flejes_y) * (long_fleje_y / 100))
            _pesos.append((n_estribos_total * num_flejes_y) * (long_fleje_y / 100) * (stirrup_area * 100) * 7.85e-3)

        despiece_data = {
            "Marca": _marcas,
            "Cantidad": _cants,
            "Diámetro": _diams,
            "Longitud (m)": _longs,
            "Longitud Total (m)": _longs_t,
            "Peso (kg)": _pesos
        }
    
    df_despiece = pd.DataFrame(despiece_data)
    st.dataframe(df_despiece.style.format({"Longitud (m)": "{:.2f}", "Longitud Total (m)": "{:.2f}", "Peso (kg)": "{:.1f}"}), 
                 use_container_width=True)
                 
    import io
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 750, "PEDIDO DE FERRETERIA - ACERO DE REFUERZO")
        c.setFont("Helvetica", 12)
        c.drawString(50, 730, "Proyecto: StructoPro - Modulo de Columnas")
        y = 680
        c.setFont("Helvetica-Bold", 10)
        headers = ["Marca", "Cantidad", "Diam", "Long.(m)", "L.Tot(m)", "Peso(kg)"]
        x_pos = [50, 180, 250, 320, 400, 480]
        for i, hdr in enumerate(headers):
            c.drawString(x_pos[i], y, hdr)
        y -= 20
        c.setFont("Helvetica", 10)
        for idx, row in df_despiece.iterrows():
            c.drawString(x_pos[0], y, str(row['Marca']))
            c.drawString(x_pos[1], y, str(row['Cantidad']))
            c.drawString(x_pos[2], y, str(row['Diámetro']))
            c.drawString(x_pos[3], y, f"{row['Longitud (m)']:.2f}")
            c.drawString(x_pos[4], y, f"{row['Longitud Total (m)']:.2f}")
            c.drawString(x_pos[5], y, f"{row['Peso (kg)']:.2f}")
            y -= 20
        c.line(50, y+10, 540, y+10)
        y -= 20
        c.setFont("Helvetica-Bold", 12)
        c.drawString(400, y, f"TOTAL: {df_despiece['Peso (kg)'].sum():.2f} kg")
        c.save()
        pdf_bytes = buffer.getvalue()
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.download_button(
            label=" Descargar Resumen para Ferretería (PDF)",
            data=pdf_bytes,
            file_name="pedido_ferreteria_columnas.pdf",
            mime="application/pdf",
            type="primary"
        )
    except ImportError:
        pass
    

    fig_bars, ax_bars = plt.subplots(figsize=(6, 4))
    fig_bars.patch.set_facecolor('#1e1e2e')
    for _ax in fig_bars.get_axes(): _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    _palette_map = {'L1A': '#ffa07a', 'L1': '#ff6b35', 'E1': '#00d4ff',
                    'GX': '#4caf50', 'GY': '#81c784', 'Espiral': '#a78bfa'}
    _colors_bar = []
    for _mk in df_despiece['Marca']:
        _c = '#b0b0b0'
        for _k, _v in _palette_map.items():
            if str(_mk).startswith(_k): _c = _v; break
        _colors_bar.append(_c)
    _n_items = len(df_despiece)
    fig_bars.set_size_inches(max(5, _n_items * 1.5), 4)
    _brs = ax_bars.bar(df_despiece['Marca'], df_despiece['Peso (kg)'],
                       color=_colors_bar, edgecolor='#555', linewidth=0.8)
    for _b in _brs:
        _hh = _b.get_height()
        ax_bars.text(_b.get_x() + _b.get_width()/2., _hh * 1.01,
                     f'{_hh:.1f} kg', ha='center', va='bottom',
                     fontsize=8, color='#cdd6f4', fontweight='bold')
    ax_bars.set_xlabel(_t('Elemento', 'Element'), color='#cdd6f4')
    ax_bars.set_ylabel(_t('Peso (kg)', 'Weight (kg)'), color='#cdd6f4')
    ax_bars.set_title(_t('Distribución de Pesos de Acero', 'Steel Weight Distribution'),
                      color='white', fontsize=11, fontweight='bold')
    ax_bars.tick_params(axis='x', rotation=30, colors='#cdd6f4', labelsize=8)
    ax_bars.tick_params(axis='y', colors='#cdd6f4')
    ax_bars.set_facecolor('#14142a')
    ax_bars.grid(True, axis='y', alpha=0.3, color='#555')
    fig_bars.patch.set_facecolor('#1e1e2e')
    fig_bars.tight_layout()
    st.pyplot(fig_bars)
    plt.close(fig_bars)
    
    with st.expander(_t("Dibujo de Figurado para Taller", "Shop Drawing Details"), expanded=False):
        st.markdown(_t("Formas reales de las barras con ganchos y dimensiones.", "Actual bar shapes with hooks and dimensions."))
        hook_len_cm = 12 * rebar_diam / 10
        if es_circular:
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1 = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            st.pyplot(fig_l1)
            plt.close(fig_l1)
            fig_spiral = draw_spiral(D, paso_espiral, stirrup_diam, _bar_label(stirrup_diam))
            st.pyplot(fig_spiral)
            plt.close(fig_spiral)
        else:
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1 = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            st.pyplot(fig_l1)
            plt.close(fig_l1)
            inside_b = b - 2 * recub_cm
            inside_h = h - 2 * recub_cm
            hook_len_est = 12 * stirrup_diam / 10
            # BUG-03 FIX — coordenadas barras con variables correctas e indentación fija
            # Coordenadas reales de barras para el figurado — usa num_filas_h / num_filas_v
            # directamente (no la estimación proporcional que causaba posiciones incorrectas)
            _rv_cm = recub_cm + stirrup_diam / 10.0 + rebar_diam / 20.0
            _bx2c  = b / 2.0 - _rv_cm       # semiancho al centro de varilla extrema
            _hy2c  = h / 2.0 - _rv_cm       # semialto  al centro de varilla extrema
            # Usar num_filas_h (barras sup/inf) y num_filas_v (barras laterales) reales
            _nxb   = max(2, num_filas_h)     # barras en caras sup e inf
            _nyb   = max(2, num_filas_v)     # barras en caras laterales
            # Generar posiciones en sistema centrado
            _xs_c  = [(-_bx2c + 2*_bx2c*i/(_nxb-1)) for i in range(_nxb)] if _nxb > 1 else [0.0]
            _ys_c  = [(-_hy2c + 2*_hy2c*i/(_nyb-1)) for i in range(_nyb)] if _nyb > 1 else [0.0]
            # Cara inferior + cara superior + caras laterales (sin repetir esquinas)
            _bcoords_raw = (
                [(x, -_hy2c) for x in _xs_c] +   # cara inferior
                [(x,  _hy2c) for x in _xs_c] +   # cara superior
                [(-_bx2c, y) for y in _ys_c[1:-1]] +  # cara izq (sin esquinas)
                [( _bx2c, y) for y in _ys_c[1:-1]]    # cara der (sin esquinas)
            )
            # Transformar sistema centrado → origen núcleo (0..inside_b, 0..inside_h)
            # que es el sistema de coordenadas de draw_stirrup_with_ties
            _bcoords_draw = [
                (round(x + inside_b / 2.0, 4), round(y + inside_h / 2.0, 4))
                for (x, y) in _bcoords_raw
            ] if _bcoords_raw else None
            # Eliminar duplicados (esquinas contadas dos veces) preservando orden
            if _bcoords_draw:
                _seen = set()
                _bcoords_draw = [
                    p for p in _bcoords_draw
                    if not (p in _seen or _seen.add(p))
                ]
            fig_e1 = draw_stirrup_with_ties(
                b, h, recub_cm, hook_len_est,
                bar_diam_mm=stirrup_diam,
                n_ties_x=num_flejes_x if not es_circular else 0,
                n_ties_y=num_flejes_y if not es_circular else 0,
                nivel_sismico_str=nivel_sismico,
                bar_name=_bar_label(stirrup_diam),
                bar_coords=_bcoords_draw,
                long_bar_diam_mm=rebar_diam)
            st.pyplot(fig_e1)
            plt.close(fig_e1)

    with st.expander(_t("Presupuesto APU", "APU Budget"), expanded=False):
        st.markdown(_t("Ingrese precios unitarios para calcular el costo total.", "Enter unit prices to calculate total cost."))
        # ── Jornal actualizado desde SMLMV oficial (calcular_jornales_todos) ──
        try:
            import sys as _sys2
            if "utils" not in _sys2.path: _sys2.path.append("utils")
            from utils.smlmv_colombia import calcular_jornales_todos, obtener_smlmv, CUADRILLAS_ESTANDAR
            _jornales   = calcular_jornales_todos()
            _smlmv_info = obtener_smlmv()
            _smlmv_val  = _smlmv_info["valor"]
            _smlmv_src  = _smlmv_info["fuente"]
            # Opciones de cargo para selector
            _cargos_ops = {
                "ayudante":         f"Ayudante  — ${_jornales['ayudante']['jornal_con_prestaciones']:,.0f}/día",
                "oficial":          f"Oficial   — ${_jornales['oficial']['jornal_con_prestaciones']:,.0f}/día",
                "maestro":          f"Maestro   — ${_jornales['maestro']['jornal_con_prestaciones']:,.0f}/día",
                "cuadrilla_vaciado": f"Cuadrilla vaciado (1M+2Of+4Ay) — precio por m³ MO"
            }
            _jornal_default = _jornales["oficial"]["jornal_con_prestaciones"]
            _smlmv_ok = True
        except Exception as _e_smlmv:
            _jornales, _smlmv_val, _smlmv_src = {}, 0, "sin conexión"
            _cargos_ops = {"oficial": "Oficial (manual)"}
            _jornal_default = 0.0   # Prompt V11: nunca hardcodear precios
            st.info(_t(
                'Configure el jornal en APU Mercado segun su pais y moneda. '
                'El costo de mano de obra se calculara con valor 0 hasta que lo configure.',
                'Set the daily wage in APU Market for your country and currency. '
                'Labor cost will be 0 until configured.'))
            _smlmv_ok = False

        if _smlmv_ok:
            st.caption(f" SMLMV {_smlmv_info['anio']}: **${_smlmv_val:,.0f}** — Fuente: *{_smlmv_src}*")

        with st.form(key="col_apu_form"):
            if "col_apu_moneda" not in st.session_state: st.session_state["col_apu_moneda"] = "COP"
            moneda = st.text_input(_t("Moneda", "Currency"), value=st.session_state["col_apu_moneda"])
            col_apu1, col_apu2 = st.columns(2)
            with col_apu1:
                if "col_apu_cemento" not in st.session_state: st.session_state["col_apu_cemento"] = 28000.0
                if "col_apu_acero" not in st.session_state: st.session_state["col_apu_acero"] = 4500.0
                if "col_apu_arena" not in st.session_state: st.session_state["col_apu_arena"] = 120000.0
                if "col_apu_grava" not in st.session_state: st.session_state["col_apu_grava"] = 130000.0
                precio_cemento = st.number_input(_t("Precio por bulto cemento", "Price per cement bag"), value=st.session_state["col_apu_cemento"], step=1000.0)
                precio_acero = st.number_input(_t("Precio por kg acero", "Price per kg steel"), value=st.session_state["col_apu_acero"], step=100.0)
                precio_arena = st.number_input(_t("Precio por m³ arena", "Price per m³ sand"), value=st.session_state["col_apu_arena"], step=5000.0)
                precio_grava = st.number_input(_t("Precio por m³ grava", "Price per m³ gravel"), value=st.session_state["col_apu_grava"], step=5000.0)
                if "col_apu_agua" not in st.session_state: st.session_state["col_apu_agua"] = 3500.0
                if "col_apu_encofrado" not in st.session_state: st.session_state["col_apu_encofrado"] = 45000.0
                precio_agua = st.number_input(_t("Precio agua (m³)", "Water price /m³"), value=st.session_state["col_apu_agua"], step=500.0)
                precio_encofrado = st.number_input(_t("Precio encofrado (m²)", "Formwork price /m²"), value=st.session_state["col_apu_encofrado"], step=2000.0)
            with col_apu2:
                # Selector de tipo de cargo MO
                _cargo_sel = st.selectbox(
                    _t("Tipo de mano de obra", "Labor type"),
                    options=list(_cargos_ops.keys()),
                    format_func=lambda k: _cargos_ops[k],
                    index=list(_cargos_ops.keys()).index("oficial") if "oficial" in _cargos_ops else 0,
                    key="col_apu_cargo_sel"
                )
                # Jornal auto desde SMLMV según cargo seleccionado
                if _smlmv_ok and _cargo_sel in _jornales:
                    _jornal_auto = _jornales[_cargo_sel]["jornal_con_prestaciones"]
                else:
                    _jornal_auto = _jornal_default
                if "col_apu_mo" not in st.session_state:
                    st.session_state["col_apu_mo"] = _jornal_auto
                precio_mo = st.number_input(
                    _t("Costo mano de obra (día) — SMLMV auto", "Labor cost per day — SMLMV auto"),
                    value=float(_jornal_auto),
                    step=1000.0,
                    help=_t(f"Calculado: SMLMV × escala × factor prest. / 26 días  |  Fuente: {_smlmv_src}",
                            f"Calculated from SMLMV scale × prest. factor / 26 days  |  Source: {_smlmv_src}")
                )
                if "col_apu_aui" not in st.session_state: st.session_state["col_apu_aui"] = 30.0
                pct_aui = st.number_input(_t("% A.I.U.", "% A.I.U."), value=st.session_state["col_apu_aui"], step=5.0) / 100.0
            st.markdown("---")
            usar_premezclado = st.checkbox(
                _t("Usar concreto premezclado (omite cemento/arena/grava)", "Use ready-mix concrete (skips cement/sand/gravel)"),
                value=st.session_state.get("col_apu_premix", False), key="col_apu_premix"
            )
            precio_premix_m3 = 0.0
            if usar_premezclado:
                precio_premix_m3 = st.number_input(
                    _t("Precio concreto premezclado / m³", "Ready-mix concrete price / m³"),
                    value=st.session_state.get("col_apu_premix_p", 500000.0),
                    step=10000.0, key="col_apu_premix_p"
                )
            submitted = st.form_submit_button(_t("Calcular Presupuesto", "Calculate Budget"))
            if submitted:
                st.session_state.apu_config = {"moneda": moneda, "cemento": precio_cemento, "acero": precio_acero,
                    "arena": precio_arena, "grava": precio_grava, "costo_dia_mo": precio_mo, "pct_aui": pct_aui,
                    "premix": usar_premezclado, "precio_premix_m3": precio_premix_m3,
                    "agua": precio_agua, "encofrado": precio_encofrado}
                st.success(_t("Precios guardados.", "Prices saved."))
                st.rerun()
        if "col_apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mix = get_mix_for_fc(fc)
            bag_kg = CEMENT_BAGS.get(norma_sel, CEMENT_BAGS["NSR-10 (Colombia)"])[0]["kg"]
            bultos_col = vol_concreto_m3 * mix["cem"] / bag_kg
            _usar_premix = apu.get("premix", False)
            _precio_premix_m3 = apu.get("precio_premix_m3", 0.0)
            if _usar_premix:
                costo_cemento = 0.0
                costo_arena   = 0.0
                costo_grava   = 0.0
                costo_conc_premix = vol_concreto_m3 * _precio_premix_m3
            else:
                costo_cemento = bultos_col * apu["cemento"]
                costo_arena   = (mix["arena"] * vol_concreto_m3 / 1500) * apu["arena"]
                costo_grava   = (mix["grava"] * vol_concreto_m3 / 1600) * apu["grava"]
                costo_conc_premix = 0.0
            costo_acero = peso_total_acero_kg * apu["acero"]
            costo_mo = (peso_total_acero_kg * 0.04 + vol_concreto_m3 * 0.4) * apu["costo_dia_mo"]
            _mix_apu = get_mix_for_fc(fc)
            _litros_apu = _mix_apu.get("agua", 180) * vol_concreto_m3
            costo_agua = (_litros_apu / 1000) * apu.get("agua", 3500)
            _dim_enc_b = D if es_circular else b
            _dim_enc_h = D if es_circular else h
            _area_enc_apu = (3.14159 * _dim_enc_b / 100) * (L_col / 100) if es_circular else (2 * (_dim_enc_b + _dim_enc_h) / 100) * (L_col / 100)
            costo_encofrado = _area_enc_apu * apu.get("encofrado", 45000)
            costo_directo = costo_cemento + costo_acero + costo_arena + costo_grava + costo_conc_premix + costo_mo + costo_agua + costo_encofrado
            aiu = costo_directo * apu["pct_aui"]
            total = costo_directo + aiu
            _c1, _c2, _c3 = st.columns(3)
            _c1.metric(_t(" Total Proyecto", " Total Project"), f"{total:,.0f} {apu['moneda']}")
            _c2.metric(_t(" Costo Directo", "Direct Cost"), f"{costo_directo:,.0f} {apu['moneda']}")
            _c3.metric(_t(" Mano de Obra", "Labor"), f"{costo_mo:,.0f} {apu['moneda']}")
            import plotly.graph_objects as _go
            _items_label = (
                [_t("Concreto PM", "Ready-mix"), _t("Acero", "Steel"), _t("M.O.", "Labor"), _t("Agua","Water"), _t("Encofrado","Formwork"), "A.I.U."]
                if _usar_premix else
                [_t("Cemento", "Cement"), _t("Acero", "Steel"), _t("Arena", "Sand"),
                 _t("Grava", "Gravel"), _t("M.O.", "Labor"), _t("Agua","Water"), _t("Encofrado","Formwork"), "A.I.U."]
            )
            _items_val = (
                [costo_conc_premix, costo_acero, costo_mo, costo_agua, costo_encofrado, aiu]
                if _usar_premix else
                [costo_cemento, costo_acero, costo_arena, costo_grava, costo_mo, costo_agua, costo_encofrado, aiu]
            )
            _colors = ["#3fb950", "#79c0ff", "#ffa657", "#d2a8ff", "#58a6ff", "#f0883e"][:len(_items_label)]
            _fig_apu = _go.Figure(_go.Bar(
                x=_items_label, y=_items_val,
                marker_color=_colors,
                text=[f"{v:,.0f}" for v in _items_val],
                textposition='outside',
                textfont=dict(size=11, color='white')
            ))
            _fig_apu.update_layout(
                title=_t("Desglose de Costos — Columna", "Cost Breakdown — Column"),
                paper_bgcolor='#161b22', plot_bgcolor='#0d1117',
                font=dict(color='#cdd6f4'),
                xaxis=dict(gridcolor='#30363d'),
                yaxis=dict(gridcolor='#30363d', tickformat=',.0f'),
                margin=dict(t=40, b=20, l=20, r=20),
                height=320
            )
            st.plotly_chart(_fig_apu, use_container_width=True)
            if _usar_premix:
                st.info(_t(
                    f" Concreto premezclado: {vol_concreto_m3:.3f} m³ × {_precio_premix_m3:,.0f} {apu['moneda']}/m³ = {costo_conc_premix:,.0f} {apu['moneda']}",
                    f" Ready-mix concrete: {vol_concreto_m3:.3f} m³ × {_precio_premix_m3:,.0f} {apu['moneda']}/m³ = {costo_conc_premix:,.0f} {apu['moneda']}"
                ))
            output_excel = io.BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                if _usar_premix:
                    df_apu = pd.DataFrame({
                        "Item":     ["Concreto Premezclado", "Acero", "Mano de Obra", "Agua", "Encofrado", "A.I.U.", "TOTAL"],
                        "Cantidad": [vol_concreto_m3, peso_total_acero_kg,
                                     peso_total_acero_kg * 0.04 + vol_concreto_m3 * 0.4,
                                     round(_litros_apu, 0), round(_area_enc_apu, 2), "", ""],
                        "Unidad":   ["m³", "kg", "días", "litros", "m²", f"{apu['pct_aui']*100:.0f}%", ""],
                        "Subtotal": [costo_conc_premix, costo_acero, costo_mo, costo_agua, costo_encofrado, aiu, total]
                    })
                else:
                    df_apu = pd.DataFrame({
                        "Item":     ["Cemento", "Acero", "Arena", "Grava", "Mano de Obra", "Agua", "Encofrado", "A.I.U.", "TOTAL"],
                        "Cantidad": [bultos_col, peso_total_acero_kg, mix["arena"]*vol_concreto_m3/1500,
                                     mix["grava"]*vol_concreto_m3/1600,
                                     peso_total_acero_kg*0.04 + vol_concreto_m3*0.4,
                                     round(_litros_apu, 0), round(_area_enc_apu, 2), "", ""],
                        "Unidad":   [f"bultos ({bag_kg}kg)", "kg", "m³", "m³", "días", "litros", "m²",
                                     f"{apu['pct_aui']*100:.0f}%", ""],
                        "Subtotal": [costo_cemento, costo_acero, costo_arena, costo_grava, costo_mo, costo_agua, costo_encofrado, aiu, total]
                    })
                df_apu.to_excel(writer, index=False, sheet_name='APU')
                workbook = writer.book
                worksheet = writer.sheets['APU']
                money_fmt = workbook.add_format({'num_format': '#,##0.00'})
                worksheet.set_column('D:D', 15, money_fmt)
            output_excel.seek(0)
            st.download_button(_t("Descargar Presupuesto Excel", "Download Budget Excel"),
                               data=output_excel, file_name=f"APU_Columna_{b:.0f}x{h:.0f}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# =============================================================================
# TAB 4: MEMORIA DE CÁLCULO COMPLETA
# =============================================================================
with tab4:
    st.subheader(_t("Generar Memoria de Cálculo Completa", "Generate Complete Calculation Report"))
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        btn_docx_col = st.button(_t("Generar Memoria DOCX", "Generate DOCX Report"), type="primary")
    with col_d2:
        try:
            import numpy as _np_ifc
            if tuple(int(x) for x in _np_ifc.__version__.split(".")[:2]) >= (2,0):
                st.warning("IFC requiere numpy<2.0 para compatibilidad. Instale: pip install 'numpy<2.0'")
            import ifcopenshell
            import ifcopenshell.api
            import ifcopenshell.guid
            import uuid, time

            def _make_ifc_columna(b_cm, h_cm, D_cm, L_cm, es_circ,
                                  fc_mpa, fy_mpa, n_bars, db_mm, dst_mm,
                                  n_long_total, As_cm2, rho_pct, recub,
                                  s_conf_cm, s_bas_cm, Lo_cm,
                                  n_estrib, Pu_kN, phiPn_kN, bresler_ratio, bresler_ok,
                                  ash_ok_flag, vol_m3, peso_kg,
                                  empresa, proyecto, norma, nivel_sis,
                                  n_flejes_x=0, n_flejes_y=0,
                                   n_filas_h=2, n_filas_v=2):
                """Genera un IFC4 con IfcColumn + barras longitudinales + estribos + Pset_NSR10."""
                O = ifcopenshell.file(schema="IFC4")
                # ── Cabecera de proyecto ────────────────────────────────────


                # ── Proyecto / Sitio / Edificio / Piso ─────────────────────
                project  = ifcopenshell.api.run("root.create_entity",    O, ifc_class="IfcProject",  name=proyecto)

                # ── Contexto geométrico ─────────────────────────────────────
                context = ifcopenshell.api.run("context.add_context", O, context_type="Model")
                body    = ifcopenshell.api.run("context.add_context", O,
                    context_type="Model", context_identifier="Body",
                    target_view="MODEL_VIEW", parent=context)

                site     = ifcopenshell.api.run("root.create_entity",    O, ifc_class="IfcSite",     name="Sitio")
                building = ifcopenshell.api.run("root.create_entity",    O, ifc_class="IfcBuilding", name="Edificio")
                storey   = ifcopenshell.api.run("root.create_entity",    O, ifc_class="IfcBuildingStorey", name="Piso 1")
                _u = ifcopenshell.api.run("unit.add_si_unit", O, unit_type="LENGTHUNIT"); ifcopenshell.api.run("unit.assign_unit", O, units=[_u])
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=project, products=[site])
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=site, products=[building])
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=building, products=[storey])

                L_m = L_cm / 100.0;  b_m = b_cm / 100.0;  h_m = h_cm / 100.0
                D_m = D_cm / 100.0;  r_m = recub / 100.0

                # ── Materiales ──────────────────────────────────────────────
                mat_conc = ifcopenshell.api.run("material.add_material", O, name=f"CONCRETO_fc{fc_mpa:.0f}MPa")
                mat_acero = ifcopenshell.api.run("material.add_material", O, name=f"ACERO_fy{fy_mpa:.0f}MPa")

                # ── IfcColumn (geometría extrusión) ─────────────────────────
                column = ifcopenshell.api.run("root.create_entity", O,
                    ifc_class="IfcColumn", name=f"COL-{'CIRC' if es_circ else f'{b_cm:.0f}x{h_cm:.0f}'}")
                column.PredefinedType = "COLUMN"

                # Perfil de la sección según tipo
                import math as _m
                if es_circ:
                    profile = O.createIfcCircleProfileDef(
                        ProfileType="AREA", ProfileName=f"D{D_cm:.0f}",
                        Radius=D_m / 2)
                else:
                    profile = O.createIfcRectangleProfileDef(
                        ProfileType="AREA", ProfileName=f"{b_cm:.0f}x{h_cm:.0f}",
                        XDim=b_m, YDim=h_m)

                # Ubicación y dirección
                origin = O.createIfcCartesianPoint([0.0, 0.0, 0.0])
                z_dir  = O.createIfcDirection([0.0, 0.0, 1.0])
                x_dir  = O.createIfcDirection([1.0, 0.0, 0.0])
                placement = O.createIfcAxis2Placement3D(Location=origin, Axis=z_dir, RefDirection=x_dir)
                solid = O.createIfcExtrudedAreaSolid(
                    SweptArea=profile,
                    Position=placement,
                    ExtrudedDirection=O.createIfcDirection([0.0, 0.0, 1.0]),
                    Depth=L_m)
                shape_rep = O.createIfcShapeRepresentation(
                    ContextOfItems=body, RepresentationIdentifier="Body",
                    RepresentationType="SweptSolid", Items=[solid])
                ifcopenshell.api.run("geometry.assign_representation", O, product=column, representation=shape_rep)
                ifcopenshell.api.run("geometry.edit_object_placement", O, product=column)

                # Asignar al piso
                ifcopenshell.api.run("spatial.assign_container", O,
                    relating_structure=storey, products=[column])
                # Asignar material al concreto
                ifcopenshell.api.run("material.assign_material", O, products=[column], material=mat_conc)

                # ── Barras Longitudinales (IfcReinforcingBar) ───────────────
                import math as _m
                db_m = db_mm / 1000.0

                def _color_by_diam(db):
                    """Mapa diámetro (mm) → RGB IFC [0-1]"""
                    c = {8:(0.8,0.9,1.0), 10:(0.4,0.8,1.0), 12:(0.1,0.7,0.2),
                         16:(1.0,0.8,0.0), 19:(1.0,0.5,0.0), 20:(1.0,0.5,0.0),
                         22:(1.0,0.2,0.2), 25:(0.9,0.0,0.0), 32:(0.6,0.0,0.6)}
                    key = min(c.keys(), key=lambda k: abs(k - db))
                    return c[key]

                rebar_color = _color_by_diam(db_mm)

                offset_real_m = r_m + (dst_mm / 1000.0) + (db_m / 2.0)
                if es_circ:
                    R_bar = (D_m / 2) - offset_real_m  # radio del eje de la barra
                    bar_positions = [
                        (_m.cos(2*_m.pi*i/n_bars)*R_bar,
                         _m.sin(2*_m.pi*i/n_bars)*R_bar)
                        for i in range(n_bars)
                    ]
                else:
                    # Distribución perimetral rectangular
                    nx = max(2, int(n_filas_h))   # barras por cara b (sidebar)
                    ny = max(2, int(n_filas_v))   # barras por cara h (sidebar)
                    xs = [offset_real_m + (b_m - 2*offset_real_m)*i/(nx-1) - b_m/2 for i in range(nx)] if nx > 1 else [0.0]
                    ys = [offset_real_m + (h_m - 2*offset_real_m)*i/(ny-1) - h_m/2 for i in range(ny)] if ny > 1 else [0.0]
                    bar_positions = (
                        [(x, -h_m/2 + offset_real_m) for x in xs] +
                        [(x,  h_m/2 - offset_real_m) for x in xs] +
                        [(-b_m/2 + offset_real_m, y) for y in ys[1:-1]] +
                        [( b_m/2 - offset_real_m, y) for y in ys[1:-1]]
                    )

                long_bars = []
                for i, (bx, by) in enumerate(bar_positions):
                    bar = ifcopenshell.api.run("root.create_entity", O,
                        ifc_class="IfcReinforcingBar", name=f"L{i+1}")
                    bar.NominalDiameter = db_m
                    bar.SteelGrade      = f"fy={fy_mpa:.0f}MPa"
                    bar.BarSurface      = "TEXTURED"
                    bar.PredefinedType  = "MAIN"

                    # Geometría: línea vertical (poly-line extruida)
                    p0 = O.createIfcCartesianPoint([bx, by, 0.0])
                    p1 = O.createIfcCartesianPoint([bx, by, L_m])
                    polyline = O.createIfcPolyline(Points=[p0, p1])
                    bar_profile = O.createIfcCircleProfileDef(
                        ProfileType="AREA", Radius=db_m/2)
                    bar_place = O.createIfcAxis2Placement3D(
                        Location=O.createIfcCartesianPoint([bx, by, 0.0]),
                        Axis=O.createIfcDirection([0.0, 0.0, 1.0]),
                        RefDirection=O.createIfcDirection([1.0, 0.0, 0.0]))
                    bar_solid = O.createIfcExtrudedAreaSolid(
                        SweptArea=bar_profile, Position=bar_place,
                        ExtrudedDirection=O.createIfcDirection([0.0, 0.0, 1.0]),
                        Depth=L_m)
                    style = O.createIfcSurfaceStyleRendering(
                        SurfaceColour=O.createIfcColourRgb(Red=rebar_color[0], Green=rebar_color[1], Blue=rebar_color[2]),
                        ReflectanceMethod="FLAT")
                    surface_style = O.createIfcSurfaceStyle(
                        Name=f"AceroLong_{db_mm:.0f}mm", Side="BOTH", Styles=[style])
                    O.createIfcStyledItem(Item=bar_solid, Styles=[surface_style])
                    bar_rep = O.createIfcShapeRepresentation(
                        ContextOfItems=body, RepresentationIdentifier="Body",
                        RepresentationType="SweptSolid", Items=[bar_solid])
                    ifcopenshell.api.run("geometry.assign_representation", O, product=bar, representation=bar_rep)
                    ifcopenshell.api.run("geometry.edit_object_placement", O, product=bar)

                    ifcopenshell.api.run("spatial.assign_container", O,
                        relating_structure=storey, products=[bar])
                    ifcopenshell.api.run("material.assign_material", O, products=[bar], material=mat_acero)
                    long_bars.append(bar)

                # ── Estribos / Espiral (IfcReinforcingBar cerrado) ──────────
                dst_m = dst_mm / 1000.0
                stirrup_color = _color_by_diam(dst_mm)
                stirrup_bars = []
                
                # Sincronizar alturas IFC EXACTAMENTE con la cantidad del APU
                import math as _m
                _n_est_por_Lo = _m.ceil(Lo_cm / s_conf_cm)
                _long_zona_libre = max(0, L_cm - 2 * Lo_cm)
                _n_estribos_centro = max(0, _m.ceil(_long_zona_libre / s_bas_cm) - 1) if _long_zona_libre > 0 else 0

                y_positions = []
                _esp_real_conf_cm = Lo_cm / _n_est_por_Lo if _n_est_por_Lo > 0 else s_conf_cm
                for i in range(_n_est_por_Lo + 1):
                    y_positions.append((i * _esp_real_conf_cm) / 100.0)
                if _n_estribos_centro > 0:
                    _esp_real_centro = (_long_zona_libre + _esp_real_conf_cm) / (_n_estribos_centro + 1)
                    for i in range(1, _n_estribos_centro + 1):
                        y_positions.append((Lo_cm + i * _esp_real_centro) / 100.0)
                if _n_est_por_Lo > 0:
                    for i in range(_n_est_por_Lo):
                        y_positions.append((L_cm - Lo_cm + (i+1) * _esp_real_conf_cm) / 100.0)
                
                # ── Precomputo de geometría de grapas (se usa dentro del bucle j,y_z) ──
                import math as _mg
                if not es_circ:
                    _r_est_m = dst_m / 2.0
                    _r_var_m = db_m  / 2.0
                    _off_m   = r_m + dst_m + _r_var_m
                    _nx_fl   = max(2, int(n_filas_h))
                    _ny_fl   = max(2, int(n_filas_v))
                    _xv_max  = b_m / 2.0 - _off_m
                    _yv_max  = h_m / 2.0 - _off_m
                    _xs_fl   = ([_off_m + (b_m-2*_off_m)*_i/(_nx_fl-1) - b_m/2
                                  for _i in range(_nx_fl)] if _nx_fl > 1 else [0.0])
                    _ys_fl   = ([_off_m + (h_m-2*_off_m)*_i/(_ny_fl-1) - h_m/2
                                  for _i in range(_ny_fl)] if _ny_fl > 1 else [0.0])
                    _R_g     = _r_var_m + _r_est_m
                    _Lc_g    = max(6.0 * dst_m, 0.075)

                    def _pts_grapa_x(X_v, Y_f, R, L_c):
                        return generar_puntos_grapa_x(X_v, Y_f, R, L_c)

                    def _pts_grapa_y(X_f, Y_v, R, L_c):
                        return generar_puntos_grapa_y(X_f, Y_v, R, L_c)
                else:
                    _ys_fl = []  # circular: sin grapas

                for j, y_z in enumerate(y_positions):
                    st_bar = ifcopenshell.api.run("root.create_entity", O,
                        ifc_class="IfcReinforcingBar", name=f"E{j+1}")
                    st_bar.NominalDiameter = dst_m
                    st_bar.SteelGrade      = f"fy={fy_mpa:.0f}MPa"
                    st_bar.BarSurface      = "TEXTURED"
                    st_bar.PredefinedType  = "MAIN"
                
                    # Polilínea cerrada del estribo
                    if es_circ:
                        R_e   = (D_m / 2) - r_m
                        n_pts = 16
                        pts_est = [O.createIfcCartesianPoint([
                            _m.cos(2*_m.pi*k/n_pts)*R_e,
                            _m.sin(2*_m.pi*k/n_pts)*R_e, y_z])
                            for k in range(n_pts + 1)]
                    else:
                        bx2  = b_m / 2.0 - r_m
                        hy2  = h_m / 2.0 - r_m
                        X_c  = -bx2;  Y_c = hy2
                        _R   = max(3.0 * dst_m, 0.012)
                        _hk  = max(6.0 * db_mm/1000.0, 0.075)
                        Cx   = X_c + _R;  Cy = Y_c - _R
                        _z1  = y_z + dst_m
                        _na  = 8
                        def _arc(cx, cy, r, a0, a1, z_c):
                            return [O.createIfcCartesianPoint([
                                cx + r*_m.cos(_m.radians(a0+(a1-a0)*_ii/_na)),
                                cy + r*_m.sin(_m.radians(a0+(a1-a0)*_ii/_na)),
                                z_c]) for _ii in range(_na+1)]
                        def _pt(x, y, z_c): return O.createIfcCartesianPoint([x, y, z_c])
                        _pie1x = Cx + _R*_m.cos(_m.radians(225)); _pie1y = Cy + _R*_m.sin(_m.radians(225))
                        _pie2x = Cx + _R*_m.cos(_m.radians(45));  _pie2y = Cy + _R*_m.sin(_m.radians(45))
                        _dk = _hk * 0.7071; _nT = _na
                        pts_est = []
                        pts_est.append(_pt(_pie1x + _dk, _pie1y - _dk, _z1))
                        pts_est.append(_pt(_pie1x,       _pie1y,       _z1))
                        pts_est += _arc(Cx, Cy, _R, 225, 90, _z1)
                        pts_est += [_pt(Cx+(bx2-_R-Cx)*_ii/_nT, Y_c, _z1+(y_z-_z1)*_ii/_nT)
                                    for _ii in range(_nT+1)]
                        pts_est += _arc(bx2-_R, Y_c-_R, _R, 90, 0, y_z)
                        pts_est += [_pt(bx2, Y_c-_R+(-(Y_c-_R)-(Y_c-_R))*_ii/_nT, y_z)
                                    for _ii in range(_nT+1)]
                        pts_est += _arc(bx2-_R, -(Y_c-_R), _R, 0, -90, y_z)
                        pts_est += [_pt(bx2-_R+(-(bx2-_R)-(bx2-_R))*_ii/_nT, -Y_c, y_z)
                                    for _ii in range(_nT+1)]
                        pts_est += _arc(-(bx2-_R), -(Y_c-_R), _R, -90, -180, y_z)
                        pts_est += [_pt(X_c, -(Y_c-_R)+(Cy-(-(Y_c-_R)))*_ii/_nT, y_z+(_z1-y_z)*_ii/_nT)
                                    for _ii in range(_nT+1)]
                        pts_est += _arc(Cx, Cy, _R, 180, 45, _z1)
                        pts_est.append(_pt(_pie2x,       _pie2y,       _z1))
                        pts_est.append(_pt(_pie2x + _dk, _pie2y - _dk, _z1))
                        if j % 2 != 0:
                            pts_est = [O.createIfcCartesianPoint(
                                [-p.Coordinates[0], -p.Coordinates[1], p.Coordinates[2]])
                                for p in pts_est]
                
                    polyline_st = O.createIfcPolyline(Points=pts_est)
                    # CRÍTICO: IfcSweptDiskSolid da el VOLUMEN físico del tubo
                    st_swept = O.createIfcSweptDiskSolid(Directrix=polyline_st, Radius=dst_m/2)
                    _scol = stirrup_color
                    _in_conf_z = (y_z*100 <= Lo_cm) or (y_z*100 >= L_cm - Lo_cm)
                    if _in_conf_z:
                        _scol = stirrup_color
                    st_style_rend = O.createIfcSurfaceStyleRendering(
                        SurfaceColour=O.createIfcColourRgb(Red=_scol[0], Green=_scol[1], Blue=_scol[2]),
                        ReflectanceMethod="FLAT")
                    st_surface_style = O.createIfcSurfaceStyle(
                        Name=f"Estribo_{dst_mm:.0f}mm", Side="BOTH", Styles=[st_style_rend])
                    O.createIfcStyledItem(Item=st_swept, Styles=[st_surface_style])
                    st_rep = O.createIfcShapeRepresentation(
                        ContextOfItems=body, RepresentationIdentifier="Body",
                        RepresentationType="AdvancedSweptSolid", Items=[st_swept])
                    ifcopenshell.api.run("geometry.assign_representation", O, product=st_bar, representation=st_rep)
                    ifcopenshell.api.run("geometry.edit_object_placement", O, product=st_bar)
                    ifcopenshell.api.run("spatial.assign_container", O,
                        relating_structure=storey, products=[st_bar])
                    ifcopenshell.api.run("material.assign_material", O, products=[st_bar], material=mat_acero)
                    stirrup_bars.append(st_bar)

                    # ── Grapas al mismo nivel Z que el estribo ──────────────
                    if not es_circ and (n_flejes_x > 0 or n_flejes_y > 0):
                        def _mk_grapa(pts2d, _tag):
                            if len(pts2d) < 2: return
                            _ifc_p = [O.createIfcCartesianPoint([x, y, y_z]) for x, y in pts2d]
                            _pl    = O.createIfcPolyline(Points=_ifc_p)
                            _sol   = O.createIfcSweptDiskSolid(Directrix=_pl, Radius=dst_m/2.0)
                            O.createIfcStyledItem(Item=_sol, Styles=[st_surface_style])
                            _bf    = ifcopenshell.api.run('root.create_entity', O,
                                ifc_class='IfcReinforcingBar', name=_tag)
                            _bf.NominalDiameter = dst_m
                            _bf.SteelGrade      = f'fy={fy_mpa:.0f} MPa'
                            _bf.PredefinedType  = 'MAIN'
                            _rf = O.createIfcShapeRepresentation(ContextOfItems=body,
                                RepresentationIdentifier='Body',
                                RepresentationType='AdvancedSweptSolid', Items=[_sol])
                            ifcopenshell.api.run('geometry.assign_representation', O,
                                product=_bf, representation=_rf)
                            ifcopenshell.api.run('geometry.edit_object_placement', O, product=_bf)
                            ifcopenshell.api.run('spatial.assign_container', O,
                                relating_structure=storey, products=[_bf])
                            ifcopenshell.api.run('aggregate.assign_object', O,
                                relating_object=column, products=[_bf])
                        
                        # ── Flejes (Cross-Ties) distribuidos equitativamente (Sincronizado con 2D/3D) ──
                        # Semianchos del estribo perimetral
                        hw_m = (h_m / 2.0) - r_m
                        bw_m = (b_m / 2.0) - r_m

                        # GRAPAS IFC ANCLADAS A VARILLAS REALES (_ys_fl / _xs_fl)
                        if n_flejes_x > 0 and len(_ys_fl) > 2:
                            for idx, y_coord in enumerate(_ys_fl[1:-1]):
                                _mk_grapa(_pts_grapa_x(_xv_max, y_coord, _R_g, _Lc_g), f'FX{idx}_E{j+1}')
                                
                        if n_flejes_y > 0 and len(_xs_fl) > 2:
                            for idx, x_coord in enumerate(_xs_fl[1:-1]):
                                _mk_grapa(_pts_grapa_y(x_coord, _yv_max, _R_g, _Lc_g), f'FY{idx}_E{j+1}')

                psets_data = {
                    "Pset_ConcreteElementGeneral": {
                        "ConstructionMethod":      "CastInPlace",
                        "StructuralClass":         "Column",
                        "StrengthClass":           f"f'c {fc_mpa:.1f} MPa",
                        "ReinforcementVolumeRatio": float(rho_pct) / 100.0,
                        "ReinforcementAreaRatio":   float(As_cm2)
                    },
                    f"Pset_{norma_sel.split('(')[0].strip().replace('-','_').replace(' ','_')}": {
                        "Norma":                norma,
                        "Nivel_Sismico":        nivel_sis,
                        "fc_MPa":               float(fc_mpa),
                        "fy_MPa":               float(fy_mpa),
                        "As_cm2":               float(As_cm2),
                        "Rho_pct":              float(rho_pct),
                        "n_barras_long":        int(n_long_total),
                        "diametro_barra_mm":    float(db_mm),
                        "diametro_estribo_mm":  float(dst_mm),
                        "sep_estrib_conf_cm":   float(s_conf_cm),
                        "sep_estrib_basica_cm": float(s_bas_cm),
                        "Lo_conf_cm":           float(Lo_cm),
                        "n_estribos_total":     int(n_estrib),
                        "Pu_kN":                float(Pu_kN),
                        "phi_Pni_kN":           float(phiPn_kN),
                        "Bresler_ratio":        float(bresler_ratio),
                        "Bresler_CUMPLE":       "SI" if bresler_ok else "NO",
                        "Ash_CUMPLE":           "SI" if ash_ok_flag else "NO",
                        "Vol_concreto_m3":      float(vol_m3),
                        "Peso_acero_kg":        float(peso_kg),
                    },
                    "Pset_ColGeom": {
                        "Tipo":       "Circular" if es_circ else "Rectangular",
                        "b_cm":       float(b_cm),
                        "h_cm":       float(h_cm),
                        "D_cm":       float(D_cm) if es_circ else 0.0,
                        "L_cm":       float(L_cm),
                        "Recub_cm":   float(recub),
                    },
                }

                for pset_name, props_dict in psets_data.items():
                    pset = ifcopenshell.api.run("pset.add_pset", O,
                        product=column, name=pset_name)
                    ifcopenshell.api.run("pset.edit_pset", O,
                        pset=pset, properties=props_dict)

                # Añadir un resumen super explícito y visible en el visor IFC
                column.Name = f"COL-{'CIRC' if es_circ else f'{b_cm:.0f}x{h_cm:.0f}'} | TOTAL: {len(y_positions)} ESTRIBOS (FLEJES)"
                column.Description = f"Conexión Motor APU: Exactamente {len(y_positions)} estribos calculados y colocados geométricamente."

                # ── Guardar en buffer ───────────────────────────────────────
                import tempfile, os as _os
                with tempfile.NamedTemporaryFile(suffix='.ifc', delete=False) as tmp:
                    tmp_path = tmp.name
                O.write(tmp_path)
                with open(tmp_path, 'rb') as f_ifc:
                    buf = f_ifc.read()
                _os.unlink(tmp_path)
                return buf

            # Botón para generar — solo ejecuta la función al hacer clic
            if st.button("⬇️ " + _t("Generar y Exportar IFC4 (BIM completo)", "Generate & Export IFC4 (full BIM)"), key="col_btn_ifc"):
                _ifc_placeholder = st.empty()
                _ifc_placeholder.info("⏳ Generando modelo IFC4… puede tomar unos segundos.")
                try:
                    phiPn_max = cap_x.get('phi_Pn_max', 0) if 'cap_x' in locals() and cap_x else 0
                    buf_ifc_col = _make_ifc_columna(
                        b_cm        = b   if not es_circular else 0,
                        h_cm        = h   if not es_circular else 0,
                        D_cm        = D   if es_circular     else 0,
                        L_cm        = L_col,
                        es_circ     = es_circular,
                        fc_mpa      = fc, fy_mpa = fy,
                        n_bars      = n_barras if es_circular else n_barras_total,
                        db_mm       = rebar_diam, dst_mm = stirrup_diam,
                        n_long_total= n_barras if es_circular else n_barras_total,
                        As_cm2      = Ast, rho_pct = cuantia,
                        recub       = recub_cm,
                        s_conf_cm   = s_conf  if not es_circular else paso_espiral,
                        s_bas_cm    = s_centro if not es_circular else paso_espiral,
                        Lo_cm       = Lo_conf,
                        n_estrib    = n_estribos_total if not es_circular else 1,
                        Pu_kN       = Pu_input * factor_fuerza,
                        phiPn_kN    = phiPn_max * factor_fuerza,
                        bresler_ratio = bresler['ratio'],
                        bresler_ok  = bresler['ok'],
                        ash_ok_flag = ash_ok,
                        vol_m3      = vol_concreto_m3,
                        peso_kg     = peso_total_acero_kg,
                        empresa     = st.session_state.get("cpm_empresa","________________"),
                        proyecto    = st.session_state.get("cpm_proyecto_nombre","________________"),
                        norma       = norma_sel,
                        nivel_sis   = nivel_sismico,
                        n_flejes_x  = num_flejes_x if not es_circular else 0,
                        n_flejes_y  = num_flejes_y if not es_circular else 0,
                        n_filas_h   = num_filas_h  if not es_circular else 0,
                        n_filas_v   = num_filas_v  if not es_circular else 0,
                    )
                    _ifc_placeholder.empty()
                    if buf_ifc_col:
                        _ifc_fname = f"Columna_{'Circ_D'+str(int(D)) if es_circular else str(int(b))+'x'+str(int(h))}_IFC4.ifc"
                        st.success(_t("IFC4 generado correctamente.", "IFC4 generated successfully."))
                        st.download_button(
                            label=_t("⬇️ Descargar IFC4", "⬇️ Download IFC4"),
                            data=buf_ifc_col,
                            file_name=_ifc_fname,
                            mime="application/x-step",
                            key="col_ifc_dl")
                        st.caption(_t(
                            f"IFC4: IfcColumn + barras long. + estribos 3D + Pset_{norma_sel.split('(')[0].strip()}",
                            f"IFC4: IfcColumn + longitudinal bars + 3D stirrups + Pset_{norma_sel.split('(')[0].strip()}"))
                except Exception as _e_ifc_btn:
                    import traceback as _tb_ifc_btn
                    _ifc_placeholder.empty()
                    st.error(f"Error IFC: {_e_ifc_btn}")
                    st.code(_tb_ifc_btn.format_exc(), language='python')


        except ImportError as e_imp:
            import sys as _sys_ifc
            import importlib
            importlib.invalidate_caches()
            _pyver = f"{_sys_ifc.version_info.major}.{_sys_ifc.version_info.minor}"
            _pyexe = _sys_ifc.executable
            st.error(
                f"AVISO Error cargando **IfcOpenShell** en este entorno Python.\n\n"
                f"**Python activo:** `{_pyexe}` (versión {_pyver})\n\n"
                f"**Detalle del error (ImportError):** `{str(e_imp)}`\n\n"
                f"**Solución 1:** Si acabas de instalar mediante pip, **REINICIA EL SERVIDOR DE STREAMLIT** (`Ctrl+C` y luego `streamlit run app.py`).\n\n"
                f"**Solución 2:** Si el error persiste, verifica la instalación en una terminal.\n"
                f"```bash\n{_pyexe} -m pip install ifcopenshell\n```"
            )
        except Exception as e_ifc:
            import traceback
            st.error(f"Error IFC: {e_ifc}")
            st.code(traceback.format_exc(), language='python')
            st.info("Asegurate de que ifcopenshell este disponible en el entorno de ejecucion.")
    if btn_docx_col:
        try:
            from utils_docx import fig_to_docx_white
        except ImportError:
            # Fallback robusto: siempre retorna io.BytesIO
            import io as _io_fbk
            def fig_to_docx_white(fig):
                """Modo claro para DOCX. Restaura colores pantalla."""
                import io as _iow
                if not fig.get_axes():
                    b=_iow.BytesIO(); b.seek(0); return b
                _ofc=fig.get_facecolor()
                _oax=[(ax,ax.get_facecolor(),ax.xaxis.label.get_color(),
                       ax.yaxis.label.get_color(),ax.title.get_color() if ax.get_title() else None,
                       {sp:ax.spines[sp].get_edgecolor() for sp in ax.spines})
                      for ax in fig.get_axes()]
                fig.patch.set_facecolor('white'); fig.patch.set_alpha(1.0)
                for ax in fig.get_axes():
                    ax.set_facecolor('#f8f9fa')
                    ax.tick_params(colors='#1a1a1a')
                    ax.xaxis.label.set_color('#1a1a1a')
                    ax.yaxis.label.set_color('#1a1a1a')
                    if ax.get_title(): ax.title.set_color('#1a1a1a')
                    for sp in ax.spines.values():
                        sp.set_edgecolor('#cccccc'); sp.set_linewidth(0.8)
                buf=_iow.BytesIO()
                fig.savefig(buf,format='png',dpi=200,bbox_inches='tight',facecolor='white',transparent=False)
                buf.seek(0)
                fig.patch.set_facecolor(_ofc)
                for ax,ofc,oxl,oyl,otit,osps in _oax:
                    ax.set_facecolor(ofc); ax.xaxis.label.set_color(oxl); ax.yaxis.label.set_color(oyl)
                    if otit: ax.title.set_color(otit)
                    for sn,sc in osps.items(): ax.spines[sn].set_edgecolor(sc)
                return buf

        doc = Document()
        doc.add_heading(f"MEMORIA ESTRUCTURAL — COLUMNA {'CIRCULAR' if es_circular else 'RECTANGULAR'} ({norma_sel})", 0)
        
        # 0. Portada dinámica — marca blanca
        _emp  = st.session_state.get("cpm_empresa","________________")
        _proy = st.session_state.get("cpm_proyecto_nombre","________________")
        _ing  = st.session_state.get("cpm_ingeniero","________________")
        p0 = doc.add_paragraph()
        p0.add_run(f"EMPRESA: {_emp}\n").bold = True
        p0.add_run(f"PROYECTO: {_proy}\n").bold = True
        p0.add_run(f"INGENIERO RESPONSABLE: {_ing}\n").bold = True
        p0.add_run(f"ELEMENTO: Columna {'Circular' if es_circular else 'Rectangular'} | NORMA: {norma_sel} | SÍSMICO: {nivel_sismico}\n")
        p0.add_run(f"MATERIALES: f'c={fc:.1f} MPa, fy={fy:.1f} MPa | FECHA: {datetime.datetime.now().strftime('%d/%m/%Y')}")
        
        # 1. Parámetros de Diseño
        doc.add_heading("1. PARÁMETROS GEOMÉTRICOS Y SOLICITACIONES", level=1)
        doc.add_paragraph(f"Geometría: {'D' if es_circular else 'b'} = {D if es_circular else b:.0f} cm, {'h = '+str(h)+' cm, ' if not es_circular else ''}L = {L_col:.0f} cm\nRecubrimiento libre: {recub_cm:.2f} cm\nCargas de diseño aplicadas:\n  Pu = {Pu_input:.1f} {unidad_fuerza}\n  Mux = {Mux_input:.1f} {unidad_mom}\n  Muy = {Muy_input:.1f} {unidad_mom}")
        
        # 2. Esbeltez
        doc.add_heading("2. PRE-DIMENSIONAMIENTO Y EFECTOS DE ESBELTEZ (NSR-10 C.10.10)", level=1)
        # ── Tipo de pórtico ─────────────────────────────────────────────────────
        _tipo_portico_str = ("Marco Desplazable (Sway Frame) — Mu amplificado por análisis P-Δ global"
                             if es_desplazable else
                             "Marco No Desplazable (Non-Sway Frame) — δns calculado por el módulo")
        doc.add_paragraph(f"Tipo de pórtico: {_tipo_portico_str}")
        doc.add_paragraph(f"Relación de esbeltez (kL/r): {slenderness['kl_r']:.2f}")
        if es_desplazable:
            _p_sway = doc.add_paragraph(
                f"MARCO DESPLAZABLE: Los momentos Mux = {Mux_magnified:.2f} {unidad_mom} y "
                f"Muy = {Muy_magnified:.2f} {unidad_mom} ingresados se asumen ya amplificados "
                f"por el análisis P-Δ del modelo estructural global (ETABS/SAP2000/MIDAS). "
                f"El módulo no aplica magnificación δns adicional, conforme a "
                f"{norma_sel} C.10.10.7 / ACI 318-19 §6.6.4.6. "
                f"Clasificación de esbeltez: {slenderness['classification']}."
            )
            _p_sway.runs[0].bold = True if slenderness['kl_r'] > 22 else False
        elif slenderness['kl_r'] > 100:
            p_esb = doc.add_paragraph("[ADVERTENCIA] Columna MUY esbelta (kL/r > 100). Las ecuaciones estándar pierden validez, requiere análisis P-delta riguroso.")
            p_esb.bold = True
        elif slenderness['kl_r'] > 22:
            doc.add_paragraph(f"Columna Esbelta — δns = {slenderness['delta_ns']:.3f} ({code['ref']}, método no-sway). Momentos amplificados: Mux = {Mux_magnified:.2f} {unidad_mom} | Muy = {Muy_magnified:.2f} {unidad_mom}.").runs[0].bold=True
            if 'Pc' in dir(): doc.add_paragraph(f"Carga crítica Euler (Pc) = {Pc:.1f} kN")
            doc.add_paragraph(f"Momento magnificado Mux* = {Mux_input*slenderness['delta_ns']:.2f} {unidad_mom}")
            _r_min = (h if not es_circular else D)*0.289
            _b_min = round(k_factor*L_col/(22*0.289)*10)/10 if 'k_factor' in dir() else round(L_col/6.4,1)
            doc.add_paragraph(f"️ Sección mínima recomendada para kL/r≤22: b_min ≥ {_b_min:.1f} cm (NSR-10 C.10.10.1)").runs[0].bold=True
        else:
            doc.add_paragraph("Columna Corta — No requiere magnificación de momentos por esbeltez.")
            
        # 3. Diseño a Flexo-compresión
        doc.add_heading("3. CAPACIDAD A FLEXO-COMPRESIÓN BIAXIAL", level=1)
        doc.add_heading("3.1 Detallado de Capas Longitudinales", level=2)
        
        Ast_total = Ast
        doc.add_paragraph(f"Área Bruta (Ag): {Ag:.2f} cm² | Área de Acero Total (Ast): {Ast_total:.2f} cm²")
        # Sustitución numérica paso a paso
        _Po_calc=(0.85*fc*(Ag-Ast_total)+fy*Ast_total)/10.0  # kN
        _Pnmax_calc=0.80*_Po_calc  # estribos
        _phi_c=0.65 if not es_circular else 0.75
        _phiPnmax=_phi_c*_Pnmax_calc
        doc.add_paragraph(f"Po = 0.85·f'c·(Ag−Ast) + fy·Ast = 0.85×{fc:.1f}×({Ag:.2f}−{Ast_total:.2f}) + {fy:.1f}×{Ast_total:.2f} = {_Po_calc:.1f} kN  [NSR-10 C.10.3.6]")
        doc.add_paragraph(f"Pn,máx = 0.80×Po = 0.80×{_Po_calc:.1f} = {_Pnmax_calc:.1f} kN  (estribos rect. ACI §22.4.2.1)")
        doc.add_paragraph(f"φPn,máx = φ×Pn,máx = {_phi_c:.2f}×{_Pnmax_calc:.1f} = {_phiPnmax:.1f} kN")
        
        tb_lay = doc.add_table(rows=1, cols=5)
        tb_lay.style = 'Light Grid'
        tb_lay.rows[0].cells[0].text = 'Capa de Acero'
        tb_lay.rows[0].cells[1].text = "d' (cm)"
        tb_lay.rows[0].cells[2].text = 'As (cm²)'
        tb_lay.rows[0].cells[3].text = 'fs (MPa)'
        tb_lay.rows[0].cells[4].text = 'εs'
        
        # Estado de limit balanceado (solo como referencia estatica aprox)
        h_eff = D if es_circular else h
        cb = h_eff * 0.003 / (0.003 + fy/200000)
        
        for i, lay in enumerate(layers if not es_circular else area_capas_mock):
            rc = tb_lay.add_row().cells
            rc[0].text = f"Capa {i+1}"
            try:
                di = lay['d'] if isinstance(lay, dict) else (D/2)
                As_i = lay['As'] if isinstance(lay, dict) else lay
                es_i = 0.003 * ((cb - di) / cb) if cb>0 else 0
                fs_i = max(-fy, min(fy, es_i * 200000))
                rc[1].text = f"{di:.1f}"
                rc[2].text = f"{As_i:.2f}"
                rc[3].text = f"{fs_i:.0f} (Bal)"
                rc[4].text = f"{es_i:.4f} (Bal)"
            except: pass
            
        doc.add_heading("3.2 Verificación Biaxial (Bresler)", level=2)
        Po_calc = (0.85 * fc * (Ag/10000 - Ast_total/10000) + fy * (Ast_total/10000)) * 1000 
        doc.add_paragraph(f"Resistencia teórica pura (Po) = {Po_calc:.1f} kN")
        doc.add_paragraph(f"Método de Bresler Biaxial: {'CUMPLE' if bresler['ok'] else 'NO CUMPLE'} (Ratio = {bresler['ratio']:.3f})")
        if not bresler['ok']:
            doc.add_paragraph(f"[ADVERTENCIA] Se requiere aumentar sección o cuantía.").bold = True
            
        doc.add_heading("3.3 Envolvente P-M de Capacidad Real", level=2)
        try:
            import io as _io_pm
            buf_pm = _io_pm.BytesIO()
            # Forzar fondo blanco antes de capturar
            _orig_fc = fig_pm_2d.get_facecolor()
            fig_pm_2d.patch.set_facecolor('white')
            fig_pm_2d.patch.set_alpha(1.0)
            for _ax in fig_pm_2d.get_axes():
                _ax.set_facecolor('white')
                _ax.tick_params(colors='#1a1a1a')
                _ax.xaxis.label.set_color('#1a1a1a')
                _ax.yaxis.label.set_color('#1a1a1a')
            fig_pm_2d.savefig(buf_pm, facecolor='white', edgecolor='none', format='png', dpi=200,
                              bbox_inches='tight', transparent=False)
            buf_pm.seek(0)
            fig_pm_2d.patch.set_facecolor(_orig_fc)  # restaurar modo oscuro
            doc.add_picture(buf_pm, width=Inches(6.0))
            buf_pm.close()
        except Exception: pass

        # ── Diagrama P-M eje Y ────────────────────────────────────────────────
        doc.add_heading(_t("3.3b Diagrama P-M — Eje Y (flexion sobre b)",
                            "3.3b P-M Diagram — Y-Axis (bending about b)"), level=2)
        doc.add_paragraph(_t(
            f"Diagrama de interaccion P-M para flexion sobre el eje Y "
            f"(dimension b = {b:.1f} cm). "
            f"Punto de demanda: Pu = {Pu_input:.1f} {unidad_fuerza}, "
            f"Muy = {Muy_input:.2f} {unidad_mom}.",
            f"P-M interaction diagram for bending about Y-axis "
            f"(dimension b = {b:.1f} cm). "
            f"Demand point: Pu = {Pu_input:.1f} {unidad_fuerza}, "
            f"Muy = {Muy_input:.2f} {unidad_mom}."
        ))
        try:
            import io as _io_pmy
            _fig_pm_y = plot_pm_2d(
                cap_y,
                Pu_input / factor_fuerza,
                Muy_input / factor_fuerza,
                unidad_fuerza=unidad_fuerza,
                unidad_mom=unidad_mom,
                titulo=_t(
                    f"Diagrama P-M — Eje Y | b={b:.0f} cm  h={h:.0f} cm",
                    f"P-M Diagram — Y-Axis | b={b:.0f} cm  h={h:.0f} cm"
                ),
                factor_fuerza=factor_fuerza,
            )
            configurar_pdf_comercial(_fig_pm_y)
            _buf_pmy = _io_pmy.BytesIO()
            _fig_pm_y.patch.set_facecolor("white")
            _fig_pm_y.patch.set_alpha(1.0)
            for _axy in _fig_pm_y.get_axes():
                _axy.set_facecolor("white")
                _axy.tick_params(colors="#1a1a1a")
                _axy.xaxis.label.set_color("#1a1a1a")
                _axy.yaxis.label.set_color("#1a1a1a")
            _fig_pm_y.savefig(_buf_pmy, facecolor="white", edgecolor="none",
                               format="png", dpi=200, bbox_inches="tight",
                               transparent=False)
            _buf_pmy.seek(0)
            doc.add_picture(_buf_pmy, width=Inches(6.0))
            _buf_pmy.close()
            plt.close(_fig_pm_y)
        except Exception:
            pass
        # ── Fin diagrama eje Y ────────────────────────────────────────────────

        doc.add_heading("3.4 Superficie de Interacción 3D Biaxial", level=2)
        try:
            import plotly.io as pio
            # N4 – Inyectar el modelo 3D con layout blanco + paneles blancos
            fig_3d.update_layout(
                template="plotly_white",
                scene=dict(
                    xaxis=dict(backgroundcolor='white', gridcolor='#cccccc', showbackground=True),
                    yaxis=dict(backgroundcolor='white', gridcolor='#cccccc', showbackground=True),
                    zaxis=dict(backgroundcolor='white', gridcolor='#cccccc', showbackground=True),
                    bgcolor='white'
                ),
                paper_bgcolor='white', plot_bgcolor='white'
            )
            buf_3d = io.BytesIO()
            pio.write_image(fig_3d, buf_3d, format="png", width=600, height=450, scale=1.5)
            buf_3d.seek(0)
            doc.add_picture(buf_3d, width=Inches(5.0))
        except Exception as e: pass
            
        # 4. Diseño a Cortante
        doc.add_heading("4. DISEÑO A CORTANTE Y CONFINAMIENTO SÍSMICO", level=1)
        doc.add_paragraph("El detallado de estribos se rige por las especificaciones de grado de disipación de energía de la NSR-10.")
        
        if not es_circular:
            doc.add_paragraph(f"Ash Requerido de control: {Ash_req:.3f} cm²/sentido  |  Ash Provisto: {Ash_prov:.3f} cm²")
        else:
            doc.add_paragraph(f"Espiral requerida mínima: {rho_s_req:.4f}  |  Espiral provista: {rho_s_prov:.4f}")
        
        sconf_saf = locals().get('sconf', locals().get('s_conf', st.session_state.get('col_sconf', 0.0)))
        Lo_conf_saf = locals().get('Lo_conf', st.session_state.get('col_Loconf', 0.0))
        stirrupdiam_saf = locals().get('stirrupdiam', locals().get('stirrup_diam', st.session_state.get('col_stirrupdiam', 0)))
        s_bas_saf = locals().get('s_bas_cm', locals().get('s_bas', locals().get('s_basico', locals().get('s_centro', 0.0))))
        if s_bas_saf == 0.0:
            s_bas_saf = locals().get('s_conf', sconf_saf)  # fallback: usar s_conf si s_bas es 0
        
        doc.add_heading("4.1 Justificación de Espaciamiento", level=2)
        doc.add_paragraph(f"Longitudes límite Lo = {Lo_conf_saf:.1f} cm")
        doc.add_paragraph(f"Número total de estribos: {n_estribos_total} uds.")
        doc.add_paragraph(f"s1 (Separación confinada): Estribos {stirrupdiam_saf:.1f}mm @ {sconf_saf:.1f} cm")
        doc.add_paragraph(f"s2 (Separación centro): Estribos {stirrupdiam_saf:.1f}mm @ {s_bas_saf:.1f} cm")

        # 5. Verificaciones Normativas
        doc.add_heading(f"5. VERIFICACIONES NORMATIVAS — {norma_sel}", level=1)
        table5 = doc.add_table(rows=1, cols=5)
        table5.style = 'Table Grid'
        _hc5 = table5.rows[0].cells
        _hc5[0].text="Parámetro"; _hc5[1].text="Artículo NSR-10"
        _hc5[2].text="Valor Calculado"; _hc5[3].text="Límite Normativo"; _hc5[4].text="Estado"
        from docx.oxml.ns import qn as _qn5
        from docx.oxml import OxmlElement as _OE5
        def _shade5(cell, color):
            s=_OE5("w:shd"); s.set(_qn5("w:fill"),color)
            cell._tc.get_or_add_tcPr().append(s)
        def _add_row5(param, art, val, lim, ok):
            rc=table5.add_row().cells
            rc[0].text=str(param); rc[1].text=str(art)
            rc[2].text=str(val);   rc[3].text=str(lim)
            if ok is None:
                rc[4].text="N/A"
                _shade5(rc[4],"D9D9D9")   # gris neutro — dato no ingresado
            elif ok:
                rc[4].text="CUMPLE"
                _shade5(rc[4],"C6EFCE")   # verde
            else:
                rc[4].text="NO CUMPLE"
                _shade5(rc[4],"FFC7CE")   # rojo
        _art_conf='C.21.6.4' if 'DES' in nivel_sismico else ('C.21.3.5' if 'DMO' in nivel_sismico else 'C.11.4.1')
        _s_lim=10.0 if 'DES' in nivel_sismico else 15.0
        _add_row5("Cuantía ρ","C.10.9.1 / C.21.6.3",f"ρ={cuantia:.3f}%",f"{rho_min}%–{rho_max}%",(rho_min<=cuantia<=rho_max))
        _add_row5("Biaxial Bresler","C.10.3.6 / ACI §22.4",f"Ratio={bresler['ratio']:.3f}","≤ 1.0",bresler['ok'])
        _add_row5("Ash confinamiento",f"{_art_conf}.4","Ash_prov ≥ Ash_req","Ash_req según fórmula",ash_ok)
        _add_row5(f"Sep. conf. zona ({nivel_sismico})",_art_conf,f"s={sconf_saf:.1f} cm",f"≤ {_s_lim:.0f} cm",sconf_saf<=_s_lim)
        _tipo_esb_str = (f"Desplazable — P-Δ en Mu (kL/r={slenderness['kl_r']:.1f})"
                     if es_desplazable else
                     f"No Desplazable — δns={slenderness['delta_ns']:.3f} (kL/r={slenderness['kl_r']:.1f})")
        _add_row5("Esbeltez / Tipo Pórtico","C.10.10.1 / C.10.10.7",_tipo_esb_str,"kL/r ≤ 100",slenderness['kl_r']<=100)
        _add_row5("Diam. mín. estribo","C.21.6.4.2",f"Ø{stirrupdiam_saf:.1f} mm","≥ Ø9.53mm (No.3)",stirrupdiam_saf>=9.53)
        if not es_circular:
            _Lo_min=max(h,b,L_col/6.0,45.0)
            _add_row5("Long. confinamiento Lo","C.21.6.4.1",f"Lo={Lo_conf_saf:.1f} cm",f"≥ {_Lo_min:.1f} cm",Lo_conf_saf>=_Lo_min)
        # ── Columna fuerte / Viga débil NSR-10 C.21.6.1 (solo DES/DMO) ────
        if "DES" in nivel_sismico or "DMO" in nivel_sismico:
            _Mn_col_kNm = (float(np.max(cap_x['phi_M_n'])) * factor_fuerza
                             if cap_x and len(cap_x.get('phi_M_n', [])) > 0 else 0.0)
            _sum_Mnc    = 2.0 * _Mn_col_kNm  # 2 columnas en el nodo (sup + inf)
            _Mn_viga_key = "c_pm_Mn_viga_kNm"
            _Mn_viga_input = float(st.session_state.get(_Mn_viga_key, 0.0))
            if _Mn_viga_input > 0:
                _Rd_ratio = _sum_Mnc / _Mn_viga_input
                _Rd_ok    = _Rd_ratio >= 1.2
                _add_row5("Columna fuerte/Viga débil","C.21.6.1",
                          f"ΣMnc/ΣMnv={_Rd_ratio:.2f}","≥ 1.20",_Rd_ok)
                if not _Rd_ok:
                    doc.add_paragraph(
                        f"AVISO C.21.6.1: ΣMnc/ΣMnv={_Rd_ratio:.2f} < 1.20. "
                        "Aumente cuantía o sección (columna fuerte, viga débil)."
                    ).runs[0].bold = True
            else:
                # No se ingresó ΣMnv — omitir verificación, no mostrar NO CUMPLE
                _add_row5("Columna fuerte/Viga débil","C.21.6.1",
                          "Dato no requerido","Ingresar ΣMnv si aplica",None)
        # ── Fin columna fuerte / viga débil ──────────────────────────────
        # ── Verificación Nodo Viga-Columna NSR-10 C.21.7.4.1 ────────────────
        if "DES" in nivel_sismico or "DMO" in nivel_sismico:
            _As_vigas  = float(st.session_state.get("c_pm_As_vigas_nodo", 0.0))
            _Vu_col_nd = float(st.session_state.get("c_pm_Vu_col_nodo", 0.0))
            _conf_nd   = st.session_state.get("c_pm_nodo_conf", "3_caras")
            if _As_vigas > 0:
                _nodo = verificar_nodo_viga_columna(
                    b if not es_circular else D,
                    h if not es_circular else D,
                    fc, fy, _As_vigas, _Vu_col_nd, _conf_nd
                )
                _nodo_label = {
                    "4_caras": "Nodo interior (4 caras)",
                    "3_caras": "Nodo borde (3 caras)",
                    "otros":   "Nodo esquina"
                }.get(_conf_nd, _conf_nd)
                _add_row5(
                    f"Nodo V-C ({_nodo_label})", "C.21.7.4.1",
                    f"Vu_j={_nodo['Vu_j_kN']:.1f}kN | φVn={_nodo['phi_Vn_kN']:.1f}kN "
                    f"(γ={_nodo['gamma']:.2f})",
                    "Vu_j ≤ φVn",
                    _nodo["ok"]
                )
                if not _nodo["ok"]:
                    doc.add_paragraph(
                        f"AVISO C.21.7.4.1: Cortante nodo Vu_j={_nodo['Vu_j_kN']:.1f}kN > "
                        f"φVn={_nodo['phi_Vn_kN']:.1f}kN (ratio={_nodo['ratio']:.2f}). "
                        "Aumente dimensiones de la columna o confinamiento del nodo."
                    ).runs[0].bold = True
            else:
                # No se ingresó As vigas — omitir verificación, no mostrar NO CUMPLE
                _add_row5("Nodo Viga-Columna","C.21.7.4.1",
                          "Dato no requerido","Ingresar As vigas si aplica",None)
        # ── Fin verificación nodo ─────────────────────────────────────────────
        if not bresler['ok']:
            doc.add_paragraph(f"AVISO NO CUMPLE Biaxial: ratio={bresler['ratio']:.3f}. Aumente sección o cuantía.").runs[0].bold=True
        if not (rho_min<=cuantia<=rho_max):
            doc.add_paragraph(f"AVISO NO CUMPLE Cuantía ρ={cuantia:.3f}%. Ajuste barras longitudinales.").runs[0].bold=True
                # 6. Cuantificación y Dosificación
        doc.add_heading("6. CANTIDADES DE OBRA, DOSIFICACIONES Y DESPIECE", level=1)
        doc.add_paragraph(f"Volumen total de concreto: {vol_concreto_m3:.4f} m³")
        doc.add_paragraph(f"Peso total Acero: {peso_total_acero_kg:.1f} kg (Cuantía volc: {peso_total_acero_kg/vol_concreto_m3:.1f} kg/m³)")
        
        # 2. Dosificación aproximada por m3
        p_dosif = doc.add_paragraph(f"\nDOSIFICACIÓN RECOMENDADA EN OBRA PARA f'c={fc:.1f} MPa (Aproximación por m3):\n")
        p_dosif.runs[0].bold = True
        if fc >= 21:
            cemento_kg = 350 * vol_concreto_m3
            arena_m3 = 0.56 * vol_concreto_m3
            grava_m3 = 0.84 * vol_concreto_m3
            agua_lt = 180 * vol_concreto_m3
            sacos = cemento_kg / 50
            doc.add_paragraph(f"• Cemento (50kg): {cemento_kg:.1f} kg (~{sacos:.1f} sacos)")
            doc.add_paragraph(f"• Arena: {arena_m3:.2f} m³")
            doc.add_paragraph(f"• Triturado/Gravilla: {grava_m3:.2f} m³")
            doc.add_paragraph(f"• Agua: {agua_lt:.1f} Litros")
        else:
            doc.add_paragraph("Requiere diseño de mezcla especializado por debajo de 21 MPa.")
        
        # Gráfica de costos APU en modo claro
        try:
            import matplotlib.pyplot as _plt_apu, io as _io_apu
            apu = st.session_state.get("apu_config", {})
            if apu:
                _usar_premix = apu.get("premix", False)
                _cost_acero = peso_total_acero_kg * apu.get("acero", 4200)
                _cost_mo = (peso_total_acero_kg * 0.04 + vol_concreto_m3 * 0.4) * apu.get("costo_dia_mo", 90000)
                _area_enc = (3.14159 * D / 100) * (L_col / 100) if es_circular else (2 * (b + h) / 100) * (L_col / 100)
                _cost_enc = _area_enc * apu.get("encofrado", 45000)
                if _usar_premix:
                    _cats = ["Concreto PM", "Acero", "Mano Obra", "Encofrado"]
                    _cost_conc = vol_concreto_m3 * apu.get("precio_premix_m3", 550000)
                    _vals_apu = [_cost_conc, _cost_acero, _cost_mo, _cost_enc]
                else:
                    _cats = ["Cemento", "Agregados", "Acero", "Mano Obra", "Encofrado"]
                    mix = get_mix_for_fc(fc)
                    bag_kg = 50.0  # Asumimos 50kg genérico para la gráfica
                    _cost_cem = (vol_concreto_m3 * mix["cem"] / bag_kg) * apu.get("cemento", 32000)
                    _cost_agregados = (mix["arena"] * vol_concreto_m3 / 1500) * apu.get("arena", 60000) + (mix["grava"] * vol_concreto_m3 / 1600) * apu.get("grava", 65000)
                    _vals_apu = [_cost_cem, _cost_agregados, _cost_acero, _cost_mo, _cost_enc]
            else:
                _cats=["Concreto","Acero Long.","Estribos"]
                _vals_apu=[
                    vol_concreto_m3*550000,
                    peso_acero_long_kg*4200,
                    peso_total_estribos_kg*4200
                ]
            _mon=st.session_state.get("col_apu_moneda","COP $")
            _fig_apu,_ax_apu=_plt_apu.subplots(figsize=(6,3))
            _fig_apu.patch.set_facecolor('white'); _ax_apu.set_facecolor('#f8f9fa')
            _colores_bar = ['#1565C0','#C62828','#EF6C00','#2E7D32','#6A1B9A','#00838F']
            _bars=_ax_apu.bar(_cats,_vals_apu,color=_colores_bar[:len(_cats)],edgecolor='white',width=0.55)
            _ax_apu.set_title(f"Desglose de Costos APU ({_mon})",color='#1a1a1a',fontsize=10,fontweight='bold')
            _ax_apu.tick_params(colors='#1a1a1a',labelsize=8)
            for sp in _ax_apu.spines.values(): sp.set_edgecolor('#cccccc')
            for _b in _bars:
                _ax_apu.text(_b.get_x()+_b.get_width()/2.,_b.get_height()*1.01,
                    f"{_b.get_height():,.0f}",ha='center',va='bottom',fontsize=7,color='#1a1a1a')
            _buf_apu=_io_apu.BytesIO()
            _fig_apu.savefig(_buf_apu,format='png',dpi=180,bbox_inches='tight',facecolor='white',transparent=False)
            _buf_apu.seek(0); _plt_apu.close(_fig_apu)
            doc.add_picture(_buf_apu,width=Inches(5.5))
        except Exception as _e_apu:
            doc.add_paragraph(f"(Gráfica APU no disponible: {_e_apu})")
        doc.add_paragraph(f"\n3. TABLA DE FIGURADOS Y EMPALMES:\n").runs[0].bold = True
        # Longitud de desarrollo NSR-10: ld_basica con minimo normativo
        ld_mm = lambda d: max(
            (1.3 * 0.02 * fy) / (math.sqrt(fc)) * d,  # formula NSR-10
            300.0,                                       # min 30 cm = 300 mm
            30.0 * d / 10.0                              # min 30db (d en mm, resultado mm)
        )
        _ld_cm = ld_mm(rebar_diam) / 10.0
        _6db_cm = 6.0 * stirrupdiam_saf / 10.0
        _ext_gancho = max(_6db_cm, 7.5)  # NSR-10: min 75mm
        doc.add_paragraph(f"• Ld/Traslapo Barra Principal (Ø{rebar_diam:.1f}mm): {_ld_cm:.0f} cm (Tracción A)")
        doc.add_paragraph(f"• Gancho Sísmico Estribo (Ø{stirrupdiam_saf:.1f}mm): A 135°, Extensión libre 6db={_6db_cm:.1f} cm (mín {_ext_gancho:.1f} cm)")

        
        # 7. Detalles Generales (Imágenes DOCX)
        doc.add_heading("7. PLANOS Y DETALLES CONSTRUCTIVOS", level=1)
        if 'fig_sec' in locals():
            try:
                import io as _io_sec
                buf_sec = _io_sec.BytesIO()
                _orig_sec = fig_sec.get_facecolor()
                fig_sec.patch.set_facecolor('white')
                fig_sec.patch.set_alpha(1.0)
                for _ax in fig_sec.get_axes():
                    _ax.set_facecolor('white')
                fig_sec.savefig(buf_sec, facecolor='white', edgecolor='none', format='png', dpi=200,
                                bbox_inches='tight', transparent=False)
                buf_sec.seek(0)
                fig_sec.patch.set_facecolor(_orig_sec)
                doc.add_picture(buf_sec, width=Inches(4.5))
                buf_sec.close()
            except Exception: pass
                
        try:
            import io as _io_l1
            hook_len_cm = 12 * rebar_diam / 10
            straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
            fig_l1_mem = draw_longitudinal_bar(long_bar_m*100, straight_len_cm,
                                               hook_len_cm, rebar_diam, _bar_label(rebar_diam))
            configurar_pdf_comercial(fig_l1_mem)
            buf_l1 = _io_l1.BytesIO()
            fig_l1_mem.savefig(buf_l1, facecolor='white', format='png', dpi=200,
                               bbox_inches='tight', transparent=False)
            buf_l1.seek(0)   # <-- seek obligatorio para que DOCX lo lea
            plt.close(fig_l1_mem)
            doc.add_picture(buf_l1, width=Inches(5))
            buf_l1.close()
        except Exception as _e_l1:
            doc.add_paragraph(f'[Barra longitudinal: {_e_l1}]')

        if not es_circular:
            try:
                import io as _io_e1
                inside_b = b - 2 * recub_cm
                inside_h = h - 2 * recub_cm
                hook_len_est = 12 * stirrupdiam_saf / 10
                fig_e1_mem = draw_stirrup_with_ties(
                    b, h, recub_cm, hook_len_est,
                    bar_diam_mm=stirrupdiam_saf,
                    n_ties_x=num_flejes_x if not es_circular else 0,
                    n_ties_y=num_flejes_y if not es_circular else 0,
                    nivel_sismico_str=nivel_sismico,
                    bar_name=_bar_label(stirrupdiam_saf))
                configurar_pdf_comercial(fig_e1_mem)
                buf_e1 = _io_e1.BytesIO()
                fig_e1_mem.savefig(buf_e1, facecolor='white', format='png', dpi=200,
                                   bbox_inches='tight', transparent=False)
                buf_e1.seek(0)   # <-- seek obligatorio
                plt.close(fig_e1_mem)
                doc.add_picture(buf_e1, width=Inches(3.5))
                buf_e1.close()
            except Exception as _e_e1:
                doc.add_paragraph(f'[Estribo: {_e_e1}]')
            
        doc.add_heading("8. REFERENCIAS NORMATIVAS", level=1)
        _refs = [
            (f"{code['ref']} — Flexocompresión",   "Hipótesis de diseño — distribución lineal de deformaciones"),
            (f"{code['ref']} — Cuantías",           "Límites de cuantía: ρ_min=1%, ρ_max según nivel sísmico y norma activa"),
            (f"{code['ref']} — Esbeltez",           "Efectos de esbeltez — método de magnificación de momentos δns (no-sway)"),
            (f"{code['ref']} — Factor δns",         "Factor δns = Cm/(1 − Pu/0.75Pc) ≥ 1.0  |  Para sway: usar análisis P-Δ directo"),
            ("NSR-10 C.21.6.3",  "Cuantía máxima zona de rótula plástica — DES: ρ ≤ 4%"),
            ("NSR-10 C.21.6.4",  "Confinamiento en columnas — zona de rótula plástica"),
            ("NSR-10 C.21.6.4.1","Lo_min = max(h, b, Lu/6, 45 cm)"),
            ("NSR-10 C.21.6.4.2","Diámetro mínimo estribo: No. 3 (Ø9.53 mm)"),
            ("NSR-10 C.21.6.4.3","Separación zona confinada: min(8db, 24dst, b/3, 15cm, 6db)"),
            ("NSR-10 C.21.6.4.4","Área mínima estribos Ash — Ec. C.21-4 y C.21-5"),
            ("ACI 318-19 §22.4.2.1","Pn,max = 0.80·Po (estribos rectangulares)"),
            ("ACI 318-19 §22.4.2.2","Pn,max = 0.85·Po (espiral continua)"),
            ("ACI 318-19 §25.7.2","Cálculo Ash para confinamiento con hoop ties"),
            ("ACI 318-19 §25.7.3","Requisitos para refuerzo en espiral continua"),
            ("Bresler (1960)",    "Método recíproco biaxial: 1/Pni = 1/Pnx + 1/Pny − 1/Po"),
            ("ICONTEC NTC 2289", "Dibujo técnico — plumas, escalas y rótulos para planos de construcción"),
        ]
        for _art, _desc in _refs:
            _p = doc.add_paragraph(style='List Bullet')
            _p.add_run(f"{_art}: ").bold = True
            _p.add_run(_desc)

        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)

        st.success(_t("Memoria Generada (Estándar N4 Light-Mode)", "Report generated successfully."))
        st.download_button(label=_t("Descargar Memoria DOCX", "Download DOCX Report"),
                           data=doc_mem, file_name=f"Memoria_Columna_{b:.0f}x{h:.0f}_N4.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.subheader(_t("Exportar Verificaciones a Excel", "Export Verifications to Excel"))
    if st.button(_t("Exportar a Excel", "Export to Excel")):
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            df_verif = pd.DataFrame({
                "Verificación": ["Cuantía", "Biaxial", "Esbeltez", "Ash", "Lo_conf", "Separación"],
                "Valor": [f"{cuantia:.2f}%", f"{bresler['ratio']:.3f}", f"{slenderness['kl_r']:.1f}", 
                          f"{Ash_prov:.3f}/{Ash_req:.3f}" if not es_circular else f"{rho_s_prov:.4f}/{rho_s_req:.4f}",
                          f"{Lo_conf:.1f} cm requerido" if not es_circular else "N/A", f"{s_conf:.1f}" if not es_circular else f"{paso_espiral:.1f}"],
                "Límite": [f"{rho_min}% - {rho_max}%", "≤ 1.0", "≤ 22", "≥ 1.0", "≥ max(b,h,L/6,45)", "≤ 15/20"],
                "Cumple": ["SÍ" if rho_min <= cuantia <= rho_max else "NO",
                           "SÍ" if bresler['ok'] else "NO",
                           "SÍ" if slenderness['kl_r'] <= 22 else "NO",
                           "SÍ" if ash_ok else "NO",
                           "SÍ (por diseño)" if not es_circular else "SÍ",
                           "SÍ" if s_conf <= 15.0 or (es_dmo and s_conf <= 15.0) else "NO"]
            })
            df_verif.to_excel(writer, sheet_name='Verificaciones', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Verificaciones']
            green_format = workbook.add_format({'bg_color': '#C6EFCE', 'font_color': '#006100'})
            red_format = workbook.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006'})
            for row, val in enumerate(df_verif["Cumple"], start=2):
                cell = f"D{row}"
                if val == "SÍ":
                    worksheet.write(cell, val, green_format)
                else:
                    worksheet.write(cell, val, red_format)
            df_despiece.to_excel(writer, sheet_name='Despiece', index=False)
            df_cant = pd.DataFrame({
                "Material": ["Concreto", "Acero Longitudinal", "Acero Estribos", "Acero Total"],
                "Cantidad": [vol_concreto_m3, peso_acero_long_kg, peso_total_estribos_kg, peso_total_acero_kg],
                "Unidad": ["m³", "kg", "kg", "kg"]
            })
            df_cant.to_excel(writer, sheet_name='Cantidades', index=False)
        excel_buffer.seek(0)
        st.download_button(label=_t("Descargar Excel", "Download Excel"),
                           data=excel_buffer, file_name=f"Verificaciones_Columna_{b:.0f}x{h:.0f}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                           
    st.markdown("---")
    st.subheader("Transferir al Presupuesto WBS")
    st.caption("Envia el volumen de concreto y el acero calculado directamente al arbol WBS de la Calculadora de Materiales.")
    
    _wbs_list = st.session_state.get("presupuesto_wbs", [])
    _apus_dict = st.session_state.get("catalogo_apus", {})
    lista_caps = [c["Capitulo"] for c in _wbs_list]
    
    if not lista_caps:
        st.warning("Aun no hay capitulos en el Presupuesto WBS. Ve a la pagina 'Calculadora de Materiales' -> Tab '1. Presupuesto (WBS)' y crea un capitulo primero.")
    elif not _apus_dict:
        st.warning("El catalogo de APUs esta vacio. Ve a 'Calculadora de Materiales' -> Tab '3. Gestor de APUs'.")
    else:
        col_wbs1, col_wbs2 = st.columns([1, 1])
        
        with col_wbs1:
            st.markdown("**1. Enviar Volumen de Concreto**")
            st.metric("Volumen calculado", f"{vol_concreto_m3:.3f} m³")
            dest_cap_c = st.selectbox("Capitulo destino", lista_caps, key="wbs_cap_col_c")
            dest_apu_c = st.selectbox("APU a aplicar", list(_apus_dict.keys()),
                                       format_func=lambda x: _apus_dict[x]["Nombre"], key="wbs_apu_col_c")
            btn_wbs_c = st.button("Enviar Concreto al WBS", type="primary", use_container_width=True, key="btn_send_conc_col")

        with col_wbs2:
            st.markdown("**2. Enviar Acero de Refuerzo**")
            st.metric("Acero total calculado", f"{peso_total_acero_kg:.1f} kg")
            dest_cap_s = st.selectbox("Capitulo destino", lista_caps, key="wbs_cap_col_s")
            dest_apu_s = st.selectbox("APU a aplicar", list(_apus_dict.keys()),
                                       format_func=lambda x: _apus_dict[x]["Nombre"], key="wbs_apu_col_s")
            btn_wbs_s = st.button("Enviar Acero al WBS", type="primary", use_container_width=True, key="btn_send_acero_col")

        if btn_wbs_c:
            _transferido = False
            for cap in st.session_state["presupuesto_wbs"]:
                if cap["Capitulo"] == dest_cap_c:
                    nid = f"{len(cap['Items'])+1:03d}"
                    cap["Items"].append({"ID_Item": nid, "ID_APU": dest_apu_c,
                                         "Cantidad": float(vol_concreto_m3), "Predecesora": ""})
                    _transferido = True
                    break
            if _transferido:
                st.success(f"Concreto ({vol_concreto_m3:.3f} m3) enviado al capitulo '{dest_cap_c}'. Ve a la Calculadora de Materiales -> Tab WBS para verlo.")
            else:
                st.error("No se encontro el capitulo seleccionado en session_state.")

        if btn_wbs_s:
            _transferido = False
            for cap in st.session_state["presupuesto_wbs"]:
                if cap["Capitulo"] == dest_cap_s:
                    nid = f"{len(cap['Items'])+1:03d}"
                    cap["Items"].append({"ID_Item": nid, "ID_APU": dest_apu_s,
                                         "Cantidad": float(peso_total_acero_kg), "Predecesora": ""})
                    _transferido = True
                    break
            if _transferido:
                st.success(f"Acero ({peso_total_acero_kg:.1f} kg) enviado al capitulo '{dest_cap_s}'. Ve a la Calculadora de Materiales -> Tab WBS para verlo.")
            else:
                st.error("No se encontro el capitulo seleccionado en session_state.")

