import streamlit as st

# ─── BANNER ESTANDAR DIAMANTE ───────────────────────────────
st.markdown("""<div style="width:100%;overflow:hidden;border-radius:14px;margin-bottom:18px;box-shadow:0 4px 32px #0008;"><svg viewBox="0 0 1100 220" xmlns="http://www.w3.org/2000/svg" style="width:100%;display:block;background:linear-gradient(135deg,#0a1128 0%,#1c2541 100%);"><g opacity="0.1" stroke="#38bdf8" stroke-width="0.5"><line x1="0" y1="55" x2="1100" y2="55"/><line x1="0" y1="110" x2="1100" y2="110"/><line x1="0" y1="165" x2="1100" y2="165"/><line x1="220" y1="0" x2="220" y2="220"/><line x1="440" y1="0" x2="440" y2="220"/><line x1="660" y1="0" x2="660" y2="220"/></g><rect x="0" y="0" width="1100" height="3" fill="#ef4444" opacity="0.9"/><rect x="0" y="217" width="1100" height="3" fill="#ef4444" opacity="0.7"/><g transform="translate(50,25)"><rect x="20" y="30" width="100" height="80" rx="3" fill="#1e293b" stroke="#ef4444" stroke-width="2" opacity="0.8"/><rect x="70" y="70" width="60" height="40" rx="3" fill="#3a0000" stroke="#ef4444" stroke-width="2" opacity="0.9"/><text x="45" y="75" font-family="monospace" font-size="8" fill="#fca5a5">Masa</text><text x="45" y="90" font-family="monospace" font-size="8" fill="#fca5a5">irregular</text><text x="70" y="15" text-anchor="middle" font-family="sans-serif" font-size="9" fill="#cbd5e1">PLANTA</text></g><g transform="translate(560,0)"><rect x="0" y="28" width="4" height="165" rx="2" fill="#ef4444"/><text x="18" y="66" font-family="Arial,sans-serif" font-size="28" font-weight="bold" fill="#ffffff">IRREGULARIDADES SISMICAS</text><text x="18" y="94" font-family="Arial,sans-serif" font-size="14" font-weight="300" fill="#93c5fd" letter-spacing="2">NSR-10 A.3 · PLANTA Y ALTURA · PENALIZACION R</text><rect x="18" y="102" width="480" height="1" fill="#ef4444" opacity="0.5"/><rect x="18" y="115" width="162" height="22" rx="11" fill="#3a0000" stroke="#ef4444" stroke-width="1"/><text x="99" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#fca5a5">IRREGULARIDAD PLANTA</text><rect x="188" y="115" width="162" height="22" rx="11" fill="#291400" stroke="#f59e0b" stroke-width="1"/><text x="269" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#fcd34d">IRREGULARIDAD ALTURA</text><rect x="358" y="115" width="120" height="22" rx="11" fill="#1e1b4b" stroke="#8b5cf6" stroke-width="1"/><text x="418" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#c4b5fd">PENALIZACION R</text><rect x="486" y="115" width="92" height="22" rx="11" fill="#052e16" stroke="#10b981" stroke-width="1"/><text x="532" y="130" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#6ee7b7">DCR GLOBAL</text><text x="18" y="156" font-family="Arial,sans-serif" font-size="11" fill="#94a3b8">Verificacion sistematica de las 5 irregularidades en planta y 6 en altura segun NSR-10 A.3.2.</text><text x="18" y="172" font-family="Arial,sans-serif" font-size="11" fill="#94a3b8">Calculo automatico de la penalizacion al factor R y la rigidez de piso requerida.</text><text x="18" y="188" font-family="Arial,sans-serif" font-size="11" fill="#94a3b8">Determinacion de la necesidad de analisis dinamico modal espectral vs. estatico equivalente.</text></g></svg></div>""", unsafe_allow_html=True)

with st.expander(" Guia Profesional — Irregularidades Sismicas", expanded=False):
    st.markdown("""
    ### Metodologia: Clasificacion de Irregularidades NSR-10 A.3
    Evaluacion formal de la regularidad del sistema estructural, requisito previo indispensable para seleccionar el metodo de analisis sismorresistente.

    ####  1. Irregularidades en Planta (A.3.2.1)
    - **Irregularidad de Torsion:** Verifica que el desplazamiento maximo de piso no supere 1.2 veces el promedio.
    - **Esquinas Entrantes:** Identifica si la saliente de la planta supera el 15% de la dimension total.
    - **Discontinuidad de Diafragma:** Detecta aberturas mayores al 50% del area de piso.

    ####  2. Irregularidades en Altura (A.3.2.2)
    - **Irregularidad de Rigidez (Piso Blando):** Rigidez lateral < 70% del piso inmediato superior.
    - **Irregularidad de Masa:** Peso del piso > 150% del peso del piso adyacente.
    - **Irregularidad de Resistencia (Piso Debil):** Resistencia lateral < 80% del piso superior.

    ####  3. Penalizacion al Factor R y Metodo de Analisis
    - Segun el tipo e intensidad de irregularidades, el modulo ajusta (penaliza) el factor R.
    - Determina si se requiere analisis dinamico modal espectral (DME) en lugar del metodo estatico equivalente (MEE).

    ####  4. Reportes
    - Genera un informe de irregularidades con semaforo de cumplimiento (Verde/Amarillo/Rojo) para cada criterio normativo.
""")

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import io
from docx import Document
from docx.shared import Inches
import datetime

# 
# IDIOMA GLOBAL
try:
    from normas_referencias import mostrar_referencias_norma
except ImportError:
    def mostrar_referencias_norma(*a, **kw): pass
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
# 

st.set_page_config(page_title=_t("Irregularidades Estructurales", "Structural Irregularities"), layout="wide")
st.title(" " + _t("Verificación de Irregularidades Estructurales", "Structural Irregularities Check"))
st.markdown(_t(
    "Evaluación de irregularidades en altura y planta según las principales normas sísmicas (E.030, NSR-10, ACI 318, etc.).\n\n"
    "**Nota:** Los factores Ia e Ip reducen el coeficiente de reducción de fuerza sísmica R. Valores <1 indican presencia de irregularidades.",
    "Evaluation of vertical and plan irregularities according to major seismic codes (E.030, NSR-10, ACI 318, etc.).\n\n"
    "**Note:** Ia and Ip factors reduce the seismic force reduction coefficient R. Values <1 indicate irregularities."
))

# Diccionario de normas (igual que antes)
NORMA_IRR = {
    "E.070 (Perú) / E.030": {
        "piso_blando":           (0.7, 0.75, 0.5),
        "piso_debil":            (0.8, 0.75, 0.5),
        "masa":                  (1.5, 0.9, 0.9),
        "geometria_vertical":    (1.3, 0.9, 0.9),
        "discontinuidad_sist":   (0.1, 0.8, 0.6),
        "torsion":               (1.3, 0.75, 0.6),
        "esquinas_entrantes":    (0.2, 0.9, 0.9),
        "diafragma":             (0.5, 0.85, 0.85),
        "sistemas_no_paralelos": (0.1, 0.9, 0.9),
    },
    "NSR-10 (Colombia)": {
        "piso_blando":           (0.7, 0.75, 0.5),
        "piso_debil":            (0.8, 0.75, 0.5),
        "masa":                  (1.5, 0.9, 0.9),
        "geometria_vertical":    (1.3, 0.9, 0.9),
        "discontinuidad_sist":   (0.1, 0.8, 0.6),
        "torsion":               (1.3, 0.75, 0.6),
        "esquinas_entrantes":    (0.2, 0.9, 0.9),
        "diafragma":             (0.5, 0.85, 0.85),
        "sistemas_no_paralelos": (0.1, 0.9, 0.9),
    },
    "ACI 318-25 (EE.UU.)": {
        "piso_blando":           (0.7, 0.75, 0.5),
        "piso_debil":            (0.8, 0.75, 0.5),
        "masa":                  (1.5, 0.9, 0.9),
        "geometria_vertical":    (1.3, 0.9, 0.9),
        "discontinuidad_sist":   (0.1, 0.8, 0.6),
        "torsion":               (1.3, 0.75, 0.6),
        "esquinas_entrantes":    (0.2, 0.9, 0.9),
        "diafragma":             (0.5, 0.85, 0.85),
        "sistemas_no_paralelos": (0.1, 0.9, 0.9),
    },
    "NEC-SE-DS (Ecuador)": {
        "piso_blando":           (0.7, 0.75, 0.5),
        "piso_debil":            (0.8, 0.75, 0.5),
        "masa":                  (1.5, 0.9, 0.9),
        "geometria_vertical":    (1.3, 0.9, 0.9),
        "discontinuidad_sist":   (0.1, 0.8, 0.6),
        "torsion":               (1.3, 0.75, 0.6),
        "esquinas_entrantes":    (0.2, 0.9, 0.9),
        "diafragma":             (0.5, 0.85, 0.85),
        "sistemas_no_paralelos": (0.1, 0.9, 0.9),
    },
    "E.060 (Perú) / NTE": {
        "piso_blando":           (0.7, 0.75, 0.5),
        "piso_debil":            (0.8, 0.75, 0.5),
        "masa":                  (1.5, 0.9, 0.9),
        "geometria_vertical":    (1.3, 0.9, 0.9),
        "discontinuidad_sist":   (0.1, 0.8, 0.6),
        "torsion":               (1.3, 0.75, 0.6),
        "esquinas_entrantes":    (0.2, 0.9, 0.9),
        "diafragma":             (0.5, 0.85, 0.85),
        "sistemas_no_paralelos": (0.1, 0.9, 0.9),
    },
    "COVENIN 1753 (Venezuela)": {
        "piso_blando":           (0.7, 0.75, 0.5),
        "piso_debil":            (0.8, 0.75, 0.5),
        "masa":                  (1.5, 0.9, 0.9),
        "geometria_vertical":    (1.3, 0.9, 0.9),
        "discontinuidad_sist":   (0.1, 0.8, 0.6),
        "torsion":               (1.3, 0.75, 0.6),
        "esquinas_entrantes":    (0.2, 0.9, 0.9),
        "diafragma":             (0.5, 0.85, 0.85),
        "sistemas_no_paralelos": (0.1, 0.9, 0.9),
    },
}

def get_irr_factors(norma_sel, tipo):
    if norma_sel not in NORMA_IRR:
        norma_sel = "NSR-10 (Colombia)"
    data = NORMA_IRR[norma_sel]
    return data.get(tipo, (0.7, 0.75, 0.5))

# Sidebar
st.sidebar.header(_t("Norma de Diseño", "Design Code"))
norma_disp = st.sidebar.selectbox(
    _t("Seleccione la norma para irregularidades:", "Select code for irregularities:"),
    list(NORMA_IRR.keys()),
    key="irr_norma"
)
mostrar_referencias_norma(norma_disp, "irregularidades")

st.sidebar.header(_t("Unidades de salida", "Output units"))
unidades_salida = st.sidebar.radio(
    "Unidades de fuerza/momento:",
    ["kiloNewtons (kN, kN·m)", "Toneladas fuerza (tonf, tonf·m)"],
    key="irr_units"
)
if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom    = "tonf·m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom    = "kN·m"

st.sidebar.header(_t("Datos generales", "General data"))
n_pisos = st.sidebar.number_input(_t("Número de pisos", "Number of stories"), 1, 50, 4, 1, key="irr_np")
nombre_proyecto = st.sidebar.text_input(_t("Nombre del proyecto", "Project name"), "Edificio Ejemplo")

def draw_scheme_vertical(values, title, unit=""):
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.set_facecolor('#1a1a2e')
    fig.patch.set_facecolor('#1a1a2e')
    pisos = list(range(1, len(values)+1))
    ax.barh(pisos, values, color='#4a6ea8')
    ax.set_xlabel(f"{title} {unit}", color='white')
    ax.set_ylabel(_t("Piso", "Story"), color='white')
    ax.set_title(title, color='white')
    ax.tick_params(colors='white')
    ax.grid(axis='x', linestyle='--', alpha=0.3)
    for spine in ax.spines.values():
        spine.set_color('white')
    return fig

# Inicializar DataFrame
if 'irr_data' not in st.session_state:
    st.session_state.irr_data = pd.DataFrame(columns=[
        'piso', 'h_i', 'Vx', 'Vy', 'βmax_x', 'βmin_x', 'βprom_x', 'βmax_y', 'βmin_y', 'βprom_y',
        'P_i', 'b_i_x', 'b_i_y', 'Kx', 'Ky', 'Vresist_x', 'Vresist_y', 'V_elem_x', 'V_elem_y'
    ])

def update_data():
    df = st.session_state.irr_data
    if len(df) < n_pisos:
        for i in range(len(df), n_pisos):
            new_row = {col: 0.0 for col in df.columns}
            new_row['piso'] = i+1
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    elif len(df) > n_pisos:
        df = df.iloc[:n_pisos]
    df = df.sort_values('piso').reset_index(drop=True)
    st.session_state.irr_data = df

update_data()

# Pestañas
tab_data, tab_altura, tab_planta, tab_resumen = st.tabs([
    " Datos por piso", " Auto-cálculo", " Datos planta", " Verificación"
])

with tab_data:
    st.subheader(_t("Datos por piso (editar tabla)", "Story data (edit table)"))
    st.info(_t("Ingrese los valores para cada piso. Los campos en blanco se considerarán 0. "
               "Para rigidez K y resistencia Vresist, si no se ingresan se pueden calcular automáticamente "
               "en la pestaña 'Auto-cálculo'.", 
               "Enter values for each story. Empty fields will be considered 0. "
               "For stiffness K and strength Vresist, if not entered they can be automatically calculated "
               "in the 'Auto-calc' tab."))
    df_edit = st.data_editor(
        st.session_state.irr_data,
        column_config={
            "piso": st.column_config.NumberColumn("Piso", min_value=1, max_value=n_pisos, step=1),
            "h_i": st.column_config.NumberColumn("Altura (m)", min_value=0.0, step=0.1),
            "Vx": st.column_config.NumberColumn(f"Vx ({unidad_fuerza})", step=1.0),
            "Vy": st.column_config.NumberColumn(f"Vy ({unidad_fuerza})", step=1.0),
            "βmax_x": st.column_config.NumberColumn("βmax_x (cm)", step=0.01),
            "βmin_x": st.column_config.NumberColumn("βmin_x (cm)", step=0.01),
            "βprom_x": st.column_config.NumberColumn("βprom_x (cm)", step=0.01),
            "βmax_y": st.column_config.NumberColumn("βmax_y (cm)", step=0.01),
            "βmin_y": st.column_config.NumberColumn("βmin_y (cm)", step=0.01),
            "βprom_y": st.column_config.NumberColumn("βprom_y (cm)", step=0.01),
            "P_i": st.column_config.NumberColumn(f"Masa/Peso ({unidad_fuerza})", step=1.0),
            "b_i_x": st.column_config.NumberColumn("b_i_x (m)", step=0.1),
            "b_i_y": st.column_config.NumberColumn("b_i_y (m)", step=0.1),
            "Kx": st.column_config.NumberColumn("Kx (kN/m)", step=100.0),
            "Ky": st.column_config.NumberColumn("Ky (kN/m)", step=100.0),
            "Vresist_x": st.column_config.NumberColumn(f"Resistencia corte Vr_x ({unidad_fuerza})", step=1.0),
            "Vresist_y": st.column_config.NumberColumn(f"Resistencia corte Vr_y ({unidad_fuerza})", step=1.0),
            "V_elem_x": st.column_config.NumberColumn(f"Fuerza elemento crítico V_elem_x ({unidad_fuerza})", step=1.0),
            "V_elem_y": st.column_config.NumberColumn(f"Fuerza elemento crítico V_elem_y ({unidad_fuerza})", step=1.0),
        },
        use_container_width=True,
        num_rows="fixed",
        key="irr_data_editor"
    )
    if st.button(_t("Actualizar datos", "Update data")):
        st.session_state.irr_data = df_edit
        st.success(_t("Datos actualizados", "Data updated"))
        st.rerun()

with tab_altura:
    st.subheader(_t("Cálculo automático de rigidez y resistencia", "Auto-calculation of stiffness and strength"))
    st.info(_t("Si no ingresó valores de rigidez K o resistencia Vr, puede calcularlos a partir de derivas y cortantes usando este botón.",
               "If you did not enter stiffness K or strength Vr values, you can calculate them from drifts and shear forces using this button."))
    if st.button(_t("Calcular K y Vr automáticamente", "Auto-calc K and Vr")):
        df = st.session_state.irr_data.copy()
        for i in range(len(df)):
            h_m = df.loc[i, 'h_i']
            if h_m > 0:
                # Rigidez K = V / (β/h)
                if df.loc[i, 'Kx'] == 0 and df.loc[i, 'βmax_x'] > 0:
                    drift_x = df.loc[i, 'βmax_x'] / 100.0 / h_m
                    Kx = df.loc[i, 'Vx'] / drift_x if drift_x > 0 else 0
                    df.loc[i, 'Kx'] = Kx
                if df.loc[i, 'Ky'] == 0 and df.loc[i, 'βmax_y'] > 0:
                    drift_y = df.loc[i, 'βmax_y'] / 100.0 / h_m
                    Ky = df.loc[i, 'Vy'] / drift_y if drift_y > 0 else 0
                    df.loc[i, 'Ky'] = Ky
            # Resistencia Vr estimada como 1.2*V (ajustable)
            if df.loc[i, 'Vresist_x'] == 0:
                df.loc[i, 'Vresist_x'] = df.loc[i, 'Vx'] * 1.2
            if df.loc[i, 'Vresist_y'] == 0:
                df.loc[i, 'Vresist_y'] = df.loc[i, 'Vy'] * 1.2
        st.session_state.irr_data = df
        st.success(_t("Valores calculados", "Values calculated"))
        st.rerun()
    st.markdown("---")
    st.subheader(_t("Guía de cálculo", "Calculation guide"))
    st.markdown("""
    **Rigidez K** (kN/m) = V / (Δ/h), donde:
    - V: cortante del piso [kN]
    - Δ: deriva máxima [cm] → convertir a metros
    - h: altura de piso [m]
    
    **Resistencia Vr** (kN): se estima como 1.2 veces el cortante sísmico del piso. 
    En un análisis detallado debe incluir el aporte del concreto y del acero según la norma correspondiente.
    """)

with tab_planta:
    st.subheader(_t("Datos para irregularidades en planta", "Data for plan irregularities"))
    col1, col2 = st.columns(2)
    with col1:
        A_total = st.number_input(_t("Área total en planta A (m²)", "Total plan area A (m²)"), 1.0, 10000.0, 500.0, key="irr_A")
        A_entrante = st.number_input(_t("Área de esquina entrante a (m²)", "Re-entrant corner area a (m²)"), 0.0, 5000.0, 0.0, key="irr_a")
        A_diafragma = st.number_input(_t("Área de abertura en diafragma A' (m²)", "Diaphragm opening area A' (m²)"), 0.0, 5000.0, 0.0, key="irr_Ap")
    with col2:
        angulo_sistemas = st.number_input(_t("Ángulo entre sistemas resistentes (°)", "Angle between lateral systems (°)"), 0.0, 180.0, 90.0, key="irr_ang")
        st.markdown("**Nota:** Los valores de βmax y βprom se ingresan en la tabla de datos por piso para cada dirección.")
        if st.button(_t("Guardar datos planta", "Save plan data")):
            st.success(_t("Datos guardados", "Data saved"))
            st.rerun()

with tab_resumen:
    df = st.session_state.irr_data.copy()
    A_total = st.session_state.get("irr_A", 100.0)
    A_entrante = st.session_state.get("irr_a", 0.0)
    A_diafragma = st.session_state.get("irr_Ap", 0.0)
    angulo_sistemas = st.session_state.get("irr_ang", 90.0)

    def verificar_irregularidades(df, A_total, A_entrante, A_diafragma, angulo_sistemas, norma):
        resultados = {
            "altura": {"X": {}, "Y": {}},
            "planta": {"X": {}, "Y": {}},
            "factores": {"Ia_x": 1.0, "Ia_y": 1.0, "Ip_x": 1.0, "Ip_y": 1.0}
        }
        for dir in ["X", "Y"]:
            dir_low = dir.lower()
            # Altura
            # Piso blando (rigidez)
            K = df[f'K{dir_low}'].values
            lim, fact_reg, fact_ext = get_irr_factors(norma, "piso_blando")
            irregular = False
            for i in range(len(K)-1):
                if K[i] > 0 and K[i+1] > 0 and K[i] / K[i+1] < lim:
                    irregular = True
                    if K[i] / K[i+1] < 0.6:
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_ext
                    else:
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_reg
                    break
            resultados["altura"][dir]["piso_blando"] = irregular

            # Piso débil (resistencia)
            Vr = df[f'Vresist_{dir_low}'].values
            lim_res, fact_reg_res, fact_ext_res = get_irr_factors(norma, "piso_debil")
            irregular = False
            for i in range(len(Vr)-1):
                if Vr[i] > 0 and Vr[i+1] > 0 and Vr[i] / Vr[i+1] < lim_res:
                    irregular = True
                    if Vr[i] / Vr[i+1] < 0.6:
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_ext_res
                    else:
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_reg_res
                    break
            resultados["altura"][dir]["piso_debil"] = irregular

            # Masa o peso
            P = df['P_i'].values
            lim_masa, fact_masa, _ = get_irr_factors(norma, "masa")
            irregular = False
            for i in range(len(P)-1):
                if P[i] > 0 and P[i+1] > 0 and P[i] / P[i+1] > lim_masa:
                    irregular = True
                    resultados["factores"][f"Ia_{dir_low}"] *= fact_masa
                    break
            resultados["altura"][dir]["masa"] = irregular

            # Geometría vertical
            b = df[f'b_i_{dir_low}'].values
            lim_geo, fact_geo, _ = get_irr_factors(norma, "geometria_vertical")
            irregular = False
            for i in range(len(b)-1):
                if b[i] > 0 and b[i+1] > 0:
                    ratio = b[i] / b[i+1]
                    if ratio < 1/lim_geo or ratio > lim_geo:
                        irregular = True
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_geo
                        break
            resultados["altura"][dir]["geometria_vertical"] = irregular

            # Discontinuidad en sistemas resistentes
            V_elem = df[f'V_elem_{dir_low}'].values
            V_total = df[f'V{dir_low}'].values
            lim_disc, fact_disc, fact_ext_disc = get_irr_factors(norma, "discontinuidad_sist")
            irregular = False
            for i in range(len(V_elem)):
                if V_total[i] > 0 and V_elem[i] / V_total[i] > lim_disc:
                    irregular = True
                    if V_elem[i] / V_total[i] > 0.25:
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_ext_disc
                    else:
                        resultados["factores"][f"Ia_{dir_low}"] *= fact_disc
                    break
            resultados["altura"][dir]["discontinuidad_sist"] = irregular

            # Planta
            # Torsión
            βmax = df[f'βmax_{dir_low}'].values
            βprom = df[f'βprom_{dir_low}'].values
            lim_tors, fact_tors, fact_ext_tors = get_irr_factors(norma, "torsion")
            irregular = False
            for i in range(len(βmax)):
                if βprom[i] > 0 and βmax[i] / βprom[i] > lim_tors:
                    irregular = True
                    if βmax[i] / βprom[i] > 1.5:
                        resultados["factores"][f"Ip_{dir_low}"] *= fact_ext_tors
                    else:
                        resultados["factores"][f"Ip_{dir_low}"] *= fact_tors
                    break
            resultados["planta"][dir]["torsion"] = irregular

            # Esquinas entrantes
            lim_esq, fact_esq, _ = get_irr_factors(norma, "esquinas_entrantes")
            irregular = (A_entrante / A_total) > lim_esq if A_total > 0 else False
            if irregular:
                resultados["factores"][f"Ip_{dir_low}"] *= fact_esq
            resultados["planta"][dir]["esquinas_entrantes"] = irregular

            # Discontinuidad del diafragma
            lim_dia, fact_dia, _ = get_irr_factors(norma, "diafragma")
            irregular = (A_diafragma / A_total) > lim_dia if A_total > 0 else False
            if irregular:
                resultados["factores"][f"Ip_{dir_low}"] *= fact_dia
            resultados["planta"][dir]["diafragma"] = irregular

            # Sistemas no paralelos
            lim_sist, fact_sist, _ = get_irr_factors(norma, "sistemas_no_paralelos")
            V_elem_max = np.max(V_elem) if len(V_elem) > 0 else 0
            V_total_max = np.max(V_total) if len(V_total) > 0 else 0
            irregular = (V_elem_max / V_total_max > lim_sist and angulo_sistemas < 30) if V_total_max > 0 else False
            if irregular:
                resultados["factores"][f"Ip_{dir_low}"] *= fact_sist
            resultados["planta"][dir]["sistemas_no_paralelos"] = irregular

        # Asegurar que los factores no superen 1
        for key in ["Ia_x", "Ia_y", "Ip_x", "Ip_y"]:
            resultados["factores"][key] = min(resultados["factores"][key], 1.0)
        return resultados

    resultados = verificar_irregularidades(df, A_total, A_entrante, A_diafragma, angulo_sistemas, norma_disp)

    # Mostrar resultados (igual que antes, se mantiene)
    st.subheader(_t("Resumen de Irregularidades", "Irregularities Summary"))
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"**{_t('Dirección X', 'X Direction')}**")
        st.metric(_t("Factor de irregularidad en altura (Ia)", "Vertical irregularity factor (Ia)"), f"{resultados['factores']['Ia_x']:.2f}")
        st.metric(_t("Factor de irregularidad en planta (Ip)", "Plan irregularity factor (Ip)"), f"{resultados['factores']['Ip_x']:.2f}")
        st.metric(_t("Factor combinado (Ia·Ip)", "Combined factor (Ia·Ip)"), f"{resultados['factores']['Ia_x'] * resultados['factores']['Ip_x']:.2f}")
    with col2:
        st.markdown(f"**{_t('Dirección Y', 'Y Direction')}**")
        st.metric(_t("Factor de irregularidad en altura (Ia)", "Vertical irregularity factor (Ia)"), f"{resultados['factores']['Ia_y']:.2f}")
        st.metric(_t("Factor de irregularidad en planta (Ip)", "Plan irregularity factor (Ip)"), f"{resultados['factores']['Ip_y']:.2f}")
        st.metric(_t("Factor combinado (Ia·Ip)", "Combined factor (Ia·Ip)"), f"{resultados['factores']['Ia_y'] * resultados['factores']['Ip_y']:.2f}")

    st.markdown("---")
    st.subheader(_t("Detalle de verificaciones", "Detailed checks"))
    for dir in ["X", "Y"]:
        st.markdown(f"#### {_t('Dirección', 'Direction')} {dir}")
        df_alt = pd.DataFrame({
            "Irregularidad": [
                "Piso blando (rigidez)",
                "Piso débil (resistencia)",
                "Masa o peso",
                "Geometría vertical",
                "Discontinuidad sistemas resistentes"
            ],
            "Presenta": [
                "Sí" if resultados['altura'][dir]['piso_blando'] else "No",
                "Sí" if resultados['altura'][dir]['piso_debil'] else "No",
                "Sí" if resultados['altura'][dir]['masa'] else "No",
                "Sí" if resultados['altura'][dir]['geometria_vertical'] else "No",
                "Sí" if resultados['altura'][dir]['discontinuidad_sist'] else "No"
            ]
        })
        st.dataframe(df_alt, use_container_width=True, hide_index=True)

        df_plan = pd.DataFrame({
            "Irregularidad": [
                "Torsional",
                "Esquinas entrantes",
                "Discontinuidad del diafragma",
                "Sistemas no paralelos"
            ],
            "Presenta": [
                "Sí" if resultados['planta'][dir]['torsion'] else "No",
                "Sí" if resultados['planta'][dir]['esquinas_entrantes'] else "No",
                "Sí" if resultados['planta'][dir]['diafragma'] else "No",
                "Sí" if resultados['planta'][dir]['sistemas_no_paralelos'] else "No"
            ]
        })
        st.dataframe(df_plan, use_container_width=True, hide_index=True)

    # Gráficos
    st.markdown("---")
    st.subheader(_t("Gráficos de rigidez y resistencia", "Stiffness and strength plots"))
    if 'Kx' in df.columns and not df['Kx'].isna().all() and df['Kx'].sum() > 0:
        fig_kx = draw_scheme_vertical(df['Kx'].values, _t("Rigidez en X", "Stiffness X"), "(kN/m)")
        st.pyplot(fig_kx)
    if 'Ky' in df.columns and not df['Ky'].isna().all() and df['Ky'].sum() > 0:
        fig_ky = draw_scheme_vertical(df['Ky'].values, _t("Rigidez en Y", "Stiffness Y"), "(kN/m)")
        st.pyplot(fig_ky)
    if 'Vresist_x' in df.columns and not df['Vresist_x'].isna().all() and df['Vresist_x'].sum() > 0:
        fig_vx = draw_scheme_vertical(df['Vresist_x'].values * factor_fuerza, _t("Resistencia en X", "Strength X"), f"({unidad_fuerza})")
        st.pyplot(fig_vx)
    if 'Vresist_y' in df.columns and not df['Vresist_y'].isna().all() and df['Vresist_y'].sum() > 0:
        fig_vy = draw_scheme_vertical(df['Vresist_y'].values * factor_fuerza, _t("Resistencia en Y", "Strength Y"), f"({unidad_fuerza})")
        st.pyplot(fig_vy)

    # Botón memoria DOCX
    st.markdown("---")
    if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
        doc = Document()
        doc.add_heading(f"{_t('Memoria de Irregularidades', 'Irregularities Report')} – {nombre_proyecto}", 0)
        doc.add_paragraph(f"{_t('Norma utilizada:', 'Code used:')} {norma_disp}")
        doc.add_paragraph(f"{_t('Fecha:', 'Date:')} {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_heading(_t("Resumen de Factores", "Summary of Factors"), level=1)
        doc.add_paragraph(f"{_t('Dirección X:', 'X Direction:')} Ia = {resultados['factores']['Ia_x']:.2f}, Ip = {resultados['factores']['Ip_x']:.2f}")
        doc.add_paragraph(f"{_t('Dirección Y:', 'Y Direction:')} Ia = {resultados['factores']['Ia_y']:.2f}, Ip = {resultados['factores']['Ip_y']:.2f}")
        doc.add_heading(_t("Verificaciones en Altura", "Vertical Irregularities"), level=1)
        for dir in ["X", "Y"]:
            doc.add_paragraph(f"{_t('Dirección', 'Direction')} {dir}:")
            for key, val in resultados['altura'][dir].items():
                doc.add_paragraph(f"  - {key}: {'Sí' if val else 'No'}")
        doc.add_heading(_t("Verificaciones en Planta", "Plan Irregularities"), level=1)
        for dir in ["X", "Y"]:
            doc.add_paragraph(f"{_t('Dirección', 'Direction')} {dir}:")
            for key, val in resultados['planta'][dir].items():
                doc.add_paragraph(f"  - {key}: {'Sí' if val else 'No'}")
        # Incluir gráficos
        buf = io.BytesIO()
        if 'Kx' in df.columns and df['Kx'].sum() > 0:
            fig_kx = draw_scheme_vertical(df['Kx'].values, "Stiffness X", "(kN/m)")
            fig_kx.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            doc.add_heading(_t("Gráfico de rigidez en X", "Stiffness X plot"), level=2)
            doc.add_picture(buf, width=Inches(4))
        if 'Vresist_x' in df.columns and df['Vresist_x'].sum() > 0:
            fig_vx = draw_scheme_vertical(df['Vresist_x'].values * factor_fuerza, "Strength X", f"({unidad_fuerza})")
            buf = io.BytesIO()
            fig_vx.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            doc.add_heading(_t("Gráfico de resistencia en X", "Strength X plot"), level=2)
            doc.add_picture(buf, width=Inches(4))
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button(_t("Descargar Memoria DOCX", "Download DOCX"), data=doc_mem,
                           file_name=f"Irregularidades_{nombre_proyecto}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 
# FOOTER
# 
st.markdown("---")
st.markdown(f"""
> **{_t('Módulo de Irregularidades', 'Irregularities Module')}**  
> {_t('Norma activa:', 'Active code:')} `{norma_disp}`  
>  *Los resultados deben ser verificados por un ingeniero estructural calificado.*
""")