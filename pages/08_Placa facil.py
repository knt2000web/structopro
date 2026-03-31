import streamlit as st
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import pandas as pd
import math
import io
import ezdxf
from ezdxf.math import Vec2
from ezdxf import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib import colors as rl_colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Placa Fácil", "Easy Slab"), layout="wide")
st.title(_t("Placa Fácil – Sistema de Vigueta y Bloques", "Easy Slab – Joist & Block System"))
st.markdown(_t("Diseño de losas con vigueta metálica y bloques de concreto (Placa Fácil). Verificación según NSR-10 (Colombia) y normas internacionales.", 
               "Design of slabs with metal joists and concrete blocks (Easy Slab). Verification according to NSR-10 (Colombia) and international codes."))

BLOCK_DATA = {
    "nombre": "Bloquelón Santafé",
    "largo": 0.80, "ancho": 0.23, "alto": 0.08,
    "peso_unitario": 13, "rendimiento_por_m2": 5.18,
    "color": "#C19A6B", "pestana": 0.015,
}
PROFILE_DATA = {
    "nombre": "Perfil Colmena",
    "alto_total": 0.080, "ancho_alma": 0.040, "espesor_alma": 0.002,
    "ancho_ala": 0.025, "espesor_ala": 0.002,
    "peso_por_m": 9.8, "color": "#C0C0C0",
}
MESH_DATA = {
    "nombre": "Malla Electrosoldada Q2",
    "diametro": 0.00635, "espaciado_largo": 0.25,
    "espaciado_corto": 0.40, "traslapo": 0.15, "peso_por_m2": 3.0,
}
CONCRETE_DATA = {"resistencia": 21, "densidad": 2400, "cemento_por_m3": 350}
SYSTEM_DATA = {
    "peso_por_m2": 175,
    "absorcion_agua": "Alta (se debe humedecer antes de fundir)",
    "aislamiento_termico": "U ≈ 2.10 W/m²K",
}
CARGA_VIVA = {"Residencial": 1.8, "Comercial": 2.5, "Industrial": 3.5, "Oficinas": 2.4}
NORMAS_PLACA = {
    "NSR-10 (Colombia)": {"luz_max": 4.2, "topping_min": 0.04, "concreto_min": 21, "requiere_viga_borde": True, "ref": "NSR-10 Capítulo C.21 y Título E"},
    "E.060 (Perú)":       {"luz_max": 4.2, "topping_min": 0.04, "concreto_min": 21, "requiere_viga_borde": True, "ref": "E.060 / NTE E.030"},
    "ACI 318-25 (EE.UU.)":{"luz_max": 4.2, "topping_min": 0.04, "concreto_min": 21, "requiere_viga_borde": True, "ref": "ACI 318-25 Capítulo 7"},
}

def profile_inertia():
    h=PROFILE_DATA["alto_total"]; b_web=PROFILE_DATA["ancho_alma"]; t_web=PROFILE_DATA["espesor_alma"]
    b_flange=PROFILE_DATA["ancho_ala"]; t_flange=PROFILE_DATA["espesor_ala"]
    h_cm=h*100; b_web_cm=b_web*100; t_web_cm=t_web*100; b_flange_cm=b_flange*100; t_flange_cm=t_flange*100
    A_web=b_web_cm*t_web_cm; A_flange=b_flange_cm*t_flange_cm
    y_web=h_cm/2; y_flange=t_flange_cm/2
    I_web=(b_web_cm*t_web_cm**3)/12; I_flange=(b_flange_cm*t_flange_cm**3)/12
    d_web=y_web-(h_cm/2); d_flange=y_flange-(h_cm/2)
    I_total=(I_web+A_web*d_web**2)+2*(I_flange+A_flange*d_flange**2)
    return I_total/10000

# ─── SELECCIÓN AUTOMÁTICA DE VARILLAS (NSR-10) ───
def select_rebars(As_req_cm2):
    rebar_table = [
        {"nom": '#3 (3/8")', "area": 0.71},
        {"nom": '#4 (1/2")', "area": 1.29},
        {"nom": '#5 (5/8")', "area": 1.99},
        {"nom": '#6 (3/4")', "area": 2.87},
        {"nom": '#7 (7/8")', "area": 3.87},
        {"nom": '#8 (1")',   "area": 5.10},
    ]
    best = None
    for b1 in rebar_table:
        for b2 in rebar_table:
            total = b1["area"] + b2["area"]
            if total >= As_req_cm2:
                if best is None or total < best["total"]:
                    best = {"b1": b1, "b2": b2, "total": total}
    if best is None:
        best = {"b1": rebar_table[-1], "b2": rebar_table[-1], "total": rebar_table[-1]["area"]*2}
    return best

# ─── SIDEBAR ───
norma_sel = st.sidebar.selectbox(_t("Norma de diseño","Design code"), list(NORMAS_PLACA.keys()), index=0)
norma = NORMAS_PLACA[norma_sel]
st.sidebar.header(_t("Datos del proyecto","Project data"))
proyecto_nombre   = st.sidebar.text_input(_t("Nombre del proyecto","Project name"), "Placa Fácil - Ejemplo")
proyecto_direccion= st.sidebar.text_input(_t("Dirección de obra","Site address"), "Calle 123, Bogotá")
proyecto_cliente  = st.sidebar.text_input(_t("Cliente / Propietario","Client / Owner"), "Constructora XYZ")
num_pisos  = st.sidebar.number_input(_t("Número de pisos","Number of floors"), 1, 10, 1, 1)
uso_placa  = st.sidebar.selectbox(_t("Uso de la placa","Slab use"), list(CARGA_VIVA.keys()), index=0)
carga_viva_kn = CARGA_VIVA[uso_placa]
st.sidebar.header(_t("Geometría de la placa","Slab geometry"))
Lx = st.sidebar.number_input(_t("Luz X (m)","Span X (m)"), 2.0, 12.0, 6.0, 0.1)
Ly = st.sidebar.number_input(_t("Luz Y (m)","Span Y (m)"), 2.0, 12.0, 5.0, 0.1)
orientacion = st.sidebar.selectbox(_t("Dirección de los perfiles","Profile direction"), ["Paralelo a X","Paralelo a Y"], index=0)
st.sidebar.header(_t("Parámetros de diseño","Design parameters"))
espesor_torta   = st.sidebar.number_input(_t("Espesor de la torta de concreto (cm)","Concrete topping thickness (cm)"), 4.0, 10.0, 5.0, 0.5) / 100.0
fc_concreto     = st.sidebar.number_input(_t("Resistencia f'c concreto (MPa)","Concrete strength f'c (MPa)"), 18.0, 35.0, 21.0, 0.5)
perfil_espaciado= st.sidebar.number_input(_t("Separación entre perfiles (cm)","Profile spacing (cm)"), 70.0, 100.0, 89.0, 1.0) / 100.0
incluir_vigas   = st.sidebar.checkbox(_t("Incluir vigas de borde","Include edge beams"), value=True)
viga_b = st.sidebar.number_input(_t("Ancho viga borde (cm)","Edge beam width (cm)"), 10.0, 30.0, 15.0, 1.0) / 100.0
viga_h = st.sidebar.number_input(_t("Altura viga borde (cm)","Edge beam height (cm)"), 15.0, 40.0, 20.0, 1.0) / 100.0

# ─── PARÁMETROS SÍSMICOS (NSR-10 C.21) ───
st.sidebar.header(_t("Parámetros sísmicos","Seismic parameters"))
zona_sismica = st.sidebar.selectbox(_t("Zona sísmica (Aa)","Seismic zone"),
    ["I (Aa<0.10)","II (0.10≤Aa<0.20)","III (0.20≤Aa<0.30)","IV (Aa≥0.30)"], index=1)
sistema_estructural = st.sidebar.selectbox(_t("Sistema estructural","Structural system"),
    ["DMO (Desempeño Mínimo)","DES (Desempeño Especial)","DMI (Desempeño Intermedio)"], index=0)
is_high_seismic = ("IV" in zona_sismica) or ("III" in zona_sismica)
if is_high_seismic and "DMO" in sistema_estructural:
    st.sidebar.warning(_t("⚠️ Zona alta: DMO no permitido según NSR-10 C.21. Use DES o DMI.",
                          "⚠️ High seismic zone: DMO not allowed per NSR-10 C.21. Use DES or DMI."))

st.sidebar.header(_t("Factores de desperdicio","Waste factors"))
desp_bloques  = st.sidebar.number_input(_t("Bloques (%)","Blocks (%)"), 0.0, 20.0, 5.0, 1.0) / 100.0
desp_concreto = st.sidebar.number_input(_t("Concreto (%)","Concrete (%)"), 0.0, 20.0, 10.0, 1.0) / 100.0
desp_malla    = st.sidebar.number_input(_t("Malla (%)","Mesh (%)"), 0.0, 20.0, 10.0, 1.0) / 100.0
desp_perfiles = st.sidebar.number_input(_t("Perfiles (%)","Profiles (%)"), 0.0, 20.0, 5.0, 1.0) / 100.0
st.sidebar.header(_t("APU – Precios unitarios","APU – Unit prices"))
moneda         = st.sidebar.text_input(_t("Moneda","Currency"), "COP", key="apu_moneda")
precio_bloque  = st.sidebar.number_input(_t("Precio por bloque (unidad)","Price per block (unit)"), 5000.0, 15000.0, 7200.0, 100.0)
precio_perfil  = st.sidebar.number_input(_t("Precio por metro lineal de perfil","Price per linear meter of profile"), 20000.0, 50000.0, 28000.0, 1000.0)
precio_malla   = st.sidebar.number_input(_t("Precio por m² de malla","Price per m² of mesh"), 8000.0, 20000.0, 11000.0, 500.0)
precio_concreto= st.sidebar.number_input(_t("Precio por m³ de concreto","Price per m³ of concrete"), 300000.0, 600000.0, 450000.0, 10000.0)
precio_mo      = st.sidebar.number_input(_t("Costo mano de obra (día)","Labor cost (day)"), 50000.0, 150000.0, 70000.0, 5000.0)
pct_herramienta= st.sidebar.number_input(_t("% Herramienta menor (sobre MO)","Minor tool percentage"), 0.0, 20.0, 5.0, 1.0) / 100.0
pct_aui        = st.sidebar.number_input(_t("% A.I.U. (sobre costo directo)","A.I.U. percentage"), 0.0, 50.0, 30.0, 5.0) / 100.0
pct_util       = st.sidebar.number_input(_t("% Utilidad (sobre costo directo)","Profit percentage"), 0.0, 20.0, 5.0, 1.0) / 100.0
iva            = st.sidebar.number_input(_t("% IVA (sobre utilidad)","IVA on profit"), 0.0, 30.0, 19.0, 1.0) / 100.0
st.sidebar.header(_t("Datos del plano","Drawing data"))
plano_numero = st.sidebar.text_input(_t("Número de plano","Drawing number"), "PL-01")
escala_plano = st.sidebar.text_input(_t("Escala","Scale"), "1:50")
revisado     = st.sidebar.text_input(_t("Revisado por","Reviewed by"), "J. Pérez")
aprobado     = st.sidebar.text_input(_t("Aprobado por","Approved by"), "C. Giraldo")
# ─────────────────────────────────────────────
# CÁLCULOS DE CANTIDADES
# ─────────────────────────────────────────────
area_total = Lx * Ly * num_pisos
if orientacion == "Paralelo a X":
    perfil_largo = Lx; perfil_ancho = Ly
else:
    perfil_largo = Ly; perfil_ancho = Lx

n_profiles = math.ceil(perfil_ancho / perfil_espaciado) + 1
longitud_total_perfiles = n_profiles * perfil_largo
longitud_total_perfiles_desp = longitud_total_perfiles * (1 + desp_perfiles) * num_pisos

if orientacion == "Paralelo a X":
    n_hileras = n_profiles - 1
    bloques_por_hilera = math.ceil(Lx / BLOCK_DATA["largo"])
else:
    n_hileras = n_profiles - 1
    bloques_por_hilera = math.ceil(Ly / BLOCK_DATA["largo"])
n_bloques = n_hileras * bloques_por_hilera
n_bloques_desp = math.ceil(n_bloques * (1 + desp_bloques)) * num_pisos

vol_torta = area_total * espesor_torta
vol_vigas = 0
if incluir_vigas:
    vol_vigas = (2 * (Lx + Ly)) * viga_b * viga_h * num_pisos
vol_concreto_total = vol_torta + vol_vigas
vol_concreto_total_desp = vol_concreto_total * (1 + desp_concreto)
area_malla = area_total * (1 + desp_malla + 0.15)

peso_bloques_kg  = n_bloques * BLOCK_DATA["peso_unitario"]
peso_concreto_kg = vol_concreto_total * CONCRETE_DATA["densidad"]
peso_perfiles_kg = longitud_total_perfiles * PROFILE_DATA["peso_por_m"] * num_pisos
peso_total_kg    = peso_bloques_kg + peso_concreto_kg + peso_perfiles_kg
carga_muerta_kgm2 = peso_total_kg / (Lx * Ly * num_pisos) if (Lx * Ly * num_pisos) > 0 else 0

total_cemento_kg = CONCRETE_DATA["cemento_por_m3"] * vol_concreto_total_desp
bultos_cemento   = math.ceil(total_cemento_kg / 50)

if perfil_largo > 3.5 and perfil_largo <= 4.2:
    st.warning(f"⚠️ La luz de los perfiles es {perfil_largo:.2f} m. Está cerca del máximo de {norma['luz_max']} m.")
elif perfil_largo > norma["luz_max"]:
    st.error(f"❌ La luz de los perfiles ({perfil_largo:.2f} m) excede el máximo permitido ({norma['luz_max']} m). Se requiere viga intermedia.")

# ─────────────────────────────────────────────
# CÁLCULOS ESTRUCTURALES
# ─────────────────────────────────────────────
g = 9.81
Wd_kg  = carga_muerta_kgm2
Wd_kn  = Wd_kg * g / 1000
Wu_kn  = 1.2 * Wd_kn + 1.6 * carga_viva_kn
Wu_lin = Wu_kn * perfil_espaciado
Mu     = Wu_lin * perfil_largo**2 / 8

I_real    = profile_inertia()
E_concreto = 4700 * math.sqrt(fc_concreto) * 1e6  # NSR-10: Ec=4700√f'c, fc en MPa, E en Pa
Wserv_kn  = Wd_kn + carga_viva_kn
Wserv_lin = Wserv_kn * perfil_espaciado
delta_max  = perfil_largo / 360
delta_calc = (5 * Wserv_lin * perfil_largo**4) / (384 * E_concreto * I_real) if I_real > 0 else 0
cumple_deflexion = delta_calc <= delta_max

As_min     = 0.0018 * BLOCK_DATA["ancho"] * espesor_torta
As_malla   = MESH_DATA["diametro"]**2 * math.pi / 4 * 1000 / MESH_DATA["espaciado_largo"]
cumple_malla = As_malla >= As_min

Vu = Wu_lin * perfil_largo / 2
Vc = 0.17 * math.sqrt(fc_concreto * 1000) * BLOCK_DATA["ancho"] * (BLOCK_DATA["alto"] + espesor_torta) / 1000
cumple_cortante = Vu <= Vc

# ─────────────────────────────────────────────
# DISEÑO DE LA VIGA DE BORDE
# ─────────────────────────────────────────────
W_beam_kn = Wu_kn * (perfil_espaciado/2 + 0.5)
beam_span  = max(Lx, Ly)
Mu_beam    = W_beam_kn * beam_span**2 / 8
Vu_beam    = W_beam_kn * beam_span / 2

d_beam  = viga_h - 0.05
b_beam  = viga_b
Rn      = Mu_beam / (0.9 * b_beam * d_beam**2) * 1e6 if (b_beam * d_beam**2) > 0 else 0
rho     = (0.85 * fc_concreto / 420) * (1 - math.sqrt(max(0, 1 - 2 * Rn / (0.85 * fc_concreto)))) if Rn > 0 else 0
As_beam     = rho * b_beam * d_beam * 1e4
As_min_beam = 0.0018 * b_beam * viga_h * 1e4
As_beam     = max(As_beam, As_min_beam)

# Selección automática de varillas
rebar_combo  = select_rebars(As_beam)
As_prov_beam = rebar_combo["total"]
ref_beam     = f"1{rebar_combo['b1']['nom']} + 1{rebar_combo['b2']['nom']} (As={As_prov_beam:.2f} cm²)"

# Estribos con factor sísmico NSR-10 C.21
if is_high_seismic and "DMO" not in sistema_estructural:
    stirrup_factor         = 0.25
    confinement_zone_length = 2.0 * max(viga_h, 0.50)
else:
    stirrup_factor         = 0.50
    confinement_zone_length = 0.0
s_beam = round(min(stirrup_factor * d_beam, 0.30), 2)

Vc_beam            = 0.17 * math.sqrt(fc_concreto * 1000) * b_beam * d_beam / 1000
cumple_cortante_beam = Vu_beam <= Vc_beam
cumple_flexion_beam  = As_prov_beam >= As_beam

x_vals = np.linspace(0, beam_span, 100)
M_vals = W_beam_kn * x_vals * (beam_span - x_vals) / 2
V_vals = W_beam_kn * (beam_span/2 - x_vals)

# ─────────────────────────────────────────────
# SECCIÓN COMPUESTA – NSR-10 C.17
# ─────────────────────────────────────────────
beff        = perfil_espaciado
E_acero     = 200e9
n_modular   = E_acero / E_concreto
A_profile   = (PROFILE_DATA["ancho_alma"] * PROFILE_DATA["alto_total"] +
               2 * PROFILE_DATA["ancho_ala"] * PROFILE_DATA["espesor_ala"])
y_profile   = PROFILE_DATA["alto_total"] / 2.0
A_slab      = beff * espesor_torta
y_slab      = PROFILE_DATA["alto_total"] + espesor_torta / 2.0
A_profile_trans = A_profile * n_modular
y_c_comp    = (A_profile_trans * y_profile + A_slab * y_slab) / (A_profile_trans + A_slab)
I_profile_own = profile_inertia()
I_slab_own  = (beff * espesor_torta**3) / 12.0
I_comp      = (I_profile_own * n_modular + A_profile_trans * (y_profile - y_c_comp)**2 +
               I_slab_own + A_slab * (y_slab - y_c_comp)**2)
S_comp      = I_comp / y_c_comp if y_c_comp > 0 else 0.0

# ─────────────────────────────────────────────
# LONGITUD DE DESARROLLO – NSR-10 C.12.2.2
# ─────────────────────────────────────────────
fy       = 420.0
db_ref   = 0.0127
lambda_conc = 1.0
ld = (fy * db_ref) / (1.1 * lambda_conc * math.sqrt(fc_concreto)) / 1000.0
ld = max(ld, 12 * db_ref, 0.30)

# ─────────────────────────────────────────────
# VERIFICACIONES NORMATIVAS
# ─────────────────────────────────────────────
verificaciones = []
verificaciones.append({"item":"Luz máxima de perfiles","referencia":norma['ref'],
    "requerido":f"≤ {norma['luz_max']:.2f} m","calculado":f"{perfil_largo:.2f} m",
    "cumple":perfil_largo<=norma["luz_max"],"obs":"Ok" if perfil_largo<=norma["luz_max"] else "Excede → requiere viga intermedia"})
verificaciones.append({"item":"Espesor de torta de concreto","referencia":"NSR-10 C.21.6.4.1",
    "requerido":f"≥ {norma['topping_min']*100:.0f} cm","calculado":f"{espesor_torta*100:.1f} cm",
    "cumple":espesor_torta>=norma["topping_min"],"obs":"Ok" if espesor_torta>=norma["topping_min"] else "Incrementar espesor"})
verificaciones.append({"item":"Resistencia del concreto","referencia":"NSR-10 C.21.3.1",
    "requerido":f"≥ {norma['concreto_min']} MPa","calculado":f"{fc_concreto:.1f} MPa",
    "cumple":fc_concreto>=norma["concreto_min"],"obs":"Ok" if fc_concreto>=norma["concreto_min"] else "Usar concreto de mayor resistencia"})
verificaciones.append({"item":"Deflexión (L/360)","referencia":"NSR-10 C.9.5.2",
    "requerido":f"≤ {delta_max*1000:.1f} mm","calculado":f"{delta_calc*1000:.1f} mm",
    "cumple":cumple_deflexion,"obs":"Ok" if cumple_deflexion else "Aumentar peralte o reducir luz"})
verificaciones.append({"item":"Cortante en apoyo (placa)","referencia":"NSR-10 C.11",
    "requerido":f"Vu ≤ {Vc:.2f} kN","calculado":f"Vu = {Vu:.2f} kN",
    "cumple":cumple_cortante,"obs":"Ok" if cumple_cortante else "Aumentar sección de viga"})
verificaciones.append({"item":"Acero mínimo de malla","referencia":"NSR-10 C.7.12",
    "requerido":f"As ≥ {As_min*100:.2f} cm²/m","calculado":f"As = {As_malla:.2f} cm²/m",
    "cumple":cumple_malla,"obs":"Ok" if cumple_malla else "Aumentar diámetro o reducir espaciamiento"})
verificaciones.append({"item":"Viga de borde - Momento","referencia":"NSR-10 C.21.6",
    "requerido":f"As ≥ {As_beam:.2f} cm²","calculado":f"As prov = {As_prov_beam:.2f} cm²",
    "cumple":cumple_flexion_beam,"obs":"Ok" if cumple_flexion_beam else "Aumentar acero"})
verificaciones.append({"item":"Viga de borde - Cortante","referencia":"NSR-10 C.11",
    "requerido":f"Vu ≤ {Vc_beam:.2f} kN","calculado":f"Vu = {Vu_beam:.2f} kN",
    "cumple":cumple_cortante_beam,"obs":"Ok" if cumple_cortante_beam else "Aumentar sección o estribos"})
altura_total = BLOCK_DATA["alto"] + espesor_torta
verificaciones.append({"item":"Altura total de la placa","referencia":"Práctica constructiva",
    "requerido":"≥ 13 cm","calculado":f"{altura_total*100:.1f} cm",
    "cumple":altura_total>=0.13,"obs":"Ok" if altura_total>=0.13 else "Aumentar espesor de bloque o torta"})
if norma.get("requiere_viga_borde", False) and incluir_vigas:
    verificaciones.append({"item":"Vigas de borde","referencia":"NSR-10 C.21.6.4",
        "requerido":"Incluidas","calculado":"Sí","cumple":True,"obs":"Ok"})
else:
    verificaciones.append({"item":"Vigas de borde","referencia":"NSR-10 C.21.6.4",
        "requerido":"Requerido para diafragma rígido","calculado":"No incluidas",
        "cumple":False,"obs":"Se recomienda incluir vigas de borde"})
verificaciones.append({"item":"Zona sísmica y confinamiento","referencia":"NSR-10 C.21",
    "requerido":f"Estribos {'d/4' if is_high_seismic and 'DMO' not in sistema_estructural else 'd/2'} @ {s_beam*100:.0f} cm",
    "calculado":f"Zona {zona_sismica} | Sistema {sistema_estructural}",
    "cumple": not (is_high_seismic and "DMO" in sistema_estructural),
    "obs":"Confinamiento requerido" if is_high_seismic and "DMO" not in sistema_estructural else ("⚠️ Cambiar sistema" if is_high_seismic and "DMO" in sistema_estructural else "Ok")})

# ─────────────────────────────────────────────
# PRESUPUESTO APU
# ─────────────────────────────────────────────
costo_bloques  = n_bloques_desp * precio_bloque
costo_perfiles = longitud_total_perfiles_desp * precio_perfil
costo_malla    = area_malla * precio_malla
costo_concreto = vol_concreto_total_desp * precio_concreto
dias_mo        = area_total * 0.8
costo_mo       = dias_mo * precio_mo
costo_directo  = costo_bloques + costo_perfiles + costo_malla + costo_concreto + costo_mo
herramienta    = costo_mo * pct_herramienta
aiu            = costo_directo * pct_aui
utilidad       = costo_directo * pct_util
iva_util       = utilidad * iva
total_proyecto = costo_directo + herramienta + aiu + iva_util
costo_maciza   = vol_concreto_total * 1.2 * precio_concreto
ahorro         = total_proyecto - costo_maciza if total_proyecto < costo_maciza else 0
sobrecosto     = total_proyecto - costo_maciza if total_proyecto > costo_maciza else 0

actividades    = ["Instalación de perfiles","Colocación de bloques","Colocación de malla","Fundida de concreto","Curado"]
duracion_dias  = [0.5*area_total/100, 1.0*area_total/100, 0.3*area_total/100, 1.5*area_total/100, 0.5]
cronograma     = pd.DataFrame({"Actividad": actividades, "Duración (días)": [max(1, round(d)) for d in duracion_dias]})
# ─────────────────────────────────────────────
# FUNCIÓN MODELO 3D
# ─────────────────────────────────────────────
def create_3d_model(Lx, Ly, orientacion, n_profiles, perfil_espaciado, perfil_largo,
                    espesor_torta, incluir_vigas, viga_b, viga_h):
    fig = go.Figure()
    def add_prism(x0, y0, z0, dx, dy, dz, color, opacity=0.8):
        x=[x0,x0+dx,x0+dx,x0,x0,x0+dx,x0+dx,x0]; y=[y0,y0,y0+dy,y0+dy,y0,y0,y0+dy,y0+dy]; z=[z0,z0,z0,z0,z0+dz,z0+dz,z0+dz,z0+dz]
        i=[0,0,4,4,1,5,2,6,3,7,0,4]; j=[1,2,5,6,5,6,6,7,7,4,4,7]; k=[2,3,6,7,6,7,7,4,4,5,1,3]
        fig.add_trace(go.Mesh3d(x=x,y=y,z=z,i=i,j=j,k=k,color=color,opacity=opacity,showlegend=False))
    def add_profile_c(x0, y0, z0, length, color):
        w_web=PROFILE_DATA["ancho_alma"]; w_flange=PROFILE_DATA["ancho_ala"]; h=PROFILE_DATA["alto_total"]
        add_prism(x0,y0-w_web/2,z0,length,w_web,h,color,0.9)
        add_prism(x0,y0-w_flange/2,z0+h-PROFILE_DATA["espesor_ala"],length,w_flange,PROFILE_DATA["espesor_ala"],color,0.9)
        add_prism(x0,y0-w_flange/2,z0,length,w_flange,PROFILE_DATA["espesor_ala"],color,0.9)
    if orientacion == "Paralelo a X":
        y_positions = np.linspace(0, Ly, n_profiles)
        for y in y_positions:
            add_profile_c(0, y, 0, Lx, PROFILE_DATA["color"])
        for i in range(len(y_positions)-1):
            y1=y_positions[i]; y2=y_positions[i+1]
            n_blocks_x=math.ceil(Lx/BLOCK_DATA["largo"])
            for j in range(n_blocks_x):
                x1=j*BLOCK_DATA["largo"]; x2=min(x1+BLOCK_DATA["largo"],Lx)
                add_prism(x1,y1,0,x2-x1,y2-y1,BLOCK_DATA["alto"],BLOCK_DATA["color"],0.8)
    else:
        x_positions = np.linspace(0, Lx, n_profiles)
        for x in x_positions:
            add_profile_c(x, 0, 0, Ly, PROFILE_DATA["color"])
        for i in range(len(x_positions)-1):
            x1=x_positions[i]; x2=x_positions[i+1]
            n_blocks_y=math.ceil(Ly/BLOCK_DATA["largo"])
            for j in range(n_blocks_y):
                y1=j*BLOCK_DATA["largo"]; y2=min(y1+BLOCK_DATA["largo"],Ly)
                add_prism(x1,y1,0,x2-x1,y2-y1,BLOCK_DATA["alto"],BLOCK_DATA["color"],0.8)
    add_prism(0,0,BLOCK_DATA["alto"],Lx,Ly,espesor_torta,"lightgray",0.4)
    if incluir_vigas:
        for y0 in [0, Ly]:
            add_prism(0,y0-viga_b/2,0,Lx,viga_b,viga_h,"darkgray",0.7)
        for x0 in [0, Lx]:
            add_prism(x0-viga_b/2,0,0,viga_b,Ly,viga_h,"darkgray",0.7)
    spacing=MESH_DATA["espaciado_largo"]; lx=[]; ly=[]; lz=[]
    for y in np.arange(0,Ly+spacing,spacing):
        lx.extend([0,Lx,None]); ly.extend([y,y,None]); lz.extend([BLOCK_DATA["alto"]+espesor_torta+0.01]*2+[None])
    for x in np.arange(0,Lx+spacing,spacing):
        lx.extend([x,x,None]); ly.extend([0,Ly,None]); lz.extend([BLOCK_DATA["alto"]+espesor_torta+0.01]*2+[None])
    fig.add_trace(go.Scatter3d(x=lx,y=ly,z=lz,mode='lines',line=dict(color='black',width=2),showlegend=False))
    fig.update_layout(scene=dict(xaxis_title='X (m)',yaxis_title='Y (m)',zaxis_title='Z (m)',aspectmode='data',bgcolor='#1a1a2e'),
        margin=dict(l=0,r=0,b=0,t=0),height=500,plot_bgcolor='black',paper_bgcolor='#1e1e1e')
    return fig

# ─────────────────────────────────────────────
# FUNCIÓN DXF
# ─────────────────────────────────────────────
def generate_dxf(Lx, Ly, orientacion, n_profiles, perfil_espaciado, perfil_largo,
                 incluir_vigas, viga_b, viga_h, proyecto_nombre, proyecto_direccion, proyecto_cliente,
                 plano_numero, escala_plano, revisado, aprobado, ref_beam,
                 zona_sismica_param, sistema_estructural_param, s_beam_param, confinement_param):
    doc_dxf = ezdxf.new('R2010')
    doc_dxf.units = ezdxf.units.M
    msp = doc_dxf.modelspace()
    capas=[('ROTULO',7),('ROTULO_TEXTO',5),('COTAS',3),('ACHURADO',253),
           ('REFUERZO',1),('EJES',2),('CORTE',4),('NOTAS',8),('CONCRETO',7),
           ('PERFILES',4),('BLOQUES',2),('MALLA',1),('EDGE_BEAMS',6)]
    for lay,col in capas:
        if lay not in doc_dxf.layers:
            doc_dxf.layers.add(lay, color=col)
    try:
        dimstyle=doc_dxf.dimstyles.get('COTAS-50')
    except Exception:
        dimstyle=doc_dxf.dimstyles.new('COTAS-50')
        dimstyle.dxf.dimblk='_ARCHTICK'; dimstyle.dxf.dimscale=50.0
        dimstyle.dxf.dimtxt=0.15; dimstyle.dxf.dimexe=0.1; dimstyle.dxf.dimexo=0.05
    rot_x,rot_y=0,-2.8; rot_w,rot_h=12,3.0
    msp.add_lwpolyline([(rot_x,rot_y),(rot_x+rot_w,rot_y),(rot_x+rot_w,rot_y+rot_h),(rot_x,rot_y+rot_h),(rot_x,rot_y)],dxfattribs={'layer':'ROTULO','color':7})
    col_widths=[2.5,4.5,2.5,2.5]
    x_pos=rot_x
    for w in col_widths:
        x_pos+=w
        msp.add_line((x_pos,rot_y),(x_pos,rot_y+rot_h),dxfattribs={'layer':'ROTULO','color':7})
    y_pos=rot_y
    for h in [0.6,0.6,0.6,0.6,0.6]:
        y_pos+=h
        msp.add_line((rot_x,y_pos),(rot_x+rot_w,y_pos),dxfattribs={'layer':'ROTULO','color':7})
    textos_rotulo=[
        ("PROYECTO:",proyecto_nombre,"PLANO Nº:",plano_numero,"ESCALA:",escala_plano,"FECHA:",datetime.now().strftime('%d/%m/%Y')),
        ("CLIENTE:",proyecto_cliente,"HOJA:","1/1","NORMA:",norma_sel,"REVISIÓN:","00"),
        ("DIRECCIÓN:",proyecto_direccion,"ELABORÓ:","C. Giraldo","APROBÓ:",aprobado,"REVISÓ:",revisado),
        ("","","","","","","",""),
        ("LOGO:","Placa Fácil","FIRMA:","_________________","","","","")
    ]
    for i,row in enumerate(textos_rotulo):
        y=rot_y+(i+0.4)*0.6
        for j,(label,value) in enumerate([(row[k],row[k+1]) for k in range(0,len(row),2)]):
            x=rot_x+sum(col_widths[:j])+0.1
            if label: msp.add_text(label,dxfattribs={'layer':'ROTULO_TEXTO','height':0.12,'insert':(x,y)})
            if value: msp.add_text(value,dxfattribs={'layer':'ROTULO_TEXTO','height':0.12,'insert':(x+1.2,y)})
    msp.add_lwpolyline([(0,0),(Lx,0),(Lx,Ly),(0,Ly),(0,0)],dxfattribs={'layer':'CONCRETO','color':7})
    if incluir_vigas:
        for y0 in [0,Ly]:
            msp.add_lwpolyline([(0,y0-viga_b/2),(Lx,y0-viga_b/2),(Lx,y0+viga_b/2),(0,y0+viga_b/2),(0,y0-viga_b/2)],dxfattribs={'layer':'EDGE_BEAMS'})
        for x0 in [0,Lx]:
            msp.add_lwpolyline([(x0-viga_b/2,0),(x0-viga_b/2,Ly),(x0+viga_b/2,Ly),(x0+viga_b/2,0),(x0-viga_b/2,0)],dxfattribs={'layer':'EDGE_BEAMS'})
    if orientacion=="Paralelo a X":
        y_positions=np.linspace(0,Ly,n_profiles)
        for y in y_positions:
            msp.add_line((0,y),(Lx,y),dxfattribs={'layer':'PERFILES','color':4})
    else:
        x_positions=np.linspace(0,Lx,n_profiles)
        for x in x_positions:
            msp.add_line((x,0),(x,Ly),dxfattribs={'layer':'PERFILES','color':4})
    spacing_x=MESH_DATA["espaciado_largo"]; spacing_y=MESH_DATA["espaciado_corto"]
    for x in np.arange(0,Lx,spacing_x):
        msp.add_line((x,0),(x,Ly),dxfattribs={'layer':'MALLA','color':1})
    for y in np.arange(0,Ly,spacing_y):
        msp.add_line((0,y),(Lx,y),dxfattribs={'layer':'MALLA','color':1})
    for x in np.arange(0,Lx,spacing_x):
        for y in np.arange(0,Ly,spacing_y):
            msp.add_circle((x,y),0.02,dxfattribs={'layer':'MALLA','color':1})
    if orientacion=="Paralelo a X":
        y_pos_list=np.linspace(0,Ly,n_profiles)
    else:
        y_pos_list=np.linspace(0,Lx,n_profiles)
    try:
        dim_x=msp.add_linear_dim(base=(0,-0.2),p1=(0,0),p2=(Lx,0),angle=0,dimstyle='COTAS-50',dxfattribs={'layer':'COTAS'})
        dim_x.render()
        dim_y=msp.add_linear_dim(base=(-0.2,0),p1=(0,0),p2=(0,Ly),angle=90,dimstyle='COTAS-50',dxfattribs={'layer':'COTAS'})
        dim_y.render()
        if n_profiles>=2:
            y1=y_pos_list[0]; y2=y_pos_list[1]
            if orientacion=="Paralelo a X":
                dim_sp=msp.add_linear_dim(base=(Lx+0.2,y1),p1=(Lx,y1),p2=(Lx,y2),angle=90,dimstyle='COTAS-50',dxfattribs={'layer':'COTAS'})
            else:
                dim_sp=msp.add_linear_dim(base=(y1,-0.4),p1=(y1,0),p2=(y2,0),angle=0,dimstyle='COTAS-50',dxfattribs={'layer':'COTAS'})
            dim_sp.render()
    except Exception:
        pass  # Si la versión de ezdxf no soporta la API, se omiten las cotas
    off_x=Lx+1.5
    msp.add_lwpolyline([(off_x,0),(off_x+1.5,0),(off_x+1.5,espesor_torta),(off_x,espesor_torta),(off_x,0)],dxfattribs={'layer':'CONCRETO'})
    hatch_c=msp.add_hatch(color=colors.BLUE)
    hatch_c.paths.add_polyline_path([Vec2(off_x,0),Vec2(off_x+1.5,0),Vec2(off_x+1.5,espesor_torta),Vec2(off_x,espesor_torta)],is_closed=True)
    hatch_c.set_pattern_fill(name='AR-CONC',scale=0.1,color=colors.BLUE)
    msp.add_lwpolyline([(off_x,espesor_torta),(off_x+1.5,espesor_torta),(off_x+1.5,espesor_torta+BLOCK_DATA["alto"]),(off_x,espesor_torta+BLOCK_DATA["alto"]),(off_x,espesor_torta)],dxfattribs={'layer':'BLOQUES'})
    hatch_b=msp.add_hatch(color=colors.RED)
    hatch_b.paths.add_polyline_path([Vec2(off_x,espesor_torta),Vec2(off_x+1.5,espesor_torta),Vec2(off_x+1.5,espesor_torta+BLOCK_DATA["alto"]),Vec2(off_x,espesor_torta+BLOCK_DATA["alto"])],is_closed=True)
    hatch_b.set_pattern_fill(name='INSUL',scale=0.05,color=colors.RED)
    msp.add_lwpolyline([(off_x,espesor_torta+BLOCK_DATA["alto"]),(off_x+1.5,espesor_torta+BLOCK_DATA["alto"]),(off_x+1.5,espesor_torta+BLOCK_DATA["alto"]+viga_h),(off_x,espesor_torta+BLOCK_DATA["alto"]+viga_h),(off_x,espesor_torta+BLOCK_DATA["alto"])],dxfattribs={'layer':'REFUERZO'})
    notas=[
        f"PLACA FÁCIL: Perfil Colmena {PROFILE_DATA['alto_total']*100:.0f}mm + Bloquelón {BLOCK_DATA['alto']*100:.0f}mm + Torta {espesor_torta*100:.0f}mm",
        f"Concreto: f'c={fc_concreto:.0f} MPa | Malla Q2 Ø{MESH_DATA['diametro']*1000:.0f}mm @{MESH_DATA['espaciado_largo']*100:.0f}cm",
        f"Viga borde: {ref_beam}, Estribos Ø1/4\" @ {s_beam_param*100:.0f} cm",
        f"Zona sísmica: {zona_sismica_param} | Sistema: {sistema_estructural_param}",
        f"Confinamiento: {confinement_param*100:.0f} cm desde apoyo" if confinement_param>0 else "Sin zona de confinamiento especial",
        f"Norma: {norma_sel}",
    ]
    for i,nota in enumerate(notas):
        msp.add_text(nota,dxfattribs={'layer':'NOTAS','height':0.15,'insert':(0.2,rot_y-0.3-i*0.3)})
    out=io.BytesIO()
    doc_dxf.write(out)
    return out.getvalue()

# ─────────────────────────────────────────────
# FUNCIÓN MEMORIA DOCX
# ─────────────────────────────────────────────
def generate_memory():
    doc=Document()
    doc.add_heading("Memoria de Cálculo – Placa Fácil",0)
    doc.add_paragraph(f"Proyecto: {proyecto_nombre}")
    doc.add_paragraph(f"Cliente: {proyecto_cliente}")
    doc.add_paragraph(f"Dirección: {proyecto_direccion}")
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Norma aplicada: {norma_sel} – {norma['ref']}")
    doc.add_heading("1. Datos de entrada",level=1)
    doc.add_paragraph(f"Luz X: {Lx:.2f} m, Luz Y: {Ly:.2f} m")
    doc.add_paragraph(f"Orientación de perfiles: {orientacion}")
    doc.add_paragraph(f"Espesor torta: {espesor_torta*100:.1f} cm")
    doc.add_paragraph(f"Altura bloque: {BLOCK_DATA['alto']*100:.1f} cm")
    doc.add_paragraph(f"Separación perfiles: {perfil_espaciado*100:.0f} cm")
    doc.add_paragraph(f"Concreto: f'c = {fc_concreto:.1f} MPa")
    doc.add_paragraph(f"Número de pisos: {num_pisos}")
    doc.add_paragraph(f"Uso: {uso_placa} – Carga viva: {carga_viva_kn} kN/m²")
    doc.add_paragraph(f"Zona sísmica: {zona_sismica} | Sistema estructural: {sistema_estructural}")
    doc.add_heading("2. Especificaciones técnicas de materiales",level=1)
    doc.add_paragraph(f"Bloquelón: {BLOCK_DATA['largo']*100:.0f}×{BLOCK_DATA['ancho']*100:.0f}×{BLOCK_DATA['alto']*100:.0f} cm, peso unitario {BLOCK_DATA['peso_unitario']:.0f} kg")
    doc.add_paragraph(f"Perfil Colmena: alto {PROFILE_DATA['alto_total']*1000:.0f} mm, peso {PROFILE_DATA['peso_por_m']:.1f} kg/m")
    doc.add_paragraph(f"Malla electrosoldada Q2: Ø{MESH_DATA['diametro']*1000:.0f} mm @{MESH_DATA['espaciado_largo']*100:.0f} cm")
    doc.add_heading("3. Cargas de diseño",level=1)
    doc.add_paragraph(f"Carga muerta: Wd = {Wd_kn:.3f} kN/m²")
    doc.add_paragraph(f"Carga viva: Wl = {carga_viva_kn:.2f} kN/m²")
    doc.add_paragraph(f"Carga última: Wu = {Wu_kn:.3f} kN/m²")
    doc.add_heading("4. Sección compuesta y longitud de desarrollo",level=1)
    doc.add_paragraph(f"Módulo de elasticidad del concreto: Ec = {E_concreto/1e9:.2f} GPa")
    doc.add_paragraph(f"Relación modular n = Es/Ec = {n_modular:.1f}")
    doc.add_paragraph(f"Inercia sección compuesta: I_comp = {I_comp*1e8:.4f} × 10⁻⁸ m⁴")
    doc.add_paragraph(f"Centroide compuesto desde base: y_c = {y_c_comp*100:.2f} cm")
    doc.add_paragraph(f"Módulo elástico inferior: S_comp = {S_comp*1e6:.4f} × 10⁻⁶ m³")
    doc.add_paragraph(f"Longitud de desarrollo ld (NSR-10 C.12): {ld*100:.1f} cm")
    doc.add_heading("5. Resultados estructurales",level=1)
    doc.add_paragraph(f"Momento último placa: Mu = {Mu:.2f} kN·m")
    doc.add_paragraph(f"Deflexión calculada: δ = {delta_calc*1000:.2f} mm (máx = {delta_max*1000:.2f} mm)")
    doc.add_paragraph(f"Cortante en apoyo: Vu = {Vu:.2f} kN, Vc = {Vc:.2f} kN")
    doc.add_paragraph(f"Acero de malla: As = {As_malla:.2f} cm²/m (mín = {As_min*100:.2f} cm²/m)")
    doc.add_heading("6. Diseño viga de borde",level=1)
    doc.add_paragraph(f"Momento: Mu = {Mu_beam:.2f} kN·m, As requerido = {As_beam:.2f} cm²")
    doc.add_paragraph(f"Refuerzo: {ref_beam}")
    doc.add_paragraph(f"Estribos @ {s_beam*100:.0f} cm (factor sísmico: {'d/4' if is_high_seismic and 'DMO' not in sistema_estructural else 'd/2'})")
    if confinement_zone_length > 0:
        doc.add_paragraph(f"Zona de confinamiento: {confinement_zone_length*100:.0f} cm desde el apoyo (NSR-10 C.21)")
    doc.add_heading("7. Verificaciones normativas",level=1)
    for v in verificaciones:
        estado="✅ CUMPLE" if v['cumple'] else "❌ NO CUMPLE"
        doc.add_paragraph(f"{estado} – {v['item']}: {v['calculado']} ({v['referencia']})")
    buf=io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
# ─────────────────────────────────────────────
# FUNCIÓN PDF RESUMEN EJECUTIVO
# ─────────────────────────────────────────────
def generate_pdf():
    buffer = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Resumen Ejecutivo – Placa Fácil", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Proyecto: {proyecto_nombre}", styles['Normal']))
    story.append(Paragraph(f"Cliente: {proyecto_cliente}", styles['Normal']))
    story.append(Paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", styles['Normal']))
    story.append(Paragraph(f"Norma: {norma_sel}", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("1. Resultados clave", styles['Heading1']))
    story.append(Paragraph(f"Área de placa: {area_total:.2f} m²", styles['Normal']))
    story.append(Paragraph(f"Volumen concreto: {vol_concreto_total_desp:.2f} m³", styles['Normal']))
    story.append(Paragraph(f"Cemento: {bultos_cemento} bultos", styles['Normal']))
    story.append(Paragraph(f"Costo total: {moneda} {total_proyecto:,.0f}", styles['Normal']))
    story.append(Paragraph(f"Costo por m²: {moneda} {total_proyecto/area_total:,.0f}", styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("2. Verificaciones normativas", styles['Heading1']))
    data_pdf = [["Verificación", "Estado"]]
    for v in verificaciones:
        data_pdf.append([v['item'], "✔" if v['cumple'] else "✘"])
    table_pdf = Table(data_pdf, colWidths=[10*cm, 3*cm])
    table_pdf.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), rl_colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), rl_colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('GRID', (0,0), (-1,-1), 1, rl_colors.black),
    ]))
    story.append(table_pdf)
    doc_pdf.build(story)
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data

# ─────────────────────────────────────────────
# INTERFAZ PRINCIPAL – PESTAÑAS
# ─────────────────────────────────────────────
tab_res, tab_3d, tab_dxf, tab_mem, tab_qty, tab_apu, tab_resumen = st.tabs([
    "📊 Resultados", "🧊 Modelo 3D", "📏 DXF", "📄 Memoria", "📦 Cantidades", "💰 APU", "📑 Resumen Ejecutivo"
])

with tab_res:
    st.subheader("Resultados del diseño")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Área total", f"{area_total:.2f} m²")
        st.metric("N° perfiles", f"{n_profiles}")
        st.metric("Long. perfiles", f"{longitud_total_perfiles_desp:.1f} m")
        st.metric("N° bloques", f"{n_bloques_desp} und")
    with col2:
        st.metric("Vol. concreto", f"{vol_concreto_total_desp:.2f} m³")
        st.metric("Área malla", f"{area_malla:.2f} m²")
        st.metric("Carga muerta", f"{carga_muerta_kgm2:.0f} kg/m²")
        st.metric("Cemento", f"{bultos_cemento} bultos")
    with col3:
        st.metric("Mu por perfil", f"{Mu:.2f} kN·m")
        st.metric("Deflexión", f"{delta_calc*1000:.1f} mm")
        st.metric("Límite deflexión", f"{delta_max*1000:.1f} mm")
        st.metric("Vu placa", f"{Vu:.2f} kN")

    st.markdown("### Viga de borde")
    st.write(f"**Momento:** Mu = {Mu_beam:.2f} kN·m | As req = {As_beam:.2f} cm² | As prov = {As_prov_beam:.2f} cm²")
    st.write(f"**Cortante:** Vu = {Vu_beam:.2f} kN | Vc = {Vc_beam:.2f} kN")
    st.write(f"**Refuerzo:** {ref_beam}, estribos @ {s_beam*100:.0f} cm")
    st.write(f"**Zona sísmica:** {zona_sismica} | **Sistema:** {sistema_estructural}")
    if confinement_zone_length > 0:
        st.info(f"🔒 Zona de confinamiento: {confinement_zone_length*100:.0f} cm desde apoyo | Estribos @ {s_beam*100:.0f} cm (NSR-10 C.21)")

    st.markdown("### Sección compuesta y longitud de desarrollo")
    col_sc1, col_sc2, col_sc3 = st.columns(3)
    with col_sc1:
        st.metric("I_comp", f"{I_comp*1e8:.4f} ×10⁻⁸ m⁴")
        st.metric("y_c centroide", f"{y_c_comp*100:.2f} cm")
    with col_sc2:
        st.metric("S_comp (módulo elástico)", f"{S_comp*1e6:.4f} ×10⁻⁶ m³")
        st.metric("n modular (Es/Ec)", f"{n_modular:.1f}")
    with col_sc3:
        st.metric("Ec concreto", f"{E_concreto/1e9:.2f} GPa")
        st.metric("ld desarrollo (NSR-10 C.12)", f"{ld*100:.1f} cm")

    st.markdown("### Diagramas viga de borde")
    fig_diag, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
    ax1.plot(x_vals, M_vals, 'b-', linewidth=2)
    ax1.fill_between(x_vals, 0, M_vals, alpha=0.3, color='blue')
    ax1.set_xlabel("Distancia (m)"); ax1.set_ylabel("Momento (kN·m)")
    ax1.set_title("Diagrama de Momento - Viga Borde"); ax1.grid(True, alpha=0.3)
    ax2.plot(x_vals, V_vals, 'r-', linewidth=2)
    ax2.fill_between(x_vals, 0, V_vals, alpha=0.3, color='red')
    ax2.set_xlabel("Distancia (m)"); ax2.set_ylabel("Cortante (kN)")
    ax2.set_title("Diagrama de Cortante - Viga Borde"); ax2.grid(True, alpha=0.3)
    plt.tight_layout()
    st.pyplot(fig_diag)
    plt.close()

    st.markdown("### Verificaciones normativas")
    for v in verificaciones:
        if v['cumple']:
            st.success(f"✅ {v['item']}: {v['calculado']} – {v['obs']}")
        else:
            st.error(f"❌ {v['item']}: {v['calculado']} – {v['obs']}")
        st.caption(f"Referencia: {v['referencia']}")

with tab_3d:
    st.subheader("Modelo 3D de la placa")
    fig_3d = create_3d_model(Lx, Ly, orientacion, n_profiles, perfil_espaciado, perfil_largo,
                              espesor_torta, incluir_vigas, viga_b, viga_h)
    st.plotly_chart(fig_3d, use_container_width=True)

with tab_dxf:
    st.subheader("Exportar plano DXF profesional")
    st.info("El DXF incluye rótulo ICONTEC, cotas, achurados, perfil Colmena, malla y vigas de borde con zona de confinamiento.")
    if st.button("Generar archivo DXF"):
        dxf_data = generate_dxf(Lx, Ly, orientacion, n_profiles, perfil_espaciado, perfil_largo,
                                incluir_vigas, viga_b, viga_h, proyecto_nombre, proyecto_direccion, proyecto_cliente,
                                plano_numero, escala_plano, revisado, aprobado, ref_beam,
                                zona_sismica, sistema_estructural, s_beam, confinement_zone_length)
        st.download_button("📥 Descargar DXF", data=dxf_data,
                          file_name=f"PlacaFacil_{proyecto_nombre}.dxf", mime="application/dxf")

with tab_mem:
    st.subheader("Memoria de cálculo")
    if st.button("Generar memoria DOCX"):
        buf = generate_memory()
        st.download_button("📥 Descargar Memoria", data=buf,
                          file_name=f"Memoria_PlacaFacil_{proyecto_nombre}.docx",
                          mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_qty:
    st.subheader("Cantidades de materiales")
    df_qty = pd.DataFrame({
        "Material": ["Bloques", "Perfiles Colmena", "Malla electrosoldada", "Concreto", "Cemento"],
        "Cantidad": [f"{n_bloques_desp} unidades", f"{longitud_total_perfiles_desp:.1f} m",
                     f"{area_malla:.2f} m²", f"{vol_concreto_total_desp:.2f} m³", f"{bultos_cemento} bultos"],
        "Desperdicio incluido": ["Sí","Sí","Sí (traslapo)","Sí","Sí"],
    })
    st.dataframe(df_qty, use_container_width=True)
    st.markdown("#### Desglose de volumen de concreto")
    st.write(f"- Torta de concreto: {vol_torta:.3f} m³")
    st.write(f"- Vigas de borde: {vol_vigas:.3f} m³")
    st.write(f"- **Total sin desperdicio:** {vol_concreto_total:.3f} m³")
    st.write(f"- **Total con desperdicio:** {vol_concreto_total_desp:.3f} m³")
    fig_qty = go.Figure(go.Bar(
        x=["Bloques", "Perfiles", "Malla", "Concreto", "Mano de obra"],
        y=[costo_bloques, costo_perfiles, costo_malla, costo_concreto, costo_mo],
        marker_color=["#C19A6B","#C0C0C0","#4CAF50","#607D8B","#FF9800"]
    ))
    fig_qty.update_layout(title="Costo por componente", yaxis_title=f"Costo ({moneda})", height=350)
    st.plotly_chart(fig_qty, use_container_width=True)

with tab_apu:
    st.subheader("Análisis de Precios Unitarios (APU)")
    df_apu = pd.DataFrame({
        "Ítem": ["Bloques (c/desperdicio)","Perfiles (c/desperdicio)","Malla (c/traslapo)","Concreto (c/desperdicio)","Mano de obra","Herramienta menor","A.I.U.","Utilidad","IVA s/Utilidad","TOTAL"],
        "Costo": [costo_bloques, costo_perfiles, costo_malla, costo_concreto, costo_mo,
                  herramienta, aiu, utilidad, iva_util, total_proyecto],
    })
    df_apu[f"Costo ({moneda})"] = df_apu["Costo"].apply(lambda x: f"{x:,.0f}")
    st.dataframe(df_apu[["Ítem", f"Costo ({moneda})"]], use_container_width=True)
    st.metric("Costo directo", f"{moneda} {costo_directo:,.0f}")
    st.metric("TOTAL PROYECTO", f"{moneda} {total_proyecto:,.0f}")
    st.metric("Costo por m²", f"{moneda} {total_proyecto/area_total:,.0f}")
    st.markdown("#### Cronograma estimado de obra")
    st.dataframe(cronograma, use_container_width=True)
    fig_gantt = go.Figure(go.Bar(
        x=cronograma["Duración (días)"], y=cronograma["Actividad"],
        orientation='h', marker_color="#01696f"
    ))
    fig_gantt.update_layout(title="Cronograma de actividades", xaxis_title="Días", height=300)
    st.plotly_chart(fig_gantt, use_container_width=True)

with tab_resumen:
    st.subheader("Resumen Ejecutivo PDF")
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        st.metric("Costo total", f"{moneda} {total_proyecto:,.0f}")
        st.metric("Costo por m²", f"{moneda} {total_proyecto/area_total:,.0f}")
        cumplidas = sum(1 for v in verificaciones if v['cumple'])
        st.metric("Verificaciones cumplidas", f"{cumplidas}/{len(verificaciones)}")
    with col_r2:
        if sobrecosto > 0:
            st.warning(f"⚠️ Sobrecosto vs placa maciza: {moneda} {sobrecosto:,.0f}")
        else:
            st.success(f"✅ Ahorro vs placa maciza: {moneda} {ahorro:,.0f}")
        st.metric("Deflexión", f"{'✅ OK' if cumple_deflexion else '❌ EXCEDE'} — {delta_calc*1000:.1f} mm")
        st.metric("Zona sísmica", f"{zona_sismica}")
    if st.button("Generar PDF Resumen"):
        pdf_data = generate_pdf()
        st.download_button("📥 Descargar PDF Resumen", data=pdf_data,
                          file_name=f"Resumen_PlacaFacil_{proyecto_nombre}.pdf",
                          mime="application/pdf")