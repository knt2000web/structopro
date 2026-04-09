import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Albañilería / Mamposteria Confinada", "Confined Masonry"), layout="wide")

# Terminología dinámica por norma (antes de sidebar, se actualiza después)
_norma_init = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
if "NSR-10" in _norma_init or "Colombia" in _norma_init:
    _titulo_mod = " Diseño de Mampostera Confinada — NSR-10 Título D (Colombia)"
    _desc_mod = "Diseño estructural completo de muros de **mampostera confinada** con columnas de confinamiento y **vigas de amarre**. Multi-norma: **E.070 (Perú)**, **NSR-10 D (Colombia)**, **NEC-SE-MP (Ecuador)**, **COVENIN 1753 (Venezuela)**."
elif "E.070" in _norma_init or "Perú" in _norma_init:
    _titulo_mod = " Diseño de Albañilería Confinada — E.070 (Perú)"
    _desc_mod = "Diseño estructural completo de muros de **albañilería confinada** con columnas de confinamiento y vigas soleras. Multi-norma: **E.070 (Perú)**, **NSR-10 D (Colombia)**, **NEC-SE-MP (Ecuador)**, **COVENIN 1753 (Venezuela)**."
else:
    _titulo_mod = " Diseño de Albañilería / Mampostera Confinada"
    _desc_mod = "Diseño estructural completo de muros de albañilería/mampostera confinada. Multi-norma: **E.070 (Perú)**, **NSR-10 D (Colombia)**, **NEC-SE-MP (Ecuador)**, **COVENIN 1753 (Venezuela)**."
st.title(_titulo_mod)
st.markdown(_desc_mod)

# ─────────────────────────────────────────────
# BASES DE DATOS POR NORMA
# ─────────────────────────────────────────────
NORMAS_AC = {
    "E.070 (Perú)":        {"pais":"pe","phi_f":0.90,"phi_v":0.85,"phi_c":0.70,"bag_kg":42.5,
                            "fm_def":65.0,"vm_def":8.05,"E_def":32500.0,
                            "fy_def":4200,"fc_def":210,"unidad":"kg/cm²","fuerza":"ton"},
    "NSR-10 (Colombia)":   {"pais":"co","phi_f":0.90,"phi_v":0.85,"phi_c":0.65,"bag_kg":50.0,
                            "fm_def":5.5,"vm_def":0.78,"E_def":6500.0,
                            "fy_def":420,"fc_def":21,"unidad":"MPa","fuerza":"kN"},
    "NEC-SE-MP (Ecuador)": {"pais":"ec","phi_f":0.90,"phi_v":0.85,"phi_c":0.65,"bag_kg":50.0,
                            "fm_def":4.0,"vm_def":0.55,"E_def":4000.0,
                            "fy_def":420,"fc_def":21,"unidad":"MPa","fuerza":"kN"},
    "COVENIN 1753 (Venezuela)":{"pais":"ve","phi_f":0.90,"phi_v":0.80,"phi_c":0.65,"bag_kg":42.5,
                            "fm_def":3.5,"vm_def":0.50,"E_def":3500.0,
                            "fy_def":420,"fc_def":18,"unidad":"MPa","fuerza":"kN"},
}

LADRILLOS_AC = {
    "E.070 (Perú)": {
        "King Kong 18H – Tipo IV (24×13×9)":  {"L":24,"W":13,"H":9,"clase":"IV"},
        "King Kong Artesanal – Tipo I (24×13×9)":{"L":24,"W":13,"H":9,"clase":"I"},
        "Pandereta (23×11×9)":                 {"L":23,"W":11,"H":9,"clase":"II"},
        "Sílico Calcáreo (24×13×9)":           {"L":24,"W":13,"H":9,"clase":"IV"},
    },
    "NSR-10 (Colombia)": {
        "Ladrillo Portante A (29×12×9)":       {"L":29,"W":12,"H":9,"clase":"A"},
        "Ladrillo Portante B (29×12×9)":       {"L":29,"W":12,"H":9,"clase":"B"},
        "Bloque de Arcilla N5 (30×12×20)":     {"L":30,"W":12,"H":20,"clase":"N5"},
        "Bloque de Concreto (40×20×20)":       {"L":40,"W":20,"H":20,"clase":"P"},
    },
    "NEC-SE-MP (Ecuador)": {
        "Ladrillo Mambrón (25×12×7)":          {"L":25,"W":12,"H":7,"clase":"A"},
        "Bloque Hormigón (40×20×20)":          {"L":40,"W":20,"H":20,"clase":"B"},
        "Ladrillo Panelón (29×14×9)":          {"L":29,"W":14,"H":9,"clase":"A"},
    },
    "COVENIN 1753 (Venezuela)": {
        "Bloque de Arcilla (25×15×10)":        {"L":25,"W":15,"H":10,"clase":"A"},
        "Bloque de Concreto (40×20×15)":       {"L":40,"W":20,"H":15,"clase":"B"},
        "Ladrillo Artesanal (23×11×6)":        {"L":23,"W":11,"H":6,"clase":"C"},
    },
}

REBAR_US = {
    "#3 (3/8\")":  {"area":0.71,"diam_mm":9.5},
    "#4 (1/2\")":  {"area":1.29,"diam_mm":12.7},
    "#5 (5/8\")":  {"area":1.99,"diam_mm":15.9},
    "#6 (3/4\")":  {"area":2.84,"diam_mm":19.1},
    "#8 (1\")":    {"area":5.07,"diam_mm":25.4},
}
REBAR_MM = {
    "8 mm":   {"area":0.503,"diam_mm":8},
    "10 mm":  {"area":0.786,"diam_mm":10},
    "12 mm":  {"area":1.131,"diam_mm":12},
    "16 mm":  {"area":2.011,"diam_mm":16},
    "19 mm":  {"area":2.835,"diam_mm":19},
}
ESTRIBO_US = {
    "#2 (1/4\")":  {"area":0.32,"diam_mm":6.35},
    "#3 (3/8\")":  {"area":0.71,"diam_mm":9.5},
}
ESTRIBO_MM = {
    "6 mm":  {"area":0.283,"diam_mm":6},
    "8 mm":  {"area":0.503,"diam_mm":8},
    "10 mm": {"area":0.786,"diam_mm":10},
}

# ─────────────────────────────────────────────
# FUNCIONES AUXILIARES
# ─────────────────────────────────────────────
def mix_for_fc(fc_mpa):
    fc = fc_mpa * 10.1972
    if fc <= 175:
        return (280,195,820,1080)
    elif fc <= 210:
        return (350,195,780,1020)
    elif fc <= 280:
        return (420,190,720,980)
    else:
        return (490,185,680,940)

def qty_table(rows):
    df = pd.DataFrame(rows, columns=["Parámetro","Valor"])
    st.dataframe(df, use_container_width=True, hide_index=True)

def to_kgcm2(v, u):
    return float(v) if u == "kg/cm²" else float(v) * 10.1972

def to_kn(v, u):
    if u == "ton":
        return float(v) * 9.80665
    return float(v)

def to_ton(v, u):
    if u == "kN":
        return float(v) / 9.80665
    return float(v)

# ─────────────────────────────────────────────
# SIDEBAR — CONFIGURACIÓN GLOBAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙ Norma y País","⚙ Code & Country"))

global_norma = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
if global_norma not in NORMAS_AC:
    global_norma = "NSR-10 (Colombia)"

norma_ac = st.sidebar.selectbox(
    _t("Norma de Diseño","Design Code"),
    list(NORMAS_AC.keys()),
    index=list(NORMAS_AC.keys()).index(global_norma),
    key="ac_norma_selector"
)
st.session_state.ac_norma = norma_ac
nc = NORMAS_AC[norma_ac]
_iso = nc["pais"]
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagpedia.net/data/flags/mini/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{norma_ac}</span>'
    f'</div>', unsafe_allow_html=True
)

# Unidades de salida (kN / tonf)
st.sidebar.header(_t(" Unidades de salida"," Output units"))
unidades_salida = st.sidebar.radio("Unidades de fuerza/momento:", ["kiloNewtons (kN, kN·m)", "Toneladas fuerza (tonf, tonf·m)"], key="ac_output_units")
if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom    = "tonf·m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom    = "kN·m"

st.sidebar.header(_t(" Materiales"," Materials"))

# Definir rangos según unidad
if nc["unidad"] == "kg/cm²":
    fm_min, fm_max, fm_step = 10.0, 300.0, 5.0
    vm_min, vm_max, vm_step = 1.0, 50.0, 0.5
    E_min, E_max, E_step = 1000.0, 100000.0, 500.0
    fc_min, fc_max, fc_step = 100.0, 500.0, 5.0
    fy_min, fy_max, fy_step = 2000.0, 6000.0, 50.0
else:  # MPa
    fm_min, fm_max, fm_step = 1.0, 30.0, 0.5
    vm_min, vm_max, vm_step = 0.1, 5.0, 0.1
    E_min, E_max, E_step = 500.0, 50000.0, 100.0
    fc_min, fc_max, fc_step = 10.0, 100.0, 1.0
    fy_min, fy_max, fy_step = 200.0, 600.0, 10.0

fm_in  = st.sidebar.number_input(f"f'm [{nc['unidad']}]", fm_min, fm_max, float(nc["fm_def"]), fm_step, key="ac_fm")
vm_in  = st.sidebar.number_input(f"v'm [{nc['unidad']}]", vm_min, vm_max,  float(nc["vm_def"]), vm_step, key="ac_vm")
E_in   = st.sidebar.number_input(f"E (mampuesto) [{nc['unidad']}]", E_min, E_max, float(nc["E_def"]), E_step, key="ac_E")
fc_in  = st.sidebar.number_input(f"f'c concreto [{nc['unidad']}]", fc_min, fc_max, float(nc["fc_def"]), fc_step, key="ac_fc")
fy_in  = st.sidebar.number_input(f"fy acero [{nc['unidad']}]", fy_min, fy_max, float(nc["fy_def"]), fy_step, key="ac_fy")
bag_kg = nc["bag_kg"]

# Convertir a kg/cm² para cálculos internos
fm   = to_kgcm2(fm_in, nc["unidad"])
vm   = to_kgcm2(vm_in, nc["unidad"])
fc   = to_kgcm2(fc_in, nc["unidad"])
fy   = to_kgcm2(fy_in, nc["unidad"])
phi_f= nc["phi_f"]
phi_v= nc["phi_v"]
phi_c= nc["phi_c"]

st.sidebar.header(_t(" Geometría del Muro"," Wall Geometry"))
ladrillo_db = LADRILLOS_AC.get(norma_ac, LADRILLOS_AC["E.070 (Perú)"])
lad_sel = st.sidebar.selectbox(_t("Tipo de Ladrillo","Brick Type"), list(ladrillo_db.keys()), key="ac_lad")
lad = ladrillo_db[lad_sel]

L_m_in = st.sidebar.number_input(_t("Longitud del muro L [m]","Wall length L [m]"),0.5,20.0,st.session_state.get("ac_L",2.6),0.1,key="ac_L")
t_m_in = st.sidebar.number_input(_t("Espesor del muro t [m]","Wall thickness t [m]"),0.1,0.5,st.session_state.get("ac_t",0.23),0.01,key="ac_t")
h_m_in = st.sidebar.number_input(_t("Altura libre del muro h [m]","Clear height h [m]"),1.0,6.0,st.session_state.get("ac_h",2.4),0.1,key="ac_h")
L_m = L_m_in * 100
t_m = t_m_in * 100
h_m = h_m_in * 100

st.sidebar.header(_t("⚡ Cargas","⚡ Loads"))
Pm_in = st.sidebar.number_input(f"Pm (D+L) [{nc['fuerza']}]",0.0,2000.0,st.session_state.get("ac_Pm",16.0),0.5,key="ac_Pm")
Pg_in = st.sidebar.number_input(f"Pg (D+0.25L) [{nc['fuerza']}]",0.0,2000.0,st.session_state.get("ac_Pg",14.0),0.5,key="ac_Pg")
Ve_in = st.sidebar.number_input(f"Ve (sismo mod.) [{nc['fuerza']}]",0.0,2000.0,st.session_state.get("ac_Ve",8.72),0.1,key="ac_Ve")
Me_in = st.sidebar.number_input(f"Mo (momento mod.) [{nc['fuerza']}·m]",0.0,5000.0,st.session_state.get("ac_Me",42.98),0.5,key="ac_Me")
Vu_in = st.sidebar.number_input(f"Vu (sismo sev.) [{nc['fuerza']}]",0.0,2000.0,st.session_state.get("ac_Vu",26.16),0.1,key="ac_Vu")
Mu_in = st.sidebar.number_input(f"Mu (momento sev.) [{nc['fuerza']}·m]",0.0,5000.0,st.session_state.get("ac_Mu",128.93),0.5,key="ac_Mu")

if nc["fuerza"] == "kN":
    Pm = Pm_in / 9.80665
    Pg = Pg_in / 9.80665
    Ve = Ve_in / 9.80665
    Me = Me_in / 9.80665
    Vu = Vu_in / 9.80665
    Mu = Mu_in / 9.80665
else:
    Pm = Pm_in; Pg = Pg_in; Ve = Ve_in; Me = Me_in; Vu = Vu_in; Mu = Mu_in

st.sidebar.header(_t(" Varillas"," Rebar"))
default_sys = 0 if "Colombia" in norma_ac else 1
bar_sys = st.sidebar.radio(_t("Sistema:","System:"),["Pulgadas (# US)","Milímetros (mm)"], index=default_sys, horizontal=True,key="ac_bar_sys")
rebar_db  = REBAR_US if "Pulgadas" in bar_sys else REBAR_MM
estrib_db = ESTRIBO_US if "Pulgadas" in bar_sys else ESTRIBO_MM
def_rb  = list(rebar_db.keys())[1]
def_est = list(estrib_db.keys())[0]
bar_col_sel  = st.sidebar.selectbox(_t("Varilla columna confinamiento","Confinement column bar"),list(rebar_db.keys()),key="ac_bar_col",
                                     index=list(rebar_db.keys()).index(st.session_state.get("ac_bar_col",def_rb)) if st.session_state.get("ac_bar_col",def_rb) in rebar_db else 1)
bar_hor_sel  = st.sidebar.selectbox(_t("Refuerzo horizontal (muro)","Horizontal rebar (wall)"),list(rebar_db.keys()),key="ac_bar_hor",
                                     index=0)
est_sel      = st.sidebar.selectbox(_t("Estribos columna","Column stirrups"),list(estrib_db.keys()),key="ac_est",index=0)

Ab_col = rebar_db[bar_col_sel]["area"]
db_col = rebar_db[bar_col_sel]["diam_mm"]
Ab_hor = rebar_db[bar_hor_sel]["area"]
db_hor = rebar_db[bar_hor_sel]["diam_mm"]
Av_est = estrib_db[est_sel]["area"]
db_est = estrib_db[est_sel]["diam_mm"]

st.sidebar.header(_t(" Columnas Confinamiento"," Confinement Columns"))
col_w_in = st.sidebar.number_input(_t("Ancho columna b [cm]","Column width b [cm]"),10.0,60.0,25.0,1.0,key="ac_colw")
col_d_in = st.sidebar.number_input(_t("Peralte columna d [cm]","Column depth d [cm]"),10.0,60.0,23.0,1.0,key="ac_cold")
recub_col = st.sidebar.number_input(_t("Recubrimiento col. [cm]","Cover col. [cm]"),2.0,5.0,3.0,0.5,key="ac_recub")
Nc_in    = st.sidebar.number_input(_t("N° columnas confinamiento Nc","Number of columns Nc"),2,6,2,1,key="ac_Nc")

st.sidebar.header(_t(" Viga Solera"," Solera Beam"))
vs_b_in  = st.sidebar.number_input(_t("b solera [cm]","Solera width [cm]"),15.0,50.0,25.0,1.0,key="ac_vsb")
vs_h_in  = st.sidebar.number_input(_t("h solera [cm]","Solera height [cm]"),15.0,50.0,20.0,1.0,key="ac_vsh")

st.sidebar.header(_t(" Datos de Obra"," Project Data"))
L_total_obra = st.sidebar.number_input(_t("Longitud total muros [m]","Total wall length [m]"),1.0,500.0,10.0,1.0,key="ac_Ltotal")
nombre_muro  = st.sidebar.text_input(_t("Nombre del muro","Wall name"),"MX-1",key="ac_nombre")
piso_muro    = st.sidebar.text_input(_t("Piso/Nivel","Floor/Level"),"Piso 1",key="ac_piso")

# ─────────────────────────────────────────────
# CÁLCULOS PRINCIPALES
# ─────────────────────────────────────────────
sigma_m_kgcm2 = Pm * 1000 / (L_m * t_m)
Fa_inner = 0.2 * fm * (1 - (h_m / (35 * t_m))**2)
Fa_lim   = 0.15 * fm
Fa       = min(Fa_inner, Fa_lim)
ok_sigma = sigma_m_kgcm2 <= Fa

alpha_raw = (Ve * L_m_in) / Me if Me > 0 else 1.0
alpha = max(1/3, min(alpha_raw, 1.0))

Vm_calc = 0.5 * vm * alpha * t_m_in * L_m_in + 0.23 * Pg
Vm1 = Vm_calc
Ve1 = Ve

ratio_VV = Vm1 / Ve1 if Ve1 > 0 else 2.0
ratio_VV = max(2.0, min(ratio_VV, 3.0))
Vui_calc = Ve * ratio_VV
Mui_calc = Me * ratio_VV

ok_fisura = Ve <= 0.55 * Vm1
cond_mur  = "Agrietado" if not ok_fisura else "No Agrietado"

cond_Hz_Vu  = Vu > Vm1
cond_Hz_sig = sigma_m_kgcm2 >= 0.05 * fm
necesita_Hz = cond_Hz_Vu or cond_Hz_sig
rho_h_min = 0.001
s_hor = 25 if necesita_Hz else 999
rho_h_prov = Ab_hor / (s_hor * t_m) if s_hor < 999 else 0
ok_rho_h = rho_h_prov >= rho_h_min

Nc       = int(Nc_in)
Lp_mayor = L_m_in
Lm       = Lp_mayor
M_col    = Mu / Nc
F_force  = Mu / L_m_in
Pc_col   = Pg / Nc
Pt_col   = 0.25 * Pg
Vc_ext = 1.5 * Vm1 * Lm / (L_m_in * (Nc + 1))
T1_ext = F_force - Pc_col
T2_ext = F_force - Pc_col - Pt_col
C1_ext = Pc_col + F_force

d_col    = col_d_in - recub_col
Ac_col   = col_w_in * col_d_in
As_min_four_d8 = 4 * 0.503
As_min_pct     = 0.1 * fc * Ac_col / fy

if not ok_fisura:
    mu_fric  = 1.0
    Asf_col  = Vc_ext / (fy * mu_fric * phi_v / 1000)
    T_use    = max(T1_ext, T2_ext, 0.0)
    Ast_col  = T_use * 1000 / (fy * phi_f)
    As_req   = Asf_col + Ast_col
    As_req   = max(As_req, As_min_four_d8)
    design_mode = "Agrietado"
else:
    T_use    = max(T1_ext, 0.0)
    Ast_col  = T_use * 1000 / (fy * phi_f)
    Asf_col  = 0.0
    As_req   = max(Ast_col, As_min_four_d8)
    design_mode = "No Agrietado"

n_bars   = math.ceil(As_req / Ab_col)
n_bars   = max(4, n_bars)
As_prov  = n_bars * Ab_col

delta_col = 1.0
phi_comp  = 0.70
C_use     = C1_ext * 1000
if phi_comp > 0 and fc > 0:
    An_comp   = As_prov + (C_use / phi_comp - As_prov * fy) / (0.85 * delta_col * fc)
else:
    An_comp   = 1000
An_comp   = max(An_comp, 15 * t_m)

phi_cf   = 0.85
Acf_col  = Vc_ext * 1000 / (0.2 * fc * phi_cf) if fc > 0 else 0
Acf_col  = max(Acf_col, 15 * t_m)

An_final = max(An_comp, Acf_col, Ac_col)
col_b_min = col_w_in
col_d_min = math.ceil(An_final / col_b_min)
col_d_use = max(col_d_in, col_d_min)
Ac_final  = col_b_min * col_d_use

As_min_final = 0.1 * fc * Ac_final / fy
As_use       = max(As_prov, As_min_final)
n_bars_f     = math.ceil(As_use / Ab_col)
n_bars_f     = max(4, n_bars_f)
As_prov_f    = n_bars_f * Ab_col

tn_col  = col_d_use - 2 * recub_col
d_ef    = col_d_use - recub_col
ratio_Ac_An = Ac_final / An_final if An_final > 0 else 1.0
if ratio_Ac_An > 1.0:
    s1_col  = Av_est * fy / (0.3 * tn_col * fc * (ratio_Ac_An - 1))
else:
    s1_col  = 100
s2_col  = Av_est * fy / (0.12 * tn_col * fc) if tn_col > 0 else 100
s3_col  = d_ef / 4
s3_col  = max(s3_col, 5.0)
s4_col  = 10.0
s_col   = min(s1_col, s2_col, s3_col, s4_col)
s_col   = max(5.0, round(s_col, 1))

Z_col   = max(1.5 * d_ef, 45.0)

Ts_sol  = Vm1 * Lm / (2 * L_m_in)
As_sol_req = Ts_sol * 1000 / (phi_f * fy)
Acs_sol    = vs_b_in * vs_h_in
As_sol_min = max(0.1 * fc * Acs_sol / fy, 4 * 0.503)
As_sol     = max(As_sol_req, As_sol_min)
n_sol      = math.ceil(As_sol / Ab_col)
n_sol      = max(4, n_sol)
As_sol_prov= n_sol * Ab_col

# ─────────────────────────────────────────────
# CANTIDADES Y COSTOS
# ─────────────────────────────────────────────
vol_col_ml   = Ac_final / 1e4 * (h_m / 100)
n_col_total  = math.ceil(L_total_obra / L_m_in) * Nc
vol_col_total= n_col_total * vol_col_ml

vol_vs_ml   = vs_b_in * vs_h_in / 1e4
vol_vs_total= vol_vs_ml * L_total_obra
vol_conc_total = vol_col_total + vol_vs_total

long_bar_col = (h_m / 100) + 0.3
peso_bar_col = n_bars_f * long_bar_col * (Ab_col / 10000) * 7850
n_est_col    = math.ceil((h_m / 100) / (s_col / 100))
perim_est    = 2 * (col_b_min + col_d_use) / 100
peso_est_col = n_est_col * perim_est * (Av_est / 10000) * 7850
peso_col_ml  = peso_bar_col + peso_est_col
peso_col_total = peso_col_ml * n_col_total

long_bar_sol = L_total_obra + 0.5
peso_sol     = n_sol * long_bar_sol * (Ab_col / 10000) * 7850
s_est_sol    = 10.0
n_est_sol    = math.ceil(vs_b_in / s_est_sol) * math.ceil(L_total_obra / L_m_in)
peso_est_sol = n_est_sol * (2*(vs_b_in+vs_h_in)/100) * (Av_est / 10000) * 7850
peso_sol_total = peso_sol + peso_est_sol

n_hiladas_hor = math.ceil(h_m / s_hor) if necesita_Hz else 0
peso_hor_total= n_hiladas_hor * L_total_obra * (Ab_hor / 10000) * 7850

total_acero_kg = peso_col_total + peso_sol_total + peso_hor_total

junta = 1.5
lad_m2 = 1 / ((lad["L"]/100 + junta/100) * (lad["H"]/100 + junta/100))
area_muro_total = L_total_obra * h_m_in
n_ladrillos_total = math.ceil(lad_m2 * area_muro_total * 1.05)
vol_mortero_total = area_muro_total * t_m_in * 0.15

# ─────────────────────────────────────────────
# FUNCIONES DE DIBUJO PARA FIGURADO
# ─────────────────────────────────────────────
def draw_longitudinal_bar(total_len_cm, straight_len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    fig, ax = plt.subplots(figsize=(max(6, total_len_cm/20), 2))
    ax.set_aspect('equal')
    ax.plot([0, straight_len_cm], [0, 0], 'k-', linewidth=2)
    ax.plot([0, 0], [0, hook_len_cm], 'k-', linewidth=2)
    ax.plot([straight_len_cm, straight_len_cm], [0, -hook_len_cm], 'k-', linewidth=2)
    ax.annotate(f"{straight_len_cm:.0f} cm", xy=(straight_len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 12db = {hook_len_cm:.0f} cm", xy=(0, hook_len_cm/2), ha='right', fontsize=8)
    ax.annotate(f"Gancho 12db", xy=(straight_len_cm, -hook_len_cm/2), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*0.2, straight_len_cm + hook_len_cm*0.2)
    ax.set_ylim(-hook_len_cm*1.2, hook_len_cm*1.2)
    ax.axis('off')
    title = f"Varilla {bar_name or 'L1'} - Ø{bar_diam_mm:.0f} mm - Longitud total {total_len_cm:.0f} cm"
    ax.set_title(title, fontsize=9, fontweight='bold')
    return fig

def draw_stirrup(b_cm, h_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    fig, ax = plt.subplots(figsize=(max(5, b_cm/12), max(4, h_cm/12)))
    ax.set_aspect('equal')
    x0, y0 = 0, 0
    ax.plot([x0, x0+b_cm], [y0, y0], 'k-', linewidth=2.5)
    ax.plot([x0+b_cm, x0+b_cm], [y0, y0+h_cm], 'k-', linewidth=2.5)
    ax.plot([x0+b_cm, x0], [y0+h_cm, y0+h_cm], 'k-', linewidth=2.5)
    ax.plot([x0, x0], [y0+h_cm, y0], 'k-', linewidth=2.5)
    vis_hook = min(hook_len_cm, b_cm/4.0, h_cm/4.0)
    hx = vis_hook * 0.707
    hy = -vis_hook * 0.707
    ax.plot([x0, x0 + hx], [y0+h_cm, y0+h_cm + hy], 'k-', linewidth=2.5)
    ax.annotate(f"Gancho 135°", xy=(x0 + hx + 0.2, y0+h_cm + hy - 0.2), fontsize=7.5, color='darkred')
    ax.annotate(f"{b_cm:.0f} cm", xy=(b_cm/2, y0-0.8), ha='center', fontsize=9)
    ax.annotate(f"{h_cm:.0f} cm", xy=(x0-0.8, h_cm/2), ha='right', va='center', fontsize=9)
    ax.set_xlim(-hook_len_cm*0.3, b_cm + hook_len_cm*0.6)
    ax.set_ylim(-hook_len_cm*0.5, h_cm + hook_len_cm*0.9)
    ax.axis('off')
    title = f"Estribo {bar_name or 'E1'} - Ø{bar_diam_mm:.0f} mm - Perímetro {2*(b_cm+h_cm):.0f} cm"
    ax.set_title(title, fontsize=9, fontweight='bold')
    return fig

# =============================================================================
# INTERFAZ: PESTAÑAS PRINCIPALES
# =============================================================================
tab_res, tab_col, tab_sol, tab_diag, tab_3d, tab_dxf, tab_mem, tab_qty, tab_apu = st.tabs([
    " Resultados",
    " Columnas",
    " Solera",
    " Diagrama 2D",
    " 3D",
    " DXF",
    " Memoria",
    " Cantidades",
    " Presupuesto"
])

with tab_res:
    st.subheader(f"Muro: {nombre_muro} | {piso_muro}")
    ms1, ms2, ms3 = st.columns(3)
    sigma_m_out = sigma_m_kgcm2 * factor_fuerza if unidades_salida == "Toneladas fuerza (tonf, tonf·m)" else sigma_m_kgcm2
    Fa_out = Fa * factor_fuerza if unidades_salida == "Toneladas fuerza (tonf, tonf·m)" else Fa
    ms1.metric("σm / Fa [kg/cm²]", f"{sigma_m_out:.2f} / {Fa_out:.2f}",
               delta=" OK" if ok_sigma else " FALLA",
               delta_color="normal" if ok_sigma else "inverse")
    ms2.metric(f"Vm1 [{unidad_fuerza}]", f"{Vm1 * factor_fuerza:.2f}")
    ms3.metric("Condición", cond_mur, delta="Refuerzo Hz" if necesita_Hz else "Sin refuerzo Hz",
               delta_color="inverse" if necesita_Hz else "normal")
    st.markdown("---")

    if ok_sigma and ok_rho_h:
        st.success(f"**DISEÑO ESTRUCTURAL OK SEGÚN {norma_ac}** — El muro cumple con las comprobaciones normativas limitantes.")
    else:
        st.error(f"**FALLO EN DISEÑO SEGÚN {norma_ac}** — Revise las verificaciones en rojo abajo.")

    r1, r2 = st.columns(2)
    # Referencias normativas por país
    _pais = nc["pais"]
    if _pais == "pe":
        art_ref  = "Art. 7.1.1b — E.070 (Perú)"
        art_ref2 = "Art. 8.5.3 — E.070 (Perú)"
        art_ref3 = "Art. 8.5.2 — E.070 (Perú)"
        art_ref6 = "Art. 8.6.1 — E.070 (Perú)"
    elif _pais == "co":
        art_ref  = "D.5.4.1 — NSR-10 (Colombia)"
        art_ref2 = "D.5.3 — NSR-10 (Colombia)"
        art_ref3 = "D.5.2 — NSR-10 (Colombia)"
        art_ref6 = "D.5.5.2 — NSR-10 (Colombia)"
    elif _pais == "ec":
        art_ref  = "§6.3 — NEC-SE-MP (Ecuador)"
        art_ref2 = "§6.5 — NEC-SE-MP (Ecuador)"
        art_ref3 = "§6.4 — NEC-SE-MP (Ecuador)"
        art_ref6 = "§6.6 — NEC-SE-MP (Ecuador)"
    else:  # ve
        art_ref  = "§7.2 — COVENIN 1753 (Venezuela)"
        art_ref2 = "§7.4 — COVENIN 1753 (Venezuela)"
        art_ref3 = "§7.3 — COVENIN 1753 (Venezuela)"
        art_ref6 = "§7.5 — COVENIN 1753 (Venezuela)"
    with r1:
        st.markdown("#### 01 — Características del Muro")
        st.info(f"f'm={fm_in}{nc['unidad']} | v'm={vm_in}{nc['unidad']} | E={E_in}{nc['unidad']}\n"
                f"L={L_m_in}m | t={t_m_in}m | h={h_m_in}m\nLadrillo: {lad_sel}")
        st.markdown("#### 02 — Cargas")
        qty_table([
            ("Pm (D+L)", f"{Pm_in:.2f} {nc['fuerza']}"),
            ("Pg (D+0.25L)", f"{Pg_in:.2f} {nc['fuerza']}"),
            ("Ve (sismo mod.)", f"{Ve_in:.2f} {nc['fuerza']}"),
            ("Me (mom. mod.)", f"{Me_in:.2f} {nc['fuerza']}·m"),
            ("Vu (sismo sev.)", f"{Vu_in:.2f} {nc['fuerza']}"),
            ("Mu (mom. sev.)", f"{Mu_in:.2f} {nc['fuerza']}·m"),
        ])
    with r2:
        st.markdown(f"#### 03 — Esfuerzos Verticales `{art_ref}`")
        st.markdown(f"- σ_m = {sigma_m_kgcm2:.2f} kg/cm²\n"
                    f"- Fa = min({Fa_inner:.2f}, {Fa_lim:.2f}) = **{Fa:.2f} kg/cm²**")
        if ok_sigma:
            st.success(f"σ_m ({sigma_m_kgcm2:.2f}) ≤ Fa ({Fa:.2f}) **— CUMPLE**")
        else:
            st.error(f"σ_m ({sigma_m_kgcm2:.2f}) > Fa ({Fa:.2f}) **— NO CUMPLE**")
            st.warning("**Sugerencia:** Aumentar la longitud del muro (`L`), el espesor (`t`), o escoger un Ladrillo de mayor resistencia a compresión (`f'm`).")
        st.markdown(f"#### 04 — Cortante Base `{art_ref2}`")
        st.markdown(f"- α = {alpha:.3f} (Ve·L/Me) ∈ [1/3, 1.0]\n"
                    f"- **Vm = {Vm1:.2f} ton** | Vui = {Vui_calc:.2f} ton | Mui = {Mui_calc:.2f} ton·m")
        if Vui_calc <= 2.0 * Vm1:
            st.success(f"Vui ({Vui_calc:.2f}) ≤ 2·Vm ({2*Vm1:.2f}) **— CUMPLE**")
        else:
            st.error(f"Vui ({Vui_calc:.2f}) > 2·Vm ({2*Vm1:.2f}) **— SECCIÓN INSUFICIENTE**")
        st.markdown(f"#### 05 — Fisuración `{art_ref3}`")
        if ok_fisura:
            st.success(f"Ve ({Ve:.2f}) ≤ 0.55·Vm ({0.55*Vm1:.2f}) — No Agrietado **— CUMPLE**")
        else:
            st.error(f"Ve ({Ve:.2f}) > 0.55·Vm ({0.55*Vm1:.2f}) — MURO AGRIETADO")
            st.warning("**Sugerencia:** Para evitar el agrietamiento, aumente el espesor del muro (`t`), o use un Ladrillo con mayor resistencia al corte (`v'm`).")
        st.markdown(f"#### 06 — Refuerzo Horizontal `{art_ref6}`")
        if necesita_Hz:
            estado_hz = ' CUMPLE' if ok_rho_h else ' NO CUMPLE'
            st.warning(f"⚠ 1 {bar_hor_sel} @ {s_hor:.0f}cm | ρh={rho_h_prov:.5f} — {estado_hz}")
            st.caption(f"ρh_min = {rho_h_min:.4f} | ρh_prov = {rho_h_prov:.5f}")
        else:
            st.success("No requiere refuerzo horizontal según la norma aplicada")

with tab_col:
    if _pais == "pe":
        art_col = "Art. 8.6.3 — E.070 (Perú)"
    elif _pais == "co":
        art_col = "D.5.5 — NSR-10 (Colombia)"
    elif _pais == "ec":
        art_col = "§6.7 — NEC-SE-MP (Ecuador)"
    else:
        art_col = "§7.6 — COVENIN 1753 (Venezuela)"
    st.subheader(f"Columnas de Confinamiento ({art_col})")
    cc1, cc2 = st.columns(2)
    with cc1:
        st.markdown("#### 07-08 — Fuerzas de Diseño")
        qty_table([
            ("Nc", Nc), ("Lm", f"{Lm:.2f} m"),
            ("F = Mu/L", f"{F_force:.2f} ton"),
            ("Pc = Pg/Nc", f"{Pc_col:.2f} ton"),
            ("Vc ext.", f"{Vc_ext:.2f} ton"),
            ("T1 ext.", f"{T1_ext:.2f} ton"),
            ("T2 ext.", f"{T2_ext:.2f} ton"),
            ("C1 ext.", f"{C1_ext:.2f} ton"),
        ])
    with cc2:
        st.markdown(f"#### 09a — Acero Vertical ({design_mode})")
        if design_mode == "Agrietado":
            qty_table([
                ("Asf (corte-fricción μ=1)", f"{Asf_col:.2f} cm²"),
                ("Ast (tracción)", f"{Ast_col:.2f} cm²"),
                ("As req = Asf + Ast", f"{As_req:.2f} cm²"),
                ("As mín (4Ø8mm)", f"{As_min_four_d8:.2f} cm²"),
            ])
        else:
            qty_table([
                ("T usar = max(T1, 0)", f"{T_use:.2f} ton"),
                ("As req = T/(φ·fy)", f"{As_req:.2f} cm²"),
                ("As mín (4Ø8mm)", f"{As_min_four_d8:.2f} cm²"),
            ])
        st.success(f"**{n_bars_f} {bar_col_sel}** → As prov = {As_prov_f:.2f} cm²")
        st.markdown("#### 09b — Núcleo Concreto An")
        qty_table([("An compresión", f"{An_comp:.1f} cm²"), ("mín 15t", f"{15*t_m:.0f} cm²")])
        st.markdown("#### 09c — Corte-Fricción Acf")
        qty_table([("Acf", f"{Acf_col:.1f} cm²"), ("mín 15t", f"{15*t_m:.0f} cm²")])
        st.markdown("#### 09d — Dimensiones Finales")
        qty_table([
            ("b × d col.", f"{col_b_min:.0f} × {col_d_use:.0f} cm"),
            ("Ac", f"{Ac_final:.0f} cm²"),
        ])
    st.markdown("---")
    ce1, ce2 = st.columns(2)
    with ce1:
        st.markdown("#### 09e — Estribos")
        qty_table([
            ("s1", f"{s1_col:.1f} cm"), ("s2", f"{s2_col:.1f} cm"),
            ("s3 = d/4 ≥5", f"{s3_col:.1f} cm"), ("s4", "10.0 cm"),
            ("s diseño", f"{s_col:.1f} cm"),
        ])
        st.info(f"{est_sel} @ **{s_col:.1f} cm**")
    with ce2:
        st.markdown("#### 09f-g — Zona Confinada")
        st.markdown(f"Z = max(1.5·d, 45) = **{Z_col:.1f} cm**")
        st.success(f"1@5cm · 4@{s_col:.0f}cm · Rto@25cm")

with tab_sol:
    if _pais == "pe":
        art_sol = "Art. 8.6.3-b — E.070 (Perú)"
        _nombre_viga = "Viga Solera"
    elif _pais == "co":
        art_sol = "D.5.6 — NSR-10 D (Colombia)"
        _nombre_viga = "Viga de Amarre"
    elif _pais == "ec":
        art_sol = "§6.8 — NEC-SE-MP (Ecuador)"
        _nombre_viga = "Viga Solera"
    else:
        art_sol = "§7.7 — COVENIN 1753 (Venezuela)"
        _nombre_viga = "Viga Corona"
    st.subheader(f" {_nombre_viga} ({art_sol})")
    sv1, sv2 = st.columns(2)
    with sv1:
        qty_table([
            ("Ts = Vm1·Lm/(2L)", f"{Ts_sol:.2f} ton"),
            ("As req", f"{As_sol_req:.2f} cm²"),
            ("As mín", f"{As_sol_min:.2f} cm²"),
            ("As diseño", f"{As_sol:.2f} cm²"),
        ])
        st.success(f" {n_sol} {bar_col_sel} → As={As_sol_prov:.2f} cm²")
    with sv2:
        qty_table([
            ("b × h", f"{vs_b_in:.0f} × {vs_h_in:.0f} cm"),
            ("Acs", f"{Acs_sol:.0f} cm²"),
        ])
        st.info(f"{est_sel}: 1@5cm · 4@10cm · Rto@25cm")

with tab_diag:
    # Nombre de la viga según norma (también disponible fuera del tab)
    _nombre_viga_local = "Viga de Amarre" if _pais == "co" else ("Viga Corona" if _pais == "ve" else "Viga Solera")
    st.subheader(f"Diagrama 2D del Muro y Sección Transversal ({_nombre_viga_local})")
    # Figura con 4 columnas: [alzado, espacio, corte columna, corte viga]
    fig2d = plt.figure(figsize=(16, 9))
    fig2d.patch.set_facecolor('#1a1a2e')
    gs = fig2d.add_gridspec(1, 4, width_ratios=[2.5, 0.1, 1, 1], wspace=0.05)
    ax     = fig2d.add_subplot(gs[0, 0])   # Alzado
    ax_sec = fig2d.add_subplot(gs[0, 2])   # Corte columna
    ax_viga= fig2d.add_subplot(gs[0, 3])   # Corte viga
    
    ax.set_facecolor('#1a1a2e')
    ax_sec.set_facecolor('#1a1a2e')
    ax_viga.set_facecolor('#1a1a2e')
    
    Lw = L_m_in*100; Hw = h_m_in*100
    bc = col_b_min; hv = vs_h_in

    # ── MURO (ladrillos)
    ax.add_patch(patches.Rectangle((bc, 0), Lw-2*bc, Hw-hv, fc='#c8633b', ec='#8B4513', lw=0.8))
    hilada_h = (lad["H"]+junta)/10; lad_l = (lad["L"]+junta)/10
    n_hil = int((Hw-hv)/hilada_h)
    for i in range(n_hil+1):
        yh = i*hilada_h
        if yh <= Hw-hv: ax.plot([bc, Lw-bc], [yh, yh], color='#8B4513', lw=0.4, alpha=0.7)
    for row in range(n_hil):
        off = lad_l/2 if row%2==1 else 0
        for ci in range(int((Lw-2*bc)/lad_l)+2):
            x = bc + ci*lad_l + off
            if bc <= x <= Lw-bc: ax.plot([x, x], [row*hilada_h, (row+1)*hilada_h], color='#8B4513', lw=0.4, alpha=0.7)

    # ── COLUMNAS confinamiento + zonas confinadas
    for xc in [0, Lw-bc]:
        ax.add_patch(patches.Rectangle((xc, 0), bc, Hw-hv, fc='#b1b4bc', ec='white', lw=1.5))
        z_h = min(Z_col, Hw-hv)
        ax.add_patch(patches.Rectangle((xc, 0), bc, z_h, fc='none', ec='yellow', lw=1.5, ls='--'))
        ax.add_patch(patches.Rectangle((xc, Hw-hv-z_h), bc, z_h, fc='none', ec='yellow', lw=1.5, ls='--'))
        ax.text(xc+bc/2, z_h/2, f"Z.C: {Z_col:.0f}cm", color='yellow', ha='center', fontsize=6, rotation=90)

    # ── VIGA (nom. según país) en alzado + flejes esquemáticos
    ax.add_patch(patches.Rectangle((0, Hw-hv), Lw, hv, fc='#b1b4bc', ec='white', lw=2))
    # Varillas long. en viga (alzado)
    _db_cm_v = db_col/10
    for yb_v in [Hw-hv+recub_col/2, Hw-recub_col/2]:
        ax.plot([recub_col, Lw-recub_col], [yb_v, yb_v], color='#cc2222', lw=1.5)
    # Flejes en viga (alzado) cada ~15cm esquemático
    _s_fleje_v = max(vs_b_in*0.5, 10.0)
    _n_flejes = int(Lw / _s_fleje_v)
    for _fi in range(1, _n_flejes):
        _xf = _fi * _s_fleje_v
        ax.plot([_xf, _xf], [Hw-hv+recub_col, Hw-recub_col], color='#ffd700', lw=1.0, alpha=0.8)
    ax.text(Lw/2, Hw-hv/2, f"{_nombre_viga_local} {vs_b_in:.0f}×{vs_h_in:.0f}cm | {n_sol} {bar_col_sel}",
            color='black', ha='center', va='center', fontsize=8, weight='bold')
    if necesita_Hz:
        n_hz = min(n_hiladas_hor, 8)
        for i in range(n_hz):
            yh2 = i*((Hw-hv)/max(n_hz,1)) + (Hw-hv)/max(n_hz*2,2)
            ax.plot([bc, Lw-bc], [yh2, yh2], color='#4287f5', lw=2.5, alpha=0.9, ls='-')
        ax.text(Lw*0.5, (Hw-hv)*0.4, f"Ref. Hz: 1 {bar_hor_sel} @ {s_hor:.0f}cm",
                color='#ffffff', fontsize=8, ha='center', va='center', weight='bold',
                bbox=dict(boxstyle='round', fc='#4287f5', ec='#225cb2', pad=0.4))
    
    # ── Varillas longitudinales en columnas (alzado)
    db_cm = db_col/10
    xs_bars = [recub_col, bc-recub_col-db_cm]
    for xc_base in [0, Lw-bc]:
        for xb in xs_bars:
            ax.plot([xc_base+xb+db_cm/2, xc_base+xb+db_cm/2], [0, Hw-hv], color='#cc2222', lw=1.5)
            
    # ── Estribos en columnas (esquemáticos)
    for xc_base in [0, Lw-bc]:
        for y_est in np.linspace(5, Hw-hv-5, 20):
            ax.plot([xc_base+recub_col, xc_base+bc-recub_col], [y_est, y_est], color='#ffd700', lw=0.8)
            
    ax.annotate('', xy=(Lw+10,0), xytext=(Lw+10,Hw), arrowprops=dict(arrowstyle='<->', color='white'))
    ax.text(Lw+15, Hw/2, f"{h_m_in:.2f} m", color='white', va='center', fontsize=10)
    ax.annotate('', xy=(0,-15), xytext=(Lw,-15), arrowprops=dict(arrowstyle='<->', color='white'))
    ax.text(Lw/2, -22, f"L = {L_m_in} m", color='white', ha='center', fontsize=10)
    ax.text(bc/2, Hw/2, "Col. Conf.", color='black', weight='bold', ha='center', fontsize=7, rotation=90)
    ax.text(Lw-bc/2, Hw/2, "Col. Conf.", color='black', weight='bold', ha='center', fontsize=7, rotation=90)
    
    # ── CORTE TRANSVERSAL COLUMNA
    Wc = col_b_min; Hc = col_d_use
    ax_sec.add_patch(patches.Rectangle((0, 0), Wc, Hc, fc='#b1b4bc', ec='white', lw=2))
    ax_sec.text(Wc/2, Hc + 3, f"{n_bars_f} {bar_col_sel}", color='white', ha='center', weight='bold', fontsize=8)
    ax_sec.text(Wc/2, -5, f"{Wc:.0f}cm", color='white', ha='center', fontsize=8)
    ax_sec.text(-5, Hc/2, f"{Hc:.0f}cm", color='white', va='center', ha='right', fontsize=8)
    re = recub_col
    ax_sec.add_patch(patches.Rectangle((re, re), Wc-2*re, Hc-2*re, fc='none', ec='#4287f5', lw=2))
    x_pos = [re, Wc-re]; y_pos = [re, Hc-re]
    if n_bars_f == 4:
        for xb in x_pos:
            for yb in y_pos:
                ax_sec.add_patch(patches.Circle((xb, yb), db_cm/2*1.5, fc='#cc2222', ec='white', lw=0.5))
    else:
        for xb in x_pos:
            for yb in np.linspace(re, Hc-re, max(2, n_bars_f//2)):
                ax_sec.add_patch(patches.Circle((xb, yb), db_cm/2*1.5, fc='#cc2222', ec='white', lw=0.5))
    ax_sec.set_aspect('equal')
    ax_sec.set_xlim(-15, Wc+15); ax_sec.set_ylim(-15, Hc+15)
    ax_sec.axis('off')
    ax_sec.set_title("Corte Columna", color='white', fontsize=9, pad=4)

    # ── CORTE TRANSVERSAL VIGA DE AMARRE / SOLERA
    Wv = vs_b_in; Hv = vs_h_in
    ax_viga.add_patch(patches.Rectangle((0, 0), Wv, Hv, fc='#9ab0c4', ec='white', lw=2))
    re_v = recub_col
    # Estribo de la viga
    ax_viga.add_patch(patches.Rectangle((re_v, re_v), Wv-2*re_v, Hv-2*re_v, fc='none', ec='#ffd700', lw=2))
    # Gancho 135° esquemático
    _hk = re_v * 0.8
    ax_viga.plot([re_v, re_v+_hk*0.7], [Hv-re_v, Hv-re_v+_hk*0.7], color='#ffd700', lw=2)
    # Varillas longitudinales (n_sol barras, distribuída en 2 capas)
    _n_sup = n_sol // 2; _n_inf = n_sol - _n_sup
    _db_v_cm = db_col/10
    _spacing_inf = (Wv - 2*re_v) / max(_n_inf-1, 1) if _n_inf > 1 else 0
    _spacing_sup = (Wv - 2*re_v) / max(_n_sup-1, 1) if _n_sup > 1 else 0
    for _i in range(_n_inf):
        _xv = re_v + _i * _spacing_inf if _n_inf > 1 else Wv/2
        ax_viga.add_patch(patches.Circle((_xv, re_v), _db_v_cm/2*1.5, fc='#cc2222', ec='white', lw=0.5))
    for _i in range(_n_sup):
        _xv = re_v + _i * _spacing_sup if _n_sup > 1 else Wv/2
        ax_viga.add_patch(patches.Circle((_xv, Hv-re_v), _db_v_cm/2*1.5, fc='#cc2222', ec='white', lw=0.5))
    ax_viga.text(Wv/2, Hv + 3, f"{n_sol} {bar_col_sel}", color='white', ha='center', weight='bold', fontsize=8)
    ax_viga.text(Wv/2, -5, f"{Wv:.0f}cm", color='white', ha='center', fontsize=8)
    ax_viga.text(-5, Hv/2, f"{Hv:.0f}cm", color='white', va='center', ha='right', fontsize=8)
    ax_viga.text(Wv/2, Hv/2, f"{est_sel}", color='#ffd700', ha='center', va='center', fontsize=7,
                 bbox=dict(boxstyle='round,pad=0.2', fc='#1a1a2e', ec='none', alpha=0.7))
    ax_viga.set_aspect('equal')
    ax_viga.set_xlim(-15, Wv+15); ax_viga.set_ylim(-15, Hv+15)
    ax_viga.axis('off')
    ax_viga.set_title(f"Corte {_nombre_viga_local}", color='white', fontsize=9, pad=4)
    ax.set_aspect('equal')
    ax.axis('off')
    ax.set_xlim(-10, Lw+30); ax.set_ylim(-30, Hw+15)
    st.pyplot(fig2d)
    buf_2d = io.BytesIO()
    fig2d.savefig(buf_2d, format='png', dpi=150, bbox_inches='tight', facecolor='#1a1a2e')
    buf_2d.seek(0); plt.close(fig2d)
    st.markdown("---\n####  Diagramas M y V")
    fig_mv, (ax_v, ax_m) = plt.subplots(1, 2, figsize=(10, 5))
    fig_mv.patch.set_facecolor('#1a1a2e')
    for axi in [ax_v, ax_m]: axi.set_facecolor('#1a1a2e')
    y_arr = np.linspace(0, h_m_in, 100)
    Ve_disp = Ve_in if nc["fuerza"] == "kN" else Ve
    Me_disp = Me_in if nc["fuerza"] == "kN" else Me
    v_arr = Ve_disp * (1 - y_arr/h_m_in)
    m_arr = Me_disp * y_arr * (1 - y_arr/(2*h_m_in))
    for axi, data, color, xlabel in [(ax_v, v_arr, '#00d4ff', f"V [{unidad_fuerza}]"),
                                      (ax_m, m_arr, '#ff9500',  f"M [{unidad_mom}]")]:
        axi.plot(data, y_arr, color=color, lw=2.5)
        axi.fill_betweenx(y_arr, 0, data, alpha=0.3, color=color)
        axi.set_xlabel(xlabel, color='white'); axi.set_ylabel("Altura [m]", color='white')
        axi.tick_params(colors='white')
        for s in ['top','right']: axi.spines[s].set_visible(False)
        for s in ['bottom','left']: axi.spines[s].set_color('white')
    ax_v.set_title("Cortante V", color='white')
    ax_m.set_title("Momento M", color='white')
    st.pyplot(fig_mv)
    buf_mv = io.BytesIO()
    fig_mv.savefig(buf_mv, format='png', dpi=150, bbox_inches='tight', facecolor='#1a1a2e')
    buf_mv.seek(0); plt.close(fig_mv)

with tab_3d:
    st.subheader("Visualización 3D")
    Lw3=L_m_in; Hw3=h_m_in; tw3=t_m_in
    bc3=col_b_min/100; hv3=vs_h_in/100
    fig3d = go.Figure()
    fig3d.add_trace(go.Mesh3d(
        x=[bc3,Lw3-bc3,Lw3-bc3,bc3,bc3,Lw3-bc3,Lw3-bc3,bc3],
        y=[0,0,tw3,tw3,0,0,tw3,tw3],
        z=[0,0,0,0,Hw3-hv3,Hw3-hv3,Hw3-hv3,Hw3-hv3],
        i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
        opacity=0.6, color='#c8633b', name='Muro'))
    st.subheader(_t("Modelo 3D Avanzado - Muro Confinado", "Advanced 3D Model - Confined Wall"))
    fig3d = go.Figure()

    # 1. MURO DE ALBAÑILERÍA (Separado de columnas)
    fig3d.add_trace(go.Mesh3d(
        x=[bc3, Lw3-bc3, Lw3-bc3, bc3, bc3, Lw3-bc3, Lw3-bc3, bc3],
        y=[0, 0, tw3, tw3, 0, 0, tw3, tw3],
        z=[0, 0, 0, 0, Hw3-hv3, Hw3-hv3, Hw3-hv3, Hw3-hv3],
        i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
        opacity=1.0, color='#9c4a2f', name='Mampostería', flatshading=True))

    # 2. COLUMNAS CONFINANTES (Concreto transparente)
    for xc3 in [0, Lw3-bc3]:
        fig3d.add_trace(go.Mesh3d(
            x=[xc3, xc3+bc3, xc3+bc3, xc3, xc3, xc3+bc3, xc3+bc3, xc3],
            y=[0, 0, tw3, tw3, 0, 0, tw3, tw3], z=[0, 0, 0, 0, Hw3, Hw3, Hw3, Hw3],
            i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
            opacity=0.25, color='#adc4ce', name='Concreto Columna', showlegend=bool(xc3==0)))

    # 3. VIGA SOLERA (Concreto transparente)
    fig3d.add_trace(go.Mesh3d(
        x=[0, Lw3, Lw3, 0, 0, Lw3, Lw3, 0], y=[0, 0, tw3, tw3, 0, 0, tw3, tw3],
        z=[Hw3-hv3, Hw3-hv3, Hw3-hv3, Hw3-hv3, Hw3, Hw3, Hw3, Hw3],
        i=[0,0,4,4,1,5,2,6,3,7,0,4], j=[1,2,5,6,5,6,6,7,7,4,4,7], k=[2,3,6,7,6,7,7,4,4,5,1,3],
        opacity=0.25, color='#96b6c5', name='Concreto Solera', showlegend=True))

    # 4. ACERO DE COLUMNAS (Longitudinal + Estribos)
    xs3 = [recub_col/100, bc3 - recub_col/100]; ys3 = [recub_col/100, tw3 - recub_col/100]
    for xc3 in [0, Lw3-bc3]:
        for bx in xs3:
            for by in ys3: # 4 varillas esquineras por columna
                fig3d.add_trace(go.Scatter3d(
                    x=[xc3+bx, xc3+bx], y=[by, by], z=[0, Hw3],
                    mode='lines', line=dict(color='#ff2e2e', width=6),
                    name='Acero Long. Col', showlegend=bool(xc3==0 and bx==xs3[0] and by==ys3[0])))
        # Flejes (Estribos)
        n_estribos = int(Hw3 / 0.15) # @15cm visual
        if n_estribos > 1:
            for i_e in range(1, n_estribos):
                ze = i_e * 0.15
                fig3d.add_trace(go.Scatter3d(
                    x=[xc3+xs3[0], xc3+xs3[1], xc3+xs3[1], xc3+xs3[0], xc3+xs3[0]], 
                    y=[ys3[0], ys3[0], ys3[1], ys3[1], ys3[0]], z=[ze, ze, ze, ze, ze],
                    mode='lines', line=dict(color='#ffd93d', width=3),
                    name='Flejes Columna', showlegend=bool(xc3==0 and i_e==1)))

    # 5. ACERO VIGA SOLERA (Longitudinal + Estribos)
    y_sol = ys3; z_sol = [Hw3-hv3 + recub_col/100, Hw3 - recub_col/100]
    for by in y_sol:
        for bz in z_sol:
            fig3d.add_trace(go.Scatter3d(
                x=[0, Lw3], y=[by, by], z=[bz, bz],
                mode='lines', line=dict(color='#ff2e2e', width=6),
                name='Acero Long. Solera', showlegend=bool(by==y_sol[0] and bz==z_sol[0])))
    n_estribos_s = int(Lw3 / 0.15)
    if n_estribos_s > 1:
        for i_s in range(1, n_estribos_s):
            xe = i_s * 0.15
            fig3d.add_trace(go.Scatter3d(
                x=[xe, xe, xe, xe, xe], y=[y_sol[0], y_sol[1], y_sol[1], y_sol[0], y_sol[0]], 
                z=[z_sol[0], z_sol[0], z_sol[1], z_sol[1], z_sol[0]],
                mode='lines', line=dict(color='#ffd93d', width=3),
                name='Flejes Solera', showlegend=bool(i_s==1)))

    # 6. REFUERZO HORIZONTAL EN MURO
    if necesita_Hz and n_hiladas_hor > 0:
        for i in range(min(n_hiladas_hor, 10)):
            zh = i * ((Hw3-hv3)/max(min(n_hiladas_hor, 10),1)) + 0.2
            if zh < (Hw3-hv3):
                fig3d.add_trace(go.Scatter3d(
                    x=[bc3, Lw3-bc3], y=[tw3/2, tw3/2], z=[zh, zh],
                    mode='lines', line=dict(color='#00ffcc', width=4),
                    name='Ref. Hz Muro', showlegend=bool(i==0)))

    fig3d.update_layout(
        scene=dict(aspectmode='data', bgcolor='#1e1e2f',
                   xaxis=dict(title='L(m)', showgrid=False), yaxis=dict(title='t(m)', showgrid=False), zaxis=dict(title='h(m)', showgrid=False),
                   camera=dict(eye=dict(x=1.5, y=-1.5, z=1.2))),
        paper_bgcolor='#1e1e2f', font=dict(color='white'), height=700, margin=dict(l=0, r=0, t=50, b=0),
        legend=dict(x=1.05, y=0.5, bgcolor='rgba(0,0,0,0.5)', bordercolor='white', borderwidth=1),
        title=dict(text=f"Detalle 3D Muro {nombre_muro} | Armadura y Estribos", font=dict(color='white')))
    st.plotly_chart(fig3d, use_container_width=True)


with tab_dxf:
    st.subheader("Plano DXF — Planta, Elevación y Sección")
    try:
        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                 dxf_dim_horiz, dxf_dim_vert, dxf_rotulo,
                                 dxf_leyenda, dxf_rotulo_campos)
        _USE_H = True
    except ImportError:
        _USE_H = False
    doc_dxf = ezdxf.new('R2010'); msp = doc_dxf.modelspace()
    if _USE_H:
        dxf_setup(doc_dxf, 50)
        dxf_add_layers(doc_dxf)
    for lay, col_num in [('MURO',5),('COLUMNAS',4),('ACERO',1),('SOLERA',3),('TEXTO',2),('COTAS',6)]:
        if lay not in doc_dxf.layers: doc_dxf.layers.add(lay, color=col_num)
    Lw_d=L_m_in; Hw_d=h_m_in; tw_d=t_m_in
    bc_d=col_b_min/100; dc_d=col_d_use/100; hv_d=vs_h_in/100; off_y=Hw_d+2.0
    msp.add_lwpolyline([(0,0),(Lw_d,0),(Lw_d,Hw_d),(0,Hw_d),(0,0)],close=True,dxfattribs={'layer':'MURO'})
    for xc_d in [0, Lw_d-bc_d]:
        msp.add_lwpolyline([(xc_d,0),(xc_d+bc_d,0),(xc_d+bc_d,Hw_d-hv_d),(xc_d,Hw_d-hv_d),(xc_d,0)],
                           close=True,dxfattribs={'layer':'COLUMNAS'})
        msp.add_text(f"{n_bars_f} {bar_col_sel}",dxfattribs={'layer':'TEXTO','height':0.06,'insert':(xc_d+bc_d/2,Hw_d/2)})
        Zd=Z_col/100
        msp.add_lwpolyline([(xc_d,0),(xc_d+bc_d,0),(xc_d+bc_d,Zd),(xc_d,Zd),(xc_d,0)],
                           close=True,dxfattribs={'layer':'COTAS'})
        msp.add_text(f"Z={Z_col:.0f}cm",dxfattribs={'layer':'TEXTO','height':0.04,'insert':(xc_d+bc_d/2,Zd/2)})
        for i in range(min(int(Hw_d/s_col*100),20)):
            ye=i*s_col/100+s_col/200
            if ye<Hw_d-hv_d:
                msp.add_line((xc_d+0.02,ye),(xc_d+bc_d-0.02,ye),dxfattribs={'layer':'ACERO'})
    if necesita_Hz:
        for i in range(min(n_hiladas_hor,12)):
            yh_d=i*s_hor/100+s_hor/200
            if yh_d<Hw_d-hv_d: msp.add_line((bc_d,yh_d),(Lw_d-bc_d,yh_d),dxfattribs={'layer':'ACERO'})
    msp.add_lwpolyline([(0,Hw_d-hv_d),(Lw_d,Hw_d-hv_d),(Lw_d,Hw_d),(0,Hw_d),(0,Hw_d-hv_d)],
                       close=True,dxfattribs={'layer':'SOLERA'})
    msp.add_text(f"Solera {vs_b_in:.0f}x{vs_h_in:.0f}|{n_sol} {bar_col_sel}",
                 dxfattribs={'layer':'TEXTO','height':0.05,'insert':(Lw_d/2,Hw_d-hv_d/2)})
    # Planta
    msp.add_lwpolyline([(0,off_y),(Lw_d,off_y),(Lw_d,off_y+tw_d),(0,off_y+tw_d),(0,off_y)],
                       close=True,dxfattribs={'layer':'MURO'})
    for xc_d in [0,Lw_d-bc_d]:
        msp.add_lwpolyline([(xc_d,off_y),(xc_d+bc_d,off_y),(xc_d+bc_d,off_y+tw_d),(xc_d,off_y+tw_d),(xc_d,off_y)],
                           close=True,dxfattribs={'layer':'COLUMNAS'})
    # Sección columna
    ox2=Lw_d+0.6
    msp.add_lwpolyline([(ox2,0),(ox2+bc_d,0),(ox2+bc_d,dc_d),(ox2,dc_d),(ox2,0)],
                       close=True,dxfattribs={'layer':'COLUMNAS'})
    msp.add_text(f"Secc. Col {col_b_min:.0f}x{col_d_use:.0f}cm",dxfattribs={'layer':'TEXTO','height':0.05,'insert':(ox2,dc_d+0.04)})
    n2x=2; n2y=max(2,n_bars_f//2)
    for xi in range(n2x):
        for yi in range(n2y):
            xb2=ox2+recub_col/100+xi*(bc_d-2*recub_col/100)/max(n2x-1,1)
            yb2=recub_col/100+yi*(dc_d-2*recub_col/100)/max(n2y-1,1)
            msp.add_circle((xb2,yb2),db_col/2000,dxfattribs={'layer':'ACERO'})
    # Cotas y rótulo profesional
    if _USE_H:
        TH = 0.025 * 50
        dxf_dim_horiz(msp, 0, Lw_d, -0.5, f"L = {L_m_in:.2f} m", 50)
        dxf_dim_vert(msp, -0.5, 0, Hw_d, f"h = {h_m_in:.2f} m", 50)
        dxf_text(msp, Lw_d/2, Hw_d+0.3, "ELEVACION", "EJES", h=TH*1.2, ha="center")
        dxf_text(msp, Lw_d/2, off_y+tw_d+0.15, "PLANTA", "EJES", h=TH*1.2, ha="center")
        dxf_text(msp, ox2+bc_d/2, dc_d+0.25, "SECCION COLUMNA", "EJES", h=TH*0.9, ha="center")
        dxf_leyenda(msp, Lw_d+1.5, Hw_d-0.3, [
            ("MURO",      "Muro mamposteria"),
            ("COLUMNAS",  f"Columna {col_b_min:.0f}x{col_d_use:.0f}cm"),
            ("ACERO",     f"Acero {db_col:.0f}mm/Estribo {db_est:.0f}mm"),
            ("SOLERA",    f"Solera {vs_b_in:.0f}x{vs_h_in:.0f}cm"),
        ], 50)
        _cam = dxf_rotulo_campos(f"Albanileria Confinada {nombre_muro} L={L_m_in:.1f}m h={h_m_in:.1f}m", norma_ac, "001")
        dxf_rotulo(msp, _cam, 0, -4.5, rot_w=9, rot_h=3, escala=50)
    _out=io.StringIO(); doc_dxf.write(_out)
    st.download_button("Descargar DXF",data=_out.getvalue().encode('utf-8'),
                       file_name=f"AlbConfinada_{nombre_muro}.dxf",mime="application/dxf")
    st.info("Plano incluye: Elevacion con acero, Planta y Seccion columna.")

with tab_mem:
    st.subheader("Memoria de Cálculo DOCX")
    if st.button("Generar Memoria DOCX"):
        doc = Document()
        doc.add_heading(f"ALBAÑILERÍA CONFINADA — {norma_ac}", 0)
        p=doc.add_paragraph(); p.add_run(f"{nombre_muro} | {piso_muro} | {datetime.now().strftime('%d/%m/%Y')}").bold=True
        doc.add_heading("01. Datos del Muro", 1)
        for k,v in [("Norma",norma_ac),("f'm",f"{fm_in} {nc['unidad']}"),("v'm",f"{vm_in} {nc['unidad']}"),
                    ("L",f"{L_m_in} m"),("t",f"{t_m_in} m"),("h",f"{h_m_in} m"),("Ladrillo",lad_sel)]:
            doc.add_paragraph(f"{k}: {v}",style='List Bullet')
        doc.add_heading("02. Cargas", 1)
        t2=doc.add_table(rows=1,cols=2); t2.style='Table Grid'
        t2.rows[0].cells[0].text="Variable"; t2.rows[0].cells[1].text="Valor"
        for k,v in [("Pm",f"{Pm_in:.2f} {nc['fuerza']}"),("Pg",f"{Pg_in:.2f} {nc['fuerza']}"),
                    ("Ve",f"{Ve_in:.2f} {nc['fuerza']}"),("Me",f"{Me_in:.2f} {nc['fuerza']}·m"),
                    ("Vu",f"{Vu_in:.2f} {nc['fuerza']}"),("Mu",f"{Mu_in:.2f} {nc['fuerza']}·m")]:
            r=t2.add_row().cells; r[0].text=k; r[1].text=v
        doc.add_heading("03. Esfuerzos Verticales ("+art_ref+")", 1)
        doc.add_paragraph(f"σ_m = {sigma_m_kgcm2:.2f} kg/cm² | Fa = {Fa:.2f} kg/cm²")
        doc.add_paragraph(f"{'CUMPLE ' if ok_sigma else 'NO CUMPLE '}: σ_m {'≤' if ok_sigma else '>'} Fa")
        doc.add_heading("04. Cortante ("+art_ref2+")", 1)
        doc.add_paragraph(f"α={alpha:.3f} | Vm={Vm1:.2f} ton | Vui={Vui_calc:.2f} ton | Mui={Mui_calc:.2f} ton·m")
        doc.add_heading("05. Fisuración ("+art_ref3+")", 1)
        doc.add_paragraph(f"Muro {cond_mur} {'CUMPLE' if ok_fisura else 'NO CUMPLE'}")
        doc.add_heading("06. Refuerzo Horizontal", 1)
        doc.add_paragraph(f"{'SÍ' if necesita_Hz else 'NO'} requiere."
                          + (f" 1Ø{db_hor:.0f}mm@{s_hor:.0f}cm | ρh={rho_h_prov:.5f}" if necesita_Hz else ""))
        doc.add_heading("07-08. Fuerzas Columnas (Tabla 11 E.070)", 1)
        for k,v in [("Nc",Nc),("Vc ext.",f"{Vc_ext:.2f} ton"),("T1 ext.",f"{T1_ext:.2f} ton"),
                    ("C1 ext.",f"{C1_ext:.2f} ton")]:
            doc.add_paragraph(f"{k}: {v}",style='List Bullet')
        doc.add_heading("09. Columnas de Confinamiento", 1)
        doc.add_paragraph(f"Condición: {design_mode}")
        for k,v in [("As req.",f"{As_req:.2f} cm²"),
                    ("Acero",f"{n_bars_f}Ø{db_col:.0f}mm = {As_prov_f:.2f} cm²"),
                    ("b×d",f"{col_b_min:.0f}×{col_d_use:.0f} cm"),
                    ("Estribos",f"Ø{db_est:.0f}mm@{s_col:.1f}cm"),
                    ("Z confinada",f"{Z_col:.1f} cm"),
                    ("Distribución","1@5cm · 4@{:.0f}cm · Rto@25cm".format(s_col))]:
            doc.add_paragraph(f"{k}: {v}",style='List Bullet')
        doc.add_heading("10. Viga Solera", 1)
        for k,v in [("Ts",f"{Ts_sol:.2f} ton"),("Acero",f"{n_sol}Ø{db_col:.0f}mm={As_sol_prov:.2f}cm²"),
                    ("Dimensiones",f"{vs_b_in:.0f}×{vs_h_in:.0f}cm"),
                    ("Estribos",f"Ø{db_est:.0f}mm: 1@5cm, 4@10cm, Rto@25cm")]:
            doc.add_paragraph(f"{k}: {v}",style='List Bullet')
        try:
            buf_2d.seek(0); doc.add_heading("Diagrama 2D",1); doc.add_picture(buf_2d,width=Inches(5.5))
        except: pass
        try:
            buf_mv.seek(0); doc.add_heading("Diagramas M y V",1); doc.add_picture(buf_mv,width=Inches(5.5))
        except: pass
        # Incluir figurado
        doc.add_heading("11. Esquemas de Figurado", 1)
        hook_len = 12 * db_col / 10
        straight_len = (h_m/100) * 100 - 2 * hook_len
        fig_col = draw_longitudinal_bar(h_m, straight_len, hook_len, db_col, bar_name=f"Columna {bar_col_sel}")
        buf_col = io.BytesIO()
        fig_col.savefig(buf_col, format='png', dpi=150, bbox_inches='tight')
        buf_col.seek(0)
        doc.add_picture(buf_col, width=Inches(4.5))
        plt.close(fig_col)
        inside_b = col_b_min - 2*recub_col
        inside_h = col_d_use - 2*recub_col
        hook_len_est = 12 * db_est / 10
        fig_est = draw_stirrup(inside_b, inside_h, hook_len_est, db_est, bar_name=f"Estribo {est_sel}")
        buf_est = io.BytesIO()
        fig_est.savefig(buf_est, format='png', dpi=150, bbox_inches='tight')
        buf_est.seek(0)
        doc.add_picture(buf_est, width=Inches(3.5))
        plt.close(fig_est)
        if necesita_Hz:
            fig_hz = draw_longitudinal_bar(L_m_in*100, L_m_in*100 - 2*hook_len, hook_len, db_hor, bar_name=f"Ref. Hz {bar_hor_sel}")
            buf_hz = io.BytesIO()
            fig_hz.savefig(buf_hz, format='png', dpi=150, bbox_inches='tight')
            buf_hz.seek(0)
            doc.add_picture(buf_hz, width=Inches(4.5))
            plt.close(fig_hz)
        buf_doc=io.BytesIO(); doc.save(buf_doc); buf_doc.seek(0)
        st.download_button("Descargar Memoria DOCX",data=buf_doc,
                           file_name=f"Memoria_{nombre_muro}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_qty:
    st.subheader("Cantidades")
    cem_kg, agua_L, arena_kg, grava_kg = mix_for_fc(fc_in)
    bags = math.ceil(cem_kg * vol_conc_total / bag_kg)
    cq1, cq2 = st.columns(2)
    with cq1:
        st.markdown("#### Concreto")
        qty_table([
            ("N° columnas", n_col_total),
            ("Vol. col.", f"{vol_col_total:.3f} m³"),
            ("Vol. solera", f"{vol_vs_total:.3f} m³"),
            ("Vol. TOTAL", f"{vol_conc_total:.3f} m³"),
            ("Cemento", f"{bags} bultos"),
            ("Arena", f"{arena_kg*vol_conc_total:.0f} kg"),
            ("Grava", f"{grava_kg*vol_conc_total:.0f} kg"),
            ("Agua", f"{agua_L*vol_conc_total:.0f} L"),
        ])
    with cq2:
        st.markdown("#### Acero")
        qty_table([
            ("Acero col.", f"{peso_col_total:.1f} kg"),
            ("Acero solera", f"{peso_sol_total:.1f} kg"),
            ("Acero Hz. muro", f"{peso_hor_total:.1f} kg"),
            ("TOTAL ACERO", f"{total_acero_kg:.1f} kg"),
        ])
        st.markdown("#### Ladrillo")
        qty_table([
            ("Área muro", f"{area_muro_total:.2f} m²"),
            ("Ladrillos", f"{n_ladrillos_total} uds"),
            ("Mortero", f"{vol_mortero_total:.3f} m³"),
        ])
    st.markdown("---\n####  Despiece de Acero")
    despiece_data = [
        {"Marca":"C1","Tipo":"Vertical col.","Diam":f"Ø{db_col:.0f}mm",
         "Cant":n_col_total*n_bars_f,"L_unit":round(long_bar_col,2),"Peso":f"{peso_col_total*0.6:.1f} kg"},
        {"Marca":"E1","Tipo":"Estribo col.","Diam":f"Ø{db_est:.0f}mm",
         "Cant":n_col_total*n_est_col,"L_unit":round(perim_est+0.12,2),"Peso":f"{peso_col_total*0.4:.1f} kg"},
        {"Marca":"S1","Tipo":"Long. solera","Diam":f"Ø{db_col:.0f}mm",
         "Cant":n_sol*math.ceil(L_total_obra/L_m_in),"L_unit":round(L_m_in+0.5,2),"Peso":f"{peso_sol:.1f} kg"},
        {"Marca":"S2","Tipo":"Estribo solera","Diam":f"Ø{db_est:.0f}mm",
         "Cant":n_est_sol,"L_unit":round(2*(vs_b_in+vs_h_in)/100+0.12,2),"Peso":f"{peso_est_sol:.1f} kg"},
    ]
    if necesita_Hz:
        despiece_data.append({"Marca":"H1","Tipo":"Ref. Hz muro","Diam":f"Ø{db_hor:.0f}mm",
                               "Cant":n_hiladas_hor*math.ceil(L_total_obra/L_m_in),
                               "L_unit":round(L_m_in,2),"Peso":f"{peso_hor_total:.1f} kg"})
    df_desp = pd.DataFrame(despiece_data)
    df_desp.columns = ["Marca","Descripción","Diámetro","Cantidad","L.unit.(m)","Peso Total"]
    st.dataframe(df_desp, use_container_width=True, hide_index=True)
    st.metric("⚖ Total Acero", f"{total_acero_kg:.1f} kg")

    # Figurado
    with st.expander("Dibujo de Figurado para Taller", expanded=False):
        st.markdown("A continuación se muestran las formas reales de las barras para facilitar el figurado.")
        hook_len = 12 * db_col / 10
        straight_len = (h_m/100) * 100 - 2 * hook_len
        fig_col = draw_longitudinal_bar(h_m, straight_len, hook_len, db_col, bar_name=f"Columna {bar_col_sel}")
        st.pyplot(fig_col)
        inside_b = col_b_min - 2*recub_col
        inside_h = col_d_use - 2*recub_col
        hook_len_est = 12 * db_est / 10
        fig_est = draw_stirrup(inside_b, inside_h, hook_len_est, db_est, bar_name=f"Estribo {est_sel}")
        st.pyplot(fig_est)
        if necesita_Hz:
            fig_hz = draw_longitudinal_bar(L_m_in*100, L_m_in*100 - 2*hook_len, hook_len, db_hor, bar_name=f"Ref. Hz {bar_hor_sel}")
            st.pyplot(fig_hz)

with tab_apu:
    st.subheader("Presupuesto APU")
    ap1, ap2 = st.columns(2)
    with ap1:
        if "apu_config" in st.session_state:
            apu=st.session_state.apu_config; mon=apu.get("moneda","$")
            p_cem=apu.get("cemento",0); p_ace=apu.get("acero",0); p_are=apu.get("arena",0)
            st.info("Precios APU cargados de configuración global.")
        else:
            mon="$"
            p_cem=st.number_input("Precio bulto cemento",value=35000.0,step=1000.0,key="ac_pcem")
            p_ace=st.number_input("Precio kg acero",value=4500.0,step=100.0,key="ac_pace")
            p_are=st.number_input("Precio m³ arena",value=70000.0,step=5000.0,key="ac_pare")
        p_lad=st.number_input("Precio ladrillo/ud",value=1300.0,step=50.0,key="ac_plad")
        p_m3c=st.number_input("M.O. concreto /m³",value=180000.0,step=10000.0,key="ac_pm3c")
        p_m2m=st.number_input("M.O. pegue /m²",value=25000.0,step=1000.0,key="ac_pm2m")
    with ap2:
        if "apu_config" in st.session_state:
            apu=st.session_state.apu_config
            pct_aui=apu.get("pct_aui",0.30); pct_h=apu.get("pct_herramienta",0.05)
            pct_ut=apu.get("pct_util",0.05); iva=apu.get("iva",0.19)
        else:
            pct_aui=st.number_input("A.U.I. %",0.0,60.0,30.0,1.0,key="ac_aui")/100
            pct_h=st.number_input("Herramienta %",0.0,10.0,5.0,0.5,key="ac_herr")/100
            pct_ut=st.number_input("Utilidad %",0.0,20.0,5.0,0.5,key="ac_util")/100
            iva=st.number_input("IVA %",0.0,30.0,19.0,1.0,key="ac_iva")/100
    cem_kg, agua_L, arena_kg, grava_kg = mix_for_fc(fc_in)
    c_conc=vol_conc_total*(cem_kg*p_cem/bag_kg+arena_kg*p_are/1500+p_m3c)
    c_acero=total_acero_kg*p_ace; c_ladr=n_ladrillos_total*p_lad; c_pega=area_muro_total*p_m2m
    cd=c_conc+c_acero+c_ladr+c_pega
    herr=cd*pct_h; aiu=cd*pct_aui; util=cd*pct_ut; ivat=util*iva
    gran_total=cd+herr+aiu+ivat
    apu_df=pd.DataFrame({
        "Item":["Concreto","Acero","Ladrillo","Pegue M.O.","Herramienta","A.I.U.","IVA Utilidad","TOTAL"],
        f"[{mon}]":[f"{c_conc:,.0f}",f"{c_acero:,.0f}",f"{c_ladr:,.0f}",f"{c_pega:,.0f}",
                   f"{herr:,.0f}",f"{aiu:,.0f}",f"{ivat:,.0f}",f"{gran_total:,.0f}"]
    })
    st.dataframe(apu_df, use_container_width=True, hide_index=True)
    st.metric(f"Gran Total [{mon}]", f"{gran_total:,.0f}")
    out_xl=io.BytesIO()
    with pd.ExcelWriter(out_xl,engine='xlsxwriter') as wr:
        apu_df.to_excel(wr,sheet_name='Presupuesto',index=False)
        df_desp.to_excel(wr,sheet_name='Despiece Acero',index=False)
        pd.DataFrame({"Material":["Concreto(m³)","Acero(kg)","Ladrillos","Mortero(m³)","Cemento(bts)"],
                      "Cantidad":[vol_conc_total,total_acero_kg,n_ladrillos_total,vol_mortero_total,bags]
                     }).to_excel(wr,sheet_name='Cantidades',index=False)
    out_xl.seek(0)
    st.download_button("Descargar Excel",data=out_xl,
                       file_name=f"Presupuesto_{nombre_muro}.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown("---")
st.markdown(f"""
> **Albañilería Confinada — Multi-Norma**  
> Norma activa: `{norma_ac}`  
> f'c = {fc_in:.2f} {nc['unidad']} | fy = {fy_in:.0f} {nc['unidad']}  
> **Referencia:** {norma_ac}  
> ⚠ *Las herramientas son de apoyo para el diseño. Verifique siempre con la norma vigente del país.*
""")