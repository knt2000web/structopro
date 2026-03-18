import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
import plotly.graph_objects as go
from docx import Document

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────



st.title(_t("Diseño de Estructuras en Madera", "Timber Structure Design"))
st.markdown(_t("Cálculo y diseño de elementos de madera estructural (Vigas, Columnas, Uniones) y cuantificación comercial en Pies Madereros. Basado en metodologías de Esfuerzos Admisibles (Marcelo Pardo / NDS / NSR-10).", 
               "Design of structural timber elements (Beams, Columns, Connections) and commercial quantification in Board Feet. Based on Allowable Stress Design (NDS / NSR-10)."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CONFIGURACIÓN GLOBAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙️ Configuración Global", "⚙️ Global Settings"))
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Normativa Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)

st.sidebar.header(_t("🪵 Propiedades Mecánicas (MPa)", "🪵 Mechanical Properties (MPa)"))
grupo_madera_opts = ["Grupo A (Alta Dureza)", "Grupo B (Media)", "Grupo C (Madera Suave)", "Personalizado / Custom"]
grupo_madera = st.sidebar.selectbox(_t("Seleccionar Grupo de Madera:", "Select Wood Group:"), 
                                    grupo_madera_opts,
                                    index=grupo_madera_opts.index(st.session_state.get("m_grupo", grupo_madera_opts[1])),
                                    key="m_grupo")

if grupo_madera == "Grupo A (Alta Dureza)":
    E_min = 9500.0; Fb = 21.0; Fv = 1.5; Fc = 14.5; G_sp = 0.55
elif grupo_madera == "Grupo B (Media)":
    E_min = 7500.0; Fb = 15.0; Fv = 1.2; Fc = 11.0; G_sp = 0.45
elif grupo_madera == "Grupo C (Madera Suave)":
    E_min = 5500.0; Fb = 10.0; Fv = 0.8; Fc = 8.0; G_sp = 0.40
else:
    E_min = 8000.0; Fb = 15.0; Fv = 1.0; Fc = 10.0; G_sp = 0.50

val_E = st.sidebar.number_input(_t("Módulo Elasticidad E_min [MPa]", "Modulus of Elasticity E_min [MPa]"), 1000.0, 20000.0, st.session_state.get("m_E", E_min), 100.0, disabled=(grupo_madera!="Personalizado / Custom"), key="m_E")
val_Fb = st.sidebar.number_input(_t("Esfuerzo Flexión Adm. Fb [MPa]", "Allowable Bending Fb [MPa]"), 1.0, 50.0, st.session_state.get("m_Fb", Fb), 1.0, disabled=(grupo_madera!="Personalizado / Custom"), key="m_Fb")
val_Fv = st.sidebar.number_input(_t("Esfuerzo Cortante Adm. Fv [MPa]", "Allowable Shear Fv [MPa]"), 0.1, 5.0, st.session_state.get("m_Fv", Fv), 0.1, disabled=(grupo_madera!="Personalizado / Custom"), key="m_Fv")
val_Fc = st.sidebar.number_input(_t("Compresión Paralela Fc [MPa]", "Parallel Compression Fc [MPa]"), 1.0, 30.0, st.session_state.get("m_Fc", Fc), 1.0, disabled=(grupo_madera!="Personalizado / Custom"), key="m_Fc")
val_G = st.sidebar.number_input(_t("Gravedad Específica G", "Specific Gravity G"), 0.1, 1.2, st.session_state.get("m_G", G_sp), 0.05, disabled=(grupo_madera!="Personalizado / Custom"), key="m_G")

# ─────────────────────────────────────────────
# T1: CALCULADORA DE PIES MADEREROS
with st.expander(_t("📏 1. Calculadora de Pies² de Madera (Board Feet)", "📏 1. Board Feet Calculator"), expanded=False):
    st.info(_t("Ingresa el espesor y ancho en pulgadas, y el largo de la pieza (en pies o metros).", "Enter thickness/width in inches, and length."))
    
    col1, col2, col3, col4 = st.columns(4)
    with col1: qty_bf = st.number_input(_t("Cantidad de piezas", "Quantity of pieces"), 1, 10000, st.session_state.get("m_qty", 10), key="m_qty")
    with col2: thick_in = st.number_input(_t("Espesor (pulgadas) [in]", "Thickness [in]"), 0.5, 20.0, st.session_state.get("m_thick", 2.0), 0.5, key="m_thick")
    with col3: width_in = st.number_input(_t("Ancho (pulgadas) [in]", "Width [in]"), 1.0, 24.0, st.session_state.get("m_width", 4.0), 0.5, key="m_width")
    with col4:
        unit_len_opts = ["Metros [m]", "Pies [ft]"]
        unit_len = st.radio("Unidad Longitud:", unit_len_opts, 
                            index=unit_len_opts.index(st.session_state.get("m_unit_len", unit_len_opts[0])),
                            horizontal=True, key="m_unit_len")
        length_val = st.number_input(_t("Largo de la pieza", "Piece Length"), 0.5, 50.0, st.session_state.get("m_length", 3.0 if unit_len=="Metros [m]" else 10.0), 0.5, key="m_length")
        
    length_ft = length_val if "ft" in unit_len else (length_val * 3.28084)
    board_feet_per_piece = (thick_in * width_in * length_ft) / 12.0
    total_board_feet = qty_bf * board_feet_per_piece
    
    mon = st.session_state.apu_config["moneda"] if "apu_config" in st.session_state else "US$"
    precio_pt = st.number_input(f"{_t('Precio por Pie Tabular', 'Price per Board Foot')} [{mon}/pt]", value=st.session_state.get("m_price_pt", 2000.0 if mon=="COP$" else 2.50), key="m_price_pt")
    
    c_res, c_plot = st.columns([1,2])
    with c_res:
        st.metric(_t("Volumen por Pieza", "Volume per Piece"), f"{board_feet_per_piece:.2f} pt")
        st.metric(_t("Volumen Total", "Total Volume"), f"{total_board_feet:.2f} pt")
        st.metric(_t("Costo Directo del Lote", "Direct Batch Cost"), f"{mon} {total_board_feet*precio_pt:,.2f}")
    with c_plot:
        # Gráfica ilustrativa 2D
        fig, ax = plt.subplots(figsize=(6, 2))
        ax.add_patch(patches.Rectangle((0, 0), length_ft, thick_in/12.0, edgecolor='saddlebrown', facecolor='peru', lw=2))
        ax.text(length_ft/2, (thick_in/12.0)/2, f"{_t('Largo:', 'Length:')} {length_ft:.1f}'", ha='center', va='center', color='white', fontweight='bold')
        ax.text(-0.5, (thick_in/12.0)/2, f"{_t('Espesor:', 'Thickness:')}\n{thick_in}\"", ha='right', va='center', color='saddlebrown')
        ax.set_title(_t(f"Sección de Madera: {thick_in}\" x {width_in}\" x {length_ft:.1f}'", f"Lumber Section: {thick_in}\" x {width_in}\" x {length_ft:.1f}'"))
        ax.set_xlim(-1, length_ft + 1)
        ax.set_ylim(-1, (thick_in/12.0) + 1)
        ax.axis('off')
        st.pyplot(fig)

# ─────────────────────────────────────────────
# T2: DISEÑO DE VIGAS DE MADERA
with st.expander(_t("🪵 2. Diseño de Vigas de Madera (Flexión y Cortante)", "🪵 2. Timber Beam Design (Flexure and Shear)"), expanded=False):
    v1, v2, v3 = st.columns(3)
    with v1:
        span_L = st.number_input(_t("Luz Viga L [m]", "Beam Span L [m]"), 1.0, 15.0, st.session_state.get("m_v_L", 4.0), 0.1, key="m_v_L")
        w_muerta = st.number_input("WD [kN/m]", 0.0, 50.0, st.session_state.get("m_v_wd", 2.0), 0.5, key="m_v_wd")
        w_viva = st.number_input("WL [kN/m]", 0.0, 50.0, st.session_state.get("m_v_wl", 3.0), 0.5, key="m_v_wl")
        W_total = w_muerta + w_viva # kN/m
    with v2:
        b_beam = st.number_input("Base (b) [mm]", 20.0, 400.0, st.session_state.get("m_v_b", 100.0), 10.0, key="m_v_b")
        h_beam = st.number_input("Altura (h) [mm]", 50.0, 1000.0, st.session_state.get("m_v_h", 250.0), 10.0, key="m_v_h")
        Area_beam = (b_beam * h_beam) / 1e6; Sx_beam = (b_beam * h_beam**2 / 6.0) / 1e9; Ix_beam = (b_beam * h_beam**3 / 12.0) / 1e12
    with v3:
        def_opts = ["L/360", "L/240", "L/180"]
        def_limit = st.selectbox(_t("Límite de Deflexión", "Deflection Limit"), def_opts, 
                                 index=def_opts.index(st.session_state.get("m_v_def", "L/360")),
                                 key="m_v_def")
        Def_max_adm = (span_L * 1000.0) / float(def_limit.split("/")[1])
        
    M_max = (W_total * span_L**2) / 8.0; V_max = (W_total * span_L) / 2.0
    fb_act = (M_max / Sx_beam) / 1000.0; fv_act = (1.5 * V_max / Area_beam) / 1000.0
    Def_act = (5.0 * W_total * (span_L*1000.0)**4) / (384.0 * val_E * (Ix_beam*1e12))
    
    ok_flex = fb_act <= val_Fb; ok_shear = fv_act <= val_Fv; ok_def = Def_act <= Def_max_adm
    
    r1, r2, r3, r4 = st.columns([1,1,1,2])
    r1.metric("f_b Actuante", f"{fb_act:.2f} MPa"); r1.success(f"✅ OK ({val_Fb})") if ok_flex else r1.error("❌ No Aprobado")
    r2.metric("f_v Actuante", f"{fv_act:.2f} MPa"); r2.success(f"✅ OK ({val_Fv})") if ok_shear else r2.error("❌ No Aprobado")
    r3.metric("Δ Actuante", f"{Def_act:.1f} mm"); r3.success(f"✅ OK ({Def_max_adm:.1f})") if ok_def else r3.error("❌ No Aprobado")
    
    with r4:
        fig2, ax2 = plt.subplots(figsize=(6,2))
        ax2.plot([0, span_L], [0, 0], color='brown', lw=max(1, h_beam/20))
        ax2.plot(0, 0, '^', markersize=15, color='gray')
        ax2.plot(span_L, 0, '^', markersize=15, color='gray')
        # Draw distributed load
        num_arr = 10
        x_arr = np.linspace(0, span_L, num_arr)
        for xa in x_arr:
            ax2.arrow(xa, 1.0, 0, -0.6, head_width=0.1, head_length=0.2, fc='red', ec='red')
        ax2.plot([0, span_L], [1.0, 1.0], color='red', lw=2)
        ax2.text(span_L/2, 1.3, f"w = {W_total} kN/m", ha='center', color='red')
        ax2.text(span_L/2, -1.0, f"L = {span_L} m\nSección: {b_beam}x{h_beam} mm", ha='center')
        ax2.set_xlim(-0.5, span_L+0.5)
        ax2.set_ylim(-1.5, 2.0)
        ax2.axis('off')
        st.pyplot(fig2)

# ─────────────────────────────────────────────
# T3: COLUMNAS DE MADERA A COMPRESIÓN PURA
with st.expander(_t("🌲 3. Diseño de Columnas de Madera (Compresión)", "🌲 3. Timber Column Design (Compression)"), expanded=False):
    c1, c2, c3 = st.columns(3)
    with c1:
        P_ax = st.number_input("Carga Axial Actuante P [kN]", 5.0, 500.0, st.session_state.get("m_c_P", 20.0), 5.0, key="m_c_P")
        KL_col = st.number_input(_t("kL Efectiva [m]", "Effective kL [m]"), 0.5, 10.0, st.session_state.get("m_c_kL", 3.0), 0.1, key="m_c_kL")
    with c2:
        b_col = st.number_input("Base (b) [mm]", 50.0, 400.0, st.session_state.get("m_c_b", 150.0), 10.0, key="m_c_b")
        h_col = st.number_input("Altura (h) [mm]", 50.0, 400.0, st.session_state.get("m_c_h", 150.0), 10.0, key="m_c_h")
        d_min = min(b_col, h_col); Area_col = (b_col * h_col) / 1e6
    with c3:
        esbeltez_l = (KL_col * 1000.0) / d_min
        st.markdown(f"**$\lambda = kL/d$:** {esbeltez_l:.2f}")
    
    if esbeltez_l > 50:
        st.error(_t("❌ Columna muy esbelta (λ > 50). ¡Aumentar sección!", "❌ Column too slender (λ > 50)!"))
    else:
        F_cE = (0.822 * val_E) / (esbeltez_l**2)
        ratio_alpha = F_cE / val_Fc; c_factor = 0.8
        termino_rad = (1 + ratio_alpha)/(2*c_factor)
        C_p = termino_rad - math.sqrt((termino_rad**2) - (ratio_alpha/c_factor))
        F_c_prime = val_Fc * C_p
        P_admisible = (F_c_prime * 1000.0 * Area_col) # kN
        
        tc1, tc2, tc3 = st.columns([1,1,2])
        tc1.metric("P Actuante", f"{P_ax:.1f} kN")
        tc2.metric("P Admisible", f"{P_admisible:.1f} kN")
        tc2.success("✅ OK") if P_ax <= P_admisible else tc2.error("❌ No Aprobado")
        
        with tc3:
            fig3, ax3 = plt.subplots(figsize=(2,4))
            ax3.plot([0, 0], [0, KL_col], color='saddlebrown', lw=max(5, b_col/20))
            ax3.arrow(0, KL_col+0.5, 0, -0.3, head_width=0.3, head_length=0.1, fc='blue', ec='blue')
            ax3.text(0, KL_col+0.6, f"P={P_ax}kN", ha='center', color='blue')
            
            # Buckling curve
            y_curve = np.linspace(0, KL_col, 50)
            x_curve = 0.5 * np.sin(np.pi * y_curve / KL_col)
            ax3.plot(x_curve, y_curve, 'r--', lw=1, alpha=0.6)
            
            ax3.set_xlim(-1, 1)
            ax3.set_ylim(-0.5, KL_col+1)
            ax3.axis('off')
            st.pyplot(fig3)

# ─────────────────────────────────────────────
# T4: UNIONES CON CLAVOS
with st.expander(_t("🔨 4. Resistencia de Uniones con Clavos (Corte Lateral)", "🔨 4. Nail Connection (Lateral Shear)"), expanded=False):
    uc1, uc2, uc3 = st.columns([1,1,2])
    with uc1:
        clavo_opts = [2.5, 3.0, 3.4, 3.8, 4.2, 5.0, 6.0]
        diam_clavo = st.selectbox("Calibre Clavo D [mm]", clavo_opts, 
                                   index=clavo_opts.index(st.session_state.get("m_u_diam", 3.4)),
                                   key="m_u_diam")
    with uc2:
        penetracion_p = st.number_input("Penetración (p) [mm]", 10.0, 150.0, st.session_state.get("m_u_p", 40.0), 5.0, key="m_u_p")
    
    Z_lat = 16.6 * (val_G**1.8) * (diam_clavo**1.5)
    D_req_full = 10.0 * diam_clavo; D_req_min = 6.0 * diam_clavo
    
    if penetracion_p < D_req_min:
        st.error(_t("❌ Penetración < 6D.", "❌ Penetration < 6D."))
    else:
        Cd_factor = min(1.0, penetracion_p / D_req_full)
        Z_adm = (Z_lat * Cd_factor) / 10.0 # kgf approx
        
        with uc1: st.metric("Corte Admisible (Z')", f"{Z_adm:.2f} kgf/clavo")
        
        with uc3:
            fig4, ax4 = plt.subplots(figsize=(6,2))
            ax4.add_patch(patches.Rectangle((0, 0), 2, 1, facecolor='peru', edgecolor='brown'))
            ax4.add_patch(patches.Rectangle((2, 0), 2, 1, facecolor='burlywood', edgecolor='brown'))
            # Nail
            ax4.plot([1.8, 2.0+(penetracion_p/40.0)], [0.5, 0.5], 'k-', lw=3) # scaled representation
            ax4.text(2.0, 0.7, f"p={penetracion_p}mm", ha='center')
            ax4.arrow(1.0, 0.5, -0.4, 0, head_width=0.1, fc='blue', ec='blue')
            ax4.arrow(3.0, 0.5, 0.4, 0, head_width=0.1, fc='blue', ec='blue')
            ax4.text(0.5, 0.2, "Z'", color='blue')
            ax4.set_xlim(0, 4); ax4.set_ylim(-0.5, 1.5); ax4.axis('off')
            st.pyplot(fig4)

st.markdown("---")
st.subheader("💾 Exportación Integral (Viga/Columna)")
tab_3d, tab_dxf, tab_doc, tab_apu = st.tabs(["🧊 Visualización 3D", "📐 Planos DXF", "📄 Memoria DOCX", "💰 Presupuesto APU XLSX"])

with tab_3d:
    col_3d_v, col_3d_c = st.columns(2)
    with col_3d_v:
        st.write("#### Viga de Madera")
        fig3d_v = go.Figure()
        X_v = span_L
        Y_v = b_beam / 1000.0
        Z_v = h_beam / 1000.0
        x_wv = [0, X_v, X_v, 0, 0, X_v, X_v, 0]
        y_wv = [-Y_v/2, -Y_v/2, Y_v/2, Y_v/2, -Y_v/2, -Y_v/2, Y_v/2, Y_v/2]
        z_wv = [-Z_v/2, -Z_v/2, -Z_v/2, -Z_v/2, Z_v/2, Z_v/2, Z_v/2, Z_v/2]
        fig3d_v.add_trace(go.Mesh3d(x=x_wv, y=y_wv, z=z_wv, alphahull=0, opacity=0.9, color='peru', name='Viga'))
        fig3d_v.update_layout(scene=dict(aspectmode='data', xaxis_title='L (m)', yaxis_title='b (m)', zaxis_title='h (m)'), 
                              margin=dict(l=0, r=0, b=0, t=0), height=350, showlegend=False, dragmode='turntable')
        st.plotly_chart(fig3d_v, use_container_width=True)
        
    with col_3d_c:
        st.write("#### Columna de Madera")
        fig3d_c = go.Figure()
        Z_c = KL_col
        X_c = b_col / 1000.0
        Y_c = h_col / 1000.0
        x_wc = [-X_c/2, X_c/2, X_c/2, -X_c/2, -X_c/2, X_c/2, X_c/2, -X_c/2]
        y_wc = [-Y_c/2, -Y_c/2, Y_c/2, Y_c/2, -Y_c/2, -Y_c/2, Y_c/2, Y_c/2]
        z_wc = [0, 0, 0, 0, Z_c, Z_c, Z_c, Z_c]
        fig3d_c.add_trace(go.Mesh3d(x=x_wc, y=y_wc, z=z_wc, alphahull=0, opacity=0.9, color='saddlebrown', name='Columna'))
        fig3d_c.update_layout(scene=dict(aspectmode='data', xaxis_title='b (m)', yaxis_title='h (m)', zaxis_title='L (m)'), 
                              margin=dict(l=0, r=0, b=0, t=0), height=350, showlegend=False, dragmode='turntable')
        st.plotly_chart(fig3d_c, use_container_width=True)

with tab_dxf:
    col_dxf1, col_dxf2 = st.columns(2)
    doc_v = ezdxf.new('R2010'); msp_v = doc_v.modelspace()
    msp_v.add_lwpolyline([(0,0), (b_beam,0), (b_beam,h_beam), (0,h_beam), (0,0)], dxfattribs={'layer': 'MADERA_VIGA', 'color': 3})
    out_v = io.StringIO(); doc_v.write(out_v)
    col_dxf1.download_button("Descargar Viga.dxf", data=out_v.getvalue(), file_name=f"Viga_Madera_{b_beam}x{h_beam}.dxf", mime="application/dxf")

    doc_c = ezdxf.new('R2010'); msp_c = doc_c.modelspace()
    msp_c.add_lwpolyline([(0,0), (b_col,0), (b_col,h_col), (0,h_col), (0,0)], dxfattribs={'layer': 'MADERA_COLUMNA', 'color': 4})
    out_c = io.StringIO(); doc_c.write(out_c)
    col_dxf2.download_button("Descargar Columna.dxf", data=out_c.getvalue(), file_name=f"Columna_Madera_{b_col}x{h_col}.dxf", mime="application/dxf")

with tab_doc:
    doc_m = Document()
    doc_m.add_heading(f"Memoria de Estructuras en Madera - {grupo_madera}", 0)
    doc_m.add_heading("1. Diseño de Viga", level=1)
    doc_m.add_paragraph(f"Sección: {b_beam} mm x {h_beam} mm. Luz: {span_L} m.")
    doc_m.add_paragraph(f"Flexión fb={fb_act:.2f} MPa vs Fb={val_Fb} MPa. -> {'CUMPLE' if ok_flex else 'FALLA'}")
    doc_m.add_heading("2. Diseño de Columna", level=1)
    if esbeltez_l <= 50:
        doc_m.add_paragraph(f"P_adm: {P_admisible:.1f} kN (Solictado: {P_ax} kN).")
    f_doc = io.BytesIO(); doc_m.save(f_doc); f_doc.seek(0)
    st.download_button("Descargar Memoria DOCX", data=f_doc, file_name="Memoria_Madera.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_apu:
    if "apu_config" in st.session_state:
        apu = st.session_state.apu_config
        mon = apu["moneda"]
        pt_viga = (b_beam/25.4) * (h_beam/25.4) * (span_L * 3.28084) / 12.0
        pt_col  = (b_col/25.4) * (h_col/25.4) * (KL_col * 3.28084) / 12.0
        output_excel = io.BytesIO()
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            df_export = pd.DataFrame({
                "Item": ["Madera Viga", "Madera Columna", "AIU"],
                "Cantidad (Pies²)": [pt_viga, pt_col, 1],
                "Unidad": [precio_pt, precio_pt, (pt_viga+pt_col)*precio_pt*apu.get("pct_aui",0.30)]
            })
            df_export["Subtotal"] = df_export["Cantidad (Pies²)"] * df_export["Unidad"]
            df_export.to_excel(writer, index=False, sheet_name='APU Madera')
        output_excel.seek(0)
        st.download_button("📥 Descargar APU en Excel", data=output_excel, file_name="APU_Madera.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else: st.warning("Carga APU en la pestaña principal de Mercado.")
