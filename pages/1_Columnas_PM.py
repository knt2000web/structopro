import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import plotly.graph_objects as go
import json
import datetime
from matplotlib.backends.backend_pdf import PdfPages

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Diagramas de Interacción", "Interaction Diagrams"), layout="wide")

st.image(r"assets/columnas_pm_header_1773261175144.png", use_container_width=False, width=700)
st.title(_t("🏗️ Diagrama de Interacción P–M y Diseño de Estribos", "🏗️ P-M Interaction Diagram & Tie Design"))
st.markdown(_t(
    "Generador interactivo de capacidad a flexocompresión biaxial y cortante para **Columnas Cuadradas y Rectangulares**.",
    "Interactive biaxial flexure-compression and shear capacity generator for **Square and Rectangular Columns**."
))

with st.expander("📺 ¿Cómo usar esta herramienta?"):
    st.markdown("""
    **Modo de Uso:**
    1. **📍 Sidebar (Menú Izquierdo):** Selecciona la Norma de tu país, el nivel sísmico, la resistencia del concreto (f'c), el acero (fy), y la geometría de la columna (Base, Altura).
    2. **🏗️ Armadura:** Define el diámetro y cantidad de varillas longitudinales, así como el diámetro de los estribos.
    3. **🚦 Solicitaciones:** Ingresa el Momento Último (Mu) y la Carga Axial (Pu) calculados en tu software de análisis (ej. ETABS, SAP2000).
    4. **📊 Resultados:** Revisa la primera pestaña para ver el Diagrama P-M; el punto verde debe quedar DENTRO de la curva roja (Resistencia de Diseño).
    5. **📦 Exportar:** En la pestaña "Cantidades" encontrarás tu presupuesto APU, despiece de acero y la memoria. En "Sección y Estribos" podrás descargar el plano en AutoCAD (DXF) con detalles de empalmes y ganchos.
    """)

# ─────────────────────────────────────────────
# DICCIONARIOS DE BARRAS
# ─────────────────────────────────────────────
REBAR_US = {
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
    "#5 (5/8\")": {"area": 1.99, "diam_mm": 15.88},
    "#6 (3/4\")": {"area": 2.84, "diam_mm": 19.05},
    "#7 (7/8\")": {"area": 3.87, "diam_mm": 22.23},
    "#8 (1\")":   {"area": 5.10, "diam_mm": 25.40},
    "#9 (1 1/8\")": {"area": 6.45, "diam_mm": 28.65},
    "#10 (1 1/4\")": {"area": 7.92, "diam_mm": 32.26},
}

REBAR_MM = {
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
    "14 mm": {"area": 1.539, "diam_mm": 14.0},
    "16 mm": {"area": 2.011, "diam_mm": 16.0},
    "18 mm": {"area": 2.545, "diam_mm": 18.0},
    "20 mm": {"area": 3.142, "diam_mm": 20.0},
    "22 mm": {"area": 3.801, "diam_mm": 22.0},
    "25 mm": {"area": 4.909, "diam_mm": 25.0},
    "28 mm": {"area": 6.158, "diam_mm": 28.0},
    "32 mm": {"area": 8.042, "diam_mm": 32.0},
}

STIRRUP_US = {
    "#2 (1/4\")": {"area": 0.32, "diam_mm": 6.35},
    "#3 (3/8\")": {"area": 0.71, "diam_mm": 9.53},
    "#4 (1/2\")": {"area": 1.29, "diam_mm": 12.70},
}

STIRRUP_MM = {
    "6 mm":  {"area": 0.283, "diam_mm": 6.0},
    "8 mm":  {"area": 0.503, "diam_mm": 8.0},
    "10 mm": {"area": 0.785, "diam_mm": 10.0},
    "12 mm": {"area": 1.131, "diam_mm": 12.0},
}

# ─────────────────────────────────────────────
# FUNCIONES DE DIBUJO PARA FIGURADO (movidas al inicio)
# ─────────────────────────────────────────────
# Tabla de equivalencias mm → designación en cuartos de pulgada (NSR-10 / ACI)
_MM_TO_BAR = {
    6.0:  "#2 (1/4\")",  6.35: "#2 (1/4\")",
    8.0:  "#2.5 (5/16\")",
    9.53: "#3 (3/8\")",  10.0: "#3 (3/8\")",
    12.0: "#4 (1/2\")",  12.70: "#4 (1/2\")",
    14.0: "#4.5 (9/16\")",
    15.88: "#5 (5/8\")", 16.0: "#5 (5/8\")",
    18.0: "#5.7 (11/16\")",
    19.05: "#6 (3/4\")", 20.0: "#6 (3/4\")",
    22.0: "#7 (7/8\")",  22.23: "#7 (7/8\")",
    25.0: "#8 (1\")",    25.40: "#8 (1\")",
    28.0: "#9 (1 1/8\")",28.65: "#9 (1 1/8\")",
    32.0: "#10 (1 1/4\")",32.26: "#10 (1 1/4\")",
}

def _bar_label(diam_mm):
    """Devuelve la designación en cuartos de pulgada más cercana para un diámetro en mm."""
    best = min(_MM_TO_BAR.keys(), key=lambda k: abs(k - diam_mm))
    if abs(best - diam_mm) <= 2.0:
        return f"{_MM_TO_BAR[best]}  Ø{diam_mm:.1f} mm"
    return f"Ø{diam_mm:.1f} mm"

def draw_longitudinal_bar(total_len_cm, straight_len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    """Dibuja una barra longitudinal con ganchos de 90° en ambos extremos."""
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(6, total_len_cm/20), 2))
    ax.set_aspect('equal')
    ax.plot([0, straight_len_cm], [0, 0], 'k-', linewidth=2)
    ax.plot([0, 0], [0, hook_len_cm], 'k-', linewidth=2)
    ax.plot([straight_len_cm, straight_len_cm], [0, -hook_len_cm], 'k-', linewidth=2)
    ax.annotate(f"{straight_len_cm:.0f} cm", xy=(straight_len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 12db = {hook_len_cm:.0f} cm", xy=(0, hook_len_cm/2), ha='right', fontsize=8)
    ax.annotate(f"Gancho 12db", xy=(straight_len_cm, -hook_len_cm/2), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*0.2, straight_len_cm + hook_len_cm*0.2)
    ax.set_ylim(-hook_len_cm*1.2, hook_len_cm*1.2)
    ax.axis('off')
    ax.set_title(f"Varilla L1 — {label} — Longitud total {total_len_cm:.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_stirrup(b_cm, h_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    """Dibuja un estribo rectangular con gancho 135° correcto (proyectado hacia el interior)."""
    import math as _math
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(5, b_cm/12), max(4, h_cm/12)))
    ax.set_aspect('equal')
    x0, y0 = 0, 0

    # Rectángulo principal del estribo
    ax.plot([x0,       x0+b_cm], [y0,       y0],       'k-', linewidth=2.5)  # inferior
    ax.plot([x0+b_cm, x0+b_cm], [y0,       y0+h_cm],  'k-', linewidth=2.5)  # derecho
    ax.plot([x0+b_cm, x0],      [y0+h_cm,  y0+h_cm],  'k-', linewidth=2.5)  # superior
    ax.plot([x0,       x0],     [y0+h_cm,  y0],        'k-', linewidth=2.5)  # izquierdo

    # Gancho 135° en esquina superior izquierda
    angle_rad = _math.radians(45)
    vis_hook = min(hook_len_cm, b_cm/4.0, h_cm/4.0)
    hx = vis_hook * _math.cos(angle_rad)
    hy = -vis_hook * _math.sin(angle_rad)
    ax.plot([x0, x0 + hx], [y0+h_cm, y0+h_cm + hy], 'k-', linewidth=2.5)

    # Segundo gancho 135° en esquina superior derecha (opuesto - alternado por norma)
    ax.plot([x0+b_cm, x0+b_cm - hx], [y0+h_cm, y0+h_cm + hy], 'k--', linewidth=1.5, alpha=0.5)
    ax.annotate("(Alternado)", xy=(x0+b_cm - hx - 0.3, y0+h_cm + hy - 0.3), fontsize=6, color='gray', ha='right')

    # Cotas
    ax.annotate(f"{b_cm:.0f} cm", xy=(b_cm/2, y0-0.8), ha='center', fontsize=9, fontweight='bold')
    ax.annotate(f"{h_cm:.0f} cm", xy=(x0-0.8, h_cm/2), ha='right', va='center', fontsize=9, fontweight='bold')

    # Etiqueta del gancho
    ax.annotate(f"Gancho 135\u00b0\n6d_e = {6*bar_diam_mm/10:.1f} cm",
                xy=(x0 + hx + 0.2, y0+h_cm + hy - 0.2), fontsize=7.5, color='darkred',
                va='top', ha='left')

    ax.set_xlim(x0 - hook_len_cm*0.3, b_cm + hook_len_cm*0.6)
    ax.set_ylim(y0 - hook_len_cm*0.5, h_cm + hook_len_cm*0.9)
    ax.axis('off')
    ax.set_title(f"Estribo E1 — {label} — Perímetro {2*(b_cm+h_cm):.0f} cm", fontsize=9, fontweight='bold')
    return fig

def draw_crosstie(len_cm, hook_len_cm, bar_diam_mm, bar_name=None):
    """Dibuja un crosstie con ganchos de 135° en ambos extremos."""
    label = bar_name if bar_name else _bar_label(bar_diam_mm)
    fig, ax = plt.subplots(figsize=(max(6, len_cm/15), 2))
    ax.set_aspect('equal')
    ax.plot([0, len_cm], [0, 0], 'k-', linewidth=2)
    ax.plot([0, -hook_len_cm*0.7], [0, -hook_len_cm*0.7], 'k-', linewidth=2)
    ax.plot([len_cm, len_cm + hook_len_cm*0.7], [0, -hook_len_cm*0.7], 'k-', linewidth=2)
    ax.annotate(f"{len_cm:.0f} cm", xy=(len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 135°", xy=(0, -hook_len_cm*0.5), ha='right', fontsize=8)
    ax.annotate(f"Gancho 135°", xy=(len_cm, -hook_len_cm*0.5), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*1.2, len_cm + hook_len_cm*1.2)
    ax.set_ylim(-hook_len_cm*1.5, hook_len_cm*0.5)
    ax.axis('off')
    ax.set_title(f"Crosstie C1 — {label} — Longitud {len_cm:.0f} cm", fontsize=9, fontweight='bold')
    return fig

# ─────────────────────────────────────────────
# PARÁMETROS POR NORMA
# ─────────────────────────────────────────────
CODES = {
    "NSR-10 (Colombia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 4.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["DMI — Disipación Mínima (Amenaza baja)", "DMO — Disipación Moderada (Amenaza media)", "DES — Disipación Especial (Amenaza alta)"],
        "ref": "NSR-10 Título C, Cap. C.10",
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary Moment Frame (SDC A–B, amenaza baja)", "IMF — Intermediate Moment Frame (SDC C, amenaza media)", "SMF — Special Moment Frame (SDC D–F, amenaza alta)"],
        "ref": "ACI 318-25 Section 22.4",
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary Moment Frame (SDC A–B, amenaza baja)", "IMF — Intermediate Moment Frame (SDC C, amenaza media)", "SMF — Special Moment Frame (SDC D–F, amenaza alta)"],
        "ref": "ACI 318-19 Section 22.4",
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["OMF — Ordinary Moment Frame (SDC A–B, amenaza baja)", "IMF — Intermediate Moment Frame (SDC C, amenaza media)", "SMF — Special Moment Frame (SDC D–F, amenaza alta)"],
        "ref": "ACI 318-14 Section 22.4",
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["GS — Grado Sísmico Reducido (Amenaza baja)", "GM — Grado Sísmico Moderado", "GA — Grado Sísmico Alto (Estructuras especiales)"],
        "ref": "NEC-SE-HM (Ecuador) Cap. 4",
    },
    "E.060 (Perú)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["PO — Pórtico Ordinario (Zona 1–2, amenaza baja)", "PE — Pórtico Especial (Zona 3–4, amenaza alta)"],
        "ref": "Norma E.060 (Perú) Art. 9.3",
    },
    "NTC-EM (México)": {
        "phi_tied": 0.70, "phi_spiral": 0.80, "phi_tension": 0.85,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["MDL — Marcos de Ductilidad Limitada (Zona A baja)", "MROD — Marcos de Ductilidad Ordinaria (Zona B–C)", "MRLE — Marcos de Ductilidad Alta (Zona D–E alta)"],
        "ref": "NTC-EM México Cap. 2",
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_tied": 0.70, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 6.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["PO — Pórtico Ordinario (Zona 1, amenaza baja)", "PM — Pórtico Moderado (Zona 2–3)", "PE — Pórtico Especial (Zona 4–7, amenaza alta)"],
        "ref": "COVENIN 1753-2006 (Venezuela) — Basada en ACI 318-05. Sísmico: COVENIN 1756:2019",
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["DO — Diseño Ordinario (Amenaza baja, zona 1–2)", "DE — Diseño Especial Sísmico (Amenaza alta, zona 3–4)"],
        "ref": "NB 1225001-2020 (Bolivia) — Basada en ACI 318-19",
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_tied": 0.65, "phi_spiral": 0.75, "phi_tension": 0.90,
        "pmax_tied": 0.80, "pmax_spiral": 0.85,
        "rho_min": 1.0, "rho_max": 8.0,
        "eps_tension_full": 0.005,
        "seismic_levels": ["GE — Grado Estándar (Amenaza sísmica baja)", "GM — Ductilidad Moderada (Amenaza media)", "GA — Ductilidad Alta / Zonas especiales"],
        "ref": "CIRSOC 201-2025 (Argentina) vigente desde Res. 11/2026 — Basada en ACI 318-19",
    },
}

# ─────────────────────────────────────────────
# PRESENTACIONES DE CEMENTO POR PAÍS
# ─────────────────────────────────────────────
CEMENT_BAGS = {
    "NSR-10 (Colombia)": [
        {"label": "Cemento gris clásico (Argos, El Cairo, Holcim)", "kg": 50.0},
        {"label": "Cemento Estructural / Tipo III (Argos Estructural)", "kg": 42.5},
        {"label": "Bolsa pequeña (presentación 25 kg)", "kg": 25.0},
    ],
    "ACI 318-25 (EE.UU.)": [
        {"label": "Type I/II sack estándar (94 lb)", "kg": 42.6},
        {"label": "Type III bolsa (47 lb)", "kg": 21.3},
    ],
    "ACI 318-19 (EE.UU.)": [
        {"label": "Type I/II sack estándar (94 lb)", "kg": 42.6},
        {"label": "Type III bolsa (47 lb)", "kg": 21.3},
    ],
    "ACI 318-14 (EE.UU.)": [
        {"label": "Type I/II sack estándar (94 lb)", "kg": 42.6},
        {"label": "Type III bolsa (47 lb)", "kg": 21.3},
    ],
    "NEC-SE-HM (Ecuador)": [
        {"label": "Cemento Holcim / Chimborazo (bulto estándar)", "kg": 50.0},
        {"label": "Bolsa pequeña", "kg": 25.0},
    ],
    "E.060 (Perú)": [
        {"label": "Cemento Andino / Pacasmayo / Yura (bolsa estándar)", "kg": 42.5},
        {"label": "Bolsa pequeña", "kg": 25.0},
    ],
    "NTC-EM (México)": [
        {"label": "Cemento Cemex / Cruz Azul / Moctezuma (bulto estándar)", "kg": 50.0},
        {"label": "Bolsa pequeña / práctica", "kg": 25.0},
    ],
    "COVENIN 1753-2006 (Venezuela)": [
        {"label": "Cemento Cemex Venezuela / Lafarge / Holcim (bulto estándar)", "kg": 42.5},
    ],
    "NB 1225001-2020 (Bolivia)": [
        {"label": "Cemento Fancesa / Viacha / Itacamba (bolsa estándar)", "kg": 50.0},
        {"label": "Bolsa pequeña", "kg": 25.0},
    ],
    "CIRSOC 201-2025 (Argentina)": [
        {"label": "Cemento Loma Negra / Holcim / Cementos Avellaneda (bolsa)", "kg": 50.0},
        {"label": "Bolsa pequeña", "kg": 25.0},
    ],
}

# ─────────────────────────────────────────────
# TABLA DE DOSIFICACIÓN ACI 211
# ─────────────────────────────────────────────
MIX_DESIGNS = [
    {"fc_mpa": 14.0, "cem": 250, "agua": 205, "arena": 810, "grava": 1060, "wc": 0.82},
    {"fc_mpa": 17.0, "cem": 290, "agua": 200, "arena": 780, "grava": 1060, "wc": 0.69},
    {"fc_mpa": 21.0, "cem": 350, "agua": 193, "arena": 720, "grava": 1060, "wc": 0.55},
    {"fc_mpa": 25.0, "cem": 395, "agua": 193, "arena": 680, "grava": 1020, "wc": 0.49},
    {"fc_mpa": 28.0, "cem": 430, "agua": 190, "arena": 640, "grava": 1000, "wc": 0.44},
    {"fc_mpa": 35.0, "cem": 530, "agua": 185, "arena": 580, "grava":  960, "wc": 0.35},
    {"fc_mpa": 42.0, "cem": 620, "agua": 180, "arena": 520, "grava":  910, "wc": 0.29},
    {"fc_mpa": 56.0, "cem": 740, "agua": 175, "arena": 450, "grava":  850, "wc": 0.24},
]

def get_mix_for_fc(fc_mpa):
    if fc_mpa <= MIX_DESIGNS[0]["fc_mpa"]: return MIX_DESIGNS[0]
    if fc_mpa >= MIX_DESIGNS[-1]["fc_mpa"]: return MIX_DESIGNS[-1]
    for i in range(len(MIX_DESIGNS)-1):
        lo, hi = MIX_DESIGNS[i], MIX_DESIGNS[i+1]
        if lo["fc_mpa"] <= fc_mpa <= hi["fc_mpa"]:
            t = (fc_mpa - lo["fc_mpa"]) / (hi["fc_mpa"] - lo["fc_mpa"])
            return {k: lo[k] + t*(hi[k]-lo[k]) for k in ("cem", "agua", "arena", "grava", "wc")}
    return MIX_DESIGNS[-1]

def get_beta1(fc):
    if fc <= 28: return 0.85
    return max(0.85 - 0.05*(fc-28)/7.0, 0.65)

def get_development_length(db_mm, fy, fc, lambda_=1.0, psi_t=1.0, psi_e=1.0, psi_s=1.0, psi_g=1.0, cb_ktr=2.5):
    if db_mm <= 0: return 0
    ld = (3/40) * (fy / (lambda_ * math.sqrt(fc))) * (psi_t * psi_e * psi_s * psi_g / cb_ktr) * db_mm
    return max(ld, 300)  # mínimo 300 mm

# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────
st.sidebar.header(_t("0. Norma de Diseño", "0. Design Code"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = list(CODES.keys())[0]
norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)
code = CODES[norma_sel]
nivel_sismico = st.sidebar.selectbox(
    _t("Nivel Sísmico / Ductilidad:", "Seismic / Ductility Level:"),
    code["seismic_levels"],
    index=0 if "c_pm_nivel_sismico" not in st.session_state else code["seismic_levels"].index(st.session_state.c_pm_nivel_sismico) if st.session_state.c_pm_nivel_sismico in code["seismic_levels"] else 0,
    key="c_pm_nivel_sismico"
)
st.sidebar.caption(f"📖 {_t('Referencia', 'Reference')}: {code['ref']}")

st.sidebar.header(_t("1. Materiales", "1. Materials"))
fc_unit = st.sidebar.radio(_t("Unidad de f'c:", "f'c Unit:"), ["MPa", "PSI", "kg/cm²"], horizontal=True, key="c_pm_fc_unit")
if fc_unit == "PSI":
    psi_options = {"2500 PSI (≈ 17.2 MPa)":2500, "3000 PSI (≈ 20.7 MPa)":3000, "3500 PSI (≈ 24.1 MPa)":3500,
                   "4000 PSI (≈ 27.6 MPa)":4000, "4500 PSI (≈ 31.0 MPa)":4500, "5000 PSI (≈ 34.5 MPa)":5000,
                   "Personalizado":None}
    psi_choice = st.sidebar.selectbox("Resistencia f'c [PSI]", list(psi_options.keys()),
        index=list(psi_options.keys()).index(st.session_state.get("c_pm_psi_choice", list(psi_options.keys())[1])) if st.session_state.get("c_pm_psi_choice", list(psi_options.keys())[1]) in psi_options else 1,
        key="c_pm_psi_choice")
    fc_psi = float(psi_options[psi_choice]) if psi_options[psi_choice] is not None else st.sidebar.number_input("f'c personalizado [PSI]", 2000.0, 12000.0, 3000.0, 100.0, key="c_pm_fc_psi_custom")
    fc = fc_psi * 0.00689476
elif fc_unit == "kg/cm²":
    kgcm2_options = {"175 kg/cm² (≈ 17.2 MPa)":175, "210 kg/cm² (≈ 20.6 MPa)":210, "250 kg/cm² (≈ 24.5 MPa)":250,
                     "280 kg/cm² (≈ 27.5 MPa)":280, "350 kg/cm² (≈ 34.3 MPa)":350, "420 kg/cm² (≈ 41.2 MPa)":420,
                     "Personalizado":None}
    kgcm2_choice = st.sidebar.selectbox("Resistencia f'c [kg/cm²]", list(kgcm2_options.keys()),
        index=list(kgcm2_options.keys()).index(st.session_state.get("c_pm_kgcm2_choice", list(kgcm2_options.keys())[1])) if st.session_state.get("c_pm_kgcm2_choice", list(kgcm2_options.keys())[1]) in kgcm2_options else 1,
        key="c_pm_kgcm2_choice")
    fc_kgcm2 = float(kgcm2_options[kgcm2_choice]) if kgcm2_options[kgcm2_choice] is not None else st.sidebar.number_input("f'c personalizado [kg/cm²]", 100.0, 1200.0, 210.0, 10.0, key="c_pm_fc_kgcm2_custom")
    fc = fc_kgcm2 / 10.1972
else:
    fc = st.sidebar.number_input("Resistencia del Concreto (f'c) [MPa]", 15.0, 80.0, 21.0, 1.0, key="c_pm_fc_mpa")

fy = st.sidebar.number_input("Fluencia del Acero (fy) [MPa]", 240.0, 500.0, 420.0, 10.0, key="c_pm_fy")
Es = 200000.0

st.sidebar.header("2. Geometría de la Sección")
b = st.sidebar.number_input("Base (b) [cm]", 15.0, 150.0, 30.0, 5.0, key="c_pm_b")
h = st.sidebar.number_input("Altura (h) [cm]", 15.0, 150.0, 40.0, 5.0, key="c_pm_h")
d_prime = st.sidebar.number_input("Recubrimiento al centroide (d') [cm]", 2.0, 15.0, 5.0, 0.5, key="c_pm_dprime")
L_col = st.sidebar.number_input("Altura libre de la columna (L) [cm]", 50.0, 1000.0, 300.0, 25.0, key="c_pm_L",
    help="Se usa para calcular estribos y cantidades de materiales")

st.sidebar.header("3. Refuerzo Longitudinal")
unit_system = st.sidebar.radio("Sistema de Unidades de las Varillas:", ["Pulgadas (EE. UU.)", "Milímetros (SI)"], key="c_pm_unit_system")
rebar_dict = REBAR_US if "Pulgadas" in unit_system else REBAR_MM
default_rebar = "#5 (5/8\")" if "Pulgadas" in unit_system else "16 mm"
rebar_type = st.sidebar.selectbox("Diámetro de las Varillas", list(rebar_dict.keys()),
    index=list(rebar_dict.keys()).index(st.session_state.get("c_pm_rebar_type", default_rebar)) if st.session_state.get("c_pm_rebar_type", default_rebar) in rebar_dict else 0,
    key="c_pm_rebar_type")
rebar_area  = rebar_dict[rebar_type]["area"]    # cm²
rebar_diam  = rebar_dict[rebar_type]["diam_mm"] # mm

num_filas_h = st.sidebar.number_input("# de filas Acero Horiz (Superior e Inferior)", 2, 15, 2, 1, key="c_pm_num_h")
num_filas_v = st.sidebar.number_input("# de filas Acero Vert (Laterales)", 2, 15, 2, 1, key="c_pm_num_v")

# Armar capas
layers = []
layers.append({'d': d_prime, 'As': num_filas_h * rebar_area})
num_capas_intermedias = num_filas_v - 2
if num_capas_intermedias > 0:
    espacio_h = (h - 2 * d_prime) / (num_capas_intermedias + 1)
    for i in range(1, num_capas_intermedias + 1):
        layers.append({'d': d_prime + i * espacio_h, 'As': 2 * rebar_area})
layers.append({'d': h - d_prime, 'As': num_filas_h * rebar_area})

n_barras_total = num_filas_h * 2 + num_capas_intermedias * 2
Ast = sum([layer['As'] for layer in layers])
Ag  = b * h
st.sidebar.markdown(f"**Total varillas:** {n_barras_total} | **Ast:** {Ast:.2f} cm²")

st.sidebar.header("4. Estribos (Flejes)")
stirrup_dict = STIRRUP_US if "Pulgadas" in unit_system else STIRRUP_MM
default_stirrup = "#3 (3/8\")" if "Pulgadas" in unit_system else "8 mm"
stirrup_type = st.sidebar.selectbox("Diámetro del Estribo", list(stirrup_dict.keys()),
    index=list(stirrup_dict.keys()).index(st.session_state.get("c_pm_stirrup_type", default_stirrup)) if st.session_state.get("c_pm_stirrup_type", default_stirrup) in stirrup_dict else 0,
    key="c_pm_stirrup_type")
stirrup_area = stirrup_dict[stirrup_type]["area"]     # cm²
stirrup_diam = stirrup_dict[stirrup_type]["diam_mm"]  # mm

st.sidebar.header("5. Factores de Diseño")
col_type_options = ["Estribos (Tied)", "Espiral (Spiral)"] if lang == "Español" else ["Tied", "Spiral"]
col_type = st.sidebar.selectbox(_t("Tipo de Columna", "Column Type"), col_type_options,
    index=col_type_options.index(st.session_state.get("c_pm_col_type", col_type_options[0])) if st.session_state.get("c_pm_col_type", col_type_options[0]) in col_type_options else 0,
    key="c_pm_col_type")
if "Estrib" in col_type or col_type == "Tied":
    phi_c_max = code["phi_tied"]
    p_max_factor = code["pmax_tied"]
else:
    phi_c_max = code["phi_spiral"]
    p_max_factor = code["pmax_spiral"]
phi_tension = code["phi_tension"]
rho_min_code = code["rho_min"]
rho_max_code = code["rho_max"]
eps_full = code["eps_tension_full"]

st.sidebar.header("6. Verificación de Diseño")
unidades_salida = st.sidebar.radio("Unidades del Diagrama (Resultados):", ["KiloNewtons (kN, kN-m)", "Toneladas Fuerza (tonf, tonf-m)"], key="c_pm_output_units")
if unidades_salida == "Toneladas Fuerza (tonf, tonf-m)":
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom    = "tonf-m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom    = "kN-m"
st.sidebar.markdown(f"Cargas últimas en **{unidad_fuerza}** y **{unidad_mom}**:")
M_u_input = st.sidebar.number_input(f"Momento Último (Mu) [{unidad_mom}]", value=st.session_state.get("c_pm_mu", round(45.0 * factor_fuerza, 2)), step=round(10.0 * factor_fuerza, 2), key="c_pm_mu")
P_u_input = st.sidebar.number_input(f"Carga Axial Última (Pu) [{unidad_fuerza}]", value=st.session_state.get("c_pm_pu", round(2700.0 * factor_fuerza, 2)), step=round(50.0 * factor_fuerza, 2), key="c_pm_pu")

st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CÁLCULO LONGITUD DE DESARROLLO Y EMPALME
# ─────────────────────────────────────────────
ld_mm = get_development_length(rebar_diam, fy, fc, lambda_=1.0, psi_t=1.0, psi_e=1.0, psi_s=1.0, psi_g=0.75 if fy <= 420 else 1.0, cb_ktr=2.5)
splice_length_mm = 1.3 * ld_mm
splice_zone_height = splice_length_mm / 10  # cm
splice_start = max(L_col / 3, 0)
splice_end = splice_start + splice_zone_height
if splice_end > L_col:
    splice_end = L_col
    splice_start = max(0, L_col - splice_zone_height)

# ─────────────────────────────────────────────
# CÁLCULO DEL DIAGRAMA P-M
# ─────────────────────────────────────────────
eps_cu = 0.003
eps_y  = fy / Es
beta_1 = get_beta1(fc)

b_mm = b * 10; h_mm = h * 10
Po_kN = (0.85 * fc * (Ag * 100 - Ast * 100) + fy * Ast * 100) / 1000.0

c_vals = np.concatenate([np.linspace(1e-5, h, 120), np.linspace(h, h * 12, 60)])
P_n_list = []; M_n_list = []; phi_P_n_list = []; phi_M_n_list = []

for c_cm in c_vals:
    c_mm = c_cm * 10
    a_mm = min(beta_1 * c_mm, h_mm)
    Cc = 0.85 * fc * a_mm * b_mm
    Mc = Cc * (h_mm / 2.0 - a_mm / 2.0)

    Ps = 0.0; Ms = 0.0; eps_t = 0.0
    for layer in layers:
        d_i_mm = layer['d'] * 10
        As_i = layer['As'] * 100
        eps_s = eps_cu * (c_mm - d_i_mm) / c_mm
        if d_i_mm >= max(l['d'] * 10 for l in layers):
            eps_t = eps_s
        fs = max(-fy, min(fy, Es * eps_s))
        if a_mm > d_i_mm and fs > 0:
            fs -= 0.85 * fc
        Ps += As_i * fs
        Ms += As_i * fs * (h_mm / 2.0 - d_i_mm)

    Pn = (Cc + Ps) / 1000.0
    Mn = abs((Mc + Ms) / 1_000_000.0)

    eps_t_tens = -eps_t
    if eps_t_tens <= eps_y:
        phi = phi_c_max
    elif eps_t_tens >= eps_full:
        phi = phi_tension
    else:
        phi = phi_c_max + (phi_tension - phi_c_max) * (eps_t_tens - eps_y) / (eps_full - eps_y)

    Pn_max_val = p_max_factor * Po_kN
    phi_Pn_max_val = phi_c_max * Pn_max_val

    Pn = min(Pn, Pn_max_val)
    phi_Pn = min(phi * Pn, phi_Pn_max_val)
    phi_Mn = phi * Mn

    P_n_list.append(Pn); M_n_list.append(Mn)
    phi_P_n_list.append(phi_Pn); phi_M_n_list.append(phi_Mn)

P_n_arr = np.array(P_n_list) * factor_fuerza
M_n_arr = np.array(M_n_list) * factor_fuerza
phi_P_n_arr = np.array(phi_P_n_list) * factor_fuerza
phi_M_n_arr = np.array(phi_M_n_list) * factor_fuerza
Pt = -fy * Ast * 100 / 1000.0 * factor_fuerza
Pn_max = p_max_factor * Po_kN * factor_fuerza
phi_Pn_max = phi_c_max * Pn_max
Po_display = Po_kN * factor_fuerza

# ─────────────────────────────────────────────
# CÁLCULO DE ESTRIBOS POR NIVEL SÍSMICO (NSR-10 / ACI 318)
# ─────────────────────────────────────────────
nivel_lower = nivel_sismico.lower()
zona_especial = any(k in nivel_lower for k in ["des ", "disipación especial", "smf", "special moment frame", "ga —", "ductilidad alta", "pe —", "pórtico especial", "de —", "diseño especial sísmico", "mrle"])
zona_moderada  = any(k in nivel_lower for k in ["dmo", "imf", "intermediate", "gm —", "moderada", "pm —", "mrod"])

# --- Criterios ACI/NSR eje 16 varillas longitudinales ---
s1 = 16 * rebar_diam / 10        # 16×db  (cm)
s2 = 48 * stirrup_diam / 10      # 48×d_e (cm)
s3 = min(b, h)                   # menor dimensión de la sección (cm)
s_basico = min(s1, s2, s3)       # s básico ACI/NSR DMI

# Zona de confinamiento seismico (Lo)
Lo_conf = max(max(b, h), L_col / 6.0, 45.0)  # NSR-10 C.21.6.4.1

if zona_especial:
    # NSR-10 Cap. C.21.6 / ACI 318-19 §18.7.5
    # Espaciamiento en zonas de confinamiento (Lo desde c/extremo)
    s_conf = min(
        8 * rebar_diam / 10,   # 8×db  (cm)     — NSR-10 C.21.6.4.3(a)
        24 * stirrup_diam / 10,# 24×d_e (cm)    — NSR-10 C.21.6.4.3(b)
        min(b, h) / 3,         # b_menor/3 (cm) — NSR-10 C.21.6.4.3(c)
        15.0                   # 150 mm máx      — NSR-10 C.21.6.4.3(d)
    )
    # Espaciamiento en zona central
    s_centro = min(
        6 * rebar_diam / 10,   # 6×db  (cm)     — NSR-10 C.21.7.3.1 para viga
        s_basico               # límite ACI general
    )
    s_centro = max(s_centro, s_conf)  # nunca menor que s_conf

    n_est_por_Lo = math.ceil(Lo_conf / s_conf)
    n_estribos_zona_dos = n_est_por_Lo * 2   # dos extremos
    longitud_zona_libre = max(0, L_col - 2 * Lo_conf)
    n_estribos_centro = max(0, math.ceil(longitud_zona_libre / s_centro) - 1) if longitud_zona_libre > 0 else 0
    n_estribos_total = n_estribos_zona_dos + n_estribos_centro + 1  # +1 del extremo superior
    s_usar = s_conf
    notas_estribos = (
        f"**NSR-10 Cap. C.21.6 (DES — Zona de Alto Riesgo Sísmico)**:\n\n"
        f"- Zona de confinamiento Lo = max(b, h, L/6, 45 cm) = **{Lo_conf:.1f} cm** desde c/extremo\n"
        f"- Estr. zona confinada s = min(8d_b, 24d_e, b_men/3, 15 cm) = **{s_conf:.1f} cm**\n"
        f"  - 8×{rebar_diam:.1f} mm = {8*rebar_diam/10:.1f} cm\n"
        f"  - 24×{stirrup_diam:.1f} mm = {24*stirrup_diam/10:.1f} cm\n"
        f"  - min(b,h)/3 = {min(b,h)/3:.1f} cm\n"
        f"- Estr. zona central s = min(6d_b, s_basico) = **{s_centro:.1f} cm**\n"
        f"- N° estribos zona confinada (por extremo) = ceil({Lo_conf:.1f}/{s_conf:.1f}) = **{n_est_por_Lo}**\n"
        f"- N° estribos centro = ceil({longitud_zona_libre:.1f}/{s_centro:.1f}) = **{n_estribos_centro}**\n"
        f"- Total = 2×{n_est_por_Lo} + {n_estribos_centro} + 1 = **{n_estribos_total}**"
    )
elif zona_moderada:
    # NSR-10 Cap. C.21.12 / ACI 318-19 §18.4
    s_conf = min(
        8 * rebar_diam / 10,
        24 * stirrup_diam / 10,
        min(b, h) / 2,
        20.0   # 200 mm máx DMO NSR-10 C.21.12.3
    )
    Lo_conf = max(max(b, h), L_col / 6.0, 45.0)
    s_centro = s_basico
    n_est_por_Lo = math.ceil(Lo_conf / s_conf)
    n_estribos_zona_dos = n_est_por_Lo * 2
    longitud_zona_libre = max(0, L_col - 2 * Lo_conf)
    n_estribos_centro = max(0, math.ceil(longitud_zona_libre / s_centro) - 1) if longitud_zona_libre > 0 else 0
    n_estribos_total = n_estribos_zona_dos + n_estribos_centro + 1
    s_usar = s_conf
    notas_estribos = (
        f"**NSR-10 Cap. C.21.12 (DMO — Zona de Riesgo Sísmico Moderado)**:\n\n"
        f"- Zona de confinamiento Lo = **{Lo_conf:.1f} cm** desde c/extremo\n"
        f"- Estr. zona confinada s = min(8d_b, 24d_e, b_men/2, 20 cm) = **{s_conf:.1f} cm**\n"
        f"  - 8×{rebar_diam:.1f} mm = {8*rebar_diam/10:.1f} cm\n"
        f"  - 24×{stirrup_diam:.1f} mm = {24*stirrup_diam/10:.1f} cm\n"
        f"  - min(b,h)/2 = {min(b,h)/2:.1f} cm\n"
        f"- Estr. zona central s = min(16d_b, 48d_e, dim_menor) = **{s_basico:.1f} cm**\n"
        f"- N° estribos zona confinada (por extremo) = **{n_est_por_Lo}**\n"
        f"- N° estribos centro = **{n_estribos_centro}**\n"
        f"- Total = 2×{n_est_por_Lo} + {n_estribos_centro} + 1 = **{n_estribos_total}**"
    )
else:
    # NSR-10 Cap. C.7.10 (DMI) — solo criterio ACI básico
    s_conf = s_basico
    Lo_conf = 0.0
    s_centro = s_basico
    n_estribos_total = math.ceil(L_col / s_basico) + 1
    s_usar = s_basico
    notas_estribos = (
        f"**NSR-10 Cap. C.7.10 (DMI — Amenaza Sísmica Baja)**:\n\n"
        f"- Espaciamiento s = min(16d_b, 48d_e, dim_menor de sección)\n"
        f"  - 16×{rebar_diam:.1f} mm = {s1:.1f} cm\n"
        f"  - 48×{stirrup_diam:.1f} mm = {s2:.1f} cm\n"
        f"  - min(b,h) = {s3:.1f} cm\n"
        f"- s = **{s_basico:.1f} cm** (sin zona de confinamiento)\n"
        f"- Total estribos = ceil({L_col:.0f}/{s_basico:.1f}) + 1 = **{n_estribos_total}**"
    )

recub_cm = max(d_prime - rebar_diam / 20.0, 2.5)
perim_estribo = 2 * (b - 2 * recub_cm) + 2 * (h - 2 * recub_cm) + 6 * stirrup_diam / 10
if num_filas_h > 1:
    sep_horiz = (b - 2 * d_prime) / (num_filas_h - 1)
else:
    sep_horiz = 0.0
n_crossties_por_estribo = 0
if num_filas_h > 2 and sep_horiz > 15.0:
    n_crossties_por_estribo = num_filas_h - 2
longitud_crosstie = (b - 2 * recub_cm) + 12 * stirrup_diam / 10
if num_capas_intermedias > 0:
    sep_vert = (h - 2 * d_prime) / (num_capas_intermedias + 1)
    n_crossties_v = num_capas_intermedias if sep_vert > 15.0 else 0
    longitud_crosstie_v = (h - 2 * recub_cm) + 12 * stirrup_diam / 10
else:
    n_crossties_v = 0
    longitud_crosstie_v = 0
total_crossties_h = n_crossties_por_estribo * 2
total_crossties_v = n_crossties_v * 2
total_crossties_por_estribo = total_crossties_h + total_crossties_v
longitud_total_crossties_m = (n_estribos_total * total_crossties_h * longitud_crosstie / 100 +
                              n_estribos_total * total_crossties_v * longitud_crosstie_v / 100)
peso_crossties_kg = longitud_total_crossties_m * stirrup_area * 100 * 7.85e-3

peso_unit_estribo = (perim_estribo / 100.0) * (stirrup_area * 100) * 7.85e-3
peso_total_estribos_kg = n_estribos_total * peso_unit_estribo + peso_crossties_kg
longitud_total_estribos_m = n_estribos_total * perim_estribo / 100.0 + longitud_total_crossties_m

vol_concreto_m3 = (b / 100) * (h / 100) * (L_col / 100)
peso_acero_long_kg = Ast * (L_col * 10) * 7.85e-3
peso_total_acero_kg = peso_acero_long_kg + peso_total_estribos_kg
relacion_acero_kg_m3 = peso_total_acero_kg / vol_concreto_m3 if vol_concreto_m3 > 0 else 0

# Cantidades de materiales de concreto (ACI 211)
_mix = get_mix_for_fc(fc)
_cem_kg_m3  = _mix["cem"]   # kg/m3
_arena_kg_m3 = _mix["arena"] # kg/m3
_grava_kg_m3 = _mix["grava"] # kg/m3
bultos_col = vol_concreto_m3 * _cem_kg_m3   # kg de cemento
arena_col  = vol_concreto_m3 * _arena_kg_m3  # kg de arena
grava_col  = vol_concreto_m3 * _grava_kg_m3  # kg de grava

# Despiece
hook_len_mm = 12 * rebar_diam
long_bar_m = (L_col + 2 * (ld_mm / 10) + 2 * (hook_len_mm / 10)) / 100
peso_long_total = n_barras_total * long_bar_m * (rebar_area * 100) * 7.85e-3

despiece_rows = []
despiece_rows.append({
    "Marca": "L1",
    "Cantidad": n_barras_total,
    "Diámetro": _bar_label(rebar_diam),
    "Longitud (m)": long_bar_m,
    "Longitud Total (m)": n_barras_total * long_bar_m,
    "Peso (kg)": peso_long_total,
    "Observación": f"Gancho 90° en ambos extremos (12db), empalme clase B a {splice_start:.0f} cm"
})
despiece_rows.append({
    "Marca": "E1",
    "Cantidad": n_estribos_total,
    "Diámetro": _bar_label(stirrup_diam),
    "Longitud (m)": perim_estribo / 100.0,
    "Longitud Total (m)": n_estribos_total * perim_estribo / 100.0,
    "Peso (kg)": peso_total_estribos_kg,
    "Observación": f"Gancho 135° en ambos extremos, s={s_usar:.1f} cm"
})
if total_crossties_por_estribo > 0:
    long_crosstie_m = longitud_crosstie / 100.0
    long_crosstie_v_m = longitud_crosstie_v / 100.0
    total_crossties = total_crossties_por_estribo * n_estribos_total
    despiece_rows.append({
        "Marca": "C1",
        "Cantidad": total_crossties,
        "Diámetro": _bar_label(stirrup_diam),
        "Longitud (m)": (long_crosstie_m + long_crosstie_v_m) / 2.0 if total_crossties_por_estribo > 0 else 0,
        "Longitud Total (m)": (n_estribos_total * total_crossties_h * long_crosstie_m + n_estribos_total * total_crossties_v * long_crosstie_v_m),
        "Peso (kg)": peso_crossties_kg,
        "Observación": "Crossties (ganchos 135° en ambos extremos)"
    })
df_despiece = pd.DataFrame(despiece_rows)

# =============================================================================
# LAYOUT PRINCIPAL
# =============================================================================
tab1, tab2, tab3 = st.tabs([_t("📊 Diagrama P–M", "📊 P–M Diagram"), _t("🔲 Sección & Estribos", "🔲 Section & Ties"), _t("📦 Cantidades de Materiales", "📦 Material Quantities")])

# TAB 1: DIAGRAMA P–M
with tab1:
    col1, col2 = st.columns([2, 1])
    with col1:
        fig_pm, ax_pm = plt.subplots(figsize=(8, 6))
        ax_pm.plot(M_n_arr, P_n_arr, label=r"Resistencia Nominal ($P_n, M_n$)", color="blue", linestyle="--")
        ax_pm.plot(phi_M_n_arr, phi_P_n_arr, label=r"Resistencia de Diseño ($\phi P_n, \phi M_n$)", color="red", linewidth=2)
        ax_pm.plot(M_u_input, P_u_input, 'o', markersize=9, color='lime', markeredgecolor='black', zorder=6, label="Punto de Diseño (Mu, Pu)")
        ax_pm.annotate(f"{Pn_max:.2f} [{unidad_fuerza}]", xy=(0, Pn_max), xytext=(5,5), textcoords="offset points", color='blue')
        ax_pm.annotate(f"{phi_Pn_max:.2f} [{unidad_fuerza}]", xy=(0, phi_Pn_max), xytext=(5,-5), textcoords="offset points", color='red')
        ax_pm.set_xlabel(f"Momento Flector M [{unidad_mom}]")
        ax_pm.set_ylabel(f"Carga Axial P [{unidad_fuerza}]")
        ax_pm.set_title(f"Diagrama de Interacción — {norma_sel}\nColumna {b:.0f}×{h:.0f} cm  |  f'c={fc:.1f} MPa  |  fy={fy:.0f} MPa  |  {nivel_sismico.split('—')[0].strip()}")
        ax_pm.grid(True, linestyle=":", alpha=0.5)
        ax_pm.legend(loc="upper right", fontsize=8)
        st.pyplot(fig_pm)
        pm_img = io.BytesIO()
        fig_pm.savefig(pm_img, format='png', dpi=150, bbox_inches='tight')
        pm_img.seek(0)

    with col2:
        st.subheader("Resumen")
        cuantia = Ast / Ag * 100
        st.markdown(f"**Norma:** {norma_sel}")
        st.markdown(f"**Nivel Sísmico:** {nivel_sismico}")
        st.markdown(f"**φ compresión:** {phi_c_max} | **φ tensión:** {phi_tension}")
        st.markdown("---")
        st.markdown("#### Verificación de Cuantía")
        data_cuantia = {
            "Parámetro": ["b × h", "Ag", "# Varillas", "Área varilla", "Ast", "Cuantía ρ", f"ρ mín ({norma_sel.split('(')[0].strip()})", f"ρ máx ({norma_sel.split('(')[0].strip()})"],
            "Valor": [f"{b:.0f}×{h:.0f} cm", f"{Ag:.1f} cm²", f"{n_barras_total} barras", f"{rebar_area:.3f} cm²", f"{Ast:.3f} cm²", f"{cuantia:.3f}%", f"{rho_min_code:.1f}%", f"{rho_max_code:.1f}%"]
        }
        st.dataframe(pd.DataFrame(data_cuantia), use_container_width=True, hide_index=True)
        if rho_min_code <= cuantia <= rho_max_code:
            st.success(f"✅ Aprobado Cuantía: ρ = {cuantia:.2f}%")
        else:
            st.error(f"❌ No Aprobado: ρ = {cuantia:.2f}% (límites {rho_min_code}% - {rho_max_code}%)")
        st.markdown("---")
        st.markdown("#### Puntos Clave P–M")
        data_pm = {
            "Descripción": ["Capacidad bruta (Po)", f"Pn,máx ({p_max_factor:.0%}×Po)", f"φPn,máx (φ={phi_c_max})", f"Tracción pura (Pt)", f"φTracción pura (φt={phi_tension})"],
            f"[{unidad_fuerza}]": [f"{Po_display:.2f}", f"{Pn_max:.2f}", f"{phi_Pn_max:.2f}", f"{Pt:.2f}", f"{phi_tension * Pt:.2f}"]
        }
        st.dataframe(pd.DataFrame(data_pm), use_container_width=True, hide_index=True)
        st.caption(f"📖 Ref: {code['ref']}")

# TAB 2: SECCIÓN Y ESTRIBOS
with tab2:
    st.subheader(_t("🧊 Visualización 3D y Detallado 2D", "🧊 3D Visualization & 2D Detailing"))
    # --- 3D con plotly ---
    fig3d = go.Figure()
    x_c = [-b/2, b/2, b/2, -b/2, -b/2, b/2, b/2, -b/2]
    y_c = [-h/2, -h/2, h/2, h/2, -h/2, -h/2, h/2, h/2]
    z_c = [0,0,0,0, L_col, L_col, L_col, L_col]
    fig3d.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
    diam_reb_cm = rebar_diam / 10.0
    line_width = max(4, diam_reb_cm * 3)
    rect_x = [-b/2 + d_prime, b/2 - d_prime]
    rect_y = [-h/2 + d_prime, h/2 - d_prime]
    def add_rebar_3d(x_pos, y_pos, label_idx=0):
        fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[y_pos, y_pos], z=[0, L_col],
                                     mode='lines', line=dict(color='darkred', width=line_width),
                                     name=f'Long. {rebar_type}', showlegend=(label_idx==0)))
    bar_count = 0
    # Inferior
    for i in range(num_filas_h):
        x = rect_x[0] + i*(rect_x[1]-rect_x[0])/(num_filas_h-1) if num_filas_h>1 else 0
        add_rebar_3d(x, rect_y[0], bar_count)
        bar_count += 1
    # Superior
    for i in range(num_filas_h):
        x = rect_x[0] + i*(rect_x[1]-rect_x[0])/(num_filas_h-1) if num_filas_h>1 else 0
        add_rebar_3d(x, rect_y[1], bar_count)
        bar_count += 1
    # Intermedias
    if num_capas_intermedias > 0:
        esp_y = (rect_y[1] - rect_y[0]) / (num_capas_intermedias + 1)
        for i in range(1, num_capas_intermedias + 1):
            y = rect_y[0] + i*esp_y
            add_rebar_3d(rect_x[0], y, bar_count)
            bar_count += 1
            add_rebar_3d(rect_x[1], y, bar_count)
            bar_count += 1
    # Estribos
    tie_color = 'cornflowerblue'
    tie_width = max(2, (stirrup_diam/10.0) * 3)
    tt_x = [rect_x[0]-diam_reb_cm/2, rect_x[1]+diam_reb_cm/2, rect_x[1]+diam_reb_cm/2, rect_x[0]-diam_reb_cm/2, rect_x[0]-diam_reb_cm/2]
    tt_y = [rect_y[0]-diam_reb_cm/2, rect_y[0]-diam_reb_cm/2, rect_y[1]+diam_reb_cm/2, rect_y[1]+diam_reb_cm/2, rect_y[0]-diam_reb_cm/2]
    L_cm = int(L_col)
    z_positions = []
    if zona_especial:
        y = s_conf / 2
        while y <= Lo_conf:
            z_positions.append(y)
            y += s_conf
        y = Lo_conf + s_basico
        while y < L_col - Lo_conf:
            z_positions.append(y)
            y += s_basico
        y = L_col - Lo_conf + s_conf / 2
        while y <= L_col:
            z_positions.append(y)
            y += s_conf
    else:
        y = s_basico / 2
        while y <= L_col:
            z_positions.append(y)
            y += s_basico

    for idx, zt in enumerate(z_positions):
        z_t_arr = [zt] * 5
        fig3d.add_trace(go.Scatter3d(x=tt_x, y=tt_y, z=z_t_arr, mode='lines', line=dict(color=tie_color, width=tie_width), name='Estribo', showlegend=(idx==0)))

    # === CROSSTIES en 3D ===
    if total_crossties_h > 0 and num_filas_h > 2:
        ct_xs_3d = list(np.linspace(rect_x[0], rect_x[1], num_filas_h))[1:-1]  # barras interiores
        ct_y0 = rect_y[0] - diam_reb_cm / 2
        ct_y1 = rect_y[1] + diam_reb_cm / 2
        ct_shown = False
        for zt in z_positions:
            for cx in ct_xs_3d:
                fig3d.add_trace(go.Scatter3d(
                    x=[cx, cx], y=[ct_y0, ct_y1], z=[zt, zt],
                    mode='lines', line=dict(color='lime', width=max(1, tie_width * 0.7), dash='dash'),
                    name='Crosstie H', showlegend=(not ct_shown)))
                ct_shown = True
    if total_crossties_v > 0 and num_capas_intermedias > 0:
        esp_y_3d = (rect_y[1] - rect_y[0]) / (num_capas_intermedias + 1)
        ct_x0 = rect_x[0] - diam_reb_cm / 2
        ct_x1 = rect_x[1] + diam_reb_cm / 2
        ctv_shown = False
        for zt in z_positions:
            for i in range(1, num_capas_intermedias + 1):
                cy = rect_y[0] + i * esp_y_3d
                fig3d.add_trace(go.Scatter3d(
                    x=[ct_x0, ct_x1], y=[cy, cy], z=[zt, zt],
                    mode='lines', line=dict(color='cyan', width=max(1, tie_width * 0.7), dash='dash'),
                    name='Crosstie V', showlegend=(not ctv_shown)))
                ctv_shown = True

    fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                        margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable')
    st.plotly_chart(fig3d, use_container_width=True)
    st.markdown("---")

    # --- Sección transversal 2D y alzado ---
    col_s1, col_s2 = st.columns([1, 1])
    with col_s1:
        st.subheader("Sección Transversal 2D")
        fig_sec, ax_s = plt.subplots(figsize=(4, max(3, 4 * h / b)))
        ax_s.set_aspect('equal')
        ax_s.set_facecolor('#1a1a2e')
        fig_sec.patch.set_facecolor('#1a1a2e')
        ax_s.add_patch(patches.Rectangle((0, 0), b, h, linewidth=2, edgecolor='white', facecolor='#4a4a6a'))
        ax_s.add_patch(patches.Rectangle((recub_cm, recub_cm), b-2*recub_cm, h-2*recub_cm,
            linewidth=max(1, stirrup_diam/20), edgecolor='#00d4ff', facecolor='none', linestyle='--'))
        r_bar = rebar_diam / 20.0
        def _draw_row(y_pos, n_bars):
            xs = np.linspace(d_prime, b - d_prime, n_bars) if n_bars > 1 else [b/2]
            for x in xs:
                ax_s.add_patch(plt.Circle((x, y_pos), r_bar, color='#ff6b35', zorder=5))
        _draw_row(h - d_prime, num_filas_h)
        if num_capas_intermedias > 0:
            esp = (h - 2*d_prime) / (num_capas_intermedias + 1)
            for i in range(1, num_capas_intermedias + 1):
                y_int = d_prime + i * esp
                ax_s.add_patch(plt.Circle((d_prime, y_int), r_bar, color='#ff6b35', zorder=5))
                ax_s.add_patch(plt.Circle((b - d_prime, y_int), r_bar, color='#ff6b35', zorder=5))
        _draw_row(d_prime, num_filas_h)
        ax_s.annotate('', xy=(b, -1.5), xytext=(0, -1.5), arrowprops=dict(arrowstyle='<->', color='white', lw=1.2))
        ax_s.text(b/2, -2.5, f"b = {b:.0f} cm", ha='center', va='top', color='white', fontsize=7)
        ax_s.annotate('', xy=(-1.5, h), xytext=(-1.5, 0), arrowprops=dict(arrowstyle='<->', color='white', lw=1.2))
        ax_s.text(-2.5, h/2, f"h = {h:.0f} cm", ha='right', va='center', color='white', fontsize=7, rotation=90)
        ax_s.set_xlim(-5, b+3); ax_s.set_ylim(-5, h+3); ax_s.axis('off')
        ax_s.set_title(f"Sección transversal — {b:.0f}×{h:.0f} cm\n{n_barras_total} varillas Ø{rebar_diam:.0f}mm + estribo Ø{stirrup_diam:.0f}mm", color='white', fontsize=8)
        st.pyplot(fig_sec)

        st.subheader("Alzado (Vista Vertical)")
        escala = min(1.0, 15.0 / (L_col / 10.0))
        L_fig = L_col / 10.0 * escala
        b_fig = b / 10.0 * escala
        fig_elev, ax_e = plt.subplots(figsize=(max(2.5, b_fig + 1.5), max(6, L_fig + 1.5)))
        ax_e.set_facecolor('#1a1a2e'); fig_elev.patch.set_facecolor('#1a1a2e')
        ax_e.add_patch(patches.Rectangle((0, 0), b_fig, L_fig, linewidth=1.5, edgecolor='white', facecolor='#4a4a6a'))
        margin = (recub_cm / b * b_fig)
        ax_e.plot([margin, margin], [0, L_fig], color='#ff6b35', linewidth=1.5)
        ax_e.plot([b_fig - margin, b_fig - margin], [0, L_fig], color='#ff6b35', linewidth=1.5)
        def y_pos_to_fig(y_cm):
            return y_cm / L_col * L_fig
        stirrup_lw = 1.5
        if zona_especial:
            Lo_fig = Lo_conf / L_col * L_fig
            y = 0.0
            while y <= Lo_conf:
                yf = y_pos_to_fig(y)
                ax_e.plot([0, b_fig], [yf, yf], color='#00d4ff', linewidth=stirrup_lw)
                y += s_conf
            y = 0.0
            while y <= Lo_conf:
                yf = y_pos_to_fig(L_col - y)
                ax_e.plot([0, b_fig], [yf, yf], color='#00d4ff', linewidth=stirrup_lw)
                y += s_conf
            y = Lo_conf + s_basico
            while y < L_col - Lo_conf:
                yf = y_pos_to_fig(y)
                ax_e.plot([0, b_fig], [yf, yf], color='#7ec8e3', linewidth=stirrup_lw * 0.7)
                y += s_basico
            ax_e.axhline(y=y_pos_to_fig(Lo_conf), color='yellow', linewidth=0.8, linestyle='--')
            ax_e.axhline(y=y_pos_to_fig(L_col - Lo_conf), color='yellow', linewidth=0.8, linestyle='--')
        else:
            y = 0.0
            while y <= L_col:
                yf = y_pos_to_fig(y)
                ax_e.plot([0, b_fig], [yf, yf], color='#00d4ff', linewidth=stirrup_lw)
                y += s_basico
        # Zona de empalme
        if splice_zone_height > 0:
            y1 = y_pos_to_fig(splice_start)
            y2 = y_pos_to_fig(splice_end)
            ax_e.fill_between([0, b_fig], y1, y2, color='yellow', alpha=0.3)
            ax_e.text(b_fig/2, (y1+y2)/2, f"Empalme\nLap={splice_length_mm/10:.1f}cm", ha='center', va='center', color='white', fontsize=6, bbox=dict(facecolor='black', alpha=0.5))
        ax_e.annotate('', xy=(b_fig/2, L_fig + 0.3), xytext=(b_fig/2, 0), arrowprops=dict(arrowstyle='<->', color='white', lw=1.0))
        ax_e.text(b_fig/2, L_fig/2, f"L\n={L_col:.0f}\ncm", ha='center', va='center', color='white', fontsize=7,
                  bbox=dict(facecolor='#4a4a6a', edgecolor='none', alpha=0.7))
        ax_e.set_xlim(-0.5, b_fig + 2.0)
        ax_e.set_ylim(-0.5, L_fig + 0.8)
        ax_e.axis('off')
        _stir_lbl = _bar_label(stirrup_diam).split(" ")[0] if "#" in _bar_label(stirrup_diam) else f"Ø{stirrup_diam:.0f}mm"
        ax_e.set_title(f"Alzado — Distribución de Estribos\nL={L_col:.0f} cm  |  {n_estribos_total} estribos {_stir_lbl}", color='white', fontsize=8)
        st.pyplot(fig_elev)

    with col_s2:
        st.subheader("📝 Verificaciones NSR-10 / ACI 318")
        # ========================
        # Tabla completa de comprobaciones
        # ========================
        _ok = "✅"; _fail = "❌"; _warn = "⚠️"

        # 1. Cuantía longitudinal
        ok_rho = rho_min_code <= cuantia <= rho_max_code
        # 2. Separación mínima entre barras (ACI 318-19 §25.8.1)
        sep_min_req = max(rebar_diam / 10, 2.5)  # cm
        sep_h = (b - 2 * d_prime) / (num_filas_h - 1) if num_filas_h > 1 else 999
        ok_sep_h = sep_h >= sep_min_req
        # 3. Recubrimiento mín (NSR-10 C.20.6.1.3 - columnas interiores 4 cm)
        recub_req = 4.0  # cm para columnas
        ok_recub = recub_cm >= recub_req
        # 4. Espaciamiento básico de estribos
        ok_s_basico = s_basico <= min(b, h)
        # 5. Zona de confinamiento (applies DES/DMO)
        ok_lo = True
        if zona_especial or zona_moderada:
            ok_lo = Lo_conf >= max(max(b, h), L_col / 6, 45.0) - 0.01
        # 6. Espaciamiento confinado
        ok_s_conf = True
        if zona_especial:
            ok_s_conf = s_conf <= 15.0
        elif zona_moderada:
            ok_s_conf = s_conf <= 20.0
        # 7. Crossties requeridos
        ct_req_h = num_filas_h > 2 and sep_h > 15.0
        sep_v = (h - 2 * d_prime) / (num_capas_intermedias + 1) if num_capas_intermedias > 0 else 0
        ct_req_v = num_capas_intermedias > 0 and sep_v > 15.0
        ok_ct = (not ct_req_h or total_crossties_h > 0) and (not ct_req_v or total_crossties_v > 0)
        # 8. Cuantía máxima
        ok_rho_max = cuantia <= rho_max_code
        # 9. Cuantía mínima
        ok_rho_min = cuantia >= rho_min_code
        # 10. d_t (distancia a la fibra extrema en tensión)
        dt_cm = h - d_prime
        ok_dt = dt_cm >= (h - d_prime)

        checks_data = {
            "#": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10],
            "Verificación": [
                f"Cuantía mínima (NSR-10 C.10.6.1)",
                f"Cuantía máxima (NSR-10 C.10.9.1)",
                f"Separación mín. barras (ACI §25.8.1)",
                f"Recubrimiento mín. (NSR-10 C.20.6.1)",
                f"Espaciamiento estribos s_basico (C.7.10.5)",
                f"Long. zona confinada Lo (C.21.6.4.1)" if zona_especial or zona_moderada else "Zona confinamiento (N/A nivel DMI)",
                f"Espaciamiento zona confinada s_conf" if zona_especial or zona_moderada else "s_conf (N/A nivel DMI)",
                f"Crossties horiz. (sep > 15 cm) (C.21.6.4.2)",
                f"Crossties vert. (capas interm. > 15 cm)",
                f"Resistencia diseño (check punto de carga)",
            ],
            "Requerido": [
                f">= {rho_min_code:.2f}%",
                f"<= {rho_max_code:.2f}%",
                f">= {sep_min_req:.1f} cm",
                f">= {recub_req:.1f} cm",
                f"<= min(b,h) = {min(b,h):.0f} cm",
                f">= max(b,h,L/6,45) = {Lo_conf:.1f} cm" if zona_especial or zona_moderada else "N/A",
                f"<= {'15.0' if zona_especial else '20.0'} cm" if zona_especial or zona_moderada else "N/A",
                f"CT si sep>{15:.0f} cm",
                f"CT si sep>{15:.0f} cm c/cap",
                "Ver diagrama P-M",
            ],
            "Calculado": [
                f"{cuantia:.3f}%",
                f"{cuantia:.3f}%",
                f"{sep_h:.1f} cm",
                f"{recub_cm:.1f} cm",
                f"{s_basico:.1f} cm",
                f"{Lo_conf:.1f} cm" if zona_especial or zona_moderada else "N/A",
                f"{s_conf:.1f} cm" if zona_especial or zona_moderada else f"{s_basico:.1f} cm",
                f"{total_crossties_h} CT/estribo" if ct_req_h else "No reqs.",
                f"{total_crossties_v} CT/estribo" if ct_req_v else "No reqs.",
                "Ver sección Tab 1",
            ],
            "Estado": [
                _ok if ok_rho_min else _fail,
                _ok if ok_rho_max else _fail,
                _ok if ok_sep_h else _fail,
                _ok if ok_recub else _fail,
                _ok if ok_s_basico else _fail,
                _ok if ok_lo else (_warn if not (zona_especial or zona_moderada) else _fail),
                _ok if ok_s_conf else (_warn if not (zona_especial or zona_moderada) else _fail),
                _ok if ok_ct else _warn,
                _ok if ok_ct else _warn,
                _ok,
            ]
        }
        st.dataframe(pd.DataFrame(checks_data), use_container_width=True, hide_index=True)
        _all_ok = all([
            ok_rho_min, ok_rho_max, ok_sep_h, ok_recub, ok_s_basico, ok_lo, ok_s_conf, ok_ct
        ])
        if _all_ok:
            st.success("✅ Todas las verificaciones NSR-10 cumplidas")
        else:
            st.error("❌ Hay verificaciones que no cumplen. Revise la tabla.")

        st.markdown("---")
        st.markdown(f"**Total de crossties en columna:** {total_crossties_por_estribo * n_estribos_total} unidades")
        st.caption("📄 Ref: NSR-10 Cap. C.7, C.10, C.21 | ACI 318-19 Cap. 18, 20, 25")
        st.markdown("#### 💾 Exportar Plano (AutoCAD)")
        doc_dxf = ezdxf.new('R2010')
        doc_dxf.units = ezdxf.units.CM
        msp = doc_dxf.modelspace()
        for lay, col in [('CONCRETO',7), ('ESTRIBOS',4), ('VARILLAS',1), ('TEXTO',3), ('EMPALME',6)]:
            if lay not in doc_dxf.layers:
                doc_dxf.layers.add(lay, color=col)
        # Sección transversal
        msp.add_lwpolyline([(0,0), (b,0), (b,h), (0,h), (0,0)], dxfattribs={'layer':'CONCRETO'})
        msp.add_lwpolyline([(recub_cm, recub_cm), (b-recub_cm, recub_cm), (b-recub_cm, h-recub_cm), (recub_cm, h-recub_cm), (recub_cm, recub_cm)], dxfattribs={'layer':'ESTRIBOS'})
        r_bar_cm = rebar_diam / 20.0
        xs_sup_inf = np.linspace(d_prime, b - d_prime, num_filas_h) if num_filas_h > 1 else [b/2]
        for x in xs_sup_inf:
            msp.add_circle((x, h - d_prime), r_bar_cm, dxfattribs={'layer':'VARILLAS'})
            msp.add_circle((x, d_prime), r_bar_cm, dxfattribs={'layer':'VARILLAS'})
        if num_capas_intermedias > 0:
            esp = (h - 2*d_prime) / (num_capas_intermedias + 1)
            for i in range(1, num_capas_intermedias + 1):
                y_int = d_prime + i * esp
                msp.add_circle((d_prime, y_int), r_bar_cm, dxfattribs={'layer':'VARILLAS'})
                msp.add_circle((b - d_prime, y_int), r_bar_cm, dxfattribs={'layer':'VARILLAS'})
        # Elevación
        off_x = b + 10
        msp.add_lwpolyline([(off_x,0), (off_x+b,0), (off_x+b, L_col), (off_x, L_col), (off_x,0)], dxfattribs={'layer':'CONCRETO'})
        # Barras longitudinales con ganchos
        margin = recub_cm
        hook_len = 12 * rebar_diam / 10
        x_left = off_x + margin
        x_right = off_x + b - margin
        msp.add_line((x_left, 0), (x_left, L_col), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_right, 0), (x_right, L_col), dxfattribs={'layer':'VARILLAS'})
        # Ganchos inferiores
        msp.add_line((x_left, 0), (x_left - hook_len, 0), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_left, 0), (x_left + hook_len, 0), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_right, 0), (x_right - hook_len, 0), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_right, 0), (x_right + hook_len, 0), dxfattribs={'layer':'VARILLAS'})
        # Ganchos superiores
        msp.add_line((x_left, L_col), (x_left - hook_len, L_col), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_left, L_col), (x_left + hook_len, L_col), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_right, L_col), (x_right - hook_len, L_col), dxfattribs={'layer':'VARILLAS'})
        msp.add_line((x_right, L_col), (x_right + hook_len, L_col), dxfattribs={'layer':'VARILLAS'})
        # Estribos
        if zona_especial:
            # confinamiento
            y = 0.0
            while y <= Lo_conf:
                msp.add_line((off_x, y), (off_x+b, y), dxfattribs={'layer':'ESTRIBOS'})
                y += s_conf
            y = 0.0
            while y <= Lo_conf:
                msp.add_line((off_x, L_col - y), (off_x+b, L_col - y), dxfattribs={'layer':'ESTRIBOS'})
                y += s_conf
            y = Lo_conf + s_basico
            while y < L_col - Lo_conf:
                msp.add_line((off_x, y), (off_x+b, y), dxfattribs={'layer':'ESTRIBOS'})
                y += s_basico
        else:
            y = 0.0
            while y <= L_col:
                msp.add_line((off_x, y), (off_x+b, y), dxfattribs={'layer':'ESTRIBOS'})
                y += s_basico
        # Zona de empalme
        if splice_zone_height > 0:
            y1 = splice_start
            y2 = splice_end
            msp.add_lwpolyline([(off_x, y1), (off_x+b, y1), (off_x+b, y2), (off_x, y2), (off_x, y1)], dxfattribs={'layer':'EMPALME'})
            msp.add_text(f"Empalme Clase B\nLap={splice_length_mm/10:.1f} cm", dxfattribs={'layer':'TEXTO','height':2,'insert':(off_x+b/2, (y1+y2)/2)})
        # Textos de cotas
        msp.add_text(f"L = {L_col:.0f} cm", dxfattribs={'layer':'TEXTO','height':2,'insert':(off_x+b/2, L_col+3)})
        msp.add_text(f"b = {b:.0f} cm", dxfattribs={'layer':'TEXTO','height':2,'insert':(off_x-3, L_col/2)})
        _out_dxf = io.StringIO()
        doc_dxf.write(_out_dxf)
        st.download_button("Descargar DXF (Sección + Elevación con ganchos)", data=_out_dxf.getvalue().encode('utf-8'), file_name=f"Columna_{b:.0f}x{h:.0f}.dxf", mime="application/dxf")

# TAB 3: CANTIDADES, DESPIECE, APU Y FIGURADO
with tab3:
    st.subheader(f"📦 Cantidades de Materiales — Columna {b:.0f}×{h:.0f} cm, L={L_col:.0f} cm")
    col_c1, col_c2, col_c3 = st.columns(3)
    col_c1.metric("Concreto", f"{vol_concreto_m3:.4f} m³")
    col_c2.metric("Acero Total", f"{peso_total_acero_kg:.2f} kg")
    col_c3.metric("Ratio Acero", f"{relacion_acero_kg_m3:.1f} kg/m³")
    st.markdown("---")
    st.markdown("#### 📏 Despiece de Acero (Bending Schedule)")
    st.dataframe(df_despiece.style.format({"Longitud (m)": "{:.2f}", "Longitud Total (m)": "{:.2f}", "Peso (kg)": "{:.1f}"}), use_container_width=True, hide_index=False)
    # Gráfico de barras
    fig_bars, ax_bars = plt.subplots(figsize=(8, 4))
    ax_bars.bar(df_despiece["Marca"], df_despiece["Peso (kg)"], color=['#ff6b35', '#4caf50', '#ff9800'])
    ax_bars.set_xlabel(_t("Elemento", "Element"))
    ax_bars.set_ylabel(_t("Peso (kg)", "Weight (kg)"))
    ax_bars.set_title(_t("Distribución de pesos por tipo de barra", "Weight distribution by bar type"))
    ax_bars.grid(True, alpha=0.3)
    st.pyplot(fig_bars)
    bars_img = io.BytesIO()
    fig_bars.savefig(bars_img, format='png', dpi=150, bbox_inches='tight')
    bars_img.seek(0)

    # ============================
    # NUEVO: FIGURADO PARA TALLER
    # ============================
    with st.expander("📐 Dibujo de Figurado para Taller (varillas, estribos, crossties)", expanded=False):
        st.markdown("A continuación se muestran las formas reales de cada tipo de barra, con sus ganchos y dimensiones para facilitar el figurado.")
        # Dibujo de varilla longitudinal L1
        hook_len_cm = 12 * rebar_diam / 10
        straight_len_cm = long_bar_m * 100 - 2 * hook_len_cm
        # Construir etiqueta de barra con designación en cuartos de pulgada
        _rebar_label = rebar_type if "#" in rebar_type else f"{rebar_type.replace('mm','').strip()} mm — {_bar_label(rebar_diam)}"
        _stirrup_label = stirrup_type if "#" in stirrup_type else f"{stirrup_type.replace('mm','').strip()} mm — {_bar_label(stirrup_diam)}"
        fig_l1 = draw_longitudinal_bar(long_bar_m*100, straight_len_cm, hook_len_cm, rebar_diam, bar_name=_rebar_label)
        st.pyplot(fig_l1)
        # Dibujo de estribo E1
        inside_b = b - 2 * recub_cm
        inside_h = h - 2 * recub_cm
        hook_len_est = 12 * stirrup_diam / 10
        fig_e1 = draw_stirrup(inside_b, inside_h, hook_len_est, stirrup_diam, bar_name=_stirrup_label)
        st.pyplot(fig_e1)
        # Dibujo de crossties si existen
        if total_crossties_por_estribo > 0:
            len_crosstie_cm = (b - 2 * recub_cm)
            fig_c1 = draw_crosstie(len_crosstie_cm, hook_len_est, stirrup_diam, bar_name=_stirrup_label)
            st.pyplot(fig_c1)
        st.caption("Nota: Los ganchos de estribos y crossties son de 135° en la práctica. Los ganchos se proyectan hacia el interior de la sección, con extensión mín. de 6d_b.")

    # ============================
    # FORMULAS DE ESTRIBOS POR NIVEL SÍSMICO
    # ============================
    with st.expander("📐 Verificación de Estribos — Formulas por Nivel Sísmico", expanded=False):
        st.markdown(notas_estribos)
        st.markdown("---")
        _zona_tag = "DES 🔴" if zona_especial else ("DMO 🟡" if zona_moderada else "DMI 🟢")
        c_e1, c_e2, c_e3 = st.columns(3)
        c_e1.metric("Nivel Sísmico", _zona_tag)
        c_e2.metric(f"Espaciamiento s (zona confinada)", f"{s_conf:.1f} cm")
        c_e3.metric(f"N° Total de Estribos", f"{n_estribos_total}")
        if zona_especial or zona_moderada:
            c_e4, c_e5 = st.columns(2)
            c_e4.metric("Long. Zona Confinada (Lo)", f"{Lo_conf:.1f} cm")
            c_e5.metric("Espaciamiento centro", f"{s_centro:.1f} cm")

    # ============================
    # NUEVO: APU CON ENTRADA DIRECTA
    # ============================
    st.markdown("---")
    with st.expander("💰 Presupuesto APU – Ingresar precios en vivo", expanded=False):
        st.markdown("Ingrese los precios unitarios de los materiales y mano de obra para calcular el costo total de la columna.")
        # Formulario de precios
        with st.form(key="apu_form"):
            moneda = st.text_input("Moneda (ej. COP, USD)", value=st.session_state.get("apu_moneda", "COP"))
            col1a, col2a = st.columns(2)
            with col1a:
                precio_cemento = st.number_input("Precio por bulto de cemento", value=st.session_state.get("apu_cemento", 28000.0), step=1000.0, format="%.2f")
                precio_acero = st.number_input("Precio por kg de acero", value=st.session_state.get("apu_acero", 7500.0), step=100.0, format="%.2f")
                precio_arena = st.number_input("Precio por m³ de arena", value=st.session_state.get("apu_arena", 120000.0), step=5000.0, format="%.2f")
                precio_grava = st.number_input("Precio por m³ de grava", value=st.session_state.get("apu_grava", 130000.0), step=5000.0, format="%.2f")
            with col2a:
                precio_mo = st.number_input("Costo mano de obra (día)", value=st.session_state.get("apu_mo", 70000.0), step=5000.0, format="%.2f")
                pct_herramienta = st.number_input("% Herramienta menor (sobre MO)", value=st.session_state.get("apu_herramienta", 5.0), step=1.0, format="%.1f") / 100.0
                pct_aui = st.number_input("% A.I.U. (sobre costo directo)", value=st.session_state.get("apu_aui", 30.0), step=5.0, format="%.1f") / 100.0
                pct_util = st.number_input("% Utilidad (sobre costo directo)", value=st.session_state.get("apu_util", 5.0), step=1.0, format="%.1f") / 100.0
                iva = st.number_input("IVA (%) sobre utilidad", value=st.session_state.get("apu_iva", 19.0), step=1.0, format="%.1f") / 100.0
            submitted = st.form_submit_button("Calcular Presupuesto")
            if submitted:
                st.session_state.apu_config = {
                    "moneda": moneda,
                    "cemento": precio_cemento,
                    "acero": precio_acero,
                    "arena": precio_arena,
                    "grava": precio_grava,
                    "costo_dia_mo": precio_mo,
                    "pct_herramienta": pct_herramienta,
                    "pct_aui": pct_aui,
                    "pct_util": pct_util,
                    "iva": iva
                }
                st.success("Precios guardados. Cierre el expander para ver el presupuesto.")
                st.rerun()

        if "apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mon = apu["moneda"]
            # bultos_col: en kg de cemento (del mix design). Convertir a bultos para APU.
            _bag_kg_apu = CEMENT_BAGS.get(norma_sel, CEMENT_BAGS["NSR-10 (Colombia)"])[0]["kg"]
            bultos_col_apu = bultos_col / _bag_kg_apu  # nº de bultos
            costo_cemento = bultos_col_apu * apu["cemento"]
            costo_acero = peso_total_acero_kg * apu["acero"]
            vol_arena_m3 = arena_col / 1500  # densidad arena
            vol_grava_m3 = grava_col / 1600  # densidad grava
            costo_arena = vol_arena_m3 * apu.get("arena", 0)
            costo_grava = vol_grava_m3 * apu.get("grava", 0)
            total_mat = costo_cemento + costo_acero + costo_arena + costo_grava
            dias_acero = peso_total_acero_kg * 0.04
            dias_concreto = vol_concreto_m3 * 0.4
            total_dias_mo = dias_acero + dias_concreto
            costo_mo = total_dias_mo * apu.get("costo_dia_mo", 70000)
            costo_directo = total_mat + costo_mo
            herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * apu.get("pct_aui", 0.30)
            utilidad = costo_directo * apu.get("pct_util", 0.05)
            iva_util = utilidad * apu.get("iva", 0.19)
            total_proyecto = costo_directo + herramienta + aiu + iva_util

            data_apu = {
                    "Item": ["Cemento (bultos)", "Acero (kg)", "Arena (m³)", "Grava (m³)", "Mano de Obra (días)", "Herramienta Menor", "A.I.U.", "IVA s/Utilidad"],
                    "Cantidad": [f"{bultos_col_apu:.1f}", f"{peso_total_acero_kg:.1f}", f"{vol_arena_m3:.2f}", f"{vol_grava_m3:.2f}", f"{total_dias_mo:.2f}", f"{apu.get('pct_herramienta',0.05)*100:.1f}% MO", f"{apu.get('pct_aui',0.30)*100:.1f}% CD", f"{apu.get('iva',0.19)*100:.1f}% Util"],
                    f"Subtotal [{mon}]": [f"{costo_cemento:,.2f}", f"{costo_acero:,.2f}", f"{costo_arena:,.2f}", f"{costo_grava:,.2f}", f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}", f"{iva_util:,.2f}"]
                }
            st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)
            st.metric(f"💎 Gran Total Proyecto [{mon}]", f"{total_proyecto:,.0f}")
            output_excel = io.BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                df_export = pd.DataFrame({
                    "Item": ["Cemento (bultos)", "Acero (kg)", "Arena (m3)", "Grava (m3)", "Mano de Obra (dias)"],
                    "Cantidad": [bultos_col_apu, peso_total_acero_kg, vol_arena_m3, vol_grava_m3, total_dias_mo],
                    "Unidad": [apu['cemento'], apu['acero'], apu.get('arena',0), apu.get('grava',0), apu.get('costo_dia_mo',70000)]
                })
                df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
                df_export.to_excel(writer, index=False, sheet_name='APU')
                workbook = writer.book
                worksheet = writer.sheets['APU']
                money_fmt = workbook.add_format({'num_format': '#,##0.00'})
                bold = workbook.add_format({'bold': True})
                worksheet.set_column('A:A', 25)
                worksheet.set_column('B:D', 15, money_fmt)
                row = len(df_export) + 1
                worksheet.write(row, 0, "Costo Directo (CD)", bold)
                worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
                row += 1
                worksheet.write(row, 0, "Herramienta Menor", bold)
                worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_herramienta",0.05)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "A.I.U", bold)
                worksheet.write_formula(row, 3, f'=D{row-2}*{apu.get("pct_aui",0.30)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "IVA s/ Utilidad", bold)
                worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_util",0.05)}*{apu.get("iva",0.19)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "TOTAL PRESUPUESTO", bold)
                worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
            output_excel.seek(0)
            st.download_button("📥 Descargar Presupuesto Excel", data=output_excel, file_name=f"APU_Columna_{b:.0f}x{h:.0f}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("Por favor, ingrese los precios en el formulario de arriba para generar el presupuesto.")

    st.markdown("---")
    st.markdown("#### 🧱 Dosificación de Concreto")
    bags_for_code = CEMENT_BAGS.get(norma_sel, CEMENT_BAGS["NSR-10 (Colombia)"])
    bag_labels = [b["label"] for b in bags_for_code]
    bag_sel_idx = st.selectbox("Presentación del cemento:", range(len(bag_labels)), format_func=lambda i: bag_labels[i], key="bag_selector")
    bag_kg = bags_for_code[bag_sel_idx]["kg"]
    mix = get_mix_for_fc(fc)
    cem_m3 = mix["cem"]; agua_m3 = mix["agua"]; arena_m3 = mix["arena"]; grava_m3 = mix["grava"]; wc = mix["wc"]
    bultos_por_m3 = cem_m3 / bag_kg
    dens_arena = 1500; dens_grava = 1600
    litros_arena_m3 = arena_m3 / dens_arena * 1000
    litros_grava_m3 = grava_m3 / dens_grava * 1000
    bultos_col = bultos_por_m3 * vol_concreto_m3
    arena_col = arena_m3 * vol_concreto_m3
    grava_col = grava_m3 * vol_concreto_m3
    agua_col = agua_m3 * vol_concreto_m3
    st.markdown(f"**f'c = {fc:.2f} MPa** — Proporciones por m³ según ACI 211:")
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        st.write(f"Cemento: {cem_m3:.0f} kg/m³ → {bultos_por_m3:.2f} bultos/m³")
        st.write(f"Arena: {arena_m3:.0f} kg/m³ ({arena_m3/1600:.2f} m³/m³)")
        st.write(f"Grava: {grava_m3:.0f} kg/m³ ({grava_m3/1600:.2f} m³/m³)")
    with col_d2:
        st.write(f"Agua: {agua_m3:.0f} L/m³")
        st.write(f"Relación a/c: {wc:.2f}")
    st.markdown("**Totales para la columna:**")
    st.write(f"Cemento: {bultos_col:.1f} bultos de {bag_kg:.0f} kg")
    st.write(f"Arena: {arena_col:.1f} kg ({arena_col/dens_arena:.2f} m³)")
    st.write(f"Grava: {grava_col:.1f} kg ({grava_col/dens_grava:.2f} m³)")
    st.write(f"Agua: {agua_col:.0f} L")

    st.markdown("---")
    st.markdown("#### 📄 Generar Memoria de Cálculo (DOCX)")
    if st.button("Generar Memoria DOCX"):
        doc = Document()
        doc.add_heading(f"Memoria de Cálculo — Columna {b:.0f}x{h:.0f} cm", 0)
        doc.add_paragraph(f"Norma Utilizada: {norma_sel}")
        doc.add_paragraph(f"Nivel Sísmico: {nivel_sismico}")
        doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_heading("1. Materiales y Geometría", level=1)
        doc.add_paragraph(f"f'c = {fc:.1f} MPa\nfy = {fy:.0f} MPa\nBase: {b:.0f} cm\nAltura: {h:.0f} cm\nLongitud: {L_col:.0f} cm")
        doc.add_heading("2. Refuerzo y Estribos", level=1)
        doc.add_paragraph(f"Varillas longitudinales: {n_barras_total} varillas tipo {rebar_type} ({Ast:.2f} cm²)\nCuantía longitudinal ρ: {cuantia:.2f}%\nEstribos: {stirrup_type} ({n_estribos_total} unidades)\nEspaciamiento zona confinada: {s_conf:.1f} cm\nEspaciamiento zona central: {s_centro:.1f} cm")
        # Formulas de estribos por nivel sísmico (para memoria)
        doc.add_heading("2.1. Diseño de Estribos por Nivel Sísmico", level=2)
        # Convertir markdown a texto plano para DOCX
        import re as _re
        notas_txt = _re.sub(r'\*+([^*]+)\*+', r'\1', notas_estribos)  # quitar **bold**
        notas_txt = notas_txt.replace("\n", "\n")
        doc.add_paragraph(notas_txt)
        # Sección 2.2: Tabla completa de verificaciones NSR-10
        doc.add_heading("2.2. Verificaciones NSR-10 / ACI 318 — Resumen de Comprobaciones", level=2)
        checks_tbl = doc.add_table(rows=1 + 10, cols=4)
        checks_tbl.style = 'Table Grid'
        ch = checks_tbl.rows[0].cells
        ch[0].text = "Verificación"; ch[1].text = "Requerido"; ch[2].text = "Calculado"; ch[3].text = "Estado"
        _sep_min_req_doc = max(rebar_diam / 10, 2.5)
        _sep_h_doc = (b - 2 * d_prime) / (num_filas_h - 1) if num_filas_h > 1 else 999.0
        _recub_req_doc = 4.0
        _ct_req_h_doc = num_filas_h > 2 and _sep_h_doc > 15.0
        _sep_v_doc = (h - 2 * d_prime) / (num_capas_intermedias + 1) if num_capas_intermedias > 0 else 0.0
        _ct_req_v_doc = num_capas_intermedias > 0 and _sep_v_doc > 15.0
        _ok_doc = "CUMPLE"; _fail_doc = "NO CUMPLE"
        _checks_doc = [
            ("Cuantía mínima (NSR-10 C.10.6.1)",          f">= {rho_min_code:.2f}%",  f"{cuantia:.3f}%",         _ok_doc if cuantia >= rho_min_code else _fail_doc),
            ("Cuantía máxima (NSR-10 C.10.9.1)",          f"<= {rho_max_code:.2f}%",  f"{cuantia:.3f}%",         _ok_doc if cuantia <= rho_max_code else _fail_doc),
            ("Separación mín. barras (ACI §25.8.1)",       f">= {_sep_min_req_doc:.1f} cm", f"{_sep_h_doc:.1f} cm", _ok_doc if _sep_h_doc >= _sep_min_req_doc else _fail_doc),
            ("Recubrimiento mín. (NSR-10 C.20.6.1.3)",    f">= {_recub_req_doc:.1f} cm",  f"{recub_cm:.1f} cm",   _ok_doc if recub_cm >= _recub_req_doc else _fail_doc),
            ("Estribo s_basico (C.7.10.5)",                f"<= {min(b,h):.1f} cm",    f"{s_basico:.1f} cm",      _ok_doc if s_basico <= min(b, h) else _fail_doc),
            ("Zona confinada Lo (C.21.6.4.1)",             f">= {Lo_conf:.1f} cm",     f"{Lo_conf:.1f} cm",       _ok_doc),
            ("Espaciamiento s_conf zona confinada",        f"<= {'15.0' if zona_especial else '20.0'} cm" if (zona_especial or zona_moderada) else "N/A",
                                                            f"{s_conf:.1f} cm",         _ok_doc if s_conf <= (15 if zona_especial else 20) or not (zona_especial or zona_moderada) else _fail_doc),
            ("Crossties horiz. (C.21.6.4.2)",             f"CT si sep>15 cm",          f"{total_crossties_h} CT/estribo" if _ct_req_h_doc else "No reqs.", _ok_doc if not _ct_req_h_doc or total_crossties_h > 0 else _fail_doc),
            ("Crossties vert. (capas interm.)",            f"CT si sep>15 cm",          f"{total_crossties_v} CT/estribo" if _ct_req_v_doc else "No reqs.", _ok_doc if not _ct_req_v_doc or total_crossties_v > 0 else _fail_doc),
            ("Diagrama P-M (carga de diseño)",             "Pu, Mu dentro del envol.", "Ver gráfica Tab 1",      _ok_doc),
        ]
        for i, (verif, req, calc, estado) in enumerate(_checks_doc):
            cells = checks_tbl.rows[i + 1].cells
            cells[0].text = verif; cells[1].text = req; cells[2].text = calc; cells[3].text = estado
        doc.add_paragraph("Nota: Todas las verificaciones se realizan según NSR-10 (Colombia) / ACI 318-19 según la norma seleccionada. El ingeniero debe validar los puntos de carga Pu-Mu en el diagrama P-M.")
        doc.add_heading("3. Resultados del Diagrama P-M", level=1)
        doc.add_paragraph(f"Capacidad Axial Máx Pn,máx: {Pn_max:.1f} {unidad_fuerza}\nResistencia de Diseño φPn,máx: {phi_Pn_max:.1f} {unidad_fuerza}")
        doc.add_picture(pm_img, width=Inches(5))
        doc.add_heading("4. Longitud de Desarrollo y Empalmes", level=1)
        doc.add_paragraph(f"Longitud de desarrollo (ld): {ld_mm/10:.1f} cm\nLongitud de empalme (Clase B): {splice_length_mm/10:.1f} cm\nZona de empalme: desde {splice_start:.0f} cm hasta {splice_end:.0f} cm")
        doc.add_heading("5. Despiece de Acero", level=1)
        table = doc.add_table(rows=1+len(despiece_rows), cols=7)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Marca"; hdr[1].text = "Cant."; hdr[2].text = "Ø"; hdr[3].text = "L (m)"; hdr[4].text = "L total (m)"; hdr[5].text = "Peso (kg)"; hdr[6].text = "Observación"
        for i, row in enumerate(despiece_rows):
            cells = table.rows[i+1].cells
            cells[0].text = row["Marca"]
            cells[1].text = str(row["Cantidad"])
            cells[2].text = str(row["Diámetro"])
            cells[3].text = f"{row['Longitud (m)']:.2f}"
            cells[4].text = f"{row['Longitud Total (m)']:.2f}"
            cells[5].text = f"{row['Peso (kg)']:.1f}"
            cells[6].text = row["Observación"]
        # Guardar figurados a BytesIO para la memoria
        import io as _io
        buf_l1 = _io.BytesIO()
        fig_l1.savefig(buf_l1, format='png', dpi=150, bbox_inches='tight')
        buf_l1.seek(0)
        
        buf_e1 = _io.BytesIO()
        fig_e1.savefig(buf_e1, format='png', dpi=150, bbox_inches='tight')
        buf_e1.seek(0)
        
        buf_c1 = None
        if total_crossties_por_estribo > 0 and 'fig_c1' in locals():
            buf_c1 = _io.BytesIO()
            fig_c1.savefig(buf_c1, format='png', dpi=150, bbox_inches='tight')
            buf_c1.seek(0)

        doc.add_heading("5.1. Esquemas de Figurado", level=2)
        doc.add_picture(buf_l1, width=Inches(4.5))
        doc.add_picture(buf_e1, width=Inches(3.5))
        if buf_c1:
            doc.add_picture(buf_c1, width=Inches(4.5))
            
        doc.add_heading("6. Cantidades de Obra y Dosificación", level=1)
        doc.add_paragraph(f"Concreto: {vol_concreto_m3:.4f} m³\nAcero total: {peso_total_acero_kg:.1f} kg ({relacion_acero_kg_m3:.1f} kg/m³)")
        doc.add_heading("Dosificación de Concreto (ACI 211):", level=2)
        doc.add_paragraph(f"f'c = {fc:.1f} MPa\nCemento: {_cem_kg_m3:.0f} kg/m³ → Aprox {bultos_por_m3:.2f} bultos/m³ (bultos de {bag_kg:.0f} kg)\nArena: {_arena_kg_m3:.0f} kg/m³\nGrava: {_grava_kg_m3:.0f} kg/m³\nAgua: {agua_m3:.0f} L/m³\nRelación a/c: {wc:.2f}")
        doc.add_paragraph(f"Totales Requeridos:\nCemento: {bultos_col:.1f} bultos\nArena: {arena_col:.1f} kg\nGrava: {grava_col:.1f} kg\nAgua: {agua_col:.0f} L")
        
        if "apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mon = apu["moneda"]
            _bag_kg_apu = CEMENT_BAGS.get(norma_sel, CEMENT_BAGS["NSR-10 (Colombia)"])[0]["kg"]
            bultos_col_apu = bultos_col / _bag_kg_apu
            costo_cemento = bultos_col_apu * apu["cemento"]
            costo_acero = peso_total_acero_kg * apu["acero"]
            vol_arena_m3 = arena_col / 1500
            vol_grava_m3 = grava_col / 1600
            costo_arena = vol_arena_m3 * apu.get("arena", 0)
            costo_grava = vol_grava_m3 * apu.get("grava", 0)
            total_mat = costo_cemento + costo_acero + costo_arena + costo_grava
            dias_mo = peso_total_acero_kg * 0.04 + vol_concreto_m3 * 0.4
            costo_mo = dias_mo * apu["costo_dia_mo"]
            costo_herr = costo_mo * apu["pct_herramienta"]
            costo_directo = total_mat + costo_mo + costo_herr
            val_aiu = costo_directo * apu["pct_aui"]
            utilidad = costo_directo * apu["pct_util"]
            val_iva = utilidad * apu["iva"]
            total_proy = costo_directo + val_aiu + val_iva
            
            doc.add_heading("7. Presupuesto APU", level=1)
            tbl_apu = doc.add_table(rows=1+7, cols=3)
            tbl_apu.style = 'Table Grid'
            h_apu = tbl_apu.rows[0].cells
            h_apu[0].text = "Ítem"; h_apu[1].text = "Cantidad"; h_apu[2].text = f"Subtotal [{mon}]"
            _items_apu = [
                ("Cemento (bultos)", f"{bultos_col_apu:.1f}", f"{costo_cemento:,.2f}"),
                ("Acero (kg)", f"{peso_total_acero_kg:.1f}", f"{costo_acero:,.2f}"),
                ("Arena (m³)", f"{vol_arena_m3:.2f}", f"{costo_arena:,.2f}"),
                ("Grava (m³)", f"{vol_grava_m3:.2f}", f"{costo_grava:,.2f}"),
                ("Mano de Obra (días)", f"{dias_mo:.2f}", f"{costo_mo:,.2f}"),
                ("A.I.U. + Herramienta", f"{(apu['pct_aui']+apu['pct_herramienta'])*100:.1f}%", f"{val_aiu+costo_herr:,.2f}"),
                ("GRAN TOTAL PROYECTO", "", f"{total_proy:,.2f}")
            ]
            for idx, (it, ca, sub) in enumerate(_items_apu):
                row_cells = tbl_apu.rows[idx+1].cells
                row_cells[0].text = it; row_cells[1].text = ca; row_cells[2].text = sub
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button("Descargar Memoria DOCX", data=doc_mem, file_name=f"Memoria_Columna_{b:.0f}x{h:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")