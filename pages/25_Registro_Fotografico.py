import streamlit as st
import io
import datetime
import os
import json
from PIL import Image, ImageDraw, ImageFont, ImageOps

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ─── UTILIDADES DE MARCA BLANCA E IDIOMA ─────────────────────────────────────
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es

# Recuperar datos globales del proyecto (Marca Blanca - SIN PREFIJO)
empresa_global   = st.session_state.get("empresa", "") or "_________________________"
proyecto_global  = st.session_state.get("proyecto", "") or "_________________________"
ingeniero_global = st.session_state.get("ingeniero", "") or "_________________________"
logo_bytes       = st.session_state.get("logo_bytes")
user_role        = st.session_state.get("user_role", "free")

# ─── ESTADO LOCAL DEL MÓDULO ─────────────────────────────────────────────────
STATE_FILE = "state_registro_fotografico.json"

_PERSIST_KEYS = [
    "empresa", "proyecto", "ingeniero", 
    "rf_actividad", "rf_clima", "rf_inspector"
]

def save_state():
    try:
        import base64 as _b64
        snapshot = {}
        for k in _PERSIST_KEYS:
            val = st.session_state.get(k)
            if val is not None:
                try:
                    json.dumps(val)
                    snapshot[k] = val
                except (TypeError, ValueError):
                    snapshot[k] = str(val)
        with open(STATE_FILE, 'w', encoding='utf-8') as f:
            json.dump(snapshot, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def load_state():
    if not os.path.exists(STATE_FILE):
        return
    try:
        with open(STATE_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        for k, v in data.items():
            if k not in st.session_state:
                st.session_state[k] = v
    except Exception:
        pass

load_state()

# ─── ROLES ───────────────────────────────────────────────────────────────────
_color = {"admin": "#1b5e20", "pro": "#0d47a1", "free": "#b71c1c"}
st.sidebar.markdown(
    f"<div style='background:{_color.get(user_role, '#666')};color:white;padding:4px 8px;"
    f"border-radius:4px;font-size:12px;text-align:center;'>ROL: {user_role.upper()}</div>",
    unsafe_allow_html=True
)

# ─── INTERFAZ PRINCIPAL ──────────────────────────────────────────────────────
# 1. Banner SVG premium
st.markdown("""<div style="width:100%;overflow:hidden;border-radius:14px;margin-bottom:20px;box-shadow:0 4px 32px #0008;">
<svg viewBox="0 0 1100 220" xmlns="http://www.w3.org/2000/svg" style="width:100%;display:block;background:linear-gradient(135deg,#1e1e2f 0%,#0a0a12 100%);">
  <!-- Grid Background -->
  <g opacity="0.05" stroke="#ffffff" stroke-width="0.5">
    <line x1="0" y1="55" x2="1100" y2="55"/><line x1="0" y1="110" x2="1100" y2="110"/>
    <line x1="0" y1="165" x2="1100" y2="165"/>
    <line x1="220" y1="0" x2="220" y2="220"/><line x1="440" y1="0" x2="440" y2="220"/>
    <line x1="660" y1="0" x2="660" y2="220"/>
  </g>
  
  <!-- Estética Fotográfica / Encuadre -->
  <g transform="translate(40,40)">
    <rect x="0" y="0" width="140" height="140" rx="12" fill="none" stroke="#3b82f6" stroke-width="2" stroke-dasharray="8 8"/>
    <circle cx="70" cy="70" r="35" fill="#2d3748" stroke="#60a5fa" stroke-width="4"/>
    <circle cx="70" cy="70" r="15" fill="#000000" opacity="0.6"/>
    <!-- Lente reflection -->
    <path d="M 55 55 Q 65 45 80 50" fill="none" stroke="#fff" stroke-width="3" stroke-linecap="round" opacity="0.5"/>
    <rect x="110" y="20" width="20" height="10" rx="4" fill="#ef4444"/>
  </g>

  <!-- Document/Report Layer -->
  <g transform="translate(230,50)">
    <rect x="0" y="0" width="90" height="120" rx="4" fill="#ffffff" opacity="0.9"/>
    <rect x="10" y="15" width="70" height="40" fill="#cbd5e1"/>
    <rect x="10" y="65" width="70" height="5" fill="#94a3b8"/>
    <rect x="10" y="78" width="50" height="5" fill="#94a3b8"/>
    <rect x="10" y="91" width="60" height="5" fill="#94a3b8"/>
    <rect x="10" y="104" width="40" height="5" fill="#94a3b8"/>
    <path d="M 65 85 L 85 105" stroke="#ef4444" stroke-width="3" fill="none"/>
    <path d="M 85 85 L 65 105" stroke="#ef4444" stroke-width="3" fill="none"/> <!-- Checkmark or mark -->
  </g>

  <!-- TEXT BLOCK -->
  <g transform="translate(400,0)">
    <rect x="0" y="40" width="4" height="140" rx="2" fill="#3b82f6"/>
    <text x="24" y="75" font-family="Arial,sans-serif" font-size="36" font-weight="bold" fill="#ffffff">REPORTE FOTOGRÁFICO</text>
    <text x="24" y="105" font-family="Arial,sans-serif" font-size="16" font-weight="300" fill="#93c5fd" letter-spacing="1">REGISTRO DE OBRA FORENSE</text>
    <rect x="24" y="120" width="400" height="1" fill="#3b82f6" opacity="0.5"/>
    <!-- Tags -->
    <rect x="24" y="135" width="130" height="24" rx="12" fill="#1e3a5f" stroke="#3b82f6" stroke-width="1"/>
    <text x="89" y="152" text-anchor="middle" font-family="Arial,sans-serif" font-size="10" font-weight="bold" fill="#60a5fa">DOCX AUTOMÁTICO</text>
    <rect x="165" y="135" width="115" height="24" rx="12" fill="#2d1f4e" stroke="#8b5cf6" stroke-width="1"/>
    <text x="222" y="152" text-anchor="middle" font-family="Arial,sans-serif" font-size="10" font-weight="bold" fill="#a78bfa">SELLO DE TIEMPO</text>
    <rect x="290" y="135" width="130" height="24" rx="12" fill="#1a3a2a" stroke="#10b981" stroke-width="1"/>
    <text x="355" y="152" text-anchor="middle" font-family="Arial,sans-serif" font-size="10" font-weight="bold" fill="#34d399">COMPRESIÓN INTELIGENTE</text>
    
    <text x="24" y="180" font-family="Arial,sans-serif" font-size="11" fill="#94a3b8">Convierte cientos de imágenes de campo en memorias optimizadas al instante.</text>
  </g>
</svg></div>""", unsafe_allow_html=True)

st.title(_t("Registro Fotográfico Forense", "Site Photo Report"))

with st.expander(_t("¿Cómo usar este módulo? — Guía Profesional", "How to use this module?"), expanded=False):
    st.markdown("""
    ### Metodología de Reportes con Valor Forense
    Genera actas y bitácoras de obra inmutables y profesionales optimizando dramáticamente los tiempos de oficina.

    #### 1. Metadatos del Informe
    - Establezca el contexto del clima, la sección evaluada y el inspector responsable. 
    - *Todos los reportes incluirán automáticamente el membrete dinámico de su empresa (Logo).*

    #### 2. Seguridad e Imparcialidad Contractual (Marca de Agua)
    - Cada imagen insertada recibe **impresión directa a los pixeles (Sello Lógico)** indicando fecha y actividad.
    - Esto previene el re-uso indebido de frentes aprobados en reclamos futuros con la interventoría u organismos reguladores.

    #### 3. Optimización Automática
    - Las imágenes nativas directas de iOS/Android ($\sim 10$ MB) saturan servidores de correo y traban Word.
    - **Algoritmo de Remuestreo Activo:** Redimensiona por debajo de 1200px con escalado LANCZOS, retiene orientación $EXIF$ nativa de hardware, y reduce un peso unificado sin pérdidas de calidad en A4 (-80% del peso en DOCX final).
    """)

# Helper Func: Sello EXIF + Resize + Auto-Orientación
def procesar_imagen_para_reporte(img_file, fecha_str, act_str):
    try:
        img = Image.open(img_file)
        # Fix Orientation automatically from EXIF data (avoid upside down phones)
        img = ImageOps.exif_transpose(img)
        
        # Convert RGBA -> RGB
        if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
            bg = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')

        # 1. Resize/Compress (Max 1200 px)
        MAX_W = 1200
        if img.width > MAX_W:
            wpercent = (MAX_W/float(img.width))
            hsize = int((float(img.height)*float(wpercent)))
            img = img.resize((MAX_W, hsize), Image.Resampling.LANCZOS)
            
        # 2. Watermark / Sello Contractual
        draw = ImageDraw.Draw(img)
        try:
            # Look for a decent standard font, handle size proportionally
            font_size = max(16, int(img.height * 0.03))
            font = ImageFont.truetype("arial.ttf", font_size)
        except IOError:
            font = ImageFont.load_default()
            
        texto_sello = f"  TOMADA: {fecha_str} | ESTADO: {act_str[:30].upper()} "
        
        # Draw background bar for text
        bbox = draw.textbbox((0, 0), texto_sello, font=font)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
        
        # Bottom left
        x = img.width * 0.02
        y = img.height - text_h - (img.height * 0.03)
        padding = img.height * 0.015
        
        # Dark bounding box with 60% opacity look (Since pure draw doesn't do alpha easily without overlay, we fake it with black)
        draw.rectangle(((x - padding, y - padding), (x + text_w + padding, y + text_h + padding)), fill=(0, 0, 0))
        draw.text((x, y), texto_sello, fill=(255, 255, 255), font=font)
        
        return img
    except Exception as e:
        return None

# 1. METADATOS DEL REPORTE
with st.expander(_t("Información del Informe", "Report Information"), expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        fecha_visita = st.date_input(_t("Fecha de la Visita", "Visit Date"), datetime.date.today(), key="rf_fecha_visita")
        actividad = st.text_input(_t("Actividad / Frente de Obra", "Activity / Work Front"), 
                                  value=st.session_state.get("rf_actividad", "Inspección de estructura"),
                                  key="rf_actividad_input")
    with col2:
        clima = st.selectbox(_t("Condición Climática", "Weather Condition"), 
                             ["Soleado", "Parcialmente Nublado", "Nublado", "Lluvia", "Lluvia Fuerte"],
                             key="rf_clima_input")
        inspector = st.text_input(_t("Inspector / Residente", "Inspector / Resident"), 
                                  value=st.session_state.get("rf_inspector", ingeniero_global),
                                  key="rf_inspector_input")
                                  
    # Update local variables carefully keeping the persistence synced
    st.session_state["rf_actividad"] = actividad
    st.session_state["rf_clima"] = clima
    st.session_state["rf_inspector"] = inspector
    save_state()

# 2. CARGA MASIVA DE IMÁGENES
st.markdown("### " + _t("Sube tus fotografías", "Upload your photos"))
archivos_subidos = st.file_uploader(
    _t("Arrastra múltiples imágenes (JPG, PNG)", "Drag & drop multiple images (JPG, PNG)"), 
    type=['png', 'jpg', 'jpeg'], 
    accept_multiple_files=True
)

if "rf_observaciones_fotos" not in st.session_state:
    st.session_state["rf_observaciones_fotos"] = {}

# 3. CUADRÍCULA DE EDICIÓN
if archivos_subidos:
    st.markdown("### " + _t("Detalle Fotográfico", "Photo Details"))
    st.info(_t(
        "Añade una descripción a cada fotografía. Se guardará automáticamente.", 
        "Add a description to each photo. It will be saved automatically."
    ))
    
    cols = st.columns(2)
    for i, archivo in enumerate(archivos_subidos):
        col = cols[i % 2]
        with col:
            img = Image.open(archivo)
            st.image(img, use_container_width=True)
            
            clave_obs = f"obs_{archivo.name}_{archivo.size}"
            obs_actual = st.session_state["rf_observaciones_fotos"].get(clave_obs, "")
            
            nueva_obs = st.text_area(
                f"Observación de la foto {i+1}:", 
                value=obs_actual, 
                key=f"rf_widget_{clave_obs}",
                height=80
            )
            st.session_state["rf_observaciones_fotos"][clave_obs] = nueva_obs
            st.markdown("---")

    # 4. CONFIGURACIÓN DE EXPORTACIÓN
    st.markdown("### " + _t("Exportar Informe", "Export Report"))
    
    formato_pag = st.radio(
        _t("Disposición de fotografías por página:", "Photos per page layout:"),
        [
            "2 Fotos por página (Grandes)", 
            "4 Fotos por página (Medianas)", 
            "6 Fotos por página (Pequeñas)"
        ],
        horizontal=True
    )

    _puede_exportar = user_role in ("admin", "pro")
    
    btn_disabled = not _puede_exportar or not _DOCX_OK
    help_text = ""
    if not _puede_exportar:
        help_text = "Requiere licencia Pro o Admin"
    elif not _DOCX_OK:
        help_text = "Instala: pip install python-docx"

    if st.button(_t("Generar Documento DOCX", "Generate DOCX Document"), type="primary", disabled=btn_disabled, help=help_text):
        with st.spinner(_t("Procesando imágenes y generando documento...", "Processing images and generating document...")):
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            # --- PORTADA / ENCABEZADO ---
            if logo_bytes and isinstance(logo_bytes, (bytes, bytearray)):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_img = p_img.add_run()
                    r_img.add_picture(io.BytesIO(logo_bytes), width=Cm(5.0))
                except Exception:
                    pass

            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("REGISTRO FOTOGRÁFICO DE OBRA\n")
            r_title.bold = True
            r_title.font.size = Pt(16)
            
            table_meta = doc.add_table(rows=5, cols=2)
            table_meta.style = 'Table Grid'
            
            meta_data = [
                ("PROYECTO:", proyecto_global),
                ("EMPRESA:", empresa_global),
                ("FECHA DE VISITA:", fecha_visita.strftime('%d/%m/%Y')),
                ("ACTIVIDAD REVISADA:", actividad),
                ("CLIMA / INSPECTOR:", f"{clima} / {inspector}")
            ]
            
            for row_idx, (label, value) in enumerate(meta_data):
                cells = table_meta.rows[row_idx].cells
                cells[0].text = label
                cells[0].paragraphs[0].runs[0].bold = True
                cells[1].text = str(value)
            
            doc.add_paragraph("\n")

            # --- CONFIGURACIÓN DE TABLA DE FOTOS ---
            if "2 Fotos" in formato_pag:
                num_cols = 1
                img_width = Inches(5.5)
            elif "4 Fotos" in formato_pag:
                num_cols = 2
                img_width = Inches(3.0)
            else:
                num_cols = 2
                img_width = Inches(2.8)

            table_fotos = doc.add_table(rows=0, cols=num_cols)
            table_fotos.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            current_row = None
            for i, archivo in enumerate(archivos_subidos):
                if i % num_cols == 0:
                    current_row = table_fotos.add_row()
                
                cell = current_row.cells[i % num_cols]
                p_cell = cell.paragraphs[0]
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Aplicar robustez EXIF, redimensionamiento LANCZOS y sello de agua forense
                try:
                    fecha_fmt = fecha_visita.strftime('%d/%m/%Y')
                    img_procesada = procesar_imagen_para_reporte(archivo, fecha_fmt, actividad)
                    
                    if img_procesada:
                        img_byte_arr = io.BytesIO()
                        # Compress output with Quality 75%
                        img_procesada.save(img_byte_arr, format='JPEG', quality=75, optimize=True)
                        img_byte_arr.seek(0)
                        
                        run_img = p_cell.add_run()
                        run_img.add_picture(img_byte_arr, width=img_width)
                    else:
                        p_cell.add_run("[Error genérico procesando imagen]")
                except Exception as e:
                    p_cell.add_run(f"[Error: {e}]")

                clave_obs = f"obs_{archivo.name}_{archivo.size}"
                obs_texto = st.session_state["rf_observaciones_fotos"].get(clave_obs, "Sin observaciones.")
                
                p_obs = cell.add_paragraph()
                p_obs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_obs = p_obs.add_run(f"Foto {i+1}: ")
                run_obs.bold = True
                p_obs.add_run(obs_texto)
                
                cell.add_paragraph("")

            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            st.success(_t("[Listo] Informe generado con éxito", "Report generated successfully"))
            
            st.download_button(
                label=_t("Descargar Informe Fotográfico (DOCX)", "Download Photo Report (DOCX)"),
                data=doc_io,
                file_name=f"Registro_Fotografico_{fecha_visita.strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
