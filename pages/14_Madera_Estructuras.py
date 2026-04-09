import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Estructuras en Madera", "Timber Structures"), layout="wide")
st.title(_t("Diseño de Estructuras en Madera", "Timber Structure Design"))
st.markdown(_t("Cálculo y diseño de elementos de madera estructural (Vigas, Columnas, Uniones) y cuantificación comercial en Pies Madereros. Basado en metodologías de Esfuerzos Admisibles (NDS / NSR-10).", 
               "Design of structural timber elements (Beams, Columns, Connections) and commercial quantification in Board Feet. Based on Allowable Stress Design (NDS / NSR-10)."))

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    <br><br>
    <i>⚠ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CONFIGURACIÓN GLOBAL
# ─────────────────────────────────────────────
st.sidebar.header(_t("⚙ Configuración Global", "⚙ Global Settings"))
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagpedia.net/data/flags/mini/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Normativa Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)

st.sidebar.header(_t(" Propiedades Mecánicas (MPa)", " Mechanical Properties (MPa)"))
grupo_madera_opts = [_t("Grupo A (Alta Dureza)", "Group A (High Density)"),
                     _t("Grupo B (Media)", "Group B (Medium)"),
                     _t("Grupo C (Madera Suave)", "Group C (Softwood)"),
                     _t("Personalizado / Custom", "Custom")]
grupo_madera = st.sidebar.selectbox(_t("Seleccionar Grupo de Madera:", "Select Wood Group:"), 
                                    grupo_madera_opts,
                                    index=grupo_madera_opts.index(st.session_state.get("m_grupo", grupo_madera_opts[1])),
                                    key="m_grupo")

if grupo_madera == _t("Grupo A (Alta Dureza)", "Group A (High Density)"):
    E_min = 9500.0; Fb = 21.0; Fv = 1.5; Fc = 14.5; G_sp = 0.55
elif grupo_madera == _t("Grupo B (Media)", "Group B (Medium)"):
    E_min = 7500.0; Fb = 15.0; Fv = 1.2; Fc = 11.0; G_sp = 0.45
elif grupo_madera == _t("Grupo C (Madera Suave)", "Group C (Softwood)"):
    E_min = 5500.0; Fb = 10.0; Fv = 0.8; Fc = 8.0; G_sp = 0.40
else:
    E_min = 8000.0; Fb = 15.0; Fv = 1.0; Fc = 10.0; G_sp = 0.50

val_E = st.sidebar.number_input(_t("Módulo Elasticidad E_min [MPa]", "Modulus of Elasticity E_min [MPa]"), 1000.0, 20000.0, st.session_state.get("m_E", E_min), 100.0, disabled=(grupo_madera!=_t("Personalizado / Custom", "Custom")), key="m_E")
val_Fb = st.sidebar.number_input(_t("Esfuerzo Flexión Adm. Fb [MPa]", "Allowable Bending Fb [MPa]"), 1.0, 50.0, st.session_state.get("m_Fb", Fb), 1.0, disabled=(grupo_madera!=_t("Personalizado / Custom", "Custom")), key="m_Fb")
val_Fv = st.sidebar.number_input(_t("Esfuerzo Cortante Adm. Fv [MPa]", "Allowable Shear Fv [MPa]"), 0.1, 5.0, st.session_state.get("m_Fv", Fv), 0.1, disabled=(grupo_madera!=_t("Personalizado / Custom", "Custom")), key="m_Fv")
val_Fc = st.sidebar.number_input(_t("Compresión Paralela Fc [MPa]", "Parallel Compression Fc [MPa]"), 1.0, 30.0, st.session_state.get("m_Fc", Fc), 1.0, disabled=(grupo_madera!=_t("Personalizado / Custom", "Custom")), key="m_Fc")
val_G = st.sidebar.number_input(_t("Gravedad Específica G", "Specific Gravity G"), 0.1, 1.2, st.session_state.get("m_G", G_sp), 0.05, disabled=(grupo_madera!=_t("Personalizado / Custom", "Custom")), key="m_G")

# ─────────────────────────────────────────────
# T1: CALCULADORA DE PIES MADEREROS
# ─────────────────────────────────────────────
with st.expander(_t(" 1. Calculadora de Pies² de Madera (Board Feet)", " 1. Board Feet Calculator"), expanded=False):
    st.info(_t("Ingresa el espesor y ancho en pulgadas, y el largo de la pieza (en pies o metros).", "Enter thickness/width in inches, and length."))
    
    col1, col2, col3, col4 = st.columns(4)
    with col1: qty_bf = st.number_input(_t("Cantidad de piezas", "Quantity of pieces"), 1, 10000, st.session_state.get("m_qty", 10), key="m_qty")
    with col2: thick_in = st.number_input(_t("Espesor (pulgadas) [in]", "Thickness [in]"), 0.5, 20.0, st.session_state.get("m_thick", 2.0), 0.5, key="m_thick")
    with col3: width_in = st.number_input(_t("Ancho (pulgadas) [in]", "Width [in]"), 1.0, 24.0, st.session_state.get("m_width", 4.0), 0.5, key="m_width")
    with col4:
        unit_len_opts = [_t("Metros [m]", "Meters [m]"), _t("Pies [ft]", "Feet [ft]")]
        unit_len = st.radio(_t("Unidad Longitud:", "Length Unit:"), unit_len_opts, 
                            index=unit_len_opts.index(st.session_state.get("m_unit_len", unit_len_opts[0])),
                            horizontal=True, key="m_unit_len")
        length_val = st.number_input(_t("Largo de la pieza", "Piece Length"), 0.5, 50.0, st.session_state.get("m_length", 3.0 if unit_len==_t("Metros [m]", "Meters [m]") else 10.0), 0.5, key="m_length")
        
    length_ft = length_val if _t("Pies", "Feet") in unit_len else (length_val * 3.28084)
    board_feet_per_piece = (thick_in * width_in * length_ft) / 12.0
    total_board_feet = qty_bf * board_feet_per_piece
    
    mon = st.session_state.apu_config["moneda"] if "apu_config" in st.session_state else "COP$"
    precio_pt = st.number_input(f"{_t('Precio por Pie Tabular', 'Price per Board Foot')} [{mon}/pt]", value=st.session_state.get("m_price_pt", 2000.0 if mon=="COP$" else 2.50), key="m_price_pt")
    
    c_res, c_plot = st.columns([1,2])
    with c_res:
        st.metric(_t("Volumen por Pieza", "Volume per Piece"), f"{board_feet_per_piece:.2f} pt")
        st.metric(_t("Volumen Total", "Total Volume"), f"{total_board_feet:.2f} pt")
        st.metric(_t("Costo Directo del Lote", "Direct Batch Cost"), f"{mon} {total_board_feet*precio_pt:,.2f}")
    with c_plot:
        fig, ax = plt.subplots(figsize=(6, 2))
        ax.add_patch(patches.Rectangle((0, 0), length_ft, thick_in/12.0, edgecolor='saddlebrown', facecolor='peru', lw=2))
        ax.text(length_ft/2, (thick_in/12.0)/2, f"{_t('Largo:', 'Length:')} {length_ft:.1f}'", ha='center', va='center', color='white', fontweight='bold')
        ax.text(-0.5, (thick_in/12.0)/2, f"{_t('Espesor:', 'Thickness:')}\n{thick_in}\"", ha='right', va='center', color='saddlebrown')
        ax.set_title(_t(f"Sección de Madera: {thick_in}\" x {width_in}\" x {length_ft:.1f}'", f"Lumber Section: {thick_in}\" x {width_in}\" x {length_ft:.1f}'"))
        ax.set_xlim(-1, length_ft + 1)
        ax.set_ylim(-1, (thick_in/12.0) + 1)
        ax.axis('off')
        st.pyplot(fig)

# ─────────────────────────────────────────────
# T2: DISEÑO DE VIGAS DE MADERA
# ─────────────────────────────────────────────
with st.expander(_t(" 2. Diseño de Vigas de Madera (Flexión y Cortante)", " 2. Timber Beam Design (Flexure and Shear)"), expanded=False):
    v1, v2, v3 = st.columns(3)
    with v1:
        span_L = st.number_input(_t("Luz Viga L [m]", "Beam Span L [m]"), 1.0, 15.0, st.session_state.get("m_v_L", 4.0), 0.1, key="m_v_L")
        w_muerta = st.number_input("WD [kN/m]", 0.0, 50.0, st.session_state.get("m_v_wd", 2.0), 0.5, key="m_v_wd")
        w_viva = st.number_input("WL [kN/m]", 0.0, 50.0, st.session_state.get("m_v_wl", 3.0), 0.5, key="m_v_wl")
        W_total = w_muerta + w_viva # kN/m
    with v2:
        b_beam = st.number_input("Base (b) [mm]", 20.0, 400.0, st.session_state.get("m_v_b", 100.0), 10.0, key="m_v_b")
        h_beam = st.number_input("Altura (h) [mm]", 50.0, 1000.0, st.session_state.get("m_v_h", 250.0), 10.0, key="m_v_h")
        Area_beam = (b_beam * h_beam) / 1e6; Sx_beam = (b_beam * h_beam**2 / 6.0) / 1e9; Ix_beam = (b_beam * h_beam**3 / 12.0) / 1e12
    with v3:
        def_opts = ["L/360", "L/240", "L/180"]
        def_limit = st.selectbox(_t("Límite de Deflexión", "Deflection Limit"), def_opts, 
                                 index=def_opts.index(st.session_state.get("m_v_def", "L/360")),
                                 key="m_v_def")
        Def_max_adm = (span_L * 1000.0) / float(def_limit.split("/")[1])
        
    M_max = (W_total * span_L**2) / 8.0; V_max = (W_total * span_L) / 2.0
    fb_act = (M_max / Sx_beam) / 1000.0; fv_act = (1.5 * V_max / Area_beam) / 1000.0
    Def_act = (5.0 * W_total * (span_L*1000.0)**4) / (384.0 * val_E * (Ix_beam*1e12))
    
    ok_flex = fb_act <= val_Fb; ok_shear = fv_act <= val_Fv; ok_def = Def_act <= Def_max_adm
    
    r1, r2, r3, r4 = st.columns([1,1,1,2])
    r1.metric("f_b Actuante", f"{fb_act:.2f} MPa")
    if ok_flex: r1.success(f" OK ({val_Fb})") 
    else: r1.error(" No Aprobado")
    
    r2.metric("f_v Actuante", f"{fv_act:.2f} MPa")
    if ok_shear: r2.success(f" OK ({val_Fv})") 
    else: r2.error(" No Aprobado")
    
    r3.metric("Δ Actuante", f"{Def_act:.1f} mm")
    if ok_def: r3.success(f" OK ({Def_max_adm:.1f})") 
    else: r3.error(" No Aprobado")
    
    with r4:
        fig2, ax2 = plt.subplots(figsize=(6,2))
        ax2.plot([0, span_L], [0, 0], color='brown', lw=max(1, h_beam/20))
        ax2.plot(0, 0, '^', markersize=15, color='gray')
        ax2.plot(span_L, 0, '^', markersize=15, color='gray')
        num_arr = 10
        x_arr = np.linspace(0, span_L, num_arr)
        for xa in x_arr:
            ax2.arrow(xa, 1.0, 0, -0.6, head_width=0.1, head_length=0.2, fc='red', ec='red')
        ax2.plot([0, span_L], [1.0, 1.0], color='red', lw=2)
        ax2.text(span_L/2, 1.3, f"w = {W_total} kN/m", ha='center', color='red')
        ax2.text(span_L/2, -1.0, f"L = {span_L} m\nSección: {b_beam}x{h_beam} mm", ha='center')
        ax2.set_xlim(-0.5, span_L+0.5)
        ax2.set_ylim(-1.5, 2.0)
        ax2.axis('off')
        st.pyplot(fig2)

# ─────────────────────────────────────────────
# T3: COLUMNAS DE MADERA A COMPRESIÓN PURA
# ─────────────────────────────────────────────
with st.expander(_t(" 3. Diseño de Columnas de Madera (Compresión)", " 3. Timber Column Design (Compression)"), expanded=False):
    c1, c2, c3 = st.columns(3)
    with c1:
        P_ax = st.number_input("Carga Axial Actuante P [kN]", 5.0, 500.0, st.session_state.get("m_c_P", 20.0), 5.0, key="m_c_P")
        KL_col = st.number_input(_t("kL Efectiva [m]", "Effective kL [m]"), 0.5, 10.0, st.session_state.get("m_c_kL", 3.0), 0.1, key="m_c_kL")
    with c2:
        b_col = st.number_input("Base (b) [mm]", 50.0, 400.0, st.session_state.get("m_c_b", 150.0), 10.0, key="m_c_b")
        h_col = st.number_input("Altura (h) [mm]", 50.0, 400.0, st.session_state.get("m_c_h", 150.0), 10.0, key="m_c_h")
        d_min = min(b_col, h_col); Area_col = (b_col * h_col) / 1e6
    with c3:
        esbeltez_l = (KL_col * 1000.0) / d_min
        st.markdown(rf"**$\lambda = kL/d$:** {esbeltez_l:.2f}")
    
    if esbeltez_l > 50:
        st.error(_t(" Columna muy esbelta (λ > 50). ¡Aumentar sección!", " Column too slender (λ > 50)!"))
    else:
        F_cE = (0.822 * val_E) / (esbeltez_l**2)
        ratio_alpha = F_cE / val_Fc; c_factor = 0.8
        termino_rad = (1 + ratio_alpha)/(2*c_factor)
        C_p = termino_rad - math.sqrt((termino_rad**2) - (ratio_alpha/c_factor))
        F_c_prime = val_Fc * C_p
        P_admisible = (F_c_prime * 1000.0 * Area_col) # kN
        
        tc1, tc2, tc3 = st.columns([1,1,2])
        tc1.metric("P Actuante", f"{P_ax:.1f} kN")
        tc2.metric("P Admisible", f"{P_admisible:.1f} kN")
        if P_ax <= P_admisible: tc2.success(" OK") 
        else: tc2.error(" No Aprobado")
        
        with tc3:
            fig3, ax3 = plt.subplots(figsize=(2,4))
            ax3.plot([0, 0], [0, KL_col], color='saddlebrown', lw=max(5, b_col/20))
            ax3.arrow(0, KL_col+0.5, 0, -0.3, head_width=0.3, head_length=0.1, fc='blue', ec='blue')
            ax3.text(0, KL_col+0.6, f"P={P_ax}kN", ha='center', color='blue')
            y_curve = np.linspace(0, KL_col, 50)
            x_curve = 0.5 * np.sin(np.pi * y_curve / KL_col)
            ax3.plot(x_curve, y_curve, 'r--', lw=1, alpha=0.6)
            ax3.set_xlim(-1, 1)
            ax3.set_ylim(-0.5, KL_col+1)
            ax3.axis('off')
            st.pyplot(fig3)

# ─────────────────────────────────────────────
# T4: UNIONES CON CLAVOS (NDS)
# ─────────────────────────────────────────────
with st.expander(_t(" 4. Resistencia de Uniones con Clavos (Corte Lateral)", " 4. Nail Connection (Lateral Shear)"), expanded=False):
    st.info(_t("Según NDS 2018, la resistencia de diseño lateral para clavos se ajusta por factores de duración de carga, humedad, temperatura y grupo.", "Based on NDS 2018, design lateral resistance for nails is adjusted by load duration, moisture, temperature and group factors."))
    uc1, uc2, uc3 = st.columns([1,1,2])
    with uc1:
        clavo_opts = [2.5, 3.0, 3.4, 3.8, 4.2, 5.0, 6.0]
        diam_clavo = st.selectbox("Calibre Clavo D [mm]", clavo_opts, 
                                   index=clavo_opts.index(st.session_state.get("m_u_diam", 3.4)),
                                   key="m_u_diam")
    with uc2:
        penetracion_p = st.number_input("Penetración (p) [mm]", 10.0, 150.0, st.session_state.get("m_u_p", 40.0), 5.0, key="m_u_p")
        factor_cd = st.selectbox(_t("Factor de duración de carga (Cd)", "Load duration factor (Cd)"), 
                                 [_t("Carga permanente (Cd=0.9)", "Permanent (Cd=0.9)"),
                                  _t("Carga viva (Cd=1.0)", "Live load (Cd=1.0)"),
                                  _t("Carga de nieve (Cd=1.15)", "Snow load (Cd=1.15)"),
                                  _t("Carga sísmica/viento (Cd=1.6)", "Seismic/wind (Cd=1.6)")],
                                 index=1, key="m_u_cd")
        cd_val = float(factor_cd.split("=")[1].replace(")", "").strip())
    with uc3:
        # Factores adicionales
        factor_cm = st.selectbox(_t("Factor de humedad (Cm)", "Moisture factor (Cm)"), 
                                 [_t("Madera seca (Cm=1.0)", "Dry wood (Cm=1.0)"),
                                  _t("Madera húmeda (Cm=0.7)", "Wet wood (Cm=0.7)")],
                                 index=0, key="m_u_cm")
        cm_val = 1.0 if "seca" in factor_cm else 0.7
        factor_ct = st.selectbox(_t("Factor de temperatura (Ct)", "Temperature factor (Ct)"), 
                                 [_t("Normal (Ct=1.0)", "Normal (Ct=1.0)"),
                                  _t("Altas temperaturas (Ct=0.8)", "High temp (Ct=0.8)")],
                                 index=0, key="m_u_ct")
        ct_val = 1.0 if "Normal" in factor_ct else 0.8
        
    Z_lat = 16.6 * (val_G**1.8) * (diam_clavo**1.5)  # NDS fórmula
    D_req_full = 10.0 * diam_clavo; D_req_min = 6.0 * diam_clavo
    
    if penetracion_p < D_req_min:
        st.error(_t(" Penetración < 6D. No cumple requerimiento mínimo.", " Penetration < 6D. Minimum requirement not met."))
        Z_adm = 0.0
    else:
        Cd_factor = min(1.0, penetracion_p / D_req_full)
        Z_ajustado = Z_lat * cd_val * cm_val * ct_val * Cd_factor
        Z_adm = (Z_ajustado) / 10.0  # convertir a kgf (1 kgf ≈ 9.81 N)
    
    st.metric(_t("Corte Admisible (Z')", "Allowable Shear (Z')"), f"{Z_adm:.2f} kgf/clavo" if Z_adm>0 else "N/A")
    
    # Esquema gráfico
    fig4, ax4 = plt.subplots(figsize=(6,2))
    ax4.add_patch(patches.Rectangle((0, 0), 2, 1, facecolor='peru', edgecolor='brown'))
    ax4.add_patch(patches.Rectangle((2, 0), 2, 1, facecolor='burlywood', edgecolor='brown'))
    # Clavo
    ax4.plot([1.8, 2.0+(penetracion_p/40.0)], [0.5, 0.5], 'k-', lw=3)
    ax4.text(2.0, 0.7, f"p={penetracion_p}mm", ha='center')
    ax4.arrow(1.0, 0.5, -0.4, 0, head_width=0.1, fc='blue', ec='blue')
    ax4.arrow(3.0, 0.5, 0.4, 0, head_width=0.1, fc='blue', ec='blue')
    ax4.text(0.5, 0.2, "Z'", color='blue')
    ax4.set_xlim(0, 4); ax4.set_ylim(-0.5, 1.5); ax4.axis('off')
    st.pyplot(fig4)

# ─────────────────────────────────────────────
# EXPORTACIÓN INTEGRAL (DXF, DOCX, APU)
# ─────────────────────────────────────────────
st.markdown("---")
st.subheader(_t(" Exportación Integral", " Comprehensive Export"))

# Preparar datos comunes para despiece y APU
pt_viga = (b_beam/25.4) * (h_beam/25.4) * (span_L * 3.28084) / 12.0
pt_col  = (b_col/25.4) * (h_col/25.4) * (KL_col * 3.28084) / 12.0
total_pt = pt_viga + pt_col

# Tabla de despiece
despiece = pd.DataFrame([
    {"Elemento": _t("Viga", "Beam"), "Sección (mm)": f"{b_beam:.0f}x{h_beam:.0f}", "Longitud (m)": f"{span_L:.2f}", "Pies Tabulares": f"{pt_viga:.2f}"},
    {"Elemento": _t("Columna", "Column"), "Sección (mm)": f"{b_col:.0f}x{h_col:.0f}", "Longitud (m)": f"{KL_col:.2f}", "Pies Tabulares": f"{pt_col:.2f}"},
    {"Elemento": _t("TOTAL", "TOTAL"), "Sección (mm)": "", "Longitud (m)": "", "Pies Tabulares": f"{total_pt:.2f}"}
])

tab_despiece, tab_3d, tab_dxf, tab_doc, tab_apu = st.tabs([
    " " + _t("Despiece de Madera", "Timber Cutting List"),
    " " + _t("Visualización 3D", "3D Visualization"),
    " " + _t("Planos DXF", "DXF Drawings"),
    " " + _t("Memoria DOCX", "DOCX Report"),
    " " + _t("Presupuesto APU", "APU Budget")
])

with tab_despiece:
    st.markdown(_t("#### Cantidades de Madera (Pies Tabulares)", "#### Timber Quantities (Board Feet)"))
    st.dataframe(despiece, use_container_width=True, hide_index=True)
    # Gráfico de barras
    fig_pie, ax_pie = plt.subplots(figsize=(6, 3))
    ax_pie.bar([_t("Viga", "Beam"), _t("Columna", "Column")], [pt_viga, pt_col], color=['peru', 'saddlebrown'])
    ax_pie.set_ylabel(_t("Pies Tabulares", "Board Feet"))
    ax_pie.set_title(_t("Volumen de madera por elemento", "Timber volume per element"))
    st.pyplot(fig_pie)

with tab_3d:
    col3d1, col3d2 = st.columns(2)
    with col3d1:
        st.write("#### Viga de Madera")
        fig3d_v = go.Figure()
        X_v = span_L
        Y_v = b_beam / 1000.0
        Z_v = h_beam / 1000.0
        x_wv = [0, X_v, X_v, 0, 0, X_v, X_v, 0]
        y_wv = [-Y_v/2, -Y_v/2, Y_v/2, Y_v/2, -Y_v/2, -Y_v/2, Y_v/2, Y_v/2]
        z_wv = [-Z_v/2, -Z_v/2, -Z_v/2, -Z_v/2, Z_v/2, Z_v/2, Z_v/2, Z_v/2]
        fig3d_v.add_trace(go.Mesh3d(x=x_wv, y=y_wv, z=z_wv, alphahull=0, opacity=0.9, color='peru', name='Viga'))
        # Añadir líneas de textura (vetas)
        for i in range(1, 5):
            x = i * X_v / 5
            fig3d_v.add_trace(go.Scatter3d(x=[x, x], y=[-Y_v/2, Y_v/2], z=[-Z_v/2, -Z_v/2], mode='lines', line=dict(color='brown', width=1), showlegend=False))
            fig3d_v.add_trace(go.Scatter3d(x=[x, x], y=[-Y_v/2, Y_v/2], z=[Z_v/2, Z_v/2], mode='lines', line=dict(color='brown', width=1), showlegend=False))
        fig3d_v.update_layout(scene=dict(aspectmode='data', xaxis_title='L (m)', yaxis_title='b (m)', zaxis_title='h (m)'), 
                              margin=dict(l=0, r=0, b=0, t=0), height=350, showlegend=False, dragmode='turntable',
                              paper_bgcolor='#1a1a2e', scene_bgcolor='#1a1a2e')
        st.plotly_chart(fig3d_v, use_container_width=True)
        
    with col3d2:
        st.write("#### Columna de Madera")
        fig3d_c = go.Figure()
        Z_c = KL_col
        X_c = b_col / 1000.0
        Y_c = h_col / 1000.0
        x_wc = [-X_c/2, X_c/2, X_c/2, -X_c/2, -X_c/2, X_c/2, X_c/2, -X_c/2]
        y_wc = [-Y_c/2, -Y_c/2, Y_c/2, Y_c/2, -Y_c/2, -Y_c/2, Y_c/2, Y_c/2]
        z_wc = [0, 0, 0, 0, Z_c, Z_c, Z_c, Z_c]
        fig3d_c.add_trace(go.Mesh3d(x=x_wc, y=y_wc, z=z_wc, alphahull=0, opacity=0.9, color='saddlebrown', name='Columna'))
        # Vetas
        for i in range(1, 5):
            z = i * Z_c / 5
            fig3d_c.add_trace(go.Scatter3d(x=[-X_c/2, X_c/2], y=[-Y_c/2, -Y_c/2], z=[z, z], mode='lines', line=dict(color='brown', width=1), showlegend=False))
            fig3d_c.add_trace(go.Scatter3d(x=[-X_c/2, X_c/2], y=[Y_c/2, Y_c/2], z=[z, z], mode='lines', line=dict(color='brown', width=1), showlegend=False))
        fig3d_c.update_layout(scene=dict(aspectmode='data', xaxis_title='b (m)', yaxis_title='h (m)', zaxis_title='L (m)'), 
                              margin=dict(l=0, r=0, b=0, t=0), height=350, showlegend=False, dragmode='turntable',
                              paper_bgcolor='#1a1a2e', scene_bgcolor='#1a1a2e')
        st.plotly_chart(fig3d_c, use_container_width=True)

with tab_dxf:
    try:
        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                 dxf_dim_horiz, dxf_dim_vert, dxf_rotulo, dxf_rotulo_campos)
        _USE_H_mad = True
    except ImportError:
        _USE_H_mad = False
    col_dxf1, col_dxf2 = st.columns(2)

    # ── Viga DXF (elevación) ──
    doc_v = ezdxf.new('R2010'); doc_v.units = ezdxf.units.M
    if _USE_H_mad:
        dxf_setup(doc_v, 50); dxf_add_layers(doc_v)
    msp_v = doc_v.modelspace()
    for lay, col_c in [('MADERA',3), ('COTAS',2), ('TEXTO',1)]:
        if lay not in doc_v.layers: doc_v.layers.add(lay, color=col_c)
    msp_v.add_lwpolyline([(0,0), (span_L,0), (span_L, h_beam/1000), (0, h_beam/1000), (0,0)],
                         close=True, dxfattribs={'layer':'MADERA'})
    if _USE_H_mad:
        TH = 0.025*50
        dxf_dim_horiz(msp_v, 0, span_L, -0.2, f"L = {span_L:.2f} m", 50)
        dxf_dim_vert(msp_v, span_L+0.15, 0, h_beam/1000, f"h = {h_beam:.0f} mm", 50)
        dxf_text(msp_v, span_L/2, h_beam/1000+0.15, "ELEVACION VIGA MADERA", "EJES", h=TH*1.1, ha="center")
        _cam_v = dxf_rotulo_campos(f"Viga Madera {b_beam:.0f}x{h_beam:.0f}mm L={span_L:.1f}m", norma_sel, "001")
        dxf_rotulo(msp_v, _cam_v, 0, -4.5, rot_w=9, rot_h=3, escala=50)
    else:
        msp_v.add_text(f"L = {span_L:.2f} m", dxfattribs={'layer':'TEXTO','height':0.05,'insert':(span_L/2, -0.1)})
        msp_v.add_text(f"h = {h_beam:.0f} mm", dxfattribs={'layer':'TEXTO','height':0.05,'insert':(span_L+0.1, h_beam/2000)})
    _out_v = io.StringIO()
    doc_v.write(_out_v)
    col_dxf1.download_button(_t(" DXF Viga", " Beam DXF"),
                             data=_out_v.getvalue().encode("utf-8"),
                             file_name=f"Viga_Madera_{b_beam:.0f}x{h_beam:.0f}.dxf", mime="application/dxf")

    # ── Columna DXF (elevación) ──
    doc_c = ezdxf.new('R2010'); doc_c.units = ezdxf.units.M
    if _USE_H_mad:
        dxf_setup(doc_c, 50); dxf_add_layers(doc_c)
    msp_c = doc_c.modelspace()
    for lay, col_c in [('MADERA',4), ('COTAS',2), ('TEXTO',1)]:
        if lay not in doc_c.layers: doc_c.layers.add(lay, color=col_c)
    msp_c.add_lwpolyline([(0,0), (b_col/1000,0), (b_col/1000, KL_col), (0, KL_col), (0,0)],
                         close=True, dxfattribs={'layer':'MADERA'})
    if _USE_H_mad:
        dxf_dim_vert(msp_c, b_col/1000+0.15, 0, KL_col, f"H = {KL_col:.2f} m", 50)
        dxf_dim_horiz(msp_c, 0, b_col/1000, -0.2, f"b = {b_col:.0f} mm", 50)
        dxf_text(msp_c, b_col/2000, KL_col+0.2, "ELEVACION COLUMNA MADERA", "EJES", h=TH*1.1, ha="center")
        _cam_c = dxf_rotulo_campos(f"Columna Madera {b_col:.0f}x{h_col:.0f}mm H={KL_col:.1f}m", norma_sel, "001")
        dxf_rotulo(msp_c, _cam_c, 0, -4.5, rot_w=9, rot_h=3, escala=50)
    else:
        msp_c.add_text(f"H = {KL_col:.2f} m", dxfattribs={'layer':'TEXTO','height':0.05,'insert':(b_col/2000, KL_col+0.1)})
        msp_c.add_text(f"b = {b_col:.0f} mm", dxfattribs={'layer':'TEXTO','height':0.05,'insert':(b_col/1000+0.05, KL_col/2)})
    _out_c = io.StringIO()
    doc_c.write(_out_c)
    col_dxf2.download_button(_t(" DXF Columna", " Column DXF"),
                             data=_out_c.getvalue().encode("utf-8"),
                             file_name=f"Columna_Madera_{b_col:.0f}x{h_col:.0f}.dxf", mime="application/dxf")

with tab_doc:
    if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
        doc = Document()
        doc.add_heading(_t(f"Memoria de Estructuras en Madera - {grupo_madera}", f"Timber Structures Report - {grupo_madera}"), 0)
        doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
        doc.add_paragraph(_t(f"Norma Activa: {norma_sel}", f"Active Code: {norma_sel}"))
        doc.add_heading(_t("1. Propiedades de la madera", "1. Timber properties"), level=1)
        doc.add_paragraph(f"E_min = {val_E:.0f} MPa, Fb = {val_Fb:.1f} MPa, Fv = {val_Fv:.1f} MPa, Fc = {val_Fc:.1f} MPa, G = {val_G:.2f}")
        doc.add_heading(_t("2. Diseño de viga", "2. Beam design"), level=1)
        doc.add_paragraph(f"Sección: {b_beam:.0f} x {h_beam:.0f} mm, Luz: {span_L:.2f} m, Carga total: {W_total:.2f} kN/m")
        doc.add_paragraph(f"fb = {fb_act:.2f} MPa {'≤' if ok_flex else '>'} Fb = {val_Fb:.1f} MPa → {'CUMPLE' if ok_flex else 'NO CUMPLE'}")
        doc.add_paragraph(f"fv = {fv_act:.2f} MPa {'≤' if ok_shear else '>'} Fv = {val_Fv:.1f} MPa → {'CUMPLE' if ok_shear else 'NO CUMPLE'}")
        doc.add_paragraph(f"Deflexión = {Def_act:.1f} mm ≤ {Def_max_adm:.1f} mm → {'CUMPLE' if ok_def else 'NO CUMPLE'}")
        doc.add_heading(_t("3. Diseño de columna", "3. Column design"), level=1)
        if esbeltez_l <= 50:
            doc.add_paragraph(f"Sección: {b_col:.0f} x {h_col:.0f} mm, kL = {KL_col:.2f} m, λ = {esbeltez_l:.1f}")
            doc.add_paragraph(f"P admisible = {P_admisible:.1f} kN, P actuante = {P_ax:.1f} kN → {'CUMPLE' if P_ax <= P_admisible else 'NO CUMPLE'}")
        else:
            doc.add_paragraph("Columna demasiado esbelta (λ > 50). No cumple.")
        doc.add_heading(_t("4. Unión con clavos", "4. Nail connection"), level=1)
        if Z_adm > 0:
            doc.add_paragraph(f"Clavo Ø{diam_clavo:.1f} mm, penetración {penetracion_p:.0f} mm")
            doc.add_paragraph(f"Resistencia ajustada Z' = {Z_adm:.2f} kgf/clavo")
        else:
            doc.add_paragraph("La penetración no cumple el mínimo (6D).")
        doc.add_heading(_t("5. Cantidades de madera", "5. Timber quantities"), level=1)
        doc.add_paragraph(f"Viga: {pt_viga:.2f} pt")
        doc.add_paragraph(f"Columna: {pt_col:.2f} pt")
        doc.add_paragraph(f"Total: {total_pt:.2f} pies tabulares")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(_t(" Descargar Memoria", " Download Report"), data=buf, file_name="Memoria_Madera.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_apu:
    if "apu_config" in st.session_state:
        apu = st.session_state.apu_config
        mon = apu["moneda"]
        # Costos de madera
        costo_madera = total_pt * precio_pt
        # Mano de obra (estimación simple: 1 oficial + 1 ayudante por día, rendimiento 30 pt/día)
        rendimiento_dia = st.number_input(_t("Rendimiento (pt/día)", "Output (bf/day)"), 10.0, 100.0, 30.0, 5.0, key="rend_mad")
        dias = total_pt / rendimiento_dia
        costo_mo = dias * apu.get("costo_dia_mo", 69333.33) * 2  # oficial + ayudante
        # Materiales y costos directos
        total_mat = costo_madera
        costo_directo = total_mat + costo_mo
        herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
        aiu = costo_directo * apu.get("pct_aui", 0.30)
        utilidad = costo_directo * apu.get("pct_util", 0.05)
        iva = utilidad * apu.get("iva", 0.19)
        total_proyecto = costo_directo + herramienta + aiu + iva
        
        st.markdown(_t("###  Presupuesto Estimado", "###  Estimated Budget"))
        data_apu = {
            _t("Item", "Item"): [_t("Madera (pies tabulares)", "Timber (board feet)"),
                                 _t("Mano de Obra (días)", "Labor (days)"),
                                 _t("Herramienta Menor", "Minor Tools"),
                                 _t("A.I.U.", "A.I.U."),
                                 _t("IVA s/Utilidad", "VAT on Profit")],
            _t("Cantidad", "Quantity"): [f"{total_pt:.2f} pt", f"{dias:.2f}", f"{apu.get('pct_herramienta',0.05)*100:.1f}% MO", f"{apu.get('pct_aui',0.30)*100:.1f}% CD", f"{apu.get('iva',0.19)*100:.1f}% Util"],
            f"Subtotal [{mon}]": [f"{costo_madera:,.2f}", f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}", f"{iva:,.2f}"]
        }
        st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)
        st.metric(f" Gran Total Proyecto [{mon}]", f"{total_proyecto:,.0f}")
        st.info("ℹ Ve a **APU Mercado** para descargar los costos en tiempo real aquí mismo.")
        
        # Excel APU
        output_excel = io.BytesIO()
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            df_export = pd.DataFrame({
                "Item": ["Madera (pt)", "Mano de Obra (días)"],
                "Cantidad": [total_pt, dias],
                "Unidad": [precio_pt, apu.get("costo_dia_mo", 69333.33)*2]
            })
            df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
            df_export.to_excel(writer, index=False, sheet_name='APU Madera')
            workbook = writer.book
            worksheet = writer.sheets['APU Madera']
            money_fmt = workbook.add_format({'num_format': '#,##0.00'})
            worksheet.set_column('A:A', 25)
            worksheet.set_column('B:D', 15, money_fmt)
        output_excel.seek(0)
        st.download_button(_t(" Descargar Presupuesto Excel", " Download Budget Excel"), data=output_excel, file_name="APU_Madera.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info(_t(" Ve a la página 'APU Mercado' para cargar los costos en vivo.", " Go to the 'Market APU' page to load live costs."))