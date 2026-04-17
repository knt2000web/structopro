try:
    from entregables_ui import mostrar_entregables
except ImportError:
    mostrar_entregables = None
import streamlit as st

#  Utilidad: Color AutoCAD segun # de cuartos de pulgada 
from normas_referencias import mostrar_referencias_norma
def _color_acero_dxf(db_mm: float) -> int:
    """Retorna color AutoCAD (1-255) segun diametro nominal en mm (ASTM/NTC)."""
    if   db_mm <  7.5: return 2   # #2 1/4"   - Amarillo
    elif db_mm < 11.1: return 3   # #3 3/8"   - Verde
    elif db_mm < 14.3: return 4   # #4 1/2"   - Cian
    elif db_mm < 17.5: return 5   # #5 5/8"   - Azul
    elif db_mm < 20.6: return 6   # #6 3/4"   - Magenta
    elif db_mm < 23.8: return 30  # #7 7/8"   - Naranja
    elif db_mm < 27.0: return 1   # #8 1"     - Rojo
    else:              return 10  # #9+ 1-1/8" - Verde oscuro (pesado)
# 
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
from docx import Document
from utils_docx import fig_to_docx_white
from docx.shared import Pt, Cm, Inches
import ezdxf
from datetime import datetime
import io
import plotly.graph_objects as go
try:
    import requests
except ImportError:
    requests = None
import json
try:
    import ifc_export
except ImportError:
    ifc_export = None

# 
# PERSISTENCIA SUPABASE (mismo patrón que Columnas)
# 

def get_supabase_rest_info():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return url, key
    except Exception:
        return None, None

def guardar_proyecto_supabase_vigas(nombre, estado_dict):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_vigas.json"
            db = {}
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            db[f"[VIGAS] {nombre.strip()}"] = {"nombre_proyecto": f"[VIGAS] {nombre.strip()}", "estado_json": json.dumps(estado_dict)}
            with open(db_path, "w", encoding="utf-8") as f: json.dump(db, f)
            return True, " Proyecto guardado (Local)"
        except Exception as e:
            return False, f"? Error guardado local: {e}"

    headers = {
        "apikey": key, "Authorization": f"Bearer {key}",
        "Content-Type": "application/json", "Prefer": "resolution=merge-duplicates"
    }
    payload = {
        "nombre_proyecto": f"[VIGAS] {nombre}",
        "user_id": st.session_state.get("user_id", "anonimo"),
        "estado_json": json.dumps(estado_dict),
    }
    try:
        endpoint = f"{url}/rest/v1/proyectos?on_conflict=nombre_proyecto"
        res = requests.post(endpoint, headers=headers, json=payload)
        if res.status_code in [200, 201, 204]:
            return True, " Proyecto guardado en la nube"
        else:
            return False, f"? Error API: {res.text}"
    except Exception as e:
        return False, f"? Error al guardar: {e}"

def cargar_proyecto_supabase_vigas(nombre):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_vigas.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                match = db.get(f"[VIGAS] {nombre}")
                if match:
                    estado = json.loads(match["estado_json"])
                    for k, v in estado.items(): st.session_state[k] = v
                    return True, f" Proyecto '{nombre}' cargado (Local)"
            return False, f"? No se encontró el proyecto '{nombre}' localmente"
        except Exception as e:
            return False, f"? Excepción al cargar local: {e}"

    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        endpoint = f"{url}/rest/v1/proyectos?nombre_proyecto=eq.[VIGAS] {nombre}&select=*"
        res = requests.get(endpoint, headers=headers)
        if res.status_code == 200:
            data = res.json()
            if data and len(data) > 0:
                estado = json.loads(data[0]["estado_json"])
                for k, v in estado.items():
                    st.session_state[k] = v
                return True, f" Proyecto '{nombre}' cargado"
            else:
                return False, f"? No se encontró el proyecto '{nombre}'"
        else:
            return False, f"? Error al cargar (API): {res.text}"
    except Exception as e:
        return False, f"? Excepción al cargar: {e}"

def listar_proyectos_supabase_vigas():
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = "db_proyectos_vigas.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                return sorted([k.replace("[VIGAS] ", "") for k in db.keys() if str(k).startswith("[VIGAS]")])
            return []
        except Exception:
            return []

    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        endpoint = f"{url}/rest/v1/proyectos?nombre_proyecto=like.[VIGAS]*&select=nombre_proyecto"
        res = requests.get(endpoint, headers=headers)
        if res.status_code == 200:
            data = res.json()
            nombres = [item["nombre_proyecto"].replace("[VIGAS] ", "") for item in data if "nombre_proyecto" in item]
            return sorted(nombres)
        return []
    except Exception:
        return []

def capturar_estado_vigas():
    """Captura todas las claves relevantes del modulo de Vigas y Losas."""
    claves = [
        "vfcmpa", "vfy", "vfcpsi", "vfckgcm2", "vbarsys", "vnormasel", "vnivelsis", "vmtipo", "vmwu", "vmmuizq", "vmmucen", "vmmuder",
        # Viga Rectangular
        "vr_b", "vr_h", "vr_dp", "vr_mu", "vr_L", "vr_bar",
        # Viga T
        "vt_bf", "vt_bw", "vt_hf", "vt_ht", "vt_dp", "vt_mu", "vt_L", "vt_bar",
        # Cortante
        "cv_bw", "cv_d", "cv_vu", "cv_L", "cv_h", "cv_st", "cv_ramas",
        # Deflexiones
        "de_b", "de_h", "de_dp", "de_as", "de_L", "de_wD", "de_wL", "de_cond",
        # Losa
        "ls_ln", "ls_h", "ls_cov", "ls_wD", "ls_wL", "ls_bar", "ls_apoyo",
        # Longitud desarrollo
        "ld_bar", "ld_psit", "ld_psie",
        # Sismo y Punzonamiento (Punto 6)
        "vc_bcol", "vc_wu", "vcmuizq", "vcmuder", "vcmucen", "vc_mu_izq", "vc_mu_der", "vc_mu_cen", "cv_wu", "cv_mu_izq", "cv_mu_der", "cv_mu_cen", "cv_s_diseno", "pz_c1", "pz_c2", "pz_h", "pz_cov", "pz_vu", "pz_tipo",
        # DXF
        "dxf_empresa", "dxf_proyecto", "dxf_plano", "dxf_elaboro", "dxf_reviso", "dxf_aprobo",
    ]
    return {k: st.session_state[k] for k in claves if k in st.session_state}


# 
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# 

st.set_page_config(page_title=_t("Suite Hormigón Armado", "Reinforced Concrete Suite"), layout="wide")
st.markdown("""<div style="width:100%;overflow:hidden;border-radius:14px;margin-bottom:20px;box-shadow:0 4px 32px #0008;">
<svg viewBox="0 0 1100 220" xmlns="http://www.w3.org/2000/svg" style="width:100%;display:block;background:linear-gradient(135deg,#060f0a 0%,#0d2630 100%);">
  <g opacity="0.06" stroke="#34d399" stroke-width="0.5">
    <line x1="0" y1="55" x2="1100" y2="55"/><line x1="0" y1="110" x2="1100" y2="110"/>
    <line x1="0" y1="165" x2="1100" y2="165"/>
    <line x1="220" y1="0" x2="220" y2="220"/><line x1="440" y1="0" x2="440" y2="220"/>
    <line x1="660" y1="0" x2="660" y2="220"/>
  </g>
  <rect x="0" y="0" width="1100" height="3" fill="#10b981" opacity="0.9"/>
  <rect x="0" y="217" width="1100" height="3" fill="#3b82f6" opacity="0.7"/>

  <!--  VIGA RECTANGULAR (perfil)  -->
  <g transform="translate(30,40)">
    <text x="75" y="-8" text-anchor="middle" font-family="monospace" font-size="9" fill="#6b7280" letter-spacing="1.5">VIGA RECTANGULAR</text>
    <rect x="0" y="18" width="150" height="75" rx="2" fill="#2d3748" stroke="#4a5568" stroke-width="1.5"/>
    <rect x="0" y="75" width="150" height="18" rx="1" fill="#1a2535"/>
    <!-- Parabola momento + -->
    <path d="M0,18 Q75,-4 150,18" fill="none" stroke="#f59e0b" stroke-width="1.4" stroke-dasharray="4,3" opacity="0.7"/>
    <!-- Stirrups -->
    <rect x="9" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.9"/>
    <rect x="27" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.7"/>
    <rect x="45" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.6"/>
    <rect x="63" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.6"/>
    <rect x="81" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.6"/>
    <rect x="99" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.6"/>
    <rect x="117" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.7"/>
    <rect x="135" y="20" width="2" height="71" rx="1" fill="#00d4ff" opacity="0.9"/>
    <!-- Top bars -->
    <circle cx="11" cy="25" r="4.5" fill="#e8a838"/><circle cx="139" cy="25" r="4.5" fill="#e8a838"/>
    <!-- Bottom bars (tension) -->
    <circle cx="11" cy="87" r="5.5" fill="#e8a838"/><circle cx="35" cy="87" r="5.5" fill="#e8a838"/>
    <circle cx="115" cy="87" r="5.5" fill="#e8a838"/><circle cx="139" cy="87" r="5.5" fill="#e8a838"/>
    <line x1="0" y1="108" x2="150" y2="108" stroke="#60a5fa" stroke-width="0.8"/>
    <text x="75" y="120" text-anchor="middle" font-family="monospace" font-size="10" fill="#93c5fd">b x h</text>
    <text x="75" y="132" text-anchor="middle" font-family="monospace" font-size="9" fill="#64748b">L (luz libre)</text>
  </g>

  <!--  VIGA T  -->
  <g transform="translate(205,38)">
    <text x="75" y="-8" text-anchor="middle" font-family="monospace" font-size="9" fill="#6b7280" letter-spacing="1.5">VIGA T — LOSA COLABORANTE</text>
    <rect x="0" y="18" width="150" height="22" rx="2" fill="#2d3748" stroke="#4a5568" stroke-width="1.5"/>
    <rect x="38" y="40" width="74" height="58" rx="2" fill="#243040" stroke="#4a5568" stroke-width="1.5"/>
    <line x1="0" y1="12" x2="150" y2="12" stroke="#10b981" stroke-width="0.8" stroke-dasharray="4,2" opacity="0.7"/>
    <text x="75" y="9" text-anchor="middle" font-family="monospace" font-size="9" fill="#34d399">beff</text>
    <!-- Flange bars -->
    <circle cx="10" cy="25" r="3.5" fill="#e8a838"/><circle cx="28" cy="25" r="3.5" fill="#e8a838"/>
    <circle cx="44" cy="25" r="4" fill="#e8a838"/><circle cx="106" cy="25" r="4" fill="#e8a838"/>
    <circle cx="122" cy="25" r="3.5" fill="#e8a838"/><circle cx="140" cy="25" r="3.5" fill="#e8a838"/>
    <!-- Web stirrups -->
    <rect x="42" y="42" width="2" height="54" rx="1" fill="#00d4ff" opacity="0.9"/>
    <rect x="56" y="42" width="2" height="54" rx="1" fill="#00d4ff" opacity="0.7"/>
    <rect x="70" y="42" width="2" height="54" rx="1" fill="#00d4ff" opacity="0.7"/>
    <rect x="84" y="42" width="2" height="54" rx="1" fill="#00d4ff" opacity="0.7"/>
    <rect x="98" y="42" width="2" height="54" rx="1" fill="#00d4ff" opacity="0.9"/>
    <rect x="108" y="42" width="2" height="54" rx="1" fill="#00d4ff" opacity="0.9"/>
    <!-- Bottom bars -->
    <circle cx="48" cy="91" r="5.5" fill="#e8a838"/><circle cx="75" cy="91" r="5.5" fill="#e8a838"/>
    <circle cx="102" cy="91" r="5.5" fill="#e8a838"/>
    <text x="75" y="112" text-anchor="middle" font-family="monospace" font-size="9" fill="#93c5fd">bw (alma)</text>
  </g>

  <!--  LOSA 1D  -->
  <g transform="translate(390,62)">
    <text x="80" y="-18" text-anchor="middle" font-family="monospace" font-size="9" fill="#6b7280" letter-spacing="1.5">LOSA UNIDIRECCIONAL</text>
    <!-- Load arrows -->
    <g stroke="#ef4444" stroke-width="1.2" opacity="0.8">
      <line x1="0" y1="-18" x2="160" y2="-18"/>
      <line x1="20" y1="-18" x2="20" y2="-4"/><polygon points="20,-4 17,-11 23,-11" fill="#ef4444"/>
      <line x1="50" y1="-18" x2="50" y2="-4"/><polygon points="50,-4 47,-11 53,-11" fill="#ef4444"/>
      <line x1="80" y1="-18" x2="80" y2="-4"/><polygon points="80,-4 77,-11 83,-11" fill="#ef4444"/>
      <line x1="110" y1="-18" x2="110" y2="-4"/><polygon points="110,-4 107,-11 113,-11" fill="#ef4444"/>
      <line x1="140" y1="-18" x2="140" y2="-4"/><polygon points="140,-4 137,-11 143,-11" fill="#ef4444"/>
    </g>
    <text x="80" y="-22" text-anchor="middle" font-family="monospace" font-size="8" fill="#f87171">wu (carga diseno)</text>
    <!-- Slab -->
    <rect x="0" y="0" width="160" height="38" rx="2" fill="#2d3748" stroke="#4a5568" stroke-width="1.5"/>
    <!-- Main bars bottom -->
    <circle cx="15" cy="31" r="4" fill="#e8a838"/><circle cx="40" cy="31" r="4" fill="#e8a838"/>
    <circle cx="65" cy="31" r="4" fill="#e8a838"/><circle cx="90" cy="31" r="4" fill="#e8a838"/>
    <circle cx="115" cy="31" r="4" fill="#e8a838"/><circle cx="140" cy="31" r="4" fill="#e8a838"/>
    <!-- Temp bars top -->
    <circle cx="27" cy="7" r="2.5" fill="#b45309"/><circle cx="52" cy="7" r="2.5" fill="#b45309"/>
    <circle cx="77" cy="7" r="2.5" fill="#b45309"/><circle cx="102" cy="7" r="2.5" fill="#b45309"/>
    <circle cx="127" cy="7" r="2.5" fill="#b45309"/>
    <!-- h annotation -->
    <line x1="168" y1="0" x2="168" y2="38" stroke="#60a5fa" stroke-width="0.8"/>
    <text x="175" y="23" font-family="monospace" font-size="10" fill="#93c5fd">h</text>
    <!-- Deflection -->
    <path d="M0,55 Q80,82 160,55" fill="none" stroke="#f59e0b" stroke-width="1.2" stroke-dasharray="4,3" opacity="0.6"/>
    <text x="80" y="82" text-anchor="middle" font-family="monospace" font-size="8" fill="#fbbf24">deflexion</text>
  </g>

  <!--  TEXT BLOCK  -->
  <g transform="translate(622,0)">
    <rect x="0" y="28" width="4" height="165" rx="2" fill="#10b981"/>
    <text x="18" y="66" font-family="Arial,sans-serif" font-size="30" font-weight="bold" fill="#ffffff">VIGAS Y LOSAS</text>
    <text x="18" y="92" font-family="Arial,sans-serif" font-size="17" font-weight="300" fill="#6ee7b7" letter-spacing="2">DISENO SISMICO NSR-10</text>
    <rect x="18" y="100" width="440" height="1" fill="#10b981" opacity="0.5"/>
    <!-- Tags -->
    <rect x="18" y="111" width="116" height="22" rx="11" fill="#0c2a1a" stroke="#10b981" stroke-width="1"/>
    <text x="76" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#34d399">FLEXION + CORTANTE</text>
    <rect x="140" y="111" width="100" height="22" rx="11" fill="#0a1e30" stroke="#3b82f6" stroke-width="1"/>
    <text x="190" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#60a5fa">VIGA T / LOSA 1D</text>
    <rect x="246" y="111" width="106" height="22" rx="11" fill="#1c1416" stroke="#ef4444" stroke-width="1"/>
    <text x="299" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#f87171">SISMICA DMO / DES</text>
    <rect x="358" y="111" width="88" height="22" rx="11" fill="#1a150a" stroke="#f59e0b" stroke-width="1"/>
    <text x="402" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#fbbf24">PLANO DXF / IFC</text>
    <!-- Description -->
    <text x="18" y="156" font-family="Arial,sans-serif" font-size="11" fill="#64748b">Diseno estructural de vigas rect., en T y losas 1D segun</text>
    <text x="18" y="172" font-family="Arial,sans-serif" font-size="11" fill="#64748b">NSR-10 C.21 — Cortante, punzonamiento, deflexiones,</text>
    <text x="18" y="188" font-family="Arial,sans-serif" font-size="11" fill="#64748b">despiece BBS, planos ICONTEC y modelos BIM IFC4.</text>
  </g>
</svg></div>""", unsafe_allow_html=True)

st.title(_t("Suite de Diseño — Vigas y Losas", "Design Suite — Beams & Slabs"))

#  PERSISTENCIA DE MÓDULO VÍA URL (sobrevive F5) 
_modulos_disponibles = [
    " Tabla de Secciones de Acero de Refuerzo",
    " Diseño Completo de Viga (Flujo Guiado)",
    " Diseño a Flexión — Viga Rectangular",
    " Diseño a Flexión — Viga T",
    " Diseño a Cortante — Vigas de Concreto",
    " Resistencia a Cortante por Punzonamiento — Losas",
    " Inercia Fisurada y Deflexiones en Vigas",
    " Diseño de Losa en Una Dirección",
    " Longitud de Desarrollo y Empalmes",
    " Diseño Sísmico Integral y Plano DXF (Viga DMO / DES)",
    " Cuadro de Mando General",
]

_modulo_desde_url = st.query_params.get("modulo", None)
_idx_default = 1
if _modulo_desde_url and _modulo_desde_url in _modulos_disponibles:
    _idx_default = _modulos_disponibles.index(_modulo_desde_url)

modulo_sel = st.selectbox(
    "Navegador de Módulos (Vigas y Losas):",
    _modulos_disponibles,
    index=_idx_default,
    key="moduloactivo"
)
st.query_params["modulo"] = modulo_sel
st.markdown("---")

st.markdown(_t("Herramientas de diseño de concreto reforzado según **10 normativas internacionales**.", "Reinforced concrete design tools based on **10 international codes**."))

# 
# FUNCIONES DE DIBUJO PARA FIGURADO (MEJORADAS)
# 
def draw_longitudinal_bar(total_len_cm, straight_len_cm, hook_len_cm, bar_diam_mm):
    """
    Dibuja una barra longitudinal con ganchos de 90° en ambos extremos.
    total_len_cm : longitud total de la barra (incluyendo ganchos)
    straight_len_cm : longitud recta entre ganchos
    hook_len_cm : longitud de cada gancho (12db)
    bar_diam_mm : diámetro de la barra (para escala)
    """
    fig, ax = plt.subplots(figsize=(max(6, total_len_cm/20), 2))
    fig.patch.set_facecolor('#1e1e2e')
    ax_arr = fig.get_axes()
    for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    # Línea central
    ax.plot([0, straight_len_cm], [0, 0], 'k-', linewidth=2)
    # Gancho izquierdo (90° hacia arriba)
    ax.plot([0, 0], [0, hook_len_cm], 'k-', linewidth=2)
    # Gancho derecho (90° hacia abajo)
    ax.plot([straight_len_cm, straight_len_cm], [0, -hook_len_cm], 'k-', linewidth=2)
    # Cotas
    ax.annotate(f"{straight_len_cm:.0f} cm", xy=(straight_len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 12db = {hook_len_cm:.0f} cm", xy=(0, hook_len_cm/2), ha='right', fontsize=8)
    ax.annotate(f"Gancho 12db", xy=(straight_len_cm, -hook_len_cm/2), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*0.2, straight_len_cm + hook_len_cm*0.2)
    ax.set_ylim(-hook_len_cm*1.2, hook_len_cm*1.2)
    ax.axis('off')
    ax.set_title(f"Varilla longitudinal - Ø{bar_diam_mm:.0f} mm - Longitud total {total_len_cm:.0f} cm", fontsize=9)
    return fig

def draw_stirrup_beam(b_cm, h_cm, hook_len_cm, bar_diam_mm):
    """
    Dibuja un estribo rectangular con ganchos de 135° en una esquina.
    b_cm, h_cm : dimensiones interiores del estribo (medidas entre caras internas)
    hook_len_cm : longitud de proyección del gancho (aprox. 6db o 12db, solo visual)
    bar_diam_mm : diámetro de la barra (para escala)
    """
    fig, ax = plt.subplots(figsize=(max(5, b_cm/15), max(5, h_cm/15)))
    fig.patch.set_facecolor('#1e1e2e')
    ax_arr = fig.get_axes()
    for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    ax.set_aspect('equal')
    # Rectángulo exterior (interior real)
    x0, y0 = 0, 0
    ax.plot([x0, x0+b_cm], [y0, y0], 'k-', linewidth=2)          # base inferior
    ax.plot([x0+b_cm, x0+b_cm], [y0, y0+h_cm], 'k-', linewidth=2) # lado derecho
    ax.plot([x0+b_cm, x0], [y0+h_cm, y0+h_cm], 'k-', linewidth=2) # base superior
    ax.plot([x0, x0], [y0+h_cm, y0], 'k-', linewidth=2)          # lado izquierdo
    # Gancho de 135° hacia el núcleo confinado (Punto 5 Trigonometría)
    import math
    angle_rad = math.radians(45) # 135 grados hacia adentro es equivalente a 45 grados desde el eje X/Y hacia el centro rect
    
    # Gancho desde la esquina superior izquierda (x0, y0+h_cm) hacia adentro
    vis_hook = min(hook_len_cm, b_cm/4.0, h_cm/4.0)
    hx = vis_hook * math.cos(angle_rad)
    hy = -vis_hook * math.sin(angle_rad)
    ax.plot([x0, x0 + hx], [y0+h_cm, y0+h_cm + hy], 'k-', linewidth=2)
    
    # Gancho simulado desde la otra esquina o rama (usualmente se dibuja uno continuo o dos en la misma esquina, simularé la 2da vuelta del estribo)
    ax.plot([x0+b_cm, x0+b_cm - hx], [y0+h_cm, y0+h_cm + hy], 'k--', linewidth=1.5, alpha=0.5)
    
    # Cotas
    ax.annotate(f"{b_cm:.0f} cm", xy=(b_cm/2, -0.5), ha='center', fontsize=8)
    ax.annotate(f"{h_cm:.0f} cm", xy=(-0.5, h_cm/2), ha='right', va='center', fontsize=8)
    ax.annotate(f"Gancho 135°", xy=(x0 - hook_len_cm*0.5, y0 - hook_len_cm*0.7), ha='right', fontsize=8)
    ax.set_xlim(-hook_len_cm*1.2, b_cm + hook_len_cm*0.5)
    ax.set_ylim(-hook_len_cm*1.2, h_cm + hook_len_cm*0.5)
    ax.axis('off')
    ax.set_title(f"Estribo - Ø{bar_diam_mm:.0f} mm - Perímetro {2*(b_cm+h_cm):.0f} cm", fontsize=9)
    return fig

# 
# UNIDADES DE SALIDA
# 
st.sidebar.header(_t("Unidades de salida", "Output units"))
unidades_salida = st.sidebar.radio("Unidades de fuerza/momento:", ["kiloNewtons (kN, kN·m)", "Toneladas fuerza (tonf, tonf·m)"], key="v_output_units")
if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom    = "tonf·m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom    = "kN·m"

# 
# APU CON ENTRADA DIRECTA (GLOBAL)
# 
with st.expander("APU – Precios en vivo (materiales y mano de obra)", expanded=False):
    st.markdown("Ingrese los precios unitarios de los materiales y mano de obra para calcular el costo total de las estructuras.")
    with st.form(key="apu_form_global"):
        moneda = st.text_input("Moneda (ej. COP, USD)", value=st.session_state.get("apu_moneda_global", "COP"))
        col1a, col2a = st.columns(2)
        with col1a:
            precio_cemento = st.number_input("Precio por bulto de cemento", value=st.session_state.get("apu_cemento_global", 28000.0), step=1000.0, format="%.2f")
            precio_acero = st.number_input("Precio por kg de acero", value=st.session_state.get("apu_acero_global", 7500.0), step=100.0, format="%.2f")
            precio_arena = st.number_input("Precio por m³ de arena", value=st.session_state.get("apu_arena_global", 120000.0), step=5000.0, format="%.2f")
            precio_grava = st.number_input("Precio por m³ de grava", value=st.session_state.get("apu_grava_global", 130000.0), step=5000.0, format="%.2f")
        with col2a:
            precio_mo = st.number_input("Costo mano de obra (día)", value=st.session_state.get("apu_mo_global", 70000.0), step=5000.0, format="%.2f")
            pct_herramienta = st.number_input("% Herramienta menor (sobre MO)", value=st.session_state.get("apu_herramienta_global", 5.0), step=1.0, format="%.1f") / 100.0
            pct_aui = st.number_input("% A.I.U. (sobre costo directo)", value=st.session_state.get("apu_aui_global", 30.0), step=5.0, format="%.1f") / 100.0
            pct_util = st.number_input("% Utilidad (sobre costo directo)", value=st.session_state.get("apu_util_global", 5.0), step=1.0, format="%.1f") / 100.0
            iva = st.number_input("IVA (%) sobre utilidad", value=st.session_state.get("apu_iva_global", 19.0), step=1.0, format="%.1f") / 100.0
        submitted = st.form_submit_button("Guardar precios")
        if submitted:
            st.session_state.apu_config = {
                "moneda": moneda,
                "cemento": precio_cemento,
                "acero": precio_acero,
                "arena": precio_arena,
                "grava": precio_grava,
                "costo_dia_mo": precio_mo,
                "pct_herramienta": pct_herramienta,
                "pct_aui": pct_aui,
                "pct_util": pct_util,
                "iva": iva
            }
            st.success("Precios guardados. Ahora se mostrarán los presupuestos en las secciones de cantidades.")
            st.rerun()

# 
# PIE DE P?GINA / DERECHOS RESERVADOS (en sidebar)
# 
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    <br><br>
    <i> ? Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ??????????????????????????????????????????
# CODES DICT (COMPLETO)
# ??????????????????????????????????????????
CODES = {
    "NSR-10 (Colombia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["DMI — Disipación Mínima", "DMO — Disipación Moderada", "DES — Disipación Especial"],
        "ref": "NSR-10 Título C (C.9, C.11, C.21)",
        "bag_kg": 50.0,
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["OMF (SDC A–B)", "IMF (SDC C)", "SMF (SDC D–F)"],
        "ref": "ACI 318-25 (Sections 9, 22, 25)",
        "bag_kg": 42.6,
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["OMF (SDC A–B)", "IMF (SDC C)", "SMF (SDC D–F)"],
        "ref": "ACI 318-19 (Sections 9, 22, 25)",
        "bag_kg": 42.6,
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["OMF (SDC A–B)", "IMF (SDC C)", "SMF (SDC D–F)"],
        "ref": "ACI 318-14 (Sections 9, 22, 25)",
        "bag_kg": 42.6,
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["GS — Grado Reducido", "GM — Grado Moderado", "GA — Grado Alto"],
        "ref": "NEC-SE-HM Ecuador (Cap. 4)",
        "bag_kg": 50.0,
    },
    "E.060 (Perú)": {
        "phi_flex": 0.90, "phi_shear": 0.85, "phi_comp": 0.70,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["PO — Pórtico Ordinario (Z1–Z2)", "PE — Pórtico Especial (Z3–Z4)"],
        "ref": "Norma E.060 Perú (Arts. 9, 11, 21)",
        "bag_kg": 42.5,
    },
    "NTC-EM (México)": {
        "phi_flex": 0.85, "phi_shear": 0.80, "phi_comp": 0.70,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["MDL — Ductilidad Limitada", "MROD — Ductilidad Ordinaria", "MRLE — Ductilidad Alta"],
        "ref": "NTC-EM México 2017 (Cap. 2, 4)",
        "bag_kg": 50.0,
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.70,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["PO — Pórtico Ordinario", "PM — Pórtico Moderado", "PE — Pórtico Especial"],
        "ref": "COVENIN 1753-2006 Venezuela",
        "bag_kg": 42.5,
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["DO — Diseño Ordinario", "DE — Diseño Especial Sísmico"],
        "ref": "NB 1225001-2020 Bolivia (ACI 318-19)",
        "bag_kg": 50.0,
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["GE — Grado Estándar", "GM — Ductilidad Moderada", "GA — Ductilidad Alta"],
        "ref": "CIRSOC 201-2025 Argentina (basada en ACI 318-19)",
        "bag_kg": 50.0,
    },
}

# Rebar tables
REBAR_US = {"#3 (Ø9.5mm)":0.71,"#4 (Ø12.7mm)":1.29,"#5 (Ø15.9mm)":1.99,"#6 (Ø19.1mm)":2.84,"#7 (Ø22.2mm)":3.87,"#8 (Ø25.4mm)":5.10,"#9 (Ø28.7mm)":6.45,"#10 (Ø32.3mm)":7.92}
REBAR_MM = {"8mm":0.503,"10mm":0.785,"12mm":1.131,"14mm":1.539,"16mm":2.011,"18mm":2.545,"20mm":3.142,"22mm":3.801,"25mm":4.909,"28mm":6.158,"32mm":8.042}
DIAM_US = {"#3 (Ø9.5mm)":9.53,"#4 (Ø12.7mm)":12.7,"#5 (Ø15.9mm)":15.88,"#6 (Ø19.1mm)":19.05,"#7 (Ø22.2mm)":22.23,"#8 (Ø25.4mm)":25.4,"#9 (Ø28.7mm)":28.65,"#10 (Ø32.3mm)":32.26}
DIAM_MM = {"8mm":8,"10mm":10,"12mm":12,"14mm":14,"16mm":16,"18mm":18,"20mm":20,"22mm":22,"25mm":25,"28mm":28,"32mm":32}

# Funciones auxiliares
def get_beta1(fc):
    if fc <= 28: return 0.85
    b = 0.85 - 0.05*(fc-28)/7.0
    return max(b, 0.65)

def get_rho_min(fc, fy, norm):
    return max(0.25*math.sqrt(fc)/fy, 1.4/fy)

def get_rho_max_beam(fc, fy, beta1, eps_cu=0.003, nivel_sis="DMO", eps_t_min=0.005):
    """
    Retorna (rho_max, rho_bal). 
    NOTA: La reducción de phi para zonas de transición (0.004 <= et < 0.005) 
    debe hacerse en el chequeo estructural posterior evaluando la deformación real.
    """
    rho_bal = (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+fy/200000))
    rho_max_norma = (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+0.004))
    if "DES" in nivel_sis or "DMO" in nivel_sis:
        rho_max = min(rho_max_norma, 0.025)
    else:
        rho_max = rho_max_norma
    return rho_max, rho_bal

def mix_for_fc(fc):
    table = [
        (14,250,205,810,1060),(17,290,200,780,1060),(21,350,193,720,1060),
        (25,395,193,680,1020),(28,430,190,640,1000),(35,530,185,580,960),
        (42,620,180,520,910),(56,740,175,450,850),
    ]
    if fc <= table[0][0]: return table[0][1:]
    if fc >= table[-1][0]: return table[-1][1:]
    for i in range(len(table)-1):
        lo,hi = table[i],table[i+1]
        if lo[0] <= fc <= hi[0]:
            t = (fc-lo[0])/(hi[0]-lo[0])
            return tuple(lo[j]+t*(hi[j]-lo[j]) for j in range(1,5))
    return table[-1][1:]

def sec_dark_fig(w, h, title=""):
    fig, ax = plt.subplots(figsize=(max(3,w/h*3), 3))
    fig.patch.set_facecolor('#1e1e2e')
    ax_arr = fig.get_axes()
    for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    fig.patch.set_facecolor('#1a1a2e')
    ax.set_facecolor('#1a1a2e')
    ax.add_patch(patches.Rectangle((0,0),w,h,linewidth=2,edgecolor='white',facecolor='#4a4a6a'))
    ax.set_xlim(-w*0.15, w*1.15); ax.set_ylim(-h*0.15, h*1.15)
    ax.axis('off')
    ax.set_title(title, color='white', fontsize=8)
    return fig, ax

def sec_light_fig(w, h, title=""):
    fig, ax = plt.subplots(figsize=(max(3,w/h*3), 3))
    fig.patch.set_facecolor('#1e1e2e')
    ax_arr = fig.get_axes()
    for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.add_patch(patches.Rectangle((0,0),w,h,linewidth=2,edgecolor='black',facecolor='#f0f0f0'))
    ax.set_xlim(-w*0.15, w*1.15); ax.set_ylim(-h*0.15, h*1.15)
    ax.axis('off')
    ax.set_title(title, color='black', fontsize=10)
    return fig, ax

def sec_light_fig_t(bf, bw, hf, ht, title=""):
    fig, ax = plt.subplots(figsize=(5,4))
    fig.patch.set_facecolor('#1e1e2e')
    ax_arr = fig.get_axes()
    for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.add_patch(patches.Rectangle(((bf-bw)/2, 0), bw, ht-hf, linewidth=1.5, edgecolor='black', facecolor='#e0e0e0'))
    ax.add_patch(patches.Rectangle((0, ht-hf), bf, hf, linewidth=1.5, edgecolor='black', facecolor='#d0d0d0'))
    ax.set_xlim(-5, bf+5); ax.set_ylim(-5, ht+5)
    ax.axis('off')
    ax.set_title(title, color='black', fontsize=10)
    return fig, ax


def qty_table(rows):
    st.dataframe(pd.DataFrame(rows, columns=["Concepto","Valor"]), use_container_width=True, hide_index=True)

def render_apu_breakdown(vol_m3, peso_kg, fc_m, num_bars_str=""):
    """
    Renders APU detail given concrete volume, steel weight and concrete strength.
    """
    if "apu_config" not in st.session_state:
        st.info("Configure los precios en el menú lateral '? Materiales Globales'para ver el presupuesto.")
        return
        
    apu = st.session_state.apu_config
    mon = apu.get("moneda", "$")
    
    st.markdown("---")
    st.success("**Análisis de Precios Unitarios (APU) Aplicado.**")
    
    m = mix_for_fc(fc_m)
    bag_kg = st.session_state.get("cemento_kg", 50.0)
    bags = (m[0] * vol_m3) / bag_kg
    
    c_cem = bags * apu.get("cemento", 0)
    c_ace = peso_kg * apu.get("acero", 0)
    vol_are = (m[2] * vol_m3) / 1600
    vol_gra = (m[3] * vol_m3) / 1600
    
    c_are = vol_are * apu.get("arena", 0)
    c_gra = vol_gra * apu.get("grava", 0)
    total_mat = c_cem + c_ace + c_are + c_gra
    
    total_dias_mo = (peso_kg * 0.04) + (vol_m3 * 0.4)
    costo_mo = total_dias_mo * apu.get("costo_dia_mo", 70000)
    
    costo_directo = total_mat + costo_mo
    herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
    aiu = costo_directo * apu.get("pct_aui", 0.30)
    utilidad = costo_directo * apu.get("pct_util", 0.05)
    iva_v = utilidad * apu.get("iva", 0.19)
    gran_total = costo_directo + herramienta + aiu + iva_v
    
    c1, c2, c3 = st.columns(3)
    c1.metric(f" ? Costo Directo ({mon})", f"{costo_directo:,.2f}")
    c2.metric(f" Mano de Obra ({mon})", f"{costo_mo:,.2f}")
    c3.metric(f" Gran Total ({mon})", f"{gran_total:,.2f}")
    
    with st.expander("Ver desglose detallado del APU"):
        st.markdown(f"**Volumen Concreto:** {vol_m3:.3f} m³ | **Acero:** {peso_kg:.1f} kg {num_bars_str}")
        df_apu = pd.DataFrame([
            ("Cemento (bultos)", c_cem),
            ("Acero (kg)", c_ace),
            ("Arena (m³)", c_are),
            ("Grava (m³)", c_gra),
            ("Mano de Obra (días)", costo_mo),
            ("COSTO DIRECTO", costo_directo),
            ("Herramienta Menor", herramienta),
            ("A.I.U.", aiu),
            ("Utilidad", utilidad),
            ("IVA (sobre utilidad)", iva_v),
            ("GRAN TOTAL", gran_total)
        ], columns=["Concepto", f"Costo ({mon})"])
        
        df_apu[f"Costo ({mon})"] = df_apu[f"Costo ({mon})"].apply(lambda x: f"{x:,.2f}")
        st.dataframe(df_apu, use_container_width=True, hide_index=True)

def add_historial_diseno(modulo, elemento, estado, norma):
    if "historial_disenos" not in st.session_state:
        st.session_state.historial_disenos = []
    
    hora_str = datetime.now().strftime("%H:%M:%S")
    entry = {
        "Módulo": modulo,
        "Elemento / Geometría": elemento,
        "Estado": estado,
        "Norma": norma,
        "Hora": hora_str
    }
    st.session_state.historial_disenos.append(entry)
    st.toast("Elemento guardado en el Cuadro de Mando General")

# ??????????????????????????????????????????
# GLOBAL SIDEBAR
# ??????????????????????????????????????????
st.sidebar.header(_t("? Norma de Diseño", "? Design Code"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = list(CODES.keys())[0]
norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagpedia.net/data/flags/mini/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>',
    unsafe_allow_html=True
)
code = CODES[norma_sel]
nivel_sis = st.sidebar.selectbox(
    _t("Nivel Sísmico:", "Seismic Level:"),
    code["seismic_levels"],
    index=code["seismic_levels"].index(st.session_state.v_nivel_sis) if "v_nivel_sis" in st.session_state and st.session_state.v_nivel_sis in code["seismic_levels"] else 0,
    key="v_nivel_sis"
)
mostrar_referencias_norma(norma_sel, "vigas_losas")
st.sidebar.markdown(f" `{code['ref']}`")
st.sidebar.markdown(f"**φ flex:** {code['phi_flex']} | **φ cort:** {code['phi_shear']}")

st.sidebar.header(_t("? Materiales Globales", "? Global Materials"))
fc_unit = st.sidebar.radio(_t("Unidad f'c:", "f'c Unit:"), ["MPa","PSI","kg/cm²"], horizontal=True, key="v_fc_unit")
if fc_unit == "PSI":
    psi_options = ["2500","3000","3500","4000","4500","5000"]
    psi_v = st.sidebar.selectbox("f'c [PSI]:", psi_options,
                                 index=psi_options.index(st.session_state.v_fc_psi) if "v_fc_psi" in st.session_state and st.session_state.v_fc_psi in psi_options else 1,
                                 key="v_fc_psi")
    fc = float(psi_v)*0.00689476
    st.sidebar.info(f"f'c = {psi_v} PSI → **{fc:.2f} MPa**")
elif fc_unit == "kg/cm²":
    kg_options = ["175","210","250","280","350","420"]
    kg_v = st.sidebar.selectbox("f'c [kg/cm²]:", kg_options,
                                index=kg_options.index(st.session_state.v_fc_kgcm2) if "v_fc_kgcm2" in st.session_state and st.session_state.v_fc_kgcm2 in kg_options else 1,
                                key="v_fc_kgcm2")
    fc = float(kg_v)/10.1972
    st.sidebar.info(f"f'c = {kg_v} kg/cm² → **{fc:.2f} MPa**")
else:
    fc = st.sidebar.number_input("f'c [MPa]:", 15.0, 80.0, st.session_state.get("v_fc_mpa", 21.0), 1.0, key="v_fc_mpa")

fy = st.sidebar.number_input("fy [MPa]:", 200.0, 500.0, st.session_state.get("v_fy", 420.0), 10.0, key="v_fy")
Es = 200000.0
Ec = 4700*math.sqrt(fc)
beta1 = get_beta1(fc)
rho_min = get_rho_min(fc, fy, norma_sel)
rho_max, rho_bal = get_rho_max_beam(fc, fy, beta1, nivel_sis=nivel_sis)

bar_sys = st.sidebar.radio("Sistema Varillas:", ["Pulgadas (# US)","Milímetros (mm)"], horizontal=True, key="v_bar_sys")
rebar_dict = REBAR_US if "Pulgadas" in bar_sys else REBAR_MM
diam_dict  = DIAM_US  if "Pulgadas" in bar_sys else DIAM_MM

phi_f = code["phi_flex"]
phi_v = code["phi_shear"]
lam   = code["lambda"]
bag_kg = code["bag_kg"]

st.sidebar.markdown("---")
st.sidebar.caption(f"Ec = {Ec:.0f} MPa  |  β? = {beta1:.3f}  |  f'c = {fc:.2f} MPa")

# ??????????????????????????????????????????
# 1. TABLA DE ACERO
# ??????????????????????????????????????????



# 
# PANEL DE BIENVENIDA — info normativa visible siempre en cada módulo
# 
def _panel_normativo(titulo, descripcion, inputs_requeridos, fc, fy,
                     rhomin, rhomax, phif, phiv, normasel, nivelsis, coderef):
    st.markdown(
        f"""<div style="background:#1a2a1a;border-left:4px solid #4caf50;
        border-radius:8px;padding:14px 18px;margin-bottom:16px;">
        <span style="color:#81c784;font-weight:700;font-size:15px;"> {titulo}</span><br>
        <span style="color:#aaa;font-size:13px;">{descripcion}</span>
        </div>""",
        unsafe_allow_html=True,
    )
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("φ flexión", f"{phif:.2f}", help=coderef)
    c2.metric("φ cortante", f"{phiv:.2f}", help=coderef)
    c3.metric("ρ mín", f"{rhomin:.4f}", help="Cuantía mínima longitudinal")
    c4.metric("ρ máx", f"{rhomax:.4f}", help="Cuantía máxima (εt ≥ 0.005)")
    c5.metric("fc / fy", f"{fc:.0f} / {fy:.0f} MPa", help="Materiales activos")
    st.caption(f" Norma: **{normasel}** | Nivel sísmico: **{nivelsis}** | Ref: {coderef}")
    if inputs_requeridos:
        with st.expander(" Datos requeridos para este módulo", expanded=False):
            for item in inputs_requeridos:
                st.markdown(f"- {item}")
    st.markdown("---")

if modulo_sel == " Tabla de Secciones de Acero de Refuerzo":
    #  PANEL 
    _panel_normativo("Tabla de Secciones de Acero de Refuerzo",
        "Referencia rápida de áreas, diámetros y pesos de varillas en sistema US y SI.",
        ["No requiere datos — es una tabla de consulta."],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(_t(f"**Referencia:** {code['ref']}", f"**Reference:** {code['ref']}"))
    rows_us, rows_mm = [], []
    for k,a in REBAR_US.items():
        d_bar = DIAM_US[k]
        rows_us.append({"Barra":k,"Ø (mm)":f"{d_bar:.2f}","?rea (cm²)":f"{a:.3f}","Peso (kg/m)":f"{a * 0.785:.3f}"})
    for k,a in REBAR_MM.items():
        d_bar = DIAM_MM[k]
        rows_us_mm = {"Barra (SI)":k,"Ø (mm)":f"{d_bar:.0f}","?rea (cm²)":f"{a:.4f}","Peso (kg/m)":f"{a * 0.785:.3f}"}
        rows_mm.append(rows_us_mm)
    c1,c2 = st.columns(2)
    with c1:
        st.markdown("##### Sistema US (pulgadas)")
        st.dataframe(pd.DataFrame(rows_us), use_container_width=True, hide_index=True)
    with c2:
        st.markdown("##### Sistema SI (milímetros)")
        st.dataframe(pd.DataFrame(rows_mm), use_container_width=True, hide_index=True)


# 
# FLUJO MAESTRO: DISEÑO COMPLETO DE VIGA
# 
if modulo_sel == " Diseño Completo de Viga (Flujo Guiado)":
    #  PANEL 
    _panel_normativo("Diseño Completo de Viga — Flujo Guiado",
        "Unifica Flexión, Cortante, Deflexiones y Memoria en un solo flujo paso a paso.",
        ["Paso 1: Geometría (bw, h, d', L, tipo sección)",
         "Paso 2: Cargas (modo automático wD/wL o manual Mu/Vu)",
         "Paso 3-5: Revisión en módulos individuales abajo",
         "Paso 6: Memoria integrada y descarga DOCX"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown("Este asistente unifica todos los cálculos (Flexión, Cortante, Deflexiones, Memoria y Planos) en un solo flujo.")
    
    vm_tabs = st.tabs([
        "1. Geometría", 
        "2. Cargas", 
        "3. Flexión Rectangular", 
        "4. Flexión Viga T", 
        "5. Cortante", 
        "6. Deflexiones", 
        "7. Resultados y Memoria"
    ])
    
    # 
    # PASO 1 — GEOMETRÍA Y MATERIALES
    # 
    with vm_tabs[0]:
        st.subheader("Geometría y Materiales de la Viga")
        c1p1, c2p1, c3p1 = st.columns(3)
        with c1p1:
            vm_b  = st.number_input("Ancho alma bw [cm]", 10.0, 150.0, st.session_state.get("vm_b", 30.0), 5.0, key="vm_b")
            vm_h  = st.number_input("Altura total h [cm]", 20.0, 200.0, st.session_state.get("vm_h", 50.0), 5.0, key="vm_h")
            vm_dp = st.number_input("Recubrim. d' [cm]", 2.0, 15.0, st.session_state.get("vm_dp", 5.0), 0.5, key="vm_dp")
        with c2p1:
            vm_L  = st.number_input("Longitud de la viga [m]", 1.0, 30.0, st.session_state.get("vm_L", 5.0), 0.5, key="vm_L")
            vm_tipo = st.radio("Tipo de sección", ["Rectangular", "Viga T"], horizontal=True, key="vm_tipo")
            if vm_tipo == "Viga T":
                vm_bf = st.number_input("Ancho del ala bf [cm]", 20.0, 300.0, st.session_state.get("vm_bf", 80.0), 5.0, key="vm_bf")
                vm_hf = st.number_input("Espesor del ala hf [cm]", 5.0, 40.0, st.session_state.get("vm_hf", 12.0), 1.0, key="vm_hf")
                # Validaciones bf Viga T (NSR-10 CR.8.12)
                limite_bf_vm = min((vm_L * 100) / 4, vm_b + 16 * vm_hf)
                if vm_bf > limite_bf_vm:
                    st.warning(f"?? El ancho del ala ingresado (bf = {vm_bf} cm) excede el máximo efectivo. "
                               f"Se usará el límite bf_efectivo = {limite_bf_vm:.1f} cm (min[L/4, bw + 16hf]).")
                    vm_bf = limite_bf_vm
            else:
                vm_bf = vm_b
                vm_hf = vm_h
        with c3p1:
            # Estribo principal
            _est_opts_vm = ["Ø6mm","Ø8mm","Ø10mm","Ø12mm","#2","#3","#4"]
            vm_est = st.selectbox("Estribo transversal", _est_opts_vm,
                                  index=_est_opts_vm.index(st.session_state.get("vm_est","Ø8mm")) if st.session_state.get("vm_est","Ø8mm") in _est_opts_vm else 1,
                                  key="vm_est")
            vm_ramas = st.number_input("# Ramas del estribo", 2, 6, st.session_state.get("vm_ramas", 2), 1, key="vm_ramas")
            # Barra longitudinal
            varillas_vm = list(rebar_dict.keys())
            _def_vm = "#4 (Ø12.7mm)" if "Pulgadas" in bar_sys else "12mm"
            _idx_vm = varillas_vm.index(_def_vm) if _def_vm in varillas_vm else 1
            vm_bar = st.selectbox("Barra longitudinal", varillas_vm,
                                  index=varillas_vm.index(st.session_state.get("vm_bar", _def_vm)) if st.session_state.get("vm_bar", _def_vm) in varillas_vm else _idx_vm,
                                  key="vm_bar")

        #  Sincronizar session_state con módulos individuales 
        vm_d = vm_h - vm_dp
        # Bus de datos → actualiza automáticamente los módulos de abajo
        st.session_state["vr_b"]  = vm_b;  st.session_state["vr_h"]  = vm_h
        st.session_state["vr_dp"] = vm_dp; st.session_state["vr_L"]  = vm_L
        st.session_state["vt_bf"] = min(vm_bf, 300.0)
        st.session_state["vt_bw"] = vm_b
        # Limitar hf
        if vm_tipo == "Viga T":
            st.session_state["vt_hf"] = min(vm_hf, 40.0)
        st.session_state["vt_h"]  = vm_h
        st.session_state["vt_dp"] = vm_dp; st.session_state["vt_L"]  = vm_L
        st.session_state["cv_bw"] = vm_b;  st.session_state["cv_d"]  = vm_d
        st.session_state["cv_h"]  = vm_h;  st.session_state["cv_L"]  = vm_L
        st.session_state["de_b"]  = vm_b;  st.session_state["de_h"]  = vm_h
        st.session_state["de_dp"] = vm_dp; st.session_state["de_L"]  = vm_L
        
        st.session_state["vr_bar"] = vm_bar
        st.session_state["vt_bar"] = vm_bar

        st.success(f"Datos propagados: bw={vm_b:.0f} cm | h={vm_h:.0f} cm | d'={vm_dp:.0f} cm | L={vm_L:.1f} m | Sección: {vm_tipo}")

    # 
    # PASO 2 — CARGAS
    # 
    with vm_tabs[1]:
        st.subheader(" Solicitaciones de Diseño")

        modo_carga = st.radio(
            "Modo de ingreso de cargas:",
            [" Automático (calcula Mu desde wD+wL)", " Manual (ingresar Mu directamente)"],
            horizontal=True, key="vm_modo_carga"
        )

        if "Automático" in modo_carga:
            st.markdown("##### Cargas distribuidas")
            c_auto1, c_auto2 = st.columns(2)
            with c_auto1:
                vm_wD      = st.number_input("wD — Carga muerta [kN/m]", 0.0, 500.0, st.session_state.get("vm_wD", 12.0), 1.0, key="vm_wD")
                vm_wL_defl = st.number_input("wL — Carga viva  [kN/m]", 0.0, 500.0, st.session_state.get("vm_wL_defl", 8.0), 1.0, key="vm_wL_defl")
                vm_wS      = st.number_input("wS — Sismo equiv. [kN/m]", 0.0, 200.0, 0.0, 1.0, key="vm_wS")
            with c_auto2:
                combo_sel = st.selectbox(
                    "Combinación de carga (C.9.2)",
                    ["1.2D + 1.6L (governa flexión)", "1.4D", "1.2D + 1.0E + 1.0L (sísmica)", "0.9D + 1.0E"],
                    key="vm_combo"
                )
                cond_apoyo = st.selectbox(
                    "Condición de apoyos",
                    ["Ambos extremos continuos", "Un extremo continuo / otro discontinuo", "Ambos extremos discontinuos (simplemente apoyada)"],
                    key="vm_cond_apoyo"
                )

            if combo_sel == "1.4D":
                vm_wu = 1.4 * vm_wD
            elif combo_sel == "1.2D + 1.0E + 1.0L (sísmica)":
                vm_wu = 1.2 * vm_wD + 1.0 * vm_wS + 1.0 * vm_wL_defl
            elif combo_sel == "0.9D + 1.0E":
                vm_wu = 0.9 * vm_wD + 1.0 * vm_wS
            else:
                vm_wu = 1.2 * vm_wD + 1.6 * vm_wL_defl

            L2 = vm_L ** 2
            if "Ambos extremos continuos" in cond_apoyo:
                coef_izq, coef_cen, coef_der = 1/11, 1/16, 1/11
            elif "Un extremo continuo" in cond_apoyo:
                coef_izq, coef_cen, coef_der = 1/24, 1/14, 1/10
            else:
                coef_izq, coef_cen, coef_der = 0.0, 1/8, 0.0

            vm_mu_izq = coef_izq * vm_wu * L2 * factor_fuerza
            vm_mu_cen = coef_cen * vm_wu * L2 * factor_fuerza
            vm_mu_der = coef_der * vm_wu * L2 * factor_fuerza
            vm_vu_max = (vm_wu * vm_L / 2) * factor_fuerza

            st.markdown("---")
            st.markdown(f"**Wu calculado:** `{vm_wu:.2f} kN/m`")
            mc1, mc2, mc3, mc4 = st.columns(4)
            mc1.metric(f"Mu⁻ Izq [{unidad_mom}]", f"{vm_mu_izq:.1f}")
            mc2.metric(f"Mu⁺ Centro [{unidad_mom}]", f"{vm_mu_cen:.1f}")
            mc3.metric(f"Mu⁻ Der [{unidad_mom}]", f"{vm_mu_der:.1f}")
            mc4.metric(f"Vu máx [{unidad_fuerza}]", f"{vm_vu_max:.1f}")
            st.caption("**NSR-10 C.8.3.3** — Coeficientes de momentos para vigas continuas.")

        else:
            vm_wD      = st.session_state.get("vm_wD", 12.0)
            vm_wL_defl = st.session_state.get("vm_wL_defl", 8.0)
            c1p2, c2p2 = st.columns(2)
            with c1p2:
                vm_mu_izq = st.number_input(f"Mu⁻ Apoyo Izquierdo [{unidad_mom}]", 0.0, 50000.0, st.session_state.get("vm_mu_izq", 80.0), 5.0, key="vm_mu_izq")
                vm_mu_cen = st.number_input(f"Mu⁺ Centro de Luz [{unidad_mom}]",    0.0, 50000.0, st.session_state.get("vm_mu_cen", 60.0), 5.0, key="vm_mu_cen")
                vm_mu_der = st.number_input(f"Mu⁻ Apoyo Derecho [{unidad_mom}]",    0.0, 50000.0, st.session_state.get("vm_mu_der", 80.0), 5.0, key="vm_mu_der")
            with c2p2:
                vm_vu_max  = st.number_input(f"Vu máximo (apoyo) [{unidad_fuerza}]", 0.1, 10000.0, st.session_state.get("vm_vu_max", 80.0), 5.0, key="vm_vu_max")
                vm_wu      = st.number_input(f"Wu factorizada [kN/m]",              0.0, 2000.0, st.session_state.get("vm_wu", 20.0), 2.0, key="vm_wu")
                vm_wL_defl = st.number_input("wL [kN/m] (deflexión)", 0.0, 500.0, st.session_state.get("vm_wL_defl", 8.0), 1.0, key="vm_wL_defl")
                vm_wD      = st.number_input("wD [kN/m] (deflexión)", 0.0, 500.0, st.session_state.get("vm_wD", 12.0), 1.0, key="vm_wD")

        st.session_state["vr_mu"] = max(vm_mu_izq, vm_mu_cen, vm_mu_der)
        st.session_state["vt_mu"] = st.session_state["vr_mu"]
        st.session_state["cv_vu"] = vm_vu_max
        st.session_state["de_wD"] = vm_wD
        st.session_state["de_wL"] = vm_wL_defl

        st.markdown(f"**Mu máximo:** {st.session_state['vr_mu']:.2f} {unidad_mom} | **Vu:** {vm_vu_max:.2f} {unidad_fuerza}")

    # 
    # PASO 3 — RESULTADOS (COMPILACIÓN DE LOS DEMÁS MÓDULOS)
    # 
    with vm_tabs[2]:
        st.info(" **Flexión Rectangular** — Los datos de geometría ingresados arriba "
                "se propagan automáticamente al módulo **'Diseño a Flexión — Viga Rectangular'**. "
                "Selecciónalo en el selector de módulos para ver el cálculo completo, "
                "o desplázate hacia abajo si está expandido en esta página.")
        st.metric("Mu máximo activo", f"{st.session_state.get('vr_mu', 0.0):.2f} {unidad_mom}")
        st.metric("Sección propagada", f"bw={st.session_state.get('vr_b',0):.0f} × h={st.session_state.get('vr_h',0):.0f} cm")
    with vm_tabs[3]:
        st.info(" **Flexión Viga T** — Los datos se propagan al módulo "
                "'Diseño a Flexión — Viga T'. Selecciónalo en el navegador de módulos.")
        st.metric("bf propagado", f"{st.session_state.get('vt_bf',0):.0f} cm")
        st.metric("bw propagado", f"{st.session_state.get('vt_bw',0):.0f} cm")
    with vm_tabs[4]:
        st.info(" **Cortante** — Los datos se propagan al módulo "
                "'Diseño a Cortante — Vigas de Concreto'. Selecciónalo en el navegador.")
        st.metric("Vu máximo activo", f"{st.session_state.get('cv_vu', 0.0):.2f} {unidad_fuerza}")
    with vm_tabs[5]:
        st.info(" **Deflexiones** — Los datos se propagan al módulo "
                "'Inercia Fisurada y Deflexiones'. Selecciónalo en el navegador.")
        st.metric("wD activo", f"{st.session_state.get('de_wD', 0.0):.2f} kN/m")
        st.metric("wL activo", f"{st.session_state.get('de_wL', 0.0):.2f} kN/m")
    with vm_tabs[6]:
        st.subheader("Memoria Integrada de Viga")
        st.info("Se recomienda revisar los resultados de los módulos de abajo antes de imprimir la memoria completa. Asegúrese de haber calculado correctamente:")
        st.markdown("- **Flexión** (Módulo 1 o 2)\n- **Cortante** (Módulo 4)\n- **D. Sísmico** (Módulo 6 si aplica)")
        
        # Lógica de descarga unificada recopilando session_state
        if st.button("Generar Memoria Unificada de Diseño", key="btn_memoria_vm"):
            try:
                vm_doc = Document()
                vm_doc.add_heading("Memoria de Cálculo Estructural - Diseño de Viga", 0)
                vm_doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                vm_doc.add_paragraph(f"Norma Activa: {norma_sel} | Nivel de Diseño Sísmico: {nivel_sis}")
                
                vm_doc.add_heading("1. Materiales", level=1)
                vm_doc.add_paragraph(f"f'c = {fc:.1f} MPa ({fc_psi:.0f} psi)\nfy = {fy:.0f} MPa ({st.session_state.get('v_fy', 0)} psi)")
                
                vm_doc.add_heading("2. Geometría", level=1)
                vm_doc.add_paragraph(f"Sección: {vm_tipo}\nAncho alma (bw) = {vm_b:.1f} cm\nAltura (h) = {vm_h:.1f} cm")
                if vm_tipo == "Viga T":
                    vm_doc.add_paragraph(f"Ancho ala (bf) = {vm_bf:.1f} cm\nEspesor ala (hf) = {vm_hf:.1f} cm")
                vm_doc.add_paragraph(f"Luz libre = {vm_L:.2f} m\nRecubrimiento (d') = {vm_dp:.1f} cm")
                
                vm_doc.add_heading("3. Solicitaciones", level=1)
                vm_doc.add_paragraph(f"Mu Izquierdo = {vm_mu_izq:.2f} {unidad_mom}\nMu Centro = {vm_mu_cen:.2f} {unidad_mom}\nMu Derecho = {vm_mu_der:.2f} {unidad_mom}")
                vm_doc.add_paragraph(f"Cortante máximo Vu = {vm_vu_max:.2f} {unidad_fuerza}\nCarga Distribuida Wu = {vm_wu:.2f} kN/m")
                
                vm_doc.add_heading("4. Diseño Longitudinal (As provisto)", level=1)
                vm_doc.add_paragraph(f"Cálculos de Acero Long.\nApoyo Izquierdo: {vm_mu_izq:.2f} kN·m -> As,req = {st.session_state.get('v_As_req_izq', 0):.2f} cm²\nCentro: {vm_mu_cen:.2f} kN·m -> As,req = {st.session_state.get('v_As_req_cen', 0):.2f} cm²\nApoyo Derecho: {vm_mu_der:.2f} kN·m -> As,req = {st.session_state.get('v_As_req_der', 0):.2f} cm²")
                
                vm_doc.add_heading("5. Diseño Transversal (Cortante)", level=1)
                vm_doc.add_paragraph(f"Estribo de diseño base = {vm_est} con {vm_ramas} ramas.")
                
                vm_doc_mem = io.BytesIO()
                vm_doc.save(vm_doc_mem)
                vm_doc_mem.seek(0)
                
                st.download_button(
                    "Descargar Archivo DOCX", 
                    data=vm_doc_mem, 
                    file_name=f"Memoria_Completa_{vm_tipo}_{vm_b:.0f}x{vm_h:.0f}.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_down_memoria_vm"
                )
                st.success("¡Memoria generada exitosamente! Presione el botón arriba para descargarla.")
            except Exception as e:
                st.error(f"Error generando memoria: {e}")



# ??????????????????????????????????????????
# 2. VIGA RECTANGULAR — FLEXIÓN
# ??????????????????????????????????????????
if modulo_sel == " Diseño a Flexión — Viga Rectangular":
    #  PANEL 
    _panel_normativo("Diseño a Flexión — Viga Rectangular",
        "Calcula el acero longitudinal As requerido para una sección simplemente reforzada.",
        ["Base b (cm)", "Altura total h (cm)", "Recubrimiento d' (cm)",
         "Momento último Mu (kN·m o tonf·m)", "Longitud de viga L (m)",
         "Diámetro de varilla longitudinal"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(_t(f"**Método (sección simplemente reforzada)** | Norma: `{code['ref']}`", f"**Method (singly reinforced section)** | Code: `{code['ref']}`"))
    st.info(_t("**Modo de uso:** Ingresa la base, altura y recubrimiento de la viga. Añade el momento flector (Mu) a soportar. Luego selecciona el diámetro de varilla y la App te dirá si cumple flexión, calculará la cantidad requerida de acero, y generará las cantidades y precios de todo el pórtico.", " **How to use:** Enter beam base, height and cover. Add ultimate moment (Mu). Then select rebar diameter and the App will check flexure, calculate required steel, quantities and prices."))
    c1,c2 = st.columns(2)
    with c1:
        b_vr = st.number_input("Ancho b [cm]", 15.0, 150.0, st.session_state.get("vr_b", 30.0), 5.0, key="vr_b")
        h_vr = st.number_input("Alto h [cm]", 20.0, 200.0, st.session_state.get("vr_h", 50.0), 5.0, key="vr_h")
        dp_vr = st.number_input("Recubrim. d' [cm]", 2.0, 15.0, st.session_state.get("vr_dp", 5.0), 0.5, key="vr_dp")
        Mu_vr = st.number_input(f"Momento último Mu [{unidad_mom}]", 0.1, 10000.0, st.session_state.get("vr_mu", 80.0), 5.0, key="vr_mu")
    with c2:
        L_vr = st.number_input("Longitud viga [m]", 1.0, 30.0, st.session_state.get("vr_L", 5.0), 0.5, key="vr_L")
        varillas_vr = list(rebar_dict.keys())
        # Mínimo práctico: #4 (12.7mm) en US  |  12mm en SI  — el #3/8mm se reserva para estribos
        _def_vr = "#4 (Ø12.7mm)" if "Pulgadas" in bar_sys else "12mm"
        _def_idx_vr = varillas_vr.index(_def_vr) if _def_vr in varillas_vr else 1
        bar_vr = st.selectbox(
            "Varilla longitudinal (mín. recomendado #4 / 12mm — el #3 se usa solo para estribos):",
            varillas_vr,
            index=varillas_vr.index(st.session_state.vr_bar) if "vr_bar" in st.session_state and st.session_state.vr_bar in varillas_vr else _def_idx_vr,
            key="vr_bar")
        Ab_vr = rebar_dict[bar_vr]; db_vr = diam_dict[bar_vr]

    d_vr = h_vr - dp_vr
    d_mm = d_vr*10; b_mm = b_vr*10
    # Convertir Mu a kN·m si es necesario
    if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
        Mu_vr_kN = Mu_vr / factor_fuerza
    else:
        Mu_vr_kN = Mu_vr
    Mu_Nmm = Mu_vr_kN * 1e6
    if d_mm > 0 and b_mm > 0:
        Rn = Mu_Nmm / (phi_f * b_mm * d_mm**2)
        disc = 1 - 2*Rn/(0.85*fc)
        if disc < 0:
            st.error("? Sección insuficiente – aumente b o h")
        else:
            rho_calc = (0.85*fc/fy)*(1 - math.sqrt(disc))
            rho_use = max(rho_calc, rho_min)
            As_req_cm2 = rho_use * b_vr * d_vr
            n_bars = math.ceil(As_req_cm2 / Ab_vr)

            #  ESPACIAMIENTO Y FILAS (NSR-10 C.7.6.1 / C.7.6.2) 
            db_e_vr   = 9.5      # mm — diám. estribo estándar Ø3/8"
            s_min_vr  = max(25.0, db_vr)           # mm — sep. min. entre barras
            sep_filas_mm_vr = 25.0                  # mm — sep. vert. entre capas
            # Máx. barras en 1ª fila dentro del ancho disponible
            _ancho_disp_mm = b_vr*10 - 2*dp_vr*10 - 2*db_e_vr - db_vr
            n_max_f1_vr = max(1, int(_ancho_disp_mm / (db_vr + s_min_vr)) + 1)
            
            d_eff_mm = d_mm
            dos_filas_vr = False

            if n_bars > n_max_f1_vr:
                for _iter in range(4):
                    n_f1_vr = n_max_f1_vr
                    n_f2_vr = n_bars - n_f1_vr
                    
                    y1_vr = dp_vr*10
                    y2_vr = dp_vr*10 + db_vr + sep_filas_mm_vr
                    y_cg_vr = (n_f1_vr*Ab_vr*y1_vr + n_f2_vr*Ab_vr*y2_vr) / (n_bars*Ab_vr)
                    d_eff_mm_nuevo = h_vr*10 - y_cg_vr
                    
                    if abs(d_eff_mm_nuevo - d_eff_mm) < 1.0:
                        d_eff_mm = d_eff_mm_nuevo
                        break
                        
                    d_eff_mm = d_eff_mm_nuevo
                    Rn_iter = Mu_Nmm / (phi_f * b_mm * d_eff_mm**2)
                    disc_iter = 1 - 2*Rn_iter/(0.85*fc)
                    if disc_iter < 0:
                        break
                    
                    rho_calc_iter = (0.85*fc/fy)*(1 - math.sqrt(disc_iter))
                    rho_use_iter = max(rho_calc_iter, rho_min)
                    As_req_cm2_iter = rho_use_iter * b_vr * (d_eff_mm / 10)
                    n_bars = math.ceil(As_req_cm2_iter / Ab_vr)
                    if n_bars <= n_max_f1_vr:
                        # El recalculo con d_eff redujo el requerimiento a 1 fila.
                        # Advertimos al usuario y salimos — no forzar barra extra.
                        dos_filas_vr = False
                        n_f1_vr = n_bars
                        n_f2_vr = 0
                        d_eff_mm = d_mm  # volver al d original (1 fila)
                        st.info("ℹ Tras iterar con d_efectivo, el acero requerido cabe en **1 fila** con el d'original.")
                        break
                        
                dos_filas_vr = True
            else:
                n_f1_vr, n_f2_vr = n_bars, 0
                y1_vr = dp_vr * 10
                y2_vr = y1_vr
                d_eff_mm = d_mm
                dos_filas_vr = False

            As_prov  = n_bars * Ab_vr
            rho_prov = As_prov / (b_vr * (d_eff_mm / 10))

            a_mm     = As_prov*100*fy/(0.85*fc*b_mm)
            phi_Mn_kNm = phi_f * As_prov*100*fy*(d_eff_mm - a_mm/2)/1e6
            ok_flex    = phi_Mn_kNm >= Mu_vr_kN
            ok_rho_min = rho_prov >= rho_min
            ok_rho_max = rho_prov <= rho_max
            tab_r, tab_s, tab_3d, tab_q = st.tabs([f"Resultados {''if (ok_flex and ok_rho_min and ok_rho_max) else '?'}"," Sección 2D"," Visualización 3D"," Cantidades"])
            with tab_r:
                st.markdown(f"**Factor de reducción φ = {phi_f}** (flexión) | Norma: `{code['ref']}`")
                st.markdown("""**Verificación fundamental:** La resistencia a flexión provista **φMn** debe ser mayor o igual al momento último demandado **Mu**.
> φMn ≥ Mu   (la viga resiste sin colapsar)""")
                st.latex(r"\\phi M_n = \\phi \cdot A_s \cdot f_y \\left(d - \frac{a}{2}\right)")
                rows = [
                    (" b × h — Base y altura de la viga", f"{b_vr:.0f} × {h_vr:.0f} cm"),
                    (" d — Peralte efectivo real (centroide del acero)" + ("   2 filas" if dos_filas_vr else ""), f"{d_eff_mm/10:.2f} cm" + (f" (original: {d_vr:.1f} cm)" if dos_filas_vr else "")),
                    (" Rn — Resistencia unitaria requerida (Mu / φ·b·d²)", f"{Rn:.3f} MPa"),
                    (" ρ calculado — Cuantía de acero que necesita la sección", f"{rho_calc*100:.4f}%"),
                    (" ρ mínimo — Cuantía mínima exigida por la norma (evita falla frágil)", f"{rho_min*100:.4f}%"),
                    (" ρ máximo — Cuantía máxima (garantiza falla dúctil con aviso)", f"{rho_max*100:.4f}%"),
                    ("ρ provisto", f"{(As_prov/(b_vr*(d_eff_mm/10)))*100:.4f}%"),
                    (" As requerido — Área de acero necesaria para resistir Mu", f"{As_req_cm2:.3f} cm²"),
                    (f" Varillas seleccionadas ({bar_vr}) — Cantidad y área provista", f"{n_bars} barras → As provisto = {As_prov:.3f} cm²"),
                    (" a — Profundidad del bloque de compresión equivalente (Whitney)", f"{a_mm:.1f} mm"),
                    (f" φMn — Momento resistente provisto [{unidad_mom}]", f"{phi_Mn_kNm*factor_fuerza:.2f}"),
                    (f" Mu — Momento último demandado (carga de diseño) [{unidad_mom}]", f"{Mu_vr:.2f}"),
                    (" Verificación Flexión (φMn ≥ Mu)" if ok_flex else "? Verificación Flexión (φMn < Mu)",
                     " CUMPLE — La viga resiste el momento de diseño" if ok_flex else f"? DEFICIENTE — φMn={phi_Mn_kNm:.2f} < Mu={Mu_vr_kN:.2f} → Aumente sección o acero"),
                    (" Cuantía mínima (? ≥ ?_min)" if ok_rho_min else "? Cuantía mínima (? < ?_min)",
                     " CUMPLE" if ok_rho_min else "? NO CUMPLE — Aumente el área de acero"),
                    (" Cuantía máxima (? ≤ ?_max)" if ok_rho_max else "? Cuantía máxima (? > ?_max)",
                     " CUMPLE" if ok_rho_max else "? EXCEDE M?XIMO — Sección sobrearmada, amplíe la sección"),
                ]
                qty_table(rows)

                #  Verificación de Ancho / Filas 
                if dos_filas_vr:
                    st.warning(
                        f" ? **2 FILAS REQUERIDAS (NSR-10 C.7.6.1 / C.7.6.2):** "
                        f"Las {n_bars} varillas {bar_vr} no caben en 1 fila (máx. {n_max_f1_vr} por fila). "
                        f"Se usa: **{n_f1_vr} barras en fila 1** + **{n_f2_vr} barras en fila 2** "
                        f"(sep. vertical = {sep_filas_mm_vr:.0f} mm). "
                        f"**d efectivo real = {d_eff_mm/10:.2f} cm** (reducido por centroide del acero)."
                    )
                else:
                    _bw_min_req = (2*dp_vr*10 + n_f1_vr*db_vr + (n_f1_vr-1)*s_min_vr) / 10
                    st.success(f"1 fila: {n_f1_vr} varillas {bar_vr} | d efectivo = {d_eff_mm/10:.2f} cm | bw mín = {_bw_min_req:.1f} cm")

                #  Diagrama de Envolvente φMn(As(x)) vs Mu 
                st.markdown("---")
                st.subheader("Diagrama de Capacidad vs Demanda")
                st.caption("Relación entre la capacidad de la sección (φMn) y el momento demandado (Mu)")

                # Construir envolvente: variando As de 0 a As_prov (0..n_bars varillas)
                _n_pts = 30
                _as_vals = np.linspace(0.01, As_prov * 2.5, _n_pts)  # rango de As
                _phi_mn_env = []
                for _as in _as_vals:
                    _a_i = _as * 100 * fy / (0.85 * fc * b_mm)
                    _phi_mn_env.append(phi_f * _as * 100 * fy * (d_mm - _a_i / 2) / 1e6 * factor_fuerza)

                fig_env = go.Figure()
                # Capacidad φMn vs As
                fig_env.add_trace(go.Scatter(
                    x=[_as for _as in _as_vals], y=_phi_mn_env,
                    name=f'φMn (capacidad)', line=dict(color='#00d4ff', width=2.5)
                ))
                # Línea de Mu horizontal
                fig_env.add_hline(
                    y=Mu_vr, line_dash='dash', line_color='#ff6b35', line_width=2,
                    annotation_text=f'Mu = {Mu_vr:.2f} {unidad_mom}', annotation_position='right'
                )
                # Punto diseño
                fig_env.add_trace(go.Scatter(
                    x=[As_prov], y=[phi_Mn_kNm * factor_fuerza],
                    mode='markers+text',
                    marker=dict(size=12, color='#ff6b35' if not ok_flex else '#44ff88', symbol='star'),
                    text=[f'φMn={phi_Mn_kNm*factor_fuerza:.2f}'],
                    textposition='top center', name='Punto diseño'
                ))
                # Línea vertical As provisto
                fig_env.add_vline(
                    x=As_prov, line_dash='dot', line_color='#44ff88', line_width=1.5,
                    annotation_text=f'As = {As_prov:.2f} cm²', annotation_position='top right'
                )
                fig_env.update_layout(
                    xaxis_title=f'As provisto (cm²) — {n_bars} var. {bar_vr}',
                    yaxis_title=f'φMn ({unidad_mom})',
                    paper_bgcolor='#1a1a2e', plot_bgcolor='#16213e',
                    font=dict(color='white'), height=320,
                    legend=dict(bgcolor='rgba(0,0,0,0.5)', font=dict(color='white')),
                    margin=dict(l=0, r=0, t=10, b=0)
                )
                st.plotly_chart(fig_env, use_container_width=True)

                if ok_flex and ok_rho_min and ok_rho_max:
                    st.success(f"Diseño Aprobado: φMn = {phi_Mn_kNm*factor_fuerza:.2f} {unidad_mom}  ≥  Mu = {Mu_vr:.2f} {unidad_mom}")
                else:
                    msg = []
                    if not ok_flex: msg.append(f"φMn={phi_Mn_kNm*factor_fuerza:.2f} < {Mu_vr:.2f}")
                    if not ok_rho_min: msg.append(f"?={rho:.5f} < ?_min={rho_min:.5f}")
                    if not ok_rho_max: msg.append(f"?={rho:.5f} > ?_max={rho_max:.5f}")
                    _detalle = " | ".join(msg) if msg else "Cuantía fuera de rango"
                    st.error(f"? Diseño No Aprobado — {_detalle} → Aumente la sección o el acero")
                st.info("**¿El acero calculado es inferior o superior?** Si el Mu ingresado viene de una combinación con **momento positivo** (vano central), el acero corresponde al **refuerzo inferior** (zona en tensión debajo). Si Mu viene de un **momento negativo** (apoyo o empotramiento), el acero es el **refuerzo superior**.")

            with tab_s:
                _titulo_2d = f"Sección {b_vr:.0f}×{h_vr:.0f} cm"
                if dos_filas_vr:
                    _titulo_2d += f" — 2 filas ({n_f1_vr}+{n_f2_vr})"
                fig, ax = sec_dark_fig(b_vr, h_vr, _titulo_2d)
                recub = max(dp_vr - db_vr/20, 0.5)
                r_bar = db_vr / 20
                ax.add_patch(patches.Rectangle((recub,recub),b_vr-2*recub,h_vr-2*recub,
                                               linewidth=1.5,edgecolor='#00d4ff',facecolor='none',linestyle='--'))
                # Fila 1 (inferior)
                y_f1_cm = dp_vr
                xs_f1 = np.linspace(dp_vr, b_vr-dp_vr, n_f1_vr) if n_f1_vr > 1 else [b_vr/2]
                for x in xs_f1:
                    ax.add_patch(plt.Circle((x, y_f1_cm), r_bar, color='#ff6b35', zorder=5))
                # Fila 2 (si aplica)
                if dos_filas_vr and n_f2_vr > 0:
                    y_f2_cm = dp_vr + db_vr/10 + sep_filas_mm_vr/10
                    xs_f2 = np.linspace(dp_vr, b_vr-dp_vr, n_f2_vr) if n_f2_vr > 1 else [b_vr/2]
                    for x in xs_f2:
                        ax.add_patch(plt.Circle((x, y_f2_cm), r_bar, color='#ffa040', zorder=5))
                    # línea centroide real
                    ax.axhline(y=h_vr - d_eff_mm/10, color='lime', linestyle=':', linewidth=1, alpha=0.7)
                    ax.text(b_vr*0.02, h_vr - d_eff_mm/10 + 0.3, f"d={d_eff_mm/10:.1f}cm",
                            color='lime', fontsize=6, va='bottom')
                ax.annotate('',xy=(b_vr,-0.8*h_vr/h_vr),xytext=(0,-0.8*h_vr/h_vr),arrowprops=dict(arrowstyle='<->',color='white'))
                ax.text(b_vr/2,-h_vr*0.12,f"b={b_vr:.0f}cm",ha='center',va='top',color='white',fontsize=7)
                ax.annotate('',xy=(-0.8,h_vr),xytext=(-0.8,0),arrowprops=dict(arrowstyle='<->',color='white'))
                ax.text(-h_vr*0.15,h_vr/2,f"h={h_vr:.0f}cm",ha='right',va='center',color='white',fontsize=7,rotation=90)
                st.pyplot(fig)
                plt.close(fig)
                _cap = f"{n_f1_vr} barras (fila 1) + {n_f2_vr} barras (fila 2) {bar_vr} | As={As_prov:.3f} cm² | d_eff={d_eff_mm/10:.2f} cm" if dos_filas_vr else f"{n_bars} varillas {bar_vr} | As={As_prov:.3f} cm²"
                st.caption(_cap)

            with tab_3d:
                st.subheader("Visualización 3D de Viga Rectangular")
                fig3d = go.Figure()
                L_mm_3d = L_vr * 100
                x_c = [-b_vr/2, b_vr/2, b_vr/2, -b_vr/2, -b_vr/2, b_vr/2, b_vr/2, -b_vr/2]
                y_c = [0, 0, h_vr, h_vr, 0, 0, h_vr, h_vr]
                z_c = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
                fig3d.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
                diam_reb_cm = db_vr / 10.0
                line_width = max(4, diam_reb_cm * 3)
                # Fila 1 → barras inferiores
                x_ini_3dvr = -b_vr/2 + dp_vr;  x_fin_3dvr = b_vr/2 - dp_vr
                xs_f1_3d = np.linspace(x_ini_3dvr, x_fin_3dvr, n_f1_vr) if n_f1_vr > 1 else [0.0]
                y_pos_f1 = y1_vr / 10
                for idx, x_pos in enumerate(xs_f1_3d):
                    fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[y_pos_f1, y_pos_f1], z=[0, L_mm_3d],
                                                mode='lines', line=dict(color='darkred', width=line_width),
                                                name=f'Varilla {bar_vr} (fila 1)', showlegend=(idx==0)))
                # Fila 2 → si hay segunda capa
                if dos_filas_vr and n_f2_vr > 0:
                    xs_f2_3d = np.linspace(x_ini_3dvr, x_fin_3dvr, n_f2_vr) if n_f2_vr > 1 else [0.0]
                    y_pos_f2 = y2_vr / 10
                    for idx, x_pos in enumerate(xs_f2_3d):
                        fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[y_pos_f2, y_pos_f2], z=[0, L_mm_3d],
                                                    mode='lines', line=dict(color='orangered', width=line_width),
                                                    name=f'Varilla {bar_vr} (fila 2)', showlegend=(idx==0)))
                tie_color = 'cornflowerblue'
                tie_width = max(2, (9.5/10.0) * 3)
                # Bug Fix: usar el s_diseno calculado en el módulo Cortante (cv_s_diseno)
                # NO usar cv_sdiseno que era el key incorrecto — el key real es cv_s_diseno
                _default_sep = int(round(st.session_state.get('cv_s_diseno', 15)))
                sep_ties = st.slider("Separación Estribos (cm)", 5, 50, max(5, min(50, _default_sep)), 1, key="vr_sep_tie",
                                     help="Valor por defecto tomado del módulo Cortante. Ajuste manualmente si lo desea.")
                tx = [-b_vr/2 + dp_vr/2, b_vr/2 - dp_vr/2, b_vr/2 - dp_vr/2, -b_vr/2 + dp_vr/2, -b_vr/2 + dp_vr/2]
                ty = [dp_vr/2, dp_vr/2, h_vr - dp_vr/2, h_vr - dp_vr/2, dp_vr/2]
                L_cm = int(L_mm_3d)
                tx_all, ty_all, tz_all = [], [], []
                for zt in range(5, L_cm, sep_ties):
                    for px, py in zip(tx, ty):
                        tx_all.append(px); ty_all.append(py); tz_all.append(zt)
                    tx_all.append(None); ty_all.append(None); tz_all.append(None)
                for px, py in zip(tx, ty):
                    tx_all.append(px); ty_all.append(py); tz_all.append(L_cm - 5)
                tx_all.append(None); ty_all.append(None); tz_all.append(None)
                fig3d.add_trace(go.Scatter3d(x=tx_all, y=ty_all, z=tz_all, mode='lines', 
                                             line=dict(color=tie_color, width=tie_width), name='Estribos', showlegend=True))
                fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                                    margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable')
                st.plotly_chart(fig3d, use_container_width=True)

            with tab_q:
                vol_horm = b_vr/100*h_vr/100*L_vr
                peso_long = As_prov * L_vr * 0.785
                m = mix_for_fc(fc)
                bags = m[0]/bag_kg * vol_horm
                rows_q = [
                    ("Concreto (b×h×L)", f"{vol_horm:.4f} m³"),
                    (f"Acero longitudinal ({n_bars} barras)", f"{peso_long:.2f} kg"),
                    (f"Cemento ({bag_kg:.0f} kg/bulto, f'c={fc:.1f} MPa)", f"{bags:.1f} bultos = {m[0]*vol_horm:.0f} kg"),
                    ("Arena", f"{(m[2]*vol_horm)/1600:.2f} m³"),
                    ("Grava", f"{(m[3]*vol_horm)/1600:.2f} m³"),
                    ("Agua", f"{m[1]*vol_horm:.0f} L"),
                    ("Referencia", code["ref"]),
                ]
                qty_table(rows_q)

                render_apu_breakdown(vol_horm, peso_long, fc, f"({n_bars} barras longitudinales)")

                with st.expander("Dibujo de Figurado para Taller", expanded=False):
                    st.markdown("A continuación se muestran las formas reales de las barras para facilitar el figurado.")
                    hook_len_cm = 12 * db_vr / 10
                    straight_len_cm = L_vr * 100 - 2 * hook_len_cm
                    fig_l1 = draw_longitudinal_bar(L_vr*100, straight_len_cm, hook_len_cm, db_vr)
                    st.pyplot(fig_l1)
                    plt.close(fig_l1)
                    recub_est = max(dp_vr, 2.5)
                    inside_b = b_vr - 2*recub_est
                    inside_h = h_vr - 2*recub_est
                    hook_len_est = 12 * 9.5 / 10  # aprox. 12db para estribo (visual)
                    fig_e1 = draw_stirrup_beam(inside_b, inside_h, hook_len_est, 9.5)
                    st.pyplot(fig_e1)
                    plt.close(fig_e1)
                    st.caption("Nota: Los ganchos de estribos son de 135° en la práctica. En el dibujo se representa de forma esquemática con líneas inclinadas.")

                # MEMORIA DOCX para Viga Rectangular
                col_d1, col_d2, col_d3 = st.columns(3)
                with col_d1:
                    btn_docx_vr = st.button("Generar Memoria DOCX (Viga Rectangular)")
                with col_d2:
                    if st.button("Enviar a Cuadro de Mando", key="cmd_vr"):
                        add_historial_diseno("Flexión Viga Rect.", f"b={b_vr:.0f} h={h_vr:.0f} cm", " APROBADO" if (ok_flex and ok_rho_min and ok_rho_max) else "? NO CUMPLE", code['ref'])
                with col_d3:
                    try:
                        n1 = n_f1_vr if dos_filas_vr else int(n_bars)
                        n2 = n_f2_vr if dos_filas_vr else 0
                        buf_ifc = ifc_export.ifc_viga_rectangular(
                            b_vr, h_vr, L_vr, fc, fy, int(n_bars), bar_vr, As_prov,
                            db_vr, d_eff_mm/10., dp_vr, Mu_vr, phi_Mn_kNm,
                            dos_filas_vr, n1, n2, sep_filas_mm_vr, norma_sel, "Proyecto NSR-10",
                            db_est_mm=db_e_vr, sep_est_mm=sep_ties*10.0
                        )
                        st.download_button("Exportar IFC (BIM)", data=buf_ifc, file_name=f"Viga_Rect_{b_vr:.0f}x{h_vr:.0f}.ifc", mime="application/x-step", key="ifc_vr")
                    except Exception as e:
                        st.error(f"Error IFC: {e}")

                if btn_docx_vr:
                    doc = Document()
                    doc.add_heading(f"MEMORIA ESTRUCTURAL — VIGA RECTANGULAR {b_vr:.0f}×{h_vr:.0f} cm ({norma_sel})", 0)

                    # 0. Portada
                    p0 = doc.add_paragraph()
                    p0.add_run(f"Proyecto: \nElemento: Viga Rectangular\nNorma: {norma_sel}\nNivel Sísmico: {nivel_sis}\nMateriales: Concreto f'c = {fc:.1f} MPa, Acero fy = {fy:.0f} MPa\nFecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}").bold = True

                    # 1. Parámetros de Diseño
                    doc.add_heading("1. PARÁMETROS DE DISEÑO", level=1)
                    doc.add_paragraph(f"Geometría: b = {b_vr:.0f} cm, h = {h_vr:.0f} cm, Recubrimiento al eje (dp) = {dp_vr:.1f} cm")
                    doc.add_paragraph(f"d efectivo real = {d_eff_mm/10:.2f} cm" + (f" (centroide 2 filas)" if dos_filas_vr else ""))
                    doc.add_paragraph(f"Cargas y solicitaciones:\n  Momento de Diseño (Mu) = {Mu_vr_kN:.2f} kN·m")

                    # 2. Pre-dimensionamiento y Esbeltez
                    doc.add_heading("2. PRE-DIMENSIONAMIENTO Y VERIFICACIÓN", level=1)
                    if dos_filas_vr:
                        doc.add_paragraph(f"Armado en 2 FILAS necesario: {n_f1_vr} en fila 1 + {n_f2_vr} en fila 2")
                        doc.add_paragraph(f"Separación vertical entre capas: {sep_filas_mm_vr:.0f} mm (mín 25 mm)")
                    else:
                        _sep_min_doc = max(25, db_vr)
                        _bw_min_doc = (2*dp_vr*10 + n_bars*db_vr + (n_bars-1)*_sep_min_doc) / 10
                        doc.add_paragraph(f"Verificación de Ancho Mínimo: bw_min = {_bw_min_doc:.1f} cm")
                        doc.add_paragraph(f"b = {b_vr:.0f} cm ≥ bw_min = {_bw_min_doc:.1f} cm — {'CUMPLE' if b_vr >= _bw_min_doc else 'NO CUMPLE'}")

                    # 3. Diseño a Flexión / Flexocompresión
                    doc.add_heading("3. DISEÑO A FLEXIÓN", level=1)
                    doc.add_paragraph(f"Índice de Resistencia Rn = Mu / (φ·b·d²) = {Rn:.3f} MPa (φ = {phi_f})")
                    doc.add_paragraph(f"Cuantía provista (ρ) = {rho_prov*100:.3f}%  (Límites: ρ_min = {rho_min*100:.3f}%, ρ_max = {rho_max*100:.3f}%)")
                    doc.add_paragraph(f"Acero Requerido/Provisto: {n_bars} varillas {bar_vr} → As = {As_prov:.3f} cm² (Profundidad bloque a = {a_mm:.1f} mm)")
                    doc.add_paragraph(f"Momento Resistente φMn = {phi_Mn_kNm:.2f} kN·m  ≥  Mu = {Mu_vr_kN:.2f} kN·m")
                    
                    doc.add_heading("3.1 Diagrama Capacidad vs Demanda", level=2)
                    fig_mn, ax_mn = plt.subplots(figsize=(5, 3))
                    fig_mn.patch.set_facecolor('white')
                    ax_mn.plot(_as_vals, _phi_mn_env, color='steelblue', lw=2, label='φMn (As)')
                    ax_mn.axhline(Mu_vr, color='red', linestyle='--', lw=1.5, label=f'Mu={Mu_vr:.2f} {unidad_mom}')
                    ax_mn.scatter([As_prov], [phi_Mn_kNm * factor_fuerza], color='green', zorder=5, s=60)
                    ax_mn.set_xlabel('As provisto (cm²)'); ax_mn.set_ylabel(f'φMn ({unidad_mom})')
                    ax_mn.legend(fontsize=8); ax_mn.grid(alpha=0.3)
                    fig_mn.tight_layout()
                    buf_mn = io.BytesIO()
                    fig_mn.savefig(buf_mn, format='png', dpi=150, bbox_inches='tight')
                    buf_mn.seek(0)
                    plt.close(fig_mn)
                    from docx.shared import Inches
                    doc.add_picture(buf_mn, width=Inches(4.5))

                    # 4. Diseño a Cortante / Confinamiento
                    doc.add_heading("4. LONGITUDES DE DESARROLLO (ld y ldh)", level=1)
                    _psi_t = 1.0; _psi_e = 1.0; _lambda = 1.0
                    _ld_mm = max(300, (fy * _psi_t * _psi_e) / (1.1 * _lambda * math.sqrt(fc)) * db_vr / 100)
                    _ldh_mm = max(150, (0.24 * fy * _psi_e) / (_lambda * math.sqrt(fc)) * db_vr / 100 * 10)
                    doc.add_paragraph(f"ld (recta) = {_ld_mm/10:.1f} cm  |  ldh (gancho) = {_ldh_mm/10:.1f} cm")

                    # 5. Verificaciones Normativas
                    doc.add_heading("5. VERIFICACIONES NORMATIVAS", level=1)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Parámetro Verificado'
                    hdr_cells[1].text = 'Estado'
                    hdr_cells[2].text = 'Valor'
                    
                    def _add_row(param, status, val):
                        row_cells = table.add_row().cells
                        row_cells[0].text = param
                        row_cells[1].text = "CUMPLE" if status else "NO CUMPLE"
                        row_cells[2].text = val
                    
                    _add_row("Resistencia a Flexión", ok_flex, f"φMn={phi_Mn_kNm:.2f} ≥ Mu={Mu_vr_kN:.2f}")
                    _add_row("Cuantía de Acero", (ok_rho_min and ok_rho_max), f"ρ={rho_prov*100:.2f}% (Entre mín y máx)")
                    _add_row("Ancho Mínimo de Viga", _ok_bw, f"b={b_vr}cm ≥ bw_min={getattr(sys.modules[__name__], '_bw_min_doc', b_vr):.1f}cm")

                    # 6. Cuantificación de Materiales
                    doc.add_heading("6. CUANTIFICACIÓN DE MATERIALES", level=1)
                    doc.add_paragraph(f"Volumen concreto: {vol_horm:.4f} m³")
                    doc.add_paragraph(f"Acero longitudinal superior/inferior principal: {peso_long:.2f} kg  ({n_bars} var. {bar_vr} × {L_vr:.2f} m)")
                    doc.add_paragraph(f"Mezcla sugerida:\n  Cemento: {bags:.1f} bultos\n  Arena: {(m[2]*vol_horm)/1600:.2f} m³\n  Grava: {(m[3]*vol_horm)/1600:.2f} m³\n  Agua: {m[1]*vol_horm:.0f} L")

                    # 7. Detalles Generales
                    doc.add_heading("7. PLANOS Y DETALLES EMBEBIDOS", level=1)
                    _titulo_2d_doc = f"Sección Viga {b_vr:.0f}×{h_vr:.0f} cm" + (f" (2 filas)" if dos_filas_vr else "")
                    fig2d, ax2d = sec_light_fig(b_vr, h_vr, _titulo_2d_doc)
                    rbar = db_vr / 20
                    y_f1_cm = dp_vr
                    xsf1 = np.linspace(dp_vr, b_vr - dp_vr, n_f1_vr) if n_f1_vr > 1 else [b_vr/2]
                    for x in xsf1: ax2d.add_patch(plt.Circle((x, y_f1_cm), rbar, color='red', zorder=5))
                    if dos_filas_vr and n_f2_vr > 0:
                        y_f2_cm = dp_vr + db_vr/10 + sep_filas_mm_vr/10
                        xs_f2 = np.linspace(dp_vr, b_vr - dp_vr, n_f2_vr) if n_f2_vr > 1 else [b_vr/2]
                        for x in xs_f2: ax2d.add_patch(plt.Circle((x, y_f2_cm), rbar, color='red', zorder=5))
                    buf2d = io.BytesIO()
                    fig2d.savefig(buf2d, format='png', dpi=150, bbox_inches='tight')
                    buf2d.seek(0)
                    plt.close(fig2d)
                    doc.add_picture(buf2d, width=Inches(3))

                    # 8. Referencias
                    doc.add_heading("8. REFERENCIAS", level=1)
                    doc.add_paragraph("NSR-10 / ACI 318 - Disposiciones de diseño a flexión, anclaje y ductilidad.")
                    doc.add_paragraph("\n\n_________________________________________\nFirma Ing. Responsable")
                    doc.add_paragraph("Matrícula Profesional: _______________")

                    doc_mem = io.BytesIO()
                    doc.save(doc_mem)
                    doc_mem.seek(0)
                    st.download_button("Descargar Memoria DOCX", data=doc_mem, file_name=f"Memoria_Viga_{b_vr:.0f}x{h_vr:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")




# ??????????????????????????????????????????
# 3. VIGA T — FLEXIÓN
# ??????????????????????????????????????????
if modulo_sel == " Diseño a Flexión — Viga T":
    #  PANEL 
    _panel_normativo("Diseño a Flexión — Viga T",
        "Diseño con ala colaborante. Verifica el ancho efectivo bf según NSR-10 C.R.8.12.",
        ["Ancho ala bf (cm) — se limita automáticamente",
         "Ancho alma bw (cm)", "Espesor ala hf (cm)", "Altura total h (cm)",
         "Momento último Mu", "Diámetro de varilla"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(_t(f"**Viga T — Sección compuesta** | Norma: `{code['ref']}`", f"**T-Beam — Composite section** | Code: `{code['ref']}`"))
    st.info(_t("**Modo de uso:** Configura el ancho del ala (bf) y del alma (bw), más los espesores y el Momento Último (Mu). El algoritmo deducirá si la viga se comporta como Rectangular (compresión solo en el ala) o como Verdadera Viga T (eje neutro en el alma).", " **How to use:** Set flange width (bf), web width (bw), thicknesses, and Ultimate Moment (Mu). The algorithm deduces if it behaves as a Rectangular or a True T-Beam."))
    c1,c2 = st.columns(2)
    with c1:
        bf_vt = st.number_input("Ancho del ala bf [cm]", 20.0, 300.0, st.session_state.get("vt_bf", 80.0), 5.0, key="vt_bf", help="Ancho efectivo de la aleta de la viga T según NSR-10 C.8.12.")
        bw_vt = st.number_input("Ancho del alma bw [cm]", 10.0, 80.0, st.session_state.get("vt_bw", 25.0), 5.0, key="vt_bw", help="Ancho del nervio de la viga T donde se colocarán los estribos.")
        hf_vt = st.number_input("Espesor ala hf [cm]", 5.0, 40.0, st.session_state.get("vt_hf", 12.0), 1.0, key="vt_hf", help="Espesor de la torta superior o loseta de la aleta.")
        ht_vt = st.number_input("Alto total h [cm]", 20.0, 200.0, st.session_state.get("vt_h", 60.0), 5.0, key="vt_h", help="Altura total de la viga, incluyendo la aleta (hf) y el alma.")
    with c2:
        dp_vt = st.number_input("Recubrimiento d' [cm]", 2.0, 15.0, st.session_state.get("vt_dp", 6.0), 0.5, key="vt_dp")
        Mu_vt = st.number_input(f"Mu [{unidad_mom}]", 0.1, 15000.0, st.session_state.get("vt_mu", 200.0), 10.0, key="vt_mu")
        L_vt  = st.number_input("Longitud [m]", 1.0, 30.0, st.session_state.get("vt_L", 6.0), 0.5, key="vt_L")
        varillas_vt = list(rebar_dict.keys())
        # ?ndice por defecto: #4 (12.7mm) en US, 12mm en SI
        _def_vt = "#4 (Ø12.7mm)" if "Pulgadas" in bar_sys else "12mm"
        _def_idx_vt = varillas_vt.index(_def_vt) if _def_vt in varillas_vt else 1
        bar_vt = st.selectbox("Varilla:", varillas_vt, 
                              index=varillas_vt.index(st.session_state.vt_bar) if "vt_bar" in st.session_state and st.session_state.vt_bar in varillas_vt else _def_idx_vt,
                              key="vt_bar")
        Ab_vt = rebar_dict[bar_vt]; db_vt = diam_dict[bar_vt]

    d_vt = ht_vt - dp_vt
    if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
        Mu_vt_kN = Mu_vt / factor_fuerza
    else:
        Mu_vt_kN = Mu_vt
        
    # Validar ancho de ala efectivo (Punto 7 NSR-10 C.8.10.2 aproximado - BLOQUEO)
    L_cm = L_vt * 100
    limite_bf = min(L_cm / 4, bw_vt + 16 * hf_vt)
    if bf_vt > limite_bf:
        st.warning(f" **Atención Normativa (C.8.10.2):** El ancho de ala $b_f$ ({bf_vt} cm) excedía el límite normativo: min($L/4={L_cm/4:.0f}$, $b_w+16h_f={bw_vt+16*hf_vt:.0f}$) = {limite_bf:.0f} cm. Se calculó usando el límite estricto de {limite_bf:.0f} cm.")
        bf_vt = limite_bf
        
    bf_mm = bf_vt*10; bw_mm = bw_vt*10; hf_mm = hf_vt*10; d_mm_vt = d_vt*10

    Rn_t = Mu_vt_kN*1e6 / (phi_f * bf_mm * d_mm_vt**2)
    disc_t = 1 - 2*Rn_t/(0.85*fc)

    if disc_t < 0:
        st.error("? Sección insuficiente. Aumente bf o h.")
    else:
        rho_bf = (0.85*fc/fy)*(1 - math.sqrt(max(disc_t,0)))
        As_rect = rho_bf * bf_vt * d_vt
        a_r = As_rect*100*fy/(0.85*fc*bf_mm)
        is_T = a_r > hf_mm
        if not is_T:
            n_bt = math.ceil(As_rect/Ab_vt)
            As_prov_vt = n_bt*Ab_vt
            a_final = As_prov_vt*100*fy/(0.85*fc*bf_mm)
            phi_Mn_vt = phi_f*As_prov_vt*100*fy*(d_mm_vt-a_final/2)/1e6
            sec_type = "Rectangular (a ≤ hf)"
        else:
            Asf = 0.85*fc*(bf_mm-bw_mm)*hf_mm/fy
            Mnf = Asf*fy*(d_mm_vt-hf_mm/2)/1e6
            Mn_web = Mu_vt_kN/phi_f - Mnf
            if Mn_web < 0:
                Mn_web = 0
            Rn_w = Mn_web*1e6/(bw_mm*d_mm_vt**2) if Mn_web>0 else 0
            disc_w = 1 - 2*Rn_w/(0.85*fc)
            rho_w = max((0.85*fc/fy)*(1-math.sqrt(max(disc_w,0))),0)
            Asw_mm2 = rho_w*bw_mm*d_mm_vt
            As_total_mm2 = Asf + Asw_mm2
            n_bt = math.ceil(As_total_mm2/100/Ab_vt)
            As_prov_vt = n_bt*Ab_vt
            a_web = (As_prov_vt*100 - Asf)*fy/(0.85*fc*bw_mm) if (As_prov_vt*100-Asf)>0 else 0
            Mn_web_p = (As_prov_vt*100-Asf)*fy*(d_mm_vt-a_web/2) if (As_prov_vt*100-Asf)>0 else 0
            phi_Mn_vt = phi_f*(Mnf + Mn_web_p/1e6)
            sec_type = "T verdadera (a > hf)"

        ok_vt = phi_Mn_vt >= Mu_vt_kN
        if not is_T:
            # Bug Fix: para sección rectangular usar bw (ancho del alma), no bf
            rho_prov_vt = As_prov_vt / (bw_vt * d_vt)
        else:
            # NSR-10: Para Viga T verdadera la ductilidad (p_max) se restringe
            # solo sobre el acero que obra en el alma (A_sw) sobre bw
            As_weq_cm2 = (As_prov_vt*100 - Asf)/100 if (As_prov_vt*100 - Asf) > 0 else 0
            rho_prov_vt = As_weq_cm2 / (bw_vt * d_vt)
        ok_rho_max_T = rho_prov_vt <= rho_max
        
        # Validar ancho de ala efectivo (Punto 12 NSR-10 C.8.10.2 aproximado)
        L_cm = L_vt * 100
        limite_bf = min(L_cm / 4, bw_vt + 16 * hf_vt)
        
        if bf_vt > limite_bf:
            st.warning(f" **Atención Normativa (C.8.10.2):** El ancho de ala $b_f$ ({bf_vt} cm) es mayor que el límite normativo sin considerar separación entre vigas: min($L/4={L_cm/4:.0f}$, $b_w+16h_f={bw_vt+16*hf_vt:.0f}$) = {limite_bf:.0f} cm.")
            # Aplicar límite al bus de datos y local
            st.session_state['vt_bf'] = float(limite_bf)
            # Reasignamos el bf usado en visualizaciones siguientes, pero el cálculo ya está hecho
            bf_vt = limite_bf

        #  ESPACIAMIENTO M?NIMO (ACI 318-25 §25.2.1) 
        # sep_min = max(db, 25 mm) → en cm: max(db_vt_cm, 2.5)
        db_vt_cm = db_vt / 10
        sep_min_cm = max(db_vt_cm, 2.5)

        # Ancho disponible en el alma (interior entre recubrimientos)
        ancho_neto_bw = bw_vt - 2 * dp_vt   # cm

        # Máximo número de varillas que caben en 1 fila con sep_min:
        # ancho ≥ n * db_vt_cm + (n-1) * sep_min_cm
        # n ≤ (ancho_neto + sep_min) / (db_vt_cm + sep_min)
        n_max_fit = max(1, int((ancho_neto_bw + sep_min_cm) / (db_vt_cm + sep_min_cm)))

        # ¿Cuántas filas necesitamos?
        n_filas_req = math.ceil(n_bt / n_max_fit)
        n_bt_draw = min(n_bt, n_max_fit)   # barras dibujadas en la 1era fila

        #  CENTROIDE REAL DEL ACERO (NSR-10 C.7.6.2) 
        sep_filas_vt_mm = 25.0  # mm — separación vertical mínima entre capas
        if n_filas_req > 1:
            n_f1_vt = n_bt_draw
            n_f2_vt = n_bt - n_f1_vt
            y1_vt_mm = dp_vt * 10
            y2_vt_mm = dp_vt * 10 + db_vt + sep_filas_vt_mm
            y_cg_vt  = (n_f1_vt * Ab_vt * y1_vt_mm + n_f2_vt * Ab_vt * y2_vt_mm) / (n_bt * Ab_vt)
            d_vt     = ht_vt - y_cg_vt / 10   # cm — peralte efectivo real
            d_mm_vt  = d_vt * 10               # mm
            dos_filas_vt_flag = True
        else:
            dos_filas_vt_flag = False

        # Verificación bw_min (para la primera fila)
        bw_min_req_vt = 2 * dp_vt + n_bt_draw * db_vt_cm + (n_bt_draw - 1) * sep_min_cm
        ok_bw_vt = bw_vt >= bw_min_req_vt

        # Confinamiento Sísmico (DES / DMO)
        Ln_vt_cm = L_vt * 100
        Lo_min = max(2 * ht_vt, Ln_vt_cm / 4, 45.0)
        s_conf_max = min(d_vt / 4, 6 * db_vt_cm, 15.0)

        tab_r,tab_s,tab_3d,tab_q = st.tabs([f"Resultados {''if (ok_vt and ok_bw_vt and ok_rho_max_T) else '?'}"," Sección 2D"," Visualización 3D"," Cantidades"])
        with tab_r:
            st.markdown(f"**Tipo de sección:** {sec_type} | **φ={phi_f}**")
            rows_vt = [
                ("bf × bw × hf × h", f"{bf_vt:.0f} × {bw_vt:.0f} × {hf_vt:.0f} × {ht_vt:.0f} cm"),
                ("d efectivo", f"{d_vt:.1f} cm"),
                ("Comportamiento", sec_type),
                (f"Varillas ({bar_vt})", f"{n_bt} barras — As prov = {As_prov_vt:.3f} cm²"),
                (f"Ancho Alma bw_mín", f"Requerido: {bw_min_req_vt:.1f} cm (Provisto: {bw_vt:.0f} cm)"),
                (f"Sep. mín entre varillas (ACI §25.2.1)", f"max(db, 25mm) = {sep_min_cm*10:.0f} mm"),
                (f"Máx varillas por fila en bw={bw_vt:.0f}cm", f"{n_max_fit} barras"),
                (f"Filas de acero necesarias", f"{n_filas_req} {'fila' if n_filas_req==1 else 'filas'}"),
                (f"φMn calculado [{unidad_mom}]", f"{phi_Mn_vt*factor_fuerza:.2f}"),
                (f"Mu solicitado [{unidad_mom}]", f"{Mu_vt:.2f}"),
                ("? provisto", f"{(As_prov_vt/(bw_vt*d_vt))*100:.4f}%"),
                ("? máximo", f"{rho_max*100:.4f}%"),
                ("Validación ?", " CUMPLE" if ok_rho_max_T else "? EXCEDE (Sobrearmada)"),
                ("Validación bw", " CUMPLE" if ok_bw_vt else "? DEFICIENTE (Alma muy angosta)"),
                ("Estado Global", " CUMPLE" if (ok_vt and ok_bw_vt) else "? DEFICIENTE"),
            ]
            qty_table(rows_vt)

            if n_filas_req > 1:
                st.warning(
                    f" ? **{n_bt} varillas {bar_vt} no caben en 1 fila** en un alma de bw={bw_vt:.0f} cm "
                    f"(sep. mínima {sep_min_cm*10:.0f} mm). "
                    f"Máximo por fila: **{n_max_fit}** barras. "
                    f"Se necesitan **{n_filas_req} capas de armadura** o aumente bw "
                    f"a {math.ceil(2*dp_vt + n_bt*db_vt_cm + (n_bt-1)*sep_min_cm)} cm."
                )

            if ok_vt and ok_bw_vt:
                st.success(f"φMn = {phi_Mn_vt*factor_fuerza:.2f} {unidad_mom} ≥ Mu = {Mu_vt:.2f} {unidad_mom}")
            elif not ok_bw_vt:
                st.error(f"? El acero no cabe en el alma configurada. Aumente bw a mínimo {math.ceil(bw_min_req_vt)} cm.")
            else:
                st.error(f"φMn = {phi_Mn_vt*factor_fuerza:.2f} {unidad_mom} < Mu = {Mu_vt:.2f} {unidad_mom}")

            st.info("**¿Acero Inferior o Superior?** Si ingresa un Momento **Positivo** (Mu), el área calculada corresponde al refuerzo en la zona traccionada (usualmente **acero inferior**). Para momento **Negativo** en un apoyo continuo, la tracción está arriba por lo que el resultado corresponde al **acero superior**.")

            with st.expander("**Detallado Sísmico: Zonas de Rótula Plástica (DES/DMO)**", expanded=False):
                st.markdown("> *Requisitos para elementos a flexión en sistemas de Resistencia Sísmica.*")
                st.info("**Nota Técnica (Lo_conf):** La Longitud de Confinamiento ($L_o$) es la zona adyacente a los nudos donde pueden formarse rótulas plásticas durante un sismo severo. En esta región se exige una mayor densidad de estribos cerrados de confinamiento para garantizar la ductilidad y evitar el pandeo de barras longitudinales (ref. NSR-10 C.21 / ACI 318 Cap.18).")
                c_s1, c_s2 = st.columns(2)
                with c_s1:
                    st.metric("Lo (Long. de confinamiento)", f"≥ {Lo_min:.0f} cm", delta="2h, Ln/4, 45cm", delta_color="off")
                with c_s2:
                    st.metric("s_conf (Separación máx)", f"≤ {s_conf_max:.1f} cm", delta="d/4, 6db, 15cm", delta_color="off")
                st.caption(f"**Nota:** El primer estribo debe colocarse a no más de 5 cm de la cara del apoyo (Norma {code['ref']}).")

        with tab_s:
            fig, ax = plt.subplots(figsize=(5,4))
            fig.patch.set_facecolor('#1e1e2e')
            ax_arr = fig.get_axes()
            for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
            fig.patch.set_facecolor('#1a1a2e'); ax.set_facecolor('#1a1a2e')
            ax.add_patch(patches.Rectangle(((bf_vt-bw_vt)/2, 0), bw_vt, ht_vt-hf_vt, linewidth=1.5, edgecolor='white', facecolor='#4a4a6a'))
            ax.add_patch(patches.Rectangle((0, ht_vt-hf_vt), bf_vt, hf_vt, linewidth=1.5, edgecolor='white', facecolor='#3a3a5a'))
            r_v = db_vt / 20
            # Dibujar solo las barras que caben en la primera fila
            x_ini = (bf_vt - bw_vt) / 2 + dp_vt
            x_fin = (bf_vt + bw_vt) / 2 - dp_vt
            if n_bt_draw > 1:
                xs_v = np.linspace(x_ini, x_fin, n_bt_draw)
            else:
                xs_v = [(x_ini + x_fin) / 2]
            for x in xs_v:
                ax.add_patch(plt.Circle((x, dp_vt), r_v, color='#ff6b35', zorder=5))
            # Indicar si hay más filas
            if n_filas_req > 1:
                n_restantes = n_bt - n_bt_draw
                ax.text(bf_vt/2, dp_vt + r_v*3 + 0.5,
                        f"+{n_restantes} en {n_filas_req-1} fila(s) adicional(es)",
                        ha='center', va='bottom', color='#ffcc00', fontsize=7, style='italic')
            ax.set_xlim(-5, bf_vt+5); ax.set_ylim(-5, ht_vt+5)
            ax.axis('off')
            title_str = f"Viga T: bf={bf_vt:.0f} bw={bw_vt:.0f} hf={hf_vt:.0f} h={ht_vt:.0f} cm\n{n_bt}×{bar_vt}"
            if n_filas_req > 1:
                title_str += f" ({n_bt_draw}/fila, {n_filas_req} filas)"
            ax.set_title(title_str, color='white', fontsize=8)
            st.pyplot(fig)
            plt.close(fig)

        with tab_3d:
            st.subheader("Visualización 3D de Viga T")
            fig3d = go.Figure()
            L_mm_3d = L_vt * 100
            # Alma
            x_w = [-bw_vt/2, bw_vt/2, bw_vt/2, -bw_vt/2, -bw_vt/2, bw_vt/2, bw_vt/2, -bw_vt/2]
            y_w = [0, 0, ht_vt-hf_vt, ht_vt-hf_vt, 0, 0, ht_vt-hf_vt, ht_vt-hf_vt]
            z_w = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
            fig3d.add_trace(go.Mesh3d(x=x_w, y=y_w, z=z_w, alphahull=0, opacity=0.15, color='gray', name='Concreto (Alma)'))
            # Ala
            x_f = [-bf_vt/2, bf_vt/2, bf_vt/2, -bf_vt/2, -bf_vt/2, bf_vt/2, bf_vt/2, -bf_vt/2]
            y_f = [ht_vt-hf_vt, ht_vt-hf_vt, ht_vt, ht_vt, ht_vt-hf_vt, ht_vt-hf_vt, ht_vt, ht_vt]
            z_f = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
            fig3d.add_trace(go.Mesh3d(x=x_f, y=y_f, z=z_f, alphahull=0, opacity=0.15, color='gray', name='Concreto (Ala)'))
            #  Varillas INFERIORES (tensión) — respeta espaciamiento mínimo 
            diam_reb_cm = db_vt / 10.0
            line_width = max(4, diam_reb_cm * 3)
            x_ini_3d = -bw_vt / 2 + dp_vt
            x_fin_3d =  bw_vt / 2 - dp_vt
            for fila in range(n_filas_req):
                n_en_fila = n_bt_draw if fila < n_filas_req - 1 else (n_bt - fila * n_max_fit)
                n_en_fila = max(1, n_en_fila)
                if n_en_fila > 1:
                    xs_fila = np.linspace(x_ini_3d, x_fin_3d, n_en_fila)
                else:
                    xs_fila = [0.0]
                y_fila = dp_vt + fila * (diam_reb_cm + 0.5)   # cada fila sube 1 diámetro + holgura
                for idx, x_pos in enumerate(xs_fila):
                    fig3d.add_trace(go.Scatter3d(
                        x=[x_pos, x_pos], y=[y_fila, y_fila], z=[0, L_mm_3d],
                        mode='lines', line=dict(color='darkred', width=line_width),
                        name=f'Varilla inf. {bar_vt}', showlegend=(fila == 0 and idx == 0)))
            #  Varillas SUPERIORES (compresión / montaje: 2 barras en esquinas del alma) 
            y_sup = ht_vt - dp_vt  # y en el alma, cerca del ala
            xs_sup = [-bw_vt/2 + dp_vt, bw_vt/2 - dp_vt]
            for idx, x_pos in enumerate(xs_sup):
                fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[y_sup, y_sup], z=[0, L_mm_3d],
                                            mode='lines', line=dict(color='orange', width=max(3, diam_reb_cm*2)),
                                            name='Varilla sup. (compresión)', showlegend=(idx==0)))
            #  Estribos (alma) 
            tie_color = 'cornflowerblue'
            tie_width = max(2, (9.5/10.0) * 3)
            sep_ties = st.slider("Separación Estribos (cm) ", 5, 50, int(st.session_state.get('cv_s_diseno', 15)), 1, key="vt_sep_tie")
            tx = [-bw_vt/2 + dp_vt/2, bw_vt/2 - dp_vt/2, bw_vt/2 - dp_vt/2, -bw_vt/2 + dp_vt/2, -bw_vt/2 + dp_vt/2]
            ty = [dp_vt/2, dp_vt/2, ht_vt - dp_vt/2, ht_vt - dp_vt/2, dp_vt/2]
            L_cm = int(L_mm_3d)
            tx_all, ty_all, tz_all = [], [], []
            for zt in range(5, L_cm, sep_ties):
                for px, py in zip(tx, ty):
                    tx_all.append(px); ty_all.append(py); tz_all.append(zt)
                tx_all.append(None); ty_all.append(None); tz_all.append(None)
            for px, py in zip(tx, ty):
                tx_all.append(px); ty_all.append(py); tz_all.append(L_cm - 5)
            tx_all.append(None); ty_all.append(None); tz_all.append(None)
            fig3d.add_trace(go.Scatter3d(x=tx_all, y=ty_all, z=tz_all, mode='lines', 
                                         line=dict(color=tie_color, width=tie_width), name='Estribos Alma', showlegend=True))
            fig3d.update_layout(
                scene=dict(aspectmode='data', xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                legend=dict(bgcolor='rgba(0,0,0,0.5)', font=dict(color='white'), x=0.01, y=0.99),
                margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable'
            )
            st.plotly_chart(fig3d, use_container_width=True)

        with tab_q:
            vol_t = (bf_vt*hf_vt + bw_vt*(ht_vt-hf_vt))/10000 * L_vt
            peso_t = As_prov_vt * L_vt * 0.785
            m = mix_for_fc(fc)
            bags = m[0]*vol_t/bag_kg
            qty_table([("Concreto Viga T", f"{vol_t:.4f} m³"),
                       (f"Acero ({n_bt} barras)", f"{peso_t:.2f} kg"),
                       (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{m[0]*vol_t/bag_kg:.1f} bultos"),
                       ("Referencia", code["ref"])])

            render_apu_breakdown(vol_t, peso_t, fc, f"({n_bt} barras)")

            with st.expander("Dibujo de Figurado para Taller (Viga T)", expanded=False):
                hook_len_cm = 12 * db_vt / 10
                straight_len_cm = L_vt * 100 - 2 * hook_len_cm
                fig_l1 = draw_longitudinal_bar(L_vt*100, straight_len_cm, hook_len_cm, db_vt)
                st.pyplot(fig_l1)
                plt.close(fig_l1)
                recub_est = max(dp_vt, 2.5)
                inside_b = bw_vt - 2*recub_est
                inside_h = ht_vt - 2*recub_est
                hook_len_est = 12 * 9.5 / 10
                fig_e1 = draw_stirrup_beam(inside_b, inside_h, hook_len_est, 9.5)
                st.pyplot(fig_e1)
                plt.close(fig_e1)

            # MEMORIA DOCX y EXPORTACIÓN IFC para Viga T
            col_t1, col_t2, col_t3 = st.columns(3)
            with col_t1:
                btn_docx_t = st.button("Generar Memoria DOCX (Viga T)")
            with col_t2:
                if st.button("Enviar a Cuadro de Mando", key="cmd_vt"):
                    add_historial_diseno("Flexión Viga T", f"bf={bf_vt:.0f} bw={bw_vt:.0f} h={ht_vt:.0f} cm", " APROBADO" if (ok_vt and ok_bw_vt and ok_rho_max_T) else "? NO CUMPLE", code['ref'])
            with col_t3:
                try:
                    buf_ifc_t = ifc_export.ifc_viga_t(
                        bf_vt, bw_vt, hf_vt, ht_vt, L_vt, fc, fy, int(n_bt), bar_vt, db_vt,
                        As_prov_vt, d_vt, dp_vt, Mu_vt, phi_Mn_kNm_vt, norma_sel, "Proyecto NSR-10",
                        db_est_mm=db_e_vr, sep_est_mm=sep_ties*10.0
                    )
                    st.download_button("Exportar IFC (BIM)", data=buf_ifc_t, file_name=f"Viga_T_{bf_vt:.0f}x{ht_vt:.0f}.ifc", mime="application/x-step", key="ifc_vt")
                except Exception as e:
                    st.error(f"Error IFC: {e}")
                    
            if btn_docx_t:
                doc = Document()
                doc.add_heading(f"Memoria de Cálculo Viga T: bf={bf_vt:.0f} bw={bw_vt:.0f} h={ht_vt:.0f} cm", 0)
                doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_heading("1. Materiales", level=1)
                doc.add_paragraph(f"Concreto: f'c = {fc:.1f} MPa")
                doc.add_paragraph(f"Acero: fy = {fy:.0f} MPa")
                doc.add_paragraph(f"Norma: {norma_sel}")
                doc.add_heading("2. Geometría", level=1)
                doc.add_paragraph(f"Ancho ala (bf) = {bf_vt:.0f} cm, Ancho alma (bw) = {bw_vt:.0f} cm")
                doc.add_paragraph(f"Espesor ala (hf) = {hf_vt:.0f} cm, Altura total (h) = {ht_vt:.0f} cm")
                doc.add_paragraph(f"d efectivo = {d_vt:.1f} cm")
                doc.add_heading("3. Refuerzo y Comportamiento", level=1)
                doc.add_paragraph(f"Comportamiento: {sec_type}")
                doc.add_paragraph(f"Refuerzo longitudinal: {n_bt} varillas {bar_vt} → As = {As_prov_vt:.3f} cm²")
                doc.add_heading("4. Verificaciones Normativas", level=1)
                checks = [
                    ("Resistencia a flexión (φMn ≥ Mu)", "CUMPLE" if ok_vt else "NO CUMPLE"),
                    ("Cuantía máxima (? ≤ ?_max)", "CUMPLE" if ok_rho_max_T else "NO CUMPLE"),
                    ("Ancho mínimo (bw ≥ bw_min_req)", "CUMPLE" if ok_bw_vt else "NO CUMPLE"),
                ]
                for desc, res in checks:
                    doc.add_paragraph(f"{desc}: {res}")
                doc.add_paragraph(f"φMn = {phi_Mn_vt*factor_fuerza:.2f} {unidad_mom} | Mu = {Mu_vt:.2f} {unidad_mom}")
                doc.add_paragraph(f"? = {(As_prov_vt/(bw_vt*d_vt))*100:.3f}% | ?_max = {rho_max*100:.3f}%")
                doc.add_paragraph(f"bw provisto = {bw_vt:.0f} cm | bw mínimo = {bw_min_req_vt:.1f} cm")
                
                doc.add_heading("5. Detallado Sísmico — Zona de Confinamiento", level=1)
                doc.add_paragraph(
                    f"NSR-10 C.21.5.3.1 (o ACI 318 §18.6.4): En vigas de pórticos especiales (DMRS/DMES), "
                    f"se requieren estribos cerrados en una longitud Lo desde la cara del apoyo. "
                    f"La norma exige: Lo_conf = max(2·h, Ln/4, 45 cm) = {Lo_min:.1f} cm."
                )
                doc.add_paragraph(
                    f"NSR-10 C.21.5.3.2: La separación de estribos en esa zona no debe exceder "
                    f"s_conf = min(d/4, 6·db, 15 cm) = {s_conf_max:.1f} cm. "
                    f"Fuera de la zona confinada se conserva el espaciado de diseño a cortante."
                )
                doc.add_paragraph(f"Valores calculados: Lo_conf = {Lo_min:.1f} cm  |  s_conf_max = {s_conf_max:.1f} cm")

                doc.add_heading("6. Detalle de Sección Transversal", level=1)
                fig_wt, ax_wt = sec_light_fig_t(bf_vt, bw_vt, hf_vt, ht_vt, f"Sección Viga T: bf={bf_vt:.0f} bw={bw_vt:.0f}")
                r_v = db_vt_cm / 2
                xs_v = np.linspace((bf_vt-bw_vt)/2+dp_vt, (bf_vt+bw_vt)/2-dp_vt, max(n_bt,2))
                for x in xs_v[:n_bt]:
                    ax_wt.add_patch(plt.Circle((x, dp_vt), r_v, color='red', zorder=5))
                    ax_wt.text(x, dp_vt - max(ht_vt*0.1, 4), "As", color='red', fontsize=8, ha='center', va='top')
                buf_vt = io.BytesIO()
                fig_wt.savefig(buf_vt, format='png', dpi=150, bbox_inches='tight')
                buf_vt.seek(0)
                plt.close(fig_wt)
                from docx.shared import Inches
                doc.add_picture(buf_vt, width=Inches(3.0))

                # --- Gráfico Mn vs As Viga T ---
                doc.add_heading("7. Diagrama Capacidad vs Demanda", level=1)
                _as_vals_t = np.linspace(0.01, max(As_prov_vt * 2.5, 0.1), 30)
                _phi_mn_env_t = []
                for _as in _as_vals_t:
                    _a_i = _as * 100 * fy / (0.85 * fc * bf_mm)
                    _phi_mn_env_t.append(phi_f * _as * 100 * fy * (d_mm_vt - _a_i / 2) / 1e6 * factor_fuerza)
                fig_mn_t, ax_mn_t = plt.subplots(figsize=(5, 3))
                fig_mn_t.patch.set_facecolor('#1e1e2e')
                ax_arr = fig_mn_t.get_axes()
                for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
                fig_mn_t.patch.set_facecolor('white')
                ax_mn_t.plot(_as_vals_t, _phi_mn_env_t, color='steelblue', lw=2, label='φMn (As)')
                ax_mn_t.axhline(Mu_vt_kN * factor_fuerza, color='red', linestyle='--', lw=1.5, label=f'Mu={Mu_vt_kN * factor_fuerza:.2f} {unidad_mom}')
                ax_mn_t.scatter([As_prov_vt], [phi_Mn_vt * factor_fuerza], color='green', zorder=5, s=60)
                ax_mn_t.set_xlabel('As provisto (cm²)'); ax_mn_t.set_ylabel(f'φMn ({unidad_mom})')
                ax_mn_t.legend(fontsize=8); ax_mn_t.grid(alpha=0.3)
                fig_mn_t.tight_layout()
                buf_mn_t = io.BytesIO()
                fig_mn_t.savefig(buf_mn_t, format='png', dpi=150, bbox_inches='tight')
                buf_mn_t.seek(0)
                plt.close(fig_mn_t)
                doc.add_picture(buf_mn_t, width=Inches(4.5))

                doc.add_heading("8. Cantidades", level=1)
                doc.add_paragraph(f"Volumen concreto: {vol_t:.4f} m³")
                doc.add_paragraph(f"Acero: {peso_t:.2f} kg ({n_bt} var. {bar_vt})")
                
                doc_mem = io.BytesIO()
                doc.save(doc_mem)
                doc_mem.seek(0)
                st.download_button("Descargar Memoria Viga T", data=doc_mem, file_name=f"Memoria_VigaT_{bf_vt:.0f}x{bw_vt:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ??????????????????????????????????????????
# 4. CORTANTE EN VIGAS
# ??????????????????????????????????????????
if modulo_sel == " Diseño a Cortante — Vigas de Concreto":
    #  PANEL 
    _panel_normativo("Diseño a Cortante — Vigas de Concreto",
        "Calcula Vc, Vs y la separación de estribos para la fuerza cortante última Vu.",
        ["Ancho alma bw (cm)", "Peralte efectivo d (cm)",
         "Cortante último Vu (kN o tonf)", "Longitud viga L (m)",
         "Diámetro y número de ramas del estribo"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(_t(f"**Diseño de estribos a cortante** | Norma: `{code['ref']}`", f"**Shear stirrup design** | Code: `{code['ref']}`"))
    st.info(_t("**Modo de uso:** Ingresa la Fuerza Cortante Factorizada (Vu) para una sección de viga dada. La app determinará la contribución del concreto (φVc) y calculará el refuerzo transversal requerido en número de estribos y separación (s).", " **How to use:** Enter Factored Shear Force (Vu). The app calculates concrete contribution (φVc) and required transverse reinforcement (stirrups spacing & amount)."))
    c1,c2 = st.columns(2)
    with c1:
        bw_cv = st.number_input("bw [cm]", 10.0, 100.0, st.session_state.get("cv_bw", 25.0), 5.0, key="cv_bw")
        d_cv  = st.number_input("d peralte efectivo [cm]", 10.0, 200.0, st.session_state.get("cv_d", 45.0), 5.0, key="cv_d")
        Vu_cv_input = st.number_input(f"Vu [{unidad_fuerza}]", 0.1, 5000.0, st.session_state.get("cv_vu", 80.0), 5.0, key="cv_vu")
        if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
            Vu_cv = Vu_cv_input / factor_fuerza
        else:
            Vu_cv = Vu_cv_input
        L_cv  = st.number_input("Longitud viga [m]", 1.0, 30.0, st.session_state.get("cv_L", 5.0), 0.5, key="cv_L")
    with c2:
        h_cv  = st.number_input("h total [cm]", 20.0, 200.0, st.session_state.get("cv_h", 50.0), 5.0, key="cv_h")
        est_opts = ["Ø6mm","Ø8mm","Ø10mm","Ø12mm","#2","#3","#4"]
        st_bar_cv = st.selectbox("Estribo:", est_opts, 
                                 index=est_opts.index(st.session_state.cv_st) if "cv_st" in st.session_state and st.session_state.cv_st in est_opts else 1,
                                 key="cv_st")
        st_area = {"Ø6mm":0.283,"Ø8mm":0.503,"Ø10mm":0.785,"Ø12mm":1.131,"#2":0.32,"#3":0.71,"#4":1.29}[st_bar_cv]
        n_ramas = st.number_input("# Ramas del estribo", 2, 6, st.session_state.get("cv_ramas", 2), 1, key="cv_ramas")
        Av_cv = st_area * n_ramas
        diam_est = {"Ø6mm":6,"Ø8mm":8,"Ø10mm":10,"Ø12mm":12,"#2":6.35,"#3":9.53,"#4":12.70}[st_bar_cv]

    st.markdown("---")
    vc_type = st.radio("Fórmula de contribución del Concreto ($V_c$):", 
                       ["Simplificada ($0.17\\lambda\\sqrt{f\\'c}·b_w d$)", "Detallada (NSR-10 C.11.2.1.1)"], 
                       horizontal=True)
    if "Detallada" in vc_type:
        cc1, cc2 = st.columns(2)
        with cc1:
            Mu_cv_input = st.number_input(f"Mu concomitante al Vu [{unidad_mom}]", 0.0, 5000.0, 80.0, 5.0, help="Momento factorizado en la misma sección donde ocurre Vu.")
            Mu_cv = Mu_cv_input / factor_fuerza if "Toneladas" in unidades_salida else Mu_cv_input
        with cc2:
            As_cv = st.number_input("As en tracción provisto [cm²]", 0.1, 200.0, 10.0, 1.0, help="Acero longitudinal en la zona de tracción.")

    bw_mm_cv = bw_cv*10; d_mm_cv = d_cv*10
    
    if "Simplificada" in vc_type:
        Vc_N = 0.17 * lam * math.sqrt(fc) * bw_mm_cv * d_mm_cv
    else:
        rho_w_cv = As_cv / (bw_cv * d_cv) if "As_cv" in locals() and bw_cv > 0 and d_cv > 0 else 0
        if "Mu_cv" not in locals() or Mu_cv <= 0:
            Vd_m = 1.0 # Limite conservador norma
        else:
            Vd_m = min((Vu_cv * d_cv / 100) / Mu_cv, 1.0)
            
        Vc_det_N = (0.16 * lam * math.sqrt(fc) + 17 * rho_w_cv * Vd_m) * bw_mm_cv * d_mm_cv
        Vc_max_N = 0.29 * lam * math.sqrt(fc) * bw_mm_cv * d_mm_cv
        Vc_N = min(Vc_det_N, Vc_max_N)

    Vc_kN = Vc_N/1000
    phi_Vc = phi_v * Vc_kN
    Vs_req_kN = max(0, Vu_cv/phi_v - Vc_kN)
    need_design = Vu_cv > phi_Vc/2

    if Vs_req_kN > 0:
        s_calc_mm = Av_cv*100*fy*d_mm_cv/(Vs_req_kN*1000)
    else:
        s_calc_mm = min(d_mm_cv/2, 600)

    Vs_lim = 0.33*math.sqrt(fc)*bw_mm_cv*d_mm_cv/1000
    if Vs_req_kN > Vs_lim:
        s_max_mm = min(d_mm_cv/4, 300)
    else:
        s_max_mm = min(d_mm_cv/2, 600)
    s_diseno_mm = min(s_calc_mm, s_max_mm)
    s_diseno_cm = s_diseno_mm/10
    st.session_state['cv_s_diseno'] = s_diseno_cm

    #  Alerta: Vu < φVc/2  solo estribos constructivos necesarios
    if not need_design:
        st.info(
            f"ℹ? **Vu = {Vu_cv_input:.2f} {unidad_fuerza} < φVc/2 = {phi_Vc/2*factor_fuerza:.2f} {unidad_fuerza}**: "
            f"No se requieren estribos por cálculo. Solo se necesitan **estribos constructivos** "
            f"con s_max = {s_diseno_cm:.1f} cm (NSR-10 C.11.5.5.1 / ACI 318 §26.7.2)."
        )
    else:
        st.info(f"φVc = {phi_Vc*factor_fuerza:.2f} {unidad_fuerza} | Vs requerido = {Vs_req_kN*factor_fuerza:.2f} {unidad_fuerza} | s diseño = {s_diseno_cm:.1f} cm")
    n_estribos = math.ceil(L_cv*100/s_diseno_cm) + 1
    Vs_prov_kN = Av_cv*100*fy*d_mm_cv/(s_diseno_mm*1000)
    phi_Vn_kN = phi_v*(Vc_kN + Vs_prov_kN)
    ok_cv = phi_Vn_kN >= Vu_cv

    if s_diseno_cm < 5:
        st.warning(" ? La separación de estribos es menor a 5 cm. Considere aumentar el diámetro de los estribos o el número de ramas.")
    elif s_diseno_cm < 7.5:
        st.info("ℹ? La separación de estribos es menor a 7.5 cm. Verifique que sea constructivamente viable.")

    tab_r,tab_s,tab_q = st.tabs([f"Resultados {''if ok_cv else '?'}"," Sección"," Cantidades"])
    with tab_r:
        st.markdown(f"**φ cortante = {phi_v}** | Norma: `{code['ref']}`")
        st.markdown(r"**Verificación Normativa:** $\\phi V_n = \\phi (V_c + V_s) \\ge V_u$")
        st.latex(r"V_s = \frac{A_v f_{yt} d}{s}")
        Vs_max_kN = 0.66*math.sqrt(fc)*bw_mm_cv*d_mm_cv/1000
        rows_cv = [
            ("bw × d", f"{bw_cv:.0f} × {d_cv:.0f} cm"),
            (f"Vc (concreto) [{unidad_fuerza}]", f"{Vc_kN*factor_fuerza:.2f}"),
            (f"φVc [{unidad_fuerza}]", f"{phi_Vc*factor_fuerza:.2f}"),
            (f"Vu [{unidad_fuerza}]", f"{Vu_cv_input:.2f}"),
            (f"Vs requerido [{unidad_fuerza}]", f"{Vs_req_kN*factor_fuerza:.2f}"),
            (f"Av ({n_ramas} ramas {st_bar_cv})", f"{Av_cv:.3f} cm²"),
            ("s calculado", f"{s_calc_mm:.0f} mm = {s_calc_mm/10:.1f} cm"),
            ("s máx (norma)", f"{s_max_mm:.0f} mm = {s_max_mm/10:.1f} cm"),
            ("s de diseño", f"**{s_diseno_cm:.1f} cm**"),
            (f"Vs provisto [{unidad_fuerza}]", f"{Vs_prov_kN*factor_fuerza:.2f}"),
            (f"φVn = φ(Vc+Vs) [{unidad_fuerza}]", f"{phi_Vn_kN*factor_fuerza:.2f}"),
            (f"Vs máx permitido [{unidad_fuerza}]", f"{Vs_max_kN*factor_fuerza:.2f}"),
            ("Estado", " CUMPLE" if ok_cv else "? DEFICIENTE"),
        ]
        qty_table(rows_cv)
        if Vs_req_kN > Vs_max_kN:
            ratio_vs = Vs_req_kN / Vs_max_kN
            bw_new_mm = (Vs_req_kN * 1000) / (0.66 * math.sqrt(fc) * d_mm_cv)
            st.error(f"⚠️ **FALLA POR APLASTAMIENTO DEL ALMA — NSR-10 C.11.4.7.9:** $V_s$ requerido ({Vs_req_kN*factor_fuerza:.2f} {unidad_fuerza}) supera el límite máximo transversal $V_{{s,max}}$ = {Vs_max_kN*factor_fuerza:.2f} {unidad_fuerza} (Ratio = **{ratio_vs:.2f}**).\n\n"
                     f"El concreto fallará por compresión diagonal antes de que el acero fluya. **Aumente la sección.**\n\n"
                     f"**Soluciones propuestas:**\n"
                     f"• Aumentar el ancho del alma **bw** ≥ **{math.ceil(bw_new_mm/10)} cm**.\n"
                     f"• Aumentar el peralte efectivo **d** o la altura total **h**.\n"
                     f"• Aumentar la resistencia del concreto **f'c**.")
            st.stop()
        elif ok_cv:
            st.success(f"Aprobado Cortante: \u03c6Vn = {phi_Vn_kN*factor_fuerza:.2f} {unidad_fuerza} >= Vu = {Vu_cv_input:.2f} {unidad_fuerza} — Estribo {st_bar_cv} @ {s_diseno_cm:.1f} cm")
        else:
            st.error(f"No Aprobado por Cortante: \u03c6Vn = {phi_Vn_kN*factor_fuerza:.2f} {unidad_fuerza} < Vu = {Vu_cv_input:.2f} {unidad_fuerza}")

    with tab_s:
        #  VISUALIZACIÓN 3D — Sección Cortante 
        recub_cv = (h_cv - d_cv) * 0.5          # recubrimiento estimado (cm)
        dp_cv_cm = recub_cv + diam_est / 20      # dist. al eje del estribo (cm)

        fig3d_cv = go.Figure()
        Lcm_cv = L_cv * 100                     # longitud en cm

        #  Sólido de concreto translúcido 
        x0, x1 = 0.0, bw_cv
        y0, y1 = 0.0, h_cv
        vx = [x0, x1, x1, x0, x0, x1, x1, x0]
        vy = [y0, y0, y1, y1, y0, y0, y1, y1]
        vz = [0,  0,  0,  0, Lcm_cv, Lcm_cv, Lcm_cv, Lcm_cv]
        fig3d_cv.add_trace(go.Mesh3d(
            x=vx, y=vy, z=vz,
            i=[0, 0, 0, 1, 4, 4, 4, 5],
            j=[1, 2, 4, 5, 5, 6, 7, 6],
            k=[2, 3, 5, 6, 6, 7, 3, 7],
            opacity=0.10, color='gray', name='Concreto'))

        #  Estribos espaciados a s_diseno_cm 
        ex1, ex2 = dp_cv_cm, bw_cv - dp_cv_cm
        ey1, ey2 = dp_cv_cm, h_cv - dp_cv_cm
        tx_rect = [ex1, ex2, ex2, ex1, ex1, None]
        ty_rect = [ey1, ey1, ey2, ey2, ey1, None]
        txall, tyall, tzall = [], [], []
        z_pos = s_diseno_cm
        while z_pos <= Lcm_cv - s_diseno_cm / 2:
            txall += tx_rect; tyall += ty_rect; tzall += [z_pos]*5 + [None]
            z_pos += s_diseno_cm
        # Primer y último estribo
        for zz in [5.0, Lcm_cv - 5.0]:
            txall += tx_rect; tyall += ty_rect; tzall += [zz]*5 + [None]

        if txall:
            fig3d_cv.add_trace(go.Scatter3d(
                x=txall, y=tyall, z=tzall,
                mode='lines', line=dict(color='cyan', width=3),
                name=f'Estribos {st_bar_cv} @ {s_diseno_cm:.1f}cm'))

        #  Barras longitudinales (4 esquinas) 
        for bx, by in [(dp_cv_cm, dp_cv_cm),
                       (bw_cv - dp_cv_cm, dp_cv_cm),
                       (dp_cv_cm, h_cv - dp_cv_cm),
                       (bw_cv - dp_cv_cm, h_cv - dp_cv_cm)]:
            fig3d_cv.add_trace(go.Scatter3d(
                x=[bx, bx], y=[by, by], z=[0, Lcm_cv],
                mode='lines', line=dict(color='orange', width=5),
                showlegend=False))

        fig3d_cv.update_layout(
            scene=dict(
                aspectmode='data',
                xaxis_title='bw (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)',
                xaxis=dict(color='white'), yaxis=dict(color='white'), zaxis=dict(color='white'),
                bgcolor='#0e1117'),
            paper_bgcolor='#0e1117',
            font_color='white',
            margin=dict(l=0, r=0, t=35, b=0),
            height=430,
            title=dict(
                text=f"Cortante 3D — {bw_cv:.0f}×{h_cv:.0f} cm | s={s_diseno_cm:.1f}cm | Vu={Vu_cv_input:.1f}{unidad_fuerza} | φVn={phi_Vn_kN*factor_fuerza:.1f}{unidad_fuerza}",
                font=dict(size=12, color='white')))
        st.plotly_chart(fig3d_cv, use_container_width=True)

    with tab_q:
        perim_cv = 2*(bw_cv-2*recub_cv) + 2*(h_cv-2*recub_cv) + 6*diam_est/10
        vol_beam_cv = bw_cv/100*h_cv/100*L_cv
        peso_est_cv = n_estribos * (perim_cv / 100.0) * st_area * 0.785
        m = mix_for_fc(fc)
        bags = m[0]*vol_beam_cv/bag_kg
        qty_table([
            (f"Estribos {st_bar_cv} @ {s_diseno_cm:.1f}cm", f"{n_estribos} estribos"),
            ("Peso estribos", f"{peso_est_cv:.2f} kg"),
            ("Longitud total estribos", f"{n_estribos*perim_cv/100:.2f} m"),
            ("Concreto viga", f"{vol_beam_cv:.4f} m³"),
            (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{bags:.1f} bultos"),
            ("Referencia", code["ref"]),
        ])

        render_apu_breakdown(vol_beam_cv, peso_est_cv, fc, f"({n_estribos} estribos)")

        with st.expander("Dibujo de Estribo para Taller", expanded=False):
            recub_est = max(recub_cv, 2.5)
            inside_b = bw_cv - 2*recub_est
            inside_h = h_cv - 2*recub_est
            hook_len_est = 12 * diam_est / 10
            fig_est = draw_stirrup_beam(inside_b, inside_h, hook_len_est, diam_est)
            st.pyplot(fig_est)
            plt.close(fig_est)
            st.caption("Estribo con ganchos de 135° (representación esquemática).")

        # MEMORIA DOCX para Cortante
        if st.button("Generar Memoria Cortante (DOCX)"):
            doc = Document()
            doc.add_heading("Memoria de Diseño a Cortante", 0)
            doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_heading("1. Datos de la Sección", level=1)
            doc.add_paragraph(f"bw = {bw_cv:.0f} cm, d = {d_cv:.0f} cm")
            doc.add_paragraph(f"Concreto: f'c = {fc:.1f} MPa")
            doc.add_paragraph(f"Acero estribos: {st_bar_cv} (Av = {Av_cv:.3f} cm², {n_ramas} ramas)")
            doc.add_heading("2. Verificaciones Normativas", level=1)
            checks = [
                (f"Resistencia a cortante (φVn ≥ Vu)", "CUMPLE" if ok_cv else "NO CUMPLE"),
                ("Vs ≤ Vs,max", "CUMPLE" if Vs_req_kN <= Vs_max_kN else "NO CUMPLE"),
                ("Separación ≤ s_max", "CUMPLE" if s_diseno_mm <= s_max_mm else "NO CUMPLE"),
            ]
            for desc, res in checks:
                doc.add_paragraph(f"{desc}: {res}")
            doc.add_paragraph(f"φVn = {phi_Vn_kN*factor_fuerza:.2f} {unidad_fuerza} | Vu = {Vu_cv_input:.2f} {unidad_fuerza}")
            doc.add_paragraph(f"Separación de diseño: s = {s_diseno_cm:.1f} cm")
            # Artículos normativos específicos por norma
            _art_cortante = {
                "NSR-10":    "NSR-10 C.11.1/C.11.5: El diseño transversal debe asegurar Vu ≤ φ(Vc + Vs), donde Vs asume estribos debidamente espaciados (s_max ≤ d/2).",
                "ACI 318-19":"ACI 318-19 §22.5: Resistencia a cortante (Vu ≤ φVn) y §9.7.6.2.2 sobre el espaciamiento transversal que garanticen interceptar fisuras.",
                "ACI 318-25":"ACI 318-25: Restricción de Vu ≤ φVn y comprobación volumétrica de confinamiento transversal.",
                "NTE E.060": "NTE E.060-2009 Art. 11.1: Revisión a fuerza cortante. La sección debe resistir combinando resistencia del concreto y el acero.",
                "NEC-15":    "NEC-15 Cap 4 Art 20: Verificación de esfuerzo cortante en el alma de los elementos estructurales perimetrales o de piso.",
                "NMX-C":     "RCDF/NMX-C: Revisión transversal para proveer integridad estructural por corte bajo solicitaciones máximas.",
            }
            _art_cv = _art_cortante.get(norma_sel, code['ref'])
            doc.add_paragraph(f"Referencia normativa: {_art_cv}")
            doc.add_paragraph(f"Norma seleccionada: {norma_sel} — Código base: {code['ref']}")

            # Sección transversal 2D en DOCX Cortante
            doc.add_heading("3. Detalle de Sección Transversal", level=1)
            from docx.shared import Inches as _Inches
            fig_cv_doc, ax_cv_doc = sec_light_fig(bw_cv, h_cv, f"Sección Cortante {bw_cv:.0f}×{h_cv:.0f} cm")
            # Estribo
            rec_cv_d = max(recub_cv, 2.5)
            ax_cv_doc.add_patch(patches.Rectangle(
                (rec_cv_d, rec_cv_d), bw_cv-2*rec_cv_d, h_cv-2*rec_cv_d,
                linewidth=2, edgecolor='blue', facecolor='none', linestyle='-'))
            # 4 barras en esquinas (representación)
            r_b_cv = 0.8
            for xb, yb in [(rec_cv_d, rec_cv_d),(bw_cv-rec_cv_d, rec_cv_d),
                           (rec_cv_d, h_cv-rec_cv_d),(bw_cv-rec_cv_d, h_cv-rec_cv_d)]:
                ax_cv_doc.add_patch(plt.Circle((xb, yb), r_b_cv, color='red', zorder=5))
            # Anotación d
            ax_cv_doc.annotate('', xy=(bw_cv*1.1, h_cv-rec_cv_d), xytext=(bw_cv*1.1, h_cv),
                               arrowprops=dict(arrowstyle='<->', color='black'))
            ax_cv_doc.text(bw_cv*1.15, h_cv - rec_cv_d/2, f"d={d_cv:.0f}cm", fontsize=7, va='center')
            buf_cv_doc = io.BytesIO()
            fig_cv_doc.savefig(buf_cv_doc, format='png', dpi=150, bbox_inches='tight')
            buf_cv_doc.seek(0); plt.close(fig_cv_doc)
            doc.add_picture(buf_cv_doc, width=_Inches(3.0))
            doc.add_paragraph(f"Estribo {st_bar_cv} @ {s_diseno_cm:.1f} cm | Av = {Av_cv:.3f} cm² | {n_ramas} ramas")

            doc_mem = io.BytesIO()
            doc.save(doc_mem)
            doc_mem.seek(0)
            st.download_button("Descargar Memoria Cortante", data=doc_mem, file_name=f"Memoria_Cortante_{bw_cv:.0f}x{d_cv:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ??????????????????????????????????????????
# 5. PUNZONAMIENTO EN LOSAS
# ??????????????????????????????????????????
if modulo_sel == " Resistencia a Cortante por Punzonamiento — Losas":
    #  PANEL 
    _panel_normativo("Cortante por Punzonamiento — Losas",
        "Verifica la resistencia al punzonamiento alrededor de columnas según NSR-10 C.11.12.",
        ["Dimensiones de columna c1, c2 (cm)",
         "Espesor de losa h (cm)", "Recubrimiento (cm)",
         "Carga última Vu (kN)"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(_t(f"**Verificación de punzonamiento** (slab-column connection) | Norma: `{code['ref']}`", f"**Punching shear check** (slab-column connection) | Code: `{code['ref']}`"))
    c1,c2 = st.columns(2)
    with c1:
        c1p = st.number_input("Dimensión columna c1 [cm]", 15.0, 100.0, st.session_state.get("pz_c1", 30.0), 5.0, key="pz_c1", help="Dimensión de la columna paralela a la dirección del momento o vano analizado.")
        c2p = st.number_input("Dimensión columna c2 [cm]", 15.0, 100.0, st.session_state.get("pz_c2", 30.0), 5.0, key="pz_c2", help="Dimensión de la columna ortogonal a c1.")
        h_pz = st.number_input("Espesor losa h [cm]", 10.0, 60.0, st.session_state.get("pz_h", 20.0), 1.0, key="pz_h", help="Espesor total de la placa o del ábaco macizo alrededor de la columna.")
    with c2:
        cov_pz = st.number_input("Recubrimiento [cm]", 1.5, 5.0, st.session_state.get("pz_cov", 2.5), 0.5, key="pz_cov")
        Vu_pz_input = st.number_input(f"Vu en columna [{unidad_fuerza}]", 10.0, 10000.0, st.session_state.get("pz_vu", 500.0), 50.0, key="pz_vu", help="Reacción axial mayorada proveniente de la losa que se apoya sobre la columna.")
        if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
            Vu_pz = Vu_pz_input / factor_fuerza
        else:
            Vu_pz = Vu_pz_input
        pz_opts = ["Interior (αs=40)","Borde (αs=30)","Esquina (αs=20)"]
        tipo_col = st.selectbox("Posición columna:", pz_opts, 
                                index=pz_opts.index(st.session_state.pz_tipo) if "pz_tipo" in st.session_state and st.session_state.pz_tipo in pz_opts else 0,
                                key="pz_tipo")
    alpha_s = {"Interior (αs=40)":40,"Borde (αs=30)":30,"Esquina (αs=20)":20}[tipo_col]

    d_pz = (h_pz - cov_pz)*10
    c1_mm = c1p*10; c2_mm = c2p*10
    bo_mm = 2*(c1_mm+d_pz) + 2*(c2_mm+d_pz)
    beta_pz = max(c1p,c2p)/min(c1p,c2p)
    if beta_pz > 2.0:
        st.warning(f" ? **Atención Normativa:** β = {beta_pz:.2f} > 2.0. La columna es significativamente rectangular; la ecuación que rige típicamente será Vc1 (penaliza la capacidad).")
    Vc1_N = (0.17+0.33/beta_pz)*lam*math.sqrt(fc)*bo_mm*d_pz
    Vc2_N = (0.083+0.083*alpha_s*d_pz/bo_mm)*lam*math.sqrt(fc)*bo_mm*d_pz
    Vc3_N = 0.33*lam*math.sqrt(fc)*bo_mm*d_pz
    Vc_pz_N = min(Vc1_N, Vc2_N, Vc3_N)
    phi_Vc_pz = phi_v*Vc_pz_N/1000
    ok_pz = phi_Vc_pz >= Vu_pz

    h_min_req = h_pz
    if not ok_pz:
        for h_test in range(int(h_pz) + 1, 300):
            d_t = (h_test - cov_pz) * 10
            bo_t = 2*(c1_mm+d_t) + 2*(c2_mm+d_t)
            Vc1_t = (0.17+0.33/beta_pz)*lam*math.sqrt(fc)*bo_t*d_t
            Vc2_t = (0.083+0.083*alpha_s*d_t/bo_t)*lam*math.sqrt(fc)*bo_t*d_t
            Vc3_t = 0.33*lam*math.sqrt(fc)*bo_t*d_t
            Vc_pz_t = min(Vc1_t, Vc2_t, Vc3_t)
            if phi_v * Vc_pz_t / 1000 >= Vu_pz:
                h_min_req = h_test
                break

    tab_r,tab_q = st.tabs([f"Resultados {''if ok_pz else '?'}"," Cantidades"])
    with tab_r:
        qty_table([
            ("d efectivo losa", f"{d_pz:.0f} mm = {d_pz/10:.1f} cm"),
            ("β = c_max/c_min", f"{beta_pz:.2f}"),
            ("bo (perímetro crítico)", f"{bo_mm:.0f} mm = {bo_mm/10:.1f} cm"),
            (f"Vc1 (β-fórmula) [{unidad_fuerza}]", f"{Vc1_N/1000*factor_fuerza:.2f}"),
            (f"Vc2 (αs-fórmula) [{unidad_fuerza}]", f"{Vc2_N/1000*factor_fuerza:.2f}"),
            (f"Vc3 (simplificada) [{unidad_fuerza}]", f"{Vc3_N/1000*factor_fuerza:.2f}"),
            (f"Vc diseño = min(Vc1,Vc2,Vc3) [{unidad_fuerza}]", f"{Vc_pz_N/1000*factor_fuerza:.2f}"),
            (f"φ Vc [{unidad_fuerza}]", f"{phi_Vc_pz*factor_fuerza:.2f}"),
            (f"Vu solicitado [{unidad_fuerza}]", f"{Vu_pz_input:.2f}"),
            ("Estado", " CUMPLE" if ok_pz else f"? REFORZAR / Aumentar h a {h_min_req} cm mín."),
        ])
        if ok_pz: st.success(f"φVc = {phi_Vc_pz*factor_fuerza:.2f} {unidad_fuerza} ≥ Vu = {Vu_pz_input:.2f} {unidad_fuerza} — Ref: {code['ref']}")
        else: st.error(f"φVc = {phi_Vc_pz*factor_fuerza:.2f} {unidad_fuerza} < Vu = {Vu_pz_input:.2f} {unidad_fuerza} — Ref: {code['ref']}")
        if not ok_pz:
            st.error(f"? **FALLA POR PUNZONAMIENTO:** El cortante solicitante Vu excede la resistencia del concreto φVc.\n\n"
                     f"**¿QUÉ AUMENTAR? Soluciones propuestas:**\n"
                     f"1. Aumentar el espesor de la losa **h** a por lo menos **{h_min_req} cm**.\n"
                     f"2. Aumentar la resistencia del concreto **f'c**.\n"
                     f"3. Aumentar las dimensiones de la columna o diseñar un ábaco / capitel.")
    with tab_q:
        st.info("**Nota Técnica (Msc):** Si existe **transferencia de momento** excéntrico (Msc) entre la losa y la columna (por ej. fuerzas sísmicas o desequilibrios), una fracción $\\gamma_v M_{sc}$ debe resistirse por excentricidad del cortante, lo que incrementa los esfuerzos $v_u$ en los extremos del perímetro crítico (ref. NSR-10 C.11.11.7 / ACI 318 22.6.4.3).")
        qty_table([("Referencia ACI", "ACI 318-25 Tabla 22.6.5.2"),
                   ("Referencia Norma", code["ref"]),
                   ("Nota","Para casos con Mu en columna verificar momento excéntrico")])

        vol_pz_m3 = (c1p/100 + d_pz/1000) * (c2p/100 + d_pz/1000) * (h_pz/100)
        render_apu_breakdown(vol_pz_m3, 0, fc, "(Concreto en zona de falla)")

    # MEMORIA DOCX para Punzonamiento
    col_p1, col_p2 = st.columns(2)
    with col_p1:
        btn_docx_pz = st.button("Generar Memoria Punzonamiento (DOCX)")
    with col_p2:
        if st.button("Enviar a Cuadro de Mando", key="cmd_pz"):
            add_historial_diseno("Punzonamiento", f"Col: {c1p:.0f}×{c2p:.0f} | h={h_pz:.0f} cm", " APROBADO" if ok_pz else "? EXCEDE", code['ref'])

    if btn_docx_pz:
        doc = Document()
        doc.add_heading("Memoria de Verificación a Punzonamiento", 0)
        doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_heading("1. Datos", level=1)
        doc.add_paragraph(f"Columna: c1={c1p:.0f} cm, c2={c2p:.0f} cm")
        doc.add_paragraph(f"Losa: h={h_pz:.0f} cm, d={d_pz:.0f} mm, recubrimiento={cov_pz:.1f} cm")
        doc.add_paragraph(f"Concreto: f'c={fc:.1f} MPa")
        doc.add_heading("2. Parámetros Geométricos Críticos", level=1)
        Ao_m2 = ((c1p + d_pz/10)/100) * ((c2p + d_pz/10)/100)
        doc.add_paragraph(f"Perímetro crítico (bo): bo = 2·(c1+d) + 2·(c2+d) = {bo_mm:.0f} mm")
        doc.add_paragraph(f"?rea crítica (Ao): Ao = (c1+d)·(c2+d) = {Ao_m2:.3f} m²")
        
        # Diagrama de Punzonamiento
        import matplotlib.pyplot as plt
        import matplotlib.patches as patches
        import io
        from docx.shared import Inches
        fig_bo, ax_bo = plt.subplots(figsize=(4, 4))
        fig_bo.patch.set_facecolor('#1e1e2e')
        ax_arr = fig_bo.get_axes()
        for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
        ax_bo.set_aspect('equal')
        # Columna
        ax_bo.add_patch(patches.Rectangle((-c1p/2, -c2p/2), c1p, c2p, color='gray', alpha=0.5, label='Columna'))
        # Perímetro Crítico (a d/2 de las caras)
        dpz_cm = d_pz / 10
        w_bo = c1p + dpz_cm
        h_bo = c2p + dpz_cm
        ax_bo.add_patch(patches.Rectangle((-w_bo/2, -h_bo/2), w_bo, h_bo, fill=False, edgecolor='red', lw=2, linestyle='--', label='Perímetro Crítico (bo)'))
        ax_bo.set_title("Sección en Planta - Perímetro Crítico")
        ax_bo.set_xlabel("cm"); ax_bo.set_ylabel("cm")
        ax_bo.legend(loc='lower center', bbox_to_anchor=(0.5, -0.3), ncol=2)
        fig_bo.tight_layout()
        buf_bo = io.BytesIO()
        fig_bo.savefig(buf_bo, format='png', dpi=150, bbox_inches='tight')
        buf_bo.seek(0)
        plt.close(fig_bo)
        doc.add_picture(buf_bo, width=Inches(3.5))
        doc.add_heading("3. Resistencias al Cortante (Norma)", level=1)
        doc.add_paragraph(f"Vc1 (β = {beta_pz:.2f}): {Vc1_N/1000*factor_fuerza:.2f} {unidad_fuerza}  --  (0.17+0.33/β)·λ·√f'c·bo·d")
        doc.add_paragraph(f"Vc2 (α = {alpha_s:.0f}): {Vc2_N/1000*factor_fuerza:.2f} {unidad_fuerza}  --  (0.083+0.083·α·d/bo)·λ·√f'c·bo·d")
        doc.add_paragraph(f"Vc3 (Límite max): {Vc3_N/1000*factor_fuerza:.2f} {unidad_fuerza}  --  0.33·λ·√f'c·bo·d")
        doc.add_paragraph(f"Vc de diseño = min(Vc1, Vc2, Vc3) = {Vc_pz_N/1000*factor_fuerza:.2f} {unidad_fuerza}")
        doc.add_heading("4. Verificaciones Normativas", level=1)
        checks = [
            ("Resistencia a punzonamiento (φVc ≥ Vu)", "CUMPLE" if ok_pz else "NO CUMPLE"),
        ]
        for desc, res in checks:
            doc.add_paragraph(f"{desc}: {res}")
        doc.add_paragraph(f"φVc = {phi_Vc_pz*factor_fuerza:.2f} {unidad_fuerza} | Vu = {Vu_pz_input:.2f} {unidad_fuerza}")
        
        _art_pz = {
            "NSR-10": "NSR-10 C.11.11.2: La losa debe dimensionarse en dos direcciones garantizando que Vu ≤ φVc en el perímetro crítico (bo) ubicado a d/2 de la columna.",
            "ACI 318-19": "ACI 318-19 §22.6.5: Resistencia a punzonamiento debe chequear el esfuerzo bidireccional máximo a d/2.",
            "ACI 318-25": "ACI 318-25: Revisión del corte bidireccional y refuerzo si el concreto no excede Vu/φ.",
            "NTE E.060": "NTE E.060-2009 Art. 11.12: Cortante por penetración o punzamiento. Revisión del perímetro crítico.",
            "NEC-15": "NEC-15: Resistencia a la acción combinada de momento y fuerza cortante bidireccional.",
            "NMX-C": "NMX-C: Efectos de cortante cerca del apoyo o áreas de cargas concentradas.",
        }.get(norma_sel, "Sección de Punzonamiento")
        doc.add_paragraph(f"Referencia normativa: {_art_pz} (Código base: {code['ref']})")
        
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button("Descargar Memoria Punzonamiento", data=doc_mem, file_name=f"Memoria_Punzonamiento_{c1p:.0f}x{c2p:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ??????????????????????????????????????????
# 6. INERCIA FISURADA + DEFLEXIONES
# ??????????????????????????????????????????
if modulo_sel == " Inercia Fisurada y Deflexiones en Vigas":
    #  PANEL 
    _panel_normativo("Inercia Fisurada y Deflexiones en Vigas",
        "Calcula Ie (Branson), deflexión inmediata y diferida. Verifica L/360 y L/480.",
        ["Base b (cm)", "Altura h (cm)", "Recubrimiento d' (cm)",
         "Acero provisto As (cm²)", "Longitud L (m)",
         "Cargas de servicio wD y wL (kN/m)"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(f"**Branson (1965) — ACI 318** | {_t('Norma', 'Code')}: `{code['ref']}`")
    c1,c2 = st.columns(2)
    with c1:
        b_de = st.number_input("b [cm]", 10.0, 150.0, st.session_state.get("de_b", 25.0), 5.0, key="de_b")
        h_de = st.number_input("h [cm]", 15.0, 200.0, st.session_state.get("de_h", 50.0), 5.0, key="de_h")
        dp_de = st.number_input("d' [cm]", 2.0, 15.0, st.session_state.get("de_dp", 6.0), 0.5, key="de_dp")
        As_de = st.number_input("As provisto [cm²]", 0.5, 100.0, st.session_state.get("de_as", 5.0), 0.5, key="de_as")
    with c2:
        L_de = st.number_input("Luz libre [m]", 1.0, 20.0, st.session_state.get("de_L", 5.0), 0.5, key="de_L")
        wD_de = st.number_input("Carga muerta wD [kN/m]", 0.0, 200.0, st.session_state.get("de_wD", 15.0), 1.0, key="de_wD")
        wL_de = st.number_input("Carga viva wL [kN/m]", 0.0, 200.0, st.session_state.get("de_wL", 10.0), 1.0, key="de_wL")
        cond_opts = ["Simplemente apoyada","Continua un extremo","Continua dos extremos"]
        cond_de = st.selectbox("Condición de apoyo:", cond_opts, 
                               index=cond_opts.index(st.session_state.de_cond) if "de_cond" in st.session_state and st.session_state.de_cond in cond_opts else 0,
                               key="de_cond")

    d_de = h_de - dp_de
    d_de_mm = d_de*10; b_de_mm = b_de*10; As_de_mm2 = As_de*100
    n_de = Es/Ec
    A_ = b_de_mm/2; B_ = n_de*As_de_mm2; C_ = -n_de*As_de_mm2*d_de_mm
    x_de = (-B_ + math.sqrt(B_**2 - 4*A_*C_))/(2*A_)
    Ig_mm4 = b_de_mm*(h_de*10)**3/12
    Icr_mm4 = b_de_mm*x_de**3/3 + n_de*As_de_mm2*(d_de_mm-x_de)**2
    yt_mm = h_de*10/2
    fr = 0.62*lam*math.sqrt(fc)
    Mcr_Nmm = fr*Ig_mm4/yt_mm
    Mcr_kNm = Mcr_Nmm/1e6

    coef = {"Simplemente apoyada":8,"Continua un extremo":10,"Continua dos extremos":16}[cond_de]
    Ma_D_kNm = wD_de*L_de**2/coef
    Ma_DL_kNm = (wD_de+wL_de)*L_de**2/coef

    def Ie(Ma_kNm, Mcr_kNm, Ig, Icr):
        if Ma_kNm <= 0: return Ig
        ratio = min(Mcr_kNm/Ma_kNm, 1.0)
        return min(ratio**3*Ig + (1-ratio**3)*Icr, Ig)

    Ie_D = Ie(Ma_D_kNm, Mcr_kNm, Ig_mm4, Icr_mm4)
    Ie_DL = Ie(Ma_DL_kNm, Mcr_kNm, Ig_mm4, Icr_mm4)

    fact_defl = {"Simplemente apoyada": 5/384, "Continua un extremo": 1/185, "Continua dos extremos": 1/250}.get(cond_de, 5/384)
    L_mm = L_de*1000
    defl_D_mm = fact_defl*wD_de*(L_mm**4)/(Ec*Ie_D)
    defl_DL_mm = fact_defl*(wD_de+wL_de)*(L_mm**4)/(Ec*Ie_DL)
    defl_L_mm = defl_DL_mm - defl_D_mm
    lim_L480 = L_mm/480; lim_L240 = L_mm/240
    ok_defl_L = defl_L_mm <= lim_L480
    ok_defl_total = defl_DL_mm <= lim_L240

    tab_r,tab_q = st.tabs([f"Resultados {''if (ok_defl_total and ok_defl_L) else '?'}"," Cantidades"])
    with tab_r:
        st.markdown(f"**Ec = {Ec:.0f} MPa** | **n = {n_de:.2f}** | **fr = {fr:.3f} MPa**")
        qty_table([
            ("Ig (inercia bruta)", f"{Ig_mm4:.3e} mm? = {Ig_mm4/1e4:.1f} cm?"),
            ("Eje neutro fisurado (x)", f"{x_de:.1f} mm"),
            ("Icr (inercia fisurada)", f"{Icr_mm4:.3e} mm? = {Icr_mm4/1e4:.1f} cm?"),
            ("fr (módulo de rotura)", f"{fr:.3f} MPa"),
            ("Mcr (momento de agrietamiento)", f"{Mcr_kNm:.2f} kN·m"),
            ("Ma (D)", f"{Ma_D_kNm:.2f} kN·m"),
            ("Ma (D+L)", f"{Ma_DL_kNm:.2f} kN·m"),
            ("Ie (D)", f"{Ie_D:.3e} mm?"),
            ("Ie (D+L)", f"{Ie_DL:.3e} mm?"),
            ("Δ carga muerta D", f"{defl_D_mm:.2f} mm"),
            ("Δ carga viva L", f"{defl_L_mm:.2f} mm"),
            ("Límite L/480 (carga viva)", f"{lim_L480:.1f} mm"),
            ("Δ_L vs L/480", " CUMPLE" if ok_defl_L else "? EXCEDE"),
            ("Límite L/240 (total)", f"{lim_L240:.1f} mm"),
            ("Δ_DL vs L/240", " CUMPLE" if ok_defl_total else "? EXCEDE"),
        ])
        st.caption(f" {code['ref']} | ACI 318-25 Tabla 24.2.2")
    with tab_q:
        qty_table([("Referencia deflexiones","ACI 318-25 Sección 24.2"),
                   ("Norma aplicada", code["ref"])])

    # MEMORIA DOCX para Deflexiones
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        btn_docx_d = st.button("Generar Memoria Deflexiones (DOCX)")
    with col_d2:
        if st.button("Enviar a Cuadro de Mando", key="cmd_de"):
            add_historial_diseno("Deflexiones", f"L={L_de:.2f} m | h={h_de:.0f} cm", " APROBADO" if ok_defl_total else "? EXCEDE", code['ref'])

    if btn_docx_d:
        doc = Document()
        doc.add_heading("Memoria de Deflexiones en Vigas", 0)
        doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_heading("1. Datos", level=1)
        doc.add_paragraph(f"Sección: b={b_de:.0f} cm, h={h_de:.0f} cm, d={d_de:.1f} cm")
        doc.add_paragraph(f"Acero longitudinal: As={As_de:.2f} cm²")
        doc.add_paragraph(f"Luz: L={L_de:.2f} m")
        doc.add_paragraph(f"Cargas: wD={wD_de:.2f} kN/m, wL={wL_de:.2f} kN/m")
        doc.add_heading("2. Propiedades de Sección (Método de Branson)", level=1)
        doc.add_paragraph(f"Inercia Bruta (Ig): {Ig_mm4/1e4:.0f} cm?")
        doc.add_paragraph(f"Inercia Fisurada (Icr): {Icr_mm4/1e4:.0f} cm?")
        doc.add_paragraph(f"Momento de Agrietamiento (Mcr): {Mcr_kNm:.2f} kN·m")
        doc.add_paragraph(f"Momentos Actuantes: Ma(D)={Ma_D_kNm:.2f} kN·m | Ma(D+L)={Ma_DL_kNm:.2f} kN·m")
        doc.add_paragraph(f"Inercia Efectiva Ie(D) = (Mcr/Ma)³·Ig + [1-(Mcr/Ma)³]·Icr = {Ie_D/1e4:.0f} cm?")
        doc.add_paragraph(f"Inercia Efectiva Ie(D+L) = {Ie_DL/1e4:.0f} cm?")
        doc.add_heading("3. Verificaciones Normativas", level=1)
        checks = [
            ("Deflexión por carga viva ≤ L/480", "CUMPLE" if ok_defl_L else "NO CUMPLE"),
            ("Deflexión total ≤ L/240", "CUMPLE" if ok_defl_total else "NO CUMPLE"),
        ]
        for desc, res in checks:
            doc.add_paragraph(f"{desc}: {res}")
        doc.add_paragraph(f"Δ viva = {defl_L_mm:.2f} mm (límite {lim_L480:.1f} mm)")
        doc.add_paragraph(f"Δ total = {defl_DL_mm:.2f} mm (límite {lim_L240:.1f} mm)")
        
        _art_def = {
            "NSR-10": "NSR-10 C.9.5.2 (Deflexiones elásticas y fórmula de Branson)",
            "ACI 318-19": "ACI 318-19 §24.2.3 (Inercia efectiva Ie)",
            "ACI 318-25": "ACI 318-25 §24.2 (Deflexiones e Ie)",
            "NTE E.060": "NTE E.060-2009 Art. 9.5.2 (Cálculo de deflexiones)",
            "NEC-15": "NEC-15 Cap. 4, Art. 18 (Límites de servicio)",
            "NMX-C": "RCDF-2017 / NMX-C Art. 8.3 (Control de deflexiones)",
        }.get(norma_sel, "Sección 24.2")
        doc.add_paragraph(f"Referencia normativa: {_art_def} (Código base: {code['ref']})")
        doc.add_paragraph(
            f"NSR-10 C.9.5.2 (Branson 1965): La inercia efectiva se calcula como "
            f"Ie = (Mcr/Ma)³·Ig + [1-(Mcr/Ma)³]·Icr ≤ Ig. "
            f"Los límites de deflexión son: Δ_L ≤ L/480 para cargas vivas, Δ_total ≤ L/240."
        )

        #  Gráfico Ie vs Ma (Branson) 
        doc.add_heading("4. Gráfico Ie vs Ma — Curva de Branson", level=1)
        _ma_pts = np.linspace(0.01, max(Ma_DL_kNm * 2, Mcr_kNm * 2), 60)
        _ie_pts = [Ie(m, Mcr_kNm, Ig_mm4, Icr_mm4) / 1e4 for m in _ma_pts]
        fig_br, ax_br = plt.subplots(figsize=(5, 3))
        fig_br.patch.set_facecolor('#1e1e2e')
        ax_arr = fig_br.get_axes()
        for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
        ax_br.plot(_ma_pts, _ie_pts, color='steelblue', lw=2, label='Ie(Ma)')
        ax_br.axhline(Ig_mm4 / 1e4, color='gray', linestyle=':', lw=1, label=f'Ig = {Ig_mm4/1e4:.0f} cm?')
        ax_br.axhline(Icr_mm4 / 1e4, color='brown', linestyle=':', lw=1, label=f'Icr = {Icr_mm4/1e4:.0f} cm?')
        ax_br.axvline(Mcr_kNm, color='orange', linestyle='--', lw=1, label=f'Mcr = {Mcr_kNm:.2f} kN·m')
        ax_br.axvline(Ma_DL_kNm, color='red', linestyle='--', lw=1, label=f'Ma(D+L) = {Ma_DL_kNm:.2f} kN·m')
        ax_br.set_xlabel('Ma (kN·m)'); ax_br.set_ylabel('Ie (cm?)')
        ax_br.legend(fontsize=7); ax_br.grid(alpha=0.3); fig_br.tight_layout()
        buf_br = io.BytesIO(); fig_br.savefig(buf_br, format='png', dpi=150, bbox_inches='tight'); buf_br.seek(0)
        plt.close(fig_br)
        from docx.shared import Inches as _InchesD
        doc.add_picture(buf_br, width=_InchesD(4.5))

        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button("Descargar Memoria Deflexiones", data=doc_mem, file_name="Memoria_Deflexiones.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ??????????????????????????????????????????
# 7. LOSA EN UNA DIRECCIÓN
# ??????????????????????????????????????????
if modulo_sel == " Diseño de Losa en Una Dirección":
    st.markdown(_t(f"**Diseño por franja de 1 metro** | Norma: `{code['ref']}`", f"**1-meter strip design** | Code: `{code['ref']}`"))
    c1,c2 = st.columns(2)
    with c1:
        ln_ls = st.number_input("Luz libre ln [m]", 1.0, 10.0, st.session_state.get("ls_ln", 3.5), 0.25, key="ls_ln")
        h_ls  = st.number_input("Espesor losa h [cm]", 8.0, 40.0, st.session_state.get("ls_h", 15.0), 1.0, key="ls_h")
        cov_ls = st.number_input("Recubrimiento [cm]", 1.5, 5.0, st.session_state.get("ls_cov", 2.5), 0.5, key="ls_cov")
        wD_ls = st.number_input("Carga muerta total (D) [kN/m²]", 1.0, 50.0, st.session_state.get("ls_wD", 7.0), 0.5, key="ls_wD")
    with c2:
        wL_ls = st.number_input("Carga viva (L) [kN/m²]", 0.5, 30.0, st.session_state.get("ls_wL", 5.0), 0.5, key="ls_wL")
        varillas_ls = list(rebar_dict.keys())
        bar_ls = st.selectbox("Varilla losa:", varillas_ls, 
                              index=varillas_ls.index(st.session_state.ls_bar) if "ls_bar" in st.session_state and st.session_state.ls_bar in varillas_ls else 0,
                              key="ls_bar")
        Ab_ls = rebar_dict[bar_ls]; db_ls = diam_dict[bar_ls]
        apoyo_opts = ["Simplemente apoyada","Continua 2 extremos","Voladizo"]
        apoyo_ls = st.selectbox("Condición:", apoyo_opts, 
                                index=apoyo_opts.index(st.session_state.ls_apoyo) if "ls_apoyo" in st.session_state and st.session_state.ls_apoyo in apoyo_opts else 0,
                                key="ls_apoyo")

    wu_ls = 1.2*wD_ls + 1.6*wL_ls
    b_ls = 100
    d_ls = h_ls - cov_ls - db_ls/20
    coef_ls = {"Simplemente apoyada":8,"Continua 2 extremos":16,"Voladizo":2}[apoyo_ls]
    Mu_ls_kNm = wu_ls*ln_ls**2/coef_ls

    d_ls_mm = d_ls*10; b_ls_mm = b_ls*10
    Rn_ls = Mu_ls_kNm*1e6/(phi_f*b_ls_mm*d_ls_mm**2)
    disc_ls = 1-2*Rn_ls/(0.85*fc)
    if disc_ls < 0:
        d_min_mm = math.sqrt((Mu_ls_kNm * 1e6) / (phi_f * b_ls_mm * 0.425 * fc))
        h_sug_cm = math.ceil(d_min_mm / 10 + cov_ls + db_ls/20)
        st.error(f"? Losa muy delgada o carga muy alta. El concreto a compresión es insuficiente para resistir el momento. $\\rightarrow$ **Sugerencia:** Aumente $h$ al menos a **{h_sug_cm} cm**.")
        st.stop()
    else:
        rho_ls = (0.85*fc/fy)*(1-math.sqrt(disc_ls))
        rho_use_ls = max(rho_ls, rho_min)
        As_req_ls = rho_use_ls*b_ls*d_ls
        s_bar_ls = Ab_ls/As_req_ls*100
        s_max_ls = min(3*h_ls, 45)
        s_use_ls = min(s_bar_ls, s_max_ls)
        As_prov_ls = Ab_ls/(s_use_ls/100)
        a_ls_mm = As_prov_ls*100*fy/(0.85*fc*b_ls_mm)
        phi_Mn_ls = phi_f*As_prov_ls*100*fy*(d_ls_mm-a_ls_mm/2)/1e6
        # rho_temp depends on fy (NSR-10 C.7.12.2)
        rho_temp_ls = 0.0020 if fy <= 280 else (0.0018 * 420/fy if fy > 420 else 0.0018)
        As_temp = rho_temp_ls*b_ls*h_ls
        s_temp = min(Ab_ls/As_temp*100, 5*h_ls, 45)

        ok_ls = phi_Mn_ls >= Mu_ls_kNm
        tab_r,tab_s,tab_g,tab_3d,tab_q = st.tabs([f"Resultados {''if ok_ls else '?'}"," Sección 2D"," Gráficos M/V"," 3D"," Cantidades"])
        with tab_r:
            qty_table([
                ("wu factorizada", f"{wu_ls:.2f} kN/m²"),
                ("Mu (franja 1m)", f"{Mu_ls_kNm:.2f} kN·m"),
                ("h losa / d efectivo", f"{h_ls:.0f} / {d_ls:.1f} cm"),
                ("As requerido", f"{As_req_ls:.3f} cm²/m"),
                (f"Espaciado varilla {bar_ls}", f"{s_bar_ls:.1f} cm → usar **{s_use_ls:.1f} cm**"),
                ("As provisto", f"{As_prov_ls:.3f} cm²/m"),
                ("φMn / Mu", f"{phi_Mn_ls:.2f} / {Mu_ls_kNm:.2f} kN·m/m"),
                ("Estado Flexión", " CUMPLE" if ok_ls else "? DEFICIENTE"),
                ("As temperatura/retracción", f"{As_temp:.3f} cm²/m"),
                (f"Varilla temp {bar_ls}", f"@ {s_temp:.1f} cm"),
            ])
            if ok_ls: st.success(f"Losa OK — {bar_ls} @ {s_use_ls:.1f} cm (As principal)")
            else: st.error(f"Losa DEFICIENTE — {bar_ls} @ {s_use_ls:.1f} cm (As principal)")
            st.info("**Acero Inferior vs. Superior:** Este diseño automático a partir de la luz libre (L) estima el máximo momento positivo de la franja. El resultado `As principal` mostrado es el **acero inferior**. Para el **acero superior** necesario en los nudos continuos, considere diseñar a flexión una viga de b=100m ingresando el Mu- de los apoyos.")
        
        with tab_s:
            SCALE = 2.5  # Factor: eje Y va de 0 a h_ls*2.5, eje X va de 0 a 40 (=100cm/2.5)
            fig_s, ax_s = sec_dark_fig(40, h_ls*SCALE, f"Sección Losa — h={h_ls:.0f}cm")
            r_ls = db_ls / 10 * SCALE / 2  # radio escalado (db en mm → cm /10, luego *SCALE)
            r_ls = max(r_ls, 0.5)  # mínimo visible
            y_bar_s = cov_ls * SCALE / 100 * 100 + r_ls  # cov_ls [cm] * SCALE → unidades del eje
            # El ancho del gráfico es 40 unidades = 100 cm → 1 cm = 0.40 unidades
            x_scale = 40 / 100  # 0.40 unidades/cm
            x_start = cov_ls * x_scale + r_ls
            x_end = 40 - cov_ls * x_scale - r_ls
            s_use_scaled = s_use_ls * x_scale
            if s_use_scaled > 0:
                for xi in np.arange(x_start, x_end + 0.01, s_use_scaled):
                    ax_s.add_patch(plt.Circle((xi, y_bar_s), r_ls, color='#ff6b35', zorder=5))
            st.pyplot(fig_s)
            plt.close(fig_s)
        
        with tab_g:
            st.subheader("Gráficos de Cortante (V) y Momento Flector (M)")
            x_vals = np.linspace(0, ln_ls, 100)
            if apoyo_ls == "Simplemente apoyada":
                V_vals = wu_ls * ln_ls / 2 - wu_ls * x_vals
                M_vals = wu_ls * ln_ls * x_vals / 2 - wu_ls * x_vals**2 / 2
            elif apoyo_ls == "Continua 2 extremos":
                V_vals = wu_ls * ln_ls / 2 - wu_ls * x_vals
                M_vals = wu_ls * ln_ls * x_vals / 2 - wu_ls * x_vals**2 / 2 - wu_ls * ln_ls**2 / 12
            else: # Voladizo
                V_vals = -wu_ls * x_vals 
                M_vals = -wu_ls * x_vals**2 / 2
            
            fig_mv, (ax_v, ax_m) = plt.subplots(2, 1, figsize=(6, 5), sharex=True)
            fig_mv.patch.set_facecolor('#1e1e2e')
            ax_arr = fig_mv.get_axes()
            for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
            fig_mv.patch.set_facecolor('#1a1a2e')
            ax_v.set_facecolor('#1a1a2e'); ax_m.set_facecolor('#1a1a2e')
            
            ax_v.plot(x_vals, V_vals, color='#00d4ff', lw=2)
            ax_v.fill_between(x_vals, 0, V_vals, color='#00d4ff', alpha=0.3)
            ax_v.axhline(0, color='white', lw=1)
            ax_v.set_ylabel("Cortante (kN)", color='white')
            ax_v.tick_params(colors='white')
            
            ax_m.plot(x_vals, M_vals, color='#ff6b35', lw=2)
            ax_m.fill_between(x_vals, 0, M_vals, color='#ff6b35', alpha=0.3)
            ax_m.axhline(0, color='white', lw=1)
            ax_m.set_ylabel("Momento (kN·m/m)", color='white')
            ax_m.set_xlabel("Distancia x (m)", color='white')
            ax_m.tick_params(colors='white')
            fig_mv.tight_layout()
            st.pyplot(fig_mv)
            plt.close(fig_mv)

        with tab_3d:
            st.subheader("Losa 3D (Franja de 1m)")
            fig3_ls = go.Figure()
            L_ls_cm = ln_ls * 100
            x_l = [-50, 50, 50, -50, -50, 50, 50, -50]
            y_l = [0, 0, h_ls, h_ls, 0, 0, h_ls, h_ls]
            z_l = [0, 0, 0, 0, L_ls_cm, L_ls_cm, L_ls_cm, L_ls_cm]
            fig3_ls.add_trace(go.Mesh3d(x=x_l, y=y_l, z=z_l, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
            
            d_real_ls = cov_ls + db_ls/20
            y_bar = d_real_ls if apoyo_ls != "Voladizo" else h_ls - d_real_ls
            n_bars_1m = int(100/s_use_ls)
            if n_bars_1m < 1: n_bars_1m = 1
            xs_ls = np.linspace(-50 + s_use_ls/2, 50 - s_use_ls/2, n_bars_1m)
            line_w_ls = max(3, db_ls/10 * 3)
            for idx, xb in enumerate(xs_ls):
                fig3_ls.add_trace(go.Scatter3d(x=[xb, xb], y=[y_bar, y_bar], z=[0, L_ls_cm],
                                              mode='lines', line=dict(color='darkred', width=line_w_ls),
                                              name=f'Principal {bar_ls}', showlegend=(idx==0)))
            
            y_temp = y_bar + db_ls/10 if apoyo_ls != "Voladizo" else y_bar - db_ls/10
            line_w_t = max(2, db_ls/10 * 2)
            n_temp = int(L_ls_cm / s_temp)
            if n_temp > 0:
                z_temp_vals = np.linspace(s_temp, L_ls_cm-s_temp, n_temp)
                tx_ls, ty_ls, tz_ls = [], [], []
                for zt in z_temp_vals:
                    tx_ls.extend([-50, 50, None])
                    ty_ls.extend([y_temp, y_temp, None])
                    tz_ls.extend([zt, zt, None])
                fig3_ls.add_trace(go.Scatter3d(x=tx_ls, y=ty_ls, z=tz_ls, mode='lines', 
                                               line=dict(color='cornflowerblue', width=line_w_t), name='Temperatura', showlegend=True))
                                               
            fig3_ls.update_layout(scene=dict(aspectmode='data', xaxis_title='Ancho (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                                margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable')
            st.plotly_chart(fig3_ls, use_container_width=True)

        with tab_q:
            area_losa = ln_ls*1
            vol_ls = area_losa*h_ls/100
            peso_flex_ls = As_prov_ls * ln_ls * 0.785
            peso_temp_ls = As_temp * ln_ls * 0.785
            m = mix_for_fc(fc)
            bags = m[0]*vol_ls/bag_kg
            qty_table([
                ("Concreto (1m de ancho)", f"{vol_ls:.4f} m³"),
                (f"Acero flexión {bar_ls} @ {s_use_ls:.1f}cm", f"{peso_flex_ls:.2f} kg/m"),
                (f"Acero temp {bar_ls} @ {s_temp:.1f}cm", f"{peso_temp_ls:.2f} kg/m"),
                (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{m[0]*vol_ls/bag_kg:.1f} bultos"),
                ("Referencia", code["ref"]),
            ])

            render_apu_breakdown(vol_ls, peso_flex_ls + peso_temp_ls, fc, "(Refuerzo principal y temp.)")

            with st.expander("Dibujo de Figurado para Taller (Varillas de losa)", expanded=False):
                hook_len_cm = 12 * db_ls / 10
                straight_len_cm = ln_ls * 100 - 2 * hook_len_cm
                fig_l1 = draw_longitudinal_bar(ln_ls*100, straight_len_cm, hook_len_cm, db_ls)
                st.pyplot(fig_l1)
                plt.close(fig_l1)
                st.caption("Varilla de refuerzo principal con ganchos de 90° en extremos (para losa simplemente apoyada o continua).")

            # MEMORIA DOCX para Losa
            col_l1, col_l2, col_l3 = st.columns(3)
            with col_l1:
                btn_docx_ls = st.button("Generar Memoria Losa (DOCX)")
            with col_l2:
                if st.button("Enviar a Cuadro de Mando", key="cmd_ls"):
                    add_historial_diseno("Losa Aligerada", f"{ancho_nerv_ls}×{h_ls} cm (s={separa_ls} cm)", " APROBADO" if ok_ls else "? NO CUMPLE", code['ref'])
            with col_l3:
                try:
                    buf_ifc_ls = ifc_export.ifc_losa(
                        h_ls, ln_ls, 1.0, fc, fy, bar_ls, db_ls, min(As_req_ls, As_prov_ls),
                        s_fin_ls, cov_ls, norma_sel, "Proyecto NSR-10"
                    )
                    st.download_button("Exportar IFC (BIM)", data=buf_ifc_ls, file_name=f"Losa_{h_ls:.0f}cm.ifc", mime="application/x-step", key="ifc_losa")
                except Exception as e:
                    st.error(f"Error IFC: {e}")
                    
            if btn_docx_ls:
                # N4: Crear figura de Momentos/Cortante en modo claro para DOCX
                # (los fill_between con alpha sobre fondo oscuro producen manchas negras en Word)
                import io as _io_mv
                fig_mv_doc, (ax_v_doc, ax_m_doc) = plt.subplots(2, 1, figsize=(6, 5), sharex=True)
                fig_mv_doc.patch.set_facecolor('white')
                ax_v_doc.set_facecolor('white')
                ax_m_doc.set_facecolor('white')
                # Colores sólidos sin alpha para evitar manchas negras en PDF/Word
                ax_v_doc.plot(x_vals, V_vals, color='#0077bb', lw=2)
                ax_v_doc.fill_between(x_vals, 0, V_vals, color='#aad4f5')  # azul pálido sólido
                ax_v_doc.axhline(0, color='#333333', lw=1)
                ax_v_doc.set_ylabel("Cortante (kN)", color='#1a1a1a')
                ax_v_doc.tick_params(colors='#1a1a1a')
                ax_v_doc.spines['bottom'].set_color('#555555')
                ax_v_doc.spines['left'].set_color('#555555')
                ax_m_doc.plot(x_vals, M_vals, color='#cc4400', lw=2)
                ax_m_doc.fill_between(x_vals, 0, M_vals, color='#f5c0a0')  # naranja pálido sólido
                ax_m_doc.axhline(0, color='#333333', lw=1)
                ax_m_doc.set_ylabel("Momento (kN·m/m)", color='#1a1a1a')
                ax_m_doc.set_xlabel("Distancia x (m)", color='#1a1a1a')
                ax_m_doc.tick_params(colors='#1a1a1a')
                ax_m_doc.spines['bottom'].set_color('#555555')
                ax_m_doc.spines['left'].set_color('#555555')
                fig_mv_doc.tight_layout()
                buf_mv = _io_mv.BytesIO()
                fig_mv_doc.savefig(buf_mv, format='png', bbox_inches='tight', dpi=150, facecolor='white')
                buf_mv.seek(0)
                plt.close(fig_mv_doc)

                doc = Document()
                doc.add_heading("Memoria de Diseño de Losa en Una Dirección", 0)
                doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_heading("1. Datos", level=1)
                doc.add_paragraph(f"Luz libre: ln = {ln_ls:.2f} m")
                doc.add_paragraph(f"Espesor losa: h = {h_ls:.1f} cm")
                doc.add_paragraph(f"Recubrimiento: d' = {cov_ls:.1f} cm → d = {d_ls:.1f} cm")
                doc.add_paragraph(f"Materiales: f'c = {fc:.1f} MPa, fy = {fy:.0f} MPa")
                doc.add_paragraph(f"Cargas: wD = {wD_ls:.2f} kN/m², wL = {wL_ls:.2f} kN/m² → wu = {wu_ls:.2f} kN/m²")
                doc.add_heading("2. Refuerzo", level=1)
                doc.add_paragraph(f"Momento último Mu = {Mu_ls_kNm:.2f} kN·m/m")
                doc.add_paragraph(f"Armadura principal: {bar_ls} @ {s_use_ls:.1f} cm → As = {As_prov_ls:.3f} cm²/m")
                doc.add_paragraph(f"Armadura de temperatura: {bar_ls} @ {s_temp:.1f} cm → As_temp = {As_temp:.3f} cm²/m")
                doc.add_heading("3. Verificaciones Normativas", level=1)
                checks = [
                    ("Resistencia a flexión (φMn ≥ Mu)", "CUMPLE" if ok_ls else "NO CUMPLE"),
                    ("Espaciamiento ≤ 3h y ≤ 45 cm", "CUMPLE" if s_use_ls <= s_max_ls else "NO CUMPLE"),
                    ("As_temp ≥ ρ_min_temp·b·h", "CUMPLE" if As_temp >= rho_temp_ls*b_ls*h_ls else "NO CUMPLE"),
                ]
                for desc, res in checks:
                    doc.add_paragraph(f"{desc}: {res}")
                
                doc.add_heading("4. Diagramas de Cortante y Momento", level=1)
                from docx.shared import Inches
                doc.add_picture(buf_mv, width=Inches(5.0))

                doc.add_heading("5. Detalle de Sección (1m de ancho)", level=1)
                fig_sdoc, ax_sdoc = sec_light_fig(100, h_ls, f"Losa Aligerada h={h_ls:.1f} cm")
                r_ls_doc = db_ls / 20
                y_bar_doc = cov_ls
                x_start_doc = cov_ls + r_ls_doc
                x_end_doc = 100 - cov_ls - r_ls_doc
                if s_use_ls > 0:
                    for xi in np.arange(x_start_doc, x_end_doc + 0.01, s_use_ls):
                        ax_sdoc.add_patch(plt.Circle((xi, y_bar_doc), max(r_ls_doc, 0.5), color='red', zorder=5))
                buf_ls_sec = io.BytesIO()
                fig_sdoc.savefig(buf_ls_sec, format='png', dpi=150, bbox_inches='tight')
                buf_ls_sec.seek(0)
                plt.close(fig_sdoc)
                doc.add_picture(buf_ls_sec, width=Inches(4.5))
                doc.add_paragraph(f"φMn = {phi_Mn_ls:.2f} kN·m/m | Mu = {Mu_ls_kNm:.2f} kN·m/m")
                # Artículo normativo dinámico para losa
                _art_losa = {
                    "NSR-10":     "NSR-10 C.9.5: El espesor debe mitigar deflexiones no permitidas. C.7.12: Colocar refuerzo mínimo (0.0018 o más) para controlar efectos de temperatura y contracción de secado.",
                    "ACI 318-19": "ACI 318-19 §24.2: Espesores requeridos. §24.4: Refuerzo bidireccional para retracción transversal a la flexión.",
                    "ACI 318-25": "ACI 318-25: Límite de flecha (L/240) y acero de contracción en placas continuas.",
                    "NTE E.060":  "NTE E.060-2009 Art 9.5 y 7.12: Se requiere peralte capaz de controlar fisuras y acero transversal mínimo contra temperatura.",
                    "NEC-15":     "NEC-15: Revisión de deflexión instantánea y diferida; y área mínima transversal.",
                    "NMX-C":      "RCDF/NMX-C: Cuantías mínimas garantizando ductilidad y acero de retracción en losas.",
                }.get(norma_sel, code['ref'])
                doc.add_paragraph(f"Referencia normativa: {_art_losa}")
                doc.add_paragraph(f"Norma: {norma_sel} — Código base: {code['ref']}")
                doc.add_heading("4. Diagramas de Momento y Cortante", level=1)
                doc.add_picture(buf_mv, width=Inches(5))
                
                doc.add_heading("5. Cantidades", level=1)
                doc.add_paragraph(f"Concreto: {vol_ls:.4f} m³ por metro de ancho")
                doc.add_paragraph(f"Acero principal: {peso_flex_ls:.2f} kg/m")
                doc.add_paragraph(f"Acero temperatura: {peso_temp_ls:.2f} kg/m")
                doc_mem = io.BytesIO()
                doc.save(doc_mem)
                doc_mem.seek(0)
                st.download_button("Descargar Memoria Losa", data=doc_mem, file_name=f"Memoria_Losa_{ln_ls:.2f}m.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ??????????????????????????????????????????
# 8. LONGITUD DE DESARROLLO (sin cambios)
# ??????????????????????????????????????????
if modulo_sel == " Longitud de Desarrollo y Empalmes":
    st.markdown(f"**Barras rectas a tracción** | Norma: `{code['ref']}`")
    c1,c2 = st.columns(2)
    with c1:
        varillas_ld = list(rebar_dict.keys())
        bar_ld = st.selectbox("Varilla:", varillas_ld, 
                              index=varillas_ld.index(st.session_state.ld_bar) if "ld_bar" in st.session_state and st.session_state.ld_bar in varillas_ld else 0,
                              key="ld_bar")
        db_ld = diam_dict[bar_ld]
        psit_opts = ["1.3 — Barra superior (>30cm betón fresco abajo)","1.0 — Otras posiciones"]
        psi_t = st.selectbox("ψt (posición):", psit_opts, 
                             index=psit_opts.index(st.session_state.ld_psit) if "ld_psit" in st.session_state and st.session_state.ld_psit in psit_opts else 1,
                             key="ld_psit")
        psie_opts = ["1.5 — Ep. y >3db o <6mm cub.","1.2 — Otros epoxy","1.0 — Sin epoxy"]
        psi_e = st.selectbox("ψe (epoxy):", psie_opts, 
                             index=psie_opts.index(st.session_state.ld_psie) if "ld_psie" in st.session_state and st.session_state.ld_psie in psie_opts else 2,
                             key="ld_psie")
    with c2:
        psis_opts = ["0.8 — Barras ≤ #6 ó ≤19mm","1.0 — Barras > #6 ó >19mm"]
        psi_s = st.selectbox("ψs (tamaño):", psis_opts, 
                             index=psis_opts.index(st.session_state.ld_psis) if "ld_psis" in st.session_state and st.session_state.ld_psis in psis_opts else (0 if db_ld <= 19 else 1),
                             key="ld_psis")
        psig_opts = ["0.75 — fy ≤ 420 MPa (ACI 318-19+)","1.0 — fy > 420 MPa"]
        psi_g = st.selectbox("ψg (resistencia):", psig_opts, 
                             index=psig_opts.index(st.session_state.ld_psig) if "ld_psig" in st.session_state and st.session_state.ld_psig in psig_opts else (0 if fy <= 420 else 1),
                             key="ld_psig")
        cb_ld = st.number_input("cb (recubrim. al centro barra) [mm]", 20.0, 100.0, st.session_state.get("ld_cb", 40.0), 2.5, key="ld_cb")
        Ktr_ld = st.number_input("Ktr (refuerzo transversal, 0 si no hay) [mm]", 0.0, 50.0, st.session_state.get("ld_ktr", 0.0), 1.0, key="ld_ktr")

    psit_v = float(psi_t.split("—")[0])
    psie_v = float(psi_e.split("—")[0])
    psis_v = float(psi_s.split("—")[0])
    psig_v = float(psi_g.split("—")[0])

    cb_ktr_db = min((cb_ld+Ktr_ld)/db_ld, 2.5)
    psi_prod = min(psit_v * psie_v, 1.7) * psis_v * psig_v
    ld_mm = (3/40)*(fy/lam/math.sqrt(fc))*(psi_prod/cb_ktr_db)*db_ld
    ld_mm = max(ld_mm, 300)

    ls_A = 1.0*ld_mm
    ls_B = 1.3*ld_mm
    ldh_mm = max((0.02*1*1*1*1*fy)/(lam*math.sqrt(fc))*db_ld, 8*db_ld, 150)

    tab_r, = st.tabs(["Resultados"])
    with tab_r:
        st.markdown(f"**φ no aplica para longitud de desarrollo** | Ref: `{code['ref']}`")
        qty_table([
            ("Varilla", f"{bar_ld} — db = {db_ld:.2f} mm"),
            ("f'c / fy", f"{fc:.1f} MPa / {fy:.0f} MPa"),
            ("(cb+Ktr)/db", f"{cb_ktr_db:.3f} {' ≤2.5' if cb_ktr_db<=2.5 else '→ limitado a 2.5'}"),
            ("ψt × ψe × ψs × ψg", f"{psit_v}×{psie_v}×{psis_v}×{psig_v} = {psit_v*psie_v*psis_v*psig_v:.3f}"),
            (" Nota: ψt×ψe ≤ 1.7", "" if psit_v*psie_v<=1.7 else " ? Limitar a 1.7"),
            ("ld (barra recta en tensión)", f"**{ld_mm:.0f} mm = {ld_mm/10:.1f} cm**"),
            ("Empalme Clase A (ld×1.0)", f"{ls_A:.0f} mm = {ls_A/10:.1f} cm"),
            ("Empalme Clase B (ld×1.3)", f"{ls_B:.0f} mm = {ls_B/10:.1f} cm"),
            ("ldh (gancho estándar 90°)", f"{ldh_mm:.0f} mm = {ldh_mm/10:.1f} cm"),
            ("Referencia", f"{code['ref']} / ACI 318-25 Sección 25.4"),
        ])
        if psit_v*psie_v > 1.7:
            st.warning(" ? El producto ψt×ψe no puede exceder 1.7 (ACI 318-25 25.4.2.5)")
        st.success(f"ld = {ld_mm:.0f} mm | Empalme B = {ls_B:.0f} mm")

# ??????????????????????????????????????????
# 9. DISEÑO S?SMICO Y DXF VIGA CONTINUA
# ??????????????????????????????????????????
if modulo_sel == " Diseño Sísmico Integral y Plano DXF (Viga DMO / DES)":
    #  PANEL 
    _panel_normativo("Diseño Sísmico Integral — Viga DMO / DES",
        "Verificaciones C.21.5 NSR-10: cuantía, momento mínimo, cortante plástico Vp y plano DXF.",
        ["Geometría completa de la viga (bw, h, L, bcol)",
         "Momentos en apoyos (Mu izq/der superior e inferior) y centro",
         "Cortante último Vu", "Barras superiores e inferiores en cada zona",
         "Nivel sísmico DMO o DES activo en el sidebar"],
        fc, fy, rho_min, rho_max, phi_f, phi_v, norma_sel, nivel_sis, code['ref'])
    st.markdown(_t(f"**Diseño completo de un vano con rótulas plásticas y detallado CAD** | Norma: `{code['ref']}`", f"**Full span design with plastic hinges & CAD detail** | Code: `{code['ref']}`"))
    
    # Reconocer nivel de disipación
    es_des = nivel_sis in ["DES — Disipación Especial", "SMF (SDC D–F)", "PE — Pórtico Especial (Z3–Z4)", "MRLE — Ductilidad Alta", "GA — Grado Alto", "PE — Pórtico Especial", "DE — Diseño Especial Sísmico", "GA — Ductilidad Alta"]
    es_dmo = nivel_sis in ["DMO — Disipación Moderada", "IMF (SDC C)", "PM — Pórtico Moderado", "GM — Ductilidad Moderada", "GM — Grado Moderado", "MROD — Ductilidad Ordinaria"]
    es_sismico = es_des or es_dmo

    if not es_sismico:
        st.info("ℹ? El nivel de amenaza seleccionado corresponde a **Disipación Mínima u Ordinaria**. Las reglas especiales de nudos, cortante Vp y confinamiento C.21.5 no son estrictamente obligatorias, pero el módulo diseñará basándose en el estándar de pórticos de momento.")
    else:
        st.warning(f"**Nivel {nivel_sis} Activado:** Se evaluarán todas las reglas estrictas de geometría, nudos y capacidad ($V_p$) aplicables a zonas sísmicas.")

    #  BLOQUE 4: TABLA DE PAR?METROS Y β? EXPL?CITO 
    beta1_tab = 0.85 if fc <= 28 else max(0.65, 0.85 - 0.05*(fc - 28)/7)
    Ec_tab    = 4700 * math.sqrt(fc)
    eps_ty    = fy / 200000
    eps_min_t = eps_ty + 0.003
    c_max_t   = (0.003 / (0.003 + eps_min_t)) * 1.0
    a_max_t   = beta1_tab * c_max_t

    if fc <= 28:
        beta1_texto = f"f'c ≤ 28 MPa → β? = 0.85"
    elif fc < 55:
        beta1_texto = f"28 < f'c < 55 MPa → β? = 0.85 − 0.05·(f'c−28)/7 = {beta1_tab:.4f}"
    else:
        beta1_texto = f"f'c ≥ 55 MPa → β? = 0.65"

    st.markdown("####  Parámetros de Diseño")
    col_p1, col_p2 = st.columns(2)

    with col_p1:
        st.markdown(f"""
<div style="background:#1c1c2e; border-radius:8px; padding:12px; font-family:monospace; font-size:13px; color:white; line-height:2.0;">
  <b>Materiales</b><br>
  f'c = {fc:.1f} MPa<br>
  fy  = {fy:.0f} MPa<br>
  Es  = 200,000 MPa<br>
  Ec  = {Ec_tab:.0f} MPa<br>
  λ   = 1.0 (concreto normal)<br>
  φ   = {phi_f:.2f} (flexión)
</div>
""", unsafe_allow_html=True)

    with col_p2:
        st.markdown(f"""
<div style="background:#1c1c2e; border-radius:8px; padding:12px; font-family:monospace; font-size:13px; color:white; line-height:2.0;">
  <b>Parámetros de ductilidad</b><br>
  β?  = {beta1_tab:.4f} &nbsp; <span style="color:#aaa;font-size:11px;">({beta1_texto})</span><br>
  εᵤ  = 0.003 (deformación última concreto)<br>
  εᵧ  = fy/Es = {eps_ty:.5f}<br>
  εₘᵢₙ = εᵧ + 0.003 = {eps_min_t:.5f} <span style="color:#aaa;font-size:11px;">(sección dúctil)</span><br>
  c_máx/d = ε_cu/(ε_cu+εₘᵢₙ) = {c_max_t:.4f}<br>
  a_máx/d = β?·c_máx/d = {a_max_t:.4f}
</div>
""", unsafe_allow_html=True)

    st.caption("NSR-10 C.10.3.3 — La deformación neta en tracción del acero más alejado εₜ debe ser ≥ εₘᵢₙ = εᵧ + 0.003 para garantizar comportamiento dúctil.")
    st.markdown("---")

    st.subheader("1. Geometría y Solicitaciones")
    c1, c2, c3 = st.columns(3)
    with c1:
        b_vc = st.number_input("Ancho viga bw [cm]", 15.0, 150.0, st.session_state.get("vc_b", 30.0), 5.0, key="vc_b")
        h_vc = st.number_input("Alto viga h [cm]", 20.0, 200.0, st.session_state.get("vc_h", 50.0), 5.0, key="vc_h")
        dp_vc = st.number_input("Recubrimiento al centro de la barra d' [cm]", 3.0, 15.0, st.session_state.get("vc_dp", 6.0), 0.5, key="vc_dp")
    with c2:
        L_vc = st.number_input("Luz Libre Ln [m] (cara a cara)", 1.0, 20.0, st.session_state.get("vc_L", 5.0), 0.1, key="vc_L")
        bcol_vc = st.number_input("Ancho Columna de Apoyo bcol [cm]", 15.0, 200.0, st.session_state.get("vc_bcol", 40.0), 5.0, key="vc_bcol")
        Wu_vc_input = st.number_input(f"Carga Gravitacional Wu factorizada [{unidad_fuerza}/m]", 0.0, 1000.0, st.session_state.get("vc_wu", 30.0), 5.0, key="vc_wu")
        Wu_vc = Wu_vc_input / factor_fuerza if unidades_salida == "Toneladas fuerza (tonf, tonf·m)" else Wu_vc_input
    with c3:
        st.markdown("**Momentos Sísmicos (Absolutos)**")
        Mu_izq_neg_in = st.number_input(f"Mu Nudo Izq (Negativo) [{unidad_mom}]", 0.0, 10000.0, st.session_state.get("vc_mu_izq", 150.0), 10.0, key="vc_mu_izq")
        Mu_der_neg_in = st.number_input(f"Mu Nudo Der (Negativo) [{unidad_mom}]", 0.0, 10000.0, st.session_state.get("vc_mu_der", 120.0), 10.0, key="vc_mu_der")
        Mu_cen_pos_in = st.number_input(f"Mu Centro Vano (Positivo) [{unidad_mom}]", 0.0, 10000.0, st.session_state.get("vc_mu_cen", 90.0), 10.0, key="vc_mu_cen")
        Mu_izq_neg = Mu_izq_neg_in / factor_fuerza if unidades_salida == "Toneladas fuerza (tonf, tonf·m)" else Mu_izq_neg_in
        Mu_der_neg = Mu_der_neg_in / factor_fuerza if unidades_salida == "Toneladas fuerza (tonf, tonf·m)" else Mu_der_neg_in
        Mu_cen_pos = Mu_cen_pos_in / factor_fuerza if unidades_salida == "Toneladas fuerza (tonf, tonf·m)" else Mu_cen_pos_in

    #  BLOQUE 1: DIAGRAMA ENVOLVENTE DE MOMENTOS 
    st.markdown("---")
    st.subheader("Diagrama de Envolvente de Momentos")
    st.caption("Distribución esquemática de momentos últimos a lo largo del vano")

    fig_env, ax_env = plt.subplots(figsize=(9, 3.2))
    fig_env.patch.set_facecolor('#1e1e2e')
    ax_arr = fig_env.get_axes()
    for _ax in ax_arr: _ax.set_facecolor('#14142a'); _ax.tick_params(colors='#cdd6f4'); _ax.xaxis.label.set_color('#cdd6f4'); _ax.yaxis.label.set_color('#cdd6f4')
    fig_env.patch.set_facecolor('#1a1a2e')
    ax_env.set_facecolor('#1a1a2e')

    hscale = max(abs(Mu_izq_neg_in), abs(Mu_der_neg_in), abs(Mu_cen_pos_in))
    hscale = hscale if hscale > 0 else 1.0
    def norm_env(v): return abs(v) / hscale * 1.2

    # Zonas superiores negativas — ROJO
    ax_env.fill_between(
        [0, L_vc*0.25], [norm_env(Mu_izq_neg_in), 0],
        color='#e74c3c', alpha=0.88)
    ax_env.fill_between(
        [L_vc*0.75, L_vc], [0, norm_env(Mu_der_neg_in)],
        color='#e74c3c', alpha=0.88)
    ax_env.fill_between(
        [L_vc*0.25, L_vc*0.75],
        [0, 0], [-norm_env(Mu_cen_pos_in)*0.10, -norm_env(Mu_cen_pos_in)*0.10],
        color='#e74c3c', alpha=0.4)

    # Zonas inferiores positivas — AMARILLO
    ax_env.fill_between(
        [0, L_vc*0.25], [0, -norm_env(Mu_izq_neg_in)*0.45],
        color='#f1c40f', alpha=0.88)
    ax_env.fill_between(
        [L_vc*0.25, L_vc*0.50], [-norm_env(Mu_izq_neg_in)*0.45, -norm_env(Mu_cen_pos_in)],
        color='#f1c40f', alpha=0.88)
    ax_env.fill_between(
        [L_vc*0.50, L_vc*0.75], [-norm_env(Mu_cen_pos_in), -norm_env(Mu_der_neg_in)*0.45],
        color='#f1c40f', alpha=0.88)
    ax_env.fill_between(
        [L_vc*0.75, L_vc], [-norm_env(Mu_der_neg_in)*0.45, 0],
        color='#f1c40f', alpha=0.88)

    # Línea base y apoyos
    ax_env.axhline(0, color='white', linewidth=1.4, alpha=0.7)
    ax_env.plot(0,     0, 'w^', markersize=10, zorder=5)
    ax_env.plot(L_vc, 0, 'w^', markersize=10, zorder=5)

    # Etiquetas superiores (negativos)
    ax_env.text(L_vc*0.02,  norm_env(Mu_izq_neg_in)*1.06,
                f"Mu_top_izq\n{Mu_izq_neg_in:.2f} {unidad_mom}",
                color='white', fontsize=7.5, ha='left', va='bottom', fontweight='bold')
    ax_env.text(L_vc*0.50,  norm_env(Mu_cen_pos_in)*0.12,
                "Mu_top_cen",
                color='#e07070', fontsize=7, ha='center', va='bottom', style='italic')
    ax_env.text(L_vc*0.98,  norm_env(Mu_der_neg_in)*1.06,
                f"Mu_top_der\n{Mu_der_neg_in:.2f} {unidad_mom}",
                color='white', fontsize=7.5, ha='right', va='bottom', fontweight='bold')

    # Etiquetas inferiores (positivos)
    ax_env.text(L_vc*0.02,  -norm_env(Mu_izq_neg_in)*0.44,
                "Mu_bot_izq",
                color='#f1c40f', fontsize=7, ha='left', va='top', style='italic')
    ax_env.text(L_vc*0.50,  -norm_env(Mu_cen_pos_in)*1.08,
                f"Mu_bot_cen\n{Mu_cen_pos_in:.2f} {unidad_mom}",
                color='white', fontsize=7.5, ha='center', va='top', fontweight='bold')
    ax_env.text(L_vc*0.98,  -norm_env(Mu_der_neg_in)*0.44,
                "Mu_bot_der",
                color='#f1c40f', fontsize=7, ha='right', va='top', style='italic')

    ax_env.set_xlim(-0.08, L_vc*1.08)
    ax_env.set_xlabel("Posición a lo largo del vano (m)", color='white', fontsize=8)
    ax_env.tick_params(colors='white')
    ax_env.set_yticks([])
    for spine in ax_env.spines.values():
        spine.set_visible(False)
    from matplotlib.patches import Patch
    ax_env.legend(
        handles=[Patch(color='#e74c3c', label='M negativo (superior)'),
                 Patch(color='#f1c40f', label='M positivo (inferior)')],
        loc='lower right', fontsize=7.5, facecolor='#1a1a2e',
        labelcolor='white', framealpha=0.6)
    fig_env.tight_layout()
    st.pyplot(fig_env)
    plt.close(fig_env)
    st.caption("Fig. N°01 — Diagrama esquemático de momentos flectores últimos.")
    st.markdown("---")

    st.subheader("2. Armadura Propuesta")
    c4, c5, c6 = st.columns(3)
    varillas_vc = list(rebar_dict.keys())
    _def_vc = "#5 (Ø15.9mm)" if "Pulgadas" in bar_sys else "16mm"
    _def_idx_vc = varillas_vc.index(_def_vc) if _def_vc in varillas_vc else 2
    
    with c4:
        st.markdown("**Acero Nudo Izquierdo**")
        nb_izq_sup = st.number_input("Cant. Barras Superiores Izq", 2, 20, 3, key="nb_izq_sup")
        bar_izq_sup = st.selectbox("Ø Barras Sup. Izq", varillas_vc, index=_def_idx_vc, key="bar_izq_sup")
        nb_izq_inf = st.number_input("Cant. Barras Inferiores Izq", 2, 20, 2, key="nb_izq_inf")
        bar_izq_inf = st.selectbox("Ø Barras Inf. Izq", varillas_vc, index=_def_idx_vc, key="bar_izq_inf")
    with c5:
        st.markdown("**Acero Nudo Derecho**")
        nb_der_sup = st.number_input("Cant. Barras Superiores Der", 2, 20, 3, key="nb_der_sup")
        bar_der_sup = st.selectbox("Ø Barras Sup. Der", varillas_vc, index=_def_idx_vc, key="bar_der_sup")
        nb_der_inf = st.number_input("Cant. Barras Inferiores Der", 2, 20, 2, key="nb_der_inf")
        bar_der_inf = st.selectbox("Ø Barras Inf. Der", varillas_vc, index=_def_idx_vc, key="bar_der_inf")
    with c6:
        st.markdown("**Acero Centro**")
        nb_cen_inf = st.number_input("Cant. Barras Inferiores Cen", 2, 20, 3, key="nb_cen_inf")
        bar_cen_inf = st.selectbox("Ø Barras Inf. Cen", varillas_vc, index=_def_idx_vc, key="bar_cen_inf")
        st.markdown("**Estribos**")
        est_opts_vc = ["Ø8mm","Ø10mm","Ø12mm","#3","#4"]
        st_bar_vc = st.selectbox("Ø Estribo", est_opts_vc, index=3 if "Pulgadas"in bar_sys else 1, key="st_bar_vc")
        n_ramas_vc = st.number_input("Ramas de estribo", 2, 6, 2, 1, key="n_ramas_vc")

    # --- C?LCULOS INTERNOS ---
    d_vc = h_vc - dp_vc
    d_vc_mm = d_vc * 10
    b_vc_mm = b_vc * 10

    # Funciones de utilidad para cálculo robusto
    def calc_As(n, bar_str): return n * rebar_dict[bar_str]
    def calc_phi_Mn(As_prov, b_mm, d_mm):
        a_mm = As_prov * 100 * fy / (0.85 * fc * b_mm)
        return phi_f * As_prov * 100 * fy * (d_mm - a_mm / 2) / 1e6
    def calc_Mpr(As_prov, b_mm, d_mm):
        # Mpr usa 1.25fy y phi=1.0
        a_mm_pr = As_prov * 100 * (1.25 * fy) / (0.85 * fc * b_mm)
        return 1.0 * As_prov * 100 * (1.25 * fy) * (d_mm - a_mm_pr / 2) / 1e6

    # 1. As provistos
    As_izq_sup = calc_As(nb_izq_sup, bar_izq_sup)
    As_izq_inf = calc_As(nb_izq_inf, bar_izq_inf)
    As_der_sup = calc_As(nb_der_sup, bar_der_sup)
    As_der_inf = calc_As(nb_der_inf, bar_der_inf)
    As_cen_inf = calc_As(nb_cen_inf, bar_cen_inf)
    
    # 2. Resistencias nominales reducidas (φMn)
    phiMn_izq_neg = calc_phi_Mn(As_izq_sup, b_vc_mm, d_vc_mm)
    phiMn_der_neg = calc_phi_Mn(As_der_sup, b_vc_mm, d_vc_mm)
    phiMn_cen_pos = calc_phi_Mn(As_cen_inf, b_vc_mm, d_vc_mm)

    # 3. Momentos Probables de los Nudos (Mpr) para plastificación
    Mpr_izq = calc_Mpr(As_izq_sup, b_vc_mm, d_vc_mm)
    Mpr_der = calc_Mpr(As_der_sup, b_vc_mm, d_vc_mm)

    # 4. Cálculo Cortante Probable Ve (Isostático gravitacional + Hiperestático plástico)
    # Ve = Wu*L/2 +- (Mpr_izq + Mpr_der)/L
    # Por normatividad sísmica, consideramos el caso más crítico.
    V_grav_isostatico = Wu_vc * L_vc / 2
    V_plastico = ((Mpr_izq + Mpr_der) * factor_fuerza) / L_vc
    Ve_cortante_diseno = max(abs(V_plastico + V_grav_isostatico), abs(V_plastico - V_grav_isostatico)) # P-EC-4: Considera sismo en ambas direcciones

    # 5. Geometría
    chk_b_min = b_vc >= 25.0
    chk_bh_ratio = (b_vc / h_vc) >= 0.3
    chk_b_max = b_vc <= (bcol_vc + 3 * h_vc)
    geo_ok = chk_b_min and chk_bh_ratio and chk_b_max

    # 6. Reglas Longitudinales de Nudos (DMO/DES)
    # a. As_inf >= 50% As_sup en cualquiera de los nudos (C.21.5.2.1)
    # b. As_cualquiera >= 25% As_max_sup en toda la luz
    As_max_nudos = max(As_izq_sup, As_der_sup)
    chk_izq_inf_50 = As_izq_inf >= 0.50 * As_izq_sup
    chk_der_inf_50 = As_der_inf >= 0.50 * As_der_sup
    chk_25_pct = min(As_izq_sup, As_izq_inf, As_der_sup, As_der_inf, As_cen_inf) >= 0.25 * As_max_nudos
    chk_2_continuas = nb_izq_inf >= 2 and nb_der_inf >= 2 and nb_cen_inf >= 2 # Continuidad de barras inferiores
    
    # Cuantía Máxima DES 0.025
    rho_max_des = 0.025
    rho_prov_max = max(As_izq_sup, As_izq_inf, As_der_sup, As_der_inf, As_cen_inf) / (b_vc * d_vc)
    chk_rho_des = rho_prov_max <= rho_max_des if es_des else True

    # 7. Reglas al Cortante y Confinamiento
    st_area_vc = {"Ø8mm":0.503,"Ø10mm":0.785,"Ø12mm":1.131,"#3":0.71,"#4":1.29}[st_bar_vc]
    db_est_vc = {"Ø8mm":8,"Ø10mm":10,"Ø12mm":12,"#3":9.53,"#4":12.70}[st_bar_vc]
    Av_vc_total = st_area_vc * n_ramas_vc

    # Vc=0 rule (C.21.5.4.2)
    # Vc=0 ssi V_plastico >= 0.5 * Ve
    aplica_Vc_cero = es_sismico and (V_plastico >= 0.5 * Ve_cortante_diseno)
    st.session_state["v_aplica_vccero_sismo"] = aplica_Vc_cero
    Vc_vc = 0.17 * lam * math.sqrt(fc) * b_vc_mm * d_vc_mm / 1000 if not aplica_Vc_cero else 0.0
    phiVc_vc = phi_v * Vc_vc

    Vs_req_vc = max(0, Ve_cortante_diseno / phi_v - Vc_vc)
    if Vs_req_vc > 0:
        s_calc_vc = Av_vc_total * 100 * fy * d_vc_mm / (Vs_req_vc * 1000)
    else:
        s_calc_vc = 9999.0

    # S_max confinamiento (C.21.5.3.2)
    db_long_min = min(diam_dict[bar_izq_sup], diam_dict[bar_izq_inf], diam_dict[bar_der_sup], diam_dict[bar_der_inf])
    if es_des:
        s_max_conf = min(d_vc_mm / 4, 6 * db_long_min, 150)
    elif es_dmo:
        s_max_conf = min(d_vc_mm / 4, 8 * db_long_min, 24 * db_est_vc, 300)
    else:
        s_max_conf = min(d_vc_mm / 2, 600)

    s_diseno_conf = min(s_calc_vc, s_max_conf)
    s_diseno_conf_cm = math.floor(s_diseno_conf / 10)
    if s_diseno_conf_cm < 5: s_diseno_conf_cm = 5

    st.markdown("---")
    st.subheader("Reporte de Verificaciones (NSR-10 / ACI 318)")
    
    t1, t2, t3 = st.tabs(["Chequeos Sísmicos", "Momentos y Flexión", " Cortante Plástico ($V_p$)"])
    with t1:
        st.markdown("**(A) Geometría**")
        st.write(f"- bw ≥ 25 cm: {'CUMPLE' if chk_b_min else f'NO CUMPLE ({b_vc} cm)'}")
        st.write(f"- bw / h ≥ 0.3: {'CUMPLE' if chk_bh_ratio else f'NO CUMPLE ({b_vc/h_vc:.2f})'}")
        st.write(f"- bw ≤ b_col + 3h: {'CUMPLE' if chk_b_max else f'NO CUMPLE ({bcol_vc+3*h_vc} cm)'}")
        
        st.markdown(f"**(B) Requisitos Cuanía Longitudinal ({'DMO/DES' if es_sismico else 'General'})**")
        if es_sismico:
            st.write(f"- As_inf ≥ 0.50 As_sup (Nudo Izquierdo): {' CUMPLE' if chk_izq_inf_50 else f'NO CUMPLE ({As_izq_inf:.2f} < {0.5*As_izq_sup:.2f})'}")
            st.write(f"- As_inf ≥ 0.50 As_sup (Nudo Derecho): {' CUMPLE' if chk_der_inf_50 else f'NO CUMPLE ({As_der_inf:.2f} < {0.5*As_der_sup:.2f})'}")
            st.write(f"- Min 2 barras continuas cara a cara: {'CUMPLE' if chk_2_continuas else 'NO CUMPLE'}")
            st.write(f"- ρ_max ≤ 0.025 (DES): {' CUMPLE' if chk_rho_des else 'NO CUMPLE (EXCEDE)'}")
        else:
            st.info("No aplican reglas estrictas C.21.5 de pórticos.")

    with t2:
        rows_flex = [
            ("Nudo Izquierdo (-)", f"φMn = {phiMn_izq_neg*factor_fuerza:.2f}", f"Mu = {Mu_izq_neg*factor_fuerza:.2f}", " OK" if phiMn_izq_neg>=Mu_izq_neg else "? FALLA"),
            ("Centro Vano (+)",   f"φMn = {phiMn_cen_pos*factor_fuerza:.2f}", f"Mu = {Mu_cen_pos*factor_fuerza:.2f}", " OK" if phiMn_cen_pos>=Mu_cen_pos else "? FALLA"),
            ("Nudo Derecho (-)",  f"φMn = {phiMn_der_neg*factor_fuerza:.2f}", f"Mu = {Mu_der_neg*factor_fuerza:.2f}", " OK" if phiMn_der_neg>=Mu_der_neg else "? FALLA"),
        ]
        st.dataframe(pd.DataFrame(rows_flex, columns=["Zona", f"Capacidad [{unidad_mom}]", f"Demanda [{unidad_mom}]", "Estado"]), use_container_width=True, hide_index=True)

        #  Sugerencias de solución para zonas con FALLA 
        _fallos = []
        _casos = [
            ("Nudo Izquierdo (-)", phiMn_izq_neg, Mu_izq_neg, "los estribos superiores izquierdos (Moi- o M_izq-)"),
            ("Centro Vano (+)",   phiMn_cen_pos, Mu_cen_pos, "el refuerzo inferior en el vano central"),
            ("Nudo Derecho (-)",  phiMn_der_neg, Mu_der_neg, "los estribos superiores derechos (Mod- o M_der-)"),
        ]
        for zona, phi_mn, mu_d, desc_zona in _casos:
            if phi_mn < mu_d:
                _fallos.append((zona, phi_mn, mu_d, desc_zona))

        if _fallos:
            st.markdown("---")
            st.markdown("###  ? Sugerencias para corregir las FALLAs")
            for zona, phi_mn, mu_d, desc_zona in _fallos:
                _deficit = (mu_d - phi_mn) * factor_fuerza
                _ratio   = mu_d / phi_mn if phi_mn > 0 else float('inf')
                # Estimación del As adicional necesario (aprox.)
                _d_mm     = (h_vc - dp_vc) * 10          # mm
                _a_actual = (As_izq_sup if "Izq" in zona else (As_der_sup if "Der" in zona else As_cen_inf)) * 100 * fy / (0.85 * fc * b_vc * 10)
                _As_req_aprox = mu_d * 1e6 / (phi_f * fy * (_d_mm - _a_actual / 2) * 100)  # cm²
                _As_actual    = (As_izq_sup if "Izq" in zona else (As_der_sup if "Der" in zona else As_cen_inf))

                with st.expander(f"**{zona}** — φMn = {phi_mn*factor_fuerza:.2f} < Mu = {mu_d*factor_fuerza:.2f} {unidad_mom} (déficit: {_deficit:.2f})", expanded=True):
                    st.error(
                        f"**La capacidad es insuficiente en {zona}.**  \n"
                        f"Necesitas **{_ratio:.2f}× más momento resistente** en {desc_zona}."
                    )
                    st.markdown("**Opciones para resolver (de menor a mayor impacto):**")
                    col_s1, col_s2 = st.columns(2)
                    with col_s1:
                        st.markdown(
                            f"**1?⃣ Aumentar el número de varillas**  \n"
                            f"As actual ≈ {_As_actual:.2f} cm²  \n"
                            f"As requerido ≈ **{max(_As_req_aprox, _As_actual*_ratio):.2f} cm²**  \n"
                            f"→ Agrega varillas del mismo diámetro o usa barras de mayor calibre en la selección de arriba."
                        )
                    with col_s2:
                        st.markdown(
                            f"**2?⃣ Aumentar el peralte h**  \n"
                            f"h actual = {h_vc:.0f} cm  \n"
                            f"→ Incrementar h ≈ **{h_vc * math.sqrt(_ratio):.0f} cm** reduciría el Rn.  \n\n"
                            f"**3?⃣ Reducir el momento de demanda Mu**  \n"
                            f"→ Verificar la combinación de cargas o el modelo estructural."
                        )


    with t3:
        #  BLOQUE 2: BARRA RESUMEN TIPO ETABS 
        st.markdown("#### ? Distribución de Acero — Resumen tipo ETABS")
        st.caption("?rea de acero requerida por zona del vano (superior e inferior)")

        def as_req_zona(Mu_kNm, b_cm, d_cm, fc_mpa, fy_mpa, phi=0.9):
            if Mu_kNm <= 0: return 0.0
            b_mm_req = b_cm * 10
            d_mm_req = d_cm * 10
            Mu_Nmm = abs(Mu_kNm) * 1e6
            Rn = Mu_Nmm / (phi * b_mm_req * d_mm_req**2)
            disc = 1 - (2 * Rn) / (0.85 * fc_mpa)
            if disc < 0: return None
            rho = (0.85 * fc_mpa / fy_mpa) * (1 - math.sqrt(disc))
            As = rho * b_cm * d_cm
            As_min = max(0.25 * math.sqrt(fc_mpa) / fy_mpa, 1.4 / fy_mpa) * b_cm * d_cm
            return max(As, As_min)

        d_vc_cm = h_vc - dp_vc
        b_vc_cm = b_vc

        As_top_izq = as_req_zona(Mu_izq_neg,  b_vc_cm, d_vc_cm, fc, fy) or 0.0
        As_top_cen = as_req_zona(abs(Mu_cen_pos)*0.15, b_vc_cm, d_vc_cm, fc, fy) or 0.0
        As_top_der = as_req_zona(Mu_der_neg,  b_vc_cm, d_vc_cm, fc, fy) or 0.0
        As_bot_izq = as_req_zona(Mu_izq_neg * 0.5, b_vc_cm, d_vc_cm, fc, fy) or 0.0
        As_bot_cen = as_req_zona(Mu_cen_pos,  b_vc_cm, d_vc_cm, fc, fy) or 0.0
        As_bot_der = as_req_zona(Mu_der_neg * 0.5, b_vc_cm, d_vc_cm, fc, fy) or 0.0

        As_top_izq_p = As_izq_sup
        As_top_der_p = As_der_sup
        As_bot_izq_p = As_izq_inf
        As_bot_cen_p = As_cen_inf
        As_bot_der_p = As_der_inf

        def color_cell(req, prov):
            if prov >= req: return f"background:#1a6b2a; color:white; padding:6px 10px; border-radius:6px; font-weight:bold;"
            else: return f"background:#7b1a1a; color:white; padding:6px 10px; border-radius:6px; font-weight:bold;"

        html_barra = f"""
        <div style="font-family:monospace; margin:10px 0;">
          <!-- FILA SUPERIOR -->
          <div style="display:flex; align-items:center; margin-bottom:4px;">
            <span style="width:70px; color:#aaa; font-size:11px;">SUPERIOR</span>
            <div style="flex:1; display:flex; align-items:center; gap:0;">
              <div style="text-align:center; flex:1;">
                <span style="{color_cell(As_top_izq, As_top_izq_p)}">{As_top_izq_p:.2f} cm²</span>
                <div style="color:#aaa; font-size:10px;">req: {As_top_izq:.2f}</div>
              </div>
              <div style="flex:2; height:4px; background:#e74c3c; margin:0 4px;"></div>
              <div style="text-align:center; flex:1;">
                <span style="{color_cell(As_top_cen, As_top_izq_p*0.25)}">{As_top_izq_p*0.25:.2f} cm²</span>
                <div style="color:#aaa; font-size:10px;">req: {As_top_cen:.2f}</div>
              </div>
              <div style="flex:2; height:4px; background:#e74c3c; margin:0 4px;"></div>
              <div style="text-align:center; flex:1;">
                <span style="{color_cell(As_top_der, As_top_der_p)}">{As_top_der_p:.2f} cm²</span>
                <div style="color:#aaa; font-size:10px;">req: {As_top_der:.2f}</div>
              </div>
            </div>
          </div>
          <!-- BARRA CENTRAL -->
          <div style="display:flex; align-items:center; margin: 2px 0;">
            <span style="width:70px;"></span>
            <div style="flex:1; height:6px; background:linear-gradient(90deg,#555,#888,#555); border-radius:3px;"></div>
          </div>
          <!-- FILA INFERIOR -->
          <div style="display:flex; align-items:center; margin-top:4px;">
            <span style="width:70px; color:#aaa; font-size:11px;">INFERIOR</span>
            <div style="flex:1; display:flex; align-items:center; gap:0;">
              <div style="text-align:center; flex:1;">
                <span style="{color_cell(As_bot_izq, As_bot_izq_p)}">{As_bot_izq_p:.2f} cm²</span>
                <div style="color:#aaa; font-size:10px;">req: {As_bot_izq:.2f}</div>
              </div>
              <div style="flex:2; height:4px; background:#f1c40f; margin:0 4px;"></div>
              <div style="text-align:center; flex:1;">
                <span style="{color_cell(As_bot_cen, As_bot_cen_p)}">{As_bot_cen_p:.2f} cm²</span>
                <div style="color:#aaa; font-size:10px;">req: {As_bot_cen:.2f}</div>
              </div>
              <div style="flex:2; height:4px; background:#f1c40f; margin:0 4px;"></div>
              <div style="text-align:center; flex:1;">
                <span style="{color_cell(As_bot_der, As_bot_der_p)}">{As_bot_der_p:.2f} cm²</span>
                <div style="color:#aaa; font-size:10px;">req: {As_bot_der:.2f}</div>
              </div>
            </div>
          </div>
          <div style="margin-top:8px; color:#888; font-size:10px;">
              Verde = As provisto ≥ As requerido &nbsp;|&nbsp;   Rojo = deficiente, revisar varillas
          </div>
        </div>
        """
        st.markdown(html_barra, unsafe_allow_html=True)
        st.markdown("---")

        #  BLOQUE 3: CONDICIÓN DOBLE ARMADURA POR ZONA 
        st.markdown("####  ? Verificación de Doble Armadura por Zona")
        st.caption("Chequeo si el bloque de compresiones 'a'supera el límite simplemente reforzado")

        beta1_vc = 0.85 if fc <= 28 else max(0.65, 0.85 - 0.05*(fc - 28)/7)
        eps_cu   = 0.003
        eps_min  = (fy / (200000 if 'ksi' not in str(unidad_mom) else 29000)) + 0.003
        c_max    = (eps_cu / (eps_cu + eps_min)) * d_vc_cm  # cm
        a_max    = beta1_vc * c_max  # cm

        def chk_doble_arm(Mu_kNm, etiqueta):
            if Mu_kNm <= 0: return etiqueta, "—", "—", " No aplica"
            b_mm = b_vc_cm * 10
            d_mm = d_vc_cm * 10
            Mu_Nmm = abs(Mu_kNm) * 1e6
            disc = d_mm**2 - (2 * Mu_Nmm) / (0.85 * fc * 0.9 * b_mm)
            if disc < 0: return etiqueta, "∞", f"{a_max:.2f}", " REQUIERE A's compresión"
            a_calc = (d_mm - math.sqrt(disc)) / 10  # cm
            if a_calc > a_max: return etiqueta, f"{a_calc:.2f}", f"{a_max:.2f}", " REQUIERE A's compresión"
            else: return etiqueta, f"{a_calc:.2f}", f"{a_max:.2f}", " Simplemente reforzada"

        filas_doble = [
            chk_doble_arm(Mu_izq_neg,  "Sup. Izquierda"),
            chk_doble_arm(Mu_der_neg,  "Sup. Derecha"),
            chk_doble_arm(Mu_cen_pos,  "Inf. Centro"),
            chk_doble_arm(Mu_izq_neg*0.5, "Inf. Izquierda"),
            chk_doble_arm(Mu_der_neg*0.5, "Inf. Derecha"),
        ]

        filas_html = [f'''<tr style="background:{'#1a3a1a' if '' in f[3] else '#3a1a1a'}; color:white; border-bottom:1px solid #444;">
  <td style="padding:7px;">{f[0]}</td>
  <td style="padding:7px; text-align:center;">{f[1]}</td>
  <td style="padding:7px; text-align:center;">{f[2]}</td>
  <td style="padding:7px; text-align:center;">{f[3]}</td>
</tr>''' for f in filas_doble]
        tabla_str = "\n".join(filas_html)

        st.markdown(f"""<table style="width:100%; border-collapse:collapse; font-size:13px; font-family:monospace;">
  <thead>
    <tr style="background:#2c2c4e; color:white;">
      <th style="padding:8px; text-align:left;">Zona</th>
      <th style="padding:8px; text-align:center;">a calculado (cm)</th>
      <th style="padding:8px; text-align:center;">a_máx (cm)</th>
      <th style="padding:8px; text-align:center;">Condición</th>
    </tr>
  </thead>
  <tbody>
{tabla_str}
  </tbody>
</table>""", unsafe_allow_html=True)
        st.caption(f"β1 = {beta1_vc:.3f} | c_máx = {c_max:.2f} cm | a_máx = {a_max:.2f} cm — "
                   f"NSR-10 C.10.3.3: La deformación neta en tracción εₜ ≥ εₘᵢₙ = {eps_min:.4f} "
                   f"garantiza falla dúctil. Si a > a_máx se requiere acero en compresión A's.")

        hay_doble = any("REQUIERE" in f[3] for f in filas_doble)
        if hay_doble:
            st.error("Una o más zonas requieren acero en compresión (doble armadura). Agregue barras superiores adicionales en esas zonas o aumente la sección.")
        else:
            st.success("Todas las zonas son simplemente reforzadas. No se requiere A's en compresión.")

        st.markdown("---")
        st.markdown("####  Verificación de Cortante Plástico")
        st.markdown(f"- $M_{{pr,izq}}$ (1.25fy, $\\phi=1$) = **{Mpr_izq*factor_fuerza:.2f}** {unidad_mom}")
        st.markdown(f"- $M_{{pr,der}}$ (1.25fy, $\\phi=1$) = **{Mpr_der*factor_fuerza:.2f}** {unidad_mom}")
        st.markdown(f"- Cortante Isostático ($W_u L / 2$) = **{V_grav_isostatico*factor_fuerza:.2f}** {unidad_fuerza}")
        st.markdown(f"- Cortante Plástico $(M_{{pr,izq}} + M_{{pr,der}})/L_n$ = **{V_plastico*factor_fuerza:.2f}** {unidad_fuerza}")
        st.markdown(f"- **Cortante de Diseño $V_e$** = **{Ve_cortante_diseno*factor_fuerza:.2f}** {unidad_fuerza}")
        if aplica_Vc_cero:
            st.error("**ALERTA S?SMICA:** El cortante inducido por el sismo supera el 50% del cortante total de diseño $V_e$. La normativa exige penalizar al concreto asumiendo **Vc = 0** en la zona de confinamiento de la rótula plástica.")
        else:
            st.success(f"Vc = {Vc_vc*factor_fuerza:.2f} {unidad_fuerza}. No se exige Vc=0.")
            
        st.markdown(f"**Diseño de Estribos en Zona Confinamiento (2h = {2*h_vc:.0f} cm)**")
        st.write(f"Separación máxima por norma ($s_{{max}}$) = {s_max_conf/10:.1f} cm")
        st.write(f"Separación requerida por $V_e$ = {s_calc_vc/10:.1f} cm")
        st.info(f"Se usarán **Estribos {st_bar_vc} ({n_ramas_vc} ramas) @ {s_diseno_conf_cm:.0f} cm** en zona confinada.")

    # --- RESULTADOS ADICIONALES (3D, APU, Memoria) ---
    st.markdown("---")
    res_t1, res_t2, res_t3 = st.tabs(["Visualización 3D", "Cantidades y APU", " Memoria de Cálculo"])
    
    with res_t1:
        fig3d = go.Figure()
        L_mm_3d = L_vc * 1000
        b_mm_3d = b_vc * 10
        h_mm_3d = h_vc * 10
        dp_mm = dp_vc * 10

        # Concreto (sólido semitransparente)
        x_c = [0, b_mm_3d, b_mm_3d, 0, 0, b_mm_3d, b_mm_3d, 0]
        y_c = [0, 0, h_mm_3d, h_mm_3d, 0, 0, h_mm_3d, h_mm_3d]
        z_c = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
        fig3d.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.12, color='gray', name='Concreto'))

        #  Barras longitudinales superiores 
        xs_sup = [dp_mm + i * (b_mm_3d - 2*dp_mm) / max(nb_izq_sup - 1, 1) for i in range(nb_izq_sup)]
        for idx, xb in enumerate(xs_sup):
            fig3d.add_trace(go.Scatter3d(
                x=[xb, xb], y=[h_mm_3d - dp_mm, h_mm_3d - dp_mm], z=[0, L_mm_3d],
                mode='lines', line=dict(color='orange', width=5),
                name=f'Acero Sup ({bar_izq_sup})', showlegend=(idx == 0)))

        #  Barras longitudinales inferiores 
        xs_inf = [dp_mm + i * (b_mm_3d - 2*dp_mm) / max(nb_cen_inf - 1, 1) for i in range(nb_cen_inf)]
        for idx, xb in enumerate(xs_inf):
            fig3d.add_trace(go.Scatter3d(
                x=[xb, xb], y=[dp_mm, dp_mm], z=[0, L_mm_3d],
                mode='lines', line=dict(color='red', width=5),
                name=f'Acero Inf ({bar_cen_inf})', showlegend=(idx == 0)))

        #  Flejes (estribos) 
        # Posición interior del estribo (cara interior del recubrimiento)
        cov_est = dp_mm - db_est_vc / 2  # distancia al borde exterior del estribo
        xe1, xe2 = cov_est, b_mm_3d - cov_est
        ye1, ye2 = cov_est, h_mm_3d - cov_est
        s_fleje_mm = s_diseno_conf_cm * 10  # espaciado en mm

        tx_st, ty_st, tz_st = [], [], []
        z_pos = s_fleje_mm
        first_st = True
        while z_pos <= L_mm_3d - s_fleje_mm / 2:
            # Rectángulo cerrado del estribo: 5 puntos (cerrando el loop) + None separador
            tx_st += [xe1, xe2, xe2, xe1, xe1, None]
            ty_st += [ye1, ye1, ye2, ye2, ye1, None]
            tz_st += [z_pos]*5 + [None]
            z_pos += s_fleje_mm

        if tx_st:
            fig3d.add_trace(go.Scatter3d(
                x=tx_st, y=ty_st, z=tz_st,
                mode='lines',
                line=dict(color='cyan', width=3),
                name=f'Flejes {st_bar_vc} @ {s_diseno_conf_cm:.0f}cm',
                showlegend=True))

        fig3d.update_layout(
            scene=dict(aspectmode='data',
                       xaxis_title='b (mm)', yaxis_title='h (mm)', zaxis_title='L (mm)'),
            margin=dict(l=0, r=0, t=0, b=0), height=450)
        st.plotly_chart(fig3d, use_container_width=True)

    with res_t2:
        total_estribos = math.ceil((L_vc * 100) / s_diseno_conf_cm) + 1
        _perim_est_apu = 2*(b_vc - 2*dp_vc) + 2*(h_vc - 2*dp_vc) + 6*db_est_vc/10
        _peso_est_apu = total_estribos * (_perim_est_apu/100) * st_area_vc * 0.785
        _peso_sup_apu = nb_izq_sup * (L_vc + 2*(bcol_vc/100)) * rebar_dict[bar_izq_sup]*100 * 7.85e-3
        _peso_inf_apu = nb_cen_inf * (L_vc + 2*(bcol_vc/100)) * rebar_dict[bar_cen_inf]*100 * 7.85e-3
        _peso_total_sism = _peso_est_apu + _peso_sup_apu + _peso_inf_apu
        _vol_sism = (b_vc/100) * (h_vc/100) * L_vc
        
        render_apu_breakdown(_vol_sism, _peso_total_sism, fc, f"({nb_izq_sup} Sup, {nb_cen_inf} Inf, {total_estribos} Estribos)")

    with res_t3:
        st.info("La memoria de cálculo consolida las verificaciones estáticas, dinámicas y sísmicas del marco estructural.")

        if st.button("Generar Memoria Exhaustiva (DOCX)"):

            from datetime import datetime

            import io

            import os

            import tempfile

            from docx.shared import Inches, Pt

            from docx.enum.text import WD_ALIGN_PARAGRAPH


            doc = Document()


            # --- TITULO PRINCIPAL ---

            title = doc.add_heading(f"Memoria de Diseño Estructural — Viga {b_vc:.0f}x{h_vc:.0f} cm", 0)

            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Fecha de Generación: {datetime.now().strftime('%d/%m/%Y %H:%M')} | Norma: {norma_sel}").alignment = WD_ALIGN_PARAGRAPH.CENTER


            # --- 1. GEOMETRÍA Y REQUISITOS ---

            doc.add_heading("1. Geometría y Verificaciones de Sección", level=1)

            p1 = doc.add_paragraph()

            p1.add_run(f"La sección analizada corresponde a una viga de sección rectangular predimensionada con Base (bw) = {b_vc} cm, Peralte (h) = {h_vc} cm, y Luz Libre (L) = {L_vc} m. ")

            p1.add_run(f"La resistencia a la compresión del concreto especificada es f'c = {fc} MPa y la fluencia del acero fy = {fy} MPa.\n\n")


            p1.add_run("Chequeos Normativos Dimensionales:\n").bold = True

            p1.add_run(f"• Ancho mínimo bw >= 25 cm -> {'CUMPLE' if chk_b_min else 'NO CUMPLE'} ({b_vc} cm)\n")

            p1.add_run(f"• Relación base/altura bw/h >= 0.3 -> {'CUMPLE' if chk_bh_ratio else 'NO CUMPLE'} ({b_vc/h_vc:.2f})\n")

            p1.add_run(f"• Ancho máximo apoyos bw <= b_col + 3h -> {'CUMPLE' if chk_b_max else 'NO CUMPLE'}")


            # --- 2. CAPACIDAD A FLEXIÓN ---

            doc.add_heading("2. Resistencia a Flexión Integrada", level=1)

            doc.add_paragraph("A continuación se presenta la verificación de capacidades resistentes teóricas (φMn) frente a los momentos flectores máximos (Mu) para las zonas críticas.")


            t_flex = doc.add_table(rows=1, cols=4)

            t_flex.style = 'Light Grid'

            hdr_cells = t_flex.rows[0].cells

            hdr_cells[0].text = 'Zona de Viga'

            hdr_cells[1].text = f'Capacidad φMn [{unidad_mom}]'

            hdr_cells[2].text = f'Demanda Mu [{unidad_mom}]'

            hdr_cells[3].text = 'Estado de Diseño'


            records_flex = [
                ("Nudo Izquierdo (-)", phiMn_izq_neg*factor_fuerza, Mu_izq_neg*factor_fuerza),
                ("Centro Vano (+)", phiMn_cen_pos*factor_fuerza, Mu_cen_pos*factor_fuerza),
                ("Nudo Derecho (-)", phiMn_der_neg*factor_fuerza, Mu_der_neg*factor_fuerza)
            ]

            for zone, cap, dem in records_flex:
                row = t_flex.add_row().cells
                row[0].text = zone
                row[1].text = f"{cap:.2f}"
                row[2].text = f"{dem:.2f}"
                row[3].text = "CUMPLE" if cap >= dem else "INSUFICIENTE"

            # --- 3. REQUISITOS SÍSMICOS LONGITUDINALES ---
            doc.add_heading("3. Integridad Sísmica Longitudinal (DMO/DES)", level=1)
            p_sism = doc.add_paragraph()
            p_sism.add_run(f"• As Nudo Izquierdo: Inferior {As_izq_inf:.2f} cm² >= 50% Superior ({0.5*As_izq_sup:.2f} cm²) -> {'CUMPLE' if chk_izq_inf_50 else 'NO CUMPLE'}\\n")
            p_sism.add_run(f"• As Nudo Derecho: Inferior {As_der_inf:.2f} cm² >= 50% Superior ({0.5*As_der_sup:.2f} cm²) -> {'CUMPLE' if chk_der_inf_50 else 'NO CUMPLE'}\\n")

            p_sism.add_run(f"• Cuantía máxima en tracción (ρ <= 0.025) -> {'CUMPLE' if chk_rho_des else 'NO CUMPLE'}\n")

            p_sism.add_run(f"• Barras continuas a lo largo de la luz -> {'CUMPLE' if chk_2_continuas else 'NO CUMPLE'}")


            # --- 4. CORTANTE PLÁSTICO Y CONFINAMIENTO ---

            doc.add_heading("4. Diseño a Cortante Plástico", level=1)

            doc.add_paragraph(f"Se calculan los momentos probables Mpr asumiendo una fluencia de 1.25fy en las barras tractivas y φ=1.0 para derivar el cortante sísmico hiperestático.")


            t_shear = doc.add_table(rows=3, cols=2)

            t_shear.style = 'Light Grid'

            t_shear.rows[0].cells[0].text = "Momento Probable Izquierdo Mpr (-)"

            t_shear.rows[0].cells[1].text = f"{Mpr_izq*factor_fuerza:.2f} {unidad_mom}"

            t_shear.rows[1].cells[0].text = "Momento Probable Derecho Mpr (-)"

            t_shear.rows[1].cells[1].text = f"{Mpr_der*factor_fuerza:.2f} {unidad_mom}"

            t_shear.rows[2].cells[0].text = "Cortante Mínimo de Diseño (Ve o Vp)"

            t_shear.rows[2].cells[1].text = f"{Ve_cortante_diseno*factor_fuerza:.2f} {unidad_fuerza}"


            p_conf = doc.add_paragraph(f"\nLongitud de Confinamiento Hacia Ambos Nudos L_o = {zona_conf:.2f} m.\n")

            p_conf.add_run(f"Disposición Transversal: Estribos {st_bar_vc} ({n_ramas_vc} Ramas) @ {math.floor(s_diseno_conf_cm)} cm.").bold = True


            # --- 5. RENDERIZADO 3D ---

            doc.add_heading("5. Modelo 3D del Ensamblaje Estructural", level=1)

            doc.add_paragraph("Proyección de la jaula de armaduras (Armado Longitudinal y Estribado Transversal) calculada y cuantificada.")


            try:
                import kaleido
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                    fig3d_img_path = tmp_img.name
                fig3d.write_image(fig3d_img_path, format="png", width=800, height=500, scale=1.5)
                doc.add_picture(fig3d_img_path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                os.remove(fig3d_img_path)
            except Exception as e:
                doc.add_paragraph(f"[Nota] No se pudo adjuntar el renderizado volumétrico Plotly debido a ausencia del motor headless kaleido o error de render: {e}")


            # --- EXPORT ---

            doc_mem = io.BytesIO()

            doc.save(doc_mem)

            doc_mem.seek(0)


            st.session_state['vigas_docx_buf'] = doc_mem

            st.success("Memoria Exhaustiva (DOCX) acoplada. Ubíquela en el Centro de Entregables inferior para su descarga.")
    st.markdown("---")
    st.subheader("Exportacion de Planos DXF e IFC (Viga DMO/DES) — Protocolo ICONTEC")

    col_d1, col_d2, col_d3 = st.columns(3)
    with col_d1:
        dxf_empresa  = str(st.text_input("EMPRESA",  "INGENIERIA ESTRUCTURAL S.A.S.", key="vg_emp")).strip() or "S/D"
        dxf_proyecto = str(st.text_input("PROYECTO", "EDIFICIO RESIDENCIAL",         key="vg_proy")).strip() or "S/D"
    with col_d2:
        dxf_plano   = st.text_input("N PLANO",  "E-02",            key="vg_pla")
        dxf_elaboro = st.text_input("Elaboro", "ING. DISENO",      key="vg_ela")
    with col_d3:
        dxf_reviso  = st.text_input("Reviso",  "ING. REVISOR",     key="vg_rev")
        dxf_aprobo  = st.text_input("Aprobo",  "DIRECTOR PROYECTO", key="vg_apr")

    # Selector de tipo de seccion y tamano de papel
    col_sel1, col_sel2 = st.columns(2)
    with col_sel1:
        tipo_viga_dxf = st.selectbox("Seccion para Despiece DXF:", ["Viga Rectangular", "Viga T (con aletas)"], key="vg_tipo_sec")
    with col_sel2:
        papel_opciones_v = {
            "Carta  (216 x 279 mm)":       (21.6,  27.9,  "CARTA"),
            "Oficio (216 x 330 mm)":       (21.6,  33.0,  "OFICIO"),
            "Medio Pliego (500 x 707 mm)": (50.0,  70.7,  "MEDIO PLIEGO"),
            "Pliego       (707 x 1000 mm)":(70.7, 100.0,  "PLIEGO"),
        }
        papel_sel_v = st.selectbox("Tamano de Papel (ICONTEC)", list(papel_opciones_v.keys()), index=2, key="vg_papel")
        ANCHO_V, ALTO_V, PAPEL_LBL_V = papel_opciones_v[papel_sel_v]

    col_x1, col_x2 = st.columns(2)
    with col_x1:
        flag_dxf = st.button("Generar Plano DXF ICONTEC - Viga", use_container_width=True, key="vg_btn_dxf")
    with col_x2:
        flag_ifc = st.button("Generar Modelo IFC4 (BIM) - Viga", use_container_width=True, key="vg_btn_ifc")

    # ── IFC4 VIGAS ───────────────────────────────────────────────────────────
    if flag_ifc:
        try:
            import ifcopenshell
            import ifcopenshell.api
            import math as _m
            import tempfile, os as _os2

            def _make_ifc_viga(bw_cm, h_cm, L_m, bf_cm, hf_cm, es_t,
                               fc_mpa, fy_mpa,
                               n_sup, db_sup_mm, n_inf, db_inf_mm, dst_mm,
                               As_sup_cm2, As_inf_cm2, recub,
                               s_conf_cm2, s_centro_cm2, zona_conf_cm2,
                               Mu_kNm, phi_Mn_kNm, Vu_kN, phi_Vn_kN,
                               vol_m3, peso_kg,
                               empresa, proyecto, norma, nivel_sis):
                O = ifcopenshell.file(schema="IFC4")


                project  = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcProject",  name=proyecto)
                context = ifcopenshell.api.run("context.add_context", O, context_type="Model")
                body    = ifcopenshell.api.run("context.add_context", O,
                    context_type="Model", context_identifier="Body",
                    target_view="MODEL_VIEW", parent=context)

                site     = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcSite",     name="Sitio")
                building = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcBuilding", name="Edificio")
                storey   = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcBuildingStorey", name="Piso 1")
                _u = ifcopenshell.api.run("unit.add_si_unit", O, unit_type="LENGTHUNIT"); ifcopenshell.api.run("unit.assign_unit", O, units=[_u])
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=project, products=[site])
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=site, products=[building])
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=building, products=[storey])

                bw_m = bw_cm/100; h_m = h_cm/100; bf_m = bf_cm/100; hf_m = hf_cm/100
                r_m  = recub/100

                mat_conc  = ifcopenshell.api.run("material.add_material", O, name=f"CONCRETO_fc{fc_mpa:.0f}MPa")
                mat_steel = ifcopenshell.api.run("material.add_material", O, name=f"ACERO_fy{fy_mpa:.0f}MPa")

                # Perfil de la seccion
                if es_t:
                    # Viga T: perfil compuesto = alma + ala
                    profile = O.createIfcIShapeProfileDef(
                        ProfileType="AREA", ProfileName=f"VigaT_{bw_cm:.0f}x{h_cm:.0f}",
                        OverallWidth=bf_m, OverallDepth=h_m,
                        WebThickness=bw_m, FlangeThickness=hf_m)
                else:
                    profile = O.createIfcRectangleProfileDef(
                        ProfileType="AREA", ProfileName=f"VigaRect_{bw_cm:.0f}x{h_cm:.0f}",
                        XDim=bw_m, YDim=h_m)

                # Viga a lo largo de eje X
                origin  = O.createIfcCartesianPoint([0.0, 0.0, 0.0])
                x_dir   = O.createIfcDirection([1.0, 0.0, 0.0])
                z_dir   = O.createIfcDirection([0.0, 0.0, 1.0])
                y_dir   = O.createIfcDirection([0.0, 1.0, 0.0])
                place3d = O.createIfcAxis2Placement3D(Location=origin, Axis=z_dir, RefDirection=x_dir)
                solid   = O.createIfcExtrudedAreaSolid(
                    SweptArea=profile, Position=place3d,
                    ExtrudedDirection=O.createIfcDirection([1.0, 0.0, 0.0]),
                    Depth=L_m)
                shape_rep = O.createIfcShapeRepresentation(
                    ContextOfItems=body, RepresentationIdentifier="Body",
                    RepresentationType="SweptSolid", Items=[solid])
                beam = ifcopenshell.api.run("root.create_entity", O,
                    ifc_class="IfcBeam", name=f"VIG-{'T' if es_t else 'R'}_{bw_cm:.0f}x{h_cm:.0f}")
                beam.PredefinedType = "BEAM"
                ifcopenshell.api.run("geometry.assign_representation", O, product=beam, representation=shape_rep)
                ifcopenshell.api.run("geometry.edit_object_placement", O, product=beam)

                ifcopenshell.api.run("spatial.assign_container", O, relating_structure=storey, products=[beam])
                ifcopenshell.api.run("material.assign_material", O, products=[beam], material=mat_conc)

                def _color_diam(d):
                    c = {8:(0.8,0.9,1.0), 10:(0.4,0.8,1.0), 12:(0.1,0.7,0.2),
                         16:(1.0,0.8,0.0), 19:(1.0,0.5,0.0), 20:(1.0,0.5,0.0),
                         22:(1.0,0.2,0.2), 25:(0.9,0.0,0.0), 32:(0.6,0.0,0.6)}
                    return c.get(min(c.keys(), key=lambda k: abs(k-d)), (1,1,1))

                # Barras longitudinales superiores
                db_sup_m = db_sup_mm/1000; db_inf_m = db_inf_mm/1000
                y_sup_m  = h_m/2 - r_m;   y_inf_m  = -h_m/2 + r_m

                for tag, n_b, db_m, y_pos in [
                    ("LS", n_sup, db_sup_m, y_sup_m),
                    ("LI", n_inf, db_inf_m, y_inf_m),
                ]:
                    xs = ([bw_m * i/(n_b-1) - bw_m/2 + r_m for i in range(n_b)]
                          if n_b > 1 else [0.0])
                    for i, xb in enumerate(xs):
                        bar = ifcopenshell.api.run("root.create_entity", O,
                            ifc_class="IfcReinforcingBar", name=f"{tag}{i+1}")
                        bar.NominalDiameter = db_m
                        bar.SteelGrade      = f"fy={fy_mpa:.0f}MPa"
                        bar.BarSurface      = "TEXTURED"
                        bar.PredefinedType  = "MAIN"

                        p0 = O.createIfcCartesianPoint([0.0,  xb, y_pos])
                        p1 = O.createIfcCartesianPoint([L_m,  xb, y_pos])
                        pl = O.createIfcPolyline(Points=[p0, p1])
                        b_prof = O.createIfcCircleProfileDef(ProfileType="AREA", Radius=db_m/2)
                        b_place = O.createIfcAxis2Placement3D(
                            Location=O.createIfcCartesianPoint([0.0, xb, y_pos]),
                            Axis=O.createIfcDirection([1.0, 0.0, 0.0]),
                            RefDirection=O.createIfcDirection([0.0, 0.0, 1.0]))
                        b_solid = O.createIfcExtrudedAreaSolid(
                            SweptArea=b_prof, Position=b_place,
                            ExtrudedDirection=O.createIfcDirection([1.0, 0.0, 0.0]),
                            Depth=L_m)
                        b_rep = O.createIfcShapeRepresentation(
                            ContextOfItems=body, RepresentationIdentifier="Body",
                            RepresentationType="SweptSolid", Items=[b_solid])
                        
                        rebar_color = _color_diam(db_m * 1000)
                        style = O.createIfcSurfaceStyleRendering(
                            SurfaceColour=O.createIfcColourRgb(Red=rebar_color[0], Green=rebar_color[1], Blue=rebar_color[2]),
                            ReflectanceMethod="FLAT")
                        surface_style = O.createIfcSurfaceStyle(
                            Name=f"AceroLong_{db_m*1000:.0f}mm", Side="BOTH", Styles=[style])
                        O.createIfcStyledItem(Item=b_solid, Styles=[surface_style])

                        ifcopenshell.api.run("geometry.assign_representation", O, product=bar, representation=b_rep)
                        ifcopenshell.api.run("geometry.edit_object_placement", O, product=bar)
                        ifcopenshell.api.run("spatial.assign_container", O, relating_structure=storey, products=[bar])
                        ifcopenshell.api.run("material.assign_material", O, products=[bar], material=mat_steel)

                # Estribos: polilínea cerrada en plano YZ a cada posicion X
                dst_m = dst_mm/1000
                zona_c_m = zona_conf_cm2/100

                def _stirrup_positions(L_total, s_conf, s_cen, zona):
                    pos = []
                    x = s_conf/100/2
                    while x <= L_total:
                        pos.append(x)
                        sep = s_conf/100 if (x < zona or x > L_total - zona) else s_cen/100
                        x += sep
                    return pos

                st_positions = _stirrup_positions(L_m, s_conf_cm2, s_centro_cm2, zona_c_m)

                bx2 = bw_m/2 - r_m;  hy2 = h_m/2 - r_m
                for j, x_pos in enumerate(st_positions):
                    st_bar = ifcopenshell.api.run("root.create_entity", O,
                        ifc_class="IfcReinforcingBar", name=f"E{j+1}")
                    st_bar.NominalDiameter = dst_m
                    st_bar.SteelGrade      = f"fy={fy_mpa:.0f}MPa"
                    st_bar.BarSurface      = "TEXTURED"
                    st_bar.PredefinedType  = "MAIN"

                    hk = max(dst_m * 6, 0.075)
                    hk_dx = hk * _m.cos(_m.radians(45))
                    hk_dy = hk * _m.sin(_m.radians(45))
                    pts = [
                        O.createIfcCartesianPoint([x_pos + dst_m/2, -bx2 + hk_dx, -hy2 + hk_dy]),
                        O.createIfcCartesianPoint([x_pos, -bx2, -hy2]),
                        O.createIfcCartesianPoint([x_pos,  bx2, -hy2]),
                        O.createIfcCartesianPoint([x_pos,  bx2,  hy2]),
                        O.createIfcCartesianPoint([x_pos, -bx2,  hy2]),
                        O.createIfcCartesianPoint([x_pos, -bx2, -hy2]),
                        O.createIfcCartesianPoint([x_pos - dst_m/2, -bx2 + hk_dx, -hy2 + hk_dy]),
                    ]
                    polyline_st = O.createIfcPolyline(Points=pts)
                    # Volumen físico para que Revit lo reconozca como Armadura
                    st_swept = O.createIfcSweptDiskSolid(Directrix=polyline_st, Radius=dst_m/2)
                    
                    stirrup_color = _color_diam(dst_m * 1000)
                    st_style_rend = O.createIfcSurfaceStyleRendering(
                        SurfaceColour=O.createIfcColourRgb(Red=stirrup_color[0], Green=stirrup_color[1], Blue=stirrup_color[2]),
                        ReflectanceMethod="FLAT")
                    st_surface_style = O.createIfcSurfaceStyle(
                        Name=f"Estribo_{dst_m*1000:.0f}mm", Side="BOTH", Styles=[st_style_rend])
                    O.createIfcStyledItem(Item=st_swept, Styles=[st_surface_style])

                    st_rep = O.createIfcShapeRepresentation(
                        ContextOfItems=body, RepresentationIdentifier="Body",
                        RepresentationType="AdvancedSweptSolid", Items=[st_swept])
                    ifcopenshell.api.run("geometry.assign_representation", O, product=st_bar, representation=st_rep)
                    ifcopenshell.api.run("geometry.edit_object_placement", O, product=st_bar)
                    ifcopenshell.api.run("spatial.assign_container", O, relating_structure=storey, products=[st_bar])
                    ifcopenshell.api.run("material.assign_material", O, products=[st_bar], material=mat_steel)

                # Pset_NSR10_Viga
                pset = ifcopenshell.api.run("pset.add_pset", O, product=beam, name="Pset_NSR10_Viga")
                ifcopenshell.api.run("pset.edit_pset", O, pset=pset, properties={
                    "Norma":                norma,
                    "Nivel_Sismico":        nivel_sis,
                    "fc_MPa":               float(fc_mpa),
                    "fy_MPa":               float(fy_mpa),
                    "bw_cm":                float(bw_cm),
                    "h_cm":                 float(h_cm),
                    "L_m":                  float(L_m),
                    "recub_cm":             float(recub),
                    "n_barras_sup":         int(n_sup),
                    "db_sup_mm":            float(db_sup_mm),
                    "As_sup_cm2":           float(As_sup_cm2),
                    "n_barras_inf":         int(n_inf),
                    "db_inf_mm":            float(db_inf_mm),
                    "As_inf_cm2":           float(As_inf_cm2),
                    "db_estribo_mm":        float(dst_mm),
                    "s_conf_cm":            float(s_conf_cm2),
                    "s_centro_cm":          float(s_centro_cm2),
                    "zona_conf_cm":         float(zona_conf_cm2),
                    "n_estribos":           len(st_positions),
                    "Mu_kNm":               float(Mu_kNm),
                    "phi_Mn_kNm":           float(phi_Mn_kNm),
                    "phi_Mn_CUMPLE":        "SI" if phi_Mn_kNm >= Mu_kNm else "NO",
                    "Vu_kN":                float(Vu_kN),
                    "phi_Vn_kN":            float(phi_Vn_kN),
                    "phi_Vn_CUMPLE":        "SI" if phi_Vn_kN >= Vu_kN else "NO",
                    "Vol_concreto_m3":      float(vol_m3),
                    "Peso_acero_kg":        float(peso_kg),
                })

                # Pset_ConcreteElementGeneral (estándar OpenBIM IFC4)
                pset_ceg = ifcopenshell.api.run("pset.add_pset", O, product=beam, name="Pset_ConcreteElementGeneral")
                as_total_cm2 = round(As_sup_cm2 + As_inf_cm2, 2)
                ag_cm2 = round(bw_cm * h_cm, 2)
                rho_vol_pct = round(as_total_cm2 / ag_cm2 * 100, 3) if ag_cm2 > 0 else 0.0
                ifcopenshell.api.run("pset.edit_pset", O, pset=pset_ceg, properties={
                    "ConstructionMethod":       "CastInPlace",
                    "StructuralClass":          "Beam",
                    "StrengthClass":            f"f'c {fc_mpa:.1f} MPa",
                    "CompressiveStrength":      float(fc_mpa),
                    "YieldStress":              float(fy_mpa),
                    "ReinforcementVolumeRatio": float(rho_vol_pct) / 100.0,
                    "ReinforcementAreaRatio":   float(as_total_cm2),
                })

                with tempfile.NamedTemporaryFile(suffix='.ifc', delete=False) as tmp2:
                    tmp_path2 = tmp2.name
                O.write(tmp_path2)
                with open(tmp_path2, 'rb') as f_ifc:
                    buf_ifc_v = f_ifc.read()
                _os2.unlink(tmp_path2)
                return buf_ifc_v, len(st_positions)

            # Recopilar parametros
            _es_t_v = (tipo_viga_dxf == "Viga T (con aletas)")
            _bw   = bw_vt if _es_t_v else b_vc
            _h    = ht_vt if _es_t_v else h_vc
            _bf   = bf_vt if _es_t_v else b_vc
            _hf   = hf_vt if _es_t_v else 0
            _nb_s = nb_izq_sup
            _db_s = diam_dict.get(bar_izq_sup, 12.0)
            _nb_i = nb_cen_inf
            _db_i = diam_dict.get(bar_cen_inf, 12.0)
            _dst  = diam_dict.get(st_bar_vc, 10.0)
            _s1   = s_diseno_conf_cm
            _sc   = min(d_vc_mm/2/10, 60.0)
            _zc   = 2 * _h
            _Mu   = Mux_pos_max
            _phiMn= phiMn_cen_pos
            _Vu   = Vu_cv_input
            _phiVn= phi_Vn_kN
            _vol  = ((_bw * _h) / 10000) * L_vc
            _pkg  = (nb_izq_sup * (L_vc*100 + 30)/100 * rebar_dict.get(bar_izq_sup, 0.001)*100*7.85e-3 +
                     nb_cen_inf * (L_vc*100 + 30)/100 * rebar_dict.get(bar_cen_inf, 0.001)*100*7.85e-3)

            buf_ifc_viga, n_est_ifc = _make_ifc_viga(
                bw_cm=_bw, h_cm=_h, L_m=L_vc, bf_cm=_bf, hf_cm=_hf, es_t=_es_t_v,
                fc_mpa=fc, fy_mpa=fy,
                n_sup=_nb_s, db_sup_mm=_db_s, n_inf=_nb_i, db_inf_mm=_db_i, dst_mm=_dst,
                As_sup_cm2=As_req_izq_sup, As_inf_cm2=As_req_cen_inf, recub=r_vc,
                s_conf_cm2=_s1, s_centro_cm2=_sc, zona_conf_cm2=_zc,
                Mu_kNm=_Mu, phi_Mn_kNm=_phiMn, Vu_kN=_Vu, phi_Vn_kN=_phiVn,
                vol_m3=_vol, peso_kg=_pkg,
                empresa=dxf_empresa, proyecto=dxf_proyecto,
                norma=norma_sel, nivel_sis=nivel_sis)

            _ifc_fname_v = f"Viga_{'T' if _es_t_v else 'Rect'}_{_bw:.0f}x{_h:.0f}_IFC4.ifc"
            st.download_button(
                label=_t("Exportar IFC4 (BIM completo)", "Export IFC4 (full BIM)"),
                data=buf_ifc_viga,
                file_name=_ifc_fname_v,
                mime="application/x-step",
                key="ifc_viga_btn")
            st.caption(f"IFC4: IfcBeam + {_nb_s+_nb_i} barras long. + {n_est_ifc} estribos 3D + Pset_NSR10_Viga")
        except ImportError as e_imp_v:
            import sys as _sys_ifc_v
            import importlib
            importlib.invalidate_caches()
            _pyver_v = f"{_sys_ifc_v.version_info.major}.{_sys_ifc_v.version_info.minor}"
            _pyexe_v = _sys_ifc_v.executable
            st.error(
                f"⚠️ Error cargando **IfcOpenShell** en este entorno Python.\n\n"
                f"**Python activo:** `{_pyexe_v}` (versión {_pyver_v})\n\n"
                f"**Detalle del error:** `{str(e_imp_v)}`\n\n"
                f"**Solución 1:** Si acabas de instalar mediante pip, **REINICIA EL SERVIDOR DE STREAMLIT** (`Ctrl+C` y luego `streamlit run app.py`).\n\n"
                f"**Solución 2:** Si el error persiste, verifica la instalación en una terminal.\n"
                f"```bash\n{_pyexe_v} -m pip install ifcopenshell\n```"
            )
        except Exception as e_ifc_v:
            import traceback
            st.error(f"Error IFC viga: {e_ifc_v}")
            st.code(traceback.format_exc(), language='python')

    # ── DXF VIGAS (Protocolo ICONTEC Completo) ──────────────────────────────
    if flag_dxf:
        import ezdxf
        from ezdxf.enums import TextEntityAlignment
        from datetime import datetime as _dt_v

        doc_dxf = ezdxf.new('R2010', setup=True)
        doc_dxf.units = ezdxf.units.CM

        # Lineweights ICONTEC (int x 100)
        LW_V = {'CONCRETO':50,'ACERO_LONG':35,'ACERO_TRANS':25,'DOBLEZ':25,
                'COTAS':18,'TEXTO':18,'EJES':13,'ROTULO':35,'MARGEN':50}
        COL_V = {'CONCRETO':7,'ACERO_LONG':1,'ACERO_TRANS':4,'DOBLEZ':6,
                 'COTAS':2,'TEXTO':7,'EJES':8,'ROTULO':8,'MARGEN':7}
        for lay, lw in LW_V.items():
            if lay not in doc_dxf.layers:
                l2 = doc_dxf.layers.add(lay, color=COL_V[lay])
                l2.dxf.lineweight = lw
            else:
                doc_dxf.layers.get(lay).dxf.lineweight = lw

        msp = doc_dxf.modelspace()

        if 'ROMANS' not in doc_dxf.styles:
            try:    doc_dxf.styles.new('ROMANS', dxfattribs={'font':'romans.shx'})
            except: doc_dxf.styles.new('ROMANS', dxfattribs={'font':'txt.shx'})

        # Dimensiones
        if tipo_viga_dxf == "Viga Rectangular":
            H_cm = h_vc;  B_cm = b_vc
            bf_cm_d = b_vc;  hf_cm_d = 0
            nb_inf_d = nb_cen_inf;  nb_sup_d = nb_izq_sup
            tipo_sec_d = "RECTANGULAR"
            area_conc_d = B_cm * H_cm
            db_sup_d = diam_dict.get(bar_izq_sup, 12.0)
            db_inf_d = diam_dict.get(bar_cen_inf, 12.0)
        else:
            H_cm = ht_vt;  B_cm = bw_vt
            bf_cm_d = bf_vt;  hf_cm_d = hf_vt
            nb_inf_d = max(2, n_bt);  nb_sup_d = 2
            tipo_sec_d = "T"
            area_conc_d = B_cm*(H_cm-hf_cm_d) + bf_cm_d*hf_cm_d
            db_sup_d = diam_dict.get(bar_izq_sup, 12.0)
            db_inf_d = diam_dict.get(bar_cen_inf, 12.0)

        L_cm_d = L_vc * 100

        # Escala automatica
        ESCALAS_V = [200, 100, 50, 25, 20, 10]
        MARGEN_V  = 1.0
        ROT_H_V   = 4.0
        AREA_W_V  = ANCHO_V - 2*MARGEN_V
        AREA_H_V  = ALTO_V  - 2*MARGEN_V - ROT_H_V - 0.5

        escala_den_v = 200
        for den in reversed(ESCALAS_V):
            drawn_l = L_cm_d / den
            drawn_h = H_cm   / den
            if drawn_l <= AREA_W_V * 0.60 and drawn_h <= AREA_H_V * 0.50:
                escala_den_v = den
                break
        ESCALA_V   = 1.0 / escala_den_v
        ESCALA_LBL_V = f"1:{escala_den_v}"

        # Marco exterior
        msp.add_lwpolyline(
            [(MARGEN_V, MARGEN_V), (ANCHO_V-MARGEN_V, MARGEN_V),
             (ANCHO_V-MARGEN_V, ALTO_V-MARGEN_V),
             (MARGEN_V, ALTO_V-MARGEN_V), (MARGEN_V, MARGEN_V)],
            dxfattribs={'layer':'MARGEN'})

        # ── ALZADO LONGITUDINAL ─────────────────────────────────────────────
        L_d  = L_cm_d * ESCALA_V
        H_d  = H_cm   * ESCALA_V
        Bcol_d = bcol_vc * ESCALA_V

        ALZ_X0 = MARGEN_V + 1.0
        ALZ_Y0 = MARGEN_V + ROT_H_V + 1.5

        # Concreto viga
        msp.add_lwpolyline(
            [(ALZ_X0, ALZ_Y0), (ALZ_X0+L_d, ALZ_Y0),
             (ALZ_X0+L_d, ALZ_Y0+H_d), (ALZ_X0, ALZ_Y0+H_d), (ALZ_X0, ALZ_Y0)],
            dxfattribs={'layer':'CONCRETO'})

        # Columnas de apoyo
        for xc in [ALZ_X0, ALZ_X0+L_d]:
            msp.add_lwpolyline(
                [(xc-Bcol_d, ALZ_Y0-H_d*0.5), (xc+Bcol_d, ALZ_Y0-H_d*0.5),
                 (xc+Bcol_d, ALZ_Y0+H_d*1.5), (xc-Bcol_d, ALZ_Y0+H_d*1.5)],
                dxfattribs={'layer':'CONCRETO'})

        # Eje centroidal
        msp.add_line(
            (ALZ_X0-Bcol_d-0.3, ALZ_Y0+H_d/2),
            (ALZ_X0+L_d+Bcol_d+0.3, ALZ_Y0+H_d/2),
            dxfattribs={'layer':'EJES', 'linetype':'DASHDOT'})

        # Barras longitudinales
        rec_d = dp_vc * ESCALA_V
        y_sup_d = ALZ_Y0 + H_d - rec_d
        y_inf_d = ALZ_Y0 + rec_d
        gancho_d = min(0.8, 15 * ESCALA_V)

        # Superior con gancho 90
        msp.add_lwpolyline(
            [(ALZ_X0-Bcol_d+0.3, y_sup_d-gancho_d),
             (ALZ_X0-Bcol_d+0.3, y_sup_d),
             (ALZ_X0+L_d+Bcol_d-0.3, y_sup_d),
             (ALZ_X0+L_d+Bcol_d-0.3, y_sup_d-gancho_d)],
            dxfattribs={'layer':'ACERO_LONG','color':_color_acero_dxf(db_sup_d)})
        # Inferior
        msp.add_lwpolyline(
            [(ALZ_X0-Bcol_d+0.3, y_inf_d+gancho_d),
             (ALZ_X0-Bcol_d+0.3, y_inf_d),
             (ALZ_X0+L_d+Bcol_d-0.3, y_inf_d),
             (ALZ_X0+L_d+Bcol_d-0.3, y_inf_d+gancho_d)],
            dxfattribs={'layer':'ACERO_LONG','color':_color_acero_dxf(db_inf_d)})

        # Etiquetas barras
        def _bar_label_v(d_mm):
            if "Pulgadas" in bar_sys:
                for k, v in diam_dict.items():
                    if abs(v - d_mm) < 0.1: return k
            return f"O{d_mm:.0f}mm"

        msp.add_text(f"{nb_sup_d} {_bar_label_v(db_sup_d)} (sup)",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.22,
                        'insert':(ALZ_X0+L_d*0.25, y_sup_d+0.15)})
        msp.add_text(f"{nb_inf_d} {_bar_label_v(db_inf_d)} (inf)",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.22,
                        'insert':(ALZ_X0+L_d*0.25, y_inf_d-0.35)})

        # Estribos en alzado
        s_conf_d   = s_diseno_conf_cm
        s_cen_d    = min(d_vc_mm/2/10, 60.0)
        zona_d     = 2 * H_cm
        db_est_mm_d = diam_dict.get(st_bar_vc, 10.0)

        x_est_d = s_conf_d/2 * ESCALA_V
        cant_izq_d = 0
        while x_est_d <= zona_d*ESCALA_V:
            msp.add_line(
                (ALZ_X0+x_est_d, ALZ_Y0+rec_d),
                (ALZ_X0+x_est_d, ALZ_Y0+H_d-rec_d),
                dxfattribs={'layer':'ACERO_TRANS','color':_color_acero_dxf(db_est_mm_d)})
            x_est_d += s_conf_d*ESCALA_V;  cant_izq_d += 1

        x_cen_d = zona_d*ESCALA_V + s_cen_d*ESCALA_V;  cant_cen_d = 0
        while x_cen_d <= (L_d - zona_d*ESCALA_V):
            msp.add_line(
                (ALZ_X0+x_cen_d, ALZ_Y0+rec_d),
                (ALZ_X0+x_cen_d, ALZ_Y0+H_d-rec_d),
                dxfattribs={'layer':'ACERO_TRANS','color':_color_acero_dxf(db_est_mm_d)})
            x_cen_d += s_cen_d*ESCALA_V;  cant_cen_d += 1

        x_der_d = L_d - zona_d*ESCALA_V;  cant_der_d = 0
        while x_der_d <= L_d - s_conf_d/2*ESCALA_V:
            msp.add_line(
                (ALZ_X0+x_der_d, ALZ_Y0+rec_d),
                (ALZ_X0+x_der_d, ALZ_Y0+H_d-rec_d),
                dxfattribs={'layer':'ACERO_TRANS','color':_color_acero_dxf(db_est_mm_d)})
            x_der_d += s_conf_d*ESCALA_V;  cant_der_d += 1

        total_est_d = cant_izq_d + cant_cen_d + cant_der_d

        # Cotas de estribos
        y_cota_d = ALZ_Y0 - 0.9
        for x1, x2, txt in [
            (0, zona_d*ESCALA_V, f"{cant_izq_d}E O{db_est_mm_d:.0f}@{s_conf_d:.0f}cm"),
            (zona_d*ESCALA_V, L_d-zona_d*ESCALA_V, f"{cant_cen_d}E O{db_est_mm_d:.0f}@{s_cen_d:.0f}cm"),
            (L_d-zona_d*ESCALA_V, L_d, f"{cant_der_d}E O{db_est_mm_d:.0f}@{s_conf_d:.0f}cm"),
        ]:
            msp.add_line((ALZ_X0+x1, y_cota_d+0.2), (ALZ_X0+x2, y_cota_d+0.2), dxfattribs={'layer':'COTAS'})
            for xv in [x1, x2]:
                msp.add_line((ALZ_X0+xv, ALZ_Y0), (ALZ_X0+xv, y_cota_d-0.1), dxfattribs={'layer':'COTAS'})
            msp.add_text(txt, dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20,
                'insert':(ALZ_X0+(x1+x2)/2, y_cota_d-0.15),
                'align_point':(ALZ_X0+(x1+x2)/2, y_cota_d-0.15),'halign':1,'valign':2})

        # Cota luz total
        y_luz = ALZ_Y0 + H_d + 0.7
        msp.add_line((ALZ_X0, y_luz), (ALZ_X0+L_d, y_luz), dxfattribs={'layer':'COTAS'})
        for xv in [ALZ_X0, ALZ_X0+L_d]:
            msp.add_line((xv, ALZ_Y0+H_d), (xv, y_luz+0.2), dxfattribs={'layer':'COTAS'})
        msp.add_text(f"L = {L_vc:.2f} m",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.24,
                        'insert':(ALZ_X0+L_d/2, y_luz+0.12),
                        'align_point':(ALZ_X0+L_d/2, y_luz+0.12),'halign':1,'valign':2})

        # Cota h y b
        xc_alz = ALZ_X0 - 0.8
        msp.add_line((xc_alz, ALZ_Y0), (xc_alz, ALZ_Y0+H_d), dxfattribs={'layer':'COTAS'})
        for yv in [ALZ_Y0, ALZ_Y0+H_d]:
            msp.add_line((ALZ_X0, yv), (xc_alz-0.15, yv), dxfattribs={'layer':'COTAS'})
        msp.add_text(f"h={H_cm:.0f}cm",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20,
                        'insert':(xc_alz-0.12, ALZ_Y0+H_d/2),
                        'align_point':(xc_alz-0.12, ALZ_Y0+H_d/2),'halign':1,'valign':2,'rotation':90})

        # ── SECCION TRANSVERSAL ─────────────────────────────────────────────
        SEC_X0_V = ALZ_X0 + L_d + Bcol_d + 2.0
        SEC_Y0_V = ALZ_Y0
        SEC_AW   = min(AREA_W_V * 0.25, 8.0)
        SEC_AH   = AREA_H_V * 0.55

        dim_bw = B_cm;  dim_h = H_cm
        esc_sec_v = min((SEC_AW-1.5)/dim_bw, (SEC_AH-2.0)/dim_h, 1.5)
        sw_v = dim_bw * esc_sec_v;  sh_v = dim_h * esc_sec_v
        ox_v = SEC_X0_V + (SEC_AW-sw_v)/2;  oy_v = SEC_Y0_V + (SEC_AH-sh_v)/2

        msp.add_text(f"SECCION {tipo_sec_d}  ({ESCALA_LBL_V})",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.26,
                        'insert':(ox_v+sw_v/2, oy_v+sh_v+0.6),
                        'align_point':(ox_v+sw_v/2, oy_v+sh_v+0.6),'halign':1,'valign':2})

        if tipo_sec_d == "T":
            bf_d = bf_cm_d * esc_sec_v;  hf_d = hf_cm_d * esc_sec_v
            flange_x = ox_v - (bf_d - sw_v)/2
            msp.add_lwpolyline([
                (ox_v, oy_v), (ox_v+sw_v, oy_v),
                (ox_v+sw_v, oy_v+sh_v-hf_d),
                (flange_x+bf_d, oy_v+sh_v-hf_d),
                (flange_x+bf_d, oy_v+sh_v),
                (flange_x, oy_v+sh_v),
                (flange_x, oy_v+sh_v-hf_d),
                (ox_v, oy_v+sh_v-hf_d), (ox_v, oy_v)],
                dxfattribs={'layer':'CONCRETO'})
        else:
            msp.add_lwpolyline(
                [(ox_v,oy_v),(ox_v+sw_v,oy_v),(ox_v+sw_v,oy_v+sh_v),(ox_v,oy_v+sh_v),(ox_v,oy_v)],
                dxfattribs={'layer':'CONCRETO'})

        # Estribo cerrado con gancho 135° sismico NSR-10 — esquina inf-izq
        re_sv = dp_vc * esc_sec_v
        xm_v, xM_v = ox_v+re_sv, ox_v+sw_v-re_sv
        ym_v, yM_v = oy_v+re_sv, oy_v+sh_v-re_sv
        hl_v = min(0.5, re_sv * 0.8)   # longitud visual del gancho
        _hgx = hl_v * 0.707;  _hgy = hl_v * 0.707  # 45 grados
        pts_est_v = [
            (xm_v + _hgx, ym_v + _hgy),  # P1: cola entrada → interior ✓
            (xm_v, ym_v),                 # P2: esquina inf-izq (vértice del gancho)
            (xm_v, yM_v),                 # P3: esquina sup-izq
            (xM_v, yM_v),                 # P4: esquina sup-der
            (xM_v, ym_v),                 # P5: esquina inf-der
            (xm_v, ym_v),                 # P6: cierre (vuelve inf-izq)
            (xm_v + _hgx, ym_v + _hgy),  # P7: cola salida → interior ✓
        ]
        msp.add_lwpolyline(pts_est_v,
            dxfattribs={'layer':'ACERO_TRANS','color':_color_acero_dxf(db_est_mm_d)})

        # Barras en seccion
        rb_v = dp_vc * esc_sec_v
        r_bar_v = db_sup_d / 20 * esc_sec_v

        xs_sup_v = ([ox_v+rb_v+(sw_v-2*rb_v)*i/(nb_sup_d-1) for i in range(nb_sup_d)]
                    if nb_sup_d > 1 else [ox_v+sw_v/2])
        xs_inf_v = ([ox_v+rb_v+(sw_v-2*rb_v)*i/(nb_inf_d-1) for i in range(nb_inf_d)]
                    if nb_inf_d > 1 else [ox_v+sw_v/2])
        for xb in xs_sup_v:
            msp.add_circle((xb, oy_v+sh_v-rb_v), r_bar_v,
                dxfattribs={'layer':'ACERO_LONG','color':_color_acero_dxf(db_sup_d)})
        for xb in xs_inf_v:
            msp.add_circle((xb, oy_v+rb_v), db_inf_d/20*esc_sec_v,
                dxfattribs={'layer':'ACERO_LONG','color':_color_acero_dxf(db_inf_d)})

        # Cotas seccion b y h
        yc_sec_v = oy_v - 0.6
        msp.add_line((ox_v, yc_sec_v), (ox_v+sw_v, yc_sec_v), dxfattribs={'layer':'COTAS'})
        for xv in [ox_v, ox_v+sw_v]:
            msp.add_line((xv, oy_v), (xv, yc_sec_v-0.15), dxfattribs={'layer':'COTAS'})
        msp.add_text(f"b={B_cm:.0f}cm",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20,
                        'insert':(ox_v+sw_v/2, yc_sec_v-0.22),
                        'align_point':(ox_v+sw_v/2, yc_sec_v-0.22),'halign':1,'valign':2})

        xc_sec_v = ox_v+sw_v+0.6
        msp.add_line((xc_sec_v, oy_v), (xc_sec_v, oy_v+sh_v), dxfattribs={'layer':'COTAS'})
        for yv in [oy_v, oy_v+sh_v]:
            msp.add_line((ox_v+sw_v, yv), (xc_sec_v+0.15, yv), dxfattribs={'layer':'COTAS'})
        msp.add_text(f"h={H_cm:.0f}cm",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20,
                        'insert':(xc_sec_v+0.12, oy_v+sh_v/2),
                        'align_point':(xc_sec_v+0.12, oy_v+sh_v/2),'halign':1,'valign':2,'rotation':90})

        # ── TABLA DE DESPIECE ───────────────────────────────────────────────
        DESPIECE_X = SEC_X0_V + SEC_AW + 1.0
        DESPIECE_Y = ALZ_Y0 + AREA_H_V
        COL_WS_V = [3.0, 2.5, 1.8, 2.5, 2.5, 2.5]   # 6 cols
        ROW_H_V  = 0.60
        TAB_TW_V = sum(COL_WS_V)
        HDRS_V   = ["MARCA","DIAM.","CANT.","L (m)","FORMA","PESO kg"]

        _long_sup_v = (L_cm_d + 2*bcol_vc - 10 + 2*15) / 100
        _long_inf_v = (L_cm_d + 2*bcol_vc - 10 + 2*15) / 100
        _perim_est_v = 2*(B_cm-2*dp_vc) + 2*(H_cm-2*dp_vc) + 6*db_est_mm_d/10
        _long_est_v  = _perim_est_v/100 + 2*(13*db_est_mm_d/10)/100
        _peso_sup_v  = nb_sup_d * _long_sup_v * rebar_dict.get(bar_izq_sup, 0.001)*100*7.85e-3
        _peso_inf_v  = nb_inf_d * _long_inf_v * rebar_dict.get(bar_cen_inf, 0.001)*100*7.85e-3
        _peso_est_v  = total_est_d * _long_est_v * st_area_vc * 0.785
        _vol_c_v     = area_conc_d / 10000 * (L_cm_d/100)

        filas_des_v = [
            ("L1-Sup", _bar_label_v(db_sup_d), str(nb_sup_d),
             f"{_long_sup_v:.2f}", "Recta+gancho90", f"{_peso_sup_v:.1f}"),
            ("L2-Inf", _bar_label_v(db_inf_d), str(nb_inf_d),
             f"{_long_inf_v:.2f}", "Recta+gancho90", f"{_peso_inf_v:.1f}"),
            ("E1-Est", _bar_label_v(db_est_mm_d), str(total_est_d),
             f"{_long_est_v:.2f}", "Cerrado 135(sis.)", f"{_peso_est_v:.1f}"),
        ]

        # Dibujar tabla
        msp.add_text("DESPIECE DE ACERO — ICONTEC 2289",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.26,
                        'insert':(DESPIECE_X+TAB_TW_V/2, DESPIECE_Y+0.25),
                        'align_point':(DESPIECE_X+TAB_TW_V/2, DESPIECE_Y+0.25),
                        'halign':1,'valign':2})
        y_tab = DESPIECE_Y
        cx_tab = DESPIECE_X
        for hdr, cw in zip(HDRS_V, COL_WS_V):
            msp.add_lwpolyline(
                [(cx_tab,y_tab-ROW_H_V),(cx_tab+cw,y_tab-ROW_H_V),
                 (cx_tab+cw,y_tab),(cx_tab,y_tab),(cx_tab,y_tab-ROW_H_V)],
                dxfattribs={'layer':'ROTULO'})
            msp.add_text(hdr,dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.18,
                'insert':(cx_tab+cw/2,y_tab-ROW_H_V/2),
                'align_point':(cx_tab+cw/2,y_tab-ROW_H_V/2),'halign':1,'valign':2})
            cx_tab += cw
        y_tab -= ROW_H_V

        for row in filas_des_v:
            cx_tab = DESPIECE_X
            for val, cw in zip(row, COL_WS_V):
                msp.add_lwpolyline(
                    [(cx_tab,y_tab-ROW_H_V),(cx_tab+cw,y_tab-ROW_H_V),
                     (cx_tab+cw,y_tab),(cx_tab,y_tab),(cx_tab,y_tab-ROW_H_V)],
                    dxfattribs={'layer':'ROTULO'})
                msp.add_text(val,dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.18,
                    'insert':(cx_tab+cw/2,y_tab-ROW_H_V/2),
                    'align_point':(cx_tab+cw/2,y_tab-ROW_H_V/2),'halign':1,'valign':2})
                cx_tab += cw
            y_tab -= ROW_H_V

        # Totales
        tot_acero_v = _peso_sup_v + _peso_inf_v + _peso_est_v
        for labels in [
            [("TOTAL ACERO","","","","",f"{tot_acero_v:.1f} kg")],
            [("CONCRETO",f"fc={fc:.0f}MPa",f"{_vol_c_v:.3f}m3","","","")],
        ]:
            cx_tab = DESPIECE_X
            for val, cw in zip(labels[0], COL_WS_V):
                msp.add_lwpolyline(
                    [(cx_tab,y_tab-ROW_H_V),(cx_tab+cw,y_tab-ROW_H_V),
                     (cx_tab+cw,y_tab),(cx_tab,y_tab),(cx_tab,y_tab-ROW_H_V)],
                    dxfattribs={'layer':'ROTULO'})
                msp.add_text(val,dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.18,
                    'insert':(cx_tab+cw/2,y_tab-ROW_H_V/2),
                    'align_point':(cx_tab+cw/2,y_tab-ROW_H_V/2),'halign':1,'valign':2})
                cx_tab += cw
            y_tab -= ROW_H_V

        # Diagrama de ganchos
        DOBL_XV  = DESPIECE_X
        DOBL_YV  = y_tab - 1.0
        msp.add_text("Gancho 135 sismico (estribo):",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20,'insert':(DOBL_XV, DOBL_YV)})
        hl_d_v = 6*db_est_mm_d/10 * 0.25
        msp.add_lwpolyline([
            (DOBL_XV+0.2, DOBL_YV-0.5), (DOBL_XV+0.2, DOBL_YV-1.1),
            (DOBL_XV+0.2+hl_d_v, DOBL_YV-1.1+hl_d_v)],
            dxfattribs={'layer':'DOBLEZ'})
        msp.add_text(f"ext.={6*db_est_mm_d:.0f}mm",
            dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.17,'insert':(DOBL_XV+0.5,DOBL_YV-0.8)})

        DOBL_XV2 = DOBL_XV + TAB_TW_V*0.5
        msp.add_text("Gancho 90 barra long.:",
            dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.20,'insert':(DOBL_XV2, DOBL_YV)})
        hl_d_v2 = 12*db_sup_d/10 * 0.20
        msp.add_lwpolyline([
            (DOBL_XV2+0.2, DOBL_YV-0.3), (DOBL_XV2+0.2, DOBL_YV-0.9),
            (DOBL_XV2+0.2+hl_d_v2, DOBL_YV-0.9)],
            dxfattribs={'layer':'DOBLEZ'})
        msp.add_text(f"ext.={12*db_sup_d:.0f}mm",
            dxfattribs={'layer':'COTAS','style':'ROMANS','height':0.17,'insert':(DOBL_XV2+0.5,DOBL_YV-0.7)})

        # ── ROTULO ICONTEC ──────────────────────────────────────────────────
        ROT_X_V = MARGEN_V;  ROT_Y_V = MARGEN_V
        ROT_W_V = ANCHO_V - 2*MARGEN_V
        msp.add_lwpolyline(
            [(ROT_X_V, ROT_Y_V), (ROT_X_V+ROT_W_V, ROT_Y_V),
             (ROT_X_V+ROT_W_V, ROT_Y_V+ROT_H_V),
             (ROT_X_V, ROT_Y_V+ROT_H_V), (ROT_X_V, ROT_Y_V)],
            dxfattribs={'layer':'ROTULO'})

        celdas_rot_v = [
            ("EMPRESA",  dxf_empresa,  0.0,          2.5,  ROT_W_V*0.46, 1.0),
            ("PROYECTO", dxf_proyecto, 0.0,          1.5,  ROT_W_V*0.46, 1.0),
            ("CONTENIDO",f"Viga {tipo_sec_d} — Despiece ICONTEC", 0.0, 0.5, ROT_W_V*0.46, 1.0),
            ("N. PLANO", dxf_plano,    ROT_W_V*0.46, 2.5,  ROT_W_V*0.18, 1.0),
            ("ESCALA",   ESCALA_LBL_V, ROT_W_V*0.46, 1.5,  ROT_W_V*0.18, 1.0),
            ("FECHA",    _dt_v.now().strftime("%d/%m/%Y"), ROT_W_V*0.46, 0.5, ROT_W_V*0.18, 1.0),
            ("NORMA",    st.session_state.get('norma_sel','NSR-10')[:10], ROT_W_V*0.64, 2.5,  ROT_W_V*0.12, 1.0),
            ("REVISION", "0",          ROT_W_V*0.76, 2.5,  ROT_W_V*0.08, 1.0),
            ("HOJA",     "1/1",        ROT_W_V*0.76, 1.5,  ROT_W_V*0.08, 1.0),
            ("PAPEL",    PAPEL_LBL_V,  ROT_W_V*0.64, 1.5,  ROT_W_V*0.12, 1.0),
            ("ELABORO",  dxf_elaboro,  ROT_W_V*0.84, 2.5,  ROT_W_V*0.08, 1.0),
            ("REVISO",   dxf_reviso,   ROT_W_V*0.84, 1.5,  ROT_W_V*0.08, 1.0),
            ("APROBO",   dxf_aprobo,   ROT_W_V*0.92, 2.5,  ROT_W_V*0.08, 1.0),
            ("ACERO kg", f"{tot_acero_v:.1f}", ROT_W_V*0.76, 0.5, ROT_W_V*0.12, 2.0),
            ("CONC. m3", f"{_vol_c_v:.3f}",   ROT_W_V*0.88, 0.5, ROT_W_V*0.12, 2.0),
        ]
        for etiq, valor, xr, yr, cw2, ch2 in celdas_rot_v:
            cx2_v = ROT_X_V + xr;  cy2_v = ROT_Y_V + yr
            msp.add_lwpolyline(
                [(cx2_v,cy2_v),(cx2_v+cw2,cy2_v),(cx2_v+cw2,cy2_v+ch2),(cx2_v,cy2_v+ch2),(cx2_v,cy2_v)],
                dxfattribs={'layer':'ROTULO'})
            msp.add_text(etiq,
                dxfattribs={'layer':'TEXTO','style':'ROMANS','height':0.13,
                            'insert':(cx2_v+0.07, cy2_v+ch2-0.17),'color':8})
            msp.add_text(valor,
                dxfattribs={'layer':'TEXTO','style':'ROMANS',
                            'height':0.28 if etiq in("EMPRESA","PROYECTO") else 0.20,
                            'insert':(cx2_v+cw2/2, cy2_v+ch2/2-0.08),
                            'align_point':(cx2_v+cw2/2, cy2_v+ch2/2-0.08),
                            'halign':1,'valign':2})

        # Linea sobre rotulo
        msp.add_line(
            (MARGEN_V, ROT_H_V+MARGEN_V),
            (ANCHO_V-MARGEN_V, ROT_H_V+MARGEN_V),
            dxfattribs={'layer':'MARGEN'})

        # ── EXPORTAR ─────────────────────────────────────────────────────────
        import tempfile, os as _os_v, io as _io_v
        with tempfile.NamedTemporaryFile(suffix='.dxf', delete=False) as tmp_v:
            tmp_path_v = tmp_v.name
        doc_dxf.saveas(tmp_path_v)
        with open(tmp_path_v, 'rb') as fv:
            dxf_bytes_v = fv.read()
        _os_v.unlink(tmp_path_v)

        nombre_dxf_v = f"Viga_{tipo_sec_d}_{B_cm:.0f}x{H_cm:.0f}_{PAPEL_LBL_V.replace(' ','_')}.dxf"
        st.download_button(
            label=_t("Descargar DXF ICONTEC - Viga", "Download DXF ICONTEC - Beam"),
            data=dxf_bytes_v, file_name=nombre_dxf_v, mime="application/dxf",
            key="dxf_viga_dl")

        try:
            import matplotlib; matplotlib.use('Agg')
            import matplotlib.pyplot as _mpdf_v
            from ezdxf.addons.drawing import RenderContext, Frontend
            from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
            fig_w_v = ANCHO_V * 0.3937;  fig_h_v = ALTO_V * 0.3937
            fig_pdf_v, ax_pdf_v = _mpdf_v.subplots(figsize=(fig_w_v, fig_h_v))
            fig_pdf_v.patch.set_facecolor('white'); ax_pdf_v.set_facecolor('white')
            
            from ezdxf.addons.drawing.config import Configuration, BackgroundPolicy
            _config_v = Configuration.defaults().with_changes(background_policy=BackgroundPolicy.WHITE)
            Frontend(RenderContext(doc_dxf), MatplotlibBackend(ax_pdf_v), config=_config_v).draw_layout(msp, finalize=True)
            
            ax_pdf_v.set_aspect('equal')
            ax_pdf_v.axis('off')
            _cx_v = ANCHO_V / 2
            _cy_v = ALTO_V  / 2
            ax_pdf_v.set_xlim(_cx_v - ANCHO_V/2 - 0.5, _cx_v + ANCHO_V/2 + 0.5)
            ax_pdf_v.set_ylim(_cy_v - ALTO_V /2 - 0.5, _cy_v + ALTO_V /2 + 0.5)
            bio_pdf_v = _io_v.BytesIO()
            fig_pdf_v.savefig(bio_pdf_v, format='pdf', bbox_inches='tight', dpi=150, facecolor='white')
            bio_pdf_v.seek(0);  _mpdf_v.close(fig_pdf_v)
            st.download_button(
                label=_t("Descargar PDF Imprimible - Viga", "Download Printable PDF - Beam"),
                data=bio_pdf_v.getvalue(),
                file_name=nombre_dxf_v.replace('.dxf','.pdf'),
                mime="application/pdf",
                key="pdf_Viga_dl")
            st.success(_t(
                f"Plano generado | Papel: {PAPEL_LBL_V} | Escala: {ESCALA_LBL_V} | Lineweights ICONTEC",
                f"Plot generated | Paper: {PAPEL_LBL_V} | Scale: {ESCALA_LBL_V} | ICONTEC lineweights"))
        except Exception as e_pdf:
            st.warning(f"PDF no disponible: {e_pdf}")




# ??????????????????????????????????????????
# FOOTER
# ??????????????????????????????????????????
st.markdown("---")
st.markdown(f"""
> **Suite de Hormigón Armado — Multi-Norma**  
> Norma activa: `{norma_sel}` | Nivel sísmico: `{nivel_sis}`  
> f'c = {fc:.2f} MPa | fy = {fy:.0f} MPa | Ec = {Ec:.0f} MPa | β? = {beta1:.3f}  
> **Referencia:** {code['ref']}  
>  ? *Las herramientas son de apoyo para el diseño. Verifique siempre con la norma vigente del país.*
""")

# ??????????????????????????????????????????
# PERSISTENCIA SUPABASE — Sidebar Guardar/Cargar
# ??????????????????????????????????????????
st.sidebar.markdown("---")
st.sidebar.markdown("### ?? Guardar / Cargar Proyecto Vigas")

# --- URL Query Params: sincronizar nombre del proyecto con la URL ---
_qp = st.query_params
if "proyecto_vigas" in _qp and "nombre_proyecto_vigas" not in st.session_state:
    st.session_state["nombre_proyecto_vigas"] = _qp["proyecto_vigas"]

nombre_producido_v = st.session_state.get("nombre_proyecto_vigas", "")

st.sidebar.markdown("**Nuevo Proyecto / Guardar**")
nombre_proy_v = st.sidebar.text_input("Nombre para guardar", value=nombre_producido_v, key="input_guardar_proy_v")

if st.sidebar.button(" Guardar Proyecto Vigas", use_container_width=True):
    if nombre_proy_v:
        ok, msg = guardar_proyecto_supabase_vigas(nombre_proy_v, capturar_estado_vigas())
        if ok:
            st.session_state["nombre_proyecto_vigas"] = nombre_proy_v
            # Sincronizar URL con el nombre del proyecto
            st.query_params["proyecto_vigas"] = nombre_proy_v
            st.sidebar.success(msg)
            st.rerun()
        else:
            st.sidebar.error(msg)
    else:
        st.sidebar.warning("Escribe un nombre de proyecto")

st.sidebar.markdown("**Cargar Proyecto Existente**")
lista_proy_v = listar_proyectos_supabase_vigas()

if lista_proy_v:
    idx_default_v = 0
    if nombre_producido_v in lista_proy_v:
        idx_default_v = lista_proy_v.index(nombre_producido_v)
    nombre_proy_cargar_v = st.sidebar.selectbox(
        "Selecciona un proyecto", lista_proy_v, index=idx_default_v, key="select_cargar_proy_v"
    )
    if st.sidebar.button(" Cargar Proyecto Vigas", use_container_width=True):
        ok, msg = cargar_proyecto_supabase_vigas(nombre_proy_cargar_v)
        if ok:
            st.session_state["nombre_proyecto_vigas"] = nombre_proy_cargar_v
            st.query_params["proyecto_vigas"] = nombre_proy_cargar_v
            st.sidebar.success(msg)
            st.rerun()
        else:
            st.sidebar.error(msg)
else:
    st.sidebar.info("No hay proyectos de Vigas en la nube.")

# 
# MÓDULO: CUADRO DE MANDO GENERAL (M3 — Historial Diseños)
# 
if modulo_sel == " Cuadro de Mando General":
    st.subheader("Cuadro de Mando General — Historial de Diseños")
    historial = st.session_state.get("historial_disenos", [])

    if not historial:
        st.info("ℹ No hay diseños registrados aún. Usa el botón ** Enviar a Cuadro de Mando** en cada módulo para guardar resultados.")
    else:
        df_hist = pd.DataFrame(historial)

        # Métricas resumen
        total = len(df_hist)
        aprobados  = len(df_hist[df_hist["Estado"].str.contains("APROBADO|", na=False)])
        rechazados = total - aprobados
        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Total Diseños", total)
        col_m2.metric(" Aprobados", aprobados, delta=f"{aprobados/total*100:.0f}%")
        col_m3.metric(" No Cumplen", rechazados, delta=f"-{rechazados/total*100:.0f}%", delta_color="inverse")

        st.markdown("---")

        # Tabla principal con colores de estado
        st.markdown("####  Registro de Diseños")
        def _color_estado(val):
            if "APROBADO" in str(val) or "" in str(val):
                return "background-color:#1a4a2e; color:#4ade80"
            elif "NO CUMPLE" in str(val) or "EXCEDE" in str(val):
                return "background-color:#4a1a1a; color:#f87171"
            return ""

        st.dataframe(
            df_hist.style.map(_color_estado, subset=["Estado"]),
            use_container_width=True, hide_index=True
        )

        st.markdown("---")

        # Gráfico Pie de distribución
        col_g1, col_g2 = st.columns([1, 2])
        with col_g1:
            import plotly.graph_objects as go_cmd
            fig_pie = go_cmd.Figure(go_cmd.Pie(
                labels=[" Aprobados", " No Cumplen"],
                values=[aprobados, rechazados],
                marker=dict(colors=["#22c55e", "#ef4444"]),
                hole=0.4,
                textinfo="label+percent",
                hovertemplate="%{label}: %{value}<extra></extra>"
            ))
            fig_pie.update_layout(
                height=280, margin=dict(l=10, r=10, t=20, b=10),
                paper_bgcolor="rgba(0,0,0,0)", font_color="white",
                showlegend=False
            )
            st.plotly_chart(fig_pie, use_container_width=True)

        with col_g2:
            # Barra de módulos
            modulos_count = df_hist.groupby("Módulo").size().reset_index(name="N")
            modulos_count["Color"] = modulos_count["N"].apply(lambda x: "#60a5fa")
            fig_bar = go_cmd.Figure(go_cmd.Bar(
                x=modulos_count["N"], y=modulos_count["Módulo"],
                orientation="h",
                marker_color="#60a5fa",
                text=modulos_count["N"], textposition="outside",
                hovertemplate="%{y}: %{x} diseños<extra></extra>"
            ))
            fig_bar.update_layout(
                height=280, margin=dict(l=10, r=10, t=20, b=10),
                paper_bgcolor="rgba(0,0,0,0)", font_color="white",
                xaxis=dict(showgrid=False), yaxis=dict(showgrid=False),
                title_text="Diseños por Módulo"
            )
            st.plotly_chart(fig_bar, use_container_width=True)

        st.markdown("---")

        # Export CSV y limpiar
        col_exp1, col_exp2 = st.columns(2)
        with col_exp1:
            csv_hist = df_hist.to_csv(index=False).encode("utf-8")
            st.download_button("Exportar Historial CSV", data=csv_hist,
                               file_name=f"historial_disenos_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                               mime="text/csv")
        with col_exp2:
            if st.button("Limpiar Historial", type="secondary"):
                st.session_state.historial_disenos = []
                st.rerun()

#  PANEL GLOBAL DE ENTREGABLES Y NORMAS 
if 'mostrar_entregables' in locals() and mostrar_entregables:
    # Capturar buffers de memoria si existen
    b_docx = st.session_state.get("vigas_docx_buf", None)
    b_xlsx = st.session_state.get("vigas_xlsx_buf", None)
    b_dxf = st.session_state.get("vigas_dxf_buf", None)
    b_ifc = st.session_state.get("vigas_ifc_buf", None)

    mostrar_entregables(
        norma_sel=st.session_state.get("norma_sel", "NSR-10 (Colombia)"),
        modulo_key="vigas_losas",
        titulo="Diseño de Vigas y Losas",
        docx_buf=b_docx,
        excel_buf=b_xlsx,
        dxf_buf=b_dxf,
        ifc_buf=b_ifc,
        docx_name="Memoria_Vigas.docx",
        excel_name="Cantidades_Vigas.xlsx",
        dxf_name="Planos_Vigas.dxf",
        ifc_name="Modelo_Vigas.ifc"
    )
