import streamlit as st
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es

st.set_page_config(page_title=_t("Diseño Sísmico", "Seismic Design"), layout="wide")
st.image(r"assets/seismic_header_1773257220819.png", use_container_width=True)
st.title(_t("Diseño Sísmico y Espectros", "Seismic Design and Spectra"))
st.markdown(_t("Análisis dinámico simplificado (frecuencia natural 1 GDL) y Generación de Espectros de Respuesta de Diseño para diversas normativas de América (NSR, E.030, NB 1225001, ASCE 7, etc).", 
               "Simplified dynamic analysis (1 DOF natural frequency) and Design Response Spectrum Generation for various American codes (NSR, E.030, NB 1225001, ASCE 7, etc)."))

# ─────────────────────────────────────────────
# CONFIGURACIÓN LATERAL
norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(f'<div style="background:#1e3a1e;border-radius:6px;padding:8px;margin-bottom:10px;"><img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;"><span style="color:#7ec87e;font-weight:600;">{_t("Normativa Activa:","Code:")} {norma_sel}</span></div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# 1. CÁLCULO DEL ESPECTRO (función genérica)
# =============================================================================
def compute_spectrum(norma, params):
    """
    Devuelve (T_domain, Sa_vals, params_texto, key_periods) donde key_periods = {'T0', 'Tc', 'TL', 'max_Sa'}
    """
    T_domain = np.linspace(0.01, 5.0, 500)
    Sa_vals = np.zeros_like(T_domain)
    key_periods = {'T0':0, 'Tc':0, 'TL':0, 'max_Sa':0}
    params_texto = []

    if "NSR" in norma:
        # NSR-10 (Colombia)
        Aa = params['Aa']
        Av = params['Av']
        Fa = params['Fa']
        Fv = params['Fv']
        I = params['I']
        Tc = 0.48 * (Av * Fv) / (Aa * Fa)
        T0 = 0.1 * Tc
        TL = 2.4 * Fv  # aproximación (según NSR-10)
        key_periods = {'T0': T0, 'Tc': Tc, 'TL': TL, 'max_Sa': 2.5 * Aa * Fa * I}
        for i, t in enumerate(T_domain):
            if t < T0:
                sa = Aa * Fa * (1.0 + (t/T0)*(2.5 - 1.0))
            elif t <= Tc:
                sa = 2.5 * Aa * Fa
            elif t <= TL:
                sa = 1.2 * Av * Fv / t
            else:
                sa = 1.2 * Av * Fv * TL / (t**2)
            Sa_vals[i] = sa * I
        params_texto = [f"{_t('Norma', 'Code')}: NSR-10 (Colombia)", f"Aa = {Aa:.3f}, Av = {Av:.3f}", 
                        f"Fa = {Fa:.2f}, Fv = {Fv:.2f}", f"I = {I:.2f}", f"Tc = {Tc:.3f} s", f"T0 = {T0:.3f} s", f"TL = {TL:.2f} s"]

    elif "E.060" in norma or "Perú" in norma:
        # E.030 (Perú)
        Z = params['Z']
        S = params['S']
        Tp = params['Tp']
        Tl = params['Tl']
        U = params['U']
        R = params['R']
        key_periods = {'T0': 0.2*Tp, 'Tc': Tp, 'TL': Tl, 'max_Sa': 2.5 * Z * U * S / R}
        for i, t in enumerate(T_domain):
            if t < 0.2*Tp:
                C = 1 + 1.5 * (t / (0.2*Tp))
            elif t <= Tp:
                C = 2.5
            elif t <= Tl:
                C = 2.5 * (Tp / t)
            else:
                C = 2.5 * ((Tp * Tl) / (t**2))
            Sa_vals[i] = (Z * U * C * S) / R
        params_texto = [f"{_t('Norma', 'Code')}: E.030 (Perú)", f"Z = {Z:.3f}", f"S = {S:.2f}", f"Tp = {Tp:.3f} s", f"Tl = {Tl:.3f} s", f"U = {U:.2f}", f"R = {R:.2f}"]

    elif "1225001" in norma or "Bolivia" in norma:
        # NB 1225001 (Bolivia) - simplificado
        Z = params['Z']
        Fa = params['Fa']
        Tc = params['Tc']
        key_periods = {'T0': 0.2*Tc, 'Tc': Tc, 'TL': 2.0, 'max_Sa': 2.5 * Z * Fa}
        for i, t in enumerate(T_domain):
            if t <= Tc:
                sa = 2.5 * Z * Fa
            else:
                sa = 2.5 * Z * Fa * (Tc / t)
            Sa_vals[i] = sa
        params_texto = [f"{_t('Norma', 'Code')}: NB 1225001 (Bolivia)", f"Z = {Z:.3f}", f"Fa = {Fa:.2f}", f"Tc = {Tc:.3f} s"]

    else:
        # ASCE 7 / ACI 318 (EE.UU.)
        S_DS = params['S_DS']
        S_D1 = params['S_D1']
        TL = params.get('TL', 4.0)
        T0 = 0.2 * (S_D1 / S_DS)
        Ts = S_D1 / S_DS
        key_periods = {'T0': T0, 'Tc': Ts, 'TL': TL, 'max_Sa': S_DS}
        for i, t in enumerate(T_domain):
            if t < T0:
                sa = S_DS * (0.4 + 0.6 * (t / T0))
            elif t <= Ts:
                sa = S_DS
            elif t <= TL:
                sa = S_D1 / t
            else:
                sa = (S_D1 * TL) / (t**2)
            Sa_vals[i] = sa
        params_texto = [f"{_t('Norma', 'Code')}: {norma_sel}", f"S_DS = {S_DS:.3f} g", f"S_D1 = {S_D1:.3f} g", f"T0 = {T0:.3f} s", f"Ts = {Ts:.3f} s", f"TL = {TL:.2f} s"]

    return T_domain, Sa_vals, params_texto, key_periods

# =============================================================================
# 2. INTERFAZ DE PARÁMETROS
# =============================================================================
st.header(_t("📊 2. Espectro Sísmico de Diseño", "📊 2. Seismic Design Response Spectrum"))

with st.container():
    col_left, col_right = st.columns([1, 2])

    with col_left:
        st.subheader(_t("Parámetros", "Parameters"))
        if "NSR" in norma_sel:
            st.write("**Norma: NSR-10 (Colombia)**")
            ciudad = st.selectbox(_t("Ciudad Principal:", "Main City:"), 
                                   ["Bogotá", "Medellín", "Cali", "Barranquilla", "Bucaramanga", _t("Otra / Custom", "Other / Custom")],
                                   index=["Bogotá", "Medellín", "Cali", "Barranquilla", "Bucaramanga", "Otra / Custom"].index(st.session_state.get("s_ciudad", "Bogotá")),
                                   key="s_ciudad")
            if ciudad == "Bogotá": Aa_def, Av_def = 0.15, 0.20
            elif ciudad == "Medellín": Aa_def, Av_def = 0.15, 0.20
            elif ciudad == "Cali": Aa_def, Av_def = 0.25, 0.25
            elif ciudad == "Barranquilla": Aa_def, Av_def = 0.10, 0.10
            elif ciudad == "Bucaramanga": Aa_def, Av_def = 0.25, 0.25
            else: Aa_def, Av_def = 0.20, 0.20
            Aa = st.number_input("Aa (Aceleración pico efectiva)", 0.05, 0.50, st.session_state.get("s_Aa", Aa_def), 0.05, key="s_Aa")
            Av = st.number_input("Av (Velocidad pico efectiva)", 0.05, 0.50, st.session_state.get("s_Av", Av_def), 0.05, key="s_Av")
            perfil = st.selectbox("Perfil de Suelo (NSR-10)", ["A", "B", "C", "D", "E"], index=["A","B","C","D","E"].index(st.session_state.get("s_perfil", "D")), key="s_perfil")
            Fa_vals = {"A":0.8, "B":1.0, "C":1.2, "D":1.4, "E":2.5}
            Fv_vals = {"A":0.8, "B":1.0, "C":1.6, "D":2.0, "E":3.2}
            Fa = Fa_vals[perfil]; Fv = Fv_vals[perfil]
            I = st.selectbox("Grupo de Uso (Importancia I)", [1.0, 1.1, 1.25, 1.5], index=[1.0,1.1,1.25,1.5].index(st.session_state.get("s_I", 1.0)), key="s_I")
            params = {"Aa": Aa, "Av": Av, "Fa": Fa, "Fv": Fv, "I": I}
            st.caption(_t("Fa, Fv según Tabla A.2.4-3 y A.2.4-4 NSR-10", "Fa, Fv from Tables A.2.4-3 and A.2.4-4 NSR-10"))

        elif "E.060" in norma_sel or "Perú" in norma_sel:
            st.write("**Norma: E.030 (Perú)**")
            ciudad_peru = st.selectbox("Zona / Ciudad:", ["Zona 4 (Lima, Costa)", "Zona 3 (Arequipa)", "Zona 2 (Cusco)", "Zona 1 (Selva)"],
                                        index=["Zona 4 (Lima, Costa)", "Zona 3 (Arequipa)", "Zona 2 (Cusco)", "Zona 1 (Selva)"].index(st.session_state.get("s_peru_ciudad", "Zona 4 (Lima, Costa)")),
                                        key="s_peru_ciudad")
            Z = 0.45 if "4" in ciudad_peru else 0.35 if "3" in ciudad_peru else 0.25 if "2" in ciudad_peru else 0.10
            suelo_peru = st.selectbox("Perfil de Suelo (S0, S1, S2, S3)", ["S0", "S1", "S2", "S3"],
                                       index=["S0","S1","S2","S3"].index(st.session_state.get("s_peru_suelo", "S2")),
                                       key="s_peru_suelo")
            S_dict = {"S0":0.8, "S1":1.0, "S2":1.15, "S3":1.4}
            Tp_dict = {"S0":0.3, "S1":0.4, "S2":0.6, "S3":1.0}
            Tl_dict = {"S0":3.0, "S1":2.5, "S2":2.0, "S3":1.6}
            S = S_dict[suelo_peru]; Tp = Tp_dict[suelo_peru]; Tl = Tl_dict[suelo_peru]
            U = st.selectbox("Uso o Importancia U (A=1.5, B=1.3, C=1.0)", [1.5, 1.3, 1.0],
                              index=[1.5,1.3,1.0].index(st.session_state.get("s_peru_U", 1.0)),
                              key="s_peru_U")
            R = st.number_input("Coeficiente R (Reducción sísmica) [Usa 1.0 para Elástico]", 1.0, 10.0, st.session_state.get("s_peru_R", 1.0), 0.5, key="s_peru_R")
            params = {"Z": Z, "S": S, "Tp": Tp, "Tl": Tl, "U": U, "R": R}

        elif "1225001" in norma_sel or "Bolivia" in norma_sel:
            st.write("**Norma: GBDS NB 1225001-2020 (Bolivia)**")
            zona_bol = st.selectbox("Zona Sísmica:", ["1 (Baja)", "2 (Media - ej. La Paz)", "3 (Alta)", "4 (Muy Alta)", "5 (Severa)"],
                                     index=["1 (Baja)","2 (Media - ej. La Paz)","3 (Alta)","4 (Muy Alta)","5 (Severa)"].index(st.session_state.get("s_bol_zona", "2 (Media - ej. La Paz)")),
                                     key="s_bol_zona")
            Z = [0.05, 0.10, 0.15, 0.20, 0.25][int(zona_bol.split(" ")[0])-1]
            suelo_bol = st.selectbox("Clase de Sitio", ["A (Roca dura)", "B (Roca)", "C (Suelo Muy Denso)", "D (Suelo Rígido)", "E (Suelo Blando)"],
                                      index=["A (Roca dura)","B (Roca)","C (Suelo Muy Denso)","D (Suelo Rígido)","E (Suelo Blando)"].index(st.session_state.get("s_bol_suelo", "D (Suelo Rígido)")),
                                      key="s_bol_suelo")
            # Factores simplificados (deberían ser más detallados)
            Fa_bol = {"A":0.9, "B":1.0, "C":1.1, "D":1.2, "E":1.4}
            Fa = Fa_bol[suelo_bol[0]]
            Tc = st.number_input("Período característico Tc [s]", 0.2, 2.0, 0.5, 0.05, key="s_bol_Tc")
            params = {"Z": Z, "Fa": Fa, "Tc": Tc}

        else:
            st.write("**Norma General / ASCE 7 (EE.UU.)**")
            Ss = st.number_input("S_S (Short Period Acc)", 0.1, 3.0, st.session_state.get("s_Ss", 1.0), 0.1, key="s_Ss")
            S1 = st.number_input("S_1 (1-sec Period Acc)", 0.05, 1.5, st.session_state.get("s_S1", 0.4), 0.1, key="s_S1")
            Fa = st.number_input("F_a (Corto Periodo)", 0.5, 3.0, st.session_state.get("s_Fa", 1.1), 0.1, key="s_Fa")
            Fv = st.number_input("F_v (Largo Periodo)", 0.5, 3.5, st.session_state.get("s_Fv", 1.6), 0.1, key="s_Fv")
            S_DS = (2.0/3.0) * Ss * Fa
            S_D1 = (2.0/3.0) * S1 * Fv
            TL = st.number_input("Período largo TL [s] (por defecto 4.0)", 2.0, 8.0, 4.0, 0.5, key="s_TL")
            params = {"S_DS": S_DS, "S_D1": S_D1, "TL": TL}
            st.caption(f"S_DS = {S_DS:.3f} g,  S_D1 = {S_D1:.3f} g")

        # Estructura (1 GDL)
        st.markdown("---")
        st.subheader(_t("🏢 Estructura (1 GDL)", "🏢 Structure (1 DOF)"))
        peso_W = st.number_input(_t("Peso de la estructura W [kN]", "Structure Weight W [kN]"), 10.0, 50000.0, st.session_state.get("s_W", 1000.0), 10.0, key="s_W")
        rigidez_k = st.number_input(_t("Rigidez Lateral k [kN/m]", "Lateral Stiffness k [kN/m]"), 100.0, 500000.0, st.session_state.get("s_k", 15000.0), 100.0, key="s_k")
        masa_m = peso_W / 9.81
        omega_n = math.sqrt(rigidez_k / masa_m)
        freq_n = omega_n / (2*math.pi)
        period_T = 1.0 / freq_n if freq_n != 0 else 0
        st.caption(f"Masa = {masa_m:.2f} t, ω_n = {omega_n:.2f} rad/s, f_n = {freq_n:.3f} Hz, T = {period_T:.3f} s")

    with col_right:
        # Calcular espectro
        T_domain, Sa_vals, params_texto, key_periods = compute_spectrum(norma_sel, params)
        # Obtener Sa en el período estructural
        if period_T <= T_domain[-1]:
            Sa_est = np.interp(period_T, T_domain, Sa_vals)
        else:
            Sa_est = 0.0
        # Cortante basal
        V_base = Sa_est * peso_W  # Sa en g → multiplicar por peso (kN)
        # Gráfico mejorado con Plotly
        fig_spec = go.Figure()
        fig_spec.add_trace(go.Scatter(x=T_domain, y=Sa_vals, mode='lines', name=_t('Espectro de Diseño', 'Design Spectrum'),
                                      line=dict(color='red', width=3)))
        # Línea del período estructural
        fig_spec.add_trace(go.Scatter(x=[period_T, period_T], y=[0, max(Sa_vals)*1.1], mode='lines',
                                      name=_t('Período Estructural', 'Structural Period'),
                                      line=dict(color='blue', width=2, dash='dash')))
        # Líneas verticales para T0, Tc, TL
        if key_periods['T0'] > 0:
            fig_spec.add_vline(x=key_periods['T0'], line_width=1, line_dash="dot", line_color="gray",
                               annotation_text=f"T₀ = {key_periods['T0']:.3f}s", annotation_position="top")
        if key_periods['Tc'] > 0:
            fig_spec.add_vline(x=key_periods['Tc'], line_width=1, line_dash="dot", line_color="gray",
                               annotation_text=f"Tc = {key_periods['Tc']:.3f}s", annotation_position="top")
        if key_periods['TL'] > 0:
            fig_spec.add_vline(x=key_periods['TL'], line_width=1, line_dash="dot", line_color="gray",
                               annotation_text=f"TL = {key_periods['TL']:.3f}s", annotation_position="top")
        # Anotación del punto del período estructural
        fig_spec.add_annotation(x=period_T, y=Sa_est, text=f"T = {period_T:.3f}s<br>Sa = {Sa_est:.3f}g",
                                showarrow=True, arrowhead=2, arrowcolor="blue", ax=30, ay=-30)
        # Layout mejorado
        fig_spec.update_layout(
            title=_t("Espectro de Respuesta de Diseño", "Design Response Spectrum"),
            xaxis_title=_t("Período T (s)", "Period T (s)"),
            yaxis_title=_t("Pseudo-Aceleración Sa (g)", "Pseudo-Acceleration Sa (g)"),
            template="plotly_dark",
            hovermode="x unified",
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
            margin=dict(l=40, r=40, t=60, b=40)
        )
        st.plotly_chart(fig_spec, use_container_width=True)

        # Mostrar resultados adicionales
        col_a, col_b = st.columns(2)
        col_a.metric(_t("Aceleración en T", "Acceleration at T"), f"{Sa_est:.3f} g")
        col_b.metric(_t("Cortante Basal V", "Base Shear V"), f"{V_base:.1f} kN")

        # Tabla de parámetros
        with st.expander(_t("📋 Parámetros adoptados", "📋 Adopted Parameters")):
            df_params = pd.DataFrame(params_texto, columns=["Descripción"])
            st.dataframe(df_params, use_container_width=True, hide_index=True)

# =============================================================================
# 3. EXPORTACIONES
# =============================================================================
st.markdown("---")
st.subheader(_t("💾 Exportación Integral", "💾 Comprehensive Export"))

tab_dxf, tab_doc, tab_xls = st.tabs([
    "📐 " + _t("Exportar Gráfico a DXF", "Export Graph to DXF"),
    "📄 " + _t("Memoria DOCX", "DOCX Report"),
    "💰 " + _t("Curva XLSX / Presupuesto APU", "XLSX Curve / APU Budget")
])

with tab_dxf:
    st.write(_t("Genera un archivo DXF con la curva espectral (coordenadas normalizadas en metros).", 
                "Generates a DXF file with the spectral curve (normalized coordinates in meters)."))
    col_dxf1, col_dxf2 = st.columns([1,3])
    with col_dxf1:
        escala_x = st.number_input(_t("Escala horizontal (m/s)", "Horizontal scale (m/s)"), 0.5, 10.0, 5.0, 0.5, key="dxf_scale_x")
        escala_y = st.number_input(_t("Escala vertical (m/g)", "Vertical scale (m/g)"), 0.5, 10.0, 2.0, 0.5, key="dxf_scale_y")
    with col_dxf2:
        if st.button(_t("Generar DXF", "Generate DXF")):
            doc_dxf = ezdxf.new('R2010')
            doc_dxf.units = ezdxf.units.M
            msp = doc_dxf.modelspace()
            # Capas
            for lay, col in [('EJE_X',7), ('EJE_Y',7), ('ESPECTRO',1), ('TEXTO',3)]:
                if lay not in doc_dxf.layers:
                    doc_dxf.layers.add(lay, color=col)
            # Ejes
            msp.add_lwpolyline([(0,0), (escala_x,0)], dxfattribs={'layer':'EJE_X'})
            msp.add_lwpolyline([(0,0), (0,escala_y)], dxfattribs={'layer':'EJE_Y'})
            # Curva
            points = []
            for t, sa in zip(T_domain, Sa_vals):
                x = t * escala_x / max(T_domain)
                y = sa * escala_y / max(Sa_vals) if max(Sa_vals)>0 else 0
                points.append((x, y))
            msp.add_lwpolyline(points, dxfattribs={'layer':'ESPECTRO', 'color':1})
            # Texto con parámetros
            texto = "\n".join(params_texto[:5])
            msp.add_text(texto, dxfattribs={'layer':'TEXTO','height':0.1,'insert':(0.1, escala_y-0.2)})
            out = io.BytesIO()
            doc_dxf.write(out)
            st.download_button(_t("📥 Descargar Espectro.dxf", "📥 Download Spectrum.dxf"), data=out.getvalue(),
                               file_name=f"Espectro_{norma_sel[:5]}.dxf", mime="application/dxf")

with tab_doc:
    if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
        doc = Document()
        doc.add_heading(_t(f"Estudio Sísmico – {norma_sel}", f"Seismic Study – {norma_sel}"), 0)
        doc.add_paragraph(_t(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}"))
        doc.add_paragraph(_t(f"Elaborado con: StructuroPro – Módulo Sísmico", "Prepared with: StructuroPro – Seismic Module"))

        doc.add_heading(_t("1. Parámetros del Espectro", "1. Spectrum Parameters"), level=1)
        for pt in params_texto:
            doc.add_paragraph(pt)
        doc.add_heading(_t("2. Períodos Característicos", "2. Characteristic Periods"), level=1)
        doc.add_paragraph(f"T₀ = {key_periods['T0']:.3f} s")
        doc.add_paragraph(f"Tc = {key_periods['Tc']:.3f} s")
        doc.add_paragraph(f"TL = {key_periods['TL']:.3f} s")
        doc.add_heading(_t("3. Estructura (1 GDL)", "3. Structure (1 DOF)"), level=1)
        doc.add_paragraph(f"{_t('Peso W:', 'Weight W:')} {peso_W:.1f} kN")
        doc.add_paragraph(f"{_t('Rigidez k:', 'Stiffness k:')} {rigidez_k:.1f} kN/m")
        doc.add_paragraph(f"{_t('Período T:', 'Period T:')} {period_T:.3f} s")
        doc.add_heading(_t("4. Resultados", "4. Results"), level=1)
        doc.add_paragraph(f"{_t('Aceleración espectral Sa(T):', 'Spectral acceleration Sa(T):')} {Sa_est:.3f} g")
        doc.add_paragraph(f"{_t('Cortante basal V:', 'Base shear V:')} {V_base:.1f} kN")
        doc.add_heading(_t("5. Curva Espectral", "5. Spectral Curve"), level=1)
        # Insertar la figura (usamos la misma que ya tenemos, pero la guardamos como imagen)
        import io
        import plotly.io as pio
        img_bytes = pio.to_image(fig_spec, format='png', width=800, height=500)
        img_buffer = io.BytesIO(img_bytes)
        doc.add_picture(img_buffer, width=Inches(5.5))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(_t("📥 Descargar Memoria DOCX", "📥 Download DOCX Report"), data=buf,
                           file_name="Estudio_Sismico.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_xls:
    st.write(_t("Exporta la curva espectral (T vs Sa) en formato Excel para su uso en software de análisis estructural.", 
                "Exports the spectral curve (T vs Sa) in Excel format for use in structural analysis software."))
    # Coordenadas
    df_spectrum = pd.DataFrame({"Periodo T (s)": T_domain, "Aceleración Sa (g)": Sa_vals})
    st.dataframe(df_spectrum.head(20), use_container_width=True)
    # APU
    if "apu_config" in st.session_state:
        apu = st.session_state.apu_config
        mon = apu["moneda"]
        honorarios = st.number_input(_t(f"Honorarios Globales por Estudio Sísmico [{mon}]", f"Global Fee for Seismic Study [{mon}]"), 500.0, 50000.0, 1500.0, 100.0, key="s_honorarios")
        output_excel = io.BytesIO()
        with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
            df_spectrum.to_excel(writer, index=False, sheet_name='Espectro')
            df_apu = pd.DataFrame({
                "Item": [_t("Estudio de Vulnerabilidad Sísmica y Espectro", "Seismic Vulnerability Study and Spectrum"),
                         _t("A.I.U. (Administración, Imprevistos, Utilidad)", "A.I.U. (Management, Contingency, Profit)")],
                "Cantidad": [1, 1],
                "Costo Unitario": [honorarios, honorarios * apu.get("pct_aui", 0.30)]
            })
            df_apu["Subtotal"] = df_apu["Cantidad"] * df_apu["Costo Unitario"]
            df_apu.to_excel(writer, index=False, sheet_name='APU_Sismica')
            workbook = writer.book
            worksheet = writer.sheets['APU_Sismica']
            money_fmt = workbook.add_format({'num_format': '#,##0.00'})
            worksheet.set_column('A:A', 40)
            worksheet.set_column('B:D', 15, money_fmt)
        output_excel.seek(0)
        st.download_button(_t("📥 Descargar Archivo Excel", "📥 Download Excel File"), data=output_excel,
                           file_name="Espectro_Sismico.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info(_t("💡 Ve a la página 'APU Mercado' para cargar los costos en vivo y activar el presupuesto.", 
                   "💡 Go to the 'Market APU' page to load live costs and enable the budget."))