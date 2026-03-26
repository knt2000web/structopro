import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
from docx import Document
from docx.shared import Pt, Cm, Inches
import ezdxf
from datetime import datetime
import io
import plotly.graph_objects as go

# ─────────────────────────────────────────────
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# ─────────────────────────────────────────────

st.set_page_config(page_title=_t("Suite Hormigón Armado", "Reinforced Concrete Suite"), layout="wide")
st.image(r"assets/beam_header_1773257190287.png", use_container_width=True)
st.title(_t("Suite de Diseño — Vigas y Losas", "Design Suite — Beams & Slabs"))
st.markdown(_t("Herramientas de diseño de concreto reforzado según **10 normativas internacionales**.", "Reinforced concrete design tools based on **10 international codes**."))

# ─────────────────────────────────────────────
# FUNCIONES DE DIBUJO PARA FIGURADO (MEJORADAS)
# ─────────────────────────────────────────────
def draw_longitudinal_bar(total_len_cm, straight_len_cm, hook_len_cm, bar_diam_mm):
    """
    Dibuja una barra longitudinal con ganchos de 90° en ambos extremos.
    total_len_cm : longitud total de la barra (incluyendo ganchos)
    straight_len_cm : longitud recta entre ganchos
    hook_len_cm : longitud de cada gancho (12db)
    bar_diam_mm : diámetro de la barra (para escala)
    """
    fig, ax = plt.subplots(figsize=(max(6, total_len_cm/20), 2))
    ax.set_aspect('equal')
    # Línea central
    ax.plot([0, straight_len_cm], [0, 0], 'k-', linewidth=2)
    # Gancho izquierdo (90° hacia arriba)
    ax.plot([0, 0], [0, hook_len_cm], 'k-', linewidth=2)
    # Gancho derecho (90° hacia abajo)
    ax.plot([straight_len_cm, straight_len_cm], [0, -hook_len_cm], 'k-', linewidth=2)
    # Cotas
    ax.annotate(f"{straight_len_cm:.0f} cm", xy=(straight_len_cm/2, 0.3), ha='center', fontsize=8)
    ax.annotate(f"Gancho 12db = {hook_len_cm:.0f} cm", xy=(0, hook_len_cm/2), ha='right', fontsize=8)
    ax.annotate(f"Gancho 12db", xy=(straight_len_cm, -hook_len_cm/2), ha='left', fontsize=8)
    ax.set_xlim(-hook_len_cm*0.2, straight_len_cm + hook_len_cm*0.2)
    ax.set_ylim(-hook_len_cm*1.2, hook_len_cm*1.2)
    ax.axis('off')
    ax.set_title(f"Varilla longitudinal - Ø{bar_diam_mm:.0f} mm - Longitud total {total_len_cm:.0f} cm", fontsize=9)
    return fig

def draw_stirrup_beam(b_cm, h_cm, hook_len_cm, bar_diam_mm):
    """
    Dibuja un estribo rectangular con ganchos de 135° en una esquina.
    b_cm, h_cm : dimensiones interiores del estribo (medidas entre caras internas)
    hook_len_cm : longitud de proyección del gancho (aprox. 6db o 12db, solo visual)
    bar_diam_mm : diámetro de la barra (para escala)
    """
    fig, ax = plt.subplots(figsize=(max(5, b_cm/15), max(5, h_cm/15)))
    ax.set_aspect('equal')
    # Rectángulo exterior (interior real)
    x0, y0 = 0, 0
    ax.plot([x0, x0+b_cm], [y0, y0], 'k-', linewidth=2)          # base inferior
    ax.plot([x0+b_cm, x0+b_cm], [y0, y0+h_cm], 'k-', linewidth=2) # lado derecho
    ax.plot([x0+b_cm, x0], [y0+h_cm, y0+h_cm], 'k-', linewidth=2) # base superior
    ax.plot([x0, x0], [y0+h_cm, y0], 'k-', linewidth=2)          # lado izquierdo
    # Gancho de 135° en esquina inferior izquierda (dos segmentos)
    hook_x1 = x0 - hook_len_cm * 0.6
    hook_y1 = y0 - hook_len_cm * 0.6
    ax.plot([x0, hook_x1], [y0, hook_y1], 'k-', linewidth=2)
    hook_x2 = hook_x1 - hook_len_cm * 0.4
    hook_y2 = hook_y1 - hook_len_cm * 0.2
    ax.plot([hook_x1, hook_x2], [hook_y1, hook_y2], 'k-', linewidth=2)
    # Cotas
    ax.annotate(f"{b_cm:.0f} cm", xy=(b_cm/2, -0.5), ha='center', fontsize=8)
    ax.annotate(f"{h_cm:.0f} cm", xy=(-0.5, h_cm/2), ha='right', va='center', fontsize=8)
    ax.annotate(f"Gancho 135°", xy=(x0 - hook_len_cm*0.5, y0 - hook_len_cm*0.7), ha='right', fontsize=8)
    ax.set_xlim(-hook_len_cm*1.2, b_cm + hook_len_cm*0.5)
    ax.set_ylim(-hook_len_cm*1.2, h_cm + hook_len_cm*0.5)
    ax.axis('off')
    ax.set_title(f"Estribo - Ø{bar_diam_mm:.0f} mm - Perímetro {2*(b_cm+h_cm):.0f} cm", fontsize=9)
    return fig

# ─────────────────────────────────────────────
# UNIDADES DE SALIDA
# ─────────────────────────────────────────────
st.sidebar.header(_t("Unidades de salida", "Output units"))
unidades_salida = st.sidebar.radio("Unidades de fuerza/momento:", ["kiloNewtons (kN, kN·m)", "Toneladas fuerza (tonf, tonf·m)"], key="v_output_units")
if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
    factor_fuerza = 0.1019716
    unidad_fuerza = "tonf"
    unidad_mom    = "tonf·m"
else:
    factor_fuerza = 1.0
    unidad_fuerza = "kN"
    unidad_mom    = "kN·m"

# ─────────────────────────────────────────────
# APU CON ENTRADA DIRECTA (GLOBAL)
# ─────────────────────────────────────────────
with st.expander("💰 APU – Precios en vivo (materiales y mano de obra)", expanded=False):
    st.markdown("Ingrese los precios unitarios de los materiales y mano de obra para calcular el costo total de las estructuras.")
    with st.form(key="apu_form_global"):
        moneda = st.text_input("Moneda (ej. COP, USD)", value=st.session_state.get("apu_moneda_global", "COP"))
        col1a, col2a = st.columns(2)
        with col1a:
            precio_cemento = st.number_input("Precio por bulto de cemento", value=st.session_state.get("apu_cemento_global", 28000.0), step=1000.0, format="%.2f")
            precio_acero = st.number_input("Precio por kg de acero", value=st.session_state.get("apu_acero_global", 7500.0), step=100.0, format="%.2f")
            precio_arena = st.number_input("Precio por m³ de arena", value=st.session_state.get("apu_arena_global", 120000.0), step=5000.0, format="%.2f")
            precio_grava = st.number_input("Precio por m³ de grava", value=st.session_state.get("apu_grava_global", 130000.0), step=5000.0, format="%.2f")
        with col2a:
            precio_mo = st.number_input("Costo mano de obra (día)", value=st.session_state.get("apu_mo_global", 70000.0), step=5000.0, format="%.2f")
            pct_herramienta = st.number_input("% Herramienta menor (sobre MO)", value=st.session_state.get("apu_herramienta_global", 5.0), step=1.0, format="%.1f") / 100.0
            pct_aui = st.number_input("% A.I.U. (sobre costo directo)", value=st.session_state.get("apu_aui_global", 30.0), step=5.0, format="%.1f") / 100.0
            pct_util = st.number_input("% Utilidad (sobre costo directo)", value=st.session_state.get("apu_util_global", 5.0), step=1.0, format="%.1f") / 100.0
            iva = st.number_input("IVA (%) sobre utilidad", value=st.session_state.get("apu_iva_global", 19.0), step=1.0, format="%.1f") / 100.0
        submitted = st.form_submit_button("Guardar precios")
        if submitted:
            st.session_state.apu_config = {
                "moneda": moneda,
                "cemento": precio_cemento,
                "acero": precio_acero,
                "arena": precio_arena,
                "grava": precio_grava,
                "costo_dia_mo": precio_mo,
                "pct_herramienta": pct_herramienta,
                "pct_aui": pct_aui,
                "pct_util": pct_util,
                "iva": iva
            }
            st.success("Precios guardados. Ahora se mostrarán los presupuestos en las secciones de cantidades.")
            st.rerun()

# ─────────────────────────────────────────────
# PIE DE PÁGINA / DERECHOS RESERVADOS (en sidebar)
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    Ing. Msc. César Augusto Giraldo Chaparro<br><br>
    <i>⚠️ Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════
# CODES DICT (COMPLETO)
# ══════════════════════════════════════════
CODES = {
    "NSR-10 (Colombia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["DMI — Disipación Mínima", "DMO — Disipación Moderada", "DES — Disipación Especial"],
        "ref": "NSR-10 Título C (C.9, C.11, C.21)",
        "bag_kg": 50.0,
    },
    "ACI 318-25 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["OMF (SDC A–B)", "IMF (SDC C)", "SMF (SDC D–F)"],
        "ref": "ACI 318-25 (Sections 9, 22, 25)",
        "bag_kg": 42.6,
    },
    "ACI 318-19 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["OMF (SDC A–B)", "IMF (SDC C)", "SMF (SDC D–F)"],
        "ref": "ACI 318-19 (Sections 9, 22, 25)",
        "bag_kg": 42.6,
    },
    "ACI 318-14 (EE.UU.)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["OMF (SDC A–B)", "IMF (SDC C)", "SMF (SDC D–F)"],
        "ref": "ACI 318-14 (Sections 9, 22, 25)",
        "bag_kg": 42.6,
    },
    "NEC-SE-HM (Ecuador)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["GS — Grado Reducido", "GM — Grado Moderado", "GA — Grado Alto"],
        "ref": "NEC-SE-HM Ecuador (Cap. 4)",
        "bag_kg": 50.0,
    },
    "E.060 (Perú)": {
        "phi_flex": 0.90, "phi_shear": 0.85, "phi_comp": 0.70,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["PO — Pórtico Ordinario (Z1–Z2)", "PE — Pórtico Especial (Z3–Z4)"],
        "ref": "Norma E.060 Perú (Arts. 9, 11, 21)",
        "bag_kg": 42.5,
    },
    "NTC-EM (México)": {
        "phi_flex": 0.85, "phi_shear": 0.80, "phi_comp": 0.70,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["MDL — Ductilidad Limitada", "MROD — Ductilidad Ordinaria", "MRLE — Ductilidad Alta"],
        "ref": "NTC-EM México 2017 (Cap. 2, 4)",
        "bag_kg": 50.0,
    },
    "COVENIN 1753-2006 (Venezuela)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.70,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["PO — Pórtico Ordinario", "PM — Pórtico Moderado", "PE — Pórtico Especial"],
        "ref": "COVENIN 1753-2006 Venezuela",
        "bag_kg": 42.5,
    },
    "NB 1225001-2020 (Bolivia)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["DO — Diseño Ordinario", "DE — Diseño Especial Sísmico"],
        "ref": "NB 1225001-2020 Bolivia (ACI 318-19)",
        "bag_kg": 50.0,
    },
    "CIRSOC 201-2025 (Argentina)": {
        "phi_flex": 0.90, "phi_shear": 0.75, "phi_comp": 0.65,
        "lambda": 1.0, "beta1_hi": 0.65, "eps_cu": 0.003,
        "rho_min_factor": 1.4,
        "seismic_levels": ["GE — Grado Estándar", "GM — Ductilidad Moderada", "GA — Ductilidad Alta"],
        "ref": "CIRSOC 201-2025 Argentina (basada en ACI 318-19)",
        "bag_kg": 50.0,
    },
}

# Rebar tables
REBAR_US = {"#3 (Ø9.5mm)":0.71,"#4 (Ø12.7mm)":1.29,"#5 (Ø15.9mm)":1.99,"#6 (Ø19.1mm)":2.84,"#7 (Ø22.2mm)":3.87,"#8 (Ø25.4mm)":5.10,"#9 (Ø28.7mm)":6.45,"#10 (Ø32.3mm)":7.92}
REBAR_MM = {"8mm":0.503,"10mm":0.785,"12mm":1.131,"14mm":1.539,"16mm":2.011,"18mm":2.545,"20mm":3.142,"22mm":3.801,"25mm":4.909,"28mm":6.158,"32mm":8.042}
DIAM_US = {"#3 (Ø9.5mm)":9.53,"#4 (Ø12.7mm)":12.7,"#5 (Ø15.9mm)":15.88,"#6 (Ø19.1mm)":19.05,"#7 (Ø22.2mm)":22.23,"#8 (Ø25.4mm)":25.4,"#9 (Ø28.7mm)":28.65,"#10 (Ø32.3mm)":32.26}
DIAM_MM = {"8mm":8,"10mm":10,"12mm":12,"14mm":14,"16mm":16,"18mm":18,"20mm":20,"22mm":22,"25mm":25,"28mm":28,"32mm":32}

# Funciones auxiliares
def get_beta1(fc):
    if fc <= 28: return 0.85
    b = 0.85 - 0.05*(fc-28)/7.0
    return max(b, 0.65)

def get_rho_min(fc, fy, norm):
    return max(0.25*math.sqrt(fc)/fy, 1.4/fy)

def get_rho_max_beam(fc, fy, beta1, eps_cu=0.003, eps_t_min=0.005):
    rho_bal = (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+fy/200000))
    rho_max = (0.85*fc*beta1/fy)*(eps_cu/(eps_cu+eps_t_min))
    return rho_max

def mix_for_fc(fc):
    table = [
        (14,250,205,810,1060),(17,290,200,780,1060),(21,350,193,720,1060),
        (25,395,193,680,1020),(28,430,190,640,1000),(35,530,185,580,960),
        (42,620,180,520,910),(56,740,175,450,850),
    ]
    if fc <= table[0][0]: return table[0][1:]
    if fc >= table[-1][0]: return table[-1][1:]
    for i in range(len(table)-1):
        lo,hi = table[i],table[i+1]
        if lo[0] <= fc <= hi[0]:
            t = (fc-lo[0])/(hi[0]-lo[0])
            return tuple(lo[j]+t*(hi[j]-lo[j]) for j in range(1,5))
    return table[-1][1:]

def sec_dark_fig(w, h, title=""):
    fig, ax = plt.subplots(figsize=(max(3,w/h*3), 3))
    fig.patch.set_facecolor('#1a1a2e')
    ax.set_facecolor('#1a1a2e')
    ax.add_patch(patches.Rectangle((0,0),w,h,linewidth=2,edgecolor='white',facecolor='#4a4a6a'))
    ax.set_xlim(-w*0.15, w*1.15); ax.set_ylim(-h*0.15, h*1.15)
    ax.axis('off')
    ax.set_title(title, color='white', fontsize=8)
    return fig, ax

def qty_table(rows):
    st.dataframe(pd.DataFrame(rows, columns=["Concepto","Valor"]), use_container_width=True, hide_index=True)

# ══════════════════════════════════════════
# GLOBAL SIDEBAR
# ══════════════════════════════════════════
st.sidebar.header(_t("🌍 Norma de Diseño", "🌍 Design Code"))
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = list(CODES.keys())[0]
norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagcdn.com/24x18/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>',
    unsafe_allow_html=True
)
code = CODES[norma_sel]
nivel_sis = st.sidebar.selectbox(_t("Nivel Sísmico:", "Seismic Level:"), code["seismic_levels"],
                                 index=code["seismic_levels"].index(st.session_state.v_nivel_sis) if "v_nivel_sis" in st.session_state and st.session_state.v_nivel_sis in code["seismic_levels"] else 0,
                                 key="v_nivel_sis")
st.sidebar.markdown(f"📖 `{code['ref']}`")
st.sidebar.markdown(f"**φ flex:** {code['phi_flex']} | **φ cort:** {code['phi_shear']}")

st.sidebar.header(_t("⚙️ Materiales Globales", "⚙️ Global Materials"))
fc_unit = st.sidebar.radio(_t("Unidad f'c:", "f'c Unit:"), ["MPa","PSI","kg/cm²"], horizontal=True, key="v_fc_unit")
if fc_unit == "PSI":
    psi_options = ["2500","3000","3500","4000","4500","5000"]
    psi_v = st.sidebar.selectbox("f'c [PSI]:", psi_options,
                                 index=psi_options.index(st.session_state.v_fc_psi) if "v_fc_psi" in st.session_state and st.session_state.v_fc_psi in psi_options else 1,
                                 key="v_fc_psi")
    fc = float(psi_v)*0.00689476
    st.sidebar.info(f"f'c = {psi_v} PSI → **{fc:.2f} MPa**")
elif fc_unit == "kg/cm²":
    kg_options = ["175","210","250","280","350","420"]
    kg_v = st.sidebar.selectbox("f'c [kg/cm²]:", kg_options,
                                index=kg_options.index(st.session_state.v_fc_kgcm2) if "v_fc_kgcm2" in st.session_state and st.session_state.v_fc_kgcm2 in kg_options else 1,
                                key="v_fc_kgcm2")
    fc = float(kg_v)/10.1972
    st.sidebar.info(f"f'c = {kg_v} kg/cm² → **{fc:.2f} MPa**")
else:
    fc = st.sidebar.number_input("f'c [MPa]:", 15.0, 80.0, st.session_state.get("v_fc_mpa", 21.0), 1.0, key="v_fc_mpa")

fy = st.sidebar.number_input("fy [MPa]:", 200.0, 500.0, st.session_state.get("v_fy", 420.0), 10.0, key="v_fy")
Es = 200000.0
Ec = 4700*math.sqrt(fc)
beta1 = get_beta1(fc)
rho_min = get_rho_min(fc, fy, norma_sel)
rho_max = get_rho_max_beam(fc, fy, beta1)

bar_sys = st.sidebar.radio("Sistema Varillas:", ["Pulgadas (# US)","Milímetros (mm)"], horizontal=True, key="v_bar_sys")
rebar_dict = REBAR_US if "Pulgadas" in bar_sys else REBAR_MM
diam_dict  = DIAM_US  if "Pulgadas" in bar_sys else DIAM_MM

phi_f = code["phi_flex"]
phi_v = code["phi_shear"]
lam   = code["lambda"]
bag_kg = code["bag_kg"]

st.sidebar.markdown("---")
st.sidebar.caption(f"Ec = {Ec:.0f} MPa  |  β₁ = {beta1:.3f}  |  f'c = {fc:.2f} MPa")

# ══════════════════════════════════════════
# 1. TABLA DE ACERO
# ══════════════════════════════════════════
with st.expander(_t("🔩 Tabla de Secciones de Acero de Refuerzo", "🔩 Rebar Area Table")):
    st.markdown(_t(f"**Referencia:** {code['ref']}", f"**Reference:** {code['ref']}"))
    rows_us, rows_mm = [], []
    for k,a in REBAR_US.items():
        d_bar = DIAM_US[k]
        rows_us.append({"Barra":k,"Ø (mm)":f"{d_bar:.2f}","Área (cm²)":f"{a:.3f}","Peso (kg/m)":f"{a * 0.785:.3f}"})
    for k,a in REBAR_MM.items():
        d_bar = DIAM_MM[k]
        rows_us_mm = {"Barra (SI)":k,"Ø (mm)":f"{d_bar:.0f}","Área (cm²)":f"{a:.4f}","Peso (kg/m)":f"{a * 0.785:.3f}"}
        rows_mm.append(rows_us_mm)
    c1,c2 = st.columns(2)
    with c1:
        st.markdown("##### Sistema US (pulgadas)")
        st.dataframe(pd.DataFrame(rows_us), use_container_width=True, hide_index=True)
    with c2:
        st.markdown("##### Sistema SI (milímetros)")
        st.dataframe(pd.DataFrame(rows_mm), use_container_width=True, hide_index=True)

# ══════════════════════════════════════════
# 2. VIGA RECTANGULAR — FLEXIÓN
# ══════════════════════════════════════════
with st.expander(_t("📐 Diseño a Flexión — Viga Rectangular", "📐 Flexural Design — Rectangular Beam")):
    st.markdown(_t(f"**Método (sección simplemente reforzada)** | Norma: `{code['ref']}`", f"**Method (singly reinforced section)** | Code: `{code['ref']}`"))
    st.info(_t("📺 **Modo de uso:** Ingresa la base, altura y recubrimiento de la viga. Añade el momento flector (Mu) a soportar. Luego selecciona el diámetro de varilla y la App te dirá si cumple flexión, calculará la cantidad requerida de acero, y generará las cantidades y precios de todo el pórtico.", "📺 **How to use:** Enter beam base, height and cover. Add ultimate moment (Mu). Then select rebar diameter and the App will check flexure, calculate required steel, quantities and prices."))
    c1,c2 = st.columns(2)
    with c1:
        b_vr = st.number_input("Ancho b [cm]", 15.0, 150.0, st.session_state.get("vr_b", 30.0), 5.0, key="vr_b")
        h_vr = st.number_input("Alto h [cm]", 20.0, 200.0, st.session_state.get("vr_h", 50.0), 5.0, key="vr_h")
        dp_vr = st.number_input("Recubrim. d' [cm]", 2.0, 15.0, st.session_state.get("vr_dp", 5.0), 0.5, key="vr_dp")
        Mu_vr = st.number_input(f"Momento último Mu [{unidad_mom}]", 0.1, 10000.0, st.session_state.get("vr_mu", 80.0), 5.0, key="vr_mu")
    with c2:
        L_vr = st.number_input("Longitud viga [m]", 1.0, 30.0, st.session_state.get("vr_L", 5.0), 0.5, key="vr_L")
        varillas_vr = list(rebar_dict.keys())
        # Mínimo práctico: #4 (12.7mm) en US  |  12mm en SI  — el #3/8mm se reserva para estribos
        _def_vr = "#4 (Ø12.7mm)" if "Pulgadas" in bar_sys else "12mm"
        _def_idx_vr = varillas_vr.index(_def_vr) if _def_vr in varillas_vr else 1
        bar_vr = st.selectbox(
            "Varilla longitudinal (mín. recomendado #4 / 12mm — el #3 se usa solo para estribos):",
            varillas_vr,
            index=varillas_vr.index(st.session_state.vr_bar) if "vr_bar" in st.session_state and st.session_state.vr_bar in varillas_vr else _def_idx_vr,
            key="vr_bar")
        Ab_vr = rebar_dict[bar_vr]; db_vr = diam_dict[bar_vr]

    d_vr = h_vr - dp_vr
    d_mm = d_vr*10; b_mm = b_vr*10
    # Convertir Mu a kN·m si es necesario
    if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
        Mu_vr_kN = Mu_vr / factor_fuerza
    else:
        Mu_vr_kN = Mu_vr
    Mu_Nmm = Mu_vr_kN * 1e6
    if d_mm > 0 and b_mm > 0:
        Rn = Mu_Nmm / (phi_f * b_mm * d_mm**2)
        disc = 1 - 2*Rn/(0.85*fc)
        if disc < 0:
            st.error("❌ Sección insuficiente – aumente b o h")
        else:
            rho_calc = (0.85*fc/fy)*(1 - math.sqrt(disc))
            rho_use = max(rho_calc, rho_min)
            As_req_cm2 = rho_use * b_vr * d_vr
            n_bars = math.ceil(As_req_cm2 / Ab_vr)
            As_prov = n_bars * Ab_vr
            rho_prov = As_prov / (b_vr * d_vr)

            a_mm = As_prov*100*fy/(0.85*fc*b_mm)
            phi_Mn_kNm = phi_f * As_prov*100*fy*(d_mm - a_mm/2)/1e6
            ok_flex = phi_Mn_kNm >= Mu_vr_kN
            ok_rho_min = rho_prov >= rho_min
            ok_rho_max = rho_prov <= rho_max

            tab_r, tab_s, tab_3d, tab_q = st.tabs(["📊 Resultados","🔲 Sección 2D","🧊 Visualización 3D","📦 Cantidades"])
            with tab_r:
                st.markdown(f"**Factor de reducción φ = {phi_f}** (flexión) | Norma: `{code['ref']}`")
                st.markdown("""**Verificación fundamental:** La resistencia a flexión provista **φMn** debe ser mayor o igual al momento último demandado **Mu**.
> φMn ≥ Mu  ✔ (la viga resiste sin colapsar)""")
                st.latex(r"\phi M_n = \phi \cdot A_s \cdot f_y \left(d - \frac{a}{2}\right)")
                rows = [
                    ("📐 b × h — Base y altura de la viga", f"{b_vr:.0f} × {h_vr:.0f} cm"),
                    ("📏 d — Peralte efectivo (altura hasta centroide del acero)", f"{d_vr:.1f} cm"),
                    ("🔢 Rn — Resistencia unitaria requerida (Mu / φ·b·d²)", f"{Rn:.3f} MPa"),
                    ("📊 ρ calculado — Cuantía de acero que necesita la sección", f"{rho_calc*100:.4f}%"),
                    ("⬇ ρ mínimo — Cuantía mínima exigida por la norma (evita falla frágil)", f"{rho_min*100:.4f}%"),
                    ("⬆ ρ máximo — Cuantía máxima (garantiza falla dúctil con aviso)", f"{rho_max*100:.4f}%"),
                    ("🔩 As requerido — Área de acero necesaria para resistir Mu", f"{As_req_cm2:.3f} cm²"),
                    (f"🔩 Varillas seleccionadas ({bar_vr}) — Cantidad y área provista", f"{n_bars} barras → As provisto = {As_prov:.3f} cm²"),
                    ("📦 a — Profundidad del bloque de compresión equivalente (Whitney)", f"{a_mm:.1f} mm"),
                    (f"✅ φMn — Momento resistente provisto [{unidad_mom}]", f"{phi_Mn_kNm*factor_fuerza:.2f}"),
                    (f"🎯 Mu — Momento último demandado (carga de diseño) [{unidad_mom}]", f"{Mu_vr:.2f}"),
                    ("📋 Verificación Flexión (φMn ≥ Mu)" if ok_flex else "❌ Verificación Flexión (φMn < Mu)",
                     "✅ CUMPLE — La viga resiste el momento de diseño" if ok_flex else f"❌ DEFICIENTE — φMn={phi_Mn_kNm:.2f} < Mu={Mu_vr_kN:.2f} → Aumente sección o acero"),
                    ("📋 Cuantía mínima (ρ ≥ ρ_min)" if ok_rho_min else "❌ Cuantía mínima (ρ < ρ_min)",
                     "✅ CUMPLE" if ok_rho_min else "❌ NO CUMPLE — Aumente el área de acero"),
                    ("📋 Cuantía máxima (ρ ≤ ρ_max)" if ok_rho_max else "❌ Cuantía máxima (ρ > ρ_max)",
                     "✅ CUMPLE" if ok_rho_max else "❌ EXCEDE MÁXIMO — Sección sobrearmada, amplíe la sección"),
                ]
                qty_table(rows)
                if ok_flex and ok_rho_min and ok_rho_max:
                    st.success(f"✅ Diseño Aprobado: φMn = {phi_Mn_kNm*factor_fuerza:.2f} {unidad_mom}  ≥  Mu = {Mu_vr:.2f} {unidad_mom}")
                else:
                    st.error("❌ Diseño No Aprobado — φMn < Mu o cuantía fuera de rango → Revisar sección o aumentar acero")
                st.info("💡 **¿El acero calculado es inferior o superior?** Si el Mu ingresado viene de una combinación con **momento positivo** (vano central), el acero corresponde al **refuerzo inferior** (zona en tensión debajo). Si Mu viene de un **momento negativo** (apoyo o empotramiento), el acero es el **refuerzo superior**.")

            with tab_s:
                fig, ax = sec_dark_fig(b_vr, h_vr, f"Sección {b_vr:.0f}×{h_vr:.0f} cm")
                recub = max(dp_vr - db_vr/20, 0.5)
                r_bar = db_vr/20
                ax.add_patch(patches.Rectangle((recub,recub),b_vr-2*recub,h_vr-2*recub,linewidth=1.5,edgecolor='#00d4ff',facecolor='none',linestyle='--'))
                xs = [b_vr/2] if n_bars == 1 else np.linspace(dp_vr, b_vr-dp_vr, n_bars) if n_bars>1 else [b_vr/2]
                for x in xs[:n_bars]:
                    ax.add_patch(plt.Circle((x, dp_vr), r_bar, color='#ff6b35', zorder=5))
                ax.annotate('',xy=(b_vr,-0.8*h_vr/h_vr),xytext=(0,-0.8*h_vr/h_vr),arrowprops=dict(arrowstyle='<->',color='white'))
                ax.text(b_vr/2,-h_vr*0.12,f"b={b_vr:.0f}cm",ha='center',va='top',color='white',fontsize=7)
                ax.annotate('',xy=(-0.8,h_vr),xytext=(-0.8,0),arrowprops=dict(arrowstyle='<->',color='white'))
                ax.text(-h_vr*0.15,h_vr/2,f"h={h_vr:.0f}cm",ha='right',va='center',color='white',fontsize=7,rotation=90)
                st.pyplot(fig)
                st.caption(f"{n_bars} varillas {bar_vr} en tensión | As={As_prov:.3f} cm²")

            with tab_3d:
                st.subheader("🧊 Visualización 3D de Viga Rectangular")
                fig3d = go.Figure()
                L_mm_3d = L_vr * 100
                x_c = [-b_vr/2, b_vr/2, b_vr/2, -b_vr/2, -b_vr/2, b_vr/2, b_vr/2, -b_vr/2]
                y_c = [0, 0, h_vr, h_vr, 0, 0, h_vr, h_vr]
                z_c = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
                fig3d.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
                diam_reb_cm = db_vr / 10.0
                line_width = max(4, diam_reb_cm * 3)
                xs = np.linspace(-b_vr/2 + dp_vr, b_vr/2 - dp_vr, max(n_bars, 2)) if n_bars > 1 else [0]
                for idx, x_pos in enumerate(xs[:n_bars]):
                    fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[dp_vr, dp_vr], z=[0, L_mm_3d],
                                                mode='lines', line=dict(color='darkred', width=line_width),
                                                name=f'Varilla {bar_vr}', showlegend=(idx==0)))
                tie_color = 'cornflowerblue'
                tie_width = max(2, (9.5/10.0) * 3)
                sep_ties = st.slider("Separación Estribos (cm)", 5, 50, int(st.session_state.get('cv_s_diseno', 15)), 1, key="vr_sep_tie")
                tx = [-b_vr/2 + dp_vr/2, b_vr/2 - dp_vr/2, b_vr/2 - dp_vr/2, -b_vr/2 + dp_vr/2, -b_vr/2 + dp_vr/2]
                ty = [dp_vr/2, dp_vr/2, h_vr - dp_vr/2, h_vr - dp_vr/2, dp_vr/2]
                L_cm = int(L_mm_3d)
                tx_all, ty_all, tz_all = [], [], []
                for zt in range(15, L_cm, sep_ties):
                    tx_all.extend(tx + [None])
                    ty_all.extend(ty + [None])
                    tz_all.extend([zt]*5 + [None])
                fig3d.add_trace(go.Scatter3d(x=tx_all, y=ty_all, z=tz_all, mode='lines', 
                                             line=dict(color=tie_color, width=tie_width), name='Estribos', showlegend=True))
                fig3d.update_layout(scene=dict(aspectmode='data', xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                                    margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable')
                st.plotly_chart(fig3d, use_container_width=True)

            with tab_q:
                vol_horm = b_vr/100*h_vr/100*L_vr
                peso_long = As_prov * L_vr * 0.785
                m = mix_for_fc(fc)
                bags = m[0]/bag_kg * vol_horm
                rows_q = [
                    ("Concreto (b×h×L)", f"{vol_horm:.4f} m³"),
                    (f"Acero longitudinal ({n_bars} barras)", f"{peso_long:.2f} kg"),
                    (f"Cemento ({bag_kg:.0f} kg/bulto, f'c={fc:.1f} MPa)", f"{bags:.1f} bultos = {m[0]*vol_horm:.0f} kg"),
                    ("Arena", f"{m[2]*vol_horm:.0f} kg"),
                    ("Grava", f"{m[3]*vol_horm:.0f} kg"),
                    ("Agua", f"{m[1]*vol_horm:.0f} L"),
                    ("Referencia", code["ref"]),
                ]
                qty_table(rows_q)

                if "apu_config" in st.session_state:
                    apu = st.session_state.apu_config
                    mon = apu.get("moneda", "$")
                    st.markdown("---")
                    st.success("✅ **Precios actualizados del mercado aplicados (APU).**")
                    c_cem = bags * apu.get("cemento", 0)
                    c_ace = peso_long * apu.get("acero", 0)
                    vol_arena_m3 = (m[2]*vol_horm)/1600
                    vol_grava_m3 = (m[3]*vol_horm)/1600
                    c_are = vol_arena_m3 * apu.get("arena", 0)
                    c_gra = vol_grava_m3 * apu.get("grava", 0)
                    total_mat = c_cem + c_ace + c_are + c_gra
                    total_dias_mo = (peso_long * 0.04) + (vol_horm * 0.4)
                    costo_mo = total_dias_mo * apu.get("costo_dia_mo", 70000)
                    costo_directo = total_mat + costo_mo
                    herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
                    aiu = costo_directo * apu.get("pct_aui", 0.30)
                    utilidad = costo_directo * apu.get("pct_util", 0.05)
                    iva_v = utilidad * apu.get("iva", 0.19)
                    gran_total = costo_directo + herramienta + aiu + iva_v
                    st.metric(f"💎 Gran Total Proyecto ({mon})", f"{gran_total:,.2f}")
                    st.caption("Incluye Materiales, Mano de Obra, Herramienta, A.I.U. e IVA s/Utilidad.")

                    st.markdown("#### 📏 Despiece de Acero (Cantidades y Costos)")
                    despiece_rows = [
                        {"Elemento": "Longitudinal (Rect)", "Barra": bar_vr, "Cant": n_bars, "L. Unitaria (m)": L_vr, "Peso Total (kg)": peso_long, "Costo Total": c_ace}
                    ]
                    df_desp = pd.DataFrame(despiece_rows)
                    st.dataframe(df_desp.style.format({"Peso Total (kg)": "{:.2f}", "Costo Total": "{:,.2f}", "L. Unitaria (m)": "{:.2f}"}), use_container_width=True, hide_index=True)

                    fig_chart = go.Figure()
                    fig_chart.add_trace(go.Bar(name='Cantidad', x=["Cemento","Acero","Arena","Grava"], y=[bags, peso_long, vol_arena_m3, vol_grava_m3], yaxis='y', marker_color='#4a4a6a'))
                    fig_chart.add_trace(go.Bar(name='Costo', x=["Cemento","Acero","Arena","Grava"], y=[c_cem, c_ace, c_are, c_gra], yaxis='y2', marker_color='#ff6b35'))
                    fig_chart.update_layout(
                        title=dict(text='Cantidades vs Costos', font=dict(color='white')),
                        yaxis=dict(title=dict(text='Cantidades', font=dict(color='#4a4a6a')), tickfont=dict(color='#4a4a6a')),
                        yaxis2=dict(title=dict(text='Costo', font=dict(color='#ff6b35')), tickfont=dict(color='#ff6b35'), overlaying='y', side='right'),
                        barmode='group', paper_bgcolor='#1a1a2e', plot_bgcolor='#1a1a2e',
                        legend=dict(x=0.01, y=0.99, bgcolor='rgba(255,255,255,0.1)', font=dict(color='white')),
                        margin=dict(l=0, r=0, t=40, b=0), height=300
                    )
                    st.plotly_chart(fig_chart, use_container_width=True)

                    with st.expander("📐 Dibujo de Figurado para Taller", expanded=False):
                        st.markdown("A continuación se muestran las formas reales de las barras para facilitar el figurado.")
                        hook_len_cm = 12 * db_vr / 10
                        straight_len_cm = L_vr * 100 - 2 * hook_len_cm
                        fig_l1 = draw_longitudinal_bar(L_vr*100, straight_len_cm, hook_len_cm, db_vr)
                        st.pyplot(fig_l1)
                        recub_est = max(dp_vr, 2.5)
                        inside_b = b_vr - 2*recub_est
                        inside_h = h_vr - 2*recub_est
                        hook_len_est = 12 * 9.5 / 10  # aprox. 12db para estribo (visual)
                        fig_e1 = draw_stirrup_beam(inside_b, inside_h, hook_len_est, 9.5)
                        st.pyplot(fig_e1)
                        st.caption("Nota: Los ganchos de estribos son de 135° en la práctica. En el dibujo se representa de forma esquemática con líneas inclinadas.")
                else:
                    st.info("💡 Configure los precios en el expander '💰 APU – Precios en vivo' al inicio de la página para ver el presupuesto.")

                # MEMORIA DOCX para Viga Rectangular
                if st.button("📄 Generar Memoria DOCX (Viga Rectangular)"):
                    doc = Document()
                    doc.add_heading(f"Memoria de Cálculo Viga: {b_vr:.0f}x{h_vr:.0f} cm", 0)
                    doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                    doc.add_heading("1. Materiales", level=1)
                    doc.add_paragraph(f"Concreto: f'c = {fc:.1f} MPa")
                    doc.add_paragraph(f"Acero: fy = {fy:.0f} MPa")
                    doc.add_paragraph(f"Norma: {norma_sel}")
                    doc.add_paragraph(f"Nivel Sísmico: {nivel_sis}")

                    doc.add_heading("2. Geometría y Refuerzo", level=1)
                    doc.add_paragraph(f"b = {b_vr:.0f} cm, h = {h_vr:.0f} cm, d = {d_vr:.1f} cm")
                    doc.add_paragraph(f"Refuerzo longitudinal: {n_bars} varillas {bar_vr} → As = {As_prov:.3f} cm²")

                    doc.add_heading("3. Verificaciones Normativas", level=1)
                    checks = [
                        ("Resistencia a flexión (φMn ≥ Mu)", "CUMPLE" if ok_flex else "NO CUMPLE"),
                        ("Cuantía mínima (ρ ≥ ρ_min)", "CUMPLE" if ok_rho_min else "NO CUMPLE"),
                        ("Cuantía máxima (ρ ≤ ρ_max)", "CUMPLE" if ok_rho_max else "NO CUMPLE"),
                    ]
                    for desc, res in checks:
                        doc.add_paragraph(f"{desc}: {res}")
                    doc.add_paragraph(f"φMn = {phi_Mn_kNm*factor_fuerza:.2f} {unidad_mom} | Mu = {Mu_vr:.2f} {unidad_mom}")
                    doc.add_paragraph(f"ρ = {rho_prov*100:.3f}% | ρ_min = {rho_min*100:.3f}% | ρ_max = {rho_max*100:.3f}%")
                    doc.add_paragraph(f"Referencia: {code['ref']}")

                    doc.add_heading("4. Cantidades de Obra", level=1)
                    doc.add_paragraph(f"Volumen concreto: {vol_horm:.4f} m³")
                    doc.add_paragraph(f"Acero longitudinal: {peso_long:.2f} kg")
                    doc.add_paragraph(f"Cemento: {bags:.1f} bultos de {bag_kg:.0f} kg")
                    doc.add_paragraph(f"Arena: {m[2]*vol_horm:.0f} kg")
                    doc.add_paragraph(f"Grava: {m[3]*vol_horm:.0f} kg")
                    doc.add_paragraph(f"Agua: {m[1]*vol_horm:.0f} L")

                    doc_mem = io.BytesIO()
                    doc.save(doc_mem)
                    doc_mem.seek(0)
                    st.download_button("Descargar Memoria DOCX", data=doc_mem, file_name=f"Memoria_Viga_{b_vr:.0f}x{h_vr:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════
# 3. VIGA T — FLEXIÓN
# ══════════════════════════════════════════
with st.expander(_t("📐 Diseño a Flexión — Viga T", "📐 Flexural Design — T-Beam")):
    st.markdown(_t(f"**Viga T — Sección compuesta** | Norma: `{code['ref']}`", f"**T-Beam — Composite section** | Code: `{code['ref']}`"))
    st.info(_t("📺 **Modo de uso:** Configura el ancho del ala (bf) y del alma (bw), más los espesores y el Momento Último (Mu). El algoritmo deducirá si la viga se comporta como Rectangular (compresión solo en el ala) o como Verdadera Viga T (eje neutro en el alma).", "📺 **How to use:** Set flange width (bf), web width (bw), thicknesses, and Ultimate Moment (Mu). The algorithm deduces if it behaves as a Rectangular or a True T-Beam."))
    c1,c2 = st.columns(2)
    with c1:
        bf_vt = st.number_input("Ancho del ala bf [cm]", 20.0, 300.0, st.session_state.get("vt_bf", 80.0), 5.0, key="vt_bf")
        bw_vt = st.number_input("Ancho del alma bw [cm]", 10.0, 80.0, st.session_state.get("vt_bw", 25.0), 5.0, key="vt_bw")
        hf_vt = st.number_input("Espesor ala hf [cm]", 5.0, 40.0, st.session_state.get("vt_hf", 12.0), 1.0, key="vt_hf")
        ht_vt = st.number_input("Alto total h [cm]", 20.0, 200.0, st.session_state.get("vt_h", 60.0), 5.0, key="vt_h")
    with c2:
        dp_vt = st.number_input("Recubrimiento d' [cm]", 2.0, 15.0, st.session_state.get("vt_dp", 6.0), 0.5, key="vt_dp")
        Mu_vt = st.number_input(f"Mu [{unidad_mom}]", 0.1, 15000.0, st.session_state.get("vt_mu", 200.0), 10.0, key="vt_mu")
        L_vt  = st.number_input("Longitud [m]", 1.0, 30.0, st.session_state.get("vt_L", 6.0), 0.5, key="vt_L")
        varillas_vt = list(rebar_dict.keys())
        # Índice por defecto: #4 (12.7mm) en US, 12mm en SI
        _def_vt = "#4 (Ø12.7mm)" if "Pulgadas" in bar_sys else "12mm"
        _def_idx_vt = varillas_vt.index(_def_vt) if _def_vt in varillas_vt else 1
        bar_vt = st.selectbox("Varilla:", varillas_vt, 
                              index=varillas_vt.index(st.session_state.vt_bar) if "vt_bar" in st.session_state and st.session_state.vt_bar in varillas_vt else _def_idx_vt,
                              key="vt_bar")
        Ab_vt = rebar_dict[bar_vt]; db_vt = diam_dict[bar_vt]

    d_vt = ht_vt - dp_vt
    if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
        Mu_vt_kN = Mu_vt / factor_fuerza
    else:
        Mu_vt_kN = Mu_vt
    bf_mm = bf_vt*10; bw_mm = bw_vt*10; hf_mm = hf_vt*10; d_mm_vt = d_vt*10

    Rn_t = Mu_vt_kN*1e6 / (phi_f * bf_mm * d_mm_vt**2)
    disc_t = 1 - 2*Rn_t/(0.85*fc)

    if disc_t < 0:
        st.error("❌ Sección insuficiente. Aumente bf o h.")
    else:
        rho_bf = (0.85*fc/fy)*(1 - math.sqrt(max(disc_t,0)))
        As_rect = rho_bf * bf_vt * d_vt
        a_r = As_rect*100*fy/(0.85*fc*bf_mm)
        is_T = a_r > hf_mm
        if not is_T:
            n_bt = math.ceil(As_rect/Ab_vt)
            As_prov_vt = n_bt*Ab_vt
            a_final = As_prov_vt*100*fy/(0.85*fc*bf_mm)
            phi_Mn_vt = phi_f*As_prov_vt*100*fy*(d_mm_vt-a_final/2)/1e6
            sec_type = "Rectangular (a ≤ hf)"
        else:
            Asf = 0.85*fc*(bf_mm-bw_mm)*hf_mm/fy
            Mnf = Asf*fy*(d_mm_vt-hf_mm/2)/1e6
            Mn_web = Mu_vt_kN/phi_f - Mnf
            if Mn_web < 0:
                Mn_web = 0
            Rn_w = Mn_web*1e6/(bw_mm*d_mm_vt**2) if Mn_web>0 else 0
            disc_w = 1 - 2*Rn_w/(0.85*fc)
            rho_w = max((0.85*fc/fy)*(1-math.sqrt(max(disc_w,0))),0)
            Asw_mm2 = rho_w*bw_mm*d_mm_vt
            As_total_mm2 = Asf + Asw_mm2
            n_bt = math.ceil(As_total_mm2/100/Ab_vt)
            As_prov_vt = n_bt*Ab_vt
            a_web = (As_prov_vt*100 - Asf)*fy/(0.85*fc*bw_mm) if (As_prov_vt*100-Asf)>0 else 0
            Mn_web_p = (As_prov_vt*100-Asf)*fy*(d_mm_vt-a_web/2) if (As_prov_vt*100-Asf)>0 else 0
            phi_Mn_vt = phi_f*(Mnf + Mn_web_p/1e6)
            sec_type = "T verdadera (a > hf)"

        ok_vt = phi_Mn_vt >= Mu_vt_kN
        rho_prov_vt = As_prov_vt / (bw_vt * d_vt)
        ok_rho_max_T = rho_prov_vt <= rho_max

        tab_r,tab_s,tab_3d,tab_q = st.tabs(["📊 Resultados","🔲 Sección 2D","🧊 Visualización 3D","📦 Cantidades"])
        with tab_r:
            st.markdown(f"**Tipo de sección:** {sec_type} | **φ={phi_f}**")
            rows_vt = [
                ("bf × bw × hf × h", f"{bf_vt:.0f} × {bw_vt:.0f} × {hf_vt:.0f} × {ht_vt:.0f} cm"),
                ("d efectivo", f"{d_vt:.1f} cm"),
                ("Comportamiento", sec_type),
                (f"Varillas ({bar_vt})", f"{n_bt} barras — As prov = {As_prov_vt:.3f} cm²"),
                (f"φMn calculado [{unidad_mom}]", f"{phi_Mn_vt*factor_fuerza:.2f}"),
                (f"Mu solicitado [{unidad_mom}]", f"{Mu_vt:.2f}"),
                ("ρ provisto", f"{(As_prov_vt/(bw_vt*d_vt))*100:.4f}%"),
                ("ρ máximo", f"{rho_max*100:.4f}%"),
                ("Validación ρ", "✅ CUMPLE" if ok_rho_max_T else "❌ EXCEDE (Sobrearmada)"),
                ("Estado", "✅ CUMPLE" if ok_vt else "❌ DEFICIENTE"),
            ]
            qty_table(rows_vt)
            (st.success if ok_vt else st.error)(f"φMn = {phi_Mn_vt*factor_fuerza:.2f} {unidad_mom} {'≥' if ok_vt else '<'} Mu = {Mu_vt:.2f} {unidad_mom}")
            st.info("💡 **¿Acero Inferior o Superior?** Si ingresa un Momento **Positivo** (Mu), el área calculada corresponde al refuerzo en la zona traccionada (usualmente **acero inferior**). Para momento **Negativo** en un apoyo continuo, la tracción está arriba por lo que el resultado corresponde al **acero superior**.")

        with tab_s:
            fig, ax = plt.subplots(figsize=(5,4))
            fig.patch.set_facecolor('#1a1a2e'); ax.set_facecolor('#1a1a2e')
            ax.add_patch(patches.Rectangle(((bf_vt-bw_vt)/2, 0), bw_vt, ht_vt-hf_vt, linewidth=1.5, edgecolor='white', facecolor='#4a4a6a'))
            ax.add_patch(patches.Rectangle((0, ht_vt-hf_vt), bf_vt, hf_vt, linewidth=1.5, edgecolor='white', facecolor='#3a3a5a'))
            r_v = db_vt/20
            xs_v = np.linspace((bf_vt-bw_vt)/2+dp_vt, (bf_vt+bw_vt)/2-dp_vt, max(n_bt,2))
            for x in xs_v[:n_bt]:
                ax.add_patch(plt.Circle((x, dp_vt), r_v, color='#ff6b35', zorder=5))
            ax.set_xlim(-5, bf_vt+5); ax.set_ylim(-5, ht_vt+5)
            ax.axis('off')
            ax.set_title(f"Viga T: bf={bf_vt:.0f} bw={bw_vt:.0f} hf={hf_vt:.0f} h={ht_vt:.0f} cm\n{n_bt}×{bar_vt}", color='white', fontsize=8)
            st.pyplot(fig)

        with tab_3d:
            st.subheader("🧊 Visualización 3D de Viga T")
            fig3d = go.Figure()
            L_mm_3d = L_vt * 100
            # Alma
            x_w = [-bw_vt/2, bw_vt/2, bw_vt/2, -bw_vt/2, -bw_vt/2, bw_vt/2, bw_vt/2, -bw_vt/2]
            y_w = [0, 0, ht_vt-hf_vt, ht_vt-hf_vt, 0, 0, ht_vt-hf_vt, ht_vt-hf_vt]
            z_w = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
            fig3d.add_trace(go.Mesh3d(x=x_w, y=y_w, z=z_w, alphahull=0, opacity=0.15, color='gray', name='Concreto (Alma)'))
            # Ala
            x_f = [-bf_vt/2, bf_vt/2, bf_vt/2, -bf_vt/2, -bf_vt/2, bf_vt/2, bf_vt/2, -bf_vt/2]
            y_f = [ht_vt-hf_vt, ht_vt-hf_vt, ht_vt, ht_vt, ht_vt-hf_vt, ht_vt-hf_vt, ht_vt, ht_vt]
            z_f = [0, 0, 0, 0, L_mm_3d, L_mm_3d, L_mm_3d, L_mm_3d]
            fig3d.add_trace(go.Mesh3d(x=x_f, y=y_f, z=z_f, alphahull=0, opacity=0.15, color='gray', name='Concreto (Ala)'))
            # ── Varillas INFERIORES (tensión) ──
            diam_reb_cm = db_vt / 10.0
            line_width = max(4, diam_reb_cm * 3)
            xs_v = np.linspace(-bw_vt/2 + dp_vt, bw_vt/2 - dp_vt, max(n_bt, 2)) if n_bt > 1 else [0]
            for idx, x_pos in enumerate(xs_v[:n_bt]):
                fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[dp_vt, dp_vt], z=[0, L_mm_3d],
                                            mode='lines', line=dict(color='darkred', width=line_width),
                                            name=f'Varilla inf. {bar_vt}', showlegend=(idx==0)))
            # ── Varillas SUPERIORES (compresión / montaje: 2 barras en esquinas del alma) ──
            y_sup = ht_vt - dp_vt  # y en el alma, cerca del ala
            xs_sup = [-bw_vt/2 + dp_vt, bw_vt/2 - dp_vt]
            for idx, x_pos in enumerate(xs_sup):
                fig3d.add_trace(go.Scatter3d(x=[x_pos, x_pos], y=[y_sup, y_sup], z=[0, L_mm_3d],
                                            mode='lines', line=dict(color='orange', width=max(3, diam_reb_cm*2)),
                                            name='Varilla sup. (compresión)', showlegend=(idx==0)))
            # ── Estribos (alma) ──
            tie_color = 'cornflowerblue'
            tie_width = max(2, (9.5/10.0) * 3)
            sep_ties = st.slider("Separación Estribos (cm) ", 5, 50, int(st.session_state.get('cv_s_diseno', 15)), 1, key="vt_sep_tie")
            tx = [-bw_vt/2 + dp_vt/2, bw_vt/2 - dp_vt/2, bw_vt/2 - dp_vt/2, -bw_vt/2 + dp_vt/2, -bw_vt/2 + dp_vt/2]
            ty = [dp_vt/2, dp_vt/2, ht_vt - dp_vt/2, ht_vt - dp_vt/2, dp_vt/2]
            L_cm = int(L_mm_3d)
            tx_all, ty_all, tz_all = [], [], []
            for zt in range(15, L_cm, sep_ties):
                tx_all.extend(tx + [None])
                ty_all.extend(ty + [None])
                tz_all.extend([zt]*5 + [None])
            fig3d.add_trace(go.Scatter3d(x=tx_all, y=ty_all, z=tz_all, mode='lines', 
                                         line=dict(color=tie_color, width=tie_width), name='Estribos Alma', showlegend=True))
            fig3d.update_layout(
                scene=dict(aspectmode='data', xaxis_title='b (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                legend=dict(bgcolor='rgba(0,0,0,0.5)', font=dict(color='white'), x=0.01, y=0.99),
                margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable'
            )
            st.plotly_chart(fig3d, use_container_width=True)

        with tab_q:
            vol_t = (bf_vt*hf_vt + bw_vt*(ht_vt-hf_vt))/10000 * L_vt
            peso_t = As_prov_vt * L_vt * 0.785
            m = mix_for_fc(fc)
            bags = m[0]*vol_t/bag_kg
            qty_table([("Concreto Viga T", f"{vol_t:.4f} m³"),
                       (f"Acero ({n_bt} barras)", f"{peso_t:.2f} kg"),
                       (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{m[0]*vol_t/bag_kg:.1f} bultos"),
                       ("Referencia", code["ref"])])

            if "apu_config" in st.session_state:
                apu = st.session_state.apu_config
                mon = apu.get("moneda", "$")
                st.markdown("---")
                st.success("✅ **Precios actualizados del mercado aplicados (APU).**")
                c_cem = bags * apu.get("cemento", 0)
                c_ace = peso_t * apu.get("acero", 0)
                vol_arena_m3 = (m[2]*vol_t)/1600
                vol_grava_m3 = (m[3]*vol_t)/1600
                c_are = vol_arena_m3 * apu.get("arena", 0)
                c_gra = vol_grava_m3 * apu.get("grava", 0)
                total_mat = c_cem + c_ace + c_are + c_gra
                total_dias_mo = (peso_t * 0.04) + (vol_t * 0.4)
                costo_mo = total_dias_mo * apu.get("costo_dia_mo", 70000)
                costo_directo = total_mat + costo_mo
                herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
                aiu = costo_directo * apu.get("pct_aui", 0.30)
                utilidad = costo_directo * apu.get("pct_util", 0.05)
                iva_v = utilidad * apu.get("iva", 0.19)
                gran_total = costo_directo + herramienta + aiu + iva_v
                st.metric(f"💎 Gran Total Proyecto ({mon})", f"{gran_total:,.2f}")

                with st.expander("📐 Dibujo de Figurado para Taller (Viga T)", expanded=False):
                    hook_len_cm = 12 * db_vt / 10
                    straight_len_cm = L_vt * 100 - 2 * hook_len_cm
                    fig_l1 = draw_longitudinal_bar(L_vt*100, straight_len_cm, hook_len_cm, db_vt)
                    st.pyplot(fig_l1)
                    recub_est = max(dp_vt, 2.5)
                    inside_b = bw_vt - 2*recub_est
                    inside_h = ht_vt - 2*recub_est
                    hook_len_est = 12 * 9.5 / 10
                    fig_e1 = draw_stirrup_beam(inside_b, inside_h, hook_len_est, 9.5)
                    st.pyplot(fig_e1)
            else:
                st.info("💡 Configure los precios en el expander '💰 APU – Precios en vivo' para ver el presupuesto.")

            # MEMORIA DOCX para Viga T
            if st.button("📄 Generar Memoria DOCX (Viga T)"):
                doc = Document()
                doc.add_heading(f"Memoria de Cálculo Viga T: bf={bf_vt:.0f} bw={bw_vt:.0f} h={ht_vt:.0f} cm", 0)
                doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_heading("1. Materiales", level=1)
                doc.add_paragraph(f"Concreto: f'c = {fc:.1f} MPa")
                doc.add_paragraph(f"Acero: fy = {fy:.0f} MPa")
                doc.add_paragraph(f"Norma: {norma_sel}")
                doc.add_heading("2. Geometría", level=1)
                doc.add_paragraph(f"Ancho ala (bf) = {bf_vt:.0f} cm, Ancho alma (bw) = {bw_vt:.0f} cm")
                doc.add_paragraph(f"Espesor ala (hf) = {hf_vt:.0f} cm, Altura total (h) = {ht_vt:.0f} cm")
                doc.add_paragraph(f"d efectivo = {d_vt:.1f} cm")
                doc.add_heading("3. Refuerzo y Comportamiento", level=1)
                doc.add_paragraph(f"Comportamiento: {sec_type}")
                doc.add_paragraph(f"Refuerzo longitudinal: {n_bt} varillas {bar_vt} → As = {As_prov_vt:.3f} cm²")
                doc.add_heading("4. Verificaciones Normativas", level=1)
                checks = [
                    ("Resistencia a flexión (φMn ≥ Mu)", "CUMPLE" if ok_vt else "NO CUMPLE"),
                    ("Cuantía máxima (ρ ≤ ρ_max)", "CUMPLE" if ok_rho_max_T else "NO CUMPLE"),
                ]
                for desc, res in checks:
                    doc.add_paragraph(f"{desc}: {res}")
                doc.add_paragraph(f"φMn = {phi_Mn_vt*factor_fuerza:.2f} {unidad_mom} | Mu = {Mu_vt:.2f} {unidad_mom}")
                doc.add_paragraph(f"ρ = {(As_prov_vt/(bw_vt*d_vt))*100:.3f}% | ρ_max = {rho_max*100:.3f}%")
                doc.add_paragraph(f"Referencia: {code['ref']}")
                doc.add_heading("5. Cantidades", level=1)
                doc.add_paragraph(f"Volumen concreto: {vol_t:.4f} m³")
                doc.add_paragraph(f"Acero: {peso_t:.2f} kg")
                doc_mem = io.BytesIO()
                doc.save(doc_mem)
                doc_mem.seek(0)
                st.download_button("Descargar Memoria Viga T", data=doc_mem, file_name=f"Memoria_VigaT_{bf_vt:.0f}x{bw_vt:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════
# 4. CORTANTE EN VIGAS
# ══════════════════════════════════════════
with st.expander(_t("⚡ Diseño a Cortante — Vigas de Concreto", "⚡ Shear Design — Concrete Beams")):
    st.markdown(_t(f"**Diseño de estribos a cortante** | Norma: `{code['ref']}`", f"**Shear stirrup design** | Code: `{code['ref']}`"))
    st.info(_t("📺 **Modo de uso:** Ingresa la Fuerza Cortante Factorizada (Vu) para una sección de viga dada. La app determinará la contribución del concreto (φVc) y calculará el refuerzo transversal requerido en número de estribos y separación (s).", "📺 **How to use:** Enter Factored Shear Force (Vu). The app calculates concrete contribution (φVc) and required transverse reinforcement (stirrups spacing & amount)."))
    c1,c2 = st.columns(2)
    with c1:
        bw_cv = st.number_input("bw [cm]", 10.0, 100.0, st.session_state.get("cv_bw", 25.0), 5.0, key="cv_bw")
        d_cv  = st.number_input("d peralte efectivo [cm]", 10.0, 200.0, st.session_state.get("cv_d", 45.0), 5.0, key="cv_d")
        Vu_cv_input = st.number_input(f"Vu [{unidad_fuerza}]", 0.1, 5000.0, st.session_state.get("cv_vu", 80.0), 5.0, key="cv_vu")
        if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
            Vu_cv = Vu_cv_input / factor_fuerza
        else:
            Vu_cv = Vu_cv_input
        L_cv  = st.number_input("Longitud viga [m]", 1.0, 30.0, st.session_state.get("cv_L", 5.0), 0.5, key="cv_L")
    with c2:
        h_cv  = st.number_input("h total [cm]", 20.0, 200.0, st.session_state.get("cv_h", 50.0), 5.0, key="cv_h")
        est_opts = ["Ø6mm","Ø8mm","Ø10mm","Ø12mm","#2","#3","#4"]
        st_bar_cv = st.selectbox("Estribo:", est_opts, 
                                 index=est_opts.index(st.session_state.cv_st) if "cv_st" in st.session_state and st.session_state.cv_st in est_opts else 1,
                                 key="cv_st")
        st_area = {"Ø6mm":0.283,"Ø8mm":0.503,"Ø10mm":0.785,"Ø12mm":1.131,"#2":0.32,"#3":0.71,"#4":1.29}[st_bar_cv]
        n_ramas = st.number_input("# Ramas del estribo", 2, 6, st.session_state.get("cv_ramas", 2), 1, key="cv_ramas")
        Av_cv = st_area * n_ramas
        diam_est = {"Ø6mm":6,"Ø8mm":8,"Ø10mm":10,"Ø12mm":12,"#2":6.35,"#3":9.53,"#4":12.70}[st_bar_cv]

    bw_mm_cv = bw_cv*10; d_mm_cv = d_cv*10
    Vc_N = 0.17*lam*math.sqrt(fc)*bw_mm_cv*d_mm_cv
    Vc_kN = Vc_N/1000
    phi_Vc = phi_v * Vc_kN
    Vs_req_kN = max(0, Vu_cv/phi_v - Vc_kN)
    need_design = Vu_cv > phi_Vc/2

    if Vs_req_kN > 0:
        s_calc_mm = Av_cv*100*fy*d_mm_cv/(Vs_req_kN*1000)
    else:
        s_calc_mm = min(d_mm_cv/2, 600)

    Vs_lim = 0.33*math.sqrt(fc)*bw_mm_cv*d_mm_cv/1000
    if Vs_req_kN > Vs_lim:
        s_max_mm = min(d_mm_cv/4, 300)
    else:
        s_max_mm = min(d_mm_cv/2, 600)
    s_diseno_mm = min(s_calc_mm, s_max_mm)
    s_diseno_cm = s_diseno_mm/10
    st.session_state['cv_s_diseno'] = s_diseno_cm

    n_estribos = math.ceil(L_cv*100/s_diseno_cm) + 1
    Vs_prov_kN = Av_cv*100*fy*d_mm_cv/(s_diseno_mm*1000)
    phi_Vn_kN = phi_v*(Vc_kN + Vs_prov_kN)
    ok_cv = phi_Vn_kN >= Vu_cv

    if s_diseno_cm < 5:
        st.warning("⚠️ La separación de estribos es menor a 5 cm. Considere aumentar el diámetro de los estribos o el número de ramas.")
    elif s_diseno_cm < 7.5:
        st.info("ℹ️ La separación de estribos es menor a 7.5 cm. Verifique que sea constructivamente viable.")

    tab_r,tab_s,tab_q = st.tabs(["📊 Resultados","🔲 Sección","📦 Cantidades"])
    with tab_r:
        st.markdown(f"**φ cortante = {phi_v}** | Norma: `{code['ref']}`")
        st.markdown(r"**Verificación Normativa:** $\phi V_n = \phi (V_c + V_s) \ge V_u$")
        st.latex(r"V_s = \frac{A_v f_{yt} d}{s}")
        Vs_max_kN = 0.66*math.sqrt(fc)*bw_mm_cv*d_mm_cv/1000
        rows_cv = [
            ("bw × d", f"{bw_cv:.0f} × {d_cv:.0f} cm"),
            (f"Vc (concreto) [{unidad_fuerza}]", f"{Vc_kN*factor_fuerza:.2f}"),
            (f"φVc [{unidad_fuerza}]", f"{phi_Vc*factor_fuerza:.2f}"),
            (f"Vu [{unidad_fuerza}]", f"{Vu_cv_input:.2f}"),
            (f"Vs requerido [{unidad_fuerza}]", f"{Vs_req_kN*factor_fuerza:.2f}"),
            (f"Av ({n_ramas} ramas {st_bar_cv})", f"{Av_cv:.3f} cm²"),
            ("s calculado", f"{s_calc_mm:.0f} mm = {s_calc_mm/10:.1f} cm"),
            ("s máx (norma)", f"{s_max_mm:.0f} mm = {s_max_mm/10:.1f} cm"),
            ("s de diseño", f"**{s_diseno_cm:.1f} cm**"),
            (f"Vs provisto [{unidad_fuerza}]", f"{Vs_prov_kN*factor_fuerza:.2f}"),
            (f"φVn = φ(Vc+Vs) [{unidad_fuerza}]", f"{phi_Vn_kN*factor_fuerza:.2f}"),
            (f"Vs máx permitido [{unidad_fuerza}]", f"{Vs_max_kN*factor_fuerza:.2f}"),
            ("Estado", "✅ CUMPLE" if ok_cv else "❌ DEFICIENTE"),
        ]
        qty_table(rows_cv)
        if Vs_req_kN > Vs_max_kN:
            st.error(f"❌ Vs requerido excede $V_{{s,max}}$ = {Vs_max_kN*factor_fuerza:.2f} {unidad_fuerza}. $\\rightarrow$ **Aumentar la sección (bw o h)**")
        elif ok_cv:
            st.success(f"✅ Aprobado Cortante: $\\phi V_n = {phi_Vn_kN*factor_fuerza:.2f}$ {unidad_fuerza} $\\ge V_u = {Vu_cv_input:.2f}$ {unidad_fuerza} — Estribo {st_bar_cv} @ {s_diseno_cm:.1f} cm")
        else:
            st.error(f"❌ No Aprobado por Cortante: $\\phi V_n = {phi_Vn_kN*factor_fuerza:.2f}$ {unidad_fuerza} $< V_u = {Vu_cv_input:.2f}$ {unidad_fuerza}")

    with tab_s:
        fig, ax = sec_dark_fig(bw_cv, h_cv, f"Sección Cortante {bw_cv:.0f}×{h_cv:.0f} cm")
        recub_cv = (h_cv-d_cv)*0.5
        ax.add_patch(patches.Rectangle((recub_cv,recub_cv),bw_cv-2*recub_cv,h_cv-2*recub_cv,linewidth=2,edgecolor='#00d4ff',facecolor='none',linestyle='--'))
        ax.text(bw_cv/2,h_cv/2,f"s={s_diseno_cm:.1f}cm\nVu={Vu_cv_input:.0f}{unidad_fuerza}",ha='center',va='center',color='white',fontsize=8)
        st.pyplot(fig)

    with tab_q:
        perim_cv = 2*(bw_cv-2*recub_cv) + 2*(h_cv-2*recub_cv) + 6*diam_est/10
        vol_beam_cv = bw_cv/100*h_cv/100*L_cv
        peso_est_cv = n_estribos * (perim_cv / 100.0) * st_area * 0.785
        m = mix_for_fc(fc)
        bags = m[0]*vol_beam_cv/bag_kg
        qty_table([
            (f"Estribos {st_bar_cv} @ {s_diseno_cm:.1f}cm", f"{n_estribos} estribos"),
            ("Peso estribos", f"{peso_est_cv:.2f} kg"),
            ("Longitud total estribos", f"{n_estribos*perim_cv/100:.2f} m"),
            ("Concreto viga", f"{vol_beam_cv:.4f} m³"),
            (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{bags:.1f} bultos"),
            ("Referencia", code["ref"]),
        ])

        if "apu_config" in st.session_state:
            apu = st.session_state.apu_config
            mon = apu.get("moneda", "$")
            st.markdown("---")
            st.success("✅ **Precios actualizados del mercado aplicados (APU).**")
            c_cem = bags * apu.get("cemento", 0)
            c_ace = peso_est_cv * apu.get("acero", 0)
            vol_arena_m3 = (m[2]*vol_beam_cv)/1600
            vol_grava_m3 = (m[3]*vol_beam_cv)/1600
            c_are = vol_arena_m3 * apu.get("arena", 0)
            c_gra = vol_grava_m3 * apu.get("grava", 0)
            total_mat = c_cem + c_ace + c_are + c_gra
            total_dias_mo = (peso_est_cv * 0.04) + (vol_beam_cv * 0.4)
            costo_mo = total_dias_mo * apu.get("costo_dia_mo", 70000)
            costo_directo = total_mat + costo_mo
            herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * apu.get("pct_aui", 0.30)
            utilidad = costo_directo * apu.get("pct_util", 0.05)
            iva_v = utilidad * apu.get("iva", 0.19)
            gran_total = costo_directo + herramienta + aiu + iva_v
            st.metric(f"💎 Gran Total Proyecto ({mon})", f"{gran_total:,.2f}")

            with st.expander("📐 Dibujo de Estribo para Taller", expanded=False):
                recub_est = max(recub_cv, 2.5)
                inside_b = bw_cv - 2*recub_est
                inside_h = h_cv - 2*recub_est
                hook_len_est = 12 * diam_est / 10
                fig_est = draw_stirrup_beam(inside_b, inside_h, hook_len_est, diam_est)
                st.pyplot(fig_est)
                st.caption("Estribo con ganchos de 135° (representación esquemática).")
        else:
            st.info("💡 Configure los precios en el expander '💰 APU – Precios en vivo' para ver el presupuesto.")

        # MEMORIA DOCX para Cortante
        if st.button("📄 Generar Memoria Cortante (DOCX)"):
            doc = Document()
            doc.add_heading("Memoria de Diseño a Cortante", 0)
            doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_heading("1. Datos de la Sección", level=1)
            doc.add_paragraph(f"bw = {bw_cv:.0f} cm, d = {d_cv:.0f} cm")
            doc.add_paragraph(f"Concreto: f'c = {fc:.1f} MPa")
            doc.add_paragraph(f"Acero estribos: {st_bar_cv} (Av = {Av_cv:.3f} cm², {n_ramas} ramas)")
            doc.add_heading("2. Verificaciones Normativas", level=1)
            checks = [
                (f"Resistencia a cortante (φVn ≥ Vu)", "CUMPLE" if ok_cv else "NO CUMPLE"),
                ("Vs ≤ Vs,max", "CUMPLE" if Vs_req_kN <= Vs_max_kN else "NO CUMPLE"),
                ("Separación ≤ s_max", "CUMPLE" if s_diseno_mm <= s_max_mm else "NO CUMPLE"),
            ]
            for desc, res in checks:
                doc.add_paragraph(f"{desc}: {res}")
            doc.add_paragraph(f"φVn = {phi_Vn_kN*factor_fuerza:.2f} {unidad_fuerza} | Vu = {Vu_cv_input:.2f} {unidad_fuerza}")
            doc.add_paragraph(f"Separación de diseño: s = {s_diseno_cm:.1f} cm")
            doc.add_paragraph(f"Referencia: {code['ref']}")
            doc_mem = io.BytesIO()
            doc.save(doc_mem)
            doc_mem.seek(0)
            st.download_button("Descargar Memoria Cortante", data=doc_mem, file_name=f"Memoria_Cortante_{bw_cv:.0f}x{d_cv:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════
# 5. PUNZONAMIENTO EN LOSAS
# ══════════════════════════════════════════
with st.expander(_t("⚡ Resistencia a Cortante por Punzonamiento — Losas", "⚡ Punching Shear Resistance — Slabs")):
    st.markdown(_t(f"**Verificación de punzonamiento** (slab-column connection) | Norma: `{code['ref']}`", f"**Punching shear check** (slab-column connection) | Code: `{code['ref']}`"))
    c1,c2 = st.columns(2)
    with c1:
        c1p = st.number_input("Dimensión columna c1 [cm]", 15.0, 100.0, st.session_state.get("pz_c1", 30.0), 5.0, key="pz_c1")
        c2p = st.number_input("Dimensión columna c2 [cm]", 15.0, 100.0, st.session_state.get("pz_c2", 30.0), 5.0, key="pz_c2")
        h_pz = st.number_input("Espesor losa h [cm]", 10.0, 60.0, st.session_state.get("pz_h", 20.0), 1.0, key="pz_h")
    with c2:
        cov_pz = st.number_input("Recubrimiento [cm]", 1.5, 5.0, st.session_state.get("pz_cov", 2.5), 0.5, key="pz_cov")
        Vu_pz_input = st.number_input(f"Vu en columna [{unidad_fuerza}]", 10.0, 10000.0, st.session_state.get("pz_vu", 500.0), 50.0, key="pz_vu")
        if unidades_salida == "Toneladas fuerza (tonf, tonf·m)":
            Vu_pz = Vu_pz_input / factor_fuerza
        else:
            Vu_pz = Vu_pz_input
        pz_opts = ["Interior (αs=40)","Borde (αs=30)","Esquina (αs=20)"]
        tipo_col = st.selectbox("Posición columna:", pz_opts, 
                                index=pz_opts.index(st.session_state.pz_tipo) if "pz_tipo" in st.session_state and st.session_state.pz_tipo in pz_opts else 0,
                                key="pz_tipo")
    alpha_s = {"Interior (αs=40)":40,"Borde (αs=30)":30,"Esquina (αs=20)":20}[tipo_col]

    d_pz = (h_pz - cov_pz)*10
    c1_mm = c1p*10; c2_mm = c2p*10
    bo_mm = 2*(c1_mm+d_pz) + 2*(c2_mm+d_pz)
    beta_pz = max(c1p,c2p)/min(c1p,c2p)
    Vc1_N = (0.17+0.33/beta_pz)*lam*math.sqrt(fc)*bo_mm*d_pz
    Vc2_N = (0.083+0.083*alpha_s*d_pz/bo_mm)*lam*math.sqrt(fc)*bo_mm*d_pz
    Vc3_N = 0.33*lam*math.sqrt(fc)*bo_mm*d_pz
    Vc_pz_N = min(Vc1_N, Vc2_N, Vc3_N)
    phi_Vc_pz = phi_v*Vc_pz_N/1000
    ok_pz = phi_Vc_pz >= Vu_pz

    h_min_req = h_pz
    if not ok_pz:
        for h_test in range(int(h_pz) + 1, 300):
            d_t = (h_test - cov_pz) * 10
            bo_t = 2*(c1_mm+d_t) + 2*(c2_mm+d_t)
            Vc1_t = (0.17+0.33/beta_pz)*lam*math.sqrt(fc)*bo_t*d_t
            Vc2_t = (0.083+0.083*alpha_s*d_t/bo_t)*lam*math.sqrt(fc)*bo_t*d_t
            Vc3_t = 0.33*lam*math.sqrt(fc)*bo_t*d_t
            Vc_pz_t = min(Vc1_t, Vc2_t, Vc3_t)
            if phi_v * Vc_pz_t / 1000 >= Vu_pz:
                h_min_req = h_test
                break

    tab_r,tab_q = st.tabs(["📊 Resultados","📦 Cantidades"])
    with tab_r:
        qty_table([
            ("d efectivo losa", f"{d_pz:.0f} mm"),
            ("β = c_max/c_min", f"{beta_pz:.2f}"),
            ("bo (perímetro crítico)", f"{bo_mm:.0f} mm = {bo_mm/10:.1f} cm"),
            (f"Vc1 (β-fórmula) [{unidad_fuerza}]", f"{Vc1_N/1000*factor_fuerza:.2f}"),
            (f"Vc2 (αs-fórmula) [{unidad_fuerza}]", f"{Vc2_N/1000*factor_fuerza:.2f}"),
            (f"Vc3 (simplificada) [{unidad_fuerza}]", f"{Vc3_N/1000*factor_fuerza:.2f}"),
            (f"Vc diseño = min(Vc1,Vc2,Vc3) [{unidad_fuerza}]", f"{Vc_pz_N/1000*factor_fuerza:.2f}"),
            (f"φ Vc [{unidad_fuerza}]", f"{phi_Vc_pz*factor_fuerza:.2f}"),
            (f"Vu solicitado [{unidad_fuerza}]", f"{Vu_pz_input:.2f}"),
            ("Estado", "✅ CUMPLE" if ok_pz else f"❌ REFORZAR / Aumentar h a {h_min_req} cm mín."),
        ])
        (st.success if ok_pz else st.error)(f"φVc = {phi_Vc_pz*factor_fuerza:.2f} {unidad_fuerza} {'≥' if ok_pz else '<'} Vu = {Vu_pz_input:.2f} {unidad_fuerza} — Ref: {code['ref']}")
        if not ok_pz:
            st.error(f"❌ **FALLA POR PUNZONAMIENTO:** El cortante solicitante Vu excede la resistencia del concreto φVc.\n\n"
                     f"**¿QUÉ AUMENTAR? Soluciones propuestas:**\n"
                     f"1. Aumentar el espesor de la losa **h** a por lo menos **{h_min_req} cm**.\n"
                     f"2. Aumentar la resistencia del concreto **f'c**.\n"
                     f"3. Aumentar las dimensiones de la columna o diseñar un ábaco / capitel.")
    with tab_q:
        qty_table([("Referencia ACI", "ACI 318-25 Tabla 22.6.5.2"),
                   ("Referencia Norma", code["ref"]),
                   ("Nota","Para casos con Mu en columna verificar momento excéntrico")])

    # MEMORIA DOCX para Punzonamiento
    if st.button("📄 Generar Memoria Punzonamiento (DOCX)"):
        doc = Document()
        doc.add_heading("Memoria de Verificación a Punzonamiento", 0)
        doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_heading("1. Datos", level=1)
        doc.add_paragraph(f"Columna: c1={c1p:.0f} cm, c2={c2p:.0f} cm")
        doc.add_paragraph(f"Losa: h={h_pz:.0f} cm, d={d_pz:.0f} mm, recubrimiento={cov_pz:.1f} cm")
        doc.add_paragraph(f"Concreto: f'c={fc:.1f} MPa")
        doc.add_heading("2. Verificaciones Normativas", level=1)
        checks = [
            ("Resistencia a punzonamiento (φVc ≥ Vu)", "CUMPLE" if ok_pz else "NO CUMPLE"),
        ]
        for desc, res in checks:
            doc.add_paragraph(f"{desc}: {res}")
        doc.add_paragraph(f"φVc = {phi_Vc_pz*factor_fuerza:.2f} {unidad_fuerza} | Vu = {Vu_pz_input:.2f} {unidad_fuerza}")
        doc.add_paragraph(f"Referencia: {code['ref']} / ACI 318-25 Sección 22.6")
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button("Descargar Memoria Punzonamiento", data=doc_mem, file_name=f"Memoria_Punzonamiento_{c1p:.0f}x{c2p:.0f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════
# 6. INERCIA FISURADA + DEFLEXIONES
# ══════════════════════════════════════════
with st.expander(_t("📏 Inercia Fisurada y Deflexiones en Vigas", "📏 Cracked Inertia & Deflections")):
    st.markdown(f"**Branson (1965) — ACI 318** | {_t('Norma', 'Code')}: `{code['ref']}`")
    c1,c2 = st.columns(2)
    with c1:
        b_de = st.number_input("b [cm]", 10.0, 150.0, st.session_state.get("de_b", 25.0), 5.0, key="de_b")
        h_de = st.number_input("h [cm]", 15.0, 200.0, st.session_state.get("de_h", 50.0), 5.0, key="de_h")
        dp_de = st.number_input("d' [cm]", 2.0, 15.0, st.session_state.get("de_dp", 6.0), 0.5, key="de_dp")
        As_de = st.number_input("As provisto [cm²]", 0.5, 100.0, st.session_state.get("de_as", 5.0), 0.5, key="de_as")
    with c2:
        L_de = st.number_input("Luz libre [m]", 1.0, 20.0, st.session_state.get("de_L", 5.0), 0.5, key="de_L")
        wD_de = st.number_input("Carga muerta wD [kN/m]", 0.0, 200.0, st.session_state.get("de_wD", 15.0), 1.0, key="de_wD")
        wL_de = st.number_input("Carga viva wL [kN/m]", 0.0, 200.0, st.session_state.get("de_wL", 10.0), 1.0, key="de_wL")
        cond_opts = ["Simplemente apoyada","Continua un extremo","Continua dos extremos"]
        cond_de = st.selectbox("Condición de apoyo:", cond_opts, 
                               index=cond_opts.index(st.session_state.de_cond) if "de_cond" in st.session_state and st.session_state.de_cond in cond_opts else 0,
                               key="de_cond")

    d_de = h_de - dp_de
    d_de_mm = d_de*10; b_de_mm = b_de*10; As_de_mm2 = As_de*100
    n_de = Es/Ec
    A_ = b_de_mm/2; B_ = n_de*As_de_mm2; C_ = -n_de*As_de_mm2*d_de_mm
    x_de = (-B_ + math.sqrt(B_**2 - 4*A_*C_))/(2*A_)
    Ig_mm4 = b_de_mm*(h_de*10)**3/12
    Icr_mm4 = b_de_mm*x_de**3/3 + n_de*As_de_mm2*(d_de_mm-x_de)**2
    yt_mm = h_de*10/2
    fr = 0.62*lam*math.sqrt(fc)
    Mcr_Nmm = fr*Ig_mm4/yt_mm
    Mcr_kNm = Mcr_Nmm/1e6

    coef = {"Simplemente apoyada":8,"Continua un extremo":10,"Continua dos extremos":16}[cond_de]
    Ma_D_kNm = wD_de*L_de**2/coef
    Ma_DL_kNm = (wD_de+wL_de)*L_de**2/coef

    def Ie(Ma_kNm, Mcr_kNm, Ig, Icr):
        if Ma_kNm <= 0: return Ig
        ratio = min(Mcr_kNm/Ma_kNm, 1.0)
        return min(ratio**3*Ig + (1-ratio**3)*Icr, Ig)

    Ie_D = Ie(Ma_D_kNm, Mcr_kNm, Ig_mm4, Icr_mm4)
    Ie_DL = Ie(Ma_DL_kNm, Mcr_kNm, Ig_mm4, Icr_mm4)

    fact_defl = {"Simplemente apoyada":5/384,"Continua un extremo":1/185,"Continua dos extremos":1/384}[cond_de]
    L_mm = L_de*1000
    defl_D_mm = fact_defl*wD_de*(L_mm**4)/(Ec*Ie_D)
    defl_DL_mm = fact_defl*(wD_de+wL_de)*(L_mm**4)/(Ec*Ie_DL)
    defl_L_mm = defl_DL_mm - defl_D_mm
    lim_L480 = L_mm/480; lim_L240 = L_mm/240
    ok_defl_L = defl_L_mm <= lim_L480
    ok_defl_total = defl_DL_mm <= lim_L240

    tab_r,tab_q = st.tabs(["📊 Resultados","📦 Cantidades"])
    with tab_r:
        st.markdown(f"**Ec = {Ec:.0f} MPa** | **n = {n_de:.2f}** | **fr = {fr:.3f} MPa**")
        qty_table([
            ("Ig (inercia bruta)", f"{Ig_mm4:.3e} mm⁴ = {Ig_mm4/1e4:.1f} cm⁴"),
            ("Eje neutro fisurado (x)", f"{x_de:.1f} mm"),
            ("Icr (inercia fisurada)", f"{Icr_mm4:.3e} mm⁴ = {Icr_mm4/1e4:.1f} cm⁴"),
            ("fr (módulo de rotura)", f"{fr:.3f} MPa"),
            ("Mcr (momento de agrietamiento)", f"{Mcr_kNm:.2f} kN·m"),
            ("Ma (D)", f"{Ma_D_kNm:.2f} kN·m"),
            ("Ma (D+L)", f"{Ma_DL_kNm:.2f} kN·m"),
            ("Ie (D)", f"{Ie_D:.3e} mm⁴"),
            ("Ie (D+L)", f"{Ie_DL:.3e} mm⁴"),
            ("Δ carga muerta D", f"{defl_D_mm:.2f} mm"),
            ("Δ carga viva L", f"{defl_L_mm:.2f} mm"),
            ("Límite L/480 (carga viva)", f"{lim_L480:.1f} mm"),
            ("Δ_L vs L/480", "✅ CUMPLE" if ok_defl_L else "❌ EXCEDE"),
            ("Límite L/240 (total)", f"{lim_L240:.1f} mm"),
            ("Δ_DL vs L/240", "✅ CUMPLE" if ok_defl_total else "❌ EXCEDE"),
        ])
        st.caption(f"📖 {code['ref']} | ACI 318-25 Tabla 24.2.2")
    with tab_q:
        qty_table([("Referencia deflexiones","ACI 318-25 Sección 24.2"),
                   ("Norma aplicada", code["ref"])])

    # MEMORIA DOCX para Deflexiones
    if st.button("📄 Generar Memoria Deflexiones (DOCX)"):
        doc = Document()
        doc.add_heading("Memoria de Deflexiones en Vigas", 0)
        doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_heading("1. Datos", level=1)
        doc.add_paragraph(f"Sección: b={b_de:.0f} cm, h={h_de:.0f} cm, d={d_de:.1f} cm")
        doc.add_paragraph(f"Acero longitudinal: As={As_de:.2f} cm²")
        doc.add_paragraph(f"Luz: L={L_de:.2f} m")
        doc.add_paragraph(f"Cargas: wD={wD_de:.2f} kN/m, wL={wL_de:.2f} kN/m")
        doc.add_heading("2. Verificaciones Normativas", level=1)
        checks = [
            ("Deflexión por carga viva ≤ L/480", "CUMPLE" if ok_defl_L else "NO CUMPLE"),
            ("Deflexión total ≤ L/240", "CUMPLE" if ok_defl_total else "NO CUMPLE"),
        ]
        for desc, res in checks:
            doc.add_paragraph(f"{desc}: {res}")
        doc.add_paragraph(f"Δ viva = {defl_L_mm:.2f} mm (límite {lim_L480:.1f} mm)")
        doc.add_paragraph(f"Δ total = {defl_DL_mm:.2f} mm (límite {lim_L240:.1f} mm)")
        doc.add_paragraph(f"Referencia: {code['ref']} / ACI 318-25 Sección 24.2")
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button("Descargar Memoria Deflexiones", data=doc_mem, file_name="Memoria_Deflexiones.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════
# 7. LOSA EN UNA DIRECCIÓN
# ══════════════════════════════════════════
with st.expander(_t("🏗️ Diseño de Losa en Una Dirección", "🏗️ One-Way Slab Design")):
    st.markdown(_t(f"**Diseño por franja de 1 metro** | Norma: `{code['ref']}`", f"**1-meter strip design** | Code: `{code['ref']}`"))
    c1,c2 = st.columns(2)
    with c1:
        ln_ls = st.number_input("Luz libre ln [m]", 1.0, 10.0, st.session_state.get("ls_ln", 3.5), 0.25, key="ls_ln")
        h_ls  = st.number_input("Espesor losa h [cm]", 8.0, 40.0, st.session_state.get("ls_h", 15.0), 1.0, key="ls_h")
        cov_ls = st.number_input("Recubrimiento [cm]", 1.5, 5.0, st.session_state.get("ls_cov", 2.5), 0.5, key="ls_cov")
        wD_ls = st.number_input("Carga muerta total (D) [kN/m²]", 1.0, 50.0, st.session_state.get("ls_wD", 7.0), 0.5, key="ls_wD")
    with c2:
        wL_ls = st.number_input("Carga viva (L) [kN/m²]", 0.5, 30.0, st.session_state.get("ls_wL", 5.0), 0.5, key="ls_wL")
        varillas_ls = list(rebar_dict.keys())
        bar_ls = st.selectbox("Varilla losa:", varillas_ls, 
                              index=varillas_ls.index(st.session_state.ls_bar) if "ls_bar" in st.session_state and st.session_state.ls_bar in varillas_ls else 0,
                              key="ls_bar")
        Ab_ls = rebar_dict[bar_ls]; db_ls = diam_dict[bar_ls]
        apoyo_opts = ["Simplemente apoyada","Continua 2 extremos","Voladizo"]
        apoyo_ls = st.selectbox("Condición:", apoyo_opts, 
                                index=apoyo_opts.index(st.session_state.ls_apoyo) if "ls_apoyo" in st.session_state and st.session_state.ls_apoyo in apoyo_opts else 0,
                                key="ls_apoyo")

    wu_ls = 1.2*wD_ls + 1.6*wL_ls
    b_ls = 100
    d_ls = h_ls - cov_ls - db_ls/20
    coef_ls = {"Simplemente apoyada":8,"Continua 2 extremos":16,"Voladizo":2}[apoyo_ls]
    Mu_ls_kNm = wu_ls*ln_ls**2/coef_ls

    d_ls_mm = d_ls*10; b_ls_mm = b_ls*10
    Rn_ls = Mu_ls_kNm*1e6/(phi_f*b_ls_mm*d_ls_mm**2)
    disc_ls = 1-2*Rn_ls/(0.85*fc)
    if disc_ls < 0:
        st.error("❌ Losa muy delgada o carga muy alta — aumente h")
    else:
        rho_ls = (0.85*fc/fy)*(1-math.sqrt(disc_ls))
        rho_use_ls = max(rho_ls, rho_min)
        As_req_ls = rho_use_ls*b_ls*d_ls
        s_bar_ls = Ab_ls/As_req_ls*100
        s_max_ls = min(3*h_ls, 45)
        s_use_ls = min(s_bar_ls, s_max_ls)
        As_prov_ls = Ab_ls/(s_use_ls/100)
        a_ls_mm = As_prov_ls*100*fy/(0.85*fc*b_ls_mm)
        phi_Mn_ls = phi_f*As_prov_ls*100*fy*(d_ls_mm-a_ls_mm/2)/1e6
        As_temp = 0.0018*b_ls*h_ls
        s_temp = min(Ab_ls/As_temp*100, 5*h_ls, 45)

        ok_ls = phi_Mn_ls >= Mu_ls_kNm
        tab_r,tab_s,tab_g,tab_3d,tab_q = st.tabs(["📊 Resultados","🔲 Sección 2D","📈 Gráficos M/V","🧊 3D","📦 Cantidades"])
        with tab_r:
            qty_table([
                ("wu factorizada", f"{wu_ls:.2f} kN/m²"),
                ("Mu (franja 1m)", f"{Mu_ls_kNm:.2f} kN·m"),
                ("h losa / d efectivo", f"{h_ls:.0f} / {d_ls:.1f} cm"),
                ("As requerido", f"{As_req_ls:.3f} cm²/m"),
                (f"Espaciado varilla {bar_ls}", f"{s_bar_ls:.1f} cm → usar **{s_use_ls:.1f} cm**"),
                ("As provisto", f"{As_prov_ls:.3f} cm²/m"),
                ("φMn / Mu", f"{phi_Mn_ls:.2f} / {Mu_ls_kNm:.2f} kN·m/m"),
                ("Estado Flexión", "✅ CUMPLE" if ok_ls else "❌ DEFICIENTE"),
                ("As temperatura/retracción", f"{As_temp:.3f} cm²/m"),
                (f"Varilla temp {bar_ls}", f"@ {s_temp:.1f} cm"),
            ])
            (st.success if ok_ls else st.error)(f"Losa {'OK' if ok_ls else 'DEFICIENTE'} — {bar_ls} @ {s_use_ls:.1f} cm (As principal)")
            st.info("💡 **Acero Inferior vs. Superior:** Este diseño automático a partir de la luz libre (L) estima el máximo momento positivo de la franja. El resultado `As principal` mostrado es el **acero inferior**. Para el **acero superior** necesario en los nudos continuos, considere diseñar a flexión una viga de b=100m ingresando el Mu- de los apoyos.")
        
        with tab_s:
            fig_s, ax_s = sec_dark_fig(40, h_ls*2.5, f"Sección Losa — h={h_ls:.0f}cm")
            r_ls = db_ls/20
            for xi in np.arange(db_ls/20+cov_ls*0.3, 38, s_use_ls*0.3):
                ax_s.add_patch(plt.Circle((xi, cov_ls*0.3+r_ls), r_ls, color='#ff6b35', zorder=5))
            st.pyplot(fig_s)
        
        with tab_g:
            st.subheader("Gráficos de Cortante (V) y Momento Flector (M)")
            x_vals = np.linspace(0, ln_ls, 100)
            if apoyo_ls == "Simplemente apoyada":
                V_vals = wu_ls * ln_ls / 2 - wu_ls * x_vals
                M_vals = wu_ls * ln_ls * x_vals / 2 - wu_ls * x_vals**2 / 2
            elif apoyo_ls == "Continua 2 extremos":
                V_vals = wu_ls * ln_ls / 2 - wu_ls * x_vals
                M_vals = wu_ls * ln_ls * x_vals / 2 - wu_ls * x_vals**2 / 2 - wu_ls * ln_ls**2 / 12
            else: # Voladizo
                V_vals = -wu_ls * x_vals 
                M_vals = -wu_ls * x_vals**2 / 2
            
            fig_mv, (ax_v, ax_m) = plt.subplots(2, 1, figsize=(6, 5), sharex=True)
            fig_mv.patch.set_facecolor('#1a1a2e')
            ax_v.set_facecolor('#1a1a2e'); ax_m.set_facecolor('#1a1a2e')
            
            ax_v.plot(x_vals, V_vals, color='#00d4ff', lw=2)
            ax_v.fill_between(x_vals, 0, V_vals, color='#00d4ff', alpha=0.3)
            ax_v.axhline(0, color='white', lw=1)
            ax_v.set_ylabel("Cortante (kN)", color='white')
            ax_v.tick_params(colors='white')
            
            ax_m.plot(x_vals, M_vals, color='#ff6b35', lw=2)
            ax_m.fill_between(x_vals, 0, M_vals, color='#ff6b35', alpha=0.3)
            ax_m.axhline(0, color='white', lw=1)
            ax_m.set_ylabel("Momento (kN·m/m)", color='white')
            ax_m.set_xlabel("Distancia x (m)", color='white')
            ax_m.tick_params(colors='white')
            fig_mv.tight_layout()
            st.pyplot(fig_mv)

        with tab_3d:
            st.subheader("🧊 Losa 3D (Franja de 1m)")
            fig3_ls = go.Figure()
            L_ls_cm = ln_ls * 100
            x_l = [-50, 50, 50, -50, -50, 50, 50, -50]
            y_l = [0, 0, h_ls, h_ls, 0, 0, h_ls, h_ls]
            z_l = [0, 0, 0, 0, L_ls_cm, L_ls_cm, L_ls_cm, L_ls_cm]
            fig3_ls.add_trace(go.Mesh3d(x=x_l, y=y_l, z=z_l, alphahull=0, opacity=0.15, color='gray', name='Concreto'))
            
            d_real_ls = cov_ls + db_ls/20
            y_bar = d_real_ls if apoyo_ls != "Voladizo" else h_ls - d_real_ls
            n_bars_1m = int(100/s_use_ls)
            if n_bars_1m < 1: n_bars_1m = 1
            xs_ls = np.linspace(-50 + s_use_ls/2, 50 - s_use_ls/2, n_bars_1m)
            line_w_ls = max(3, db_ls/10 * 3)
            for idx, xb in enumerate(xs_ls):
                fig3_ls.add_trace(go.Scatter3d(x=[xb, xb], y=[y_bar, y_bar], z=[0, L_ls_cm],
                                              mode='lines', line=dict(color='darkred', width=line_w_ls),
                                              name=f'Principal {bar_ls}', showlegend=(idx==0)))
            
            y_temp = y_bar + db_ls/10 if apoyo_ls != "Voladizo" else y_bar - db_ls/10
            line_w_t = max(2, db_ls/10 * 2)
            n_temp = int(L_ls_cm / s_temp)
            if n_temp > 0:
                z_temp_vals = np.linspace(s_temp, L_ls_cm-s_temp, n_temp)
                tx_ls, ty_ls, tz_ls = [], [], []
                for zt in z_temp_vals:
                    tx_ls.extend([-50, 50, None])
                    ty_ls.extend([y_temp, y_temp, None])
                    tz_ls.extend([zt, zt, None])
                fig3_ls.add_trace(go.Scatter3d(x=tx_ls, y=ty_ls, z=tz_ls, mode='lines', 
                                               line=dict(color='cornflowerblue', width=line_w_t), name='Temperatura', showlegend=True))
                                               
            fig3_ls.update_layout(scene=dict(aspectmode='data', xaxis_title='Ancho (cm)', yaxis_title='h (cm)', zaxis_title='L (cm)'),
                                margin=dict(l=0, r=0, b=0, t=0), height=450, dragmode='turntable')
            st.plotly_chart(fig3_ls, use_container_width=True)

        with tab_q:
            area_losa = ln_ls*1
            vol_ls = area_losa*h_ls/100
            peso_flex_ls = As_prov_ls * ln_ls * 0.785
            peso_temp_ls = As_temp * ln_ls * 0.785
            m = mix_for_fc(fc)
            bags = m[0]*vol_ls/bag_kg
            qty_table([
                ("Concreto (1m de ancho)", f"{vol_ls:.4f} m³"),
                (f"Acero flexión {bar_ls} @ {s_use_ls:.1f}cm", f"{peso_flex_ls:.2f} kg/m"),
                (f"Acero temp {bar_ls} @ {s_temp:.1f}cm", f"{peso_temp_ls:.2f} kg/m"),
                (f"Cemento ({bag_kg:.0f}kg/bulto)", f"{m[0]*vol_ls/bag_kg:.1f} bultos"),
                ("Referencia", code["ref"]),
            ])

            if "apu_config" in st.session_state:
                apu = st.session_state.apu_config
                mon = apu.get("moneda", "$")
                st.markdown("---")
                st.success("✅ **Precios actualizados del mercado aplicados (APU).**")
                c_cem = bags * apu.get("cemento", 0)
                c_ace = (peso_flex_ls + peso_temp_ls) * apu.get("acero", 0)
                vol_arena_m3 = (m[2]*vol_ls)/1600
                vol_grava_m3 = (m[3]*vol_ls)/1600
                c_are = vol_arena_m3 * apu.get("arena", 0)
                c_gra = vol_grava_m3 * apu.get("grava", 0)
                total_mat = c_cem + c_ace + c_are + c_gra
                total_dias_mo = ((peso_flex_ls + peso_temp_ls) * 0.04) + (vol_ls * 0.4)
                costo_mo = total_dias_mo * apu.get("costo_dia_mo", 70000)
                costo_directo = total_mat + costo_mo
                herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
                aiu = costo_directo * apu.get("pct_aui", 0.30)
                utilidad = costo_directo * apu.get("pct_util", 0.05)
                iva_v = utilidad * apu.get("iva", 0.19)
                gran_total = costo_directo + herramienta + aiu + iva_v
                st.metric(f"💎 Gran Total Proyecto ({mon})", f"{gran_total:,.2f}")

                with st.expander("📐 Dibujo de Figurado para Taller (Varillas de losa)", expanded=False):
                    hook_len_cm = 12 * db_ls / 10
                    straight_len_cm = ln_ls * 100 - 2 * hook_len_cm
                    fig_l1 = draw_longitudinal_bar(ln_ls*100, straight_len_cm, hook_len_cm, db_ls)
                    st.pyplot(fig_l1)
                    st.caption("Varilla de refuerzo principal con ganchos de 90° en extremos (para losa simplemente apoyada o continua).")
            else:
                st.info("💡 Configure los precios en el expander '💰 APU – Precios en vivo' para ver el presupuesto.")

            # MEMORIA DOCX para Losa
            if st.button("📄 Generar Memoria Losa (DOCX)"):
                # Para la memoria redibujamos fig_mv con fondo blanco para impresión
                import matplotlib.pyplot as plt
                fig_mv_w, (ax_v_w, ax_m_w) = plt.subplots(2, 1, figsize=(6, 5), sharex=True)
                fig_mv_w.patch.set_facecolor('white')
                ax_v_w.set_facecolor('white'); ax_m_w.set_facecolor('white')
                ax_v_w.plot(x_vals, V_vals, color='blue', lw=2)
                ax_v_w.fill_between(x_vals, 0, V_vals, color='blue', alpha=0.1)
                ax_v_w.axhline(0, color='black', lw=1)
                ax_v_w.set_ylabel("Cortante (kN)")
                ax_m_w.plot(x_vals, M_vals, color='red', lw=2)
                ax_m_w.fill_between(x_vals, 0, M_vals, color='red', alpha=0.1)
                ax_m_w.axhline(0, color='black', lw=1)
                ax_m_w.set_ylabel("Momento (kN·m/m)")
                ax_m_w.set_xlabel("Distancia x (m)")
                fig_mv_w.tight_layout()
                import io as _io
                buf_mv = _io.BytesIO()
                fig_mv_w.savefig(buf_mv, format='png', dpi=150, bbox_inches='tight')
                buf_mv.seek(0)

                doc = Document()
                doc.add_heading("Memoria de Diseño de Losa en Una Dirección", 0)
                doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_heading("1. Datos", level=1)
                doc.add_paragraph(f"Luz libre: ln = {ln_ls:.2f} m")
                doc.add_paragraph(f"Espesor losa: h = {h_ls:.1f} cm")
                doc.add_paragraph(f"Recubrimiento: d' = {cov_ls:.1f} cm → d = {d_ls:.1f} cm")
                doc.add_paragraph(f"Materiales: f'c = {fc:.1f} MPa, fy = {fy:.0f} MPa")
                doc.add_paragraph(f"Cargas: wD = {wD_ls:.2f} kN/m², wL = {wL_ls:.2f} kN/m² → wu = {wu_ls:.2f} kN/m²")
                doc.add_heading("2. Refuerzo", level=1)
                doc.add_paragraph(f"Momento último Mu = {Mu_ls_kNm:.2f} kN·m/m")
                doc.add_paragraph(f"Armadura principal: {bar_ls} @ {s_use_ls:.1f} cm → As = {As_prov_ls:.3f} cm²/m")
                doc.add_paragraph(f"Armadura de temperatura: {bar_ls} @ {s_temp:.1f} cm → As_temp = {As_temp:.3f} cm²/m")
                doc.add_heading("3. Verificaciones Normativas", level=1)
                checks = [
                    ("Resistencia a flexión (φMn ≥ Mu)", "CUMPLE" if ok_ls else "NO CUMPLE"),
                    ("Espaciamiento ≤ 3h y ≤ 45 cm", "CUMPLE" if s_use_ls <= s_max_ls else "NO CUMPLE"),
                    ("As_temp ≥ 0.0018·b·h", "CUMPLE" if As_temp >= 0.0018*b_ls*h_ls else "NO CUMPLE"),
                ]
                for desc, res in checks:
                    doc.add_paragraph(f"{desc}: {res}")
                doc.add_paragraph(f"φMn = {phi_Mn_ls:.2f} kN·m/m | Mu = {Mu_ls_kNm:.2f} kN·m/m")
                doc.add_paragraph(f"Referencia: {code['ref']}")
                
                doc.add_heading("4. Diagramas de Momento y Cortante", level=1)
                doc.add_picture(buf_mv, width=Inches(5))
                
                doc.add_heading("5. Cantidades", level=1)
                doc.add_paragraph(f"Concreto: {vol_ls:.4f} m³ por metro de ancho")
                doc.add_paragraph(f"Acero principal: {peso_flex_ls:.2f} kg/m")
                doc.add_paragraph(f"Acero temperatura: {peso_temp_ls:.2f} kg/m")
                doc_mem = _io.BytesIO()
                doc.save(doc_mem)
                doc_mem.seek(0)
                st.download_button("Descargar Memoria Losa", data=doc_mem, file_name=f"Memoria_Losa_{ln_ls:.2f}m.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ══════════════════════════════════════════
# 8. LONGITUD DE DESARROLLO (sin cambios)
# ══════════════════════════════════════════
with st.expander("🔗 Longitud de Desarrollo y Empalmes"):
    st.markdown(f"**Barras rectas a tracción** | Norma: `{code['ref']}`")
    c1,c2 = st.columns(2)
    with c1:
        varillas_ld = list(rebar_dict.keys())
        bar_ld = st.selectbox("Varilla:", varillas_ld, 
                              index=varillas_ld.index(st.session_state.ld_bar) if "ld_bar" in st.session_state and st.session_state.ld_bar in varillas_ld else 0,
                              key="ld_bar")
        db_ld = diam_dict[bar_ld]
        psit_opts = ["1.3 — Barra superior (>30cm betón fresco abajo)","1.0 — Otras posiciones"]
        psi_t = st.selectbox("ψt (posición):", psit_opts, 
                             index=psit_opts.index(st.session_state.ld_psit) if "ld_psit" in st.session_state and st.session_state.ld_psit in psit_opts else 1,
                             key="ld_psit")
        psie_opts = ["1.5 — Ep. y >3db o <6mm cub.","1.2 — Otros epoxy","1.0 — Sin epoxy"]
        psi_e = st.selectbox("ψe (epoxy):", psie_opts, 
                             index=psie_opts.index(st.session_state.ld_psie) if "ld_psie" in st.session_state and st.session_state.ld_psie in psie_opts else 2,
                             key="ld_psie")
    with c2:
        psis_opts = ["0.8 — Barras ≤ #6 ó ≤19mm","1.0 — Barras > #6 ó >19mm"]
        psi_s = st.selectbox("ψs (tamaño):", psis_opts, 
                             index=psis_opts.index(st.session_state.ld_psis) if "ld_psis" in st.session_state and st.session_state.ld_psis in psis_opts else (0 if db_ld <= 19 else 1),
                             key="ld_psis")
        psig_opts = ["0.75 — fy ≤ 420 MPa (ACI 318-19+)","1.0 — fy > 420 MPa"]
        psi_g = st.selectbox("ψg (resistencia):", psig_opts, 
                             index=psig_opts.index(st.session_state.ld_psig) if "ld_psig" in st.session_state and st.session_state.ld_psig in psig_opts else (0 if fy <= 420 else 1),
                             key="ld_psig")
        cb_ld = st.number_input("cb (recubrim. al centro barra) [mm]", 20.0, 100.0, st.session_state.get("ld_cb", 40.0), 2.5, key="ld_cb")
        Ktr_ld = st.number_input("Ktr (refuerzo transversal, 0 si no hay) [mm]", 0.0, 50.0, st.session_state.get("ld_ktr", 0.0), 1.0, key="ld_ktr")

    psit_v = float(psi_t.split("—")[0])
    psie_v = float(psi_e.split("—")[0])
    psis_v = float(psi_s.split("—")[0])
    psig_v = float(psi_g.split("—")[0])

    cb_ktr_db = min((cb_ld+Ktr_ld)/db_ld, 2.5)
    psi_prod = min(psit_v * psie_v, 1.7) * psis_v * psig_v
    ld_mm = (3/40)*(fy/lam/math.sqrt(fc))*(psi_prod/cb_ktr_db)*db_ld
    ld_mm = max(ld_mm, 300)

    ls_A = 1.0*ld_mm
    ls_B = 1.3*ld_mm
    ldh_mm = max((0.02*1*1*1*1*fy)/(lam*math.sqrt(fc))*db_ld, 8*db_ld, 150)

    tab_r, = st.tabs(["📊 Resultados"])
    with tab_r:
        st.markdown(f"**φ no aplica para longitud de desarrollo** | Ref: `{code['ref']}`")
        qty_table([
            ("Varilla", f"{bar_ld} — db = {db_ld:.2f} mm"),
            ("f'c / fy", f"{fc:.1f} MPa / {fy:.0f} MPa"),
            ("(cb+Ktr)/db", f"{cb_ktr_db:.3f} {'✅ ≤2.5' if cb_ktr_db<=2.5 else '→ limitado a 2.5'}"),
            ("ψt × ψe × ψs × ψg", f"{psit_v}×{psie_v}×{psis_v}×{psig_v} = {psit_v*psie_v*psis_v*psig_v:.3f}"),
            ("✅ Nota: ψt×ψe ≤ 1.7", "✅" if psit_v*psie_v<=1.7 else "⚠️ Limitar a 1.7"),
            ("ld (barra recta en tensión)", f"**{ld_mm:.0f} mm = {ld_mm/10:.1f} cm**"),
            ("Empalme Clase A (ld×1.0)", f"{ls_A:.0f} mm = {ls_A/10:.1f} cm"),
            ("Empalme Clase B (ld×1.3)", f"{ls_B:.0f} mm = {ls_B/10:.1f} cm"),
            ("ldh (gancho estándar 90°)", f"{ldh_mm:.0f} mm = {ldh_mm/10:.1f} cm"),
            ("Referencia", f"{code['ref']} / ACI 318-25 Sección 25.4"),
        ])
        if psit_v*psie_v > 1.7:
            st.warning("⚠️ El producto ψt×ψe no puede exceder 1.7 (ACI 318-25 25.4.2.5)")
        st.success(f"✅ ld = {ld_mm:.0f} mm | Empalme B = {ls_B:.0f} mm")

# ══════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════
st.markdown("---")
st.markdown(f"""
> **Suite de Hormigón Armado — Multi-Norma**  
> Norma activa: `{norma_sel}` | Nivel sísmico: `{nivel_sis}`  
> f'c = {fc:.2f} MPa | fy = {fy:.0f} MPa | Ec = {Ec:.0f} MPa | β₁ = {beta1:.3f}  
> **Referencia:** {code['ref']}  
> ⚠️ *Las herramientas son de apoyo para el diseño. Verifique siempre con la norma vigente del país.*
""")