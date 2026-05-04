import streamlit as st

#  Utilidad: Color AutoCAD segun # de cuartos de pulgada 
from normas_referencias import mostrar_referencias_norma
def _color_acero_dxf(db_mm: float) -> int:
    """Retorna color AutoCAD (1-255) segun diametro nominal en mm (ASTM/NTC)."""
    if   db_mm <  7.5: return 6   # #2 1/4"   - Magenta (Prompt v12: no usar color 2 para aceros <16mm)
    elif db_mm < 11.1: return 3   # #3 3/8"   - Verde
    elif db_mm < 14.3: return 4   # #4 1/2"   - Cian
    elif db_mm < 17.5: return 5   # #5 5/8"   - Azul
    elif db_mm < 20.6: return 6   # #6 3/4"   - Magenta
    elif db_mm < 23.8: return 30  # #7 7/8"   - Naranja
    elif db_mm < 27.0: return 1   # #8 1"     - Rojo
    else:              return 10  # #9+ 1-1/8" - Verde oscuro (pesado)
# 
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import plotly.graph_objects as go
import datetime as _dt
try:
    import ifc_export          # BIM export (opcional — requiere ifcopenshell)
    _IFC_AVAILABLE = True
except ImportError:
    ifc_export = None
    _IFC_AVAILABLE = False

# 
# IDIOMA GLOBAL

# ═══════════════════════════════════════════════════════════════════
#  PERSISTENCIA LOCAL v6.1 — 09_Zapatas
# ═══════════════════════════════════════════════════════════════════
import json as _json_persist
import os   as _os_persist
import base64 as _b64_persist

_STATE_FILE = "zapatas_state.json"

_PERSIST_KEYS = [
    # Identidad global
    "empresa", "proyecto", "ingeniero",
    "dxfelaboro", "dxfreviso", "dxfaprobo",
    "normasel", "idioma",
    "userrole", "user_role",   # ambas claves para compatibilidad entre modulos
    # Geotecnia / zapata aislada
    "zfc", "zfy", "zphi", "cohval", "cohunit",
    "gamval", "gamunit", "zgsat",
    "zap_b", "zap_l", "zhzcp", "zap_df", "zbcolcp",
    "zQact", "znf", "zfs", "zspt",
    "zbousB", "zbousL", "zbousZ",
    "zH_horiz", "zdelta",
    "zopthmin", "zopthmax",
    "znombreregistro",
    # Diseño estructural
    "poscoliso",
    # APU
    "apuconfig",
]

def _save_state_zap():
    try:
        snap = {}
        for k in _PERSIST_KEYS:
            val = st.session_state.get(k)
            if val is None:
                continue
            if k == "logobytes" and isinstance(val, (bytes, bytearray)):
                snap[k] = {"type": "bytes_b64", "data": _b64_persist.b64encode(val).decode()}
            else:
                try:
                    _json_persist.dumps(val)
                    snap[k] = val
                except (TypeError, ValueError):
                    snap[k] = str(val)
        with open(_STATE_FILE, "w", encoding="utf-8") as _f:
            _json_persist.dump(snap, _f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def _load_state_zap():
    if not _os_persist.path.exists(_STATE_FILE):
        return
    try:
        with open(_STATE_FILE, "r", encoding="utf-8") as _f:
            data = _json_persist.load(_f)
        for k, v in data.items():
            if k not in st.session_state:
                if isinstance(v, dict) and v.get("type") == "bytes_b64":
                    st.session_state[k] = _b64_persist.b64decode(v["data"])
                else:
                    st.session_state[k] = v
    except Exception:
        pass

_load_state_zap()

lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es
# 


# ═══════════════════════════════════════════════════════════════════
#  IDENTIDAD GLOBAL v6.1
# ═══════════════════════════════════════════════════════════════════
def _get_identity():
    """Lee campos de identidad desde session_state con fallback a guión bajo."""
    empresa   = st.session_state.get("empresa",   "") or "________________"
    proyecto  = st.session_state.get("proyecto",  "") or "________________"
    ingeniero = st.session_state.get("ingeniero", "") or "________________"
    elab = st.session_state.get("dxfelaboro", st.session_state.get("dxf_elaboro", "")) or "________________"
    rev  = st.session_state.get("dxfreviso",  st.session_state.get("dxf_reviso",  "")) or "________________"
    apb  = st.session_state.get("dxfaprobo",  st.session_state.get("dxf_aprobo",  "")) or "________________"
    logo = st.session_state.get("logobytes",  st.session_state.get("logo_bytes"))
    return empresa, proyecto, ingeniero, elab, rev, apb, logo


def _rol_badge():
    """Muestra badge de rol en sidebar y devuelve (rol, puede_exportar).
    Lee 'user_role' (clave de Columnas) o 'userrole' (clave legacy de Zapatas)
    para que el rol establecido en cualquier modulo sea reconocido globalmente.
    """
    # Prioridad: user_role (Columnas/global) → userrole (legacy Zapatas) → free
    rol = (
        st.session_state.get("user_role")
        or st.session_state.get("userrole")
        or "free"
    )
    # Sincronizar ambas claves para que otros modulos vean el mismo rol
    st.session_state["user_role"] = rol
    st.session_state["userrole"]  = rol

    color = {"admin": "#1b5e20", "pro": "#0d47a1", "free": "#b71c1c"}.get(rol, "#b71c1c")
    st.sidebar.markdown(
        f"<span style='background:{color};color:white;padding:2px 8px;"
        f"border-radius:4px;font-size:11px'>{rol.upper()}</span>",
        unsafe_allow_html=True,
    )
    puede_exportar = rol in ("admin", "pro")
    return rol, puede_exportar


# ═══════════════════════════════════════════════════════════════════
#  HELPER: figura matplotlib con fondo blanco para DOCX
# ═══════════════════════════════════════════════════════════════════
def _sanitize_dxf_text(txt: str) -> str:
    """Elimina Em-Dashes y caracteres no válidos para DXF (Prompt Maestro v12)."""
    return str(txt).replace('\u2014', '-').replace('\u2013', '-').replace('--', '-').replace('\u2019', "'")


def _fig_to_docx_white(fig):
    """Devuelve bytes PNG con fondo blanco, apto para embeber en DOCX."""
    orig_fc   = fig.get_facecolor()
    orig_axes = [(ax, ax.get_facecolor()) for ax in fig.get_axes()]
    fig.patch.set_facecolor("white")
    for ax, _ in orig_axes:
        ax.set_facecolor("#f8f9fa")
        ax.tick_params(colors="#1a1a1a")
        ax.xaxis.label.set_color("#1a1a1a")
        ax.yaxis.label.set_color("#1a1a1a")
        for spine in ax.spines.values():
            spine.set_edgecolor("#cccccc")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, facecolor="white",
                transparent=False, bbox_inches="tight")
    buf.seek(0)
    img_bytes = buf.read()
    fig.patch.set_facecolor(orig_fc)
    for ax, fc in orig_axes:
        ax.set_facecolor(fc)
    return img_bytes

st.set_page_config(page_title=_t("Zapatas y Suelos", "Footings and Soils"), layout="wide")

# 
# PERSISTENCIA SUPABASE
# 
import requests
import json
import streamlit as st

def get_supabase_rest_info():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return url, key
    except Exception as e:
        return None, None

def guardar_proyecto_supabase(nombre, estado_dict, ref):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = f"db_proyectos_{ref.lower()}.json"
            db = {}
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            db[f"[{ref.upper()}] {nombre}"] = {"nombre_proyecto": f"[{ref.upper()}] {nombre}", "estado_json": json.dumps(estado_dict)}
            with open(db_path, "w", encoding="utf-8") as f: json.dump(db, f)
            return True, "Proyecto guardado (Local)"
        except Exception as e:
            return False, f"Error local: {e}"
    
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "Prefer": "resolution=merge-duplicates"
    }
    payload = {
        "nombre_proyecto": f"[{ref.upper()}] {nombre}",
        "user_id": st.session_state.get("user_id", "anonimo"),
        "estado_json": json.dumps(estado_dict),
    }
    try:
        endpoint = f"{url}/rest/v1/proyectos?on_conflict=nombre_proyecto"
        res = requests.post(endpoint, headers=headers, json=payload, timeout=5)
        if res.status_code in [200, 201, 204]: return True, "Proyecto en la nube"
        return False, f"Error API: {res.text}"
    except Exception as e: return False, f"Error API: {e}"

def cargar_proyecto_supabase(nombre, ref):
    url, key = get_supabase_rest_info()
    full_name = f"[{ref.upper()}] {nombre}" if not nombre.startswith("[") else nombre
    if not url or not key:
        try:
            import os
            db_path = f"db_proyectos_{ref.lower()}.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                match = db.get(full_name) or db.get(nombre)
                if match:
                    estado = json.loads(match["estado_json"])
                    for k, v in estado.items(): st.session_state[k] = v
                    return True, "Proyecto cargado (Local)"
            return False, "No encontrado"
        except Exception as e: return False, str(e)
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        endpoint = f"{url}/rest/v1/proyectos?nombre_proyecto=eq.{full_name}&select=*"
        res = requests.get(endpoint, headers=headers, timeout=5)
        if res.status_code == 200:
            data = res.json()
            if data:
                estado = json.loads(data[0]["estado_json"])
                for k, v in estado.items(): st.session_state[k] = v
                return True, "Proyecto cargado"
        return False, "No encontrado"
    except Exception as e: return False, str(e)

def listar_proyectos_supabase(ref):
    url, key = get_supabase_rest_info()
    if not url or not key:
        try:
            import os
            db_path = f"db_proyectos_{ref.lower()}.json"
            if os.path.exists(db_path):
                with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
                return sorted([k.replace(f"[{ref.upper()}] ", "") for k in db.keys()])
            return []
        except: return []
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        endpoint = f"{url}/rest/v1/proyectos?select=nombre_proyecto"
        res = requests.get(endpoint, headers=headers, timeout=5)
        if res.status_code == 200:
            return sorted([i["nombre_proyecto"].replace(f"[{ref.upper()}] ", "") for i in res.json() if f"[{ref.upper()}]" in i.get("nombre_proyecto","")])
        return []
    except: return []

def capturar_estado_module(ref):
    if ref == "zapatas":
        match = ["z_", "cp_", "apu_", "z_tipo_idx", "fc_basico", "fy_basico"]
    else:
        match = ["pi_", "p_", "apu_", "fc_pi"]
    return {k: st.session_state[k] for k in st.session_state if any(k.startswith(m) for m in match)}

# 


# 1. Banner SVG premium
st.markdown("""<div style="width:100%;overflow:hidden;border-radius:14px;margin-bottom:20px;box-shadow:0 4px 32px #0008;">
<svg viewBox="0 0 1100 220" xmlns="http://www.w3.org/2000/svg" style="width:100%;display:block;background:linear-gradient(135deg,#110a01 0%,#241503 100%);">
  <g opacity="0.08" stroke="#fbbf24" stroke-width="0.5">
    <line x1="0" y1="55" x2="1100" y2="55"/><line x1="0" y1="110" x2="1100" y2="110"/>
    <line x1="0" y1="165" x2="1100" y2="165"/>
    <line x1="220" y1="0" x2="220" y2="220"/><line x1="440" y1="0" x2="440" y2="220"/>
    <line x1="660" y1="0" x2="660" y2="220"/>
  </g>
  <rect x="0" y="0" width="1100" height="3" fill="#f59e0b" opacity="0.9"/>
  <rect x="0" y="217" width="1100" height="3" fill="#fbbf24" opacity="0.7"/>
  <!-- ZAPATA CONCENTRICA -->
  <g transform="translate(60,60)">
    <text x="50" y="-20" text-anchor="middle" font-family="sans-serif" font-size="10" font-weight="600" fill="#d1d5db" letter-spacing="1">AISLADA</text>
    <!-- Columna -->
    <rect x="35" y="0" width="30" height="40" fill="#2d3748" stroke="#4a5568" stroke-width="2"/>
    <path d="M35,10 L65,10 M35,20 L65,20 M35,30 L65,30" stroke="#f59e0b" stroke-width="1.5" opacity="0.8"/>
    <!-- Zapata -->
    <path d="M10,40 L90,40 L100,70 L0,70 Z" fill="#1e293b" stroke="#64748b" stroke-width="2" stroke-linejoin="round"/>
    <line x1="5" y1="65" x2="95" y2="65" stroke="#38bdf8" stroke-width="2"/>
    <circle cx="20" cy="65" r="2.5" fill="#f59e0b"/><circle cx="40" cy="65" r="2.5" fill="#f59e0b"/><circle cx="60" cy="65" r="2.5" fill="#f59e0b"/><circle cx="80" cy="65" r="2.5" fill="#f59e0b"/>
    <!-- Presion -->
    <g stroke="#ef4444" stroke-width="1.2">
      <line x1="15" y1="85" x2="15" y2="75"/><polygon points="15,75 12,80 18,80" fill="#ef4444"/>
      <line x1="50" y1="85" x2="50" y2="75"/><polygon points="50,75 47,80 53,80" fill="#ef4444"/>
      <line x1="85" y1="85" x2="85" y2="75"/><polygon points="85,75 82,80 88,80" fill="#ef4444"/>
    </g>
    <text x="50" y="98" text-anchor="middle" font-family="monospace" font-size="9" fill="#f87171">qmax</text>
  </g>
  <!-- ZAPATA MEDIANERA & VIGA -->
  <g transform="translate(250,60)">
    <text x="85" y="-20" text-anchor="middle" font-family="sans-serif" font-size="10" font-weight="600" fill="#d1d5db" letter-spacing="1">MEDIANERA + AMARRE</text>
    <!-- Zapata 1 (Med) -->
    <rect x="0" y="0" width="30" height="40" fill="#2d3748" stroke="#4a5568" stroke-width="2"/>
    <path d="M0,40 L70,40 L70,70 L0,70 Z" fill="#1e293b" stroke="#64748b" stroke-width="2"/>
    <!-- Viga -->
    <rect x="30" y="45" width="90" height="20" fill="#334155" stroke="#475569" stroke-width="1.5" stroke-dasharray="2,2"/>
    <!-- Zapata 2 (Ais) -->
    <rect x="120" y="0" width="30" height="40" fill="#2d3748" stroke="#4a5568" stroke-width="2"/>
    <path d="M100,40 L170,40 L180,70 L90,70 Z" fill="#1e293b" stroke="#64748b" stroke-width="2"/>
    <!-- Reb -->
    <line x1="5" y1="65" x2="65" y2="65" stroke="#38bdf8" stroke-width="1.5"/>
    <line x1="100" y1="65" x2="170" y2="65" stroke="#38bdf8" stroke-width="1.5"/>
    <line x1="30" y1="60" x2="120" y2="60" stroke="#f59e0b" stroke-width="1.5"/>
    <!-- Linea de Propiedad -->
    <line x1="-5" y1="-10" x2="-5" y2="100" stroke="#f43f5e" stroke-width="1.5" stroke-dasharray="5,3"/>
    <text x="-12" y="10" text-anchor="end" font-family="monospace" font-size="8" fill="#fb7185" transform="rotate(-90,-12,10)">LINDERO</text>
  </g>
  <!-- TEXT BLOCK -->
  <g transform="translate(560,0)">
    <rect x="0" y="28" width="4" height="165" rx="2" fill="#f59e0b"/>
    <text x="18" y="66" font-family="Arial,sans-serif" font-size="30" font-weight="bold" fill="#ffffff">ZAPATAS Y GEOTECNIA</text>
    <text x="18" y="92" font-family="Arial,sans-serif" font-size="17" font-weight="300" fill="#fcd34d" letter-spacing="2">CIMENTACIONES NSR-10 / ACI 318</text>
    <rect x="18" y="100" width="480" height="1" fill="#f59e0b" opacity="0.5"/>
    <!-- Tags -->
    <rect x="18" y="111" width="130" height="22" rx="11" fill="#291400" stroke="#f59e0b" stroke-width="1"/>
    <text x="83" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#fbbf24">CAPACIDAD PORTANTE</text>
    <rect x="156" y="111" width="90" height="22" rx="11" fill="#1e1b2e" stroke="#8b5cf6" stroke-width="1"/>
    <text x="201" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#a78bfa">PUNZONAMIENTO</text>
    <rect x="254" y="111" width="106" height="22" rx="11" fill="#1c1416" stroke="#ef4444" stroke-width="1"/>
    <text x="307" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#f87171">FLEXO-COMPRESION</text>
    <rect x="368" y="111" width="88" height="22" rx="11" fill="#0c1f2e" stroke="#38bdf8" stroke-width="1"/>
    <text x="412" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#7dd3fc">PLANO DXF / IFC</text>
    <!-- Description -->
    <text x="18" y="156" font-family="Arial,sans-serif" font-size="11" fill="#9ca3af">Analisis de interaccion suelo-estructura para zapatas aisladas y medianeras.</text>
    <text x="18" y="172" font-family="Arial,sans-serif" font-size="11" fill="#9ca3af">Calculo Terzaghi, Meyerhof o Vesic con exportacion BIM de refuerzo tridimensional,</text>
    <text x="18" y="188" font-family="Arial,sans-serif" font-size="11" fill="#9ca3af">verificando excentricidades biaxiales, presiones admisibles, momentos y cortante.</text>
  </g>
</svg></div>""", unsafe_allow_html=True)

# 2. Panel global "Guía Completa de Uso"
with st.expander(" ¿Cómo usar este módulo? — Guía Completa de Uso", expanded=False):
    st.markdown('''
    ### Guia Completa — Zapatas y Geotecnia

    Konte calcula y diseña zapatas **Aisladas**, **Medianeras** y sistemas **Medianera + Viga de Amarre**,
    cubriendo desde la capacidad portante del suelo hasta el despiece de acero 3D, la memoria DOCX y el plano DXF/IFC.
    El modulo es **MULTINORMA**: selecciona la norma activa en la barra lateral y todos los parametros
    (phi, rho_min, rho_max, factores de capacidad, recubrimientos) se ajustan automaticamente.

    ---

    #### Normas soportadas
    | Norma | Pais | rho_min | phi_flexion | phi_cortante |
    |---|---|---|---|---|
    | NSR-10 | Colombia | 0.0018 | 0.90 | 0.75 |
    | ACI 318-25 / 19 / 14 | EE.UU. / Mexico / Peru | 0.0018 | 0.90 | 0.75 |
    | E.060 | Peru | 0.0018 | 0.90 | 0.85 |
    | NEC-SE-HM | Ecuador | 0.0018 | 0.90 | 0.75 |
    | NTC-EM | Mexico | 0.0020 | 0.90 | 0.75 |
    | COVENIN 1753 | Venezuela | 0.0018 | 0.90 | 0.75 |
    | NB 1225001 | Bolivia | 0.0018 | 0.90 | 0.75 |
    | CIRSOC 201 | Argentina | 0.0018 | 0.90 | 0.75 |

    ---

    #### Paso 1 — Configuracion Global (Barra Lateral)
    - **Norma activa**: cambia phi, rho_min, rho_max y recubrimientos simultaneamente.
    - **Materiales**: fc (MPa) y fy (MPa) de la zapata.
    - **Refuerzo**: sistema pulgadas (US) o mm y diametro de varilla.
    - **Identidad**: empresa, ingeniero, logotipo se estampan en todos los entregables.

    #### Paso 2 — Tipo de Zapata
    | Tipo | Cuando usarlo |
    |---|---|
    | 1. Aislada Centrica | Columna centrada sin restriccion de lindero |
    | 2. Medianera (Lindero) | Columna junto a un lindero — excentricidad forzada en dir. B |
    | 3. Esquina (Dos linderos) | Columna en esquina — excentricidad biaxial |
    | 4. Con Viga de Amarre | Sistema medianera + zapata interior + viga strap beam |

    #### Paso 3 — Flujo de Calculo (Zapata Aislada)
    1. Ingresa B, L, H (m), Df, pedestal cx/cy y cargas Pu, Mu_B, Mu_L.
    2. El modulo verifica capacidad portante (Terzaghi, Meyerhof o Vesic) con correccion por nivel freatico.
    3. Verifica presiones qmax, excentricidad (e < B/6), cortante unidireccional y punzonamiento.
    4. Diseña el acero por flexion en ambas direcciones (simple o doble parrilla automaticamente).
    5. Genera entregables: Resultados | Plano DXF | Cantidades APU.

    #### Paso 4 — Sistema Medianera + Viga de Amarre (Tipo 4)
    Tabs disponibles: **Geometria** | **Viga de Amarre** | **Vista 3D** | **Plano DXF** | **Memoria DOCX**

    ---

    #### Cuando se activa Doble Parrilla de Acero (automatico)

    La zapata normalmente usa **una sola parrilla inferior** (zona de traccion). El modulo activa
    **doble parrilla** (inferior + superior comprimida) cuando detecta alguna de estas condiciones:

    **Condicion 1 — Cuantia maxima superada (NSR-10 C.10.3.5 / ACI 318 Table 21.2.2):**
    La cuantia requerida supera rho_max = 0.75 * rho_bal, donde:
    rho_bal = (0.85 * beta1 * fc / fy) * 600 / (600 + fy)
    beta1 = max(0.65, 0.85 - 0.05*(fc-28)/7) — aplica igual para todas las normas soportadas.

    **Condicion 2 — Discriminante negativo:**
    El espesor H propuesto es insuficiente para resistir Mu como seccion simple
    (discriminante 1 - 2*Rn/(0.85*fc) < 0).

    **Calculo cuando se activa doble parrilla:**
    - As_inf (cara inferior, traccion): rho = rho_max → As_inf = rho_max * b * d
    - As_sup (cara superior, compresion): As_sup = delta_Mu / (phi * fy * (d - d'))
      donde delta_Mu = Mu - phi*Mn_max_simple y d' = recubrimiento + db/2
    - Separacion minima entre capas: NSR-10 C.7.6.1 / ACI 318-25.2.1 (min 2.5 cm o 1.33*Dmax_agr)

    **Indicadores visuales de doble parrilla:**
    - Tabla de resultados: Estado muestra "DOBLE PARRILLA" en lugar de "OK"
    - Aviso amarillo con detalle de As_inf y As_sup por direccion
    - Vista 3D: barras superiores en ROSA (dir. B) y MAGENTA (dir. L) con linea discontinua
    - Plano DXF: capa ACERO_SUP (color magenta), lineas punteadas en planta y solidas en cortes
    - Memoria DOCX: secciones 5.1 Dir B y 5.2 Dir L con calculo completo de ambas parrillas

    **Recomendacion practica:** Si se activa doble parrilla, aumentar primero H (espesor)
    o fc — en la mayoria de casos esto elimina la necesidad sin incrementar el acero total.

    ---

    #### Verificaciones Normativas Cubiertas
    - NSR-10 C.10.2.7 / ACI 318 Table 22.2.2 — beta1 y bloque de compresion
    - NSR-10 C.10.3.5 / ACI 318 Table 21.2.2 — rho_max = 0.75*rho_bal
    - NSR-10 C.11.11 / ACI 318-22.5 — Cortante unidireccional (viga ancha)
    - NSR-10 C.11.12 / ACI 318-22.6 — Punzonamiento bidireccional
    - NSR-10 C.12.2 / ACI 318-25.5 — Longitud de desarrollo recto
    - NSR-10 C.12.5 / ACI 318-25.3 — Gancho estandar 90 grados
    - NSR-10 C.7.6.1 / ACI 318-25.2.1 — Separacion minima entre capas (doble parrilla)
    - NSR-10 H.3 — Cimentacion superficial vs profunda
    - Terzaghi 1943 / Meyerhof 1963 / Vesic 1973 — Capacidad portante
    - Bowles 1996 — Coeficiente de balasto ks (Winkler)
    ''')


# 
# PIE DE PÁGINA / DERECHOS RESERVADOS
# 
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    <br><br>
    <i> Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# 
# CONFIGURACIÓN GENERAL Y APU
# 
st.sidebar.header(_t(" Configuración Global", " Global Settings"))

# ── Roles v6.1 ──────────────────────────────────────────────────────────────
_rol, _puede_exportar = _rol_badge()
if _rol == "free":
    st.sidebar.warning(
        "Modo Profesional requerido para exportar DOCX, DXF e IFC. "
        "Contacta al administrador para activar tu licencia."
    )
elif _rol == "admin":
    st.sidebar.info("Modo Administrador activo — todos los entregables habilitados.")
# ─────────────────────────────────────────────────────────────────────────────
if "norma_sel" not in st.session_state:
    st.session_state.norma_sel = "NSR-10 (Colombia)"

norma_sel = st.session_state.norma_sel
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")
st.sidebar.markdown(
    f'<div style="background:#1e3a1e;border-radius:6px;padding:8px 12px;margin-bottom:4px;">'
    f'<img src="https://flagpedia.net/data/flags/mini/{_iso}.png" style="vertical-align:middle;margin-right:8px;">'
    f'<span style="color:#7ec87e;font-weight:600;font-size:13px;">{_t("Norma Activa:","Active Code:")} {norma_sel}</span>'
    f'</div>', unsafe_allow_html=True
)

# Parametros normativos base
# ── Factores phi por norma (NSR-10/ACI/E.060/NEC/NTC/COVENIN/CIRSOC) ───────
_PHI = {
    "NSR-10 Colombia":             {"v": 0.75, "f": 0.90},
    "ACI 318-25 EE.UU.":          {"v": 0.75, "f": 0.90},
    "ACI 318-19 EE.UU.":          {"v": 0.75, "f": 0.90},
    "ACI 318-14 EE.UU.":          {"v": 0.75, "f": 0.90},
    "E.060 Peru":                  {"v": 0.85, "f": 0.90},
    "NEC-SE-HM Ecuador":           {"v": 0.75, "f": 0.90},
    "NTC-EM México":               {"v": 0.75, "f": 0.90},
    "COVENIN 1753-2006 Venezuela": {"v": 0.75, "f": 0.90},
    "NB 1225001-2020 Bolivia":     {"v": 0.75, "f": 0.90},
    "CIRSOC 201-2025 Argentina":   {"v": 0.75, "f": 0.90},
}
_phi_vals = _PHI.get(norma_sel, {"v": 0.75, "f": 0.90})
phi_v = _phi_vals["v"]  # Cortante
phi_f = _phi_vals["f"]  # Flexión

fc_basico = st.sidebar.number_input(_t("f'c Zapata [MPa]", "f'c Footing [MPa]"), 15.0, 50.0, st.session_state.get("z_fc", 21.0), 1.0, key="z_fc")
fy_basico = st.sidebar.number_input(_t("fy Acero [MPa]", "fy Steel [MPa]"), 240.0, 500.0, st.session_state.get("z_fy", 420.0), 10.0, key="z_fy")

#  CONVERSOR GLOBAL DE UNIDADES DE SUELO ← Visible en toda la página
st.sidebar.markdown("---")
st.sidebar.header(_t(" Conversor Unidades de Suelo", " Soil Units Converter"))
_cu = st.sidebar.selectbox(_t("Unidad a convertir:", "Unit to convert:"), 
    ["kPa → ...", "ton/m² → ...", "kg/cm² → ...", "MPa → ...", "psi → ..."], key="conv_unit_global")
_cv = st.sidebar.number_input(_t("Valor:", "Value:"), value=1.0, step=0.1, key="conv_val_global")
_ck = {
    "kPa → ...": _cv, "ton/m² → ...": _cv * 9.80665,
    "kg/cm² → ...": _cv * 98.0665, "MPa → ...": _cv * 1000.0,
    "psi → ...": _cv * 6.89476,
}.get(_cu, _cv)
st.sidebar.markdown(f"""
<div style="background:#223;padding:10px;border-radius:5px;text-align:center;">
    <small>{_cu.split('→')[0].strip()}</small><br>
    <b>{_cv:.2f}</b><br>
    <span style="color:#0f0;">⬇</span><br>
    <b style="color:#4caf50;">{_ck:.2f} kPa</b><br><small>({_ck/9.80665:.2f} t/m² | {_ck/98.0665:.2f} kg/cm²)</small>
</div>
""", unsafe_allow_html=True)

try:
    mostrar_referencias_norma(norma_sel, "zapatas")
except NameError:
    pass
    st.sidebar.subheader(" Guardar / Cargar Proyecto")
    nombre_producido = st.session_state.get(f"np_{"zapatas"}", "")
    
    st.sidebar.markdown("**Nuevo Proyecto / Guardar**")
    nombre_proy_guardar = st.sidebar.text_input("Nombre para guardar", value=nombre_producido, key=f"input_guardar_{"zapatas"}")
    
    if st.sidebar.button(" Guardar Proyecto", use_container_width=True, key=f"btn_save_{"zapatas"}"):
        if nombre_proy_guardar:
            ok, msg = guardar_proyecto_supabase(nombre_proy_guardar, capturar_estado_module("zapatas"), "zapatas")
            if ok:
                st.session_state[f"np_{"zapatas"}"] = nombre_proy_guardar
                st.sidebar.success(msg)
            else:
                st.sidebar.error(msg)
        else:
            st.sidebar.warning("Escribe un nombre de proyecto")

    st.sidebar.markdown("**Cargar Proyecto Existente**")
    lista_proyectos = listar_proyectos_supabase("zapatas")
    if lista_proyectos:
        idx_def = lista_proyectos.index(nombre_producido) if nombre_producido in lista_proyectos else 0
        nombre_proy_cargar = st.sidebar.selectbox("Selecciona un proyecto", lista_proyectos, index=idx_def, key=f"sel_load_{"zapatas"}")
        
        def on_cargar_click():
            proy = st.session_state[f"sel_load_{"zapatas"}"]
            if proy:
                ok, msg = cargar_proyecto_supabase(proy, "zapatas")
                st.session_state[f"__msg_cargar_{"zapatas"}"] = (ok, msg)
                if ok: st.session_state[f"np_{"zapatas"}"] = proy

        st.sidebar.button(" Cargar", on_click=on_cargar_click, use_container_width=True, key=f"btn_load_{"zapatas"}")

        if f"__msg_cargar_{"zapatas"}" in st.session_state:
            ok, msg = st.session_state.pop(f"__msg_cargar_{"zapatas"}")
            if ok: st.sidebar.success(msg)
            else: st.sidebar.error(msg)

st.sidebar.markdown(f"""
| | |
|:---|---:|
| **kPa** | `{_ck:.2f}` |
| **ton/m²** | `{_ck/9.80665:.3f}` |
| **kg/cm²** | `{_ck/98.0665:.4f}` |
| **MPa** | `{_ck/1000:.5f}` |
| **psi** | `{_ck/6.89476:.2f}` |
""")

if any(n in norma_sel for n in ["Colombia", "EE.UU.", "Perú", "México", "Venezuela"]):
    REBAR_DICT = {
        "N.3 (3/8\")": {"area": 0.71, "db": 9.5},
        "N.4 (1/2\")": {"area": 1.29, "db": 12.7},
        "N.5 (5/8\")": {"area": 1.99, "db": 15.9},
        "N.6 (3/4\")": {"area": 2.84, "db": 19.1},
        "N.7 (7/8\")": {"area": 3.87, "db": 22.2},
        "N.8 (1\")":   {"area": 5.10, "db": 25.4},
    }
    def_idx = 1 # N.4
else:
    REBAR_DICT = {
        "10 mm": {"area": 0.785, "db": 10.0},
        "12 mm": {"area": 1.131, "db": 12.0},
        "14 mm": {"area": 1.539, "db": 14.0},
        "16 mm": {"area": 2.011, "db": 16.0},
        "18 mm": {"area": 2.545, "db": 18.0},
        "20 mm": {"area": 3.142, "db": 20.0},
        "22 mm": {"area": 3.801, "db": 22.0},
        "25 mm": {"area": 4.909, "db": 25.0},
    }
    def_idx = 1 # 12 mm

# 
# HELPER GLOBAL: BOUSSINESQ INFLUENCE FACTOR (version escalar y vectorizada)
# 
def I_z_bous(m, n):
    V1 = m**2 + n**2 + 1
    V2 = m**2 * n**2
    term1 = (2*m*n*np.sqrt(V1)) / (V1 + V2) if (V1 + V2) != 0 else 0
    term2 = (V1 + 1) / V1 if V1 != 0 else 0
    angulo = np.arctan2(2*m*n*np.sqrt(V1), (V1 - V2))
    # Manejo de ángulo para cuadrantes
    angulo = np.where(V1 - V2 < 0, angulo + np.pi, angulo)
    return (1 / (4 * np.pi)) * (term1 * term2 + angulo)

# Versión vectorizada para arrays de numpy
def I_z_bous_vec(m_arr, n_arr):
    V1 = m_arr**2 + n_arr**2 + 1
    V2 = m_arr**2 * n_arr**2
    term1 = (2*m_arr*n_arr*np.sqrt(V1)) / (V1 + V2)
    term1 = np.where((V1 + V2) != 0, term1, 0.0)
    term2 = (V1 + 1) / V1
    term2 = np.where(V1 != 0, term2, 0.0)
    angulo = np.arctan2(2*m_arr*n_arr*np.sqrt(V1), (V1 - V2))
    angulo = np.where(V1 - V2 < 0, angulo + np.pi, angulo)
    return (1 / (4 * np.pi)) * (term1 * term2 + angulo)

# 
# OPCIÓN B — FUNCIONES AUXILIARES GLOBALES (Fases 1-5)
# 

#  TABLA MULTINORMA VIGA DE AMARRE 
_STRAP_NORM = {
    "NSR-10 (Colombia)":          {"nombre": "Viga de amarre",    "art": "C.15.13.3",     "h_factor": 20, "F_ax_pct": 0.025, "idioma": "es"},
    "ACI 318-25 (EE.UU.)":        {"nombre": "Strap beam",        "art": "§13.3.3",       "h_factor": 15, "F_ax_pct": 0.025, "idioma": "en"},
    "ACI 318-19 (EE.UU.)":        {"nombre": "Strap beam",        "art": "§13.3.3",       "h_factor": 15, "F_ax_pct": 0.025, "idioma": "en"},
    "ACI 318-14 (EE.UU.)":        {"nombre": "Strap beam",        "art": "§13.3.3",       "h_factor": 15, "F_ax_pct": 0.025, "idioma": "en"},
    "E.060 (Perú)":               {"nombre": "Viga de conexión",  "art": "§15.13",        "h_factor": 20, "F_ax_pct": 0.025, "idioma": "es"},
    "NTC-EM (México)":            {"nombre": "Viga de liga",      "art": "§5.4",          "h_factor": 20, "F_ax_pct": 0.050, "idioma": "es"},
    "NEC-SE-HM (Ecuador)":        {"nombre": "Viga de amarre",    "art": "§5.5",          "h_factor": 20, "F_ax_pct": 0.025, "idioma": "es"},
    "COVENIN 1753-2006 (Venezuela)":{"nombre": "Viga de conexión","art": "§C.15.13",      "h_factor": 20, "F_ax_pct": 0.025, "idioma": "es"},
    "NB 1225001-2020 (Bolivia)":  {"nombre": "Viga de amarre",    "art": "§15.13",        "h_factor": 20, "F_ax_pct": 0.025, "idioma": "es"},
    "CIRSOC 201-2025 (Argentina)":{"nombre": "Viga de conexión",  "art": "§C.15.13",      "h_factor": 20, "F_ax_pct": 0.025, "idioma": "es"},
}

def _get_strap_norm(norma):
    """Retorna los parámetros normativos de la viga de amarre según la norma activa."""
    for key in _STRAP_NORM:
        if key in norma or norma in key:
            return _STRAP_NORM[key]
    return _STRAP_NORM["NSR-10 (Colombia)"]  # fallback

#  F1: DISTRIBUCIÓN DE PRESIONES 
def calcular_distribucion_presiones(P, M_B, M_L, B, L):
    """
    Retorna qu_max, qu_min, tipo_dist, e_B, e_L para cualquier tipo de zapata.
    Convención ACI 318 / NSR-10:
      e_B = |M_B / P|  → excentricidad en dirección B (M_B voltea sobre el eje perpendicular a B)
      e_L = |M_L / P|  → excentricidad en dirección L
    """
    A = B * L
    Ix = B * L**3 / 12.0  # inercia respecto al eje paralelo a B (para M_B)
    Iy = L * B**3 / 12.0  # inercia respecto al eje paralelo a L (para M_L)
    # EC-1 CORREGIDO: M_B → e_B,  M_L → e_L  (no cruzado)
    e_B = abs(M_B / P) if P > 0 else 0.0   # excentricidad en dir. B (causada por M_B)
    e_L = abs(M_L / P) if P > 0 else 0.0   # excentricidad en dir. L (causada por M_L)
    # Presiones en las 4 esquinas con signo correcto
    def q_esq(sx, sy):
        # sx multiplica la distancia en dir. B (±B/2), sy en dir. L (±L/2)
        return P/A + (M_B * sx * B/2) / Iy + (M_L * sy * L/2) / Ix
    corners = [q_esq(sx, sy) for sx in [1, -1] for sy in [1, -1]]
    qu_max = max(corners); qu_min = min(corners)
    if qu_min >= 0:
        tipo_dist = "trapezoidal"
    elif e_B > B/6 or e_L > L/6:
        tipo_dist = "triangular"
    else:
        tipo_dist = "con_tension"
    return qu_max, qu_min, tipo_dist, e_B, e_L, Ix, Iy, A

#  F2: ÁREA EFECTIVA MEYERHOF 
def calcular_area_efectiva_meyerhof(B, L, e_B, e_L):
    """Retorna B', L', A' para el caso qu_min < 0."""
    B_prima = max(0.1, B - 2 * e_B)
    L_prima = max(0.1, L - 2 * e_L)
    A_prima = B_prima * L_prima
    return B_prima, L_prima, A_prima

#  F3: MAPA DE CALOR 3D PRESIONES 
def render_mapa_calor_presiones(B, L, P, M_B, M_L, Ix, Iy, A, qu_max_val, qu_min_val, c1_col_cm, c2_col_cm, titulo="", q_adm=None):
    """Genera go.Figure con surface 3D de qu(x,y)."""
    _nx, _ny = 30, 30
    _xg = np.linspace(-B/2, B/2, _nx)
    _yg = np.linspace(-L/2, L/2, _ny)
    _Xm, _Ym = np.meshgrid(_xg, _yg)
    _Qm = P/A + (M_L * _Xm / Iy) + (M_B * _Ym / Ix)
    _ten = bool(np.any(_Qm < 0))
    _Qp = np.where(_Qm < 0, 0, _Qm)
    fig = go.Figure()
    _rango = max(float(np.max(_Qp)) - float(np.min(_Qp[_Qp > 0]) if np.any(_Qp > 0) else 0), 1.0)
    _cmin  = max(0.0, float(np.max(_Qp)) - _rango * 4)
    if _rango < 5.0:
        _cmin = 0.0

    fig.add_trace(go.Surface(x=_xg, y=_yg, z=_Qp, colorscale="RdYlGn_r",
                             cmin=_cmin, cmax=float(np.max(_Qp))*1.02,
                             colorbar=dict(title=dict(text="qu [kPa]", font=dict(color="white")), tickfont=dict(color="white")),
                             opacity=0.9, name="qu(x,y)"))
    
    rango_real = float(np.max(_Qp)) - float(np.min(_Qp[_Qp > 0])) if np.any(_Qp > 0) else 0
    if rango_real > 1.0:
        _imax = np.unravel_index(np.argmax(_Qm), _Qm.shape)
        fig.add_trace(go.Scatter3d(x=[_xg[_imax[1]]], y=[_yg[_imax[0]]], z=[float(_Qm[_imax])],
                                   mode="markers+text", marker=dict(color="#ff4444", size=8, symbol="diamond"),
                                   text=[f"qu_max={qu_max_val:.1f}kPa"], textfont=dict(color="white"), name=f"qu_max"))
    _c1m = c1_col_cm/100; _c2m = c2_col_cm/100
    _cx = [-_c1m/2, _c1m/2, _c1m/2, -_c1m/2, -_c1m/2]
    _cy = [-_c2m/2, -_c2m/2, _c2m/2, _c2m/2, -_c2m/2]
    _i_col = np.argmin(np.abs(_xg - 0))
    _j_col = np.argmin(np.abs(_yg - 0))
    _z_col = [float(_Qp[_j_col, _i_col]) * 1.005] * 5
    fig.add_trace(go.Scatter3d(x=_cx, y=_cy, z=_z_col, mode="lines",
                               line=dict(color="white", width=4), name="Columna"))
    if q_adm is not None and q_adm > 0:
        _z_adm = np.full(_Xm.shape, q_adm)
        fig.add_trace(go.Surface(x=_xg, y=_yg, z=_z_adm, colorscale=[[0, 'blue'], [1, 'blue']],
                                 opacity=0.35, showscale=False, name="q_adm Límite"))
    _titulo = titulo or (f"qu(x,y) | qu_max={qu_max_val:.1f} | qu_min={qu_min_val:.1f} kPa" +
                         (" |  LEVANTAMIENTO" if _ten else " |  Compresión total"))
    fig.update_layout(scene=dict(xaxis_title="B[m]", yaxis_title="L[m]", zaxis_title="qu[kPa]",
                                 bgcolor="#0f1117",
                                 xaxis=dict(gridcolor="#333"), yaxis=dict(gridcolor="#333"),
                                 zaxis=dict(gridcolor="#333"),
                                 camera=dict(eye=dict(x=1.8, y=1.8, z=1.6))),
                      paper_bgcolor="#0f1117", font=dict(color="white"),
                      margin=dict(l=0, r=0, b=0, t=40), height=480,
                      title=dict(text=_titulo, font=dict(color="#ffdd44" if _ten else "#44ff88")),
                      showlegend=True)
    return fig, _ten

#  F4: CÁLCULO VIGA DE AMARRE 
def calcular_viga_amarre(R_strap_kN, L_libre_m, b_viga_cm, h_viga_cm, gamma_c=24.0):
    """
    Genera el diagrama de fuerzas, V(x) y M(x) de la viga de amarre.
    R_strap_kN : reacción que la viga recibe de la zapata exterior (kN)
    Retorna    : x_arr, V_arr, M_arr, M_max, V_max, pp_kNm, R_int, advertencia_Rint
    Nota       : gamma_c no incluye el peso del suelo sobre la viga. Si la viga
                 va enterrada con relleno encima, incrementar gamma_c_efectivo fuera de
                 esta función antes de llamarla.
    """
    pp_viga = gamma_c * (b_viga_cm/100) * (h_viga_cm/100)  # kN/m (solo hormigón)
    pp_tot  = pp_viga * L_libre_m
    R_int   = R_strap_kN - pp_tot  # reacción de la zapata interior
    # EC-2 CORREGIDO: bandera de advertencia si R_int < 0 (levantamiento potencial)
    _adv_Rint = R_int < 0
    n = 200
    x_arr = np.linspace(0, L_libre_m, n)
    V_arr = R_strap_kN - pp_viga * x_arr          # V(x)
    M_arr = R_strap_kN * x_arr - pp_viga * x_arr**2 / 2.0  # M(x)
    M_max = float(np.max(np.abs(M_arr)))
    V_max = float(np.max(np.abs(V_arr)))
    return x_arr, V_arr, M_arr, M_max, V_max, pp_viga, R_int, _adv_Rint

#  F5: DISEÑO VIGA DE AMARRE 
def disenar_viga_amarre(M_max_kNm, V_max_kN, F_ax_kN, b_v_cm, h_v_cm,
                        fc_MPa, fy_MPa, L_libre_m, P_total_kN, norma,
                        recub_v=4.0, bar_long="N.5 (5/8\")", bar_est="N.3 (3/8\")",
                        rebar_dict=None):
    """
    Diseña la viga de amarre: As_sup, As_inf, estribos, verificación axial.
    Retorna dict con todos los resultados.
    """
    if rebar_dict is None:
        rebar_dict = {"N.3 (3/8\")": {"area": 0.71, "db": 9.5},
                      "N.4 (1/2\")": {"area": 1.29, "db": 12.7},
                      "N.5 (5/8\")": {"area": 1.99, "db": 15.9},
                      "N.6 (3/4\")": {"area": 2.84, "db": 19.1}}
    sn = _get_strap_norm(norma)
    phi_f = 0.90; phi_v_s = 0.75
    if "E.060" in norma: phi_v_s = 0.85

    db_long = rebar_dict.get(bar_long, {"db": 15.9})["db"]
    db_est  = rebar_dict.get(bar_est, {"db": 9.5})["db"]
    As_long_unit = rebar_dict.get(bar_long, {"area": 1.99})["area"]
    As_est_unit  = rebar_dict.get(bar_est, {"area": 0.71})["area"]

    d_v = h_v_cm - recub_v - db_est/10.0 - db_long/20.0   # peralte efectivo cm
    # EC-3 CORREGIDO: d_v ≤ 0 es error de diseño, no se silencia con fallback
    if d_v <= 0:
        # Retornamos dict de error que render_medianera() detecta y muestra con st.error
        return {
            "_error": True,
            "_msg": (f" **Peralte efectivo d_v = {d_v:.1f} cm ≤ 0** con h={h_v_cm:.0f} cm, "
                     f"recub={recub_v:.1f} cm, Ø_est={db_est:.1f}mm, Ø_long={db_long:.1f}mm. "
                     "Aumente h o reduzca el recubrimiento."),
        }

    # Dimensión mínima normativa
    h_min_cm = L_libre_m * 100 / sn["h_factor"]
    ok_hmin = h_v_cm >= h_min_cm

    # Diseño flexión — As sup (momento positivo máximo) y As inf (mínimo)
    def _As_flexion(Mu_kNm, b_cm, d_cm):
        if Mu_kNm <= 0: return 0.0, True
        Rn = (Mu_kNm * 1e6) / (phi_f * (b_cm*10) * (d_cm*10)**2)  # MPa
        disc = 1 - 2*Rn / (0.85 * fc_MPa)
        disc = max(disc, 0.0)
        rho = (0.85 * fc_MPa / fy_MPa) * (1 - math.sqrt(disc))
        rho_use = max(rho, 0.0018)
        As = rho_use * (b_cm*10) * (d_cm*10) / 100.0  # cm²
        return As, disc > 0

    As_sup, ok_sup = _As_flexion(M_max_kNm, b_v_cm, d_v)
    As_inf_min = max(0.0018 * (b_v_cm*10) * (d_v*10) / 100.0,
                     (0.25 * math.sqrt(fc_MPa) / fy_MPa) * (b_v_cm*10) * (d_v*10) / 100.0)
    As_inf = As_inf_min

    # Acero mínimo 4 barras en esquinas (NSR-10 C.15.13)
    n_long_min = 4
    As_min_4barras = n_long_min * As_long_unit
    As_sup = max(As_sup, As_min_4barras)
    As_inf = max(As_inf, As_min_4barras)
    n_sup = math.ceil(As_sup / As_long_unit)
    n_inf = math.ceil(As_inf / As_long_unit)

    # Fuerza axial sísmica
    F_ax_req = sn["F_ax_pct"] * P_total_kN
    ok_Fax = F_ax_kN <= F_ax_req  # la viga debe soportar en tensión/compresión

    # Diseño cortante
    phi_Vc = phi_v_s * 0.17 * math.sqrt(fc_MPa) * (b_v_cm*10) * (d_v*10) / 1000.0  # kN
    Vs_req = max(0, V_max_kN / phi_v_s - phi_Vc / phi_v_s)
    if Vs_req > 0:
        # s = Av * fy * d / Vs (2 ramas)
        s_est_cm = (2 * As_est_unit * fy_MPa * (d_v*10)) / (Vs_req * 1000.0 / 10.0)
        s_est_cm = min(s_est_cm, d_v/2, 30.0)  # máx. d/2 y 30 cm
    else:
        s_est_cm = min(d_v/2, 30.0)
    s_est_cm = max(s_est_cm, 5.0)
    ok_cort = phi_Vc >= V_max_kN

    # Longitud de zona crítica (estribos más cerrados)
    L_crit_cm = max(2 * h_v_cm, 45.0)
    s_crit_cm = min(s_est_cm, h_v_cm/4, 6*(db_long/10.0), 15.0)

    phi_Mn_sup = phi_f * (As_sup * fy_MPa * (d_v*10) * (1 - (As_sup*fy_MPa)/(1.7*fc_MPa*(b_v_cm*10))) ) / 1e6  # kN·m
    ok_flex = phi_Mn_sup >= M_max_kNm

    return {
        "sn": sn, "d_v": d_v, "h_min_cm": h_min_cm, "ok_hmin": ok_hmin,
        "As_sup": As_sup, "n_sup": n_sup, "As_inf": As_inf, "n_inf": n_inf,
        "phi_Vc": phi_Vc, "Vs_req": Vs_req, "s_est_cm": s_est_cm,
        "s_crit_cm": s_crit_cm, "L_crit_cm": L_crit_cm,
        "ok_cort": ok_cort, "ok_flex": ok_flex, "ok_Fax": ok_Fax,
        "F_ax_req": F_ax_req, "phi_Mn_sup": phi_Mn_sup,
        "bar_long": bar_long, "bar_est": bar_est,
        "As_long_unit": As_long_unit, "As_est_unit": As_est_unit,
        "db_long": db_long, "db_est": db_est,
    }

#  F6: DIAGRAMA V(x) / M(x) PLOTLY 
def render_diagrama_VM(x_arr, V_arr, M_arr, R_strap, R_int, pp_viga, L_libre):
    """Genera figura Plotly con dos paneles: V(x) arriba y M(x) abajo."""
    from plotly.subplots import make_subplots
    fig = make_subplots(rows=2, cols=1, shared_xaxes=True,
                        subplot_titles=("Diagrama de Cortante V(x) [kN]",
                                        "Diagrama de Momento M(x) [kN·m]"),
                        vertical_spacing=0.1)
    # V(x)
    fig.add_trace(go.Scatter(x=x_arr, y=V_arr, fill='tozeroy', fillcolor='rgba(0,200,255,0.15)',
                             line=dict(color='#00c8ff', width=2), name="V(x)"), row=1, col=1)
    fig.add_hline(y=0, line_color="white", line_width=0.5, row=1, col=1)
    # M(x)
    fig.add_trace(go.Scatter(x=x_arr, y=M_arr, fill='tozeroy', fillcolor='rgba(255,165,0,0.15)',
                             line=dict(color='#ffa500', width=2), name="M(x)"), row=2, col=1)
    fig.add_hline(y=0, line_color="white", line_width=0.5, row=2, col=1)
    # Anotaciones de reacciones
    fig.add_vline(x=0, line_color="#7fff00", line_width=1.5, line_dash="dot")
    fig.add_vline(x=L_libre, line_color="#ff6b6b", line_width=1.5, line_dash="dot")
    # FI-7 CORREGIDO: anotaciones con xref/yref en paper para posición fija y visible
    fig.add_annotation(x=0.02, y=0.95, xref="paper", yref="paper",
                       text=f"↑ R_ext = {R_strap:.1f} kN",
                       showarrow=False, font=dict(color="#7fff00", size=11),
                       bgcolor="rgba(0,0,0,0.5)", align="left")
    _Rint_lbl = f"↑ R_int = {R_int:.1f} kN" if R_int >= 0 else f" R_int = {R_int:.1f} kN (levant.)"
    _Rint_col = "#ff6b6b" if R_int >= 0 else "#ffaa00"
    fig.add_annotation(x=0.98, y=0.95, xref="paper", yref="paper",
                       text=_Rint_lbl,
                       showarrow=False, font=dict(color=_Rint_col, size=11),
                       bgcolor="rgba(0,0,0,0.5)", align="right")
    fig.update_layout(paper_bgcolor="#0f1117", plot_bgcolor="#0f1117",
                      font=dict(color="white"), showlegend=False,
                      margin=dict(l=20, r=20, t=40, b=20), height=400,
                      xaxis2=dict(title="x [m]"),
                      yaxis=dict(gridcolor="#222"), yaxis2=dict(gridcolor="#222"))
    return fig

#  F7: VISTA 3D PLOTLY — SISTEMA COMPLETO 
def render_3d_sistema(zap_ext, zap_int, viga_d, dis_d, fc_col=None):
    """
    zap_ext: dict con B, L, H, x0_col, y0_col (centro zapata exterior en coords globales)
    zap_int: dict igual para zapata interior
    viga_d:  dict con b, h, L_libre, x0, y0
    dis_d:   resultado de disenar_viga_amarre()
    """
    fig = go.Figure()
    # Helper: bloque cúbico translúcido
    def _box(ox, oy, oz, w, d, h, color, name, opacity=0.25):
        x = [ox, ox+w, ox+w, ox, ox, ox+w, ox+w, ox]
        y = [oy, oy, oy+d, oy+d, oy, oy, oy+d, oy+d]
        z = [oz]*4 + [oz+h]*4
        fig.add_trace(go.Mesh3d(x=x, y=y, z=z, alphahull=0, opacity=opacity,
                                color=color, name=name, showlegend=True))
    # Zapata exterior
    zE = zap_ext
    _box(zE["x0"]-zE["B"]/2, zE["y0"]-zE["L"]/2, 0,
         zE["B"], zE["L"], zE["H"]/100, "steelblue", "Zapata Exterior")
    # Zapata interior
    zI = zap_int
    _box(zI["x0"]-zI["B"]/2, zI["y0"]-zI["L"]/2, 0,
         zI["B"], zI["L"], zI["H"]/100, "#4CAF50", "Zapata Interior")
    # Columnas
    Hz_E = zE["H"]/100; Hz_I = zI["H"]/100
    _box(zE["x0"]-zE.get("c1",40)/200, zE["y0"]-zE.get("c2",40)/200, Hz_E,
         zE.get("c1",40)/100, zE.get("c2",40)/100, 1.0, "slategray", "Col. Exterior", opacity=0.6)
    _box(zI["x0"]-zI.get("c1",40)/200, zI["y0"]-zI.get("c2",40)/200, Hz_I,
         zI.get("c1",40)/100, zI.get("c2",40)/100, 1.0, "#888", "Col. Interior", opacity=0.6)
    # Viga de amarre — EM-3 CORREGIDO: ancla en la superficie de la zapata, no en borde de columna
    # La viga de amarre está empotrada en cada zapata, no en las columnas.
    # Geométricamente va de x0_zapata_ext hasta x1_zapata_int.
    vg = viga_d
    _hv = vg["h"]/100; _bv = vg["b"]/100
    _xv0 = zE["x0"] + zE["B"]/2          # borde derecho de la zapata exterior (ancla)
    _xv1 = zI["x0"] - zI["B"]/2          # borde izquierdo de la zapata interior (ancla)
    _yv0 = -_bv/2
    _box(_xv0, _yv0, max(Hz_E, Hz_I), max(_xv1-_xv0, 0.01), _bv, _hv, "#FF9800", "Viga de Amarre", opacity=0.75)
    # Acero longitudinal de la viga (líneas)
    _rec = vg.get("recub", 4)/100
    _n_sup = dis_d["n_sup"]; _n_inf = dis_d["n_inf"]
    _db_l = dis_d["db_long"]/1000
    _z_base = max(Hz_E, Hz_I) + _rec + _db_l/2
    _z_sup  = max(Hz_E, Hz_I) + _hv - _rec - _db_l/2
    for _i, _ny in enumerate(np.linspace(-_bv/2+_rec+_db_l, _bv/2-_rec-_db_l, max(_n_sup,2))):
        _show = _i == 0
        fig.add_trace(go.Scatter3d(x=[_xv0, _xv1], y=[_ny, _ny], z=[_z_sup, _z_sup],
                                   mode="lines", line=dict(color="#ff6b35", width=4),
                                   name="Acero Sup. Viga" if _show else None, showlegend=_show, legendgroup="Asup"))
        fig.add_trace(go.Scatter3d(x=[_xv0, _xv1], y=[_ny, _ny], z=[_z_base, _z_base],
                                   mode="lines", line=dict(color="#ffd54f", width=4),
                                   name="Acero Inf. Viga" if _show else None, showlegend=_show, legendgroup="Ainf"))
    # Estribos (marcos rectangulares cada s_est_cm)
    _s_est = dis_d["s_est_cm"]/100
    _xv_range = np.arange(_xv0 + _s_est, _xv1, _s_est)
    _is_first_est = True
    for _xe in _xv_range:
        _est_x = [_xe]*5; _est_y = [-_bv/2+_rec, _bv/2-_rec, _bv/2-_rec, -_bv/2+_rec, -_bv/2+_rec]
        _est_z = [_z_base, _z_base, _z_sup, _z_sup, _z_base]
        fig.add_trace(go.Scatter3d(x=_est_x, y=_est_y, z=_est_z,
                                   mode="lines", line=dict(color="#aaa", width=2),
                                   name="Estribos Viga" if _is_first_est else None,
                                   showlegend=_is_first_est, legendgroup="est"))
        _is_first_est = False

    # FI-3: go.Surface de qu(x,y) bajo cada zapata (mapa de presiones integrado en 3D)
    def _add_pressure_surface(zap, label, colorscale="RdYlGn_r"):
        _cx = zap["x0"]; _cy = zap["y0"]
        _B = zap["B"];   _L = zap["L"]
        _nx, _ny_g = 20, 20
        _xg = np.linspace(_cx - _B/2, _cx + _B/2, _nx)
        _yg = np.linspace(_cy - _L/2, _cy + _L/2, _ny_g)
        _Xm, _Ym = np.meshgrid(_xg, _yg)
        # Distribución trapezoidal lineal entre qu_max (borde lindero) y qu_min/2 (borde opuesto)
        _qu_max = zap.get("qu_max", 150.0)
        _qu_min = max(zap.get("qu_min", 80.0), 0.0)
        # Gradiente en X desde qu_max hasta qu_min
        _Qm = _qu_max + (_qu_min - _qu_max) * (_Xm - (_cx - _B/2)) / _B
        _Qm = np.clip(_Qm, 0, None)
        # Escalar a -0.05m para que quede justo bajo la zapata (visual)
        _Z_surf = np.full_like(_Qm, -0.05)
        # Codificar qu como colormap
        fig.add_trace(go.Surface(
            x=_xg, y=_yg, z=_Z_surf,
            surfacecolor=_Qm, colorscale=colorscale,
            cmin=0, cmax=float(np.max(_Qm))*1.1,
            showscale=False, opacity=0.85,
            name=f"Presiones {label}", showlegend=True,
            hovertemplate=f"<b>{label}</b><br>qu = %{{customdata:.1f}} kPa<extra></extra>",
            customdata=_Qm,
        ))
    # Añadir superficie de presiones para cada zapata con sus qu del session_state
    _qu_data_e = fc_col or {}  # se pasa como fc_col cuando hay datos
    _zE_qu = {**zE, "qu_max": zE.get("qu_max", 150.0), "qu_min": zE.get("qu_min", 60.0)}
    _zI_qu = {**zI, "qu_max": zI.get("qu_max", 120.0), "qu_min": zI.get("qu_min", 60.0)}
    _add_pressure_surface(_zE_qu, "Z.Ext")
    _add_pressure_surface(_zI_qu, "Z.Int")

    fig.update_layout(scene=dict(aspectmode='data',
                                 xaxis_title='X [m]', yaxis_title='Y [m]', zaxis_title='Z [m]',
                                 bgcolor='#0f1117',
                                 xaxis=dict(showgrid=True, gridcolor='#333'),
                                 yaxis=dict(showgrid=True, gridcolor='#333'),
                                 zaxis=dict(showgrid=True, gridcolor='#333')),
                      margin=dict(l=0, r=0, b=0, t=40), height=600,
                      paper_bgcolor='#0f1117', font=dict(color='white'),
                      title=dict(text="Sistema Zapata Medianera + Viga de Amarre — Vista 3D",
                                 font=dict(color='white')),
                      showlegend=True, dragmode='turntable')
    return fig

#  F8: DXF — SISTEMA COMPLETO 
def generar_dxf_sistema(zap_ext_d, zap_int_d, viga_d, dis_d, norma, fc, fy, papel_w, papel_h, papel_lbl):
    """Genera plano DXF ICONTEC con planta general, y cortes."""
    import datetime as _dt2
    sn = _get_strap_norm(norma)
    doc = ezdxf.new('R2010', setup=True)
    doc.units = ezdxf.units.CM
    
    LW = {'ZAPATA_EXT':40, 'ZAPATA_INT':40, 'VIGA_STRAP':50, 'ACERO_LONG':40, 'ACERO_EST':25, 'COTAS':25, 'TEXTO':25, 'ROTULO':35, 'MARGEN':70}
    COL = {'ZAPATA_EXT':3, 'ZAPATA_INT':5, 'VIGA_STRAP':1, 'ACERO_LONG':6, 'ACERO_EST':7, 'COTAS':2, 'TEXTO':7, 'ROTULO':8, 'MARGEN':7}
    for lay, lw in LW.items():
        doc.layers.new(lay, dxfattribs={'color':COL[lay], 'lineweight':lw})
    doc.styles.new('ROMANS', dxfattribs={'font':'Arial.ttf'})
    msp = doc.modelspace()
    
    # Dibujar elementos virtualmente (unidades CM reales)
    def _rect(ox, oy, w, h, lay):
        msp.add_lwpolyline([(ox,oy),(ox+w,oy),(ox+w,oy+h),(ox,oy+h),(ox,oy)], dxfattribs={'layer':lay})
    def _txt(x, y, txt, lay, h=2.5, ha='center', va='center'):
        msp.add_text(txt, dxfattribs={'layer':lay,'style':'ROMANS','height':h,
                                      'insert':(x,y),'align_point':(x,y),
                                      'halign':{'left':0,'center':1,'right':2}[ha],
                                      'valign': 2 if va=='center' else 0})
                                      
    Be_cm = zap_ext_d["B_cm"]; Le_cm = zap_ext_d["L_cm"]
    Bi_cm = zap_int_d["B_cm"]; Li_cm = zap_int_d["L_cm"]
    sep_cm = viga_d["L_libre_m"]*100; bv_cm = viga_d["b_cm"]; hv_cm = viga_d["h_cm"]
    
    # ── PLANTA ──
    _rect(0, -Le_cm/2, Be_cm, Le_cm, 'ZAPATA_EXT')
    _rect(Be_cm + sep_cm, -Li_cm/2, Bi_cm, Li_cm, 'ZAPATA_INT')
    _rect(Be_cm, -bv_cm/2, sep_cm, bv_cm, 'VIGA_STRAP')
    _txt(Be_cm/2, 0, f"Z.EXT {Be_cm:.0f}x{Le_cm:.0f}", 'TEXTO', 4.0)
    _txt(Be_cm + sep_cm + Bi_cm/2, 0, f"Z.INT {Bi_cm:.0f}x{Li_cm:.0f}", 'TEXTO', 4.0)
    _txt(Be_cm + sep_cm/2, 0, f"VIGA {bv_cm:.0f}x{hv_cm:.0f}", 'TEXTO', 4.0)
    _txt((Be_cm+sep_cm+Bi_cm)/2, max(Le_cm, Li_cm)/2 + 20, "VISTA EN PLANTA", 'TEXTO', 5.0)

    # ── CORTE LONGITUDINAL ──
    oy_el = -max(Le_cm, Li_cm)/2 - 100
    hz_ext = zap_ext_d["H_cm"]; hz_int = zap_int_d["H_cm"]
    _rect(0, oy_el, Be_cm, hz_ext, 'ZAPATA_EXT')
    _rect(Be_cm + sep_cm, oy_el, Bi_cm, hz_int, 'ZAPATA_INT')
    _rect(Be_cm, oy_el + max(hz_ext, hz_int), sep_cm, hv_cm, 'VIGA_STRAP')
    
    rec_v = viga_d.get("recub_cm", 4)
    y_sup = oy_el + max(hz_ext, hz_int) + hv_cm - rec_v
    y_inf = oy_el + max(hz_ext, hz_int) + rec_v
    msp.add_line((Be_cm+rec_v, y_sup), (Be_cm+sep_cm-rec_v, y_sup), dxfattribs={'layer':'ACERO_LONG'})
    msp.add_line((Be_cm+rec_v, y_inf), (Be_cm+sep_cm-rec_v, y_inf), dxfattribs={'layer':'ACERO_LONG'})
    
    s_est = dis_d['s_est_cm']
    import numpy as np
    for _xe in np.arange(Be_cm + s_est, Be_cm + sep_cm, s_est):
        msp.add_line((_xe, y_inf), (_xe, y_sup), dxfattribs={'layer':'ACERO_EST'})
    _txt((Be_cm+sep_cm+Bi_cm)/2, oy_el - 30, f"CORTE LONGITUDINAL", 'TEXTO', 5.0)

    # ── CORTE TRANSVERSAL ──
    ox_tr = Be_cm*2 + sep_cm + Bi_cm + 80
    oy_tr = oy_el + hz_ext
    _rect(ox_tr, oy_tr, bv_cm, hv_cm, 'VIGA_STRAP')
    msp.add_lwpolyline([(ox_tr+rec_v, oy_tr+rec_v), (ox_tr+bv_cm-rec_v, oy_tr+rec_v), 
                        (ox_tr+bv_cm-rec_v, oy_tr+hv_cm-rec_v), (ox_tr+rec_v, oy_tr+hv_cm-rec_v)], 
                        close=True, dxfattribs={'layer':'ACERO_EST'})
    _txt(ox_tr + bv_cm/2, oy_tr - 30, f"CORTE TRANS.", 'TEXTO', 5.0)

    # ESCALADO
    dim_w = ox_tr + bv_cm + 50
    dim_h = (max(Le_cm, Li_cm)/2 + 50) + abs(oy_el - 60)
    escala_den = 50
    for den in [100, 75, 50, 25, 20]:
        if dim_w / den <= papel_w*0.8 and dim_h / den <= papel_h*0.8:
            escala_den = den; break
    ESC = 1.0 / escala_den

    # ── MARGEN Y ROTULO N4 ──
    msp.add_lwpolyline([(1,1), (papel_w-1,1), (papel_w-1,papel_h-1), (1,papel_h-1), (1,1)], dxfattribs={'layer':'MARGEN'})
    msp.add_lwpolyline([(1,1), (papel_w-1,1), (papel_w-1,5), (1,5), (1,1)], dxfattribs={'layer':'ROTULO'})
    _txt(papel_w/2, 3.5, f"ZAPATA MEDIANERA + {sn['nombre']} | {norma}", 'TEXTO', 0.5)
    _txt(papel_w/2, 2.0, f"Papel: {papel_lbl} | Escala Vista Inicial: 1:{escala_den}", 'TEXTO', 0.4)

    # Convertir a bloque y escalar
    doc.blocks.new(name='VISTAS_SISTEMA')
    blk = doc.blocks.get('VISTAS_SISTEMA')
    for e in list(msp):
        if e.dxf.layer not in ['MARGEN', 'ROTULO', 'TEXTO']:
            blk.add_entity(e.copy())
            e.destroy()
        elif e.dxf.layer == 'TEXTO' and e.dxf.height >= 1.0:
            blk.add_entity(e.copy())
            e.destroy()
            
    ox_b = (papel_w - dim_w*ESC)/2 + 10
    oy_b = 6 + (papel_h - 6 - dim_h*ESC)/2 + abs(oy_el - 60)*ESC
    msp.add_blockref('VISTAS_SISTEMA', insert=(ox_b, oy_b), dxfattribs={'xscale': ESC, 'yscale': ESC})

    return doc

def generar_docx_ampliado(zap_ext_d, zap_int_d, viga_d, dis_d, dist_d, norma, fig_vm=None):
    """Genera Document() de python-docx Estándar Diamante N4."""
    from docx import Document as _Doc
    from docx.shared import Inches as _In, Pt as _Pt
    import datetime as _dt_dxf
    sn = _get_strap_norm(norma)
    doc = _Doc()
    
    doc.add_heading(f"MEMORIA ESTRUCTURAL — SISTEMA {sn['nombre'].upper()} ({norma})", 0)

    # 0. Portada — identidad global v6.1
    _emp_d, _proy_d, _ing_d, _elab_d, _rev_d, _apb_d, _logo_d = _get_identity()
    p0 = doc.add_paragraph()
    p0.add_run(
        f"Empresa: {_emp_d}  |  Proyecto: {_proy_d}  |  Ing.: {_ing_d}\n"
        f"Elemento: Zapata Medianera + {sn['nombre']}\n"
        f"Norma: {norma} — Art. {sn['art']}\n"
        f"Elaboró: {_elab_d}   Revisó: {_rev_d}   Aprobó: {_apb_d}\n"
        f"Materiales: Concreto f'c = {viga_d.get('fc',21)} MPa, Acero fy = {viga_d.get('fy',420)} MPa\n"
        f"Fecha: {_dt_dxf.datetime.now().strftime('%d/%m/%Y')}"
    ).bold = True
    if _logo_d:
        try:
            import io as _io_logo
            doc.add_picture(_io_logo.BytesIO(_logo_d), width=Inches(1.5))
        except Exception:
            pass

    # 1. Parámetros de Diseño y Geometría
    doc.add_heading("1. PARÁMETROS DE DISEÑO", level=1)
    doc.add_paragraph(f"  Geometría Zapata Exterior: B={zap_ext_d['B_cm']:.0f} cm × L={zap_ext_d['L_cm']:.0f} cm × H={zap_ext_d['H_cm']:.0f} cm")
    doc.add_paragraph(f"  Geometría Zapata Interior: B={zap_int_d['B_cm']:.0f} cm × L={zap_int_d['L_cm']:.0f} cm × H={zap_int_d['H_cm']:.0f} cm")
    doc.add_paragraph(f"  Viga de Amarre: {viga_d['b_cm']:.0f} cm × {viga_d['h_cm']:.0f} cm × L_libre={viga_d['L_libre_m']:.2f} m")

    # 2. Distribución de Presiones
    doc.add_heading("2. DISTRIBUCIÓN DE PRESIONES", level=1)
    doc.add_paragraph(f"  qu_max = {dist_d.get('qu_max',0):.2f} kPa  |  qu_min = {dist_d.get('qu_min',0):.2f} kPa")
    doc.add_paragraph(f"  Excentricidades resultantes: e_B = {dist_d.get('e_B',0):.3f} m  |  e_L = {dist_d.get('e_L',0):.3f} m")
    mey = dist_d.get('meyerhof')
    if mey:
        doc.add_paragraph(f"  Área Efectiva Meyerhof: B'={mey['B_prima']:.2f} m × L'={mey['L_prima']:.2f} m → qu_eff={mey['qu_eff']:.1f} kPa")

    # 3. Diseño Flexo-compresión Viga Strap
    doc.add_heading(f"3. DISEÑO DE {sn['nombre'].upper()} A FLEXIÓN", level=1)
    doc.add_paragraph(f"  Peralte mínimo normativo: L/{sn['h_factor']} = {dis_d['h_min_cm']:.1f} cm  →  {'CUMPLE' if dis_d['ok_hmin'] else 'NO CUMPLE'}")
    doc.add_paragraph(f"  Acero longitudinal Superior: {dis_d['n_sup']}x {dis_d['bar_long']}  (As = {dis_d['As_sup']:.2f} cm²)")
    doc.add_paragraph(f"  Acero longitudinal Inferior: {dis_d['n_inf']}x {dis_d['bar_long']}  (As = {dis_d['As_inf']:.2f} cm²)")
    doc.add_paragraph(f"  Fuerza axial sísmica requerida: {dis_d['F_ax_req']:.1f} kN")

    # 4. Diseño a Cortante
    doc.add_heading("4. DISEÑO A CORTANTE Y ESTRIBOS", level=1)
    doc.add_paragraph(f"  Resistencia del concreto (φVc): {dis_d['phi_Vc']:.1f} kN  →  {'CUMPLE cortante' if dis_d['ok_cort'] else 'NO CUMPLE cortante'}")
    doc.add_paragraph(f"  Distribución de Estribos: {dis_d['bar_est']} @ {dis_d['s_crit_cm']:.0f} cm en zona crítica, {dis_d['s_est_cm']:.0f} cm en el resto.")

    # 5. Verificaciones Normativas
    doc.add_heading("5. VERIFICACIONES DE SISTEMA ZAPATA-VIGA", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parámetro Verificado'
    hdr_cells[1].text = 'Estado'
    hdr_cells[2].text = 'Valor'
    
    def _add_row(param, status, val):
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = "CUMPLE" if status else "NO CUMPLE"
        row_cells[2].text = val
    
    _add_row("Peralte h mínimo", dis_d['ok_hmin'], f"h={viga_d['h_cm']}cm ≥ {dis_d['h_min_cm']:.1f}cm")
    _add_row("Diseño a Cortante", dis_d['ok_cort'], f"φVc={dis_d['phi_Vc']:.0f}kN")
    _add_row("Fuerza axial", dis_d['ok_Fax'], f"Fax={dis_d['F_ax_req']:.0f}kN")

    # 6. Cuantificación
    doc.add_heading("6. CUANTIFICACIÓN DE VIGA", level=1)
    doc.add_paragraph("Volúmenes y aceros reportados por tabulación separada en la UI.")

    # 7. Detalles Generales
    doc.add_heading("7. DIAGRAMAS Y PLANOS", level=1)
    if fig_vm is not None:
        try:
            _img_bytes = _fig_to_docx_white(fig_vm)
            doc.add_paragraph("Diagrama de Cortante y Momento (Viga):")
            doc.add_picture(io.BytesIO(_img_bytes), width=_In(5.5))
        except Exception as e:
            doc.add_paragraph(f"(Error al embeber imagen: {e})")

    # 8. Referencias
    doc.add_heading("8. REFERENCIAS CITADAS", level=1)
    doc.add_paragraph("NSR-10 Título C y E.")
    doc.add_paragraph("\n\n_________________________________________\nFirma Ing. Responsable")
    doc.add_paragraph("Matrícula Profesional: _______________")

    return doc


# 
# T1: ESFUERZOS EN EL SUELO (BOUSSINESQ)
# 
with st.expander(_t("1. Esfuerzos en masa de suelo debajo de zapata", "1. Soil Stresses under Footing (Boussinesq)"), expanded=False):
    st.info(_t("**Modo de uso:** Ingresa las dimensiones de la zapata y la carga aplicada. El programa usa la solución de Boussinesq (integración de carga rectangular) para encontrar el incremento de esfuerzo vertical a cierta profundidad Z debajo del centro de la zapata.", " **How to use:** Enter footing dimensions and load. Uses Boussinesq method to find vertical stress increment at depth Z."))
    c1, c2 = st.columns(2)
    with c1:
        #  Selector de Unidades para la Carga 
        _P_unidades = {
            "kN":     1.0,
            "kgf":    0.00980665,
            "tf (ton-fuerza)": 9.80665,
            "kip":    4.44822,
            "lb (lbf)": 0.00444822,
        }
        _P_unit = st.selectbox("Unidad de carga P", list(_P_unidades.keys()), key="z_bous_Punit")
        _factor = _P_unidades[_P_unit]
        # Límites dinámicos según unidad
        _P_max = {
            "kN": 10000.0, "kgf": 1_000_000.0, "tf (ton-fuerza)": 1000.0,
            "kip": 2248.0, "lb (lbf)": 2_248_000.0,
        }
        _P_def = {
            "kN": 1000.0, "kgf": 100_000.0, "tf (ton-fuerza)": 100.0,
            "kip": 224.8, "lb (lbf)": 224_800.0,
        }
        P_bous_raw = st.number_input(
            f"Carga en Zapata P [{_P_unit}]",
            min_value=0.1, max_value=_P_max[_P_unit],
            value=_P_def[_P_unit], step=_P_max[_P_unit]/100,
            key="z_bous_P"
        )
        P_bous = P_bous_raw * _factor   # siempre en kN para los cálculos
        if _P_unit != "kN":
            st.caption(f"≡ **{P_bous:,.2f} kN** | {P_bous/9.80665:.2f} tf | {P_bous*224.809:.0f} lb")
        B_bous = st.number_input("Ancho B [m]", 0.5, 10.0, st.session_state.get("z_bous_B", 2.0), 0.1, key="z_bous_B")
        L_bous = st.number_input("Largo L [m]", 0.5, 10.0, st.session_state.get("z_bous_L", 2.0), 0.1, key="z_bous_L")
    with c2:
        Z_bous = st.number_input("Profundidad de análisis Z [m]", 0.1, 20.0, st.session_state.get("z_bous_Z", 2.0), 0.5, key="z_bous_Z")
        q_0 = P_bous / (B_bous * L_bous) # Esfuerzo de contacto kN/m2
        st.markdown(_t(f"**Esfuerzo de contacto ($q_0$):** {q_0:.2f} kPa = {q_0/9.80665:.3f} t/m² = {q_0/98.0665:.4f} kg/cm²",
                       f"**Contact Stress ($q_0$):** {q_0:.2f} kPa = {q_0/9.80665:.3f} t/m² = {q_0/98.0665:.4f} kg/cm²"))
        
    # Fadum/Boussinesq bajo el centro: dividimos el rectangulo en 4 rectangulos de B/2 x L/2
    m = (B_bous/2.0) / Z_bous
    n = (L_bous/2.0) / Z_bous
    delta_sigma_z = 4 * q_0 * I_z_bous(m, n)
    
    st.success(f"Incremento de esfuerzo vertical bajo el centro a Z={Z_bous}m: **Δσ_z = {delta_sigma_z:.2f} kPa**")
    
    # --- Gráficos Analíticos Boussinesq ---
    fig_b, (ax_v, ax_h) = plt.subplots(1, 2, figsize=(12, 4.5))
    fig_b.patch.set_facecolor("#0f1117")
    ax_v.set_facecolor("#161b22")
    ax_h.set_facecolor("#161b22")

    # 1. Perfil Vertical bajo el centro
    z_max = max(B_bous*4.0, 10.0)
    zs = np.linspace(0.1, z_max, 100)
    sigmas_v = [4 * q_0 * I_z_bous((B_bous/2.0)/z, (L_bous/2.0)/z) for z in zs]
    
    q_10 = q_0 * 0.10
    # Encontrar Z_influencia donde delta_sigma <= 10% q_0
    z_inf = next((z for z, sig in zip(zs, sigmas_v) if sig <= q_10), None)

    ax_v.plot(sigmas_v, zs, color="#00ffcc", lw=2.5, label="Δσ_z bajo el centro")
    ax_v.axvline(q_10, color="#ffdd44", linestyle="-.", lw=1.5, label="10% q₀ (Influencia)")
    if z_inf:
        ax_v.axhline(z_inf, color="#ff9500", linestyle=":", lw=1.5, label=f"Z inf ≈ {z_inf:.1f} m")

    ax_v.invert_yaxis()
    ax_v.set_xlabel("Incremento de Esfuerzo Δσ_z [kPa]", color="white")
    ax_v.set_ylabel("Profundidad Z [m]", color="white")
    ax_v.set_title("Distribución Vertical", color="white", fontsize=11)
    
    # Ajuste dinámico del rango X: un poco más del maximo de los valores relevantes
    # O ignorar el valor exacto en z=0.1 si es igual a q_0, enfocar en el tramo util.
    max_sigma_plot = max(sigmas_v)
    ax_v.set_xlim(0, max_sigma_plot * 1.1)
    
    ax_v.tick_params(colors="white")
    ax_v.grid(True, linestyle=":", alpha=0.3, color="white")
    ax_v.legend(loc="lower right", facecolor="#161b22", labelcolor="white", edgecolor="none")

    # 2. Perfiles Horizontales (Transversales al eje Y)
    def sigma_x(x_arr, z_val):
        # Esfuerzo usando Superposición Múltiple de rectángulos (Fadum corner formula)
        sigmas_h = []
        y1, y2 = L_bous/2.0, -L_bous/2.0
        for x in x_arr:
            x1, x2 = B_bous/2.0 - x, -B_bous/2.0 - x
            I1 = I_z_bous_vec(abs(x1)/z_val, abs(y1)/z_val) * np.sign(x1) * np.sign(y1)
            I2 = I_z_bous_vec(abs(x2)/z_val, abs(y1)/z_val) * np.sign(x2) * np.sign(y1)
            I3 = I_z_bous_vec(abs(x1)/z_val, abs(y2)/z_val) * np.sign(x1) * np.sign(y2)
            I4 = I_z_bous_vec(abs(x2)/z_val, abs(y2)/z_val) * np.sign(x2) * np.sign(y2)
            sigmas_h.append(q_0 * abs(I1 - I2 - I3 + I4))
        return sigmas_h

    xs = np.linspace(-B_bous*2, B_bous*2, 100)
    profundidades_plot = [B_bous, B_bous*2, B_bous*3]
    colores = ["#ff5252", "#ffd740", "#69f0ae"]
    
    for zd, col in zip(profundidades_plot, colores):
        sig_x = sigma_x(xs, zd)
        ax_h.plot(xs, sig_x, color=col, lw=2, label=f"Z = {zd:.1f} m")

    # Dibujar contorno de la zapata en Z=0
    ax_h.plot([-B_bous/2, B_bous/2], [q_0, q_0], color="white", lw=2.5, linestyle="-", label="Base Zapata")

    ax_h.set_xlabel("Distancia X [m]", color="white")
    ax_h.set_ylabel("Δσ_z [kPa]", color="white")
    ax_h.set_title("Perfiles Horizontales Transversales", color="white", fontsize=11)
    ax_h.set_xlim(-B_bous*2, B_bous*2)
    ax_h.set_ylim(0, q_0 * 1.1)
    ax_h.tick_params(colors="white")
    ax_h.grid(True, linestyle=":", alpha=0.3, color="white")
    ax_h.legend(loc="upper right", facecolor="#161b22", labelcolor="white", edgecolor="none")

    fig_b.tight_layout()
    st.pyplot(fig_b)
    plt.close(fig_b)

# 
# T2: CAPACIDAD PORTANTE DE SUELO (TERZAGHI) + ASENTAMIENTOS
# 
with st.expander(_t("2. Capacidad Portante de Suelo (Terzaghi) y Asentamientos", " 2. Bearing Capacity (Terzaghi) and Settlements"), expanded=False):
    st.info(_t(
        " **Modo de uso:** Ingresa φ, c, γ y la geometría de la zapata. "
        "El módulo calcula la capacidad última de Terzaghi con influencia del NF, "
        "grafica el diagrama Vesic (1973) para tipo de falla y el bulbo de presiones, "
        "y opcionalmente estima el asentamiento elástico inmediato.",
        " **How to use:** Enter φ, c, γ and footing geometry. Module calculates "
        "Terzaghi ultimate capacity with water-table correction, Vesic failure-type chart, "
        "pressure bulb, and optionally estimates immediate elastic settlement."
    ))

    #  CONVERSOR DE UNIDADES 
    with st.container():
        st.markdown("** Conversor Rápido de Resistencia de Suelo**")
        uc1, uc2, uc3 = st.columns([1, 1, 2])
        with uc1:
            q_conv_unit = st.selectbox(_t("Unidad de entrada:", "Input Unit:"),
                ["kPa (kN/m²)", "ton/m²", "kg/cm²", "MPa", "psi", "kN/m²"],
                key="q_conv_unit_terz")
        with uc2:
            q_conv_val = st.number_input(_t("Valor:", "Value:"), value=1.0, step=0.1, key="q_conv_val_terz")
        with uc3:
            _conv = {"kPa (kN/m²)": q_conv_val, "ton/m²": q_conv_val*9.80665,
                     "kg/cm²": q_conv_val*98.0665, "MPa": q_conv_val*1000.0,
                     "psi": q_conv_val*6.89476, "kN/m²": q_conv_val}.get(q_conv_unit, q_conv_val)
            st.markdown(f"| Unidad | Valor |\n|--------|-------|\n"
                        f"| **kPa** | `{_conv:.3f}` |\n"
                        f"| **ton/m²** | `{_conv/9.80665:.3f}` |\n"
                        f"| **kg/cm²** | `{_conv/98.0665:.4f}` |\n"
                        f"| **MPa** | `{_conv/1000.0:.5f}` |\n"
                        f"| **psi** | `{_conv/6.89476:.2f}` |")
    st.divider()

    #  ENTRADAS 
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("#####  Parámetros Geotécnicos")
        phi_ang  = st.number_input(_t("Ángulo de fricción φ [°]","Friction angle φ [°]"),
                                   0.0, 50.0, st.session_state.get("z_phi", 30.0), 1.0, key="z_phi")
        coh_unit = st.selectbox(_t("Unidad cohesión:","Cohesion Unit:"),
                                ["kPa", "kg/cm²", "ton/m²"], key="coh_u")
        coh_val  = st.number_input(f"Cohesión c [{coh_unit}]", 0.0, 200.0,
                                   st.session_state.get("coh_val", 5.0 if coh_unit=="kPa" else 0.05),
                                   0.5 if coh_unit=="kPa" else 0.01, key="coh_val")
        coh_c    = coh_val if coh_unit=="kPa" else (coh_val*98.0665 if coh_unit=="kg/cm²" else coh_val*9.80665)
        gam_unit = st.selectbox(_t("Unidad γ húmedo:","γ moist Unit:"),
                                ["kN/m³", "ton/m³", "kg/m³"], key="gam_u")
        gam_val  = st.number_input(f"Peso esp. húmedo γ [{gam_unit}]", 10.0,
                                   25.0 if gam_unit != "kg/m³" else 2500.0,
                                   st.session_state.get("gam_val", 18.0 if gam_unit != "kg/m³" else 1800.0),
                                   0.5, key="gam_val")
        gamma_s  = gam_val if gam_unit=="kN/m³" else (gam_val*9.80665 if gam_unit=="ton/m³" else gam_val*0.00980665)
        gamma_sat_t = st.number_input("Peso esp. saturado γ_sat [kN/m³]", 16.0, 25.0, 20.0, 0.5, key="z_gsat")
        gamma_w  = 9.81
    with c2:
        st.markdown("#####  Geometría y Carga")
        forma_zap = st.selectbox(_t("Forma de zapata","Footing shape"),
                                 ["Cuadrada", "Continua (Muro)", "Circular"] if lang=="Español"
                                 else ["Square", "Continuous (Wall)", "Circular"], key="z_shape")
        B_cp  = st.number_input(_t("Ancho B [m]","Width B [m]"), 0.5, 10.0,
                                st.session_state.get("zap_b", 2.0), 0.1, key="zap_b")
        L_cp  = st.number_input(_t("Largo L [m]","Length L [m]"), 0.5, 100.0,
                                st.session_state.get("zap_l", 2.0), 0.1, key="zap_l")
        Hz_cp = st.number_input(_t("Altura zapata Hz [m]","Footing height Hz [m]"),
                                0.1, 2.0, 0.5, 0.05, key="z_hz_cp")
        Df_cp = st.number_input(_t("Profundidad Df [m]","Depth Df [m]"), 0.0, 10.0,
                                st.session_state.get("zap_df", 1.5), 0.1, key="zap_df")
        b_col_cp = st.number_input(_t("Lado columna b [m]","Column side b [m]"), 0.1, 2.0, 0.4, 0.05, key="z_bcol_cp")
        Q_act  = st.number_input("Carga vertical actuante Q [kN]", 10.0, 50000.0, 800.0, 100.0, key="z_Qact")
    with c3:
        st.markdown("#####  Nivel Freático y FS")
        NF_prof = st.number_input("NF — Profundidad nivel freático [m]", 0.0, 20.0, 1.0, 0.5, key="z_nf")
        FS_terz = st.number_input(_t("Factor de Seguridad (FS)","Safety Factor (FS)"),
                                  1.0, 5.0, st.session_state.get("z_fs", 3.0), 0.1, key="z_fs")
        N_spt   = st.number_input("N60 campo (SPT)", 0, 100, 14, 1, key="z_spt")
        st.caption("ℹ N60 clasifica perfil (Vesic)")
        metodo_cap = st.radio("Método de Capacidad", ["Terzaghi", "Meyerhof", "Vesic"], horizontal=True)

    # Parámetros para asentamiento elástico movidos abajo para unificar IO
    pass

    Hzap = Hz_cp * 100.0   # conversión m → cm para cálculos de acero y exportación




    #  CÁLCULO GEOTÉCNICO 
    phi_rad = math.radians(phi_ang)

    if metodo_cap == "Terzaghi":
        # Factores capacidad — Terzaghi
        if phi_ang == 0:
            Nc, Nq, Ngamma = 5.7, 1.0, 0.0
        else:
            a_t    = math.exp((0.75*math.pi - phi_rad/2)*math.tan(phi_rad))
            Nq     = (a_t**2) / (2*math.cos(math.radians(45) + phi_rad/2)**2)
            Nc     = (Nq - 1) / math.tan(phi_rad)
            Ngamma = 2*(Nq+1)*math.tan(phi_rad) / (1 + 0.4*math.sin(4*phi_rad))

        # Factores de forma (Terzaghi simplificado)
        if forma_zap in ["Cuadrada", "Square"]:
            sc, sq, sgamma = 1.3, 1.0, 0.8
        elif forma_zap in ["Circular"]:
            sc, sq, sgamma = 1.3, 1.0, 0.6
        else:
            sc, sq, sgamma = 1.0, 1.0, 1.0
        dc, dq, dgamma = 1.0, 1.0, 1.0
        ic_v = iq_v = ig_v = 1.0  # Terzaghi: sin factores de inclinacion
            
    elif metodo_cap == "Meyerhof":
        # Factores de Capacidad — Meyerhof (1963)
        if phi_ang == 0:
            Nc, Nq, Ngamma = 5.14, 1.0, 0.0
            sc = 1 + 0.2 * (B_cp/L_cp)
            sq, sgamma = 1.0, 1.0
            dc = 1 + 0.2 * (Df_cp/B_cp)
            dq, dgamma = 1.0, 1.0
        else:
            Nq     = math.exp(math.pi * math.tan(phi_rad)) * (math.tan(math.pi/4 + phi_rad/2)**2)
            Nc     = (Nq - 1) / math.tan(phi_rad)
            Ngamma = (Nq - 1) * math.tan(1.4 * phi_rad)
            Kp     = math.tan(math.pi/4 + phi_rad/2)**2
            sc     = 1 + 0.2 * Kp * (B_cp / L_cp)
            sq     = 1 + 0.1 * Kp * (B_cp / L_cp) if phi_ang >= 10 else 1.0
            sgamma = sq
            dc     = 1 + 0.2 * math.sqrt(Kp) * (Df_cp / B_cp)
            dq     = 1 + 0.1 * math.sqrt(Kp) * (Df_cp / B_cp) if phi_ang >= 10 else 1.0
            dgamma = dq
        # Vesic: inclinacion neutra
        ic_v = iq_v = ig_v = 1.0

    else:
        # ──────────────────────────────────────────────────────────────────────
        # PASO 5 — MÉTODO VESIC (1973/1975) — NSR-10 H.3 / Bowles (1996) §3-5
        # Ref: Vesic (1973) "Analysis of Ultimate Loads of Shallow Foundations"
        #      JSMFD, ASCE 99(1):45-73
        # q_ult = c·Nc·Fcs·Fcd·Fci + q·Nq·Fqs·Fqd·Fqi + 0.5·γ·B·Nγ·Fγs·Fγd·Fγi
        # ──────────────────────────────────────────────────────────────────────

        # ── Factores de capacidad de carga (Vesic = Hansen excepto Nγ) ────
        if phi_ang == 0:
            Nc, Nq, Ngamma = 5.14, 1.0, 0.0
        else:
            Nq     = math.exp(math.pi * math.tan(phi_rad)) * (math.tan(math.pi/4 + phi_rad/2)**2)
            Nc     = (Nq - 1) / math.tan(phi_rad)
            Ngamma = 2.0 * (Nq + 1.0) * math.tan(phi_rad)   # Vesic 1973 — difiere de Hansen

        # ── Factores de forma (Fcs, Fqs, Fγs) Vesic 1973 ─────────────────
        _BL = B_cp / L_cp if L_cp > 0 else 1.0
        if phi_ang == 0:
            Fcs    = 1.0 + _BL * (Nq / Nc) if Nc > 0 else 1.0
            Fqs    = 1.0
            Fgs    = max(1.0 - 0.4 * _BL, 0.60)
        else:
            Fcs    = 1.0 + _BL * (Nq / Nc) if Nc > 0 else 1.0
            Fqs    = 1.0 + _BL * math.tan(phi_rad)
            Fgs    = max(1.0 - 0.4 * _BL, 0.60)

        # ── Factores de profundidad (Fcd, Fqd, Fγd) Hansen/Vesic ─────────
        _rat = Df_cp / B_cp if B_cp > 0 else 0.0
        _k   = _rat if _rat <= 1.0 else math.atan(_rat)   # atan para Df/B > 1
        _t2  = math.tan(phi_rad) * (1 - math.sin(phi_rad))**2
        Fqd  = 1.0 + 2.0 * _t2 * _k
        Fgd  = 1.0
        if phi_ang == 0:
            Fcd = 1.0 + 0.4 * _k
        else:
            _dnom = Nq * math.tan(phi_rad)
            Fcd  = Fqd - (1 - Fqd) / _dnom if _dnom > 0 else Fqd

        # ── Input del ángulo de inclinación de la carga ψ ─────────────────
        _psi_inp = st.number_input(
            _t("Angulo inclinacion carga psi [grados] (0 = vertical)",
               "Load inclination angle psi [deg] (0 = vertical)"),
            0.0, 45.0,
            st.session_state.get("z_vesic_psi", 0.0), 1.0,
            key="z_vesic_psi",
            help=_t(
                "Factor de inclinacion Vesic 1973. psi=0 para carga vertical pura (Fci=Fqi=Fgi=1).",
                "Vesic 1973 inclination factor. psi=0 → Fci=Fqi=Fgi=1."
            )
        )

        # ── Factores de inclinación (Fci, Fqi, Fγi) Vesic 1973 ───────────
        _psi_rad = math.radians(float(st.session_state.get("z_vesic_psi", 0.0)))
        _n_v = (2.0 + _BL) / (1.0 + _BL)   # factor n de Vesic
        if _psi_rad == 0.0 or phi_ang == 0:
            Fci, Fqi, Fgi = 1.0, 1.0, 1.0
        else:
            Fqi = max((1.0 - math.tan(_psi_rad))**_n_v, 0.0)
            Fgi = max((1.0 - math.tan(_psi_rad))**(_n_v + 1), 0.0)
            _dnom_ci = Nc * math.tan(phi_rad) if (Nc > 0 and phi_ang > 0) else 1.0
            Fci = Fqi - (1.0 - Fqi) / _dnom_ci if _dnom_ci > 0 else 0.0
            Fci = max(Fci, 0.0)

        # Alias unificados para bloque de q_ult y tabla de resultados
        sc, sq, sgamma = Fcs, Fqs, Fgs
        dc, dq, dgamma = Fcd, Fqd, Fgd
        ic_v, iq_v, ig_v = Fci, Fqi, Fgi

    # Corrección por nivel freático (Casos I / II / III)
    gamma_prime = gamma_sat_t - gamma_w
    if NF_prof <= Df_cp:                         # Caso I: NF sobre el desplante
        q_sob    = gamma_s*NF_prof + gamma_prime*(Df_cp - NF_prof)
        gamma_eff = gamma_prime
        caso_nf  = "I"
    elif NF_prof < Df_cp + B_cp:                 # Caso II: NF entre Df y Df+B
        z2        = NF_prof - Df_cp
        gamma_eff = gamma_s + (z2/B_cp)*(gamma_prime - gamma_s) if B_cp > 0 else gamma_s
        q_sob    = gamma_s*Df_cp
        caso_nf  = "II"
    else:                                        # Caso III: NF muy profundo
        q_sob    = gamma_s*Df_cp
        gamma_eff = gamma_s
        caso_nf  = "III"

    # q_ult unificado — para Vesic incluye factores de inclinacion Fci/Fqi/Fgi
    # Para Terzaghi y Meyerhof ic_v=iq_v=ig_v=1 (neutros)
    q_ult  = (sc*dc*ic_v*coh_c*Nc
             + sq*dq*iq_v*q_sob*Nq
             + sgamma*dgamma*ig_v*0.5*gamma_eff*B_cp*Ngamma)
    q_adm  = q_ult / FS_terz
    # ── Defaults estructurales (sobreescritos en pestaña Diseño Estructural) ──
    # fc_basico y fy_basico ya definidos en sidebar
    recub_z    = float(st.session_state.get("z_recub_val", 7.5))
    c1_col     = float(st.session_state.get("z_c1col_val", 40.0))
    c2_col     = float(st.session_state.get("z_c2col_val", 40.0))
    bar_z      = st.session_state.get("z_barz_val", list(REBAR_DICT.keys())[3])
    if bar_z not in REBAR_DICT: bar_z = list(REBAR_DICT.keys())[3]
    db_bar_z   = REBAR_DICT[bar_z]["db"]
    A_bar_z    = REBAR_DICT[bar_z]["area"] * 100  # mm²
    H_zap      = float(st.session_state.get("z_hzap_val", Hzap))
    d_z        = max(1.0, H_zap - recub_z - db_bar_z / 10.0)  # cm peralte efectivo
    d_avg      = d_z   # peralte promedio en cm
    pos_col_iso = st.session_state.get("pos_col_iso", "Interior")
    gamma_prom  = float(st.session_state.get("z_gamma_prom", 20.0))
    Df_z        = float(st.session_state.get("z_df_z", 1.0))
    s_c_uso     = float(st.session_state.get("z_sc_uso", 3.0))
    q_adm_z     = float(q_adm)   # valor calculado; sobreescrito por widget ISO
    qu_max     = float(q_adm);  qu_min = 0.0   # presiones mayoradas (default hasta calcular ISO)
    q_net       = max(0.1, q_adm - gamma_prom * Df_z)
    # Defaults de resultados estructurales (sobreescritos en pestaña Diseño)
    ok_1way_B  = False;  ok_1way_L  = False;  ok_punz    = False
    nbarras_B  = 0;      nbarras_L  = 0
    n_barras_B = 0;      n_barras_L = 0
    sep_B      = 0.0;    sep_L      = 0.0
    Mu_B       = 0.0;    Mu_L       = 0.0
    Vu_1B      = 0.0;    Vu_1L      = 0.0;    Vu_punz    = 0.0
    phi_Vc_1B  = 0.0;    phi_Vc_1L  = 0.0;    phi_Vc_P   = 0.0
    Vu_1way_B     = 0.0
    Vu_1way_L     = 0.0
    phi_Vc_1way_B = 0.0
    phi_Vc_1way_L = 0.0
    Vu_punz       = 0.0
    phi_Vc_P      = 0.0
    ok_1way_L     = False
    ok_punz       = False
    bo_perim   = 0.0;    Mu_flex_B  = 0.0;    Mu_flex_L  = 0.0
    rho_B         = 0.0
    rho_L         = 0.0
    rho_use_B     = 0.0
    rho_use_L     = 0.0
    disc_B_val    = 0.0
    disc_L_val    = 0.0
    As_req_L      = 0.0
    n_barras_L    = 0
    sep_L         = 0.0
    Mu_flex_B     = 0.0
    Mu_flex_L     = 0.0
    disc_L        = 0.0
    L_use         = L_cp
    lv_l          = 0.0
    rho_use_L     = 0.0
    L_disp_min_cm = 0.0; _ld_req_cm = 0.0
    ldh_req_cm = 0.0;    ok_ldh     = False
    _ok_ld     = False;   _ld_disp_min_cm = 0.0
    _e_apu_z   = ''    # excepción gráfica APU (default vacío)
    _caso_fav_B = False;  _caso_fav_L = False
    _factor_ld  = 1.7    # factor desarrollo recto (caso general NSR-10)
    # Defaults cuantificación APU (calculados en tab_apu, usados en bloque PDF)
    _vol_conc_def  = B_cp * L_cp * (Hzap / 100.0)
    bultos_zap     = _vol_conc_def * 350 / 50.0          # bultos 50 kg
    vol_arena_z    = _vol_conc_def * 0.55                 # m³
    vol_grava_z    = _vol_conc_def * 0.80                 # m³
    litros_agua    = _vol_conc_def * 185.0                # litros
    vol_excavacion = (B_cp + 0.5) * (L_cp + 0.5) * Df_cp # m³
    peso_total_acero_zap = 0.0                            # kg (actualizado en tab_apu)
    vol_concreto_zap     = _vol_conc_def
    # Aliases para bloque PDF/DOCX y diagrama N-M
    Mu_flex_B  = Mu_B;   Mu_flex_L  = Mu_L
    momento_dir_B = Mu_B
    momento_dir_L = Mu_L
    rho_z_min  = 0.0018   # NSR-10 cuantía mínima zapata
    As_min     = 0.0
    vol_c      = 0.0;    peso_L = 0.0;  peso_B = 0.0
    kg_m3      = 2400.0
    MsvcB      = 0.0;    MsvcL  = 0.0
    _it        = 0
    L_disp_min_cm = 0.0;  L_disp_B_cm = 0.0;  L_disp_L_cm = 0.0
    _ld_disp_B_cm = 0.0;  _ld_disp_L_cm = 0.0
    lv_b       = 0.0;    lv_l       = 0.0
    d_z_m      = d_z / 100.0
    B_use      = B_cp;   L_use      = L_cp
    q_sob      = 0.0;    gamma_eff  = gamma_s
    caso_nf    = 1
    As_req_B   = 0.0;    As_req_L   = 0.0
    As_req_B     = 0.0;    As_req_L     = 0.0
    disc_B     = 0.0;    disc_L     = 0.0
    # Defaults adicionales para secciones posteriores
    _best      = None    # resultado optimizador
    _qu        = 0.0     # presión actuante optimizador
    file_name  = "zapata.dxf"
    mime       = "application/dxf"
    ok_volc    = False;  ok_desl = False
    _bo        = 0.0;    _bo1_eff = 0.0;  _bo2_eff = 0.0
    sep_max    = 45.0    # cm separación máxima NSR-10
    fontSize   = 10;     spaceBefore = 4;  colWidths = None
    tmp_path   = ""
    ldh_req_cm = 0.0;  ok_ldh  = False;  _ld_req_cm = 0.0
    _bz        = bar_z   # alias barra zapata en exportación
    pedestal   = False   # flag pedestal
    x1         = 0.0;    y1 = 0.0   # coords diagrama N-M
    borderwidth= 1       # estilo tabla
    ks         = 0.0     # módulo de balasto
    H_horiz    = float(st.session_state.get("z_Hhoriz", 0.0))
    FS_desl    = 999.0;  FS_volc = 999.0
    M_volc     = 0.0;    ok_volc = False;  ok_desl = False
    FR         = 0.0;    M_estab = 0.0;    arm_volc = 0.0
    delta_ang  = float(st.session_state.get("z_delta", 0.0))
    Pp_pasivo  = 0.0
    W_total    = 0.0;    P_svc   = float(Q_act)
    area_ok    = True    # default hasta que pestaña ISO calcule
    # Defaults para bloque PDF/DOCX (calculados en pestaña diseño)
    P_serv     = float(Q_act);  M_serv  = 0.0;  M_ult   = 0.0
    qa         = q_adm;         e_serv  = 0.0
    qmax_s     = q_adm;         qmin_s  = 0.0
    qmax_u     = q_adm * 1.3;   qmin_u  = 0.0
    P_svc      = float(Q_act);  P_ult   = float(Q_act) * 1.4
    q_act      = float(Q_act) / max(0.01, B_cp * L_cp)
    z_mid      = 0.0     # profundidad media capa asentamiento
    delta_color= "#14b8a6"  # color semáforo
    
    Q_ult  = q_ult * B_cp * L_cp
    q_act  = Q_act / (B_cp * L_cp)
    
    FS_calc = q_ult / q_act if q_act > 0 else 999.0
    cumplio = q_act <= q_adm

    #  RESULTADOS 
    st.divider()
    st.markdown(f"####  Resultados de Capacidad Portante ({metodo_cap})")
    col_c1, col_c2, col_c3 = st.columns(3)
    col_c1.metric("Esfuerzo Admisible (q_adm)", f"{q_adm:.2f} kPa")
    col_c2.metric("Esfuerzo Actuante (q_act)", f"{q_act:.2f} kPa", delta=f"{q_adm - q_act:.2f} kPa (Margen)", delta_color="normal")
    fs_calc_disp = "∞" if q_act == 0 else f"{FS_calc:.2f}"
    col_c3.metric("FS Calculado (q_ult / q_act)", fs_calc_disp, delta=f"Requerido: {FS_terz:.2f}", delta_color="off")
    
    if cumplio:
        st.success(f"**Verificación Exitosa:** El esfuerzo actuante ({q_act:.2f} kPa) es MENOR o IGUAL al esfuerzo admisible ({q_adm:.2f} kPa).")
    else:
        st.error(f"**Falla por Capacidad:** El esfuerzo actuante ({q_act:.2f} kPa) es MAYOR al esfuerzo admisible ({q_adm:.2f} kPa). Aumenta la sección de la zapata o reduce la carga.")

    col_r1, col_r2 = st.columns(2)
    with col_r1:
        _rows_factores = [
            {"Parámetro": "Método",                    "Valor": metodo_cap},
            {"Parámetro": "Nc / Nq / Nγ",              "Valor": f"{Nc:.3f} / {Nq:.3f} / {Ngamma:.3f}"},
            {"Parámetro": "sc / sq / sγ  (forma)",     "Valor": f"{sc:.4f} / {sq:.4f} / {sgamma:.4f}"},
            {"Parámetro": "dc / dq / dγ  (profundidad)","Valor": f"{dc:.4f} / {dq:.4f} / {dgamma:.4f}"},
            {"Parámetro": "ic / iq / iγ  (inclinación)","Valor": f"{ic_v:.4f} / {iq_v:.4f} / {ig_v:.4f}"},
            {"Parámetro": "Caso NF",                   "Valor": f"Caso {caso_nf} | NF={NF_prof} m"},
            {"Parámetro": "q sobrecarga",              "Valor": f"{q_sob:.2f} kPa"},
            {"Parámetro": "γ' efectivo",               "Valor": f"{gamma_eff:.2f} kN/m³"},
            {"Parámetro": "q_ult  [kPa]",              "Valor": f"{q_ult:.2f}"},
            {"Parámetro": "q_adm  [kPa]",              "Valor": f"{q_adm:.2f}"},
        ]
        st.dataframe(pd.DataFrame(_rows_factores), use_container_width=True, hide_index=True)
        if metodo_cap in ("Vesic", "Meyerhof"):
            _T1 = sc*dc*ic_v*coh_c*Nc
            _T2 = sq*dq*iq_v*q_sob*Nq
            _T3 = sgamma*dgamma*ig_v*0.5*gamma_eff*B_cp*Ngamma
            st.caption(
                f"Desglose q_ult = "
                f"**{_T1:.1f}** (cohesión) + **{_T2:.1f}** (sobrecarga) + "
                f"**{_T3:.1f}** (peso propio) = **{_T1+_T2+_T3:.1f} kPa**"
            )
    with col_r2:
        st.markdown("**Conversión Admisible:**")
        m1,m2,m3 = st.columns(3)
        m1.metric("kPa (kN/m²)",    f"{q_adm:.1f}")
        m2.metric("ton/m²", f"{q_adm/9.80665:.2f}")
        m3.metric("kg/cm²", f"{q_adm/98.0665:.3f}")


    # ── PASO 5E: Comparativa de métodos Terzaghi / Meyerhof / Vesic ─────────
    with st.expander(
        _t("Comparativa de Métodos: Terzaghi vs Meyerhof vs Vesic",
           "Methods Comparison: Terzaghi vs Meyerhof vs Vesic"),
        expanded=False
    ):
        st.markdown(_t(
            "Cálculo automático de q_ult con los tres métodos usando los mismos parámetros de entrada.",
            "Automatic q_ult calculation with all three methods using the same input parameters."
        ))
        # Calcular los tres métodos
        def _qult_metodo(metodo, phi_a, coh, gamms, q_sb, geff, B, L, Df):
            import math as _m
            _pr = _m.radians(phi_a)
            # Factores N
            if phi_a == 0:
                Nc, Nq, Ng = 5.14, 1.0, 0.0
            else:
                Nq = _m.exp(_m.pi * _m.tan(_pr)) * (_m.tan(_m.pi/4 + _pr/2)**2)
                Nc = (Nq - 1) / _m.tan(_pr)
                if metodo == "Terzaghi":
                    _at = _m.exp((0.75*_m.pi - _pr/2)*_m.tan(_pr))
                    Nq  = (_at**2) / (2*_m.cos(_m.radians(45) + _pr/2)**2)
                    Nc  = (Nq - 1) / _m.tan(_pr) if phi_a > 0 else 5.7
                    Ng  = 2*(Nq+1)*_m.tan(_pr) / (1 + 0.4*_m.sin(4*_pr))
                elif metodo == "Meyerhof":
                    Ng = (Nq - 1) * _m.tan(1.4 * _pr)
                else:  # Vesic
                    Ng = 2.0 * (Nq + 1.0) * _m.tan(_pr)
            if phi_a == 0:
                Nc, Nq, Ng = 5.14, 1.0, 0.0

            # Factores de forma
            _BL = B/L if L > 0 else 1.0
            if metodo == "Terzaghi":
                sc, sq, sg = 1.3, 1.0, 0.8  # cuadrada/continua simplificado
            elif metodo == "Meyerhof":
                if phi_a == 0:
                    sc, sq, sg = 1 + 0.2*_BL, 1.0, 1.0
                else:
                    _Kp = _m.tan(_m.pi/4 + _pr/2)**2
                    sc  = 1 + 0.2*_Kp*_BL
                    sq  = 1 + 0.1*_Kp*_BL if phi_a >= 10 else 1.0
                    sg  = sq
            else:  # Vesic
                _Nq2 = _m.exp(_m.pi*_m.tan(_pr))*(_m.tan(_m.pi/4+_pr/2)**2) if phi_a>0 else 1.0
                _Nc2 = (_Nq2-1)/_m.tan(_pr) if phi_a>0 else 5.14
                sc = 1 + _BL*(_Nq2/_Nc2) if _Nc2 > 0 else 1.0
                sq = 1 + _BL*_m.tan(_pr) if phi_a > 0 else 1.0
                sg = max(1 - 0.4*_BL, 0.6)

            # Factores de profundidad
            _rat = Df/B if B > 0 else 0.0
            if metodo == "Terzaghi":
                dc, dq, dg = 1.0, 1.0, 1.0
            elif metodo == "Meyerhof":
                if phi_a == 0:
                    dc, dq, dg = 1+0.2*_rat, 1.0, 1.0
                else:
                    _Kp = _m.tan(_m.pi/4+_pr/2)**2
                    dc  = 1+0.2*_m.sqrt(_Kp)*_rat
                    dq  = 1+0.1*_m.sqrt(_Kp)*_rat if phi_a>=10 else 1.0
                    dg  = dq
            else:  # Vesic
                _k = _rat if _rat <= 1.0 else _m.atan(_rat)
                _t2 = _m.tan(_pr)*(1-_m.sin(_pr))**2
                dq  = 1+2*_t2*_k
                dg  = 1.0
                dc  = dq - (1-dq)/(Nq*_m.tan(_pr)) if (phi_a>0 and Nq*_m.tan(_pr)>0) else (1+0.4*_k)

            qu = sc*dc*coh*Nc + sq*dq*q_sb*Nq + sg*dg*0.5*geff*B*Ng
            return qu, Nc, Nq, Ng, sc, sq, sg, dc, dq, dg

        _filas_cmp = []
        for _met in ["Terzaghi", "Meyerhof", "Vesic"]:
            _qu, _Nc, _Nq, _Ng, _sc, _sq, _sg, _dc, _dq, _dg = _qult_metodo(
                _met, phi_ang, coh_c, gamma_s, q_sob, gamma_eff,
                B_cp, L_cp, Df_cp
            )
            _filas_cmp.append({
                "Método": _met,
                "Nq": f"{_Nq:.3f}",
                "Nγ": f"{_Ng:.3f}",
                "sc/sq/sγ": f"{_sc:.3f}/{_sq:.3f}/{_sg:.3f}",
                "dc/dq/dγ": f"{_dc:.3f}/{_dq:.3f}/{_dg:.3f}",
                "q_ult [kPa]": f"{_qu:.1f}",
                "q_adm (FS=3) [kPa]": f"{_qu/FS_terz:.1f}",
                "Activo": "→ EN USO" if _met == metodo_cap else ""
            })
        st.dataframe(pd.DataFrame(_filas_cmp), use_container_width=True, hide_index=True)
        st.caption(_t(
            "Nγ Vesic = 2(Nq+1)tanφ | Nγ Meyerhof = (Nq-1)tan(1.4φ) | Terzaghi: sin factores d/s",
            "Nγ Vesic = 2(Nq+1)tanφ | Nγ Meyerhof = (Nq-1)tan(1.4φ) | Terzaghi: no d/s factors"
        ))

    
    # ── PASO 9A: Optimizador automático B×L (costo mínimo) ──────────────────
    with st.expander(
        _t("Optimizador B×L Automático (Costo Mínimo Concreto)",
           "Auto B×L Optimizer (Min Concrete Cost)"),
        expanded=False
    ):
        st.markdown(_t(
            "Dado q_adm, la carga de servicio Ps y el rango de dimensiones, "
            "calcula todas las combinaciones factibles y ordena por volumen mínimo de concreto.",
            "Given q_adm, service load Ps and dimension range, calculates all feasible "
            "combinations and sorts by minimum concrete volume."
        ))
        _opt_c1, _opt_c2 = st.columns(2)
        with _opt_c1:
            _opt_Ps    = st.number_input(_t("Carga de servicio Ps (kN)","Service load Ps (kN)"),
                10.0, 50000.0, float(Q_act), 50.0, key="opt_Ps")
            _opt_qadm  = st.number_input(_t("q_adm (kPa)","q_adm (kPa)"),
                10.0, 2000.0, float(q_adm), 10.0, key="opt_qadm")
            _opt_H     = st.number_input(_t("Altura zapata H (cm)","Footing height H (cm)"),
                20.0, 150.0, float(Hzap), 5.0, key="opt_H")
        with _opt_c2:
            _opt_Bmin  = st.number_input(_t("B mínimo (m)","B min (m)"),
                0.5, 5.0, max(0.5, round(B_cp - 0.5, 1)), 0.1, key="opt_Bmin")
            _opt_Bmax  = st.number_input(_t("B máximo (m)","B max (m)"),
                1.0, 12.0, min(12.0, round(B_cp + 1.5, 1)), 0.1, key="opt_Bmax")
            _opt_paso  = st.select_slider(_t("Paso de búsqueda (m)","Search step (m)"),
                options=[0.05, 0.10, 0.20, 0.25, 0.50], value=0.10, key="opt_paso")
            _opt_cuad  = st.checkbox(_t("Solo zapatas cuadradas (B=L)","Square only (B=L)"),
                value=False, key="opt_cuad")

        if st.button(_t("Calcular Optimizador","Run Optimizer"),
                     key="opt_run", use_container_width=True):
            import numpy as _np_opt
            _B_range = _np_opt.arange(_opt_Bmin, _opt_Bmax + _opt_paso/2, _opt_paso)
            _rows_opt = []
            for _B in _B_range:
                _B = round(float(_B), 3)
                _L_range2 = [_B] if _opt_cuad else _np_opt.arange(
                    _B, min(_B * 2.5, _opt_Bmax) + _opt_paso/2, _opt_paso)
                for _L in _L_range2:
                    _L = round(float(_L), 3)
                    _q_act_o = _opt_Ps / (_B * _L)
                    if _q_act_o > _opt_qadm:
                        continue   # no cumple portante
                    _vol = _B * _L * (_opt_H / 100.0)
                    _FS_o= q_ult / _q_act_o if _q_act_o > 0 else 999.0
                    _rows_opt.append({
                        "B (m)": _B, "L (m)": _L,
                        "Área (m²)": round(_B * _L, 3),
                        "q_act (kPa)": round(_q_act_o, 1),
                        "Margen q (kPa)": round(_opt_qadm - _q_act_o, 1),
                        "FS calc.": round(_FS_o, 2) if _FS_o < 999 else "∞",
                        "Vol. concreto (m³)": round(_vol, 3),
                        "B/L": round(_B/_L, 2),
                        "Estado": "OK" if _q_act_o <= _opt_qadm else "FALLA",
                    })

            if _rows_opt:
                _df_opt = pd.DataFrame(_rows_opt).sort_values("Vol. concreto (m³)")
                _df_opt = _df_opt.reset_index(drop=True)
                _df_opt.index += 1

                # Semáforo estado
                def _color_opt(val):
                    if val == "OK":
                        return "background-color:#d4edda;color:#155724;font-weight:bold"
                    return "background-color:#f8d7da;color:#721c24;font-weight:bold"

                st.dataframe(
                    _df_opt.style.map(_color_opt, subset=["Estado"]),
                    use_container_width=True, hide_index=False
                )
                # Métricas del óptimo
                _best = _df_opt.iloc[0]
                _c_o1, _c_o2, _c_o3, _c_o4 = st.columns(4)
                _c_o1.metric(_t("Geometría óptima","Optimal geometry"),
                             f"{_best['B (m)']:.2f} × {_best['L (m)']:.2f} m")
                _c_o2.metric(_t("Vol. mínimo","Min volume"),
                             f"{_best['Vol. concreto (m³)']:.3f} m³")
                _c_o3.metric(_t("q_act óptima","Optimal q_act"),
                             f"{_best['q_act (kPa)']:.1f} kPa")
                _c_o4.metric(_t("Combinaciones válidas","Valid combinations"),
                             len(_df_opt))

                # CSV del optimizador
                _csv_opt = _df_opt.to_csv().encode("utf-8-sig")
                st.download_button(
                    _t("Descargar CSV Optimizador","Download Optimizer CSV"),
                    data=_csv_opt,
                    file_name=f"Optimizador_Zapata_{_opt_Ps:.0f}kN.csv",
                    mime="text/csv", key="opt_csv", use_container_width=True
                )
                st.caption(_t(
                    f"Ordenado por volumen mínimo. H fijo = {_opt_H:.0f} cm. "
                    f"q_ult = {q_ult:.1f} kPa (método {metodo_cap}).",
                    f"Sorted by minimum volume. H fixed = {_opt_H:.0f} cm. "
                    f"q_ult = {q_ult:.1f} kPa ({metodo_cap} method)."
                ))
            else:
                st.warning(_t(
                    "No hay combinaciones válidas en el rango definido. "
                    "Aumenta B_max o revisa q_adm.",
                    "No valid combinations in the defined range. "
                    "Increase B_max or check q_adm."
                ))

    # ── P3-A: Comparador de Alternativas con Semáforo ────────────────────────
    with st.expander(
        _t("Comparador de Alternativas (Tabla Semáforo B₁/B₂/B₃)",
           "Alternatives Comparison (Traffic-Light Table B₁/B₂/B₃)"),
        expanded=False
    ):
        st.markdown(_t(
            "Define hasta 4 geometrías alternativas de zapata y compara simultáneamente "
            "todas las verificaciones con semáforo visual. Útil para decisiones rápidas "
            "en reuniones de diseño o propuestas al cliente.",
            "Define up to 4 alternative footing geometries and compare all verifications "
            "simultaneously with a traffic-light table. Useful for quick design decisions."
        ))

        _alt_n = st.slider(_t("Número de alternativas","Number of alternatives"), 2, 4, 3,
                           key="alt_n")
        _alt_cols = st.columns(_alt_n)
        _alt_geoms = []
        for _ia in range(_alt_n):
            with _alt_cols[_ia]:
                st.markdown(f"**Alt. {_ia+1}**")
                _aB = st.number_input(f"B{_ia+1} (m)", 0.5, 10.0,
                    round(B_cp + (_ia - 1) * 0.25, 2), 0.05, key=f"alt_B_{_ia}")
                _aL = st.number_input(f"L{_ia+1} (m)", 0.5, 12.0,
                    round(L_cp + (_ia - 1) * 0.25, 2), 0.05, key=f"alt_L_{_ia}")
                _aH = st.number_input(f"H{_ia+1} (cm)", 20.0, 150.0,
                    float(Hzap), 5.0, key=f"alt_H_{_ia}")
                _alt_geoms.append({"B": _aB, "L": _aL, "H": _aH})

        # Calcular todas las alternativas
        _alt_rows = []
        for _ia, _g in enumerate(_alt_geoms):
            _B, _L, _H = _g["B"], _g["L"], _g["H"]
            _Hm  = _H / 100.0
            _qAct= Q_act / (_B * _L)
            _ok_q= _qAct <= q_adm
            _d_e = _H - recub_z - db_bar_z/10/2   # cm
            _Vu1B= max(0, _qAct * _B * max(0, (_B/2 - c1_col/200 - _d_e/100)))  # kN
            _phiVcB = 0.75 * (0.17 * (fc_basico**0.5)) * (_B*100) * _d_e / 1000
            _ok_1B  = _Vu1B <= _phiVcB
            _Vu1L= max(0, _qAct * _L * max(0, (_L/2 - c2_col/200 - _d_e/100)))
            _phiVcL = 0.75 * (0.17 * (fc_basico**0.5)) * (_L*100) * _d_e / 1000
            _ok_1L  = _Vu1L <= _phiVcL
            _bo  = 2 * ((c1_col + _d_e) + (c2_col + _d_e))
            _VuP = max(0, _qAct * (_B*_L - ((c1_col/100+_d_e/100)*(c2_col/100+_d_e/100))))
            _phiVcP = 0.75 * 0.33 * (fc_basico**0.5) * (_bo/100) * (_d_e/100) * 1000
            _ok_P   = _VuP <= _phiVcP
            _vol = _B * _L * _Hm
            _FS_a= q_ult / _qAct if _qAct > 0 else 999.0
            _all_ok = all([_ok_q, _ok_1B, _ok_1L, _ok_P])
            _alt_rows.append({
                "Alternativa": f"Alt. {_ia+1}  {_B:.2f}×{_L:.2f}×{_H:.0f}cm",
                "B×L (m)": f"{_B:.2f}×{_L:.2f}",
                "H (cm)": _H,
                "Área (m²)": round(_B*_L, 3),
                "Vol. concreto (m³)": round(_vol, 3),
                "q_act (kPa)": round(_qAct, 1),
                "FS": round(_FS_a, 2) if _FS_a < 999 else "∞",
                "Portante": "OK" if _ok_q else "FALLA",
                "Cortante B": "OK" if _ok_1B else "FALLA",
                "Cortante L": "OK" if _ok_1L else "FALLA",
                "Punzonamiento": "OK" if _ok_P else "FALLA",
                "Global": "✅ CUMPLE" if _all_ok else "❌ FALLA",
            })

        _df_alt = pd.DataFrame(_alt_rows)

        def _color_alt(val):
            if val in ("OK", "✅ CUMPLE"):
                return "background-color:#d4edda;color:#155724;font-weight:bold"
            elif val in ("FALLA", "❌ FALLA"):
                return "background-color:#f8d7da;color:#721c24;font-weight:bold"
            return ""

        _cols_sem = ["Portante","Cortante B","Cortante L","Punzonamiento","Global"]
        st.dataframe(
            _df_alt.style.map(_color_alt, subset=_cols_sem),
            use_container_width=True, hide_index=True
        )

        # Recomendación automática
        _best_alts = [r for r in _alt_rows if r["Global"] == "✅ CUMPLE"]
        if _best_alts:
            _best_vol_alt = min(_best_alts, key=lambda x: x["Vol. concreto (m³)"])
            st.success(_t(
                f"Alternativa óptima: **{_best_vol_alt['B×L (m)']} m** con "
                f"{_best_vol_alt['Vol. concreto (m³)']:.3f} m³ — menor costo entre las que CUMPLEN.",
                f"Optimal alternative: **{_best_vol_alt['B×L (m)']} m** with "
                f"{_best_vol_alt['Vol. concreto (m³)']:.3f} m³ — lowest cost among compliant."
            ))
        else:
            st.warning(_t(
                "Ninguna alternativa cumple todas las verificaciones. Aumenta H o B.",
                "No alternative passes all verifications. Increase H or B."
            ))

        _csv_alt = _df_alt.to_csv(index=False).encode("utf-8-sig")
        st.download_button(_t("CSV Comparativa","Comparison CSV"),
            _csv_alt, f"Comparativa_Zapata_{B_cp:.1f}x{L_cp:.1f}.csv",
            "text/csv", key="alt_csv", use_container_width=True)

    # ── PASO 9B: Perfil estratigráfico multicapa + Steinbrenner ─────────────
    with st.expander(
        _t("Perfil Estratigráfico y Asentamiento (Steinbrenner / Elsticos)",
           "Soil Profile and Settlement (Steinbrenner / Elastic)"),
        expanded=False
    ):
        st.markdown(_t(
            "Define hasta 5 capas de suelo. Calcula el asentamiento elástico inmediato "
            "por el método de Steinbrenner (1934) con factor de forma Iz por profundidad. "
            "Ref: Bowles (1996) Cap. 5, NSR-10 H.3.",
            "Define up to 5 soil layers. Calculates immediate elastic settlement using "
            "Steinbrenner (1934) with depth influence factor Iz. "
            "Ref: Bowles (1996) Cap. 5, NSR-10 H.3."
        ))

        _nlayers = st.slider(_t("Número de capas","Number of layers"), 1, 5, 3, key="stb_nlayers")
        st.markdown(_t("**Parámetros por capa:**","**Layer parameters:**"))

        _layer_data = []
        _cols_hdr = st.columns([1.5, 1.5, 1.5, 1.5, 1.5])
        for _ch, _lbl in zip(_cols_hdr,
            [_t("Capa","Layer"), _t("Espesor H (m)","Thickness H (m)"),
             _t("Es (MPa)","Es (MPa)"), "ν (Poisson)",
             _t("Tipo","Type")]):
            _ch.caption(_lbl)

        for _il in range(_nlayers):
            _r = st.columns([1.5, 1.5, 1.5, 1.5, 1.5])
            _lbl_c = _r[0].text_input(f"Capa {_il+1}", value=f"Capa {_il+1}",
                key=f"stb_nom_{_il}", label_visibility="collapsed")
            _H_c   = _r[1].number_input(f"H{_il+1}", 0.1, 20.0,
                [2.0, 3.0, 5.0, 4.0, 6.0][_il] if _il < 5 else 2.0,
                0.1, key=f"stb_H_{_il}", label_visibility="collapsed")
            _Es_c  = _r[2].number_input(f"Es{_il+1}", 1.0, 500.0,
                [15.0, 25.0, 40.0, 60.0, 80.0][_il] if _il < 5 else 20.0,
                1.0, key=f"stb_Es_{_il}", label_visibility="collapsed")
            _nu_c  = _r[3].number_input(f"nu{_il+1}", 0.1, 0.499,
                [0.35, 0.30, 0.28, 0.25, 0.22][_il] if _il < 5 else 0.30,
                0.01, key=f"stb_nu_{_il}", label_visibility="collapsed")
            _tipo_c= _r[4].selectbox(f"T{_il+1}",
                [_t("Arcilla","Clay"), _t("Arena","Sand"),
                 _t("Limo","Silt"), _t("Roca","Rock")],
                key=f"stb_tipo_{_il}", label_visibility="collapsed")
            _layer_data.append({
                "nombre": _lbl_c, "H": _H_c, "Es": _Es_c * 1000.0,  # kPa
                "nu": _nu_c, "tipo": _tipo_c
            })

        # ── Cálculo Steinbrenner ─────────────────────────────────────────
        # Se_total = sum_capas [ (q * B * (1-ν²) / Es) * Iz_m * F1 ]
        # Iz: Factor de influencia de Steinbrenner en el centroide de cada capa
        import math as _math_stb
        import numpy as _np_stb
        import plotly.graph_objects as _go_stb

        _q_stb  = q_act         # esfuerzo aplicado kPa
        _B_stb  = B_cp          # m
        _L_stb  = L_cp          # m
        _Df_stb = Df_cp         # m

        def _iz_steinbrenner(z, B, L):
            """Factor de influencia Steinbrenner (1934) en profundidad z bajo la esquina."""
            _m = L / B if B > 0 else 1.0
            _n = z / B if B > 0 else 0.0
            if _n == 0:
                return 0.25
            _A1 = _n / (1 + _n**2)**0.5 * 1 / (1 + _m**2 + _n**2)**0.5
            _A2 = _m / (1 + _n**2 + _m**2)**0.5
            _B1 = _math_stb.atan2(
                _m * _n,
                _math_stb.sqrt(1 + _m**2 + _n**2) * 1.0
            )
            _F1 = (1 / (2 * _math_stb.pi)) * (_math_stb.atan2(_m * _n,
                    _np_stb.sqrt(1 + _m**2 + _n**2)) +
                   _m * _n * (1/(1 + _n**2)**0.5 / (1 + _m**2 + _n**2)**0.5 +
                               1/(1 + _m**2 + _n**2)**0.5))
            return max(0.0, min(0.5, _F1))

        # Calcular asentamiento por capas
        _asen_rows = []
        _z_acum = 0.0
        _Se_total_mm = 0.0
        _z_plot = [0.0]
        _Se_plot = [0.0]

        for _il, _layer in enumerate(_layer_data):
            _z_mid  = _z_acum + _layer["H"] / 2.0
            _Iz_mid = _iz_steinbrenner(_z_mid, _B_stb / 2, _L_stb / 2)  # esquina → centro: /2
            # Fórmula simplificada: Se_i = (q_net * B * Iz * (1-ν²)) / Es_i
            _Se_i   = (_q_stb * _B_stb * _Iz_mid * (1 - _layer["nu"]**2)) / _layer["Es"]
            _Se_i_mm= _Se_i * 1000.0  # mm
            _Se_total_mm += _Se_i_mm
            _asen_rows.append({
                "Capa": _layer["nombre"],
                "Tipo": _layer["tipo"],
                "H (m)": _layer["H"],
                "Es (MPa)": round(_layer["Es"]/1000, 1),
                "ν": _layer["nu"],
                "z_mid (m)": round(_z_mid, 2),
                "Iz": round(_Iz_mid, 4),
                "Se_i (mm)": round(_Se_i_mm, 2),
                "Se acum. (mm)": round(_Se_total_mm, 2),
            })
            _z_acum += _layer["H"]
            _z_plot.append(_z_acum)
            _Se_plot.append(_Se_total_mm)

        # Tabla de capas
        _df_stb = pd.DataFrame(_asen_rows)
        st.dataframe(_df_stb, use_container_width=True, hide_index=True)

        # Métricas
        _LIMIT_NSR = 25.0  # mm — NSR-10 H.3 asentamiento máximo
        _ok_asen   = _Se_total_mm <= _LIMIT_NSR
        _s1, _s2, _s3 = st.columns(3)
        _s1.metric(_t("Asentamiento total","Total settlement"),
                   f"{_Se_total_mm:.1f} mm",
                   delta=f"Límite NSR-10: {_LIMIT_NSR:.0f} mm",
                   delta_color="normal" if _ok_asen else "inverse")
        _s2.metric(_t("q neto aplicado","Applied net q"), f"{_q_stb:.1f} kPa")
        _s3.metric(_t("Capas analizadas","Layers analyzed"), _nlayers)

        if _ok_asen:
            st.success(_t(
                f"CUMPLE — Asentamiento {_Se_total_mm:.1f} mm ≤ {_LIMIT_NSR:.0f} mm (NSR-10 H.3).",
                f"OK — Settlement {_Se_total_mm:.1f} mm ≤ {_LIMIT_NSR:.0f} mm (NSR-10 H.3)."
            ))
        else:
            st.error(_t(
                f"NO CUMPLE — Asentamiento {_Se_total_mm:.1f} mm > {_LIMIT_NSR:.0f} mm. "
                "Considerar pilotaje o mejoramiento de suelo.",
                f"FAIL — Settlement {_Se_total_mm:.1f} mm > {_LIMIT_NSR:.0f} mm. "
                "Consider piling or soil improvement."
            ))

        # Perfil estratigráfico + curva asentamiento
        _fig_stb = _go_stb.Figure()
        # Bandas de capas
        _z_c = 0.0
        _colores_stb = ["#c7a98a","#d4b896","#e8d5b7","#b8cce4","#8db3d4"]
        for _il, _layer in enumerate(_layer_data):
            _fig_stb.add_shape(type="rect",
                x0=0, x1=_Se_plot[-1]*1.2 + 5,
                y0=-(_z_c + _layer["H"]), y1=-_z_c,
                fillcolor=_colores_stb[_il % len(_colores_stb)],
                opacity=0.35, line=dict(color="#888",width=0.5))
            _fig_stb.add_annotation(
                x=(_Se_plot[-1]*1.2 + 5)*0.85,
                y=-(_z_c + _layer["H"]/2),
                text=f"{_layer['nombre']}<br>Es={_layer['Es']/1000:.0f}MPa",
                showarrow=False, font=dict(size=9, color="white"))
            _z_c += _layer["H"]

        # Curva de asentamiento acumulado
        _fig_stb.add_trace(_go_stb.Scatter(
            x=_Se_plot, y=[-z for z in _z_plot],
            mode="lines+markers",
            line=dict(color="#f59e0b", width=2.5),
            marker=dict(color="#f59e0b", size=7),
            name=_t("Se acum. (mm)","Cumulative Se (mm)")
        ))
        # Línea límite NSR-10
        _fig_stb.add_vline(x=_LIMIT_NSR, line=dict(
            color="#ef4444", width=1.5, dash="dash"),
            annotation_text="NSR-10 25mm",
            annotation_font=dict(color="#ef4444", size=9))

        _fig_stb.update_layout(
            xaxis=dict(title=_t("Asentamiento acumulado (mm)","Cumulative settlement (mm)"),
                       gridcolor="#334155", color="white", rangemode="nonnegative"),
            yaxis=dict(title=_t("Profundidad (m)","Depth (m)"),
                       gridcolor="#334155", color="white"),
            paper_bgcolor="#0f1117", plot_bgcolor="#161b22",
            font=dict(color="white", size=11),
            legend=dict(bgcolor="#1e293b", bordercolor="#334155",
                        font=dict(color="white", size=10)),
            margin=dict(l=50, r=20, t=40, b=40),
            title=dict(
                text=_t(
                    f"Perfil Estratigráfico — Se_total = {_Se_total_mm:.1f} mm",
                    f"Soil Profile — Se_total = {_Se_total_mm:.1f} mm"
                ),
                font=dict(color="white", size=12), x=0.5
            )
        )
        st.plotly_chart(_fig_stb, use_container_width=True)
        st.caption(_t(
            "Método Steinbrenner (1934) — Factor Iz por capa centroidal. "
            "Ref: Bowles (1996) Cap. 5 | NSR-10 H.3 | Límite: 25 mm total.",
            "Steinbrenner (1934) method — Iz factor at layer centroid. "
            "Ref: Bowles (1996) Ch.5 | NSR-10 H.3 | Limit: 25 mm total."
        ))

    # ── PASO 7A: Historial de iteraciones por proyecto ───────────────────────
    with st.expander(
        _t("Historial de Iteraciones (B×L×H)", "Iteration History (B×L×H)"),
        expanded=False
    ):
        st.markdown(_t(
            "Cada vez que calculas, la iteración se guarda automáticamente. "
            "Compara geometrías y costos sin perder el diseño previo.",
            "Each calculation is saved automatically. "
            "Compare geometries and costs without losing previous designs."
        ))

        # ── Guardar iteración actual ─────────────────────────────────────
        import datetime as _dt_hist
        if "zap_historial" not in st.session_state:
            st.session_state["zap_historial"] = []

        _iter_actual = {
            "Iter": len(st.session_state["zap_historial"]) + 1,
            "Timestamp": _dt_hist.datetime.now().strftime("%H:%M:%S"),
            "B (m)": round(B_cp, 2),
            "L (m)": round(L_cp, 2),
            "H (cm)": round(Hzap, 1),
            "q_ult (kPa)": round(q_ult, 1),
            "q_adm (kPa)": round(q_adm, 1),
            "q_act (kPa)": round(q_act, 1),
            "FS calc.": round(FS_calc, 2) if FS_calc < 999 else "∞",
            "Cap. portante": "OK" if cumplio else "FALLA",
            "Cortante B": "OK" if ok_1way_B else "FALLA",
            "Cortante L": "OK" if ok_1way_L else "FALLA",
            "Punzonamiento": "OK" if ok_punz else "FALLA",
            "Barras B": f"{nbarras_B}×{bar_z}@{sep_B:.0f}cm",
            "Barras L": f"{nbarras_L}×{bar_z}@{sep_L:.0f}cm",
            "Vol. concreto (m³)": round(B_cp * L_cp * (Hzap/100), 3),
            "Método": metodo_cap,
        }

        # Botón guardar iteración
        _col_h1, _col_h2, _col_h3 = st.columns([2, 2, 1])
        with _col_h1:
            if st.button(
                _t("Guardar iteración actual", "Save current iteration"),
                key="zap_hist_save", use_container_width=True
            ):
                # Evitar duplicados exactos consecutivos
                if (not st.session_state["zap_historial"] or
                    st.session_state["zap_historial"][-1]["B (m)"] != _iter_actual["B (m)"] or
                    st.session_state["zap_historial"][-1]["L (m)"] != _iter_actual["L (m)"] or
                    st.session_state["zap_historial"][-1]["H (cm)"] != _iter_actual["H (cm)"]):
                    st.session_state["zap_historial"].append(_iter_actual)
                    st.success(_t(
                        f"Iteración {_iter_actual['Iter']} guardada.",
                        f"Iteration {_iter_actual['Iter']} saved."
                    ))
                else:
                    st.info(_t(
                        "Geometría idéntica a la iteración previa. Cambia B, L o H primero.",
                        "Identical geometry to previous iteration. Change B, L or H first."
                    ))
        with _col_h2:
            if st.button(
                _t("Limpiar historial", "Clear history"),
                key="zap_hist_clear", use_container_width=True
            ):
                st.session_state["zap_historial"] = []
                st.rerun()
        with _col_h3:
            st.caption(_t(
                f"{len(st.session_state['zap_historial'])} iter. guardadas",
                f"{len(st.session_state['zap_historial'])} saved iters."
            ))

        # ── Tabla comparativa ────────────────────────────────────────────
        if st.session_state["zap_historial"]:
            _df_hist = pd.DataFrame(st.session_state["zap_historial"])
            # Semáforo por columna de estado
            def _color_estado(val):
                if val == "OK":
                    return "background-color:#d4edda;color:#155724;font-weight:bold"
                elif val == "FALLA":
                    return "background-color:#f8d7da;color:#721c24;font-weight:bold"
                return ""
            _cols_estado = ["Cap. portante","Cortante B","Cortante L","Punzonamiento"]
            _styled = _df_hist.style.map(_color_estado, subset=_cols_estado)
            st.dataframe(_styled, use_container_width=True, hide_index=True)

            # Resumen mínimo/máximo
            _best_fs = _df_hist["FS calc."].replace("∞", 9999).astype(float).max()
            _best_vol = _df_hist["Vol. concreto (m³)"].min()
            _c_s1, _c_s2, _c_s3 = st.columns(3)
            _c_s1.metric(_t("Mejor FS calculado","Best FS"), f"{_best_fs:.2f}")
            _c_s2.metric(_t("Menor vol. concreto","Min concrete vol."), f"{_best_vol:.3f} m³")
            _c_s3.metric(_t("Iteraciones guardadas","Saved iterations"),
                         len(st.session_state["zap_historial"]))

            # ── PASO 7B: Exportar CSV ────────────────────────────────────
            st.markdown("---")
            st.markdown(_t("#### Exportar Resultados CSV","#### Export Results CSV"))
            _csv_buf = _df_hist.to_csv(index=False).encode("utf-8-sig")
            _col_csv1, _col_csv2 = st.columns(2)
            with _col_csv1:
                st.download_button(
                    _t("Descargar CSV Historial","Download History CSV"),
                    data=_csv_buf,
                    file_name=f"Historial_Zapata_{B_cp:.1f}x{L_cp:.1f}.csv",
                    mime="text/csv",
                    key="zap_hist_csv",
                    use_container_width=True
                )
            with _col_csv2:
                # CSV iteración actual (snapshot instantáneo)
                _csv_snap = pd.DataFrame([_iter_actual]).to_csv(index=False).encode("utf-8-sig")
                st.download_button(
                    _t("CSV Iteración Actual","Current Iteration CSV"),
                    data=_csv_snap,
                    file_name=f"Zapata_{B_cp:.1f}x{L_cp:.1f}_snap.csv",
                    mime="text/csv",
                    key="zap_snap_csv",
                    use_container_width=True
                )
        else:
            st.info(_t(
                "No hay iteraciones guardadas. Presiona 'Guardar iteración actual' para comenzar.",
                "No saved iterations. Press 'Save current iteration' to start."
            ))

    # ── Defaults doble parrilla (se sobreescriben al calcular) ──────────────────
    if 'dp_B' not in dir():
        dp_B = False; dp_L = False; _doble_parrilla = False
        As_req_B_sup = 0.0; As_req_L_sup = 0.0
        n_barras_B_sup = 0; n_barras_L_sup = 0
        sep_B_sup = 0.0; sep_L_sup = 0.0
        Mn_max_B = 0.0; Mn_max_L = 0.0
        aviso_dp_B = ""; aviso_dp_L = ""
        _rho_max_z = 0.0; _rho_bal_z = 0.0; _rho_min_zap = 0.0018

    # ── PASO 7C: Visualización 3D armadura real (parrilla Plotly) ────────────
    with st.expander(
        _t("Visualización 3D Armadura (Parrilla Real)", "3D Reinforcement View (Real Grid)"),
        expanded=False
    ):
        import plotly.graph_objects as _go3d
        import numpy as _np3d

        _B3  = B_cp
        _L3  = L_cp
        _H3  = Hzap / 100.0      # m
        _rec = recub_z / 100.0   # m
        _db  = db_bar_z / 1000.0 # m
        _d3  = d_avg / 100.0     # m
        _nB  = int(nbarras_B)
        _nL  = int(nbarras_L)
        _sep_B3 = _sep_B = _B3 / (_nB + 1) if _nB > 0 else _B3
        _sep_L3 = _sep_L = _L3 / (_nL + 1) if _nL > 0 else _L3

        _traces = []

        # ── Contorno zapata (bloque concreto transparente) ───────────────
        _verts = [
            [0,0,0],[_B3,0,0],[_B3,_L3,0],[0,_L3,0],  # base
            [0,0,_H3],[_B3,0,_H3],[_B3,_L3,_H3],[0,_L3,_H3]  # tapa
        ]
        _i_box = [0,0,1,1,4,4,5,5,0,0,3,3]
        _j_box = [1,3,2,5,5,7,6,6,4,1,7,2]
        _k_box = [3,2,6,6,7,6,7,7,5,5,4,6]  # simplificado — usar surface edges
        _xs_b = [v[0] for v in _verts]
        _ys_b = [v[1] for v in _verts]
        _zs_b = [v[2] for v in _verts]
        # Aristas de la caja
        _edges_box = [
            (0,1),(1,2),(2,3),(3,0),   # base
            (4,5),(5,6),(6,7),(7,4),   # tapa
            (0,4),(1,5),(2,6),(3,7)    # laterales
        ]
        for _e in _edges_box:
            _x0,_y0,_z0 = _verts[_e[0]]
            _x1,_y1,_z1 = _verts[_e[1]]
            _traces.append(_go3d.Scatter3d(
                x=[_x0,_x1], y=[_y0,_y1], z=[_z0,_z1],
                mode="lines",
                line=dict(color="#64748b", width=2),
                showlegend=False, hoverinfo="skip"
            ))

        # ── Barras dir. B (paralelas al eje X, z = rec + db/2) ───────────
        _z_barB = _rec + _db / 2
        for _j in range(_nB):
            _y_j = (_j + 1) * _sep_B3
            _x_arr = _np3d.linspace(_rec, _B3 - _rec, 50)
            _traces.append(_go3d.Scatter3d(
                x=_x_arr,
                y=_np3d.full(50, _y_j),
                z=_np3d.full(50, _z_barB),
                mode="lines",
                line=dict(color="#f59e0b", width=4),
                name=_t(f"Barra B dir. {bar_z}", f"Bar B dir. {bar_z}") if _j == 0 else None,
                showlegend=_j == 0,
                hovertemplate=f"Dir. B — {bar_z}<br>y={_y_j:.2f}m<extra></extra>"
            ))

        # ── Barras dir. L (paralelas al eje Y, z = rec + 1.5*db) ────────
        _z_barL = _rec + 1.5 * _db
        for _i in range(_nL):
            _x_i = (_i + 1) * _sep_L3
            _y_arr = _np3d.linspace(_rec, _L3 - _rec, 50)
            _traces.append(_go3d.Scatter3d(
                x=_np3d.full(50, _x_i),
                y=_y_arr,
                z=_np3d.full(50, _z_barL),
                mode="lines",
                line=dict(color="#38bdf8", width=4),
                name=_t(f"Barra L dir. {bar_z}", f"Bar L dir. {bar_z}") if _i == 0 else None,
                showlegend=_i == 0,
                hovertemplate=f"Dir. L — {bar_z}<br>x={_x_i:.2f}m<extra></extra>"
            ))

        # ── Pedestal (columna) ───────────────────────────────────────────
        _cx1 = (_B3 - c1_col/100) / 2
        _cx2 = (_B3 + c1_col/100) / 2
        _cy1 = (_L3 - c2_col/100) / 2
        _cy2 = (_L3 + c2_col/100) / 2
        _H_ped = _H3 + 0.40   # 40 cm de pedestal visible
        _edges_ped = [
            (_cx1,_cy1,_H3),(_cx2,_cy1,_H3),
            (_cx2,_cy2,_H3),(_cx1,_cy2,_H3),(_cx1,_cy1,_H3),  # base
            (_cx1,_cy1,_H_ped),(_cx2,_cy1,_H_ped),
            (_cx2,_cy2,_H_ped),(_cx1,_cy2,_H_ped),(_cx1,_cy1,_H_ped),  # tapa
        ]
        _traces.append(_go3d.Scatter3d(
            x=[p[0] for p in _edges_ped],
            y=[p[1] for p in _edges_ped],
            z=[p[2] for p in _edges_ped],
            mode="lines",
            line=dict(color="#94a3b8", width=3),
            name=_t("Pedestal","Column"), showlegend=True,
            hoverinfo="skip"
        ))
        for _xp, _yp in [(_cx1,_cy1),(_cx2,_cy1),(_cx2,_cy2),(_cx1,_cy2)]:
            _traces.append(_go3d.Scatter3d(
                x=[_xp,_xp], y=[_yp,_yp], z=[_H3,_H_ped],
                mode="lines", line=dict(color="#94a3b8", width=3),
                showlegend=False, hoverinfo="skip"
            ))

        # ── Barras SUPERIORES dir. B (doble parrilla) — cara comprimida ─────
        if dp_B and n_barras_B_sup > 0:
            _z_barB_sup = _H3 - _rec - _db / 2
            _sep_B_sup3 = (_L3 - 2*_rec) / max(n_barras_B_sup - 1, 1)
            for _j in range(n_barras_B_sup):
                _y_j_s = _rec + _j * _sep_B_sup3
                _x_arr_s = _np3d.linspace(_rec, _B3 - _rec, 50)
                _traces.append(_go3d.Scatter3d(
                    x=_x_arr_s,
                    y=_np3d.full(50, _y_j_s),
                    z=_np3d.full(50, _z_barB_sup),
                    mode="lines",
                    line=dict(color="#f43f5e", width=4, dash="dash"),
                    name=_t(f"As_SUP dir.B ({bar_z})", f"As_TOP dir.B ({bar_z})") if _j == 0 else None,
                    showlegend=_j == 0,
                    hovertemplate=f"SUPERIOR B — {bar_z}<br>As_sup={As_req_B_sup:.2f}cm²<extra></extra>"
                ))

        # ── Barras SUPERIORES dir. L (doble parrilla) — cara comprimida ─────
        if dp_L and n_barras_L_sup > 0:
            _z_barL_sup = _H3 - _rec - 1.5 * _db
            _sep_L_sup3 = (_B3 - 2*_rec) / max(n_barras_L_sup - 1, 1)
            for _i in range(n_barras_L_sup):
                _x_i_s = _rec + _i * _sep_L_sup3
                _y_arr_s = _np3d.linspace(_rec, _L3 - _rec, 50)
                _traces.append(_go3d.Scatter3d(
                    x=_np3d.full(50, _x_i_s),
                    y=_y_arr_s,
                    z=_np3d.full(50, _z_barL_sup),
                    mode="lines",
                    line=dict(color="#e879f9", width=4, dash="dash"),
                    name=_t(f"As_SUP dir.L ({bar_z})", f"As_TOP dir.L ({bar_z})") if _i == 0 else None,
                    showlegend=_i == 0,
                    hovertemplate=f"SUPERIOR L — {bar_z}<br>As_sup={As_req_L_sup:.2f}cm²<extra></extra>"
                ))


        _fig3d = _go3d.Figure(data=_traces)
        _fig3d.update_layout(
            scene=dict(
                xaxis=dict(title=f"B = {_B3:.2f} m", backgroundcolor="#161b22",
                           gridcolor="#334155", zerolinecolor="#475569"),
                yaxis=dict(title=f"L = {_L3:.2f} m", backgroundcolor="#161b22",
                           gridcolor="#334155", zerolinecolor="#475569"),
                zaxis=dict(title=f"H = {Hzap:.0f} cm", backgroundcolor="#161b22",
                           gridcolor="#334155", zerolinecolor="#475569"),
                bgcolor="#161b22",
                aspectmode="data",
                camera=dict(eye=dict(x=1.6, y=1.2, z=1.0))
            ),
            paper_bgcolor="#0f1117",
            plot_bgcolor="#0f1117",
            font=dict(color="white", size=11),
            legend=dict(
                bgcolor="#1e293b", bordercolor="#334155", borderwidth=1,
                font=dict(color="white", size=10)
            ),
            margin=dict(l=0, r=0, t=40, b=0),
            title=dict(
                text=_t(
                    f"Parrilla de Armado 3D — {_nB} barras dir.B + {_nL} barras dir.L ({bar_z})"
                    + (f" | DOBLE PARRILLA: As_SUP_B={As_req_B_sup:.1f}cm² As_SUP_L={As_req_L_sup:.1f}cm²" if _doble_parrilla else ""),
                    f"3D Reinforcement Grid — {_nB} bars dir.B + {_nL} bars dir.L ({bar_z})"
                    + (f" | DOUBLE REBAR: As_TOP_B={As_req_B_sup:.1f}cm² As_TOP_L={As_req_L_sup:.1f}cm²" if _doble_parrilla else "")
                ),
                font=dict(color="white", size=13), x=0.5
            )
        )
        st.plotly_chart(_fig3d, use_container_width=True)
        st.caption(_t(
            f"Amarillo = dir. B inf ({_nB}×{bar_z} @ {sep_B:.1f}cm)  |  "
            f"Celeste = dir. L inf ({_nL}×{bar_z} @ {sep_L:.1f}cm)  |  "
            f"Recubrimiento = {recub_z:.1f} cm  |  d_prom = {d_avg:.1f} cm"
            + (f"  |  ROSA = As_SUP B ({n_barras_B_sup}×{bar_z})  MAGENTA = As_SUP L ({n_barras_L_sup}×{bar_z})" if _doble_parrilla else ""),
            f"Yellow = dir. B inf ({_nB}×{bar_z} @ {sep_B:.1f}cm)  |  "
            f"Cyan = dir. L inf ({_nL}×{bar_z} @ {sep_L:.1f}cm)  |  "
            f"Cover = {recub_z:.1f} cm  |  d_avg = {d_avg:.1f} cm"
            + (f"  |  PINK = As_TOP B ({n_barras_B_sup}×{bar_z})  MAGENTA = As_TOP L ({n_barras_L_sup}×{bar_z})" if _doble_parrilla else "")
        ))

    # ── P3-C: Chequeo de Grupo de Zapatas (Traslape Zonas Influencia) ────────
    with st.expander(
        _t("Chequeo Grupo de Zapatas (NSR-10 H.3 — Traslape Bulbos)",
           "Footing Group Check (NSR-10 H.3 — Pressure Bulb Overlap)"),
        expanded=False
    ):
        st.markdown(_t(
            "Verifica que las zonas de influencia de presiones (bulbos) entre zapatas "
            "adyacentes no se traslapen. Criterio NSR-10 H.3: separación libre s ≥ B "
            "(zapata cuadrada) o s ≥ B/2 (continua). Visualización 2D del grupo.",
            "Checks that pressure bulb zones between adjacent footings do not overlap. "
            "NSR-10 H.3 criterion: clear spacing s ≥ B (isolated) or s ≥ B/2 (strip). "
            "2D plan view of the group."
        ))

        _grp_n = st.slider(_t("Número de zapatas en grupo","Number of footings in group"),
                           2, 8, 4, key="grp_n")

        st.markdown(_t("**Posición y dimensión de cada zapata (centros en planta):**",
                       "**Position and dimension of each footing (centers in plan):**"))

        _grp_cols = st.columns(4)
        for _gh, _lbl in zip(_grp_cols, ["Zapata", "X centro (m)", "Y centro (m)", "B (m)"]):
            _gh.caption(_lbl)

        _grp_data = []
        _defaults_x = [0.0, 4.0, 8.0, 12.0, 0.0, 4.0, 8.0, 12.0]
        _defaults_y = [0.0, 0.0, 0.0,  0.0,  5.0, 5.0, 5.0,  5.0]
        for _ig in range(_grp_n):
            _gc = st.columns(4)
            _gc[0].caption(f"Z{_ig+1}")
            _gx = _gc[1].number_input(f"X{_ig+1}", -50.0, 200.0,
                _defaults_x[_ig] if _ig < 8 else float(_ig*4),
                0.5, key=f"grp_x_{_ig}", label_visibility="collapsed")
            _gy = _gc[2].number_input(f"Y{_ig+1}", -50.0, 200.0,
                _defaults_y[_ig] if _ig < 8 else 0.0,
                0.5, key=f"grp_y_{_ig}", label_visibility="collapsed")
            _gb = _gc[3].number_input(f"B{_ig+1} (m)", 0.5, 10.0,
                float(B_cp), 0.1, key=f"grp_b_{_ig}", label_visibility="collapsed")
            _grp_data.append({"id": f"Z{_ig+1}", "x": _gx, "y": _gy, "B": _gb})

        _sep_crit = st.radio(
            _t("Criterio separación mínima","Min spacing criterion"),
            [_t("s ≥ B (aisladas NSR-10 H.3)","s ≥ B (isolated NSR-10 H.3)"),
             _t("s ≥ B/2 (continua/medianera)","s ≥ B/2 (strip/edge)")],
            key="grp_crit", horizontal=True
        )
        _factor_sep = 1.0 if "B/2" not in _sep_crit else 0.5

        # ── Calcular pares y traslapes ────────────────────────────────────
        import plotly.graph_objects as _go_grp
        import math as _math_grp
        _ncr_grp = []
        _fig_grp  = _go_grp.Figure()

        # Dibujar cada zapata (rectángulo + zona de influencia)
        for _ig, _z in enumerate(_grp_data):
            _cx, _cy, _bz = _z["x"], _z["y"], _z["B"]
            # Zapata
            _x_z = [_cx-_bz/2, _cx+_bz/2, _cx+_bz/2, _cx-_bz/2, _cx-_bz/2]
            _y_z = [_cy-_bz/2, _cy-_bz/2, _cy+_bz/2, _cy+_bz/2, _cy-_bz/2]
            _fig_grp.add_trace(_go_grp.Scatter(
                x=_x_z, y=_y_z, fill="toself",
                fillcolor="rgba(56,189,248,0.3)",
                line=dict(color="#38bdf8", width=2),
                name=_z["id"], showlegend=True,
                hovertemplate=f"{_z['id']}: B={_bz:.2f}m @ ({_cx:.1f},{_cy:.1f})<extra></extra>"
            ))
            # Etiqueta
            _fig_grp.add_annotation(
                x=_cx, y=_cy, text=_z["id"],
                showarrow=False, font=dict(color="white", size=11, family="Arial Black"))
            # Zona de influencia (círculo = B de influencia)
            _theta_arr = [i * 3.14159 / 18 for i in range(37)]
            _r_inf = _bz * _factor_sep + _bz / 2
            _fig_grp.add_trace(_go_grp.Scatter(
                x=[_cx + _r_inf * _math_grp.cos(t) for t in _theta_arr],
                y=[_cy + _r_inf * _math_grp.sin(t) for t in _theta_arr],
                mode="lines",
                line=dict(color="#f59e0b", width=1, dash="dot"),
                showlegend=False, hoverinfo="skip"
            ))

        # Verificar cada par
        _pares_rows = []
        for _i in range(len(_grp_data)):
            for _j in range(_i+1, len(_grp_data)):
                _zi, _zj = _grp_data[_i], _grp_data[_j]
                _dist = _math_grp.sqrt((_zi["x"]-_zj["x"])**2 + (_zi["y"]-_zj["y"])**2)
                _s_libre = _dist - (_zi["B"] + _zj["B"]) / 2
                _s_min   = max(_zi["B"], _zj["B"]) * _factor_sep
                _ok_par  = _s_libre >= _s_min
                _pares_rows.append({
                    "Par": f"{_zi['id']} — {_zj['id']}",
                    "Dist. entre centros (m)": round(_dist, 3),
                    "s libre (m)": round(_s_libre, 3),
                    "s mínima (m)": round(_s_min, 3),
                    "Margen (m)": round(_s_libre - _s_min, 3),
                    "Estado": "OK" if _ok_par else "TRASLAPE",
                })
                if not _ok_par:
                    # Línea roja de traslape
                    _fig_grp.add_trace(_go_grp.Scatter(
                        x=[_zi["x"], _zj["x"]], y=[_zi["y"], _zj["y"]],
                        mode="lines",
                        line=dict(color="#ef4444", width=3),
                        showlegend=False,
                        hovertemplate=(
                            f"{_zi['id']}↔{_zj['id']}: "
                            f"s_libre={_s_libre:.2f}m < s_min={_s_min:.2f}m<extra></extra>"
                        )
                    ))
                else:
                    _fig_grp.add_trace(_go_grp.Scatter(
                        x=[_zi["x"], _zj["x"]], y=[_zi["y"], _zj["y"]],
                        mode="lines",
                        line=dict(color="#22c55e", width=1.5, dash="dot"),
                        showlegend=False, hoverinfo="skip"
                    ))

        _fig_grp.update_layout(
            xaxis=dict(title="X (m)", gridcolor="#334155", color="white",
                       scaleanchor="y", scaleratio=1),
            yaxis=dict(title="Y (m)", gridcolor="#334155", color="white"),
            paper_bgcolor="#0f1117", plot_bgcolor="#161b22",
            font=dict(color="white", size=11),
            legend=dict(bgcolor="#1e293b", bordercolor="#334155",
                        font=dict(color="white", size=10)),
            margin=dict(l=40, r=20, t=50, b=40),
            title=dict(
                text=_t("Planta Grupo de Zapatas — Bulbos de Presión",
                        "Footing Group Plan — Pressure Bulbs"),
                font=dict(color="white", size=12), x=0.5
            )
        )
        st.plotly_chart(_fig_grp, use_container_width=True)

        # Tabla de pares
        _df_pares = pd.DataFrame(_pares_rows)
        def _color_par(val):
            if val == "OK":
                return "background-color:#d4edda;color:#155724;font-weight:bold"
            elif val == "TRASLAPE":
                return "background-color:#f8d7da;color:#721c24;font-weight:bold"
            return ""
        st.dataframe(_df_pares.style.map(_color_par, subset=["Estado"]),
                     use_container_width=True, hide_index=True)

        _n_traslapes = sum(1 for r in _pares_rows if r["Estado"] == "TRASLAPE")
        _gp1, _gp2, _gp3 = st.columns(3)
        _gp1.metric(_t("Pares verificados","Pairs checked"), len(_pares_rows))
        _gp2.metric(_t("Traslapes detectados","Overlaps detected"), _n_traslapes,
                    delta_color="inverse" if _n_traslapes > 0 else "off")
        _gp3.metric(_t("Pares OK","Pairs OK"), len(_pares_rows) - _n_traslapes)

        if _n_traslapes > 0:
            st.error(_t(
                f"{_n_traslapes} par(es) presentan traslape de zonas de influencia. "
                "Aumentar separación entre zapatas o revisar ubicación (NSR-10 H.3).",
                f"{_n_traslapes} pair(s) show pressure bulb overlap. "
                "Increase spacing or revise footing locations (NSR-10 H.3)."
            ))
        else:
            st.success(_t(
                "Sin traslapes. Todas las zapatas cumplen separación mínima (NSR-10 H.3).",
                "No overlaps. All footings meet minimum spacing (NSR-10 H.3)."
            ))

        st.caption(_t(
            f"Criterio: s_libre ≥ {_factor_sep:.1f}·B_max  |  "
            "Líneas rojas = traslape | Líneas verdes = OK | Círculos punteados = zona influencia",
            f"Criterion: s_free ≥ {_factor_sep:.1f}·B_max  |  "
            "Red = overlap | Green = OK | Dotted circles = influence zone"
        ))

    # ── PASO 9C: Alertas NCR (No Conformidades) exportables ────────────────
    with st.expander(
        _t("Alertas de No Conformidad (NCR) — Auditoría",
           "Non-Conformance Alerts (NCR) — Audit"),
        expanded=False
    ):
        st.markdown(_t(
            "Tabla automática de todas las verificaciones con estado FALLA. "
            "Exportable en CSV para actas de interventoría, auditoría o revisión interna.",
            "Automatic table of all FAIL verifications. "
            "Exportable as CSV for site inspection records, audits or internal review."
        ))
        import datetime as _dt_ncr

        _ncr_rows = []
        def _ncr(codigo, desc, demanda, capacidad, unidad, cumple, norma):
            if not cumple:
                _ncr_rows.append({
                    "Código NCR": f"NCR-{len(_ncr_rows)+1:03d}",
                    "Fecha": _dt_ncr.date.today().strftime("%Y-%m-%d"),
                    "Norma": norma,
                    "Verificación": desc,
                    "Demanda": f"{demanda:.2f} {unidad}",
                    "Capacidad": f"{capacidad:.2f} {unidad}",
                    "Déficit": f"{demanda - capacidad:.2f} {unidad}",
                    "Severidad": "ALTA" if (demanda - capacidad)/max(capacidad,0.001) > 0.20
                                 else "MEDIA",
                    "Acción sugerida": (
                        "Aumentar B o L" if "Portante" in desc else
                        "Aumentar peralte H" if "Cortante" in desc else
                        "Aumentar armado" if "Flexión" in desc else
                        "Aumentar H o usar gancho" if "Anclaje" in desc else
                        "Revisar geometría"
                    ),
                    "Estado": "ABIERTA"
                })

        # Geotécnica
        _ncr("GEO-01", "Capacidad portante",
             q_act, q_adm, "kPa", cumplio, f"NSR-10 H.3 / {metodo_cap}")
        _ncr("GEO-02", "Factor de seguridad FS",
             FS_terz, FS_calc if FS_calc < 999 else FS_terz + 1, "", FS_calc >= FS_terz,
             "NSR-10 H.3")

        # Estabilidad
        if H_horiz > 0:
            _ncr("EST-01", "Estabilidad volcamiento", FS_terz, FS_volc, "",
                 ok_volc, "NSR-10 H.4.1")
            _ncr("EST-02", "Estabilidad deslizamiento", FS_terz, FS_desl, "",
                 ok_desl, "NSR-10 H.4.2")

        # Estructural
        _ncr("EST-03", "Cortante 1D dir. B (C.11.3)",
             Vu_1way_B, phi_Vc_1way_B, "kN", ok_1way_B, "NSR-10 C.11.3 / ACI 318")
        _ncr("EST-04", "Cortante 1D dir. L (C.11.3)",
             Vu_1way_L, phi_Vc_1way_L, "kN", ok_1way_L, "NSR-10 C.11.3 / ACI 318")
        _ncr("EST-05", "Punzonamiento bo (C.11.12)",
             Vu_punz, phi_Vc_P, "kN", ok_punz, "NSR-10 C.11.12 / ACI 318")
        _ncr("EST-06", "Flexión dir. B (C.10.2)",
             As_req_B, disc_B * A_bar_z if disc_B > 0 else 0.0, "cm²",
             disc_B > 0, "NSR-10 C.10.2 / ACI 318")
        _ncr("EST-07", "Flexión dir. L (C.10.2)",
             As_req_L, disc_L * A_bar_z if disc_L > 0 else 0.0, "cm²",
             disc_L > 0, "NSR-10 C.10.2 / ACI 318")
        _ncr("ANC-01", "Anclaje gancho 90° (C.12.5)",
             ldh_req_cm, L_disp_min_cm, "cm", ok_ldh, "NSR-10 C.12.5")
        _ncr("ANC-02", "Desarrollo recto ld (C.12.2)",
             _ld_req_cm, L_disp_min_cm, "cm", _ok_ld, "NSR-10 C.12.2")

        # Cuantía
        _rho_min_zap = max(0.0018, 0.25 * (fc_basico**0.5) / fy_basico)
        _rho_B = As_req_B / (B_cp * 100 * d_avg) if (B_cp * d_avg) > 0 else 0
        _rho_L = As_req_L / (L_cp * 100 * d_avg) if (L_cp * d_avg) > 0 else 0
        _ncr("CUA-01", "Cuantía mínima dir. B (C.7.12)",
             _rho_min_zap * 1000, _rho_B * 1000, "‰",
             _rho_B >= _rho_min_zap, "NSR-10 C.7.12")
        _ncr("CUA-02", "Cuantía mínima dir. L (C.7.12)",
             _rho_min_zap * 1000, _rho_L * 1000, "‰",
             _rho_L >= _rho_min_zap, "NSR-10 C.7.12")

        # Separación máxima barras
        _sep_max_zap = min(45.0, 3 * Hzap)  # cm NSR-10 C.7.12.2
        _ncr("SEP-01", "Separación máxima barras dir. B (C.7.12.2)",
             sep_B, _sep_max_zap, "cm", sep_B <= _sep_max_zap, "NSR-10 C.7.12.2")
        _ncr("SEP-02", "Separación máxima barras dir. L (C.7.12.2)",
             sep_L, _sep_max_zap, "cm", sep_L <= _sep_max_zap, "NSR-10 C.7.12.2")

        # Asentamiento (si se calculó Steinbrenner)
        _Se_val = float(st.session_state.get("_se_total_session", 0.0))
        if _Se_val > 0:
            _ncr("GEO-03", "Asentamiento elástico (Steinbrenner)",
                 _Se_val, 25.0, "mm", _Se_val <= 25.0, "NSR-10 H.3")

        # Mostrar resultado
        if _ncr_rows:
            _df_ncr = pd.DataFrame(_ncr_rows)
            # Color por severidad
            def _color_ncr_sev(val):
                if val == "ALTA":
                    return "background-color:#f8d7da;color:#721c24;font-weight:bold"
                elif val == "MEDIA":
                    return "background-color:#fff3cd;color:#856404;font-weight:bold"
                return ""
            st.dataframe(
                _df_ncr.style.map(_color_ncr_sev, subset=["Severidad"]),
                use_container_width=True, hide_index=True
            )
            _n_alta  = (_df_ncr["Severidad"] == "ALTA").sum()
            _n_media = (_df_ncr["Severidad"] == "MEDIA").sum()
            _na1, _na2, _na3 = st.columns(3)
            _na1.metric(_t("NCR Alta severidad","High severity NCR"),
                        _n_alta, delta=_t("Acción inmediata","Immediate action"),
                        delta_color="inverse" if _n_alta > 0 else "off")
            _na2.metric(_t("NCR Media severidad","Medium severity NCR"), _n_media)
            _na3.metric(_t("Total NCR abiertas","Total open NCR"), len(_ncr_rows))

            _csv_ncr = _df_ncr.to_csv(index=False).encode("utf-8-sig")
            st.download_button(
                _t("Descargar NCR CSV (Auditoría)","Download NCR CSV (Audit)"),
                data=_csv_ncr,
                file_name=f"NCR_Zapata_{B_cp:.1f}x{L_cp:.1f}.csv",
                mime="text/csv", key="ncr_csv", use_container_width=True
            )
        else:
            st.success(_t(
                "Sin No Conformidades. Todas las verificaciones CUMPLEN.",
                "No Non-Conformances. All verifications PASS."
            ))

    # ── PASO 8: Diagrama N-M biaxial columna-zapata ──────────────────────────
    with st.expander(
        _t("Diagrama N-M Biaxial Columna-Zapata (NSR-10 C.10)",
           "N-M Biaxial Interaction Diagram Column-Footing (NSR-10 C.10)"),
        expanded=False
    ):
        st.markdown(_t(
            "Verifica que el par (Pu, Mu) de diseño queda DENTRO del diagrama de interacción "
            "de la columna que transmite cargas a la zapata. Método simplificado rectangular "
            "NSR-10 C.10.2 / ACI 318.",
            "Checks that the design pair (Pu, Mu) lies INSIDE the column interaction diagram "
            "that transmits loads to the footing. Simplified rectangular NSR-10 C.10.2 / ACI 318."
        ))

        _nm_c1, _nm_c2, _nm_c3 = st.columns(3)
        with _nm_c1:
            st.markdown("**Sección columna**")
            _nm_b  = st.number_input(_t("b columna (cm)","Column b (cm)"),
                10.0, 150.0, float(c1_col), 5.0, key="nm_b_col")
            _nm_h  = st.number_input(_t("h columna (cm)","Column h (cm)"),
                10.0, 150.0, float(c2_col), 5.0, key="nm_h_col")
            _nm_rec= st.number_input(_t("Recubrimiento (cm)","Cover (cm)"),
                2.0, 8.0, 4.0, 0.5, key="nm_rec_col")
        with _nm_c2:
            st.markdown("**Acero longitudinal columna**")
            _nm_bar_opts = {
                "#3 (9.5mm)":0.71,"#4 (12.7mm)":1.27,"#5 (15.9mm)":1.98,
                "#6 (19.1mm)":2.84,"#7 (22.2mm)":3.87,"#8 (25.4mm)":5.07,
                "#9 (28.6mm)":6.45,"#10 (32mm)":8.19,"#11 (35.8mm)":10.06,
            }
            _nm_bar = st.selectbox(_t("Varilla columna","Column bar"),
                list(_nm_bar_opts.keys()), index=4, key="nm_bar_col")
            _nm_Ab  = _nm_bar_opts[_nm_bar]
            _nm_nb  = st.number_input(_t("N° barras total","Total bars"),
                4, 32, 8, 2, key="nm_nb_col")
            _nm_As_col = _nm_nb * _nm_Ab   # cm²
            _nm_rho    = _nm_As_col / (_nm_b * _nm_h)
        with _nm_c3:
            st.markdown("**Cargas de diseño**")
            _nm_Pu = st.number_input(_t("Pu (kN)","Pu (kN)"),
                0.0, 50000.0, float(P_ult), 10.0, key="nm_Pu")
            _nm_Mu = st.number_input(_t("Mu (kN·m)","Mu (kN·m)"),
                0.0, 10000.0, float(max(Mu_flex_B, Mu_flex_L)), 5.0, key="nm_Mu")
            _nm_fc = st.number_input("f'c columna (MPa)",
                17.0, 60.0, float(fc_basico), 1.0, key="nm_fc_col")
            _nm_fy = st.number_input("fy columna (MPa)",
                240.0, 560.0, float(fy_basico), 10.0, key="nm_fy_col")

        st.divider()

        # ── Cálculo diagrama de interacción simplificado ─────────────────
        import numpy as _np_nm
        import plotly.graph_objects as _go_nm

        _b_cm   = _nm_b
        _h_cm   = _nm_h
        _d_pr   = _h_cm - _nm_rec - 1.6   # d' capa compresión (cm)
        _d_efn  = _h_cm - _nm_rec - 1.6   # d efectivo (cm)
        _A_s    = _nm_As_col / 2           # mitad en cada cara
        _beta1  = max(0.65, 0.85 - 0.05 * max(0.0, (_nm_fc - 28) / 7))
        _e_cu   = 0.003
        _Es     = 200000.0  # MPa
        _phi_c  = 0.65      # NSR-10 C.9.3.2 columnas con estribos
        _phi_b  = 0.90      # flexión pura

        # Puntos característicos del diagrama
        _pts_N = []  # kN
        _pts_M = []  # kN·m

        # 1. Compresión pura (Punto A) — NSR-10 ec. C.10-1
        _P0 = 0.85 * _nm_fc * (_b_cm * _h_cm - _nm_As_col) + _nm_fy * _nm_As_col  # kN (cm²→kN: /10)
        _P0_kN = _P0 / 10.0
        _Pmax  = 0.80 * _phi_c * _P0_kN   # Pmax NSR-10 C.10.3.6
        _pts_N.append(_Pmax);  _pts_M.append(0.0)

        # 2. Barrido c desde h (compresión) a 0 (tensión)
        _c_vals = _np_nm.linspace(_h_cm * 0.95, 0.01, 120)
        for _cv in _c_vals:
            _a  = _beta1 * _cv
            _a  = min(_a, _h_cm)
            # Deformación acero compresión (cara superior)
            _eps_s2 = _e_cu * (_cv - _d_pr) / _cv
            _fs2    = max(-_nm_fy, min(_nm_fy, _eps_s2 * _Es))
            # Deformación acero tensión (cara inferior)
            _eps_s1 = _e_cu * (_d_efn - _cv) / _cv
            _fs1    = max(-_nm_fy, min(_nm_fy, _eps_s1 * _Es))
            # Fuerzas
            _Cc = 0.85 * _nm_fc * _a * _b_cm / 10.0   # kN
            _Cs = _A_s * _fs2 / 10.0                   # kN
            _Ts = _A_s * _fs1 / 10.0                   # kN
            _P  = _phi_c * (_Cc + _Cs - _Ts)           # kN
            _M  = _phi_c * (
                _Cc * (_h_cm/2 - _a/2) +
                _Cs * (_h_cm/2 - _d_pr) +
                _Ts * (_d_efn - _h_cm/2)
            ) / 100.0  # kN·m
            _pts_N.append(_P)
            _pts_M.append(abs(_M))

        # 3. Tensión pura (Punto D)
        _Pt = -_phi_b * _nm_fy * _nm_As_col / 10.0  # kN (negativo)
        _pts_N.append(_Pt); _pts_M.append(0.0)

        # Cerrar diagrama (simetría lado tensión → M negativo espejado)
        _pts_N_full = _pts_N + list(reversed(_pts_N))
        _pts_M_full = _pts_M + [-m for m in reversed(_pts_M)]

        # ── Punto de balanceo (deform. simultánea εcu + εy) ──────────────
        _c_bal   = _e_cu / (_e_cu + _nm_fy / _Es) * _d_efn
        _a_bal   = _beta1 * _c_bal
        _Cc_bal  = 0.85 * _nm_fc * _a_bal * _b_cm / 10.0
        _Ts_bal  = _A_s * _nm_fy / 10.0
        _Cs_bal  = _A_s * max(-_nm_fy, min(_nm_fy,
                    _e_cu * (_c_bal - _d_pr) / _c_bal * _Es)) / 10.0
        _Pb      = _phi_c * (_Cc_bal + _Cs_bal - _Ts_bal)
        _Mb      = _phi_c * (
            _Cc_bal * (_h_cm/2 - _a_bal/2) +
            _Cs_bal * (_h_cm/2 - _d_pr) +
            _Ts_bal * (_d_efn - _h_cm/2)
        ) / 100.0

        # ── Plotly ───────────────────────────────────────────────────────
        _inside = False
        # Verificar si (Pu, Mu) está dentro del diagrama
        # Simplificado: buscar N en la curva y verificar M_cap > Mu
        try:
            _Pu_kN = _nm_Pu
            _Mu_kNm= _nm_Mu
            _idx_close = int(_np_nm.argmin(_np_nm.abs(_np_nm.array(_pts_N) - _Pu_kN)))
            _M_cap_pu  = _pts_M[_idx_close]
            _inside    = _Mu_kNm <= _M_cap_pu and _Pu_kN >= _Pt
        except Exception:
            _inside = False

        _fig_nm = _go_nm.Figure()

        # Envolvente de interacción
        _fig_nm.add_trace(_go_nm.Scatter(
            x=_pts_M_full, y=_pts_N_full,
            mode="lines", fill="toself",
            fillcolor="rgba(30,80,120,0.25)",
            line=dict(color="#38bdf8", width=2),
            name=_t("Diagrama N-M","N-M Diagram")
        ))

        # Punto de balanceo
        _fig_nm.add_trace(_go_nm.Scatter(
            x=[_Mb, -_Mb], y=[_Pb, _Pb],
            mode="markers+text",
            marker=dict(color="#f59e0b", size=10, symbol="diamond"),
            text=["Pb","Pb"],
            textposition="top center",
            textfont=dict(color="#f59e0b", size=10),
            name=_t("Punto balanceo","Balance point")
        ))

        # Pmax
        _fig_nm.add_trace(_go_nm.Scatter(
            x=[0], y=[_Pmax],
            mode="markers+text",
            marker=dict(color="#a78bfa", size=10, symbol="square"),
            text=["Pmax"],
            textposition="top right",
            textfont=dict(color="#a78bfa", size=10),
            name=_t("Pmax (0.80φPo)","Pmax")
        ))

        # Punto de diseño (Pu, Mu)
        _color_pt = "#22c55e" if _inside else "#ef4444"
        _label_pt = _t("CUMPLE","OK") if _inside else _t("NO CUMPLE","FAIL")
        _fig_nm.add_trace(_go_nm.Scatter(
            x=[_nm_Mu, -_nm_Mu], y=[_nm_Pu, _nm_Pu],
            mode="markers+text",
            marker=dict(color=_color_pt, size=14, symbol="circle",
                        line=dict(color="white", width=2)),
            text=[f"({_nm_Mu:.0f},{_nm_Pu:.0f})", ""],
            textposition="top right",
            textfont=dict(color=_color_pt, size=11, family="Arial Black"),
            name=f"(Pu,Mu) — {_label_pt}"
        ))

        _fig_nm.update_layout(
            xaxis=dict(title=_t("M (kN·m)","M (kN·m)"),
                       gridcolor="#334155", zerolinecolor="#64748b", color="white"),
            yaxis=dict(title=_t("N (kN)","N (kN)"),
                       gridcolor="#334155", zerolinecolor="#64748b", color="white"),
            paper_bgcolor="#0f1117", plot_bgcolor="#161b22",
            font=dict(color="white", size=11),
            legend=dict(bgcolor="#1e293b", bordercolor="#334155",
                        borderwidth=1, font=dict(color="white", size=10)),
            margin=dict(l=50, r=20, t=50, b=40),
            title=dict(
                text=_t(
                    f"Diagrama N-M — Col {_nm_b:.0f}×{_nm_h:.0f}cm | "
                    f"{_nm_nb}×{_nm_bar} | f'c={_nm_fc:.0f} MPa | fy={_nm_fy:.0f} MPa",
                    f"N-M Diagram — Col {_nm_b:.0f}×{_nm_h:.0f}cm | "
                    f"{_nm_nb}×{_nm_bar} | f'c={_nm_fc:.0f} MPa | fy={_nm_fy:.0f} MPa"
                ),
                font=dict(color="white", size=12), x=0.5
            ),
            shapes=[dict(
                type="line", x0=0, x1=0,
                y0=min(_pts_N_full)*1.05, y1=max(_pts_N_full)*1.05,
                line=dict(color="#475569", width=1, dash="dot")
            )]
        )
        st.plotly_chart(_fig_nm, use_container_width=True)

        # Métricas
        _mc1, _mc2, _mc3, _mc4 = st.columns(4)
        _mc1.metric("Pmax (kN)", f"{_Pmax:.1f}")
        _mc2.metric(_t("Pb balanceo (kN)","Pb balance (kN)"), f"{_Pb:.1f}")
        _mc3.metric(_t("M_cap @ Pu (kN·m)","M_cap @ Pu"), f"{_M_cap_pu:.1f}")
        _mc4.metric(_t("Estado (Pu,Mu)","State (Pu,Mu)"), _label_pt,
                    delta=f"M_dem={_nm_Mu:.1f} vs cap={_M_cap_pu:.1f}",
                    delta_color="normal" if _inside else "inverse")

        # Info ρ
        st.caption(
            f"ρ = As/(b·h) = {_nm_As_col:.2f}/({_nm_b:.0f}×{_nm_h:.0f}) = "
            f"{_nm_rho*100:.2f}%  "
            f"[NSR-10 C.10.9.1: 1% ≤ ρ ≤ 8%]  →  "
            f"{'OK' if 0.01 <= _nm_rho <= 0.08 else 'FUERA DE RANGO'}"
        )

        if not _inside:
            st.error(_t(
                f"El par (Pu={_nm_Pu:.0f} kN, Mu={_nm_Mu:.0f} kN·m) queda FUERA del diagrama. "
                "Aumentar sección o cuantía de la columna.",
                f"The pair (Pu={_nm_Pu:.0f} kN, Mu={_nm_Mu:.0f} kN·m) is OUTSIDE the diagram. "
                "Increase column section or reinforcement ratio."
            ))
        else:
            st.success(_t(
                f"CUMPLE — (Pu={_nm_Pu:.0f} kN, Mu={_nm_Mu:.0f} kN·m) dentro del diagrama N-M.",
                f"OK — (Pu={_nm_Pu:.0f} kN, Mu={_nm_Mu:.0f} kN·m) inside N-M diagram."
            ))

    # (Detector de pilotes removido de aquí, mantenemos la versión completa más abajo)
    # ITEMS 2 & 3: Calculadora ks (coeficiente de balasto) y Rigidez Lambda
    with st.expander("Coeficiente de Balasto ks (Winkler) y Rigidez Relativa"):
        col_ks1, col_ks2 = st.columns(2)
        with col_ks1:
            Es_mpa = st.number_input("Módulo elástico suelo Es [MPa]", 1.0, 200.0, 20.0, 1.0)
            nu_ks  = st.number_input("Relación de Poisson ν", 0.10, 0.49, 0.35, 0.01)
        with col_ks2:
            # Bowles (1996)
            ks = (Es_mpa * 1000.0) / (B_cp * (1.0 - nu_ks**2)) if B_cp > 0 else 0
            st.metric("ks estimado [kN/m³]", f"{ks:,.0f}")
            st.caption("Ref: Bowles (1996) — válido para zapatas rígidas")
        
        # Rigidez rel. (Lambda)
        import math
        Ec_zap = 4700 * math.sqrt(fc_basico) * 1000 if 'fc_basico' in globals() else 21000000  # kPa
        Iz_zap = (B_cp * L_cp**3) / 12.0
        
        if Iz_zap > 0 and Ec_zap > 0:
            lambda_rel = (ks * B_cp * L_cp**4) / (4.0 * Ec_zap * Iz_zap)
            if lambda_rel < 0.1:
                tipo_rig = " **RÍGIDA** — Distribución lineal válida "
            elif lambda_rel < 1.5:
                tipo_rig = " **SEMI-RÍGIDA** — Considerar análisis Winkler detallado"
            else:
                tipo_rig = " **FLEXIBLE** — Distribución real ≠ trapezoidal. Usar losa sobre resortes."
            st.metric("λ (rigidez relativa)", f"{lambda_rel:.4f}")
            st.markdown(tipo_rig)

    #  ASENTAMIENTO ELÁSTICO (opcional) 
    with st.expander("Asentamiento elástico inmediato", expanded=True):
        usar_asent = st.checkbox("Calcular asentamiento inmediato", value=False)
        if usar_asent:
            c_e1, c_e2 = st.columns(2)
            E_suelo = c_e1.number_input("Módulo elasticidad E [MPa]", 1.0, 500.0, 20.0, 1.0)
            nu_suelo = c_e2.number_input("Relación de Poisson ν", 0.1, 0.5, 0.35, 0.01)
            H_estrato = st.number_input("Profundidad del estrato compressible H [m]", 1.0, 50.0, 10.0, 1.0, key="h_estrato")
            
            st.divider()
            st.markdown("####  Resultado Asentamiento Elástico")
            E_kPa = E_suelo * 1000
            m = L_cp / B_cp
            n = H_estrato / B_cp
            
            def I_f_center(m, n):
                return min(1.2, 0.8 + 0.05 * (m - 1))
                
            I_center = I_f_center(m, n)
            asentamiento = q_adm * B_cp * (1 - nu_suelo**2) / E_kPa * I_center * 1000  # en mm
            st.metric("Asentamiento inmediato estimado", f"{asentamiento:.1f} mm")
            st.caption(" Cálculo aproximado usando factor de influencia empírico. Para diseño detallado, usar métodos más refinados (e.g., Steinbrenner completo).")
            # Guardar asentamiento para uso en memoria
            st.session_state['asentamiento'] = asentamiento
        else:
            st.session_state['asentamiento'] = 0.0

        # CRITERIO NSR-10 H.3 — ¿Cuándo se recomienda pilotes? (Detector Inteligente)
        _necesita_pilotes = False
        _razones_pil = []
        if q_adm < 50:
            _razones_pil.append(f"q_adm = {q_adm:.0f} kPa < 50 kPa — suelo muy blando")
        if 'asentamiento' in locals() and asentamiento > 25:
            _razones_pil.append(f"Asentamiento estimado {asentamiento:.0f} mm > 25 mm")
        if N_spt < 5:
            _razones_pil.append(f"N60 = {N_spt} golpes — suelo muy suelto/blando")
        if B_cp > 0 and (Df_cp > B_cp * 3):
            _razones_pil.append(f"Df/B = {Df_cp/B_cp:.1f} > 3 — zapata muy profunda")

        if _razones_pil:
            _necesita_pilotes = True
            st.warning(" **Considera cimentación profunda (pilotes):**\n" + 
                       "\n".join(f"- {r}" for r in _razones_pil))
            st.info("*Este módulo calcula zapatas superficiales. Para pilotes, ver módulo **10_Pilotes** (próximamente).*")

    #  GRÁFICA: TIPO DE FALLA (Vesic 1973 / SPT) 
    st.divider()
    st.markdown("####  Diagrama Tipo de Falla — Vesic (1973)")
    Dr_pct = min(100.0, 1.08 * math.sqrt(max(N_spt, 0) / 60.0) * 100)
    Df_B   = Df_cp / B_cp if B_cp > 0 else 0.0

    fig_ves, ax_ves = plt.subplots(figsize=(8, 5))
    ax_ves.set_facecolor("#0f1117"); fig_ves.patch.set_facecolor("#0f1117")
    Dr_arr = np.linspace(0, 100, 300)
    # Límites aproximados de Vesic (1973) expresados como -Df/B* vs Dr
    c1v = -0.6 + 0.006*Dr_arr             # Punzonamiento → Local
    c2v = -2.5 + 0.030*Dr_arr             # Local → General
    ax_ves.plot(Dr_arr, c1v, color="#4fc3f7", lw=1.8, label="Límite punz./local")
    ax_ves.plot(Dr_arr, c2v, color="#81d4fa", lw=1.8, label="Límite local/general")
    ax_ves.fill_between(Dr_arr, c2v, -5, alpha=0.35, color="#e53935")
    ax_ves.fill_between(Dr_arr, c1v, c2v, alpha=0.30, color="#fb8c00")
    ax_ves.fill_between(Dr_arr, 0, c1v, alpha=0.25, color="#43a047")
    ax_ves.text(10, -0.3, "Falla por\nPunzonamiento", color="white", fontsize=12, fontweight="bold", ha="center")
    ax_ves.text(50, -1.5, "Falla local\npor corte",   color="white", fontsize=12, fontweight="bold", ha="center")
    ax_ves.text(80, -3.5, "Falla general\npor corte", color="white", fontsize=12, fontweight="bold", ha="center")
    # Miniatura fundación
    rec_w = 15; rec_h = 0.3
    ax_ves.add_patch(patches.Rectangle((2, -Df_B - rec_h), rec_w, rec_h,
                                       fc="#888", ec="white", lw=0.8))
    ax_ves.annotate(f"B={B_cp}m\nHz={Hz_cp}m", (2+rec_w/2, -Df_B - rec_h*2),
                    color="yellow", fontsize=7, ha="center")
    # Punto actual
    ax_ves.plot(Dr_pct, -Df_B, "D", color="cyan", ms=12, zorder=6,
                label=f"Dr≈{Dr_pct:.0f}% | Df/B={Df_B:.2f}")
    ax_ves.set_xlim(0, 100); ax_ves.set_ylim(-5, 0)
    ax_ves.set_xlabel("Densidad relativa Dr (%)", color="white")
    ax_ves.set_ylabel("-Df / B*", color="white")
    ax_ves.set_title(f"Tipo de Falla — N60={N_spt} campo → Dr≈{Dr_pct:.0f}%  |  Df/B={Df_B:.2f}", color="white")
    ax_ves.tick_params(colors="white"); ax_ves.spines[:].set_color("#444")
    ax_ves.legend(loc="lower right", fontsize=8, facecolor="#111", labelcolor="white")
    st.pyplot(fig_ves)
    plt.close(fig_ves)

    #  GRÁFICA: BULBO DE PRESIONES / MECANISMO DE FALLA (VECTORIZADO) 
    st.divider()
    st.markdown("####  Mecanismo de Falla y Bulbo de Presiones (Terzaghi)")
    fig_tb, ax_tb = plt.subplots(figsize=(10, 7))
    ax_tb.set_facecolor("#0f1117"); fig_tb.patch.set_facecolor("#0f1117")

    half_B = B_cp / 2.0
    zap_bot = -(Df_cp + Hz_cp)

    # === MECANISMO DE FALLA DE TERZAGHI (Prandtl) ===
    # Ángulos de cuña
    alpha = math.pi/4 + phi_rad/2
    H_tri = half_B * math.tan(alpha)
    pole_x = half_B
    pole_y = zap_bot
    theta_0 = math.atan2(-H_tri, -half_B)
    sweep = math.pi/2
    theta_end = theta_0 + sweep

    # Zona I — Triángulo activo central (Rosa oscuro)
    tri_pts = [(-half_B, zap_bot), (half_B, zap_bot), (0, zap_bot - H_tri)]
    ax_tb.add_patch(plt.Polygon(tri_pts, fc="#e57373", ec="black", lw=1.2, alpha=0.9, zorder=5, label="Zona I (Activa)"))

    if phi_rad >= 0.0:
        # Geometría del espiral (Arco)
        theta_arr = np.linspace(theta_0, theta_end, 40)
        r_arr = (half_B / math.cos(alpha)) * np.exp((theta_arr - theta_0) * math.tan(phi_rad))
        x_spiral = pole_x + r_arr * np.cos(theta_arr)
        y_spiral = pole_y + r_arr * np.sin(theta_arr)
        end_x = x_spiral[-1]
        end_y = y_spiral[-1]

        beta_passive = math.pi/4 - phi_rad/2
        # Prevenir error si beta_passive se hace 0 (para arcilla phi=0, beta_passive=45 deg, siempre >0)
        dist_x = abs(end_y - pole_y) / math.tan(beta_passive)
        x_top = end_x + dist_x

        # Zona II (Radial - Amarillo) - Derecha
        poly_Z2_R = [(pole_x, pole_y)] + list(zip(x_spiral, y_spiral))
        ax_tb.add_patch(plt.Polygon(poly_Z2_R, fc="#ffe082", ec="black", lw=1.2, alpha=0.85, zorder=4, label="Zona II (Radial)"))
        # Zona III (Pasiva - Naranja claro) - Derecha
        poly_Z3_R = [(pole_x, pole_y), (end_x, end_y), (x_top, pole_y)]
        ax_tb.add_patch(plt.Polygon(poly_Z3_R, fc="#ffb74d", ec="black", lw=1.2, alpha=0.85, zorder=3, label="Zona III (Pasiva)"))

        # Zonas II y III - Izquierda (Espejo simétrico)
        poly_Z2_L = [(-x, y) for (x, y) in poly_Z2_R]
        ax_tb.add_patch(plt.Polygon(poly_Z2_L, fc="#ffe082", ec="black", lw=1.2, alpha=0.85, zorder=4))
        poly_Z3_L = [(-x, y) for (x, y) in poly_Z3_R]
        ax_tb.add_patch(plt.Polygon(poly_Z3_L, fc="#ffb74d", ec="black", lw=1.2, alpha=0.85, zorder=3))

        # Sobrecarga de fosa (Suelo arriba de Df - Relleno pasivo)
        if zap_bot < 0:
            ax_tb.add_patch(patches.Rectangle((-x_top, zap_bot), 2*x_top, -zap_bot,
                                              fc="#ffebee", ec="black", lw=1, alpha=0.6, zorder=2, label="Sobrecarga"))
            # Extender limits del eje X para que abarque hasta la cuna pasiva
            ax_tb.set_xlim(-x_top * 1.2, x_top * 1.2)

    # Zapata
    ax_tb.add_patch(patches.Rectangle((-half_B, zap_bot), B_cp, Hz_cp,
                                      fc="#546e7a", ec="white", lw=1.5, zorder=4,
                                      label=f"Zapata {B_cp}×{L_cp}m"))
    # Columna
    ax_tb.add_patch(patches.Rectangle((-b_col_cp/2, 0), b_col_cp, -Df_cp + Hz_cp*0.1,
                                      fc="#78909c", ec="white", lw=1, zorder=4))

    # Bulbo de presiones Boussinesq (Teoría elástica 2D - Suma de Cuadrantes)
    _xg = np.linspace(-B_cp*3.5, B_cp*3.5, 200)
    _zg = np.linspace(zap_bot, zap_bot - B_cp*3.5, 150)
    Xg, Zg = np.meshgrid(_xg, _zg)
    q0_bulbo = q_ult
    dz = np.abs(Zg - zap_bot)
    dz = np.where(dz < 0.05, 0.05, dz)  # evitar división por cero
    
    def I_corner(X, Y, Z_arr):
        m = np.abs(X) / Z_arr
        n = np.abs(Y) / Z_arr
        m = np.clip(m, 1e-6, 1e6)
        n = np.clip(n, 1e-6, 1e6)
        return I_z_bous_vec(m, n) * np.sign(X) * np.sign(Y)

    # Suma algebraica para (x, y=0)
    x1 = half_B - Xg
    x2 = -half_B - Xg
    y1 = L_cp / 2.0
    y2 = -L_cp / 2.0

    I1 = I_corner(x1, y1, dz)
    I2 = I_corner(x2, y1, dz)
    I3 = I_corner(x1, y2, dz)
    I4 = I_corner(x2, y2, dz)
    
    sigma_arr = q0_bulbo * np.abs(I1 - I2 - I3 + I4)

    # Renderizado suave y profesional
    levels = np.linspace(q0_bulbo*0.02, q0_bulbo*0.95, 50)
    cs = ax_tb.contourf(Xg, Zg, sigma_arr, levels=levels, cmap="YlOrBr", alpha=0.9, zorder=1)
    cs_lines = ax_tb.contour(Xg, Zg, sigma_arr, levels=levels[::3], colors='white', linewidths=0.5, zorder=1, alpha=0.6)
    ax_tb.clabel(cs_lines, inline=True, fontsize=9, fmt='%1.0f')
    cbar = fig_tb.colorbar(cs, ax=ax_tb, label="Δσz [kPa]", shrink=0.7)
    cbar.ax.yaxis.set_tick_params(color="white")
    plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color="white")

    # Línea de exploración (si existe en sesión)
    if "z_exploracion" in st.session_state:
        z_ex = st.session_state.z_exploracion
        ax_tb.axhline(-(Df_cp + z_ex), color="cyan", lw=1.8, linestyle="--", label=f"Prof. exploración = {z_ex:.1f}m")

    # Nivel freático
    if NF_prof < Df_cp + B_cp*2.5:
        ax_tb.axhline(-NF_prof, color="#29b6f6", lw=1.5, linestyle=":",
                      label=f" NF = {NF_prof}m")
    ax_tb.axhline(zap_bot, color="#888", lw=0.8, linestyle="--")
    ax_tb.axhline(0, color="#555", lw=0.5)
    ax_tb.set_xlim(-B_cp*4, B_cp*4)
    ax_tb.set_ylim(zap_bot - B_cp*2.8, 0.5)
    ax_tb.set_xlabel("Distancia [m]", color="white")
    ax_tb.set_ylabel("Profundidad [m]", color="white")
    ax_tb.set_title(f"Bulbo de Presiones + Falla Terzaghi — q_ult={q_ult:.1f} kPa | q_adm={q_adm:.1f} kPa",
                    color="white", fontsize=11)
    ax_tb.tick_params(colors="white"); ax_tb.spines[:].set_color("#444")
    ax_tb.legend(loc="upper right", fontsize=8, facecolor="#111", labelcolor="white")
    st.pyplot(fig_tb)
    plt.close(fig_tb)

# 
# T4: PROFUNDIDAD MÍNIMA EXPLORACIÓN
# 
with st.expander(_t("4. Profundidad Mínima de Exploración de Subsuelo (NSR-10)", " 4. Minimum Subsurface Exploration Depth")):
    st.info(_t(
        " **Modo de uso:** Sincronizado automáticamente con las dimensiones (B, L) y carga (Q_actuante) ingresadas en el panel principal. "
        "Según la norma **NSR-10 (Título H.3.2.3)**, las perforaciones deben alcanzar una profundidad donde el incremento de "
        "esfuerzo vertical (Δσ_z) sea menor o igual al **10% del esfuerzo de contacto (q_0)** transmitido por la estructura.",
        " **How to use:** Automatically synced with main B, L and Q inputs. NSR-10 requires exploring "
        "until the stress increment drops below 10% of the contact stress (q_0)."
    ))
    
    q0_ex = Q_act / (B_cp * L_cp)
    st.markdown(f"**Dimensiones Base:** $B={B_cp}$ m, $L={L_cp}$ m | **Carga Transferida:** $Q={Q_act}$ kN | **Esfuerzo de Contacto:** $q_0 = {q0_ex:.2f}$ kPa")
    
    # Biseccion para encontrar Z donde delta_sigma_z / q0_ex = 0.10
    z_low = 0.1
    z_high = 50.0
    for _ in range(30):
        z_mid = (z_low + z_high)/2
        m_ex = (B_cp/2.0) / z_mid
        n_ex = (L_cp/2.0) / z_mid
        rat = 4 * I_z_bous(m_ex, n_ex)
        if rat > 0.10:
            z_low = z_mid
        else:
            z_high = z_mid
            
    st.success(f"La profundidad mínima normativa de exploración (donde Δσ_z = 10% de q0) es: **Z_exploración = {z_mid:.2f} m** bajo el nivel de fundación.")
    
    # Gráfica ilustrativa del descenso del esfuerzo
    fig_ex, ax_ex = plt.subplots(figsize=(8, 5))
    fig_ex.patch.set_facecolor("#0f1117")
    ax_ex.set_facecolor("#161b22")
    
    z_plot = np.linspace(0.1, z_mid * 1.5, 100)
    sig_plot = [4 * q0_ex * I_z_bous((B_cp/2.0)/z, (L_cp/2.0)/z) for z in z_plot]
    
    ax_ex.plot(sig_plot, z_plot, color="#69f0ae", lw=2.5, label="Δσ_z bajo el centro")
    if "z_exploracion" in st.session_state:
        zx = st.session_state.z_exploracion
        sz = 4 * q0_ex * I_z_bous((B_cp/2.0)/zx, (L_cp/2.0)/zx)
        ax_ex.plot(sz, zx, marker="*", color="yellow", markersize=14, zorder=10, label=f"Z={zx:.1f}m ({sz:.1f} kPa)")
        ax_ex.axhline(zx, color="yellow", linestyle=":", alpha=0.5)
    ax_ex.axvline(q0_ex * 0.10, color="#ffdd44", linestyle="-.", lw=1.5, label=f"10% q₀ ({q0_ex*0.1:.1f} kPa)")
    ax_ex.axhline(z_mid, color="#ff9500", linestyle=":", lw=1.5, label=f"Z_exploración = {z_mid:.1f} m")
    ax_ex.fill_betweenx(z_plot, sig_plot, 0, where=(z_plot <= z_mid), color="white", alpha=0.08)
    
    ax_ex.invert_yaxis()
    ax_ex.set_xlabel("Incremento de Esfuerzo Δσ_z [kPa]", color="white", fontsize=10)
    ax_ex.set_ylabel("Profundidad Z [m]", color="white", fontsize=10)
    ax_ex.set_title("Atenuación Analítica del Esfuerzo (Criterio 10% NSR-10)", color="white", fontsize=11)
    ax_ex.set_xlim(0, max(sig_plot) * 1.1)
    ax_ex.tick_params(colors="white", labelsize=9)
    ax_ex.grid(True, linestyle=":", alpha=0.3, color="white")
    ax_ex.legend(loc="lower right", facecolor="#161b22", labelcolor="white", edgecolor="none")
    
    fig_ex.tight_layout()
    st.pyplot(fig_ex)
    plt.close(fig_ex)
    
    # Guardar en sesión para usarlo en gráficos
    st.session_state.z_exploracion = z_mid

# 
# T3: DISEÑO ESTRUCTURAL DE ZAPATA + DIBUJADOR 3000 (con biaxialidad)
# 
with st.expander(_t("3. Diseño Estructural de Zapata Prismática y Dibujador 3000", "3. Footing Structural Design & DXF Drafter"), expanded=True):
    st.markdown(f"**Norma Estructural activa:** `{norma_sel}`")

    #  SELECTOR DE TIPO DE ZAPATA (Opción B) 
    _TIPOS_ZAP = [
        "1. Aislada Céntrica",
        "2. Medianera (Lindero)",
        "3. Esquina (Dos linderos)",
        "4. Con Viga de Amarre / Strap Beam",
    ]
    tipo_zap = st.selectbox(
        _t(" Tipo de Zapata:", " Footing Type:"),
        _TIPOS_ZAP,
        index=st.session_state.get("z_tipo_idx", 0),
        key="z_tipo_zap",
        help=_t("Selecciona el tipo de zapata para activar el flujo de cálculo correspondiente.",
                "Select the footing type to activate the corresponding calculation flow.")
    )
    st.session_state["z_tipo_idx"] = _TIPOS_ZAP.index(tipo_zap)

    # Derivar los flags de modo
    _modo_aislada   = ("1." in tipo_zap)
    _modo_medianera = ("2." in tipo_zap)
    _modo_esquina   = ("3." in tipo_zap)
    _modo_strap     = ("4." in tipo_zap or _modo_medianera or _modo_esquina)

    #  TIPO 2, 3, 4: Redirigir a render_medianera() 
    if not _modo_aislada:
        render_medianera(norma_sel, fc_basico, fy_basico, REBAR_DICT, def_idx, phi_v, phi_f, tipo_zap)
        # Los tipos 2/3/4 tienen su propio flujo completo, no continúan el bloque T3
    else:
        #  TIPO 1: Flujo original de zapata aislada céntrica 
        st.info(_t("**Modo: Zapata Aislada Céntrica** — Selecciona el método de ingreso de cargas.",
                   " **Mode: Isolated Centered Footing** — Select the loads input method."))

        st.write("---")
        st.markdown("### 1. Cargas Actuantes en la Base de la Columna")
        modo_cargas = st.radio("Método de Ingreso de Cargas:", 
                     ["Envolventes Finales (ETABS / SAP2000)", "Generar Combinaciones (Ingreso D, L, E)"], 
                     horizontal=True, key="z_modo_cargas")

        if "Combinaciones" in modo_cargas:
            cD, cL, cE = st.columns(3)
            with cD:
                PD = st.number_input("Carga Muerta P_D [kN]", value=500.0, step=50.0)
                MD_B = st.number_input("Momento M_D (dir B) [kN·m]", value=35.0)
                MD_L = st.number_input("Momento M_D (dir L) [kN·m]", value=29.0)
            with cL:
                PL = st.number_input("Carga Viva P_L [kN]", value=350.0, step=50.0)
                ML_B = st.number_input("Momento M_L (dir B) [kN·m]", value=18.0)
                ML_L = st.number_input("Momento M_L (dir L) [kN·m]", value=16.0)
            with cE:
                PE = st.number_input("Sismo P_E [kN]", value=80.0, step=10.0)
                ME_B = st.number_input("Momento M_E (dir B) [kN·m]", value=55.0)
                ME_L = st.number_input("Momento M_E (dir L) [kN·m]", value=61.0)
                
            P_svc = max(PD + PL, PD + PE, PD + 0.75*PL + 0.75*PE)
            M_svc_B = max(abs(MD_B)+abs(ML_B), abs(MD_B)+abs(ME_B), abs(MD_B)+0.75*abs(ML_B)+0.75*abs(ME_B))
            M_svc_L = max(abs(MD_L)+abs(ML_L), abs(MD_L)+abs(ME_L), abs(MD_L)+0.75*abs(ML_L)+0.75*abs(ME_L))

            P_ult = max(1.2*PD + 1.6*PL, 1.2*PD + 1.0*PL + 1.0*PE)
            M_ult_B = max(1.2*abs(MD_B) + 1.6*abs(ML_B), 1.2*abs(MD_B) + 1.0*abs(ML_B) + 1.0*abs(ME_B))
            M_ult_L = max(1.2*abs(MD_L) + 1.6*abs(ML_L), 1.2*abs(MD_L) + 1.0*abs(ML_L) + 1.0*abs(ME_L))
            
            st.info(f"**Envolventes Calculadas:** Ps={P_svc:.1f} kN | Pu={P_ult:.1f} kN | Ms_B={M_svc_B:.1f} kN·m | Mu_B={M_ult_B:.1f} kN·m", icon="💡")

        else:
            cC1, cC2, _ = st.columns([1.5, 1.5, 1])
            with cC1:
                P_svc = st.number_input("Carga Axial Svc. Ps [kN]", value=800.0, step=50.0)
                M_svc_B = st.number_input("Mom. Svc. Ms (dir. B) [kN·m]", value=0.0, step=10.0)
                M_svc_L = st.number_input("Mom. Svc. Ms (dir. L) [kN·m]", value=0.0, step=10.0)
            with cC2:
                P_ult = st.number_input("Carga Axial Fact. Pu [kN]", value=1120.0, step=50.0)
                M_ult_B = st.number_input("Mom. Fact. Mu (dir. B) [kN·m]", value=0.0, step=10.0)
                M_ult_L = st.number_input("Mom. Fact. Mu (dir. L) [kN·m]", value=0.0, step=10.0)

        st.write("---")
        st.markdown("### 2. Geometría, Suelo y Diseño Estructural")
        colG, colE = st.columns(2)
            
        with colG:
            st.write(_t("#### Geometría y Suelo", "#### Geometry and Soil"))
            q_unit = st.selectbox(_t("Unidad q_adm:", "q_adm Unit:"), ["kPa (kN/m²)", "ton/m²", "kg/cm²", "MPa"], index=0)
            
            default_q = 200.0
            step_q = 10.0
            if q_unit == "ton/m²":
                default_q, step_q = 20.0, 1.0
            elif q_unit == "kg/cm²":
                default_q, step_q = 2.0, 0.1
            elif q_unit == "MPa":
                default_q, step_q = 0.2, 0.05
                
            q_val_input = st.number_input(_t("Capacidad Portante q_adm", "Bearing Capacity q_adm"), value=default_q, step=step_q)
            
            if q_unit == "kPa (kN/m²)":
                q_adm_z = q_val_input
            elif q_unit == "ton/m²":
                q_adm_z = q_val_input * 9.80665
            elif q_unit == "MPa":
                q_adm_z = q_val_input * 1000.0
            else: # kg/cm²
                q_adm_z = q_val_input * 98.0665

            st.info(f"**Equivalencia:** `{q_adm_z:.1f} kPa` | `{q_adm_z/9.80665:.2f} ton/m²` | `{q_adm_z/98.0665:.3f} kg/cm²` | `{q_adm_z/1000:.4f} MPa`")
            
            cc1, cc2 = st.columns(2)
            with cc1:
                c1_col = st.number_input(_t("Columna c1 (dir. B) [cm]", "Col c1 (B dir.) [cm]"), min_value=5.0, value=40.0, step=5.0, key="z_c1col_val")
                pos_col_iso = st.selectbox("Posición", ["Interior", "Borde (eje X)", "Borde (eje Y)", "Esquina"], key="pos_col_iso")
                gamma_prom = st.number_input(_t("γ_promedio [kN/m³]", "γ_avg [kN/m³]"), value=20.0)
            with cc2:
                c2_col = st.number_input(_t("Columna c2 (dir. L) [cm]", "Col c2 (L dir.) [cm]"), min_value=5.0, value=40.0, step=5.0, key="z_c2col_val")
                Df_z = st.number_input(_t("Desplante Df [m]", "Depth Df [m]"), value=1.0, step=0.1)
                s_c_uso = st.number_input("Sobrecarga s/c [kPa]", value=3.0, step=0.5, help="Sobrecarga externa s/c (equivale aprox a 300 kg/m²)")

        with colE:
            st.write("#### Diseño Estructural")
            H_zap = st.number_input("Espesor H propuesto [cm]", value=50.0, step=5.0, key="z_hzap_val")
            # ── Recubrimiento mínimo por norma ─────────────────────────────────────────
            _COVER_MIN_ZAP = {
                'NSR-10 Colombia': 7.5, 'ACI 318-25 EE.UU.': 7.5,
                'ACI 318-19 EE.UU.': 7.5, 'ACI 318-14 EE.UU.': 7.5,
                'E.060 Peru': 7.5, 'NEC-SE-HM Ecuador': 7.5,
                'NTC-EM México': 7.0, 'COVENIN 1753-2006 Venezuela': 7.5,
                'NB 1225001-2020 Bolivia': 7.5, 'CIRSOC 201-2025 Argentina': 7.5,
            }
            _recub_default = _COVER_MIN_ZAP.get(norma_sel, 7.5)
            recub_z = st.number_input("Recubrimiento al suelo [cm]", value=float(_recub_default), step=0.5, help=f"Min. normativo: {_recub_default} cm — {norma_sel}", key="z_recub_val")
            bar_z = st.selectbox("Varilla a utilizar:", list(REBAR_DICT.keys()), index=def_idx, key="z_barz_val")
            A_bar_z = REBAR_DICT[bar_z]["area"] * 100 # mm2
            db_bar_z = REBAR_DICT[bar_z]["db"] # mm
        
    #  Validaciones robustas 
    if c1_col <= 0 or c2_col <= 0:
        st.error("Las dimensiones de la columna (c1, c2) deben ser mayores a 0 cm.")
        st.stop()
    # Verificar espesor mínimo
    d_min = recub_z + db_bar_z/10.0 + 2  # cm
    if H_zap < d_min:
        st.warning(f" El espesor H={H_zap} cm es menor que el mínimo recomendado (recubrimiento + db/10 + 2 cm = {d_min:.1f} cm). Aumente H.")
    # Verificar q_net positivo
    q_net = q_adm_z - gamma_prom * Df_z
    if q_net <= 0:
        st.error(f"El esfuerzo neto disponible q_net = {q_net:.2f} kPa es negativo o nulo. Reduzca Df o aumente q_adm.")
        st.stop()

    # Paso 1: Dimensionamiento en planta (considerando excentricidades)
    # Para momento combinado, la presión máxima se da por flexión biaxial:
    # q_max = P/A ± (Mx * y / Ix) ± (My * x / Iy)
    # Primero obtenemos dimensiones mínimas basadas en q_net (sin momento)
    Area_req = P_svc / q_net
    # Proponemos zapata cuadrada si los momentos son bajos
    L_req = math.sqrt(Area_req * (c2_col/c1_col) if c1_col>0 else Area_req)
    B_req = Area_req / L_req if L_req > 0 else 0
    B_zap = math.ceil(B_req * 20) / 50.0
    L_zap = math.ceil(L_req * 20) / 50.0
    st.markdown(f"**Dimensiones mínimas sin excentricidad:** B = {B_zap:.2f} m, L = {L_zap:.2f} m")
    
    cB, cL = st.columns(2)
    B_use = cB.number_input("B usado para cálculo [m]", value=max(2.0, B_zap), step=0.1)
    L_use = cL.number_input("L usado para cálculo [m]", value=max(2.0, L_zap), step=0.1)

    #  Presión de contacto con flexión biaxial 
    A_use = B_use * L_use
    
    A_req = getattr(st.session_state, 'z_A_req_cache', max(B_zap * L_zap, 0.1)) # fallback safely
    try: A_req = P_svc / q_adm_z if q_adm_z > 0 else 0.1
    except: pass
    if B_use * L_use < A_req:
        st.error(
            f" **ÁREA INSUFICIENTE** — Se requieren {A_req:.2f} m² "
            f"pero se proveen {B_use*L_use:.2f} m². "
            f"Aumenta B o L."
        )
    
    Ix = (B_use * L_use**3) / 12   # momento de inercia respecto al eje X (paralelo a B)
    Iy = (L_use * B_use**3) / 12   # momento de inercia respecto al eje Y (paralelo a L)
    # Coordenadas de las esquinas (x,y) con origen en el centro de la zapata
    corners = [(-B_use/2, -L_use/2), ( B_use/2, -L_use/2),
               ( B_use/2,  L_use/2), (-B_use/2,  L_use/2)]
    q_corners = []
    for x, y in corners:
        q = P_ult/A_use + (M_ult_L * x / Iy) + (M_ult_B * y / Ix)
        q_corners.append(q)
    qu_max = max(q_corners)
    qu_min = min(q_corners)
    # Presión promedio sobre el área crítica (se usa el promedio de las máximas y la presión en la zona de cortante)
    # Para simplificar, usamos el promedio de qu_max y max(qu_min,0)
    qu_avg = (qu_max + max(qu_min, 0)) / 2.0

    #  B4: ÁREA EFECTIVA MEYERHOF (cuando qu_min < 0) 
    # Si hay zona en tensión, la superposición simple sobreestima la presión
    # real. Se aplica el método de área efectiva: B' = B-2eB, L' = L-2eL
    _ecc_B = abs(M_svc_L / P_svc) if P_svc > 0 else 0.0   # excentricidad dir.B [m]
    _ecc_L = abs(M_svc_B / P_svc) if P_svc > 0 else 0.0   # excentricidad dir.L [m]
    _usar_meyerhof = qu_min < 0 and P_ult > 0
    if _usar_meyerhof:
        _B_prima = max(0.1, B_use - 2 * _ecc_B)
        _L_prima = max(0.1, L_use - 2 * _ecc_L)
        _A_prima = _B_prima * _L_prima
        qu_avg = P_ult / _A_prima        # QA-1: qu_avg mutado antes de usarse en punzonamiento y cortante
        qu_max = qu_avg * 1.25           # distribución trapezoidal aproximada
        qu_min = 0.0                     # por definición en el área efectiva
        _meyerhof_info = (f" **Área Efectiva Meyerhof aplicada** — qu_min < 0 → "
                          f"eB={_ecc_B:.2f} m, eL={_ecc_L:.2f} m | "
                          f"B'={_B_prima:.2f} m, L'={_L_prima:.2f} m | "
                          f"A'={_A_prima:.2f} m² | qu_eff={qu_avg:.1f} kPa")
    else:
        _meyerhof_info = None

    # Peralte efectivo
    d_z = H_zap - recub_z - (db_bar_z/10.0)
    d_avg = d_z   # actualizar con valor real calculado
    if d_z <= 0:
        st.error(f"EC-3: Peralte efectivo d_z = {d_z:.2f} cm ≤ 0. Aumente H o reduzca recub. / diámetro.")
        st.stop()
    d_z_m = d_z / 100.0

    #  CORTANTE UNIDIRECCIONAL (Viga) con integración exacta de presión 
    lv_b = (B_use - c1_col/100.0) / 2.0
    lv_l = (L_use - c2_col/100.0) / 2.0

    def q_at(x, y):
        return P_ult/A_use + (M_ult_L * x / Iy) + (M_ult_B * y / Ix)

    # Cortante en dirección B
    x_corte_b = lv_b - d_z_m
    Vu_1way_B = 0.0
    if x_corte_b > 0:
        y_vals = np.linspace(-L_use/2, L_use/2, 50)
        x_vals = np.linspace(x_corte_b, lv_b, 50)
        dx = (lv_b - x_corte_b) / 50
        dy = L_use / 50
        XV, YV = np.meshgrid(x_vals, y_vals)
        Q_xy = q_at(XV, YV)
        Vu_1way_B = np.sum(np.where(Q_xy > 0, Q_xy, 0)) * dx * dy

    phi_Vc_1way_B = phi_v * 0.17 * 1.0 * math.sqrt(fc_basico) * (L_use * 1000) * (d_z * 10) / 1000.0  # kN
    ok_1way_B = phi_Vc_1way_B >= Vu_1way_B

    # Cortante en dirección L
    y_corte_l = lv_l - d_z_m
    Vu_1way_L = 0.0
    if y_corte_l > 0:
        x_vals = np.linspace(-B_use/2, B_use/2, 50)
        y_vals = np.linspace(y_corte_l, lv_l, 50)
        dx = B_use / 50
        dy = (lv_l - y_corte_l) / 50
        XV, YV = np.meshgrid(x_vals, y_vals)
        Q_xy = q_at(XV, YV)
        Vu_1way_L = np.sum(np.where(Q_xy > 0, Q_xy, 0)) * dx * dy

    phi_Vc_1way_L = phi_v * 0.17 * 1.0 * math.sqrt(fc_basico) * (B_use * 1000) * (d_z * 10) / 1000.0  # kN
    ok_1way_L = phi_Vc_1way_L >= Vu_1way_L

    #  PUNZONAMIENTO (bidireccional) con presión promedio 
    bo_1 = c1_col/100.0 + d_z_m
    bo_2 = c2_col/100.0 + d_z_m
    if pos_col_iso == "Interior":
        bo_perim = 2 * (bo_1 + bo_2)
        alpha_s = 40
        Area_punz = bo_1 * bo_2
    elif pos_col_iso == "Borde (eje X)":
        # Columna en el borde X: d_z_m/2 en c2
        _bo2_eff = bo_2 - d_z_m/2.0
        bo_perim = 2*bo_1 + _bo2_eff
        alpha_s = 30
        Area_punz = bo_1 * _bo2_eff
    elif pos_col_iso == "Borde (eje Y)":
        # Columna en el borde Y: d_z_m/2 en c1
        _bo1_eff = bo_1 - d_z_m/2.0
        bo_perim = 2*bo_2 + _bo1_eff
        alpha_s = 30
        Area_punz = _bo1_eff * bo_2
    else: # Esquina
        _bo1_eff = bo_1 - d_z_m/2.0
        _bo2_eff = bo_2 - d_z_m/2.0
        bo_perim = _bo1_eff + _bo2_eff
        alpha_s = 20
        Area_punz = _bo1_eff * _bo2_eff
    # Presión promedio en el área crítica (se usa qu_avg)
    Vu_punz = P_ult - qu_avg * Area_punz

    _min_col = min(c1_col, c2_col)
    beta_c = max(c1_col, c2_col) / _min_col if _min_col > 0 else 1.0
    try:
        Vc1 = 0.33 * math.sqrt(fc_basico)
        Vc2 = 0.17 * (1 + 2/beta_c) * math.sqrt(fc_basico)
        Vc3 = 0.083 * (2 + alpha_s * (d_z*10) / (bo_perim*1000)) * math.sqrt(fc_basico)
        vc_min_MPa = min(Vc1, Vc2, Vc3)
    except ZeroDivisionError:
        vc_min_MPa = 0.0
        st.warning(" División por cero al calcular resistencia a punzonamiento — revise dimensiones.")
    phi_Vc_punz = phi_v * vc_min_MPa * (bo_perim * 1000) * (d_z * 10) / 1000.0  # kN
    ok_punz = phi_Vc_punz >= Vu_punz

    #  FLEXIÓN (momentos con integración exacta) 
    # Momento en dirección B (respecto al eje paralelo a L)
    # Se integra presión * brazo a lo largo del voladizo
    # Para simplificar, integramos numéricamente
    def momento_dir_B():
        # Integrar (q(x,y) * (distancia desde la cara de la columna)) sobre el área del voladizo
        # La distancia desde la cara de la columna es (x + lv_b) con x desde -lv_b a 0 (coordenada local)
        # Usamos coordenadas globales: x desde -B_use/2 hasta -B_use/2 + lv_b ? Mejor usar x desde -lv_b a 0 con origen en cara
        # pero la presión es función de x e y globales.
        # La distancia desde la cara es (x_cara) que es lv_b - (x_global + B_use/2)? Mejor usar integración directa.
        # Usamos un método numérico: sobre el voladizo en dirección B, x_global desde -B_use/2 hasta -B_use/2+lv_b (lado izquierdo)
        # y_global desde -L_use/2 hasta L_use/2. El brazo de momento es la distancia desde la cara de la columna:
        # brazo = (x_global + B_use/2) ? No, la cara de la columna está a -B_use/2 + c1_col/200? Confuso.
        # Para simplificar, usamos la fórmula clásica con presión promedio, pero mejoramos con distribución lineal.
        # Dado que ya tenemos presión variable, podemos usar una integración numérica simple.
        # Tomamos puntos en el voladizo.
        x_min = -B_use/2
        x_cara_col = -c1_col/200.0  # coordenada de la cara de la columna (lado izquierdo)
        # El voladizo izquierdo va desde x_cara_col hasta x_min? No, el voladizo es desde x_cara_col hasta -B_use/2? Revisar.
        # En realidad la columna está centrada, entonces la cara izquierda está en -c1_col/200.
        # El voladizo izquierdo va desde x = -B_use/2 hasta x = -c1_col/200.
        x_left = -B_use/2
        x_right_face = -c1_col/200.0
        if x_right_face <= x_left:
            return 0.0
        n_x = 20
        n_y = 20
        x_vals = np.linspace(x_left, x_right_face, n_x)
        y_vals = np.linspace(-L_use/2, L_use/2, n_y)
        dx = (x_right_face - x_left) / n_x
        dy = L_use / n_y
        Mu_B = 0.0
        for xi in x_vals:
            # brazo desde la cara de la columna
            lever = x_right_face - xi
            for yi in y_vals:
                q_xy = q_at(xi, yi)
                if q_xy > 0:
                    Mu_B += q_xy * lever * dx * dy
        return Mu_B

    Mu_flex_B = momento_dir_B()
    # Similar para dirección L (voladizo en Y)
    def momento_dir_L():
        y_bot = -L_use/2
        y_face = -c2_col/200.0
        if y_face <= y_bot:
            return 0.0
        n_x = 20
        n_y = 20
        x_vals = np.linspace(-B_use/2, B_use/2, n_x)
        y_vals = np.linspace(y_bot, y_face, n_y)
        dx = B_use / n_x
        dy = (y_face - y_bot) / n_y
        Mu_L = 0.0
        for yi in y_vals:
            lever = y_face - yi
            for xi in x_vals:
                q_xy = q_at(xi, yi)
                if q_xy > 0:
                    Mu_L += q_xy * lever * dx * dy
        return Mu_L

    Mu_flex_L = momento_dir_L()

    # Diseño a flexión — Simple o Doble Parrilla (NSR-10 C.10.3.5 / ACI 318-22 Table 21.2.2)
    # ── Cuantía mínima por norma ────────────────────────────────────────────────
    _RHO_MIN_ZAP = {
        'NSR-10 Colombia': 0.0018, 'ACI 318-25 EE.UU.': 0.0018,
        'ACI 318-19 EE.UU.': 0.0018, 'ACI 318-14 EE.UU.': 0.0018,
        'E.060 Peru': 0.0018, 'NEC-SE-HM Ecuador': 0.0018,
        'NTC-EM México': 0.0020, 'COVENIN 1753-2006 Venezuela': 0.0018,
        'NB 1225001-2020 Bolivia': 0.0018, 'CIRSOC 201-2025 Argentina': 0.0018,
    }
    _rho_min_zap = _RHO_MIN_ZAP.get(norma_sel, 0.0018)
    # ── ρ_bal, ρ_max (0.75·ρ_bal) y Mn_max de sección simple ────────────────────
    _beta1_z   = max(0.65, 0.85 - 0.05 * (fc_basico - 28) / 7.0)
    _rho_bal_z = (0.85 * _beta1_z * fc_basico / fy_basico) * (600.0 / (600.0 + fy_basico))
    _rho_max_z = 0.75 * _rho_bal_z   # NSR-10 C.10.3.5 / ACI 318 Table 21.2.2

    def _flex_zapata(Mu_kNm, ancho_m, d_cm):
        """Diseño flexional. Retorna (rho_inf, As_inf, As_sup, disc, dp, Mn_max, aviso)."""
        b_mm = ancho_m * 1000.0;  d_mm = d_cm * 10.0
        d_prime_cm = recub_z + db_bar_z / 20.0        # recub sup + r de barra (cm)
        arm_dp_m   = max((d_cm - d_prime_cm) / 100.0, 0.01)  # brazo As_inf–As_sup (m)
        Mn_max_kNm = (phi_f * _rho_max_z * fy_basico * (d_cm / 100.0) *
                      (1.0 - _rho_max_z * fy_basico / (1.7 * fc_basico))) * ancho_m * 1000.0
        dp = False;  As_sup = 0.0;  aviso = ""
        try:
            Rn   = (Mu_kNm * 1e6) / (phi_f * b_mm * d_mm**2)
            disc = 1.0 - 2.0 * Rn / (0.85 * fc_basico)
            if disc < 0.0:
                dp = True;  disc = 0.0
            rho_req = (0.85 * fc_basico / fy_basico) * (1.0 - math.sqrt(disc))
            if rho_req > _rho_max_z:
                dp = True
        except (ZeroDivisionError, ValueError):
            rho_req = _rho_max_z;  disc = 0.0;  dp = True
        if dp:
            delta_Mu = max(Mu_kNm - Mn_max_kNm, 0.0)
            As_sup   = (delta_Mu * 1e3) / (phi_f * fy_basico * arm_dp_m * 1e4)  # cm²
            rho_inf  = _rho_max_z
            aviso    = (f"DOBLE PARRILLA: Mu={Mu_kNm:.1f} > Mn_max={Mn_max_kNm:.1f} kN·m. "
                        f"As_sup={As_sup:.2f} cm² (cara superior comprimida, d'={d_prime_cm:.1f} cm).")
        else:
            rho_inf = max(rho_req, _rho_min_zap)
        As_inf = rho_inf * (ancho_m * 100.0) * d_cm
        return rho_inf, As_inf, As_sup, disc, dp, Mn_max_kNm, aviso

    rho_use_B, As_req_B_inf, As_req_B_sup, disc_B, dp_B, Mn_max_B, aviso_dp_B = _flex_zapata(Mu_flex_B, L_use, d_z)
    As_req_B      = As_req_B_inf
    n_barras_B    = math.ceil(As_req_B / REBAR_DICT[bar_z]["area"])
    sep_B         = (L_use*100 - 2*recub_z) / max(1, n_barras_B - 1)
    n_barras_B_sup = math.ceil(As_req_B_sup / REBAR_DICT[bar_z]["area"]) if dp_B else 0
    sep_B_sup     = (L_use*100 - 2*recub_z) / max(1, n_barras_B_sup - 1) if n_barras_B_sup > 1 else 0.0

    rho_use_L, As_req_L_inf, As_req_L_sup, disc_L, dp_L, Mn_max_L, aviso_dp_L = _flex_zapata(Mu_flex_L, B_use, d_z)
    As_req_L      = As_req_L_inf
    n_barras_L    = math.ceil(As_req_L / REBAR_DICT[bar_z]["area"])
    sep_L         = (B_use*100 - 2*recub_z) / max(1, n_barras_L - 1)
    n_barras_L_sup = math.ceil(As_req_L_sup / REBAR_DICT[bar_z]["area"]) if dp_L else 0
    sep_L_sup     = (B_use*100 - 2*recub_z) / max(1, n_barras_L_sup - 1) if n_barras_L_sup > 1 else 0.0

    _doble_parrilla = dp_B or dp_L

    # Para compatibilidad con el resto del código
    Mu_flex = Mu_flex_B
    disc_z = disc_B
    As_req_total = As_req_B
    n_barras_Z = n_barras_B
    separacion_S = sep_B
    # Alias para historial, PDF y exportación
    nbarras_B  = n_barras_B
    As_req_B     = As_req_B
    As_req_L     = As_req_L
    disc_B     = disc_B if 'disc_B' in dir() else 0.0
    disc_L     = disc_L if 'disc_L' in dir() else 0.0
    nbarras_L  = n_barras_L
    Mu_B       = Mu_flex_B
    Mu_L       = Mu_flex_L
    Vu_1B      = Vu_1way_B if 'Vu_1way_B' in dir() else 0.0
    Vu_1L      = Vu_1way_L if 'Vu_1way_L' in dir() else 0.0
    phi_Vc_1B  = phi_Vc_1way_B if 'phi_Vc_1way_B' in dir() else 0.0
    phi_Vc_1L  = phi_Vc_1way_L if 'phi_Vc_1way_L' in dir() else 0.0
    separacion_S = sep_B

    # --- CÁLCULO DE GANCHOS Y DOBLECES (NSR-10 / ACI 318) ---
    # Diámetro mínimo de doblez (D_doblez)
    db_mm = db_bar_z
    if db_mm <= 25.4: # Hasta #8 (1")
        D_doblez_mm = 6 * db_mm
    elif db_mm <= 35.8: # #9 a #11
        D_doblez_mm = 8 * db_mm
    else: # #14 a #18
        D_doblez_mm = 10 * db_mm
    D_doblez_cm = D_doblez_mm / 10.0

    # Longitud de extensión del gancho a 90° (L_ext)
    # Norma ACI/NSR-10: max(12*db, 150mm) usualmente para estribos, pero para anclaje en tracción (gancho estándar 90°) es 12*db
    L_ext_gancho_mm = 12 * db_mm
    L_ext_gancho_cm = L_ext_gancho_mm / 10.0
    
    # Radios para dibujo
    radio_doblez_cm = D_doblez_cm / 2.0
    
    # Longitudes de desarrollo disponibles
    ldh_disp_B = (B_use*100 - c1_col)/2 - recub_z
    ldh_disp_L = (L_use*100 - c2_col)/2 - recub_z
    
    # Altura disponible para el gancho
    h_gancho_disp = H_zap - 2*recub_z
    L_gancho_real_cm = min(h_gancho_disp, L_ext_gancho_cm + radio_doblez_cm + db_mm/10.0) # Lo que cabe en la zapata


    tab_res, tab_dwg, tab_apu = st.tabs(["Resultados del Diseño", "Plano 3000 (DXF)", " Cantidades APU"])
    
    with tab_res:
        st.markdown(f"**Revisión Estructural: f'c = {fc_basico} MPa | fy = {fy_basico} MPa**")
        
        #  RECOMENDACIÓN DE ACERO Y ALERTA 

        warning_bar = []
        if sep_B > sep_max or sep_L > sep_max:
            warning_bar.append(f" La separación ({max(sep_B, sep_L):.1f} cm) excede la máxima permitida normativa ({sep_max} cm). ¡Usa una varilla más pequeña (ej. #4 o #5)!")
        if sep_B < 10.0 or sep_L < 10.0:
            warning_bar.append(f" La separación ({min(sep_B, sep_L):.1f} cm) es muy estrecha para fundir bien (< 10 cm). Considera usar una varilla de mayor diámetro.")
        if warning_bar:
            st.error("\n\n".join(warning_bar))
            
        c2d1, c2d2 = st.columns([1, 1])
        with c2d1:
            st.markdown(r"**1. Cortante 1D Dir. B:** $\phi V_c \ge V_u$")
            if ok_1way_B:
                st.success(f"OK Dir B: $\\phi V_c = {phi_Vc_1way_B:.1f}$ $\\ge {Vu_1way_B:.1f}$ kN")
            else:
                st.error(f"FALLA Dir B: $\\phi V_c = {phi_Vc_1way_B:.1f}$ $< {Vu_1way_B:.1f}$ kN")
                
            st.markdown(r"**1b. Cortante 1D Dir. L:** $\phi V_c \ge V_u$")
            if ok_1way_L:
                st.success(f"OK Dir L: $\\phi V_c = {phi_Vc_1way_L:.1f}$ $\\ge {Vu_1way_L:.1f}$ kN")
            else:
                st.error(f"FALLA Dir L: $\\phi V_c = {phi_Vc_1way_L:.1f}$ $< {Vu_1way_L:.1f}$ kN $\\rightarrow$ **Aumentar H**")

        with c2d2:
            st.markdown(r"**2. Punzonamiento:** $\phi V_c \ge V_{up}$")
            if ok_punz:
                st.success(f"OK Punzonamiento: $\\phi V_c = {phi_Vc_punz:.1f}$ $\\ge {Vu_punz:.1f}$ kN")
            else:
                st.error(f"FALLA Punzonamiento: $\\phi V_c = {phi_Vc_punz:.1f}$ $< {Vu_punz:.1f}$ kN $\\rightarrow$ **Aumentar H o Columna**")
        
        st.markdown("---")
        
        estado_tension = " Compresión Total (e ≤ B/6)" if qu_min >= 0 else " Levantamiento"
        
        # Cálculo de longitud de desarrollo l_dh (NSR-10 C.12.5) para gancho 90°
        # l_dh = (0.24 * fy / sqrt(f'c)) * db_mm  [resultado en mm]
        ldh_req_mm = (0.24 * fy_basico / math.sqrt(fc_basico)) * db_bar_z
        ldh_req_cm = max(15.0, 8 * (db_bar_z/10.0), ldh_req_mm / 10.0)
        # Longitud disponible en la zapata desde la cara de la columna (en la dirección más crítica)
        L_disp_B_cm = ((B_use - c1_col/100.0) / 2.0) * 100.0 - recub_z
        L_disp_L_cm = ((L_use - c2_col/100.0) / 2.0) * 100.0 - recub_z
        L_disp_min_cm = min(L_disp_B_cm, L_disp_L_cm)
        ok_ldh = L_disp_min_cm >= ldh_req_cm

        # ── Longitud de desarrollo recta ld (NSR-10 C.12.2.2 / ACI 318-19 Table 25.5.2.1)
        # Condición: barra inferior (tensión), espaciamiento >= 2db y recubrimiento >= db
        # Caso favorable (s >= 2db y recub >= db):
        #   ld = (fy / (3.5 * sqrt(fc))) * db   [mm]  NSR-10 C.12.2.3 caso 1
        # Caso general (sin condición de espacio):
        #   ld = (fy / (1.7 * sqrt(fc))) * db   [mm]  NSR-10 C.12.2.3 caso 2
        # Factor de reducción por recubrimiento y espaciamiento (lambda_ld)
        _cond_esp_B = (sep_B >= 2 * (db_bar_z / 10.0))   # espaciamiento >= 2db (cm)
        _cond_esp_L = (sep_L >= 2 * (db_bar_z / 10.0))
        _cond_rec   = (recub_z >= (db_bar_z / 10.0))     # recubrimiento >= db
        _caso_fav_B = _cond_esp_B and _cond_rec
        _caso_fav_L = _cond_esp_L and _cond_rec
        # Usar el caso más crítico (desfavorable de ambas direcciones)
        _factor_ld = 3.5 if (_caso_fav_B and _caso_fav_L) else 1.7
        # ld mínima normativa: max(300 mm, fórmula)
        _ld_req_mm_B = (fy_basico / (_factor_ld * math.sqrt(fc_basico))) * db_bar_z
        _ld_req_mm_L = _ld_req_mm_B   # misma varilla, misma fórmula
        _ld_req_cm   = max(30.0, _ld_req_mm_B / 10.0)   # min 300 mm = 30 cm (NSR-10 C.12.1)
        # Longitud disponible en la zapata para desarrollo recto
        # (desde cara de columna hasta borde menos recubrimiento)
        _ld_disp_B_cm = L_disp_B_cm   # ya calculada = (B-c1)/2 * 100 - recub
        _ld_disp_L_cm = L_disp_L_cm
        _ld_disp_min_cm = min(_ld_disp_B_cm, _ld_disp_L_cm)
        _ok_ld = _ld_disp_min_cm >= _ld_req_cm
        _ld_caso_txt = _t(
            f"Caso {'favorable' if (_caso_fav_B and _caso_fav_L) else 'general'} (s{'≥' if (_caso_fav_B and _caso_fav_L) else '<'}2db)",
            f"{'Favorable' if (_caso_fav_B and _caso_fav_L) else 'General'} case (s{'≥' if (_caso_fav_B and _caso_fav_L) else '<'}2db)"
        )

        data_res = [
            {"Revisión": "Geometría Propuesta", "Solicitado": f"Area Req = {Area_req:.2f} m²", "Capacidad/Provisto": f"Area Prov. = {A_use:.2f} m²", "Estado": " OK" if A_use>=Area_req else " Subdimensionado"},
            {"Revisión": "Pres. Contacto qu", "Solicitado": f"Max: {qu_max:.2f} kPa", "Capacidad/Provisto": f"Min: {qu_min:.2f} kPa", "Estado": estado_tension},
            {"Revisión": "Flexión Dir. B", "Solicitado": f"Mu_B = {Mu_flex_B:.1f} kN·m | Mn_max_simple = {Mn_max_B:.1f} kN·m", "Capacidad/Provisto": f"As_inf = {As_req_B:.1f} cm² ({n_barras_B} {bar_z} c/{sep_B:.1f}cm)" + (f" | As_SUP = {As_req_B_sup:.2f} cm² ({n_barras_B_sup} {bar_z} c/{sep_B_sup:.1f}cm)" if dp_B else ""), "Estado": "DOBLE PARRILLA" if dp_B else (" OK" if disc_B > 0 else "Rompe en compresion")},
            {"Revisión": "Flexión Dir. L", "Solicitado": f"Mu_L = {Mu_flex_L:.1f} kN·m | Mn_max_simple = {Mn_max_L:.1f} kN·m", "Capacidad/Provisto": f"As_inf = {As_req_L:.1f} cm² ({n_barras_L} {bar_z} c/{sep_L:.1f}cm)" + (f" | As_SUP = {As_req_L_sup:.2f} cm² ({n_barras_L_sup} {bar_z} c/{sep_L_sup:.1f}cm)" if dp_L else ""), "Estado": "DOBLE PARRILLA" if dp_L else (" OK" if disc_L > 0 else "Rompe en compresion")},
            {
                "Revisión": _t("Anclaje Gancho 90° (NSR-10 C.12.5)", "Hook Anchorage 90° (C.12.5)"),
                "Solicitado": f"ldh req: {ldh_req_cm:.1f} cm",
                "Capacidad/Provisto": f"L_disp: {L_disp_min_cm:.1f} cm",
                "Estado": "OK - CUMPLE" if ok_ldh else "NO CUMPLE - Aumentar B o L"
            },
            {
                "Revisión": _t(
                    f"Desarrollo Recto ld (NSR-10 C.12.2) — {_ld_caso_txt}",
                    f"Straight Development ld (ACI 318 25.5.2) — {_ld_caso_txt}"
                ),
                "Solicitado": f"ld req: {_ld_req_cm:.1f} cm  (min 30 cm)",
                "Capacidad/Provisto": f"L_disp: {_ld_disp_min_cm:.1f} cm  (B: {_ld_disp_B_cm:.1f} | L: {_ld_disp_L_cm:.1f})",
                "Estado": "OK - CUMPLE" if _ok_ld else "NO CUMPLE - Aumentar B o L"
            },
        ]
        st.table(pd.DataFrame(data_res))
        # ── Aviso Doble Parrilla ─────────────────────────────────────────────────────
        if _doble_parrilla:
            st.warning(
                _t(
                    "DOBLE PARRILLA requerida: el momento de diseno supera la capacidad ",
                    "de la seccion simplemente reforzada (rho > 0.75*rho_bal). ",
                    "Se agrego acero de compresion As_sup en la cara SUPERIOR de la zapata. ",
                    "Verifique separacion entre capas: NSR-10 C.7.6.1 min max(2.5cm, 1.33*Dmax_agr).",
                    "DOUBLE REINFORCEMENT required: Mu > Mn_max singly-reinforced. ",
                    "Compression steel As_sup added to TOP face. ",
                    "Check layer spacing: ACI 318-22 25.2.1 min max(1in, 1.33*Dmax_agr).",
                )
            )
            if aviso_dp_B: st.info(f"Dir. B: {aviso_dp_B}")
            if aviso_dp_L: st.info(f"Dir. L: {aviso_dp_L}")

        # ── Sección visual detalle de anclaje (Paso 2 — NSR-10 C.12.2 / C.12.5)
        with st.expander(_t(
            "Detalle de Anclaje — Gancho 90° vs Desarrollo Recto (NSR-10 C.12)",
            "Anchorage Detail — 90° Hook vs Straight Development (ACI 318 Ch.25)"
        ), expanded=False):
            _col_anc1, _col_anc2 = st.columns(2)
            with _col_anc1:
                st.markdown(_t("**Gancho Estándar 90° (NSR-10 C.12.5)**", "**Standard 90° Hook (ACI 318 25.4.3)**"))
                st.markdown(f"""
| Parámetro | Valor |
|---|---|
| Diámetro barra | {db_bar_z:.1f} mm ({bar_z}) |
| Diámetro doblez mín. | {D_doblez_cm*10:.1f} mm |
| Extensión recta 12db | {L_ext_gancho_cm:.1f} cm |
| **ldh requerida** | **{ldh_req_cm:.1f} cm** |
| L_disp dir. B | {L_disp_B_cm:.1f} cm |
| L_disp dir. L | {L_disp_L_cm:.1f} cm |
| **Verificación** | **{"OK - CUMPLE" if ok_ldh else "NO CUMPLE"}** |
""")
            with _col_anc2:
                st.markdown(_t("**Desarrollo Recto ld (NSR-10 C.12.2)**", "**Straight Development ld (ACI 318 25.5.2)**"))
                st.markdown(f"""
| Parámetro | Valor |
|---|---|
| Diámetro barra | {db_bar_z:.1f} mm ({bar_z}) |
| fc / fy | {fc_basico:.0f} / {fy_basico:.0f} MPa |
| {_ld_caso_txt} | factor = {_factor_ld:.1f} |
| **ld requerida** | **{_ld_req_cm:.1f} cm** (mín 30 cm) |
| L_disp dir. B | {_ld_disp_B_cm:.1f} cm |
| L_disp dir. L | {_ld_disp_L_cm:.1f} cm |
| **Verificación** | **{"OK - CUMPLE" if _ok_ld else "NO CUMPLE - Aumentar B o L"}** |
""")
            # Recomendación automática
            if not ok_ldh and not _ok_ld:
                st.error(_t(
                    f"AVISO: Ni el gancho 90° ni el desarrollo recto cumplen con la longitud disponible. "
                    f"Aumente B o L en al menos {max(ldh_req_cm, _ld_req_cm) - _ld_disp_min_cm:.1f} cm.",
                    f"WARNING: Neither hook nor straight development meet available length. "
                    f"Increase B or L by at least {max(ldh_req_cm, _ld_req_cm) - _ld_disp_min_cm:.1f} cm."
                ))
            elif not _ok_ld and ok_ldh:
                st.warning(_t(
                    f"AVISO: El desarrollo recto NO cumple ({_ld_req_cm:.1f} cm > {_ld_disp_min_cm:.1f} cm). "
                    f"El gancho 90° SI cumple. Se recomienda usar gancho.",
                    f"WARNING: Straight development does NOT meet ({_ld_req_cm:.1f} cm > {_ld_disp_min_cm:.1f} cm). "
                    f"90° hook PASSES. Recommend using hook."
                ))
            elif _ok_ld and ok_ldh:
                st.success(_t(
                    "CUMPLE: Tanto el gancho 90° como el desarrollo recto son viables para esta geometría.",
                    "PASS: Both 90° hook and straight development are feasible for this geometry."
                ))


        area_ok = A_use >= Area_req
        if not area_ok:
            st.error(
                f" **ÁREA INSUFICIENTE** — Requerida: {Area_req:.2f} m² | "
                f"Provista: {A_use:.2f} m². Aumenta B o L."
            )

        # B4 aviso si se activó el área efectiva Meyerhof
        if _meyerhof_info:
            st.warning(_meyerhof_info)

        # 
        # A1 — FS VOLCAMIENTO Y DESLIZAMIENTO (NSR-10 H.4)
        # 
        st.markdown("---")
        st.markdown("###  Estabilidad Global — Volcamiento y Deslizamiento (NSR-10 H.4)")
        with st.expander("Parámetros de estabilidad", expanded=True):
            _c_fs1, _c_fs2, _c_fs3 = st.columns(3)
            H_horiz = _c_fs1.number_input("Carga horizontal H [kN]", 0.0, 5000.0,
                                          st.session_state.get("z_H_horiz", 0.0), 10.0, key="z_H_horiz")
            delta_ang = _c_fs2.number_input("Ángulo fricción cimentación δ [°] (≈ ²⁄₃φ)",
                                            0.0, 45.0, st.session_state.get("z_delta", phi_ang*2/3), 1.0, key="z_delta")
            # [fix_z_delta] eliminado: st.session_state["z_delta"] = delta_ang  (el widget key='z_delta' ya gestiona session_state)
            Pp_pasivo = _c_fs3.number_input("Empuje pasivo Ep [kN] (si hay muro adyacente)", 0.0, 5000.0, 0.0, 10.0)

        # Carga vertical total de servicio (peso zapata + suelo + carga)
        W_total = P_svc + (B_use * L_use * Df_z * gamma_prom) + (B_use * L_use * s_c_uso)   # kN (peso real zap + relleno + sobrecarga)

        # --- FS VOLCAMIENTO (referencia al borde de la zapata) ---
        # Momento estabilizador = P_svc * B/2  (carga vertical actúa en el centro)
        # Momento volcador     = H_horiz * (Df_z + H_zap/100)
        M_estab = W_total * (B_use / 2.0)                  # kN·m (respecto al borde más comprometido)
        arm_volc = Df_z + H_zap / 100.0                    # m  (brazo desde el punto de giro)
        M_volc   = H_horiz * arm_volc if H_horiz > 0 else 0.0
        FS_volc  = M_estab / M_volc if M_volc > 0 else 999.0
        ok_volc  = FS_volc >= 1.5

        # --- FS DESLIZAMIENTO ---
        delta_rad = math.radians(delta_ang)
        FR        = W_total * math.tan(delta_rad) + Pp_pasivo  # kN (fuerza resistente)
        FS_desl   = FR / H_horiz if H_horiz > 0 else 999.0
        ok_desl   = FS_desl >= 1.5

        _col_v, _col_d = st.columns(2)
        with _col_v:
            st.markdown("**Volcamiento:**")
            fs_v_disp = "∞" if M_volc == 0 else f"{FS_volc:.2f}"
            st.metric("FS Volcamiento", fs_v_disp, delta=f"Req. ≥ 1.50",
                      delta_color="normal" if ok_volc else "inverse")
            if ok_volc:
                st.success(f"FSv = {fs_v_disp} ≥ 1.50  —  M_estab={M_estab:.1f} / M_volc={M_volc:.1f} kN·m")
            else:
                st.error(f"FSv = {fs_v_disp} < 1.50  —  Aumentar B o Df, o reducir H_horiz")
        with _col_d:
            st.markdown("**Deslizamiento:**")
            fs_d_disp = "∞" if H_horiz == 0 else f"{FS_desl:.2f}"
            st.metric("FS Deslizamiento", fs_d_disp, delta=f"Req. ≥ 1.50",
                      delta_color="normal" if ok_desl else "inverse")
            if ok_desl:
                st.success(f"FSd = {fs_d_disp} ≥ 1.50  —  FR={FR:.1f} / H={H_horiz:.1f} kN")
            elif H_horiz == 0:
                st.info("ℹ Ingresa H horizontal > 0 para calcular deslizamiento")
            else:
                st.error(f"FSd = {fs_d_disp} < 1.50  —  Considerar llave de corte o aumentar Df")

        # Actualizar data_res con estabilidad
        st.caption(f"δ = {delta_ang:.1f}° | W_total = {W_total:.1f} kN | b. estabilizador = {B_use/2.0:.2f} m | brazo volcador(H) = {arm_volc:.2f} m")

        # 
        # A2 — MAPA DE CALOR PLOTLY 3D — qu(x,y)
        # 
        st.markdown("---")
        st.markdown("###  Distribución de Presiones bajo la Zapata — qu(x,y)")
        with st.expander("Mapa de calor 3D de esfuerzos de contacto", expanded=True):
            _nx, _ny = 30, 30
            _x_grid = np.linspace(-B_use/2, B_use/2, _nx)
            _y_grid = np.linspace(-L_use/2, L_use/2, _ny)
            _Xm, _Ym = np.meshgrid(_x_grid, _y_grid)
            _Qm = P_ult / A_use + (M_ult_L * _Xm / Iy) + (M_ult_B * _Ym / Ix)

            # Si hay zona negativa → zona de levantamiento (clip a 0 para representación)
            _tiene_tension = bool(np.any(_Qm < 0))
            _Qm_plot = np.where(_Qm < 0, 0, _Qm)   # clip visual

            _fig_heat = go.Figure()

            # Superficie 3D
            _r2 = max(float(np.max(_Qm_plot)) - (float(np.min(_Qm_plot[_Qm_plot > 0])) if np.any(_Qm_plot > 0) else 0), 1.0)
            _cmin2 = max(0.0, float(np.max(_Qm_plot)) - _r2 * 4)
            if _r2 < 5.0:
                _cmin2 = 0.0

            _fig_heat.add_trace(go.Surface(
                x=_x_grid, y=_y_grid, z=_Qm_plot,
                colorscale="RdYlGn_r",
                cmin=_cmin2, cmax=float(np.max(_Qm_plot)) * 1.02,
                colorbar=dict(title=dict(text="qu [kPa]", font=dict(color="white")),
                              tickfont=dict(color="white")),
                opacity=0.9,
                name="qu(x,y)"
            ))

            # Punto qu_max
            rango_real_iso = float(np.max(_Qm_plot)) - float(np.min(_Qm_plot[_Qm_plot > 0])) if np.any(_Qm_plot > 0) else 0
            if rango_real_iso > 1.0:
                _imax = np.unravel_index(np.argmax(_Qm), _Qm.shape)
                _fig_heat.add_trace(go.Scatter3d(
                    x=[_x_grid[_imax[1]]], y=[_y_grid[_imax[0]]], z=[float(_Qm[_imax])],
                    mode="markers+text",
                    marker=dict(color="#ff4444", size=8, symbol="diamond"),
                    text=[f"qu_max={qu_max:.1f} kPa"],
                    textfont=dict(color="white"),
                    name=f"qu_max = {qu_max:.1f} kPa"
                ))

            # Columna proyectada
            _c1m = c1_col / 100.0; _c2m = c2_col / 100.0
            _cx = [-_c1m/2, _c1m/2, _c1m/2, -_c1m/2, -_c1m/2]
            _cy = [-_c2m/2, -_c2m/2, _c2m/2, _c2m/2, -_c2m/2]
            _i_col2 = np.argmin(np.abs(_x_grid - 0))
            _j_col2 = np.argmin(np.abs(_y_grid - 0))
            _z_col = [float(_Qm_plot[_j_col2, _i_col2]) * 1.005] * 5
            _fig_heat.add_trace(go.Scatter3d(
                x=_cx, y=_cy, z=_z_col,
                mode="lines",
                line=dict(color="white", width=4),
                name="Columna"
            ))

            _fig_heat.update_layout(
                scene=dict(
                    xaxis_title="B [m]", yaxis_title="L [m]", zaxis_title="qu [kPa]",
                    bgcolor="#0f1117",
                    xaxis=dict(gridcolor="#333"), yaxis=dict(gridcolor="#333"),
                    zaxis=dict(gridcolor="#333"),
                    camera=dict(eye=dict(x=1.8, y=1.8, z=1.6))
                ),
                paper_bgcolor="#0f1117", font=dict(color="white"),
                margin=dict(l=0, r=0, b=0, t=40), height=500,
                title=dict(text=f"qu(x,y) — qu_max={qu_max:.1f} kPa | qu_min={qu_min:.1f} kPa"
                               + (" |  ZONA EN LEVANTAMIENTO" if _tiene_tension else " |  Compresión total"),
                           font=dict(color="#ffdd44" if _tiene_tension else "#44ff88")),
                showlegend=True,
            )
            st.plotly_chart(_fig_heat, use_container_width=True)
            if abs(qu_max - qu_min) < 1.0:
                st.info(
                    " **Distribución uniforme** — qu = "
                    f"{qu_max:.1f} kPa constante en toda la huella. "
                    "El gráfico 3D cobra valor cuando ingresas momentos Mu ≠ 0 "
                    "para ver la variación de presiones."
                )
            else:
                if not _tiene_tension:
                    st.warning(
                        f" **Distribución trapezoidal** — "
                        f"qu_max = {qu_max:.1f} kPa | qu_min = {qu_min:.1f} kPa. "
                        f"Excentricidad activa en la zapata."
                    )
                else:
                    st.error(
                        f" **Levantamiento parcial** — qu_min = {qu_min:.1f} kPa < 0. "
                        "Parte de la zapata no está en contacto con el suelo. Considere aumentar dimensiones "
                        "o activar el método de Área Efectiva Meyerhof (Bloque B4 disponible abajo)."
                    )

        # 
        # A3 — DISEÑO AUTOMÁTICO ITERATIVO DE H
        # 
        st.markdown("---")
        st.markdown("###  Optimizador de Espesor H (Diseño Automático Iterativo)")
        with st.expander("Encontrar H mínimo que cumple cortante y punzonamiento", expanded=False):
            st.info("El optimizador busca el menor H [cm] (en pasos de 5 cm) que satisface simultáneamente "
                    "Cortante 1D (Dir.B), Cortante 1D (Dir.L) y Punzonamiento, con los materiales y cargas actuales.")
            _H_min_check = st.number_input("H mínimo a considerar [cm]", 20.0, 80.0, 20.0, 5.0, key="z_opt_hmin")
            _H_max_check = st.number_input("H máximo a considerar [cm]", 50.0, 200.0, 120.0, 5.0, key="z_opt_hmax")

            if st.button(" Ejecutar optimizador de H", type="primary"):
                _H_vals   = np.arange(_H_min_check, _H_max_check + 5, 5)
                _res_iter = []
                _H_optimo = None

                for _Hi in _H_vals:
                    _di  = _Hi - recub_z - (db_bar_z / 10.0)          # peralte efectivo cm
                    _dim = _di / 100.0                                   # en metros
                    if _di <= 0:
                        continue

                    # Cortante 1D Dir.B
                    _xc_b = lv_b - _dim
                    _Vu_b = 0.0
                    if _xc_b > 0:
                        _yv = np.linspace(-L_use/2, L_use/2, 50)
                        _xv = np.linspace(_xc_b, lv_b, 50)
                        _dx = (lv_b - _xc_b) / 50; _dy = L_use / 50
                        _XV, _YV = np.meshgrid(_xv, _yv)
                        _Q_xy = q_at(_XV, _YV)
                        _Vu_b = np.sum(np.where(_Q_xy > 0, _Q_xy, 0)) * _dx * _dy
                    _phiVc_b = phi_v * 0.17 * math.sqrt(fc_basico) * (L_use*1000) * (_di*10) / 1000.0

                    # Cortante 1D Dir.L
                    _yc_l = lv_l - _dim
                    _Vu_l = 0.0
                    if _yc_l > 0:
                        _xv2 = np.linspace(-B_use/2, B_use/2, 50)
                        _yv2 = np.linspace(_yc_l, lv_l, 50)
                        _dx2 = B_use / 50; _dy2 = (_yc_l and (lv_l - _yc_l) / 50)
                        _XV2, _YV2 = np.meshgrid(_xv2, _yv2)
                        _Q_xy2 = q_at(_XV2, _YV2)
                        _Vu_l = np.sum(np.where(_Q_xy2 > 0, _Q_xy2, 0)) * _dx2 * _dy2
                    _phiVc_l = phi_v * 0.17 * math.sqrt(fc_basico) * (B_use*1000) * (_di*10) / 1000.0

                    # Punzonamiento
                    _bo1 = c1_col/100.0 + _dim; _bo2 = c2_col/100.0 + _dim
                    if pos_col_iso == "Interior":
                        _bo = 2*(_bo1 + _bo2); _alphas = 40
                        _Ap = _bo1 * _bo2
                    elif pos_col_iso == "Borde (eje X)":
                        _bo2_eff = _bo2 - _dim/2.0
                        _bo = 2*_bo1 + _bo2_eff; _alphas = 30
                        _Ap = _bo1 * _bo2_eff
                    elif pos_col_iso == "Borde (eje Y)":
                        _bo1_eff = _bo1 - _dim/2.0
                        _bo = 2*_bo2 + _bo1_eff; _alphas = 30
                        _Ap = _bo1_eff * _bo2
                    else: # Esquina
                        _bo1_eff = _bo1 - _dim/2.0
                        _bo2_eff = _bo2 - _dim/2.0
                        _bo = _bo1_eff + _bo2_eff; _alphas = 20
                        _Ap = _bo1_eff * _bo2_eff
                    _Vup = P_ult - qu_avg * _Ap
                    _betac = max(c1_col, c2_col) / max(min(c1_col, c2_col), 1e-3)
                    _Vc3p = 0.083*(2 + _alphas*(_di*10)/(_bo*1000))*math.sqrt(fc_basico)
                    _vc_p = min(0.33*math.sqrt(fc_basico), 0.17*(1+2/_betac)*math.sqrt(fc_basico), _Vc3p)
                    _phiVc_p = phi_v * _vc_p * (_bo*1000) * (_di*10) / 1000.0

                    _ok_b = _phiVc_b >= _Vu_b
                    _ok_l = _phiVc_l >= _Vu_l
                    _ok_p = _phiVc_p >= _Vup
                    _cumple = _ok_b and _ok_l and _ok_p

                    _res_iter.append({
                        "H [cm]": int(_Hi), "d [cm]": f"{_di:.1f}",
                        "φVc_1B [kN]": f"{_phiVc_b:.1f}", "Vu_1B [kN]": f"{_Vu_b:.1f}",
                        "φVc_1L [kN]": f"{_phiVc_l:.1f}", "Vu_1L [kN]": f"{_Vu_l:.1f}",
                        "φVc_P [kN]": f"{_phiVc_p:.1f}", "Vu_P [kN]": f"{_Vup:.1f}",
                        "Estado": " CUMPLE" if _cumple else " FALLA"
                    })
                    if _cumple and _H_optimo is None:
                        _H_optimo = _Hi

                st.dataframe(pd.DataFrame(_res_iter), use_container_width=True, hide_index=True)

                if _H_optimo:
                    st.success(f"**H óptimo = {_H_optimo:.0f} cm** — Es el menor espesor que cumple todos los criterios "
                               f"estructurales con los materiales y cargas actuales.")
                    st.info(f"Ahora puede escribir **H = {_H_optimo:.0f} cm** en el input del panel principal y recalcular.")
                else:
                    st.error(f"Ningún H en el rango [{_H_min_check:.0f}–{_H_max_check:.0f}] cm cumple todos los criterios. "
                             f"Considere aumentar B×L o f'c.")

        # Grafico 2D esquemático
        with st.expander("Ver Esquema de Armado 2D", expanded=False):
            fig_2d, ax_2d = plt.subplots(figsize=(5,5))
            fig_2d.patch.set_facecolor("#0f1117"); ax_2d.set_facecolor("#161b22")
            zap_rect = plt.Rectangle((-B_use/2, -L_use/2), B_use, L_use, fill=True, color="#2c3e50", ec="#3498db", lw=2)
            ax_2d.add_patch(zap_rect)
            col_rect = plt.Rectangle((-c1_col/200, -c2_col/200), c1_col/100, c2_col/100, fill=True, color="#7f8c8d", ec="white")
            ax_2d.add_patch(col_rect)
            # Acero dir B (paralelo a L)
            for r_b in np.linspace(-B_use/2 + recub_z/100, B_use/2 - recub_z/100, int(n_barras_B)):
                ax_2d.plot([r_b, r_b], [-L_use/2+0.05, L_use/2-0.05], color="#e74c3c", alpha=0.5, lw=1)
            # Acero dir L (paralelo a B)
            for r_l in np.linspace(-L_use/2 + recub_z/100, L_use/2 - recub_z/100, int(n_barras_L)):
                ax_2d.plot([-B_use/2+0.05, B_use/2-0.05], [r_l, r_l], color="#f1c40f", alpha=0.5, lw=1)
                
            ax_2d.set_xlim(-B_use/2 - 0.5, B_use/2 + 0.5)
            ax_2d.set_ylim(-L_use/2 - 0.5, L_use/2 + 0.5)
            ax_2d.set_aspect('equal')
            ax_2d.axis('off')
            ax_2d.set_title(f"Armadura Inferior\n{int(n_barras_B)} {bar_z} (B) / {int(n_barras_L)} {bar_z} (L)", color="white", fontsize=10)
            st.pyplot(fig_2d)
            plt.close(fig_2d)
        
        #  MEMORIA DE CÁLCULO COMPLETA 
        doc_zap = Document()
        doc_zap.add_heading(f"MEMORIA ESTRUCTURAL — ZAPATA {tipo_zap} ({norma_sel})", 0)

        # 0. Portada — identidad global v6.1
        import datetime
        _emp_z, _proy_z, _ing_z, _elab_z, _rev_z, _apb_z, _logo_z = _get_identity()
        p0 = doc_zap.add_paragraph()
        _ns = st.session_state.get("nivel_sismico", "N/A - General")
        p0.add_run(
            f"Empresa: {_emp_z}  |  Proyecto: {_proy_z}  |  Ing.: {_ing_z}\n"
            f"Elemento: Zapata {tipo_zap} — Norma: {norma_sel}\n"
            f"Nivel Sísmico: {_ns}\n"
            f"Elaboró: {_elab_z}   Revisó: {_rev_z}   Aprobó: {_apb_z}\n"
            f"Materiales: Concreto f'c = {fc_basico:.1f} MPa, Acero fy = {fy_basico:.1f} MPa\n"
            f"Fecha: {__import__('datetime').datetime.now().strftime('%d/%m/%Y')}"
        ).bold = True
        if _logo_z:
            try:
                doc_zap.add_picture(io.BytesIO(_logo_z), width=Inches(1.5))
            except Exception:
                pass

        # 1. Parámetros de Diseño
        doc_zap.add_heading("1. PARÁMETROS DE DISEÑO", level=1)
        doc_zap.add_paragraph(f"Geometría: B = {B_use:.2f} m, L = {L_use:.2f} m, H = {H_zap:.0f} cm\nRecubrimiento: {recub_z} cm\nPedestal: c1 = {c1_col} cm, c2 = {c2_col} cm")
        doc_zap.add_paragraph(f"Cargas de servicio aplicadas:\n  P_serv = {P_serv:.1f} kN\n  M_serv = {M_serv:.1f} kN·m")
        doc_zap.add_paragraph(f"Cargas últimas de diseño:\n  Pu = {P_ult:.1f} kN\n  Mu = {M_ult:.1f} kN·m\n  Vu = {Vu_punz:.1f} kN")

        # 2. Esfuerzos y Presiones
        doc_zap.add_heading("2. DISTRIBUCIÓN DE PRESIONES Y ESTABILIDAD", level=1)
        doc_zap.add_paragraph(f"Capacidad del suelo (qa): {qa:.1f} kPa\nExcentricidad (e): {e_serv:.3f} m (L/6 = {L_use/6:.3f} m)")
        doc_zap.add_paragraph(f"Presiones Máximas (Servicio): q_max = {qmax_s:.1f} kPa, q_min = {qmin_s:.1f} kPa")
        doc_zap.add_paragraph(f"Presiones de Diseño (Mayoradas): q_max_u = {qmax_u:.1f} kPa, q_min_u = {qmin_u:.1f} kPa")
        if area_ok and (ok_1way_B and ok_1way_L):
            doc_zap.add_paragraph("✔ Geometría en planta CUMPLE admitiendo las presiones bajo qa.").bold = True
        else:
            doc_zap.add_paragraph("NO CUMPLE ÁREA INSUFICIENTE para la capacidad del suelo q_a ó Volcamiento.").bold = True

        # 3. Diseño a Cortante (Unidireccional)
        # 2.5 Capacidad portante — método activo con todos los factores
        doc_zap.add_heading("2.5 CAPACIDAD PORTANTE — MÉTODO " + metodo_cap.upper(), level=2)
        _ic_doc  = ic_v if 'ic_v' in dir() else 1.0
        _iq_doc  = iq_v if 'iq_v' in dir() else 1.0
        _ig_doc  = ig_v if 'ig_v' in dir() else 1.0
        _psi_doc = float(st.session_state.get("z_vesic_psi", 0.0))
        doc_zap.add_paragraph(
            f"Referencia: NSR-10 Titulo H.3  |  "
            f"Vesic (1973) JSMFD ASCE 99(1):45-73  |  Bowles (1996) Sec.3-5\n"
            f"Parametros geotecnicos:\n"
            f"  phi = {phi_ang:.1f} deg  |  c = {coh_c:.2f} kPa  |  "
            f"gamma = {gamma_s:.1f} kN/m3  |  gamma_sat = {gamma_sat_t:.1f} kN/m3\n"
            f"  Df = {Df_cp:.2f} m  |  B = {B_cp:.2f} m  |  L = {L_cp:.2f} m  |  "
            f"NF = {NF_prof:.2f} m  |  Caso NF = {caso_nf}\n"
            f"Factores de capacidad ({metodo_cap}):\n"
            f"  Nc = {Nc:.4f}  |  Nq = {Nq:.4f}  |  Ngamma = {Ngamma:.4f}\n"
            f"Factores de forma:\n"
            f"  sc = {sc:.4f}  |  sq = {sq:.4f}  |  sgamma = {sgamma:.4f}\n"
            f"Factores de profundidad:\n"
            f"  dc = {dc:.4f}  |  dq = {dq:.4f}  |  dgamma = {dgamma:.4f}\n"
            f"Factores de inclinacion (psi = {_psi_doc:.1f} deg):\n"
            f"  ic = {_ic_doc:.4f}  |  iq = {_iq_doc:.4f}  |  igamma = {_ig_doc:.4f}\n"
            f"Correccion nivel freatico:\n"
            f"  q_sob = {q_sob:.2f} kPa  |  gamma_eff = {gamma_eff:.2f} kN/m3\n"
            f"Resultado:\n"
            f"  q_ult = sc*dc*ic*c*Nc + sq*dq*iq*q*Nq + sgamma*dgamma*igamma*0.5*gamma*B*Ngamma\n"
            f"  q_ult = {sc*dc*_ic_doc*coh_c*Nc:.2f} + "
            f"{sq*dq*_iq_doc*q_sob*Nq:.2f} + "
            f"{sgamma*dgamma*_ig_doc*0.5*gamma_eff*B_cp*Ngamma:.2f} "
            f"= {q_ult:.2f} kPa\n"
            f"  q_adm = q_ult / FS = {q_ult:.2f} / {FS_terz:.1f} = {q_adm:.2f} kPa\n"
            f"  q_act = Q / (B*L) = {Q_act:.1f} / ({B_cp:.2f}*{L_cp:.2f}) = {q_act:.2f} kPa\n"
            f"  FS calculado = q_ult / q_act = {FS_calc:.2f}  "
            f"{'CUMPLE' if cumplio else 'NO CUMPLE -- Ajustar B o L'}"
        )


        doc_zap.add_heading("3. DISEÑO A CORTANTE UNIDIRECCIONAL", level=1)
        doc_zap.add_paragraph(f"Peralte efectivo promedio (d): {d_avg:.1f} cm")
        doc_zap.add_paragraph(f"Cortante Actuante Mayorado (Vu): {Vu_1way_B:.1f} kN")
        doc_zap.add_paragraph(f"Resistencia del Concreto (φVc): {phi_Vc_1way_B:.1f} kN")
        if (ok_1way_B and ok_1way_L):
            doc_zap.add_paragraph("✔ CUMPLE diseño por Cortante.").bold = True
        else:
            doc_zap.add_paragraph("NO CUMPLE NO CUMPLE Cortante — Aumentar el espesor H de la zapata.").bold = True

        # 4. Diseño a Punzonamiento (Bidireccional)
        doc_zap.add_heading("4. DISEÑO A PUNZONAMIENTO", level=1)
        doc_zap.add_paragraph(f"Perímetro crítico (bo): {bo_perim:.1f} cm")
        doc_zap.add_paragraph(f"Cortante Punzonante Actuante (Vu_p): {Vu_punz:.1f} kN")
        doc_zap.add_paragraph(f"Resistencia del Concreto (φvc_p_total): {phi_Vc_P:.1f} kN")
        if ok_punz:
            doc_zap.add_paragraph("✔ CUMPLE diseño por Punzonamiento.").bold = True
        else:
            doc_zap.add_paragraph("NO CUMPLE NO CUMPLE Punzonamiento — Aumentar H o la dimensión de la columna.").bold = True

        # 5. Diseño a Flexión y Cuantías
        doc_zap.add_heading("5. DISEÑO A FLEXIÓN Y REFUERZO", level=1)
        doc_zap.add_paragraph(
            f"beta1={_beta1_z:.3f} | rho_bal={_rho_bal_z:.5f} | rho_max=0.75*rho_bal={_rho_max_z:.5f}"
        )
        doc_zap.add_heading("5.1 Dirección B (barras paralelas a L)", level=2)
        doc_zap.add_paragraph(
            f"Mu_B={Mu_flex_B:.2f} kN·m  |  Mn_max(simple)={Mn_max_B:.2f} kN·m"
        )
        _d_prime_z_cm = recub_z + db_bar_z / 20.0
        if dp_B:
            doc_zap.add_paragraph(
                f"DOBLE PARRILLA DIR B: Mu_B={Mu_flex_B:.1f} kN·m > Mn_max={Mn_max_B:.1f} kN·m\n"
                f"  As_inf (tension, inferior) = {As_req_B:.2f} cm2  ->  {n_barras_B} {bar_z}  @  {sep_B:.1f} cm\n"
                f"  As_sup (compresion, superior) = {As_req_B_sup:.2f} cm2  ->  {n_barras_B_sup} {bar_z}  @  {sep_B_sup:.1f} cm\n"
                f"  d-prima (recub sup+r) = {_d_prime_z_cm:.1f} cm\n"
                f"  NSR-10 C.10.3.5 / ACI 318-22 Table 21.2.2 / NSR-10 C.7.6.1 sep min capas"
            )
        else:
            doc_zap.add_paragraph(
                f"Parrilla SIMPLE Dir B: As_inf={As_req_B:.2f} cm2  ->  {n_barras_B} {bar_z}  @  {sep_B:.1f} cm"
            )
        doc_zap.add_heading("5.2 Dirección L (barras paralelas a B)", level=2)
        doc_zap.add_paragraph(
            f"Mu_L={Mu_flex_L:.2f} kN·m  |  Mn_max(simple)={Mn_max_L:.2f} kN·m"
        )
        if dp_L:
            doc_zap.add_paragraph(
                f"DOBLE PARRILLA DIR L: Mu_L={Mu_flex_L:.1f} kN·m > Mn_max={Mn_max_L:.1f} kN·m\n"
                f"  As_inf (tension, inferior) = {As_req_L:.2f} cm2  ->  {n_barras_L} {bar_z}  @  {sep_L:.1f} cm\n"
                f"  As_sup (compresion, superior) = {As_req_L_sup:.2f} cm2  ->  {n_barras_L_sup} {bar_z}  @  {sep_L_sup:.1f} cm\n"
                f"  d-prima (recub sup+r) = {_d_prime_z_cm:.1f} cm\n"
                f"  NSR-10 C.10.3.5 / ACI 318-22 Table 21.2.2"
            )
        else:
            doc_zap.add_paragraph(
                f"Parrilla SIMPLE Dir L: As_inf={As_req_L:.2f} cm2  ->  {n_barras_L} {bar_z}  @  {sep_L:.1f} cm"
            )
        doc_zap.add_paragraph(f"Armado propuesto: {bar_z} @ {sep_L:.1f} cm (Dir L), {bar_z} @ {sep_B:.1f} cm (Dir B)")

        # 6. Cuantificación de Materiales
        # ── Sección 5.5: Anclaje (Paso 2) ──────────────────────────────────
        doc_zap.add_heading("5.5. LONGITUD DE ANCLAJE — NSR-10 C.12", level=2)
        doc_zap.add_paragraph(
            f"Barras de refuerzo: {bar_z}  |  db = {db_bar_z:.1f} mm  |  "
            f"fc = {fc_basico:.0f} MPa  |  fy = {fy_basico:.0f} MPa"
        )
        doc_zap.add_paragraph(
            f"GANCHO 90 grados (NSR-10 C.12.5):\n"
            f"  ldh requerida = max(15 cm, 8db, 0.24*fy/sqrt(fc)*db) = {ldh_req_cm:.1f} cm\n"
            f"  L_disponible zapata (dir. critica) = {L_disp_min_cm:.1f} cm\n"
            f"  Verificacion: {'OK - CUMPLE' if ok_ldh else 'NO CUMPLE — Aumentar B o L'}"
        )
        doc_zap.add_paragraph(
            f"DESARROLLO RECTO ld (NSR-10 C.12.2.2 / ACI 318-19 Table 25.5.2.1):\n"
            f"  Condicion de espaciamiento: {'favorable (s>=2db)' if (_caso_fav_B and _caso_fav_L) else 'general (s<2db)'}\n"
            f"  Factor: {_factor_ld:.1f}  |  ld = fy/(factor*sqrt(fc))*db\n"
            f"  ld requerida = max(30 cm, formula) = {_ld_req_cm:.1f} cm\n"
            f"  L_disponible dir. B = {_ld_disp_B_cm:.1f} cm  |  dir. L = {_ld_disp_L_cm:.1f} cm\n"
            f"  Verificacion: {'OK - CUMPLE' if _ok_ld else 'NO CUMPLE — Aumentar B o L'}"
        )
        _anclaje_recom = (
            "Se recomienda gancho estandar 90 grados (NSR-10 C.12.5) para acomodar las barras "
            "dentro de la zapata con la geometria propuesta."
            if ok_ldh and not _ok_ld else (
            "Tanto el gancho 90 grados como el desarrollo recto son viables."
            if ok_ldh and _ok_ld else
            "AVISO: Verificar la longitud de anclaje. Considere aumentar B, L o usar gancho 135 grados."
            )
        )
        doc_zap.add_paragraph(f"Recomendacion: {_anclaje_recom}")
        # ─────────────────────────────────────────────────────────────────────

        doc_zap.add_heading("6. CUANTIFICACIÓN DE MATERIALES", level=1)
        doc_zap.add_paragraph(f"Volumen de concreto: {vol_c:.3f} m³")
        doc_zap.add_paragraph(f"Peso del acero (Dir L): {peso_L:.1f} kg")
        doc_zap.add_paragraph(f"Peso del acero (Dir B): {peso_B:.1f} kg")
        doc_zap.add_paragraph(f"Peso total Acero: {peso_L + peso_B:.1f} kg ({kg_m3:.1f} kg/m³)")

        # 7. Detalles Generales
        doc_zap.add_heading("7. PLANOS Y DETALLES", level=1)
        doc_zap.add_paragraph("El plano en detalle puede ser exportado desde la suite estructural utilizando el estándar DXF-ICONTEC.")
        
        # 8. Referencias Citadas
        doc_zap.add_heading("8. REFERENCIAS CITADAS", level=1)
        doc_zap.add_paragraph("NSR-10 Título C.15: Zapatas")
        doc_zap.add_paragraph("NSR-10 Título C.11: Cortante y Torsión")
        doc_zap.add_paragraph("\n\n_________________________________________\nFirma Ing. Responsable")
        doc_zap.add_paragraph("Matrícula Profesional: _______________")

        # 7.1 Desglose de Costos APU
        doc_zap.add_heading("7.1 Desglose de Costos APU", level=2)
        try:
            import matplotlib.pyplot as _plt_apu_z, io as _io_apu_z
            _apu_z = st.session_state.get("apu_config", {})
            if _apu_z:
                _usar_premix_z = _apu_z.get("premix", False)
                _cost_acero_z = (peso_L + peso_B) * _apu_z.get("acero", 0)
                _cost_mo_z = ((peso_L + peso_B) * 0.04 + vol_c * 0.4) * _apu_z.get("costo_dia_mo", 0)
                _area_enc_z = 2 * (B_use + L_use) * (H_zap / 100.0)  # m2 encofrado lateral
                _cost_enc_z = _area_enc_z * _apu_z.get("encofrado", 0)
                if _usar_premix_z:
                    _cost_conc_z = vol_c * _apu_z.get("precio_premix_m3", 0)
                    _vals_apu_z = [_cost_conc_z, _cost_acero_z, _cost_mo_z, _cost_enc_z]
                    _cats_z = ["Concreto", "Acero", "M.O.", "Encofrado"]
                else:
                    _cost_cem_z = bultos_zap * _apu_z.get("cemento", 0)
                    _cost_agr_z = vol_arena_z * _apu_z.get("arena", 0) + vol_grava_z * _apu_z.get("grava", 0)
                    _vals_apu_z = [_cost_cem_z, _cost_agr_z, _cost_acero_z, _cost_mo_z, _cost_enc_z]
                    _cats_z = ["Cemento", "Agregados", "Acero", "M.O.", "Encofrado"]
                _colores_z = ['#4f46e5', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#ec4899']
                _mon_z = _apu_z.get("moneda", "COP $")
                _fig_az, _ax_az = _plt_apu_z.subplots(figsize=(6, 3))
                _fig_az.patch.set_facecolor('white'); _ax_az.set_facecolor('#f8f9fa')
                _bars_z = _ax_az.bar(_cats_z, _vals_apu_z, color=_colores_z[:len(_cats_z)], edgecolor='white', width=0.55)
                _ax_az.set_title(f"Desglose de Costos APU — Zapata ({_mon_z})", color='#1a1a1a', fontsize=10, fontweight='bold')
                _ax_az.tick_params(colors='#1a1a1a', labelsize=8)
                for sp in _ax_az.spines.values(): sp.set_edgecolor('#cccccc')
                for _b in _bars_z:
                    _ax_az.text(_b.get_x() + _b.get_width()/2., _b.get_height()*1.01,
                                f"{_b.get_height():,.0f}", ha='center', va='bottom', fontsize=7, color='#333333')
                _buf_az = _io_apu_z.BytesIO()
                _fig_az.savefig(_buf_az, format='png', dpi=180, facecolor='white', transparent=False)
                _buf_az.seek(0); _plt_apu_z.close(_fig_az)
                doc_zap.add_picture(_buf_az, width=Inches(5.5))
        except Exception as _e_apu_z:
            doc_zap.add_paragraph(f"(Gráfica APU no disponible: {_e_apu_z})")

        f_zap_io = io.BytesIO()
        doc_zap.save(f_zap_io)
        f_zap_io.seek(0)

        # ── DESCARGA DOCX + PDF NATIVO (ReportLab) ──────────────────────────
        st.markdown("---")
        st.markdown("#### Exportar Memoria de Cálculo")
        _col_exp1, _col_exp2 = st.columns(2)

        with _col_exp1:
            st.download_button(
                _t("Descargar Memoria DOCX", "Download DOCX Report"),
                data=f_zap_io,
                file_name=f"Memoria_Zapata_{tipo_zap[:1]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="za_dl_docx",
                use_container_width=True,
                disabled=not _puede_exportar,
                help=_t("Requiere licencia Pro o Admin.","Requires Pro or Admin license.") if not _puede_exportar else None
            )

        with _col_exp2:
            if st.button(
                _t("Generar PDF Memoria", "Generate PDF Report"),
                key="za_btn_pdf",
                use_container_width=True,
                disabled=not _puede_exportar,
                help=_t("Requiere licencia Pro o Admin.","Requires Pro or Admin license.") if not _puede_exportar else None
            ):
                try:
                    from reportlab.lib.pagesizes import letter, A4
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import cm, mm
                    from reportlab.lib import colors
                    from reportlab.platypus import (
                        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                        HRFlowable, KeepTogether
                    )
                    from reportlab.platypus import Image as RLImage
                    import io as _io_pdf
                    import datetime as _dt_pdf

                    _empresa, _proyecto, _ingeniero, _elab, _rev, _apb, _logo = _get_identity()
                    _pdf_buf = _io_pdf.BytesIO()
                    _doc_pdf = SimpleDocTemplate(
                        _pdf_buf, pagesize=A4,
                        leftMargin=2*cm, rightMargin=2*cm,
                        topMargin=2.5*cm, bottomMargin=2*cm
                    )

                    _styles  = getSampleStyleSheet()
                    _title_s = ParagraphStyle("TitleKonte",
                        parent=_styles["Title"], fontSize=14, leading=18,
                        textColor=colors.HexColor("#1b3a5c"), spaceAfter=4)
                    _h1_s    = ParagraphStyle("H1Konte",
                        parent=_styles["Heading1"], fontSize=11, leading=14,
                        textColor=colors.HexColor("#1b3a5c"),
                        borderPad=2, spaceBefore=10, spaceAfter=4)
                    _h2_s    = ParagraphStyle("H2Konte",
                        parent=_styles["Heading2"], fontSize=10, leading=13,
                        textColor=colors.HexColor("#2d6a4f"), spaceBefore=6, spaceAfter=3)
                    _body_s  = ParagraphStyle("BodyKonte",
                        parent=_styles["Normal"], fontSize=9, leading=13, spaceAfter=3)
                    _caption = ParagraphStyle("Caption",
                        parent=_styles["Normal"], fontSize=8, leading=10,
                        textColor=colors.grey, spaceAfter=2)
                    _ok_s    = ParagraphStyle("OK",
                        parent=_styles["Normal"], fontSize=9,
                        textColor=colors.HexColor("#2d6a4f"))
                    _fail_s  = ParagraphStyle("FAIL",
                        parent=_styles["Normal"], fontSize=9,
                        textColor=colors.HexColor("#c0392b"))

                    _story = []

                    # ── Portada ─────────────────────────────────────────────
                    # Logo si existe
                    if _logo:
                        try:
                            _logo_buf = _io_pdf.BytesIO(_logo)
                            _logo_img = RLImage(_logo_buf, width=3*cm, height=1.5*cm)
                            _story.append(_logo_img)
                            _story.append(Spacer(1, 4*mm))
                        except Exception:
                            pass

                    _story.append(Paragraph(
                        f"MEMORIA ESTRUCTURAL — ZAPATA {tipo_zap.upper()}", _title_s))
                    _story.append(Paragraph(f"Norma: {norma_sel}", _body_s))
                    _story.append(HRFlowable(width="100%", thickness=1,
                                             color=colors.HexColor("#1b3a5c"), spaceAfter=4))
                    _tbl_portada = Table([
                        ["Empresa:",   _empresa,  "Proyecto:", _proyecto],
                        ["Ingeniero:", _ingeniero,"Fecha:",    _dt_pdf.date.today().strftime("%Y-%m-%d")],
                        ["Elaboró:",   _elab,     "Revisó:",   _rev],
                        ["Aprobó:",    _apb,      "Norma:",    norma_sel],
                    ], colWidths=[2.5*cm, 5.5*cm, 2.5*cm, 5.5*cm])
                    _tbl_portada.setStyle(TableStyle([
                        ("FONTNAME",  (0,0),(-1,-1), "Helvetica"),
                        ("FONTSIZE",  (0,0),(-1,-1), 8),
                        ("FONTNAME",  (0,0),(0,-1),  "Helvetica-Bold"),
                        ("FONTNAME",  (2,0),(2,-1),  "Helvetica-Bold"),
                        ("BACKGROUND",(0,0),(-1,-1), colors.HexColor("#f0f4f8")),
                        ("ROWBACKGROUNDS",(0,0),(-1,-1),[colors.HexColor("#f0f4f8"),colors.white]),
                        ("GRID",      (0,0),(-1,-1), 0.3, colors.HexColor("#bbcdd8")),
                        ("TOPPADDING",(0,0),(-1,-1), 3),
                        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                    ]))
                    _story.append(_tbl_portada)
                    _story.append(Spacer(1, 6*mm))

                    # ── 1. Parámetros geotécnicos ───────────────────────────
                    _story.append(Paragraph("1. PARÁMETROS GEOTÉCNICOS", _h1_s))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8"), spaceAfter=3))
                    _geo_data = [
                        ["Parámetro", "Valor", "Parámetro", "Valor"],
                        ["φ (ángulo fricción)", f"{phi_ang:.1f}°",
                         "c (cohesión)", f"{coh_c:.2f} kPa"],
                        ["γ húmedo", f"{gamma_s:.1f} kN/m³",
                         "γ_sat", f"{gamma_sat_t:.1f} kN/m³"],
                        ["Df (desplante)", f"{Df_cp:.2f} m",
                         "NF (nivel freático)", f"{NF_prof:.2f} m"],
                        ["B × L", f"{B_cp:.2f} × {L_cp:.2f} m",
                         "Caso NF", f"Caso {caso_nf}"],
                        ["Método capacidad", metodo_cap,
                         "FS requerido", f"{FS_terz:.1f}"],
                        ["Nc / Nq / Nγ", f"{Nc:.3f} / {Nq:.3f} / {Ngamma:.3f}",
                         "sc/sq/sγ", f"{sc:.3f}/{sq:.3f}/{sgamma:.3f}"],
                        ["dc/dq/dγ", f"{dc:.3f}/{dq:.3f}/{dgamma:.3f}",
                         "ic/iq/iγ", f"{ic_v:.3f}/{iq_v:.3f}/{ig_v:.3f}"],
                        ["q_sob (sobrecarga)", f"{q_sob:.2f} kPa",
                         "γ_eff", f"{gamma_eff:.2f} kN/m³"],
                        ["q_ult", f"{q_ult:.2f} kPa",
                         "q_adm", f"{q_adm:.2f} kPa"],
                        ["q_act", f"{q_act:.2f} kPa",
                         "FS calculado", f"{FS_calc:.2f}"],
                    ]
                    _tbl_geo = Table(_geo_data,
                        colWidths=[4*cm, 3.8*cm, 4*cm, 3.8*cm])
                    _tbl_geo.setStyle(TableStyle([
                        ("FONTNAME",  (0,0),(-1,0),  "Helvetica-Bold"),
                        ("FONTSIZE",  (0,0),(-1,-1), 8),
                        ("BACKGROUND",(0,0),(-1,0),  colors.HexColor("#1b3a5c")),
                        ("TEXTCOLOR", (0,0),(-1,0),  colors.white),
                        ("ROWBACKGROUNDS",(0,1),(-1,-1),
                         [colors.white, colors.HexColor("#f0f4f8")]),
                        ("GRID",      (0,0),(-1,-1), 0.3, colors.HexColor("#bbcdd8")),
                        ("FONTNAME",  (0,1),(0,-1),  "Helvetica-Bold"),
                        ("FONTNAME",  (2,1),(2,-1),  "Helvetica-Bold"),
                        ("TOPPADDING",(0,0),(-1,-1), 3),
                        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                    ]))
                    _story.append(_tbl_geo)

                    # Estado capacidad portante
                    _story.append(Spacer(1, 3*mm))
                    if cumplio:
                        _story.append(Paragraph(
                            f"CUMPLE — q_act = {q_act:.2f} kPa ≤ q_adm = {q_adm:.2f} kPa  (FS = {FS_calc:.2f})",
                            _ok_s))
                    else:
                        _story.append(Paragraph(
                            f"NO CUMPLE — q_act = {q_act:.2f} kPa > q_adm = {q_adm:.2f} kPa. "
                            "Aumentar B o L.", _fail_s))
                    _story.append(Spacer(1, 4*mm))

                    # ── 2. Geometría y cargas ───────────────────────────────
                    _story.append(Paragraph("2. GEOMETRÍA Y CARGAS", _h1_s))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8"), spaceAfter=3))
                    _b_use = float(st.session_state.get("zBuse", B_cp))
                    _l_use = float(st.session_state.get("zLuse", L_cp))
                    _h_zap = float(st.session_state.get("zhzcp", 0.5))
                    _geo2 = [
                        ["Dimensión", "Valor", "Carga", "Valor"],
                        ["B (ancho)", f"{_b_use:.2f} m",
                         "P servicio (Ps)", f"{P_svc:.1f} kN"],
                        ["L (largo)", f"{_l_use:.2f} m",
                         "M dir. B (Svc)", f"{0.0:.1f} kN·m"],
                        ["H (espesor)", f"{_h_zap*100:.0f} cm",
                         "M dir. L (Svc)", f"{0.0:.1f} kN·m"],
                        ["Columna b×c", f"{c1_col:.0f}×{c2_col:.0f} cm",
                         "Pu (mayorado)", f"{P_ult:.1f} kN"],
                    ]
                    _tbl_geo2 = Table(_geo2, colWidths=[4*cm, 3.8*cm, 4*cm, 3.8*cm])
                    _tbl_geo2.setStyle(TableStyle([
                        ("FONTNAME", (0,0),(-1,0), "Helvetica-Bold"),
                        ("FONTSIZE", (0,0),(-1,-1), 8),
                        ("BACKGROUND",(0,0),(-1,0), colors.HexColor("#2d6a4f")),
                        ("TEXTCOLOR",(0,0),(-1,0), colors.white),
                        ("ROWBACKGROUNDS",(0,1),(-1,-1),
                         [colors.white, colors.HexColor("#f0f4f8")]),
                        ("GRID",(0,0),(-1,-1), 0.3, colors.HexColor("#bbcdd8")),
                        ("FONTNAME",(0,1),(0,-1), "Helvetica-Bold"),
                        ("FONTNAME",(2,1),(2,-1), "Helvetica-Bold"),
                        ("TOPPADDING",(0,0),(-1,-1), 3),
                        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                    ]))
                    _story.append(_tbl_geo2)
                    _story.append(Spacer(1, 4*mm))

                    # ── 3. Diseño estructural ───────────────────────────────
                    _story.append(Paragraph("3. DISEÑO ESTRUCTURAL (NSR-10 C.15 / ACI 318)", _h1_s))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8"), spaceAfter=3))
                    _dis_data = [
                        ["Verificación", "Capacidad", "Demanda", "Estado"],
                        ["Cortante 1D dir. B (C.11.3)",
                         f"φVc = {phi_Vc_1way_B:.1f} kN",
                         f"Vu = {Vu_1way_B:.1f} kN",
                         "CUMPLE" if ok_1way_B else "NO CUMPLE"],
                        ["Cortante 1D dir. L (C.11.3)",
                         f"φVc = {phi_Vc_1way_L:.1f} kN",
                         f"Vu = {Vu_1way_L:.1f} kN",
                         "CUMPLE" if ok_1way_L else "NO CUMPLE"],
                        ["Punzonamiento bo (C.11.12)",
                         f"φVc_p = {phi_Vc_P:.1f} kN",
                         f"Vu_p = {Vu_punz:.1f} kN",
                         "CUMPLE" if ok_punz else "NO CUMPLE"],
                        ["Flexión dir. B (C.10.2)",
                         f"As = {As_req_B:.2f} cm²",
                         f"Mu = {Mu_flex_B:.1f} kN·m",
                         "CUMPLE" if disc_B > 0 else "NO CUMPLE"],
                        ["Flexión dir. L (C.10.2)",
                         f"As = {As_req_L:.2f} cm²",
                         f"Mu = {Mu_flex_L:.1f} kN·m",
                         "CUMPLE" if disc_L > 0 else "NO CUMPLE"],
                        ["Anclaje gancho 90° (C.12.5)",
                         f"L_disp = {L_disp_min_cm:.1f} cm",
                         f"ldh = {ldh_req_cm:.1f} cm",
                         "CUMPLE" if ok_ldh else "NO CUMPLE"],
                        ["Desarrollo recto ld (C.12.2)",
                         f"L_disp = {L_disp_min_cm:.1f} cm",
                         f"ld = {_ld_req_cm:.1f} cm",
                         "CUMPLE" if _ok_ld else "NO CUMPLE"],
                    ]
                    _dis_tbl = Table(_dis_data,
                        colWidths=[5.2*cm, 3.8*cm, 3.2*cm, 2.8*cm])
                    _dis_tbl.setStyle(TableStyle([
                        ("FONTNAME",  (0,0),(-1,0),  "Helvetica-Bold"),
                        ("FONTSIZE",  (0,0),(-1,-1), 8),
                        ("BACKGROUND",(0,0),(-1,0),  colors.HexColor("#1b3a5c")),
                        ("TEXTCOLOR", (0,0),(-1,0),  colors.white),
                        ("ROWBACKGROUNDS",(0,1),(-1,-1),
                         [colors.white, colors.HexColor("#f0f4f8")]),
                        ("GRID",(0,0),(-1,-1), 0.3, colors.HexColor("#bbcdd8")),
                        ("TOPPADDING",(0,0),(-1,-1), 3),
                        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                        # Color estado: verde CUMPLE, rojo NO CUMPLE
                        *[("TEXTCOLOR", (3,i+1),(3,i+1),
                           colors.HexColor("#2d6a4f") if _dis_data[i+1][3]=="CUMPLE"
                           else colors.HexColor("#c0392b"))
                          for i in range(len(_dis_data)-1)],
                        *[("FONTNAME", (3,i+1),(3,i+1), "Helvetica-Bold")
                          for i in range(len(_dis_data)-1)],
                    ]))
                    _story.append(_dis_tbl)
                    _story.append(Spacer(1, 4*mm))

                    # ── 4. Armado (cantidades) ──────────────────────────────
                    _story.append(Paragraph("4. ARMADO PROPUESTO", _h1_s))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8"), spaceAfter=3))
                    _arm_data = [
                        ["Elemento", "Resultado"],
                        ["Varilla utilizada", f"{bar_z} — db = {db_bar_z:.1f} mm  Ab = {A_bar_z:.2f} cm²"],
                        ["Barras dir. B (inf)", f"{nbarras_B} × {bar_z}  @  {sep_B:.1f} cm  (As_inf = {As_req_B:.2f} cm²)" + (f"  |  As_SUP: {n_barras_B_sup}×{bar_z} @{sep_B_sup:.0f}cm ({As_req_B_sup:.2f}cm²)" if dp_B else "")],
                        ["Barras dir. L (inf)", f"{nbarras_L} × {bar_z}  @  {sep_L:.1f} cm  (As_inf = {As_req_L:.2f} cm²)" + (f"  |  As_SUP: {n_barras_L_sup}×{bar_z} @{sep_L_sup:.0f}cm ({As_req_L_sup:.2f}cm²)" if dp_L else "")],
                        ["DOBLE PARRILLA", ("Dir B + Dir L" if (dp_B and dp_L) else ("Dir B" if dp_B else ("Dir L" if dp_L else "NO"))) + (f" — rho_max={_rho_max_z:.5f}" if _doble_parrilla else "")] if _doble_parrilla else ["Tipo Armado", "Parrilla Simple (una capa)"],
                        ["Recubrimiento", f"{recub_z:.1f} cm  (d efectivo ≈ {d_avg:.1f} cm)"],
                        ["Peralte H", f"{Hzap:.0f} cm"],
                        ["Anclaje gancho 90°", f"ldh = {ldh_req_cm:.1f} cm  |  L_disp = {L_disp_min_cm:.1f} cm  "
                         f"{'CUMPLE' if ok_ldh else 'NO CUMPLE'}"],
                        ["Desarrollo recto ld", f"ld = {_ld_req_cm:.1f} cm  |  L_disp = {L_disp_min_cm:.1f} cm  "
                         f"{'CUMPLE' if _ok_ld else 'NO CUMPLE'}"],
                        ["Concreto fc", f"{fc_basico:.0f} MPa"],
                        ["Acero fy", f"{fy_basico:.0f} MPa"],
                        ["Volumen concreto (est.)",
                         f"{_b_use * _l_use * _h_zap:.3f} m³"],
                    ]
                    _arm_tbl = Table(_arm_data, colWidths=[5*cm, 11*cm])
                    _arm_tbl.setStyle(TableStyle([
                        ("FONTNAME",  (0,0),(-1,0),  "Helvetica-Bold"),
                        ("FONTSIZE",  (0,0),(-1,-1), 8),
                        ("BACKGROUND",(0,0),(-1,0),  colors.HexColor("#2d6a4f")),
                        ("TEXTCOLOR", (0,0),(-1,0),  colors.white),
                        ("ROWBACKGROUNDS",(0,1),(-1,-1),
                         [colors.white, colors.HexColor("#f0f4f8")]),
                        ("GRID",(0,0),(-1,-1), 0.3, colors.HexColor("#bbcdd8")),
                        ("FONTNAME",  (0,1),(0,-1),  "Helvetica-Bold"),
                        ("TOPPADDING",(0,0),(-1,-1), 3),
                        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                    ]))
                    _story.append(_arm_tbl)
                    _story.append(Spacer(1, 4*mm))

                    # ── 5. Estabilidad ──────────────────────────────────────
                    _story.append(Paragraph("5. ESTABILIDAD GLOBAL (NSR-10 H.4)", _h1_s))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8"), spaceAfter=3))
                    _est_data = [
                        ["Verificación", "FS Calculado", "FS Requerido", "Estado"],
                        ["Volcamiento",
                         f"{FS_volc:.2f}" if M_volc > 0 else "N/A",
                         "1.50",
                         ("CUMPLE" if ok_volc else "NO CUMPLE") if M_volc > 0 else "Sin carga H"],
                        ["Deslizamiento",
                         f"{FS_desl:.2f}" if H_horiz > 0 else "N/A",
                         "1.50",
                         ("CUMPLE" if ok_desl else "NO CUMPLE") if H_horiz > 0 else "Sin carga H"],
                    ]
                    _est_tbl = Table(_est_data, colWidths=[5.5*cm, 3.5*cm, 3.5*cm, 3.5*cm])
                    _est_tbl.setStyle(TableStyle([
                        ("FONTNAME",  (0,0),(-1,0),  "Helvetica-Bold"),
                        ("FONTSIZE",  (0,0),(-1,-1), 8),
                        ("BACKGROUND",(0,0),(-1,0),  colors.HexColor("#1b3a5c")),
                        ("TEXTCOLOR", (0,0),(-1,0),  colors.white),
                        ("ROWBACKGROUNDS",(0,1),(-1,-1),
                         [colors.white, colors.HexColor("#f0f4f8")]),
                        ("GRID",(0,0),(-1,-1), 0.3, colors.HexColor("#bbcdd8")),
                        ("TOPPADDING",(0,0),(-1,-1), 3),
                        ("BOTTOMPADDING",(0,0),(-1,-1), 3),
                    ]))
                    _story.append(_est_tbl)
                    _story.append(Spacer(1, 6*mm))

                    # ── 6. Firma ────────────────────────────────────────────
                    _story.append(Paragraph("6. FIRMA RESPONSABLE", _h1_s))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8"), spaceAfter=6))
                    _story.append(Spacer(1, 8*mm))
                    _firma_tbl = Table([
                        [f"{'_'*30}", "", f"{'_'*30}"],
                        [_ingeniero, "", _apb],
                        ["Diseñó / Calculó", "", "Aprobó"],
                        ["Matrícula Prof.: _______________", "",
                         "Matrícula Prof.: _______________"],
                    ], colWidths=[7*cm, 2*cm, 7*cm])
                    _firma_tbl.setStyle(TableStyle([
                        ("FONTNAME",  (0,0),(-1,-1), "Helvetica"),
                        ("FONTSIZE",  (0,0),(-1,-1), 8),
                        ("FONTNAME",  (0,1),(0,1),   "Helvetica-Bold"),
                        ("FONTNAME",  (2,1),(2,1),   "Helvetica-Bold"),
                        ("ALIGN",     (0,0),(-1,-1), "CENTER"),
                        ("TOPPADDING",(0,0),(-1,-1), 2),
                    ]))
                    _story.append(_firma_tbl)

                    # ── Pie de página (nota al pie del story) ───────────────
                    _story.append(Spacer(1, 4*mm))
                    _story.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.HexColor("#bbcdd8")))
                    _story.append(Paragraph(
                        f"Generado por Konte / StructoPro  |  "
                        f"Norma: {norma_sel}  |  "
                        f"Método: {metodo_cap}  |  "
                        f"Fecha: {__import__('datetime').date.today().strftime('%Y-%m-%d')}  |  "
                        "Documento de uso exclusivo del proyecto indicado.",
                        _caption
                    ))

                    # ── Construir PDF ────────────────────────────────────────
                    _doc_pdf.build(_story)
                    _pdf_buf.seek(0)
                    st.success(_t("PDF generado correctamente.", "PDF generated successfully."))
                    st.download_button(
                        _t("Descargar PDF Memoria", "Download PDF Report"),
                        data=_pdf_buf,
                        file_name=f"Memoria_Zapata_{tipo_zap[:1]}_{norma_sel[:3]}.pdf",
                        mime="application/pdf",
                        key="za_dl_pdf",
                        use_container_width=True
                    )

                except ImportError:
                    st.error(_t(
                        "ReportLab no está instalado. Agrega 'reportlab' a requirements.txt.",
                        "ReportLab is not installed. Add 'reportlab' to requirements.txt."
                    ))
                except Exception as _e_pdf:
                    st.error(f"Error generando PDF: {_e_pdf}")

        # ── P3-B: Ficha Técnica Unificada PDF+DXF+IFC en un clic ─────────────
        st.markdown("---")
        st.markdown(
            _t("#### 📋 Ficha Técnica Unificada (PDF + DXF + IFC)",
               "#### 📋 Unified Technical Sheet (PDF + DXF + IFC)")
        )
        st.info(_t(
            "Genera y descarga en un solo flujo: **Memoria PDF** (ReportLab) + "
            "**Plano DXF** (ezdxf) + **Modelo IFC** (si ifcopenshell disponible). "
            "Formato de entrega completo para interventoría.",
            "Generates and downloads in one flow: **PDF Report** (ReportLab) + "
            "**DXF Drawing** (ezdxf) + **IFC Model** (if ifcopenshell available). "
            "Complete delivery package for site inspection."
        ))
        _ficha_c1, _ficha_c2, _ficha_c3 = st.columns(3)
        _ficha_incl_pdf = _ficha_c1.checkbox(
            _t("Incluir PDF Memoria","Include PDF Report"), True, key="ficha_pdf")
        _ficha_incl_dxf = _ficha_c2.checkbox(
            _t("Incluir DXF Plano","Include DXF Drawing"), True, key="ficha_dxf")
        _ficha_incl_ifc = _ficha_c3.checkbox(
            _t("Incluir IFC Modelo","Include IFC Model"),
            _IFC_AVAILABLE, key="ficha_ifc")

        if st.button(
            _t("Generar Ficha Técnica Unificada","Generate Unified Technical Sheet"),
            key="ficha_btn", use_container_width=True,
            disabled=not _puede_exportar,
            help=_t("Requiere licencia Pro o Admin.",
                    "Requires Pro or Admin license.") if not _puede_exportar else None
        ):
            _ficha_ok = []
            _ficha_err = []

            # ── PDF ──────────────────────────────────────────────────────
            if _ficha_incl_pdf:
                try:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import cm
                    from reportlab.lib import colors
                    from reportlab.platypus import (
                        SimpleDocTemplate, Paragraph, Spacer, Table,
                        TableStyle as _TblSt, HRFlowable
                    )
                    import io as _io_f
                    import datetime as _dt_f
                    _empresa_f, _proy_f, _ing_f, _el_f, _rv_f, _ap_f, _lg_f = _get_identity()
                    _pdf_f = _io_f.BytesIO()
                    _doc_f = SimpleDocTemplate(_pdf_f, pagesize=A4,
                        leftMargin=2*cm, rightMargin=2*cm,
                        topMargin=2.5*cm, bottomMargin=2*cm)
                    _st_f  = getSampleStyleSheet()
                    _h1_f  = ParagraphStyle("H1F", parent=_st_f["Heading1"],
                        fontSize=11, textColor=colors.HexColor("#1b3a5c"),
                        spaceBefore=8, spaceAfter=4)
                    _bd_f  = ParagraphStyle("BF", parent=_st_f["Normal"], fontSize=9)
                    _stry  = []
                    _stry.append(Paragraph(
                        f"FICHA TÉCNICA — ZAPATA {tipo_zap.upper()} ({norma_sel})",
                        ParagraphStyle("TF", parent=_st_f["Title"], fontSize=14,
                                       textColor=colors.HexColor("#1b3a5c"))))
                    _stry.append(Spacer(1, 3*cm))
                    _stry.append(HRFlowable(width="100%", thickness=1,
                                            color=colors.HexColor("#1b3a5c")))
                    _stry.append(Spacer(1, 0.3*cm))
                    _tbl_f = Table([
                        ["Empresa:", _empresa_f, "Proyecto:", _proy_f],
                        ["Ingeniero:", _ing_f, "Fecha:", _dt_f.date.today().strftime("%Y-%m-%d")],
                        ["Dimensiones:", f"B={B_cp:.2f}m × L={L_cp:.2f}m × H={Hzap:.0f}cm",
                         "Método:", metodo_cap],
                        ["q_adm:", f"{q_adm:.1f} kPa", "q_act:", f"{q_act:.1f} kPa"],
                        ["q_ult:", f"{q_ult:.1f} kPa", "FS calc.:",
                         f"{FS_calc:.2f}" if FS_calc < 999 else "∞"],
                        ["Armado B:", f"{nbarras_B}×{bar_z}@{sep_B:.0f}cm",
                         "Armado L:", f"{nbarras_L}×{bar_z}@{sep_L:.0f}cm"],
                        ["Estado global:", "CUMPLE" if (cumplio and ok_1way_B and
                            ok_1way_L and ok_punz) else "NO CUMPLE", "", ""],
                    ], colWidths=[3*cm, 5.5*cm, 3*cm, 5.5*cm])
                    _tbl_f.setStyle(_TblSt([
                        ("FONTSIZE",(0,0),(-1,-1),9),
                        ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),
                        ("FONTNAME",(2,0),(2,-1),"Helvetica-Bold"),
                        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#bbcdd8")),
                        ("ROWBACKGROUNDS",(0,0),(-1,-1),
                         [colors.HexColor("#f0f4f8"),colors.white]),
                        ("TOPPADDING",(0,0),(-1,-1),3),
                        ("BOTTOMPADDING",(0,0),(-1,-1),3),
                    ]))
                    _stry.append(_tbl_f)
                    _stry.append(Spacer(1, 0.5*cm))
                    _stry.append(Paragraph(
                        f"Norma: {norma_sel}  |  Generado por Konte/StructoPro  |  "
                        f"{_dt_f.date.today().strftime('%Y-%m-%d')}",
                        ParagraphStyle("Cap", parent=_st_f["Normal"],
                                       fontSize=8, textColor=colors.grey)))
                    _doc_f.build(_stry)
                    _pdf_f.seek(0)
                    st.download_button(
                        _t("Descargar Ficha PDF","Download PDF Sheet"),
                        data=_pdf_f,
                        file_name=f"FichaTecnica_Zapata_{B_cp:.1f}x{L_cp:.1f}.pdf",
                        mime="application/pdf",
                        key="ficha_dl_pdf",
                        use_container_width=True
                    )
                    _ficha_ok.append("PDF")
                except Exception as _e_fp:
                    _ficha_err.append(f"PDF: {_e_fp}")

            # ── DXF (reutilizar bytes ya generados si existen) ────────────
            if _ficha_incl_dxf:
                _dxf_prev = st.session_state.get("_last_dxf_bytes")
                if _dxf_prev:
                    st.download_button(
                        _t("Descargar Ficha DXF","Download DXF Drawing"),
                        data=_dxf_prev,
                        file_name=f"FichaTecnica_Zapata_{B_cp:.1f}x{L_cp:.1f}.dxf",
                        mime="application/octet-stream",
                        key="ficha_dl_dxf", use_container_width=True
                    )
                    _ficha_ok.append("DXF")
                else:
                    st.caption(_t(
                        "DXF: Genera primero el plano DXF en la sección 'Planos DXF ICONTEC' "
                        "para incluirlo en la ficha.",
                        "DXF: First generate the DXF drawing in 'ICONTEC DXF Plans' section "
                        "to include it in the sheet."
                    ))

            # ── IFC ───────────────────────────────────────────────────────
            if _ficha_incl_ifc and _IFC_AVAILABLE:
                _ifc_prev = st.session_state.get("_last_ifc_bytes")
                if _ifc_prev:
                    st.download_button(
                        _t("Descargar Ficha IFC","Download IFC Model"),
                        data=_ifc_prev,
                        file_name=f"FichaTecnica_Zapata_{B_cp:.1f}x{L_cp:.1f}.ifc",
                        mime="application/octet-stream",
                        key="ficha_dl_ifc", use_container_width=True
                    )
                    _ficha_ok.append("IFC")
                elif not _IFC_AVAILABLE:
                    st.caption(_t("IFC no disponible (ifcopenshell no instalado).",
                                  "IFC not available (ifcopenshell not installed)."))
                else:
                    st.caption(_t(
                        "IFC: Genera primero el modelo IFC en la sección BIM.",
                        "IFC: First generate the IFC model in the BIM section."
                    ))

            # Resumen
            if _ficha_ok:
                st.success(_t(
                    f"Ficha unificada generada: {' + '.join(_ficha_ok)}",
                    f"Unified sheet generated: {' + '.join(_ficha_ok)}"
                ))
            if _ficha_err:
                for _fe in _ficha_err:
                    st.error(f"Error: {_fe}")

        #  EXPORTACIÓN DXF (Protocolo Diamante N4)
        import datetime as _dt_dxf
        st.write("#### Planos DXF ICONTEC")
        papel_opc = {"Carta (21x28cm)": (21.6,27.9,"CARTA"), "Oficio (21x33cm)": (21.6,33.0,"OFICIO"), "Medio Pliego (50x70cm)": (50.0,70.7,"MEDIO PLIEGO"), "Pliego (70x100cm)": (70.7,100.0,"PLIEGO")}
        papel_sel = st.selectbox("Formato de Papel (Zapata Aislada):", list(papel_opc.keys()), key="za_papel")
        W_P, H_P, LBL_P = papel_opc[papel_sel]

        _dxf_bytes = None
        if st.button("Generar Plano DXF ICONTEC - Zapata Aislada"):
            doc_dxf = ezdxf.new('R2010', setup=True)
            doc_dxf.units = ezdxf.units.CM
            LW = {'CONCRETO':70, 'ACERO_L':40, 'ACERO_B':40, 'ACERO_SUP':40, 'COTAS':25, 'TEXTO':25, 'ROTULO':35, 'MARGEN':70}
            COL = {'CONCRETO':7, 'ACERO_L':3, 'ACERO_B':1, 'ACERO_SUP':6, 'COTAS':2, 'TEXTO':7, 'ROTULO':8, 'MARGEN':7}
            for lay, lw in LW.items():
                doc_dxf.layers.new(lay, dxfattribs={'color':COL[lay], 'lineweight':lw})
            doc_dxf.styles.new('ROMANS', dxfattribs={'font':'Arial.ttf'})
            
            msp = doc_dxf.modelspace()
            
            B_cm = B_use * 100; L_cm = L_use * 100; Hc = H_zap; rec = recub_z
            db_cm = db_bar_z / 10.0; radio_dob = D_doblez_cm / 2.0; L_ext = L_ext_gancho_cm
            c1c = c1_col; c2c = c2_col
            
            # Dibujar en Model Space (unidades reales CM)
            def _rect(ox, oy, w, h, lay):
                msp.add_lwpolyline([(ox,oy),(ox+w,oy),(ox+w,oy+h),(ox,oy+h),(ox,oy)], dxfattribs={'layer': lay})
            
            def _text(x, y, txt, lay, h=2.5, ha='left', va='center'):
                halign = {'left':0, 'center':1, 'right':2}[ha]
                msp.add_text(txt, dxfattribs={'layer': lay, 'style': 'ROMANS', 'height': h, 'insert': (x,y), 'align_point': (x,y), 'halign': halign, 'valign': 2 if va=='center' else 0})
                
            def _hook_arc(ox, oy, side, lay):
                r = radio_dob
                if side == 'left':
                    msp.add_arc((ox + r, oy + r), r, start_angle=180, end_angle=270, dxfattribs={'layer': lay})
                    msp.add_line((ox, oy + r), (ox, oy + r + L_ext), dxfattribs={'layer': lay})
                else:
                    msp.add_arc((ox - r, oy + r), r, start_angle=270, end_angle=360, dxfattribs={'layer': lay})
                    msp.add_line((ox, oy + r), (ox, oy + r + L_ext), dxfattribs={'layer': lay})

            ex_A = 0; ey_A = 0
            _rect(ex_A, ey_A, B_cm, Hc, 'CONCRETO')
            _rect(ex_A + (B_cm - c1c) / 2, ey_A + Hc, c1c, 50, 'CONCRETO')
            y_bar_A = ey_A + rec + db_cm/2
            for j in range(int(n_barras_L)): msp.add_circle((ex_A + rec + j * sep_L, y_bar_A), db_cm/2, dxfattribs={'layer':'ACERO_L'})
            msp.add_line((ex_A + rec + radio_dob, y_bar_A), (ex_A + B_cm - rec - radio_dob, y_bar_A), dxfattribs={'layer':'ACERO_B'})
            _hook_arc(ex_A + rec, y_bar_A, 'left', 'ACERO_B'); _hook_arc(ex_A + B_cm - rec, y_bar_A, 'right', 'ACERO_B')
            _text(ex_A + B_cm/2, ey_A - 15, "CORTE A-A (Dir. B)", 'TEXTO', h=3.0, ha='center')
            # ── Acero SUPERIOR Dir. B (Corte A-A) — solo si doble parrilla ──────────
            if dp_B and n_barras_B_sup > 0:
                _y_sup_AA = ey_A + Hc - rec - db_cm/2
                for _j in range(int(n_barras_B_sup)):
                    _xs = ex_A + rec + _j * sep_B_sup if n_barras_B_sup > 1 else ex_A + B_cm/2
                    msp.add_circle((_xs, _y_sup_AA), db_cm/2, dxfattribs={'layer':'ACERO_SUP'})
                msp.add_line((ex_A + rec + radio_dob, _y_sup_AA), (ex_A + B_cm - rec - radio_dob, _y_sup_AA),
                             dxfattribs={'layer':'ACERO_SUP'})
                _text(ex_A + B_cm/2, _y_sup_AA + db_cm + 2,
                      f"As_SUP={As_req_B_sup:.1f}cm2 ({n_barras_B_sup}{bar_z}@{sep_B_sup:.0f}cm)",
                      'TEXTO', h=2.0, ha='center')


            ex_B = B_cm + 60; ey_B = 0
            _rect(ex_B, ey_B, L_cm, Hc, 'CONCRETO')
            _rect(ex_B + (L_cm - c2c) / 2, ey_B + Hc, c2c, 50, 'CONCRETO')
            y_bar_B = ey_B + rec + db_cm/2
            for i in range(int(n_barras_B)): msp.add_circle((ex_B + rec + i * sep_B, y_bar_B), db_cm/2, dxfattribs={'layer':'ACERO_B'})
            y_bar_L = y_bar_B + db_cm
            msp.add_line((ex_B + rec + radio_dob, y_bar_L), (ex_B + L_cm - rec - radio_dob, y_bar_L), dxfattribs={'layer':'ACERO_L'})
            _hook_arc(ex_B + rec, y_bar_L, 'left', 'ACERO_L'); _hook_arc(ex_B + L_cm - rec, y_bar_L, 'right', 'ACERO_L')
            _text(ex_B + L_cm/2, ey_B - 15, "CORTE B-B (Dir. L)", 'TEXTO', h=3.0, ha='center')

            px = 0; py = Hc + 60
            _rect(px, py, B_cm, L_cm, 'CONCRETO')
            _rect(px + (B_cm - c1c)/2, py + (L_cm - c2c)/2, c1c, c2c, 'CONCRETO')
            for i in range(int(n_barras_B)):
                yi = py + rec + i * sep_B
                msp.add_line((px + rec, yi), (px + B_cm - rec, yi), dxfattribs={'layer':'ACERO_B'})
            for j in range(int(n_barras_L)):
                xj = px + rec + j * sep_L
                msp.add_line((xj, py + rec), (xj, py + L_cm - rec), dxfattribs={'layer':'ACERO_L'})
            _text(px + B_cm/2, py - 15, "PLANTA ZAPATA", 'TEXTO', h=3.0, ha='center')
            
            # ── Acero SUPERIOR Dir. L (Corte B-B) — solo si doble parrilla ──────────
            if dp_L and n_barras_L_sup > 0:
                _y_sup_BB = ey_B + Hc - rec - db_cm/2
                for _i in range(int(n_barras_L_sup)):
                    _xs = ex_B + rec + _i * sep_L_sup if n_barras_L_sup > 1 else ex_B + L_cm/2
                    msp.add_circle((_xs, _y_sup_BB), db_cm/2, dxfattribs={'layer':'ACERO_SUP'})
                msp.add_line((ex_B + rec + radio_dob, _y_sup_BB), (ex_B + L_cm - rec - radio_dob, _y_sup_BB),
                             dxfattribs={'layer':'ACERO_SUP'})
                _text(ex_B + L_cm/2, _y_sup_BB + db_cm + 2,
                      f"As_SUP={As_req_L_sup:.1f}cm2 ({n_barras_L_sup}{bar_z}@{sep_L_sup:.0f}cm)",
                      'TEXTO', h=2.0, ha='center')

            # Mover dibujos a un bloque (virtual) y escalarlo para encajar en el papel
            escala_den = 50
            dim_w = max(B_cm+60+L_cm, B_cm)
            dim_h = (Hc + 60 + L_cm)
            for den in [100, 75, 50, 25, 20]:
                if dim_w / den <= W_P*0.75 and dim_h / den <= H_P*0.75:
                    escala_den = den; break
            
            # Margen y Rotulo Protocolo N4
            msp.add_lwpolyline([(0.5,0.5), (W_P-0.5,0.5), (W_P-0.5,H_P-0.5), (0.5,H_P-0.5), (0.5,0.5)], dxfattribs={'layer':'MARGEN'})
            msp.add_lwpolyline([(0.5,0.5), (W_P-0.5,0.5), (W_P-0.5,4.5), (0.5,4.5), (0.5,0.5)], dxfattribs={'layer':'ROTULO'})
            _text(W_P/2, 2.5, f"ZAPATA AISLADA {B_use:.2f}x{L_use:.2f}m - H={H_zap:.0f}cm | {norma_sel}", 'TEXTO', h=0.4, ha='center')
            _text(W_P/2, 1.5, f"Papel: {LBL_P} | Escala Vista: 1:{escala_den}", 'TEXTO', h=0.3, ha='center')
            
            # Crear un nuevo bloque y mover entidades
            # (Por simplicidad en streamlit-ezdxf al vuelo, redibujaremos con escala)
            doc_dxf.blocks.new(name='VISTAS_ZAP')
            blk = doc_dxf.blocks.get('VISTAS_ZAP')
            
            # ── Acero SUPERIOR en planta (líneas punteadas) — solo si doble parrilla ─
            if _doble_parrilla:
                if dp_B and n_barras_B_sup > 0:
                    for _i in range(int(n_barras_B_sup)):
                        _yi_s = py + rec + _i * sep_B_sup if n_barras_B_sup > 1 else py + L_cm/2
                        _line = msp.add_line((px + rec, _yi_s), (px + B_cm - rec, _yi_s),
                                             dxfattribs={'layer':'ACERO_SUP'})
                        _line.dxf.linetype = 'DASHED'
                if dp_L and n_barras_L_sup > 0:
                    for _j in range(int(n_barras_L_sup)):
                        _xj_s = px + rec + _j * sep_L_sup if n_barras_L_sup > 1 else px + B_cm/2
                        _line = msp.add_line((_xj_s, py + rec), (_xj_s, py + L_cm - rec),
                                             dxfattribs={'layer':'ACERO_SUP'})
                        _line.dxf.linetype = 'DASHED'
            # Transferimos desde msp a blk y borramos de msp
            for e in list(msp):
                if e.dxf.layer not in ['MARGEN', 'ROTULO', 'TEXTO']:
                    blk.add_entity(e.copy())
                    e.destroy()
            for e in list(msp): # Limpiar textos h=3.0 que son del dibujo
                if e.dxf.layer == 'TEXTO' and e.dxf.height > 1.0:
                    blk.add_entity(e.copy())
                    e.destroy()
            
            # Insertar el bloque escalado en el centro
            ESC = 1.0 / escala_den
            ox = (W_P - dim_w*ESC)/2
            oy = 5 + (H_P - 6 - dim_h*ESC)/2
            msp.add_blockref('VISTAS_ZAP', insert=(ox, oy), dxfattribs={'xscale': ESC, 'yscale': ESC})

            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.dxf', delete=False) as tmp:
                tmp_path = tmp.name
            doc_dxf.saveas(tmp_path)
            with open(tmp_path, 'rb') as f:
                _dxf_bytes = f.read()
            os.unlink(tmp_path)
            st.success("DXF generado correctamente con Estándar Diamante.")

        #  EXPORTACIÓN IFC AUTÓNOMA 
        st.markdown("---")
        st.markdown("####  Exportar Modelo BIM (.ifc)")
        st.caption("Implementación Nativa IFC4 con Parametrización NSR-10/ACI")
        
        def _make_ifc_zapata_aislada(B_m, L_m, H_m, c1_m, c2_m, fc_mpa, fy_mpa,
                                     bar_B, n_B, bar_L, n_L, recub_cm,
                                     Pu_kN, norma, proyecto,
                                     dp_B=False, dp_L=False,
                                     n_B_sup=0, n_L_sup=0):
            """IFC4 completo: hormigón (zapata + pedestal) + refuerzo inferior y superior."""
            try:
                import ifcopenshell
                import ifcopenshell.api
            except ImportError:
                return None

            O = ifcopenshell.file(schema="IFC4")
            db_m = REBAR_DICT[bar_B]["db"] / 1000.0   # diámetro varilla en m

            def p4(x, y, z): return O.createIfcCartesianPoint((float(x), float(y), float(z)))
            def ax2(pt, z_dir=(0., 0., 1.), x_dir=(1., 0., 0.)):
                return O.createIfcAxis2Placement3D(
                    pt, O.createIfcDirection(z_dir), O.createIfcDirection(x_dir))
            def make_rect_solid(w, d, h, orig=(0., 0., 0.)):
                prof = O.createIfcRectangleProfileDef("AREA", None, ax2(p4(*orig)), w, d)
                return O.createIfcExtrudedAreaSolid(prof, ax2(p4(*orig)), O.createIfcDirection((0., 0., 1.)), h)

            # ── Jerarquía del proyecto ──────────────────────────────────────────
            project  = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcProject",  name=proyecto)
            context  = ifcopenshell.api.run("context.add_context", O, context_type="Model")
            body     = ifcopenshell.api.run("context.add_context", O, context_type="Model",
                            context_identifier="Body", target_view="MODEL_VIEW", parent_context=context)
            site     = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcSite",     name="Sitio")
            building = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcBuilding", name="Edificio")
            _u = ifcopenshell.api.run("unit.add_si_unit", O, unit_type="LENGTHUNIT")
            ifcopenshell.api.run("unit.assign_unit", O, units=[_u])
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=project,  products=[site])
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=site,     products=[building])

            # ── Zapata ──────────────────────────────────────────────────────────
            zapata = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcFooting",
                                         name="Zapata Aislada")
            zapata.PredefinedType = "PAD_FOOTING"
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=building, products=[zapata])
            shape_zap = O.createIfcShapeRepresentation(body, "Body", "SweptSolid",
                                                       [make_rect_solid(B_m, L_m, H_m)])
            ifcopenshell.api.run("geometry.assign_representation", O, product=zapata, representation=shape_zap)
            ifcopenshell.api.run("geometry.edit_object_placement", O, product=zapata)

            # ── Pedestal ────────────────────────────────────────────────────────
            pedestal = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcColumn",
                                            name="Pedestal / Columna")
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=building, products=[pedestal])
            shape_ped = O.createIfcShapeRepresentation(body, "Body", "SweptSolid",
                                                       [make_rect_solid(c1_m, c2_m, H_m * 0.5,
                                                                        orig=(B_m/2 - c1_m/2,
                                                                              L_m/2 - c2_m/2, H_m))])
            ifcopenshell.api.run("geometry.assign_representation", O, product=pedestal, representation=shape_ped)
            ifcopenshell.api.run("geometry.edit_object_placement", O, product=pedestal)

            # ── Función auxiliar: crear barra cilíndrica IfcReinforcingBar ──────
            def add_rebar(name, x0, y0, z0, x1, y1, z1, diameter_m, predefined="BOTTOM"):
                seg = O.createIfcPolyline([p4(x0, y0, z0), p4(x1, y1, z1)])
                rep = O.createIfcShapeRepresentation(body, "Body", "Curve3D", [seg])
                rb  = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcReinforcingBar",
                                           name=name)
                rb.NominalDiameter  = float(diameter_m)
                rb.BarLength        = float(((x1-x0)**2+(y1-y0)**2+(z1-z0)**2)**0.5)
                rb.PredefinedType   = "LIGATURE" if "EST" in name.upper() else "LONGITUDINAL"
                ifcopenshell.api.run("aggregate.assign_object", O, relating_object=zapata, products=[rb])
                ifcopenshell.api.run("geometry.assign_representation", O, product=rb, representation=rep)
                ifcopenshell.api.run("geometry.edit_object_placement", O, product=rb)
                return rb

            # ── Barras inferiores Dir B (paralelas a L) ─────────────────────────
            rec_m  = recub_cm / 100.0
            z_inf  = rec_m + db_m / 2.0
            z_sup  = H_m - rec_m - db_m / 2.0
            y_step = (L_m - 2 * rec_m) / max(n_B - 1, 1)
            for i in range(n_B):
                y_pos = rec_m + i * y_step
                add_rebar(f"B_INF_{i+1}", 0, y_pos, z_inf, B_m, y_pos, z_inf, db_m)

            # ── Barras inferiores Dir L (paralelas a B) ─────────────────────────
            x_step = (B_m - 2 * rec_m) / max(n_L - 1, 1)
            for i in range(n_L):
                x_pos = rec_m + i * x_step
                add_rebar(f"L_INF_{i+1}", x_pos, 0, z_inf, x_pos, L_m, z_inf, db_m)

            # ── Barras superiores Dir B — solo si doble parrilla ───────────────
            if dp_B and n_B_sup > 0:
                y_step_s = (L_m - 2 * rec_m) / max(n_B_sup - 1, 1)
                for i in range(n_B_sup):
                    y_pos = rec_m + i * y_step_s
                    add_rebar(f"B_SUP_{i+1}", 0, y_pos, z_sup, B_m, y_pos, z_sup, db_m)

            # ── Barras superiores Dir L — solo si doble parrilla ───────────────
            if dp_L and n_L_sup > 0:
                x_step_s = (B_m - 2 * rec_m) / max(n_L_sup - 1, 1)
                for i in range(n_L_sup):
                    x_pos = rec_m + i * x_step_s
                    add_rebar(f"L_SUP_{i+1}", x_pos, 0, z_sup, x_pos, L_m, z_sup, db_m)

            # ── Propiedades Pset normativa ───────────────────────────────────────
            pset = ifcopenshell.api.run("pset.add_pset", O, product=zapata,
                                        name="Pset_StructoPro_Zapata")
            props = {
                "Norma":          norma,
                "f_c_MPa":        float(fc_mpa),
                "fy_MPa":         float(fy_mpa),
                "Acero_Dir_B_inf": f"{n_B} {bar_B}",
                "Acero_Dir_L_inf": f"{n_L} {bar_L}",
                "Carga_Pu_kN":    float(Pu_kN),
                "Recubrimiento_cm": float(recub_cm),
                "Doble_Parrilla": "SÍ" if (dp_B or dp_L) else "NO",
            }
            if dp_B and n_B_sup > 0:
                props["Acero_Dir_B_sup"] = f"{n_B_sup} {bar_B}"
            if dp_L and n_L_sup > 0:
                props["Acero_Dir_L_sup"] = f"{n_L_sup} {bar_L}"
            ifcopenshell.api.run("pset.edit_pset", O, pset=pset, properties=props)

            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.ifc', delete=False) as tmp:
                tmp_path = tmp.name
            O.write(tmp_path)
            with open(tmp_path, 'rb') as f:
                b_ifc = f.read()
            os.unlink(tmp_path)
            return b_ifc

        buf_ifc = None
        if st.button("Generar Modelo BIM (IFC) - Zapata Aislada"):
            with st.spinner("Generando IFC4 nativo..."):
                buf_ifc = _make_ifc_zapata_aislada(
                    B_use, L_use, H_zap/100, c1_col/100, c2_col/100, fc_basico, fy_basico,
                    bar_z, int(n_barras_B), bar_z, int(n_barras_L), recub_z, P_ult, norma_sel,
                    st.session_state.get("proyecto", "") or "Proyecto",
                    dp_B=dp_B, dp_L=dp_L,
                    n_B_sup=int(n_barras_B_sup), n_L_sup=int(n_barras_L_sup)
                )

                if buf_ifc:
                    st.success("IFC4 generado con éxito.")
                else:
                    st.error("Error al generar IFC.")

        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            if _puede_exportar:
                st.download_button("Memoria DOCX", data=f_zap_io,
                               file_name=f"Memoria_Zapata_{B_use:.1f}x{L_use:.1f}m.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True, disabled=not area_ok)
            else:
                st.button("Memoria DOCX", disabled=True, help="Requiere licencia Pro o Admin.", use_container_width=True)
        with col_m2:
            if _puede_exportar:
                st.download_button("Descargar Plano DXF (ICONTEC)", data=_dxf_bytes,
                               file_name=f"Zapata_Aislada_{B_use:.1f}x{L_use:.1f}.dxf" if _dxf_bytes else "Plano.dxf",
                               mime="application/dxf",
                               use_container_width=True, disabled=(_dxf_bytes is None))
            else:
                st.button("Descargar Plano DXF (ICONTEC)", disabled=True,
                          help="Requiere licencia Pro o Admin.", use_container_width=True)
        with col_m3:
            if _puede_exportar:
                st.download_button("Descargar Modelo BIM (IFC)", data=buf_ifc,
                               file_name=f"Zapata_Aislada_{B_use:.1f}x{L_use:.1f}.ifc" if buf_ifc else "Modelo.ifc",
                               mime="application/x-step",
                               use_container_width=True, disabled=(buf_ifc is None))
            else:
                st.button("Descargar Modelo BIM (IFC)", disabled=True,
                          help="Requiere licencia Pro o Admin.", use_container_width=True)

    with tab_dwg:
        st.subheader("Visualización 3D de la Fundación con Acero de Refuerzo")
        fig3d = go.Figure()

        Hz_m  = H_zap / 100.0
        rec_m = recub_z / 100.0
        db_m  = db_bar_z / 1000.0
        z_bar = rec_m + db_m / 2

        # Zapata
        x_z = [-B_use/2, B_use/2, B_use/2, -B_use/2, -B_use/2, B_use/2, B_use/2, -B_use/2]
        y_z = [-L_use/2, -L_use/2, L_use/2, L_use/2, -L_use/2, -L_use/2, L_use/2, L_use/2]
        z_z = [0]*4 + [Hz_m]*4
        fig3d.add_trace(go.Mesh3d(x=x_z, y=y_z, z=z_z, alphahull=0, opacity=0.25,
                                  color='steelblue', name='Zapata', showlegend=True))

        # Columna
        c1_m = c1_col/100.0; c2_m = c2_col/100.0
        x_c = [-c1_m/2, c1_m/2, c1_m/2, -c1_m/2, -c1_m/2, c1_m/2, c1_m/2, -c1_m/2]
        y_c = [-c2_m/2, -c2_m/2, c2_m/2, c2_m/2, -c2_m/2, -c2_m/2, c2_m/2, c2_m/2]
        z_c = [Hz_m]*4 + [Hz_m + 1.0]*4
        fig3d.add_trace(go.Mesh3d(x=x_c, y=y_c, z=z_c, alphahull=0, opacity=0.5,
                                  color='slategray', name='Columna', showlegend=True))

        # Medidas para ganchos en 3D
        _r_m = radio_doblez_cm / 100.0          # radio de curva en metros
        _ext_m = L_ext_gancho_cm / 100.0         # extensión recta en metros
        _hook_h = (math.pi * _r_m / 2) + _ext_m  # altura total del gancho (arco + recta)

        # Varillas Dir. B (barras corren en X=dirección B, se distribuyen en Y=dirección L)
        _ys_barB = np.linspace(-L_use/2 + rec_m, L_use/2 - rec_m, n_barras_B)
        _show_B = True
        _show_Bh = True
        for yi in _ys_barB:
            # Tramo recto horizontal
            fig3d.add_trace(go.Scatter3d(
                x=[-B_use/2 + rec_m, B_use/2 - rec_m], y=[yi, yi], z=[z_bar, z_bar],
                mode='lines', line=dict(color='#ff6b35', width=5),
                name='Acero Dir.B' if _show_B else None, showlegend=_show_B, legendgroup='aB'))
            _show_B = False
            # Gancho extremo -X (sube verticalmente)
            fig3d.add_trace(go.Scatter3d(
                x=[-B_use/2 + rec_m, -B_use/2 + rec_m], y=[yi, yi],
                z=[z_bar, z_bar + _hook_h],
                mode='lines', line=dict(color='#ff9a6c', width=3, dash='dot'),
                name='Ganchos Dir.B' if _show_Bh else None, showlegend=_show_Bh, legendgroup='aBh'))
            # Gancho extremo +X
            fig3d.add_trace(go.Scatter3d(
                x=[B_use/2 - rec_m, B_use/2 - rec_m], y=[yi, yi],
                z=[z_bar, z_bar + _hook_h],
                mode='lines', line=dict(color='#ff9a6c', width=3, dash='dot'),
                name=None, showlegend=False, legendgroup='aBh'))
            _show_Bh = False

        # Varillas Dir. L (barras corren en Y=dirección L, se distribuyen en X=dirección B)
        _xs_barL = np.linspace(-B_use/2 + rec_m, B_use/2 - rec_m, n_barras_L)
        _z_barL  = z_bar + db_m
        _hook_h_L = _hook_h + db_m   # ligeramente más alto porque van encima
        _show_L = True
        _show_Lh = True
        for xi in _xs_barL:
            # Tramo recto horizontal
            fig3d.add_trace(go.Scatter3d(
                x=[xi, xi], y=[-L_use/2 + rec_m, L_use/2 - rec_m], z=[_z_barL, _z_barL],
                mode='lines', line=dict(color='#ffd54f', width=5),
                name='Acero Dir.L' if _show_L else None, showlegend=_show_L, legendgroup='aL'))
            _show_L = False
            # Gancho extremo -Y (sube verticalmente)
            fig3d.add_trace(go.Scatter3d(
                x=[xi, xi], y=[-L_use/2 + rec_m, -L_use/2 + rec_m],
                z=[_z_barL, _z_barL + _hook_h_L],
                mode='lines', line=dict(color='#ffe57f', width=3, dash='dot'),
                name='Ganchos Dir.L' if _show_Lh else None, showlegend=_show_Lh, legendgroup='aLh'))
            # Gancho extremo +Y
            fig3d.add_trace(go.Scatter3d(
                x=[xi, xi], y=[L_use/2 - rec_m, L_use/2 - rec_m],
                z=[_z_barL, _z_barL + _hook_h_L],
                mode='lines', line=dict(color='#ffe57f', width=3, dash='dot'),
                name=None, showlegend=False, legendgroup='aLh'))
            _show_Lh = False

        fig3d.update_layout(
            scene=dict(aspectmode='data', xaxis_title='B (m)', yaxis_title='L (m)', zaxis_title='Z (m)',
                       bgcolor='#0f1117',
                       xaxis=dict(showgrid=True, gridcolor='#333'),
                       yaxis=dict(showgrid=True, gridcolor='#333'),
                       zaxis=dict(showgrid=True, gridcolor='#333')),
            margin=dict(l=0, r=0, b=0, t=30), height=550,
            showlegend=True, dragmode='turntable', paper_bgcolor='#0f1117', font=dict(color='white'),
            title=dict(text=f"Zapata {B_use:.2f}x{L_use:.2f}m | H={H_zap:.0f}cm | "
                           f"Dir.B: {n_barras_B}×{bar_z} c/{sep_B:.1f}cm | Dir.L: {n_barras_L}×{bar_z} c/{sep_L:.1f}cm",
                       font=dict(color='white')))
        st.plotly_chart(fig3d, use_container_width=True)
        
        st.markdown("---")
        st.write("#### Geometría de Zapata 2D")
        fig_z, ax_z = plt.subplots(figsize=(6, 4))
        ax_z.set_facecolor('#1a1a2e'); fig_z.patch.set_facecolor('#1a1a2e')
        ax_z.add_patch(patches.Rectangle((0,0), B_use*100, H_zap, linewidth=2, edgecolor='darkgray', facecolor='#4a4a6a'))
        pos_x_col = (B_use*100 - c1_col) / 2
        ax_z.add_patch(patches.Rectangle((pos_x_col, H_zap), c1_col, 50, linewidth=2, edgecolor='white', facecolor='#6a6a8a'))
        for i in range(n_barras_Z):
            xi = recub_z + i * separacion_S
            ax_z.add_patch(plt.Circle((xi, recub_z), db_bar_z/10, color='#ff6b35', zorder=5))
        # Dibujar perfil del doblez para la varilla principal (Dir B)
        # Tramo horizontal inferior
        _x_start = recub_z + radio_doblez_cm
        _x_end = B_use*100 - recub_z - radio_doblez_cm
        _y_bar = recub_z
        ax_z.add_patch(patches.Rectangle((_x_start, _y_bar - db_bar_z/20.0), _x_end - _x_start, db_bar_z/10.0, color='#ffd54f', zorder=4))
        
        # Gancho Izquierdo
        arc_izq = patches.Arc((_x_start, _y_bar + radio_doblez_cm), D_doblez_cm, D_doblez_cm, angle=180, theta1=90, theta2=180, color='#ffd54f', lw=3, zorder=4)
        ax_z.add_patch(arc_izq)
        ax_z.plot([recub_z, recub_z], [_y_bar + radio_doblez_cm, _y_bar + radio_doblez_cm + L_ext_gancho_cm], color='#ffd54f', lw=3, zorder=4)
        
        # Gancho Derecho
        arc_der = patches.Arc((_x_end, _y_bar + radio_doblez_cm), D_doblez_cm, D_doblez_cm, angle=270, theta1=90, theta2=180, color='#ffd54f', lw=3, zorder=4)
        ax_z.add_patch(arc_der)
        ax_z.plot([B_use*100 - recub_z, B_use*100 - recub_z], [_y_bar + radio_doblez_cm, _y_bar + radio_doblez_cm + L_ext_gancho_cm], color='#ffd54f', lw=3, zorder=4)
        
        # Cota del doblez
        ax_z.annotate(f"{L_ext_gancho_cm:.1f}cm", xy=(recub_z, _y_bar + radio_doblez_cm + L_ext_gancho_cm/2), xytext=(-15, _y_bar + radio_doblez_cm + L_ext_gancho_cm/2), arrowprops=dict(arrowstyle="->", color='yellow'), color='yellow', fontsize=8, va='center')

        ax_z.text(B_use*100/2, H_zap/2, f"{n_barras_Z} varillas {bar_z} L={L_use}m\nSep:{separacion_S:.1f}cm\nGancho: 90°", color='white', ha='center', va='center')
        ax_z.set_xlim(-20, B_use*100+20)
        ax_z.set_ylim(-10, H_zap+70)
        ax_z.axis('off')
        st.pyplot(fig_z)
        plt.close(fig_z)
        
        st.markdown("---")
        st.write("####  Vista 3D Nativa (Volumen de Concreto)")
        fig_vol = go.Figure()
        # Base zapata (prisma rectangular)
        bx = B_use; ly = L_use; hz = H_zap / 100.0
        # Definir los 8 vertices
        vx = [-bx/2, bx/2, bx/2, -bx/2, -bx/2, bx/2, bx/2, -bx/2]
        vy = [-ly/2, -ly/2, ly/2, ly/2, -ly/2, -ly/2, ly/2, ly/2]
        vz = [0, 0, 0, 0, hz, hz, hz, hz]
        fig_vol.add_trace(go.Mesh3d(x=vx, y=vy, z=vz, alphahull=0, opacity=0.75, color='#a8e6cf', name='Zapata'))
        
        # Extrusión central (columna)
        c1 = c1_col / 100.0; c2 = c2_col / 100.0; hc = 1.0 # 1m de pedestal
        cx = [-c1/2, c1/2, c1/2, -c1/2, -c1/2, c1/2, c1/2, -c1/2]
        cy = [-c2/2, -c2/2, c2/2, c2/2, -c2/2, -c2/2, c2/2, c2/2]
        cz = [hz, hz, hz, hz, hz+hc, hz+hc, hz+hc, hz+hc]
        fig_vol.add_trace(go.Mesh3d(x=cx, y=cy, z=cz, alphahull=0, opacity=0.9, color='lightgray', name='Columna'))
        
        fig_vol.update_layout(scene=dict(aspectmode='data',
                                         xaxis=dict(showgrid=True, gridcolor='#333'),
                                         yaxis=dict(showgrid=True, gridcolor='#333'),
                                         zaxis=dict(showgrid=True, gridcolor='#333')),
                              paper_bgcolor='#0f1117', margin=dict(l=0,r=0,b=0,t=0), height=400)
        st.plotly_chart(fig_vol, use_container_width=True)


    with tab_apu:
        #  Cantidades base 
        vol_excavacion   = (B_use + 0.5) * (L_use + 0.5) * Df_z
        vol_concreto_zap = B_use * L_use * (H_zap / 100.0)

        # Longitud de varilla con gancho real (arco + extensión)
        _area_m2  = REBAR_DICT[bar_z]["area"] * 1e-4
        _long_gancho_m  = ((math.pi * radio_doblez_cm / 2) + L_ext_gancho_cm) / 100.0
        _long_var_B = L_use - 2*(recub_z/100.0) + 2*_long_gancho_m
        _long_var_L = B_use - 2*(recub_z/100.0) + 2*_long_gancho_m
        _kg_por_m   = REBAR_DICT[bar_z]["area"] * 1e-4 * 7850          # kg/m por varilla
        peso_barras_B_apu = n_barras_B * _long_var_B * _kg_por_m
        peso_barras_L_apu = n_barras_L * _long_var_L * _kg_por_m
        peso_total_acero_zap = peso_barras_B_apu + peso_barras_L_apu

        # Proporciones concreto (ACI 211 aprox. para fc 21 MPa)
        pct_arena_apu = 0.55; pct_grava_apu = 0.80
        bultos_zap    = vol_concreto_zap * 350 / 50.0               # bultos de 50 kg
        vol_arena_z   = vol_concreto_zap * pct_arena_apu             # m³
        vol_grava_z   = vol_concreto_zap * pct_grava_apu             # m³
        litros_agua   = vol_concreto_zap * 185.0                     # l/m³  (rel a/c ≈0.53)

        #  SECCIÓN 1: Resumen de materiales 
        st.markdown("###  Resumen de Materiales — Quantiy Take-Off")
        cols_m = st.columns(4)
        cols_m[0].metric(" Excavación", f"{vol_excavacion:.2f} m³")
        cols_m[1].metric(" Vol. Concreto", f"{vol_concreto_zap:.2f} m³")
        cols_m[2].metric(" Acero Total", f"{peso_total_acero_zap:.1f} kg")
        cols_m[3].metric(" Cuantía", f"{peso_total_acero_zap/vol_concreto_zap:.1f} kg/m³")

        st.markdown("####  Ingredientes de Concreto")
        df_mat = pd.DataFrame([
            {"Material": " Cemento",        "Cantidad": f"{bultos_zap:.1f}",    "Unidad": "bultos (50 kg)"},
            {"Material": " Arena",          "Cantidad": f"{vol_arena_z:.3f}",   "Unidad": "m³"},
            {"Material": " Gravilla",        "Cantidad": f"{vol_grava_z:.3f}",   "Unidad": "m³"},
            {"Material": " Agua",            "Cantidad": f"{litros_agua:.0f}",   "Unidad": "litros"},
            {"Material": " Acero refuerzo",  "Cantidad": f"{peso_total_acero_zap:.1f}", "Unidad": "kg"},
        ])
        # Integrar precios guardados en base global o usar defaults directos si no se ha visitado el APU Mercado
        _apu = st.session_state.get("apu_config", {})
        if not _apu:
            st.info(_t(
                'AVISO: Configure los precios en el modulo APU Mercado para obtener costos reales.',
                'NOTICE: Configure prices in the APU Market module to obtain real costs.'
            ))
        _mon = _apu.get("moneda", "COP")
        _p_cem = _apu.get('cemento', 0)
        _p_ace = _apu.get('acero', 0)
        _p_are = _apu.get('arena', 0)
        _p_gra = _apu.get('grava', 0)
        c_excav_u = _apu.get('costo_excav_m3', 0)
        _has_prices = bool(_apu)  # Solo calcular desglose si hay precios configurados en APU Mercado

        _c_cem  = bultos_zap * _p_cem
        _c_ace  = peso_total_acero_zap * _p_ace
        _c_are  = vol_arena_z * _p_are
        _c_gra  = vol_grava_z * _p_gra
        _c_exc  = vol_excavacion * c_excav_u if _has_prices else 0.0
        
        _total_mat = _c_exc + _c_cem + _c_are + _c_gra + _c_ace

        # Calcular Gran Total si hay precios
        _gran_total = 0.0
        if _has_prices:
            total_dias_mo = (peso_total_acero_zap * 0.04) + (vol_concreto_zap * 0.4) + (vol_excavacion * 0.3)
            costo_mo = total_dias_mo * _apu.get('costo_dia_mo', 0)
            costo_directo = _total_mat + costo_mo
            herramienta = costo_mo * _apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * _apu.get("pct_aui", 0.30)
            utilidad = costo_directo * _apu.get("pct_util", 0.05)
            iva = utilidad * _apu.get("iva", 0.19)
            _gran_total = costo_directo + herramienta + aiu + iva

            col_msg, col_metric = st.columns([2, 1])
            col_msg.success(f" Precios actualizados del scraping — {_mon}")
            col_metric.metric(f" Gran Total Proyecto [{_mon}]", f"{_gran_total:,.0f}")
            
            #  DESGLOSE TRANSPARENTE DEL GRAN TOTAL 
            st.markdown("####  Desglose del Presupuesto")
            _pct_he  = _apu.get('pct_herramienta', 0.05)*100
            _pct_aiu = _apu.get('pct_aui', 0.30)*100
            _pct_uti = _apu.get('pct_util', 0.05)*100
            _pct_iva = _apu.get('iva', 0.19)*100
            _desglose = [
                {"Concepto": "① Materiales Directos",  "Base de cálculo": "Excavación + Concreto + Acero", "Importe": f"{_total_mat:,.0f} {_mon}"},
                {"Concepto": "② Mano de Obra",         "Base de cálculo": f"{total_dias_mo:.2f} días × {_apu.get('costo_dia_mo', 0):,.0f}/día", "Importe": f"{costo_mo:,.0f} {_mon}"},
                {"Concepto": " COSTO DIRECTO (CD)",   "Base de cálculo": "① + ②",                         "Importe": f"{costo_directo:,.0f} {_mon}"},
                {"Concepto": f"③ Herramienta Menor",   "Base de cálculo": f"{_pct_he:.1f}% × MO",          "Importe": f"{herramienta:,.0f} {_mon}"},
                {"Concepto": f"④ A.I.U. ({_pct_aiu:.0f}%)", "Base de cálculo": f"{_pct_aiu:.1f}% × CD",   "Importe": f"{aiu:,.0f} {_mon}"},
                {"Concepto": f"⑤ IVA s/ Utilidad",    "Base de cálculo": f"{_pct_iva:.1f}% × Utilidad ({_pct_uti:.1f}% CD)", "Importe": f"{iva:,.0f} {_mon}"},
                {"Concepto": " GRAN TOTAL",           "Base de cálculo": "CD + ③ + ④ + ⑤",                "Importe": f"**{_gran_total:,.0f} {_mon}**"},
            ]
            st.dataframe(pd.DataFrame(_desglose), use_container_width=True, hide_index=True)
        else:
            st.info("ℹ Ve a **APU Mercado** para descargar los costos en tiempo real aquí mismo.")



        # TABLA: Materiales con costos
        _mat_rows = [
            {"Material": " Excavación",    "Cantidad": f"{vol_excavacion:.2f}", "Unidad": "m³",           "Precio Unit.": f"{c_excav_u:,.0f} {_mon}" if _has_prices else "—", "Subtotal": f"{_c_exc:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": " Cemento",        "Cantidad": f"{bultos_zap:.1f}",    "Unidad": "bultos (50kg)", "Precio Unit.": f"{_p_cem:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_cem:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": " Arena",          "Cantidad": f"{vol_arena_z:.3f}",   "Unidad": "m³",           "Precio Unit.": f"{_p_are:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_are:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": " Gravilla",        "Cantidad": f"{vol_grava_z:.3f}",   "Unidad": "m³",           "Precio Unit.": f"{_p_gra:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_gra:,.0f} {_mon}" if _has_prices else "—"},
            {"Material": " Agua",            "Cantidad": f"{litros_agua:.0f}",   "Unidad": "litros",       "Precio Unit.": "—", "Subtotal": "—"},
            {"Material": " Acero refuerzo",  "Cantidad": f"{peso_total_acero_zap:.1f}", "Unidad": "kg",  "Precio Unit.": f"{_p_ace:,.0f} {_mon}" if _has_prices else "—",   "Subtotal": f"{_c_ace:,.0f} {_mon}" if _has_prices else "—"},
        ]
        st.dataframe(pd.DataFrame(_mat_rows), use_container_width=True, hide_index=True)

        if _has_prices:
            _total_mat = _c_exc + _c_cem + _c_are + _c_gra + _c_ace
            st.metric(f"Total Materiales [{_mon}]", f"{_total_mat:,.0f}")

        #  SECCIÓN 2: Despiece de Acero con costos 
        st.markdown("####  Despiece de Acero de Refuerzo")
        db_mm_apu    = REBAR_DICT[bar_z]["db"]
        area_cm2_apu = REBAR_DICT[bar_z]["area"]
        _c_ace_B = peso_barras_B_apu * _p_ace
        _c_ace_L = peso_barras_L_apu * _p_ace

        _row_B = {
            "Dir.": "B  ( sobre L →)",
            "Varilla": bar_z, "db [mm]": f"{db_mm_apu:.1f}",
            "N° Barras": n_barras_B, "Sep. [cm]": f"{sep_B:.1f}",
            "L gancho [cm]": f"{L_ext_gancho_cm:.1f}",
            "L total [m]": f"{_long_var_B:.3f}",
            "kg/m": f"{_kg_por_m:.3f}", "Peso [kg]": f"{peso_barras_B_apu:.2f}",
        }
        _row_L = {
            "Dir.": "L  ( sobre B →)",
            "Varilla": bar_z, "db [mm]": f"{db_mm_apu:.1f}",
            "N° Barras": n_barras_L, "Sep. [cm]": f"{sep_L:.1f}",
            "L gancho [cm]": f"{L_ext_gancho_cm:.1f}",
            "L total [m]": f"{_long_var_L:.3f}",
            "kg/m": f"{_kg_por_m:.3f}", "Peso [kg]": f"{peso_barras_L_apu:.2f}",
        }
        _row_tot = {
            "Dir.": " TOTAL ",
            "Varilla": "", "db [mm]": "",
            "N° Barras": n_barras_B + n_barras_L, "Sep. [cm]": "",
            "L gancho [cm]": "",
            "L total [m]": f"{n_barras_B*_long_var_B + n_barras_L*_long_var_L:.2f}",
            "kg/m": "", "Peso [kg]": f"{peso_total_acero_zap:.2f}",
        }
        if _has_prices:
            _row_B[f"Costo [{_mon}]"] = f"{_c_ace_B:,.0f}"
            _row_L[f"Costo [{_mon}]"] = f"{_c_ace_L:,.0f}"
            _row_tot[f"Costo [{_mon}]"] = f"{(_c_ace_B+_c_ace_L):,.0f}"

        st.dataframe(pd.DataFrame([_row_B, _row_L, _row_tot]), use_container_width=True, hide_index=True)
        st.caption(f"Gancho 90° ACI/NSR-10: D_doblez mín = {radio_doblez_cm*2:.1f} cm | ext. = {L_ext_gancho_cm:.1f} cm")

        #  GRÁFICO: Despiece visual 
        st.markdown("####  Diagrama de Despiece")
        _items_g   = ["Excavación", "Cemento", "Arena", "Gravilla", "Acero Dir.B", "Acero Dir.L"]
        _qty_g     = [vol_excavacion, bultos_zap, vol_arena_z, vol_grava_z, peso_barras_B_apu, peso_barras_L_apu]
        _units_g   = ["m³", "bultos", "m³", "m³", "kg", "kg"]
        _colors_g  = ["#5c8a5a","#e8c07d","#c2a06b","#9b7b5c","#ff6b35","#ffd54f"]

        
        if _has_prices:
            _cost_g = [_c_exc, _c_cem, _c_are, _c_gra, _c_ace_B, _c_ace_L]
            datos = list(zip(_cost_g, _qty_g, _items_g, _colors_g, _units_g))
            datos.sort(reverse=True)
            _cost_g, _qty_g, _items_g, _colors_g, _units_g = zip(*datos)
        else:
            datos = list(zip(_qty_g, _items_g, _colors_g, _units_g))
            datos.sort(reverse=True)
            _qty_g, _items_g, _colors_g, _units_g = zip(*datos)

        _fig_desp = go.Figure()
        _fig_desp.add_trace(go.Bar(
            x=_items_g, y=_qty_g,
            marker_color=_colors_g,
            text=[f"{q:.2f} {u}" for q, u in zip(_qty_g, _units_g)],
            textposition="outside",
            name="Cantidad"
        ))

        if _has_prices:
            _fig_desp.add_trace(go.Bar(
                x=_items_g, y=_cost_g,
                marker_color=_colors_g, opacity=0.5,
                text=[f"{v:,.0f} {_mon}" for v in _cost_g],
                textposition="outside",
                name=f"Costo [{_mon}]",
                yaxis="y2"
            ))
            _fig_desp.update_layout(yaxis2=dict(title=f"Costo [{_mon}]", overlaying="y", side="right", showgrid=False, tickfont=dict(color="#aaa")))

        _fig_desp.update_layout(
            paper_bgcolor="#0f1117", plot_bgcolor="#0f1117",
            font=dict(color="white"), barmode="overlay",
            xaxis=dict(showgrid=False),
            yaxis=dict(title="Cantidad", showgrid=True, gridcolor="#222"),
            legend=dict(bgcolor="#111", font=dict(color="white")),
            margin=dict(l=20, r=20, t=30, b=20), height=380,
            title=dict(text=f"Despiece Zapata {B_use:.2f}x{L_use:.2f}m — {bar_z}", font=dict(color="white"))
        )
        st.plotly_chart(_fig_desp, use_container_width=True)



        if "apu_config" in st.session_state:
            st.markdown("---")
            st.markdown("###  Presupuesto Estimado (Promedio de Fuentes Regionales)")
            apu = st.session_state.get("apu_config", {})
            if not apu:
                st.info(_t(
                    'AVISO: Configure los precios en el modulo APU Mercado.',
                    'NOTICE: Configure prices in APU Market module.'
                ))
                apu = {}
            mon = apu.get("moneda", "")
            c_excav = vol_excavacion * 25000
            
            bultos_zap  = vol_concreto_zap * 350 / 50.0
            pct_arena = apu.get("pct_arena_mezcla", 0.55)
            pct_grava = apu.get("pct_grava_mezcla", 0.80)
            vol_arena_z = vol_concreto_zap * pct_arena
            vol_grava_z = vol_concreto_zap * pct_grava

            c_cem = bultos_zap * apu["cemento"]
            c_ace = peso_total_acero_zap * apu["acero"]
            c_are = vol_arena_z * apu["arena"]
            c_gra = vol_grava_z * apu["grava"]
            total_mat = c_cem + c_ace + c_are + c_gra + c_excav
            
            total_dias_mo = (peso_total_acero_zap * 0.04) + (vol_concreto_zap * 0.4) + (vol_excavacion * 0.3)
            costo_mo = total_dias_mo * apu.get('costo_dia_mo', 0)
            
            costo_directo = total_mat + costo_mo
            herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
            aiu = costo_directo * apu.get("pct_aui", 0.30)
            utilidad = costo_directo * apu.get("pct_util", 0.05)
            iva = utilidad * apu.get("iva", 0.19)
            total_proyecto = costo_directo + herramienta + aiu + iva
            
            data_zap_apu = {
                "Concepto": [
                    "① Excavación", "② Cemento", "③ Acero Refuerzo", "④ Arena", "⑤ Grava",
                    " SUBTOTAL MATERIALES DIRECTOS",
                    "⑥ Mano de Obra",
                    " COSTO DIRECTO (CD)",
                    f"⑦ Herramienta Menor ({apu.get('pct_herramienta', 0.05)*100:.1f}%)",
                    f"⑧ A.I.U. ({apu.get('pct_aui', 0.30)*100:.0f}%)",
                    f"⑨ IVA s/ Utilidad ({apu.get('iva', 0.19)*100:.0f}%)",
                    " GRAN TOTAL DEL PROYECTO",
                ],
                "Base de cálculo": [
                    f"{vol_excavacion:.2f} m³ × 25,000/m³",
                    f"{bultos_zap:.1f} blt × {apu['cemento']:,.0f}/blt",
                    f"{peso_total_acero_zap:.1f} kg × {apu['acero']:,.0f}/kg",
                    f"{vol_arena_z:.2f} m³ × {apu['arena']:,.0f}/m³",
                    f"{vol_grava_z:.2f} m³ × {apu['grava']:,.0f}/m³",
                    "① + ② + ③ + ④ + ⑤",
                    f"{total_dias_mo:.2f} días × {apu.get('costo_dia_mo', 0):,.0f}/día",
                    "Mat. + MO",
                    f"{apu.get('pct_herramienta', 0.05)*100:.1f}% × Mano de Obra",
                    f"{apu.get('pct_aui', 0.30)*100:.0f}% × Costo Directo",
                    f"{apu.get('iva', 0.19)*100:.0f}% × Utilidad ({apu.get('pct_util', 0.05)*100:.0f}% CD)",
                    "CD + ⑦ + ⑧ + ⑨",
                ],
                f"Importe [{mon}]": [
                    f"{c_excav:,.0f}", f"{c_cem:,.0f}", f"{c_ace:,.0f}", f"{c_are:,.0f}", f"{c_gra:,.0f}",
                    f"{total_mat:,.0f}",
                    f"{costo_mo:,.0f}",
                    f"{costo_directo:,.0f}",
                    f"{herramienta:,.0f}",
                    f"{aiu:,.0f}",
                    f"{iva:,.0f}",
                    f"{total_proyecto:,.0f}",
                ]
            }
            st.dataframe(pd.DataFrame(data_zap_apu), use_container_width=True, hide_index=True)
            st.info(f"**Diferencia Gran Total vs Materiales:** {total_proyecto - total_mat:,.0f} {mon} = Mano de Obra ({costo_mo:,.0f}) + Herramienta ({herramienta:,.0f}) + AIU ({aiu:,.0f}) + IVA ({iva:,.0f})")

            
            # Excel APU Export
            output_excel = io.BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                df_export = pd.DataFrame({
                    "Item": ["Excavación", "Cemento", "Acero", "Arena", "Grava", "Mano de Obra"],
                    "Cantidad": [vol_excavacion, bultos_zap, peso_total_acero_zap, vol_arena_z, vol_grava_z, total_dias_mo],
                    "Unidad": [25000, apu['cemento'], apu['acero'], apu['arena'], apu['grava'], apu.get('costo_dia_mo', 0)]
                })
                df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
                df_export.to_excel(writer, index=False, sheet_name='APU')
                workbook = writer.book
                worksheet = writer.sheets['APU']
                money_fmt = workbook.add_format({'num_format': '#,##0.00'})
                bold = workbook.add_format({'bold': True})
                worksheet.set_column('A:A', 25)
                worksheet.set_column('B:D', 15, money_fmt)
                
                row = len(df_export) + 1
                worksheet.write(row, 0, "Costo Directo (CD)", bold)
                worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
                row += 1
                worksheet.write(row, 0, "Herramienta Menor", bold)
                worksheet.write_formula(row, 3, f'=D7*{apu.get("pct_herramienta", 0.05)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "A.I.U", bold)
                worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui", 0.30)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "IVA s/ Utilidad", bold)
                _row_util = row - 1
                worksheet.write_formula(row, 3, f'=D{_row_util}*{apu.get("iva", 0.19)}', money_fmt)
                row += 1
                worksheet.write(row, 0, "TOTAL PRESUPUESTO", bold)
                worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
                
            output_excel.seek(0)
            st.download_button(label="Descargar Presupuesto Excel (.xlsx)", data=output_excel, 
                               file_name=f"APU_Zapata_{B_use}x{L_use}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("Ve a la página 'APU Mercado'para cargar los costos base de agregados, acero y cemento y que tu presupuesto se genere automáticamente.")

    # 
    # B5 — CUADRO DE MANDO MULTI-ZAPATA + C6 ASENTAMIENTO DIFERENCIAL
    # 
    st.markdown("---")
    
    #  FI-2: Factor de Interacción Biaxial (IFC) 
    st.write("#### Factor de Interacción (Carga y Momento)")
    # FI-2 φPn corregido: columna (no zapata) — NSR-10 C.10 / ACI 318 §22.4.2
    _Ag_col_mm2 = (c1_col * 10) * (c2_col * 10)     # área bruta columna en mm²
    _As_col_tot = max(0.01 * _Ag_col_mm2, 400)       # cuantía mín. 1% → mm²
    phi_Pn = 0.65 * 0.80 * (0.85 * fc_basico * (_Ag_col_mm2 - _As_col_tot)
             + fy_basico * _As_col_tot) / 1000.0     # kN
    phi_Mnx = phi_f * As_req_B * 1e-4 * fy_basico * d_z_m * 1000
    phi_Mny = phi_f * As_req_L * 1e-4 * fy_basico * d_z_m * 1000
    
    ifc = (P_ult/max(phi_Pn,1)) + (abs(M_ult_B)/max(phi_Mnx,1)) + (abs(M_ult_L)/max(phi_Mny,1))
    
    _col_ifc1, _col_ifc2 = st.columns(2)
    _col_ifc1.metric("IFC (P + Mx + My)", f"{ifc:.3f}", delta="≤ 1.0", delta_color="inverse")
    if ifc > 1.0:
        _col_ifc2.error(f" IFC = {ifc:.3f} > 1.0 — Diseño INSEGURO para compresión biaxial.")
    else:
        _col_ifc2.success(f" IFC = {ifc:.3f} ≤ 1.0")

    st.markdown("##  Cuadro de Mando del Proyecto — Registro de Zapatas")
    st.caption("Registra varias zapatas del proyecto para comparación y verificación de asentamiento diferencial.")

    #  Inicializar lista en SessionState 
    if "zapatas_proyecto" not in st.session_state:
        st.session_state["zapatas_proyecto"] = []

    #  Botón para agregar la zapata actual 
    _nombre_zap = st.text_input("Nombre / Etiqueta de esta zapata (ej. Z-1, Eje A-3)",
                                value=f"Z-{len(st.session_state['zapatas_proyecto'])+1}",
                                key="z_nombre_registro")
    _col_btn1, _col_btn2 = st.columns([1, 3])
    if _col_btn1.button(" Agregar al proyecto", type="primary"):
        # FI-5 y FI-1: Actualización con nuevos campos
        _is_ok = (ok_1way_B and ok_1way_L and ok_punz and disc_B > 0 and disc_L > 0 and ifc <= 1.0)
        _entrada = {
            "Nombre":     _nombre_zap,
            "Tipo":       "Aislada",
            "Pos. Col.":  pos_col_iso if 'pos_col_iso' in dir() else "—",
            "B [m]":      round(B_use, 2),
            "L [m]":      round(L_use, 2),
            "H [cm]":     int(H_zap),
            "Pu [kN]":    round(P_ult, 1),
            "IFC":        round(ifc, 3) if 'ifc' in dir() else 0.0,
            "DCR_pun":    round(Vu_punz/phi_Vc_punz, 3) if 'phi_Vc_punz' in dir() and phi_Vc_punz > 0 else 9.99,
            "q_adm [kPa]":round(q_adm_z, 1),
            "qu_max [kPa]":round(qu_max, 2),
            "qu_min [kPa]":round(qu_min, 2),
            "FSv":        round(FS_volc, 2) if 'FS_volc' in dir() else "—",
            "FSd":        round(FS_desl, 2) if 'FS_desl' in dir() else "—",
            "As_B [cm²]": round(As_req_B, 1),
            "As_L [cm²]": round(As_req_L, 1),
            "Estado": " OK" if _is_ok else " Revisar",
        }
        # evitar duplicado por nombre
        _nombres_existentes = [z["Nombre"] for z in st.session_state["zapatas_proyecto"]]
        if _nombre_zap in _nombres_existentes:
            _idx = _nombres_existentes.index(_nombre_zap)
            st.session_state["zapatas_proyecto"][_idx] = _entrada
            st.success(f"Zapata **{_nombre_zap}** actualizada en el registro.")
            _save_state_zap()  # Persistencia v6.1
        else:
            st.session_state["zapatas_proyecto"].append(_entrada)
            st.success(f"Zapata **{_nombre_zap}** agregada al proyecto ({len(st.session_state['zapatas_proyecto'])} total).")
            _save_state_zap()  # Persistencia v6.1

    if _col_btn2.button(" Limpiar registro"):
        st.session_state["zapatas_proyecto"] = []
        st.info("Registro limpiado.")

    #  Tabla comparativa 
    if st.session_state["zapatas_proyecto"]:
        _df_proy = pd.DataFrame(st.session_state["zapatas_proyecto"])
        st.markdown("###  Comparativo de Zapatas del Proyecto")
        st.dataframe(_df_proy, use_container_width=True, hide_index=True)

        # Exportar a Excel
        _buf_proy = io.BytesIO()
        with pd.ExcelWriter(_buf_proy, engine="xlsxwriter") as _wr:
            _df_proy.to_excel(_wr, index=False, sheet_name="Zapatas")
            _wb2 = _wr.book
            _ws2 = _wr.sheets["Zapatas"]
            _hdr_fmt = _wb2.add_format({"bold": True, "bg_color": "#1a3a5c", "font_color": "white"})
            for _ci, _col in enumerate(_df_proy.columns):
                _ws2.write(0, _ci, _col, _hdr_fmt)
                _ws2.set_column(_ci, _ci, max(12, len(str(_col))+2))
        _buf_proy.seek(0)
        st.download_button("Exportar tabla de proyecto a Excel",
                           data=_buf_proy,
                           file_name="Cuadro_Mando_Zapatas.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        #  C6 — ASENTAMIENTO DIFERENCIAL 
        if len(st.session_state["zapatas_proyecto"]) >= 2:
            st.markdown("---")
            st.markdown("###  C6 — Verificación de Asentamiento Diferencial (NSR-10 H.3)")
            _asen_vals = [z["Asentam. [mm]"] for z in st.session_state["zapatas_proyecto"]]
            _noms_vals = [z["Nombre"]        for z in st.session_state["zapatas_proyecto"]]

            # Separación entre zapatas (para calcular límite angular)
            _sep_viga = st.number_input("Separación entre zapatas adyacentes L_viga [m] (luz entre ejes)",
                                        1.0, 30.0, 6.0, 0.5, key="z_sep_viga")
            _delta_adm = _sep_viga * 1000 / 300.0   # mm  (NSR-10 H.3: δ ≤ L/300)

            _filas_d = []
            for _i in range(len(_asen_vals) - 1):
                _delta_s = abs(_asen_vals[_i+1] - _asen_vals[_i])
                _ok_d    = _delta_s <= _delta_adm
                _filas_d.append({
                    "Par": f"{_noms_vals[_i]} — {_noms_vals[_i+1]}",
                    "s₁ [mm]": f"{_asen_vals[_i]:.1f}",
                    "s₂ [mm]": f"{_asen_vals[_i+1]:.1f}",
                    "Δs [mm]": f"{_delta_s:.1f}",
                    f"Máx. NSR-10 [mm] (L/{300})": f"{_delta_adm:.1f}",
                    "Estado": " OK" if _ok_d else f" Excede ({_delta_s:.1f} > {_delta_adm:.1f})"
                })

            st.dataframe(pd.DataFrame(_filas_d), use_container_width=True, hide_index=True)

            _max_delta = max(abs(_asen_vals[i+1] - _asen_vals[i]) for i in range(len(_asen_vals)-1))
            if _max_delta <= _delta_adm:
                st.success(f"**Asentamiento diferencial máximo = {_max_delta:.1f} mm ≤ {_delta_adm:.1f} mm** (L_viga/{300}) — NSR-10 H.3 satisfecho.")
            else:
                st.error(f"**Δs_max = {_max_delta:.1f} mm > {_delta_adm:.1f} mm** — El asentamiento diferencial excede el límite NSR-10 H.3. "
                         "Considere zapatas más rígidas, aumentar B×L o usar vigas de cimentación.")
            st.caption(f"Nota: los asentamientos se leen del campo 'Asentam. [mm]'registrado para cada zapata. "
                       "Si el valor es 0, calcule el asentamiento en la pestaña de Geotecnia antes de agregar la zapata al registro.")
    else:
        st.info("ℹ Agrega al menos una zapata con el botón ** Agregar al proyecto** para activar el cuadro de mando.")


# 
# FUNCIÓN PRINCIPAL: render_medianera()
# Tipos 2 (Medianera/Lindero), 3 (Esquina) y 4 (Con Viga de Amarre)
# 
def render_medianera(norma, fc, fy, rebar_dict, def_idx, phi_v, phi_f, tipo_zap):
    """Flujo completo de diseño para zapata medianera, de esquina y con viga de amarre."""
    sn = _get_strap_norm(norma)
    _es_medianera = "2." in tipo_zap
    _es_esquina   = "3." in tipo_zap
    _es_strap     = "4." in tipo_zap or _es_medianera or _es_esquina

    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#1a2a4a,#0d1b35);border-radius:10px;
    padding:16px 20px;margin-bottom:16px;border-left:4px solid #4db8ff;">
    <h4 style="color:#4db8ff;margin:0 0 6px 0;"> {tipo_zap}</h4>
    <p style="color:#aad4ff;margin:0;font-size:13px;">
    Norma activa: <strong>{norma}</strong> — {sn['nombre']} ({sn['art']})
    </p>
    </div>
    """, unsafe_allow_html=True)

    #  TABS DE TRABAJO 
    tab_geo, tab_viga, tab_3d, tab_dxf, tab_docx = st.tabs([
        " Geometría y Presiones",
        f" {sn['nombre']}",
        " Vista 3D",
        " DXF",
        " Memoria",
    ])

    # 
    # TAB 1: GEOMETRÍA Y DISTRIBUCIÓN DE PRESIONES
    # 
    with tab_geo:
        st.subheader("Geometría y Cargas")
        cA, cB, cC = st.columns(3)
        with cA:
            st.markdown("##### Zapata Exterior (lindero)")
            Be = st.number_input("Ancho B_ext [m]", 0.5, 10.0,
                                 st.session_state.get("zm_Be", 2.0), 0.1, key="zm_Be")
            Le = st.number_input("Largo L_ext [m]", 0.5, 10.0,
                                 st.session_state.get("zm_Le", 2.0), 0.1, key="zm_Le")
            He = st.number_input("Espesor H_ext [cm]", 20.0, 120.0,
                                 st.session_state.get("zm_He", 50.0), 5.0, key="zm_He")
            c1e = st.number_input("Col. c1_ext [cm]", 15.0, 80.0, 40.0, 5.0, key="zm_c1e")
            c2e = st.number_input("Col. c2_ext [cm]", 15.0, 80.0, 40.0, 5.0, key="zm_c2e")
            Pe  = st.number_input("Carga Axial P_ext [kN]", 50.0, 5000.0, 500.0, 50.0, key="zm_Pe")
            Me_B = st.number_input("Momento M_ext dir.B [kN·m]", 0.0, 1000.0, 0.0, 10.0, key="zm_MeB")
            Me_L = st.number_input("Momento M_ext dir.L [kN·m]", 0.0, 1000.0, 0.0, 10.0, key="zm_MeL")

        with cB:
            st.markdown("##### Zapata Interior (campo libre)")
            Bi = st.number_input("Ancho B_int [m]", 0.5, 10.0,
                                 st.session_state.get("zm_Bi", 2.0), 0.1, key="zm_Bi")
            Li = st.number_input("Largo L_int [m]", 0.5, 10.0,
                                 st.session_state.get("zm_Li", 2.0), 0.1, key="zm_Li")
            Hi = st.number_input("Espesor H_int [cm]", 20.0, 120.0,
                                 st.session_state.get("zm_Hi", 50.0), 5.0, key="zm_Hi")
            c1i = st.number_input("Col. c1_int [cm]", 15.0, 80.0, 40.0, 5.0, key="zm_c1i")
            c2i = st.number_input("Col. c2_int [cm]", 15.0, 80.0, 40.0, 5.0, key="zm_c2i")
            Pi  = st.number_input("Carga Axial P_int [kN]", 50.0, 5000.0, 700.0, 50.0, key="zm_Pi")
            Mi_B = st.number_input("Momento M_int dir.B [kN·m]", 0.0, 1000.0, 0.0, 10.0, key="zm_MiB")
            Mi_L = st.number_input("Momento M_int dir.L [kN·m]", 0.0, 1000.0, 0.0, 10.0, key="zm_MiL")

        with cC:
            st.markdown("##### Sistema y Suelo")
            q_adm_m = st.number_input("q_adm [kPa]", 50.0, 600.0, 150.0, 10.0, key="zm_qadm")
            L_libre = st.number_input("Distancia libre entre zapatas [m]", 1.0, 20.0,
                                      st.session_state.get("zm_Llibre", 4.0), 0.5, key="zm_Llibre")
            dist_lindero = st.number_input("Dist. borde col. ext → lindero [cm]", 0.0, 200.0, 0.0, 5.0,
                                           key="zm_distlind",
                                           help="Si = 0: la columna coincide con el lindero. "
                                                "La excentricidad se calculará automáticamente.")
            gamma_c_m = st.number_input("γ_concreto [kN/m³]", 22.0, 26.0, 24.0, 0.5, key="zm_gam")

        with st.expander("Optimizador automático B×L (Medianera)", expanded=False):
            st.caption("Busca la combinación mínima Be×Le / Bi×Li que cumpla q_adm con carga dada.")
            _q_adm_opt = st.number_input("q_adm objetivo [kPa]", 50.0, 600.0,
                                         st.session_state.get("zm_qadm", 150.0), 10.0, key="opt_qadm_zm")
            _Pe_opt = st.number_input("Carga exterior P_ext [kN]", 50.0, 5000.0, 500.0, 50.0, key="opt_Pe_zm")
            _Pi_opt = st.number_input("Carga interior P_int [kN]", 50.0, 5000.0, 700.0, 50.0, key="opt_Pi_zm")
            _paso   = st.selectbox("Paso de búsqueda [m]", [0.05, 0.10, 0.25], index=1, key="opt_paso_zm")
            if st.button("Optimizar dimensiones", key="opt_btn_zm"):
                import itertools as _it
                import numpy as _np_opt
                _dims = _np_opt.arange(0.5, 6.0 + _paso, _paso)
                _best = None
                for _b, _l in _it.product(_dims, _dims):
                    _qu = _Pe_opt / (_b * _l)
                    if _qu <= _q_adm_opt:
                        _area = _b * _l
                        if _best is None or _area < _best[2]:
                            _best = (_b, _l, _area, _qu)
                if _best:
                    st.success(f"Dimensión óptima exterior: B = **{_best[0]:.2f} m** × L = **{_best[1]:.2f} m**")
                    st.metric("q_neto", f"{_best[3]:.1f} kPa", delta=f"{_q_adm_opt - _best[3]:.1f} kPa de reserva")
                    _b2, _l2 = _best[0], _best[1]
                    _qu2 = _Pi_opt / (_b2 * _l2)
                    _ok2 = _qu2 <= _q_adm_opt
                    st.info(f"Interior (misma dimensión {_b2:.2f}×{_l2:.2f}m): qu = {_qu2:.1f} kPa  {'' if _ok2 else ' aumentar'}")
                else:
                    st.error("No se encontró dimensión dentro del rango 0.5–6.0 m. Revisa q_adm o las cargas.")

        st.divider()
        #  CÁLCULO DISTRIBUCIONES DE PRESIONES 
        st.subheader("Distribución de Presiones")
        col_pe, col_pi = st.columns(2)

        def _mostrar_presiones(label, P, M_B, M_L, B, L, c1_cm, c2_cm, prefix, q_adm):
            qu_max, qu_min, tipo_dist, e_B, e_L, Ix, Iy, A = calcular_distribucion_presiones(
                P, M_B, M_L, B, L)
            meyerhof_data = None
            if qu_min < 0:
                B_p, L_p, A_p = calcular_area_efectiva_meyerhof(B, L, e_B, e_L)
                qu_eff = P / A_p
                meyerhof_data = {"B_prima": B_p, "L_prima": L_p, "qu_eff": qu_eff}
                st.warning(f" **{label}**: qu_min = {qu_min:.1f} kPa < 0 — "
                           f"Área Efectiva Meyerhof → B'={B_p:.2f}m × L'={L_p:.2f}m → **qu_eff = {qu_eff:.1f} kPa**")
                qu_avg_uso = qu_eff
            elif tipo_dist == "triangular":
                # EM-2 CORREGIDO: para distribución triangular el valor representativo
                # es qu_max * (2/3), no el promedio aritmético (que subestima la carga)
                # Referencia: Bowles (1996) §8.4 — área del triángulo de presiones
                qu_avg_uso = qu_max * (2.0 / 3.0)
            else:
                qu_avg_uso = (qu_max + qu_min) / 2

            ok_q = qu_avg_uso <= q_adm
            semaf = "" if ok_q else ""
            _d_lbl = {"trapezoidal":"Trapezoidal","triangular":"Triangular (parcial levant.)","con_tension":"Con tensión"}
            st.markdown(f"""
| Parámetro | Valor |
|---|---|
| **qu_max** | `{qu_max:.1f} kPa` |
| **qu_min** | `{qu_min:.1f} kPa` |
| **Distribución** | {_d_lbl.get(tipo_dist, tipo_dist)} |
| **e_B** | `{e_B:.3f} m` |
| **e_L** | `{e_L:.3f} m` |
| **q_adm** | `{q_adm:.1f} kPa` |
| **Verificación** | {semaf} `qu_avg={qu_avg_uso:.1f}` ≤ `{q_adm:.1f}` kPa |
""")
            # Mapa de calor
            fig_hm, _ten = render_mapa_calor_presiones(B, L, P, M_B, M_L, Ix, Iy, A,
                                                       qu_max, qu_min, c1_cm, c2_cm,
                                                       titulo=f"qu(x,y) — {label}", q_adm=q_adm)
            st.plotly_chart(fig_hm, use_container_width=True)
        if abs(qu_max - qu_min) < 1.0:
            st.info(
                " **Distribución uniforme** — qu = "
                f"{qu_max:.1f} kPa constante en toda la huella. "
                "El gráfico 3D cobra valor cuando ingresas momentos Mu ≠ 0 "
                "para ver la variación de presiones."
            )
        else:
            if qu_min >= 0:
                st.warning(
                    f" **Distribución trapezoidal** — "
                    f"qu_max = {qu_max:.1f} kPa | qu_min = {qu_min:.1f} kPa. "
                    f"Excentricidad activa en la zapata."
                )
            else:
                st.error(
                    f" **Levantamiento parcial** — qu_min = {qu_min:.1f} kPa < 0. "
                    "Parte de la zapata no está en contacto con el suelo."
                )
            return {"qu_max": qu_max, "qu_min": qu_min, "tipo": tipo_dist,
                    "e_B": e_B, "e_L": e_L, "Ix": Ix, "Iy": Iy, "A": A,
                    "qu_avg": qu_avg_uso, "meyerhof": meyerhof_data}

        with col_pe:
            st.markdown("**Zapata Exterior**")
            dist_e = _mostrar_presiones("Zapata Exterior", Pe, Me_B, Me_L,
                                        Be, Le, c1e, c2e, "ext", q_adm_m)
        with col_pi:
            st.markdown("**Zapata Interior**")
            dist_i = _mostrar_presiones("Zapata Interior", Pi, Mi_B, Mi_L,
                                        Bi, Li, c1i, c2i, "int", q_adm_m)

        # Excentricidad y reacción de la viga de amarre
        # R_strap = M_exc / L_centro_a_centro (equilibrio de momentos)
        # Para medianera: la excentricidad es e_B respecto al centroide de la zapata
        e_exc = dist_e["e_B"] + (Be/2 - c1e/200)  # excentricidad total desde col al centroide
        # ── Excentricidad y reacción de viga(s) de amarre ──────────────────
        # Para medianera (tipo 2): excentricidad solo en dir. B → 1 viga
        # Para esquinera (tipo 3): excentricidades en dir. B y L → 2 vigas perp.
        R_strap_kN   = Pe * e_exc / L_libre if L_libre > 0 else 0.0

        if _es_esquina:
            # Segunda excentricidad: desde centroide columna ext al centroide zapata en dir. L
            e_exc_L     = dist_e["e_L"] + (Le/2 - c2e/200)
            L_libre_L   = st.number_input(
                _t("Distancia libre entre zapatas dir. L [m]", "Free span between footings dir. L [m]"),
                1.0, 20.0, st.session_state.get("zm_Llibre_L", L_libre), 0.5, key="zm_Llibre_L"
            )
            R_strap_L_kN = Pe * e_exc_L / L_libre_L if L_libre_L > 0 else 0.0
            st.info(
                f"**Esquinera — Dos vigas de amarre:**"
                f"- Viga dir. B: R_strap_B = {Pe:.1f} × {e_exc:.3f} / {L_libre:.1f} "
                f"= **{R_strap_kN:.2f} kN**"
                f"- Viga dir. L: R_strap_L = {Pe:.1f} × {e_exc_L:.3f} / {L_libre_L:.1f} "
                f"= **{R_strap_L_kN:.2f} kN**"
            )
        else:
            R_strap_L_kN = 0.0
            L_libre_L    = L_libre
            st.info(
                f"**Reaccion en viga de amarre:** R_strap = Pe x e / L_libre = "
                f"{Pe:.1f} x {e_exc:.3f} / {L_libre:.1f} = **{R_strap_kN:.2f} kN**"
            )

        # Guardar datos de estado para otras tabs
        st.session_state["zm_dist_e"]       = dist_e
        st.session_state["zm_dist_i"]       = dist_i
        st.session_state["zm_R_strap"]      = R_strap_kN
        st.session_state["zm_R_strap_L"]    = R_strap_L_kN
        st.session_state["zm_Llibre_L"]     = L_libre_L

    # 
    # TAB 2: DISEÑO DE LA VIGA DE AMARRE
    # 
    with tab_viga:
        st.subheader(f"Diseño de {sn['nombre']} — {norma}")
        st.markdown(f"> Artículo de referencia: **{sn['art']}** | "
                    f"Dimensión mínima: L / {sn['h_factor']} | "
                    f"Fuerza axial sísmica: {sn['F_ax_pct']*100:.1f}% × ΣP")

        cV1, cV2, cV3 = st.columns(3)
        with cV1:
            bv = st.number_input(f"Ancho {sn['nombre']} b [cm]", 20.0, 60.0, 30.0, 5.0, key="zm_bv")
            hv = st.number_input(f"Alto {sn['nombre']} h [cm]", 30.0, 120.0, 50.0, 5.0, key="zm_hv")
            recub_v = st.number_input("Recubrimiento [cm]", 2.5, 8.0, 4.0, 0.5, key="zm_recubv")
        with cV2:
            bar_long_v = st.selectbox("Varilla longitudinal:", list(rebar_dict.keys()),
                                      index=min(def_idx+1, len(rebar_dict)-1), key="zm_bar_long_v")
            bar_est_v  = st.selectbox("Varilla estribo:", list(rebar_dict.keys()),
                                      index=0, key="zm_bar_est_v")
        with cV3:
            # MC-1 CORREGIDO: R_strap se propaga automáticamente desde la pestaña Geom.
            _R_auto = float(st.session_state.get("zm_R_strap", 50.0))
            st.metric("R_strap calculado [kN]", f"{_R_auto:.2f}",
                      help="Calculado automáticamente en 'Geometría y Presiones'. "
                           "Si cambias los inputs, recalculará al volver a esa pestaña.")
            R_strap_kN_ui = st.number_input(
                "Sobreescribir R_strap [kN] (opcional):",
                0.1, 3000.0, _R_auto, 10.0, key="zm_R_strap_ui",
                help="Deja el valor calculado o ajusta manualmente si es necesario.")
            P_total_kN = st.number_input("ΣP columnas (para F_ax sísmica) [kN]:",
                                          100.0, 10000.0,
                                          float(Pe if 'Pe' in dir() else 1000.0),
                                          100.0, key="zm_Ptotal")

        # Calcular V(x) y M(x) — EC-2: ahora retorna 8 valores (incluye _adv_Rint)
        x_arr, V_arr, M_arr, M_max, V_max, pp_viga, R_int, _adv_Rint = calcular_viga_amarre(
            R_strap_kN_ui, L_libre, bv, hv, gamma_c_m)
        # EC-2: mostrar advertencia de levantamiento si R_int < 0
        if _adv_Rint:
            st.error(
                f" **Levantamiento en zapata interior: R_int = {R_int:.2f} kN < 0** — "
                "El peso propio de la viga supera la reacción de amarre. "
                "Aumento de R_strap o reducción de L_libre requeridos."
            )
        # MC-4: aviso diferenciado para NTC-EM (México) con F_ax = 5%
        if "NTC" in norma:
            st.info(" **NTC-EM México:** La fuerza axial sísmica requerida es **F_ax = 5% × ΣP**, "
                    "el doble del estándar ACI/NSR-10 (2.5%). Revisar el diseño.")
        F_ax_sism = sn["F_ax_pct"] * P_total_kN

        # Diseñar la viga — EC-3: capturar dict de error si d_v <= 0
        dis = disenar_viga_amarre(
            M_max, V_max, F_ax_sism, bv, hv, fc, fy, L_libre, P_total_kN, norma,
            recub_v=recub_v, bar_long=bar_long_v, bar_est=bar_est_v, rebar_dict=rebar_dict)
        if dis.get("_error"):
            st.error(dis["_msg"])
            st.stop()

        # Tabla de verificaciones — semáforo
        st.subheader(" /  Tabla de Verificaciones")
        _v_rows = [
            ("Dimensión mínima h", f"{hv:.0f} cm", f"L/{sn['h_factor']} = {dis['h_min_cm']:.1f} cm", dis["ok_hmin"]),
            ("Flexión sup. φMn+", f"{dis['phi_Mn_sup']:.1f} kN·m", f"M_max = {M_max:.2f} kN·m", dis["ok_flex"]),
            ("Cortante φVc", f"{dis['phi_Vc']:.1f} kN", f"V_max = {V_max:.2f} kN", dis["ok_cort"]),
            ("F. axial sísmica (req.)", f"{F_ax_sism:.1f} kN", f"{sn['F_ax_pct']*100:.1f}% × ΣP = {dis['F_ax_req']:.1f} kN", dis["ok_Fax"]),
        ]
        _cols = st.columns([2.5, 2, 2.5, 0.8])
        _cols[0].markdown("**Verificación**"); _cols[1].markdown("**Calculado**")
        _cols[2].markdown("**Límite**"); _cols[3].markdown("**OK**")
        for _vr, _vc, _vl, _vo in _v_rows:
            _c = st.columns([2.5, 2, 2.5, 0.8])
            _ic = "" if _vo else ""
            _c[0].markdown(_vr); _c[1].markdown(f"`{_vc}`")
            _c[2].markdown(f"`{_vl}`"); _c[3].markdown(f"**{_ic}**")

        st.divider()
        # Tabla de armado
        st.subheader(f"Armado de {sn['nombre']}")
        _a_col1, _a_col2 = st.columns(2)
        with _a_col1:
            st.markdown(f"""
| Elemento | Resultado |
|---|---|
| **Barras superiores** | **{dis['n_sup']} × {dis['bar_long']}** = {dis['As_sup']:.2f} cm² |
| **Barras inferiores** | **{dis['n_inf']} × {dis['bar_long']}** = {dis['As_inf']:.2f} cm² |
| **d efectivo** | {dis['d_v']:.1f} cm |
| **Estribos** | {dis['bar_est']} **@ {dis['s_crit_cm']:.0f} cm** (zona crít. {dis['L_crit_cm']:.0f} cm) |
| **Estribos (resto)** | {dis['bar_est']} **@ {dis['s_est_cm']:.0f} cm** |
""")
        with _a_col2:
            st.markdown(f"""
| Referencia Normativa | |
|---|---|
| **Norma** | {norma} |
| **Artículo** | {sn['art']} |
| **h_min** | L/{sn['h_factor']} = {dis['h_min_cm']:.1f} cm |
| **F_ax mínima** | {sn['F_ax_pct']*100:.1f}% × ΣP = {dis['F_ax_req']:.1f} kN |
""")

        # Diagrama V(x) / M(x)
        st.subheader("Diagrama de Esfuerzos Internos")
        fig_vm = render_diagrama_VM(x_arr, V_arr, M_arr,
                                    R_strap_kN_ui, R_int, pp_viga, L_libre)
        st.plotly_chart(fig_vm, use_container_width=True)

        # Guardar resultados intermedios en sesión
        st.session_state["zm_dis"] = dis
        st.session_state["zm_fig_vm"] = fig_vm
        st.session_state["zm_viga_d"] = {
            "b_cm": bv, "h_cm": hv, "b": bv, "h": hv, "L_libre_m": L_libre,
            "recub_cm": recub_v, "recub": recub_v, "fc": fc, "fy": fy,
        }


        # ── Segunda viga de amarre (solo si esquinera tipo 3) ─────────
        if _es_esquina:
            st.divider()
            st.subheader(_t(
                f"Segunda Viga de Amarre — dir. L ({sn['nombre']})",
                f"Second Strap Beam — dir. L ({sn['nombre']})"
            ))
            _R_strap_L = float(st.session_state.get("zm_R_strap_L", 0.0))
            _L_libre_L = float(st.session_state.get("zm_Llibre_L", L_libre))
            st.metric("R_strap dir. L [kN]", f"{_R_strap_L:.2f}",
                      help=_t("Calculado automaticamente en 'Geometria y Presiones'.",
                               "Auto-calculated in 'Geometry & Pressures' tab."))

            _cVL1, _cVL2 = st.columns(2)
            with _cVL1:
                bv_L = st.number_input(
                    _t("Ancho viga L b [cm]","Strap beam dir. L width [cm]"),
                    20.0, 60.0, bv, 5.0, key="zm_bvL")
                hv_L = st.number_input(
                    _t("Alto viga L h [cm]","Strap beam dir. L height [cm]"),
                    30.0, 120.0, hv, 5.0, key="zm_hvL")
            with _cVL2:
                bar_long_vL = st.selectbox(
                    _t("Varilla longitudinal dir. L","Long. rebar dir. L"),
                    list(rebar_dict.keys()),
                    index=min(def_idx+1, len(rebar_dict)-1), key="zm_bar_long_vL")
                bar_est_vL = st.selectbox(
                    _t("Varilla estribo dir. L","Stirrup rebar dir. L"),
                    list(rebar_dict.keys()), index=0, key="zm_bar_est_vL")

            _xL, _VL, _ML, _MmaxL, _VmaxL, _ppL, _RintL, _advL = calcular_viga_amarre(
                _R_strap_L, _L_libre_L, bv_L, hv_L, gamma_c_m)

            if _advL:
                st.error(
                    f"Levantamiento en zapata interior dir. L: "
                    f"R_int = {_RintL:.2f} kN < 0. Aumentar R_strap_L o reducir L_libre_L."
                )

            _P_total_L = float(st.session_state.get("zm_Ptotal", Pe + Pi))
            _F_ax_L    = sn["F_ax_pct"] * _P_total_L

            dis_L = disenar_viga_amarre(
                _MmaxL, _VmaxL, _F_ax_L, bv_L, hv_L, fc, fy,
                _L_libre_L, _P_total_L, norma,
                recub_v=recub_v, bar_long=bar_long_vL,
                bar_est=bar_est_vL, rebar_dict=rebar_dict)

            if dis_L.get("_error"):
                st.error(dis_L["_msg"])
            else:
                st.markdown(f"""
| Elemento | Resultado |
|---|---|
| **Barras superiores** | **{dis_L['n_sup']} x {dis_L['bar_long']}** = {dis_L['As_sup']:.2f} cm² |
| **Barras inferiores** | **{dis_L['n_inf']} x {dis_L['bar_long']}** = {dis_L['As_inf']:.2f} cm² |
| **d efectivo** | {dis_L['d_v']:.1f} cm |
| **Estribos zona critica** | {dis_L['bar_est']} @ {dis_L['s_crit_cm']:.0f} cm (L_crit = {dis_L['L_crit_cm']:.0f} cm) |
| **Estribos resto** | {dis_L['bar_est']} @ {dis_L['s_est_cm']:.0f} cm |
""")
                _ok_rows_L = [
                    (dis_L["ok_hmin"], f"h_min = {dis_L['h_min_cm']:.1f} cm"),
                    (dis_L["ok_flex"], f"phi_Mn = {dis_L['phi_Mn_sup']:.1f} kNm >= M_max = {_MmaxL:.1f} kNm"),
                    (dis_L["ok_cort"], f"phi_Vc = {dis_L['phi_Vc']:.1f} kN >= V_max = {_VmaxL:.1f} kN"),
                    (dis_L["ok_Fax"], f"F_ax = {_F_ax_L:.1f} kN"),
                ]
                for _ok, _txt in _ok_rows_L:
                    if _ok:
                        st.success(f"CUMPLE — {_txt}")
                    else:
                        st.error(f"NO CUMPLE — {_txt}. Aumente b, h o varilla.")

                # Diagrama V/M viga L
                _fig_vmL = render_diagrama_VM(
                    _xL, _VL, _ML, _R_strap_L, _RintL, _ppL, _L_libre_L)
                st.plotly_chart(_fig_vmL, use_container_width=True)

                # Persistir
                st.session_state["zm_dis_L"]   = dis_L
                st.session_state["zm_viga_dL"] = {
                    "b_cm": bv_L, "h_cm": hv_L, "b": bv_L, "h": hv_L,
                    "L_libre_m": _L_libre_L, "recub_cm": recub_v,
                    "recub": recub_v, "fc": fc, "fy": fy,
                }


    # 
    # TAB 3: VISTA 3D PLOTLY
    # 
    with tab_3d:
        st.subheader("Vista 3D del Sistema")
        dis_3d  = st.session_state.get("zm_dis")
        viga_3d = st.session_state.get("zm_viga_d")
        if dis_3d and viga_3d:
            zap_ext_3d = {"x0": 0.0, "y0": 0.0, "B": Be, "L": Le, "H": He, "c1": c1e, "c2": c2e}
            if _es_esquina:
                # Zapata interior en diagonal (esquina) — dos vigas perpendiculares
                _L_libre_L3d = float(st.session_state.get("zm_Llibre_L", L_libre))
                zap_int_3d = {
                    "x0": Be + L_libre, "y0": Le + _L_libre_L3d,
                    "B": Bi, "L": Li, "H": Hi, "c1": c1i, "c2": c2i
                }
                st.info(_t(
                    "Vista 3D — Esquinera: zapata exterior en origen, interior en diagonal. "
                    "Las dos vigas de amarre se muestran en el modelo.",
                    "3D View — Corner footing: exterior at origin, interior at diagonal. "
                    "Both strap beams are shown."
                ))
            else:
                zap_int_3d = {"x0": Be + L_libre, "y0": 0.0, "B": Bi, "L": Li, "H": Hi, "c1": c1i, "c2": c2i}
            fig_3d = render_3d_sistema(zap_ext_3d, zap_int_3d, viga_3d, dis_3d)
            st.plotly_chart(fig_3d, use_container_width=True)
            st.caption(_t(
                "Arrastra para rotar | Doble clic para resetear vista | Rueda para zoom",
                "Drag to rotate | Double-click to reset | Scroll to zoom"
            ))
        else:
            st.info("ℹ Completa primero la pestaña **'Geometría y Presiones'** y el **'Diseño de Viga'** para generar la vista 3D.")

    # 
    # TAB 4: EXPORTACIÓN DXF
    # 
    with tab_dxf:
        st.subheader("Plano DXF ICONTEC - Sistema y Detalles")
        dis_dxf  = st.session_state.get("zm_dis")
        viga_dxf = st.session_state.get("zm_viga_d")
        if dis_dxf and viga_dxf:
            papel_opc_s = {"Carta (21x28cm)": (21.6,27.9,"CARTA"), "Oficio (21x33cm)": (21.6,33.0,"OFICIO"), "Pliego (70x100cm)": (70.7,100.0,"PLIEGO")}
            papel_sel_s = st.selectbox("Formato de Papel (Sistema):", list(papel_opc_s.keys()), key="zs_papel")
            W_P_S, H_P_S, LBL_P_S = papel_opc_s[papel_sel_s]
            
            zap_ext_dxf = {"B_cm": Be*100, "L_cm": Le*100, "H_cm": He}
            zap_int_dxf = {"B_cm": Bi*100, "L_cm": Li*100, "H_cm": Hi}
            
            _buf_dxf = None
            if st.button("Generar DXF ICONTEC del Sistema", key="zm_gen_dxf", type="primary"):
                with st.spinner("Generando planos DXF..."):
                    try:
                        doc_dxf = generar_dxf_sistema(zap_ext_dxf, zap_int_dxf, viga_dxf, dis_dxf, norma, fc, fy, W_P_S, H_P_S, LBL_P_S)
                        import os, tempfile
                        with tempfile.NamedTemporaryFile(suffix='.dxf', delete=False) as tmp:
                            tmp_path = tmp.name
                        doc_dxf.saveas(tmp_path)
                        with open(tmp_path, 'rb') as f:
                            _buf_dxf = f.read()
                        os.unlink(tmp_path)
                        st.success("DXF generado con Estándar Diamante N4.")
                    except Exception as _e_dxf:
                        st.error(f"Error generando DXF: {_e_dxf}")
            
            if _buf_dxf:
                if _puede_exportar:
                    st.download_button(
                        f"Descargar DXF — {sn['nombre']}",
                        data=_buf_dxf,
                        file_name=f"zapata_sistema_{sn['nombre'].replace(' ','_').lower()}.dxf",
                        mime="application/dxf",
                        key="zm_dl_dxf"
                    )
                else:
                    st.button("Descargar DXF Sistema", disabled=True,
                              help="Requiere licencia Pro o Admin.")

        else:
            st.warning(" Completa primero la pestaña **Geometría y Presiones** y el diseño de la viga.")
        
        # ── EXPORTACIÓN IFC AUTÓNOMA STRAP ──
        st.markdown("---")
        st.markdown("####  Exportar Modelo BIM (.ifc)")
        st.caption("Implementación Nativa IFC4 para Sistema Medianero/Correa")
        
        def _make_ifc_sistema(z_ext, z_int, vig, fc, fy, norma, proy_nombre):
            try:
                import ifcopenshell
                import ifcopenshell.api
            except ImportError:
                return None
            
            O = ifcopenshell.file(schema="IFC4")

            
            def p4(x,y,z): return O.createIfcCartesianPoint((float(x),float(y),float(z)))
            def ax2(pt, z_dir=(0.,0.,1.), x_dir=(1.,0.,0.)):
                return O.createIfcAxis2Placement3D(pt, O.createIfcDirection(z_dir), O.createIfcDirection(x_dir))
            
            proj = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcProject", name=proy_nombre)
            ctx  = ifcopenshell.api.run("context.add_context", O, context_type="Model")
            body = ifcopenshell.api.run("context.add_context", O, context_type="Model", context_identifier="Body", target_view="MODEL_VIEW", parent=ctx)
            site = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcSite", name="Sitio")
            bldg = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcBuilding", name="Fundacion")
            _u = ifcopenshell.api.run("unit.add_si_unit", O, unit_type="LENGTHUNIT"); ifcopenshell.api.run("unit.assign_unit", O, units=[_u])
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=proj, products=[site])
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=site, products=[bldg])
            
            # Z.EXT
            z1 = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcFooting", name="Zapata Exterior")
            z1.PredefinedType = "PAD_FOOTING"
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=bldg, products=[z1])
            prf1 = O.createIfcRectangleProfileDef('AREA', None, ax2(p4(0,0,0)), z_ext['B_cm']/100, z_ext['L_cm']/100)
            ex1  = O.createIfcExtrudedAreaSolid(prf1, ax2(p4(0,0,0)), O.createIfcDirection((0.,0.,1.)), z_ext['H_cm']/100)
            ifcopenshell.api.run("geometry.assign_representation", O, product=z1, representation=O.createIfcShapeRepresentation(body, 'Body', 'SweptSolid', [ex1]))
            ifcopenshell.api.run("geometry.edit_object_placement", O, product=z1)
            
            # Z.INT
            z2 = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcFooting", name="Zapata Interior")
            z2.PredefinedType = "PAD_FOOTING"
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=bldg, products=[z2])
            prf2 = O.createIfcRectangleProfileDef('AREA', None, ax2(p4(0,0,0)), z_int['B_cm']/100, z_int['L_cm']/100)
            sep_m = vig["L_libre_m"]
            ex2  = O.createIfcExtrudedAreaSolid(prf2, ax2(p4(z_ext['B_cm']/100 + sep_m,0,0)), O.createIfcDirection((0.,0.,1.)), z_int['H_cm']/100)
            ifcopenshell.api.run("geometry.assign_representation", O, product=z2, representation=O.createIfcShapeRepresentation(body, 'Body', 'SweptSolid', [ex2]))
            ifcopenshell.api.run("geometry.edit_object_placement", O, product=z2)
            
            # VIGA
            bv = vig['b_cm']/100; hv = vig['h_cm']/100
            vg = ifcopenshell.api.run("root.create_entity", O, ifc_class="IfcBeam", name="Viga de Amarre")
            ifcopenshell.api.run("aggregate.assign_object", O, relating_object=bldg, products=[vg])
            prfv = O.createIfcRectangleProfileDef('AREA', None, ax2(p4(0,0,0)), bv, hv)
            exv  = O.createIfcExtrudedAreaSolid(prfv, ax2(p4(z_ext['B_cm']/100,0,max(z_ext['H_cm']/100, z_int['H_cm']/100)), x_dir=(0.,1.,0.), z_dir=(1.,0.,0.)), O.createIfcDirection((0.,0.,1.)), sep_m)
            ifcopenshell.api.run("geometry.assign_representation", O, product=vg, representation=O.createIfcShapeRepresentation(body, 'Body', 'SweptSolid', [exv]))
            ifcopenshell.api.run("geometry.edit_object_placement", O, product=vg)
            
            import os, tempfile
            with tempfile.NamedTemporaryFile(suffix='.ifc', delete=False) as tmp:
                tmp_path = tmp.name
            O.write(tmp_path)
            with open(tmp_path, 'rb') as f:
                b_ifc = f.read()
            os.unlink(tmp_path)
            return b_ifc

        _buf_ifc_s = None
        if dis_dxf and viga_dxf:
            if st.button("Generar BIM IFC del sistema", key="zm_btn_ifc"):
                with st.spinner("Construyendo geometrías 3D IFC4..."):
                    _buf_ifc_s = _make_ifc_sistema(zap_ext_dxf, zap_int_dxf, viga_dxf, fc, fy, norma, "Sistema Medianero")
                    if _buf_ifc_s:
                        st.success("IFC4 nativo generado con éxito.")
                    else:
                        st.error("Error al generar IFC.")
                        
            if _buf_ifc_s:
                if _puede_exportar:
                    st.download_button("Descargar Modelo BIM (IFC)", data=_buf_ifc_s,
                                       file_name="Sistema_Medianero.ifc", mime="application/x-step")
                else:
                    st.button("Descargar Modelo BIM (IFC)", disabled=True,
                              help="Requiere licencia Pro o Admin.")

    # ═══════════════════════════════════════════════════════════════════
    # PASO 3 — DISEÑO ESTRUCTURAL ZAPATAS DEL SISTEMA (NSR-10 C.15 / ACI 318 Ch.13)
    # ═══════════════════════════════════════════════════════════════════
    with st.expander(
        _t("Diseño Estructural de Zapatas del Sistema (NSR-10 C.15 / ACI 318 Ch.13)",
           "Footing Structural Design (NSR-10 C.15 / ACI 318 Ch.13)"),
        expanded=True
    ):
        _col_dis1, _col_dis2 = st.columns(2)
        with _col_dis1:
            _recub_zap = st.number_input(
                _t("Recubrimiento al suelo [cm]", "Cover to soil [cm]"),
                3.0, 10.0, 7.5, 0.5, key="zm_recub_zap"
            )
            _bar_zap = st.selectbox(
                _t("Varilla zapata", "Footing rebar"),
                list(rebar_dict.keys()), index=def_idx, key="zm_bar_zap"
            )
        with _col_dis2:
            _db_zap  = rebar_dict[_bar_zap]["db"]    # mm
            _Ab_zap  = rebar_dict[_bar_zap]["area"]  # cm²
            st.markdown(
                f"db = **{_db_zap:.1f} mm** | Ab = **{_Ab_zap:.2f} cm²**"
            )

        # ─── Función de diseño estructural de una zapata ──────────────
        def _diseno_zapata_sistema(B, L, H, c1, c2, P_svc, M_B_svc, M_L_svc,
                                   recub, db_mm, Ab_cm2, fc_mpa, fy_mpa,
                                   phi_v_loc, phi_f_loc, q_adm_loc, etiqueta):
            import math as _m
            d_cm = H - recub - db_mm / 20.0
            if d_cm <= 0:
                return {"error": True,
                        "msg": f"Peralte efectivo d = {d_cm:.1f} cm <= 0 en {etiqueta}. "
                               "Aumente H o reduzca recubrimiento."}
            A_use = B * L
            P_ult = 1.4 * P_svc
            qu_u  = P_ult / A_use

            # Cortante 1D
            av_B = max((B - c1/100.0)/2.0 - d_cm/100.0, 0.001)
            av_L = max((L - c2/100.0)/2.0 - d_cm/100.0, 0.001)
            Vu_1B      = qu_u * av_B * L
            Vu_1L      = qu_u * av_L * B
            phi_Vc_B   = phi_v_loc * 0.17 * _m.sqrt(fc_mpa) * (L*100) * d_cm / 1000
            phi_Vc_L   = phi_v_loc * 0.17 * _m.sqrt(fc_mpa) * (B*100) * d_cm / 1000
            ok_1B = phi_Vc_B >= Vu_1B
            ok_1L = phi_Vc_L >= Vu_1L

            # Punzonamiento 2D
            b_pun = c1/100.0 + d_cm/100.0
            l_pun = c2/100.0 + d_cm/100.0
            bo_cm = 2*(b_pun + l_pun)*100
            Vu_p  = P_ult - qu_u * b_pun * l_pun
            beta_c = max(c1,c2)/min(c1,c2) if min(c1,c2)>0 else 1.0
            vc1 = 0.17*(1+2/beta_c)*_m.sqrt(fc_mpa)
            vc2 = 0.083*(2+40*d_cm/bo_cm)*_m.sqrt(fc_mpa)
            vc3 = 0.33*_m.sqrt(fc_mpa)
            vc_p      = min(vc1, vc2, vc3)
            phi_Vc_p  = phi_v_loc * vc_p * bo_cm * d_cm / 1000
            ok_pun    = phi_Vc_p >= Vu_p

            # Flexión biaxial
            vB  = max((B - c1/100.0)/2.0, 0.001)
            vL  = max((L - c2/100.0)/2.0, 0.001)
            Mu_B = qu_u * vB**2 / 2.0 * L
            Mu_L = qu_u * vL**2 / 2.0 * B

            def _as_flex(Mu_kNm, b_cm_ancho, d_cm_ef):
                if Mu_kNm <= 0:
                    return max(0.0018*b_cm_ancho*10*d_cm_ef*10/100, 0.0), True
                Rn   = Mu_kNm*1e6 / (phi_f_loc * b_cm_ancho*10 * (d_cm_ef*10)**2)
                disc = max(1 - 2*Rn/(0.85*fc_mpa), 0.0)
                rho  = max(0.85*fc_mpa/fy_mpa*(1-_m.sqrt(disc)), 0.0018)
                return rho * b_cm_ancho*10 * d_cm_ef*10 / 100, disc > 0

            As_B, ok_dB = _as_flex(Mu_B, L*100, d_cm)
            As_L, ok_dL = _as_flex(Mu_L, B*100, d_cm)
            n_B  = max(2, math.ceil(As_B / Ab_cm2))
            n_L  = max(2, math.ceil(As_L / Ab_cm2))
            sep_B = (L*100 - 2*recub)/(n_B-1) if n_B>1 else (L*100-2*recub)
            sep_L = (B*100 - 2*recub)/(n_L-1) if n_L>1 else (B*100-2*recub)

            # Anclaje
            ldh_mm  = 0.24*fy_mpa/_m.sqrt(fc_mpa)*db_mm
            ldh_cm  = max(15.0, 8*db_mm/10.0, ldh_mm/10.0)
            Ldisp_B = (B-c1/100.0)/2.0*100.0 - recub
            Ldisp_L = (L-c2/100.0)/2.0*100.0 - recub
            Ldisp   = min(Ldisp_B, Ldisp_L)
            ok_ldh  = Ldisp >= ldh_cm
            _fac_ld = 3.5 if (recub>=db_mm/10.0 and sep_B>=2*db_mm/10.0 and sep_L>=2*db_mm/10.0) else 1.7
            ld_cm   = max(30.0, fy_mpa/(_fac_ld*_m.sqrt(fc_mpa))*db_mm/10.0)
            ok_ld   = Ldisp >= ld_cm

            return dict(
                error=False, etiqueta=etiqueta,
                d_cm=d_cm, A_use=A_use, qu_u=qu_u, P_ult=P_ult,
                Vu_1B=Vu_1B, phi_Vc_B=phi_Vc_B, ok_1B=ok_1B,
                Vu_1L=Vu_1L, phi_Vc_L=phi_Vc_L, ok_1L=ok_1L,
                Vu_p=Vu_p, phi_Vc_p=phi_Vc_p, ok_pun=ok_pun,
                bo_cm=bo_cm, vc_p=vc_p,
                Mu_B=Mu_B, As_B=As_B, n_B=n_B, sep_B=sep_B, ok_dB=ok_dB,
                Mu_L=Mu_L, As_L=As_L, n_L=n_L, sep_L=sep_L, ok_dL=ok_dL,
                ldh_cm=ldh_cm, ld_cm=ld_cm,
                Ldisp_B=Ldisp_B, Ldisp_L=Ldisp_L, Ldisp=Ldisp,
                ok_ldh=ok_ldh, ok_ld=ok_ld, _fac_ld=_fac_ld,
            )

        _dis_e = _diseno_zapata_sistema(
            Be, Le, He, c1e, c2e, Pe, Me_B, Me_L,
            _recub_zap, _db_zap, _Ab_zap, fc, fy, phi_v, phi_f,
            q_adm_m, _t("Exterior (Lindero)", "Exterior (Property Line)")
        )
        _dis_i = _diseno_zapata_sistema(
            Bi, Li, Hi, c1i, c2i, Pi, Mi_B, Mi_L,
            _recub_zap, _db_zap, _Ab_zap, fc, fy, phi_v, phi_f,
            q_adm_m, _t("Interior (Campo Libre)", "Interior (Free Field)")
        )

        # ─── Tabla de verificaciones por zapata ──────────────────────
        def _tabla_zap(d):
            if d.get("error"):
                st.error(d["msg"]); return
            st.markdown(
                f"### {d['etiqueta']}  |  "
                f"d = {d['d_cm']:.1f} cm  |  Pu = {d['P_ult']:.1f} kN  |  "
                f"qu_u = {d['qu_u']:.1f} kPa"
            )
            rows = [
                (_t("Cortante 1D dir. B (C.11.3)","1-Way Shear dir. B"),
                 f"phi_Vc={d['phi_Vc_B']:.1f} kN", f"Vu={d['Vu_1B']:.1f} kN",  d["ok_1B"]),
                (_t("Cortante 1D dir. L (C.11.3)","1-Way Shear dir. L"),
                 f"phi_Vc={d['phi_Vc_L']:.1f} kN", f"Vu={d['Vu_1L']:.1f} kN",  d["ok_1L"]),
                (_t("Punzonamiento (C.11.12)","Punching Shear (C.11.12)"),
                 f"phi_Vc_p={d['phi_Vc_p']:.1f} kN  bo={d['bo_cm']:.0f} cm",
                 f"Vu_p={d['Vu_p']:.1f} kN", d["ok_pun"]),
                (_t("Flexion dir. B (C.10.2)","Flexure dir. B (C.10.2)"),
                 f"As={d['As_B']:.2f} cm2 -> {d['n_B']} {_bar_zap} @ {d['sep_B']:.1f} cm",
                 f"Mu={d['Mu_B']:.1f} kNm", d["ok_dB"]),
                (_t("Flexion dir. L (C.10.2)","Flexure dir. L (C.10.2)"),
                 f"As={d['As_L']:.2f} cm2 -> {d['n_L']} {_bar_zap} @ {d['sep_L']:.1f} cm",
                 f"Mu={d['Mu_L']:.1f} kNm", d["ok_dL"]),
                (_t("Anclaje Gancho 90 (C.12.5)","Hook 90 deg (C.12.5)"),
                 f"L_disp={d['Ldisp']:.1f} cm", f"ldh req={d['ldh_cm']:.1f} cm", d["ok_ldh"]),
                (_t("Desarrollo Recto ld (C.12.2)","Straight Dev. ld (C.12.2)"),
                 f"L_disp={d['Ldisp']:.1f} cm",
                 f"ld req={d['ld_cm']:.1f} cm (fac={d['_fac_ld']:.1f})", d["ok_ld"]),
            ]
            hdr = st.columns([3.2, 3, 2.5, 1])
            hdr[0].markdown("**Verificacion**"); hdr[1].markdown("**Capacidad**")
            hdr[2].markdown("**Demanda**");      hdr[3].markdown("**OK**")
            for vr, vc, vl, vo in rows:
                c = st.columns([3.2, 3, 2.5, 1])
                c[0].markdown(vr); c[1].markdown(f"`{vc}`"); c[2].markdown(f"`{vl}`")
                c[3].markdown("**:green[SI]**" if vo else "**:red[NO]**")
            fallos = [r[0] for r in rows if not r[3]]
            if fallos:
                st.error(_t(f"NO CUMPLE — {len(fallos)} verificacion(es): {', '.join(fallos)}. Aumente H o B/L.",
                            f"FAIL — {len(fallos)} check(s): {', '.join(fallos)}. Increase H or B/L."))
            else:
                st.success(_t(f"CUMPLE — {d['etiqueta']}: todas las verificaciones satisfechas.",
                              f"PASS — {d['etiqueta']}: all checks passed."))
            st.divider()

        _tabla_zap(_dis_e)
        _tabla_zap(_dis_i)

        # Persistir para Memoria
        st.session_state["zm_dis_e"]    = _dis_e
        st.session_state["zm_dis_i"]    = _dis_i
        st.session_state["zm_bar_zap"]  = _bar_zap
        st.session_state["zm_recub_zap"]= _recub_zap

    # ═══════════════════════════════════════════════════════════════════
    # TAB 5 — MEMORIA DE CÁLCULO DOCX (SISTEMA MEDIANERO / ESQUINERO)
    # ═══════════════════════════════════════════════════════════════════
    with tab_docx:
        st.subheader(_t("Memoria de Calculo — Sistema Medianero/Esquinero",
                        "Calculation Report — Strap/Corner Footing System"))
        empresa, proyecto, ingeniero, elab, rev, apb, logo = _get_identity()

        if st.button(_t("Generar Memoria DOCX del Sistema", "Generate System DOCX Report"),
                     type="primary", key="zm_btn_docx"):
            _d_e  = st.session_state.get("zm_dis_e", {})
            _d_i  = st.session_state.get("zm_dis_i", {})
            _d_v  = st.session_state.get("zm_dis",   {})
            _bz   = st.session_state.get("zm_bar_zap",   _bar_zap)
            _rz   = st.session_state.get("zm_recub_zap", _recub_zap)

            if not _d_v:
                st.warning(_t(
                    "AVISO: Completa el Diseno de Viga de Amarre antes de generar la memoria.",
                    "WARNING: Complete Strap Beam Design before generating the report."
                ))
            else:
                from docx import Document as _DocZM
                import datetime as _dtZM

                doc_m = _DocZM()
                doc_m.add_heading(
                    f"MEMORIA ESTRUCTURAL — SISTEMA {tipo_zap.upper()} ({norma})", 0)
                doc_m.add_paragraph(
                    f"Empresa: {empresa}  |  Proyecto: {proyecto}  |  "
                    f"Ingeniero: {ingeniero}  |  "
                    f"Fecha: {_dtZM.date.today().strftime('%Y-%m-%d')}")
                doc_m.add_paragraph(f"Elaboro: {elab}  |  Reviso: {rev}  |  Aprobo: {apb}")

                # 1. Parametros generales
                doc_m.add_heading("1. PARAMETROS GENERALES", level=1)
                doc_m.add_paragraph(
                    f"Norma: {norma}  |  fc = {fc:.0f} MPa  |  fy = {fy:.0f} MPa"
                    f"q_adm = {q_adm_m:.1f} kPa  |  L_libre = {L_libre:.2f} m"
                    f"Recubrimiento = {_rz:.1f} cm  |  Varilla = {_bz}")

                # 2. Geometria y cargas
                doc_m.add_heading("2. GEOMETRIA Y CARGAS", level=1)
                doc_m.add_paragraph(
                    f"ZAPATA EXTERIOR:"
                    f"  B={Be:.2f}m  L={Le:.2f}m  H={He:.0f}cm  "
                    f"Col={c1e:.0f}x{c2e:.0f}cm  P={Pe:.1f}kN"
                    f"ZAPATA INTERIOR:"
                    f"  B={Bi:.2f}m  L={Li:.2f}m  H={Hi:.0f}cm  "
                    f"Col={c1i:.0f}x{c2i:.0f}cm  P={Pi:.1f}kN")

                # 3. Viga de amarre
                doc_m.add_heading(f"3. {sn['nombre'].upper()} ({sn['art']})", level=1)
                if not _d_v.get("_error"):
                    _Mmax_doc = float(st.session_state.get("zm_Mmax", 0))
                    _Vmax_doc = float(st.session_state.get("zm_Vmax", 0))
                    doc_m.add_paragraph(
                        f"Seccion: {bv:.0f} x {hv:.0f} cm  |  "
                        f"R_strap = {R_strap_kN:.2f} kN"
                        f"Barras sup: {_d_v.get('n_sup','?')} x {_d_v.get('bar_long','?')}  |  "
                        f"Barras inf: {_d_v.get('n_inf','?')} x {_d_v.get('bar_long','?')}"
                        f"Estribos: {_d_v.get('bar_est','?')} @ {_d_v.get('s_est_cm',0):.0f} cm  "
                        f"(zona critica: @ {_d_v.get('s_crit_cm',0):.0f} cm)")

                # Segunda viga (solo esquinera tipo 3)
                if _es_esquina:
                    _d_vL  = st.session_state.get("zm_dis_L", {})
                    _vigdL = st.session_state.get("zm_viga_dL", {})
                    doc_m.add_heading(f"3b. SEGUNDA VIGA DE AMARRE — DIR. L ({sn['art']})", level=1)
                    if _d_vL and not _d_vL.get("_error"):
                        _bvL_d = _vigdL.get("b_cm", "?")
                        _hvL_d = _vigdL.get("h_cm", "?")
                        _LlibL = _vigdL.get("L_libre_m", "?")
                        doc_m.add_paragraph(
                            f"Seccion: {_bvL_d} x {_hvL_d} cm  |  "
                            f"R_strap_L = {st.session_state.get('zm_R_strap_L', 0):.2f} kN  |  "
                            f"L_libre_L = {_LlibL} m"
                            f"Barras sup: {_d_vL.get('n_sup','?')} x {_d_vL.get('bar_long','?')}  |  "
                            f"Barras inf: {_d_vL.get('n_inf','?')} x {_d_vL.get('bar_long','?')}"
                            f"Estribos: {_d_vL.get('bar_est','?')} @ {_d_vL.get('s_est_cm',0):.0f} cm  "
                            f"(zona critica: @ {_d_vL.get('s_crit_cm',0):.0f} cm)")
                    else:
                        doc_m.add_paragraph(
                            "AVISO: Ejecuta el diseno de la segunda viga (dir. L) "
                            "en la UI para obtener resultados.")

                # 4. Diseno estructural por zapata
                for _etq, _dd in [("ZAPATA EXTERIOR", _d_e), ("ZAPATA INTERIOR", _d_i)]:
                    doc_m.add_heading(f"4. DISENO ESTRUCTURAL — {_etq}", level=1)
                    if not _dd or _dd.get("error"):
                        doc_m.add_paragraph(
                            "AVISO: Ejecuta el Diseno Estructural en la UI para obtener resultados.")
                        continue
                    doc_m.add_paragraph(
                        f"d efectivo = {_dd['d_cm']:.1f} cm  |  "
                        f"Pu = {_dd['P_ult']:.1f} kN  |  qu_u = {_dd['qu_u']:.1f} kPa"
                        f"CORTANTE UNIDIRECCIONAL (NSR-10 C.11.3):"
                        f"  Dir B: phi_Vc_1way_B={_dd['phi_Vc_B']:.1f}kN  Vu={_dd['Vu_1B']:.1f}kN  "
                        f"{'OK - CUMPLE' if _dd['ok_1B'] else 'NO CUMPLE'}"
                        f"  Dir L: phi_Vc_1way_B={_dd['phi_Vc_L']:.1f}kN  Vu={_dd['Vu_1L']:.1f}kN  "
                        f"{'OK - CUMPLE' if _dd['ok_1L'] else 'NO CUMPLE'}"
                        f"PUNZONAMIENTO (NSR-10 C.11.12):"
                        f"  bo={_dd['bo_cm']:.0f}cm  phiVc_p={_dd['phi_Vc_p']:.1f}kN  "
                        f"Vu_p={_dd['Vu_p']:.1f}kN  "
                        f"{'OK - CUMPLE' if _dd['ok_pun'] else 'NO CUMPLE'}"
                        f"FLEXION (NSR-10 C.10.2):"
                        f"  Dir B: As={_dd['As_B']:.2f}cm2 -> "
                        f"{_dd['n_B']} {_bz} @ {_dd['sep_B']:.1f}cm  Mu={_dd['Mu_B']:.1f}kNm"
                        f"  Dir L: As={_dd['As_L']:.2f}cm2 -> "
                        f"{_dd['n_L']} {_bz} @ {_dd['sep_L']:.1f}cm  Mu={_dd['Mu_L']:.1f}kNm"
                        f"ANCLAJE (NSR-10 C.12):"
                        f"  Gancho 90: ldh={_dd['ldh_cm']:.1f}cm  L_disp={_dd['Ldisp']:.1f}cm  "
                        f"{'OK - CUMPLE' if _dd['ok_ldh'] else 'NO CUMPLE'}"
                        f"  Recto ld:  ld={_dd['ld_cm']:.1f}cm (fac={_dd['_fac_ld']:.1f})  "
                        f"L_disp={_dd['Ldisp']:.1f}cm  "
                        f"{'OK - CUMPLE' if _dd['ok_ld'] else 'NO CUMPLE'}")

                # 5. APU simplificado
                doc_m.add_heading("5. CANTIDADES BASICAS DE MATERIALES", level=1)
                _vol_e = Be * Le * He/100
                _vol_i = Bi * Li * Hi/100
                _vol_v = bv/100 * hv/100 * L_libre
                _vol_tot = _vol_e + _vol_i + _vol_v
                doc_m.add_paragraph(
                    f"Vol. concreto zapata exterior: {_vol_e:.3f} m3"
                    f"Vol. concreto zapata interior: {_vol_i:.3f} m3"
                    f"Vol. concreto viga de amarre:  {_vol_v:.3f} m3"
                    f"TOTAL concreto sistema:        {_vol_tot:.3f} m3")

                # Firma
                doc_m.add_heading("6. FIRMA RESPONSABLE", level=1)
                doc_m.add_paragraph(
                    '_' * 45 + '\n' + ingeniero + '\nMatricula Profesional: _______________'
                )

                _zm_io = io.BytesIO()
                doc_m.save(_zm_io)
                _zm_io.seek(0)
                st.success(_t("Memoria generada correctamente.", "Report generated successfully."))
                st.download_button(
                    _t("Descargar Memoria DOCX Sistema", "Download System DOCX Report"),
                    data=_zm_io,
                    file_name=f"Memoria_{tipo_zap[:2].strip().replace('.','')}_Sistema.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="zm_dl_docx",
                    disabled=False
                )

                # ── PDF nativo sistema medianero/esquinero ────────────────
                if st.button(
                    _t("Generar PDF Sistema", "Generate System PDF"),
                    key="zm_btn_pdf", use_container_width=True,
                    disabled=not _puede_exportar
                ):
                    try:
                        from reportlab.lib.pagesizes import A4
                        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                        from reportlab.lib.units import cm, mm
                        from reportlab.lib import colors
                        from reportlab.platypus import (
                            SimpleDocTemplate, Paragraph, Spacer, Table,
                            TableStyle, HRFlowable
                        )
                        import io as _io_pdf_zm
                        import datetime as _dt_pdf_zm

                        _d_e_pdf  = st.session_state.get("zm_dis_e", {})
                        _d_i_pdf  = st.session_state.get("zm_dis_i", {})
                        _d_v_pdf  = st.session_state.get("zm_dis",   {})
                        _bz_pdf   = st.session_state.get("zm_bar_zap", "?")
                        _empresa_z, _proy_z, _ing_z, _elab_z, _rev_z, _apb_z, _logo_z = _get_identity()

                        _pdf_zm   = _io_pdf_zm.BytesIO()
                        _doc_zm   = SimpleDocTemplate(_pdf_zm, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2*cm)
                        _st       = getSampleStyleSheet()
                        _h1z      = ParagraphStyle("H1Z", parent=_st["Heading1"],
                            fontSize=11, textColor=colors.HexColor("#1b3a5c"),
                            spaceBefore=8, spaceAfter=4)
                        _boz      = ParagraphStyle("BodyZ", parent=_st["Normal"],
                            fontSize=9, leading=13)
                        _capz     = ParagraphStyle("CapZ", parent=_st["Normal"],
                            fontSize=8, textColor=colors.grey)

                        _story_z = []

                        # Portada
                        _story_z.append(Paragraph(
                            f"MEMORIA ESTRUCTURAL — SISTEMA {tipo_zap.upper()} ({norma})",
                            ParagraphStyle("TZ", parent=_st["Title"], fontSize=13,
                                           textColor=colors.HexColor("#1b3a5c"))))
                        _story_z.append(Spacer(1, 3*mm))
                        _tpz = Table([
                            ["Empresa:", _empresa_z, "Proyecto:", _proy_z],
                            ["Ingeniero:", _ing_z, "Fecha:", _dt_pdf_zm.date.today().strftime("%Y-%m-%d")],
                        ], colWidths=[2.5*cm, 5.5*cm, 2.5*cm, 5.5*cm])
                        _tpz.setStyle(TableStyle([
                            ("FONTSIZE",(0,0),(-1,-1),8),
                            ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),
                            ("FONTNAME",(2,0),(2,-1),"Helvetica-Bold"),
                            ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#bbcdd8")),
                            ("ROWBACKGROUNDS",(0,0),(-1,-1),[colors.HexColor("#f0f4f8"),colors.white]),
                            ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
                        ]))
                        _story_z.append(_tpz)
                        _story_z.append(Spacer(1, 4*mm))

                        # Parámetros generales
                        _story_z.append(Paragraph("1. PARÁMETROS GENERALES", _h1z))
                        _story_z.append(HRFlowable(width="100%",thickness=0.5,
                            color=colors.HexColor("#bbcdd8"),spaceAfter=3))
                        _story_z.append(Paragraph(
                            f"Norma: {norma}  |  fc = {fc:.0f} MPa  |  fy = {fy:.0f} MPa  |  "
                            f"q_adm = {q_adm_m:.1f} kPa  |  L_libre = {L_libre:.2f} m  |  "
                            f"Varilla = {_bz_pdf}", _boz))
                        _story_z.append(Spacer(1, 3*mm))

                        # Zapatas
                        for _lbl_z, _B_z, _L_z, _H_z, _c1_z, _c2_z, _P_z in [
                            ("EXTERIOR", Be, Le, He, c1e, c2e, Pe),
                            ("INTERIOR", Bi, Li, Hi, c1i, c2i, Pi),
                        ]:
                            _story_z.append(Paragraph(
                                f"2. ZAPATA {_lbl_z}  (B={_B_z:.2f}m × L={_L_z:.2f}m × H={_H_z:.0f}cm  "
                                f"Col={_c1_z:.0f}×{_c2_z:.0f}cm  P={_P_z:.1f}kN)", _h1z))
                            _story_z.append(HRFlowable(width="100%",thickness=0.5,
                                color=colors.HexColor("#bbcdd8"),spaceAfter=3))
                            _dd_pdf = _d_e_pdf if _lbl_z=="EXTERIOR" else _d_i_pdf
                            if _dd_pdf and not _dd_pdf.get("error"):
                                _rows_pdf = [
                                    ["Verificación","Capacidad","Demanda","Estado"],
                                    ["Cortante 1D dir. B",
                                     f"φVc={_dd_pdf['phi_Vc_B']:.1f}kN",
                                     f"Vu={_dd_pdf['Vu_1B']:.1f}kN",
                                     "CUMPLE" if _dd_pdf["ok_1B"] else "NO CUMPLE"],
                                    ["Cortante 1D dir. L",
                                     f"φVc={_dd_pdf['phi_Vc_L']:.1f}kN",
                                     f"Vu={_dd_pdf['Vu_1L']:.1f}kN",
                                     "CUMPLE" if _dd_pdf["ok_1L"] else "NO CUMPLE"],
                                    ["Punzonamiento bo",
                                     f"φVc_p={_dd_pdf['phi_Vc_p']:.1f}kN",
                                     f"Vu_p={_dd_pdf['Vu_p']:.1f}kN",
                                     "CUMPLE" if _dd_pdf["ok_pun"] else "NO CUMPLE"],
                                    ["Flexión dir. B",
                                     f"As={_dd_pdf['As_B']:.2f}cm²",
                                     f"Mu={_dd_pdf['Mu_B']:.1f}kNm",
                                     "CUMPLE" if _dd_pdf["ok_dB"] else "NO CUMPLE"],
                                    ["Flexión dir. L",
                                     f"As={_dd_pdf['As_L']:.2f}cm²",
                                     f"Mu={_dd_pdf['Mu_L']:.1f}kNm",
                                     "CUMPLE" if _dd_pdf["ok_dL"] else "NO CUMPLE"],
                                    ["Anclaje gancho 90°",
                                     f"L_disp={_dd_pdf['Ldisp']:.1f}cm",
                                     f"ldh={_dd_pdf['ldh_cm']:.1f}cm",
                                     "CUMPLE" if _dd_pdf["ok_ldh"] else "NO CUMPLE"],
                                    ["Desarrollo recto ld",
                                     f"L_disp={_dd_pdf['Ldisp']:.1f}cm",
                                     f"ld={_dd_pdf['ld_cm']:.1f}cm",
                                     "CUMPLE" if _dd_pdf["ok_ld"] else "NO CUMPLE"],
                                ]
                                _tbl_dp = Table(_rows_pdf,
                                    colWidths=[4.5*cm,3.5*cm,3*cm,2.5*cm])
                                _tbl_dp.setStyle(TableStyle([
                                    ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
                                    ("FONTSIZE",(0,0),(-1,-1),8),
                                    ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#1b3a5c")),
                                    ("TEXTCOLOR",(0,0),(-1,0),colors.white),
                                    ("ROWBACKGROUNDS",(0,1),(-1,-1),
                                     [colors.white,colors.HexColor("#f0f4f8")]),
                                    ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#bbcdd8")),
                                    ("TOPPADDING",(0,0),(-1,-1),2),
                                    ("BOTTOMPADDING",(0,0),(-1,-1),2),
                                    *[("TEXTCOLOR",(3,i+1),(3,i+1),
                                       colors.HexColor("#2d6a4f") if _rows_pdf[i+1][3]=="CUMPLE"
                                       else colors.HexColor("#c0392b"))
                                      for i in range(len(_rows_pdf)-1)],
                                    *[("FONTNAME",(3,i+1),(3,i+1),"Helvetica-Bold")
                                      for i in range(len(_rows_pdf)-1)],
                                ]))
                                _story_z.append(_tbl_dp)
                            else:
                                _story_z.append(Paragraph(
                                    "AVISO: Ejecutar diseño estructural en la UI.", _boz))
                            _story_z.append(Spacer(1, 3*mm))

                        # Viga de amarre
                        _story_z.append(Paragraph(
                            f"3. {sn['nombre'].upper()} ({sn['art']})", _h1z))
                        _story_z.append(HRFlowable(width="100%",thickness=0.5,
                            color=colors.HexColor("#bbcdd8"),spaceAfter=3))
                        if _d_v_pdf and not _d_v_pdf.get("_error"):
                            _story_z.append(Paragraph(
                                f"Sección: {bv:.0f}×{hv:.0f}cm  |  R_strap={R_strap_kN:.2f}kN  |  "
                                f"Barras sup: {_d_v_pdf.get('n_sup','?')}×{_d_v_pdf.get('bar_long','?')}  |  "
                                f"Estribos: {_d_v_pdf.get('bar_est','?')} @ "
                                f"{_d_v_pdf.get('s_est_cm',0):.0f}cm", _boz))

                        # Firma
                        _story_z.append(Spacer(1, 8*mm))
                        _story_z.append(Paragraph("FIRMA RESPONSABLE", _h1z))
                        _story_z.append(HRFlowable(width="100%",thickness=0.5,
                            color=colors.HexColor("#bbcdd8"),spaceAfter=6))
                        _story_z.append(Spacer(1, 8*mm))
                        _fz = Table([
                            [f"{'_'*28}", "", f"{'_'*28}"],
                            [_ing_z, "", _apb_z],
                            ["Diseñó / Calculó", "", "Aprobó"],
                        ], colWidths=[7*cm,2*cm,7*cm])
                        _fz.setStyle(TableStyle([
                            ("FONTSIZE",(0,0),(-1,-1),8),
                            ("ALIGN",(0,0),(-1,-1),"CENTER"),
                            ("FONTNAME",(0,1),(0,1),"Helvetica-Bold"),
                            ("FONTNAME",(2,1),(2,1),"Helvetica-Bold"),
                        ]))
                        _story_z.append(_fz)
                        _story_z.append(Spacer(1,3*mm))
                        _story_z.append(Paragraph(
                            f"Konte/StructoPro  |  {norma}  |  "
                            f"{_dt_pdf_zm.date.today().strftime('%Y-%m-%d')}", _capz))

                        _doc_zm.build(_story_z)
                        _pdf_zm.seek(0)
                        st.success(_t("PDF sistema generado.", "System PDF generated."))
                        st.download_button(
                            _t("Descargar PDF Sistema", "Download System PDF"),
                            data=_pdf_zm,
                            file_name=f"Memoria_{tipo_zap[:2].strip()}_Sistema.pdf",
                            mime="application/pdf",
                            key="zm_dl_pdf"
                        )
                    except ImportError:
                        st.error(_t(
                            "ReportLab no instalado. Agrega 'reportlab' a requirements.txt.",
                            "ReportLab not installed. Add 'reportlab' to requirements.txt."
                        ))
                    except Exception as _e_pdf_zm:
                        st.error(f"Error generando PDF sistema: {_e_pdf_zm}")


# ── Separación máxima por norma ─────────────────────────────────────────────
_SEP_MAX_LIMITS = {
    'NSR-10 Colombia': 45.0, 'ACI 318-25 EE.UU.': 45.0,
    'ACI 318-19 EE.UU.': 45.0, 'ACI 318-14 EE.UU.': 45.0,
    'E.060 Peru': 45.0, 'NEC-SE-HM Ecuador': 45.0,
    'NTC-EM México': 40.0, 'COVENIN 1753-2006 Venezuela': 40.0,
    'NB 1225001-2020 Bolivia': 45.0, 'CIRSOC 201-2025 Argentina': 40.0,
}
_sep_max_abs = _SEP_MAX_LIMITS.get(norma_sel, 45.0)
sep_max = min(3 * H_zap, _sep_max_abs)

