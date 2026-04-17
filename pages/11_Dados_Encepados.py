import streamlit as st
try:
    from normas_referencias import mostrar_referencias_norma
except ImportError:
    def mostrar_referencias_norma(*a, **kw): pass

#  Utilidad: Color AutoCAD segun # de cuartos de pulgada 
def _color_acero_dxf(db_mm: float) -> int:
    """Retorna color AutoCAD (1-255) segun diametro nominal en mm (ASTM/NTC)."""
    if   db_mm <  7.5: return 2   # #2 1/4"   - Amarillo
    elif db_mm < 11.1: return 3   # #3 3/8"   - Verde
    elif db_mm < 14.3: return 4   # #4 1/2"   - Cian
    elif db_mm < 17.5: return 5   # #5 5/8"   - Azul
    elif db_mm < 20.6: return 6   # #6 3/4"   - Magenta
    elif db_mm < 23.8: return 30  # #7 7/8"   - Naranja
    elif db_mm < 27.0: return 1   # #8 1"     - Rojo
    else:              return 10  # #9+ 1-1/8" - Verde oscuro (pesado)
# 
import pandas as pd
import numpy as np
import math
import io
import sys
import os

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Plotly
try:
    import plotly.graph_objects as go
    _PLOTLY_AVAILABLE = True
except ImportError:
    _PLOTLY_AVAILABLE = False

# Librerías opcionales
try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import ezdxf
    _DXF_AVAILABLE = True
except ImportError:
    _DXF_AVAILABLE = False

try:
    import ifcopenshell
    _IFC_AVAILABLE = True
except ImportError:
    _IFC_AVAILABLE = False

# 
# IDIOMA GLOBAL
lang = st.session_state.get("idioma", "Español")
def _t(es, en): return en if lang == "English" else es
# 

# 
# PERSISTENCIA SUPABASE
# 
import requests
import json

def get_supabase_rest_info():
    try:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return url, key
    except Exception as e:
        return None, None

def guardar_proyecto_supabase(nombre, estado_dict, ref):
    url, key = get_supabase_rest_info()
    if not url or not key:
        db_path = f"db_proyectos_{ref.lower()}.json"
        db = {}
        if os.path.exists(db_path):
            with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
        db[f"[{ref.upper()}] {nombre}"] = {"nombre_proyecto": f"[{ref.upper()}] {nombre}", "estado_json": json.dumps(estado_dict)}
        with open(db_path, "w", encoding="utf-8") as f: json.dump(db, f)
        return True, " Proyecto guardado (Local)"
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Content-Type": "application/json", "Prefer": "resolution=merge-duplicates"}
    payload = {"nombre_proyecto": f"[{ref.upper()}] {nombre}", "user_id": st.session_state.get("user_id", "anonimo"), "estado_json": json.dumps(estado_dict)}
    try:
        res = requests.post(f"{url}/rest/v1/proyectos?on_conflict=nombre_proyecto", headers=headers, json=payload)
        if res.status_code in [200, 201, 204]: return True, " Proyecto en la nube"
        return False, f" Error API: {res.text}"
    except Exception as e: return False, f" Error API: {e}"

def cargar_proyecto_supabase(nombre, ref):
    url, key = get_supabase_rest_info()
    full_name = f"[{ref.upper()}] {nombre}" if not nombre.startswith("[") else nombre
    if not url or not key:
        db_path = f"db_proyectos_{ref.lower()}.json"
        if os.path.exists(db_path):
            with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            match = db.get(full_name) or db.get(nombre)
            if match:
                estado = json.loads(match["estado_json"])
                for k, v in estado.items(): st.session_state[k] = v
                return True, " Proyecto cargado (Local)"
        return False, " No encontrado"
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        res = requests.get(f"{url}/rest/v1/proyectos?nombre_proyecto=eq.{full_name}&select=*", headers=headers)
        if res.status_code == 200 and res.json():
            estado = json.loads(res.json()[0]["estado_json"])
            for k, v in estado.items(): st.session_state[k] = v
            return True, " Proyecto cargado"
        return False, " No encontrado"
    except Exception as e: return False, str(e)

def listar_proyectos_supabase(ref):
    url, key = get_supabase_rest_info()
    if not url or not key:
        db_path = f"db_proyectos_{ref.lower()}.json"
        if os.path.exists(db_path):
            with open(db_path, "r", encoding="utf-8") as f: db = json.load(f)
            return sorted([k.replace(f"[{ref.upper()}] ", "") for k in db.keys()])
        return []
    headers = {"apikey": key, "Authorization": f"Bearer {key}", "Accept": "application/json"}
    try:
        res = requests.get(f"{url}/rest/v1/proyectos?select=nombre_proyecto", headers=headers)
        if res.status_code == 200:
            return sorted([i["nombre_proyecto"].replace(f"[{ref.upper()}] ", "") for i in res.json() if f"[{ref.upper()}]" in i.get("nombre_proyecto","")])
        return []
    except: return []

def capturar_estado_module(ref):
    match = ["fc_dado", "fy_acero", "peso_conc", "recub_dado", "db_dado", "plantilla", "Pu_", "Mux_", "Muy_"]
    return {k: st.session_state[k] for k in st.session_state if any(k.startswith(m) for m in match) or k == "registro_dados"}

# 

st.set_page_config(page_title=_t("Dados de Pilotes (Encepados)", "Pile Caps"), layout="wide")

# 1. Banner SVG premium
st.markdown("""<div style="width:100%;overflow:hidden;border-radius:14px;margin-bottom:20px;box-shadow:0 4px 32px #0008;">
<svg viewBox="0 0 1100 220" xmlns="http://www.w3.org/2000/svg" style="width:100%;display:block;background:linear-gradient(135deg,#020617 0%,#0f172a 100%);">
  <g opacity="0.08" stroke="#14b8a6" stroke-width="0.5">
    <line x1="0" y1="55" x2="1100" y2="55"/><line x1="0" y1="110" x2="1100" y2="110"/>
    <line x1="0" y1="165" x2="1100" y2="165"/>
    <line x1="220" y1="0" x2="220" y2="220"/><line x1="440" y1="0" x2="440" y2="220"/>
    <line x1="660" y1="0" x2="660" y2="220"/>
  </g>
  <rect x="0" y="0" width="1100" height="3" fill="#0d9488" opacity="0.9"/>
  <rect x="0" y="217" width="1100" height="3" fill="#14b8a6" opacity="0.7"/>
  <!-- DADO PARAMETRICO -->
  <g transform="translate(60,40)">
    <text x="50" y="-15" text-anchor="middle" font-family="sans-serif" font-size="10" font-weight="600" fill="#cbd5e1" letter-spacing="1">DADO / ENCEPADO</text>
    <!-- Columna -->
    <rect x="35" y="0" width="30" height="40" fill="#334155" stroke="#475569" stroke-width="2"/>
    <path d="M35,10 L65,10 M35,20 L65,20 M35,30 L65,30" stroke="#14b8a6" stroke-width="1.5" opacity="0.8"/>
    <!-- Dado -->
    <path d="M0,40 L100,40 L110,80 L-10,80 Z" fill="#1e293b" stroke="#64748b" stroke-width="2" stroke-linejoin="round"/>
    <!-- Acero Malla -->
    <line x1="-5" y1="70" x2="105" y2="70" stroke="#f59e0b" stroke-width="1.5" stroke-dasharray="2,2"/>
    <line x1="-5" y1="75" x2="105" y2="75" stroke="#f59e0b" stroke-width="2"/>
    <!-- Pilotes Cortados -->
    <rect x="10" y="80" width="15" height="30" fill="#cbd5e1" stroke="#475569" stroke-width="1.5"/>
    <rect x="75" y="80" width="15" height="30" fill="#cbd5e1" stroke="#475569" stroke-width="1.5"/>
    <!-- Gancho ACI -->
    <path d="M25,75 L25,45 M25,45 L15,45 M85,75 L85,45 M85,45 L75,45" fill="none" stroke="#ef4444" stroke-width="1.5"/>
  </g>
  <!-- STRUT AND TIE (MECANISMO) -->
  <g transform="translate(250,50)">
    <text x="50" y="-20" text-anchor="middle" font-family="sans-serif" font-size="10" font-weight="600" fill="#cbd5e1" letter-spacing="1">STRUT &amp; TIE (STM)</text>
    <!-- Columna Nodo Top -->
    <rect x="35" y="0" width="30" height="20" fill="#334155"/>
    <!-- Nodos Pilotes Bottom -->
    <rect x="0" y="70" width="20" height="20" fill="#475569" opacity="0.5"/>
    <rect x="80" y="70" width="20" height="20" fill="#475569" opacity="0.5"/>
    <!-- Struts (Bielas compresion) -->
    <line x1="50" y1="20" x2="10" y2="70" stroke="#38bdf8" stroke-width="4"/>
    <line x1="50" y1="20" x2="90" y2="70" stroke="#38bdf8" stroke-width="4"/>
    <!-- Ties (Tirantes tension) -->
    <line x1="10" y1="70" x2="90" y2="70" stroke="#ef4444" stroke-width="3" stroke-dasharray="4,2"/>
    <!-- Texto Info -->
    <text x="110" y="25" font-family="monospace" font-size="9" fill="#7dd3fc">C: Compresión</text>
    <text x="110" y="75" font-family="monospace" font-size="9" fill="#fca5a5">T: Tensión (Acero)</text>
    <text x="50" y="45" text-anchor="middle" font-family="sans-serif" font-size="10" font-weight="bold" fill="#ffffff">&#952;</text>
  </g>
  <!-- TEXT BLOCK -->
  <g transform="translate(560,0)">
    <rect x="0" y="28" width="4" height="165" rx="2" fill="#14b8a6"/>
    <text x="18" y="66" font-family="Arial,sans-serif" font-size="30" font-weight="bold" fill="#ffffff">DADOS ESPACIALES</text>
    <text x="18" y="92" font-family="Arial,sans-serif" font-size="17" font-weight="300" fill="#5eead4" letter-spacing="2">CABEZALES DE PILOTES ACI-318</text>
    <rect x="18" y="100" width="480" height="1" fill="#14b8a6" opacity="0.5"/>
    <!-- Tags -->
    <rect x="18" y="111" width="100" height="22" rx="11" fill="#0f2e21" stroke="#10b981" stroke-width="1"/>
    <text x="68" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#34d399">PUNZONAMIENTO</text>
    <rect x="126" y="111" width="90" height="22" rx="11" fill="#301515" stroke="#ef4444" stroke-width="1"/>
    <text x="171" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#f87171">FLEXIÓN 2D</text>
    <rect x="224" y="111" width="130" height="22" rx="11" fill="#0c1a2e" stroke="#38bdf8" stroke-width="1"/>
    <text x="289" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#7dd3fc">STRUT-AND-TIE (STM)</text>
    <rect x="362" y="111" width="76" height="22" rx="11" fill="#2d1c07" stroke="#f59e0b" stroke-width="1"/>
    <text x="400" y="126" text-anchor="middle" font-family="Arial,sans-serif" font-size="9" font-weight="bold" fill="#fbbf24">APU / IFC</text>
    <!-- Description -->
    <text x="18" y="156" font-family="Arial,sans-serif" font-size="11" fill="#9ca3af">Analisis tridimensional y bidireccional rigido para cabezales de 2 a 6 pilotes.</text>
    <text x="18" y="172" font-family="Arial,sans-serif" font-size="11" fill="#9ca3af">Revisión de punzonamiento complejo integrado y alerta temprana para uso normativo</text>
    <text x="18" y="188" font-family="Arial,sans-serif" font-size="11" fill="#9ca3af">del metodo de Bielas y Tirantes (Strut-And-Tie). Incluye despiece grafico 3D.</text>
  </g>
</svg></div>""", unsafe_allow_html=True)

# 2. Panel global "Guía Profesional"
with st.expander(" ¿Cómo usar este módulo? — Guía Profesional", expanded=False):
    st.markdown('''
    ### Metodología y Flujo de Diseño (Cabezales de Pilotes)
    Konte evalúa la integridad y dimensionamiento de dados/encepados mediante análisis de rigidez bidireccional y modelado tridimensional cinemático.
    
    ####  1. Cinemática y Distribución Esfuerzos (Por Pilote)
    - Especifica configuraciones geométricas de grupos de pilotes (2 a 6 elementos).
    - Ingresa las cargas últimas ($P_u$, $M_{ux}$, $M_{uy}$). Konte resolverá la distribución tensional asumiendo el cabezal rígido.
    - **Criterio de Aceptación:** La reacción individual máxima en servicio y última no debe exceder las especificaciones del geotécnico (Pilote en rojo indica estado de falla portante).
    
    ####  2. Verificación Dinámica de Cortante (Aceros)
    - **Punzonamiento Complejo:** Se chequean los perímetros críticos ($d/2$) de columnas y pilotes de esquina/borde, ajustando la capacidad resistente concéntrica frente al espesor útil ($d$).
    - **Flexión y Cortante como Viga Ancha:** Se analizan los cuellos de botella de compresión diagonal, determinando la cuantía requerida para las parrillas de flexión X y Y.
    
    ####  3. Alerta Strut-And-Tie (STM)
    - **Aviso Normativo Crítico (ACI 318):** Si la relación a/d (distancia al pilote / peralte) es menor a $2.0$, la teoría de viga falla. El módulo audita automáticamente estas geometrías tipo "D-Region" advirtiendo el requerimiento del método de Bielas y Tirantes para el diseño seguro de compresión local.
    
    ####  4. Integración BIM y APU Estructural
    - Despliegue interactivo en 3D de parrillas bases, estribos de confinamiento, ganchos estándar ACI y arranques de columna.
    - Cuadro de APU (KonteAnalytics) con factor de desperdicio incorporado y listado paramétrico de despachos (DXF/IFC).
    ''')


# 
# SIDEBAR
# 
with st.sidebar:
    st.header(_t("Configuración Global", "Global Settings"))
    norma_sel = st.sidebar.selectbox(
        _t("Norma de diseño", "Design Code"),
        ["NSR-10 (Colombia)", "ACI 318-25 (EE.UU.)", "ACI 318-19 (EE.UU.)",
         "ACI 318-14 (EE.UU.)", "NEC-SE-HM (Ecuador)", "E.060 (Perú)",
         "NTC-EM (México)", "COVENIN 1753-2006 (Venezuela)"],
        key="dados_norma_sel"
    )
    mostrar_referencias_norma(norma_sel, "dados_encepados")
    st.markdown("---")
    st.subheader(" Guardar / Cargar Proyecto")
    nombre_producido = st.session_state.get(f"np_dados", "")
    st.markdown("**Nuevo Proyecto / Guardar**")
    nombre_proy_guardar = st.text_input("Nombre para guardar", value=nombre_producido, key="input_guardar_dados")
    if st.button(" Guardar Proyecto", use_container_width=True, key="btn_save_dados"):
        if nombre_proy_guardar:
            ok, msg = guardar_proyecto_supabase(nombre_proy_guardar, capturar_estado_module("dados"), "dados")
            if ok:
                st.session_state["np_dados"] = nombre_proy_guardar
                st.success(msg)
            else:
                st.error(msg)
        else:
            st.warning("Escribe un nombre")

    st.markdown("**Cargar Proyecto Existente**")
    lista_proyectos = listar_proyectos_supabase("dados")
    if lista_proyectos:
        idx_def = lista_proyectos.index(nombre_producido) if nombre_producido in lista_proyectos else 0
        nombre_proy_cargar = st.selectbox("Selecciona un proyecto", lista_proyectos, index=idx_def, key="sel_load_dados")
        
        def on_cargar_click():
            proy = st.session_state.get("sel_load_dados")
            if proy:
                ok, msg = cargar_proyecto_supabase(proy, "dados")
                st.session_state["__msg_cargar_dados"] = (ok, msg)
                st.session_state["np_dados"] = proy

        st.button(" Cargar", on_click=on_cargar_click, use_container_width=True, key="btn_load_dados")
        if "__msg_cargar_dados" in st.session_state:
            ok, msg = st.session_state.pop("__msg_cargar_dados")
            if ok: st.success(msg)
            else: st.error(msg)
    st.markdown("---")
    
    # 1. Configuración de Materiales
    with st.expander(_t("1. Materiales", "1. Materials"), expanded=True):
        fc_dado = st.number_input(_t("f'c Concreto Dado (MPa)", "f'c Cap Concrete (MPa)"), min_value=21.0, value=28.0, step=1.0)
        fy_acero = st.number_input(_t("fy Acero (MPa)", "fy Steel (MPa)"), min_value=240.0, value=420.0, step=10.0)
        peso_conc = st.number_input(_t("Peso Concreto (kN/m³)", "Concrete Weight (kN/m³)"), min_value=20.0, value=24.0, step=0.5)
        recub_dado = st.number_input(_t("Recubrimiento (m)", "Clear Cover (m)"), min_value=0.04, value=0.075, step=0.005)
        db_dado = st.number_input(_t("Diám. Barra principal (cm)", "Main Rebar Dia (cm)"), min_value=1.0, value=2.54, step=0.1)

    # 2. Geometría de Pilotes Internos
    with st.expander(_t("2. Pilotes Relacionados", "2. Related Piles"), expanded=True):
        D_pilote = st.number_input(_t("Diámetro del Pilote (m)", "Pile Diameter (m)"), min_value=0.20, value=0.60, step=0.05)
        Q_adm_pilote = st.number_input(_t("Q_adm Pilote Existente (kN)", "Existing Pile Q_adm (kN)"), min_value=100.0, value=1500.0, step=100.0)
        embeb_pilote = st.number_input(_t("Embebido Pilote en Dado (m)", "Pile Embedment (m)"), min_value=0.05, value=0.10, step=0.05)

    # 3. Geometría de la Columna
    with st.expander(_t("3. Geometría Columna", "3. Column Geometry"), expanded=True):
        c1_col = st.number_input(_t("Dimensión Cx (cm)", "Dimension Cx (cm)"), min_value=20.0, value=50.0, step=5.0)
        c2_col = st.number_input(_t("Dimensión Cy (cm)", "Dimension Cy (cm)"), min_value=20.0, value=50.0, step=5.0)
    
    # 4. Solicitaciones (Cargas en Base de Columna)
    with st.expander(_t("4. Solicitaciones (Diseño)", "4. Loads (Design)"), expanded=True):
        st.markdown(_t("*Cargas Últimas Mayoradas*", "*Factored Ultimate Loads*"))
        Pu = st.number_input(_t("Carga Axial Pu (kN)", "Axial Load Pu (kN)"), min_value=0.0, value=4500.0, step=100.0)
        Mux = st.number_input(_t("Momento Mux (kN.m)", "Moment Mux (kN.m)"), value=250.0, step=50.0)
        Muy = st.number_input(_t("Momento Muy (kN.m)", "Moment Muy (kN.m)"), value=150.0, step=50.0)

# 
# CUERPO PRINCIPAL (TABS)
# 
tab_geo, tab_des, tab_apu, tab_bim = st.tabs([
    _t("1. Configuración de Grupo", "1. Group Configuration"),
    _t("2. Punzonamiento y Flexión", "2. Punching & Flexure"),
    _t("3. Presupuesto (APU)", "3. Budget (APU)"),
    _t("4. Planos y BIM", "4. Drawings & BIM")
])

with tab_geo:
    st.subheader(_t("1.1 Parámetros del Encepado", "1.1 Pile Cap Parameters"))
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        plantilla = st.selectbox(_t("Disposición Paramétrica", "Parametric Layout"), [
            "2 Pilotes (Rectangular)", 
            "3 Pilotes (Triangular)", 
            "4 Pilotes (Cuadrícula)",
            "5 Pilotes (Dado + Central)",
            "6 Pilotes (Rectangular 3x2)"
        ], index=2)
        S_pilote = st.number_input(_t("Separación entre centros S (m)", "Center Spacing S (m)"), min_value=D_pilote*2, value=max(D_pilote*3, 1.0), step=0.1)
        # Selector de forma solo activo para 3 pilotes
        if plantilla.startswith("3"):
            forma_3p = st.radio(_t("Forma del Dado (3 Pilotes)", "Cap Shape (3 Piles)"),
                [_t("Triangular (Geometría real)", "Triangular (True geometry)"),
                 _t("Rectangular (Formaleta estandar)", "Rectangular (Standard formwork)")],
                index=0, horizontal=True, key="forma_3p")
            usar_triangular = forma_3p.startswith("T") or forma_3p.startswith("Tri")
        else:
            usar_triangular = False

    with col_d2:
        H_dado = st.number_input(_t("Espesor del Dado H (m)", "Cap Thickness H (m)"), min_value=0.4, value=1.0, step=0.1)
        voladizo = st.number_input(_t("Voladizo del Borde a centro de Pilote (m)", "Edge Overhang from Pile center (m)"), min_value=0.3, value=max(0.5, D_pilote/2 + 0.15), step=0.05)
    
    tipo_dado = plantilla.split(" ")[0]
    
    if tipo_dado == "2":
        pilotes = [(-S_pilote/2, 0), (S_pilote/2, 0)]
    elif tipo_dado == "3":
        R = S_pilote / math.sqrt(3)
        h_ap = R / 2
        pilotes = [(0, R), (-S_pilote/2, -h_ap), (S_pilote/2, -h_ap)]
    elif tipo_dado == "4":
        pilotes = [(-S_pilote/2, S_pilote/2), (S_pilote/2, S_pilote/2), (-S_pilote/2, -S_pilote/2), (S_pilote/2, -S_pilote/2)]
    elif tipo_dado == "5":
        pilotes = [(-S_pilote/2, S_pilote/2), (S_pilote/2, S_pilote/2), (-S_pilote/2, -S_pilote/2), (S_pilote/2, -S_pilote/2), (0, 0)]
    elif tipo_dado == "6":
        pilotes = [
            (-S_pilote, S_pilote/2), (0, S_pilote/2), (S_pilote, S_pilote/2),
            (-S_pilote, -S_pilote/2), (0, -S_pilote/2), (S_pilote, -S_pilote/2)
        ]
        
    n_pil = len(pilotes)
    xs = [p[0] for p in pilotes]
    ys = [p[1] for p in pilotes]
    
    B_sugerido = (max(xs) - min(xs)) + 2 * voladizo
    L_sugerido = (max(ys) - min(ys)) + 2 * voladizo
    
    # Contorno geométrico preciso del dado
    if tipo_dado == "3" and usar_triangular:
        dy = voladizo / math.cos(math.pi/6)
        v1 = (0, R + dy)
        v2 = (S_pilote/2 + dy*math.cos(math.pi/6), -h_ap - dy*math.sin(math.pi/6))
        v3 = (-v2[0], v2[1])
        contorno_dado = [v1, v2, v3]
        es_triangular = True
    else:
        contorno_dado = [(-B_sugerido/2, -L_sugerido/2), (B_sugerido/2, -L_sugerido/2), (B_sugerido/2, L_sugerido/2), (-B_sugerido/2, L_sugerido/2)]
        es_triangular = False

    
    st.markdown("---")
    st.subheader(_t("1.2 Cinemática Rígida y Geometría Generada", "1.2 Rigid Kinematics and Generated Geometry"))
    
    # Cinemática: Pi = P/n ± My*x / sum(x^2) ± Mx*y / sum(y^2)
    sum_x2 = sum(x**2 for x in xs)
    sum_y2 = sum(y**2 for y in ys)
    
    P_unit = Pu / n_pil
    
    res_pilotes = []
    max_pu_pilote = 0.0
    for i, (x, y) in enumerate(pilotes):
        P_my = (Muy * x / sum_x2) if sum_x2 > 0 else 0
        P_mx = (Mux * y / sum_y2) if sum_y2 > 0 else 0
        Pi = P_unit + P_my + P_mx
        
        max_pu_pilote = max(max_pu_pilote, Pi)
        estado = "OK OK" if Pi <= Q_adm_pilote else "FALLA EXCESO"
        res_pilotes.append({
            "ID": f"P{i+1}", 
            "X [m]": round(x, 2), 
            "Y [m]": round(y, 2), 
            "Carga Axial P_ui [kN]": round(Pi, 1), 
            "Estado": estado
        })
    
    df_pilotes = pd.DataFrame(res_pilotes)
    
    c_g1, c_g2 = st.columns([1.5, 2])
    with c_g1:
        st.write(f"**Dimensiones Encepado:** Ancho B={B_sugerido:.2f}m | Largo L={L_sugerido:.2f}m")
        st.write(f"**Reacción Máxima por Pilote:** {max_pu_pilote:.1f} kN (Límite: {Q_adm_pilote} kN)")
        st.dataframe(df_pilotes, use_container_width=True)
        if max_pu_pilote > Q_adm_pilote:
            st.error(_t(f"¡PELIGRO! La carga de {max_pu_pilote:.1f} kN excede la Capacidad Admisible Geotécnica. Aumente la separación del grupo o dimensiones.", f"DANGER! The load of {max_pu_pilote:.1f} kN exceeds Geotechnical Bearing Capacity."))
            
    with c_g2:
        if _PLOTLY_AVAILABLE:
            fig_geo = go.Figure()
            # Bounding box del dado (Polígono genérico)
            x_cont = [p[0] for p in contorno_dado] + [contorno_dado[0][0]]
            y_cont = [p[1] for p in contorno_dado] + [contorno_dado[0][1]]
            fig_geo.add_trace(go.Scatter(x=x_cont, y=y_cont, mode='lines', fill='toself', fillcolor='rgba(0, 255, 255, 0.05)', line=dict(color="cyan", width=2), showlegend=False, hoverinfo='skip'))
            # Columna
            fig_geo.add_shape(type="rect", x0=-c1_col/200, y0=-c2_col/200, x1=c1_col/200, y1=c2_col/200, line=dict(color="orange", width=2), fillcolor="orange", opacity=0.5)
            
            # Pilotes
            for idx, r in df_pilotes.iterrows():
                color_p = "green" if "OK" in r["Estado"] else "red"
                fig_geo.add_shape(type="circle", x0=r["X [m]"]-D_pilote/2, y0=r["Y [m]"]-D_pilote/2, x1=r["X [m]"]+D_pilote/2, y1=r["Y [m]"]+D_pilote/2, line=dict(color=color_p, width=2), fillcolor=color_p, opacity=0.3)
                fig_geo.add_annotation(x=r["X [m]"], y=r["Y [m]"], text=f"<b>{r['ID']}</b><br>{r['Carga Axial P_ui [kN]']} kN", showarrow=False, font=dict(color="white", size=10))
            
            fig_geo.update_layout(title="Distribución Cinemática en Planta", xaxis=dict(scaleanchor="y", scaleratio=1, title="X [m]"), yaxis=dict(title="Y [m]"), height=450, margin=dict(l=20, r=20, t=40, b=20), paper_bgcolor="#1e2530", plot_bgcolor="#1e2530", font=dict(color="white"))
            st.plotly_chart(fig_geo, use_container_width=True)

with tab_des:
    st.subheader(_t("Análisis ACI 318 Seccional (Cortante y Flexión)", "ACI 318 Sectional Analysis (Shear and Flexure)"))
    
    d_m = H_dado - recub_dado - (db_dado / 100.0) / 2.0
    c1_m = c1_col / 100.0
    c2_m = c2_col / 100.0
    
    c_des1, c_des2, c_des3, c_des4 = st.columns(4)
    
    #  1. PUNZONAMIENTO COLUMNA CENTRAL 
    b_o = 2 * ((c1_m + d_m) + (c2_m + d_m))
    Vu_punz = Pu
    # Reducción por pilotes dentro del perímetro crítico
    for idx, r in df_pilotes.iterrows():
        if abs(r["X [m]"]) <= (c1_m/2 + d_m/2) and abs(r["Y [m]"]) <= (c2_m/2 + d_m/2):
            Vu_punz -= r["Carga Axial P_ui [kN]"]
            
    beta_c = max(c1_col, c2_col) / max(min(c1_col, c2_col), 1)
    alpha_s = 40 # Columna Interior
    vc1 = 0.33 * math.sqrt(fc_dado)
    vc2 = 0.17 * (1 + 2/beta_c) * math.sqrt(fc_dado)
    vc3 = 0.083 * (2 + alpha_s * d_m / b_o) * math.sqrt(fc_dado)
    
    phiVc_punz = 0.75 * min(vc1, vc2, vc3) * 1000 * b_o * d_m
    ok_punz = phiVc_punz >= Vu_punz
    
    with c_des1:
        st.write("### Punz. Columna")
        st.metric("Peralte Efectivo d", f"{d_m:.3f} m")
        st.metric("Vu actuante", f"{Vu_punz:.1f} kN")
        st.write(f"**φVc Resistente:** {phiVc_punz:.1f} kN")
        if ok_punz: st.success("CUMPLE")
        else: st.error("FALLA PERALTE")

    #  2. PUNZONAMIENTO PILOTE (ESQUINA/CRÍTICO) 
    # Evaluamos el pilote con mayor carga P_ui
    Pu_pilote_max = df_pilotes["Carga Axial P_ui [kN]"].max() if n_pil > 0 else 0
    # Perímetro crítico iterativo (asumimos interior o de borde truncado)
    b_o_pil = math.pi * (D_pilote + d_m)
    # ACI 318 límite resistente a punzonamiento pilote
    phiVc_pilz = 0.75 * 0.33 * math.sqrt(fc_dado) * 1000 * b_o_pil * d_m
    ok_pilz = phiVc_pilz >= Pu_pilote_max
    
    with c_des2:
        st.write("### Punz. Pilote Crítico")
        st.metric("Pu Pilote Máx", f"{Pu_pilote_max:.1f} kN")
        st.write(f"**b_o considerado:** {b_o_pil:.2f} m")
        st.write(f"**φVc Resistente:** {phiVc_pilz:.1f} kN")
        if ok_pilz: st.success("CUMPLE")
        else: st.error("FALLA PERALTE")
        
    #  3. CORTANTE UNIDIRECCIONAL (VIGA) 
    def descuento_por_ubicacion(dist_centro, cara_col, d_efectivo, d_pil):
        dist_cara = dist_centro - cara_col
        if dist_cara >= d_efectivo + d_pil/2.0: return 1.0 # 100% genera cortante
        elif dist_cara <= d_efectivo - d_pil/2.0: return 0.0 # No genera
        else: return (dist_cara - (d_efectivo - d_pil/2.0)) / d_pil # Fracción

    Vu_vx = 0; Vu_vy = 0
    for idx, r in df_pilotes.iterrows():
        frac_x = descuento_por_ubicacion(abs(r["X [m]"]), c1_m/2, d_m, D_pilote)
        Vu_vx += r["Carga Axial P_ui [kN]"] * frac_x
        frac_y = descuento_por_ubicacion(abs(r["Y [m]"]), c2_m/2, d_m, D_pilote)
        Vu_vy += r["Carga Axial P_ui [kN]"] * frac_y
        
    phiVc_vx = 0.75 * 0.17 * math.sqrt(fc_dado) * 1000 * L_sugerido * d_m
    phiVc_vy = 0.75 * 0.17 * math.sqrt(fc_dado) * 1000 * B_sugerido * d_m
    ok_vx = phiVc_vx >= Vu_vx
    ok_vy = phiVc_vy >= Vu_vy
    
    with c_des3:
        st.write("### Cortante Unidireccional")
        st.metric("Vu (X) Reducido", f"{Vu_vx:.1f} kN")
        st.write(f"φVc Resistencia = {phiVc_vx:.1f} kN {'OK' if ok_vx else 'FALLA'}")
        st.metric("Vu (Y) Reducido", f"{Vu_vy:.1f} kN")
        st.write(f"φVc Resistencia = {phiVc_vy:.1f} kN {'OK' if ok_vy else 'FALLA'}")

    #  4. FLEXIÓN Y ACERO DE REFUERZO 
    Mu_flex_x = 0; Mu_flex_y = 0
    for idx, r in df_pilotes.iterrows():
        if abs(r["X [m]"]) > c1_m/2: Mu_flex_x += r["Carga Axial P_ui [kN]"] * (abs(r["X [m]"]) - c1_m/2)
        if abs(r["Y [m]"]) > c2_m/2: Mu_flex_y += r["Carga Axial P_ui [kN]"] * (abs(r["Y [m]"]) - c2_m/2)
        
    def req_as(Mu_kNm, b_m, d_mf):
        if Mu_kNm <= 0: return 0.0
        Mn = Mu_kNm / 0.90
        R = Mn / (b_m * d_mf**2 * 1000)
        val = 1 - 2*R/(0.85*fc_dado)
        if val < 0: return 9999.9 # Compresión falla
        rho = (0.85 * fc_dado / fy_acero) * (1 - math.sqrt(val))
        rho_min = max(0.25*math.sqrt(fc_dado)/fy_acero, 1.4/fy_acero)
        rho_use = max(rho, rho_min)
        return rho_use * (b_m*100) * (d_mf*100)
        
    As_x = req_as(Mu_flex_x, L_sugerido, d_m)
    As_y = req_as(Mu_flex_y, B_sugerido, d_m)

    #  Seleccion de barras en pulgadas (ASTM A615 / NSR-10) 
    BARRAS_IN = {
        "#3 (3/8\"": {"db_mm":9.5,  "As_cm2":0.71},
        "#4 (1/2\"": {"db_mm":12.7, "As_cm2":1.27},
        "#5 (5/8\"": {"db_mm":15.9, "As_cm2":1.99},
        "#6 (3/4\"": {"db_mm":19.1, "As_cm2":2.85},
        "#7 (7/8\"": {"db_mm":22.2, "As_cm2":3.87},
        "#8 (1\"":   {"db_mm":25.4, "As_cm2":5.07},
        "#9 (1-1/8\"":{"db_mm":28.7,"As_cm2":6.45},
        "#10 (1-1/4\"":{"db_mm":32.3,"As_cm2":8.19},
        "#11 (1-3/8\"":{"db_mm":35.8,"As_cm2":10.07},
    }
    st.markdown("---")
    st.subheader(_t("2.0 Seleccion de Armadura Longitudinal", "2.0 Reinforcement Bar Selection"))
    cb1, cb2 = st.columns(2)
    with cb1:
        bar_sel_x = st.selectbox(_t("Barra Eje X","Bar X-Dir"), list(BARRAS_IN.keys()), index=4, key="bar_sel_x")
    with cb2:
        bar_sel_y = st.selectbox(_t("Barra Eje Y","Bar Y-Dir"), list(BARRAS_IN.keys()), index=4, key="bar_sel_y")
    _bx = BARRAS_IN[bar_sel_x]; _by = BARRAS_IN[bar_sel_y]
    As_bx = _bx["As_cm2"];  As_by = _by["As_cm2"]
    bar_num_x = bar_sel_x.split("(")[0].strip()
    bar_num_y = bar_sel_y.split("(")[0].strip()
    # Numero de barras requeridas
    n_bar_x = max(2, math.ceil(As_x / As_bx))
    n_bar_y = max(2, math.ceil(As_y / As_by))
    As_prov_x = n_bar_x * As_bx
    As_prov_y = n_bar_y * As_by
    # Separacion real
    _ancho_x = B_sugerido - 2*recub_dado/100
    _ancho_y = L_sugerido - 2*recub_dado/100
    sep_x_calc = (_ancho_x / max(n_bar_x - 1, 1)) * 100  # cm
    sep_y_calc = (_ancho_y / max(n_bar_y - 1, 1)) * 100  # cm
    # Cuantia minima ACI 318 §9.6.1: rho_min = max(0.25*sqrt(fc)/fy, 1.4/fy)
    rho_min = max(0.25*math.sqrt(fc_dado)/fy_acero, 1.4/fy_acero)
    Ag_x = B_sugerido * d_m * 1e4  # cm2
    Ag_y = L_sugerido * d_m * 1e4
    As_min_x = rho_min * Ag_x
    As_min_y = rho_min * Ag_y
    # Cuantia maxima ACI 318 §9.6.3: rho_max = 0.75 * rho_bal
    rho_bal = (0.85 * 0.85 * fc_dado / fy_acero) * (600 / (600 + fy_acero))
    As_max_x = 0.75 * rho_bal * Ag_x
    As_max_y = 0.75 * rho_bal * Ag_y
    ok_As_x = (As_prov_x >= As_x) and (As_prov_x >= As_min_x) and (As_prov_x <= As_max_x)
    ok_As_y = (As_prov_y >= As_y) and (As_prov_y >= As_min_y) and (As_prov_y <= As_max_y)
    ok_sep_x = sep_x_calc <= 45  # cm ACI max
    ok_sep_y = sep_y_calc <= 45

    c_bar1, c_bar2 = st.columns(2)
    with c_bar1:
        st.markdown(f"**Eje X:** {bar_sel_x}")
        st.markdown(f"- N de barras: **{n_bar_x}**  |  Sep: **{sep_x_calc:.1f} cm** {'(OK)' if ok_sep_x else '(EXCEDE MAX 45cm)'}")
        st.markdown(f"- As req: {As_x:.1f} cm2  |  As prov: {As_prov_x:.1f} cm2 {'(OK)' if As_prov_x >= As_x else '(INSUFICIENTE)'}")
        st.markdown(f"- As min (ACI): {As_min_x:.1f} cm2  |  As max: {As_max_x:.1f} cm2")
        if ok_As_x: st.success("Armado Eje X CUMPLE")
        else: st.error("Armado Eje X FALLA — Revise numero y separacion de barras")
    with c_bar2:
        st.markdown(f"**Eje Y:** {bar_sel_y}")
        st.markdown(f"- N de barras: **{n_bar_y}**  |  Sep: **{sep_y_calc:.1f} cm** {'(OK)' if ok_sep_y else '(EXCEDE MAX 45cm)'}")
        st.markdown(f"- As req: {As_y:.1f} cm2  |  As prov: {As_prov_y:.1f} cm2 {'(OK)' if As_prov_y >= As_y else '(INSUFICIENTE)'}")
        st.markdown(f"- As min (ACI): {As_min_y:.1f} cm2  |  As max: {As_max_y:.1f} cm2")
        if ok_As_y: st.success("Armado Eje Y CUMPLE")
        else: st.error("Armado Eje Y FALLA — Revise numero y separacion de barras")

    # db_dado actualizado con la barra seleccionada para exportaciones
    db_dado = _bx["db_mm"] / 10.0  # cm

    
    with c_des4:
        st.write("### Flexión (As)")
        st.metric("Mu,x (Alineado en Y)", f"{Mu_flex_x:.1f} kNm")
        st.write(f"**As,x:** {As_x:.1f} cm² (A lo largo de L)")
        st.metric("Mu,y (Alineado en X)", f"{Mu_flex_y:.1f} kNm")
        st.write(f"**As,y:** {As_y:.1f} cm² (A lo largo de B)")
        
    # Método Puntal y Tensor (Alerta)
    st.markdown("---")
    st.write("#### ADVERTENCIA: Revisión Básica de Bielas y Tirantes (STM - ACI 318 Cap 23)")
    # Si la relación entre la distancia eje-pilar y peralte es menor a 2, rige STM
    max_dist = max([math.sqrt(r["X [m]"]**2 + r["Y [m]"]**2) for _, r in df_pilotes.iterrows()]) if n_pil > 0 else 0
    if max_dist < 2 * d_m:
        st.warning("La distancia desde los pilotes a la columna es menor a 2d. Este dado es rígido/profundo. ACI 318 exige diseño predominante usando **Bielas y Tirantes (STM)** en lugar de flexión plana clásica.")
    else:
        st.success("La distancia supera los 2d. El diseño convencional de vigas/losas (Flexión/Cortante) rige de manera precisa.")

    with tab_apu:
        st.subheader(_t("Presupuesto y Cantidades de Obra (APU)", "Budget and Bill of Quantities (APU)"))
        
        # 1. Cantidades
        vol_exc_dado = B_sugerido * L_sugerido * H_dado * 1.10 # 10% desperdicio lateral
        vol_conc_dado = B_sugerido * L_sugerido * H_dado
        area_form_dado = 2 * (B_sugerido + L_sugerido) * H_dado
        
        peso_acero_x = (As_x / 10000.0) * B_sugerido * 7850 * 1.05 # Area x long X density
        peso_acero_y = (As_y / 10000.0) * L_sugerido * 7850 * 1.05
        peso_acero_total = round(peso_acero_x + peso_acero_y, 1)

        st.markdown("##### Cuadro de Cantidades de Obra - Encepado")
        filas_cant = [
            {"Nº": 1, "Descripción": "Excavación mecánica encepado", "Unidad": "m³", "Cantidad": vol_exc_dado},
            {"Nº": 2, "Descripción": f"Concreto encepado f'c={fc_dado:.0f} MPa", "Unidad": "m³", "Cantidad": vol_conc_dado},
            {"Nº": 3, "Descripción": "Acero Refuerzo (Parrilla Inferior + Superior)", "Unidad": "kg", "Cantidad": peso_acero_total},
            {"Nº": 4, "Descripción": "Formaleta tableros metálicos", "Unidad": "m²", "Cantidad": area_form_dado},
        ]
        df_cant = pd.DataFrame(filas_cant)
        
        c_m1, c_m2, c_m3 = st.columns(3)
        c_m1.metric("Vol. Concreto Total", f"{vol_conc_dado:.2f} m³")
        c_m2.metric("Acero Requerido", f"{peso_acero_total:.1f} kg")
        c_m3.metric("Ratio Acero/Concreto", f"{peso_acero_total/(vol_conc_dado+0.001):.1f} kg/m³")
        
        st.dataframe(df_cant.style.format({"Cantidad": "{:,.2f}"}), use_container_width=True, hide_index=True)
        st.divider()

        # 2. Precios Unitarios
        st.markdown("##### Base de Precios Unitarios (COP) — Editable")
        precios_default = pd.DataFrame([
            {"Nº": 1, "Descripción": "Excavación mecánica encepado", "Unidad": "m³", "Precio Unitario (COP)": 36_000},
            {"Nº": 2, "Descripción": f"Concreto encepado f'c={fc_dado:.0f} MPa", "Unidad": "m³", "Precio Unitario (COP)": 460_000},
            {"Nº": 3, "Descripción": "Acero Refuerzo (Parrilla Inferior + Superior)", "Unidad": "kg", "Precio Unitario (COP)": 5_600},
            {"Nº": 4, "Descripción": "Formaleta tableros metálicos", "Unidad": "m²", "Precio Unitario (COP)": 40_000},
        ])
        
        df_precios = st.data_editor(
            precios_default,
            column_config={"Precio Unitario (COP)": st.column_config.NumberColumn("Precio Unitario (COP)", min_value=0, format="$ %d", required=True)},
            num_rows="fixed", use_container_width=True, hide_index=True, key="editor_precios_dados"
        )
        
        df_pres = df_cant.copy()
        df_pres["Precio Unitario (COP)"] = df_precios["Precio Unitario (COP)"].values
        df_pres["Subtotal (COP)"] = df_pres["Cantidad"] * df_pres["Precio Unitario (COP)"]
        costo_directo = df_pres["Subtotal (COP)"].sum()
        
        st.divider()
        st.markdown("##### Administración, Imprevistos y Utilidad (A.I.U.)")
        c_aiu1, c_aiu2, c_aiu3 = st.columns(3)
        with c_aiu1: pct_admin = st.slider("Administración (%)", 5, 25, 15, 1, key="aiu_adj_a")
        with c_aiu2: pct_impr = st.slider("Imprevistos (%)", 2, 10, 5, 1, key="aiu_adj_i")
        with c_aiu3: pct_util = st.slider("Utilidad (%)", 5, 20, 10, 1, key="aiu_adj_u")
        
        costo_total = costo_directo * (1 + (pct_admin + pct_impr + pct_util) / 100.0)
        
        st.markdown(f"""
        <div style="background:#1e3a5f;border:1px solid #1f6feb;border-radius:8px;padding:14px 16px; margin-top:15px">
            <div style="font-size:14px;color:#79c0ff;margin-bottom:4px">TOTAL DADO (A.I.U.={pct_admin + pct_impr + pct_util}%)</div>
            <div style="font-size:1.5rem;font-weight:700;color:#79c0ff">${costo_total:,.0f} COP</div>
            <div style="font-size:12px;color:#58a6ff;margin-top:4px">${costo_total/Pu:.0f} / kN transmitido</div>
        </div>
        """, unsafe_allow_html=True)
        
        if _PLOTLY_AVAILABLE:
            fig_donut = go.Figure(go.Pie(
                labels=["Costo Directo", "A.I.U."],
                values=[costo_directo, costo_total - costo_directo],
                hole=0.55, marker=dict(colors=["#29b6f6", "#ef5350"])
            ))
            fig_donut.update_layout(paper_bgcolor="#1e2530", font=dict(color="white"), height=300, margin=dict(t=10, b=10, l=10, r=10))
            st.plotly_chart(fig_donut, use_container_width=True)

with tab_bim:
    st.subheader(_t("Integración BIM 3D", "BIM 3D Integration"))
    
    if _PLOTLY_AVAILABLE:
        fig3d = go.Figure()
        
        # 1. Dado (Pile Cap) - Malla por Alphahull
        x_d = [p[0] for p in contorno_dado] * 2
        y_d = [p[1] for p in contorno_dado] * 2
        z_d = [-H_dado] * len(contorno_dado) + [0] * len(contorno_dado)
        
        fig3d.add_trace(go.Mesh3d(
            x=x_d, y=y_d, z=z_d,
            alphahull=0,
            color='lightblue', opacity=0.3, name='Concreto Encepado'
        ))
        
        # 2. Columna
        c1x = c1_m / 2; c2y = c2_m / 2; hc = 1.0
        x_col = [-c1x, c1x, c1x, -c1x, -c1x, c1x, c1x, -c1x]
        y_col = [-c2y, -c2y, c2y, c2y, -c2y, -c2y, c2y, c2y]
        z_col = [hc, hc, hc, hc, 0, 0, 0, 0]
        fig3d.add_trace(go.Mesh3d(
            x=x_col, y=y_col, z=z_col,
            i=[7, 0, 0, 0, 4, 4, 6, 6, 4, 0, 3, 2],
            j=[3, 4, 1, 2, 5, 6, 5, 2, 0, 1, 6, 3],
            k=[0, 7, 2, 3, 6, 7, 1, 1, 5, 5, 7, 6],
            color='#607d8b', opacity=1.0, name='Columna'
        ))

        # 3. Pilotes
        L_render = 3.0
        for idx_p, r_p in df_pilotes.iterrows():
            ppx = r_p['X [m]']; ppy = r_p['Y [m]']
            z_top_pil = -(H_dado - embeb_pilote)
            theta_p = np.linspace(0, 2*np.pi, 20)
            z_cil = np.linspace(z_top_pil, z_top_pil - L_render, 2)
            Tc, Zc = np.meshgrid(theta_p, z_cil)
            Xc = ppx + (D_pilote/2) * np.cos(Tc)
            Yc = ppy + (D_pilote/2) * np.sin(Tc)
            fig3d.add_trace(go.Surface(
                x=Xc, y=Yc, z=Zc,
                colorscale=[[0, '#9e9e9e'], [1, '#bdbdbd']],
                showscale=False, opacity=0.95, name=f'Pilote {r_p["ID"]}'
            ))

        # 4. Canasta de Acero (Ganchos ACI 318 §25.3: lhb = 12*db)
        z_inf = -H_dado + embeb_pilote + 0.03
        z_sup = -0.07
        espacio_barras = 0.20
        color_acero = '#cfd8dc'
        lh = max(0.15, 12.0 * (db_dado / 100.0))

        # Limites del contorno para generar rangos de barras
        xs_c = [p[0] for p in contorno_dado]
        ys_c = [p[1] for p in contorno_dado]
        xmin_c = min(xs_c) + recub_dado/100
        xmax_c = max(xs_c) - recub_dado/100
        ymin_c = min(ys_c) + recub_dado/100
        ymax_c = max(ys_c) - recub_dado/100

        def clip_bar_x(yp):
            """Retorna (xmin_bar, xmax_bar) para una barra horizontal en y=yp dentro del contorno"""
            if not es_triangular:
                return xmin_c, xmax_c
            xi = []
            for k in range(len(contorno_dado)):
                a = contorno_dado[k]
                b = contorno_dado[(k+1) % len(contorno_dado)]
                if min(a[1], b[1]) <= yp <= max(a[1], b[1]) and a[1] != b[1]:
                    xi.append(a[0] + (yp - a[1]) * (b[0] - a[0]) / (b[1] - a[1]))
            if len(xi) >= 2:
                return min(xi) + recub_dado/100, max(xi) - recub_dado/100
            return None, None

        def clip_bar_y(xp):
            """Retorna (ymin_bar, ymax_bar) para una barra transversal en x=xp dentro del contorno"""
            if not es_triangular:
                return ymin_c, ymax_c
            yi = []
            for k in range(len(contorno_dado)):
                a = contorno_dado[k]
                b = contorno_dado[(k+1) % len(contorno_dado)]
                if min(a[0], b[0]) <= xp <= max(a[0], b[0]) and a[0] != b[0]:
                    yi.append(a[1] + (xp - a[0]) * (b[1] - a[1]) / (b[0] - a[0]))
            if len(yi) >= 2:
                return min(yi) + recub_dado/100, max(yi) - recub_dado/100
            return None, None

        # --- Parrilla eje X (barras a lo largo de X, distribuidas en Y) ---
        for yp in np.arange(ymin_c + espacio_barras/2, ymax_c, espacio_barras):
            x0, x1 = clip_bar_x(yp)
            if x0 is None or x1 <= x0: continue
            # Inferior: barra horizontal con bastones 90° en cada extremo apuntando hacia ARRIBA
            fig3d.add_trace(go.Scatter3d(
                x=[x0, x0, x1, x1],
                y=[yp, yp, yp, yp],
                z=[z_inf + lh, z_inf, z_inf, z_inf + lh],
                mode='lines', line=dict(color=color_acero, width=3),
                marker=dict(size=0),
                showlegend=False, hoverinfo='skip'))
            # Superior: barra horizontal con bastones 90° apuntando hacia ABAJO
            fig3d.add_trace(go.Scatter3d(
                x=[x0, x0, x1, x1],
                y=[yp, yp, yp, yp],
                z=[z_sup - lh, z_sup, z_sup, z_sup - lh],
                mode='lines', line=dict(color=color_acero, width=3),
                marker=dict(size=0),
                showlegend=False, hoverinfo='skip'))

        # --- Parrilla eje Y (barras a lo largo de Y, distribuidas en X) ---
        dz_y = 0.025  # Las barras Y van por encima de las X
        for xp in np.arange(xmin_c + espacio_barras/2, xmax_c, espacio_barras):
            y0, y1 = clip_bar_y(xp)
            if y0 is None or y1 <= y0: continue
            # Inferior
            fig3d.add_trace(go.Scatter3d(
                x=[xp, xp, xp, xp],
                y=[y0, y0, y1, y1],
                z=[z_inf + dz_y + lh, z_inf + dz_y, z_inf + dz_y, z_inf + dz_y + lh],
                mode='lines', line=dict(color=color_acero, width=3),
                marker=dict(size=0),
                showlegend=False, hoverinfo='skip'))
            # Superior
            fig3d.add_trace(go.Scatter3d(
                x=[xp, xp, xp, xp],
                y=[y0, y0, y1, y1],
                z=[z_sup - dz_y - lh, z_sup - dz_y, z_sup - dz_y, z_sup - dz_y - lh],
                mode='lines', line=dict(color=color_acero, width=3),
                marker=dict(size=0),
                showlegend=False, hoverinfo='skip'))

        # --- Acero de Retraccion Perimetral (NSR-10 piel cada 30cm si H>60cm) ---
        if H_dado > 0.60:
            num_capas = int((H_dado - 0.20) / 0.30)
            if num_capas > 0:
                px_p = [p[0] * 0.93 for p in contorno_dado] + [contorno_dado[0][0] * 0.93]
                py_p = [p[1] * 0.93 for p in contorno_dado] + [contorno_dado[0][1] * 0.93]
                for c in range(1, num_capas + 1):
                    z_c = -H_dado + 0.10 + c * 0.30
                    fig3d.add_trace(go.Scatter3d(
                        x=px_p, y=py_p, z=[z_c]*len(px_p),
                        mode='lines', line=dict(color='#26c6da', width=2),
                        name="Ref. Perimetral Piel" if c == 1 else "",
                        marker=dict(size=0),
                        showlegend=(c == 1), hoverinfo='skip'))

        fig3d.update_layout(
            scene=dict(
                xaxis=dict(title="X [m]", range=[-B_sugerido*1.2, B_sugerido*1.2]),
                yaxis=dict(title="Y [m]", range=[-L_sugerido*1.2, L_sugerido*1.2]),
                zaxis=dict(title="Z [m]"),
                aspectmode='data'
            ),
            height=600, margin=dict(l=0, r=0, t=0, b=0),
            paper_bgcolor="#1e2530", font=dict(color="white")
        )
        st.plotly_chart(fig3d, use_container_width=True)
        
    else:
        st.error("Plotly no disponible.")

    # Entregables
    st.markdown("---")
    st.subheader(_t("Exportar Entregables", "Export Deliverables"))
    col_exp1, col_exp2, col_exp3 = st.columns(3)
    
    if _DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading(f"Memoria de Cálculo Estructural: Encepado de Pilotes", 0)
        doc.add_paragraph("Diseño de cabezal rígido según reglamento ACI 318 / NSR-10.")
        doc.add_paragraph("Nota: Este software procesa la cinemática rígida de placa, evaluando punzonamiento y flexiones bidireccionales.")
        
        doc.add_heading("1. Geometría y Materiales", level=2)
        shape_text = "Polígono Triangular offseteado" if tipo_dado == "3" else f"Rectangular (BxL) {B_sugerido:.2f} m x {L_sugerido:.2f} m"
        doc.add_paragraph(f"- Forma: {shape_text}\n- Espesor (H): {H_dado:.2f} m (Peralte d = {d_m:.2f} m)\n- Recubrimiento: {recub_dado} cm\n- f'c concreto: {fc_dado:.1f} MPa\n- fy acero: {fy_acero:.1f} MPa\n- Número de pilotes: {n_pil} (D={D_pilote:.2f} m)")
        
        doc.add_heading("2. Punzonamiento Columna (ACI 318 §22.6.5)", level=2)
        doc.add_paragraph(f"Corte bidireccional evaluado a d/2 de la columna plana ({c1_col}x{c2_col} cm).\n- Vu actuante (descontando pilotes internos totalmente y parcialmente): {Vu_punz:.1f} kN\n- Resistencia Nominal phiVc: {phiVc_punz:.1f} kN\n- Disposición: {'CUMPLE OK' if ok_punz else 'FALLA - REQUIERE MAYOR PERALTE O f´c'} (Vu <= phiVc)")
        
        doc.add_heading("3. Punzonamiento Pilote Crítico (ACI 318 §22.6 - NSR-10)", level=2)
        doc.add_paragraph(f"Evaluación del perímetro crítico (bo = {b_o_pil:.2f} m) sobre el pilote más exigido.\n- Pu máximo transmitido por columna a pilote: {Pu_pilote_max:.1f} kN\n- Resistencia Nominal phiVc: {phiVc_pilz:.1f} kN\n- Disposición: {'CUMPLE OK' if ok_pilz else 'FALLA'}")
        
        doc.add_heading("4. Cortante Unidireccional - Viga Ancha", level=2)
        doc.add_paragraph(f"Corte direccional evaluado a distancia d de la cara de apoyo, escalando cargas de pilotes interceptados perimetralmente.\n- Dir X: Vu = {Vu_vx:.1f} kN | phiVc = {phiVc_vx:.1f} kN -> {'CUMPLE' if ok_vx else 'FALLA'}\n- Dir Y: Vu = {Vu_vy:.1f} kN | phiVc = {phiVc_vy:.1f} kN -> {'CUMPLE' if ok_vy else 'FALLA'}")
        
        doc.add_heading("5. Flexión, Cuantía Térmica y As Requerido", level=2)
        doc.add_paragraph(f"Momentos (Mu) calculados iterativamente en la cara rígida del pedestal.\n- Eje X: Mu_x = {Mu_flex_x:.1f} kNm | Área Acero (As) Requerida = {As_x:.1f} cm²\n- Eje Y: Mu_y = {Mu_flex_y:.1f} kNm | Área Acero (As) Requerida = {As_y:.1f} cm²")
        if H_dado > 0.60:
            doc.add_paragraph(f"Verificación Térmica (ACI 318): Dado excede los 60cm. Se ordenan capas horizontales de acero de retracción intermedio en todo el recubrimiento vertical cada 30cm para garantizar cohesión del núcleo masivo.")
            
        doc.add_heading("6. Bielas y Tirantes (STM - Alerta Normativa)", level=2)
        if max_dist < 2 * d_m:
            doc.add_paragraph("ADVERTENCIA STM: Relación de Luces / Peralte (L/d < 2) indica un inminente comportamiento estructural de Viga de Gran Peralte (Deep Beam). ACI 318 exige diseño avanzado de reticulado STM en lugar de hipótesis en flexión plana. Los As reportados deben reforzarse como tiradores primarios horizontales con anclajes plenos tipo 90°.")
        else:
            doc.add_paragraph("COMPORTAMIENTO FLEXIBLE: La distancia de pilotes extremos supera los 2*d. El diseño bidireccional puro domina y la aproximación calculada es válida sin modelar STM.")
            
        doc.add_heading("7. Detalles BIM e Informes Gráficos", level=2)
        doc.add_paragraph("Para adjuntar gráficos paramétricos en posteriores exportaciones DOCX, verifique tener instalado 'kaleido' en su entorno Python. El modelo de exportación DXF integra correctamente las envolventes poligonales de diseño.")
        
            
        bio = io.BytesIO()
        doc.save(bio)
        col_exp1.download_button(_t("Descargar DOCX", "Download DOCX"), data=bio.getvalue(), file_name="Memoria_Encepado.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if _DXF_AVAILABLE:
        doc_dxf = ezdxf.new("R2010")
        doc_dxf.header["$INSUNITS"] = 6  # metros
        for _l, _c in [("DADO",3),("PILOTES",4),("COLUMNA",2),("ACERO",1),
                       ("COTAS",6),("TEXTO",7),("ROTULO",7),("MARCO",7),
                       ("CORTE",3),("FIGURADO",5),("MARCO2",15)]:
            doc_dxf.layers.add(_l, color=_c)
        msp = doc_dxf.modelspace()

        #  Normalizar recubrimiento a metros 
        _rec_m = (recub_dado if recub_dado >= 1 else recub_dado * 100) / 100
        _lh_m  = max(0.15, 12 * db_dado / 100)   # hook length: m
        _rec_cm_val = _rec_m * 100

        #  Zonas del plano 
        _PXc = B_sugerido / 2 + 1.5        # plan center X
        _PYc = L_sugerido / 2 + 17.5       # plan center Y
        # Section zone: dado at (1.5, 10) bottom-left
        _SX0 = 1.5
        _SY0 = 10.0                         # bottom of section dado
        # Figurado zone
        _FX0 = max(B_sugerido + 4.0, 9.0)  # left of figurado
        _FY0 = 26.0                         # top of figurado
        # Quantities zone
        _QX0 = _FX0
        _QY0 = 18.0                         # top of quantities
        # Title block
        _RX0 = 33.0
        _RY0 = 0.0
        _RW  = 9.0
        _RH  = 27.0

        def _P(x, y): return (x + _PXc, y + _PYc)   # translate to plan zone
        def _S(x, y): return (x + _SX0, y + _SY0)   # translate to section zone

        # 
        # ZONA 1 — PLANTA ENCEPADO
        # 
        _dado_pts_p = [_P(x, y) for (x, y) in contorno_dado]
        msp.add_lwpolyline(_dado_pts_p, close=True,
                           dxfattribs={"layer":"DADO","lineweight":50})
        # Columna
        _col_pts_p = [_P(-c1_m/2,-c2_m/2),_P(c1_m/2,-c2_m/2),
                      _P(c1_m/2,c2_m/2),_P(-c1_m/2,c2_m/2)]
        msp.add_lwpolyline(_col_pts_p, close=True,
                           dxfattribs={"layer":"COLUMNA","lineweight":40})
        # Pilotes
        for _, _rr in df_pilotes.iterrows():
            _ppx, _ppy = _rr["X [m]"], _rr["Y [m]"]
            msp.add_circle(_P(_ppx, _ppy), radius=D_pilote/2,
                           dxfattribs={"layer":"PILOTES","lineweight":30})
            msp.add_line(_P(_ppx-D_pilote/2*0.7,_ppy),_P(_ppx+D_pilote/2*0.7,_ppy),
                         dxfattribs={"layer":"PILOTES"})
            msp.add_line(_P(_ppx,_ppy-D_pilote/2*0.7),_P(_ppx,_ppy+D_pilote/2*0.7),
                         dxfattribs={"layer":"PILOTES"})
            msp.add_text(str(_rr["ID"]),
                         dxfattribs={"height":0.07,"layer":"TEXTO"}).set_placement(
                _P(_ppx + D_pilote/2 + 0.04, _ppy + 0.02))
        # Parrilla acero en planta
        _xmin_p = _PXc - B_sugerido/2 + _rec_m
        _xmax_p = _PXc + B_sugerido/2 - _rec_m
        _ymin_p = _PYc - L_sugerido/2 + _rec_m
        _ymax_p = _PYc + L_sugerido/2 - _rec_m
        _y_ac = _ymin_p + sep_y_calc/200
        while _y_ac < _ymax_p:
            msp.add_line((_xmin_p, _y_ac), (_xmax_p, _y_ac), dxfattribs={"layer":"ACERO","color":_color_acero_dxf(db_dado*10)})
            _y_ac += sep_y_calc / 100
        _x_ac = _xmin_p + sep_x_calc/200
        while _x_ac < _xmax_p:
            msp.add_line((_x_ac, _ymin_p), (_x_ac, _ymax_p), dxfattribs={"layer":"ACERO","color":_color_acero_dxf(db_dado*10)})
            _x_ac += sep_x_calc / 100
        # Cotas planta
        _th = 0.09
        msp.add_text(f"B = {B_sugerido:.2f} m",
                     dxfattribs={"height":_th,"layer":"COTAS"}).set_placement(
            (_PXc - 0.4, _PYc + L_sugerido/2 + 0.15))
        msp.add_text(f"L = {L_sugerido:.2f} m",
                     dxfattribs={"height":_th,"layer":"COTAS"}).set_placement(
            (_PXc + B_sugerido/2 + 0.12, _PYc))
        msp.add_text("A", dxfattribs={"height":0.12,"layer":"TEXTO"}).set_placement(
            (_PXc - B_sugerido/2 - 0.30, _PYc))
        msp.add_text("A", dxfattribs={"height":0.12,"layer":"TEXTO"}).set_placement(
            (_PXc + B_sugerido/2 + 0.10, _PYc))
        msp.add_text("PLANTA ENCEPADO",
                     dxfattribs={"height":0.13,"layer":"TEXTO"}).set_placement(
            (_PXc - 0.55, _PYc - L_sugerido/2 - 0.40))
        msp.add_text(f"Esc. 1:50",
                     dxfattribs={"height":0.09,"layer":"TEXTO"}).set_placement(
            (_PXc - 0.20, _PYc - L_sugerido/2 - 0.60))

        # 
        # ZONA 2 — CORTE A-A
        # 
        # Dado en corte: de (0,0) a (B, H)
        msp.add_lwpolyline([_S(0,0),_S(B_sugerido,0),
                            _S(B_sugerido,H_dado),_S(0,H_dado)],
                           close=True, dxfattribs={"layer":"CORTE","lineweight":60})
        # Columna sobre dado
        _col_c0 = B_sugerido/2 - c1_m/2
        _col_c1 = B_sugerido/2 + c1_m/2
        msp.add_lwpolyline([_S(_col_c0, H_dado),_S(_col_c1, H_dado),
                            _S(_col_c1, H_dado+0.80),_S(_col_c0, H_dado+0.80)],
                           close=True, dxfattribs={"layer":"COLUMNA","lineweight":40})
        # Pilotes en corte
        for _ppx in sorted(df_pilotes["X [m]"].unique()):
            _px0 = B_sugerido/2 + _ppx - D_pilote/2
            _px1 = B_sugerido/2 + _ppx + D_pilote/2
            msp.add_lwpolyline([_S(_px0,-3.0),_S(_px1,-3.0),
                                _S(_px1,0),_S(_px0,0)],
                               close=True, dxfattribs={"layer":"PILOTES","lineweight":30})
            msp.add_line(_S(B_sugerido/2+_ppx,-3.0),_S(B_sugerido/2+_ppx,0),
                         dxfattribs={"layer":"PILOTES"})
        # Armadura inferior
        _z_inf_c = _SY0 + _rec_m
        _x_bar_l = _SX0 + _rec_m
        _x_bar_r = _SX0 + B_sugerido - _rec_m
        msp.add_line((_x_bar_l, _z_inf_c), (_x_bar_r, _z_inf_c),
                     dxfattribs={"layer":"ACERO","lineweight":40})
        # Ganchos inf
        msp.add_line((_x_bar_l, _z_inf_c), (_x_bar_l, _z_inf_c + _lh_m),
                     dxfattribs={"layer":"ACERO","lineweight":40})
        msp.add_line((_x_bar_r, _z_inf_c), (_x_bar_r, _z_inf_c + _lh_m),
                     dxfattribs={"layer":"ACERO","lineweight":40})
        msp.add_text(f"{n_bar_x}{bar_num_x}@{sep_x_calc:.0f}cm (Inf)",
                     dxfattribs={"height":0.09,"layer":"ACERO"}).set_placement(
            (_x_bar_r + 0.10, _z_inf_c - 0.05))
        # Armadura superior
        _z_sup_c = _SY0 + H_dado - _rec_m
        msp.add_line((_x_bar_l, _z_sup_c), (_x_bar_r, _z_sup_c),
                     dxfattribs={"layer":"ACERO","lineweight":40})
        msp.add_line((_x_bar_l, _z_sup_c), (_x_bar_l, _z_sup_c - _lh_m),
                     dxfattribs={"layer":"ACERO","lineweight":40})
        msp.add_line((_x_bar_r, _z_sup_c), (_x_bar_r, _z_sup_c - _lh_m),
                     dxfattribs={"layer":"ACERO","lineweight":40})
        msp.add_text(f"{n_bar_x}{bar_num_x}@{sep_x_calc:.0f}cm (Sup)",
                     dxfattribs={"height":0.09,"layer":"ACERO"}).set_placement(
            (_x_bar_r + 0.10, _z_sup_c + 0.02))
        # Acero piel intermedio
        if H_dado > 0.60:
            _n_int = int((H_dado - 0.20) / 0.30)
            for _ci in range(1, _n_int + 1):
                _zz = _SY0 + 0.10 + _ci * 0.30
                msp.add_line((_x_bar_l, _zz), (_x_bar_r, _zz),
                             dxfattribs={"layer":"ACERO","color":_color_acero_dxf(db_dado*10)})
            msp.add_text(f"Piel: #3@30cm",
                         dxfattribs={"height":0.08,"layer":"ACERO"}).set_placement(
                (_x_bar_r + 0.10, _SY0 + H_dado / 2))
        # Cotas corte
        msp.add_text(f"H = {H_dado:.2f} m",
                     dxfattribs={"height":_th,"layer":"COTAS"}).set_placement(
            (_SX0 + B_sugerido + 0.10, _SY0 + H_dado/2))
        msp.add_text(f"B = {B_sugerido:.2f} m",
                     dxfattribs={"height":_th,"layer":"COTAS"}).set_placement(
            (_SX0 + B_sugerido/2 - 0.30, _SY0 + H_dado + 0.12))
        msp.add_text(f"Rec. = {_rec_cm_val:.1f} cm",
                     dxfattribs={"height":0.07,"layer":"COTAS"}).set_placement(
            (_x_bar_l, _z_inf_c + 0.04))
        msp.add_text("CORTE A-A",
                     dxfattribs={"height":0.13,"layer":"TEXTO"}).set_placement(
            (_SX0 + B_sugerido/2 - 0.28, _SY0 - 0.40))
        msp.add_text("Esc. 1:50",
                     dxfattribs={"height":0.09,"layer":"TEXTO"}).set_placement(
            (_SX0 + B_sugerido/2 - 0.18, _SY0 - 0.58))

        # 
        # ZONA 3 — FIGURADO DE BARRAS
        # 
        _bar_a_len = B_sugerido - 2*_rec_m   # longitud neta barra X (m)
        _bar_b_len = L_sugerido - 2*_rec_m   # longitud neta barra Y (m)
        _scale_fig = 1.0 / 20.0              # escala figurado 1:20

        def _draw_bar_fig(fx, fy, bar_len_m, hook_m, label, diam, qty,
                          direction="h"):
            """Dibuja una barra con ganchos en el figurado.
            fx, fy: origen de la figura. direction='h' horizontal.
            Devuelve el Y inferior usado."""
            _bls = bar_len_m * _scale_fig    # bar length scaled
            _hls = hook_m    * _scale_fig    # hook scaled
            th_f = 0.10
            # Outline box
            _box_w = _bls + 2 * _hls + 0.60
            msp.add_lwpolyline([(fx,fy),(fx+_box_w,fy),
                                (fx+_box_w,fy+1.8),(fx,fy+1.8)],
                               close=True, dxfattribs={"layer":"MARCO2","lineweight":15})
            # Dibujar barra: left hook + horizontal + right hook
            _y0 = fy + 1.0
            _x_hl = fx + 0.30 + _hls      # right end of left hook
            _x_hr = _x_hl + _bls          # left end of right hook
            _hh = _hls * 0.8              # hook vertical height (scaled)
            # horizontal
            msp.add_line((_x_hl, _y0), (_x_hr, _y0),
                         dxfattribs={"layer":"FIGURADO","lineweight":40})
            # left hook (up)
            msp.add_line((_x_hl, _y0), (_x_hl, _y0 + _hh),
                         dxfattribs={"layer":"FIGURADO","lineweight":40})
            # right hook (up)
            msp.add_line((_x_hr, _y0), (_x_hr, _y0 + _hh),
                         dxfattribs={"layer":"FIGURADO","lineweight":40})
            # Cotas
            msp.add_text(f"{_hls*1000/(_scale_fig*1000)*100:.0f}",
                         dxfattribs={"height":0.07,"layer":"FIGURADO"}).set_placement(
                (fx+0.30, _y0+_hh+0.04))
            msp.add_text(f"{bar_len_m*100:.0f}",
                         dxfattribs={"height":0.07,"layer":"FIGURADO"}).set_placement(
                (_x_hl + _bls/2 - 0.1, _y0 + 0.04))
            msp.add_text(f"{_hls*1000/(_scale_fig*1000)*100:.0f}",
                         dxfattribs={"height":0.07,"layer":"FIGURADO"}).set_placement(
                (_x_hr + 0.02, _y0+_hh+0.04))
            # Label
            msp.add_text(f"Marca: {label}   {diam}   N={qty}   L total={bar_len_m+2*hook_m:.2f}m",
                         dxfattribs={"height":0.10,"layer":"TEXTO"}).set_placement(
                (fx + 0.10, fy + 0.15))
            return fy   # box starts at fy

        # Título zona figurado
        msp.add_text("FIGURADO DE BARRAS (ESC. 1:20)",
                     dxfattribs={"height":0.14,"layer":"TEXTO"}).set_placement(
            (_FX0, _FY0 + 0.12))
        msp.add_line((_FX0, _FY0),(_FX0 + 20.0, _FY0), dxfattribs={"layer":"MARCO"})

        # Barra A: eje X (main bar with hooks)
        _draw_bar_fig(_FX0, _FY0 - 2.0, _bar_a_len, _lh_m,
                      "A", bar_num_x, n_bar_x * 2)
        msp.add_text(f"Barra A — Armado eje X (inf + sup): {n_bar_x*2} pzas  {bar_num_x}@{sep_x_calc:.0f}cm",
                     dxfattribs={"height":0.09,"layer":"TEXTO"}).set_placement(
            (_FX0, _FY0 - 0.30))

        # Barra B: eje Y
        _draw_bar_fig(_FX0 + 9.0, _FY0 - 2.0, _bar_b_len, _lh_m,
                      "B", bar_num_y, n_bar_y * 2)
        msp.add_text(f"Barra B — Armado eje Y (inf + sup): {n_bar_y*2} pzas  {bar_num_y}@{sep_y_calc:.0f}cm",
                     dxfattribs={"height":0.09,"layer":"TEXTO"}).set_placement(
            (_FX0 + 9.0, _FY0 - 0.30))

        # Barra C: piel (recta)
        if H_dado > 0.60:
            _n_int_fig = int((H_dado - 0.20) / 0.30)
            _piel_len = 2*(B_sugerido + L_sugerido) - 8*_rec_m
            fx_c = _FX0
            fy_c = _FY0 - 4.2
            _bls_c = min(_piel_len * _scale_fig, 5.0)
            msp.add_lwpolyline([(fx_c,fy_c),(fx_c+_bls_c+0.6,fy_c),
                                (fx_c+_bls_c+0.6,fy_c+1.4),(fx_c,fy_c+1.4)],
                               close=True, dxfattribs={"layer":"MARCO2","lineweight":15})
            msp.add_line((fx_c+0.30, fy_c+0.8), (fx_c+0.30+_bls_c, fy_c+0.8),
                         dxfattribs={"layer":"FIGURADO","lineweight":30})
            msp.add_text(f"Marca C — Piel #3@30cm: {4*_n_int_fig} pzas rectas",
                         dxfattribs={"height":0.09,"layer":"TEXTO"}).set_placement(
                (fx_c+0.10, fy_c+0.15))

        # 
        # ZONA 4 — CUADRO DE CANTIDADES
        # 
        _vol_c = B_sugerido * L_sugerido * H_dado
        _long_bx = _bar_a_len + 2 * _lh_m
        _long_by = _bar_b_len + 2 * _lh_m
        _kg_m_x = (db_dado ** 2) * 0.00785           # kg/m de acero
        _kg_bx = n_bar_x * 2 * _long_bx * _kg_m_x
        _kg_by = n_bar_y * 2 * _long_by * _kg_m_x
        _kg_tot = _kg_bx + _kg_by
        _form = 2 * (B_sugerido + L_sugerido) * H_dado

        _CX0 = _FX0
        _CY0 = _QY0           # top of table
        _CW  = 20.0
        _col_ws = [7.0, 2.5, 2.0, 2.5, 6.0]

        # Marco externo
        _table_rows = 7
        _CH = 0.45 * (_table_rows + 1) + 0.55
        msp.add_lwpolyline(
            [(_CX0,_CY0-_CH),(_CX0+_CW,_CY0-_CH),
             (_CX0+_CW,_CY0),(_CX0,_CY0)],
            close=True, dxfattribs={"layer":"MARCO","lineweight":40})
        # Titulo
        msp.add_text("CUADRO DE CANTIDADES — ACERO Y MATERIALES",
                     dxfattribs={"height":0.14,"layer":"ROTULO"}).set_placement(
            (_CX0+0.10, _CY0-0.30))
        msp.add_line((_CX0,_CY0-0.50),(_CX0+_CW,_CY0-0.50),
                     dxfattribs={"layer":"MARCO","lineweight":25})
        # Header
        msp.add_line((_CX0,_CY0-0.95),(_CX0+_CW,_CY0-0.95),
                     dxfattribs={"layer":"MARCO","lineweight":15})
        _hdrs2 = ["MATERIAL / ELEMENTO","CANT.","UND.","PESO kg","OBSERVACION"]
        _cxa = _CX0
        for _hh2, _cwa in zip(_hdrs2, _col_ws):
            msp.add_text(_hh2, dxfattribs={"height":0.10,"layer":"TEXTO"}).set_placement(
                (_cxa+0.06, _CY0-0.85))
            _cxa += _cwa
            if _cxa < _CX0+_CW:
                msp.add_line((_cxa,_CY0-0.50),(_cxa,_CY0-_CH),
                             dxfattribs={"layer":"MARCO","lineweight":10})
        # Data rows
        _tdata = [
            (f"Concreto f'c={fc_dado:.0f}MPa", f"{_vol_c:.2f}", "m³",
             "-", "Factor desperdicio: 1.05"),
            (f"Barra A: {bar_num_x} Eje X (inf+sup)", str(n_bar_x*2), "barras",
             f"{_kg_bx:.1f}", f"L={_long_bx:.2f}m/barra  sep.={sep_x_calc:.0f}cm"),
            (f"Barra B: {bar_num_y} Eje Y (inf+sup)", str(n_bar_y*2), "barras",
             f"{_kg_by:.1f}", f"L={_long_by:.2f}m/barra  sep.={sep_y_calc:.0f}cm"),
            ("TOTAL ACERO", "-", "-", f"{_kg_tot:.1f}",
             "Incluir 5% desperdicio + empalmes"),
            (f"Formaleta lateral", f"{_form:.2f}", "m²",
             "-", "Panel fenólico 1 uso"),
            (f"Pilotes D={D_pilote*100:.0f}cm", str(n_pil), "uds.",
             "-", "Suministro según diseño pilotaje"),
        ]
        _ry2 = _CY0 - 0.95
        for _row in _tdata:
            _ry2 -= 0.45
            msp.add_line((_CX0,_ry2),(_CX0+_CW,_ry2),
                         dxfattribs={"layer":"MARCO","lineweight":10})
            _cxa2 = _CX0
            for _cell, _cwa in zip(_row, _col_ws):
                msp.add_text(str(_cell), dxfattribs={"height":0.09,"layer":"TEXTO"}).set_placement(
                    (_cxa2+0.06, _ry2+0.08))
                _cxa2 += _cwa

        # 
        # ZONA 5 — ROTULO ICONTEC (franja derecha fija)
        # 
        # Marco doble
        msp.add_lwpolyline(
            [(_RX0,_RY0),(_RX0+_RW,_RY0),(_RX0+_RW,_RY0+_RH),(_RX0,_RY0+_RH)],
            close=True, dxfattribs={"layer":"MARCO","lineweight":80})
        msp.add_lwpolyline(
            [(_RX0+0.06,_RY0+0.06),(_RX0+_RW-0.06,_RY0+0.06),
             (_RX0+_RW-0.06,_RY0+_RH-0.06),(_RX0+0.06,_RY0+_RH-0.06)],
            close=True, dxfattribs={"layer":"MARCO","lineweight":20})
        # Dividers: at y = 4.5, 6.0, 7.5, 10, 23 from _RY0
        _rdivs = [4.5, 6.0, 7.5, 10.0, 23.0]
        for _rd in _rdivs:
            msp.add_line((_RX0,_RY0+_rd),(_RX0+_RW,_RY0+_rd),
                         dxfattribs={"layer":"ROTULO","lineweight":20})
        # Divider vertical (col derecha)
        msp.add_line((_RX0+_RW*0.55,_RY0+4.5),(_RX0+_RW*0.55,_RY0+10.0),
                     dxfattribs={"layer":"ROTULO","lineweight":15})
        # Texts: carefully spaced ≥0.2 from each divider
        _ritems = [
            # Encabezado principal (entre 23 y 27)
            (0.15, 26.5, 0.30, "ESTRUCTURAS DE CIMENTACION PROFUNDA"),
            (0.15, 25.8, 0.18, f"ENCEPADO TIPO {n_pil} PILOTES"),
            (0.15, 25.2, 0.14, f"Proyecto: {nombre_proyecto if 'nombre_proyecto' in dir() else 'CIMENTACIONES'}"),
            # Descripción (entre 10.0 y 23.0)
            (0.15, 22.5, 0.12, "DESCRIPCION:"),
            (0.15, 22.0, 0.11, f"Dado: {B_sugerido:.2f}m x {L_sugerido:.2f}m x H={H_dado:.2f}m"),
            (0.15, 21.5, 0.11, f"Pilotes: D={D_pilote*100:.0f}cm  S={S_pilote:.2f}m  Voladizo={voladizo:.2f}m"),
            (0.15, 21.0, 0.11, f"Armado: {bar_num_x}@{sep_x_calc:.0f}cm (X) / {bar_num_y}@{sep_y_calc:.0f}cm (Y)"),
            (0.15, 20.3, 0.11, f"As,x req={As_x:.1f} cm2  prov={As_prov_x:.1f} cm2"),
            (0.15, 19.8, 0.11, f"As,y req={As_y:.1f} cm2  prov={As_prov_y:.1f} cm2"),
            (0.15, 19.0, 0.11, "MATERIALES:"),
            (0.15, 18.5, 0.11, f"Concreto: f'c = {fc_dado:.0f} MPa"),
            (0.15, 18.0, 0.11, f"Acero: fy = {fy_acero:.0f} MPa | NSR-10 / ACI 318-19"),
            (0.15, 17.2, 0.11, "VERIFICACIONES ESTRUCTURALES:"),
            (0.15, 16.7, 0.10, f"  Punz. columna: {'CUMPLE' if ok_punz else 'FALLA'}"),
            (0.15, 16.3, 0.10, f"  Punz. pilote:  {'CUMPLE' if ok_pilz else 'FALLA'}"),
            (0.15, 15.9, 0.10, f"  Cortante X:    {'CUMPLE' if ok_vx else 'FALLA'}"),
            (0.15, 15.5, 0.10, f"  Cortante Y:    {'CUMPLE' if ok_vy else 'FALLA'}"),
            (0.15, 15.1, 0.10, f"  As,x:          {'CUMPLE' if ok_As_x else 'FALLA'}"),
            (0.15, 14.7, 0.10, f"  As,y:          {'CUMPLE' if ok_As_y else 'FALLA'}"),
            (0.15, 13.8, 0.11, "NOTAS:"),
            (0.15, 13.3, 0.09, "1. Barras segun ASTM A615 Gr60."),
            (0.15, 12.9, 0.09, "2. Recubrimiento: NSR-10 tabla C.7.7."),
            (0.15, 12.5, 0.09, "3. Concreto fluido, asentam. 10-15cm."),
            (0.15, 12.1, 0.09, "4. Ver planos de pilotes para detalles."),
            (0.15, 11.7, 0.09, "5. Acero con tracabilidad NTC2289."),
            # Zona ficha tecnica (entre 7.5 y 10.0)
            (0.15, 9.65,  0.11, "ESCALA: 1:50"),
            (0.15, 9.15,  0.11, "LAMINA: E-01"),
            (0.15, 8.65,  0.11, f"PLANO: {n_pil:02d}P"),
            # Col derecha de ficha
            (_RW*0.55+0.10, 9.65, 0.10, "REVISO: ________________"),
            (_RW*0.55+0.10, 9.15, 0.10, "CALCULO: _______________"),
            (_RW*0.55+0.10, 8.65, 0.10, "FECHA: _________________"),
            # Zona firmas (entre 6.0 y 7.5)
            (0.15, 7.15, 0.10, "REVISA Y APRUEBA:"),
            (0.15, 6.65, 0.10, "__________________________"),
            (0.15, 6.15, 0.09, "Ingeniero Proyectista"),
            # Zona normas (entre 4.5 y 6.0)
            (0.15, 5.65, 0.10, "NORMAS: NSR-10 / ACI 318-19"),
            (0.15, 5.15, 0.10, "ICONTEC NTC 1579 — Dibujo Tec."),
            # Footer (entre 0 y 4.5)
            (0.15, 4.15, 0.10, "EMPRESA:"),
            (0.15, 3.65, 0.10, "____________________________"),
            (0.15, 3.15, 0.09, "Nombre del proyecto"),
            (0.15, 2.65, 0.09, "Ciudad, Colombia"),
            (0.15, 2.15, 0.09, "____________________________"),
            (0.15, 1.65, 0.09, f"f'c={fc_dado:.0f}MPa | fy={fy_acero:.0f}MPa"),
            (0.15, 1.15, 0.09, f"V.total acero: {_kg_bx+_kg_by:.0f} kg"),
            (0.15, 0.65, 0.09, f"Vol. concreto: {_vol_c:.2f} m3"),
        ]
        for _tx, _ty, _th2, _tt in _ritems:
            msp.add_text(_tt, dxfattribs={"height":_th2,"layer":"ROTULO"}).set_placement(
                (_RX0+_tx, _RY0+_ty))

        #  DESCARGA DXF 
        bio_dxf = io.StringIO()
        doc_dxf.write(bio_dxf)
        col_exp2.download_button(_t("Descargar DXF","Download DXF"),
                                 data=bio_dxf.getvalue(),
                                 file_name="Plano_Encepado.dxf")

        #  DESCARGA PDF (via ezdxf drawing addon + matplotlib) 
        try:
            from ezdxf.addons.drawing import RenderContext, Frontend
            from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
            from ezdxf.addons.drawing.config import Configuration, BackgroundPolicy
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            _fig, _ax = plt.subplots(figsize=(33.11, 23.39))  # A1 inches
            _ax.set_aspect("equal")
            _ax.axis("off")
            _config_dp = Configuration.defaults().with_changes(background_policy=BackgroundPolicy.WHITE)
            _backend_pdf = MatplotlibBackend(_ax)
            Frontend(RenderContext(doc_dxf), _backend_pdf, config=_config_dp).draw_layout(msp, finalize=True)
            _bio_pdf = io.BytesIO()
            _fig.savefig(_bio_pdf, format="pdf", bbox_inches="tight", dpi=150)
            plt.close(_fig)
            _bio_pdf.seek(0)
            col_exp2.download_button(
                _t("Descargar PDF","Download PDF"),
                data=_bio_pdf.getvalue(),
                file_name="Plano_Encepado.pdf",
                mime="application/pdf")
        except Exception as _ep:
            col_exp2.info(f"PDF no disponible: {str(_ep)}")

    if _IFC_AVAILABLE:
        try:
            from ifc_export import ifc_dado_parametrico
            bio_ifc = ifc_dado_parametrico(
                B_dado_m=B_sugerido, L_dado_m=L_sugerido, H_dado_m=H_dado,
                df_pilotes=df_pilotes, D_pilote_m=D_pilote, embeb_m=embeb_pilote,
                fc_dado=fc_dado, c1_cm=c1_col, c2_cm=c2_col,
                As_x_cm2=As_x, As_y_cm2=As_y,
                db_mm=db_dado*10, recub_cm=recub_dado,
                sep_x_cm=sep_x_calc, sep_y_cm=sep_y_calc
            )
            col_exp3.download_button("IFC BIM 3D", data=bio_ifc.getvalue(), file_name="Modelo_Encepado.ifc")
        except Exception as e:
            col_exp3.button("IFC BIM (Error)", help=str(e))
    else:
        col_exp3.warning("ifcopenshell no disponible.")

#  REGISTRO DE DADOS (CUADRO DE MANDO) 
with st.sidebar:
    st.markdown("---")
    st.header(_t("Cuadro de Mando", "Dashboard"))
    
    if "registro_dados" not in st.session_state:
        st.session_state.registro_dados = []
        
    if st.button(_t("Guardar Diseño Actual", "Save Current Design")):
        st.session_state.registro_dados.append({
            "Plantilla": plantilla.split(" ")[0] + "P",
            "B x L x H": f"{B_sugerido:.1f}x{L_sugerido:.1f}x{H_dado:.1f}",
            "P_Max (kN)": f"{Pu_pilote_max:.0f}",
            "Punz. Col": "OK" if ok_punz else "X",
            "Punz. Pil": "OK" if ok_pilz else "X",
            "Cortante": "OK" if (ok_vx and ok_vy) else "X",
            "Flex. As": f"{As_x:.0f} / {As_y:.0f}"
        })
        st.success("Configuración Guardada")
        
    if len(st.session_state.registro_dados) > 0:
        df_reg = pd.DataFrame(st.session_state.registro_dados)
        st.dataframe(df_reg, use_container_width=True)
        if st.button(_t("Limpiar Registro", "Clear Registry")):
            st.session_state.registro_dados = []
            st.rerun()
