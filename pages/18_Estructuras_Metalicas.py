import streamlit as st
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import math
import io
import ezdxf
from docx import Document
from docx.shared import Inches, Pt
import datetime

# 
# IDIOMA GLOBAL
try:
    from normas_referencias import mostrar_referencias_norma
except ImportError:
    def mostrar_referencias_norma(*a, **kw): pass
lang = st.session_state.get("idioma", "Español")
def _t(es, en):
    return en if lang == "English" else es

norma_sel = st.session_state.get("norma_sel", "NSR-10 (Colombia)")

mostrar_referencias_norma(norma_sel, "estructuras_metalicas")
_PAIS_ISO = {"NSR-10 (Colombia)":"co","ACI 318-25 (EE.UU.)":"us","ACI 318-19 (EE.UU.)":"us","ACI 318-14 (EE.UU.)":"us","NEC-SE-HM (Ecuador)":"ec","E.060 (Perú)":"pe","NTC-EM (México)":"mx","COVENIN 1753-2006 (Venezuela)":"ve","NB 1225001-2020 (Bolivia)":"bo","CIRSOC 201-2025 (Argentina)":"ar"}
_iso = _PAIS_ISO.get(norma_sel, "un")

st.set_page_config(page_title=_t("Estructuras Metálicas", "Steel Structures"), layout="wide")
st.image(r"assets/steel_header_1773257206595.png", use_container_width=True)
st.title(_t("Diseño de Estructuras Metálicas", "Steel Structure Design"))
st.markdown(_t(f"Cálculo de Propiedades, Compresión y Flexión de Perfiles Laminados en Caliente y Conformados en Frío. Adaptado a la terminología de **{norma_sel}**.", 
               "Properties, Compression, and Flexure computation for Hot-Rolled and Cold-Formed Steel Sections."))

# 
# PIE DE PÁGINA / DERECHOS RESERVADOS
# 
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; color: gray; font-size: 11px;">
    © 2026 Todos los derechos reservados.<br>
    <b>Realizado por:</b><br>
    <br><br>
    <i> Nota Legal: Esta herramienta es un apoyo profesional. El uso de los resultados es responsabilidad exclusiva del ingeniero diseñador.</i>
</div>
""", unsafe_allow_html=True)

# Diccionario de Términos Regionales
if "NTC-EM" in norma_sel or "México" in norma_sel:
    term_W = "Viga IPR"
    term_Col = "Castillo / Columna"
    term_C = "Polín / Perfil C"
    term_Tubo = "HSS / PTR"
elif "CIRSOC 201" in norma_sel or "CIRSOC 301" in norma_sel or "Argentina" in norma_sel or "Bolivia" in norma_sel:
    term_W = "Perfil Doble T"
    term_Col = "Columna"
    term_C = "Costanera"
    term_Tubo = "Tubo Estructural"
elif "Perú" in norma_sel or "Ecuador" in norma_sel:
    term_W = "Viga Tipo I / H"
    term_Col = "Columna"
    term_C = "Correa / Perfil C"
    term_Tubo = "Tubo LAF"
else:
    term_W = "Perfil W / Tipo I"
    term_Col = "Columna / Poste"
    term_C = "Perfil C"
    term_Tubo = "Perfil Tubular Rect."

# 
# CONFIGURACIÓN MATERIAL Y UNIDADES
# 
st.sidebar.header(_t(" Preferencias", " Settings"))
sys_u = st.sidebar.radio(_t(" Sistema de Unidades", " Unit System"), 
                         ["Métrico (SI)", "Imperial (US)"], 
                         index=0 if st.session_state.get("unidades", "Métrico") == "Métrico" else 1, 
                         key="st_sys_u")
is_m = (sys_u == "Métrico (SI)")

f_mm = 1.0 if is_m else 25.4; f_m = 1.0 if is_m else 0.3048; f_MPa = 1.0 if is_m else 6.89476
f_kN = 1.0 if is_m else 4.44822; f_kNm = 1.0 if is_m else 1.35581; f_kg = 1.0 if is_m else 0.453592

t_mm = "mm" if is_m else "in"; t_m = "m" if is_m else "ft"; t_MPa = "MPa" if is_m else "ksi"
t_kN = "kN" if is_m else "kip"; t_kNm = "kN-m" if is_m else "kip-ft"; t_kg = "kg" if is_m else "lb"

def show_mm(v): return v / f_mm
def show_m(v): return v / f_m
def show_MPa(v): return v / f_MPa
def show_kN(v): return v / f_kN
def show_kNm(v): return v / f_kNm
def show_kg(v): return v / f_kg
def show_mm2(v): return v / (f_mm**2)
def show_mm3(v): return v / (f_mm**3)
def show_mm4(v): return v / (f_mm**4)
def show_kgm(v): return v / (f_kg / f_m)

st.sidebar.markdown("---")
st.sidebar.header(_t(" Materiales", " Materials"))
Fy_inp = st.sidebar.number_input(_t(f"Fluencia Fy [{t_MPa}]", f"Yield Fy [{t_MPa}]"), value=250.0 if is_m else 36.0, step=10.0 if is_m else 1.0, key="st_fy_i")
Fu_inp = st.sidebar.number_input(_t(f"Último Fu [{t_MPa}]", f"Ultimate Fu [{t_MPa}]"), value=400.0 if is_m else 58.0, step=10.0 if is_m else 1.0, key="st_fu_i")
E_inp = st.sidebar.number_input(_t(f"Elasticidad E [{t_MPa}]", f"Modulus E [{t_MPa}]"), value=200000.0 if is_m else 29000.0, step=1000.0 if is_m else 1000.0, key="st_E_i")

Fy_adm = Fy_inp * f_MPa
Fu_adm = Fu_inp * f_MPa
E_steel = E_inp * f_MPa
G_steel = E_steel / (2 * (1 + 0.3))  # Shear modulus approx
peso_esp_acero = 7850.0  # kg/m3

# Variables globales para despiece y costos (se guardan en session_state)
if "steel_despiece" not in st.session_state:
    st.session_state.steel_despiece = []
if "steel_costo_total" not in st.session_state:
    st.session_state.steel_costo_total = 0.0

def agregar_al_despiece(tipo, seccion, longitud_m, peso_kg, observacion, memoria=""):
    st.session_state.steel_despiece.append({
        "Tipo": tipo,
        "Sección": seccion,
        "Longitud (m)": longitud_m,
        "Peso (kg)": peso_kg,
        "Observación": observacion,
        "Memoria": memoria
    })
    st.session_state.steel_costo_total += peso_kg

# Funciones de dibujo 2D (se mantienen)
def plot_W(d, bf, tw, tf):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((-bf/2, d/2 - tf), bf, tf, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((-bf/2, -d/2), bf, tf, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((-tw/2, -d/2 + tf), tw, d - 2*tf, facecolor='steelblue', edgecolor='black'))
    
    ax.annotate('', xy=(bf/2*1.2, d/2), xytext=(bf/2*1.2, -d/2), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(bf/2*1.25, 0, f"d={d:.1f}", color='red', va='center')
    ax.annotate('', xy=(-bf/2, d/2*1.15), xytext=(bf/2, d/2*1.15), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(0, d/2*1.22, f"bf={bf:.1f}", color='red', ha='center')
    
    bbox_props = dict(boxstyle="round,pad=0.2", fc="#2b5b84", ec="none")
    ax.text(0, 0, f"tw={tw:.1f}", color='white', ha='center', va='center', fontsize=9, rotation=90, bbox=bbox_props)
    ax.text(0, d/2-tf/2, f"tf={tf:.1f}", color='white', ha='center', va='center', fontsize=9, bbox=bbox_props)

    lim_max = max(d, bf)
    ax.set_xlim(-lim_max*0.8, lim_max*0.8); ax.set_ylim(-lim_max*0.7, lim_max*0.7)
    ax.set_aspect('equal')
    ax.axis('off'); return fig

def plot_3d_W(d, bf, tw, tf, L_m):
    """Crea un modelo 3D del perfil W con Plotly."""
    d_cm = d / 10; bf_cm = bf / 10; tw_cm = tw / 10; tf_cm = tf / 10; L_cm = L_m * 100
    xs = [bf_cm/2, bf_cm/2, tw_cm/2, tw_cm/2, bf_cm/2, bf_cm/2,
          -bf_cm/2, -bf_cm/2, -tw_cm/2, -tw_cm/2, -bf_cm/2, -bf_cm/2, bf_cm/2]
    zs = [d_cm/2, d_cm/2-tf_cm, d_cm/2-tf_cm, -d_cm/2+tf_cm, -d_cm/2+tf_cm, -d_cm/2,
          -d_cm/2, -d_cm/2+tf_cm, -d_cm/2+tf_cm, d_cm/2-tf_cm, d_cm/2-tf_cm, d_cm/2, d_cm/2]
    fig = go.Figure()
    fig.add_trace(go.Scatter3d(x=xs, y=[0]*len(xs), z=zs, mode='lines', line=dict(color='darkred', width=4), name='Frente'))
    fig.add_trace(go.Scatter3d(x=xs, y=[L_cm]*len(xs), z=zs, mode='lines', line=dict(color='darkred', width=4), name='Posterior'))
    for i in range(12):
        fig.add_trace(go.Scatter3d(x=[xs[i], xs[i]], y=[0, L_cm], z=[zs[i], zs[i]], mode='lines', line=dict(color='gray', width=3), showlegend=False))
    fig.update_layout(scene=dict(aspectmode='data', xaxis_title='Ancho (cm)', yaxis_title='Largo (cm)', zaxis_title='Alto (cm)'), margin=dict(l=0, r=0, b=0, t=0), height=400, showlegend=False)
    return fig

def plot_T(d, bf, tw, tf):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((-bf/2, d - tf), bf, tf, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((-tw/2, 0), tw, d - tf, facecolor='steelblue', edgecolor='black'))
    
    ax.annotate('', xy=(bf/2*1.2, d), xytext=(bf/2*1.2, 0), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(bf/2*1.25, d/2, f"d={d:.1f}", color='red', va='center')
    ax.annotate('', xy=(-bf/2, d*1.15), xytext=(bf/2, d*1.15), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(0, d*1.22, f"bf={bf:.1f}", color='red', ha='center')

    bbox_props = dict(boxstyle="round,pad=0.2", fc="#2b5b84", ec="none")
    ax.text(0, (d-tf)/2, f"tw={tw:.1f}", color='white', ha='center', va='center', fontsize=9, rotation=90, bbox=bbox_props)
    ax.text(0, d-tf/2, f"tf={tf:.1f}", color='white', ha='center', va='center', fontsize=9, bbox=bbox_props)

    lim_max = max(d, bf)
    ax.set_xlim(-lim_max*0.8, lim_max*0.8); ax.set_ylim(-d*0.2, d*1.3)
    ax.set_aspect('equal')
    ax.axis('off'); return fig

def plot_L(h, b, t):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((0, 0), t, h, facecolor='steelblue', edgecolor='black'))
    ax.add_patch(patches.Rectangle((t, 0), b-t, t, facecolor='steelblue', edgecolor='black'))
    
    ax.annotate('', xy=(-h*0.15, h), xytext=(-h*0.15, 0), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(-h*0.2, h/2, f"h={h:.1f}", color='red', ha='right', va='center')
    ax.annotate('', xy=(0, -b*0.15), xytext=(b, -b*0.15), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(b/2, -b*0.25, f"b={b:.1f}", color='red', ha='center', va='top')

    bbox_props = dict(boxstyle="round,pad=0.2", fc="#2b5b84", ec="none")
    ax.text(t/2, h/2, f"t={t:.1f}", color='white', ha='center', va='center', fontsize=9, rotation=90, bbox=bbox_props)
    ax.text(b/2, t/2, f"t={t:.1f}", color='white', ha='center', va='center', fontsize=9, bbox=bbox_props)

    lim_max = max(b, h)
    ax.set_xlim(-lim_max*0.4, lim_max*1.2); ax.set_ylim(-lim_max*0.4, lim_max*1.2)
    ax.set_aspect('equal')
    ax.axis('off'); return fig

def plot_C(h, b, d_lip, t):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((0, 0), t, h, facecolor='darkgray', edgecolor='black'))
    ax.add_patch(patches.Rectangle((0, h-t), b, t, facecolor='darkgray', edgecolor='black'))
    ax.add_patch(patches.Rectangle((0, 0), b, t, facecolor='darkgray', edgecolor='black'))
    if d_lip > 0:
        ax.add_patch(patches.Rectangle((b-t, h-d_lip), t, d_lip, facecolor='darkgray', edgecolor='black'))
        ax.add_patch(patches.Rectangle((b-t, 0), t, d_lip, facecolor='darkgray', edgecolor='black'))
        
    ax.annotate('', xy=(-h*0.15, h), xytext=(-h*0.15, 0), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(-h*0.2, h/2, f"h={h:.1f}", color='red', ha='right', va='center')
    ax.annotate('', xy=(0, -b*0.2), xytext=(b, -b*0.2), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(b/2, -b*0.3, f"b={b:.1f}", color='red', ha='center', va='top')

    bbox_props = dict(boxstyle="round,pad=0.2", fc="#5a5a5a", ec="none")
    ax.text(t/2, h/2, f"t={t:.1f}", color='white', ha='center', va='center', fontsize=9, rotation=90, bbox=bbox_props)

    ax.set_xlim(-b*0.5, b*1.5); ax.set_ylim(-h*0.4, h*1.4)
    ax.set_aspect('equal')
    ax.axis('off'); return fig

def plot_Tubo(h, b, t):
    fig, ax = plt.subplots(figsize=(3, 3))
    ax.add_patch(patches.Rectangle((0, 0), b, h, facecolor='none', edgecolor='black', lw=3))
    ax.add_patch(patches.Rectangle((t, t), b-2*t, h-2*t, facecolor='white', edgecolor='black', lw=2))
    ax.add_patch(patches.Rectangle((0, 0), b, h, facecolor='darkgray', alpha=0.5))
    
    ax.annotate('', xy=(-b*0.15, h), xytext=(-b*0.15, 0), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(-b*0.2, h/2, f"h={h:.1f}", color='red', ha='right', va='center')
    ax.annotate('', xy=(0, -h*0.15), xytext=(b, -h*0.15), arrowprops=dict(arrowstyle='<->', color='red'))
    ax.text(b/2, -h*0.25, f"b={b:.1f}", color='red', ha='center', va='top')

    bbox_props = dict(boxstyle="round,pad=0.2", fc="#5a5a5a", ec="none")
    ax.text(b/2, t/2, f"t={t:.1f}", color='white', ha='center', va='center', fontsize=9, bbox=bbox_props)

    ax.set_xlim(-b*0.4, b*1.2); ax.set_ylim(-h*0.4, h*1.2)
    ax.set_aspect('equal')
    ax.axis('off'); return fig

# Pestañas principales
tab_P, tab_C, tab_F, tab_CF, tab_E = st.tabs([
    _t(f"1. Propiedades {term_W}", "1. W-Shape Properties"),
    _t("2. LAMINADOS: Compresión", "2. HOT-ROLLED: Compression"),
    _t(f"3. LAMINADOS: Flexión ({term_W})", "3. HOT-ROLLED: Flexure"),
    _t("4. CONFORMADOS EN FRÍO", "4. COLD-FORMED"),
    _t(" Exportaciones Globales", " Global Exports")
])

# 
# TAB 1: PROPIEDADES PERFIL W
# 
with tab_P:
    st.header(_t(f"Propiedades de Sección - {term_W}", f"Section Properties - {term_W}"))
    p1, p2, p3 = st.columns([1,2,1])
    with p1:
        st.subheader(f"Dimensiones [{t_mm}]")
        dw_i = st.number_input(f"Peralte total (d) [{t_mm}]", 1.0, 1500.0, 300.0 if is_m else 12.0, 1.0 if is_m else 0.5, key="st_p_d")
        bfw_i = st.number_input(f"Ancho patín (bf) [{t_mm}]", 1.0, 800.0, 150.0 if is_m else 6.0, 1.0 if is_m else 0.5, key="st_p_bf")
        tfw_i = st.number_input(f"Espesor patín (tf) [{t_mm}]", 0.1, 100.0, 10.0 if is_m else 0.5, 0.5 if is_m else 0.1, key="st_p_tf")
        tww_i = st.number_input(f"Espesor alma (tw) [{t_mm}]", 0.1, 100.0, 6.0 if is_m else 0.25, 0.5 if is_m else 0.05, key="st_p_tw")
        L_ref_i = st.number_input(f"Longitud de referencia (L) [{t_m}]", 0.1, 20.0, 3.0 if is_m else 10.0, 0.5 if is_m else 1.0, key="st_p_L")
        
        dw = dw_i * f_mm; bfw = bfw_i * f_mm; tfw = tfw_i * f_mm; tww = tww_i * f_mm; L_ref = L_ref_i * f_m
    with p2:
        # Cálculos básicos
        A_w = 2 * (bfw * tfw) + (dw - 2*tfw)*tww
        Ix_w = (tww*(dw - 2*tfw)**3)/12.0 + 2 * ((bfw*tfw**3)/12.0 + (bfw*tfw)*(dw/2.0 - tfw/2.0)**2)
        Iy_w = 2*((tfw*bfw**3)/12.0) + ((dw-2*tfw)*tww**3)/12.0
        rx_w = math.sqrt(Ix_w / A_w)
        ry_w = math.sqrt(Iy_w / A_w)
        Sx_w = Ix_w / (dw/2.0)
        Zx_w = bfw*tfw*(dw-tfw) + (tww*(dw-2*tfw)**2)/4.0
        peso_lin = (A_w / 1e6) * peso_esp_acero
        peso_total = peso_lin * L_ref

        st.write(f"**Área (A):** {show_mm2(A_w):.2f} {t_mm}²")
        st.write(f"**Momento Inercia Eje Fuerte (Ix):** {show_mm4(Ix_w):.2f} {t_mm}⁴")
        st.write(f"**Momento Inercia Eje Débil (Iy):** {show_mm4(Iy_w):.2f} {t_mm}⁴")
        st.write(f"**Radio de Giro (rx):** {show_mm(rx_w):.2f} {t_mm}  |  **(ry):** {show_mm(ry_w):.2f} {t_mm}")
        st.write(f"**Módulo Plástico Fuerte (Zx):** {show_mm3(Zx_w):.2f} {t_mm}³")
        st.write(f" **Peso Lineal Estimado:** {show_kgm(peso_lin):.2f} {t_kg}/{t_m}")
        st.write(f"**Peso total para L={L_ref_i:.2f} {t_m}:** {show_kg(peso_total):.2f} {t_kg}")

        # Opción para agregar al despiece
        if st.button(_t("Agregar este perfil al despiece", "Add this profile to the list"), key="btn_add_prop_w"):
            mem_text = f"Área (A): {show_mm2(A_w):.1f} {t_mm}²\nInercia (Ix, Iy): {show_mm4(Ix_w):.1f}, {show_mm4(Iy_w):.1f} {t_mm}⁴\nMódulo Plástico Fuerte (Zx): {show_mm3(Zx_w):.1f} {t_mm}³\nRadios de Giro (rx, ry): {show_mm(rx_w):.1f}, {show_mm(ry_w):.1f} {t_mm}\nPeso propio estimado: {show_kgm(peso_lin):.2f} {t_kg}/{t_m}\nPeso componente [{L_ref:.2f} m]: {peso_total:.2f} {t_kg}"
            agregar_al_despiece(
                tipo=f"{term_W} (propiedades)",
                seccion=f"{dw}x{bfw}x{tww}x{tfw} mm",
                longitud_m=L_ref,
                peso_kg=peso_total,
                observacion=_t("Perfil W calculado", "W-shape calculated"),
                memoria=mem_text
            )
            st.success(_t("Perfil agregado al despiece", "Profile added to list"))

    with p3:
        st.pyplot(plot_W(dw, bfw, tww, tfw), use_container_width=False)
        st.markdown("---")
        st.subheader(_t("Visualización 3D", "3D Visualization"))
        fig3d = plot_3d_W(dw, bfw, tww, tfw, L_ref)
        st.plotly_chart(fig3d, use_container_width=True)

# 
# TAB 2: LAMINADOS - COMPRESIÓN
# 
with tab_C:
    st.header(_t(f"Resistencia a Compresión Axial", f"Axial Compression Resistance"))
    comp_opts = [term_W, "Perfil T", "Perfil L (Angular)"]
    tipo_comp = st.selectbox("Seleccione Perfil Laminado en Caliente a evaluar:", comp_opts,
                             index=comp_opts.index(st.session_state.get("st_c_tipo", comp_opts[0])),
                             key="st_c_tipo")
    col_c1, col_c2, col_c3 = st.columns([1,1.5,1.5])
    with col_c1:
        st.subheader("Datos Compresión")
        L_c_i = st.number_input(f"Longitud No Arriostrada (Lc) [{t_m}]", 0.1, 20.0, 3.0 if is_m else 10.0, 0.5 if is_m else 1.0, key="st_c_L")
        P_u_i = st.number_input(f"Carga Axial Requerida Pu [{t_kN}]", 10.0, 10000.0, 500.0 if is_m else 110.0, 10.0 if is_m else 5.0, key="st_c_Pu")
        L_c_m = L_c_i * f_m
        P_u = P_u_i * f_kN
    with col_c2:
        if tipo_comp == term_W:
            d_c_i = st.number_input(f"{tipo_comp} - Peralte d [{t_mm}]", 1.0, 1000.0, 300.0 if is_m else 12.0, key="st_cw_d")
            bf_c_i = st.number_input(f"{tipo_comp} - Ancho bf [{t_mm}]", 1.0, 500.0, 150.0 if is_m else 6.0, key="st_cw_bf")
            tf_c_i = st.number_input(f"{tipo_comp} - tf [{t_mm}]", 0.1, 50.0, 10.0 if is_m else 0.5, key="st_cw_tf")
            tw_c_i = st.number_input(f"{tipo_comp} - tw [{t_mm}]", 0.1, 50.0, 6.0 if is_m else 0.25, key="st_cw_tw")
            d_c = d_c_i*f_mm; bf_c = bf_c_i*f_mm; tf_c = tf_c_i*f_mm; tw_c = tw_c_i*f_mm
            A_c = 2 * (bf_c * tf_c) + (d_c - 2*tf_c)*tw_c
            Iy_c = 2*((tf_c*bf_c**3)/12.0) + ((d_c-2*tf_c)*tw_c**3)/12.0
            r_c = math.sqrt(Iy_c / A_c)
            fig_c = plot_W(d_c, bf_c, tw_c, tf_c)
            peso_lin_c = (A_c / 1e6) * peso_esp_acero
        elif tipo_comp == "Perfil T":
            d_c_i = st.number_input(f"Perfil T - Peralte d [{t_mm}]", 1.0, 500.0, 150.0 if is_m else 6.0, key="st_ct_d")
            bf_c_i = st.number_input(f"Perfil T - Ancho bf [{t_mm}]", 1.0, 500.0, 150.0 if is_m else 6.0, key="st_ct_bf")
            tf_c_i = st.number_input(f"Perfil T - tf [{t_mm}]", 0.1, 50.0, 10.0 if is_m else 0.5, key="st_ct_tf")
            tw_c_i = st.number_input(f"Perfil T - tw [{t_mm}]", 0.1, 50.0, 8.0 if is_m else 0.3, key="st_ct_tw")
            d_c = d_c_i*f_mm; bf_c = bf_c_i*f_mm; tf_c = tf_c_i*f_mm; tw_c = tw_c_i*f_mm
            A_c = (bf_c * tf_c) + (d_c - tf_c)*tw_c
            Iy_c = ((tf_c*bf_c**3)/12.0) + ((d_c-tf_c)*tw_c**3)/12.0
            r_c = math.sqrt(Iy_c / A_c)
            fig_c = plot_T(d_c, bf_c, tw_c, tf_c)
            peso_lin_c = (A_c / 1e6) * peso_esp_acero
        else:  # Perfil L
            b_c_i = st.number_input(f"Perfil L - Altura/Base b [{t_mm}]", 1.0, 300.0, 100.0 if is_m else 4.0, key="st_cl_b")
            t_c_i = st.number_input(f"Perfil L - Espesor t [{t_mm}]", 0.1, 30.0, 10.0 if is_m else 0.5, key="st_cl_t")
            b_c = b_c_i*f_mm; t_c = t_c_i*f_mm
            A_c = (2 * b_c - t_c) * t_c
            r_c = 0.2 * b_c
            fig_c = plot_L(b_c, b_c, t_c)
            peso_lin_c = (A_c / 1e6) * peso_esp_acero
    with col_c3:
        Lc_mm = L_c_m * 1000.0
        esbeltez = Lc_mm / r_c
        Fe = (math.pi**2 * E_steel) / (max(esbeltez, 1.0)**2)
        if esbeltez <= 4.71 * math.sqrt(E_steel/Fy_adm):
            Fcr = (0.658**(Fy_adm/Fe)) * Fy_adm
        else:
            Fcr = 0.877 * Fe
        Pn = Fcr * A_c / 1000.0
        phi_Pn = 0.90 * Pn
        st.markdown(f"**Relación de Esbeltez L/r_min:** {esbeltez:.2f} (Límite ≈ 200)")
        st.markdown(f"**Esfuerzo Crítico (Fcr):** {show_MPa(Fcr):.2f} {t_MPa}")
        st.markdown(f"**Resistencia Axial de Diseño ($\phi P_n$):** <span style='color:blue;font-size:22px'>{show_kN(phi_Pn):.1f} {t_kN}</span>", unsafe_allow_html=True)
        if P_u <= phi_Pn:
            st.success(f" ¡Aprobado! (FS = {phi_Pn/P_u:.2f})")
        else:
            st.error("No Aprobado por Compresión Axial / Pandeo.")
            st.warning("**Recomendación:** Aumente el área de la sección incrementando **bf**, **d**, o sus espesores (**tf**, **tw**). También puede reducir la longitud no arriostrada **Lc**.")
        st.pyplot(fig_c, use_container_width=False)

        # Agregar al despiece si se desea
        peso_total_c = peso_lin_c * L_c_m
        if st.button(_t("Agregar este elemento al despiece", "Add this element to the list")):
            mem_text = f"Carga Axial Requerida (Pu): {P_u_i:.1f} {t_kN}\nLongitud No Arriostrada (Lc): {L_c_i:.2f} {t_m}\nRelación de Esbeltez Global (L/r_min): {esbeltez:.1f}\nEsfuerzo Crítico por Pandeo (Fcr): {show_MPa(Fcr):.1f} {t_MPa}\nÁrea Gruesa (Ag): {show_mm2(A_c):.1f} {t_mm}²\nResistencia Nominal (Pn): {show_kN(Pn):.1f} {t_kN}\nResistencia de Diseño (φPn): {show_kN(phi_Pn):.1f} {t_kN}\nEstado de Evaluación: {'CUMPLE' if P_u <= phi_Pn else 'FALLA'} (Factor de Seguridad LRFD = {phi_Pn/max(P_u, 0.001):.2f})"
            agregar_al_despiece(
                tipo=f"{tipo_comp} (compresión)",
                seccion=f"{d_c_i if 'd_c_i' in locals() else b_c_i}x{bf_c_i if 'bf_c_i' in locals() else t_c_i} [{t_mm}]",
                longitud_m=L_c_m,
                peso_kg=peso_total_c,
                observacion=_t(f"Resistencia φPn={show_kN(phi_Pn):.1f} {t_kN} vs Pu={P_u_i:.1f} {t_kN}", f"Strength φPn={show_kN(phi_Pn):.1f} {t_kN} vs Pu={P_u_i:.1f} {t_kN}"),
                memoria=mem_text
            )
            st.success(_t("Elemento agregado", "Element added"))

# 
# TAB 3: LAMINADOS - FLEXIÓN (PERFIL W)
# 
with tab_F:
    st.header(_t(f"Resistencia a Flexión - {term_W}", f"Flexural Resistance - {term_W}"))
    st.write(_t("Evaluación del momento flector considerando Pandeo Lateral Torsional (LTB).", "Bending moment evaluation considering Lateral Torsional Buckling (LTB)."))
    f_c1, f_c2, f_c3 = st.columns(3)
    with f_c1:
        Mu_i = st.number_input(f"Momento Último Solicitante Mu [{t_kNm}]", 10.0, 5000.0, 150.0 if is_m else 110.0, 10.0 if is_m else 5.0, key="st_f_Mu")
        Lb_i = st.number_input(f"Longitud No Arriostrada Lateramente Lb [{t_m}]", 0.1, 20.0, 2.0 if is_m else 6.0, 0.5 if is_m else 1.0, key="st_f_Lb")
        Cb = st.number_input("Coeficiente de Momento (Cb)", 1.0, 3.0, st.session_state.get("st_f_Cb", 1.0), 0.1, key="st_f_Cb")
        Mu = Mu_i * f_kNm
        Lb_m = Lb_i * f_m
    with f_c2:
        d_f_i = st.number_input(f"Peralte d [{t_mm}]", 1.0, 1500.0, 300.0 if is_m else 12.0, 1.0 if is_m else 0.5, key="st_fw_d")
        bf_f_i = st.number_input(f"Ancho bf [{t_mm}]", 1.0, 800.0, 150.0 if is_m else 6.0, 1.0 if is_m else 0.5, key="st_fw_bf")
        tf_f_i = st.number_input(f"Espesor tf [{t_mm}]", 0.1, 50.0, 10.0 if is_m else 0.5, 0.5 if is_m else 0.1, key="st_fw_tf")
        tw_f_i = st.number_input(f"Espesor tw [{t_mm}]", 0.1, 50.0, 6.0 if is_m else 0.25, 0.5 if is_m else 0.05, key="st_fw_tw")
        L_viga_i = st.number_input(f"Longitud de la viga (L) [{t_m}]", 0.1, 20.0, 6.0 if is_m else 20.0, 0.5 if is_m else 1.0, key="st_f_L")
        
        d_f=d_f_i*f_mm; bf_f=bf_f_i*f_mm; tf_f=tf_f_i*f_mm; tw_f=tw_f_i*f_mm; L_viga=L_viga_i*f_m
    with f_c3:
        # Propiedades para LTB
        h0 = d_f - tf_f
        cw = ((h0**2)*tf_f*bf_f**3)/24.0
        Iy_f = 2*((tf_f*bf_f**3)/12.0) + ((d_f-2*tf_f)*tw_f**3)/12.0
        J_tor = (2*bf_f*tf_f**3 + (d_f-tf_f)*tw_f**3)/3.0
        r_ts = bf_f / min(math.sqrt(12*(1+(1/6)*(d_f*tw_f)/(bf_f*tf_f))), 1000)
        Zx_f = bf_f*tf_f*(d_f-tf_f) + (tw_f*(d_f-2*tf_f)**2)/4.0
        Sx_f = (2 * ((bf_f*tf_f**3)/12.0 + (bf_f*tf_f)*(d_f/2.0 - tf_f/2.0)**2) + (tw_f*(d_f - 2*tf_f)**3)/12.0) / (d_f/2.0)
        Lp = 1.76 * r_ts * math.sqrt(E_steel/Fy_adm) / 1000.0
        Lr = np.pi * r_ts * math.sqrt(E_steel/(0.7*Fy_adm)) / 1000.0
        Mp = Fy_adm * Zx_f / 1e6
        st.write(f"**Límite Fluencia (Lp):** {show_m(Lp):.2f} {t_m}")
        st.write(f"**Límite Pandeo (Lr):** {show_m(Lr):.2f} {t_m}")
        st.write(f"**Momento Plástico (Mp):** {show_kNm(Mp):.1f} {t_kNm}")
        Lbmm = Lb_m * 1000.0
        if Lb_m <= Lp:
            Mn = Mp
            estado = "Fluencia (Zona Plástica)"
        elif Lb_m <= Lr:
            Mn = Cb * (Mp - (Mp - 0.7*Fy_adm*Sx_f/1e6) * ((Lb_m - Lp)/(Lr - Lp)))
            Mn = min(Mn, Mp)
            estado = "LTB Inelástico"
        else:
            Fcr_ltb = (Cb * np.pi**2 * E_steel / (max(Lbmm/r_ts, 1.0))**2) * math.sqrt(1 + 0.078*(J_tor/Sx_f)*(Lbmm/r_ts)**2)
            Mn = Fcr_ltb * Sx_f / 1e6
            Mn = min(Mn, Mp)
            estado = "LTB Elástico"
        phi_Mn = 0.90 * Mn
        st.markdown(f"**Zona Comportamiento:** {estado}")
        st.markdown(f"**Capacidad Momento ($\phi M_n$):** <span style='color:blue;font-size:22px'>{show_kNm(phi_Mn):.1f} {t_kNm}</span>", unsafe_allow_html=True)
        if Mu <= phi_Mn:
            st.success(f" ¡Viga Cumple a Flexión! (FS={phi_Mn/Mu:.2f})")
        else:
            st.error("Viga falla por Flexión / LTB.")
            st.warning("**Recomendación:** Para flexión, lo más eficiente es aumentar la inercia incrementando el peralte **d**. También ayuda ensanchar **bf** y aumentar espesores, o arriostrar más seguido reduciendo **Lb**.")
            
        # Add pyplot W-Shape here to visualize real-time
        fig_f = plot_W(d_f, bf_f, tw_f, tf_f)
        st.pyplot(fig_f, use_container_width=False)

        # Peso y despiece
        A_f = 2 * (bf_f * tf_f) + (d_f - 2*tf_f)*tw_f
        peso_lin_f = (A_f / 1e6) * peso_esp_acero
        peso_total_f = peso_lin_f * L_viga
        st.write(f" **Peso estimado de la viga:** {show_kg(peso_total_f):.2f} {t_kg}")

        if st.button(_t("Agregar esta viga al despiece", "Add this beam to the list")):
            mem_text = f"Momento Flector Solicitante (Mu): {Mu_i:.1f} {t_kNm}\nLongitud Total (L): {L_viga_i:.2f} {t_m}\nLongitud No Arriostrada (Lb): {show_m(Lb_m):.2f} {t_m}\nLímites LTB: Lp={show_m(Lp):.2f} {t_m}, Lr={show_m(Lr):.2f} {t_m}\nMomento Plástico (Mp): {show_kNm(Mp):.1f} {t_kNm}\nMódulo de Sección Elástico (Sx): {show_mm3(Sx_f):.1f} {t_mm}³\nZona de Comportamiento LTB: {estado}\nResistencia Nominal Equivalente (Mn): {show_kNm(Mn):.1f} {t_kNm}\nResistencia de Diseño a Flexión (φMn): {show_kNm(phi_Mn):.1f} {t_kNm}\nEstado de Evaluación: {'CUMPLE' if Mu <= phi_Mn else 'FALLA'} (Factor de Seguridad LRFD = {phi_Mn/max(Mu, 0.001):.2f})"
            agregar_al_despiece(
                tipo=f"{term_W} (flexión)",
                seccion=f"{d_f_i}x{bf_f_i}x{tw_f_i}x{tf_f_i} [{t_mm}]",
                longitud_m=L_viga,
                peso_kg=peso_total_f,
                observacion=_t(f"φMn={show_kNm(phi_Mn):.1f} {t_kNm} vs Mu={Mu_i:.1f} {t_kNm}", f"φMn={show_kNm(phi_Mn):.1f} {t_kNm} vs Mu={Mu_i:.1f} {t_kNm}"),
                memoria=mem_text
            )
            st.success(_t("Viga agregada al despiece", "Beam added to list"))

# 
# TAB 4: CONFORMADOS EN FRÍO (Cold-Formed)
# 
with tab_CF:
    st.header(_t("Perfiles Conformados en Frío (Cold-Formed)", "Cold-Formed Sections"))
    cf_opts = [f"{term_C} con labios (rígido)", f"{term_C} sin labios (U)", term_Tubo]
    tipo_cf = st.selectbox("Seleccione Perfil Conformado en Frío:", cf_opts,
                           index=cf_opts.index(st.session_state.get("st_cf_tipo", cf_opts[0])),
                           key="st_cf_tipo")
    col_cf1, col_cf2 = st.columns([1,2])
    with col_cf1:
        h_cf_i = st.number_input(f"Altura total h [{t_mm}]", 1.0, 400.0, 150.0 if is_m else 6.0, 1.0 if is_m else 0.5, key="st_cf_h")
        b_cf_i = st.number_input(f"Ancho base b [{t_mm}]", 1.0, 200.0, 50.0 if is_m else 2.0, 1.0 if is_m else 0.5, key="st_cf_b")
        t_cf_i = st.number_input(f"Espesor de lámina t [{t_mm}]", 0.1, 10.0, 2.0 if is_m else 0.08, 0.1 if is_m else 0.01, key="st_cf_t")
        L_cf_i = st.number_input(f"Longitud [{t_m}]", 0.1, 20.0, 2.0 if is_m else 6.0, 0.5 if is_m else 1.0, key="st_cf_L")
        h_cf = h_cf_i * f_mm; b_cf = b_cf_i * f_mm; t_cf = t_cf_i * f_mm; L_cf = L_cf_i * f_m
        
        if "labios" in tipo_cf:
            d_lip_i = st.number_input(f"Pestaña/Labio d [{t_mm}]", 0.0, 50.0, 15.0 if is_m else 0.6, key="st_cf_l")
            d_lip = d_lip_i * f_mm
            A_cf = (h_cf + 2*b_cf + 2*d_lip - 4*t_cf) * t_cf
            fig_cf = plot_C(h_cf, b_cf, d_lip, t_cf)
        elif tipo_cf == term_Tubo:
            d_lip = 0
            A_cf = 2*(h_cf + b_cf - 2*t_cf) * t_cf
            fig_cf = plot_Tubo(h_cf, b_cf, t_cf)
        else:
            d_lip = 0
            A_cf = (h_cf + 2*b_cf - 2*t_cf) * t_cf
            fig_cf = plot_C(h_cf, b_cf, 0, t_cf)
            
        peso_lin_cf = (A_cf / 1e6) * peso_esp_acero
        peso_total_cf = peso_lin_cf * L_cf
        st.write(f"**Área Bruta (A):** {show_mm2(A_cf):.2f} {t_mm}²")
        st.write(f" **Peso Lineal Estimado:** {show_kgm(peso_lin_cf):.2f} {t_kg}/{t_m}")
        st.write(f"**Peso total:** {show_kg(peso_total_cf):.2f} {t_kg}")
    with col_cf2:
        st.subheader("Verificación de Compresión")
        Pu_cf_i = st.number_input(f"Carga Axial Pu [{t_kN}]", 5.0, 500.0, 50.0 if is_m else 11.0, 1.0 if is_m else 0.5, key="st_cf_Pu")
        Pu_cf = Pu_cf_i * f_kN
        w_t_ratio = max(h_cf/t_cf, b_cf/t_cf)
        if w_t_ratio > 200:
            st.warning(" Relación ancho/espesor > 200. Riesgo de pandeo local muy severo.")
            Ae = 0.5 * A_cf
        elif w_t_ratio > 50:
            Ae = 0.8 * A_cf
        else:
            Ae = A_cf
        st.write(f"**Área Efectiva por Pandeo Local (Ae):** aprox {show_mm2(Ae):.2f} {t_mm}²")
        if tipo_cf == term_Tubo:
            r_cf = b_cf * 0.38
        else:
            r_cf = b_cf * 0.30
        esbeltez_cf = (L_cf*1000.0) / r_cf
        st.write(f"**Relación Esbeltez Global L/r:** {esbeltez_cf:.2f}")
        Fe_cf = (math.pi**2 * E_steel) / (max(esbeltez_cf, 1)**2)
        if esbeltez_cf <= 4.71 * math.sqrt(E_steel/Fy_adm):
            Fcr_cf = (0.658**(Fy_adm/Fe_cf)) * Fy_adm
        else:
            Fcr_cf = 0.877 * Fe_cf
        Pn_cf = (Ae * Fcr_cf) / 1000.0
        phi_Pn_cf = 0.85 * Pn_cf
        st.markdown(f"**Resistencia Axial Compresión ($\phi P_n$):** <span style='color:blue;font-size:22px'>{show_kN(phi_Pn_cf):.1f} {t_kN}</span>", unsafe_allow_html=True)
        if Pu_cf <= phi_Pn_cf:
            st.success(" ¡El perfil conformado en frío CUMPLE!")
        else:
            st.error("El perfil falla por compresión / pandeo global o local.")
            st.warning("**Recomendación:** Aumente el espesor de la lámina **t**, agrande las dimensiones **h** o **b**, o incluya labios rigidizadores para mitigar el pandeo local.")
        st.pyplot(fig_cf, use_container_width=False)

        if st.button(_t("Agregar este perfil al despiece", "Add this profile to the list"), key="btn_add_cf_section"):
            mem_text = f"Carga Axial Solicitante (Pu): {Pu_cf_i:.1f} {t_kN}\nÁrea Bruta de Sección (Ag): {show_mm2(A_cf):.1f} {t_mm}²\nÁrea Efectiva Computada por Pandeo Local (Ae): {show_mm2(Ae):.1f} {t_mm}²\nRadio de Giro Pésimo Estimado (r): {show_mm(r_cf):.1f} {t_mm}\nEsbeltez Global (L/r): {esbeltez_cf:.1f}\nEsfuerzo Crítico Relacionado (Fcr): {show_MPa(Fcr_cf):.1f} {t_MPa}\nResistencia Nominal Base (Pn): {show_kN(Pn_cf):.1f} {t_kN}\nResistencia de Diseño Perfil de Lámina (φPn): {show_kN(phi_Pn_cf):.1f} {t_kN}\nEstado de Ejecución: {'CUMPLE' if Pu_cf <= phi_Pn_cf else 'FALLA'} (FS LRFD = {phi_Pn_cf/max(Pu_cf, 0.001):.2f})"
            agregar_al_despiece(
                tipo=f"{tipo_cf} (conformado)",
                seccion=f"{h_cf_i}x{b_cf_i}x{t_cf_i} [{t_mm}]",
                longitud_m=L_cf,
                peso_kg=peso_total_cf,
                observacion=_t(f"φPn={show_kN(phi_Pn_cf):.1f} {t_kN} vs Pu={Pu_cf_i:.1f} {t_kN}", f"φPn={show_kN(phi_Pn_cf):.1f} {t_kN} vs Pu={Pu_cf_i:.1f} {t_kN}"),
                memoria=mem_text
            )
            st.success(_t("Perfil agregado al despiece", "Profile added to list"))

# 
# TAB 5: EXPORTACIONES GLOBALES
# 
with tab_E:
    st.header(_t("Exportaciones Globales", "Global Exports"))

    # Mostrar despiece acumulado
    st.subheader(_t("Despiece de Acero", "Steel Cutting List"))
    if st.session_state.steel_despiece:
        df_desp = pd.DataFrame(st.session_state.steel_despiece)
        st.dataframe(df_desp.style.format({"Longitud (m)": "{:.2f}", "Peso (kg)": "{:.2f}"}), use_container_width=True, hide_index=False)
        st.write(f"**Peso total acumulado:** {st.session_state.steel_costo_total:.2f} kg")
        # Gráfico de barras
        fig_bars, ax_bars = plt.subplots(figsize=(8, 4))
        # Agrupar por tipo
        df_group = df_desp.groupby("Tipo")["Peso (kg)"].sum().reset_index()
        ax_bars.bar(df_group["Tipo"], df_group["Peso (kg)"], color='steelblue')
        ax_bars.set_xlabel(_t("Tipo de perfil", "Profile type"))
        ax_bars.set_ylabel(_t("Peso (kg)", "Weight (kg)"))
        ax_bars.set_title(_t("Distribución de pesos por tipo", "Weight distribution by type"))
        plt.setp(ax_bars.get_xticklabels(), rotation=45, ha="right", fontsize=9)
        fig_bars.subplots_adjust(bottom=0.35)
        ax_bars.grid(True, alpha=0.3)
        st.pyplot(fig_bars)
        bars_img = io.BytesIO()
        fig_bars.savefig(bars_img, format='png', dpi=150, bbox_inches='tight')
        bars_img.seek(0)
    else:
        st.info(_t("Aún no hay elementos en el despiece. Agrega perfiles desde las pestañas anteriores.", "No elements in the list yet. Add profiles from the previous tabs."))

    st.markdown("---")
    st.subheader(_t("Memoria de Cálculo (DOCX)", " Calculation Report (DOCX)"))
    if st.button(_t("Generar Memoria DOCX", "Generate DOCX Report")):
        doc = Document()
        doc.add_heading(_t(f"Memoria de Estructuras Metálicas - {norma_sel}", f"Steel Structures Report - {norma_sel}"), 0)
        doc.add_paragraph(_t(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", f"Date: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"))
        doc.add_heading(_t("1. Materiales", "1. Materials"), level=1)
        doc.add_paragraph(f"Fy = {Fy_adm:.1f} MPa, Fu = {Fu_adm:.1f} MPa, E = {E_steel:.0f} MPa")
        doc.add_heading(_t("2. Despiece de Acero", "2. Steel Cutting List"), level=1)
        if st.session_state.steel_despiece:
            table = doc.add_table(rows=1+len(st.session_state.steel_despiece), cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = _t("Tipo", "Type")
            hdr[1].text = _t("Sección", "Section")
            hdr[2].text = _t("Longitud (m)", "Length (m)")
            hdr[3].text = _t("Peso (kg)", "Weight (kg)")
            hdr[4].text = _t("Observación", "Observation")
            for i, row in enumerate(st.session_state.steel_despiece):
                cells = table.rows[i+1].cells
                cells[0].text = row["Tipo"]
                cells[1].text = row["Sección"]
                cells[2].text = f"{row['Longitud (m)']:.2f}"
                cells[3].text = f"{row['Peso (kg)']:.1f}"
                cells[4].text = row["Observación"]
            doc.add_paragraph(f"{_t('Peso total:', 'Total weight:')} {st.session_state.steel_costo_total:.2f} kg")
            # Insertar gráfico
            doc.add_picture(bars_img, width=Inches(5))
        else:
            doc.add_paragraph(_t("No se han agregado elementos.", "No elements added."))
            
        doc.add_heading(_t("3. Memoria Detallada por Componente", "3. Detailed Component Report"), level=1)
        for i, row in enumerate(st.session_state.steel_despiece):
            doc.add_heading(f"{_t('Elemento', 'Element')} {i+1}: {row['Tipo']} - {row['Sección']}", level=2)
            for linea in str(row.get('Memoria', row['Observación'])).split('\n'):
                doc.add_paragraph(linea)
                
        doc.add_heading(_t("4. Notas de diseño", "4. Design notes"), level=1)
        doc.add_paragraph(_t("Los cálculos se realizaron según los principios de resistencia LRFD (AISC 360) aplicables.", "Calculations performed using applicable LRFD principles (AISC 360)."))
        doc_mem = io.BytesIO()
        doc.save(doc_mem)
        doc_mem.seek(0)
        st.download_button(_t("Descargar Memoria", "Download Report"), data=doc_mem, file_name="Memoria_Acero.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.subheader(_t("Presupuesto APU", "APU Budget"))
    apu = st.session_state.get("apu_config", {
        "moneda": "COP$",
        "costo_dia_mo": 69333.33,
        "pct_herramienta": 0.05,
        "pct_aui": 0.30,
        "pct_util": 0.05,
        "iva": 0.19
    })
    mon = apu["moneda"]
    precio_kg = st.number_input(_t(f"Costo por Kg de Acero Estructural Suministrado y Armado [{mon}/kg]", f"Steel supply and erection cost per kg [{mon}/kg]"),
                                value=st.session_state.get("st_price_kg", 8000.0 if mon=="COP$" else 4.0), step=100.0, key="st_price_kg")
    costo_acero = st.session_state.steel_costo_total * precio_kg
    
    # Mano de obra (rendimiento aproximado 0.05 días/kg)
    dias_mo = st.session_state.steel_costo_total * 0.05
    costo_mo = dias_mo * apu.get("costo_dia_mo", 69333.33)
    
    # Herramienta, AIU, etc.
    herramienta = costo_mo * apu.get("pct_herramienta", 0.05)
    aiu = costo_acero * apu.get("pct_aui", 0.30)
    utilidad = costo_acero * apu.get("pct_util", 0.05)
    iva = utilidad * apu.get("iva", 0.19)
    total_proyecto = costo_acero + costo_mo + herramienta + aiu + iva

    data_apu = {
        _t("Item", "Item"): [_t("Acero estructural", "Structural steel"), _t("Mano de obra (montaje)", "Labor (erection)"),
                             _t("Herramienta menor", "Minor tools"), _t("A.I.U.", "A.I.U."),
                             _t("IVA s/Utilidad", "VAT on profit"), _t("TOTAL", "TOTAL")],
        _t("Cantidad", "Quantity"): [f"{st.session_state.steel_costo_total:.1f} kg", f"{dias_mo:.2f} días", f"{apu.get('pct_herramienta',0.05)*100:.1f}% MO",
                                     f"{apu.get('pct_aui',0.30)*100:.1f}%", f"{apu.get('iva',0.19)*100:.1f}%", ""],
        f"Subtotal [{mon}]": [f"{costo_acero:,.2f}", f"{costo_mo:,.2f}", f"{herramienta:,.2f}", f"{aiu:,.2f}",
                              f"{iva:,.2f}", f"**{total_proyecto:,.2f}**"]
    }
    st.dataframe(pd.DataFrame(data_apu), use_container_width=True, hide_index=True)

    # Exportar Excel APU
    output_excel = io.BytesIO()
    with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
        df_export = pd.DataFrame({
            "Item": ["Acero estructural", "Mano de obra"],
            "Cantidad": [st.session_state.steel_costo_total, dias_mo],
            "Unidad": [precio_kg, apu.get("costo_dia_mo", 69333.33)]
        })
        df_export["Subtotal"] = df_export["Cantidad"] * df_export["Unidad"]
        df_export.to_excel(writer, index=False, sheet_name='APU Acero')
        workbook = writer.book
        worksheet = writer.sheets['APU Acero']
        money_fmt = workbook.add_format({'num_format': '#,##0.00'})
        worksheet.set_column('A:A', 25)
        worksheet.set_column('B:D', 15, money_fmt)
        row = len(df_export) + 1
        worksheet.write(row, 0, "Costo Directo (CD)", workbook.add_format({'bold': True}))
        worksheet.write_formula(row, 3, f'=SUM(D2:D{row})', money_fmt)
        row += 1
        worksheet.write(row, 0, "Herramienta Menor", workbook.add_format({'bold': True}))
        worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_herramienta",0.05)}', money_fmt)
        row += 1
        worksheet.write(row, 0, "A.I.U", workbook.add_format({'bold': True}))
        worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_aui",0.30)}', money_fmt)
        row += 1
        worksheet.write(row, 0, "IVA s/ Utilidad", workbook.add_format({'bold': True}))
        worksheet.write_formula(row, 3, f'=D{row-1}*{apu.get("pct_util",0.05)}*{apu.get("iva",0.19)}', money_fmt)
        row += 1
        worksheet.write(row, 0, "TOTAL PRESUPUESTO", workbook.add_format({'bold': True}))
        worksheet.write_formula(row, 3, f'=D{row-3}+D{row-2}+D{row-1}+D{row}', money_fmt)
    output_excel.seek(0)
    st.download_button(_t("Descargar Presupuesto Excel", "Download Budget Excel"), data=output_excel,
                       file_name="APU_Acero.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.markdown("---")
    st.subheader(_t("DXF de Sección", "Section DXF"))
    st.write(_t("Genera un archivo DXF con la sección transversal del último perfil W definido en la pestaña de propiedades.", "Generates a DXF file with the cross-section of the last W-shape defined in the properties tab."))
    try:
        from dxf_helpers import (dxf_setup, dxf_add_layers, dxf_text,
                                 dxf_dim_horiz, dxf_dim_vert, dxf_rotulo, dxf_rotulo_campos)
        _USE_H_m = True
    except ImportError:
        _USE_H_m = False
    try:
        doc_dxf = ezdxf.new('R2010')
        doc_dxf.units = ezdxf.units.M
        if _USE_H_m:
            dxf_setup(doc_dxf, 20)
            dxf_add_layers(doc_dxf)
        msp_st = doc_dxf.modelspace()
        for lay, col in [('PERFIL_W',5), ('TEXTO',3), ('COTAS',2)]:
            if lay not in doc_dxf.layers:
                doc_dxf.layers.add(lay, color=col)
        # Usar dimensiones del perfil en mm → convertir a m para el DXF en metros
        if 'dw' in locals() and 'bfw' in locals():
            sc = 1/1000  # mm → m
            pts_w = [
                (-bfw/2*sc, dw/2*sc), (bfw/2*sc, dw/2*sc),
                (bfw/2*sc, (dw/2-tfw)*sc), (tww/2*sc, (dw/2-tfw)*sc),
                (tww/2*sc, -(dw/2-tfw)*sc), (bfw/2*sc, -(dw/2-tfw)*sc),
                (bfw/2*sc, -dw/2*sc), (-bfw/2*sc, -dw/2*sc),
                (-bfw/2*sc, -(dw/2-tfw)*sc), (-tww/2*sc, -(dw/2-tfw)*sc),
                (-tww/2*sc, (dw/2-tfw)*sc), (-bfw/2*sc, (dw/2-tfw)*sc),
                (-bfw/2*sc, dw/2*sc)
            ]
            msp_st.add_lwpolyline(pts_w, dxfattribs={'layer':'PERFIL_W', 'color':5, 'closed':True})
            if _USE_H_m:
                TH = 0.025 * 20
                dxf_dim_horiz(msp_st, -bfw/2*sc, bfw/2*sc, -dw/2*sc - 0.05,
                              f"bf = {show_mm(bfw):.1f} mm", 20)
                dxf_dim_vert(msp_st, bfw/2*sc + 0.05, -dw/2*sc, dw/2*sc,
                             f"d = {show_mm(dw):.1f} mm", 20)
                dxf_text(msp_st, 0, dw/2*sc + 0.04, "SECCION PERFIL W", "EJES", h=TH*1.2, ha="center")
                _cam_m = dxf_rotulo_campos(f"Perfil W {show_mm(dw):.0f}x{show_mm(bfw):.0f}mm", norma_sel, "001")
                dxf_rotulo(msp_st, _cam_m, -bfw/2*sc, -dw/2*sc - 0.25, rot_w=max(bfw*sc*2, 0.2), rot_h=0.15, escala=20)
            out_dxf = io.StringIO()
            doc_dxf.write(out_dxf)
            st.download_button(_t("Descargar Perfil_W.dxf", "Download Perfil_W.dxf"),
                               data=out_dxf.getvalue().encode('utf-8'),
                               file_name=f"Perfil_W_{show_mm(dw):.0f}x{show_mm(bfw):.0f}.dxf", mime="application/dxf")
        else:
            st.warning(_t("Primero defina un perfil W en la pestaña de propiedades.", "First define a W-shape in the properties tab."))
    except Exception as e:
        st.error(f"Error generando DXF: {e}")